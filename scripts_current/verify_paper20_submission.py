from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

import fitz
import openpyxl
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


ROOT = Path("D:/fzyc")
OUT = ROOT / "output"
CORE = OUT / "paper20_candidate_pool_audit_20260712"
PAPER = OUT / "小论文-20_Journal_of_Cheminformatics_主文.docx"
SUPP = CORE / "supplementary"
AUDIT = OUT / "小论文-20_最终完整性审计.json"
FORBIDDEN = ["test oracle", "true oracle", "true attainable upper bound", "independent confirmation", "universal governance framework", "prospective preregistration"]


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def border_value(cell, edge: str) -> str | None:
    borders = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
    if borders is None:
        return None
    node = borders.find(qn(f"w:{edge}"))
    return node.get(qn("w:val")) if node is not None else None


def main() -> None:
    doc = Document(PAPER)
    text = "\n".join(p.text for p in doc.paragraphs)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    methods = [x for x in headings if re.match(r"2\.\d+ ", x)]
    results = [x for x in headings if re.match(r"3\.\d+ ", x)]
    discussion = [x for x in headings if re.match(r"4\.\d+ ", x)]
    table_checks = []
    for i, table in enumerate(doc.tables, start=1):
        first, last = table.rows[0], table.rows[-1]
        table_checks.append({
            "table": i, "columns": len(table.columns),
            "header_top": all(border_value(c, "top") == "single" for c in first.cells),
            "header_bottom": all(border_value(c, "bottom") == "single" for c in first.cells),
            "last_bottom": all(border_value(c, "bottom") == "single" for c in last.cells),
            "no_vertical_borders": all(border_value(c, "left") == "nil" and border_value(c, "right") == "nil" for row in table.rows for c in row.cells),
        })
    fig_files = sorted((CORE / "main_figures").glob("Figure_*.*"))
    wb = openpyxl.load_workbook(SUPP / "Additional_file_2_Supplementary_Tables_S1-S17.xlsx", read_only=True)
    supp_pdf = fitz.open(SUPP / "Additional_file_3_Supplementary_Figures_S1-S14.pdf")
    main_pdf = fitz.open(OUT / "小论文-20_Journal_of_Cheminformatics_审阅版.pdf")
    cross = __import__("pandas").read_csv(CORE / "cross_fitted_k32_minus_k4.csv")
    zero = __import__("pandas").read_csv(CORE / "multiview_zero_width_verification.csv")
    checks = {
        "main_docx_opens": True,
        "main_pdf_pages": len(main_pdf),
        "methods_subsections": len(methods), "methods_expected": 13,
        "results_subsections": len(results), "results_expected": 8,
        "discussion_subsections": len(discussion), "discussion_expected": 7,
        "tables": len(doc.tables), "max_table_columns": max(len(t.columns) for t in doc.tables),
        "three_line_tables": table_checks,
        "landscape_sections": sum(s.orientation == WD_ORIENT.LANDSCAPE for s in doc.sections),
        "main_figure_files": len(fig_files), "main_figure_stems": sorted(set(p.stem for p in fig_files)),
        "supplementary_sheets": wb.sheetnames, "supplementary_sheet_count": len(wb.sheetnames),
        "supplementary_figure_pages": len(supp_pdf),
        "replacement_character_count": text.count("\ufffd"),
        "mojibake_hits": [x for x in ["锟", "candidate緋ool", "outer綼udit", "ROC続UC", "P緂lycoprotein"] if x in text],
        "forbidden_hits": {x: text.lower().count(x) for x in FORBIDDEN if x in text.lower()},
        "figure_caption_numbers": sorted(set(map(int, re.findall(r"Figure (\d+)\.", text)))),
        "table_caption_numbers": sorted(set(map(int, re.findall(r"Table (\d+)\.", text)))),
        "reference_count": len(re.findall(r"^\d+\. ", text, flags=re.MULTILINE)),
        "cross_fitted_positive_endpoints": int((cross.k32_minus_k4_cross_fitted_gap > 0).sum()),
        "cross_fitted_negative_endpoints": cross.loc[cross.k32_minus_k4_cross_fitted_gap < 0, "task"].tolist(),
        "same_unit_positive_endpoints": int((cross.k32_minus_k4_same_unit_gap > 0).sum()),
        "exact_zero_width_multiview_endpoints": zero.loc[zero.exact_zero_width_seed_interval.astype(str).str.lower().eq("true"), "task"].tolist(),
    }
    checks["pass"] = all([
        len(methods) == 13, len(results) == 8, len(discussion) == 7,
        len(doc.tables) == 3, checks["max_table_columns"] <= 6,
        all(all(v for k, v in t.items() if k not in {"table", "columns"}) for t in table_checks),
        checks["landscape_sections"] == 0, len(fig_files) == 18,
        len(wb.sheetnames) == 17, len(supp_pdf) == 14,
        checks["replacement_character_count"] == 0, not checks["mojibake_hits"], not checks["forbidden_hits"],
        checks["figure_caption_numbers"] == [1, 2, 3, 4, 5, 6],
        checks["table_caption_numbers"] == [1, 2, 3],
        checks["cross_fitted_positive_endpoints"] == 7,
    ])
    files = [PAPER, OUT / "小论文-20_Journal_of_Cheminformatics_审阅版.pdf", SUPP / "Additional_file_1_Supplementary_Methods_and_Results.docx", SUPP / "Additional_file_1_Supplementary_Methods_and_Results.pdf", SUPP / "Additional_file_2_Supplementary_Tables_S1-S17.xlsx", SUPP / "Additional_file_3_Supplementary_Figures_S1-S14.pdf"]
    checks["key_file_hashes"] = {str(p): sha256(p) for p in files}
    AUDIT.write_text(json.dumps(checks, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(checks, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
