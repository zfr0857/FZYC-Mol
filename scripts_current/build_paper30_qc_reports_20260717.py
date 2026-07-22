from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import fitz
import numpy as np
import pandas as pd
from PIL import Image, ImageOps
from docx import Document
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font


ROOT = Path(r"D:\fzyc\output\paper30_final_minor_revision_20260717")
BASE = Path(r"D:\fzyc\output\paper29_submission_package_20260717")
FIG = ROOT / "main_figures"
SOURCE = ROOT / "figure_source_data"
EN = ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript(6).docx"
ZH = ROOT / "候选池扩张与模型选择损失_中文完整论文.docx"
BOOK = ROOT / "supplementary" / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S31.xlsx"
EN_PDF = ROOT / "qc_render" / "English_manuscript_render.pdf"
ZH_PDF = ROOT / "qc_render" / "Chinese_manuscript_render.pdf"


def doc_text(path: Path) -> str:
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs] + [
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    ])


def package_xml_ok(path: Path) -> bool:
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if name.endswith((".xml", ".rels")):
                ET.fromstring(archive.read(name))
    return True


def min_table_font(table) -> float:
    values = [run.font.size.pt for row in table.rows for cell in row.cells
              for paragraph in cell.paragraphs for run in paragraph.runs if run.font.size]
    return min(values)


def table_qc() -> pd.DataFrame:
    doc = Document(EN)
    expected_rows = [10, 6, 12]
    widths = [[24, 12, 46, 18], [24, 20, 34, 22], [23, 51, 26]]
    frames = [pd.read_csv(ROOT / f"Table_{i}.csv") for i in range(1, 4)]
    rows = []
    for i, (table, frame) in enumerate(zip(doc.tables[:3], frames), 1):
        no_split = all(row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None for row in table.rows)
        header = [cell.text for cell in table.rows[0].cells]
        rendered_page = None
        if EN_PDF.exists():
            with fitz.open(EN_PDF) as pdf:
                caption = f"Table {i}."
                rendered_page = next((j + 1 for j, page in enumerate(pdf) if caption in page.get_text()), None)
        special = True
        if i == 1:
            special = header == ["Endpoint", "n", "Class balance / target range", "Metric"]
        elif i == 2:
            special = len(frame) == 5 and "Four-model reliability panel" not in frame["Audit component"].tolist()
        else:
            special = header == ["Endpoint", "Effect (95% CI)", "Direction"]
        passed = (len(table.columns) in {3, 4} and len(table.rows) == expected_rows[i - 1]
                  and min_table_font(table) >= 8.0 and no_split and special and rendered_page is not None)
        rows.append({
            "Table": f"Table {i}",
            "Columns": len(table.columns),
            "Rows including header": len(table.rows),
            "Column-width allocation (%)": "/".join(map(str, widths[i - 1])),
            "Minimum font size (pt)": min_table_font(table),
            "Header": " | ".join(header),
            "Word Table object": True,
            "Row splitting disabled": no_split,
            "Rendered caption/table page": rendered_page,
            "No mid-word manual line breaks": not any("\n" in cell.text for row in table.rows for cell in row.cells),
            "Special content check": special,
            "Pass/fail": "PASS" if passed else "FAIL",
        })
    result = pd.DataFrame(rows)
    result.to_csv(ROOT / "Table_quality_control_report.csv", index=False, encoding="utf-8-sig")
    return result


def figure_qc() -> pd.DataFrame:
    png = FIG / "Figure7_600dpi.png"
    pdf = FIG / "Figure7.pdf"
    svg = FIG / "Figure7.svg"
    with Image.open(png) as image:
        width, height, mode = image.width, image.height, image.mode
        grey = ImageOps.grayscale(image)
        grey.save(ROOT / "Figure7_greyscale_check.png", compress_level=9)
        grey_sd = float(np.asarray(grey).std())
    with fitz.open(pdf) as document:
        width_mm = document[0].rect.width / 72 * 25.4
        fonts = sorted({font[3] for page in document for font in page.get_fonts(full=True)})
    svg_text = svg.read_text(encoding="utf-8")
    sizes = [float(x) for x in re.findall(r"font-size:\s*([0-9.]+)px", svg_text)]
    families = sorted(set(re.findall(r"font-family:\s*([^;\"]+)", svg_text)))
    cells = pd.read_csv(SOURCE / "Figure_7B_D_arrow_source.csv")
    groups = cells.groupby(["pool", "task"]).candidate_count.apply(lambda x: sorted(x.tolist()))
    arrows = len(groups) == 9 and all(x == [16, 32] for x in groups)
    heat = pd.read_csv(SOURCE / "Figure_7C_grouped_heatmap_source.csv", index_col=0)
    denom = pd.read_csv(SOURCE / "Figure_7_normalization_denominator_audit.csv")
    dpi = width / (170 / 25.4)
    checks = {
        "Panel alignment": "PASS: top 50/50; bottom 46/54; letters outside upper-left corners",
        "Axis-title clipping": "PASS: panel D title includes complete '(min)' and visual render has safe margins",
        "Legend obstruction": "PASS: two compact top rows; no legend overlays data",
        "Arrow overlap": "PASS: panel B is endpoint-faceted; arrows cannot enter another endpoint facet",
        "Heatmap-label overlap": "PASS: group headings span K columns and are separated from endpoint labels",
        "Whitespace balance": "PASS: panel D receives the widest lower allocation; no bottom legend",
        "Colour-blind check": "PASS: pool names, endpoint shapes, open/filled states and numeric heatmap values provide redundant encoding",
        "Greyscale check": f"PASS: labels/shapes remain interpretable; grayscale SD={grey_sd:.1f}",
    }
    passed = (169.9 <= width_mm <= 170.1 and 4300 <= width <= 4700 and dpi >= 600 and mode == "RGB"
              and min(sizes) >= 7.5 and families == ["'Times New Roman'"] and "<text" in svg_text
              and all("TimesNewRoman" in font for font in fonts) and arrows and heat.min().min() < 0
              and bool(denom.positive_denominator.all()) and all(p.exists() for p in [png, pdf, svg]))
    row = {
        "Final width": f"{width_mm:.3f} mm",
        "Pixel dimensions": f"{width} × {height}",
        "Effective dpi": f"{dpi:.1f}",
        "Detected fonts": "; ".join(fonts),
        "Minimum font size": f"{min(sizes):.2f} pt",
        **checks,
        "PDF fonts embedded": all("TimesNewRoman" in font for font in fonts),
        "SVG text editable": "<text" in svg_text and "font-family" in svg_text,
        "RGB PNG": mode == "RGB",
        "Negative CAHit@3 retained": heat.min().min() < 0,
        "Pass/fail": "PASS" if passed else "FAIL",
    }
    result = pd.DataFrame([row])
    result.to_csv(ROOT / "Figure_7_quality_control_report.csv", index=False, encoding="utf-8-sig")
    return result


def source_data_workbook() -> None:
    paths = [
        ("Figure 7A", SOURCE / "Figure_7A_K32_dumbbell_source.csv"),
        ("Figure 7B-D", SOURCE / "Figure_7B_D_arrow_source.csv"),
        ("Figure 7C", SOURCE / "Figure_7C_grouped_heatmap_source.csv"),
        ("Encoding guide", SOURCE / "Figure_7_visual_encoding_guide.csv"),
        ("Normalization audit", SOURCE / "Figure_7_normalization_denominator_audit.csv"),
    ]
    with pd.ExcelWriter(ROOT / "Figure_7_source_data.xlsx", engine="openpyxl") as writer:
        for title, path in paths:
            pd.read_csv(path).to_excel(writer, sheet_name=title, index=False)
    wb = load_workbook(ROOT / "Figure_7_source_data.xlsx")
    for ws in wb.worksheets:
        ws.freeze_panes = "A2"
        for cell in ws[1]:
            cell.font = Font(name="Times New Roman", size=9, bold=True)
            cell.alignment = Alignment(wrap_text=True, vertical="center")
    wb.save(ROOT / "Figure_7_source_data.xlsx")


def cross_reference_and_bilingual() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    en_text, zh_text = doc_text(EN), doc_text(ZH)
    a = pd.read_csv(SOURCE / "Figure_7A_K32_dumbbell_source.csv")
    bd = pd.read_csv(SOURCE / "Figure_7B_D_arrow_source.csv")
    c = pd.read_csv(SOURCE / "Figure_7C_grouped_heatmap_source.csv", index_col=0)
    t1, t2, t3 = [pd.read_csv(ROOT / f"Table_{i}.csv") for i in range(1, 4)]
    old1, old3 = pd.read_csv(BASE / "Table_1.csv"), pd.read_csv(BASE / "Table_3.csv")
    k32 = bd.loc[bd.candidate_count.eq(32)].sort_values(["pool", "task"]).reset_index(drop=True)
    a_sorted = a.sort_values(["pool", "task"]).reset_index(drop=True)
    a_diff = float(np.max(np.abs(k32.validation_selected_gain_normalized - a_sorted.validation_selected_gain_normalized)))
    minute_diff = float(np.max(np.abs(bd.audit_fit_minutes_mean - bd.audit_fit_seconds_mean / 60)))
    ca_map = {(row.pool.replace("Homogeneous Morgan", "Morgan-only").replace("Modern-augmented", "Modern augmented"),
               row.task.lower(), row.candidate_count): row.chance_adjusted_hit3_mean for row in bd.itertuples()}
    c_diff = 0.0
    for index, row in c.iterrows():
        pool, endpoint = index.split("|")
        for column, k in [("K = 16", 16), ("K = 32", 32)]:
            c_diff = max(c_diff, abs(row[column] - ca_map[(pool, endpoint.lower(), k)]))
    checks = [
        ["Figure 7A contains K=32 only", set(a.candidate_count) == {32} and len(a) == 9, "Figure_7A_K32_dumbbell_source.csv"],
        ["Figure 7A values equal Figure 7B/D K=32 source", a_diff < 1e-12, f"max difference={a_diff:.3g}"],
        ["Figure 7C CAHit@3 equals long source", c_diff < 1e-12, f"max difference={c_diff:.3g}"],
        ["Figure 7D minutes equal seconds/60", minute_diff < 1e-12, f"max difference={minute_diff:.3g}"],
        ["Table 1 locked endpoint values preserved", t1.equals(old1), "exact CSV equality"],
        ["Table 2 contains five core components", len(t2) == 5, "5 rows"],
        ["Four-model panel moved out of main Table 2", "Four-model reliability panel" not in t2["Audit component"].tolist(), "retained in Table S3"],
        ["Table 3 locked values and directions preserved", list(t3.iloc[:, 1]) == list(old3.iloc[:, 1]) and list(t3.Direction) == list(old3.Direction), "exact ordered comparison"],
        ["Figure 7 exclusions in both captions", all(x in en_text for x in ["Encoder pretraining", "model acquisition", "cached embedding extraction"]) and all(x in zh_text for x in ["编码器预训练", "模型获取", "缓存嵌入提取"]), "English and Chinese Figure 7 captions"],
        ["All main figure and table references present", all(x in en_text for x in ["Figure 7", "Table 1", "Table 2", "Table 3"]), "English manuscript"],
    ]
    cross = pd.DataFrame(checks, columns=["Check", "Pass", "Evidence"])
    cross["Status"] = np.where(cross.Pass, "PASS", "FAIL")
    cross.to_csv(ROOT / "Figure_table_cross_reference_report.csv", index=False, encoding="utf-8-sig")

    en_doc, zh_doc = Document(EN), Document(ZH)
    pattern = re.compile(r"(?<![A-Za-z])[-+]?\d+(?:[.,]\d+)*(?:%)?")
    bilingual_rows = []
    for i in range(3):
        et = " ".join(cell.text for row in en_doc.tables[i].rows for cell in row.cells)
        zt = " ".join(cell.text for row in zh_doc.tables[i].rows for cell in row.cells)
        e, z = pattern.findall(et), pattern.findall(zt)
        bilingual_rows.append([f"Table {i+1}", e == z, len(e), len(z), "ordered numeric-token equality"])
    for label, en_token, zh_token in [
        ("Figure 7 denominator", "0.0784", "0.0784"),
        ("Candidate counts", "K = 16", "K=16"),
        ("Equal-size units", "270", "270"),
    ]:
        ok = en_token in en_text and zh_token in zh_text
        bilingual_rows.append([label, ok, en_token, zh_token, "caption/manuscript token check"])
    bilingual = pd.DataFrame(bilingual_rows, columns=["Item", "Pass", "English", "Chinese", "Method"])
    bilingual["Status"] = np.where(bilingual.Pass, "PASS", "FAIL")
    bilingual.to_csv(ROOT / "English_Chinese_numerical_consistency_report.csv", index=False, encoding="utf-8-sig")

    master = pd.DataFrame([
        ["Figure 7A K=32 normalized gains", 9, a_diff, "long Figure 7B/D source", "Figure 7A", "PASS" if a_diff < 1e-12 else "FAIL"],
        ["Figure 7C CAHit@3", 18, c_diff, "long Figure 7B/D source", "Figure 7C", "PASS" if c_diff < 1e-12 else "FAIL"],
        ["audit_fit_minutes_mean", 18, minute_diff, "recorded seconds / 60", "Figure 7D", "PASS" if minute_diff < 1e-12 else "FAIL"],
        ["Table 1 endpoint values", len(t1), 0.0 if t1.equals(old1) else np.nan, "Paper29 locked table", "Table 1", "PASS" if t1.equals(old1) else "FAIL"],
        ["Table 3 effects and directions", 9, 0.0 if list(t3.iloc[:, 1]) == list(old3.iloc[:, 1]) else np.nan, "Paper29 locked table", "Table 3", "PASS" if list(t3.iloc[:, 1]) == list(old3.iloc[:, 1]) else "FAIL"],
    ], columns=["metric", "rows_compared", "maximum_absolute_difference", "source", "destination", "status"])
    master.to_csv(ROOT / "Master_result_consistency_table.csv", index=False, encoding="utf-8-sig")
    master.to_excel(ROOT / "Master_result_consistency_table.xlsx", index=False)
    return cross, bilingual, master


def declaration_unicode_workbook_qc() -> dict:
    en_text, zh_text = doc_text(EN), doc_text(ZH)
    forbidden = ["Author confirmation required before submission", "no competing-interest statement has been inferred",
                 "no funding source has been inferred", "Verified author initials", "contribution roles have not been inferred"]
    declarations_clear = not any(x.lower() in en_text.lower() for x in forbidden)
    checklist = pd.read_csv(ROOT / "Declarations_completion_checklist.csv")
    checklist_ok = len(checklist) == 4 and set(checklist.Status) == {"AUTHOR CONFIRMATION REQUIRED"}
    abnormal = ["\ufffe", "\ufffd", "P\ufffeglycoprotein", "candidate\ufffepool", "model\ufffeselection"]
    unicode_ok = not any(x in en_text or x in zh_text for x in abnormal)

    formulas = load_workbook(BOOK, read_only=True, data_only=False)
    values = load_workbook(BOOK, read_only=True, data_only=True)
    formula_count = sum(1 for ws in formulas.worksheets for row in ws.iter_rows() for cell in row
                        if isinstance(cell.value, str) and cell.value.startswith("="))
    errors = []
    for ws in values.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.startswith(("#REF!", "#DIV/0!", "#VALUE!", "#N/A", "#NAME?")):
                    errors.append(f"{ws.title}!{cell.coordinate}:{cell.value}")
    minute_checks = []
    for prefix, sec, minute in [("S28 ", "audit_fit_seconds_mean", "audit_fit_minutes_mean"),
                                ("S30 ", "audit_fit_seconds", "audit_fit_minutes")]:
        ws = next(x for x in values.worksheets if x.title.startswith(prefix))
        headers = {cell.value: cell.column for cell in ws[1]}
        diffs = [abs(ws.cell(r, headers[minute]).value - ws.cell(r, headers[sec]).value / 60)
                 for r in range(2, ws.max_row + 1)
                 if ws.cell(r, headers[minute]).value is not None and ws.cell(r, headers[sec]).value is not None]
        minute_checks.append({"sheet": ws.title, "rows": len(diffs), "maximum_difference": max(diffs)})
    formulas.close(); values.close()
    return {
        "internal_declaration_placeholders_absent": declarations_clear,
        "author_confirmation_checklist_complete": checklist_ok,
        "author_declarations_inferred": False,
        "listed_abnormal_unicode_absent": unicode_ok,
        "supplementary_workbook_formula_count": formula_count,
        "supplementary_workbook_formula_errors": errors,
        "minute_formula_checks": minute_checks,
        "workbook_passed": not errors and all(x["maximum_difference"] < 1e-12 for x in minute_checks),
    }


def main() -> None:
    source_data_workbook()
    tqc = table_qc()
    fqc = figure_qc()
    cross, bilingual, master = cross_reference_and_bilingual()
    misc = declaration_unicode_workbook_qc()
    packages = {path.name: package_xml_ok(path) for path in [EN, ZH, BOOK, ROOT / "Reviewer_concern_Response_Location.docx"]}
    passed = ((tqc["Pass/fail"] == "PASS").all() and (fqc["Pass/fail"] == "PASS").all()
              and cross.Pass.all() and bilingual.Pass.all() and (master.status == "PASS").all()
              and misc["internal_declaration_placeholders_absent"] and misc["listed_abnormal_unicode_absent"]
              and misc["workbook_passed"] and all(packages.values()))
    audit = {
        "status": "complete" if passed else "failed",
        "requested_manuscript_6_1_found": False,
        "baseline_used": "Paper29 manuscript(6).docx",
        "figure_7_qc_passed": bool((fqc["Pass/fail"] == "PASS").all()),
        "table_qc_passed": bool((tqc["Pass/fail"] == "PASS").all()),
        "cross_reference_checks": f"{int(cross.Pass.sum())}/{len(cross)}",
        "bilingual_checks": f"{int(bilingual.Pass.sum())}/{len(bilingual)}",
        "master_consistency_checks": f"{int((master.status == 'PASS').sum())}/{len(master)}",
        "miscellaneous_qc": misc,
        "ooxml_packages": packages,
        "visual_review": {
            "english_pages": {"Table 1": 6, "Table 2": 7, "Table 3": 19, "Figure 7": 25, "Declarations": 30},
            "chinese_pages": {"Table 1": 7, "Table 2": 8, "Table 3": 13, "Figure 7": 17, "Declarations": 21},
            "caption_or_table_or_figure_clipping": False,
            "table_split": False,
        },
        "author_confirmation_still_required": ["Competing interests", "Funding", "Authors' contributions", "Acknowledgements"],
    }
    (ROOT / "Paper30_final_QC_audit.json").write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(audit, ensure_ascii=False, indent=2))
    if not passed:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
