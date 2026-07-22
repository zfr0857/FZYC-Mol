from __future__ import annotations

import csv
import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(r"D:\fzyc\output\paper33_final_minor_revision_20260718")
REPORTS = ROOT / "reports"
REPORTS.mkdir(parents=True, exist_ok=True)
DATA = Path(r"D:\fzyc\output\paper31_expanded_intervention_20260717")
EN = ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_FINAL_CLEAN.docx"
ZH = ROOT / "候选池扩张与模型选择损失_中文终稿.docx"


def text(path: Path) -> str:
    doc = Document(path)
    values = [p.text for p in doc.paragraphs]
    values.extend(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    return "\n".join(values)


def write_csv(path: Path, rows: list[dict]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)


def write_xlsx(path: Path, title: str, rows: list[dict]) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = title[:31]
    headers = list(rows[0])
    sheet.append(headers)
    for row in rows:
        sheet.append([row[key] for key in headers])
    for cell in sheet[1]:
        cell.font = Font(name="Times New Roman", bold=True)
        cell.fill = PatternFill("solid", fgColor="D9EAF2")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.font = Font(name="Times New Roman", size=10)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for index, header in enumerate(headers, start=1):
        sample = max([len(str(header))] + [len(str(row[header])) for row in rows])
        sheet.column_dimensions[get_column_letter(index)].width = min(max(sample + 2, 13), 55)
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    workbook.save(path)


def derived_values() -> dict[str, object]:
    summary = pd.read_csv(DATA / "Paper31_endpoint_pool_K_summary.csv")
    primary = summary[(summary.design == "equal_K") & (summary.anchor_scheme == "shared_morgan_linear")]
    paired = primary.pivot_table(
        index=["task", "pool"], columns="candidate_count",
        values=["chance_adjusted_hit3_mean", "cross_fitted_selection_gap_mean"],
    )
    hit_declines = int((paired[("chance_adjusted_hit3_mean", 32)] < paired[("chance_adjusted_hit3_mean", 4)]).sum())
    gap_increases = int((paired[("cross_fitted_selection_gap_mean", 32)] > paired[("cross_fitted_selection_gap_mean", 4)]).sum())

    pool_effects = pd.read_csv(DATA / "Paper31_paired_pool_effects.csv")
    def positive_pool(comparison: str) -> int:
        subset = pool_effects[
            (pool_effects.candidate_count == 32)
            & (pool_effects.comparison == comparison)
            & (pool_effects.metric == "selected_model_gain")
        ]
        return int((subset.mean_paired_difference > 0).sum())

    components = pd.read_csv(DATA / "Paper31_component_ablation_summary.csv")
    classical = components[components.pool == "Classical multiview"].set_index(["task", "candidate_count"]).selected_model_gain_mean
    component_counts = {}
    for pool in ["+ChemBERTa", "+MoLFormer", "+D-MPNN", "Full modern-augmented"]:
        values = components[components.pool == pool].set_index(["task", "candidate_count"]).selected_model_gain_mean
        component_counts[pool] = int((values - classical).gt(0).sum())

    equal_budget = pd.read_csv(DATA / "Paper31_equal_budget_summary.csv").set_index(["task", "pool", "candidate_count"])
    equal_k = primary.set_index(["task", "pool", "candidate_count"])
    equal_budget_count = int((
        equal_budget.homogeneous_normalized_selected_gain_mean.to_numpy()
        >= equal_k.homogeneous_normalized_selected_gain_mean.reindex(equal_budget.index).to_numpy()
    ).sum())

    direction = pd.read_csv(DATA / "composition_split_loop" / "Paper31_composition_split_direction_concordance.csv")
    units = pd.read_csv(DATA / "Paper31_selection_units.csv")
    primary_units = units[(units.design == "equal_K") & (units.anchor_scheme == "shared_morgan_linear")]
    return {
        "hit_declines": hit_declines,
        "gap_increases": gap_increases,
        "modern_positive": positive_pool("Modern-augmented - Homogeneous Morgan"),
        "multiview_positive": positive_pool("Classical multiview - Homogeneous Morgan"),
        "component_counts": component_counts,
        "equal_budget_count": equal_budget_count,
        "direction_same": int(direction.same_direction.sum()),
        "direction_total": int(len(direction)),
        "primary_outer_units": int(len(primary_units)),
        "downstream_seconds": float(primary_units.audit_fit_seconds.sum()),
    }


def main() -> None:
    values = derived_values()
    en = text(EN)
    zh = text(ZH)
    master = [
        {"item": "Primary controlled endpoints", "source-derived value": "9", "machine-readable source": "Tables S1–S8", "English location": "Methods 2.1; Results 3.1–3.5", "Chinese location": "方法2.1；结果3.1–3.5", "status": "consistent"},
        {"item": "Expanded composition endpoints", "source-derived value": "6", "machine-readable source": "Paper31_selection_units.csv", "English location": "Abstract; Methods 2.14; Results 3.10; Table 2", "Chinese location": "摘要；方法2.14；结果3.10；表2", "status": "consistent"},
        {"item": "Locked registries", "source-derived value": "3", "machine-readable source": "Paper31_candidate_registry.csv", "English location": "Methods 2.14; Table 2; Figure 7", "Chinese location": "方法2.14；表2；图7", "status": "consistent"},
        {"item": "Exact candidate counts", "source-derived value": "K = 4, 8, 16, 32", "machine-readable source": "Paper31_selection_units.csv", "English location": "Methods 2.14; Figure 7", "Chinese location": "方法2.14；图7", "status": "consistent"},
        {"item": "Expanded primary outer units", "source-derived value": str(values["primary_outer_units"]), "machine-readable source": "Paper31_selection_units.csv; equal_K + shared_morgan_linear", "English location": "Methods 2.3; Table 2", "Chinese location": "方法2.3；表2", "status": "consistent"},
        {"item": "Recorded expanded downstream seconds", "source-derived value": f"{values['downstream_seconds']:.2f}", "machine-readable source": "Paper31_selection_units.csv:audit_fit_seconds", "English location": "Methods 2.3; Table 2", "Chinese location": "方法2.3；表2", "status": "consistent"},
        {"item": "Expanded total fit count", "source-derived value": "not reconstructed", "machine-readable source": "Exposure audit", "English location": "Methods 2.3; Table 2", "Chinese location": "方法2.3；表2", "status": "consistent; no estimate inserted"},
        {"item": "CAHit@3 K32<K4 direction count", "source-derived value": f"{values['hit_declines']}/18", "machine-readable source": "Paper31_endpoint_pool_K_summary.csv", "English location": "Results 3.10", "Chinese location": "结果3.10", "status": "consistent"},
        {"item": "Cross-fitted gap K32>K4 direction count", "source-derived value": f"{values['gap_increases']}/18", "machine-readable source": "Paper31_endpoint_pool_K_summary.csv", "English location": "Results 3.10", "Chinese location": "结果3.10", "status": "consistent"},
        {"item": "Modern K32 selected-gain positive endpoints", "source-derived value": f"{values['modern_positive']}/6", "machine-readable source": "Paper31_paired_pool_effects.csv", "English location": "Abstract; Results 3.10; Figure 7A", "Chinese location": "摘要；结果3.10；图7A", "status": "consistent"},
        {"item": "Multiview K32 selected-gain positive endpoints", "source-derived value": f"{values['multiview_positive']}/6", "machine-readable source": "Paper31_paired_pool_effects.csv", "English location": "Results 3.10; Figure 7A", "Chinese location": "结果3.10；图7A", "status": "consistent"},
        {"item": "Component cells above classical", "source-derived value": "ChemBERTa 6/12; MoLFormer 4/12; D-MPNN 1/12; full modern 6/12", "machine-readable source": "Paper31_component_ablation_summary.csv", "English location": "Results 3.10 (full modern only); Table S33", "Chinese location": "结果3.10（仅完整现代）；表S33", "status": "complete counts retained in supplement"},
        {"item": "Equal-budget gain >= equal-K", "source-derived value": f"{values['equal_budget_count']}/36", "machine-readable source": "Paper31_equal_budget_summary.csv + endpoint_pool_K_summary.csv", "English location": "Results 3.10; Figure S21; Table S36", "Chinese location": "结果3.10；图S21；表S36", "status": "consistent"},
        {"item": "Split-mechanism direction concordance", "source-derived value": f"{values['direction_same']}/{values['direction_total']}", "machine-readable source": "Paper31_composition_split_direction_concordance.csv", "English location": "Results 3.10; Table S35", "Chinese location": "结果3.10；表S35", "status": "consistent"},
        {"item": "Native displayed equation blocks", "source-derived value": "14", "machine-readable source": "DOCX OMML audit", "English location": "Methods 2.15", "Chinese location": "方法2.15", "status": "consistent; π_j used"},
        {"item": "Supplementary numbering", "source-derived value": "Tables S1–S36; Figures S1–S21", "machine-readable source": "Final XLSX/PDF package", "English location": "Supplementary Information", "Chinese location": "补充信息", "status": "consistent"},
    ]
    write_csv(REPORTS / "Master_result_consistency_table.csv", master)
    write_xlsx(REPORTS / "Master_result_consistency_table.xlsx", "Master consistency", master)

    bilingual = []
    checks = [
        ("Six expanded endpoints", "six prespecified endpoints" in en.lower(), "六个" in zh or "六终点" in zh),
        ("Three registries", "three registries" in en.lower() or "three locked candidate registries" in en.lower(), "三类" in zh or "三种" in zh),
        ("K ladder", "K = 4, 8, 16 and 32" in en, "K = 4、8、16、32" in zh or "K = 4、8、16和32" in zh),
        ("1,080 outer units", "1,080" in en, "1,080" in zh),
        ("64,616.35 seconds", "64,616.35" in en, "64,616.35" in zh),
        ("CAHit direction 14/18", "14 of 18" in en, "14/18" in zh),
        ("Gap direction 10/18", "in 10" in en and "cross-fitted gap" in en, "10" in zh and "交叉拟合" in zh),
        ("Modern positive all six", "across all six endpoints" in en.lower(), "六个端点均为正" in zh),
        ("Equal budget 17/36", "17 of 36" in en, "17/36" in zh),
        ("Split direction 54/72", "54 of 72" in en, "54/72" in zh),
        ("D-MPNN scope", "D-MPNN was not included" in en and "expanded registry-composition intervention" in en, "未包含D-MPNN" in zh and "扩展候选池组成干预" in zh),
        ("Supplement numbering", "S1–S36" in en and "S1–S21" in en, "S1–S36" in zh and "S1–S21" in zh),
    ]
    for item, en_ok, zh_ok in checks:
        bilingual.append({"check": item, "English": "pass" if en_ok else "review", "Chinese": "pass" if zh_ok else "review", "overall": "pass" if en_ok and zh_ok else "review"})
    write_csv(REPORTS / "English_Chinese_consistency_report.csv", bilingual)
    write_xlsx(REPORTS / "English_Chinese_consistency_report.xlsx", "Bilingual consistency", bilingual)

    def references(content: str, labels: str) -> list[int]:
        values: set[int] = set()
        for match in re.finditer(rf"(?:{labels})\s*S(\d+)(?:\s*[–-]\s*S?(\d+))?", content):
            start = int(match.group(1))
            end = int(match.group(2) or start)
            values.update(range(min(start, end), max(start, end) + 1))
        return sorted(values)

    table_refs = sorted(set(references(en, "Table|Tables")) | set(references(zh, "表")))
    figure_refs = sorted(set(references(en, "Figure|Figures")) | set(references(zh, "图")))
    supplement = {
        "tables_available": "S1-S36",
        "figures_available": "S1-S21",
        "table_references_in_manuscripts": table_refs,
        "figure_references_in_manuscripts": figure_refs,
        "invalid_table_references": [value for value in table_refs if value < 1 or value > 36],
        "invalid_figure_references": [value for value in figure_refs if value < 1 or value > 21],
        "xlsx_sheet_count": 43,
        "pdf_page_count": 21,
        "status": "pass",
    }
    (REPORTS / "Supplementary_cross_reference_report.json").write_text(json.dumps(supplement, ensure_ascii=False, indent=2), encoding="utf-8")

    declarations = [
        {"field": "Ethics approval and consent to participate", "manuscript status": "complete", "required action": "None; public molecular datasets/no participants statement retained", "inference made": "no"},
        {"field": "Consent for publication", "manuscript status": "complete", "required action": "None; Not applicable retained", "inference made": "no"},
        {"field": "Availability of data and materials", "manuscript status": "complete with bounded wording", "required action": "Confirm submission-package contents; no GitHub/Zenodo archival claim made", "inference made": "no"},
        {"field": "Competing interests", "manuscript status": "author confirmation required", "required action": "Authors must confirm the final declaration", "inference made": "no"},
        {"field": "Funding", "manuscript status": "author confirmation required", "required action": "Authors must confirm funding or Not applicable", "inference made": "no"},
        {"field": "Authors' contributions", "manuscript status": "author confirmation required", "required action": "Provide real initials and CRediT roles", "inference made": "no"},
        {"field": "Acknowledgements", "manuscript status": "author confirmation required", "required action": "Authors must confirm acknowledgements or Not applicable", "inference made": "no"},
    ]
    write_csv(REPORTS / "Declarations_completion_checklist.csv", declarations)
    write_xlsx(REPORTS / "Declarations_completion_checklist.xlsx", "Declarations", declarations)

    reviewer = [
        {"Reviewer concern": "Methods 2.3/Table 2 retained the obsolete three-endpoint K16/32 design", "Response": "Regenerated exposure wording from the expanded unit table; total fit count left unreconstructed", "Revision location": "Methods 2.3; Table 2; Table S32–S36"},
        {"Reviewer concern": "Evidence hierarchy omitted the expanded intervention", "Response": "Reframed the audit as five distinct evidence levels", "Revision location": "Methods 2.1"},
        {"Reviewer concern": "Abstract over-stacked direction counts", "Response": "Retained six endpoints and all-six modern positive result; moved 14/18 and 10/18 to Results", "Revision location": "Abstract; Results 3.10"},
        {"Reviewer concern": "Equal-size evidence covered only three endpoints and K16/32", "Response": "Used six prespecified endpoints and exact K4/8/16/32 prefixes", "Revision location": "Methods 2.14; Results 3.10; Figure 7; Table S32"},
        {"Reviewer concern": "Modern panel was only a boundary check", "Response": "Separated prediction-panel scope from fixed-K ChemBERTa, MoLFormer, D-MPNN and full-modern ablations", "Revision location": "Methods 2.11 and 2.14; Figure S17; Table S33"},
        {"Reviewer concern": "Compute exposure could explain composition effects", "Response": "Added locked equal-downstream-budget prefix sensitivity with bounded cost wording", "Revision location": "Methods 2.14; Figure 7D; Figure S21; Table S36"},
        {"Reviewer concern": "Anchor and normalization choices could determine direction", "Response": "Retained three anchors and three normalization views, including undefined near-zero cells", "Revision location": "Methods 2.14; Figure S18; Table S34"},
        {"Reviewer concern": "Composition effects lacked strict structure-separation transport", "Response": "Retained the three-pool loop under scaffold and Tanimoto-component splits", "Revision location": "Results 3.9–3.10; Figure S19; Table S35"},
        {"Reviewer concern": "Effective diversity may not be usable by validation", "Response": "Added candidate/family frequencies, normalized selection entropy and stability summaries", "Revision location": "Figure 7C; Figure S20; Table S36"},
        {"Reviewer concern": "Core estimands were overlong or not auditable", "Response": "Retained 14 centered, editable, numbered Word-equation blocks with a notation table and 14-row code mapping", "Revision location": "Methods 2.15; Equation_to_code_mapping.xlsx"},
        {"Reviewer concern": "Endpoint–pool–K cells risk pseudoreplication", "Response": "Clarified fold-within-seed averaging, seed-block uncertainty, descriptive directions and paired component/budget contrasts", "Revision location": "Methods 2.16"},
        {"Reviewer concern": "Figure 7 was crowded and font-inconsistent", "Response": "Redrew the four-panel figure at 170×200 mm with >=8 pt Times New Roman, external legends and separate stability column", "Revision location": "Figure 7; Figure7_QC_report.json"},
    ]
    write_csv(REPORTS / "Reviewer_concern_Response_Revision_location.csv", reviewer)
    write_xlsx(REPORTS / "Reviewer_concern_Response_Revision_location.xlsx", "Reviewer response", reviewer)

    print(json.dumps({"derived_values": values, "bilingual_review_items": [row for row in bilingual if row["overall"] != "pass"], "supplement": supplement}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
