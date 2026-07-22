from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from lxml import etree
from openpyxl import load_workbook
from PIL import Image


ROOT = Path(r"D:\fzyc")
RUN = ROOT / "results" / "equal_size_registry_composition_20260716" / "new_candidates"
LOGS = ROOT / "results" / "equal_size_registry_composition_20260716" / "logs"
ANALYSIS = ROOT / "output" / "paper27_equal_size_registry_composition_20260716"
REVISION = ROOT / "output" / "paper27_equal_size_registry_composition_revision_20260716"
SUPP = REVISION / "supplementary"
TASKS = ("clintox", "bace", "esol")
SEEDS = (11, 23, 37, 53, 71)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1 << 20), b""):
            digest.update(block)
    return digest.hexdigest()


def doc_text(path: Path) -> str:
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    chunks.extend(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    return "\n".join(chunks)


def valid_ooxml(path: Path) -> bool:
    try:
        with zipfile.ZipFile(path) as archive:
            etree.fromstring(archive.read("[Content_Types].xml"))
            for name in archive.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    etree.fromstring(archive.read(name))
        return True
    except (zipfile.BadZipFile, etree.XMLSyntaxError, KeyError):
        return False


def main() -> None:
    checks: list[dict[str, object]] = []

    def check(name: str, passed: bool, evidence: object) -> None:
        checks.append({"check": name, "passed": bool(passed), "evidence": evidence})

    complete = sum((RUN / task / f"seed_{seed}" / "complete.json").exists() for task in TASKS for seed in SEEDS)
    check("all 15 task-seed runs complete", complete == 15, complete)

    inner = pd.read_csv(RUN / "inner_scores.csv")
    outer = pd.read_csv(RUN / "outer_candidate_scores.csv")
    registry = pd.read_csv(RUN / "candidate_registry.csv")
    check("6,300 new candidate fits represented", len(inner) + len(outer) == 6300, {"inner": len(inner), "outer": len(outer)})
    candidate_count = registry.groupby(["task", "seed"])["candidate"].nunique().to_dict()
    check("35 new candidates registered per task-seed", all(value == 35 for value in candidate_count.values()), candidate_count)
    per_endpoint = {
        task: {"inner": int((inner["task"] == task).sum()), "outer": int((outer["task"] == task).sum())}
        for task in TASKS
    }
    check("balanced endpoint exposure", all(v == {"inner": 1575, "outer": 525} for v in per_endpoint.values()), per_endpoint)
    clean_logs = {name: (LOGS / name).stat().st_size for name in ("clintox.err.log", "bace.err.log", "esol_lsqr.err.log")}
    check("final runner error logs empty", all(size == 0 for size in clean_logs.values()), clean_logs)

    audit = json.loads((ANALYSIS / "equal_size_registry_composition_audit.json").read_text(encoding="utf-8"))
    check("analysis audit complete", audit["status"] == "complete" and all(audit["checks"].values()), audit["status"])
    check("analysis row counts exact", audit["rows"] == {"inner": 12960, "outer": 4320, "selection_units": 270, "summary": 18, "diversity": 72}, audit["rows"])
    check("all 18 endpoint-pool-K balance checks passed", len(audit["checks"]) == 18 and all(audit["checks"].values()), len(audit["checks"]))

    summary = pd.read_csv(ANALYSIS / "equal_size_endpoint_summary.csv")
    k32 = summary.loc[summary["candidate_count"].eq(32)]
    expected_pools = {"Homogeneous Morgan", "Classical multiview", "Modern-augmented"}
    check("three exact K=32 pools across endpoints", set(k32["pool"]) == expected_pools and len(k32) == 9, sorted(k32["pool"].unique()))
    paired = pd.read_csv(ANALYSIS / "equal_size_paired_pool_effects.csv")
    favourable = paired.loc[
        (paired["candidate_count"].eq(32))
        & paired["metric"].isin(["oracle_opportunity_gain", "selected_model_gain"])
    ]
    check(
        "all fixed-K composition effects improve oracle and selected gain",
        len(favourable) == 12 and (favourable["mean_paired_difference"] > 0).all(),
        favourable[["task", "comparison", "metric", "mean_paired_difference"]].to_dict("records"),
    )
    check("non-monotone diversity-gap result retained", abs(audit["association"]["relative_diversity_vs_cross_fitted_gap_spearman"] + 0.2109621885455056) < 1e-12, audit["association"])

    english = REVISION / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript.docx"
    tracked = REVISION / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript_tracked_changes.docx"
    reviewer = REVISION / "Reviewer_concern_Response_Location.docx"
    chinese = next(path for path in REVISION.glob("*.docx") if path not in {english, tracked, reviewer})
    supplement = SUPP / "Additional_file_1_Supplementary_Methods_and_Results.docx"
    en_text, zh_text, supp_text = doc_text(english), doc_text(chinese), doc_text(supplement)
    for phrase in ("2.14 Equal-size registry-composition intervention", "3.10 Equal-size composition changed", "4.8 Registry composition", "Figure 7.", "Tables S1-S31"):
        check(f"English manuscript contains {phrase}", phrase in en_text, phrase)
    for phrase in ("2.14 等规模候选池组成干预", "3.10 等规模组成改变", "4.8 注册表组成", "图7", "S1–S31"):
        check(f"Chinese manuscript contains {phrase}", phrase in zh_text, phrase)
    for phrase in ("S15 Equal-size registry-composition intervention", "Table S27.", "Table S31."):
        check(f"Supplement contains {phrase}", phrase in supp_text, phrase)

    workbook = SUPP / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S31.xlsx"
    wb = load_workbook(workbook, read_only=True, data_only=False)
    required = {"S27 Registry composition", "S28 Endpoint summary", "S29 Effective diversity", "S30 Selection units", "S31 Paired pool effects"}
    check("supplementary workbook contains S27-S31", required.issubset(wb.sheetnames), wb.sheetnames[-5:])
    check("supplementary workbook has 31 sheets", len(wb.sheetnames) == 31, len(wb.sheetnames))
    wb.close()

    figure_dir = REVISION / "main_figures"
    figure_files = [figure_dir / "Figure_7_equal_size_registry_composition.pdf", figure_dir / "Figure_7_equal_size_registry_composition.svg", figure_dir / "Figure_7_equal_size_registry_composition_600dpi.png"]
    check("Figure 7 vector and raster outputs present", all(path.is_file() and path.stat().st_size > 10000 for path in figure_files), [path.name for path in figure_files])
    width, height = Image.open(figure_files[2]).size
    check("Figure 7 high-resolution raster", width >= 4000 and height >= 2000, {"width": width, "height": height})
    check("Figure 7 SVG retains editable text", "<text" in figure_files[1].read_text(encoding="utf-8"), True)

    for path in (english, chinese, supplement, tracked, reviewer, workbook):
        check(f"valid OOXML: {path.name}", valid_ooxml(path), path.stat().st_size)
    with zipfile.ZipFile(tracked) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
        settings = etree.fromstring(archive.read("word/settings.xml"))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    insertions = len(root.xpath("//w:ins", namespaces=ns))
    deletions = len(root.xpath("//w:del", namespaces=ns))
    check("tracked document has insertions and deletions", insertions >= 20 and deletions >= 5, {"insertions": insertions, "deletions": deletions})
    check("tracked revisions enabled", bool(settings.xpath("//w:trackRevisions", namespaces=ns)), True)

    document_layout: dict[str, object] = {}
    for label, path in (("English", english), ("Chinese", chinese), ("Supplement", supplement)):
        doc = Document(path)
        printable_widths = [section.page_width - section.left_margin - section.right_margin for section in doc.sections]
        maximum_shape = max((shape.width for shape in doc.inline_shapes), default=0)
        minimum_printable = min(printable_widths)
        document_layout[label] = {
            "inline_shapes": len(doc.inline_shapes),
            "maximum_shape_inches": round(maximum_shape / 914400, 3),
            "minimum_printable_inches": round(minimum_printable / 914400, 3),
        }
        check(f"{label} inline figures fit printable page width", maximum_shape <= minimum_printable, document_layout[label])

    frame = pd.DataFrame(checks)
    checklist = REVISION / "Equal_size_registry_composition_final_checklist.csv"
    frame.to_csv(checklist, index=False)
    result = {
        "status": "complete" if frame["passed"].all() else "failed",
        "passed": int(frame["passed"].sum()),
        "total": len(frame),
        "failed_checks": frame.loc[~frame["passed"], "check"].tolist(),
        "document_layout": document_layout,
        "key_hashes": {
            "english_manuscript": sha256(english),
            "tracked_manuscript": sha256(tracked),
            "supplementary_workbook": sha256(workbook),
            "figure7_svg": sha256(figure_files[1]),
            "analysis_audit": sha256(ANALYSIS / "equal_size_registry_composition_audit.json"),
        },
    }
    (REVISION / "Equal_size_registry_composition_final_audit.json").write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False, indent=2))
    if result["status"] != "complete":
        raise SystemExit(1)


if __name__ == "__main__":
    main()
