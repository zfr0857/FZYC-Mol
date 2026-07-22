from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
IN_MD = DOCS / "初稿-4.md"
OUT_DOCX = DOCS / "初稿-4.docx"


def set_run_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_font(paragraph, size: float = 10.5, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text.strip())
    set_run_font(run, size=8.0, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def clear_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is not None:
        tc_pr.remove(borders)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tc_pr.append(borders)


def set_cell_border(cell, edge: str, size: int = 8) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    elem = borders.find(qn(f"w:{edge}"))
    if elem is None:
        elem = OxmlElement(f"w:{edge}")
        borders.append(elem)
    elem.set(qn("w:val"), "single")
    elem.set(qn("w:sz"), str(size))
    elem.set(qn("w:space"), "0")
    elem.set(qn("w:color"), "000000")


def apply_three_line_style(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    rows = table.rows
    if not rows:
        return
    for row in rows:
        for cell in row.cells:
            clear_cell_borders(cell)
    for cell in rows[0].cells:
        set_cell_border(cell, "top", 10)
        set_cell_border(cell, "bottom", 8)
    for cell in rows[-1].cells:
        set_cell_border(cell, "bottom", 10)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows: list[list[str]] = []
    for idx, line in enumerate(table_lines):
        parts = [p.strip() for p in line.strip("|").split("|")]
        if idx == 1 and all(re.fullmatch(r":?-{3,}:?", p or "---") for p in parts):
            continue
        rows.append(parts)
    max_cols = max((len(r) for r in rows), default=0)
    for row in rows:
        row.extend([""] * (max_cols - len(row)))
    return rows, i


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), value, bold=(r_idx == 0))
    apply_three_line_style(table)


def add_image(doc: Document, line: str) -> bool:
    match = re.match(r"!\[[^\]]*\]\((.*?)\)", line.strip())
    if not match:
        return False
    image_path = Path(match.group(1))
    if not image_path.exists():
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"[图像文件未找到：{image_path}]")
        set_run_font(run, size=9)
        return True
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(str(image_path), width=Cm(14.5))
    except Exception:
        run = paragraph.add_run(f"[图像插入失败：{image_path}]")
        set_run_font(run, size=9)
    return True


def add_paragraph_with_inline(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run)


def build_docx() -> None:
    markdown = IN_MD.read_text(encoding="utf-8")
    lines = markdown.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    i = 0
    title_written = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        if add_image(doc, stripped):
            i += 1
            continue
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            heading_level = min(level, 3)
            paragraph = doc.add_heading(text, level=heading_level)
            set_paragraph_font(paragraph, size=14 if heading_level == 1 else 12, bold=True)
            i += 1
            continue
        if not title_written:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped)
            set_run_font(run, size=16, bold=True)
            title_written = True
            i += 1
            continue
        add_paragraph_with_inline(doc, stripped)
        i += 1

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    build_docx()
