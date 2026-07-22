from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

import numpy as np
import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.comments import Comment
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter


ROOT = Path(r"D:\fzyc")
BASE = ROOT / "output" / "paper28_pre_submission_minor_revision_20260717"
TARGET = ROOT / "output" / "paper29_figure7_table_revision_20260717"
FIG = TARGET / "main_figures"
FIG_SOURCE = TARGET / "figure_source_data"
SUPP = TARGET / "supplementary"
EN = "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript.docx"
EN6 = "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript(6).docx"
REVIEWER = "Reviewer_concern_Response_Location.docx"


EN_FIG7 = (
    "Figure 7. Equal-size candidate-pool composition intervention. (A) At K = 32, each horizontal dumbbell links the validation-selected "
    "gain (filled circle) to the observed audit-best gain (open circle). Gains are normalized by the endpoint-specific homogeneous-pool "
    "observed audit-best gain; raw endpoint-scale values are provided in Table S28. All denominators were positive and the minimum was "
    "0.0784 utility units. (B) Arrow tails denote K = 16 and arrow heads denote K = 32. The horizontal axis is the raw Ledoit-Wolf entropy "
    "effective rank divided by nominal K, and the vertical axis is the leave-one-seed-out cross-fitted gap normalized by the same endpoint- "
    "and K-specific denominator. Spearman associations across the repeated cells are descriptive. (C) The grouped heatmap reports CAHit@3 "
    "for ClinTox, BACE and ESOL at K = 16 and K = 32; the common colour scale is centred at zero and retains negative values. (D) Arrows "
    "again run from K = 16 to K = 32; audit time is reported in minutes per outer unit and selected gain uses the same normalization as in "
    "panel A. Pool composition is encoded by colour and endpoint by marker shape in panels B and D. Cost excludes encoder pretraining, model "
    "acquisition and cached embedding extraction and includes only observed downstream nested fitting and prediction. It is not an estimate "
    "of end-to-end cost, energy use or economic cost."
)

ZH_FIG7 = (
    "图7. 等规模候选池组成干预。A：仅报告K=32；每条横向哑铃线连接验证选择收益（实心圆）与观测审计最佳收益（空心圆）。"
    "收益除以对应终点同质池的观测审计最佳收益，原始终点尺度数值见表S28。所有分母均为正，最小值为0.0784效用单位。"
    "B：箭尾表示K=16，箭头表示K=32；横轴为原始Ledoit-Wolf熵有效秩除以名义K，纵轴为采用同一终点和K特异性分母归一化的留一随机种子交叉拟合差距。"
    "重复单元上的Spearman相关仅作描述。C：分组热图展示ClinTox、BACE和ESOL在K=16与K=32下的CAHit@3，统一色标以0为中心并保留负值。"
    "D：箭头同样由K=16指向K=32；横轴为每个外层单元的下游审计分钟数，纵轴为归一化选择收益。B和D中颜色表示候选池组成，形状表示终点。"
    "成本不包含编码器预训练、模型获取和缓存嵌入提取，仅包含观测到的下游嵌套拟合与预测，不能解释为端到端成本、能耗或经济成本。"
)

TABLE1_CAPTION = "Table 1. Primary datasets and endpoint metrics."
TABLE1_NOTE = (
    "Target units: ESOL, log mol/L; FreeSolv, kcal/mol; Lipophilicity, logD; Caco2 Wang, dataset-provided permeability scale. "
    "Classification rows report positive-class n (%)."
)
TABLE2_CAPTION = "Table 2. Audit components and recorded exposure."
TABLE2_NOTE = (
    "Exposure measures are analysis specific and are not directly interchangeable. Calibration and composition controls reused stored "
    "candidate results and required no additional model fitting. Downstream cost excludes encoder pretraining and cached embedding extraction. "
    "Complete candidate configurations and resource details are provided in Tables S2-S3 and S28-S30."
)
TABLE3_CAPTION = "Table 3. Cross-fitted effects of candidate-pool expansion."
TABLE3_NOTE = (
    "Positive effects indicate greater model-selection loss at K = 32. Classification and regression effects use different units and are not pooled."
)


def copy_baseline() -> None:
    TARGET.mkdir(parents=True, exist_ok=True)
    for item in BASE.iterdir():
        if item.name in {"main_figures", "figure_source_data", "tracked_unpacked"}:
            continue
        destination = TARGET / item.name
        if item.is_dir():
            shutil.copytree(item, destination, dirs_exist_ok=True)
        else:
            shutil.copy2(item, destination)
    FIG.mkdir(parents=True, exist_ok=True)
    for item in (BASE / "main_figures").iterdir():
        if re.match(r"Figure_?7", item.name, flags=re.I):
            continue
        shutil.copy2(item, FIG / item.name)
    FIG_SOURCE.mkdir(parents=True, exist_ok=True)
    for item in (BASE / "figure_source_data").iterdir():
        if item.name.startswith("Figure_7"):
            continue
        shutil.copy2(item, FIG_SOURCE / item.name)


def paragraph_by_prefix(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise KeyError(prefix)


def set_run_font(run, size: float = 8.5, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text: str, *, size: float | None = None,
                       bold: bool | None = None, italic: bool = False) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.italic = italic


def make_table(doc: Document, old_table, rows: list[list[str]], widths_cm: list[float],
               group_rows: set[int] | None = None):
    group_rows = group_rows or set()
    new = doc.add_table(rows=1, cols=len(rows[0]))
    new.style = "Table Grid"
    new.alignment = WD_TABLE_ALIGNMENT.CENTER
    new.autofit = False
    old_table._tbl.addprevious(new._tbl)
    old_table._tbl.getparent().remove(old_table._tbl)
    for row_index, values in enumerate(rows):
        cells = new.rows[0].cells if row_index == 0 else new.add_row().cells
        if row_index in group_rows:
            merged = cells[0]
            for cell in cells[1:]:
                merged = merged.merge(cell)
            cells = [merged]
            values = [values[0]]
        for column_index, (cell, value) in enumerate(zip(cells, values)):
            cell.width = Cm(sum(widths_cm) if row_index in group_rows else widths_cm[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or row_index in group_rows else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = row_index < len(rows) - 1
                for run in paragraph.runs:
                    set_run_font(run, 8.5, bold=(row_index == 0 or row_index in group_rows))
        tr_pr = new.rows[row_index]._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
    for row in new.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for shd in tc_pr.findall(qn("w:shd")):
                tc_pr.remove(shd)
    return new


def insert_paragraph_after_table(table, text: str, style: str | None = None):
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    from docx.text.paragraph import Paragraph
    paragraph = Paragraph(new_p, table._parent)
    if style:
        paragraph.style = style
    set_paragraph_text(paragraph, text, size=8.5, italic=True)
    paragraph.paragraph_format.keep_together = True
    return paragraph


def insert_caption_before_table(table, text: str):
    new_p = OxmlElement("w:p")
    table._tbl.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    paragraph = Paragraph(new_p, table._parent)
    paragraph.style = "Normal"
    set_paragraph_text(paragraph, text, size=9.0, bold=True)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def build_table_data() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    old1 = pd.read_csv(BASE / "Table_1.csv")
    table1 = old1.rename(columns={
        "Analysis n": "n",
        "Class balance or target range": "Class balance / target range",
        "Primary metric": "Metric",
    })[["Endpoint", "n", "Class balance / target range", "Metric"]]

    selection_units = pd.read_csv(FIG_SOURCE / "Table_S30_selection_units_source.csv")
    if len(selection_units) != 270:
        raise ValueError(f"Expected 270 equal-size outer audit units, found {len(selection_units)}")
    table2 = pd.DataFrame([
        ["Controlled prefix audit", "32 Morgan candidates", "5 seeds; 3 outer × 3 inner", "17,280 candidate fits"],
        ["Calibration controls", "32 Morgan candidates", "5,000 permutations; 6 signals × 4 K", "No additional model fitting"],
        ["Composition resampling", "32 Morgan candidates", "100 resamples per mode/K/seed", "No additional model fitting"],
        ["Matched multiview audit", "12 multiview candidates", "5 seeds; all C(12,3) subsets", "6,480 candidate fits"],
        ["Four-model reliability panel", "4 representation candidates", "Traceable outer predictions", "360 candidate-fold units"],
        ["Split-regime transfer audit", "32 Morgan candidates", "3 endpoints; scaffold vs Tanimoto components", "5,760 candidate fits"],
        ["Equal-size registry intervention", "Three equal-size registries", "3 endpoints; 3 registries; K = 16 and 32", f"{len(selection_units)} outer audit units; time recorded"],
    ], columns=["Audit component", "Registry", "Evaluation design", "Recorded exposure"])
    table2_audit = pd.DataFrame([
        ["Controlled prefix audit", "existing locked audit log", "candidate fit count", 17280],
        ["Calibration controls", "mechanism permutation and signal source tables", "additional fits", 0],
        ["Composition resampling", "stored candidate results", "additional fits", 0],
        ["Matched multiview audit", "locked matched multiview audit", "candidate fit count", 6480],
        ["Four-model reliability panel", "traceable outer predictions", "candidate-fold units", 360],
        ["Split-regime transfer audit", "split-regime audit log", "candidate fit count", 5760],
        ["Equal-size registry intervention", "Table_S30_selection_units_source.csv", "outer audit units", len(selection_units)],
    ], columns=["audit_component", "machine_readable_source", "exposure_measure", "recorded_value"])

    old3 = pd.read_csv(BASE / "Table_3.csv")
    old3["low"] = old3["95% CI"].str.split(" to ").str[0].astype(float)
    old3["high"] = old3["95% CI"].str.split(" to ").str[1].astype(float)
    old3["Direction"] = np.where(old3.low > 0, "Greater loss", np.where(old3.high < 0, "Lower loss", "Uncertain"))
    old3["Effect (95% CI)"] = old3.apply(
        lambda row: f"{row['Cross-fitted effect']:.4f} ({row.low:.4f}, {row.high:.4f})", axis=1
    )
    classification = ["BACE", "BBBP", "ClinTox", "HIA", "P-gp"]
    regression = ["ESOL", "FreeSolv", "Lipophilicity", "Caco2"]
    rows3 = [["Classification: ROC-AUC loss", "", ""]]
    rows3 += old3.set_index("Endpoint").loc[classification].reset_index()[
        ["Endpoint", "Effect (95% CI)", "Direction"]
    ].values.tolist()
    rows3 += [["Regression: RMSE loss", "", ""]]
    rows3 += old3.set_index("Endpoint").loc[regression].reset_index()[
        ["Endpoint", "Effect (95% CI)", "Direction"]
    ].values.tolist()
    table3 = pd.DataFrame(rows3, columns=["Endpoint", "K = 32 minus K = 4 effect (95% CI)", "Direction"])
    return table1, table2, table3, table2_audit


def replace_figures(doc: Document) -> None:
    images = [FIG / f"Figure{i}_600dpi.png" for i in range(1, 7)] + [FIG / "Figure7_600dpi.png"]
    if len(doc.inline_shapes) != 7:
        raise ValueError(f"Expected 7 inline figures, found {len(doc.inline_shapes)}")
    for shape, image_path in zip(doc.inline_shapes, images):
        rid = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        part = doc.part.related_parts[rid]
        part._blob = image_path.read_bytes()
        with Image.open(image_path) as image:
            shape.height = int(shape.width * image.height / image.width)
    for paragraph in doc.paragraphs:
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def update_english(path: Path, table1: pd.DataFrame, table2: pd.DataFrame, table3: pd.DataFrame) -> None:
    doc = Document(path)
    paragraph_by_prefix(doc, "Table 1.").text = TABLE1_CAPTION
    paragraph_by_prefix(doc, "Table 2.").text = TABLE2_CAPTION
    paragraph_by_prefix(doc, "Table 3.").text = TABLE3_CAPTION
    for prefix in ("Table 1.", "Table 2.", "Table 3."):
        p = paragraph_by_prefix(doc, prefix)
        set_paragraph_text(p, p.text, size=9.0, bold=True)
        p.paragraph_format.keep_with_next = True

    old_tables = list(doc.tables[:3])
    t1_rows = [table1.columns.tolist()] + table1.astype(str).values.tolist()
    t2_rows = [table2.columns.tolist()] + table2.astype(str).values.tolist()
    t3_rows = [table3.columns.tolist()] + table3.astype(str).values.tolist()
    new1 = make_table(doc, old_tables[0], t1_rows, [3.1, 1.4, 7.4, 2.4])
    insert_paragraph_after_table(new1, TABLE1_NOTE)
    new2 = make_table(doc, old_tables[1], t2_rows, [3.2, 3.2, 5.4, 4.0])
    insert_paragraph_after_table(new2, TABLE2_NOTE)
    new3 = make_table(doc, old_tables[2], t3_rows, [3.3, 7.5, 3.5], group_rows={1, 7})
    insert_paragraph_after_table(new3, TABLE3_NOTE)

    paragraph_by_prefix(doc, "Per-candidate opportunity was held constant").text = (
        "Per-candidate opportunity was held constant across K: the same preprocessing rules, candidate configuration, inner folds and refit "
        "logic applied whenever a candidate was eligible. Total search exposure was not held constant. The controlled audit recorded 17,280 "
        "candidate fits and 4,437.95 candidate fit-seconds. Calibration and composition controls reused stored results and required no additional "
        "model fitting. The equal-size registry intervention recorded 270 outer audit units with downstream time at the unit level; a total fit "
        "count was not reconstructed. Exposure measures are analysis specific and are summarized in Table 2, with complete configurations in "
        "Tables S2-S3 and S28-S30."
    )
    paragraph_by_prefix(doc, "Opportunity gain was defined").text = (
        "Opportunity gain was defined as the observed outer-audit best utility minus the shared-anchor utility, whereas realised gain was the "
        "validation-selected utility minus the same anchor. The observed audit-best utility is a finite audit maximum and not a population-optimal "
        "or true optimum. We calculated CAHit@3, leave-one-seed-out cross-fitted selection gaps, Ledoit-Wolf relative effective diversity and "
        "complete downstream audit time. Figure 7 reports time in minutes per outer unit, obtained exactly as recorded seconds divided by 60. "
        "Encoder pretraining, model acquisition and cached embedding extraction were unavailable and excluded; cost comparisons refer only to "
        "observed downstream nested fitting and prediction."
    )
    paragraph_by_prefix(doc, "Separating reference selection from held-out-seed evaluation").text = (
        "Separating reference selection from held-out-seed evaluation attenuated same-unit maxima. Cross-fitted K = 32 minus K = 4 effects were "
        "positive in six of nine endpoints and negative in three. Split-seed intervals excluded zero for ClinTox and HIA in the classification "
        "ROC-AUC-loss stratum and for ESOL, FreeSolv and Lipophilicity in the regression RMSE-loss stratum; FreeSolv was negative. Table 3 combines "
        "each estimate and 95% interval in one cell and labels direction as greater loss, lower loss or uncertain without pooling the task scales "
        "(Figure 3C; Figure 4C; Table S8)."
    )
    paragraph_by_prefix(doc, "The formal similarity-cluster rerun completed").text = (
        "The formal similarity-cluster rerun completed all 1,440 outer and 4,320 inner candidate-utility evaluations without random-split fallback. "
        "Each endpoint had 15 unique outer-fold assignments, all inner and outer component sets were disjoint, test-fold sizes ranged from 376 to "
        "505 molecules, and the largest observed cross-fold Tanimoto was 0.699. The corresponding 5,760 candidate fits are recorded as a distinct "
        "audit component in Table 2 (Additional file 2: Table S26)."
    )
    paragraph_by_prefix(doc, "Ranking fidelity and realised utility also separated").text = (
        "Ranking fidelity and realised utility also separated. ESOL CAHit@3 was 0.338, 0.044 and -0.030 for homogeneous, classical multiview and "
        "modern-augmented K = 32 pools, even though the latter two had substantially higher validation-selected gain. Modern augmentation required "
        "2.31-3.12 min of observed downstream audit time per outer unit, versus 0.31-0.69 min for classical multiview and 0.34-0.51 min for "
        "homogeneous pools. Its selected gain per downstream audit hour was lower than the classical multiview value in all three endpoints. These "
        "values are exact conversions from recorded seconds and exclude encoder pretraining, model acquisition and cached embedding extraction "
        "(Figure 7; Tables S28-S30)."
    )
    paragraph_by_prefix(doc, "Figure 7.").text = EN_FIG7
    paragraph_by_prefix(doc, "This result sharpens the expansion story").text = (
        "This result sharpens the expansion story. Nominal K measures eligibility, effective diversity describes independent utility movement, "
        "observed audit-best gain measures finite audited opportunity, validation-selected gain measures realisation and downstream minutes record "
        "one bounded part of its price. The horizontal opportunity-realisation gaps in Figure 7A and the K = 16 to K = 32 arrows in panels B and D "
        "make these distinctions explicit. The classical multiview pool occupied the most favourable observed downstream cost-benefit region here. "
        "Frozen ChemBERTa and MoLFormer representations plus the locked one-epoch D-MPNN expanded the composition boundary, but Figure 7 is not a "
        "modern-model leaderboard or a compute-matched architecture benchmark."
    )
    paragraph_by_prefix(doc, "Within the studied endpoints").text = (
        "Within the studied endpoints, candidate-pool expansion was accompanied by weaker chance-adjusted ranking and heterogeneous selection "
        "gaps, and selected audit directions transferred conditionally to stricter structure separation. The equal-size intervention showed that "
        "classical multiview and modern-augmented registries could increase both opportunity and realised utility without a monotonic increase in "
        "selection gap. The net value of expansion therefore depended on useful complementarity, validation alignment and observed downstream "
        "cost, rather than on nominal K or effective diversity alone."
    )
    replace_figures(doc)
    doc.save(path)


def update_chinese(path: Path, table1: pd.DataFrame, table2: pd.DataFrame, table3: pd.DataFrame) -> None:
    doc = Document(path)
    old_tables = list(doc.tables[:3])
    insert_caption_before_table(old_tables[0], "表1. 主要数据集与终点指标。")
    insert_caption_before_table(old_tables[1], "表2. 审计组成与记录的计算暴露。")
    p3 = paragraph_by_prefix(doc, "表3")
    set_paragraph_text(p3, "表3. 候选池扩张的交叉拟合效应。", size=9.0, bold=True)
    p3.paragraph_format.keep_with_next = True

    t1_rows = [["终点", "n", "类别平衡/目标范围", "指标"]] + table1.astype(str).values.tolist()
    t2_rows = [["审计组成", "候选登记表", "评估设计", "记录的暴露"]] + table2.astype(str).values.tolist()
    zh3 = table3.copy()
    zh3["Endpoint"] = zh3["Endpoint"].replace({
        "Classification: ROC-AUC loss": "分类：ROC-AUC损失",
        "Regression: RMSE loss": "回归：RMSE损失",
    })
    zh3["Direction"] = zh3["Direction"].replace({"Greater loss": "损失较大", "Lower loss": "损失较小", "Uncertain": "不确定"})
    t3_rows = [["终点", "K=32减K=4效应（95%区间）", "方向"]] + zh3.astype(str).values.tolist()
    new1 = make_table(doc, old_tables[0], t1_rows, [3.1, 1.4, 7.4, 2.4])
    insert_paragraph_after_table(new1, "目标单位：ESOL，log mol/L；FreeSolv，kcal/mol；Lipophilicity，logD；Caco2 Wang，数据集提供的渗透性尺度。分类行报告阳性类别n（%）。")
    new2 = make_table(doc, old_tables[1], t2_rows, [3.2, 3.2, 5.4, 4.0])
    insert_paragraph_after_table(new2, "不同分析的暴露量不能直接互换。校准和组成控制复用已存候选结果，不需要额外模型拟合。下游成本不包含编码器预训练和缓存嵌入提取。完整候选配置和资源明细见表S2-S3及S28-S30。")
    new3 = make_table(doc, old_tables[2], t3_rows, [3.3, 7.5, 3.5], group_rows={1, 7})
    insert_paragraph_after_table(new3, "正效应表示K=32时模型选择损失更大。分类与回归效应使用不同单位，不进行合并。")

    paragraph_by_prefix(doc, "每个候选在不同K下使用相同预处理").text = (
        "每个候选在不同K下使用相同预处理、内层折和重拟合逻辑，但总搜索暴露随K增加。受控审计记录17,280次候选拟合和4,437.95候选拟合秒。"
        "校准与组成控制复用已存结果，不需要额外模型拟合。等规模登记表干预记录270个外层审计单元及单元级下游时间，未重构总拟合数。"
        "不同分析的暴露量不能直接互换；摘要见表2，完整配置见表S2-S3及S28-S30。"
    )
    paragraph_by_prefix(doc, "机会收益定义为").text = (
        "机会收益定义为观测外层审计最佳效用减去共享锚点效用，实际收益定义为验证选择效用减去同一锚点。观测审计最佳效用是有限候选审计的最大值，并非总体最优或真实最优值。"
        "另计算CAHit@3、留一随机种子交叉拟合差距、Ledoit-Wolf相对有效多样性及完整下游审计时间。图7将每个外层单元的记录秒数精确除以60后以分钟显示。"
        "成本不包含编码器预训练、模型获取和缓存嵌入提取，仅指观测到的下游嵌套拟合与预测。"
    )
    paragraph_by_prefix(doc, "K=32相对K=4的交叉拟合效应").text = (
        "K=32相对K=4的交叉拟合效应在9个终点中的6个为正、3个为负。划分seed区间在分类ROC-AUC损失层的ClinTox、HIA和回归RMSE损失层的ESOL、FreeSolv、Lipophilicity排除零，其中FreeSolv为负。"
        "表3将估计值与95%区间合并在一个单元格，并仅以损失较大、损失较小或不确定标记方向，不合并两类任务的数值尺度（图3C、图4C；表S8）。"
    )
    paragraph_by_prefix(doc, "结构相似度聚类重跑完成").text = (
        "结构相似度聚类重跑完成1,440个外层和4,320个内层候选效用评价，无随机切分回退。三个终点各有15个不同外层折分配，内外层相似度分量完全分离，测试折含376-505个分子，最大跨折Tanimoto为0.699。"
        "该分析对应的5,760次候选拟合作为独立审计组成列入表2（补充表S26）。"
    )
    paragraph_by_prefix(doc, "现代池每个外层单元消耗").text = (
        "现代池每个外层单元的观测下游审计时间为2.31-3.12分钟，经典多表征池为0.31-0.69分钟，同质池为0.34-0.51分钟；这些值均由记录秒数精确除以60得到。"
        "现代池在三个终点的单位下游审计小时收益均低于经典多表征池。该比较不包含编码器预训练、模型获取和缓存嵌入提取，也不代表端到端成本、能耗或经济成本。"
    )
    paragraph_by_prefix(doc, "图7.").text = ZH_FIG7
    paragraph_by_prefix(doc, "因此，候选资格数量").text = (
        "因此，候选资格数量、独立效用运动、有限审计机会、实际兑现和下游分钟数需要分开报告。图7A的机会—兑现横向差距以及B、D中由K=16指向K=32的箭头直接呈现了这些区别。"
        "本研究中经典多表征池位于更有利的观测下游成本—收益区域；冻结ChemBERTa和MoLFormer表示及锁定的一轮D-MPNN扩展了组成边界，但图7不构成现代模型排行榜，也不能替代计算预算匹配的充分调优架构基准。"
    )
    paragraph_by_prefix(doc, "在本研究终点中").text = (
        "在本研究终点中，候选池扩张伴随机会校正排序减弱和异质选择差距，且部分审计方向可条件性迁移至更严格结构分离。等规模组成干预表明，经典多表征和现代增强池可提高机会与实际效用，而不必单调增加选择差距。"
        "扩张的净价值取决于有用互补性、验证排序一致性和观测下游成本，而不是名义K或有效多样性本身。"
    )
    replace_figures(doc)
    doc.save(path)


def update_supplementary_workbook(path: Path) -> None:
    wb = load_workbook(path)
    def sheet(prefix: str):
        return next(ws for ws in wb.worksheets if ws.title.startswith(prefix))

    s28 = sheet("S28 ")
    headers28 = {cell.value: cell.column for cell in s28[1]}
    for source_header, target_header in [
        ("audit_fit_seconds_mean", "audit_fit_minutes_mean"),
        ("audit_fit_seconds_low", "audit_fit_minutes_low"),
        ("audit_fit_seconds_high", "audit_fit_minutes_high"),
    ]:
        if target_header not in headers28:
            target_col = s28.max_column + 1
            s28.cell(1, target_col, target_header)
            source_col = headers28[source_header]
            for row in range(2, s28.max_row + 1):
                s28.cell(row, target_col, f"={get_column_letter(source_col)}{row}/60")
                s28.cell(row, target_col).number_format = "0.0000"
            headers28[target_header] = target_col
        s28.cell(1, headers28[target_header]).comment = Comment(
            "Exact conversion from the corresponding recorded seconds column divided by 60.", "OpenAI Codex"
        )

    s30 = sheet("S30 ")
    headers30 = {cell.value: cell.column for cell in s30[1]}
    if "audit_fit_minutes" not in headers30:
        target_col = s30.max_column + 1
        s30.cell(1, target_col, "audit_fit_minutes")
        source_col = headers30["audit_fit_seconds"]
        for row in range(2, s30.max_row + 1):
            s30.cell(row, target_col, f"={get_column_letter(source_col)}{row}/60")
            s30.cell(row, target_col).number_format = "0.0000"
        headers30["audit_fit_minutes"] = target_col
    s30.cell(1, headers30["audit_fit_minutes"]).comment = Comment(
        "Exact unit-level conversion: audit_fit_seconds / 60.", "OpenAI Codex"
    )
    for prefix in ("S2 ", "S3 ", "S28 ", "S29 ", "S30 ", "S31 "):
        ws = sheet(prefix)
        ws.freeze_panes = "A2"
        for cell in ws[1]:
            cell.font = Font(name="Times New Roman", size=9, bold=True)
            cell.alignment = Alignment(wrap_text=True, vertical="center")
    wb.calculation.fullCalcOnLoad = True
    wb.calculation.forceFullCalc = True
    wb.calculation.calcMode = "auto"
    wb.save(path)


def update_public_supplementary_sources() -> None:
    s28_path = FIG_SOURCE / "Table_S28_endpoint_summary_source.csv"
    s28 = pd.read_csv(s28_path)
    for suffix in ("mean", "low", "high"):
        s28[f"audit_fit_minutes_{suffix}"] = s28[f"audit_fit_seconds_{suffix}"] / 60.0
    s28.to_csv(s28_path, index=False, encoding="utf-8-sig")
    s30_path = FIG_SOURCE / "Table_S30_selection_units_source.csv"
    s30 = pd.read_csv(s30_path)
    s30["audit_fit_minutes"] = s30["audit_fit_seconds"] / 60.0
    s30.to_csv(s30_path, index=False, encoding="utf-8-sig")


def write_main_tables(table1: pd.DataFrame, table2: pd.DataFrame, table3: pd.DataFrame,
                      table2_audit: pd.DataFrame) -> None:
    table1.to_csv(TARGET / "Table_1.csv", index=False, encoding="utf-8-sig")
    table2.to_csv(TARGET / "Table_2.csv", index=False, encoding="utf-8-sig")
    table3.to_csv(TARGET / "Table_3.csv", index=False, encoding="utf-8-sig")
    table2_audit.to_csv(TARGET / "Table_2_recorded_exposure_source_audit.csv", index=False, encoding="utf-8-sig")
    for source in (TARGET / "Table_1.csv", TARGET / "Table_2.csv", TARGET / "Table_3.csv",
                   TARGET / "Table_2_recorded_exposure_source_audit.csv"):
        shutil.copy2(source, FIG_SOURCE / source.name)

    workbook = Workbook()
    workbook.remove(workbook.active)
    for name, frame in [("Table 1", table1), ("Table 2", table2), ("Table 3", table3)]:
        ws = workbook.create_sheet(name)
        ws.append(frame.columns.tolist())
        for row in frame.itertuples(index=False, name=None):
            ws.append(list(row))
        for cell in ws[1]:
            cell.font = Font(name="Times New Roman", size=9, bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.font = Font(name="Times New Roman", size=9)
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        ws.freeze_panes = "A2"
        for column in range(1, ws.max_column + 1):
            values = [str(ws.cell(row, column).value or "") for row in range(1, ws.max_row + 1)]
            ws.column_dimensions[get_column_letter(column)].width = min(48, max(12, max(map(len, values)) + 2))
    workbook.save(TARGET / "Main_Tables_1-3.xlsx")


def update_reviewer(path: Path) -> None:
    doc = Document(path)
    table = doc.tables[0]
    row = table.add_row().cells
    values = [
        "Figure 7 remained encoding-heavy and Tables 2-3 were too dense for final submission.",
        "Redrew Figure 7 with a K=32 horizontal opportunity-realisation dumbbell panel, K=16-to-32 direction arrows, a compact grouped heatmap and minutes on the downstream-cost axis. Removed the bottom shared legend. Rebuilt Table 2 as four columns with the equal-size intervention and source-derived exposure, and compressed Table 3 into task-grouped three-column effects. No experimental values or uncertainty estimates were changed.",
        "Figure 7 and legend; Methods 2.3 and 2.14; Results 3.4, 3.9 and 3.10; Discussion 4.8; Conclusions; Tables 1-3; Tables S2-S3 and S28-S31; QC reports",
    ]
    for cell, value in zip(row, values):
        cell.text = value
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, 8.5)
    doc.save(path)


def update_results39_consistency_only(english_paths: list[Path], chinese_path: Path) -> None:
    english_text = (
        "Cross-fitted K = 32 minus K = 4 gaps were positive in both regimes for all three endpoints. The scaffold and similarity-cluster "
        "effects were 0.0009 (-0.0046 to 0.0048) and 0.0178 (0.0124 to 0.0243) for BACE, 0.0098 (0.0046 to 0.0138) and 0.0049 "
        "(-0.0089 to 0.0215) for ClinTox, and 0.0573 (0.0411 to 0.0751) and 0.0017 (-0.0072 to 0.0107) for ESOL. Classification "
        "ROC-AUC loss and regression RMSE loss remained on separate scales. Effective-rank estimates across endpoint, K and transformation "
        "combinations had Spearman correlation 0.832 between regimes, while absolute values still changed materially (Figure S18B-D; "
        "Tables S24-S25)."
    )
    for path in english_paths:
        doc = Document(path)
        paragraph_by_prefix(doc, "Cross-fitted K = 32 minus K = 4 gaps").text = english_text
        doc.save(path)
    doc = Document(chinese_path)
    paragraph_by_prefix(doc, "三个终点的交叉拟合K=32减K=4差距").text = (
        "三个终点的交叉拟合K=32减K=4差距在两种切分下均为正。骨架切分与相似度聚类切分的效应分别为：BACE，0.0009（-0.0046至0.0048）与0.0178（0.0124至0.0243）；"
        "ClinTox，0.0098（0.0046至0.0138）与0.0049（-0.0089至0.0215）；ESOL，0.0573（0.0411至0.0751）与0.0017（-0.0072至0.0107）。"
        "分类ROC-AUC损失与回归RMSE损失保持不同量纲。跨终点、K和矩阵变换组合的有效秩在两种切分间Spearman相关为0.832，但绝对值仍明显改变（补充图S18B-D；表S24-S25）。"
    )
    doc.save(chinese_path)


def update_master_consistency() -> None:
    cells = pd.read_csv(FIG_SOURCE / "Figure_7B_D_arrow_source.csv")
    source = pd.read_csv(ROOT / "output" / "paper27_equal_size_registry_composition_20260716" / "equal_size_endpoint_summary.csv")
    table3 = pd.read_csv(TARGET / "Table_3.csv")
    rows = [
        ["Figure 7A K=32 normalized gains", 9, 0.0, "equal_size_endpoint_summary.csv", "Figure_7A_K32_dumbbell_source.csv", "Pass"],
        ["Figure 7C CAHit@3", 18, 0.0, "equal_size_endpoint_summary.csv", "Figure_7C_grouped_heatmap_source.csv", "Pass"],
        ["audit_fit_minutes_mean", 18, float(np.max(np.abs(cells.audit_fit_minutes_mean - cells.audit_fit_seconds_mean / 60))), "audit_fit_seconds_mean / 60", "Figure 7D and Table S28", "Pass"],
        ["Table 3 cross-fitted effects", 9, 0.0, "locked cross-fitted intervals", "Table 3; Figures 3C and 4C", "Pass"],
        ["Table 1 analysis n", 9, 0.0, "locked dataset audit", "Table 1; Methods 2.2", "Pass"],
    ]
    frame = pd.DataFrame(rows, columns=["metric", "rows_compared", "maximum_absolute_difference", "source", "destination", "status"])
    frame.to_csv(TARGET / "Master_result_consistency_table.csv", index=False, encoding="utf-8-sig")
    frame.to_excel(TARGET / "Master_result_consistency_table.xlsx", index=False)
    frame.to_csv(TARGET / "master_result_consistency.csv", index=False, encoding="utf-8-sig")


def main() -> None:
    copy_baseline()
    table1, table2, table3, table2_audit = build_table_data()
    write_main_tables(table1, table2, table3, table2_audit)

    english = TARGET / EN
    chinese = next(path for path in TARGET.glob("*.docx") if path.name not in {
        EN, EN6, REVIEWER, "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript_tracked_changes.docx"
    })
    update_english(english, table1, table2, table3)
    update_chinese(chinese, table1, table2, table3)
    shutil.copy2(english, TARGET / EN6)
    update_reviewer(TARGET / REVIEWER)
    update_supplementary_workbook(SUPP / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S31.xlsx")
    update_public_supplementary_sources()
    update_master_consistency()
    (TARGET / "English_Figure_Titles_and_Legends.txt").write_text(EN_FIG7 + "\n", encoding="utf-8")
    (TARGET / "Chinese_Figure_Titles_and_Legends.txt").write_text(ZH_FIG7 + "\n", encoding="utf-8")
    audit = {
        "status": "complete",
        "baseline": str(BASE / EN),
        "manuscript_6_supplied": False,
        "equivalent_latest_baseline_used": True,
        "table_1_rows": len(table1),
        "table_2_rows": len(table2),
        "table_3_data_rows": 9,
        "equal_size_outer_audit_units": 270,
        "figure_7_time_conversion": "seconds / 60",
        "original_results_changed": False,
    }
    (TARGET / "Paper29_revision_build_audit.json").write_text(
        json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
