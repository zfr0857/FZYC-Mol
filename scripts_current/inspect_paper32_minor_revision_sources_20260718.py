from pathlib import Path

from docx import Document


ROOT = Path(r"D:\fzyc\work\paper33_final_minor_revision_20260718")
ROOT.mkdir(parents=True, exist_ok=True)
DOCS = {
    "en": Path(r"D:\fzyc\output\paper32_equation_table_format_20260718\Candidate_pool_expansion_Journal_of_Cheminformatics_final_unified_format.docx"),
    "zh": Path(r"C:\Users\Administrator\Desktop\Chinese_manuscript_final_Times_New_Roman_figures.docx"),
}


for language, path in DOCS.items():
    document = Document(path)
    with (ROOT / f"{language}_paragraphs.tsv").open("w", encoding="utf-8") as handle:
        handle.write("index\tstyle\ttext\n")
        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.replace("\t", " ").replace("\n", " ")
            handle.write(f"{index}\t{paragraph.style.name}\t{text}\n")
    with (ROOT / f"{language}_tables.txt").open("w", encoding="utf-8") as handle:
        for table_index, table in enumerate(document.tables, start=1):
            handle.write(f"TABLE {table_index}\n")
            for row in table.rows:
                handle.write("\t".join(cell.text.replace("\n", " | ") for cell in row.cells) + "\n")
            handle.write("\n")
    print(language, len(document.paragraphs), len(document.tables), len(document.inline_shapes))
