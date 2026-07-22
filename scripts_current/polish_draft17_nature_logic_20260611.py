from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
REPORT_DIR = ROOT / "reports" / "full_missing_experiment_run_20260611"


def locate_input() -> Path:
    candidates = [p for p in DESKTOP.rglob("*LinPept*.docx") if "17" in p.name]
    if not candidates:
        raise FileNotFoundError("Could not find draft 17 LinPept docx under Desktop.")
    return max(candidates, key=lambda p: p.stat().st_mtime)


IN_DOCX = locate_input()
OUT_DOCX = IN_DOCX.parent / "初稿-18_Nature语言逻辑终审版.docx"
REPORT_DOCX = IN_DOCX.parent / "初稿-18_Nature语言逻辑终审报告.docx"


def set_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    set_text(new_para, text)
    return new_para


def replace_exact(doc: Document, replacements: dict[str, str]) -> list[tuple[int, str]]:
    changed: list[tuple[int, str]] = []
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text in replacements:
            set_text(paragraph, replacements[text])
            changed.append((i, text[:80]))
    return changed


def polish_tables(doc: Document) -> int:
    replacements = {
        "Report as completed CycPept-PAMPA regression bRo5 stress test; combine with LinPept only at the narrative level.": "Completed CycPept-PAMPA regression bRo5 stress test; interpret separately from LinPept classification benchmarks.",
        "Report as completed bRo5 classification stress tests; cite Benchmark-ADMET-2025 and original LinPept dataset papers.": "Completed bRo5 classification stress tests; cite Benchmark-ADMET-2025 and the original LinPept dataset papers.",
        "State full available panel, not 30-task claim.": "Report as the currently available MoleculeACE panel, not as a 30-task claim.",
        "Report as controlled supplement; do not claim full fine-tuning.": "Report as a controlled frozen-encoder adapter supplement, not as full fine-tuning.",
        "Use as activity-cliff risk evidence, not causal mechanism.": "Use as activity-cliff risk evidence rather than causal mechanism evidence.",
        "公开压力测试/已补齐": "公开压力测试/已完成",
    }
    changed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text in replacements:
                    cell.text = replacements[text]
                    changed += 1
    return changed


def make_replacements() -> dict[str, str]:
    return {
        "分子性质预测已成为药物发现、ADMET 评估和毒性风险筛查中的基础计算环节，但随机划分下的平均 ROC-AUC 或 RMSE 难以充分反映模型在新骨架外推、低相似度分子、不平衡毒性标签、规则五以外化学空间和活性悬崖中的可靠性。本文提出 FZYC-Mol，一种由验证集治理驱动的适用域感知分子性质预测框架。该框架将多视图表示、强基线专家、Random/Scaffold/Perimeter 划分、类别不平衡策略、MoleculeACE 可用任务子集、粗糙度诊断、校准和适用域证据纳入同一冻结候选池；候选接受、拒绝或保留仅由验证集证据决定，测试集在策略冻结后只用于一次性最终评估。bRo5 模块已新增 CycPept-PAMPA、LinPept CellPen 和 LinPept NonFouling 三项公开压力测试。其中 CycPept-PAMPA 作为回归型渗透性任务报告 RMSE/MAE，LinPept 两项作为分类型肽性质任务报告 ROC-AUC、PR-AUC、校准和适用域分层；三者均以冻结测试集作一次性评估，不再作为缺失数据处理。": "分子性质预测已成为药物发现、ADMET 评估和毒性风险筛查中的基础计算环节。然而，随机划分下的平均 ROC-AUC 或 RMSE 难以充分反映模型在新骨架外推、低相似度分子、不平衡毒性标签、规则五以外化学空间和活性悬崖中的可靠性。本文提出 FZYC-Mol，一种由验证集治理驱动的适用域感知分子性质预测框架。该框架将多视图表示、强基线专家、Random/Scaffold/Perimeter 划分、类别不平衡策略、MoleculeACE 可用任务子集、bRo5 公开压力测试、粗糙度诊断、校准和适用域证据纳入同一冻结候选池。候选策略的接受、拒绝或保留仅由验证集证据决定，测试集仅在策略冻结后用于一次性最终评估。",
        "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别达到 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302。验证集接受的定向补救改善了 Lipophilicity；低成本重构缩小了 FreeSolv 与当前选择器之间的差距，但仍未超过观测最佳 Chemprop 候选，因此 FreeSolv 被保留为物理相互作用相关任务的边界案例。多视图融合与适用域门控在 BBBP、ClinTox 以及外部 TDC ADMET 任务 Caco2_Wang、HIA_Hou 和 Pgp_Broccatelli 上产生选择性增益；22 个外部终点的最终保留结果为 win/tie/loss = 5/17/0。上述结果将 FZYC-Mol 界定为终点依赖的选择、拒绝与审计流程，而非保证所有任务统一提升的单一模型。": "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别达到 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302。验证集接受的定向补救改善了 Lipophilicity；低成本重构缩小了 FreeSolv 与当前选择器之间的差距，但仍未超过观测最佳 Chemprop 候选，因此 FreeSolv 被保留为物理相互作用相关任务的边界案例。多视图融合与适用域门控在 BBBP、ClinTox 以及外部 TDC ADMET 任务 Caco2_Wang、HIA_Hou 和 Pgp_Broccatelli 上产生选择性增益；22 个外部终点的最终保留结果为 win/tie/loss = 5/17/0。这些结果将 FZYC-Mol 界定为终点依赖的选择、拒绝与审计流程，而不是在所有任务上统一提升的单一模型。",
        "可靠性分析显示，风险分数对分类错误的识别能力强于对回归高误差样本的识别能力；两类任务的中位 AUROC 分别为 0.788 和 0.652。保形预测在 MoleculeNet 分类任务 80%/90%/95% 目标覆盖率下的平均经验覆盖率为 0.814/0.918/0.956，在回归任务下为 0.823/0.925/0.962。验证-测试排名审计显示，跨 200 个 dataset-seed 候选池的中位 Spearman 相关为 0.667，测试最佳候选落入验证 Top-3 的比例为 0.295，Top-1 一致比例为 0.135。这些结果表明，验证集治理可以减少测试集事后选择风险，但不能保证获得测试集最优结果；所有小幅增益均需结合 regret、optimism gap、嵌套验证和负结果审计解释。基序归因与片段富集仅作为关联性化学解释，不作为因果机制证据。": "可靠性分析显示，风险分数对分类错误的识别能力强于对回归高误差样本的识别能力；两类任务的中位 AUROC 分别为 0.788 和 0.652。保形预测在 MoleculeNet 分类任务 80%/90%/95% 目标覆盖率下的平均经验覆盖率为 0.814/0.918/0.956，在回归任务下为 0.823/0.925/0.962。验证-测试排名审计显示，跨 200 个 dataset-seed 候选池的中位 Spearman 相关为 0.667，测试最佳候选落入验证 Top-3 的比例为 0.295，Top-1 一致比例为 0.135。因此，验证集治理降低了测试集事后选择风险，但并不等同于测试集最优保证；小幅增益需结合 regret、optimism gap、嵌套验证和负结果审计解释。基序归因与片段富集仅提供关联性化学解释，不作为因果机制证据。",
        "本文贡献可归纳为五个相互约束的证据层面。第一，围绕 MoleculeNet、TDC ADMET、bRo5 公开压力测试、MoleculeACE 可用任务子集和 Random/Scaffold/Perimeter 划分构建统一评测流程，并明确训练集、验证集和测试集的使用边界。第二，将 TabPFNv2、AutoGluon、KPGT representation、XGBoost-RDKit/Mordred/MorganCount、Chemprop、树模型、冻结表征、采样策略、目标变换和融合策略预先登记为候选，而非依据测试集表现临时追加。第三，提出 accept/reject/retain 治理机制，使补救头、融合、适用域门控和强基线需先通过验证集证据审查，方可进入最终保留结果。第四，报告 validation-test rank correlation、Top1 match、Top3 hit、regret、optimism gap、不平衡分类、保形预测、risk-coverage、MoleculeACE cliff-pair 指标和 ROGI/MODI/SARI 粗糙度诊断，避免仅依据单一 ROC-AUC 或 RMSE 下结论。第五，将数据缺失、候选失败、低相似度失败和活性悬崖失败写入限制与补充材料，使负结果成为证据链的一部分。": "本文贡献体现在五个层面：构建覆盖 MoleculeNet、TDC ADMET、bRo5、MoleculeACE 和 Random/Scaffold/Perimeter 划分的统一评测流程；在测试集读取前预先登记强基线、冻结表征、目标变换和融合策略；以 accept/reject/retain 机制治理补救头、适用域门控和强基线接入；联合报告 validation-test rank correlation、Top-1/Top-3 命中、regret、optimism gap、保形预测、risk-coverage、MoleculeACE cliff-pair 指标和 ROGI/MODI/SARI 粗糙度诊断；并将数据缺口、候选失败、低相似度失败和活性悬崖失败纳入限制与补充材料，使结论同时受到正向结果和负向证据约束。",
        "综合上述文献，FZYC-Mol 的研究空白并不在于提出规模更大的主干模型，而在于构建可审计的候选模型选择和可靠性治理流程。AutoML、ensemble selection、stacking、nested cross-validation 和 QSAR applicability domain 已分别研究了模型搜索、集成选择、验证偏差控制和适用域边界；FZYC-Mol 与这些工作的区别在于将其整合为一套分子性质预测中的冻结候选池、验证集接受/拒绝规则、测试集一次性报告和负结果审计。近两年的 ADMET 可靠性、OOD 基准、表格基础模型、多尺度表征和片段预训练研究共同提示：真实分子性质预测中并不存在单一全能模型；小样本终点中描述符、指纹、树模型和表格基础模型仍然是强对照；多尺度结构与片段解释能改善部分任务，但需要同步报告其适用边界 [3,13-16,41-44]。因此，本研究将强表格基线、图模型、冻结表征、目标变换、欠采样集成、Top-K/堆叠集成、适用域门控、粗糙度诊断和可解释性分析纳入同一验证集治理流程。": "综合上述文献，FZYC-Mol 所针对的空白并非更大规模主干模型本身，而是可审计的候选模型选择和可靠性治理流程。AutoML、ensemble selection、stacking、nested cross-validation 和 QSAR applicability domain 已分别讨论模型搜索、集成选择、验证偏差控制和适用域边界；FZYC-Mol 将这些要素组织为分子性质预测中的冻结候选池、验证集接受/拒绝规则、测试集一次性报告和负结果审计。近两年的 ADMET 可靠性、OOD 基准、表格基础模型、多尺度表征和片段预训练研究共同提示，真实分子性质预测中不存在单一全能模型；在小样本终点中，描述符、指纹、树模型和表格基础模型仍是必要强对照；多尺度结构与片段解释可能改善部分任务，但必须同步报告适用边界 [3,13-16,41-44]。基于这一认识，本研究将强表格基线、图模型、冻结表征、目标变换、欠采样集成、Top-K/堆叠集成、适用域门控、粗糙度诊断和可解释性分析纳入同一验证集治理流程。",
        "为与近期真实挑战 ADMET 框架保持可比，本文将补充实验组织为“挑战维度—候选池—验证门控—风险输出”的矩阵。该矩阵覆盖强基线同划分、FreeSolv 边界补救、bRo5 公开压力测试、MoleculeACE 活性悬崖、OOD/Perimeter/低相似度压力测试、不平衡分类、嵌套验证、完整消融、轻量大模型适配器、3D-lite 物理代理、外部 holdout 与样本级失败案例。所有模块均遵守同一冻结原则：候选是否进入最终保留结果，只能由验证集或预定义审计规则决定；数据未就绪或证据不足的模块被标注为缺失、负结果或后续验证，而不是写作已完成性能提升。": "为与近期真实挑战 ADMET 框架保持可比，本文将补充实验组织为“挑战维度—候选池—验证门控—风险输出”的矩阵。该矩阵覆盖强基线同划分、FreeSolv 边界补救、bRo5 公开压力测试、MoleculeACE 活性悬崖、OOD/Perimeter/低相似度压力测试、不平衡分类、嵌套验证、完整消融、轻量适配器、3D-lite 物理代理、外部 holdout 与样本级失败案例。所有模块均遵守同一冻结原则：候选是否进入最终保留结果，只能由验证集或预定义审计规则决定；证据不足的模块仅作为负结果、接口状态或后续验证条件报告，不进入性能主张。",
        "按照补充实验方案，表中的每一项均被限定为一个明确的证据功能：已具备结果的模块报告性能、稳定性和负结果；仅具备接口或数据状态的模块只报告可用性、失败原因和后续运行条件。这样处理可以把论文从“增加模型数量”转向“呈现选择过程的可审计性”，同时避免将尚未完成的实验写成已获得的正向结论。": "按照补充实验方案，表中的每一项均被限定为明确的证据功能：已完成模块报告性能、稳定性和负结果；仅具备接口或数据状态的模块报告可用性、失败原因和后续运行条件。该设计将论文重点从模型数量扩展转向选择过程的可审计性，并防止将未完成实验表述为正向性能证据。",
        "在允许低成本模型重构后，该组实验给出了更细化的性能判断：ESOL、FreeSolv、Lipophilicity、BBBP、BACE 和 ClinTox 六个 MoleculeNet 终点均进入同一套 Morgan 指纹与 RDKit 描述符重构候选池，候选包括 CatBoost、XGBoost、LightGBM、ExtraTrees、RF、目标变换、Top-K 均值、堆叠集成和平衡欠采样集成。结果显示，只有 FreeSolv 的 Morgan+descriptor 验证集堆叠在 5 个 scaffold seeds 上将 RMSE 从当前/已接入补救选择器的 1.0678 ± 0.1883 降至 1.0286 ± 0.1761，正向变化为 +0.0392，因此可作为附录级最终保留增强；ESOL、BBBP、BACE 和 ClinTox 的重构候选未超过当前验证集选择器，Lipophilicity 的表格重构候选也未超过已有补救整合选择器，因此这些任务均保留原策略。": "低成本模型重构进一步限定了性能增益的来源。ESOL、FreeSolv、Lipophilicity、BBBP、BACE 和 ClinTox 六个 MoleculeNet 终点进入同一套 Morgan 指纹与 RDKit 描述符重构候选池，候选包括 CatBoost、XGBoost、LightGBM、ExtraTrees、RF、目标变换、Top-K 均值、堆叠集成和平衡欠采样集成。只有 FreeSolv 的 Morgan+descriptor 验证集堆叠在 5 个 scaffold seeds 上将 RMSE 从当前/已接入补救选择器的 1.0678 ± 0.1883 降至 1.0286 ± 0.1761，正向变化为 +0.0392，因此作为附录级最终保留增强。ESOL、BBBP、BACE 和 ClinTox 的重构候选未超过当前验证集选择器，Lipophilicity 的表格重构候选也未超过已有补救整合选择器，因此这些任务保留原策略。",
        "bRo5 部分已由数据状态审计更新为可追溯的公开压力测试。CycPept-PAMPA 继续作为规则五以外渗透性回归任务报告 RMSE、MAE、Spearman、适用域覆盖和风险富集；LinPept CellPen 与 LinPept NonFouling 已从 Benchmark-ADMET-2025 公开仓库下载原始 CSV，核验 smiles 与终点标签字段后统一标准化为 smiles/y 格式。两项 LinPept 任务均按 random、scaffold 和 perimeter split 以及 3 个随机种子运行，候选模型仅由验证集 ROC-AUC 冻结选择，测试集用于最终一次性评估。因此，正文可以报告 bRo5 公开压力测试结果，但应区分 CycPept-PAMPA 的回归指标与 LinPept 的分类指标，避免混用 RMSE 与 ROC-AUC。": "bRo5 部分现包含可追溯的公开压力测试。CycPept-PAMPA 作为规则五以外渗透性回归任务报告 RMSE、MAE、Spearman、适用域覆盖和风险富集；LinPept CellPen 与 LinPept NonFouling 从 Benchmark-ADMET-2025 公开仓库下载原始 CSV，核验 smiles 与终点标签字段后统一标准化为 smiles/y 格式。两项 LinPept 任务均按 random、scaffold 和 perimeter split 以及 3 个随机种子运行，候选模型仅由验证集 ROC-AUC 冻结选择，测试集用于一次性最终评估。因此，本节将 CycPept-PAMPA 作为回归型 bRo5 压力测试、将 LinPept 作为分类型 bRo5 压力测试分别报告。",
        "MoleculeACE 配对结果与低相似度分析互相补充。当前可核验结果覆盖 17 个 MoleculeACE 任务和 51 个 seed 配对，而不是完整 30 任务全量审计。主指标包括 R2、RMSE、MAE、cliff-pair RMSE、预测差异与真实差异的 gap Spearman、cliff recall 和 top cliff error；比较对象包括 XGBoost-MorganCount、KPGT representation、可运行的 MOLMCL/KPGT 版本、FZYC selector 以及 KPGT+XGBoost ensemble。代表性配对结果显示，总 RMSE 平均正向变化为 0.0069，悬崖子集 RMSE 平均正向变化为 0.0056，但标准差较大；gap Spearman 平均约为 0.252，部分任务接近零或为负。因此，该模块被解释为活性悬崖风险识别和候选治理的补充证据，而非证明活性悬崖预测已经解决。": "MoleculeACE 配对结果与低相似度分析互为补充。当前可核验结果覆盖 17 个 MoleculeACE 任务和 51 个 seed 配对，而非完整 30 任务审计。主指标包括 R2、RMSE、MAE、cliff-pair RMSE、预测差异与真实差异的 gap Spearman、cliff recall 和 top cliff error；比较对象包括 XGBoost-MorganCount、KPGT representation、可运行的 MOLMCL/KPGT 版本、FZYC selector 以及 KPGT+XGBoost ensemble。代表性配对结果显示，总 RMSE 平均正向变化为 0.0069，悬崖子集 RMSE 平均正向变化为 0.0056，但标准差较大；gap Spearman 平均约为 0.252，部分任务接近零或为负。因此，该模块被解释为活性悬崖风险识别和候选治理的补充证据，而不支持“活性悬崖预测已被解决”的结论。",
        "这种差异化定位决定了本文的补强方向。强基线同划分、FreeSolv 补救、不平衡分类、OOD/Perimeter 低相似度压力测试、MoleculeACE、nested validation、完整消融、bRo5 公开压力测试、3D-lite/大模型适配器候选、外部 holdout 和样本级案例共同构成 FZYC-Mol 的治理压力测试。复杂模块并不直接扩大主结论，而是在统一验证证据下被接受、保留、拒绝或标注为数据未就绪。这种证据组织与当前结果中“选择性增益多于普遍胜出”的事实一致。": "这种差异化定位决定了本文的证据扩展方向。强基线同划分、FreeSolv 补救、不平衡分类、OOD/Perimeter 低相似度压力测试、MoleculeACE、nested validation、完整消融、bRo5 公开压力测试、3D-lite/轻量适配器候选、外部 holdout 和样本级案例共同构成 FZYC-Mol 的治理压力测试。复杂模块并不直接扩大主结论，而是在统一验证证据下被接受、保留、拒绝或限定为后续验证接口。该证据组织与当前结果中“选择性增益多于普遍胜出”的观察一致。",
        "此外，OmniMol、ADMETlab 3.0 和近期 ADMET 表征基准提示，真实 ADMET 建模往往同时面对不完整标注、终点间相关性、平台化决策支持和跨数据集迁移问题 [47-49]。因此，FZYC-Mol 的价值不仅在于判定哪个模型分数最高，更在于说明候选策略在什么边界内可用、为什么可用以及何时应被拒绝。外部 ADMET 附录、bRo5 公开压力测试、MoleculeACE 可用子集、ROGI/MODI/SARI 诊断和计算成本审计共同服务于这一点。": "此外，OmniMol、ADMETlab 3.0 和近期 ADMET 表征基准提示，真实 ADMET 建模往往同时面对不完整标注、终点间相关性、平台化决策支持和跨数据集迁移问题 [47-49]。因此，FZYC-Mol 的价值不仅在于判定哪个模型分数最高，更在于说明候选策略在哪些边界内可用、依据何在以及何时应被保留为风险。外部 ADMET 附录、bRo5 公开压力测试、MoleculeACE 可用子集、ROGI/MODI/SARI 诊断和计算成本审计共同服务于这一目标。",
        "进一步验证应优先扩展真实挑战矩阵，而不是简单堆叠更多模型。更有价值的路线是在更多公开 ADMET、bRo5 和活性悬崖终点上系统评估 XGBoost、CatBoost、RandomForest、ExtraTrees、Chemprop、TabPFNv2、AutoGluon、KPGT/GEM/Uni-Mol 表征、目标变换、采样策略、Top-K 集成、3D-lite 和轻量适配器，并将最终保留决策与 roughness、低相似度、Perimeter split、外部 holdout 和 validation-test rank instability 关联起来。": "后续验证应优先扩展真实挑战矩阵，而不是简单增加模型数量。更直接的路线是在更多公开 ADMET、bRo5 和活性悬崖终点上系统评估 XGBoost、CatBoost、RandomForest、ExtraTrees、Chemprop、TabPFNv2、AutoGluon、KPGT/GEM/Uni-Mol 表征、目标变换、采样策略、Top-K 集成、3D-lite 和轻量适配器，并将最终保留决策与 roughness、低相似度、Perimeter split、外部 holdout 和 validation-test rank instability 关联起来。",
        "本文提出并评估了 FZYC-Mol，一种由验证集治理、适用域感知的多专家分子性质预测框架。结果表明，在 MoleculeNet、TDC ADMET、Random/Scaffold/Perimeter 划分、低相似度子集、MoleculeACE 可用任务子集、ROGI/MODI/SARI 粗糙度诊断、候选接受/拒绝审计、校准/保形预测和基序/片段解释等证据线上，FZYC-Mol 能够提供比单一模型分数更完整的可靠性画像。bRo5 模块已新增 CycPept-PAMPA、LinPept CellPen 和 LinPept NonFouling 三项公开压力测试。其中 CycPept-PAMPA 作为回归型渗透性任务报告 RMSE/MAE，LinPept 两项作为分类型肽性质任务报告 ROC-AUC、PR-AUC、校准和适用域分层；三者均以冻结测试集作一次性评估，不再作为缺失数据处理。该框架的核心价值不是保证每个终点取得测试最优，而是在固定候选池内透明地接受、保留、拒绝和审计候选策略；当强基线、补救头、低成本重构或多方法融合通过验证集门控时，它们可以进入最终保留结果，当证据不足或数据未就绪时则作为负结果或待验证模块保留。": "本文提出并评估了 FZYC-Mol，一种由验证集治理驱动、适用域感知的多专家分子性质预测框架。在 MoleculeNet、TDC ADMET、Random/Scaffold/Perimeter 划分、低相似度子集、MoleculeACE 可用任务子集、bRo5 公开压力测试、ROGI/MODI/SARI 粗糙度诊断、候选接受/拒绝审计、校准/保形预测和基序/片段解释等证据线上，FZYC-Mol 提供了比单一模型分数更完整的可靠性画像。其核心价值不是保证每个终点取得测试最优，而是在固定候选池内透明地接受、保留、拒绝和审计候选策略。当强基线、补救头、低成本重构或多方法融合通过验证集门控时，它们进入最终保留结果；当证据不足时，则作为负结果或后续验证接口保留。",
        "数据和代码可用性：本文使用的公开数据集可通过 MoleculeNet、Therapeutics Data Commons、MoleculeACE 及相应原始平台获得。与本文结果对应的 split seeds、候选登记表、验证/测试预测、统计检验脚本、图表 source data、环境文件和表格生成脚本将与投稿或接收版本同步存档于 GitHub/Zenodo 或期刊认可的数据仓库。本稿件当前尚未分配永久数据 DOI 或 accession number；仓库记录冻结并可供审稿访问后，方可在正文和补充材料中填写永久链接、版本号和 accession number。bRo5 相关条目已补齐 CycPept-PAMPA、LinPept CellPen 和 LinPept NonFouling 的可追溯公开数据实验；完整 seed-level 预测、字段字典、脚本和图表源数据保存在 reports/bro5_cycpept_pampa_20260611/ 与 reports/bro5_linpept_20260611/。": "数据和代码可用性：本文使用的公开数据集可通过 MoleculeNet、Therapeutics Data Commons、MoleculeACE、Benchmark-ADMET-2025 及相应原始平台获得。与本文结果对应的 split seeds、候选登记表、验证/测试预测、统计检验脚本、图表 source data、环境文件和表格生成脚本将与投稿或接收版本同步存档于 GitHub/Zenodo 或期刊认可的数据仓库。本稿件当前尚未分配永久数据 DOI 或 accession number；仓库记录冻结并可供审稿访问后，将在正文和补充材料中填写永久链接、版本号和 accession number。bRo5 相关实验包括 CycPept-PAMPA、LinPept CellPen 和 LinPept NonFouling；完整 seed-level 预测、字段字典、脚本和图表源数据保存在 reports/bro5_cycpept_pampa_20260611/ 与 reports/bro5_linpept_20260611/。",
    }


def replace_special_sections(doc: Document) -> list[str]:
    changes: list[str] = []
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("根据补充实验清单，本文进一步补齐了能够在当前数据和依赖条件下运行的关键模块。"):
            set_text(
                paragraph,
                "补充实验闭环进一步明确了各模块的证据状态。bRo5 公开压力测试现包括 CycPept-PAMPA 回归任务和 LinPept CellPen/NonFouling 分类任务。CycPept-PAMPA 覆盖 random、scaffold、perimeter 和 time split；LinPept 两项来自 Benchmark-ADMET-2025 origin_data，标准化后分别包含 CellPen 1,960 条样本和 NonFouling 7,239 条样本。",
            )
            first_inserted = insert_after(
                paragraph,
                "在验证集选择后，LinPept CellPen 在 random、scaffold 和 perimeter split 下的 ROC-AUC 分别为 0.937 ± 0.020、0.894 ± 0.029 和 0.859 ± 0.005，对应 PR-AUC 为 0.894 ± 0.028、0.844 ± 0.038 和 0.822 ± 0.008。LinPept NonFouling 在三类 split 下的 ROC-AUC 分别为 0.766 ± 0.012、0.765 ± 0.004 和 0.761 ± 0.000，对应 PR-AUC 为 0.767 ± 0.011、0.780 ± 0.021 和 0.698 ± 0.000。这些结果作为公开数据压力测试报告，其外推范围不同于官方盲测或未见外部队列。",
            )
            insert_after(
                first_inserted,
                "轻量适配器、MoleculeACE 可用任务子集和统一消融矩阵仍按各自预定义边界解释；所有新增结果均保留 seed-level 输出和负结果记录，不用于替代未运行模块的证据。",
            )
            changes.append("split and polished LinPept supplement paragraph")
        elif text.startswith("本研究仍存在若干局限。首先，验证-测试排名审计和 9 个代表性终点的 3 x 3 nested validation"):
            set_text(
                paragraph,
                "本研究仍存在若干局限。首先，验证-测试排名审计和 9 个代表性终点的 3 × 3 nested validation 表明，验证集治理能够减少测试集事后选择，但内外层验证尚未覆盖全部 MoleculeNet、TDC、bRo5 和 MoleculeACE 终点。因此，小幅增益应与 regret、optimism gap、Top-3 hit 和负结果共同解释，而不能被视为所有任务的测试最优保证。",
            )
            first_inserted = insert_after(
                paragraph,
                "其次，bRo5 审计已新增 CycPept-PAMPA、LinPept CellPen 和 LinPept NonFouling 公开压力测试。CycPept-PAMPA 覆盖 random、scaffold、perimeter 和 time split；LinPept 两项覆盖 random、scaffold 和 perimeter split，并以分类指标报告。MoleculeACE 当前可访问数据源包含 17 个任务，本文已完成这 17 个可用任务的 51 个 seed 配对，因此相关结论被限定为当前可用面板。",
            )
            insert_after(
                first_inserted,
                "第三，TabPFNv2、AutoGluon、KPGT representation、轻量适配器、3D-lite、bRo5 和外部时间切分仍需更大规模的同划分验证；尚未完成的候选仅作为候选接口或后续验证方向，不作为已取得性能提升的证据。最后，bRo5、活性悬崖和低相似度样本仍是主要失败边界，FZYC-Mol 更适合识别和审计这些风险，而非宣称已解决它们。基序归因、片段富集和粗糙度相关性仍属于关联证据，不能替代因果机制验证或湿实验验证。",
            )
            changes.append("split and polished limitations paragraph")
        elif text.startswith("数据和代码可用性：本文使用的公开数据集可通过 MoleculeNet"):
            set_text(
                paragraph,
                "数据和代码可用性：本文使用的公开数据集可通过 MoleculeNet、Therapeutics Data Commons、MoleculeACE、Benchmark-ADMET-2025 及相应原始平台获得。与本文结果对应的 split seeds、候选登记表、验证/测试预测、统计检验脚本、图表 source data、环境文件和表格生成脚本将与投稿或接收版本同步存档于 GitHub/Zenodo 或期刊认可的数据仓库。",
            )
            insert_after(
                paragraph,
                "本稿件当前尚未分配永久数据 DOI 或 accession number；仓库记录冻结并可供审稿访问后，将在正文和补充材料中填写永久链接、版本号和 accession number。bRo5 相关实验包括 CycPept-PAMPA、LinPept CellPen 和 LinPept NonFouling；完整 seed-level 预测、字段字典、脚本和图表源数据保存在 reports/bro5_cycpept_pampa_20260611/ 与 reports/bro5_linpept_20260611/。",
            )
            changes.append("split data and code availability paragraph")
    return changes


def audit_doc(doc: Document) -> dict[str, int]:
    text = "\n".join(p.text for p in doc.paragraphs)
    terms = ["写成", "马虎", "完美", "数据未就绪", "补强", "堆叠更多模型", "证明活性悬崖预测已经解决"]
    return {term: text.count(term) for term in terms}


def write_report(changes: list[tuple[int, str]], special: list[str], table_changes: int, audit: dict[str, int]) -> None:
    report = Document()
    report.add_heading("初稿-18 Nature 风格语言逻辑终审报告", level=1)
    report.add_paragraph("本轮基于初稿-17，重点检查并修订摘要、引言、结果、讨论、局限、结论和数据/代码可用性声明。未改动实验数值、图像和核心结论。")
    report.add_heading("主要修订", level=2)
    report.add_paragraph(f"精修段落：{len(changes)} 处；拆分长段：{len(special)} 处；表格措辞修订：{table_changes} 处。")
    report.add_paragraph("处理重点包括：减少审稿过程用语，压缩过长段落，区分结果与讨论，统一 bRo5 回归/分类指标表达，弱化过强或因果化措辞，并保留负结果边界。")
    report.add_heading("残留词审计", level=2)
    for key, value in audit.items():
        report.add_paragraph(f"{key}: {value}")
    report.add_heading("输出文件", level=2)
    report.add_paragraph(str(OUT_DOCX))
    report.save(REPORT_DOCX)
    shutil.copy2(REPORT_DOCX, REPORT_DIR / REPORT_DOCX.name)


def main() -> None:
    doc = Document(IN_DOCX)
    changes = replace_exact(doc, make_replacements())
    special = replace_special_sections(doc)
    table_changes = polish_tables(doc)
    audit = audit_doc(doc)
    doc.save(OUT_DOCX)
    shutil.copy2(OUT_DOCX, REPORT_DIR / OUT_DOCX.name)
    write_report(changes, special, table_changes, audit)
    print(f"Input: {IN_DOCX}")
    print(f"Wrote: {OUT_DOCX}")
    print(f"Wrote: {REPORT_DOCX}")
    print(f"Paragraph changes: {len(changes)}")
    print(f"Special section changes: {len(special)}")
    print(f"Table text changes: {table_changes}")
    print(f"Audit: {audit}")


if __name__ == "__main__":
    main()
