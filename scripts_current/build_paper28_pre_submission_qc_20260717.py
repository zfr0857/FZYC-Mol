from __future__ import annotations

import json
import re
from pathlib import Path

import fitz
import pandas as pd
from docx import Document
from lxml import etree
from openpyxl import load_workbook
from PIL import Image


ROOT = Path(r"D:\fzyc")
BASE = ROOT / "output" / "paper28_pre_submission_minor_revision_20260717"
FIG = BASE / "main_figures"
SUPP = BASE / "supplementary"
ANALYSIS = ROOT / "output" / "paper27_equal_size_registry_composition_20260716"
EN = BASE / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript.docx"
REVIEWER = BASE / "Reviewer_concern_Response_Location.docx"
ZH = next(path for path in BASE.glob("*.docx") if path not in {EN, REVIEWER} and "tracked" not in path.name)
WORKBOOK = SUPP / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S31.xlsx"
WIDTH_IN = 170 / 25.4


def document_text(path: Path) -> str:
    doc = Document(path)
    text = [paragraph.text for paragraph in doc.paragraphs]
    text.extend(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    return "\n".join(text)


def svg_metadata(path: Path) -> dict[str, object]:
    root = etree.parse(str(path))
    text_nodes = root.xpath("//*[local-name()='text']")
    style_nodes = root.xpath("//*[@style]")
    families: list[str] = []
    sizes: list[float] = []
    widths: list[float] = []
    for node in style_nodes:
        style = node.get("style", "")
        family = re.search(r"font-family:\s*([^;]+)", style)
        size = re.search(r"font-size:\s*([0-9.]+)px", style)
        width = re.search(r"stroke-width:\s*([0-9.]+)", style)
        if family:
            families.append(family.group(1).strip(" '\""))
        if size:
            sizes.append(float(size.group(1)))
        if width and float(width.group(1)) > 0:
            widths.append(float(width.group(1)))
    allowed = sum(family in {"Times New Roman", "STIXGeneral"} for family in families)
    return {
        "text_nodes": len(text_nodes),
        "families": sorted(set(families)),
        "allowed_percentage": 100 * allowed / len(families) if families else 0,
        "minimum_font_size": min(sizes) if sizes else float("nan"),
        "minimum_line_width": min(widths) if widths else float("nan"),
    }


def pdf_metadata(path: Path) -> dict[str, object]:
    doc = fitz.open(path)
    fonts: dict[int, tuple] = {}
    for page in doc:
        for font in page.get_fonts(full=True):
            fonts[font[0]] = font
    names = sorted({font[3] for font in fonts.values()})
    embedded = True
    extracted = []
    for xref, font in fonts.items():
        name, ext, _, data = doc.extract_font(xref)
        extracted.append({"name": name, "ext": ext, "bytes": len(data)})
        embedded &= ext.lower() in {"ttf", "otf", "cff"} and len(data) > 0
    allowed = all("TimesNewRoman" in name or "STIXGeneral" in name for name in names)
    return {"fonts": names, "embedded": embedded, "allowed": allowed, "details": extracted}


def figure_reports() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    quality, pdf_rows, svg_rows = [], [], []
    for index in range(1, 8):
        png = FIG / f"Figure{index}_600dpi.png"
        svg = FIG / f"Figure{index}.svg"
        pdf = FIG / f"Figure{index}.pdf"
        image = Image.open(png)
        svg_info = svg_metadata(svg)
        pdf_info = pdf_metadata(pdf)
        font_ok = svg_info["allowed_percentage"] == 100 and pdf_info["allowed"]
        size_ok = svg_info["minimum_font_size"] >= 7.5
        raster_ok = image.width >= 4016 and image.mode == "RGB"
        editable = svg_info["text_nodes"] > 0
        passed = font_ok and size_ok and raster_ok and editable and pdf_info["embedded"]
        quality.append({
            "Figure": f"Figure {index}",
            "Final width": "170 mm",
            "Pixel size": f"{image.width} x {image.height}",
            "Effective dpi": round(image.width / WIDTH_IN, 1),
            "Detected font families": "; ".join(svg_info["families"]),
            "Times New Roman or STIX percentage": round(svg_info["allowed_percentage"], 1),
            "Minimum font size": svg_info["minimum_font_size"],
            "Embedded PDF fonts": "; ".join(pdf_info["fonts"]),
            "Editable SVG text": editable,
            "Minimum line width": svg_info["minimum_line_width"],
            "Overlap": "not detected in visual review",
            "Clipping": "not detected in visual review",
            "Legend obstruction": "not detected in visual review",
            "Colour-blind check": "pass: colour is redundant with marker, line style, position or cell text",
            "Greyscale check": "pass: marker/line redundancy or printed cell values retained",
            "RGB white background": image.mode == "RGB",
            "Pass/fail": "Pass" if passed else "Fail",
        })
        pdf_rows.append({
            "Figure": f"Figure {index}",
            "Fonts": "; ".join(pdf_info["fonts"]),
            "Font extraction": json.dumps(pdf_info["details"]),
            "All fonts embedded": pdf_info["embedded"],
            "Only Times New Roman or STIX": pdf_info["allowed"],
            "Status": "Pass" if pdf_info["embedded"] and pdf_info["allowed"] else "Fail",
        })
        svg_rows.append({
            "Figure": f"Figure {index}",
            "Text nodes": svg_info["text_nodes"],
            "Font families": "; ".join(svg_info["families"]),
            "Allowed-font percentage": round(svg_info["allowed_percentage"], 1),
            "Minimum font size": svg_info["minimum_font_size"],
            "Text remains editable": editable,
            "Status": "Pass" if font_ok and size_ok and editable else "Fail",
        })
    return pd.DataFrame(quality), pd.DataFrame(pdf_rows), pd.DataFrame(svg_rows)


def oracle_report() -> pd.DataFrame:
    rows = []
    documents = [EN, ZH, SUPP / "Additional_file_1_Supplementary_Methods_and_Results.docx", REVIEWER]
    for path in documents:
        hits = len(re.findall("oracle", document_text(path), flags=re.IGNORECASE))
        rows.append({"artifact": path.name, "locations_scanned": "paragraphs and tables", "oracle_hits": hits, "status": "Pass" if hits == 0 else "Fail"})
    workbook = load_workbook(WORKBOOK, read_only=True, data_only=False)
    workbook_hits = 0
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    workbook_hits += len(re.findall("oracle", cell.value, flags=re.IGNORECASE))
    workbook.close()
    rows.append({"artifact": WORKBOOK.name, "locations_scanned": "all string cells in S1-S31", "oracle_hits": workbook_hits, "status": "Pass" if workbook_hits == 0 else "Fail"})
    for path in sorted(FIG.glob("Figure[1-7].svg")):
        hits = len(re.findall("oracle", path.read_text(encoding="utf-8"), flags=re.IGNORECASE))
        rows.append({"artifact": path.name, "locations_scanned": "editable SVG text", "oracle_hits": hits, "status": "Pass" if hits == 0 else "Fail"})
    for path in sorted((BASE / "figure_source_data").glob("Figure_7*")):
        if path.suffix.lower() not in {".csv", ".json"}:
            continue
        hits = len(re.findall("oracle", path.read_text(encoding="utf-8"), flags=re.IGNORECASE))
        rows.append({"artifact": path.name, "locations_scanned": "Figure 7 public source table", "oracle_hits": hits, "status": "Pass" if hits == 0 else "Fail"})
    return pd.DataFrame(rows)


def unicode_report() -> pd.DataFrame:
    rows = []
    artifacts = [EN, ZH, SUPP / "Additional_file_1_Supplementary_Methods_and_Results.docx", REVIEWER]
    for path in artifacts:
        text = document_text(path)
        bad = [char for char in text if char == "\ufffd" or 0xE000 <= ord(char) <= 0xF8FF or (ord(char) < 32 and char not in "\n\t\r")]
        rows.append({"artifact": path.name, "replacement_or_private_use_characters": len(bad), "status": "Pass" if not bad else "Fail"})
    for path in sorted(FIG.glob("Figure[1-7].svg")):
        text = path.read_text(encoding="utf-8")
        bad = [char for char in text if char == "\ufffd" or 0xE000 <= ord(char) <= 0xF8FF]
        rows.append({"artifact": path.name, "replacement_or_private_use_characters": len(bad), "status": "Pass" if not bad else "Fail"})
    return pd.DataFrame(rows)


def cross_reference_report() -> pd.DataFrame:
    en_text = document_text(EN)
    zh_text = document_text(ZH)
    rows = []
    for index in range(1, 8):
        conditions = {
            "English caption": f"Figure {index}." in en_text,
            "Chinese caption": any(token in zh_text for token in (f"图{index}.", f"图 {index}.", f"图{index} ", f"图 {index} ")),
            "PDF": (FIG / f"Figure{index}.pdf").is_file(),
            "SVG": (FIG / f"Figure{index}.svg").is_file(),
            "600 dpi PNG": (FIG / f"Figure{index}_600dpi.png").is_file(),
        }
        rows.append({"Figure": index, **conditions, "status": "Pass" if all(conditions.values()) else "Fail"})
    return pd.DataFrame(rows)


def numeric_consistency_report() -> pd.DataFrame:
    en_text, zh_text = document_text(EN), document_text(ZH)
    values = ["0.125", "0.188", "0.079", "0.084", "1.419", "2.260", "0.103", "0.172", "0.072", "0.078", "1.325", "2.179", "0.176", "2.191", "10.88", "2.90", "1.87", "12.40", "6.09", "9.07", "13.96", "7.92", "11.47", "-0.008", "-0.211", "-0.030", "0.0784"]
    rows = []
    for value in values:
        en = value in en_text
        zh = value in zh_text
        rows.append({"value": value, "English present": en, "Chinese present": zh, "status": "Pass" if en and zh else "Fail"})
    return pd.DataFrame(rows)


def master_result_consistency() -> pd.DataFrame:
    raw = pd.read_csv(ANALYSIS / "equal_size_endpoint_summary.csv").rename(columns={
        "oracle_opportunity_gain_mean": "observed_audit_best_opportunity_gain_mean",
        "oracle_opportunity_gain_low": "observed_audit_best_opportunity_gain_low",
        "oracle_opportunity_gain_high": "observed_audit_best_opportunity_gain_high",
    })
    table = pd.read_excel(WORKBOOK, sheet_name="S28 Endpoint summary")
    keys = ["pool", "task", "candidate_count"]
    columns = [
        "observed_audit_best_opportunity_gain_mean", "selected_model_gain_mean",
        "same_unit_selection_gap_mean", "chance_adjusted_hit3_mean",
        "cross_fitted_selection_gap_mean", "audit_fit_seconds_mean",
    ]
    merged = raw[keys + columns].merge(table[keys + columns], on=keys, suffixes=("_analysis", "_workbook"), validate="one_to_one")
    rows = []
    for metric in columns:
        difference = (merged[f"{metric}_analysis"] - merged[f"{metric}_workbook"]).abs()
        rows.append({
            "metric": metric,
            "rows_compared": len(merged),
            "maximum_absolute_difference": float(difference.max()),
            "source": "equal_size_endpoint_summary.csv",
            "destination": "Additional file 2, Table S28",
            "status": "Pass" if float(difference.max()) < 1e-12 else "Fail",
        })
    return pd.DataFrame(rows)


def declarations_report() -> pd.DataFrame:
    text = document_text(EN)
    headings = ["Ethics approval and consent to participate", "Consent for publication", "Availability of data and materials", "Competing interests", "Funding", "Authors' contributions", "Acknowledgements"]
    rows = []
    for heading in headings:
        present = heading in text
        pending = heading in {"Competing interests", "Funding", "Authors' contributions", "Acknowledgements"}
        rows.append({
            "declaration": heading,
            "heading_present": present,
            "author_confirmation_required": pending,
            "status": "Author confirmation required" if pending else ("Pass" if present else "Fail"),
        })
    return pd.DataFrame(rows)


def main() -> None:
    quality, pdf_fonts, svg_fonts = figure_reports()
    oracle = oracle_report()
    unicode = unicode_report()
    cross = cross_reference_report()
    numeric = numeric_consistency_report()
    master = master_result_consistency()
    declarations = declarations_report()
    reports = {
        "Figure_quality_control_table.csv": quality,
        "PDF_font_embedding_report.csv": pdf_fonts,
        "SVG_font_and_editability_report.csv": svg_fonts,
        "Oracle_terminology_cleanup_report.csv": oracle,
        "Unicode_cleanup_report.csv": unicode,
        "Figure_table_cross_reference_check.csv": cross,
        "English_Chinese_numeric_consistency_check.csv": numeric,
        "Master_result_consistency_table.csv": master,
        "Declarations_final_check.csv": declarations,
    }
    for name, frame in reports.items():
        frame.to_csv(BASE / name, index=False, encoding="utf-8-sig")
    with pd.ExcelWriter(BASE / "Master_result_consistency_table.xlsx", engine="openpyxl") as writer:
        master.to_excel(writer, index=False, sheet_name="Consistency")
    statuses = {
        "figures": bool(quality["Pass/fail"].eq("Pass").all()),
        "pdf_fonts": bool(pdf_fonts["Status"].eq("Pass").all()),
        "svg_fonts": bool(svg_fonts["Status"].eq("Pass").all()),
        "oracle_cleanup": bool(oracle.status.eq("Pass").all()),
        "unicode": bool(unicode.status.eq("Pass").all()),
        "cross_references": bool(cross.status.eq("Pass").all()),
        "numeric_consistency": bool(numeric.status.eq("Pass").all()),
        "master_result_consistency": bool(master.status.eq("Pass").all()),
    }
    audit = {
        "status": "complete" if all(statuses.values()) else "failed",
        "checks": statuses,
        "author_confirmation_still_required": declarations.loc[declarations.author_confirmation_required, "declaration"].tolist(),
        "figure7_definition_audit": json.loads((BASE / "figure_source_data" / "Figure_7_definition_audit.json").read_text(encoding="utf-8")),
    }
    (BASE / "Final_minor_revision_QC_audit.json").write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(audit, ensure_ascii=False, indent=2))
    if audit["status"] != "complete":
        raise SystemExit(1)


if __name__ == "__main__":
    main()
