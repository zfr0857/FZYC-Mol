from __future__ import annotations

import csv
import sys
from pathlib import Path

from docx import Document


def find_plan() -> Path:
    base = Path(r"C:\Users\Administrator\Downloads")
    candidates = [p for p in base.glob("FZYC-Mol_*.docx") if p.suffix.lower() == ".docx"]
    preferred = [
        p
        for p in candidates
        if ("\u8865\u5145\u5b9e\u9a8c" in p.name and "\u8be6\u7ec6\u7248" in p.name)
        or "\u65e0\u4e71\u7801" in p.name
    ]
    if not preferred:
        preferred = candidates
    if not preferred:
        raise FileNotFoundError("No FZYC-Mol plan document found in Downloads.")
    return max(preferred, key=lambda p: p.stat().st_mtime)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    plan = find_plan()
    doc = Document(plan)
    out_dir = Path("work/draft15_audit")
    out_dir.mkdir(parents=True, exist_ok=True)

    paras = []
    for i, paragraph in enumerate(doc.paragraphs, 1):
        text = " ".join(paragraph.text.split())
        if text:
            paras.append((i, paragraph.style.name if paragraph.style else "", text))

    with (out_dir / "plan_paragraph_index.csv").open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["index", "style", "text"])
        writer.writerows(paras)

    with (out_dir / "plan_table_index.csv").open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["table", "row", "cols", "cells"])
        for t_idx, table in enumerate(doc.tables, 1):
            for r_idx, row in enumerate(table.rows, 1):
                cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
                writer.writerow([t_idx, r_idx, len(cells), " || ".join(cells)])

    print(f"Plan: {plan}")
    print(f"Paragraphs: {len(paras)}; tables: {len(doc.tables)}")
    print("\nHeadings and module cues:")
    cue_terms = [
        "bRo5",
        "MoleculeACE",
        "FreeSolv",
        "TabPFN",
        "AutoGluon",
        "Perimeter",
        "nested",
        "ablation",
        "\u6d88\u878d",
        "\u4f4e\u76f8\u4f3c",
        "\u4e0d\u5e73\u8861",
        "\u5916\u90e8",
        "\u76f2\u6d4b",
        "\u5927\u6a21\u578b",
        "\u7c97\u7cd9",
        "\u6848\u4f8b",
        "\u4f18\u5148",
        "\u4e09\u7ebf\u8868",
    ]
    for i, style, text in paras:
        if style.startswith("Heading") or any(term.lower() in text.lower() for term in cue_terms):
            print(f"P{i:04d} [{style}] {text[:900]}")

    print("\nTables:")
    for idx, table in enumerate(doc.tables, 1):
        rows = len(table.rows)
        cols = len(table.columns) if rows else 0
        first = " | ".join(cell.text.strip().replace("\n", " / ") for cell in table.rows[0].cells) if rows else ""
        print(f"T{idx:02d}: {rows}x{cols} :: {first[:500]}")
        for ridx, row in enumerate(table.rows[:4], 1):
            cells = " || ".join(cell.text.strip().replace("\n", " / ") for cell in row.cells)
            print(f"  r{ridx}: {cells[:700]}")


if __name__ == "__main__":
    main()
