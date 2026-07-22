from __future__ import annotations

import hashlib
import json
import os
import shutil
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import load_workbook
from pypdf import PdfReader, PdfWriter

ROOT=Path("D:/fzyc")
OLD=ROOT/"output"/"paper20_candidate_pool_audit_20260712"/"supplementary"
NEW=Path(os.environ.get("FZYC_ANALYSIS_OUT", ROOT/"output"/"paper21_final_reanalysis_20260713"))
MINOR=Path(os.environ.get("FZYC_MINOR_OUT", ROOT/"output"/"paper23_minor_revision_20260713"))
SUP=MINOR/"supplementary"; SUP.mkdir(parents=True,exist_ok=True)
CORE=Path(os.environ.get("FZYC_CORE_OUT", ROOT/"output"/"paper20_candidate_pool_audit_20260712"))
PREFIX=Path(os.environ.get("FZYC_PREFIX_BASE", ROOT/"results"/"paper22_combined_nested_20260713"/"prefix32"))
MULTI=Path(os.environ.get("FZYC_MULTIVIEW_BASE", ROOT/"results"/"paper22_combined_nested_20260713"/"multiview12"))
HARD=ROOT/"output"/"sci1_hardening_20260707"
DISPLAY={"bace":"BACE","bbbp":"BBBP","clintox":"ClinTox","esol":"ESOL","freesolv":"FreeSolv","lipo":"Lipophilicity","tdc_caco2_wang":"Caco2","tdc_hia_hou":"HIA","tdc_pgp_broccatelli":"P-gp"}

def read_old(sheet):
    return pd.read_excel(OLD/"Additional_file_2_Supplementary_Tables_S1-S17.xlsx",sheet_name=sheet)

def tagged(frames):
    return pd.concat([f.assign(record_type=tag) for tag,f in frames],ignore_index=True,sort=False)

def sha(path):
    h=hashlib.sha256()
    with open(path,"rb") as f:
        for b in iter(lambda:f.read(1<<20),b""): h.update(b)
    return h.hexdigest()

def add_page_number(section):
    p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); p._p.append(fld)

def build_xlsx():
    out=SUP/"Additional_file_2_Machine_readable_Supplementary_Tables_S1-S22.xlsx"
    sheets={
      "S1 Dataset provenance":read_old("S1 Dataset provenance"),
      "S2 Complete candidate registry":read_old("S2 Candidate registry"),
      "S3 Candidate configurations":read_old("S3 Config compute"),
      "S4 Split and seed audit":tagged([("summary",pd.read_csv(NEW/"seed_split_prediction_audit_summary.csv")),("outer_unit",pd.read_csv(NEW/"seed_split_prediction_audit_detail.csv")),("regression_split_manifest",pd.read_csv(PREFIX/"regression_split_manifest.csv"))]),
      "S5 Fold-level utilities":tagged([
          (f"outer_seed_{seed}",pd.read_csv(PREFIX/f"seed_{seed}"/"outer_candidate_scores.csv"))
          for seed in [11,23,37,53,71]
      ] + [("ranking_unit",pd.read_csv(NEW/"chance_adjusted_ranking_units.csv")),("ranking_seed_summary",pd.read_csv(MINOR/"ranking_metric_seed_summary.csv")),("ranking_endpoint_summary",pd.read_csv(MINOR/"ranking_metric_endpoint_summary.csv")),("ranking_main_summary",pd.read_csv(MINOR/"ranking_metric_main_summary.csv"))]),
      "S6 Effective-rank bootstrap":tagged([("main_verification",pd.read_csv(MINOR/"effective_rank_verification.csv")),("bootstrap_summary",pd.read_csv(NEW/"effective_rank_bootstrap_5000_summary.csv")),("monte_carlo_stability",pd.read_csv(NEW/"effective_rank_monte_carlo_stability.csv")),("leave_one_out",pd.read_csv(NEW/"effective_rank_leave_one_out.csv"))]),
      "S7 Reference sensitivity":pd.read_csv(NEW/"effective_rank_reference_sensitivity.csv"),
      "S8 Cross-fitted intervals":tagged([("endpoint_verification",pd.read_csv(MINOR/"cross_fitted_result_verification.csv")),("endpoint_interval",pd.read_csv(NEW/"cross_fitted_complete_intervals.csv")),("fold_effect",pd.read_csv(NEW/"cross_fitted_fold_effects.csv"))]),
      "S9 Matched-K subset results":tagged([("endpoint_verification",pd.read_csv(MINOR/"matched_k3_220_subset_verification.csv")),("subset_summary",pd.read_csv(NEW/"matched_k3_220_subset_summary.csv")),("selection_frequency",pd.read_csv(NEW/"matched_k3_220_selection_frequency.csv")),("outer_unit",pd.read_csv(NEW/"matched_k3_220_subset_units.csv"))]),
      "S10 Full multiview results":tagged([("matched_K_summary",pd.read_csv(CORE/"matched_k_multiview_summary.csv")),("matched_K_outer_unit",pd.read_csv(CORE/"matched_k_multiview_units.csv")),("candidate_registry",pd.read_csv(MULTI/"candidate_registry.csv"))]),
      "S11 Prediction corr error overlap":tagged([("effective_diversity",pd.read_csv(NEW/"prediction_level_effective_diversity.csv")),("correlation",pd.read_csv(NEW/"prediction_correlation_long.csv"))]),
      "S12 Chemical support stratification":tagged([("verification",pd.read_csv(MINOR/"chemical_support_verification.csv")),("tanimoto_support",pd.read_csv(NEW/"chemical_support_selection_audit.csv")),("scaffold_relation",pd.read_csv(NEW/"scaffold_novelty_error_complementarity.csv"))]),
      "S13 Deduplication":read_old("S12 Dedup sensitivity"),
      "S14 Conformal prediction":read_old("S13 Conformal"),
      "S15 TDC":read_old("S14 TDC endpoints"),
      "S16 MoleculeACE":read_old("S15 MoleculeACE"),
      "S17 bRo5":read_old("S16 bRo5"),
      "S18 Failed candidates":read_old("S17 Failed unavailable"),
      "S19 AutoGluon budgets":tagged([("budget_summary",pd.read_csv(ROOT/"results"/"source_data"/"autogluon_budget.csv")),("outer_fold",pd.read_csv(ROOT/"results"/"source_data"/"autogluon_budget_outer_long.csv"))]),
      "S20 Selection risk LOEO":tagged([("selection_risk",pd.read_csv(ROOT/"results"/"selection_closure"/"selection_risk_summary.csv")),("leave_one_endpoint_out",pd.read_csv(ROOT/"results"/"selection_closure"/"leave_one_endpoint_out_policy.csv")),("null_calibration",pd.read_csv(ROOT/"results"/"selection_closure"/"null_calibration_summary.csv"))]),
      "S21 Mechanism permutation":tagged([("summary",pd.read_csv(MINOR/"mechanism_permutation_null_summary.csv")),("draw",pd.read_csv(MINOR/"mechanism_permutation_null_draws.csv.gz"))]),
      "S22 Signal recovery":tagged([("summary",pd.read_csv(MINOR/"mechanism_signal_recovery_summary.csv")),("outer_unit",pd.read_csv(MINOR/"mechanism_signal_recovery_units.csv"))]),
    }
    with pd.ExcelWriter(out,engine="openpyxl") as w:
        for name,frame in sheets.items(): frame.to_excel(w,sheet_name=name[:31],index=False)
    wb=load_workbook(out)
    for ws in wb.worksheets:
        ws.freeze_panes="A2"; ws.auto_filter.ref=ws.dimensions
        for c in ws[1]: c.font=c.font.copy(bold=True)
        for col in ws.columns:
            letter=col[0].column_letter; width=min(42,max(9,max(len(str(x.value or "")) for x in col[:250])+2)); ws.column_dimensions[letter].width=width
    wb.save(out); return out

def page_pdf(fig, name):
    out=SUP/name; fig.savefig(out,bbox_inches="tight",facecolor="white"); plt.close(fig); return out

def title(fig, text):
    fig.suptitle(text,x=.05,ha="left",fontsize=11,fontweight="bold")

def build_s5():
    d=pd.read_csv(MINOR/"mechanism_permutation_null_summary.csv")
    fig,axes=plt.subplots(1,2,figsize=(7.2,4.2)); title(fig,"Figure S5. Random-rank permutation calibration")
    for ax,(metric,label,color) in zip(axes,[("chance_adjusted_hit","CAHit@3","#3F6F9F"),("normalized_mrr_gain","Normalized MRR gain","#3C8C82")]):
        q=d[d.metric.eq(metric)].sort_values("candidate_count")
        ax.fill_between(q.candidate_count,q.null_q025,q.null_q975,color="#7A838C",alpha=.18,label="Permutation 95% envelope")
        ax.plot(q.candidate_count,q.observed_endpoint_median,"o-",color=color,label="Observed endpoint median")
        ax.axhline(0,color="#202830",lw=.8); ax.set(xlabel="Candidate count, K",ylabel=label,xticks=[4,8,16,32]); ax.legend(frameon=False,fontsize=7)
    fig.tight_layout(rect=[0,0,1,.94]); return page_pdf(fig,"Supplementary_Figure_S5_current_permutation_controls.pdf")

def build_s6():
    d=pd.read_csv(MINOR/"mechanism_signal_recovery_summary.csv")
    fig,axes=plt.subplots(1,2,figsize=(7.2,4.2)); title(fig,"Figure S6. Graded signal-recovery positive control")
    colors=["#3F6F9F","#3C8C82","#CF7A3B","#76659A"]
    for k,c in zip([4,8,16,32],colors):
        q=d[d.candidate_count.eq(k)].sort_values("injected_signal")
        axes[0].plot(q.injected_signal,q.chance_adjusted_hit_median,"o-",color=c,label=f"K={k}")
        axes[1].plot(q.injected_signal,q.fixed_range_selection_loss_median,"o-",color=c,label=f"K={k}")
    axes[0].set(xlabel="Injected validation–audit signal",ylabel="Median CAHit@3",ylim=(-.08,1.05))
    axes[1].set(xlabel="Injected validation–audit signal",ylabel="Median fixed-range selection loss",ylim=(-.02,None))
    for ax in axes: ax.legend(frameon=False,ncol=2,fontsize=7)
    fig.tight_layout(rect=[0,0,1,.94]); return page_pdf(fig,"Supplementary_Figure_S6_current_signal_recovery.pdf")

def build_s9():
    d=pd.read_csv(NEW/"matched_k3_220_subset_summary.csv")
    fig,axes=plt.subplots(1,2,figsize=(7.2,5.2)); title(fig,"Figure S9. Full matched-K subset distributions")
    for ax,t,xlab in [(axes[0],"classification","Median ROC-AUC gain"),(axes[1],"regression","Median RMSE reduction")]:
        q=d[d.task_type.eq(t)]; tasks=sorted(q.task.unique()); vals=[q.loc[q.task.eq(x),"selected_model_gain_median"].dropna() for x in tasks]
        ax.boxplot(vals,vert=False,tick_labels=[DISPLAY[x] for x in tasks],showfliers=False); ax.axvline(0,color="#202830",lw=.8); ax.set(xlabel=xlab)
    fig.tight_layout(rect=[0,0,1,.95]); return page_pdf(fig,"Supplementary_Figure_S9_full_matched_K_subset_distributions.pdf")

def build_s10():
    d=pd.read_csv(NEW/"matched_k3_220_selection_frequency.csv")
    fig,axes=plt.subplots(1,2,figsize=(7.2,5.2)); title(fig,"Figure S10. Representation and learner selection frequency")
    for ax,col,label in [(axes[0],"representation","Representation"),(axes[1],"family","Learner family")]:
        p=d.pivot_table(index="task",columns=col,values="size",aggfunc="sum",fill_value=0); p=p.div(p.sum(axis=1),axis=0)
        im=ax.imshow(p,aspect="auto",vmin=0,vmax=max(.01,float(p.to_numpy().max())),cmap="Blues"); ax.set(yticks=np.arange(len(p)),yticklabels=[DISPLAY[x] for x in p.index],xticks=np.arange(len(p.columns)),xticklabels=[str(x).replace("_"," ") for x in p.columns]); plt.setp(ax.get_xticklabels(),rotation=35,ha="right"); ax.set_title(label,loc="left"); fig.colorbar(im,ax=ax,fraction=.04,pad=.03,label="Within-endpoint frequency")
    fig.tight_layout(rect=[0,0,1,.95]); return page_pdf(fig,"Supplementary_Figure_S10_representation_selection_frequency.pdf")

def build_s11():
    d=read_old("S14 TDC endpoints").sort_values("performance_delta_vs_previous")
    fig,ax=plt.subplots(figsize=(7.2,5.2)); title(fig,"Figure S11. TDC endpoint panel")
    ax.barh(np.arange(len(d)),d.performance_delta_vs_previous,color=np.where(d.performance_delta_vs_previous>=0,"#3F6F9F","#B65A55")); ax.axvline(0,color="#202830",lw=.8); ax.set(yticks=np.arange(len(d)),yticklabels=d.dataset,xlabel="Direction-normalized delta versus previous")
    fig.tight_layout(rect=[0,0,1,.95]); return page_pdf(fig,"Supplementary_Figure_S11_TDC.pdf")

def build_s12():
    d=read_old("S13 Conformal"); risk=pd.read_csv(ROOT/"results"/"source_data"/"risk_coverage_metrics.csv")
    fig,axes=plt.subplots(1,2,figsize=(7.2,5.2)); title(fig,"Figure S12. Conformal coverage and risk-ranking summaries")
    q=d[d.alpha.eq(.10)].groupby(["task_type","method"],as_index=False).agg(coverage=("mean_coverage","mean"))
    for t,c in [("classification","#3F6F9F"),("regression","#CF7A3B")]:
        z=q[q.task_type.eq(t)]; axes[0].scatter(z.coverage,z.method,label=t,color=c,s=20)
    axes[0].axvline(.9,color="#202830",ls=":",lw=.8); axes[0].set(xlabel="Mean observed coverage",ylabel="Method"); axes[0].legend(frameon=False)
    g=risk.groupby("endpoint",as_index=False).agg(aurc=("aurc","mean"),random=("random_baseline_risk","mean")); axes[1].scatter(g.random,g.aurc,color="#3C8C82"); lim=max(float(g.random.max()),float(g.aurc.max()))*1.05; axes[1].plot([0,lim],[0,lim],ls=":",color="#7A838C"); axes[1].set(xlabel="Random-baseline risk",ylabel="AURC",xlim=(0,lim),ylim=(0,lim))
    fig.tight_layout(rect=[0,0,1,.95]); return page_pdf(fig,"Supplementary_Figure_S12_conformal_and_risk_coverage.pdf")

def build_s13():
    d=read_old("S15 MoleculeACE").sort_values("direction_accuracy")
    fig,ax=plt.subplots(figsize=(7.2,5.2)); title(fig,"Figure S13. MoleculeACE activity-cliff results"); ax.barh(np.arange(len(d)),d.direction_accuracy,color="#CF7A3B"); ax.set(yticks=np.arange(len(d)),yticklabels=d.task,xlabel="Cliff-pair direction accuracy",xlim=(0,1)); fig.tight_layout(rect=[0,0,1,.95]); return page_pdf(fig,"Supplementary_Figure_S13_MoleculeACE.pdf")

def build_s14():
    d=read_old("S16 bRo5"); vals=d.test_RMSE.astype(str).str.extract(r"([0-9.]+)")[0].astype(float)
    fig,ax=plt.subplots(figsize=(7.2,5.2)); title(fig,"Figure S14. Beyond-rule-of-five migration results"); ax.bar(d.split,vals,color=["#3F6F9F","#CF7A3B","#B65A55","#76659A"][:len(d)]); ax.set(ylabel="Test RMSE"); fig.tight_layout(rect=[0,0,1,.95]); return page_pdf(fig,"Supplementary_Figure_S14_bRo5.pdf")

def build_s15():
    support=pd.read_csv(NEW/"chemical_support_selection_audit.csv"); neg=pd.read_csv(ROOT/"output"/"sci1_mechanism_uq_decision_20260707"/"clintox_minority_negative_result.csv").drop_duplicates("candidate")
    fig,axes=plt.subplots(1,2,figsize=(7.2,5.2)); title(fig,"Figure S15. Low-support and ClinTox failure cases")
    q=support[support.task_type.eq("classification")].groupby("tanimoto_bin",as_index=False).selected_performance.median(); order=[x for x in ["<0.5","0.5-0.7",">0.7"] if x in q.tanimoto_bin.values]; q=q.set_index("tanimoto_bin").loc[order]; axes[0].plot(["<0.5","0.5–<0.7","≥0.7"],q.selected_performance,"o-",color="#3F6F9F"); axes[0].set(xlabel="Maximum train-set Tanimoto",ylabel="Median ROC-AUC",ylim=(0,1))
    axes[1].barh(np.arange(len(neg)),neg.minority_recall,color="#B65A55"); axes[1].set(yticks=np.arange(len(neg)),yticklabels=neg.candidate,xlabel="ClinTox minority recall",xlim=(0,1)); fig.tight_layout(rect=[0,0,1,.95]); return page_pdf(fig,"Supplementary_Figure_S15_failure_cases.pdf")

def build_s16():
    d=pd.read_csv(CORE/"matched_k_multiview_units.csv"); d=d[d.pool_name.eq("ladder_K12")]
    fig,ax=plt.subplots(figsize=(7.2,5.2)); title(fig,"Figure S16. Multiview model-seed audit")
    for task,g in d.groupby("task"): ax.plot(sorted(g.seed.unique()),g.groupby("seed").gain_vs_morgan_k3.mean(),marker="o",label=DISPLAY[task])
    ax.set(xlabel="Model seed",ylabel="Mean paired gain"); ax.legend(frameon=False,ncol=3,fontsize=6.2); fig.tight_layout(rect=[0,0,1,.95]); return page_pdf(fig,"Supplementary_Figure_S16_multiview_seed_audit.pdf")

def build_s17():
    s=pd.read_csv(NEW/"seed_split_prediction_audit_summary.csv")
    a=pd.read_csv(NEW/"effective_rank_monte_carlo_stability.csv")
    fig,axs=plt.subplots(1,2,figsize=(7.2,3.1))
    colors=["#3F6F9F" if x=="classification" else "#CF7A3B" for x in s.task_type]
    axs[0].barh(np.arange(len(s)),s.unique_outer_split_assignments,color=colors); axs[0].set(yticks=np.arange(len(s)),yticklabels=s.endpoint.str.replace("tdc_",""),xlabel="Unique scaffold partitions",title="A  Seed/split uniqueness")
    axs[0].axvline(5,color="#7A838C",ls=":",lw=.8)
    for i,v in enumerate(s.max_abs_rerun_minus_stored): axs[0].text(max(s.unique_outer_split_assignments.iloc[i],.1)+.08,i,f"Δ≤{v:.1e}",va="center",fontsize=6)
    vals=a.abs_median_change_4000_to_5000; axs[1].hist(vals,bins=25,color="#3C8C82",alpha=.85); axs[1].axvline(vals.median(),color="#202830",lw=1,label=f"median={vals.median():.4f}"); axs[1].axvline(vals.quantile(.95),color="#76659A",lw=1,label=f"95th={vals.quantile(.95):.4f}")
    axs[1].set(xlabel="Absolute change, 4,000 to 5,000",ylabel="Analysis-metric records",title="B  Bootstrap Monte Carlo stability"); axs[1].legend(frameon=False,fontsize=7)
    fig.tight_layout(); return page_pdf(fig,"Supplementary_Figure_S17_split_and_bootstrap_audit.pdf")

def build_supp_pdf(s5, s6, pages):
    old=PdfReader(str(OLD/"Additional_file_3_Supplementary_Figures_S1-S14.pdf")); writer=PdfWriter()
    for p in old.pages[:4]: writer.add_page(p)
    for replacement in [s5,s6]:
        for p in PdfReader(str(replacement)).pages: writer.add_page(p)
    for p in old.pages[6:8]: writer.add_page(p)
    for page in pages:
        for p in PdfReader(str(page)).pages: writer.add_page(p)
    out=SUP/"Additional_file_3_Supplementary_Figures_S1-S17.pdf"
    with open(out,"wb") as f: writer.write(f)
    return out

def build_methods():
    d=Document(); d.add_heading("Supplementary Methods and Results",level=1)
    sections=[
      ("S1 Evidence hierarchy and reconstruction","The primary evidence is the frozen nine-endpoint candidate-prefix audit. Candidate-composition controls, finite-maximum simulation, the 12-candidate multiview registry and the four-model prediction panel answer narrower mechanism or boundary questions. Historical TDC, AutoGluon, conformal, MoleculeACE, bRo5, deduplication and failure-case records remain supplementary and are not treated as independent external confirmation."),
      ("S2 Seed, split and prediction-hash audit","Both classification and regression endpoints used five distinct seed-specific scaffold partitions. Regression Bemis–Murcko groups were allocated intact with seed-dependent random tie breaking and sample-count balancing; every outer-training partition contained three inner scaffold folds. The manifest records train, validation and test sample counts, scaffold counts, target mean, standard deviation and range, split hashes and no cross-fold scaffold overlap."),
      ("S3 Effective-rank bootstrap and reference sensitivity","Raw, row-centred, fixed-reference-relative and within-unit-rank matrices were analysed with empirical and Ledoit-Wolf correlation estimators, spectral-entropy rank, participation-ratio rank and median candidate correlation. At K = 32, the endpoint-median shrinkage entropy rank ranged from 2.98 for raw utility patterns to 27.14 for within-unit ranks. This range reflects matrix construction and removal of common audit-unit difficulty; it is not a count of independent predictive behaviours. A fixed master seed of 20260713 generated 5,000 hierarchical bootstrap replicates. Candidate 1, a predefined linear candidate, fixed Morgan-RF and the registry-median candidate were specified without consulting outer performance. Full draws contain 2,565,000 metric rows and are supplied separately as a compressed CSV."),
      ("S4 Chance-adjusted ranking aggregation","The unique machine-readable source contains 540 endpoint–split-seed–outer-fold–K rows: nine endpoints, five split seeds, three outer folds and K = 4, 8, 16 or 32. For each row, MRR equals 1/rank and normalized MRR gain equals (MRR − H_K/K)/(1 − H_K/K), where H_K/K is the random-order expectation. Each metric was averaged across folds within split seed, then across split seeds within endpoint; the main result is the median and IQR across the nine endpoint means. Classification and regression sensitivity summaries retain separate task strata."),
      ("S5 Mechanism-calibration controls","The random-rank negative control used 5,000 replicates and repeated the empirical fold-to-seed-to-endpoint aggregation. The graded positive control used 4,320 locked outer-candidate utilities from 135 audit units, four K values and six injected validation–audit signal levels; non-perfect cells used 500 simulations. Observed CAHit@3 exceeded the 95% permutation envelope at every K. CAHit@3 and normalized MRR gain increased monotonically, while fixed-range selection loss decreased monotonically, with injected signal at every K. Tables S21-S22 retain every summary and outer-unit value."),
      ("S5 Cross-fitted intervals","For every endpoint, the three outer folds were averaged within each of five unique split seeds before bootstrap resampling. Tables include same-unit and cross-fitted effects, split-seed bootstrap and t intervals, positive seed means, leave-one-seed ranges, their difference and fold-level effects. Classification ROC-AUC loss and regression RMSE loss are task-stratified co-primary evidence strata. Their units are not pooled and no cross-task average is calculated. Cross-fitted effects are positive for BACE, ClinTox, ESOL, Lipophilicity, Caco2 and HIA, and negative for BBBP, FreeSolv and P-gp; intervals exclude zero for ClinTox, HIA, ESOL, FreeSolv and Lipophilicity."),
      ("S6 Exhaustive matched K = 3 analysis","All C(12,3) = 220 subsets were enumerated from stored candidate-level inner and outer utilities without retraining. The endpoint remains the interpretation unit; overlapping subsets are sensitivity contrasts within one registry rather than independent experiments. Mutually exclusive composition classes, medians, IQRs, distribution ranges and within-endpoint outperforming-Morgan proportions are supplied. CAHit@3 is not estimable because both observed and random Hit@3 equal one at K = 3."),
      ("S7 Chemical-support and scaffold-boundary audit","Molecules were stratified by maximum training-set Morgan Tanimoto as <0.5, 0.5 to <0.7 and at least 0.7. Traceable selected-candidate performance, error overlap, disagreement, uncertainty–error association and false-negative or high-error enrichment are reported by support stratum. Novel-scaffold quantities are expressed relative to seen-or-related scaffolds as reference = 1. These analyses are exploratory reliability boundaries rather than independent validation."),
      ("S8 TDC endpoint analysis","Twenty-two TDC endpoint summaries retain official metric direction, the historical comparator, performance-run source, retained model and direction-normalized change. Evaluation budgets differ from the primary audit, so these rows provide contextual reliability evidence only."),
      ("S9 Automated machine-learning budgets","AutoGluon summaries cover 30-, 300- and 1,800-second budgets across nine datasets, with 27 budget summaries and 81 outer-fold records. Fit time, model count, memory, refit status and comparisons with frozen selector baselines are retained in Table S19. Unequal feature preparation, hardware and search exposure preclude a superiority claim."),
      ("S10 Selection-risk and leave-one-endpoint-out analyses","Selection-risk correlations, null calibration and nine leave-one-endpoint-out policy records are retained in Table S20. They are negative or exploratory evidence: nine endpoints cannot identify a generally transferable meta-selector, and endpoint meta-features were evaluated retrospectively."),
      ("S11 Conformal prediction and CQR","Table S14 contains 189 classification-conformal and regression-CQR summaries at 80%, 90% and 95% nominal coverage where available. Label-conditional and similarity-conditioned methods improved minority coverage in several tasks but enlarged prediction sets. Coverage remains conditional on the evaluated folds and does not guarantee screening utility after a new chemical shift."),
      ("S12 MoleculeACE, bRo5 and deduplication sensitivity","MoleculeACE includes 17 task summaries of cliff-pair direction accuracy, gap association and pair count. The bRo5 panel retains four split summaries, including random, scaffold, perimeter and time settings where available. Eighteen deduplication-policy summaries compare global, training-fold and retained-duplicate strategies. None is presented as independent confirmation of the primary candidate-expansion effect."),
      ("S13 Representation errors and failure cases","The limited prediction panel contains RDKit-RF, a GCN and frozen ChemBERTa and MoLFormer probes. Prediction and error correlations were incomplete, and training budgets were unequal. Low-support, novel-scaffold, activity-cliff, bRo5 and ClinTox minority failures delimit the evidence rather than establish a modern-model leaderboard."),
      ("S14 Source limitations","All-candidate molecule-level predictions are complete for the regression registry and for a post-hoc rerun of the classification 32-candidate registry. The classification rerun showed small stochastic refit drift (maximum absolute outer-candidate metric difference 0.0010 and two one-standard-error selection changes), so locked primary metrics remain the source of record and the new prediction exports are used only for traceability. Complete candidate-level prediction exports remain unavailable across the full multiview registry, so prediction-level diversity is not evaluated for the matched-subset analysis. The public endpoint panels are retrospective, and independently reproduced prospective validation is outside the present study scope."),
    ]
    for h,p in sections: d.add_heading(h,level=2); d.add_paragraph(p)
    table_names=["Dataset provenance and cleaning","Complete candidate registry","Candidate configurations","Split and seed audit","Fold-level utilities","Effective-rank bootstrap","Reference sensitivity","Cross-fitted intervals","Matched-K subset results","Full multiview results","Prediction correlations and error overlap","Chemical-support stratification","Deduplication","Conformal prediction","TDC","MoleculeACE","bRo5","Failed candidates","AutoGluon budgets","Selection risk and LOEO","Mechanism permutation null","Graded signal recovery"]
    d.add_heading("Supplementary table directory",level=1)
    for i,name in enumerate(table_names,1): d.add_paragraph(f"Table S{i}. {name}.")
    figure_names=["Cleaning flow","Candidate correlation matrices","Eigenvalue spectra","Raw ranking metrics","Permutation controls","Signal recovery","Full winner-optimism simulation","Fold-level audit gaps","Full matched-K subset distributions","Representation selection frequency","TDC","Conformal and risk-coverage","MoleculeACE","bRo5","Failure cases","Multiview model-seed audit","Split and bootstrap audit"]
    d.add_heading("Supplementary figure directory",level=1)
    for i,name in enumerate(figure_names,1): d.add_paragraph(f"Figure S{i}. {name}.")
    for p in d.paragraphs:
        p.paragraph_format.line_spacing=1.5
        for r in p.runs: r.font.name="Times New Roman"; r.font.size=Pt(10)
    for sec in d.sections: add_page_number(sec)
    out=SUP/"Additional_file_1_Supplementary_Methods_and_Results.docx"; d.save(out); return out

def build_response():
    rows=[
      ("Condense the structured Abstract","Rewrote the structured Abstract to 318 words, retained decision-relevant directional findings and added only the mechanism-calibration result needed to support the revised Scientific Contribution.","Abstract; Scientific Contribution"),
      ("Expand the Introduction into six paragraphs","Expanded the Introduction to six purpose-specific paragraphs covering the scientific problem, two sources of apparent candidate expansion, nested-validation safeguards, molecular-domain risks, the unresolved gap and the locked study aims without importing result values.","Introduction, paragraphs 1-6"),
      ("Realign Discussion sections 4.1-4.7","Reorganized the Discussion so that each registered heading now addresses its named topic: nominal versus effective diversity, matrix construction, chance-adjusted ranking, cross-fitting, matched-size multiview evidence, chemical-support reliability and limitations.","Discussion 4.1-4.7"),
      ("Refine Figure 6D without a composite score","Combined ClinTox discrimination and reliability in one aligned four-model dot plot rather than nested subplots. ROC-AUC and PR-AUC are shown beside minority recall, false-negative rate, conditional coverage and mean prediction-set size; the 90% coverage target is retained and no radar chart or composite score is used.","Figure 6D and legend"),
      ("Retain all Declarations headings","Restored every journal-required declaration heading. Ethics, consent and availability statements remain populated; competing interests, funding, author contributions and acknowledgements carry explicit author-confirmation placeholders because author-specific facts were not supplied.","Declarations; completion checklist"),
      ("Use one normalized-MRR source and aggregation","Designated chance_adjusted_ranking_units.csv as the unique unit-level source, recomputed MRR and H_K/K normalization, averaged folds within split seed and split seeds within endpoint, and summarized the nine endpoint means by median and IQR.","Methods 2.5; Results 3.2; Figure 3A; Table S5; verification workbook"),
      ("Calibrate ranking metrics with current-data controls","Recomputed a 5,000-replicate random-rank negative control and six-level signal-recovery positive control from the corrected nine-endpoint, five-split-seed locked outer utilities. Replaced legacy Supplementary Figures S5-S6 and added machine-readable Tables S21-S22.","Methods 2.5; Results 3.2; Figure 3A-B; Supplementary Figures S5-S6; Tables S21-S22"),
      ("Avoid a single intrinsic effective-candidate count","Reframed effective diversity as matrix dependent and reported raw, row-centred, fixed-reference-relative and within-unit-rank estimates side by side. The text no longer interprets the low raw rank as a literal count of independent predictors.","Abstract; Methods 2.6; Results 3.1; Discussion 4.1-4.2; Figure 2; Tables S6-S7"),
      ("Treat classification and regression symmetrically","Defined classification ROC-AUC loss and regression RMSE loss as task-stratified co-primary evidence strata, retained separate units and axes, and calculated no pooled cross-task effect.","Methods 2.13; Results 3.4; Table 3; Figures 3C and 4C"),
      ("Report cross-fitted direction and interval interpretation","Reported six positive and three negative endpoint effects. Intervals exclude zero for ClinTox, HIA, ESOL, FreeSolv and Lipophilicity; FreeSolv is retained as a significant negative result. Table 3 interpretations now depend on whether the interval crosses zero.","Abstract; Results 3.4; Table 3; Figure 3C; Table S8"),
      ("Separate machine-readable verification from exposition","Created a master result-consistency table and dedicated ranking, effective-rank, cross-fitted, matched-subset and chemical-support verification reports with source hashes and manuscript/figure locations.","Minor revision verification workbook and CSV/JSON reports"),
      ("Clarify scientific positioning","Reframed as a retrospective nested audit; removed predictor, universal-selector and external-validation claims.","Title; Abstract; Introduction; Discussion"),
      ("Verify whether five seeds represent distinct splits","Generated five distinct seeded scaffold partitions for both classification and regression; audited sample, target, scaffold-overlap and split-hash fields.","Methods 2.4; Table 3; Table S4"),
      ("Correct regression uncertainty","Reran the four regression endpoints on five distinct scaffold partitions and replaced model-seed sensitivity with split-seed intervals.","Methods 2.13; Table 3; Figure 3C; Table S8"),
      ("Increase effective-rank bootstrap","Completed 5,000 hierarchical replicates with fixed seed, stability tracking and all four matrix constructions.","Methods 2.6; Results 3.1; Figure 2; Table S6"),
      ("Test reference sensitivity","Added four predefined references plus leave-one-seed and leave-one-fold analyses.","Figure 2D; Table S7"),
      ("Complete cross-fitted intervals","Reported same-unit, cross-fitted and difference effects for both task strata, with split-seed intervals, positive seed means and fold effects.","Table 3; Table S8"),
      ("Separate classification and regression axes","Redrew all raw-effect panels with independent ROC-AUC and RMSE axes.","Figures 3C, 4A-B, 5D, 6B"),
      ("Complete matched K = 3 analysis","Enumerated all 220 subsets and reported composition, gain, cross-fitted gap, diversity, normalized MRR and selection frequencies.","Methods 2.10; Results 3.6; Figure 5; Table S9"),
      ("Explain CAHit@3 at K = 3","Marked it not estimable because observed and random Hit@3 are both one.","Methods 2.10; Table S9"),
      ("Add chemical-support audit","Added low/intermediate/high Tanimoto strata for performance, overlap, disagreement, uncertainty and false-negative enrichment.","Methods 2.12; Results 3.8; Figure 6B; Table S12"),
      ("Add scaffold-novelty complementarity","Added prespecified scaffold-fingerprint relatedness threshold and compared error overlap, disagreement and enrichment.","Methods 2.12; Figure 6C; Table S12"),
      ("Avoid architecture leaderboard claims","Restricted the four-model panel to fixed-configuration error structure and disclosed unequal training budgets.","Methods 2.11; Results 3.7; Limitations"),
      ("Retain ClinTox as a negative result","Reported ROC-AUC, PR-AUC, minority recall, false-negative rate and conditional coverage without deployment claims.","Results 3.8; Figure 6D"),
      ("Use six evidence-driven main figures","Redrew Figure 1 as a continuous architecture and Figures 2-6 as four-panel figures. Figure 6 retained because support and scaffold analyses completed.","Figures 1-6"),
      ("Standardize tables and manuscript format","Limited the main text to three editable three-line tables; added double spacing, continuous line numbering and automatic page numbering.","Main manuscript; Tables 1-3"),
      ("Provide Chinese counterpart","Created an independently worded Chinese academic manuscript using the same numerical evidence, tables, figures and limitations.","Chinese complete manuscript"),
    ]
    d=Document(); d.add_heading("Reviewer concern–Response–Location",0); t=d.add_table(rows=1,cols=3); t.style="Table Grid"
    for c,x in zip(t.rows[0].cells,["Reviewer concern","Response","Location"]): c.text=x
    for row in rows:
        for c,x in zip(t.add_row().cells,row): c.text=x
    out=MINOR/"Reviewer_concern_Response_Location.docx"; d.save(out); return out

def build_audit(files):
    audit={"generated_files":{Path(p).name:{"sha256":sha(p),"bytes":Path(p).stat().st_size} for p in files},
      "requires_source_data":["author names and CRediT contributions","funding statement","competing-interest statement","acknowledgements"],
      "completed_new_analysis":["post-hoc all-candidate per-molecule prediction exports for the classification 32-candidate registry, with refit-drift audit and no replacement of locked primary metrics","current-data random-rank permutation and graded signal-recovery mechanism controls"],
      "requires_new_analysis":[],
      "out_of_scope_future_work":["prospective or independently reproduced external validation","complete candidate-level prediction exports across the full multiview registry"],
      "completed_checks":{"classification_split_seeds":"5 distinct partitions per endpoint","regression_split_seeds":"5 distinct partitions per endpoint","regression_no_cross_fold_scaffold_overlap":True,"effective_rank_bootstrap":5000,"matched_K3_subsets":220,"mechanism_permutations":5000,"mechanism_signal_levels":6,"classification_regression_axes_separated":True,"main_tables_three_line":True,"manual_page_breaks":False,"main_figures":6,"supplementary_tables":22,"supplementary_figures":17}}
    out=MINOR/"Final_scientific_and_format_audit.json"; out.write_text(json.dumps(audit,ensure_ascii=False,indent=2),encoding="utf-8"); return out

def main():
    x=build_xlsx(); s5=build_s5(); s6=build_s6(); pages=[build_s9(),build_s10(),build_s11(),build_s12(),build_s13(),build_s14(),build_s15(),build_s16(),build_s17()]; pdf=build_supp_pdf(s5,s6,pages); methods=build_methods(); response=build_response()
    audit=build_audit([x,s5,s6,*pages,pdf,methods,response,NEW/"effective_rank_bootstrap_5000_draws.csv.gz"])
    print("\n".join(map(str,[methods,x,pdf,response,audit])))

if __name__=="__main__": main()
