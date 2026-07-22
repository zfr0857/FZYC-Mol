from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_academic_polished.docx"
MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_academic_polished.md"
REPORT = DOCS / "academic_language_polish_report_20260604.md"


REPLACEMENTS = [
    (
        "这样既能回应“为什么不直接训练更大模型”的问题，也能解释为什么某些 pilot 结果应保留在 appendix 而不覆盖 retained-best。",
        "这样既能回应“为何不以更大规模模型训练作为主线”的问题，也能解释为什么某些 pilot 结果应保留在 appendix，而不覆盖 retained-best。",
    ),
    (
        "ROC-AUC 高不一定意味着阳性样本召回、概率校准或筛选富集可靠。",
        "ROC-AUC 较高并不必然意味着阳性样本召回、概率校准或筛选富集可靠。",
    ),
    (
        "若测试分子与最近邻高度相似但标签差异或归一化 target jump 很大，则说明该 endpoint 具有高 roughness，仅增大模型容量不一定有效，通常需要 target transform、robust loss、局部邻域诊断或 ensemble 稳定化。",
        "若测试分子与最近邻高度相似但标签差异或归一化 target jump 显著，则说明该 endpoint 具有高 roughness。单纯增大模型容量通常不足以解决该问题，需要 target transform、robust loss、局部邻域诊断或 ensemble 稳定化。",
    ),
    (
        "可靠性结果显示，单一不确定性分数很难覆盖所有错误类型。",
        "可靠性结果显示，单一不确定性分数难以覆盖所有错误类型。",
    ),
    (
        "注：弱候选不替换原 validation selector。",
        "注：未通过验证的候选不替换原 validation selector。",
    ),
    (
        "若 endpoint 的主指标是 Spearman，模型不一定需要最小化绝对误差，而应更关注排序稳定性；若主指标是 MAE/RMSE，则应重点处理极端值和长尾 target。",
        "若 endpoint 的主指标是 Spearman，模型目标不应仅限于最小化绝对误差，而应更关注排序稳定性；若主指标是 MAE/RMSE，则应重点处理极端值和长尾 target。",
    ),
    (
        "授权完成后，可以直接使用同一脚本重跑完整 validation-only integration，再决定是否进入主文或强附录。",
        "授权完成后，可使用同一脚本重新运行完整 validation-only integration，再决定是否进入主文或强附录。",
    ),
    (
        "避免把无法完全复现 split 的模型误写成直接可比结果。",
        "避免将无法完全复现 split 的模型表述为同一条件下的直接对照。",
    ),
    (
        "这个 registry 的作用是把 same-split head-to-head 对照和 literature-reported reference 分开，避免将无法完全复现 split 的模型表述为同一条件下的直接对照。",
        "该 registry 用于区分 same-split head-to-head 对照与 literature-reported reference，避免将无法完全复现 split 的模型表述为同一条件下的可比对照。",
    ),
    (
        "直接扩大模型可能提高某些 endpoint，但也会增加算力成本、过拟合风险和验证选择复杂度。",
        "单纯扩大模型规模可能提高某些 endpoint，但也会增加算力成本、过拟合风险和验证选择复杂度。",
    ),
    (
        "后续实验应优先扩展外部 benchmark/appendix，而非首先进行大规模模型重训练。优先级最高的是 fast external appendix benchmark：在更多公开 ADMET endpoint 上跑 CatBoost/XGBoost/ExtraTrees/LGBM/RF、Top-K/stacking、target transform 和 undersampling ensemble，并把 retained-best 与 roughness 诊断绑定。这样可以提高证据完整性，同时保持算力成本可控。",
        "后续实验应优先扩展外部 benchmark/appendix，而非首先进行大规模模型重训练。首要方向是 fast external appendix benchmark：在更多公开 ADMET endpoint 上系统评估 CatBoost/XGBoost/ExtraTrees/LGBM/RF、Top-K/stacking、target transform 和 undersampling ensemble，并将 retained-best 决策与 roughness 诊断关联。该策略可在控制计算成本的同时提高证据完整性。",
    ),
    (
        "performance-mode 候选本身不会在所有 endpoint 上胜出。",
        "performance-mode 候选本身并非在所有 endpoint 上占优。",
    ),
]


def apply_replacements(text: str) -> str:
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text


def set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = text


def polish_docx() -> None:
    doc = Document(DOCX)
    for paragraph in doc.paragraphs:
        cleaned = apply_replacements(paragraph.text)
        if cleaned != paragraph.text:
            set_paragraph_text(paragraph, cleaned)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    cleaned = apply_replacements(paragraph.text)
                    if cleaned != paragraph.text:
                        set_paragraph_text(paragraph, cleaned)
    doc.save(DOCX)


def polish_markdown() -> None:
    text = MD.read_text(encoding="utf-8")
    MD.write_text(apply_replacements(text), encoding="utf-8")


def update_report() -> None:
    text = REPORT.read_text(encoding="utf-8") if REPORT.exists() else "# Academic Language Polish Report\n"
    addition = (
        "\nSecond-pass refinements:\n"
        "- Removed residual colloquial expressions such as running benchmarks, weak candidates, and direct model restart wording.\n"
        "- Rephrased model-capacity, roughness, ROC-AUC, validation-only, and retained-best statements in a more formal academic style.\n"
        "- Preserved all numerical results, tables, figures, and experimental conclusions.\n"
    )
    if "Second-pass refinements:" not in text:
        REPORT.write_text(text.rstrip() + "\n" + addition, encoding="utf-8")


def main() -> None:
    polish_docx()
    polish_markdown()
    update_report()
    print(f"polished {DOCX.relative_to(ROOT)}")
    print(f"polished {MD.relative_to(ROOT)}")
    print(f"updated {REPORT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
