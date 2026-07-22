from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(r"D:\fzyc\output\paper33_final_minor_revision_20260718")
SOURCE = ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_FINAL_CLEAN.docx"
OUT = ROOT / "revision_extracts"
OUT.mkdir(parents=True, exist_ok=True)

spec = importlib.util.spec_from_file_location("builder", Path(r"D:\fzyc\scripts\build_paper33_final_minor_revision_manuscripts_20260718.py"))
builder = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(builder)


def set_font(doc: Document) -> None:
    for style in doc.styles:
        if style.type == 1:
            style.font.name = "Times New Roman"
            rfonts = style.element.get_or_add_rPr().get_or_add_rFonts()
            rfonts.set(qn("w:ascii"), "Times New Roman")
            rfonts.set(qn("w:hAnsi"), "Times New Roman")
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def paragraphs_between(source: Document, start: str, end: str) -> list[tuple[str, str]]:
    paragraphs = source.paragraphs
    i = next(index for index, p in enumerate(paragraphs) if p.text.strip().startswith(start))
    j = next(index for index, p in enumerate(paragraphs[i + 1:], start=i + 1) if p.text.strip().startswith(end))
    return [(p.style.name, p.text) for p in paragraphs[i:j] if p.text.strip()]


def write_docx(path: Path, title: str, sections: list[list[tuple[str, str]]]) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for section in sections:
        for style, text in section:
            if style.startswith("Heading"):
                level = min(int(style.split()[-1]), 3)
                doc.add_heading(text, level=level)
            else:
                doc.add_paragraph(text)
    set_font(doc)
    doc.save(path)


def table2(source: Document) -> None:
    table = source.tables[1]
    doc = Document()
    doc.add_heading("Table 2. Audit components and recorded exposure.", level=1)
    target = doc.add_table(rows=1, cols=4)
    for c, source_cell in enumerate(table.rows[0].cells):
        target.rows[0].cells[c].text = source_cell.text
    for source_row in table.rows[1:]:
        cells = target.add_row().cells
        for c, source_cell in enumerate(source_row.cells):
            cells[c].text = source_cell.text
    builder.three_line(target)
    set_font(doc)
    doc.save(OUT / "Updated_Table2.docx")

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Table 2"
    for row in table.rows:
        sheet.append([cell.text for cell in row.cells])
    for cell in sheet[1]:
        cell.font = Font(name="Times New Roman", bold=True)
        cell.fill = PatternFill("solid", fgColor="D9EAF2")
    thin = Side(style="thin", color="000000")
    medium = Side(style="medium", color="000000")
    for row_index, row in enumerate(sheet.iter_rows(), start=1):
        for cell in row:
            cell.font = Font(name="Times New Roman", bold=row_index == 1)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(top=medium if row_index == 1 else None,
                                 bottom=medium if row_index == sheet.max_row else (thin if row_index == 1 else None))
    for index, width in enumerate([31, 28, 68, 53], start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    sheet.freeze_panes = "A2"
    workbook.save(OUT / "Updated_Table2.xlsx")


def main() -> None:
    source = Document(SOURCE)
    abstract = [
        p.text for p in source.paragraphs
        if p.text.startswith(("Background:", "Methods:", "Results:", "Conclusions:", "Scientific Contribution:"))
    ]
    (OUT / "Updated_Abstract.txt").write_text("\n\n".join(abstract) + "\n", encoding="utf-8")
    write_docx(OUT / "Updated_Abstract.docx", "Updated Abstract", [[("Normal", text) for text in abstract]])

    methods = [
        paragraphs_between(source, "2.1 ", "2.2 "),
        paragraphs_between(source, "2.3 ", "2.4 "),
        paragraphs_between(source, "2.14 ", "2.15 "),
        paragraphs_between(source, "2.16 ", "3 Results"),
    ]
    write_docx(OUT / "Updated_Methods_2.1_2.3_2.14_2.16.docx", "Updated Methods", methods)
    write_docx(OUT / "Updated_Results_3.10.docx", "Updated Results 3.10", [paragraphs_between(source, "3.10 ", "4 Discussion")])
    write_docx(
        OUT / "Updated_Discussion_4.8_and_Conclusions.docx",
        "Updated Discussion 4.8 and Conclusions",
        [paragraphs_between(source, "4.8 ", "4.9 "), paragraphs_between(source, "5 Conclusions", "Supplementary Information")],
    )
    table2(source)
    print(OUT)


if __name__ == "__main__":
    main()
