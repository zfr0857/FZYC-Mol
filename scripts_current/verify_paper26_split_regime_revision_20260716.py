from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from lxml import etree
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "results" / "split_regime_transfer_20260716" / "similarity_cluster"
ANALYSIS = ROOT / "output" / "paper26_split_regime_transfer_20260716"
BASE = ROOT / "output" / "paper26_split_regime_transfer_revision_20260716"
SUPP = BASE / "supplementary"
SEEDS = (11, 23, 37, 53, 71)
TASKS = ("clintox", "bace", "esol")


def sha(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1 << 20), b""):
            digest.update(block)
    return digest.hexdigest()


def document_text(path: Path) -> str:
    doc = Document(path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def main() -> None:
    checks = []

    def check(name: str, passed: bool, evidence: object) -> None:
        checks.append({"check": name, "passed": bool(passed), "evidence": evidence})

    inner_rows = outer_rows = policy_rows = split_rows = 0
    split_frames = []
    complete = 0
    for seed in SEEDS:
        base = RUN / f"seed_{seed}"
        inner_rows += len(pd.read_csv(base / "inner_scores.csv"))
        outer_rows += len(pd.read_csv(base / "outer_candidate_scores.csv"))
        policy_rows += len(pd.read_csv(base / "policy_detail.csv"))
        splits = pd.read_csv(base / "split_manifest.csv")
        split_rows += len(splits)
        split_frames.append(splits)
        complete += sum((base / "tasks" / task / "complete.json").exists() for task in TASKS)
    split = pd.concat(split_frames, ignore_index=True)
    check("all formal task-seed runs complete", complete == 15, complete)
    check("new inner candidate utilities", inner_rows == 4320, inner_rows)
    check("new outer candidate utilities", outer_rows == 1440, outer_rows)
    check("policy audit rows", policy_rows == 900, policy_rows)
    check("split-manifest rows", split_rows == 135, split_rows)
    check("no random split fallback", not split["outer_split_type"].str.contains("random").any() and not split["inner_split_type"].str.contains("random").any(), sorted(set(split.outer_split_type) | set(split.inner_split_type)))
    check("all component groups disjoint", bool(split["no_group_overlap"].all()), int(split["no_group_overlap"].sum()))
    maximum = float(split[["max_train_validation_tanimoto", "max_train_test_tanimoto", "max_validation_test_tanimoto"]].max().max())
    check("cross-fold Tanimoto below threshold", maximum < 0.70, maximum)
    unique_outer = split.groupby("endpoint")["outer_split_hash"].nunique().to_dict()
    check("15 unique outer-fold assignments per endpoint", all(value == 15 for value in unique_outer.values()), unique_outer)
    class_rows = split.loc[split["task_type"].eq("classification")]
    class_complete = all(
        (class_rows[f"{part}_target_min"].eq(0) & class_rows[f"{part}_target_max"].eq(1)).all()
        for part in ("train", "validation", "test")
    )
    check("both classes retained in every classification fold", class_complete, len(class_rows))

    audit = json.loads((ANALYSIS / "split_regime_transfer_audit.json").read_text(encoding="utf-8"))
    check("analysis audit complete", audit["status"] == "complete", audit["status"])
    check("CAHit direction agreement in three endpoints", audit["same_cahit_change_direction_count"] == 3, audit["same_cahit_change_direction_count"])
    check("cross-fitted direction agreement in three endpoints", audit["same_cross_effect_direction_count"] == 3, audit["same_cross_effect_direction_count"])
    check("effective-rank cross-regime correlation recorded", abs(audit["effective_diversity_spearman_between_regimes"] - 0.8316326530612245) < 1e-12, audit["effective_diversity_spearman_between_regimes"])

    en = BASE / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript.docx"
    zh = next(path for path in BASE.glob("*.docx") if path.name != en.name and "Reviewer" not in path.name and "tracked" not in path.name)
    tracked = BASE / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript_tracked_changes.docx"
    en_text, zh_text = document_text(en), document_text(zh)
    for phrase in ["2.13 Split-regime transfer audit", "3.9 Audit directions transferred", "4.7 Split-regime transfer reduced", "4.8 Limitations", "Tables S1-S26", "Figures S1-S18", "0.832"]:
        check(f"English manuscript contains {phrase}", phrase in en_text, phrase)
    for phrase in ["2.13 切分机制迁移审计", "3.9 结构分离下审计方向迁移", "4.7 切分迁移降低", "4.8 局限性", "S1–S26", "S1–S18", "0.832"]:
        check(f"Chinese manuscript contains {phrase}", phrase in zh_text, phrase)
    check("Table 2 records new exposure", "5,760" in "\n".join(cell.text for row in Document(en).tables[1].rows for cell in row.cells), "5,760 fits")

    with zipfile.ZipFile(tracked) as archive:
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        settings_xml = etree.fromstring(archive.read("word/settings.xml"))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    insertions = len(document_xml.xpath("//w:ins", namespaces=ns))
    deletions = len(document_xml.xpath("//w:del", namespaces=ns))
    check("tracked manuscript has insertions and deletions", insertions >= 30 and deletions >= 10, {"insertions": insertions, "deletions": deletions})
    check("tracked revisions enabled", bool(settings_xml.xpath("//w:trackRevisions", namespaces=ns)), True)

    workbook = SUPP / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S26.xlsx"
    sheets = load_workbook(workbook, read_only=True).sheetnames
    required_sheets = {"S23 Split-regime ranking", "S24 Split-regime cross-fit", "S25 Split-regime diversity", "S26 Similarity split audit"}
    check("supplementary workbook contains S23-S26", required_sheets.issubset(sheets), sheets[-4:])
    old_pdf = SUPP / "Additional_file_3_Supplementary_Figures_S1-S17.pdf"
    new_pdf = SUPP / "Additional_file_3_Supplementary_Figures_S1-S18.pdf"
    old_pages, new_pages = len(PdfReader(old_pdf).pages), len(PdfReader(new_pdf).pages)
    check("supplementary figure PDF adds one page", new_pages == old_pages + 1, {"old": old_pages, "new": new_pages})
    png = SUPP / "Supplementary_Figure_S18_split_regime_transfer_600dpi.png"
    width, height = Image.open(png).size
    check("S18 high-resolution raster", width >= 4000 and height >= 3000, {"width": width, "height": height})
    check("S18 vector outputs present", (SUPP / "Supplementary_Figure_S18_split_regime_transfer.svg").exists() and (SUPP / "Supplementary_Figure_S18_split_regime_transfer.pdf").exists(), True)

    rendered = BASE / "rendered"
    rendered_pages = {name: len(PdfReader(rendered / name).pages) for name in ["English_clean.pdf", "Chinese_clean.pdf", "Supplementary_methods.pdf", "English_tracked_final_view.pdf"]}
    check("all Word documents rendered", all(value > 0 for value in rendered_pages.values()), rendered_pages)

    pending = ["Competing interests", "Funding", "Authors' contributions", "Acknowledgements"]
    check("author declarations remain transparently pending", all("Author confirmation required" in en_text or heading == "Authors' contributions" for heading in pending), pending)

    frame = pd.DataFrame(checks)
    frame.to_csv(BASE / "Split_regime_transfer_final_checklist.csv", index=False)
    final = {
        "status": "complete" if frame["passed"].all() else "failed",
        "passed": int(frame["passed"].sum()), "total": len(frame),
        "failed_checks": frame.loc[~frame["passed"], "check"].tolist(),
        "key_hashes": {
            "english_manuscript": sha(en), "tracked_manuscript": sha(tracked),
            "supplementary_workbook": sha(workbook), "supplementary_figures": sha(new_pdf),
            "analysis_audit": sha(ANALYSIS / "split_regime_transfer_audit.json"),
        },
    }
    (BASE / "Split_regime_transfer_final_audit.json").write_text(json.dumps(final, indent=2), encoding="utf-8")
    print(json.dumps(final, indent=2))
    if final["status"] != "complete":
        raise SystemExit(1)


if __name__ == "__main__":
    main()
