from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from pypdf import PdfReader, PdfWriter


ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "output" / "paper26_split_regime_transfer_revision_20260716"
ANALYSIS = ROOT / "output" / "paper26_split_regime_transfer_20260716"
EN = BASE / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript.docx"
ZH = next(path for path in BASE.glob("*.docx") if path.name != EN.name and "Reviewer" not in path.name)
SUPP = BASE / "supplementary"


EN_ABSTRACT_METHODS = (
    "Methods: We conducted a retrospective nested audit across nine public molecular-property endpoints using repeated "
    "seeded scaffold partitions. Progressively expanded candidate registries were evaluated through matrix-dependent "
    "diversity estimators, chance-adjusted ranking measures, permutation and graded signal-recovery controls, leave-one-seed-out "
    "cross-fitted references and matched-size multiview comparisons. A split-regime transfer audit retrained the same 32-candidate "
    "registry on ClinTox, BACE and ESOL using Morgan-512 Tanimoto connected components as intact groups. Reliability was further "
    "examined across chemical-similarity and scaffold-support boundaries."
)
EN_ABSTRACT_RESULTS = (
    "Results: Estimated candidate diversity changed markedly after common audit-unit difficulty and utility-level shifts were "
    "removed, showing that no single effective-rank value adequately characterized the registry. Chance-adjusted validation-ranking "
    "fidelity weakened as the candidate pool expanded. Observed adjusted ranking exceeded the permutation envelope at every K, while "
    "injected validation–audit signal monotonically restored recovery and reduced selection loss. Cross-fitted selection gaps increased "
    "in six of nine endpoints but decreased in three. In the three-endpoint split-regime transfer audit, K = 32 minus K = 4 CAHit@3 "
    "changes were negative and cross-fitted gap directions were positive under both split regimes for all three endpoints, whereas "
    "effect magnitude and interval exclusion remained endpoint dependent. Effective-rank estimates had Spearman correlation 0.832 "
    "between regimes. Matched-size multiview effects and chemical-support reliability were also endpoint dependent."
)
EN_ABSTRACT_CONCLUSIONS = (
    "Conclusions: Nominal candidate count alone did not describe the effective structure or risk of molecular model selection. "
    "Candidate expansion created both representational opportunity and additional selection pressure, and the direction of selected "
    "audit quantities transferred to a stricter structure-separated split in three representative endpoints without implying invariant "
    "effect size. Molecular benchmarks should jointly report candidate eligibility, matrix-dependent diversity, chance-adjusted ranking, "
    "cross-fitted gaps, split-regime sensitivity, computational exposure and chemical-support boundaries."
)
EN_CONTRIBUTION = (
    "Scientific Contribution: This study separates nominal candidate-pool size from matrix-dependent utility-pattern diversity and "
    "chance-adjusted ranking degradation, calibrated by negative and positive controls. It combines cross-fitted selection-gap analysis "
    "with exhaustive matched-size representation comparisons, chemical-support auditing and a fully nested split-regime transfer audit "
    "that distinguishes direction transport from invariant effect magnitude. The contribution is an analysis of molecular model-selection "
    "practice rather than a new predictor, universal selector or external validation study."
)


def find(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise KeyError(prefix)


def set_text(paragraph, text: str, *, east_asia: str | None = None) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    if east_asia:
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def insert_after(paragraph, text: str, style: str, *, east_asia: str | None = None):
    new = paragraph._parent.add_paragraph(style=style)
    paragraph._p.addnext(new._p)
    set_text(new, text, east_asia=east_asia)
    return new


def add_table2_row(doc: Document) -> None:
    table = doc.tables[1]
    values = [
        "Split-regime transfer", "32 Morgan candidates",
        "3 endpoints; 5 seeds; scaffold and Tanimoto-component splits",
        "Split robustness", "5,760 new candidate fits",
    ]
    if any(row.cells[0].text.strip() == values[0] for row in table.rows):
        return
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        cell.text = value
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(8.5)


def update_english() -> None:
    doc = Document(EN)
    set_text(find(doc, "Methods:"), EN_ABSTRACT_METHODS)
    set_text(find(doc, "Results:"), EN_ABSTRACT_RESULTS)
    set_text(find(doc, "Conclusions:"), EN_ABSTRACT_CONCLUSIONS)
    set_text(find(doc, "Scientific Contribution:"), EN_CONTRIBUTION)
    set_text(
        find(doc, "Under a retrospectively locked"),
        "Under a retrospectively locked repeated nested scaffold evaluation, how did candidate-pool expansion relate to matrix-dependent "
        "utility-pattern diversity, chance-adjusted ranking fidelity, cross-fitted selection gaps and representation-composition effects, "
        "and did the principal audit directions persist under a stricter structure-separated split? We address this limited question by "
        "separating nominal K from matrix-defined diversity, calibrating ranking measures against negative and positive controls, combining "
        "cross-fitted and matched-size analyses, and adding a three-endpoint split-regime transfer audit. Candidate eligibility, failed fits, "
        "split identities, source hashes and computational exposure remain part of the audit trail. The study evaluates model-selection "
        "behaviour and does not propose a new molecular predictor or an independent external validation."
    )

    stats = find(doc, "2.13 Statistical inference")
    set_text(stats, "2.14 Statistical inference")
    anchor = find(doc, "These analyses were retained")
    heading = insert_after(anchor, "2.13 Split-regime transfer audit", "Heading 2")
    p1 = insert_after(
        heading,
        "ClinTox, BACE and ESOL represented rare-class classification, regular classification and regression, respectively. The locked "
        "32-candidate Morgan-512 registry, candidate order, K = 4, 8, 16 and 32 prefixes, seeds 11, 23, 37, 53 and 71, and three outer "
        "by three inner folds were retained. Scaffold results reused the locked source-of-record outputs. For the new structure-separated "
        "rerun, molecular pairs with Morgan-512 Tanimoto similarity at least 0.70 were connected, and each connected component was kept "
        "intact during outer and inner allocation. The threshold and allocation rule were fixed before formal model fitting after splitter "
        "feasibility checks; model performance did not inform grouping.",
        "Normal",
    )
    insert_after(
        p1,
        "A seeded greedy allocator balanced sample counts and, for classification, both class counts across intact components. Every formal "
        "inner and outer split was required to retain both classes where applicable, keep components disjoint and have maximum cross-fold "
        "Tanimoto below 0.70; random fallback was prohibited. The transfer audit added 5,760 candidate fits. Direction transport was assessed "
        "for CAHit@3 change from K = 4 to K = 32, leave-one-seed-out cross-fitted K = 32 minus K = 4 selection gaps and Ledoit–Wolf effective "
        "ranks under the four prespecified matrix transformations. This was a robustness analysis on three selected endpoints, not a new "
        "confirmatory population sample.",
        "Normal",
    )

    figure6 = find(doc, "Figure 6.")
    heading = insert_after(figure6, "3.9 Audit directions transferred but effect magnitudes changed under structure separation", "Heading 2")
    p1 = insert_after(
        heading,
        "The formal similarity-cluster rerun completed all 1,440 outer and 4,320 inner candidate-utility evaluations without random-split "
        "fallback. Each endpoint had 15 unique outer-fold assignments, all inner and outer component sets were disjoint, test-fold sizes "
        "ranged from 376 to 505 molecules, and the largest observed cross-fold Tanimoto was 0.699 (Additional file 2: Table S26).",
        "Normal",
    )
    p2 = insert_after(
        p1,
        "CAHit@3 changes from K = 4 to K = 32 were negative under scaffold and similarity-cluster splits for BACE (-0.662 and -0.009), "
        "ClinTox (-0.956 and -0.690) and ESOL (-0.883 and -0.515). Thus the direction persisted in all three endpoints, but the BACE "
        "similarity-cluster change was close to zero and normalized MRR did not decline uniformly. The result supports transport of the "
        "prespecified CAHit@3 direction, not universal degradation of every ranking metric (Additional file 3: Figure S18A).",
        "Normal",
    )
    insert_after(
        p2,
        "Cross-fitted K = 32 minus K = 4 gaps were positive in both regimes for all three endpoints. The scaffold and similarity-cluster "
        "effects were 0.0009 (-0.0046 to 0.0048) and 0.0178 (0.0124 to 0.0243) for BACE, 0.0098 (0.0049 to 0.0138) and 0.0049 "
        "(-0.0089 to 0.0215) for ClinTox, and 0.0573 (0.0411 to 0.0749) and 0.0017 (-0.0072 to 0.0107) for ESOL. Classification "
        "ROC-AUC loss and regression RMSE loss remained on separate scales. Effective-rank estimates across endpoint, K and transformation "
        "combinations had Spearman correlation 0.832 between regimes, while absolute values still changed materially (Figure S18B-D; "
        "Tables S24-S25).",
        "Normal",
    )

    limitation = find(doc, "4.7 Limitations")
    set_text(limitation, "4.8 Limitations")
    anchor = find(doc, "ClinTox illustrates why")
    heading = insert_after(anchor, "4.7 Split-regime transfer reduced, but did not remove, design dependence", "Heading 2")
    p1 = insert_after(
        heading,
        "Preserving the CAHit@3 change direction and the cross-fitted gap direction under a split that prohibited cross-fold Tanimoto of "
        "0.70 or greater weakens the explanation that the primary finding arose only from ordinary scaffold allocation. The rank correlation "
        "of matrix-dependent diversity estimates also shows that broad registry structure was partly preserved when the split mechanism changed.",
        "Normal",
    )
    insert_after(
        p1,
        "The transfer was conditional rather than invariant. BACE showed almost no CAHit@3 change under the similarity-cluster split, "
        "normalized MRR was not uniformly lower, and interval exclusion shifted across endpoints. Split mechanism therefore changed both "
        "difficulty and the estimated magnitude of selection pressure. The appropriate claim is direction transport in three representative "
        "endpoints, not a universal law across targets, temporal splits, model registries or deployment populations.",
        "Normal",
    )
    set_text(
        find(doc, "The primary registry was intentionally"),
        "The primary registry was intentionally near-duplicate, only nine endpoints entered the main audit, and effective diversity at the "
        "largest registry was estimated from 15 outer rows. The split-regime transfer audit covered only ClinTox, BACE and ESOL, used one "
        "Morgan-derived 0.70 component threshold and did not include temporal or target-aware source-selection splits. Shrinkage, hierarchical "
        "resampling and split transfer expose sensitivity but cannot replace additional independent audit units. The study was retrospective "
        "and not prospectively preregistered, and public outer folds were not an independent lockbox."
    )
    set_text(
        find(doc, "Within the studied endpoints"),
        "Within the studied endpoints, candidate-pool expansion was accompanied by weaker chance-adjusted validation ranking and heterogeneous "
        "model-selection gaps. In three representative endpoints, the K = 32 minus K = 4 CAHit@3 change remained negative and the cross-fitted "
        "gap direction remained positive under a Tanimoto-component split, but effect magnitude and interval exclusion changed. Estimated "
        "candidate diversity remained strongly matrix dependent and only partly stable across split regimes."
    )
    set_text(
        find(doc, "Molecular model-selection studies should"),
        "Molecular model-selection studies should jointly report candidate eligibility, nominal K, matrix-dependent utility-pattern diversity, "
        "chance-adjusted ranking fidelity, endpoint-specific same-unit and cross-fitted gaps, split uniqueness and split-regime sensitivity, "
        "computational exposure, failed candidates and chemical-support boundaries. These quantities support transparent audit interpretation "
        "but do not define a universal selector or a deployment-ready screening system."
    )
    set_text(find(doc, "Additional file 1."), "Additional file 1. Supplementary Methods and Results, including the split-regime transfer audit.")
    set_text(find(doc, "Additional file 2."), "Additional file 2. Machine-readable Supplementary Tables (Tables S1-S26; XLSX).")
    set_text(find(doc, "Additional file 3."), "Additional file 3. Supplementary Figures (Figures S1-S18; PDF).")
    set_text(
        find(doc, "Public dataset provenance"),
        "Public dataset provenance is listed in Additional file 2: Table S1. Derived fold-level tables, split manifests, source hashes, "
        "split-regime transfer outputs and analysis code are supplied in the accompanying submission package."
    )
    add_table2_row(doc)
    doc.core_properties.subject = "Candidate-pool expansion with split-regime transfer audit"
    doc.save(EN)


def update_chinese() -> None:
    doc = Document(ZH)
    font = "宋体"
    set_text(find(doc, "方法："), "方法：本研究使用重复种子骨架划分，对九个公开分子性质终点开展回顾性嵌套审计，并通过矩阵依赖多样性、机会校正排序、置换与分级信号恢复对照、留一种子交叉拟合参照和等规模多视图比较评价递增候选登记表。另在ClinTox、BACE和ESOL上使用Morgan-512 Tanimoto连通分量作为不可拆分组，重新训练同一32候选登记表以开展切分机制迁移审计。", east_asia=font)
    set_text(find(doc, "结果："), "结果：候选多样性估计随矩阵构造显著变化，机会校正验证排序随候选池扩大总体减弱，置换和信号恢复对照显示指标具有预期零点与恢复行为。九终点交叉拟合差距六正三负。三终点迁移审计中，CAHit@3的K=32减K=4变化在两种切分下均为负，交叉拟合差距方向在三个终点均为正，但效应大小和区间排零具有终点依赖性；两种切分下有效秩估计的Spearman相关为0.832。", east_asia=font)
    set_text(find(doc, "结论："), "结论：名义候选数量不能单独描述模型选择的有效结构或风险。三个代表性终点的关键审计方向可迁移到更严格的结构分离切分，但效应大小并不恒定。分子基准应联合报告候选资格、矩阵依赖多样性、机会校正排序、交叉拟合差距、切分机制敏感性、计算暴露和化学支持边界。", east_asia=font)
    set_text(find(doc, "科学贡献："), "科学贡献：本研究将名义候选池规模与矩阵依赖效用模式多样性和机会校正排序退化分开，并以负对照和正对照校准排序量；进一步联合交叉拟合选择差距、穷举等规模表征比较、化学支持审计和完全嵌套的切分机制迁移审计，以区分方向迁移与效应大小不变。本文贡献是对分子模型选择实践的分析，而非新的预测器、通用选择器或外部验证研究。", east_asia=font)

    stats = find(doc, "2.13 统计推断")
    set_text(stats, "2.14 统计推断", east_asia=font)
    anchor = find(doc, "外层精确Murcko")
    heading = insert_after(anchor, "2.13 切分机制迁移审计", "Heading 2", east_asia=font)
    p1 = insert_after(heading, "选择ClinTox、BACE和ESOL分别代表罕见类别分类、常规分类和回归。保持锁定的32候选Morgan-512登记表、候选顺序、K=4/8/16/32前缀、种子11/23/37/53/71以及3外层×3内层设计不变。结构分离重跑将Tanimoto≥0.70的分子对连边，并将每个连通分量作为内外层均不可拆分的组。阈值与分配规则在正式模型拟合前经切分可行性检查后固定，模型性能不参与分组。", "Normal", east_asia=font)
    insert_after(p1, "种子化贪心分配同时平衡样本量以及分类任务的两类样本量。所有正式切分均要求组间无重叠、跨折最大Tanimoto低于0.70且分类折保留两类，禁止随机回退。迁移审计新增5,760次候选拟合，并比较CAHit@3的K=32减K=4变化、留一种子交叉拟合差距和四种矩阵变换下的Ledoit–Wolf有效秩。", "Normal", east_asia=font)

    fig6 = find(doc, "图6 ")
    heading = insert_after(fig6, "3.9 结构分离下审计方向迁移但效应大小改变", "Heading 2", east_asia=font)
    p1 = insert_after(heading, "结构相似度聚类重跑完成1,440个外层和4,320个内层候选效用评价，无随机切分回退。三个终点各有15个不同外层折分配，内外层相似度分量完全分离，测试折含376–505个分子，最大跨折Tanimoto为0.699（补充表S26）。", "Normal", east_asia=font)
    p2 = insert_after(p1, "BACE、ClinTox和ESOL在骨架切分与相似度聚类切分下的CAHit@3变化分别为-0.662/-0.009、-0.956/-0.690和-0.883/-0.515。三个终点方向均保持为负，但BACE在严格切分下接近零，且标准化MRR并非一致下降，因此结论限定于预设CAHit@3方向（补充图S18A）。", "Normal", east_asia=font)
    insert_after(p2, "三个终点的交叉拟合K=32减K=4差距在两种切分下均为正，但区间排零发生变化：BACE为0.0009与0.0178，ClinTox为0.0098与0.0049，ESOL为0.0573与0.0017。分类ROC-AUC损失与回归RMSE损失保持不同量纲。跨终点、K和矩阵变换组合的有效秩在两种切分间Spearman相关为0.832，但绝对值仍有明显变化（补充图S18B–D；表S24–S25）。", "Normal", east_asia=font)

    limitation = find(doc, "4.7 局限性")
    set_text(limitation, "4.8 局限性", east_asia=font)
    anchor = find(doc, "ClinTox说明")
    heading = insert_after(anchor, "4.7 切分迁移降低但未消除设计依赖", "Heading 2", east_asia=font)
    p1 = insert_after(heading, "在禁止跨折Tanimoto≥0.70的切分下，CAHit@3变化和交叉拟合差距方向仍一致，削弱了主结果仅由普通骨架分配造成的解释；有效秩的较高秩相关也说明登记表的总体结构部分保留。", "Normal", east_asia=font)
    insert_after(p1, "迁移并非数值不变。BACE的CAHit@3变化在相似度聚类切分下接近零，标准化MRR并非一致降低，区间排零也随终点改变。因此应表述为三个代表终点的方向迁移，而非跨目标、时间切分、模型登记表或部署人群的普遍定律。", "Normal", east_asia=font)
    set_text(find(doc, "主要登记表有意包含"), "主要登记表有意包含近重复候选，主审计仅含九个终点，最大登记表的有效多样性由15个外层行估计。切分迁移只覆盖ClinTox、BACE和ESOL，使用一个Morgan派生的0.70连通分量阈值，未覆盖时间切分或目标感知源选择。收缩、层级重采样和切分迁移不能替代更多独立审计单元；本研究为回顾性分析，公开外层折也不是独立锁箱。", east_asia=font)
    set_text(find(doc, "在9个分子性质审计中"), "在9个分子性质审计中，名义候选数不能唯一描述有效搜索多样性，矩阵构造会显著改变有效秩估计。三个代表终点在Tanimoto连通分量切分下仍保持CAHit@3变化为负和交叉拟合差距为正的方向，但效应大小与区间排零改变。该结果增强了方向稳健性，同时继续限定外推范围。", east_asia=font)
    set_text(find(doc, "分子模型选择研究应同时报告"), "分子模型选择研究应同时报告候选资格、名义K、矩阵依赖效用模式多样性、机会校正排序、同一单元和交叉拟合差距、划分唯一性与切分机制敏感性、计算暴露、失败候选和化学支持边界。该审计框架提高证据透明度，但不构成通用选择器或部署级筛选系统。", east_asia=font)
    set_text(find(doc, "公开数据来源见"), "公开数据来源见补充表S1；派生逐折表、切分清单、源文件哈希、切分迁移结果和分析代码随提交包提供。", east_asia=font)
    set_text(find(doc, "Additional file 1："), "Additional file 1：补充方法与结果（含切分机制迁移审计）；Additional file 2：机器可读补充表S1–S26；Additional file 3：补充图S1–S18。", east_asia=font)
    doc.save(ZH)


def style_sheet(ws) -> None:
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    fill = PatternFill("solid", fgColor="D9E7F2")
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for column in range(1, ws.max_column + 1):
        values = [str(ws.cell(row, column).value or "") for row in range(1, min(ws.max_row, 500) + 1)]
        ws.column_dimensions[get_column_letter(column)].width = min(max(max(map(len, values), default=8) + 2, 10), 42)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def add_frame(workbook, name: str, frame: pd.DataFrame) -> None:
    if name in workbook.sheetnames:
        del workbook[name]
    ws = workbook.create_sheet(name)
    ws.append(list(frame.columns))
    for row in frame.itertuples(index=False, name=None):
        ws.append([None if pd.isna(value) else value for value in row])
    style_sheet(ws)


def update_workbook() -> Path:
    old = SUPP / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S22.xlsx"
    new = SUPP / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S26.xlsx"
    shutil.copy2(old, new)
    workbook = load_workbook(new)
    ranking_units = pd.read_csv(ANALYSIS / "split_regime_ranking_units.csv")
    ranking_summary = pd.read_csv(ANALYSIS / "split_regime_ranking_endpoint_summary.csv")
    ranking_units.insert(0, "record_level", "outer_unit")
    ranking_summary.insert(0, "record_level", "endpoint_summary")
    add_frame(workbook, "S23 Split-regime ranking", pd.concat([ranking_summary, ranking_units], ignore_index=True, sort=False))
    cross_units = pd.read_csv(ANALYSIS / "split_regime_cross_fitted_units.csv")
    cross_summary = pd.read_csv(ANALYSIS / "split_regime_cross_fitted_effects.csv")
    cross_units.insert(0, "record_level", "outer_unit")
    cross_summary.insert(0, "record_level", "endpoint_effect")
    add_frame(workbook, "S24 Split-regime cross-fit", pd.concat([cross_summary, cross_units], ignore_index=True, sort=False))
    add_frame(workbook, "S25 Split-regime diversity", pd.read_csv(ANALYSIS / "split_regime_effective_diversity.csv"))
    split_frames = []
    cluster = ROOT / "results" / "split_regime_transfer_20260716" / "similarity_cluster"
    for seed in (11, 23, 37, 53, 71):
        frame = pd.read_csv(cluster / f"seed_{seed}" / "split_manifest.csv")
        frame.insert(0, "record_level", "inner_outer_split")
        split_frames.append(frame)
    summary = pd.read_csv(ANALYSIS / "similarity_split_integrity_summary.csv")
    summary.insert(0, "record_level", "endpoint_summary")
    add_frame(workbook, "S26 Similarity split audit", pd.concat([summary, *split_frames], ignore_index=True, sort=False))
    workbook.save(new)
    return new


def update_supplementary_methods() -> None:
    path = SUPP / "Additional_file_1_Supplementary_Methods_and_Results.docx"
    doc = Document(path)
    source = find(doc, "S14 Source limitations")
    set_text(source, "S15 Source limitations")
    anchor = find(doc, "The limited prediction panel")
    heading = insert_after(anchor, "S14 Split-regime transfer audit", "Heading 2")
    insert_after(
        heading,
        "The transfer audit retained the locked candidate registry, K prefixes, five split seeds and 3 x 3 nested design for ClinTox, "
        "BACE and ESOL. New structure-separated folds used intact connected components of a Morgan-512 similarity graph with edges at "
        "Tanimoto >= 0.70. All 5,760 new fits completed without fallback. CAHit@3 K32 - K4 changes and cross-fitted gap directions agreed "
        "between scaffold and similarity-cluster regimes in all three endpoints, but magnitude and interval exclusion changed. Effective-rank "
        "estimates across endpoint, K and transformation combinations had Spearman correlation 0.832. Tables S23-S26 and Figure S18 retain "
        "the full unit-level results, split manifests and source data.",
        "Normal",
    )
    table_dir = find(doc, "Table S22.")
    p = insert_after(table_dir, "Table S23. Split-regime chance-adjusted ranking units and endpoint summaries.", "Normal")
    p = insert_after(p, "Table S24. Split-regime cross-fitted units and endpoint effects.", "Normal")
    p = insert_after(p, "Table S25. Split-regime matrix-dependent effective diversity.", "Normal")
    insert_after(p, "Table S26. Similarity-component split integrity and complete split manifest.", "Normal")
    insert_after(find(doc, "Figure S17."), "Figure S18. Split-regime transfer audit: ranking, cross-fitted gaps, split integrity and effective-rank concordance.", "Normal")
    doc.save(path)


def merge_supplementary_figures() -> Path:
    old = SUPP / "Additional_file_3_Supplementary_Figures_S1-S17.pdf"
    s18 = SUPP / "Supplementary_Figure_S18_split_regime_transfer.pdf"
    new = SUPP / "Additional_file_3_Supplementary_Figures_S1-S18.pdf"
    writer = PdfWriter()
    for path in (old, s18):
        for page in PdfReader(path).pages:
            writer.add_page(page)
    with new.open("wb") as handle:
        writer.write(handle)
    return new


def write_text_exports() -> None:
    (BASE / "Revised_Abstract.txt").write_text(
        "\n".join([EN_ABSTRACT_METHODS, EN_ABSTRACT_RESULTS, EN_ABSTRACT_CONCLUSIONS]), encoding="utf-8"
    )
    (BASE / "New_Abstract.txt").write_text(
        "\n".join([EN_ABSTRACT_METHODS, EN_ABSTRACT_RESULTS, EN_ABSTRACT_CONCLUSIONS]), encoding="utf-8"
    )
    (BASE / "Scientific_Contribution.txt").write_text(EN_CONTRIBUTION, encoding="utf-8")
    (BASE / "New_Scientific_Contribution.txt").write_text(EN_CONTRIBUTION, encoding="utf-8")


def main() -> None:
    update_english()
    update_chinese()
    workbook = update_workbook()
    update_supplementary_methods()
    figures = merge_supplementary_figures()
    write_text_exports()
    print(EN)
    print(ZH)
    print(workbook)
    print(figures)


if __name__ == "__main__":
    main()
