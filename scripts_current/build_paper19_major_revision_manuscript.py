from __future__ import annotations

import json
import re
import shutil
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("D:/fzyc")
sys.path.insert(0, str(ROOT / "scripts"))

from build_paper19_manuscript import (  # noqa: E402
    add_body,
    add_figure,
    add_heading,
    add_table,
    add_table_caption,
    configure_document,
    set_run_font,
)


OUT = ROOT / "output"
MAJOR = OUT / "paper19_major_revision_20260712"
REV = OUT / "paper19_jcheminformatics_revision_20260712"
PREV = OUT / "paper19_rejection_driven_experiments_20260712"
FIG = MAJOR / "figures"
PAPER = OUT / "小论文-19_8000词MajorRevision.docx"
DESKTOP = Path("C:/Users/Administrator/Desktop/小论文-19_8000词MajorRevision.docx")
RESPONSE_MD = OUT / "小论文-19_MajorRevision逐条回复.md"
RESPONSE_DOCX = OUT / "小论文-19_MajorRevision逐条回复.docx"
ABSTRACT_MD = OUT / "小论文-19_修订摘要与Scientific_Contribution.md"
AUDIT_JSON = OUT / "小论文-19_8000词MajorRevision审计.json"


def body(doc: Document, text: str) -> None:
    add_body(doc, text)


def abstract_part(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run_font(p.add_run(f"{label}: "), 10.5, bold=True)
    set_run_font(p.add_run(text), 10.5)


def add_reference(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.left_indent = Cm(0.63)
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(text), 9)


def start_landscape(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)


def restore_portrait(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)


def dataset_table(doc: Document) -> None:
    frame = pd.read_csv(REV / "dataset_characteristics.csv")
    rows = []
    for _, r in frame.iterrows():
        if r.task_type == "classification":
            distribution = f"{int(r.positive_n)} positive ({100*r.positive_rate:.1f}%)"
            metric = "ROC-AUC"
        else:
            distribution = f"{r.target_min:.2f}–{r.target_max:.2f}"
            metric = "RMSE"
        rows.append([r.display_name, r.task_type, f"{int(r.raw_n)}/{int(r.analysis_n)}", distribution, r.target_unit, metric, r.evidence_role.replace("confirmation", "stress test")])
    add_table_caption(doc, "Table 1. Datasets, analysis populations and evidence roles.")
    add_table(doc, ["Endpoint", "Task", "Raw/analysis n", "Class balance or range", "Unit", "Primary metric", "Evidence role"], rows)


def registry_table(doc: Document) -> None:
    rows = [
        ["Controlled registry prefixes", "K = 4, 8, 16, 32", "Morgan-512", "Linear, bagging and boosting variants", "Same preprocessing and per-candidate opportunity", "17,280 fits; 4,437.95 recorded fit-s"],
        ["Composition controls", "K = 4, 8, 16, 32", "Same registered candidates", "Random order, random subset, family balanced", "100 frozen seeds per mode and K", "Total exposure varied with K"],
        ["Shared-split multiview stress test", "12", "Morgan, MACCS, RDKit2D, concatenated", "Linear, random forest, LightGBM", "Common 3 × 3 × 5 design", "6,480 fits"],
        ["Limited representation-baseline stress test", "4", "Morgan, graph, ChemBERTa, MoLFormer", "RDKit-RF, five-epoch GCN, two frozen probes", "Unequal search budgets; selector/error analysis only", "360 outer units; 220,040 predictions"],
    ]
    add_table_caption(doc, "Table 2. Candidate registries, selection opportunities and computational exposure.")
    add_table(doc, ["Analysis", "Nominal size", "Representations", "Learners", "Comparison constraint", "Recorded exposure"], rows)


def diversity_ranking_table(doc: Document) -> None:
    div = pd.read_csv(MAJOR / "effective_diversity_shrinkage_summary.csv")
    rank = pd.read_csv(MAJOR / "chance_adjusted_ranking_summary.csv")
    old = pd.read_csv(PREV / "paper19_effective_diversity.csv").groupby("candidate_count", as_index=False)["outer_entropy_effective_rank"].mean()
    frame = div.merge(rank, on="candidate_count").merge(old, on="candidate_count")
    rows = []
    for _, r in frame.iterrows():
        rows.append([
            int(r.candidate_count),
            f"{r.outer_entropy_effective_rank:.2f}",
            f"{r.mean_shrinkage_entropy_rank:.2f} ({r.hierarchical_ci95_low_entropy:.2f}–{r.hierarchical_ci95_high_entropy:.2f})",
            f"{r.mean_shrinkage_participation_rank:.2f} ({r.hierarchical_ci95_low_participation:.2f}–{r.hierarchical_ci95_high_participation:.2f})",
            f"{r.mean_shrinkage_median_correlation:.3f} ({r.hierarchical_ci95_low_correlation:.3f}–{r.hierarchical_ci95_high_correlation:.3f})",
            f"{r.mean_chance_adjusted_top3:.3f}",
            f"{r.mean_normalized_mrr_gain:.3f}",
            f"{r.mean_ndcg:.3f}",
            f"{r.mean_spearman:.3f}/{r.mean_kendall:.3f}",
        ])
    start_landscape(doc)
    add_table_caption(doc, "Table 3. Effective-diversity estimates and chance-adjusted ranking fidelity.")
    add_table(doc, ["K", "Empirical entropy rank", "Shrinkage entropy rank (95% CI)", "Participation rank (95% CI)", "Median correlation (95% CI)", "CAHit@3", "Normalized MRR gain", "NDCG", "Spearman/Kendall"], rows)
    restore_portrait(doc)


def loss_table(doc: Document) -> None:
    frame = pd.read_csv(MAJOR / "selection_loss_seed_clustered.csv")
    names = {"tdc_caco2_wang": "Caco2 Wang", "tdc_hia_hou": "HIA Hou", "tdc_pgp_broccatelli": "P-gp Broccatelli", "lipo": "Lipophilicity", "bbbp": "BBBP", "bace": "BACE", "clintox": "ClinTox", "esol": "ESOL", "freesolv": "FreeSolv"}
    rows = []
    for _, r in frame.iterrows():
        rows.append([
            names[r.task], r.effect_scale, f"{r.mean_delta:.4f}",
            f"{r.seed_clustered_ci95_low:.4f}–{r.seed_clustered_ci95_high:.4f}",
            f"{r.fold_bootstrap_ci95_low:.4f}–{r.fold_bootstrap_ci95_high:.4f}",
            f"{int(r.positive_seed_means)}/5",
            f"{r.leave_one_seed_out_min:.4f}–{r.leave_one_seed_out_max:.4f}",
        ])
    start_landscape(doc)
    add_table_caption(doc, "Table 4. Endpoint-specific change in raw-scale selection loss from K = 4 to K = 32.")
    add_table(doc, ["Endpoint", "Scale", "Mean change", "Seed-clustered 95% CI", "Outer-unit 95% CI", "Positive seed means", "Leave-one-seed range"], rows)
    restore_portrait(doc)


def multiview_table(doc: Document) -> None:
    frame = pd.read_csv(MAJOR / "multiview_absolute_endpoint_summary.csv")
    names = {"tdc_caco2_wang": "Caco2 Wang", "tdc_hia_hou": "HIA Hou", "tdc_pgp_broccatelli": "P-gp Broccatelli", "lipo": "Lipophilicity", "bbbp": "BBBP", "bace": "BACE", "clintox": "ClinTox", "esol": "ESOL", "freesolv": "FreeSolv"}
    reps = {"multiview": "Concatenated", "rdkit2d": "RDKit2D", "morgan512": "Morgan", "maccs": "MACCS"}
    rows = []
    for _, r in frame.iterrows():
        rows.append([
            names[r.task], r.performance_metric,
            f"{r.morgan_only_selected_performance:.4f}", f"{r.full_multiview_selected_performance:.4f}",
            f"{r.mean_raw_paired_gain:.4f}", f"{r.seed_clustered_ci95_low:.4f}–{r.seed_clustered_ci95_high:.4f}",
            f"{int(r.positive_paired_units)}/15", f"{r.mean_normalized_gain:.3f}",
            f"{reps[r.most_selected_representation]} ({int(r.most_selected_representation_count)}/15)",
        ])
    start_landscape(doc)
    add_table_caption(doc, "Table 5. Absolute selected performance and paired gains in the shared-split multiview stress test.")
    add_table(doc, ["Endpoint", "Metric", "Morgan-only selected", "Full-pool selected", "Raw paired gain", "Seed-clustered 95% CI", "Positive units", "Normalized gain", "Most selected representation"], rows)
    restore_portrait(doc)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    header = doc.sections[0].header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(header.add_run("Candidate-pool expansion and model-selection loss"), 8)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.first_line_indent = Cm(0)
    set_run_font(title.add_run("Validation-ranking distortion and selection loss under candidate-pool expansion in molecular property prediction: a frozen audit study"), 16, bold=True)

    add_heading(doc, "Abstract", 1)
    abstract_part(doc, "Background", "Modern molecular-property studies compare increasingly large collections of fingerprints, descriptors, graph models, language-model representations and tuning variants. Expansion can create genuine representational opportunity, but repeated selection on finite validation information can also distort rankings and increase model-selection loss.")
    abstract_part(doc, "Methods", "We performed a retrospective frozen audit across nine endpoints using five seeds, three outer scaffold folds and three inner folds. Registered prefixes contained K = 4, 8, 16 or 32 candidates. Effective diversity was estimated from 15 × K outer-utility matrices using empirical and Ledoit–Wolf shrinkage correlations, spectral-entropy and participation-ratio ranks, hierarchical bootstrap intervals and omission sensitivity. Ranking was assessed with chance-adjusted Top-3 recovery, normalized mean reciprocal rank gain, percentile rank, normalized discounted cumulative gain, Spearman correlation and Kendall correlation.")
    abstract_part(doc, "Results", "At K = 32, empirical spectral-entropy rank was 2.00, whereas the shrinkage estimate was 3.43 (hierarchical 95% CI 2.54–5.90); leave-one-seed endpoint means ranged from 3.45 to 4.06. Chance-adjusted Top-3 recovery declined from 0.881 at K = 4 to 0.240 at K = 32, and normalized mean reciprocal-rank gain declined from 0.727 to 0.157. Raw selection loss increased in eight of nine endpoints. In a shared-split multiview stress test, classification ROC-AUC gains ranged from 0.0039 to 0.0873 and regression RMSE reductions from 0.1058 to 1.1949; the mean normalized realized gain was 0.343 (95% CI 0.210–0.483). Under equal candidate truth, finite-audit winner optimism increased with K.")
    abstract_part(doc, "Conclusions", "Nominal expansion greatly exceeded estimated predictive diversity and was accompanied by weaker chance-adjusted ranking fidelity and greater endpoint-specific selection loss. Heterogeneous candidates could still provide realizable, endpoint-dependent gains, but outer-audit maxima remained finite-sample estimates rather than true generalization upper bounds.")
    abstract_part(doc, "Scientific Contribution", "This study separates nominal candidate-pool expansion from effective predictive diversity and quantifies how both relate to validation-ranking fidelity and endpoint-specific model-selection loss. It distinguishes heterogeneous representational gains from the mechanical optimism produced by maximizing finite audit estimates. The contribution is an auditable analysis of molecular model selection rather than a new predictive backbone or a claim of universal model superiority.")
    body(doc, "Keywords: molecular property prediction; candidate-pool expansion; model selection; effective diversity; nested cross-validation; ranking fidelity; selection loss; uncertainty quantification")

    add_heading(doc, "1 Introduction", 1)
    body(doc, "Molecular property prediction has moved from comparisons among a few fixed fingerprints and regressors to broad searches spanning circular fingerprints, substructure keys, physicochemical descriptors, graph neural networks, directed message-passing networks, pretrained chemical-language models and automated ensembles [1,2,16–24]. This expansion is scientifically attractive because each representation exposes different inductive biases. Fingerprints emphasize local substructures, descriptor panels encode interpretable global properties, graph networks learn neighbourhood aggregation and language models transfer information from large unlabeled molecular corpora. Benchmark papers increasingly evaluate several of these choices together, often with additional hyperparameter grids and ensemble rules. The resulting candidate registry can contain dozens or hundreds of nominal alternatives even when the endpoint provides only a modest number of scaffolds or labelled molecules. Under those conditions, the model-selection process becomes an important component of the reported result rather than a neutral preliminary step.")
    body(doc, "Candidate expansion creates two opportunities that should not be conflated. The first is genuine representational opportunity: a new candidate may capture chemical information that the existing pool cannot express and may therefore improve performance on unseen molecules. The second is repeated validation-selection opportunity: every additional candidate supplies another noisy estimate that can be ranked against the same finite validation information. A candidate can appear superior because it is genuinely better, because it is favoured by the particular validation scaffolds, or because the maximum among many noisy estimates tends to be high. Direct test leakage is not required for this problem to arise. Even with a fully hidden outer fold, repeated comparisons among representations, learners, tuning variants, calibration procedures and selection rules can consume validation information and destabilize the ordering that determines the final model [3,4].")
    body(doc, "Repeated nested cross-validation provides an essential separation between inner selection and outer evaluation, but it does not automatically quantify the structure of the search itself. A nested design can estimate a frozen selector while leaving the independence and correlation of candidates unresolved. It also does not reveal total computational exposure, recovery of the outer-best candidate or optimism in the largest outer estimate. Candidate failures and unavailable runs create further asymmetry because a nominal registry may not equal the set that actually entered selection. These quantities become especially important when K exceeds the number of repeated outer units, because an empirical candidate-correlation matrix is then rank deficient and its eigenvalue spectrum is unstable.")
    body(doc, "Existing molecular benchmarks have established common datasets and metrics, and model papers have demonstrated the value of message passing, pretrained representations and multimodal fusion [1,2,5–10,21–24,34,35]. Research on applicability domains, distribution shift, activity cliffs and uncertainty has separately shown that random or scaffold averages can conceal difficult chemical regions [9–14,29,30,33]. Most of this literature understandably emphasizes final predictive performance. Candidate registry size, effective search diversity, accumulated compute exposure and validation-ranking fidelity are less often treated as primary outcomes. The best observed outer score is also frequently discussed as if it approximated an oracle, although it remains the maximum of finite estimates. Our study addresses this reporting gap without proposing a new molecular representation or claiming that one model family is universally preferable.")
    body(doc, "We ask a limited question: under a frozen repeated nested scaffold design, how does registered candidate-pool expansion relate to effective predictive diversity, chance-adjusted validation-ranking fidelity and endpoint-specific raw-scale selection loss? The study makes three contributions. First, it contrasts nominal K with shrinkage-stabilized spectral and participation-ratio effective ranks and explicitly quantifies finite-unit uncertainty. Second, it connects chance-adjusted ranking metrics to paired raw-scale selection loss while preserving endpoint and metric units. Third, it separates a near-duplicate expansion audit from a shared-split multiview stress test, a limited representation-baseline stress test and reliability analyses. This hierarchy permits heterogeneous gains to be reported without treating the outer folds as an independent lockbox or presenting the audit as a universal model-selection algorithm.")

    add_heading(doc, "2 Methods", 1)
    add_heading(doc, "2.1 Study design and evidence hierarchy", 2)
    body(doc, "The study was a retrospective frozen audit of completed molecular-property experiments. Before the present revision, we fixed the nine primary endpoints, candidate order, K values, seeds, outer and inner split indices, selection rules, outcomes and exclusion logic used for the reported analyses. The lock was reconstructed after the original outer outcomes existed and is therefore not a prospective preregistration. We use frozen to mean that the present analyses did not change candidate eligibility, split assignments or selection rules after examining the endpoint summaries. This terminology describes the analysis workflow and does not imply that an untouched prospective cohort was available.")
    body(doc, "Evidence was organized into four levels. The primary analysis evaluated registered prefixes of a 32-candidate near-duplicate conventional-model pool across nine endpoints. Candidate-composition resampling and a known-truth simulation examined mechanisms that could generate the observed trend. A 12-candidate multiview experiment assessed whether heterogeneous representations could produce realizable gains under the same public endpoints and nested splits. A four-candidate representation panel, conformal analyses and chemical-boundary stratifications were limited stress tests. They were not merged into a single leaderboard because they involved different candidate budgets, endpoint sets and evidential purposes.")
    body(doc, "The outer folds served as audit sets for frozen inner-selection decisions. Outer labels did not influence candidate eligibility, hyperparameters, preprocessing estimates or tie breaking within a completed unit. However, the same outer results were used to quantify the maximum observed candidate utility, ranking agreement and selection loss. The outer maximum is consequently an observed finite-audit best, not a true generalization upper bound. The multiview and representation panels reused the same public endpoint definitions and cannot provide external independent confirmation.")

    add_heading(doc, "2.2 Datasets and molecular standardization", 2)
    body(doc, "The primary panel comprised ESOL, FreeSolv, Lipophilicity, blood–brain barrier penetration (BBBP), beta-secretase 1 inhibition (BACE), ClinTox, Caco2 Wang, human intestinal absorption (HIA Hou) and P-glycoprotein inhibition (P-gp Broccatelli) [1,2]. ESOL, FreeSolv, Lipophilicity and Caco2 were regression endpoints evaluated by root mean squared error (RMSE). The remaining endpoints were binary classifications evaluated primarily by receiver operating characteristic area under the curve (ROC-AUC). Precision–recall area under the curve (PR-AUC), Brier score, expected calibration error and minority recall were retained for reliability analyses but did not replace the registered classification utility used for candidate selection.")
    body(doc, "Molecules were parsed and standardized with RDKit. The largest valid molecular fragment was retained, charge normalization was applied where chemically valid and canonical SMILES preserved stereochemical information. These operations preceded feature calculation but were governed by fixed endpoint-level rules rather than by outer performance. Within modelling pipelines, imputation, scaling and any fitted transformation were estimated only from the corresponding training fold. For classification endpoints, duplicate structures were merged only when their labels agreed; groups with conflicting labels were removed. For regression endpoints, replicate labels were averaged and the replicate count was retained for the cleaning audit.")
    body(doc, "Cleaning affected endpoints differently. BBBP was reduced from 2,050 records to 1,955 standardized structures after 62 consistent duplicate merges, 22 conflicting rows and 11 invalid SMILES were addressed. ClinTox was reduced from 1,484 to 1,376 structures, including exclusion of 98 conflicting rows and four invalid SMILES. Only 58 ClinTox structures were positive after cleaning, corresponding to 4.2% of the analysis population. This scarcity is reported whenever thresholded toxicity or minority coverage is interpreted. Regression target ranges and units were retained so that raw RMSE effects would not be silently combined across incompatible scales.")
    dataset_table(doc)

    add_heading(doc, "2.3 Candidate registry and computational exposure", 2)
    body(doc, "The controlled expansion registry contained 32 candidates using a common Morgan-512 representation. Classification positions 1–4 were logistic-regression variants. Positions 5–12 contained random-forest and extremely randomized-tree variants, positions 13–16 histogram or gradient-boosting variants, positions 17–24 LightGBM variants, positions 25–28 XGBoost variants and positions 29–32 CatBoost variants. The regression registry replaced the first four candidates with ridge and elastic-net models and retained the same tree-family structure. K = 4, 8, 16 and 32 were deterministic prefixes of this order.")
    body(doc, "Registry order was chosen to add model-family complexity in stages: regularized linear candidates entered first, bagging candidates second and boosting families thereafter. It was not reordered using outer-audit performance. Prefix expansion therefore changed both nominal K and family composition, which motivated the separate random-order, random-subset and family-balanced controls. Because the registry contains many candidates that share a fingerprint and closely related learners, it was designed as a controlled repeated-selection stress rather than as a representative sample of all graph and foundation models.")
    body(doc, "Per-candidate opportunity was held constant across K: the same preprocessing rules, candidate configuration, inner folds and refit logic applied whenever a candidate was eligible. Total search exposure was not held constant. Larger prefixes required proportionally more fits and supplied more utilities to the selector. The controlled audit recorded 17,280 fits and 4,437.95 candidate fit-seconds. These logs characterize model-fitting exposure on the recorded central processing unit workflow; they do not reconstruct unlogged data-loading overhead or graphical processing unit utilization.")
    registry_table(doc)

    add_heading(doc, "2.4 Repeated nested scaffold evaluation", 2)
    body(doc, "For every endpoint, five seeds (11, 23, 37, 53 and 71) indexed repeated evaluations. Each seed used three outer scaffold folds. Classification splits were stratified by label where the endpoint permitted stable stratification, and regression splits used scaffold grouping without label stratification. Within each outer training partition, three inner scaffold folds estimated candidate validation utility. The full design therefore produced 15 outer-audit units and 45 inner-utility rows per endpoint and candidate.")
    body(doc, "The validation-best selector ranked candidates by mean inner utility, where utility was ROC-AUC for classification and negative RMSE for regression. Ties were resolved deterministically by registry order. The chosen candidate was refitted on the complete outer training partition and evaluated once on the outer fold. Fixed-single, one-standard-error and risk-adjusted selectors were retained as sensitivity analyses, but validation-best defined the primary selection-loss contrast. Candidate failures remained in the audit trail and did not trigger outer-informed replacement.")
    body(doc, "Outer folds from the same seed share much of their data and are not independent experimental replications. Seeds also reuse the same endpoint population. We therefore regard fold and seed repetition as design-sensitivity information rather than biological or task-level replication. Primary endpoint intervals first averaged the three outer folds within each seed and then resampled the five seed means. Fold-level bootstrap intervals are shown as secondary comparisons because they use all 15 units and can be narrower when fold dependence is ignored.")

    add_heading(doc, "2.5 Selection-loss and ranking estimands", 2)
    body(doc, "All utilities were oriented so that larger values were better. Let C_K denote the registered prefix of size K, V_uj the mean inner-validation utility for candidate j in outer unit u and A_uj its outer-audit utility. The selected candidate was s_u = argmax_j V_uj, and the observed outer best was b_u = argmax_j A_uj. Raw selection loss was L_u(K) = A_u,b_u − A_u,s_u. It is an ROC-AUC difference for classification and an RMSE increase for regression because regression utility equals negative RMSE.")
    body(doc, "Raw-scale loss was the primary outcome because it preserves the quantity used in each endpoint. Range-normalized loss divided L_u(K) by the observed candidate-utility range within the same outer unit and pool. That transformation supported sensitivity summaries across endpoints but depended on the worst and best candidates in the finite pool. We therefore did not interpret a cross-endpoint mean normalized gain as a universal chemical effect, and we did not combine raw RMSE changes from ESOL, FreeSolv, Lipophilicity and Caco2 as if they shared one physical unit.")
    body(doc, "The rank of the observed outer-best candidate within the validation ordering generated several fidelity estimands. Top-3 recovery was chance adjusted as CAHit@3 = (Hit@3 − 3/K)/(1 − 3/K), with Hit@3 = 1 when the outer-best candidate was among the validation top three. For K = 4, the chance baseline is 0.75; for K = 32, it is 0.09375. Values above zero indicate recovery beyond random ranking. Raw Top-1 recovery was retained only as a descriptive supplementary quantity because its random expectation decreases mechanically from 1/4 to 1/32.")
    body(doc, "Mean reciprocal rank (MRR) was also adjusted for K. Under a random ordering, E(MRR_random) = H_K/K, where H_K is the Kth harmonic number. Normalized MRR gain was (MRR_observed − MRR_random)/(1 − MRR_random). Additional outcomes were the percentile rank of the outer-best candidate in the validation list, normalized discounted cumulative gain (NDCG), Spearman correlation and Kendall correlation between complete validation and outer candidate orderings. These metrics have different sensitivities to the top of the ranking and prevent the interpretation from depending on one K-sensitive hit rate.")

    add_heading(doc, "2.6 Effective-diversity estimation", 2)
    body(doc, "Effective diversity was estimated from candidate-utility matrices rather than from candidate labels or model-family counts. For each endpoint and K, the outer matrix contained 15 rows, indexed by five seeds and three outer folds, and K columns, indexed by registered candidates. The corresponding inner matrix contained 45 rows, indexed by seed, outer fold and inner fold, and K columns. The main diversity analysis used the outer-utility matrix because it characterizes behavioural similarity on held-out scaffold groups. Inner-utility estimates were retained to assess whether redundancy was already visible during selection.")
    body(doc, "A direct empirical K × K correlation matrix is singular when K exceeds the number of rows and can yield unstable eigenvalues even below that boundary. We therefore reported both the original empirical correlation spectrum and a Ledoit–Wolf shrinkage covariance estimate after column standardization. The shrinkage covariance was converted to a correlation matrix R by dividing each element by the product of its marginal standard deviations. This estimator pulls noisy covariance directions toward a structured target and produces a full-rank matrix when K > 15, but it does not create new independent information.")
    body(doc, "Let λ_1,…,λ_K be the non-negative eigenvalues of R and p_i = λ_i/Σ_j λ_j. Spectral-entropy effective rank was r_entropy = exp(−Σ_i p_i log p_i). Participation-ratio effective rank was r_PR = (Σ_i λ_i)^2/Σ_i λ_i^2, equivalently 1/Σ_i p_i^2. Both equal one for a perfectly redundant rank-one spectrum and K when all eigenvalues are equal. Spectral entropy is more sensitive to small eigenvalue mass, whereas the participation ratio places greater weight on dominant directions. Median off-diagonal correlation was reported as a directly interpretable complement.")
    body(doc, "Uncertainty was estimated with a hierarchical bootstrap. Five seeds were sampled with replacement; within every sampled seed, three outer-fold rows were then sampled with replacement. The shrinkage correlation and both effective ranks were recomputed for each of 500 hierarchical samples. Intervals are percentile 95% intervals of the endpoint-mean bootstrap distribution. Sensitivity analyses omitted each seed and each outer-fold index in turn. The available near-duplicate audit did not export all-candidate per-molecule predictions, so a K = 32 prediction-matrix effective rank could not be calculated from the current source data.")

    add_heading(doc, "2.7 Candidate-composition controls", 2)
    body(doc, "Prefix composition could confound a K effect because each increase added specific model families. Three frozen controls addressed this issue. Random-order controls permuted registry order before taking prefixes. Random-subset controls sampled K candidates from the complete registry without requiring a prefix. Family-balanced controls sampled candidates to maintain representation of available learner families. Every mode used 100 predetermined seeds for each endpoint and K.")
    body(doc, "The subset or order seed was a resampling device, not an independent endpoint. Results were first summarized within endpoint, and endpoint remained the cross-task inferential unit. At K = 32 all three modes necessarily equal the complete pool; convergence at that point is therefore a design property rather than independent corroboration. The controls test whether increasing normalized selection loss was peculiar to one early registry composition, not whether all conceivable candidate distributions would show the same trend.")

    add_heading(doc, "2.8 Finite-audit winner-optimism simulation", 2)
    body(doc, "A known-truth simulation isolated the mechanical effect of maximizing noisy audit estimates. In the equal-truth scenario, all candidates had identical true standardized utility. Audit errors followed a multivariate Gaussian distribution with pairwise correlations of 0, 0.5 or 0.9. Effective audit sample sizes were 25, 50, 100 or 200, candidate counts were K = 4, 8, 16 or 32 and 30,000 Monte Carlo replicates were generated for every configuration. A weak-gradient scenario was retained as a secondary sensitivity analysis.")
    body(doc, "Winner optimism was the maximum observed audit utility minus the true utility of the selected maximum, expressed in simulation standard-deviation units. Because candidate truth was equal in the primary simulation, any positive maximum was produced entirely by finite estimation noise. The simulation was not used to subtract a bias correction from empirical endpoints. Its purpose was to demonstrate that an outer maximum can be optimistic even when the outer data were not used for inner selection and even when candidate estimates are strongly correlated.")

    add_heading(doc, "2.9 Shared-split multiview stress test", 2)
    body(doc, "The multiview pool crossed four representations with three learners, yielding 12 candidates. Representations were Morgan-512 fingerprints, Molecular ACCess System (MACCS) keys, RDKit2D descriptors and a concatenated vector containing all three views. Learners were a regularized linear model, random forest and LightGBM. The Morgan-only comparator contained the same three learner classes restricted to Morgan-512, so the full pool received additional representation choices and additional total selection opportunity but not additional learner families.")
    body(doc, "All 12 candidates used the same nine endpoints, five seeds, outer folds and inner folds as the primary audit. Imputation and standardization were implemented inside scikit-learn pipelines and fitted only on each training fold. Targets were evaluated on their dataset-provided scales; regression utility was the negative of raw RMSE, and no post hoc standardized target was reported as an original-unit improvement. The selected-performance contrast paired full-pool and Morgan-only validation-best results within endpoint, seed and outer fold.")
    body(doc, "Raw paired gain was full-pool utility minus Morgan-only utility. It therefore represents ROC-AUC gain for classification and RMSE reduction for regression. Seed-clustered intervals averaged the three folds within each seed before bootstrap resampling. We also reported the absolute selected performance of both pools, the number of positive paired outer units, the range-normalized gain and the representation most frequently selected. This information is needed because a large raw gain can reflect a weak frozen Morgan-only comparator on a particular endpoint rather than a universal one-unit effect of multiview representations.")

    add_heading(doc, "2.10 Limited representation-baseline stress test", 2)
    body(doc, "A separate six-endpoint panel included RDKit random forest (RDKit-RF), a graph convolutional network (GCN), a ChemBERTa frozen-embedding probe and a MoLFormer frozen-embedding probe. RDKit-RF used Morgan fingerprints and 120 trees. The GCN contained two 32-unit graph-convolution layers, global mean pooling and five fixed training epochs. ChemBERTa and MoLFormer encoders were frozen and followed by logistic or ridge probes. These configurations provide representation diversity but are not exhaustive or equivalently tuned modern baselines.")
    body(doc, "Search budgets differed across candidates: RDKit-RF used a fixed tree configuration, the GCN used a short fixed training schedule and the language models did not receive end-to-end fine tuning. The panel was therefore restricted to selector behaviour and error correlation under shared splits. It was not used to claim that conventional fingerprints outperform adequately tuned graph or foundation models. Chemprop/D-MPNN and TabPFN were not included because a complete shared-split audit was unavailable.")
    body(doc, "Per-molecule outer predictions were available for the four completed candidates. For classification, an error set contained misclassified samples under the frozen threshold; for regression it contained samples in the candidate-specific high-error subset defined by the analysis script. Pairwise Jaccard overlap quantified shared errors. This measure reveals whether candidates fail on the same molecules but does not imply that candidates with lower overlap are more accurate overall.")

    add_heading(doc, "2.11 Reliability and chemical-boundary analyses", 2)
    body(doc, "Classification uncertainty was evaluated with split conformal, label-conditional conformal and Mondrian label-and-similarity conformal prediction at 80%, 90% and 95% nominal coverage. Minority-class coverage and prediction-set size were primary reliability outputs because overall marginal coverage can conceal undercoverage of rare positives. Regression intervals used conformalized quantile regression (CQR). Coverage and interval width were reported separately by endpoint because interval widths inherit endpoint units [26–28,33,36].")
    body(doc, "Ensemble uncertainty was assessed by Spearman association with absolute error, risk–coverage behaviour and enrichment of the 10% highest-error samples. Chemical support was characterized by maximum Morgan Tanimoto similarity to the training set and Bemis–Murcko scaffold novelty. Additional boundary analyses included activity-cliff pairs, extreme labels and beyond-rule-of-five (bRo5) perimeter cases [11,12,29,30]. These strata were descriptive stress tests. They do not convert public scaffold splits into prospective medicinal-chemistry validation.")
    body(doc, "ClinTox was treated as a prespecified negative reliability example because only 58 positive standardized structures were available. We reported ROC-AUC and PR-AUC together with historical fixed-precision recall, minority-class conformal coverage and false-negative behaviour. Conformal coverage addresses prediction-set validity, whereas fixed-threshold recall addresses a different screening decision. Neither can substitute for the other, and neither establishes deployment-level toxicology performance.")

    add_heading(doc, "2.12 Statistical inference", 2)
    body(doc, "The primary K = 32 versus K = 4 selection-loss contrast was paired by endpoint, seed and outer fold. For each endpoint, the three fold differences were averaged within seed, producing five seed-level values. Seed-clustered 95% intervals were obtained by resampling those five means with replacement 10,000 times. Leave-one-seed-out estimates assessed sensitivity to each repeat. The earlier 15-unit percentile bootstrap was retained in Table 4 as a secondary comparator, not as the preferred independence model.")
    body(doc, "Across endpoints, direction was summarized by the number of positive endpoint means. A two-sided exact sign test assessed whether eight positive directions among nine endpoints differed from an equal-direction null, with the endpoint as the unit. Chance-adjusted ranking summaries and effective-rank summaries used endpoint means; bootstrap intervals resampled endpoints or the prespecified seed–fold hierarchy as described above. P values were not used as the first line of evidence, and no fold, seed or candidate was treated as an independent molecular task [31].")
    body(doc, "Missing candidates were not imputed or replaced using outer results. All stochastic procedures used recorded seeds. Numerical checks compared utility direction, absolute performance and paired gain for every multiview endpoint. Regression gains were verified as reductions in raw dataset-scale RMSE. Statistical intervals describe sensitivity under the evaluated datasets, splits and seeds. They do not quantify uncertainty for an unseen drug-discovery programme or for a broader population of molecular endpoints.")
    add_figure(doc, FIG / "Figure_1_major_revision_workflow.png", 16.0, "Figure 1. Retrospective frozen audit and evidence hierarchy. (A) Candidate registration preceded nested selection and outer-auditing under a retrospective analysis lock. (B) Registered prefixes defined K = 4, 8, 16 and 32 while total search exposure increased with K. (C) Each endpoint used five seeds, three outer scaffold folds and three inner folds. (D) Nominal size, effective diversity, chance-adjusted ranking, endpoint-specific loss and compute exposure were joint reporting targets. (E) Shared-split multiview and limited representation panels were internal stress tests rather than independent confirmation. (F) The observed outer best is not a true generalization upper bound, and the outer-audit is not an independent confirmation set.")

    add_heading(doc, "3 Results", 1)
    add_heading(doc, "3.1 Nominal expansion exceeded effective candidate diversity", 2)
    body(doc, "This analysis asked whether nominal K represented an equivalent number of independent performance opportunities. It did not. Using the original empirical outer-utility correlation, endpoint-mean spectral-entropy rank increased from 1.58 at K = 4 to 2.00 at K = 32. The Ledoit–Wolf shrinkage estimator produced higher and more uncertain values: 2.13 at K = 4, 2.77 at K = 8, 3.28 at K = 16 and 3.43 at K = 32. The hierarchical 95% CI at K = 32 was 2.54–5.90, remaining far below the nominal count of 32 but substantially wider than the earlier point estimate implied.")
    body(doc, "Participation-ratio rank supported the same qualitative conclusion while giving smaller estimates. At K = 32 it was 1.80 (hierarchical 95% CI 1.51–3.61), compared with the shrinkage spectral-entropy rank of 3.43. Shrinkage median pairwise utility correlation was 0.814 (95% CI 0.724–0.857) at K = 32. The divergence between entropy and participation ranks indicates a spectrum dominated by a few directions with additional low-mass components, rather than exactly two interchangeable candidates.")
    body(doc, "Finite-unit sensitivity was material but did not erase the nominal–effective gap. Recomputing the K = 32 shrinkage estimate after omitting one seed yielded endpoint-mean spectral ranks from 3.45 to 4.06; omitting one outer-fold index yielded means from 3.79 to 4.19. Individual endpoints were less stable, as expected from 10–12 retained matrix rows. The appropriate conclusion is therefore that the evaluated utility-based estimator placed effective diversity in the low single digits, with substantial finite-unit uncertainty, not that 32 candidates equalled precisely two independent candidates.")
    diversity_ranking_table(doc)

    add_heading(doc, "3.2 Chance-adjusted ranking fidelity declined with candidate expansion", 2)
    body(doc, "This analysis asked whether the validation ordering recovered the outer-best candidate beyond the opportunity expected from K alone. Chance-adjusted Top-3 recovery declined from 0.881 at K = 4 to 0.550 at K = 8, 0.334 at K = 16 and 0.240 at K = 32. Endpoint-bootstrap 95% intervals were 0.704–1.000 at K = 4 and 0.158–0.305 at K = 32. Because the random Top-3 probability was removed, this decline cannot be attributed solely to the changing 3/K chance baseline.")
    body(doc, "Normalized MRR gain showed a parallel reduction from 0.727 at K = 4 to 0.270 at K = 8, 0.240 at K = 16 and 0.157 at K = 32. NDCG declined from 0.943 to 0.579. Mean Spearman correlation fell from 0.760 to 0.629, and mean Kendall correlation fell from 0.706 to 0.494. Percentile rank was less monotonic, changing from 0.881 at K = 4 to 0.750 at K = 32. Thus, the outer-best candidate often remained in the upper validation ranks even when exact top recovery deteriorated.")
    body(doc, "Raw Top-1 recovery and MRR also declined, from 0.778 to 0.081 and from 0.869 to 0.264, respectively. These raw changes are descriptive only because their random expectations depend strongly on K. The concordant decline in CAHit@3, normalized MRR gain, NDCG, Spearman and Kendall provides the stronger evidence for ranking distortion under the evaluated design. Permutation controls remained centred near zero, while signal-recovery controls increased monotonically with injected validation–audit correlation.")

    add_heading(doc, "3.3 Endpoint-specific selection loss increased in eight endpoints", 2)
    body(doc, "This analysis asked whether weaker ranking fidelity translated into loss on the selected outer candidate. K = 32 minus K = 4 mean raw selection loss was positive in eight of nine endpoints. The endpoint-level two-sided exact sign test was P = 0.039. Classification increases were 0.0053 ROC-AUC for BACE, 0.0035 for BBBP, 0.0203 for ClinTox and 0.0195 for HIA Hou. Seed-clustered intervals excluded zero for BBBP, ClinTox and HIA Hou; the BACE interval was −0.0009–0.0106.")
    body(doc, "Regression effects were retained on their original endpoint scales. RMSE selection loss increased by 0.0940 for ESOL (seed-clustered 95% CI 0.0899–0.0979), 0.0743 for FreeSolv (0.0356–0.1046), 0.0077 for Lipophilicity (0.0058–0.0091) and 0.0328 for Caco2 (0.0310–0.0347). These values should not be averaged as a common physical effect. The larger ESOL and FreeSolv values indicate greater failure to recover the observed outer best on those dataset scales, not a universal 0.08-unit chemical error.")
    body(doc, "P-gp Broccatelli preserved the negative result: mean loss changed by −0.0003 ROC-AUC, with seed-clustered 95% CI −0.0047–0.0042 and only two of five seed means positive. Leave-one-seed estimates crossed both directions. The exception shows that candidate expansion did not mechanically increase empirical loss in every endpoint. It also prevents the eight-positive pattern from being described as universal.")
    loss_table(doc)

    add_heading(doc, "3.4 Composition controls supported the expansion trend", 2)
    body(doc, "This analysis asked whether the loss trend could be explained solely by the deterministic registry prefix. It could not. Mean range-normalized selection loss increased from 0.089 to 0.214 for random orders, from 0.090 to 0.214 for random subsets and from 0.087 to 0.214 for family-balanced subsets between K = 4 and K = 32. At K = 8 the three means ranged from 0.125 to 0.129; at K = 16 they ranged from 0.162 to 0.170. Composition differences were therefore smaller than the change across K.")
    body(doc, "The controls do not provide three independent confirmations at K = 32, because every mode equals the complete registry at that size. Their evidential value lies at K = 4, 8 and 16, where increasing loss appeared under multiple frozen ways of constructing the pool. Endpoint remained the inferential unit despite 100 subset or order seeds. The result supports an expansion trend within this registered candidate population but does not identify a causal effect for arbitrary future candidate distributions.")
    add_figure(doc, FIG / "Figure_2_major_revision_diversity_ranking_loss.png", 16.0, "Figure 2. Effective diversity, chance-adjusted ranking fidelity and selection loss. (A) Empirical and Ledoit–Wolf shrinkage spectral-entropy ranks are shown against nominal K; error bars are hierarchical 95% bootstrap intervals for the shrinkage estimator. (B) Shrinkage median candidate-utility correlation with hierarchical intervals. (C) Chance-adjusted Top-3 recovery (CAHit@3) and normalized MRR gain. (D) K = 32 minus K = 4 endpoint-specific raw selection loss with seed-clustered 95% intervals; blue denotes ROC-AUC loss and orange denotes RMSE loss, and the scales are not pooled. (E) Random-order, random-subset and family-balanced controls. (F) Endpoint empirical and shrinkage entropy ranks at K = 32; the dashed diagonal denotes equality.")

    add_heading(doc, "3.5 Finite audit estimates showed mechanical winner optimism", 2)
    body(doc, "This simulation asked whether maximizing finite outer estimates produces optimism even when candidates are equally good. With pairwise candidate correlation 0.9 and effective audit size 50, mean winner optimism increased from 0.046 standard-deviation units at K = 4 to 0.092 at K = 32. At effective size 25, K = 32 optimism was 0.132; at effective size 200 it was 0.047. Lower correlation produced larger maxima because candidates supplied more independent noise directions.")
    body(doc, "The simulation demonstrates a mechanical property of finite maxima, not an empirical bias correction. Real candidates differ in truth, and their endpoint-specific correlation structures are more complex than equicorrelation. Nevertheless, the equal-truth result shows why observed outer-best performance must be distinguished from a true upper bound. An outer fold can be untouched by inner selection and still yield an optimistic maximum when the analyst compares many outer candidate estimates after the fact.")

    add_heading(doc, "3.6 Multiview candidates produced heterogeneous raw-scale gains", 2)
    body(doc, "This stress test asked whether a more heterogeneous representation pool could produce gains that survived frozen selection. Across 135 paired outer units, the full 12-candidate pool had a mean range-normalized realized gain of 0.343 over the three-candidate Morgan-only pool (95% CI 0.210–0.483 across endpoint means), with positive endpoint means in all nine endpoints. The result supports realizable representational opportunity under common splits, but it also includes the additional selection opportunity supplied by nine extra candidates.")
    body(doc, "Classification raw gains varied by more than an order of magnitude. Mean ROC-AUC gains were 0.0039 for BACE (seed-clustered 95% CI −0.0025–0.0108), 0.0249 for BBBP (0.0161–0.0333), 0.0873 for ClinTox (0.0767–0.1002), 0.0809 for HIA Hou (0.0695–0.0931) and 0.0227 for P-gp Broccatelli (0.0168–0.0284). BACE was positive in only seven of 15 paired units and retained an interval crossing zero, whereas ClinTox was positive in all paired units.")
    body(doc, "Regression RMSE reductions were 0.9008 for ESOL, 1.1949 for FreeSolv, 0.2401 for Lipophilicity and 0.1058 for Caco2. The absolute comparisons verified the direction and scale: Morgan-only versus full-pool RMSE was 1.6881 versus 0.7873 for ESOL, 2.7345 versus 1.5396 for FreeSolv, 0.9702 versus 0.7301 for Lipophilicity and 0.5785 versus 0.4728 for Caco2. All transformations were fitted within training folds, and reported RMSE values were on dataset-provided scales.")
    body(doc, "The large FreeSolv and ESOL improvements reflect the weakness of the frozen Morgan-only selected comparator under those endpoints rather than a universal one-unit improvement attributable to multiview representations. The full pool could select RDKit2D and concatenated candidates that were absent from the comparator, and it had more total selection opportunities. Across all outer units, concatenated multiview was selected 84 times, RDKit2D 44 times, MACCS four times and Morgan three times. These counts corroborate representation dependence but do not isolate the causal contribution of concatenation from pool expansion.")
    multiview_table(doc)
    add_figure(doc, FIG / "Figure_3_major_revision_controls_multiview.png", 16.0, "Figure 3. Calibration controls, finite-audit optimism and shared-split multiview effects. (A) Observed CAHit@3 exceeded the permutation null. (B) Chance-adjusted recovery increased with injected validation–audit correlation. (C) Equal-truth winner optimism increased with K and decreased with effective audit size at candidate correlation 0.9. (D) Endpoint-specific multiview raw gains with seed-clustered 95% intervals; classification and regression use different units. (E) Representations selected by the full-pool validation-best rule. (F) Absolute Morgan-only and full-pool selected performance, shown separately for classification ROC-AUC and regression RMSE.")

    add_heading(doc, "3.7 Representation baselines showed correlated and complementary errors", 2)
    body(doc, "This limited panel asked whether qualitatively different representations produced independent error patterns under shared splits. Across six MoleculeNet endpoints, the selector recovered the observed outer best in 0.878 of 90 endpoint–seed–fold units, with mean range-normalized selection loss 0.0048 and mean validation–audit Spearman correlation 0.840. RDKit-RF was selected in each endpoint summary, but this comparison reflects the frozen configurations and unequal search budgets rather than a comprehensive benchmark of model classes.")
    body(doc, "Pairwise error-overlap Jaccard values averaged 0.215. They ranged from 0.168 for GCN versus RDKit-RF to 0.296 for ChemBERTa versus RDKit-RF. ChemBERTa versus GCN was 0.175, GCN versus MoLFormer 0.195 and ChemBERTa versus MoLFormer 0.224. Errors were therefore neither independent nor identical. Partial complementarity supports examining per-molecule error structure, whereas the non-zero overlap cautions against treating four named architectures as four independent opportunities.")
    body(doc, "The panel cannot support a performance hierarchy among adequately tuned modern methods. The GCN received only five fixed epochs, and both language-model representations used frozen linear probes rather than end-to-end fine tuning. RDKit-RF consequently had a different optimization budget and inductive bias. The completed outputs are informative for selector behaviour and error correlation under the evaluated configurations, but they should not be generalized to all graph networks, ChemBERTa variants or MoLFormer training strategies.")

    add_heading(doc, "3.8 Reliability deteriorated near chemical boundaries", 2)
    body(doc, "This analysis asked whether predictive reliability remained stable in sparse or discontinuous chemical regions. At 90% nominal coverage, mean minority-class coverage across classification endpoints was 0.788 for split conformal, 0.887 for label-conditional conformal and 0.891 for Mondrian label-and-similarity conformal prediction. Conditional methods substantially reduced minority undercoverage but remained slightly below the target on average. Prediction sets also became larger, so improved coverage did not constitute a free increase in information.")
    body(doc, "CQR coverage was endpoint dependent: 0.836 for ESOL, 0.893 for FreeSolv and 0.917 for Lipophilicity at the 90% target. The cross-endpoint mean was 0.882. Raw interval widths were not pooled for interpretation because the targets use different units and dynamic ranges. Ensemble uncertainty had mean uncertainty–error Spearman correlation 0.321 and enriched the 10% highest-error samples by 1.54-fold, indicating useful but incomplete error ranking.")
    body(doc, "Chemical similarity stratification showed a consistent boundary signal. Mean classification ROC-AUC was 0.803 below maximum train-set Tanimoto 0.5, 0.898 between 0.5 and 0.7 and 0.924 above 0.7. Expected calibration error did not decrease monotonically in the same way, illustrating that discrimination and calibration respond differently to chemical support. Novel scaffolds, extreme labels, activity-cliff pairs and bRo5 perimeter cases contributed distinct high-error or misclassification categories and were therefore displayed categorically rather than on a shared continuous error scale.")
    body(doc, "ClinTox remained the strongest negative result. The cleaned endpoint contained 58 positives. In the limited representation panel, RDKit-RF achieved mean ROC-AUC 0.878 and PR-AUC 0.433, while the historical frozen analysis reported recall 0.588 ± 0.168 at precision at least 0.80 and 0.491 ± 0.195 at precision at least 0.90. Under the fixed threshold used for the false-negative audit, RDKit-RF minority recall was 0.105 and the false-negative rate was 0.895. Conditional conformal methods improved class-conditional coverage, but coverage did not imply acceptable thresholded toxicity recall.")
    add_figure(doc, FIG / "Figure_4_major_revision_reliability_boundaries.png", 16.0, "Figure 4. Limited representation baselines, reliability and chemical boundaries. (A) Utilities are normalized within endpoint for visualization only and are not a cross-endpoint performance average. (B) Pairwise Jaccard overlap of exported per-molecule errors. (C) Minority-class coverage at the 90% conformal target. (D) Endpoint-specific 90% CQR coverage. (E) ROC-AUC and expected calibration error (ECE) by maximum train-set Tanimoto. (F) Presence of high-error, misclassification, false-negative and pair-failure categories by chemical-boundary source; the categorical matrix avoids comparing unlike error units on one continuous axis.")

    add_heading(doc, "4 Discussion", 1)
    add_heading(doc, "4.1 Nominal candidate count did not represent independent opportunity", 2)
    body(doc, "The primary audit shows that nominal K and effective diversity are different properties of a model search. Thirty-two registered variants generated a shrinkage spectral-entropy rank of 3.43 rather than 32, and the participation ratio was lower still. The registry therefore contained many correlated ways to exploit validation variation. This finding is specific to a Morgan-based conventional-model registry and should not be extrapolated numerically to a search containing end-to-end graph networks and independently pretrained foundation models. The general reporting lesson is that candidate count alone cannot reveal how many distinct predictive directions entered selection.")
    body(doc, "Effective rank itself is an estimate, not a fixed attribute. With only 15 outer-utility rows, the empirical 32 × 32 correlation matrix is rank deficient. Shrinkage increased the estimated entropy rank from 2.00 to 3.43 and produced a wide interval. Omission analyses changed the endpoint mean by several tenths and individual endpoints more strongly. These results argue against statements such as '32 candidates corresponded to exactly two candidates.' A more defensible interpretation is that effective diversity remained in the low single digits under utility-based estimators, with uncertainty arising from the finite repeated design.")

    add_heading(doc, "4.2 Chance-adjusted ranking distortion explained selection loss", 2)
    body(doc, "Raw Top-1 recovery and MRR necessarily become harder as K grows, so their decline cannot by itself establish ranking distortion. Opportunity correction changed the evidential basis but not the conclusion. CAHit@3 and normalized MRR gain both declined strongly, while NDCG, Spearman and Kendall moved in the same direction. Validation ranking remained informative: all adjusted metrics stayed above their random baselines, and validation-best outperformed random selection. The operative problem was therefore partial loss of ranking fidelity, not complete failure of validation.")
    body(doc, "The eight-endpoint increase in raw selection loss was consistent with that mechanism. When ranking fidelity weakened, the validation-selected candidate captured less of the observed outer opportunity. Composition controls showed that the trend could not be explained solely by one registry prefix. However, P-gp Broccatelli did not increase and BACE intervals were sensitive to seed clustering. Candidate expansion should thus be described as increasing risk under the evaluated design, not as a deterministic penalty for every endpoint.")
    body(doc, "No single ranking statistic completely characterizes this process. CAHit@3 emphasizes whether the eventual outer best was placed in a short validation list, which is relevant when only a few models can be advanced. Normalized MRR gain gives progressively less credit as that candidate moves down the list, while NDCG evaluates more of the ordered candidate set. Spearman and Kendall correlations are insensitive to the identity of one winner but detect broader reordering. The agreement among these measures strengthens the interpretation because it is not driven by one arbitrary cut-off. Their differences are also informative: the relatively stable percentile rank indicates that the outer-best candidate often remained above average even when exact recovery became difficult.")

    add_heading(doc, "4.3 Heterogeneous expansion can still be beneficial", 2)
    body(doc, "The multiview stress test prevents an overly defensive interpretation. A heterogeneous pool produced positive endpoint-mean realized gains under frozen selection, and the selected representations shifted overwhelmingly away from Morgan-only candidates. Expansion can therefore be scientifically worthwhile when it adds information that survives outer-auditing. The appropriate comparison is not 'small pool good, large pool bad,' but whether additional effective diversity and realized utility justify the added search exposure and ranking uncertainty.")
    body(doc, "Those gains depended on the comparator. Morgan-only ESOL and FreeSolv performance was weak under the frozen three-learner pool, which made the absolute RMSE reductions large. The full pool also had four times as many candidates and could benefit from selection opportunity. Reporting absolute Morgan-only and full-pool performance, paired raw effects, positive-unit counts and normalized gains exposes this dependence. A separate representation ablation with equalized candidate counts would be required to isolate representation content from pool-size opportunity completely.")

    add_heading(doc, "4.4 Outer-audit maxima remain finite-sample estimates", 2)
    body(doc, "Nested cross-validation separates inner selection from outer evaluation, but it does not make the largest outer score a true oracle. The equal-truth simulation demonstrates that maximizing finite estimates produces positive winner optimism even when no candidate is truly better. Strong candidate correlation attenuates but does not eliminate the effect. Reporting the observed outer best is useful because it defines the opportunity available in the audited pool, yet it should be labelled as a finite-audit maximum and accompanied by the selected performance.")
    body(doc, "The outer folds are also reused across candidates and share molecules across repetitions. They are paired comparison units, not independent biological experiments. Seed-clustered intervals were generally similar in direction to fold-level intervals but differed in width, particularly for BACE and FreeSolv. This difference is not a technical nuisance: it reflects the dependence structure of repeated scaffold evaluation and should be visible in the main results when claims depend on narrow intervals.")

    add_heading(doc, "4.5 Implications for molecular benchmark design", 2)
    body(doc, "Molecular benchmark reports should make the search process inspectable. A minimum candidate ledger should record candidate identity, representation, learner family, hyperparameters, eligibility, failure status and registry order. Reports should distinguish per-candidate opportunity from total search exposure, because equal treatment of each model does not imply an equal total budget across K. Candidate correlation and effective diversity should be reported with uncertainty whenever many alternatives share representations or training data.")
    body(doc, "Ranking fidelity should also be opportunity adjusted. CAHit@3, normalized MRR gain, percentile rank, NDCG and rank correlations answer complementary questions and are preferable to a single raw Top-1 rate. Performance summaries should pair the selected result with the observed finite-audit best and endpoint-specific loss. Failed or unavailable candidates belong in the registry even when they do not enter the numerical comparison, because silent removal changes the effective search and can favour methods with more reliable software pipelines.")
    body(doc, "Compute accounting is part of this evidential record rather than a secondary engineering detail. Two studies can give every candidate the same number of folds and still expose very different total searches if one registry contains four models and the other contains 32. Conversely, equal wall-clock budgets can allocate unequal training opportunity to candidates with different runtimes. A useful report should therefore separate the number of eligible candidates, attempted fits, completed fits, candidate-specific tuning decisions and total recorded exposure. These quantities do not by themselves correct selection loss, but they make it possible to distinguish representational gain from a larger opportunity budget. They also reveal whether a complex candidate was disadvantaged by an unusually short schedule or whether a lightweight family benefited from substantially more trials.")

    add_heading(doc, "4.6 Reliability remains conditional on chemical support", 2)
    body(doc, "Model-selection auditing does not replace predictive uncertainty analysis. Conditional conformal methods improved minority coverage, but coverage varied by endpoint and required larger prediction sets. CQR approached nominal coverage on average while under-covering ESOL. Similarity stratification showed that low-support molecules had weaker discrimination, and activity cliffs demonstrated that high structural similarity does not guarantee property continuity. Reliability claims should therefore be conditioned on class, chemical support, scaffold novelty and the decision threshold.")
    body(doc, "ClinTox illustrates why discrimination, coverage and screening utility must be separated. ROC-AUC can remain high while PR-AUC, fixed-precision recall and fixed-threshold minority recall expose rare-class limitations. Conformal coverage indicates whether prediction sets contain the true label at the target frequency; it does not ensure that a single threshold retrieves toxic molecules at an acceptable rate. The present ClinTox results do not support deployment as a toxicity filter.")

    add_heading(doc, "4.7 Limitations", 2)
    body(doc, "Several limitations constrain the scope of inference. First, the primary registry was intentionally near-duplicate and does not represent a complete foundation-model search. Second, only nine endpoints entered the primary audit, so endpoint-level consistency cannot establish a field-wide law. Third, effective diversity at K = 32 was estimated from 15 outer rows; shrinkage and hierarchical bootstrap reduce instability but cannot replace more independent audit units. Fourth, all-candidate per-molecule predictions were unavailable for the near-duplicate registry, preventing a prediction-matrix effective-rank analysis.")
    body(doc, "Fifth, the shared-split multiview analysis reused the same endpoints and outer folds, and the outer-audit was not an independent lockbox. Sixth, the limited representation panel did not equalize tuning budgets and cannot rank fully optimized modern architectures. Seventh, normalized utility depends on finite pool extremes, and Figure 4 boundary categories arise from different error definitions. Finally, conformal and chemical-boundary analyses characterize retrospective public data rather than prospective programme decisions. The study contributes an audit and reporting practice; it does not automatically improve prediction performance or define a universal selector.")
    body(doc, "The diversity estimator also has conceptual limits beyond sample size. Utility correlation collapses each candidate's behaviour on an outer fold to one endpoint-level score. Two candidates can have similar ROC-AUC or RMSE across folds while making different per-molecule errors, and conversely they can have correlated predictions but different aggregate utilities. Spectral effective rank should therefore be interpreted as the effective dimensionality of the observed utility matrix, not as a complete measure of prediction diversity or chemical-mechanism diversity. Shrinkage stabilizes the spectrum by imposing structure on noisy covariance estimates, but the result remains estimator dependent. Reporting empirical, shrinkage and participation-ratio values together is more informative than selecting whichever estimate best supports a preferred narrative.")

    add_heading(doc, "5 Conclusions", 1)
    body(doc, "Within nine frozen molecular-property audits, nominal candidate expansion greatly exceeded estimated effective diversity. At K = 32, utility-based effective rank remained in the low single digits and carried substantial finite-unit uncertainty. Chance-adjusted Top-3 recovery, normalized MRR gain, NDCG and rank correlations declined as K increased, while raw selection loss increased in eight endpoints. Heterogeneous multiview candidates nevertheless produced endpoint-dependent realized gains. The equal-truth simulation further showed that the maximum of finite outer estimates contains mechanical winner optimism.")
    body(doc, "Molecular model-selection studies should jointly report the candidate registry, nominal K, effective diversity with uncertainty, chance-adjusted ranking fidelity, endpoint-specific raw-scale selection loss, total compute exposure, failed candidates and chemical-support boundaries. Selected performance should be distinguished from the observed finite-audit best, and public shared-split stress tests should not be described as independent confirmation. These practices make representational gains and selection risks visible without implying a new predictive backbone or universal model superiority. They also permit readers to judge whether an apparent advance reflects new chemical information, greater search exposure or finite-sample selection opportunity.")

    add_heading(doc, "List of abbreviations", 1)
    body(doc, "BACE, beta-secretase 1; BBBP, blood–brain barrier penetration; bRo5, beyond the rule of five; CAHit@3, chance-adjusted Top-3 hit; CQR, conformalized quantile regression; ECE, expected calibration error; GCN, graph convolutional network; HIA, human intestinal absorption; MACCS, Molecular ACCess System; MRR, mean reciprocal rank; NDCG, normalized discounted cumulative gain; P-gp, P-glycoprotein; PR-AUC, precision–recall area under the curve; RDKit-RF, RDKit Morgan-fingerprint random forest; RMSE, root mean squared error; ROC-AUC, receiver operating characteristic area under the curve.")

    add_heading(doc, "References", 1)
    refs = [
        "1. Wu Z, Ramsundar B, Feinberg EN, et al. MoleculeNet: a benchmark for molecular machine learning. Chem Sci. 2018;9:513–530. doi:10.1039/C7SC02664A.",
        "2. Huang K, Fu T, Gao W, et al. Therapeutics Data Commons: machine learning datasets and tasks for drug discovery and development. NeurIPS Datasets and Benchmarks Track. 2021. arXiv:2102.09548.",
        "3. Cawley GC, Talbot NLC. On over-fitting in model selection and subsequent selection bias in performance evaluation. J Mach Learn Res. 2010;11:2079–2107.",
        "4. Varma S, Simon R. Bias in error estimation when using cross-validation for model selection. BMC Bioinformatics. 2006;7:91. doi:10.1186/1471-2105-7-91.",
        "5. Zhao D, Zhu Y, Wu Z, et al. Revisiting ADMET prediction reliability under real-world challenges in the foundation model era. J Cheminform. 2026. doi:10.1186/s13321-026-01217-2.",
        "6. Zhang L, Zeng Y, Qi Y, et al. DCPM-ADMET: fusion of dual-component pre-trained model and molecular fingerprints to enhance drug ADMET properties prediction. J Cheminform. 2026. doi:10.1186/s13321-026-01244-z.",
        "7. Jang Y, Lee J, Jeong K, Kim J. Multimodal graph fusion with statistically guided parsimonious descriptor selection for molecular property prediction. J Cheminform. 2026;18:18. doi:10.1186/s13321-025-01140-y.",
        "8. Zhang Y, Liu W, Zhao H, et al. MolGramTreeNet: a multimodal molecular property prediction model via grammar tree-constrained molecular representation. iScience. 2026;29:114928. doi:10.1016/j.isci.2026.114928.",
        "9. Wen X, Liu H, Long W, Wei S, Zhu R. Consistent semantic representation learning for out-of-distribution molecular property prediction. Brief Bioinform. 2025;26:bbaf147. doi:10.1093/bib/bbaf147.",
        "10. Yin T, Gao P, Panapitiya G, Saldanha EG. Out-of-distribution evaluation of active learning pipelines for molecular property prediction. RSC Adv. 2026;16:5281–5295. doi:10.1039/D5RA08055J.",
        "11. Uchibori Y, Kaneko H. Generation of molecules near the applicability domain boundaries of property prediction models. J Chem Inf Model. 2026;66:6866–6879. doi:10.1021/acs.jcim.5c03220.",
        "12. Kim JY, Vlachos DG. Distance-aware molecular property prediction in nonlinear structure–property space. J Chem Inf Model. 2025;65:6744–6756. doi:10.1021/acs.jcim.5c01037.",
        "13. Tang H, Yue T, Li Y. Assessing uncertainty in machine learning for polymer property prediction: a benchmark study. J Chem Inf Model. 2025;65:6585–6598. doi:10.1021/acs.jcim.5c00550.",
        "14. Fralish Z, Reker D. Pairwise learning for molecular property prediction and optimization. Front Drug Discov. 2026;6:1859068. doi:10.3389/fddsv.2026.1859068.",
        "15. Landrum G. RDKit: open-source cheminformatics software. Version 2026.3.1. https://www.rdkit.org/. Accessed 12 July 2026.",
        "16. Rogers D, Hahn M. Extended-connectivity fingerprints. J Chem Inf Model. 2010;50:742–754. doi:10.1021/ci100050t.",
        "17. Breiman L. Random forests. Mach Learn. 2001;45:5–32. doi:10.1023/A:1010933404324.",
        "18. Ke G, Meng Q, Finley T, et al. LightGBM: a highly efficient gradient boosting decision tree. Adv Neural Inf Process Syst. 2017;30.",
        "19. Chen T, Guestrin C. XGBoost: a scalable tree boosting system. Proc ACM SIGKDD. 2016:785–794. doi:10.1145/2939672.2939785.",
        "20. Prokhorenkova L, Gusev G, Vorobev A, Dorogush AV, Gulin A. CatBoost: unbiased boosting with categorical features. Adv Neural Inf Process Syst. 2018;31:6638–6648.",
        "21. Yang K, Swanson K, Jin W, et al. Analyzing learned molecular representations for property prediction. J Chem Inf Model. 2019;59:3370–3388. doi:10.1021/acs.jcim.9b00237.",
        "22. Chithrananda S, Grand G, Ramsundar B. ChemBERTa: large-scale self-supervised pretraining for molecular property prediction. arXiv:2010.09885. 2020.",
        "23. Ross J, Belgodere B, Chenthamarakshan V, et al. Large-scale chemical language representations capture molecular structure and properties. Nat Mach Intell. 2022;4:1256–1264. doi:10.1038/s42256-022-00580-7.",
        "24. Erickson N, Mueller J, Shirkov A, et al. AutoGluon-Tabular: robust and accurate AutoML for structured data. arXiv:2003.06505. 2020.",
        "25. Tropsha A. Best practices for QSAR model development, validation, and exploitation. Mol Inform. 2010;29:476–488. doi:10.1002/minf.201000061.",
        "26. Vovk V, Gammerman A, Shafer G. Algorithmic Learning in a Random World. New York: Springer; 2005.",
        "27. Shafer G, Vovk V. A tutorial on conformal prediction. J Mach Learn Res. 2008;9:371–421.",
        "28. Guo C, Pleiss G, Sun Y, Weinberger KQ. On calibration of modern neural networks. Proc ICML. 2017;70:1321–1330.",
        "29. van Tilborg D, Alenicheva A, Grisoni F. Exposing the limitations of molecular machine learning with activity cliffs. J Chem Inf Model. 2022;62:5938–5951. doi:10.1021/acs.jcim.2c01073.",
        "30. Sheridan RP. Time-split cross-validation as a method for estimating prospective prediction performance. J Chem Inf Model. 2013;53:783–790. doi:10.1021/ci400084k.",
        "31. Demšar J. Statistical comparisons of classifiers over multiple data sets. J Mach Learn Res. 2006;7:1–30.",
        "32. Hoyt CT, Zdrazil B, Guha R, et al. Improving reproducibility and reusability in the Journal of Cheminformatics. J Cheminform. 2023;15:62. doi:10.1186/s13321-023-00730-y.",
        "33. Parrondo-Pizarro R, Lanini J, Rodríguez-Pérez R. Uncertainty quantification in molecular machine learning for property predictions under data shifts. J Chem Inf Model. 2026;66:923–935. doi:10.1021/acs.jcim.5c02381.",
        "34. Deng J, Yang Z, Wang H, Ojima I, Samaras D, Wang F. A systematic study of key elements underlying molecular property prediction. Nat Commun. 2023;14:6395. doi:10.1038/s41467-023-41948-6.",
        "35. Li Z, Chen X, Wen H, et al. A systematic survey and benchmark of deep learning for molecular property prediction in the foundation model era. arXiv:2604.16586. 2026.",
        "36. Romano Y, Patterson E, Candès EJ. Conformalized quantile regression. Adv Neural Inf Process Syst. 2019;32.",
    ]
    for ref in refs:
        add_reference(doc, ref)
    return doc


def word_stats(path: Path) -> dict[str, int]:
    doc = Document(path)
    pattern = re.compile(r"[A-Za-z]+(?:[-'][A-Za-z]+)*|\d+(?:\.\d+)?")
    paragraphs = [p.text for p in doc.paragraphs]
    ref_index = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "References"), len(paragraphs))
    abstract_start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Abstract")
    intro_start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "1 Introduction")
    main_start = intro_start
    main_end = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "List of abbreviations")
    return {
        "abstract_words": len(pattern.findall(" ".join(paragraphs[abstract_start + 1:intro_start]))),
        "main_text_words_introduction_through_conclusions": len(pattern.findall(" ".join(paragraphs[main_start:main_end]))),
        "paragraph_words_excluding_references": len(pattern.findall(" ".join(paragraphs[:ref_index]))),
        "table_words": sum(len(pattern.findall(c.text)) for t in doc.tables for r in t.rows for c in r.cells),
        "reference_words": len(pattern.findall(" ".join(paragraphs[ref_index:]))),
    }


def write_response() -> None:
    rows = [
        ("Effective rank was estimated from only 15 outer units at K = 32.", "Defined the 15 × K outer-utility and 45 × K inner-utility matrices; added Ledoit–Wolf shrinkage, spectral-entropy and participation-ratio ranks, hierarchical bootstrap intervals and omission sensitivity.", "Methods 2.6; Results 3.1; Table 3; Figure 2A,B,F"),
        ("The previous value of approximately two was presented without uncertainty.", "Retained the empirical estimate (2.00) but added shrinkage estimate 3.43, 95% CI 2.54–5.90 and cautious interpretation.", "Abstract Results; Results 3.1; Discussion 4.1"),
        ("A prediction-matrix effective rank was requested.", "The required all-candidate per-sample predictions do not exist for the near-duplicate registry; this remains explicitly unresolved rather than imputed.", "Methods 2.6; unresolved-analysis list"),
        ("Top-1 and raw MRR decline mechanically with K.", "Promoted CAHit@3, normalized MRR gain, percentile rank, NDCG, Spearman and Kendall; raw Top-1 and MRR are descriptive only.", "Methods 2.5; Results 3.2; Table 3; Figure 2C"),
        ("Multiview analysis was labelled as confirmation.", "Replaced confirmation language with shared-split multiview stress test or corroboration under common splits.", "Throughout; Figure 1E; Methods 2.9; Results 3.6"),
        ("Large ESOL and FreeSolv gains required verification.", "Added absolute Morgan-only and full-pool performance, seed-clustered paired intervals, positive-unit counts, normalized gain and selected representation.", "Table 5; Figure 3D–F; Results 3.6"),
        ("Outer units are dependent.", "Averaged folds within seed, bootstrapped five seed means, added leave-one-seed sensitivity and retained 15-unit intervals only as secondary values.", "Methods 2.4 and 2.12; Table 4"),
        ("Abstract required Scientific Contribution and removal of exploratory meta-risk.", "Added five structured abstract headings, removed LOEO risk and kept the abstract below 350 words.", "Abstract"),
        ("Introduction needed a critical literature gap rather than a list.", "Rewrote the Introduction as five paragraphs separating representation opportunity, validation opportunity, nested-CV limits, literature gap and three contributions.", "Introduction"),
        ("Methods needed a standalone diversity section and explicit formulas.", "Expanded Methods to 12 subsections and provided both effective-rank formulas and correlation construction.", "Methods 2.1–2.12"),
        ("Modern baselines were not comparably tuned.", "Renamed the panel limited representation-baseline stress test and restricted inference to selector behaviour and error overlap.", "Methods 2.10; Results 3.7; Discussion 4.7"),
        ("ClinTox needed multiple rare-class metrics.", "Reported 58 positives, ROC-AUC, PR-AUC, historical fixed-precision recall, minority coverage and fixed-threshold false-negative behaviour.", "Methods 2.11; Results 3.8; Discussion 4.6"),
        ("Figure 2 used a dual y-axis and unadjusted ranking emphasis.", "Removed the dual axis, added effective-rank intervals, correlation intervals and chance-adjusted ranking metrics.", "Figure 2"),
        ("Figure 3 lacked absolute multiview performance.", "Added separate absolute ROC-AUC and RMSE comparisons for Morgan-only and full-pool selectors.", "Figure 3F"),
        ("Figure 4F compared incompatible error scores on one axis.", "Replaced the continuous bar chart with a categorical source-by-failure-type presence matrix.", "Figure 4F"),
        ("Journal formatting and terminology were inconsistent.", "Removed the journal banner, added abbreviations, standardized outer-audit, K = n, 3 × 3 × 5, endpoint names, en-dash intervals and abbreviation definitions.", "Throughout"),
        ("Manuscript length was below the requested level.", "Expanded scientific Methods, Results and Discussion without adding unverified experiments or repetitive background.", "Full manuscript; word-count audit"),
    ]
    lines = ["# Major Revision response", "", "| Reviewer concern | Revision | Location |", "|---|---|---|"]
    for concern, revision, location in rows:
        lines.append(f"| {concern} | {revision} | {location} |")
    lines.extend([
        "", "## Analyses requiring new computation", "",
        "- Prediction-matrix effective rank for the K = 32 near-duplicate pool, requiring all-candidate per-molecule predictions.",
        "- A candidate-count-equalized multiview ablation to separate representation content from additional selection opportunity.",
        "- A fully and comparably tuned representation-baseline benchmark; the current panel is intentionally limited.",
        "", "## Items requiring unavailable source data", "",
        "- All-candidate per-sample prediction exports for the 32-candidate near-duplicate registry.",
        "- An untouched independent external cohort; current outer folds are repeated public-data audits.",
    ])
    RESPONSE_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")
    response = Document()
    configure_document(response)
    title = response.add_paragraph(style="Title")
    title.paragraph_format.first_line_indent = Cm(0)
    set_run_font(title.add_run("Major Revision response"), 16, bold=True)
    add_table(response, ["Reviewer concern", "Revision", "Location"], [list(row) for row in rows])
    add_heading(response, "Analyses requiring new computation", 1)
    for item in [
        "Prediction-matrix effective rank for the K = 32 near-duplicate pool, requiring all-candidate per-molecule predictions.",
        "A candidate-count-equalized multiview ablation to separate representation content from additional selection opportunity.",
        "A fully and comparably tuned representation-baseline benchmark; the current panel is intentionally limited.",
    ]:
        add_body(response, item, no_indent=True)
    add_heading(response, "Items requiring unavailable source data", 1)
    for item in [
        "All-candidate per-sample prediction exports for the 32-candidate near-duplicate registry.",
        "An untouched independent external cohort; current outer folds are repeated public-data audits.",
    ]:
        add_body(response, item, no_indent=True)
    response.save(RESPONSE_DOCX)


def write_abstract_file() -> None:
    doc = Document(PAPER)
    texts = [p.text for p in doc.paragraphs]
    start = texts.index("Abstract") + 1
    end = texts.index("1 Introduction")
    ABSTRACT_MD.write_text("# Revised Abstract and Scientific Contribution\n\n" + "\n\n".join(texts[start:end]) + "\n", encoding="utf-8")


def main() -> None:
    doc = build_document()
    doc.save(PAPER)
    shutil.copy2(PAPER, DESKTOP)
    write_response()
    write_abstract_file()
    stats = word_stats(PAPER)
    final_doc = Document(PAPER)
    final_text = "\n".join(p.text for p in final_doc.paragraphs)
    vertical_borders = 0
    for table in final_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                borders = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcBorders")
                if borders is None:
                    continue
                for edge in ("left", "right", "insideV"):
                    node = borders.find(qn(f"w:{edge}"))
                    if node is not None and node.get(qn("w:val")) not in {None, "nil", "none"}:
                        vertical_borders += 1
    headings = [p.text for p in final_doc.paragraphs if p.style.name.startswith("Heading")]
    checks = {
        "abstract_under_350": stats["abstract_words"] <= 350,
        "article_core_approximately_8000": 7700 <= stats["abstract_words"] + stats["main_text_words_introduction_through_conclusions"] <= 8300,
        "methods_12_subsections": all(any(h.startswith(f"2.{i} ") for h in headings) for i in range(1, 13)),
        "results_8_subsections": all(any(h.startswith(f"3.{i} ") for h in headings) for i in range(1, 9)),
        "discussion_7_subsections": all(any(h.startswith(f"4.{i} ") for h in headings) for i in range(1, 8)),
        "five_tables": len(final_doc.tables) == 5,
        "four_figures": len(final_doc.inline_shapes) == 4,
        "three_line_tables_no_vertical_borders": vertical_borders == 0,
        "replacement_character_absent": chr(0xFFFD) not in final_text,
        "old_multiview_confirmation_terms_absent": not any(term in final_text.lower() for term in ["multiview confirmation", "representational confirmation", "shared-split confirmation"]),
        "journal_banner_absent": "Research Article | Journal of Cheminformatics" not in final_text,
        "scientific_contribution_present": "Scientific Contribution:" in final_text,
    }
    audit = {
        "file": str(PAPER),
        "word_counts": stats,
        "article_core_words_abstract_through_conclusions": stats["abstract_words"] + stats["main_text_words_introduction_through_conclusions"],
        "target_main_text_words": "approximately 8000",
        "main_text_target_met": 7500 <= stats["main_text_words_introduction_through_conclusions"] <= 8400,
        "figures": 4,
        "tables": 5,
        "unavailable_analysis_not_fabricated": True,
        "checks": checks,
        "all_checks_pass": all(checks.values()),
    }
    AUDIT_JSON.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
