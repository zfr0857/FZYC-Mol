from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

import fitz
from docx import Document
from openpyxl import load_workbook


ROOT = Path(r"D:\fzyc\output\paper33_final_minor_revision_20260718")


def document_checks(path: Path, language: str) -> dict:
    doc = Document(path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    abstract = [p.text.split(":", 1)[1] for p in doc.paragraphs if p.text.startswith(("Background:", "Methods:", "Results:", "Conclusions:", "Scientific Contribution:"))]
    return {
        "tables": len(doc.tables),
        "figures": len(doc.inline_shapes),
        "native_equations": len(re.findall(r"<(?:m|ns\d+):oMath(?:\s|>)", xml)),
        "equation_numbers": all(f">({number})<" in xml for number in range(1, 15)),
        "old_270_absent": "270 outer" not in all_text and "270个" not in all_text,
        "old_design_absent": "3 endpoints; K = 16 and 32" not in all_text and "3个端点；K = 16" not in all_text,
        "new_exposure_present": "1,080" in all_text and "64,616.35" in all_text,
        "supplement_numbering_present": "S1–S36" in all_text and "S1–S21" in all_text,
        "abstract_words": len(re.findall(r"\b[\w’'-]+\b", " ".join(abstract))) if language == "en" else None,
        "pi_selection_frequency": "π" in xml,
    }


def main() -> None:
    en = ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_FINAL_CLEAN.docx"
    track = ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_FINAL_TRACK_CHANGES.docx"
    zh = ROOT / "候选池扩张与模型选择损失_中文终稿.docx"
    required = {
        "English clean DOCX": en,
        "English track-changes DOCX": track,
        "Chinese synchronized DOCX": zh,
        "Updated Abstract": ROOT / "revision_extracts" / "Updated_Abstract.docx",
        "Updated Methods": ROOT / "revision_extracts" / "Updated_Methods_2.1_2.3_2.14_2.16.docx",
        "Updated Results 3.10": ROOT / "revision_extracts" / "Updated_Results_3.10.docx",
        "Updated Discussion/Conclusions": ROOT / "revision_extracts" / "Updated_Discussion_4.8_and_Conclusions.docx",
        "Updated Table 2": ROOT / "revision_extracts" / "Updated_Table2.docx",
        "Figure 7 PDF": ROOT / "main_figures" / "Figure7.pdf",
        "Figure 7 SVG": ROOT / "main_figures" / "Figure7.svg",
        "Figure 7 PNG": ROOT / "main_figures" / "Figure7_600dpi.png",
        "Figure 7 QC": ROOT / "Figure7_QC_report.json",
        "Equation mapping": ROOT / "Equation_to_code_mapping.xlsx",
        "Supplement numbering": ROOT / "Supplementary_numbering_report.json",
        "Declarations checklist": ROOT / "reports" / "Declarations_completion_checklist.xlsx",
        "Bilingual consistency": ROOT / "reports" / "English_Chinese_consistency_report.xlsx",
        "Reviewer response matrix": ROOT / "reports" / "Reviewer_concern_Response_Revision_location.xlsx",
        "Supplementary Methods": ROOT / "supplementary" / "Additional_file_1_Supplementary_Methods_and_Results.docx",
        "Supplementary Tables": ROOT / "supplementary" / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S36.xlsx",
        "Supplementary Figures": ROOT / "supplementary" / "Additional_file_3_Supplementary_Figures_S1-S21.pdf",
        "Master consistency table": ROOT / "reports" / "Master_result_consistency_table.xlsx",
    }
    required_status = {name: path.exists() and path.stat().st_size > 0 for name, path in required.items()}

    english = document_checks(en, "en")
    chinese = document_checks(zh, "zh")
    with zipfile.ZipFile(track) as archive:
        track_xml = archive.read("word/document.xml").decode("utf-8")
    track_status = {
        "insertions": track_xml.count("<w:ins"),
        "deletions": track_xml.count("<w:del"),
        "has_revisions": "<w:ins" in track_xml and "<w:del" in track_xml,
        "native_equations": track_xml.count("<m:oMath>"),
    }

    workbook = load_workbook(required["Supplementary Tables"], read_only=True, data_only=False)
    supplementary_tables = {"sheet_count": len(workbook.sheetnames), "first": workbook.sheetnames[0], "last": workbook.sheetnames[-1]}
    workbook.close()
    pdf = fitz.open(required["Supplementary Figures"])
    supplementary_figures = {"page_count": pdf.page_count}
    pdf.close()

    figure_qc = json.loads(required["Figure 7 QC"].read_text(encoding="utf-8"))
    font_audit = json.loads((ROOT / "Figure_1-7_font_audit.json").read_text(encoding="utf-8"))
    bilingual_rows = load_workbook(required["Bilingual consistency"], read_only=True, data_only=True)["Bilingual consistency"]
    bilingual_review = [row[0].value for row in bilingual_rows.iter_rows(min_row=2) if row[3].value != "pass"]

    technical_pass = (
        all(required_status.values())
        and english["tables"] == chinese["tables"] == 4
        and english["figures"] == chinese["figures"] == 7
        and english["native_equations"] == chinese["native_equations"] == 14
        and english["equation_numbers"] and chinese["equation_numbers"]
        and english["old_270_absent"] and chinese["old_270_absent"]
        and english["old_design_absent"] and chinese["old_design_absent"]
        and english["abstract_words"] <= 350
        and track_status["has_revisions"]
        and supplementary_tables["sheet_count"] == 43
        and supplementary_figures["page_count"] == 21
        and figure_qc["overall_status"] == "pass"
        and font_audit["status"] == "passed"
        and not bilingual_review
    )
    report = {
        "technical_status": "pass" if technical_pass else "fail",
        "submission_status": "author confirmation required",
        "blocking_author_fields": ["Competing interests", "Funding", "Authors' contributions", "Acknowledgements"],
        "required_outputs": required_status,
        "english_manuscript": english,
        "chinese_manuscript": chinese,
        "track_changes": track_status,
        "supplementary_tables": supplementary_tables,
        "supplementary_figures": supplementary_figures,
        "figure7_qc": figure_qc["overall_status"],
        "figure_1_7_font_audit": font_audit["status"],
        "bilingual_review_items": bilingual_review,
        "non_inference_statement": "No author, funding, competing-interest, acknowledgement, repository or unreconstructed fit-count information was inferred.",
    }
    (ROOT / "Final_submission_validation_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        f"Technical status: {report['technical_status'].upper()}",
        "Submission status: AUTHOR CONFIRMATION REQUIRED",
        "Blocking author fields: " + "; ".join(report["blocking_author_fields"]),
        report["non_inference_statement"],
    ]
    (ROOT / "Final_submission_validation_report.txt").write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
