from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def source_docx() -> Path:
    base = Path.home() / "Desktop" / "\u4fee\u6539"
    preferred = base / "FZYC-Mol_\u521d\u7a3f4_\u8865\u5145\u5b9e\u9a8c\u5168\u8dd1\u4fee\u8ba2.docx"
    if preferred.exists():
        return preferred
    candidates = sorted(base.glob("FZYC-Mol*\u5168\u8dd1\u4fee\u8ba2.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not candidates:
        raise FileNotFoundError("Cannot find the full-run revised manuscript.")
    return candidates[0]


def target_docx(src: Path) -> Path:
    return src.with_name("FZYC-Mol_\u521d\u7a3f4_\u8865\u5145\u5b9e\u9a8c\u5168\u8dd1\u516c\u5f0f\u4fee\u8ba2.docx")


def set_run_font(run, size: float = 10, bold: bool | None = None, font: str = "Times New Roman") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style_paragraph(p, size: float = 10, bold: bool = False, font: str = "Times New Roman") -> None:
    for run in p.runs:
        set_run_font(run, size=size, bold=bold, font=font)


def insert_text(anchor, text: str, size: float = 10, bold: bool = False, style: str | None = None, align=None):
    p = anchor.insert_paragraph_before(text)
    if style:
        try:
            p.style = style
        except Exception:
            pass
    if align is not None:
        p.alignment = align
    style_paragraph(p, size=size, bold=bold)
    return p


def insert_formula(anchor, text: str, number: int) -> None:
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{text}    ({number})")
    set_run_font(run, size=9.5, font="Cambria Math")
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)


def replace_paragraph(doc: Document, old_start: str, new_text: str) -> None:
    for p in doc.paragraphs:
        if p.text.strip().startswith(old_start):
            p.text = new_text
            style_paragraph(p, size=10)
            return


def add_formula_references(doc: Document) -> None:
    replace_paragraph(
        doc,
        "\u9009\u62e9\u5668\u9996\u5148\u5728\u6bcf\u4e2a\u7ec8\u70b9\u7684\u9a8c\u8bc1\u96c6\u9884\u6d4b\u8868\u4e0a\u8ba1\u7b97\u5b98\u65b9\u4e3b\u6307\u6807",
        "\u9009\u62e9\u5668\u9996\u5148\u5728\u6bcf\u4e2a\u7ec8\u70b9\u7684\u9a8c\u8bc1\u96c6\u9884\u6d4b\u8868\u4e0a\u8ba1\u7b97\u5b98\u65b9\u4e3b\u6307\u6807\uff0c\u5e76\u6309\u516c\u5f0f (1) \u5c06\u6240\u6709\u6307\u6807\u8f6c\u4e3a\u7edf\u4e00\u7684\u201c\u8d8a\u5927\u8d8a\u597d\u201d\u5f97\u5206\uff0c\u518d\u7528\u516c\u5f0f (2) \u51bb\u7ed3\u6700\u7ec8\u5019\u9009\u3002\u82e5\u5355\u4e00\u4e13\u5bb6\u663e\u8457\u5360\u4f18\uff0c\u5219\u4fdd\u7559\u6700\u4f73\u4e13\u5bb6\uff1b\u82e5\u591a\u4e2a\u4e13\u5bb6\u5728\u9a8c\u8bc1\u96c6\u4e0a\u4e92\u8865\uff0c\u5219\u5019\u9009 Top-K \u5747\u503c\u3001\u5806\u53e0\u96c6\u6210\u6216\u4e0d\u786e\u5b9a\u6027\u52a0\u6743\u878d\u5408\uff0c\u5176\u9884\u6d4b\u5c42\u878d\u5408\u5bf9\u5e94\u516c\u5f0f (5)-(6)\u3002\u5019\u9009\u6c60\u3001\u4e3b\u6307\u6807\u3001\u5bb9\u5dee\u548c\u5e73\u5c40\u89c4\u5219\u5728\u8bfb\u53d6\u6d4b\u8bd5\u96c6\u524d\u56fa\u5b9a\uff1b\u672a\u8d85\u8fc7\u5f53\u524d\u6700\u7ec8\u4fdd\u7559\u7ed3\u679c\u7684\u8865\u6551\u5934\u3001\u76ee\u6807\u53d8\u6362\u6216\u5f3a\u57fa\u7ebf\u8fdb\u5165\u9644\u5f55\u800c\u4e0d\u6539\u53d8\u4e3b\u7ed3\u679c\u3002\u8be5\u89c4\u5219\u4f7f\u6b63\u7ed3\u679c\u3001\u8d1f\u7ed3\u679c\u548c\u672a\u63a5\u5165\u5019\u9009\u90fd\u6709\u660e\u786e\u4f4d\u7f6e\uff0c\u4e5f\u964d\u4f4e\u4e86\u6d4b\u8bd5\u96c6\u4e8b\u540e\u9009\u62e9\u504f\u5dee\u3002",
    )
    replace_paragraph(
        doc,
        "\u4e3a\u5ba1\u8ba1\u9a8c\u8bc1\u96c6\u9009\u62e9\u662f\u5426\u53ef\u80fd\u8fc7\u62df\u5408",
        "\u4e3a\u5ba1\u8ba1\u9a8c\u8bc1\u96c6\u9009\u62e9\u662f\u5426\u53ef\u80fd\u8fc7\u62df\u5408\uff0c\u672c\u6587\u989d\u5916\u8ba1\u7b97\u6bcf\u4e2a dataset-seed \u5019\u9009\u6c60\u4e2d\u9a8c\u8bc1\u96c6\u6392\u540d\u4e0e\u6d4b\u8bd5\u96c6\u6392\u540d\u7684 Spearman \u76f8\u5173\u6027\uff0c\u5e76\u8bb0\u5f55\u9a8c\u8bc1\u96c6\u7b2c\u4e00\u540d\u662f\u5426\u4e5f\u662f\u6d4b\u8bd5\u96c6\u7b2c\u4e00\u540d\u3001\u6d4b\u8bd5\u96c6\u7b2c\u4e00\u540d\u662f\u5426\u843d\u5165\u9a8c\u8bc1\u96c6 Top-3\uff0c\u4ee5\u53ca\u516c\u5f0f (3)-(4) \u5b9a\u4e49\u7684 regret/optimism gap\u3002\u8be5\u5206\u6790\u4e0d\u53c2\u4e0e\u6a21\u578b\u9009\u62e9\uff0c\u53ea\u7528\u4e8e\u8bc4\u4f30\u9009\u62e9\u5668\u98ce\u9669\uff1a\u82e5\u9a8c\u8bc1-\u6d4b\u8bd5\u6392\u540d\u76f8\u5173\u6027\u8f83\u4f4e\u3001Top-3 \u547d\u4e2d\u4e0d\u8db3\u6216\u51fa\u73b0\u8d1f\u76f8\u5173\uff0c\u76f8\u5173\u7ed3\u679c\u5728\u6b63\u6587\u4e2d\u6309\u8fb9\u754c\u6216\u8d1f\u7ed3\u679c\u89e3\u91ca\uff0c\u800c\u4e0d\u4f5c\u4e3a\u65e0\u6761\u4ef6\u6027\u80fd\u63d0\u5347\u8bc1\u636e\u3002",
    )
    replace_paragraph(
        doc,
        "\u53ef\u9760\u6027\u6a21\u5757\u5305\u62ec\u96c6\u6210\u6807\u51c6\u5dee",
        "\u53ef\u9760\u6027\u6a21\u5757\u5305\u62ec\u96c6\u6210\u6807\u51c6\u5dee\u3001\u9519\u8bef\u6a21\u578b\u3001\u9884\u6d4b\u504f\u5dee\u3001\u53cd\u5411 Tanimoto \u8ddd\u79bb\u3001\u91cd\u6784\u8bef\u5dee\u3001\u9519\u8bef-\u9002\u7528\u57df\u6df7\u5408\u6307\u6807\u3001\u98ce\u9669-\u8986\u76d6\u66f2\u7ebf\u3001\u4fdd\u5f62\u8986\u76d6\u7387\u3001\u6821\u51c6\u548c\u7c97\u7cd9\u5ea6\u4ee3\u7406\u6307\u6807\u3002\u5176\u4e2d Tanimoto \u6700\u8fd1\u90bb\u76f8\u4f3c\u5ea6\u3001\u96c6\u6210\u4e0d\u786e\u5b9a\u6027\u3001\u9519\u8bef\u6807\u7b7e\u3001\u6df7\u5408\u98ce\u9669\u548c\u5bcc\u96c6\u5ea6\u5206\u522b\u7531\u516c\u5f0f (7)-(12) \u5b9a\u4e49\u3002\u5206\u7c7b\u4efb\u52a1\u628a\u9884\u6d4b\u6807\u7b7e\u9519\u8bef\u4f5c\u4e3a\u9519\u8bef\u68c0\u51fa\u76ee\u6807\uff0c\u56de\u5f52\u4efb\u52a1\u628a\u7edd\u5bf9\u8bef\u5dee\u6700\u9ad8\u7684 20% \u6837\u672c\u4f5c\u4e3a\u9ad8\u8bef\u5dee\u76ee\u6807\uff0c\u5e76\u7528 AUROC\u3001AUPRC\u3001\u98ce\u9669-\u8986\u76d6\u66f2\u7ebf\u548c top-10% \u5bcc\u96c6\u5171\u540c\u8bc4\u4f30\u98ce\u9669\u5206\u6570\u3002\u53ef\u89e3\u91ca\u6027\u6a21\u5757\u5305\u62ec\u57fa\u5e8f\u7279\u5f81\u91cd\u8981\u6027\u3001\u7247\u6bb5\u5bcc\u96c6\u3001\u9aa8\u67b6/\u8fd1\u90bb\u6848\u4f8b\u590d\u6838\u4e0e\u9ad8\u8bef\u5dee\u6837\u672c\u5206\u6790\u3002\u672c\u6587\u53ea\u628a\u8fd9\u4e9b\u8bc1\u636e\u89e3\u91ca\u4e3a\u6a21\u578b\u4f7f\u7528\u8fb9\u754c\u548c\u5316\u5b66\u76f8\u5173\u6027\uff0c\u4e0d\u628a\u5173\u8054\u6027\u57fa\u5e8f\u89e3\u91ca\u7b49\u540c\u4e8e\u56e0\u679c\u673a\u5236\u3002",
    )
    replace_paragraph(
        doc,
        "\u5206\u7c7b\u4efb\u52a1\u4fdd\u7559 ROC-AUC \u4f5c\u4e3a MoleculeNet",
        "\u5206\u7c7b\u4efb\u52a1\u4fdd\u7559 ROC-AUC \u4f5c\u4e3a MoleculeNet \u548c\u82e5\u5e72 TDC \u7ec8\u70b9\u7684\u4e3b\u6307\u6807\uff0c\u540c\u65f6\u62a5\u544a PR-AUC\u3001Brier \u5206\u6570\u3001ECE\u3001MCC\u3001\u5e73\u8861\u51c6\u786e\u7387\u3001\u98ce\u9669-\u8986\u76d6\u66f2\u7ebf\u548c\u4fdd\u5f62\u8986\u76d6\u7387\u3002Brier \u548c ECE \u7684\u5b9a\u4e49\u89c1\u516c\u5f0f (13)\uff1b\u8fd9\u79cd\u5199\u6cd5\u662f\u4e3a\u4e86\u56de\u5e94\u7c7b\u522b\u4e0d\u5e73\u8861\u4efb\u52a1\u4e2d\u7684\u5e38\u89c1\u5ba1\u7a3f\u8d28\u7591\uff1aROC-AUC \u8f83\u9ad8\u5e76\u4e0d\u5fc5\u7136\u610f\u5473\u7740\u9633\u6027\u6837\u672c\u53ec\u56de\u3001\u6982\u7387\u6821\u51c6\u6216\u7b5b\u9009\u5bcc\u96c6\u53ef\u9760\u3002ClinTox\u3001DILI\u3001hERG \u548c CYP \u5e95\u7269\u4efb\u52a1\u56e0\u6b64\u5728\u4e0d\u5e73\u8861\u5206\u7c7b\u589e\u5f3a\u4e2d\u91cd\u70b9\u62a5\u544a\u3002",
    )
    replace_paragraph(
        doc,
        "\u4fdd\u5f62\u9884\u6d4b\u5728 80%/90%/95%",
        "\u4fdd\u5f62\u9884\u6d4b\u5728 80%/90%/95% \u76ee\u6807\u4e0b\u7684\u5206\u7c7b\u9884\u6d4b\u96c6\u4e0e\u56de\u5f52\u9884\u6d4b\u533a\u95f4\u5206\u522b\u6309\u516c\u5f0f (14) \u548c\u516c\u5f0f (15) \u8ba1\u7b97\u3002\u6b64\u5904\u8986\u76d6\u7387\u53ea\u89e3\u91ca\u4e3a\u7ed9\u5b9a\u5212\u5206\u548c\u6821\u51c6\u96c6\u4e0b\u7684\u6709\u9650\u6837\u672c\u4fdd\u8bc1\uff0c\u4e0d\u4ee3\u8868\u8de8\u5b9e\u9a8c\u6761\u4ef6\u3001\u8de8\u5b9e\u9a8c\u5ba4\u6216\u8de8\u5316\u5b66\u7a7a\u95f4\u7684\u65e0\u6761\u4ef6\u4fdd\u8bc1\u3002",
    )


def insert_formula_section(doc: Document) -> None:
    anchor = next((p for p in doc.paragraphs if p.text.strip() == "\u7ed3\u679c"), None)
    if anchor is None:
        raise RuntimeError("Cannot find Results heading anchor.")

    insert_text(anchor, "3.7 \u6570\u5b66\u5b9a\u4e49\u4e0e\u8bc4\u4ef7\u516c\u5f0f", size=11, bold=True)
    insert_text(
        anchor,
        "\u4e3a\u4f7f\u9a8c\u8bc1\u96c6\u9009\u62e9\u3001\u878d\u5408\u3001\u9002\u7528\u57df\u3001\u4e0d\u786e\u5b9a\u6027\u548c\u8865\u5145\u5b9e\u9a8c\u7684\u5b9a\u4e49\u53ef\u590d\u6838\uff0c\u672c\u6587\u5c06\u6838\u5fc3\u8ba1\u7b97\u7edf\u4e00\u8bb0\u4e3a\u4ee5\u4e0b\u516c\u5f0f\u3002\u8bbe\u7b2c t \u4e2a\u7ec8\u70b9\u7684\u6570\u636e\u4e3a D_t = D_t^train \u222a D_t^val \u222a D_t^test\uff0c\u5019\u9009\u4e13\u5bb6\u96c6\u4e3a A_t\uff0c\u4e13\u5bb6 a \u5bf9\u5206\u5b50 x_i \u7684\u56de\u5f52\u8f93\u51fa\u6216\u5206\u7c7b\u9633\u6027\u6982\u7387\u7edf\u4e00\u8bb0\u4e3a z_hat_{i,a}\u3002\u5bf9\u4e8e\u4f4e\u503c\u4f18\u5148\u7684 RMSE/MAE\uff0c\u5148\u53d6\u8d1f\u503c\u8f6c\u4e3a\u7edf\u4e00\u7684\u6b63\u5411\u5f97\u5206 Q\u3002",
        size=10,
    )

    formulas = [
        ("Q_t(a) = M_t(y^val, z_hat_a^val), if M is higher-better;  Q_t(a) = -M_t(y^val, z_hat_a^val), if M is lower-better", "\u7edf\u4e00\u6307\u6807\u65b9\u5411"),
        ("a*_t = argmax_{a in A_t} Q_t(a)", "\u9a8c\u8bc1\u96c6\u9009\u62e9\u5668"),
        ("Regret_t = Q_t^test(a^dagger_t) - Q_t^test(a*_t),  a^dagger_t = argmax_{a in A_t} Q_t^test(a)", "\u6d4b\u8bd5 oracle regret"),
        ("OptGap_t = Q_t^val(a*_t) - Q_t^test(a*_t)", "optimism gap"),
        ("T_K(t) = top-K candidates ranked by Q_t(a);  z_hat_i^mean = K^{-1} sum_{a in T_K(t)} z_hat_{i,a}", "Top-K \u5747\u503c\u878d\u5408"),
        ("w_{i,a} = exp(Q_t(a)/tau) (u_{i,a}+epsilon)^{-1} / sum_{b in T_K(t)} exp(Q_t(b)/tau) (u_{i,b}+epsilon)^{-1};  z_hat_i^uw = sum_a w_{i,a} z_hat_{i,a}", "\u4e0d\u786e\u5b9a\u6027\u52a0\u6743\u878d\u5408"),
        ("T(m_i,m_j) = |m_i AND m_j| / |m_i OR m_j|;  s_i^NN = max_{j in D^train union D^val} T(m_i,m_j);  d_i^AD = 1 - s_i^NN", "Morgan Tanimoto \u9002\u7528\u57df"),
        ("g_i = I(s_i^NN >= delta_AD);  z_hat_i^AD = g_i z_hat_i^fusion + (1-g_i) z_hat_i^safe", "\u9002\u7528\u57df\u95e8\u63a7"),
        ("u_i^ens = sqrt((K-1)^{-1} sum_{a in T_K(t)} (z_hat_{i,a} - mean_a z_hat_{i,a})^2)", "\u96c6\u6210\u4e0d\u786e\u5b9a\u6027"),
        ("e_i = I(y_hat_i != y_i) for classification;  e_i = I(|y_i-z_hat_i| >= Quantile_0.80(|y-z_hat| on validation)) for regression", "\u9519\u8bef/\u9ad8\u8bef\u5dee\u6807\u7b7e"),
        ("r_i = lambda p_i^err + (1-lambda) d_i^AD,  0 <= lambda <= 1", "\u9519\u8bef-\u9002\u7528\u57df\u6df7\u5408\u98ce\u9669"),
        ("EF_q = P(e_i=1 | r_i in top q%) / P(e_i=1)", "\u98ce\u9669\u5bcc\u96c6\u5ea6"),
        ("Brier = n^{-1} sum_i (p_i-y_i)^2;  ECE = sum_{b=1}^B |B_b|/n |acc(B_b)-conf(B_b)|", "\u6982\u7387\u6821\u51c6"),
        ("S_i^cls = 1 - p_hat_{i,y_i};  q_alpha = Quantile_ceil((n_cal+1)(1-alpha))/n_cal({S_i^cls});  C_alpha(x) = {c: 1-p_hat_c(x) <= q_alpha}", "\u5206\u7c7b\u4fdd\u5f62\u9884\u6d4b\u96c6"),
        ("S_i^reg = |y_i-z_hat_i|;  q_alpha = Quantile_ceil((n_cal+1)(1-alpha))/n_cal({S_i^reg});  I_alpha(x) = [z_hat(x)-q_alpha, z_hat(x)+q_alpha]", "\u56de\u5f52\u4fdd\u5f62\u9884\u6d4b\u533a\u95f4"),
        ("bin_i = >0.7 if s_i^NN>0.7;  bin_i = 0.5-0.7 if 0.5<=s_i^NN<=0.7;  bin_i = <0.5 if s_i^NN<0.5", "Tanimoto \u4e09\u6863\u5206\u5c42"),
        ("Delta Q_{t,c} = Q_t(c) - Q_t(Full)", "\u7edf\u4e00\u6d88\u878d\u5dee\u503c"),
        ("Delta y_{ij}=|y_i-y_j|, Delta z_hat_{ij}=|z_hat_i-z_hat_j|;  cliff(i,j)=I(T(m_i,m_j)>=theta_s and Delta y_{ij}>=theta_y)", "MoleculeACE \u6d3b\u6027\u60ac\u5d16"),
        ("rho_gap = Spearman({Delta y_{ij}}, {Delta z_hat_{ij}});  Acc_dir = n^{-1} sum_{(i,j)} I(sign(y_i-y_j)=sign(z_hat_i-z_hat_j))", "\u6d3b\u6027\u5dee\u5f02\u76f8\u5173\u4e0e\u65b9\u5411\u51c6\u786e\u7387"),
        ("RMSE = sqrt(n^{-1} sum_i (y_i-z_hat_i)^2);  MAE = n^{-1} sum_i |y_i-z_hat_i|", "\u56de\u5f52\u8bef\u5dee"),
    ]

    for idx, (formula, label) in enumerate(formulas, start=1):
        insert_text(anchor, f"\u516c\u5f0f ({idx}) {label}\uff1a", size=9, bold=True)
        insert_formula(anchor, formula, idx)

    insert_text(
        anchor,
        "\u5728\u4e0a\u8ff0\u5b9a\u4e49\u4e2d\uff0cz_hat_i^safe \u8868\u793a\u9a8c\u8bc1\u96c6\u9884\u5148\u786e\u5b9a\u7684\u4fdd\u5b88\u5019\u9009\uff08\u5982\u7b80\u5355\u5171\u8bc6\u3001\u57fa\u7ebf\u6a21\u578b\u6216\u6700\u8fd1\u90bb\u8bca\u65ad\u8f93\u51fa\uff09\uff1btau\u3001epsilon\u3001delta_AD\u3001lambda\u3001theta_s \u548c theta_y \u5747\u4ec5\u5728\u8bad\u7ec3/\u9a8c\u8bc1\u5c42\u9762\u786e\u5b9a\u3002\u6b63\u6587\u6240\u6709\u6d88\u878d\u5dee\u503c\u3001\u4fdd\u5f62\u8986\u76d6\u3001\u4f4e\u76f8\u4f3c\u5ea6\u5206\u5c42\u548c MoleculeACE \u5dee\u5f02\u76f8\u5173\u7ed3\u679c\u5747\u6309\u8fd9\u4e9b\u5b9a\u4e49\u8ba1\u7b97\u3002",
        size=10,
    )


def main() -> None:
    src = source_docx()
    out = target_docx(src)
    shutil.copy2(src, out)
    doc = Document(out)
    add_formula_references(doc)
    insert_formula_section(doc)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
