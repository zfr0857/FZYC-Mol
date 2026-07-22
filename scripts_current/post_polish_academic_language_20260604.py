from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_academic_polished.docx"
MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_academic_polished.md"


REPLACEMENTS = [
    ("表 14. Targeted improvement case studies。", "表 14. 定向改进案例。"),
    ("表 2 和图 2 显示，FZYC-Mol 的优势不是某一个专家在所有任务上稳定优于，而是", "表 2 和图 2 显示，FZYC-Mol 的优势并非来自某一专家在所有任务上的稳定优势，而是"),
    ("表 6 支持一个重要策略判断：", "表 6 表明，"),
    ("Lipophilicity 是本轮最具代表性的性能瓶颈模块修复案例。", "Lipophilicity 是最具代表性的性能瓶颈模块修复案例。"),
    ("为避免继续无边界堆叠 candidate，本轮将 selector 改进正式接入结果部分。", "为避免无边界扩展候选模型，本文将 selector 改进纳入正式结果部分。"),
    ("与 formal fixed-selector integration 放在一起解读时，3D-lite/roughness 的负结果反而加强了论文结论：", "结合固定选择器策略结果可见，3D-lite/roughness 的负结果进一步支持本文结论："),
    ("这样的写法可以避免实验结果显得零散：主文先回答 FZYC-Mol 是否在标准分子性质预测任务上有效，再回答这种有效性是否能在更真实的 ADMET 外推、不平衡分类和低相似度样本中保持，最后用负结果与案例分析说明哪些模块仍然是后续工作瓶颈。", "该组织方式使结果呈现形成清晰逻辑：首先评估 FZYC-Mol 在标准分子性质预测任务中的有效性，随后检验其在 ADMET 外推、不平衡分类和低相似度样本中的稳健性，最后通过负结果与案例分析界定后续改进空间。"),
    ("注：该表用于连接主文、附录和后续 回应材料 叙事；所有性能增强均遵守 validation-only selector 原则，未接入候选作为负结果和限制保留。", "注：该表用于连接主文与补充材料；未接入候选作为负结果或限制报告。"),
    ("本轮 pilot 的结果不建议覆盖主结果。MoleculeNet strong pilot 在 FreeSolv、Lipo 和 ClinTox 上分别选择 stacking 或 rank-fusion 候选，但均未超过已有 retained-best；因此它们更适合作为附录证据，说明本文确实比较了更强候选，但不会为了追求表面涨分而违反 retained-best gate。freesolv: retained/reference=1.0286, pilot=1.4417, decision=retain previous best；lipo: retained/reference=0.6835, pilot=0.7901, decision=retain previous best；clintox: retained/reference=0.9489, pilot=0.9186, decision=retain previous best。", "Pilot 结果未触发主结果更新。MoleculeNet strong pilot 在 FreeSolv、Lipo 和 ClinTox 上分别选择 stacking 或 rank-fusion 候选，但其性能未超过已有 retained-best，因此仅作为附录证据报告。FreeSolv、Lipo 和 ClinTox 的 pilot 结果分别为 1.4417、0.7901 和 0.9186，对应 retained/reference 为 1.0286、0.6835 和 0.9489。"),
    ("后续实验可进一步，最值得做的不是进行大规模模型重训练，而是外部 benchmark/appendix。", "后续实验应优先扩展外部 benchmark/appendix，而非首先进行大规模模型重训练。"),
    ("最值得继续做的是", "优先方向是"),
    ("反而需要", "通常需要"),
    ("人为选择有利的测试集结果", "基于测试集结果进行事后选择"),
    ("本轮 pilot", "该 pilot"),
    ("本轮", "该实验阶段"),
]


def apply_replacements(text: str) -> str:
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text


def set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = text


def main() -> None:
    doc = Document(DOCX)
    for paragraph in doc.paragraphs:
        cleaned = apply_replacements(paragraph.text)
        if cleaned != paragraph.text:
            set_paragraph_text(paragraph, cleaned)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    cleaned = apply_replacements(paragraph.text)
                    if cleaned != paragraph.text:
                        set_paragraph_text(paragraph, cleaned)
    doc.save(DOCX)

    text = MD.read_text(encoding="utf-8")
    MD.write_text(apply_replacements(text), encoding="utf-8")
    print(f"polished {DOCX.relative_to(ROOT)}")
    print(f"polished {MD.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
