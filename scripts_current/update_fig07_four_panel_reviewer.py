from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches


ROOT = Path("D:/fzyc")
OUT = ROOT / "output"
SRC = ROOT / "results" / "source_data" / "tdc_gate_audit.csv"
FIG_DIR = OUT / "paper18_fig07_four_panel_reviewer"
AUDIT = OUT / "paper18_fig07_four_panel_reviewer_audit.json"

BLUE = "#3A6EA5"
GREY = "#8E97A1"
LIGHT = "#D9DFE7"
RED = "#B86B5E"
DARK = "#1F2937"
GREEN = "#4C9A74"


CAT_LABELS = {
    "promoted_and_improved": "Promoted + improved",
    "retained_and_avoided_harm": "Retained + avoided harm",
    "inconclusive_due_to_wide_ci": "Inconclusive\n(seed interval)",
}

CAT_ORDER = [
    "promoted_and_improved",
    "retained_and_avoided_harm",
    "inconclusive_due_to_wide_ci",
]


def docx_path() -> Path:
    matches = [
        p
        for p in OUT.glob("*.docx")
        if p.name.endswith("-18.docx") and not p.name.startswith("~$")
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one 小论文-18.docx, found {matches}")
    return matches[0]


def setup_plot() -> None:
    plt.rcParams.update(
        {
            "font.family": "Arial",
            "font.size": 11,
            "axes.titlesize": 13,
            "axes.labelsize": 12,
            "xtick.labelsize": 10,
            "ytick.labelsize": 9,
            "legend.fontsize": 9,
            "axes.linewidth": 0.9,
            "savefig.dpi": 600,
            "svg.fonttype": "none",
        }
    )


def clean_label(name: str) -> str:
    return (
        name.replace("_carbonmangels", "")
        .replace("_az", " AZ")
        .replace("_astrazeneca", " AZ")
        .replace("_lombardo", " Lombardo")
        .replace("_obach", " Obach")
        .replace("_broccatelli", " Broccatelli")
        .replace("_martins", " Martins")
        .replace("_veith", " Veith")
        .replace("_wang", " Wang")
        .replace("_hou", " Hou")
        .replace("_zhu", " Zhu")
        .replace("_ma", " Ma")
        .replace("_aqsoldb", " AqSolDB")
        .replace("_", " ")
    )


def load_data() -> pd.DataFrame:
    df = pd.read_csv(SRC)
    df["category_label"] = df["gate_category"].map(CAT_LABELS)
    df["endpoint_label"] = df["endpoint"].map(clean_label)
    df["interval_low"] = df["ci_low"].astype(float)
    df["interval_high"] = df["ci_high"].astype(float)
    df["interval_half_width"] = (df["interval_high"] - df["interval_low"]) / 2
    df["interval_crosses_zero"] = (df["interval_low"] <= 0) & (df["interval_high"] >= 0)
    df["abs_delta"] = df["test_delta"].abs()
    df["seed_positive_fraction"] = df["seed_win_count"] / df["n_paired_seeds"]
    df["seed_negative_fraction"] = df["seed_loss_count"] / df["n_paired_seeds"]
    df["seed_tie_fraction"] = df["seed_tie_count"] / df["n_paired_seeds"]
    return df


def category_counts(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for cat in CAT_ORDER:
        g = df[df["gate_category"].eq(cat)]
        rows.append(
            {
                "gate_category": cat,
                "category_label": CAT_LABELS[cat],
                "retained": int((~g["promoted"]).sum()),
                "promoted": int(g["promoted"].sum()),
                "total": int(len(g)),
            }
        )
    return pd.DataFrame(rows)


def seed_summary(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for cat in CAT_ORDER:
        g = df[df["gate_category"].eq(cat)]
        rows.append(
            {
                "gate_category": cat,
                "category_label": CAT_LABELS[cat],
                "positive_seed_delta": int(g["seed_win_count"].sum()),
                "tie_seed_delta": int(g["seed_tie_count"].sum()),
                "negative_seed_delta": int(g["seed_loss_count"].sum()),
                "total_seed_summaries": int(g["n_paired_seeds"].sum()),
            }
        )
    out = pd.DataFrame(rows)
    for col in ["positive_seed_delta", "tie_seed_delta", "negative_seed_delta"]:
        out[col.replace("_delta", "_fraction")] = out[col] / out["total_seed_summaries"]
    return out


def draw_figure(df: pd.DataFrame) -> tuple[Path, Path, dict[str, float]]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    counts = category_counts(df)
    seeds = seed_summary(df)
    source = df.merge(counts[["gate_category", "retained", "promoted", "total"]], on="gate_category")
    source = source.merge(
        seeds[
            [
                "gate_category",
                "positive_seed_delta",
                "tie_seed_delta",
                "negative_seed_delta",
                "total_seed_summaries",
            ]
        ],
        on="gate_category",
    )
    source.to_csv(FIG_DIR / "fig07_tdc_gate_four_panel_reviewer_source.csv", index=False)

    fig, axes = plt.subplots(2, 2, figsize=(12.5, 8.9))
    ax1, ax2, ax3, ax4 = axes.ravel()

    labels = counts["category_label"].tolist()
    y = np.arange(len(counts))
    ax1.barh(y, counts["retained"], color=GREY, edgecolor=DARK, linewidth=0.4, label="Retained")
    ax1.barh(
        y,
        counts["promoted"],
        left=counts["retained"],
        color=BLUE,
        edgecolor=DARK,
        linewidth=0.4,
        label="Promoted",
    )
    for idx, row in counts.iterrows():
        ax1.text(row["total"] + 0.25, idx, str(row["total"]), va="center", color="#64748B")
    ax1.set_yticks(y, labels)
    ax1.invert_yaxis()
    ax1.set_xlabel("Endpoints")
    ax1.set_title("Gate audit counts")
    ax1.set_xlim(0, max(counts["total"]) + 2)
    ax1.legend(frameon=False, ncol=2, loc="upper right")

    forest = df.sort_values("test_delta", ascending=True).reset_index(drop=True)
    yy = np.arange(len(forest))
    colors = np.where(forest["promoted"], BLUE, GREY)
    ax2.hlines(yy, forest["interval_low"], forest["interval_high"], color=LIGHT, lw=2.4, zorder=1)
    ax2.scatter(forest["test_delta"], yy, c=colors, s=42, edgecolor=DARK, linewidth=0.35, zorder=2)
    ax2.axvline(0, color=DARK, lw=1.0)
    ax2.set_yticks(yy, forest["endpoint_label"])
    ax2.set_xlabel("Direction-normalized test delta")
    ax2.set_title("Endpoint-level effects")
    ax2.set_xlim(min(forest["interval_low"].min() - 0.02, -0.08), max(forest["interval_high"].max() + 0.08, 0.3))

    seeds_plot = seeds.set_index("gate_category").loc[CAT_ORDER].reset_index()
    y = np.arange(len(seeds_plot))
    left = np.zeros(len(seeds_plot))
    parts = [
        ("positive_seed_fraction", "Positive seed delta", BLUE),
        ("tie_seed_fraction", "Tie", LIGHT),
        ("negative_seed_fraction", "Negative seed delta", RED),
    ]
    for col, label, color in parts:
        ax3.barh(y, seeds_plot[col], left=left, color=color, edgecolor="white", linewidth=0.8, label=label)
        left += seeds_plot[col].to_numpy()
    for idx, row in seeds_plot.iterrows():
        ax3.text(
            1.02,
            idx,
            f"{row['positive_seed_delta']}/{row['total_seed_summaries']} +; "
            f"{row['negative_seed_delta']}/{row['total_seed_summaries']} -",
            va="center",
            ha="left",
            color="#64748B",
            fontsize=9,
            clip_on=False,
        )
    ax3.set_yticks(y, seeds_plot["category_label"])
    ax3.invert_yaxis()
    ax3.set_xlim(0, 1.0)
    ax3.set_xlabel("Fraction of paired seed summaries")
    ax3.set_title("Seed-direction audit")
    ax3.legend(frameon=False, ncol=3, loc="lower center", bbox_to_anchor=(0.5, -0.24))

    marker_colors = {
        "promoted_and_improved": BLUE,
        "retained_and_avoided_harm": GREY,
        "inconclusive_due_to_wide_ci": RED,
    }
    for cat in CAT_ORDER:
        g = df[df["gate_category"].eq(cat)]
        ax4.scatter(
            g["test_delta"],
            g["interval_half_width"],
            s=60 + 140 * np.clip(g["abs_delta"], 0, 0.2),
            color=marker_colors[cat],
            edgecolor=DARK,
            linewidth=0.35,
            alpha=0.88,
            label=CAT_LABELS[cat].replace("\n", " "),
        )
    ax4.axvline(0, color=DARK, lw=1.0)
    ax4.set_xlabel("Direction-normalized test delta")
    ax4.set_ylabel("Seed-interval half-width")
    ax4.set_title("Effect size versus uncertainty")
    ax4.legend(frameon=False, loc="upper right")

    for ax in axes.ravel():
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.grid(axis="x", color="#E7EBF0", lw=0.8)
        ax.set_axisbelow(True)

    fig.tight_layout(w_pad=2.4, h_pad=2.6)
    png = FIG_DIR / "fig07_tdc_gate_four_panel_reviewer.png"
    svg = FIG_DIR / "fig07_tdc_gate_four_panel_reviewer.svg"
    fig.savefig(png, bbox_inches="tight", facecolor="white")
    fig.savefig(svg, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    metrics = {
        "endpoints": int(len(df)),
        "promoted": int(df["promoted"].sum()),
        "retained": int((~df["promoted"]).sum()),
        "promoted_and_improved": int(df["gate_category"].eq("promoted_and_improved").sum()),
        "retained_and_avoided_harm": int(df["gate_category"].eq("retained_and_avoided_harm").sum()),
        "inconclusive": int(df["gate_category"].eq("inconclusive_due_to_wide_ci").sum()),
        "promoted_but_wide_interval": int(
            df["promoted"].eq(True).mul(df["gate_category"].eq("inconclusive_due_to_wide_ci")).sum()
        ),
        "promoted_seed_positive": int(
            df.loc[df["gate_category"].eq("promoted_and_improved"), "seed_win_count"].sum()
        ),
        "promoted_seed_total": int(
            df.loc[df["gate_category"].eq("promoted_and_improved"), "n_paired_seeds"].sum()
        ),
        "retained_avoided_negative": int(
            df.loc[df["gate_category"].eq("retained_and_avoided_harm"), "seed_loss_count"].sum()
        ),
        "retained_avoided_total": int(
            df.loc[df["gate_category"].eq("retained_and_avoided_harm"), "n_paired_seeds"].sum()
        ),
        "inconclusive_cross_zero": int(
            df.loc[df["gate_category"].eq("inconclusive_due_to_wide_ci"), "interval_crosses_zero"].sum()
        ),
        "inconclusive_total": int(df["gate_category"].eq("inconclusive_due_to_wide_ci").sum()),
        "max_positive_delta": float(df["test_delta"].max()),
        "min_delta": float(df["test_delta"].min()),
    }
    return png, svg, metrics


def replace_text(doc: Document, start: str, text: str) -> None:
    for p in doc.paragraphs:
        if p.text.strip().startswith(start):
            p.text = text
            return
    raise RuntimeError(f"Paragraph not found: {start}")


def replace_image_before_caption(doc: Document, caption_start: str, image: Path) -> None:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(caption_start):
            target = doc.paragraphs[i - 1]
            target.clear()
            target.add_run().add_picture(str(image), width=Inches(6.6))
            return
    raise RuntimeError(f"Caption not found: {caption_start}")


def update_docx(png: Path, metrics: dict[str, float]) -> tuple[Path, Path, bool]:
    target = docx_path()
    backup = OUT / f"{target.stem}_图7四联图前备份_{datetime.now():%Y%m%d_%H%M%S}.docx"
    shutil.copy2(target, backup)
    doc = Document(str(target))
    replace_text(
        doc,
        "TDC 结果强调门控规则",
        (
            "TDC 四联审计图将门控规则拆解为审稿人最关心的四个问题：冻结门控产生了哪些行动决策，"
            "晋级终点是否具有外层测试改善，保留终点是否确实避免潜在伤害，以及哪些终点仅因三种子区间过宽而证据不足。"
            f"在 22 个 TDC 终点中，冻结门控产生 {metrics['promoted']} 个 promoted 和 {metrics['retained']} 个 retained；"
            f"事后审计显示 {metrics['promoted_and_improved']} 个 promoted-and-improved、"
            f"{metrics['retained_and_avoided_harm']} 个 retained-and-avoided-harm，以及 "
            f"{metrics['inconclusive']} 个区间过宽的 inconclusive 终点。"
        ),
    )
    replace_text(
        doc,
        "TDC 部分仅使用 3 个种子",
        (
            "由于 TDC 部分仅使用 3 个种子，图7中的横向区间解释为 seed-summary interval，而不是严格抽样置信区间。"
            f"promoted-and-improved 终点的种子方向为 {metrics['promoted_seed_positive']}/{metrics['promoted_seed_total']} 正向，"
            f"retained-and-avoided-harm 终点为 {metrics['retained_avoided_negative']}/{metrics['retained_avoided_total']} 负向，"
            "支持门控在这两类终点上作出行动性决策；相反，跨零且区间较宽的终点被保留为证据不足，而不是被解释为真实无效。"
        ),
    )
    replace_text(
        doc,
        "图 7",
        (
            "图 7  TDC 门控有效性的四联审计图。左上，22 个 TDC 终点经冻结门控后的 promoted/retained 构成和事后审计类别；"
            "右上，所有终点的方向归一化测试增量及三种子 seed-summary interval；左下，各审计类别的逐种子方向支持；"
            "右下，效应大小与区间半宽的关系，用于区分可行动改善、避免伤害和证据不足。"
        ),
    )
    replace_image_before_caption(doc, "图 7", png)
    try:
        doc.save(str(target))
        return target, backup, False
    except PermissionError:
        fallback = OUT / f"{target.stem}_图7四联图.docx"
        doc.save(str(fallback))
        return fallback, backup, True


def audit_docx(target: Path) -> dict[str, object]:
    with ZipFile(target) as zf:
        bad = zf.testzip()
        media_count = sum(n.startswith("word/media/") for n in zf.namelist())
    doc = Document(str(target))
    text = "\n".join(p.text for p in doc.paragraphs)
    return {
        "target": str(target),
        "zip_ok": bad is None,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "figures": len(doc.inline_shapes),
        "media_count": media_count,
        "caption_updated": "TDC 门控有效性的四联审计图" in text,
        "reviewer_logic_updated": "审稿人最关心的四个问题" in text,
    }


def write_report(target: Path, backup: Path, metrics: dict[str, float], audit: dict[str, object]) -> Path:
    report = OUT / "小论文-18_图7四联图审稿人视角加厚报告.md"
    lines = [
        "# 小论文-18 图7四联图审稿人视角加厚报告",
        "",
        f"- 生成时间：{datetime.now().isoformat(timespec='seconds')}",
        f"- 更新文档：`{target}`",
        f"- 备份文档：`{backup}`",
        f"- 图表目录：`{FIG_DIR}`",
        "",
        "## 四联图设计逻辑",
        "",
        "- 左上：门控行动结果，显示 promoted/retained 与事后审计类别是否一致。",
        "- 右上：所有终点测试增量和 seed-summary interval，避免只展示有利终点。",
        "- 左下：三种子方向支持，直接回应样本/种子数不足的审稿质疑。",
        "- 右下：效应大小与区间半宽，解释哪些终点是证据不足而不是负结论。",
        "",
        "## 关键结果",
        "",
        f"- TDC 终点数：{metrics['endpoints']}；promoted：{metrics['promoted']}；retained：{metrics['retained']}。",
        f"- promoted-and-improved：{metrics['promoted_and_improved']}；retained-and-avoided-harm：{metrics['retained_and_avoided_harm']}；inconclusive：{metrics['inconclusive']}。",
        f"- promoted-and-improved 种子方向：{metrics['promoted_seed_positive']}/{metrics['promoted_seed_total']} 正向。",
        f"- retained-and-avoided-harm 种子方向：{metrics['retained_avoided_negative']}/{metrics['retained_avoided_total']} 负向。",
        "",
        "## 审计",
        "",
        "```json",
        json.dumps(audit, ensure_ascii=False, indent=2),
        "```",
    ]
    report.write_text("\n".join(lines), encoding="utf-8-sig")
    return report


def main() -> None:
    setup_plot()
    df = load_data()
    png, svg, metrics = draw_figure(df)
    target, backup, used_fallback = update_docx(png, metrics)
    audit = audit_docx(target)
    audit.update(
        {
            "png": str(png),
            "svg": str(svg),
            "source_data": str(FIG_DIR / "fig07_tdc_gate_four_panel_reviewer_source.csv"),
            "backup": str(backup),
            "used_fallback_docx": used_fallback,
            "metrics": metrics,
        }
    )
    audit["passed"] = (
        audit["zip_ok"]
        and audit["caption_updated"]
        and audit["reviewer_logic_updated"]
        and Path(audit["source_data"]).exists()
    )
    AUDIT.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    report = write_report(target, backup, metrics, audit)
    print(
        json.dumps(
            {
                "docx": str(target),
                "report": str(report),
                "png": str(png),
                "svg": str(svg),
                "passed": audit["passed"],
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
