from __future__ import annotations

import json
import re
import shutil
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("D:/fzyc")
sys.path.insert(0, str(ROOT / "scripts"))
from build_paper19_manuscript import add_body, add_figure, add_heading, configure_document, set_run_font  # noqa: E402

OUT = ROOT / "output"
CORE = OUT / "paper20_candidate_pool_audit_20260712"
MAJOR = OUT / "paper19_major_revision_20260712"
REV = OUT / "paper19_jcheminformatics_revision_20260712"
FIG = CORE / "main_figures"
SOURCE = OUT / "小论文-19_8000词MajorRevision.docx"
PAPER = OUT / "小论文-20_Journal_of_Cheminformatics_主文.docx"
DESKTOP = Path("C:/Users/Administrator/Desktop/小论文-20_Journal_of_Cheminformatics_主文.docx")
AUDIT = OUT / "小论文-20_主文审计.json"
TITLE = "Candidate-pool expansion, validation-ranking distortion and model-selection loss in molecular property prediction: a retrospective nested audit"
FORBIDDEN = [
    "test oracle", "true oracle", "true attainable upper bound", "independent confirmation",
    "universal governance framework", "prospective preregistration",
]
DISPLAY = {
    "bace": "BACE", "bbbp": "BBBP", "clintox": "ClinTox", "esol": "ESOL",
    "freesolv": "FreeSolv", "lipo": "Lipophilicity", "tdc_caco2_wang": "Caco2 Wang",
    "tdc_hia_hou": "HIA Hou", "tdc_pgp_broccatelli": "P-gp Broccatelli",
}


def clean(text: str) -> str:
    replacements = {
        "effective predictive diversity": "utility-pattern diversity",
        "predictive diversity": "utility-pattern diversity",
        "independent confirmation": "external validation",
        "prospective preregistration": "prospectively registered design",
        "test oracle": "observed audit-best candidate",
        "true oracle": "population-optimal candidate",
        "true generalization upper bound": "population performance bound",
        "universal selector": "generally applicable selector",
        "universal model superiority": "general model superiority",
        "15 ℅ K": "15 x K",
        "3 × 3 × 5": "3 x 3 x 5",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def body(doc: Document, text: str) -> None:
    add_body(doc, clean(text))


def source_paragraphs(source: Document, start: int, stop: int) -> list[str]:
    return [clean(p.text) for p in source.paragraphs[start:stop] if p.text.strip() and not p.style.name.startswith("Heading") and not p.text.startswith(("Table ", "Figure "))]


def abstract_part(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run_font(p.add_run(f"{label}: "), 10.5, bold=True)
    set_run_font(p.add_run(clean(text)), 10.5)


def set_cell_border(cell, top: bool = False, bottom: bool = False) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:tcBorders"))
    if old is not None:
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = OxmlElement(f"w:{edge}")
        active = (edge == "top" and top) or (edge == "bottom" and bottom)
        tag.set(qn("w:val"), "single" if active else "nil")
        if active:
            tag.set(qn("w:sz"), "8")
            tag.set(qn("w:color"), "000000")
        borders.append(tag)
    tc_pr.append(borders)


def add_three_line_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths: list[float], page_break_before: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(caption), 9.5, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")
    header = table.rows[0]
    header._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for j, value in enumerate(headers):
        header.cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(header.cells[j], top=True, bottom=True)
        paragraph = header.cells[j].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(value), 8.2, bold=True)
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cells[j], bottom=i == len(rows) - 1)
            paragraph = cells[j].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            set_run_font(paragraph.add_run(str(value)), 8.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def table1(doc: Document) -> None:
    frame = pd.read_csv(REV / "dataset_characteristics.csv")
    rows = []
    for _, r in frame.iterrows():
        if r.task_type == "classification":
            distribution = f"{int(r.positive_n)} positive ({100*r.positive_rate:.1f}%)"
            metric = "ROC-AUC"
        else:
            distribution = f"{r.target_min:.2f} to {r.target_max:.2f} {r.target_unit}"
            metric = "RMSE"
        rows.append([r.display_name, f"{int(r.analysis_n)}", distribution, metric])
    add_three_line_table(doc, "Table 1. Datasets and primary endpoint metrics.",
        ["Endpoint", "n", "Class balance or target range", "Primary metric"], rows, [3.5, 1.5, 7.0, 2.5])


def table2(doc: Document) -> None:
    rows = [
        ["Controlled prefixes (K=4/8/16/32)", "Morgan-512", "Linear; bagging; boosting", "Expansion audit", "17,280 fits"],
        ["Composition controls (K=4/8/16/32)", "Morgan-512", "Registered learners", "Registry sensitivity", "100/mode/K"],
        ["Multiview (K=3/6/9/12)", "Morgan; MACCS; RDKit2D; concat.", "Linear; RF; LightGBM", "Matched composition", "6,480 fits"],
        ["Modern-model panel (K=4)", "Morgan; GCN; ChemBERTa; MoLFormer", "RF; GCN; frozen probes", "Boundary only", "360 units"],
    ]
    add_three_line_table(doc, "Table 2. Candidate registries and computational exposure.",
        ["Analysis and K", "Representations", "Learners", "Purpose", "Fits/runs"], rows, [4.0, 4.1, 3.3, 3.5, 2.0], page_break_before=True)


def table3(doc: Document) -> None:
    frame = pd.read_csv(CORE / "cross_fitted_k32_minus_k4.csv")
    rows = []
    for _, r in frame.iterrows():
        metric = "ROC-AUC loss" if r.task_type == "classification" else "RMSE loss"
        rows.append([
            DISPLAY[r.task], metric, f"{r.k32_minus_k4_same_unit_gap:.4f}",
            f"{r.seed_clustered_ci95_low_same_unit_gap:.4f} to {r.seed_clustered_ci95_high_same_unit_gap:.4f}",
            f"{r.k32_minus_k4_cross_fitted_gap:.4f}",
        ])
    add_three_line_table(doc, "Table 3. Endpoint-specific K = 32 minus K = 4 audit-gap effects.",
        ["Endpoint", "Metric", "Same-unit effect", "Seed-clustered 95% interval", "Cross-fitted effect"], rows, [3.3, 2.8, 2.5, 4.2, 2.7], page_break_before=True)


def figure(doc: Document, number: int, stem: str, caption: str) -> None:
    add_figure(doc, FIG / f"Figure_{number}_{stem}.png", 16.1, f"Figure {number}. {clean(caption)}")


def build_document() -> Document:
    source = Document(SOURCE)
    doc = Document()
    configure_document(doc)
    header = doc.sections[0].header.paragraphs[0]
    header.clear(); header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(header.add_run("Candidate-pool expansion and model-selection loss"), 8)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run(TITLE), 16, bold=True)
    add_heading(doc, "Abstract", 1)
    abstract_part(doc, "Background", "Molecular-property studies increasingly select among large, correlated candidate registries. Candidate expansion can add chemical information, but finite validation data may rank candidates unreliably and inflate the gap between the selected model and the largest observed audit estimate.")
    abstract_part(doc, "Methods", "We performed a retrospective audit of nine endpoints under five seeds, three outer scaffold folds and three inner folds. Registered Morgan-based prefixes contained K = 4, 8, 16 or 32 candidates. Utility-pattern diversity was estimated from raw, row-centred, fixed-reference-relative and within-unit-rank matrices using empirical and Ledoit-Wolf spectral estimators, participation ratios, hierarchical bootstrap intervals and omission sensitivity. We decomposed audit-best and selected-model gains, evaluated a leave-one-seed-out cross-fitted reference, and compared matched-size multiview pools.")
    abstract_part(doc, "Results", "At K = 32, hierarchical-bootstrap median entropy ranks were 3.39 for raw utilities, 10.38 after row centring, 5.13 relative to the fixed candidate and 12.06 for within-unit ranks, all below nominal K. Chance-adjusted Hit@3 declined from 0.881 at K = 4 to 0.240 at K = 32, and normalized mean reciprocal-rank gain declined from 0.727 to 0.157. The same-unit audit gap increased in eight of nine endpoints; the cross-fitted K = 32 minus K = 4 effect was positive in seven, negative for BBBP and P-gp, and uncertain for BACE. At matched K = 3, alternative representation compositions improved eight endpoints on average but not BACE. Exact zero-width multiview seed intervals on three regression endpoints reflected repeated identical fold-level scores rather than additional precision.")
    abstract_part(doc, "Conclusions", "Within the evaluated design, nominal expansion exceeded utility-pattern diversity and coincided with weaker validation ranking and larger endpoint-dependent model-selection gaps. Heterogeneous representations could add value, but their apparent benefit combined representation composition, candidate count and finite-sample opportunity.")
    abstract_part(doc, "Scientific Contribution", "This study separates common audit-unit difficulty, chance-adjusted ranking degradation, cross-fitted selection gaps and matched-size representation effects in a repeated nested molecular benchmark. It is a model-selection audit and reporting analysis rather than a new molecular predictor or external validation study.")
    body(doc, "Keywords: molecular property prediction; candidate-pool expansion; model selection; nested cross-validation; utility-pattern diversity; ranking fidelity; audit gap; multiview representation")

    add_heading(doc, "1 Introduction", 1)
    for text in source_paragraphs(source, 9, 14): body(doc, text)

    add_heading(doc, "2 Methods", 1)
    add_heading(doc, "2.1 Study design and evidence hierarchy", 2)
    for text in source_paragraphs(source, 16, 19): body(doc, text)
    body(doc, "The present revision treated the candidate-pool audit as the only primary inferential line. TDC, automated machine-learning, conformal, activity-cliff, beyond-rule-of-five and duplicate-sensitivity analyses were assigned to Supplementary Information because they characterize transportability or reliability but do not directly identify the expansion mechanism.")

    add_heading(doc, "2.2 Datasets and molecular standardization", 2)
    for text in source_paragraphs(source, 20, 23): body(doc, text)
    table1(doc)

    add_heading(doc, "2.3 Candidate registry and computational exposure", 2)
    for text in source_paragraphs(source, 25, 28): body(doc, text)
    table2(doc)

    add_heading(doc, "2.4 Repeated nested scaffold evaluation", 2)
    for text in source_paragraphs(source, 30, 33): body(doc, text)

    add_heading(doc, "2.5 Audit-gap and ranking estimands", 2)
    for text in source_paragraphs(source, 34, 38): body(doc, text)
    body(doc, "For each outer unit, observed audit-best gain was the largest eligible outer utility minus the fixed candidate utility; selected-model gain was the validation-selected utility minus the same fixed utility. Their difference defined the incremental observed audit gap. The realization ratio divided selected-model gain by observed audit-best gain when the denominator was positive. Classification and regression were summarized separately because ROC-AUC loss and RMSE loss are not exchangeable units.")

    add_heading(doc, "2.6 Effective-diversity estimation", 2)
    body(doc, "For each endpoint and K, we formed a 15 x K matrix whose rows were seed-by-outer-fold units and whose columns were candidate utilities. Four matrices were analysed: raw utilities; row-centred utilities after subtracting each unit mean; utilities relative to candidate 1, with the zero reference column removed; and within-unit candidate ranks. Row centring targets common audit-unit difficulty, whereas the fixed-reference and rank matrices ask whether candidate contrasts persist after removing level shifts. We use utility-pattern diversity because aggregate fold utilities do not fully represent prediction-level or mechanistic diversity.")
    for text in source_paragraphs(source, 39, 43): body(doc, text)

    add_heading(doc, "2.7 Candidate-composition controls", 2)
    for text in source_paragraphs(source, 44, 46): body(doc, text)

    add_heading(doc, "2.8 Cross-fitted audit reference", 2)
    body(doc, "To reduce the same-unit advantage of defining the reference with the evaluated outer scores, we used leave-one-seed-out cross-fitting. For each endpoint, K and held-out seed, candidate mean outer utility was computed over the other four seeds. The highest mean candidate, with registry order used for ties, was fixed as the reference and evaluated on all three folds of the held-out seed. The cross-fitted gap was reference utility minus validation-selected utility; the same-unit gap used the maximum candidate utility in that held-out fold. Their difference quantifies the additional optimism induced by the same-unit maximum.")
    body(doc, "This reference is less circular than a same-unit maximum but is not a new cohort or prospective experiment. It reuses the same endpoint population and split generator. Its purpose is to test whether the K = 32 minus K = 4 direction persists when reference selection and evaluation are separated at the seed level.")

    add_heading(doc, "2.9 Finite-audit winner-optimism simulation", 2)
    for text in source_paragraphs(source, 47, 49): body(doc, text)

    add_heading(doc, "2.10 Matched-size multiview stress test", 2)
    body(doc, "The multiview registry crossed four representations (Morgan-512, MACCS, RDKit2D and their concatenation) with linear, random-forest and LightGBM learners. Incremental pools contained K = 3, 6, 9 and 12 candidates by adding representations in that order. All representation pairs at K = 6 and all triples at K = 9 were evaluated as composition sensitivities. At K = 3, Morgan across three learners was compared with six fixed alternatives that assigned MACCS, RDKit2D and concatenated features bijectively to the same three learners. Thus K and learner count were constant while representation composition changed.")
    body(doc, "Within-learner analyses selected among four representations for one learner and compared the result with that learner's Morgan candidate. The full K = 12 versus Morgan K = 3 contrast is termed the multiview-pool gain because it combines representation breadth and additional candidates. Only comparisons at the same K are termed representation-composition effects.")
    body(doc, "For every paired unit we retained seed, outer fold, selected candidates and utilities. Source-file SHA-256 values and deterministic hashes of split keys and paired scores were recorded. Candidate-level prediction exports were absent from this result directory, so a prediction hash could not be reconstructed and is identified as requiring source data.")

    add_heading(doc, "2.11 Limited representation-baseline analysis", 2)
    for text in source_paragraphs(source, 54, 57): body(doc, text)

    add_heading(doc, "2.12 Reliability and chemical-boundary analyses", 2)
    for text in source_paragraphs(source, 58, 61): body(doc, text)
    body(doc, "These analyses were retained in the main text only as boundaries on interpretation. Complete TDC, conformal, MoleculeACE, bRo5, deduplication and failure-case results are provided in Additional files 1-3.")

    add_heading(doc, "2.13 Statistical inference", 2)
    body(doc, "For endpoint-specific K = 32 minus K = 4 contrasts, outer-fold differences were first averaged within each seed. Seed-clustered 95% intervals resampled the five seed means 10,000 times. Leave-one-seed estimates assessed repeat sensitivity. Effective-diversity intervals used 100 hierarchical bootstrap replicates that sampled seeds and then outer folds with replacement; the modest replicate count reflects the computational cost of repeated covariance shrinkage and is reported explicitly.")
    body(doc, "For multiview paired gains, we report seed means, seed-clustered bootstrap intervals, t intervals across five seed means and leave-one-seed ranges. When all seed means were identical, the resulting interval was exactly or numerically zero width. Such intervals were interpreted as a property of deterministic repeated scores, not as five independent confirmations. No P values were used to convert continuous evidence into binary claims, and no multiplicity-adjusted hypothesis family was prespecified.")

    figure(doc, 1, "retrospective_nested_audit_architecture", "Retrospective nested audit architecture. The continuous workflow links data registration, candidate construction, repeated nested selection, outer auditing, cross-fitted reference construction, utility-pattern diversity, ranking fidelity, audit-gap decomposition, matched-size multiview analysis, reliability boundaries and auditable reporting.")

    add_heading(doc, "3 Results", 1)
    add_heading(doc, "3.1 Nominal expansion exceeded utility-pattern diversity", 2)
    body(doc, "The first analysis asked how much nominal K exceeded the effective dimensionality of candidate utility patterns after accounting for common outer-unit difficulty. At K = 32, the endpoint-mean raw Ledoit-Wolf entropy rank was 3.43; the hierarchical-bootstrap median was 3.39 (95% interval 2.74 to 5.58). Row centring increased the bootstrap median to 10.38 (8.63 to 12.81), showing that common endpoint-fold difficulty explained a substantial part of the raw candidate correlation. The fixed-reference-relative median was 5.13 (3.86 to 7.12), whereas within-unit ranks gave 12.06 (9.70 to 14.61). Every adjusted estimate remained below nominal K = 32.")
    body(doc, "Estimator choice changed magnitude but not the central contrast. At K = 32, raw participation-ratio rank was 1.80 and the reference-relative value was 2.89; row-centred and rank-based point estimates were higher. The raw median candidate correlation was 0.814, the reference-relative value was 0.688 and row centring reduced the median to 0.039. These changes demonstrate why unadjusted utility correlation cannot by itself be described as complete predictive diversity.")
    body(doc, "Leave-one-seed and leave-one-fold analyses retained a large nominal-effective gap but showed endpoint heterogeneity. Reference-relative K = 32 ranks were highest for HIA and lowest for ESOL, Caco2, Lipophilicity and FreeSolv. The conclusion is therefore qualitative: the evaluated registry supplied more nominal choices than distinct utility-pattern directions, while the exact effective count depended on how common audit difficulty was removed (Additional file 2: Tables S5-S7).")
    figure(doc, 2, "effective_candidate_diversity", "Effective candidate diversity after audit-difficulty adjustment. (A-C) Hierarchical-bootstrap median Ledoit-Wolf entropy ranks and 95% intervals for raw, row-centred and fixed-reference-relative matrices; the dotted line denotes nominal K. (D) Participation-ratio ranks across four transformations. (E) Median candidate correlation before and after adjustment. (F) K = 32 endpoint omission ranges for the reference-relative entropy rank. Error bars in A-C are hierarchical bootstrap intervals; horizontal lines in F are leave-one-seed and leave-one-fold ranges combined.")

    add_heading(doc, "3.2 Chance-adjusted ranking fidelity declined with K", 2)
    for text in source_paragraphs(source, 76, 79): body(doc, text)
    body(doc, "The opportunity-adjusted measures therefore identify a ranking problem beyond the mechanical decline of Top-1 accuracy with K. Validation ordering retained information above chance, but that information was insufficient to preserve the same level of winner recovery as the registry expanded.")

    add_heading(doc, "3.3 Audit-gap decomposition separated opportunity from realization", 2)
    body(doc, "At K = 32, the classification mean observed audit-best gain over the fixed candidate was 0.0311 ROC-AUC, whereas selected-model gain was 0.0173 and the incremental observed audit gap was 0.0138. At K = 4, the corresponding values were 0.0028, -0.0014 and 0.0042. For regression, the K = 32 means were 0.8078 RMSE units of observed audit-best gain, 0.7378 selected-model gain and a 0.0700 gap, compared with 0.7619, 0.7442 and 0.0178 at K = 4. The scales are reported separately and are not pooled.")
    body(doc, "The decomposition shows that expansion can raise both the opportunity represented by the largest observed utility and the amount captured by the selected model. Selection loss increases when the former rises faster than the latter. Realization ratios were unstable when the audit-best gain was near zero and are therefore secondary, machine-readable quantities rather than a headline result (Additional file 2: Table S8).")
    figure(doc, 3, "ranking_distortion_and_audit_gaps", "Chance-adjusted ranking distortion and endpoint audit gaps. (A) Chance-adjusted Hit@3. (B) Mean audit-rank percentile. (C) Normalized mean reciprocal-rank gain. (D) NDCG, Spearman and Kendall agreement. (E) Endpoint-specific K = 32 minus K = 4 same-unit audit-gap effects with seed-clustered 95% intervals. Blue denotes ROC-AUC loss and orange denotes RMSE loss; units are not pooled. (F) Candidate-composition controls.")

    add_heading(doc, "3.4 Cross-fitted effects remained positive in seven endpoints", 2)
    body(doc, "Separating reference selection from held-out-seed evaluation attenuated the expansion effect but did not remove it for most endpoints. The K = 32 minus K = 4 cross-fitted gap was positive in BACE (0.0009), ClinTox (0.0098), ESOL (0.0578), FreeSolv (0.0580), Lipophilicity (0.0072), Caco2 (0.0149) and HIA (0.0176). BBBP (-0.0005) and P-gp (-0.0009) were negative. The corresponding seed-clustered intervals excluded zero for ClinTox, ESOL, Lipophilicity, Caco2 and HIA, but not for BACE, BBBP, FreeSolv or P-gp.")
    body(doc, "The same-unit effect was positive in eight endpoints and negative only for P-gp. It exceeded the cross-fitted effect in most endpoints, consistent with optimism from selecting the maximum within the evaluated unit. BACE remained directionally positive but uncertain, and P-gp provided a reproducible negative-direction boundary. These exceptions are retained in the main table rather than relegated to sensitivity material.")
    body(doc, "Because the cross-fitted reference still uses the same public endpoint and split generator, it does not transform the audit into an external validation. It does, however, support use of model-selection loss in the title as an endpoint-qualified empirical result rather than a claim based only on same-unit maxima.")
    table3(doc)

    add_heading(doc, "3.5 Finite-audit maxima contained winner optimism", 2)
    for text in source_paragraphs(source, 92, 94): body(doc, text)
    body(doc, "At candidate correlation 0.9, winner optimism increased as effective audit size decreased and as K increased. This simulation does not estimate a correction for the empirical endpoints; it demonstrates that a positive same-unit maximum can arise mechanically even when candidate population utilities are equal.")
    figure(doc, 4, "audit_gap_decomposition", "Audit-gap decomposition and finite-sample winner optimism. (A-C) Endpoint-specific observed audit-best gain, selected-model gain and incremental gap at K = 32; blue and orange denote classification and regression utility units, which are not pooled. (D) Same-unit versus cross-fitted K = 32 minus K = 4 effects; the dotted line is equality. (E) Cross-fitted endpoint effects with seed-clustered 95% intervals. (F) Equal-truth winner-optimism simulation at candidate correlation 0.9.")

    add_heading(doc, "3.6 Matched-size multiview gains were endpoint dependent", 2)
    body(doc, "The matched K = 3 analysis changed the interpretation of the original full-pool comparison. Across six fixed alternative representation assignments, the mean representation-composition effect was positive for BBBP, ClinTox, ESOL, FreeSolv, Lipophilicity, Caco2, HIA and P-gp, but slightly negative for BACE (-0.0021 ROC-AUC). All six alternatives were positive for each of the eight responsive endpoints; only two of six were positive for BACE. Thus multiview content, not candidate count alone, contributed to the observed gains.")
    body(doc, "All representation pairs at K = 6 and triples at K = 9 showed the same broad endpoint heterogeneity. The incremental K = 3, 6, 9 and 12 ladder generally increased selected utility, but most gains occurred before the full pool for several endpoints. Within-learner comparisons also varied, indicating that representation effects interacted with learner family rather than forming a single representation ranking.")
    body(doc, "The full multiview-pool gains remained 0.0039 to 0.0873 ROC-AUC for classification and 0.1058 to 1.1949 RMSE reduction for regression. BACE crossed zero under the seed-clustered interval. ESOL, Lipophilicity and Caco2 had exactly repeated seed means, producing zero-width bootstrap and t intervals. Their fold-level selected candidates and paired scores were identical across the five seed labels; these intervals therefore reflect deterministic repeated splits, not unusually high inferential precision. Complete matched-K results are in Additional file 2: Tables S9-S10.")
    figure(doc, 5, "matched_size_multiview", "Matched-size multiview pools and endpoint-dependent gains. (A) Within-endpoint normalized effects for matched K = 3 representation compositions. (B) Within-learner effects normalized within endpoint and referenced to the same learner's Morgan candidate. (C) Within-endpoint normalized K = 3, 6, 9 and 12 ladder. (D) Raw full K = 12 versus Morgan K = 3 paired effects with seed-clustered 95% intervals. (E) Endpoint-relative selected-performance ratio, defined so that higher is better. (F) Representation-selection frequency. Raw classification and regression effects are shown only in D and are not pooled.")

    add_heading(doc, "3.7 Representation errors were correlated but complementary", 2)
    for text in source_paragraphs(source, 105, 108): body(doc, text)
    body(doc, "Prediction-level error exports were available for the limited four-model representation panel, enabling pairwise Jaccard overlap. Mean overlap ranged from 0.168 for GCN versus RDKit-RF to 0.296 for ChemBERTa versus RDKit-RF. The incomplete overlap indicates complementary errors, but unequal training budgets prevent a performance ranking of fully optimized architectures. All-candidate prediction exports were unavailable for the 32-candidate near-duplicate registry.")

    add_heading(doc, "3.8 Reliability deteriorated at chemical boundaries", 2)
    for text in source_paragraphs(source, 109, 115): body(doc, text)
    body(doc, "ClinTox remained the principal negative result. Minority recall varied from 0.107 for RDKit-RF to 0.955 for the GCN, but the latter produced many false positives and the highest 10-fold false-negative-cost utility was not equivalent to a deployable toxicity decision rule. Tanimoto strata, activity-cliff pairs, bRo5 perimeters and extreme labels identified distinct boundary mechanisms. Their complete numerical results and representative structures are supplied in Additional files 1-3.")
    figure(doc, 6, "representation_errors_and_boundaries", "Representation errors and chemical reliability boundaries. (A) Representation utilities normalized within endpoint relative to RDKit-RF. (B) Pairwise prediction-error Jaccard overlap. (C) Endpoint mean error complementarity. (D) ROC-AUC and one minus ECE across maximum train-set Tanimoto strata. (E) ClinTox minority recall and false-negative rate. (F) Presence of failure categories by source. Normalization in A is for visualization only; complete raw values are in Additional file 2.")

    add_heading(doc, "4 Discussion", 1)
    add_heading(doc, "4.1 Nominal K did not represent independent opportunity", 2)
    for text in source_paragraphs(source, 117, 119): body(doc, text)
    body(doc, "Common outer-unit difficulty was not a minor technical adjustment. Row centring raised the K = 32 entropy-rank median from 3.39 to 10.38, while reference-relative and rank transformations gave different intermediate or higher values. The stable conclusion is not a single effective count; it is that nominal K, common difficulty and candidate contrasts must be reported separately.")

    add_heading(doc, "4.2 Chance-adjusted ranking degradation accompanied selection loss", 2)
    for text in source_paragraphs(source, 120, 123): body(doc, text)
    body(doc, "Cross-fitting sharpened this interpretation. Same-unit maxima overstated the gap relative to references chosen on other seeds, yet seven endpoint directions remained positive. The negative BBBP and P-gp effects and uncertain BACE effect prevent a deterministic or universal statement. Within the studied endpoints, expansion increased model-selection loss more often than it reduced it.")

    add_heading(doc, "4.3 Matched K changed the multiview interpretation", 2)
    body(doc, source.paragraphs[124].text)
    body(doc, "The matched K = 3 alternatives showed that eight endpoints benefited from changing representation composition while holding candidate and learner counts constant. Consequently, the full-pool gain cannot be dismissed as search opportunity alone. Conversely, the K ladder and within-learner heterogeneity show that the full K = 12 contrast still combines representation breadth, learner interactions and additional selection opportunities. The precise claim is an endpoint-dependent representation-composition effect plus a broader multiview-pool gain.")

    add_heading(doc, "4.4 Same-unit maxima remain finite-sample estimates", 2)
    for text in source_paragraphs(source, 127, 129): body(doc, text)
    body(doc, "The zero-width multiview intervals provide a concrete warning. Repeating a seed label did not create new regression fold outcomes when the split generator returned identical partitions and deterministic fitted scores. Such repetitions are useful reproducibility checks, but uncertainty summaries must reveal when their effective information content is one set of folds rather than five distinct repeats.")

    add_heading(doc, "4.5 Implications for molecular benchmark reporting", 2)
    for text in source_paragraphs(source, 130, 133): body(doc, text)
    body(doc, "A compact benchmark report can therefore remain rigorous without displaying every registry row in the main article. The main text should identify the estimands, candidate opportunity, matched comparisons and endpoint exceptions; machine-readable supplementary tables should retain every candidate, fold, status and source hash.")

    add_heading(doc, "4.6 Reliability remains conditional on chemical support", 2)
    for text in source_paragraphs(source, 134, 136): body(doc, text)
    body(doc, "These boundary analyses do not rescue or overturn the candidate-expansion finding. They delimit where an apparently strong selected model can still fail, especially for rare toxicity labels, low-similarity molecules and activity cliffs. Reliability evidence should accompany, rather than be averaged into, the model-selection audit.")

    add_heading(doc, "4.7 Limitations", 2)
    for text in source_paragraphs(source, 137, 140): body(doc, text)
    body(doc, "The leave-one-seed-out reference is cross-fitted only at the repeat level and still reuses public endpoint data. Hierarchical diversity intervals used 100 replicates, sufficient for a sensitivity interval but not for fine tail estimation. The multiview source directory lacked molecule-level prediction exports and split-assignment hashes, so paired-score hashes could be verified but prediction hashes could not. Exact repeated regression scores also reduced the effective information in nominal seed replication. Finally, the matched K = 3 alternatives used fixed representation-to-learner assignments rather than exhaustive selection from every possible three-candidate subset.")
    body(doc, "Secondary TDC, automated machine-learning, conformal, MoleculeACE, bRo5, duplicate-sensitivity and failure-case analyses are retrospective boundary evidence. They should not be interpreted as new prospective cohorts or as proof that one selection strategy transfers to all molecular endpoints.")

    add_heading(doc, "5 Conclusions", 1)
    body(doc, "Within nine repeated nested molecular-property audits, nominal candidate expansion exceeded the effective dimensionality of candidate utility patterns under raw, common-difficulty-adjusted, fixed-reference and rank-based estimators. Chance-adjusted validation ranking weakened with K. Same-unit gaps increased in eight endpoints, while leave-one-seed-out cross-fitted effects remained positive in seven and retained BBBP and P-gp as negative directions. Finite-audit maxima contained winner optimism, and matched-size multiview gains were heterogeneous rather than automatic.")
    body(doc, "Molecular model-selection studies should jointly report candidate eligibility, nominal pool size, utility-pattern diversity with uncertainty, chance-adjusted ranking fidelity, endpoint-specific same-unit and cross-fitted gaps, computational exposure, failed candidates and chemical-support boundaries. Matched-size controls are needed before attributing full-pool gains to representation breadth alone, and repeated identical split outcomes should not be interpreted as independent precision.")

    add_heading(doc, "Supplementary Information", 1)
    body(doc, "Additional file 1. Supplementary Methods and Results. Extended methods, secondary analyses and boundary results.")
    body(doc, "Additional file 2. Supplementary Tables. Machine-readable endpoint, candidate, fold-level and sensitivity-analysis tables (Tables S1-S17).")
    body(doc, "Additional file 3. Supplementary Figures. Extended diagnostic, robustness and chemical-boundary figures (Figures S1-S14).")

    add_heading(doc, "List of abbreviations", 1)
    body(doc, source.paragraphs[144].text)
    add_heading(doc, "Declarations", 1)
    add_heading(doc, "Ethics approval and consent to participate", 2); body(doc, "Not applicable.")
    add_heading(doc, "Consent for publication", 2); body(doc, "Not applicable.")
    add_heading(doc, "Availability of data and materials", 2); body(doc, "Public dataset provenance is listed in Additional file 2: Table S1. Derived fold-level tables, source hashes and analysis code are supplied in the accompanying submission package.")
    add_heading(doc, "Competing interests", 2); body(doc, "No competing-interest statement was present in the source manuscript.")
    add_heading(doc, "Funding", 2); body(doc, "No funding statement was present in the source manuscript.")
    add_heading(doc, "Authors' contributions", 2); body(doc, "Author contribution roles were not present in the source manuscript.")
    add_heading(doc, "Acknowledgements", 2); body(doc, "No acknowledgement statement was present in the source manuscript.")

    add_heading(doc, "References", 1)
    for text in source_paragraphs(source, 146, len(source.paragraphs)):
        p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(-0.63); p.paragraph_format.left_indent = Cm(0.63); p.paragraph_format.space_after = Pt(2); set_run_font(p.add_run(text), 9)
    return doc


def audit_document(doc: Document) -> dict[str, object]:
    texts = [p.text for p in doc.paragraphs]
    text = "\n".join(texts)
    abstract_text = " ".join(texts[texts.index("Abstract") + 1:texts.index("1 Introduction")])
    main_start = texts.index("Abstract") + 1
    main_end = texts.index("Supplementary Information")
    main_text = " ".join(texts[main_start:main_end])
    word_count = len(re.findall(r"\b[\w'-]+\b", main_text))
    abstract_count = len(re.findall(r"\b[\w'-]+\b", abstract_text))
    figure_refs = sorted(set(int(x) for x in re.findall(r"Figure (\d+)", text)))
    table_refs = sorted(set(int(x) for x in re.findall(r"Table (\d+)", text)))
    forbidden_hits = {term: text.lower().count(term.lower()) for term in FORBIDDEN if term.lower() in text.lower()}
    max_table_columns = max(len(t.columns) for t in doc.tables)
    return {
        "title": TITLE, "abstract_words": abstract_count, "abstract_within_350": abstract_count <= 350,
        "abstract_through_conclusions_words": word_count, "tables": len(doc.tables),
        "max_table_columns": max_table_columns, "main_figures": 6,
        "figure_numbers_present": figure_refs, "table_numbers_present": table_refs,
        "forbidden_term_hits": forbidden_hits,
        "requires_source_data": [
            "Author names, affiliations, corresponding author and ORCID",
            "Competing interests, funding, author contributions and acknowledgements",
            "Multiview candidate-level prediction hash and explicit split-assignment hash",
            "Exact historical source file 小论文-18(4).docx was not available; closest archived 小论文-18.docx was not used to overwrite the main text",
        ],
    }


def main() -> None:
    doc = build_document()
    doc.save(PAPER)
    shutil.copy2(PAPER, DESKTOP)
    audit = audit_document(doc)
    AUDIT.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"paper": str(PAPER), "desktop": str(DESKTOP), **audit}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
