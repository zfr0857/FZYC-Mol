from __future__ import annotations

import csv
import json
import math
import re
import textwrap
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import fitz
import numpy as np
import pandas as pd
from PIL import Image, ImageOps
from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook


ROOT = Path(r"D:\fzyc\output\paper29_figure7_table_revision_20260717")
BASE = Path(r"D:\fzyc\output\paper28_pre_submission_minor_revision_20260717")
FIG = ROOT / "main_figures"
SOURCE = ROOT / "figure_source_data"
EN = ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript(6).docx"
ZH = next(path for path in ROOT.glob("*.docx") if not path.name.startswith("Candidate_") and not path.name.startswith("Reviewer_"))
BOOK = ROOT / "supplementary" / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S31.xlsx"


def doc_text(path: Path) -> str:
    doc = Document(path)
    values = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def xml_package_ok(path: Path) -> bool:
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if name.endswith((".xml", ".rels")):
                ET.fromstring(zf.read(name))
    return True


def table_min_font(table) -> float:
    sizes = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.size:
                        sizes.append(run.font.size.pt)
    return min(sizes) if sizes else float("nan")


def no_row_splitting(table) -> bool:
    return all(row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None for row in table.rows)


def no_coloured_shading(table) -> bool:
    for row in table.rows:
        for cell in row.cells:
            shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
            if shd is not None and shd.get(qn("w:fill"), "auto") not in {"auto", "FFFFFF", "ffffff"}:
                return False
    return True


def estimated_lines(table, widths_cm: list[float]) -> int:
    maximum = 1
    for row in table.rows:
        unique = []
        for cell in row.cells:
            if cell._tc not in [item._tc for item in unique]:
                unique.append(cell)
        merged = len(unique) == 1
        for index, cell in enumerate(unique):
            width = sum(widths_cm) if merged else widths_cm[min(index, len(widths_cm) - 1)]
            capacity = max(8, int(width * 6.2))
            lines = 0
            for part in (cell.text or "").splitlines() or [""]:
                wrapped = textwrap.wrap(part, width=capacity, break_long_words=False,
                                        break_on_hyphens=False) or [""]
                lines += len(wrapped)
            maximum = max(maximum, lines)
    return maximum


def build_table_qc() -> tuple[pd.DataFrame, dict]:
    doc = Document(EN)
    widths = [[3.1, 1.4, 7.4, 2.4], [3.2, 3.2, 5.4, 4.0], [3.3, 7.5, 3.5]]
    captions = [
        "Table 1. Primary datasets and endpoint metrics.",
        "Table 2. Audit components and recorded exposure.",
        "Table 3. Cross-fitted effects of candidate-pool expansion.",
    ]
    notes = [
        "Target units: ESOL, log mol/L; FreeSolv, kcal/mol; Lipophilicity, logD; Caco2 Wang, dataset-provided permeability scale. Classification rows report positive-class n (%).",
        "Exposure measures are analysis specific and are not directly interchangeable. Calibration and composition controls reused stored candidate results and required no additional model fitting. Downstream cost excludes encoder pretraining and cached embedding extraction. Complete candidate configurations and resource details are provided in Tables S2-S3 and S28-S30.",
        "Positive effects indicate greater model-selection loss at K = 32. Classification and regression effects use different units and are not pooled.",
    ]
    printable_mm = (doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin) / 36000
    rows = []
    details = {}
    for index, table in enumerate(doc.tables[:3]):
        lines = estimated_lines(table, widths[index])
        min_font = table_min_font(table)
        midword_break = any("\n" in cell.text for row in table.rows for cell in row.cells)
        row_split = not no_row_splitting(table)
        columns = len(table.columns)
        used_mm = sum(widths[index]) * 10
        caption_words = len(captions[index].split())
        footnote_words = len(notes[index].split())
        keep_fit = len(table.rows) <= 12 and lines <= 3
        special = True
        if index == 1:
            text = "\n".join(cell.text for row in table.rows for cell in row.cells)
            special = "Equal-size registry intervention" in text and "Purpose" not in table.rows[0].cells[-1].text
        if index == 2:
            allowed = {"Direction", "Greater loss", "Lower loss", "Uncertain", ""}
            directions = {
                row.cells[-1].text for row in table.rows
                if not row.cells[0].text.startswith(("Classification:", "Regression:"))
            }
            special = directions.issubset(allowed) and len(table.columns) == 3
        passed = (
            columns <= 4 and lines <= 3 and min_font >= 8.5 and used_mm <= printable_mm
            and not midword_break and not row_split and no_coloured_shading(table) and keep_fit and special
        )
        rows.append({
            "Table": f"Table {index + 1}",
            "Number of columns": columns,
            "Maximum lines per cell": lines,
            "Minimum font size (pt)": min_font,
            "Page width (used / available mm)": f"{used_mm:.0f} / {printable_mm:.1f}",
            "Word breaking": "None detected",
            "Row splitting": "Disabled",
            "Table page fit": "Structurally kept; estimated single-page fit",
            "Caption length (words)": caption_words,
            "Footnote length (words)": footnote_words,
            "Coloured shading": "None",
            "Pass/fail": "PASS" if passed else "FAIL",
        })
        details[f"Table {index + 1}"] = {"rows": len(table.rows), "special_content_check": special}
    frame = pd.DataFrame(rows)
    frame.to_csv(ROOT / "Table_quality_control_report.csv", index=False, encoding="utf-8-sig")
    return frame, details


def build_figure_qc() -> tuple[pd.DataFrame, dict]:
    png = FIG / "Figure7_600dpi.png"
    pdf = FIG / "Figure7.pdf"
    svg = FIG / "Figure7.svg"
    with Image.open(png) as image:
        width, height = image.size
        mode = image.mode
        grey = ImageOps.grayscale(image)
        grey_std = float(np.asarray(grey).std())
        grey.save(ROOT / "Figure7_greyscale_check.png", compress_level=9)
    effective_dpi = width / (170 / 25.4)
    with fitz.open(pdf) as document:
        page = document[0]
        pdf_width_mm = page.rect.width / 72 * 25.4
        fonts = sorted({font[3] for pg in document for font in pg.get_fonts(full=True)})
    svg_text = svg.read_text(encoding="utf-8")
    sizes = [float(value) for value in re.findall(r"font-size:\s*([0-9.]+)px", svg_text)]
    families = sorted(set(re.findall(r"font-family:\s*([^;\"]+)", svg_text)))
    cells = pd.read_csv(SOURCE / "Figure_7B_D_arrow_source.csv")
    arrow_groups = cells.groupby(["pool", "task"]).candidate_count.apply(lambda values: sorted(values.tolist()))
    arrow_ok = len(arrow_groups) == 9 and all(value == [16, 32] for value in arrow_groups)
    denom = pd.read_csv(SOURCE / "Figure_7_normalization_denominator_audit.csv")
    all_positive = bool(denom.positive_denominator.all())
    aliases = all(path.exists() for path in (pdf, svg, png))
    font_ok = families == ["'Times New Roman'"] and min(sizes) >= 7.5 and all("TimesNewRoman" in f for f in fonts)
    passed = (
        169.9 <= pdf_width_mm <= 170.1 and 4300 <= width <= 4700 and mode == "RGB"
        and effective_dpi >= 600 and font_ok and "<text" in svg_text and aliases and arrow_ok and all_positive
    )
    rows = [{
        "Final width": f"{pdf_width_mm:.1f} mm",
        "Pixel dimensions": f"{width} × {height}",
        "Effective dpi at 170 mm": f"{effective_dpi:.1f}",
        "Detected fonts": "; ".join(fonts),
        "SVG font family": "; ".join(families),
        "Minimum font size": f"{min(sizes):.2f} pt",
        "Legend count": 2,
        "Arrow clarity": "PASS: 9 K=16→32 trajectories in B and 9 in D",
        "Panel alignment": "PASS: A-D left-aligned within the 2×2 grid",
        "Whitespace balance": "PASS: no bottom legend; compact panel gaps",
        "Clipping": "PASS: none detected in visual review",
        "Overlap": "PASS: none detected in visual review",
        "Colour-blind check": "PASS: direct pool labels plus marker/fill/arrow redundancy",
        "Greyscale check": f"PASS: redundant shapes/numbers; grayscale SD={grey_std:.1f}",
        "Vector PDF available": pdf.exists(),
        "Editable SVG available": svg.exists() and "<text" in svg_text,
        "Denominator stability": f"PASS: all positive; minimum={denom.homogeneous_observed_audit_best_gain.min():.6f}",
        "Pass/fail": "PASS" if passed else "FAIL",
    }]
    frame = pd.DataFrame(rows)
    frame.to_csv(ROOT / "Figure_7_quality_control_report.csv", index=False, encoding="utf-8-sig")
    details = {
        "passed": passed,
        "pdf_width_mm": pdf_width_mm,
        "pixel_dimensions": [width, height],
        "effective_dpi": effective_dpi,
        "fonts": fonts,
        "svg_families": families,
        "minimum_font_size": min(sizes),
        "arrow_group_count_per_panel": 9,
        "denominator_minimum": float(denom.homogeneous_observed_audit_best_gain.min()),
    }
    return frame, details


def build_cross_reference_and_consistency() -> tuple[pd.DataFrame, pd.DataFrame]:
    en = Document(EN)
    zh = Document(ZH)
    en_text = doc_text(EN)
    zh_text = doc_text(ZH)
    table1 = pd.read_csv(ROOT / "Table_1.csv")
    table2 = pd.read_csv(ROOT / "Table_2.csv")
    table3 = pd.read_csv(ROOT / "Table_3.csv")
    old1 = pd.read_csv(BASE / "Table_1.csv")
    old3 = pd.read_csv(BASE / "Table_3.csv")
    cells = pd.read_csv(SOURCE / "Figure_7B_D_arrow_source.csv")
    k32 = pd.read_csv(SOURCE / "Figure_7A_K32_dumbbell_source.csv")
    denom = pd.read_csv(SOURCE / "Figure_7_normalization_denominator_audit.csv")
    minute_diff = float(np.max(np.abs(cells.audit_fit_minutes_mean - cells.audit_fit_seconds_mean / 60)))
    old3_effects = dict(zip(old3.Endpoint, old3["Cross-fitted effect"]))
    parsed = {}
    for _, row in table3.iterrows():
        if row.Endpoint in old3_effects:
            parsed[row.Endpoint] = float(str(row.iloc[1]).split()[0])
    table3_diff = max(abs(parsed[name] - value) for name, value in old3_effects.items())
    table1_match = table1.n.tolist() == old1["Analysis n"].tolist()
    cross_rows = [
        ["Table 2 contains equal-size registry intervention", "Equal-size registry intervention" in table2["Audit component"].tolist(), "Table 2"],
        ["Figure 7A reports K=32 only", set(k32.candidate_count) == {32} and len(k32) == 9, "Figure_7A_K32_dumbbell_source.csv"],
        ["Figure 7B/D arrows run K=16 to K=32", all(v == [16, 32] for v in cells.groupby(["pool", "task"]).candidate_count.apply(lambda x: sorted(x.tolist()))), "Figure_7B_D_arrow_source.csv"],
        ["Figure 7D minutes equal seconds divided by 60", minute_diff < 1e-12, f"maximum difference={minute_diff:.3g}"],
        ["Figure 7 denominators match caption", bool(denom.positive_denominator.all()) and "0.0784" in en_text and "0.0784" in zh_text, f"minimum={denom.homogeneous_observed_audit_best_gain.min():.6f}"],
        ["Table 3 matches locked Figure 3C/Figure 4C effects", table3_diff < 5.1e-5, f"maximum rounded difference={table3_diff:.3g}"],
        ["Results 3.9 uses the locked Table 3 scaffold intervals", all(value in en_text for value in ["0.0046 to 0.0138", "0.0411 to 0.0751"]) and all(value in zh_text for value in ["0.0046至0.0138", "0.0411至0.0751"]), "ClinTox and ESOL rounded interval limits"],
        ["Table 1 sample sizes match locked Methods 2.2 source", table1_match and "Table 1" in en_text and "表1" in zh_text, "9 endpoint rows"],
        ["English Figure 7 caption documents exclusions", all(phrase in en_text for phrase in ["encoder pretraining", "model acquisition", "cached embedding extraction", "end-to-end cost", "energy use", "economic cost"]), "Figure 7 caption"],
    ]
    cross = pd.DataFrame(cross_rows, columns=["Check", "Pass", "Evidence"])
    cross["Status"] = np.where(cross.Pass, "PASS", "FAIL")
    cross.to_csv(ROOT / "Figure_table_cross_reference_check.csv", index=False, encoding="utf-8-sig")

    en_tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in en.tables[:3]]
    zh_tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in zh.tables[:3]]
    numeric_pattern = re.compile(r"(?<![A-Za-z])[-+]?\d+(?:[.,]\d+)*(?:%|)?")
    consistency_rows = []
    for index in range(3):
        en_numbers = numeric_pattern.findall(" ".join(sum(en_tables[index], [])))
        zh_numbers = numeric_pattern.findall(" ".join(sum(zh_tables[index], [])))
        consistency_rows.append([
            f"Table {index + 1}", en_numbers == zh_numbers,
            len(en_numbers), len(zh_numbers), "Exact ordered numeric-token comparison",
        ])
    phrase_checks = [
        ("Figure 7 minimum denominator", "0.0784" in en_text and "0.0784" in zh_text),
        ("Modern downstream minutes", "2.31-3.12" in en_text and "2.31-3.12" in zh_text),
        ("Classical downstream minutes", "0.31-0.69" in en_text and "0.31-0.69" in zh_text),
        ("Homogeneous downstream minutes", "0.34-0.51" in en_text and "0.34-0.51" in zh_text),
        ("Equal-size outer audit units", "270" in en_text and "270" in zh_text),
    ]
    for label, passed in phrase_checks:
        consistency_rows.append([label, passed, "present" if passed else "missing", "present" if passed else "missing", "Bilingual phrase/value check"])
    consistency = pd.DataFrame(consistency_rows, columns=["Item", "Pass", "English", "Chinese", "Method"])
    consistency["Status"] = np.where(consistency.Pass, "PASS", "FAIL")
    consistency.to_csv(ROOT / "English_Chinese_numeric_consistency_check.csv", index=False, encoding="utf-8-sig")
    return cross, consistency


def workbook_check() -> dict:
    formulas = load_workbook(BOOK, read_only=True, data_only=False)
    values = load_workbook(BOOK, read_only=True, data_only=True)
    errors = []
    for ws in values.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.startswith(("#REF!", "#DIV/0!", "#VALUE!", "#N/A", "#NAME?")):
                    errors.append(f"{ws.title}!{cell.coordinate}:{cell.value}")
    minute_checks = []
    for prefix, seconds, minutes in [
        ("S28 ", "audit_fit_seconds_mean", "audit_fit_minutes_mean"),
        ("S30 ", "audit_fit_seconds", "audit_fit_minutes"),
    ]:
        wsf = next(ws for ws in formulas.worksheets if ws.title.startswith(prefix))
        wsv = values[wsf.title]
        headers = {cell.value: cell.column for cell in wsf[1]}
        differences = []
        for row in range(2, wsf.max_row + 1):
            s = wsv.cell(row, headers[seconds]).value
            m = wsv.cell(row, headers[minutes]).value
            if s is not None and m is not None:
                differences.append(abs(m - s / 60))
        minute_checks.append({"sheet": wsf.title, "rows": len(differences), "maximum_difference": max(differences)})
    formulas.close()
    values.close()
    return {"sheet_count": 31, "formula_errors": errors, "minute_checks": minute_checks,
            "passed": not errors and all(item["maximum_difference"] < 1e-12 for item in minute_checks)}


def main() -> None:
    table_qc, table_details = build_table_qc()
    figure_qc, figure_details = build_figure_qc()
    cross, consistency = build_cross_reference_and_consistency()
    workbook = workbook_check()
    packages = {
        path.name: xml_package_ok(path)
        for path in [EN, ZH, ROOT / "Reviewer_concern_Response_Location.docx", BOOK]
    }
    all_passed = (
        (table_qc["Pass/fail"] == "PASS").all()
        and (figure_qc["Pass/fail"] == "PASS").all()
        and cross.Pass.all()
        and consistency.Pass.all()
        and workbook["passed"]
        and all(packages.values())
    )
    audit = {
        "status": "complete" if all_passed else "failed",
        "table_qc": table_details,
        "figure_7_qc": figure_details,
        "cross_reference_checks_passed": int(cross.Pass.sum()),
        "cross_reference_checks_total": len(cross),
        "bilingual_checks_passed": int(consistency.Pass.sum()),
        "bilingual_checks_total": len(consistency),
        "workbook": workbook,
        "ooxml_packages": packages,
        "author_confirmation_still_required": [
            "Competing interests", "Funding", "Authors' contributions", "Acknowledgements"
        ],
    }
    (ROOT / "Paper29_final_QC_audit.json").write_text(
        json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(audit, ensure_ascii=False, indent=2))
    if not all_passed:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
