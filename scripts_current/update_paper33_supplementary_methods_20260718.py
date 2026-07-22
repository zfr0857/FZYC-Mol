from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"D:\fzyc\output\paper33_final_minor_revision_20260718")
SOURCE = Path(r"D:\fzyc\output\paper30_submission_package_20260717\supplementary\Additional_file_1_Supplementary_Methods_and_Results.docx")
TARGET = ROOT / "supplementary" / "Additional_file_1_Supplementary_Methods_and_Results.docx"


def find(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(prefix)


def replace(doc: Document, prefix: str, text: str) -> None:
    find(doc, prefix).text = text


def remove(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def replace_between(doc: Document, start: str, end: str, blocks: list[tuple[str, str]]) -> None:
    start_p = find(doc, start)
    end_p = find(doc, end)
    paragraphs = doc.paragraphs
    i = next(index for index, p in enumerate(paragraphs) if p._p is start_p._p)
    j = next(index for index, p in enumerate(paragraphs) if p._p is end_p._p)
    for paragraph in paragraphs[i + 1:j]:
        remove(paragraph)
    for style, text in blocks:
        end_p.insert_paragraph_before(text, style=style)


def set_fonts(doc: Document) -> None:
    for style in doc.styles:
        if style.type != 1:
            continue
        style.font.name = "Times New Roman"
        if style.name == "Normal":
            style.font.size = Pt(10)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            rfonts.set(qn("w:ascii"), "Times New Roman")
            rfonts.set(qn("w:hAnsi"), "Times New Roman")


def main() -> None:
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, TARGET)
    doc = Document(TARGET)
    replace(doc, "The primary evidence is", (
        "Evidence was retained in five levels that answer different questions and are not combined into one leaderboard: "
        "Level 1, the nine-endpoint controlled prefix audit; Level 2, permutation calibration, composition controls and finite-audit simulation; "
        "Level 3, the matched-size twelve-candidate multiview audit; Level 4, the six-endpoint expanded registry-composition intervention with component, budget, anchor, normalization, stability and split-mechanism sensitivity analyses; "
        "and Level 5, prediction-level reliability and chemical-support boundary evidence. Historical TDC, AutoGluon, conformal, MoleculeACE, bRo5, deduplication and failure-case records remain supplementary context rather than independent confirmation."
    ))
    replace_between(doc, "S15 Equal-size", "S16 Source limitations", [
        ("Normal", "The expanded intervention fixed six endpoints (ClinTox, BACE, BBBP, ESOL, Lipophilicity and Caco2_Wang), three locked registries and exact prefixes K = 4, 8, 16 and 32. Five split seeds and three outer folds yielded 1,080 primary outer units. Unit-level records contain 64,616.35 downstream fit/predict seconds; the total expanded-intervention fit count was not reconstructed."),
        ("Normal", "The homogeneous registry used Morgan learner/tuning variants, the classical multiview registry combined Morgan, MACCS, RDKit descriptors and concatenated representations, and the modern-augmented registry interleaved locked frozen ChemBERTa and MoLFormer representations plus a separately locked one-epoch D-MPNN. The D-MPNN candidate belonged only to this expanded intervention and not to the limited four-model prediction panel. Fixed-K component ablations at K = 16 and 32 replaced prespecified classical candidates so that K remained exact."),
        ("Normal", "Anchor sensitivity used three prespecified anchors: the shared Morgan linear candidate, a fixed Morgan random forest and the predefined registry-median candidate. Normalization sensitivity retained raw gains, endpoint-MAD normalization and paired homogeneous-audit-best normalization. The paired normalization was undefined when homogeneous finite-audit opportunity was at most 1e-12; such cells were retained as missing rather than regularized or imputed."),
        ("Normal", "For equal-downstream-budget sensitivity, each prefix was truncated at the endpoint- and K-specific median downstream time of the classical multiview registry; outer performance never determined retention. Recorded downstream time includes inner and outer downstream fit/predict time and excludes model acquisition, encoder pretraining and cached embedding extraction. The Pareto frontier was constructed descriptively by sorting eligible pool–K–design points by mean downstream time and retaining each point only when its mean normalized selected gain exceeded the running maximum. It is a bounded downstream-compute analysis, not an end-to-end efficiency claim."),
        ("Normal", "Selection stability used candidate selection frequencies π_j, normalized selection entropy, modal proportion, leave-one-seed-out agreement and adjacent-fold switches. The three-endpoint split loop repeated all three registries under seeded scaffold and Tanimoto-component splits with endpoint, seed, fold, K, selection and metric definitions held fixed. Complete registries, unit-level records, component results, anchor and normalization sensitivity, split results, stability and budget quantities are provided in Tables S32–S36 and Figures S16–S21."),
    ])
    replace(doc, "S15 Equal-size", "S15 Expanded registry-composition intervention")
    table_heading = find(doc, "Supplementary figure directory")
    table_entries = [
        "Table S32. Expanded-intervention registry, exact prefix order, representation and source.",
        "Table S33. Modern-component ablations and component selection frequencies.",
        "Table S34. Anchor and normalization sensitivity with direction concordance.",
        "Table S35. Composition-by-split loop, direction concordance and split-specific selection stability.",
        "Table S36. Selection stability, candidate frequencies, equal-budget stability and equal-budget diversity.",
    ]
    for text in table_entries:
        table_heading.insert_paragraph_before(text, style="Normal")

    replace_between(doc, "Supplementary figure directory", "Figure S16.", [
        ("Normal", "Figure S1. Cleaning flow."),
        ("Normal", "Figure S2. Candidate correlation matrices."),
        ("Normal", "Figure S3. Eigenvalue spectra."),
        ("Normal", "Figure S4. Raw ranking metrics."),
        ("Normal", "Figure S5. Permutation controls."),
        ("Normal", "Figure S6. Signal recovery."),
        ("Normal", "Figure S7. Full winner-optimism simulation."),
        ("Normal", "Figure S8. Fold-level audit gaps."),
        ("Normal", "Figure S9. Full matched-K subset distributions."),
        ("Normal", "Figure S10. Representation selection frequency."),
        ("Normal", "Figure S11. TDC."),
        ("Normal", "Figure S12. Conformal and risk-coverage."),
        ("Normal", "Figure S13. MoleculeACE."),
        ("Normal", "Figure S14. bRo5."),
        ("Normal", "Figure S15. Failure cases."),
    ])
    # Replace the three obsolete old captions and append the new final set.
    replace(doc, "Figure S16.", "Figure S16. All endpoint–pool–K composition effects.")
    replace(doc, "Figure S17.", "Figure S17. Modern-component ablations.")
    replace(doc, "Figure S18.", "Figure S18. Anchor and normalization sensitivity.")
    doc.add_paragraph("Figure S19. Composition-by-split-regime transport.", style="Normal")
    doc.add_paragraph("Figure S20. Candidate selection frequency and normalized entropy.", style="Normal")
    doc.add_paragraph("Figure S21. Equal-K and equal-downstream-budget comparison.", style="Normal")

    set_fonts(doc)
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
