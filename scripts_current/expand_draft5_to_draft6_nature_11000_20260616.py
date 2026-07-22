# -*- coding: utf-8 -*-
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports" / "draft6_nature_11000_20260616"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def locate_source() -> Path:
    candidates = [p for p in (Path.home() / "Desktop").rglob("初稿-5_表格修订版.docx")]
    if candidates:
        return max(candidates, key=lambda p: p.stat().st_mtime)
    fallback = [p for p in (Path.home() / "Desktop").rglob("初稿-5.docx")]
    if fallback:
        return max(fallback, key=lambda p: p.stat().st_mtime)
    raise FileNotFoundError("Could not locate 初稿-5_表格修订版.docx or 初稿-5.docx.")


SRC_DOCX = locate_source()
DEST_DOCX = SRC_DOCX.parent / "初稿-6.docx"
REPORT_DOCX = SRC_DOCX.parent / "初稿-6_Nature扩写与QA报告.docx"


def style_run(run) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def set_para_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    style_run(run)
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(3)


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    try:
        new_para.style = paragraph.style
    except Exception:
        pass
    set_para_text(new_para, text)
    return new_para


def find_prefix(doc: Document, prefix: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    return None


def find_exact(doc: Document, text: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def replace_prefix(doc: Document, prefix: str, text: str, changes: list[str], label: str) -> None:
    paragraph = find_prefix(doc, prefix)
    if paragraph is None:
        changes.append(f"未匹配：{label}")
        return
    set_para_text(paragraph, text)
    changes.append(label)


def add_after_prefix(doc: Document, prefix: str, additions: list[str], changes: list[str], label: str) -> None:
    paragraph = find_prefix(doc, prefix)
    if paragraph is None:
        changes.append(f"未匹配插入点：{label}")
        return
    anchor = paragraph
    for text in additions:
        anchor = insert_after(anchor, text)
    changes.append(label)


def add_after_exact(doc: Document, heading: str, additions: list[str], changes: list[str], label: str) -> None:
    paragraph = find_exact(doc, heading)
    if paragraph is None:
        changes.append(f"未匹配插入点：{label}")
        return
    anchor = paragraph
    for text in additions:
        anchor = insert_after(anchor, text)
    changes.append(label)


def set_border(element, edge: str, val: str, size: str = "8", color: str = "000000") -> None:
    tag = "w:tblBorders" if element.tag.endswith("tblPr") else "w:tcBorders"
    borders = element.find(qn(tag))
    if borders is None:
        borders = OxmlElement(tag)
        element.append(borders)
    edge_el = borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        borders.append(edge_el)
    edge_el.set(qn("w:val"), val)
    edge_el.set(qn("w:sz"), size)
    edge_el.set(qn("w:space"), "0")
    edge_el.set(qn("w:color"), color)


def keep_table_compact(doc: Document) -> None:
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        for edge in ["top", "bottom"]:
            set_border(tbl_pr, edge, "single", "12")
        for edge in ["left", "right", "insideH", "insideV"]:
            set_border(tbl_pr, edge, "nil")
        for row_idx, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))
            if row_idx == 0 and tr_pr.find(qn("w:tblHeader")) is None:
                tr_pr.append(OxmlElement("w:tblHeader"))
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                for edge in ["left", "right", "insideH", "insideV"]:
                    set_border(tc_pr, edge, "nil")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def expand_manuscript(doc: Document) -> list[str]:
    changes: list[str] = []

    replace_prefix(
        doc,
        "分子性质预测已成为药物发现和 ADMET 评估的基础工具",
        "分子性质预测已成为药物发现和 ADMET 评估的基础工具，但公开基准上的高平均分并不总能转化为可靠的化学决策。随机划分、低相似度分子、不平衡毒性标签、规则五以外化学空间和活性悬崖，都会削弱单一 ROC-AUC 或 RMSE 对真实应用风险的解释力。本研究提出 FZYC-Mol，一种由验证集治理驱动的适用域感知框架。该框架将分子图、指纹、二维描述符、片段/骨架信息和冻结预训练表征登记为候选证据，并仅允许通过验证集门控的专家模型、融合策略、补救头和适用域模块进入最终测试。",
        changes,
        "摘要：保持紧凑，不为字数牺牲 Nature 风格",
    )

    add_after_prefix(
        doc,
        "分子性质预测正在从药物发现早期筛选工具转变",
        [
            "这种可靠性问题在药物发现早期尤其突出。早期筛选通常面对大量候选分子和有限实验资源，模型错误不只是数值偏差，还会改变后续合成、实验排队和风险评估的优先级。对于溶解度、脂溶性、血脑屏障通透性、毒性和转运体相关终点，研究者往往需要在结构相似但性质差异较大的化合物之间做出选择。若模型只在随机划分下表现良好，却无法说明预测何时超出适用域，模型分数就难以转化为可执行的化学证据。",
            "因此，本研究关注的不是单一模型结构是否在某个排行榜上取得最高分，而是模型选择本身能否被审计。一个面向真实应用的分子性质预测框架应同时回答四个问题：候选模型如何产生，最终策略如何被选择，测试集是否只在策略冻结后使用，以及失败样本能否被识别并解释。将这些问题显式写入评价流程，有助于把模型开发从结果驱动的比较转向过程可复核的可靠性评估。",
        ],
        changes,
        "引言：补充真实应用风险和研究问题",
    )

    add_after_prefix(
        doc,
        "这一缺口来自评价方式与应用场景之间的错位",
        [
            "已有分子机器学习研究提供了丰富的表示学习和模型训练策略，包括基于指纹的传统机器学习、消息传递神经网络、分子语言模型和多任务预训练模型。它们在不同数据集上各有优势，但优势往往依赖终点性质、数据规模、划分方式和评价指标。对于小样本回归任务，简单树模型或指纹专家可能比复杂模型更稳定；对于类别不平衡毒性任务，ROC-AUC 的提升也可能不能反映阳性样本召回。因此，公平比较不应只比较模型名称，而应把划分、候选登记、指标方向和选择规则同时固定。",
            "另一个容易被忽视的问题是事后选择偏差。随着候选模型、目标变换、集成方式和补救头数量增加，研究者即使没有直接查看测试标签，也可能通过反复试探逐渐把测试集变成隐性开发集。FZYC-Mol 将这一风险作为方法设计对象，而不是在结果之后补充说明。候选策略必须先进入登记表，再通过验证集证据决定是否保留；没有通过验证集门控的模块即使在个别测试划分上表现较好，也不进入主性能结论。",
        ],
        changes,
        "引言：补充既有方法局限与事后选择偏差",
    )

    add_after_prefix(
        doc,
        "本研究的贡献集中在三个层面",
        [
            "这一设计使 FZYC-Mol 更接近一套模型治理流程，而不仅是一个预测器集合。它把常规性能比较、外部 ADMET 测试、适用域分析、保形预测、活性悬崖审计和负结果记录放在同一叙事中，目的是让读者看到每个结论由哪些证据支撑、哪些证据不足，以及哪些模块只适合作为后续验证方向。这样的定位也决定了本文的主张边界：FZYC-Mol 可提高部分终点的最终保留性能，并能系统暴露模型边界，但不被表述为在所有分子性质任务上都优于所有强基线的通用模型。",
        ],
        changes,
        "引言：补充贡献边界",
    )

    add_after_prefix(
        doc,
        "本研究使用 MoleculeNet、Therapeutics Data Commons ADMET",
        [
            "数据层面的设计遵循两个原则。第一，主结果只使用公开可获得、可复核的数据来源，避免把未公开数据或未完成实验写入完成性结论。第二，同一数据集在不同分析模块中保持相同任务定义，回归任务和分类任务分别使用预先定义的主指标，辅助指标只用于解释可靠性、校准或不平衡标签影响。对于 bRo5 相关数据，本研究仅将其作为规则五以外化学空间的公共压力测试，相关结论限定在可获得公开数据上的外推与适用域审计。",
            "在所有任务中，训练集、验证集和测试集承担不同功能。训练集用于拟合候选专家，验证集用于选择策略、阈值、权重和补救头，测试集只用于策略冻结后的最终评估。该划分使每个结果都能追溯到明确的决策时点。若一个模块只在测试后才显示潜在价值，它会被记录为负结果或后续验证接口，而不会被纳入主性能主张。",
        ],
        changes,
        "方法：扩展数据与测试冻结原则",
    )

    add_after_prefix(
        doc,
        "为模拟不同应用难度，本研究采用多层次划分策略",
        [
            "这些划分并非简单重复评价，而是对应不同使用情境。Random split 主要反映同分布插值能力，scaffold split 更接近新骨架外推，structure-separated split 强化结构隔离，low-similarity hard subset 则用于考察训练集近邻不足时的性能下降。若一个模型只在随机划分中表现突出，却在 scaffold 或低相似度子集上明显退化，其结果不应被解释为真实药物发现中的广泛泛化能力。",
            "划分策略还决定了适用域分析的解释方式。对于每个测试分子，本研究记录其与训练集及验证集最近邻的 Morgan Tanimoto 相似度，并结合骨架信息和风险分位解释错误分布。这样做的目的不是把相似度阈值作为绝对拒用规则，而是为读者提供一个可复核的证据层：当模型在低相似度分子上失效时，失败可以被定位到结构外推而不是简单归因于模型不足。",
        ],
        changes,
        "方法：扩展划分策略与适用域联系",
    )

    add_after_prefix(
        doc,
        "FZYC-Mol 将分子表示划分为五类",
        [
            "多视图表示的作用是覆盖不同层次的化学信息，而不是机械地增加特征数量。分子图表示适合描述原子和键的局部连接，Morgan 和 MACCS 等指纹适合捕捉可复现的局部子结构，RDKit 描述符提供物化性质与二维拓扑摘要，BRICS 与 Murcko scaffold 支持片段和骨架层面的解释，冻结预训练表征则提供从大规模分子语料中学习到的上下文表示。不同视图进入同一候选池后，只有通过验证集证据的视图或融合方式才会被保留。",
            "这种设计也为负结果提供了可解释空间。如果某一终点中图模型、指纹专家和冻结表征头产生相近性能，最终保留策略可以选择更稳定或更简单的候选；如果某一终点只有补救头在验证集上形成一致增益，该模块才会进入最终结果。换言之，多视图框架并不预设复杂模型一定更优，而是让表示选择成为可审计的经验问题。",
        ],
        changes,
        "方法：扩展多视图表示的设计理由",
    )

    add_after_prefix(
        doc,
        "候选专家池包括 RF、ExtraTrees",
        [
            "候选登记表是 FZYC-Mol 的核心对象。每个候选策略在测试前记录其输入表示、学习器、目标变换、随机种子、指标方向和选择条件。对于回归任务，主指标为 RMSE，候选比较以降低误差为目标；对于分类任务，主指标为 ROC-AUC，同时保留 PR-AUC、校准误差和固定精度召回等辅助指标。辅助指标不直接替代主指标，但用于解释不平衡标签和筛选场景中的风险。",
            "融合候选在验证集预测矩阵上构造，而不是在测试集上重新学习。Top-K 均值用于考察多个强候选的一致贡献，堆叠模型用于检验专家之间是否存在互补信息，不确定性加权用于降低高方差预测的影响，适用域门控用于在高风险样本上调整或标记输出。所有这些模块均受同一冻结规则约束，未通过验证集门控的结果只能作为负结果或后续研究线索。",
        ],
        changes,
        "方法：扩展候选登记和融合约束",
    )

    add_after_prefix(
        doc,
        "FZYC-Mol 的验证集治理流程包括六个预先定义的步骤",
        [
            "该流程的关键不是把验证集当作一个更小的测试集，而是把验证集作为策略选择的唯一依据。为减少选择偏差，候选策略一旦在验证集上被接受，其测试评估所使用的权重、阈值和规则不再调整。若验证集证据不足，即使某个复杂模块在个别数据集上有直观吸引力，也不会进入最终策略。这一规则使正结果和负结果处于同一证据框架中。",
            "在实际解释中，验证集治理仍不是万能防护。候选池越大、终点噪声越高、样本量越小，验证集排序越可能出现不稳定。因此，本文将 validation-test 排名审计、Top-3 命中率、optimism gap 和 nested validation 作为选择器本身的可靠性证据。这样，读者不仅能看到最终策略的性能，也能看到选择过程可能产生的误差范围。",
        ],
        changes,
        "方法：扩展验证集治理与选择器风险",
    )

    add_after_prefix(
        doc,
        "适用域分析使用两类信号",
        [
            "风险分数在本文中被定义为辅助审计信号，而不是新的性能指标。分类任务中，错误事件由预测标签与真实标签不一致定义；回归任务中，高误差事件由绝对误差处于最高分位的样本定义。风险分数综合相似度、模型分歧、预测偏差和重构误差等信息，用于衡量一个样本是否需要额外复核。由于这些信号与真实错误之间不一定存在稳定因果关系，本文只把风险分数解释为可操作的复核提示。",
            "risk-coverage 曲线用于描述拒用或复核策略的潜在收益。随着高风险样本被优先移出，保留样本上的平均风险或错误率可能下降；但如果风险分数与错误事件相关性不足，曲线收益就会有限。该分析使模型可靠性从单点指标扩展到决策情境：用户可以根据可接受的覆盖率和复核成本选择是否采用风险门控。",
        ],
        changes,
        "方法：扩展风险分数与 risk-coverage 解释",
    )

    add_after_prefix(
        doc,
        "分类任务通过 Brier 分数",
        [
            "保形预测用于为单点预测增加覆盖率意义。回归任务中，split conformal prediction 在校准集上估计非一致性分数，并为测试样本构造预测区间；分类任务中，非一致性分数用于构造预测集合。目标覆盖率设定为 80%、90% 和 95%，经验覆盖率用于检查预测集合或区间是否接近预设置信水平。该模块不参与候选选择，避免把可靠性输出再次转化为性能优化目标。",
            "可解释性分析同样遵循边界化原则。基序归因、BRICS 片段富集和 Murcko scaffold 分析可以帮助定位错误样本或高风险样本的结构背景，但不能直接证明某一片段导致某一性质变化。本文要求这些解释与最小支持度、效应量、p 值或 FDR 校正共同报告，并结合最近邻案例进行复核。只有当统计富集、化学合理性和样本案例一致时，解释才被视为较强的关联证据。",
        ],
        changes,
        "方法：扩展保形预测与解释边界",
    )

    add_after_exact(
        doc,
        "3 结果",
        [
            "结果部分按照证据强度递进展开。首先报告 MoleculeNet 主面板中的最终保留性能，再检验强基线同划分比较和定向补救是否改变主要结论；随后评估外部 TDC ADMET 终点、可靠性输出、保形覆盖率和验证-测试排名关系；最后通过 OOD、低相似度、活性悬崖、消融实验和失败案例界定模型边界。这样的顺序使读者先看到模型是否有效，再看到这种有效性在何种条件下成立。",
        ],
        changes,
        "结果：增加证据阶梯导语",
    )

    add_after_prefix(
        doc,
        "分类任务中，BBBP 和 ClinTox",
        [
            "这些结果显示，FZYC-Mol 的增益具有终点依赖性。ESOL 和 BACE 中，验证集选择器能够识别稳定候选；Lipophilicity 中，定向补救头形成明确改善；FreeSolv 则提醒读者，物理相互作用较强的任务可能需要更丰富的构象或能量信息。分类终点的结果同样不能只按 ROC-AUC 解释，尤其是 ClinTox 这类阳性样本稀少的任务，阳性召回和校准质量对筛选应用更关键。",
        ],
        changes,
        "结果：补充 MoleculeNet 结果模式",
    )

    add_after_prefix(
        doc,
        "为避免结论依赖弱基线模型",
        [
            "强基线同划分比较的意义在于排除一个常见替代解释：FZYC-Mol 的表现并非因为对照模型设置过弱。将表格模型、图模型、自动机器学习接口和冻结表征头纳入同一候选登记后，最终结果仍由验证集门控决定。若强基线已经足够稳定，FZYC-Mol 保留原策略；若多视图融合或补救头在验证集上形成一致优势，才进入最终保留结果。",
        ],
        changes,
        "结果：补充强基线比较解释",
    )

    add_after_prefix(
        doc,
        "在外部 TDC ADMET 任务上",
        [
            "外部 ADMET 结果进一步支持选择性增益而非普遍提升的结论。Caco2_Wang、HIA_Hou 和 Pgp_Broccatelli 等终点的改善说明，多视图融合和适用域门控在部分吸收和转运相关任务中具有价值；多数终点保留原策略则说明验证集治理不会强制复杂化模型。对于投稿而言，这一结果模式比单纯追求更多 win 更重要，因为它展示了框架如何在证据不足时拒绝过度声明。",
        ],
        changes,
        "结果：补充外部 ADMET 解释",
    )

    add_after_prefix(
        doc,
        "可靠性分析显示，风险分数识别分类错误的能力",
        [
            "该差异提示分类和回归任务的可靠性结构并不相同。分类错误通常表现为边界附近概率不稳定、模型分歧升高或适用域信号减弱，因此更容易被风险分数捕捉；回归高误差则可能由连续性质的局部跳变、测量噪声或未建模物理因素引起，单纯依赖相似度和模型分歧往往不足。本文据此将风险分数用于复核排序，而不是作为自动拒用规则。",
        ],
        changes,
        "结果：补充风险分数差异解释",
    )

    add_after_prefix(
        doc,
        "保形预测结果显示",
        [
            "保形预测结果的意义在于把模型输出从单点值扩展为带有覆盖率约束的区间或集合。对于回归任务，区间可以提示预测不确定性的范围；对于分类任务，预测集合可以提示模型是否需要在多个类别之间保留不确定性。需要注意的是，覆盖率接近目标并不意味着所有样本都同等可靠，区间宽度、类别不平衡和适用域距离仍需结合解释。",
        ],
        changes,
        "结果：补充保形预测意义边界",
    )

    add_after_prefix(
        doc,
        "验证-测试排名审计显示",
        [
            "排名审计为验证集治理提供了反身性检查。中等 Spearman 相关说明验证集排序含有有用信息，但较低的 Top-1 一致率说明验证集第一名不能被简单等同于测试集最佳。将 Top-3 命中率、regret 和 optimism gap 与 nested validation 共同报告，可以让读者判断候选选择是否稳健，也能避免把偶然测试优势包装成方法优势。",
        ],
        changes,
        "结果：补充排名审计解释",
    )

    add_after_prefix(
        doc,
        "在 OOD、Perimeter 和低相似度压力测试中",
        [
            "低相似度和活性悬崖分析强调了平均指标之外的样本级风险。即使整体 ROC-AUC 或 RMSE 较好，模型仍可能在训练集中缺少近邻、骨架发生变化或局部标签差异较大的样本上失效。MoleculeACE 的 cliff-pair 分析将这一问题具体化为相似分子之间的真实性质差异。FZYC-Mol 在此处的作用不是消除活性悬崖，而是把这类样本标记为需要更谨慎解释的区域。",
        ],
        changes,
        "结果：补充低相似度与活性悬崖边界",
    )

    add_after_prefix(
        doc,
        "消融实验和负结果审计显示",
        [
            "负结果在本文中不是附属内容，而是框架可信度的一部分。若一个模块未通过验证集门控，它说明该模块在当前数据和指标下缺乏足够证据；若一个模块只改善特定终点，它说明该模块具有条件性适用范围。将这些信息放入主文可以减少选择性报告风险，也帮助后续研究者判断哪些方向值得继续投入计算或实验资源。",
        ],
        changes,
        "结果：补充负结果的审计价值",
    )

    add_after_prefix(
        doc,
        "本研究将分子性质预测的评价重点",
        [
            "这一转变对分子机器学习评价具有方法学意义。许多公开基准天然鼓励单一数字比较，但药物化学决策往往要求同时考虑候选是否处于适用域内、预测是否经过校准、错误能否被提前标记，以及失败样本是否具有可解释的化学背景。FZYC-Mol 把这些要求组织成一条证据链，使模型评估从“谁的平均分更高”转向“哪些结论在何种条件下可信”。",
        ],
        changes,
        "讨论：补充方法学意义",
    )

    add_after_prefix(
        doc,
        "与单一主干模型不同",
        [
            "这种定位也解释了为什么本文没有把所有未通过门控的模块从叙事中删除。对于方法型论文，失败模块可以揭示设计假设的边界。3D-lite 物理代理和轻量适配器需要更严格的嵌套验证，粗糙度加权在相关性不稳定时只能作为诊断线索，bRo5 公共压力测试不能替代独立盲测。把这些边界明确写出，有助于防止读者把框架误解为一个无条件泛化的黑箱系统。",
        ],
        changes,
        "讨论：补充失败模块边界",
    )

    add_after_prefix(
        doc,
        "可靠性分析进一步限定了该框架的使用方式",
        [
            "在实际部署中，FZYC-Mol 更适合以决策卡形式输出，而不是只输出一个数值预测。一个完整的预测记录应包括点预测、适用域相似度、风险分位、是否经过校准、是否处于保形预测集合内、最近邻案例和拒用理由。这样的记录可以帮助药物化学家区分高置信候选、需要人工复核的候选和超出模型适用范围的候选。本文的实验结果为这种报告方式提供了离线证据，但尚不能替代前瞻性项目验证。",
        ],
        changes,
        "讨论：补充实践使用方式",
    )

    add_after_prefix(
        doc,
        "本研究仍存在局限",
        [
            "还有一个限制来自公开数据本身。不同数据集的标签质量、测量条件和任务定义并不完全一致，某些终点可能包含实验误差、批次差异或历史数据整合带来的噪声。FZYC-Mol 可以通过负结果、校准和风险分析暴露部分不稳定性，但不能从离线数据中消除这些来源。因此，本文的结论应被理解为公开基准和公共压力测试上的可靠性审计，而不是对真实项目成功率的直接估计。",
            "参考文献和外部数据也需要在投稿前继续核验。本文保留了 2025 和 2026 年相关文献作为研究背景，但正式投稿前仍应逐条检查 DOI、卷页、预印本状态和数据可用性。所有图表对应的 source data、随机种子级预测和候选登记表应与稿件同步归档，以便审稿人复核模型选择过程。若某项后续验证尚未完成，应继续写为未来工作或边界条件，而不能写入完成性结果。",
        ],
        changes,
        "讨论：补充数据与引用核验边界",
    )

    add_after_prefix(
        doc,
        "在这些边界内，FZYC-Mol 为 ADMET 预测",
        [
            "因此，本文的最终含义不是提出一个能够取代所有分子性质预测模型的新模型，而是提出一种更透明的报告方式。对于模型开发者，它提供候选登记、验证选择、负结果记录和可靠性审计的流程；对于使用者，它提供判断预测是否可信的辅助证据；对于审稿人，它提供检查性能主张是否越过证据边界的结构。这个层面的贡献与单一性能提升互补，也更符合高影响力期刊对可复核性和边界意识的要求。",
        ],
        changes,
        "讨论：补充总结性意义",
    )

    add_after_prefix(
        doc,
        "本研究提出 FZYC-Mol",
        [
            "与只报告最终性能的模型比较相比，FZYC-Mol 更强调选择过程的透明性。通过在同一框架中呈现强基线、选择器审计、风险覆盖、保形预测、低相似度分析和负结果，本研究使读者能够判断哪些增益受到验证集支持，哪些模块只适合作为后续验证方向。这样的证据组织方式是本文面向 Nature 风格写作保留的核心主线。",
        ],
        changes,
        "结论：补充核心主线",
    )

    return changes


def final_language_cleanup(doc: Document) -> list[str]:
    changes: list[str] = []
    replacements = {
        "validation-test 排名审计": "验证-测试排名审计",
    }
    for paragraph in doc.paragraphs:
        updated = paragraph.text
        applied: list[str] = []
        for old, new in replacements.items():
            if old in updated:
                updated = updated.replace(old, new)
                applied.append(f"{old} -> {new}")
        if applied:
            set_para_text(paragraph, updated)
            changes.append("术语统一：" + "; ".join(applied))
    return changes


def count_chinese_in_doc(doc: Document) -> int:
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    text = "\n".join(texts)
    return sum(1 for c in text if "\u4e00" <= c <= "\u9fff")


def count_body_chinese(doc: Document) -> int:
    skip = False
    texts: list[str] = []
    for paragraph in doc.paragraphs:
        t = paragraph.text.strip()
        if t == "References":
            skip = True
        if not skip:
            texts.append(t)
    return sum(1 for c in "\n".join(texts) if "\u4e00" <= c <= "\u9fff")


def collect_qa(doc: Document, before_body: int, before_total: int) -> list[str]:
    text = "\n".join(p.text for p in doc.paragraphs)
    qa = []
    qa.append(f"正文中文字符：{count_body_chinese(doc)}；扩写前正文中文字符：{before_body}。")
    qa.append(f"全文中文字符（含表格）：{count_chinese_in_doc(doc)}；扩写前全文中文字符：{before_total}。")
    qa.append(f"段落数：{len(doc.paragraphs)}；表格数：{len(doc.tables)}；图片数：{len(doc.inline_shapes)}。")
    qa.append(f"最大表格列数：{max((len(t.columns) for t in doc.tables), default=0)}。")
    qa.append(f"长表格单元格（>90 字）：{sum(1 for t in doc.tables for r in t.rows for c in r.cells if len(c.text.strip()) > 90)}。")
    qa.append(f"禁止跨页拆分表格行：{sum(1 for t in doc.tables for r in t.rows if r._tr.get_or_add_trPr().find(qn('w:cantSplit')) is not None)}。")
    for term in ["首创", "首次", "革命性", "证明了", "数据状态审计", "6 投稿前自我审查", "弱 基线模型", "backbone", "class weight", "validation metric"]:
        qa.append(f"风险词/旧表述检查 {term}：{text.count(term)} 处。")
    for heading in ["摘要", "1 引言", "2 材料与方法", "3 结果", "4 讨论", "5 结论", "List of abbreviations", "Declarations", "References"]:
        qa.append(f"章节检查 {heading}：{'存在' if heading in text else '未检出'}。")
    return qa


def build_report(changes: list[str], qa: list[str]) -> None:
    doc = Document()
    doc.add_heading("初稿-6 Nature 扩写与 QA 报告", level=1)
    doc.add_paragraph(f"源文档：{SRC_DOCX}")
    doc.add_paragraph(f"输出文档：{DEST_DOCX}")
    doc.add_heading("一条核心论点", level=2)
    doc.add_paragraph(
        "在分子性质预测中，本研究展示 FZYC-Mol 通过验证集治理、适用域审计和负结果记录，使模型选择与可靠性边界可复核；证据来自 MoleculeNet、TDC ADMET、MoleculeACE、bRo5 公共压力测试、保形预测和验证-测试排名审计；边界是其不宣称所有终点普遍提升，也不把公共压力测试写成独立盲测。"
    )
    doc.add_heading("术语账本", level=2)
    table = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["标准术语", "含义", "需避免的漂移", "处理"]):
        table.rows[0].cells[i].text = h
    rows = [
        ["FZYC-Mol", "验证集治理驱动的适用域感知框架", "模型、系统、平台混用", "正文统一为框架或 FZYC-Mol"],
        ["验证集治理", "候选策略由验证集接受、拒绝或保留", "测试后选择", "强调测试集冻结"],
        ["公共压力测试", "公开数据上的外推边界观察", "独立盲测", "bRo5 相关结论降调"],
        ["风险分数", "错误或高误差样本复核提示", "自动拒用规则", "限定为辅助可靠性证据"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_heading("主要扩写位置", level=2)
    for item in changes:
        doc.add_paragraph(item)
    doc.add_heading("QA 结果", level=2)
    for item in qa:
        doc.add_paragraph(item)
    doc.save(REPORT_DOCX)


def main() -> None:
    source = Document(SRC_DOCX)
    before_body = count_body_chinese(source)
    before_total = count_chinese_in_doc(source)

    doc = Document(SRC_DOCX)
    changes = expand_manuscript(doc)
    changes.extend(final_language_cleanup(doc))
    keep_table_compact(doc)
    qa = collect_qa(doc, before_body, before_total)

    doc.save(DEST_DOCX)
    shutil.copy2(DEST_DOCX, OUT_DIR / DEST_DOCX.name)
    build_report(changes, qa)
    shutil.copy2(REPORT_DOCX, OUT_DIR / REPORT_DOCX.name)

    print(f"Wrote {DEST_DOCX}")
    print(f"Wrote {REPORT_DOCX}")
    for item in qa:
        print(item)


if __name__ == "__main__":
    main()
