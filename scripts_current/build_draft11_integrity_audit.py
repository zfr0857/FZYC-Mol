# -*- coding: utf-8 -*-
"""Build draft 11 after an academic-integrity audit.

This pass removes or downgrades claims that are not supported by local
artifacts:
- bRo5 datasets are currently missing_data, so no bRo5 performance result is
  claimed.
- MoleculeACE evidence covers the available 17-task subset and 51 seed-pair
  comparisons, not a completed 30-task full audit.
- Data/code availability is made explicit about absent permanent identifiers.
"""

from __future__ import annotations

from pathlib import Path
import csv

from docx import Document


DESKTOP = Path(r"C:\Users\Administrator\Desktop")
SRC = max(DESKTOP.glob("FZYC-Mol_*Nature*.docx"), key=lambda p: p.stat().st_mtime)
OUT = DESKTOP / "FZYC-Mol_初稿-11.docx"
REPORT_DIR = Path("reports/academic_integrity_audit")
REPORT_DIR.mkdir(parents=True, exist_ok=True)


PARAGRAPH_REPLACEMENTS_BY_PREFIX: dict[str, str] = {
    "分子性质预测是药物发现、ADMET 评估和毒性风险筛查中的关键计算环节": (
        "分子性质预测是药物发现、ADMET 评估和毒性风险筛查中的关键计算环节。随机划分下的平均 "
        "ROC-AUC 或 RMSE 难以充分反映模型在新骨架、少样本终点、不平衡标签、超规则五化学空间、"
        "活性悬崖、实验噪声和适用域漂移条件下的可靠性。受近期真实挑战 ADMET 基准研究启发 [3]，"
        "本文提出 FZYC-Mol，一种由验证集治理的适用域感知分子性质预测框架。该框架不依赖单一大型"
        "主干模型，而是把多视图表示、强基线专家、Random/Scaffold/Perimeter 划分、类别不平衡策略、"
        "MoleculeACE 可用任务子集、粗糙度诊断、校准和适用域证据纳入冻结候选池；bRo5 模块在当前"
        "版本中仅作为数据状态审计和后续运行接口保留，不作为已完成性能评估主张。最终策略仅由验证集"
        "决定，测试集只在策略冻结后用于一次性评估。"
    ),
    "本文贡献可概括为五点": (
        "本文贡献可概括为五点。第一，围绕 MoleculeNet、TDC ADMET、bRo5 数据状态审计、"
        "MoleculeACE 可用任务子集和 Random/Scaffold/Perimeter 划分构建统一评测流程，并明确训练集、"
        "验证集和测试集的使用边界。第二，将 TabPFNv2、AutoGluon、KPGT representation、"
        "XGBoost-RDKit/Mordred/MorganCount、Chemprop、树模型、冻结表征、采样策略、目标变换和融合策略"
        "统一登记为预注册候选，而不是在测试集后追加模型。第三，提出可审计的 accept/reject/retain "
        "治理机制，使候选策略需先通过验证集证据才能进入主结果；未通过候选作为负结果保留。第四，"
        "报告 validation-test rank correlation、Top1 match、Top3 hit、regret、optimism gap、"
        "不平衡分类、保形预测、risk-coverage、MoleculeACE cliff-pair 指标和 ROGI/MODI/SARI 粗糙度诊断，"
        "避免只依据单一 ROC-AUC 或 RMSE 下结论。第五，将数据缺失、候选失败、低相似度失败和活性悬崖失败"
        "明确写入限制与补充材料，避免把未完成实验表述为正结果。"
    ),
    "为与 Zhao et al. 2026 的真实挑战框架": (
        "为与 Zhao et al. 2026 的真实挑战框架 [3] 保持可比，本文将扩展实验组织为"
        "“挑战维度—候选池—验证门控—风险输出”的矩阵。该矩阵把强基线、OOD 划分、不平衡分类、"
        "bRo5 数据状态审计、活性悬崖、粗糙度、候选接受/拒绝审计和校准/保形预测纳入同一验证集治理流程，"
        "而不是将各模型作为相互孤立的追加结果。实验 A-C 和 E-G 构成支撑主线结论的核心证据，实验 D 仅作为"
        "bRo5 后续评估的缺失数据与运行条件说明，实验 H 用于加强可靠性输出。所有模块均遵守测试集冻结原则："
        "候选是否进入最终保留结果，只由验证集或预先定义的审计规则决定。"
    ),
    "为保证扩展实验可复现": (
        "为保证扩展实验可复现，本文将扩展实验与既有结果解耦，并要求所有运行过程输出稳定命名的 CSV、"
        "字段字典、缺失数据报告和缺失依赖报告。该设计使强基线、OOD 划分、不平衡策略、bRo5 数据状态审计、"
        "MoleculeACE 可用子集审计、粗糙度诊断和候选接受/拒绝记录能够从同一结果目录追溯到数据划分、"
        "随机种子、候选模型和最终表格。具体脚本、配置和字段说明见代码可用性部分及 Supplementary Methods。"
    ),
    "bRo5 外部压力测试用于覆盖更复杂的化学空间": (
        "bRo5 部分在当前版本中仅作为数据状态审计，而不是性能结果。CycPept PAMPA、LinPept NonFouling "
        "和 LinPept CellPen 被预先列为后续压力测试数据集；然而本研究审计文件显示三者当前均为 "
        "missing_data，尚未具备同划分模型训练、RMSE/MAE、AD 覆盖率或风险富集评估条件。因此，正文不报告 "
        "bRo5 性能数值，也不将其纳入 FZYC-Mol 的已完成外部验证证据。该模块保留的作用是明确所需输入字段、"
        "运行条件和未来评估边界，避免将缺失数据误写为完成实验。"
    ),
    "为形成完整的验证闭环": (
        "为形成完整的验证闭环，本文将关键扩展证据纳入结果逻辑：先审计验证集选择是否存在排名偏差，再用"
        "配对统计和系统消融界定各模块贡献，随后在 Random/Scaffold/Perimeter 梯度、低相似度、bRo5 数据"
        "状态审计、MoleculeACE 可用 17 任务子集、不平衡分类、ROGI/MODI/SARI 粗糙度诊断、候选接受/拒绝审计、"
        "校准/保形预测和失败案例中检查适用边界。这些分析不改变测试集一次性报告原则，而是将"
        "“为什么接受、保留或拒绝某个候选”写成可追溯的证据链。"
    ),
    "MoleculeACE 配对结果与低相似度分析互相补充": (
        "MoleculeACE 配对结果与低相似度分析互相补充。当前可核验结果覆盖 17 个 MoleculeACE 任务和 "
        "51 个 seed 配对，而不是完整 30 任务全量审计。主指标包括 R2、RMSE、MAE、cliff-pair RMSE、"
        "预测差异与真实差异的 gap Spearman、cliff recall 和 top cliff error；比较对象包括 "
        "XGBoost-MorganCount、KPGT representation、可运行的 MOLMCL/KPGT 版本、FZYC selector 以及 "
        "KPGT+XGBoost ensemble。已有代表性配对结果显示，活性悬崖目标候选在 51 个 seed 配对中的总 "
        "RMSE 平均正向变化为 0.0069，悬崖子集 RMSE 平均正向变化为 0.0056，但标准差较大；预测差异与真实"
        "差异的平均 Spearman 约为 0.252，部分任务接近零或为负。因此，本文将该结果解释为悬崖风险识别和"
        "候选治理的补充证据，而不是对活性悬崖预测已经解决或 30 任务均已完成的证明。"
    ),
    "这种差异化定位决定了本文的补强方向": (
        "这种差异化定位决定了本文的补强方向。实验 A-H 将近期真实挑战研究强调的强基线、Perimeter split、"
        "类别不平衡、bRo5 数据状态、活性悬崖、粗糙度、候选接受/拒绝审计和校准/保形预测转化为 FZYC-Mol "
        "的候选治理压力测试。由此，复杂模块不被直接纳入主结论，而是在统一验证证据下被接受、保留、拒绝或"
        "标注为数据未就绪。这种证据组织与当前结果中“选择性增益多于普遍胜出”的事实一致。"
    ),
    "Zhao et al. 2026 将 ADMET 可靠性组织为四类真实挑战": (
        "Zhao et al. 2026 将 ADMET 可靠性组织为四类真实挑战：数据稀缺/OOD、类别不平衡、bRo5 化学空间"
        "和活性悬崖 [3]。该工作提供了一个高水平基准模板，强调 TabPFNv2、AutoGluon、KPGT、Uni-Mol、GEM、"
        "GNN 和传统机器学习在不同场景下的边界。与这种大规模 benchmark 不同，FZYC-Mol 的核心不是重新证明"
        "哪类模型总体最强，而是在候选池进入冻结测试之前，给出一个可审计的 accept/reject/retain 治理机制，"
        "并把负结果、roughness 风险和样本级失败案例纳入同一证据链。"
    ),
    "此外，OmniMol、ADMETlab 3.0 和近期 ADMET 表征基准提示": (
        "此外，OmniMol、ADMETlab 3.0 和近期 ADMET 表征基准提示，真实 ADMET 建模往往同时面对不完整标注、"
        "终点间相关性、平台化决策支持和跨数据集迁移问题 [47-49]。因此，FZYC-Mol 的价值不只在于判定哪个"
        "模型分数最高，更在于说明候选策略在什么边界内可用、为什么可用以及何时应被拒绝。外部 ADMET 附录、"
        "bRo5 数据缺失审计、MoleculeACE 可用子集、ROGI/MODI/SARI 诊断和计算成本审计共同服务于这一点。"
    ),
    "本研究仍存在若干局限": (
        "本研究仍存在若干局限。首先，验证-测试排名审计和 9 个代表性终点的 3×3 nested validation 表明，"
        "验证集治理可以减少测试集事后选择，但内外层验证尚未覆盖全部 MoleculeNet、TDC、bRo5 和 MoleculeACE "
        "终点，不能保证所有任务均达到测试最优；因此，小幅增益应与 regret、optimism gap、Top3 hit 和负结果"
        "共同解释。其次，bRo5 三个预设数据集在当前审计中均为 missing_data，不能作为已完成性能评估；"
        "MoleculeACE 当前可核验结果覆盖 17 个任务和 51 个 seed 配对，不能写成完整 30 任务全量结论。第三，"
        "TabPFNv2、AutoGluon、KPGT representation、bRo5 和 MoleculeACE 仍需在完整同划分结果中继续扩展，"
        "尚未完成完整同划分验证的候选仅作为候选接口或后续验证方向，不作为已取得性能提升的证据。第四，"
        "bRo5、活性悬崖和低相似度样本仍是主要失败边界，FZYC-Mol 更适合识别和审计这些风险，而不是宣称已解决它们。"
        "第五，基序归因、片段富集和粗糙度相关性仍属于关联证据，不能替代因果机制或湿实验验证。"
    ),
    "本文提出并系统评估了 FZYC-Mol": (
        "本文提出并系统评估了 FZYC-Mol，一种由验证集治理、适用域感知的多专家分子性质预测框架。结果表明，"
        "在 MoleculeNet、TDC ADMET、Random/Scaffold/Perimeter 划分、低相似度子集、MoleculeACE 可用任务子集、"
        "ROGI/MODI/SARI 粗糙度诊断、候选接受/拒绝审计、校准/保形预测和基序/片段解释等多条证据线上，"
        "FZYC-Mol 能够提供比单一模型分数更完整的可靠性画像。bRo5 当前仅完成数据状态审计，不作为已完成性能结果。"
        "其核心价值不是保证每个终点取得测试最优，而是在固定候选池内透明地接受、保留、拒绝和审计候选策略；"
        "当强基线、补救头、低成本重构或多方法融合通过验证集门控时，它们可以进入最终保留结果，当证据不足或数据"
        "未就绪时则作为负结果或待验证模块保留。"
    ),
    "数据和代码可用性：本文使用的公开数据集可通过": (
        "数据和代码可用性：本文使用的公开数据集可通过 MoleculeNet、Therapeutics Data Commons、MoleculeACE "
        "及相应原始平台获得。与本文结果对应的 split seeds、候选登记表、验证/测试预测、统计检验脚本、图表 "
        "source data、环境文件和表格生成脚本将与投稿或接收版本同步存档于 GitHub/Zenodo 或期刊认可的数据仓库。"
        "本稿件当前尚未分配永久数据 DOI 或 accession number；仓库记录冻结并可供审稿访问后，方可在正文和补充材料中"
        "填写永久链接、版本号和 accession number。bRo5 相关条目当前以 missing_data 状态报告，不作为已公开或已完成"
        "性能数据存档。"
    ),
}


INLINE_REPLACEMENTS = {
    "MoleculeACE 30 任务": "MoleculeACE 可用任务子集",
    "MoleculeACE 30任务": "MoleculeACE 可用任务子集",
    "bRo5 压力测试": "bRo5 数据状态审计",
    "bRo5/肽类压力测试": "bRo5 数据状态审计",
    "bRo5 肽类压力": "bRo5 数据状态",
    "作为bRo5": "作为 bRo5",
}


TABLE_UPDATES = {
    (3, 5, 3): "bRo5 数据是否已满足性能评估条件",
    (3, 5, 4): "CycPept PAMPA、LinPept NonFouling、LinPept CellPen；当前审计状态均为 missing_data",
    (3, 5, 5): "missing_data_report、所需字段、后续运行条件；当前不报告 RMSE/MAE 或选择器结果",
    (3, 6, 3): "活性悬崖是否从案例升级为可用任务子集证据",
    (3, 6, 4): "MoleculeACE 可用 17 任务子集；XGBoost-MorganCount、KPGT representation、FZYC selector、可运行 MOLMCL/KPGT 版本",
    (3, 7, 4): "MoleculeNet、TDC、MoleculeACE 可用子集；ROGI/MODI/SARI 或可复现代理",
    (26, 6, 4): "整合 MoleculeACE cliff objective selector 与 17 任务/51 seed 配对的 cliff subset 结果；数值悬崖预测仍作为限制。",
}


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text


def replace_paragraphs(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        for prefix, repl in PARAGRAPH_REPLACEMENTS_BY_PREFIX.items():
            if text.startswith(prefix):
                set_paragraph_text(paragraph, repl)
                changed += 1
                break
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new = text
        for old, repl in INLINE_REPLACEMENTS.items():
            new = new.replace(old, repl)
        if new != text:
            set_paragraph_text(paragraph, new)
            changed += 1
    return changed


def replace_tables(doc: Document) -> int:
    changed = 0
    for (table_no, row_no, col_no), repl in TABLE_UPDATES.items():
        table = doc.tables[table_no - 1]
        cell = table.rows[row_no - 1].cells[col_no - 1]
        para = cell.paragraphs[0]
        if para.text != repl:
            set_paragraph_text(para, repl)
            changed += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    new = text
                    for old, repl in INLINE_REPLACEMENTS.items():
                        new = new.replace(old, repl)
                    if new != text:
                        set_paragraph_text(paragraph, new)
                        changed += 1
    return changed


def write_audit_report() -> None:
    rows = [
        {
            "risk": "bRo5 overclaim",
            "evidence_file": "reports/experiment_update/bro5_data_status.csv",
            "local_evidence": "cycpept_pampa, linpept_nonfouling and linpept_cellpen are all missing_data",
            "manuscript_action": "All bRo5 performance claims were downgraded to data-status audit / future evaluation conditions.",
            "residual_action": "Provide actual bRo5 CSV/TSV files before claiming RMSE/MAE/AD/risk enrichment.",
        },
        {
            "risk": "MoleculeACE 30-task overclaim",
            "evidence_file": "reports/experiment_update/moleculeace_30_task_summary.csv; reports/remaining_missing_experiments_20260606/moleculeace_gap_correlation_summary.csv",
            "local_evidence": "Current verified data cover 17 unique tasks and 51 seed-pair comparisons.",
            "manuscript_action": "All '30-task full audit' wording was replaced with 'available 17-task subset / 51 seed pairings'.",
            "residual_action": "Run and archive all intended 30 tasks before restoring a 30-task claim.",
        },
        {
            "risk": "Unassigned repository identifiers",
            "evidence_file": "manuscript availability statement",
            "local_evidence": "No permanent DOI/accession is assigned in the draft.",
            "manuscript_action": "Availability statement now states identifiers are not yet assigned and must be filled only after repository deposition.",
            "residual_action": "Create repository record with DOI/accession, README, field dictionary and reviewer-access link.",
        },
        {
            "risk": "Citation verification manual-needed items",
            "evidence_file": "reports/academic_integrity_audit/citation_verification_draft11.csv",
            "local_evidence": "31/50 references verified automatically under strict DOI/arXiv/Crossref checks; 19 non-DOI, conference, software or title-search-ambiguous references require manual confirmation.",
            "manuscript_action": "No core claim was made to depend solely on unverified metadata; DOI/arXiv-verified recent references were retained.",
            "residual_action": "Manually verify refs 2, 11, 18-21, 23-27, 30-31, 34-35, 37, 39-41 before submission.",
        },
    ]
    with (REPORT_DIR / "draft11_academic_integrity_actions.csv").open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def main() -> None:
    doc = Document(SRC)
    p_changes = replace_paragraphs(doc)
    t_changes = replace_tables(doc)
    write_audit_report()
    doc.save(OUT)
    print(f"source={SRC}")
    print(f"output={OUT}")
    print(f"paragraph_changes={p_changes}")
    print(f"table_changes={t_changes}")
    print(f"report={REPORT_DIR / 'draft11_academic_integrity_actions.csv'}")


if __name__ == "__main__":
    main()
