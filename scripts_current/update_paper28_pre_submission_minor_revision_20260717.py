from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
from openpyxl.comments import Comment
from PIL import Image


ROOT = Path(r"D:\fzyc")
SOURCE = ROOT / "output" / "paper27_equal_size_registry_composition_revision_20260716"
TARGET = ROOT / "output" / "paper28_pre_submission_minor_revision_20260717"
FIG = TARGET / "main_figures"
ANALYSIS = ROOT / "output" / "paper27_equal_size_registry_composition_20260716"
EN_NAME = "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript.docx"
REVIEWER = "Reviewer_concern_Response_Location.docx"


EN_ABSTRACT_RESULTS = (
    "Results: Effective diversity and chance-adjusted ranking remained matrix and endpoint dependent. In the split-regime audit, the "
    "K = 32 minus K = 4 CAHit@3 direction remained negative and the cross-fitted gap direction positive under both split regimes for all "
    "three endpoints, although magnitudes changed. At matched K = 32, classical multiview and modern-augmented pools increased both observed "
    "audit-best opportunity and validation-selected gain relative to the homogeneous Morgan pool in ClinTox, BACE and ESOL. Modern augmentation "
    "had the largest raw effective rank at every endpoint, but relative entropy rank was not monotonically associated with cross-fitted gap "
    "across the 18 repeated endpoint-pool-K audit cells (descriptive Spearman rho = -0.211). ESOL showed higher selected gain despite lower "
    "CAHit@3, separating ranking recovery from realised utility. Modern augmentation required substantially more downstream audit time."
)

EN_FIG7 = (
    "Figure 7. Equal-size candidate-pool composition intervention. (A) At K = 32, observed audit-best and validation-selected gains are "
    "normalized by the endpoint-specific homogeneous-pool observed audit-best gain for visualization; raw endpoint-scale values are reported "
    "in Table S28. All denominators were positive and the minimum was 0.0784 utility units. (B) Raw entropy effective rank divided by nominal "
    "candidate count K is compared with the leave-one-seed-out cross-fitted selection gap, normalized by the same endpoint- and K-specific "
    "homogeneous-pool observed audit-best gain. (C) Chance-adjusted Hit@3 is shown for ClinTox, BACE and ESOL across the three candidate "
    "registries at K = 16 and K = 32. (D) Observed downstream audit time is compared with normalized selected gain. Colour denotes registry "
    "composition, marker shape denotes endpoint and marker size denotes candidate count. Encoder pretraining, model acquisition and cached "
    "embedding-extraction costs were unavailable and excluded. The comparison therefore concerns observed downstream nested fitting and "
    "prediction costs only."
)

ZH_FIG7 = (
    "图7. 等规模候选池组成干预。A：K=32时，将观测审计最佳收益与验证选择收益除以相应终点同质池的观测审计最佳收益以便跨终点展示；"
    "原始终点尺度数值见表S28。所有分母均为正，最小值为0.0784效用单位。B：以名义候选数K归一化的原始熵有效秩（r_entropy/K）与留一随机种子"
    "交叉拟合选择差距比较，纵轴采用相同的终点和K特异性同质池观测审计最佳收益归一化。C：展示ClinTox、BACE和ESOL在三类候选池及K=16、32下的"
    "机会校正Hit@3。D：比较观测下游审计时间和归一化选择收益。颜色表示候选池组成，形状表示终点，标记大小表示候选数。编码器预训练、模型获取和缓存嵌入提取"
    "成本不可得且未计入，因此成本比较仅涉及观测到的下游嵌套拟合与预测。"
)


def paragraph_by_prefix(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise KeyError(prefix)


def replace_embedded_figures(doc: Document) -> None:
    images = [FIG / f"Figure{i}_600dpi.png" for i in range(1, 7)] + [
        FIG / "Figure_7_equal_size_registry_composition_600dpi.png"
    ]
    if len(doc.inline_shapes) != 7:
        raise ValueError(f"Expected seven inline figures, found {len(doc.inline_shapes)}")
    for shape, image_path in zip(doc.inline_shapes, images):
        rid = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        part = doc.part.related_parts[rid]
        if part.content_type != "image/png":
            raise ValueError(f"Expected PNG image part, found {part.content_type}")
        part._blob = image_path.read_bytes()
        width, height = Image.open(image_path).size
        shape.height = int(shape.width * height / width)
    for paragraph in doc.paragraphs:
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def update_english(path: Path) -> None:
    doc = Document(path)
    paragraph_by_prefix(doc, "Results:").text = EN_ABSTRACT_RESULTS
    paragraph_by_prefix(doc, "For every endpoint, seed, outer fold, pool and K").text = (
        "Opportunity gain was defined as the observed outer-audit best utility minus the shared-anchor utility, whereas realised gain was the "
        "validation-selected utility minus the same anchor. The observed audit-best utility is a finite audit maximum and not a population-optimal "
        "or true optimum. We also calculated CAHit@3, leave-one-seed-out cross-fitted selection gaps, Ledoit-Wolf relative effective diversity "
        "under four matrix transformations, complete downstream audit time and selected gain per audit hour. Encoder pretraining, model acquisition "
        "and cached embedding-extraction costs were unavailable and excluded; cost comparisons therefore refer only to observed downstream nested "
        "fitting and prediction."
    )
    paragraph_by_prefix(doc, "At K = 32, classical multiview versus homogeneous Morgan pools").text = (
        "At K = 32, classical multiview versus homogeneous Morgan pools increased mean observed audit-best opportunity gain from 0.125 to 0.188 "
        "in ClinTox, 0.079 to 0.084 in BACE and 1.419 to 2.260 utility units in ESOL. The corresponding validation-selected gains increased from "
        "0.103 to 0.172, 0.072 to 0.078 and 1.325 to 2.179. Modern-augmented validation-selected gains were 0.176, 0.079 and 2.191. All K = 32 "
        "observed audit-best and validation-selected gain differences versus the homogeneous pool were positive in all five seed means for both "
        "composition interventions (Additional file 2: Tables S28 and S31)."
    )
    paragraph_by_prefix(doc, "Raw K = 32 entropy ranks").text = (
        "Raw K = 32 entropy ranks were 10.88, 2.90 and 1.87 for homogeneous ClinTox, BACE and ESOL pools; 12.40, 6.09 and 9.07 for classical "
        "multiview pools; and 13.96, 7.92 and 11.47 for modern-augmented pools. Relative entropy rank was not monotonically associated with "
        "same-unit or cross-fitted selection gap across the 18 endpoint-pool-K cells (Spearman rho = -0.008 and -0.211). These cells reuse three "
        "endpoints across pools and K values and are not 18 independent experiments; the correlations are descriptive. At K = 32, the modern "
        "cross-fitted gap was smaller than the homogeneous gap for ClinTox (0.0048 versus 0.0092) and ESOL (0.0105 versus 0.0578), but slightly "
        "larger for BACE (0.0035 versus 0.0025)."
    )
    paragraph_by_prefix(doc, "Ranking fidelity and realised utility also separated").text = (
        "Ranking fidelity and realised utility also separated. ESOL CAHit@3 was 0.338, 0.044 and -0.030 for homogeneous, classical multiview and "
        "modern-augmented K = 32 pools, even though the latter two had substantially higher validation-selected gain. Modern augmentation consumed "
        "139-187 seconds of observed downstream audit time per outer unit versus 18-41 seconds for classical multiview and 21-31 seconds for "
        "homogeneous pools. Its selected gain per downstream audit hour was lower than the classical multiview value in all three endpoints; this "
        "comparison excludes encoder pretraining, model acquisition and cached embedding extraction (Figure 7; Tables S28-S30)."
    )
    paragraph_by_prefix(doc, "Figure 7.").text = EN_FIG7
    paragraph_by_prefix(doc, "Holding K and folds fixed").text = (
        "Holding K and folds fixed directly addressed the concern that a Morgan-dominated registry might determine the primary result. Both "
        "classical and modern additions raised the observed finite-audit ceiling and usually allowed the validation selector to realise most of "
        "that gain. The absence of a positive monotonic diversity-gap association in 18 repeated audit cells is descriptive evidence that effective "
        "rank is a structural audit descriptor, not a causal penalty parameter. Useful complementarity does not imply that more diversity is always "
        "better; value depends on whether inner validation ranking can identify the added information with sufficient stability."
    )
    paragraph_by_prefix(doc, "This result sharpens the expansion story").text = (
        "This result sharpens the expansion story. Nominal K measures eligibility, effective diversity describes independent utility movement, "
        "observed audit-best gain measures finite audited opportunity, validation-selected gain measures realisation and downstream compute records "
        "part of its price. The classical multiview pool occupied the most favourable observed downstream cost-benefit region here. Frozen "
        "ChemBERTa and MoLFormer representations plus the locked one-epoch D-MPNN expanded the composition boundary, but Figure 7 does not constitute "
        "a modern-model leaderboard or replace a compute-matched, fully tuned architecture benchmark."
    )
    paragraph_by_prefix(doc, "Figure 1.").text = (
        "Figure 1. Retrospective nested audit of candidate-pool expansion and model-selection loss. The central module separates candidate ranking "
        "in three inner scaffold folds from one-shot auditing in three outer folds across five seeded partitions, with no outer-label feedback. "
        "Registered endpoints and prespecified candidate prefixes enter from the left. Surrounding branches summarize utility-pattern diversity, "
        "chance-adjusted ranking, same-unit and cross-fitted selection gaps, split-regime transfer, equal-size registry composition, finite-audit "
        "winner optimism, chemical-support boundaries and four-model prediction/error reliability. These streams converge on an auditable evidence "
        "map rather than a model leaderboard."
    )
    paragraph_by_prefix(doc, "Figure 6.").text = (
        "Figure 6. Prediction reliability across chemical-support boundaries. Panel A combines prediction correlations and high-error Jaccard "
        "overlaps in a symmetric four-model matrix with separate colour scales. Panel B consolidates classification ROC-AUC, regression RMSE, "
        "classification false-negative rate and classification/regression high-error enrichment into one Tanimoto-support risk matrix; cell text "
        "gives natural-scale medians and colour encodes only the within-row adverse direction. Panel C reports novel-scaffold relative changes in "
        "error overlap, model disagreement, high-error enrichment and false-negative enrichment on a log-ratio axis. Panel D separates discrimination "
        "and minority-safety measures for the four fixed-configuration models; prediction-set size is mapped from its 1-2 scale only for display."
    )
    replace_embedded_figures(doc)
    doc.save(path)
    if "oracle" in "\n".join(p.text for p in Document(path).paragraphs).lower():
        raise ValueError("English manuscript still contains prohibited terminology")


def update_chinese(path: Path) -> None:
    doc = Document(path)
    paragraph_by_prefix(doc, "结果：").text = (
        "结果：切分机制迁移中，三个终点的K=32减K=4 CAHit@3方向均保持为负，交叉拟合差距方向均为正，但效应大小改变。固定K=32时，"
        "经典多表征池和现代增强池在ClinTox、BACE和ESOL均提高观测审计最佳机会收益和验证选择收益。现代池在三个终点均具有最高原始有效秩，"
        "但18个重复终点-池-K审计单元的相对熵秩与交叉拟合差距不呈单调关系（描述性Spearman ρ=-0.211）。ESOL在选择收益提高的同时"
        "CAHit@3下降，说明排序恢复与效用兑现并不等价。现代池的观测下游审计时间明显增加。"
    )
    paragraph_by_prefix(doc, "相对于共享锚点计算").text = (
        "机会收益定义为观测外层审计最佳效用减去共享锚点效用，实际收益定义为验证选择效用减去同一锚点。观测审计最佳效用是有限候选审计的"
        "最大值，并非总体最优或真实最优值。另报告CAHit@3、留一随机种子交叉拟合差距、四种矩阵变换下的Ledoit-Wolf相对有效多样性、"
        "完整下游审计时间及单位下游审计小时收益。编码器预训练、模型获取和缓存嵌入提取成本不可得且未计入。"
    )
    paragraph_by_prefix(doc, "在K=32时，经典多表征池相对于同质Morgan池").text = (
        "在K=32时，经典多表征池相对于同质Morgan池将ClinTox、BACE和ESOL的平均观测审计最佳机会收益分别从0.125、0.079和1.419"
        "提高至0.188、0.084和2.260；验证选择收益分别从0.103、0.072和1.325提高至0.172、0.078和2.179。现代增强池的验证选择收益"
        "分别为0.176、0.079和2.191。两种组成干预相对于同质池的K=32观测审计最佳与验证选择收益差在五个随机种子均为正。"
    )
    paragraph_by_prefix(doc, "K=32原始熵秩").text = (
        "K=32原始熵秩在同质池中分别为10.88、2.90和1.87，在经典多表征池中为12.40、6.09和9.07，在现代增强池中为13.96、7.92"
        "和11.47。18个终点-池-K单元的相对熵秩与同单元或交叉拟合选择差距的Spearman相关为-0.008和-0.211。这些单元重复使用三个终点、"
        "不同候选池和K，并非18个独立实验，因此相关仅作描述。ESOL现代池CAHit@3为-0.030，但验证选择收益高于同质池，进一步表明排序命中"
        "与效用兑现并不等价。"
    )
    paragraph_by_prefix(doc, "现代池每个外层单元消耗").text = (
        "现代池每个外层单元消耗139–187秒观测下游审计时间，经典多表征池为18–41秒，同质池为21–31秒；现代池在三个终点的单位下游审计"
        "小时收益均低于经典多表征池。该比较不包含编码器预训练、模型获取和缓存嵌入提取成本。"
    )
    paragraph_by_prefix(doc, "图7.").text = ZH_FIG7
    paragraph_by_prefix(doc, "图1 ").text = (
        "图1 候选池扩张与模型选择损失的回顾性嵌套审计框架。中央模块将5个种子划分中3个内层骨架折的候选排序，与3个外层折的一次性审计"
        "严格分离，并禁止外层标签反馈。左侧输入已登记终点和预设候选前缀；周围分支汇总效用模式多样性、机会校正排序、同单元与交叉拟合"
        "选择差距、切分机制迁移、等规模注册表组成、有限审计赢家乐观偏差、化学支持边界以及四模型预测/错误可靠性。所有证据流汇入可审计"
        "证据图谱，而非模型排行榜。"
    )
    paragraph_by_prefix(doc, "图6 ").text = (
        "图6 化学支持边界下的预测可靠性。A在对称四模型矩阵中合并预测相关与高错误Jaccard重叠，并使用两个独立色标。B将分类ROC-AUC、"
        "回归RMSE、分类假阴性率以及分类和回归高错误富集整合为Tanimoto支持风险矩阵；单元格文字给出自然尺度中位数，颜色仅编码行内不利"
        "方向。C在对数比值轴上报告新骨架相对于已见或相关骨架的错误重叠、模型分歧、高错误富集和假阴性富集变化。D将四个固定配置模型的"
        "判别能力与少数类安全性指标分组展示；预测集大小仅为绘图从1–2尺度映射至0–1。"
    )
    paragraph_by_prefix(doc, "固定K与折分后").text = (
        "固定K与折分后，经典和现代候选均提高了观测有限审计性能上限，并通常允许验证选择器兑现大部分收益。18个重复审计单元中有效多样性"
        "与差距缺乏正向单调关系，这一描述性结果说明有效秩是结构审计量而不是因果惩罚参数。有效互补性不等于多样性越高越好；其价值取决于"
        "内层验证排序能否稳定识别新增信息。"
    )
    paragraph_by_prefix(doc, "因此，候选资格数量").text = (
        "因此，候选资格数量、独立效用运动、有限审计机会、实际兑现和下游计算价格需要分开报告。本研究中经典多表征池位于更有利的观测下游"
        "成本—收益区域；冻结ChemBERTa和MoLFormer表示及锁定的一轮D-MPNN扩展了组成边界，但图7不构成现代模型排行榜，也不能替代计算"
        "预算匹配的充分调优架构基准。"
    )
    replace_embedded_figures(doc)
    doc.save(path)
    if "oracle" in "\n".join(p.text for p in Document(path).paragraphs).lower():
        raise ValueError("Chinese manuscript still contains prohibited terminology")


def update_chinese_captions_only(path: Path) -> None:
    doc = Document(path)
    paragraph_by_prefix(doc, "图1 ").text = (
        "图1 候选池扩张与模型选择损失的回顾性嵌套审计框架。中央模块将5个种子划分中3个内层骨架折的候选排序，与3个外层折的一次性审计"
        "严格分离，并禁止外层标签反馈。左侧输入已登记终点和预设候选前缀；周围分支汇总效用模式多样性、机会校正排序、同单元与交叉拟合"
        "选择差距、切分机制迁移、等规模注册表组成、有限审计赢家乐观偏差、化学支持边界以及四模型预测/错误可靠性。所有证据流汇入可审计"
        "证据图谱，而非模型排行榜。"
    )
    paragraph_by_prefix(doc, "图6 ").text = (
        "图6 化学支持边界下的预测可靠性。A在对称四模型矩阵中合并预测相关与高错误Jaccard重叠，并使用两个独立色标。B将分类ROC-AUC、"
        "回归RMSE、分类假阴性率以及分类和回归高错误富集整合为Tanimoto支持风险矩阵；单元格文字给出自然尺度中位数，颜色仅编码行内不利"
        "方向。C在对数比值轴上报告新骨架相对于已见或相关骨架的错误重叠、模型分歧、高错误富集和假阴性富集变化。D将四个固定配置模型的"
        "判别能力与少数类安全性指标分组展示；预测集大小仅为绘图从1–2尺度映射至0–1。"
    )
    doc.save(path)
    (TARGET / "Chinese_Figure_Titles_and_Legends.txt").write_text(
        "\n\n".join(paragraph.text for paragraph in doc.paragraphs if any(paragraph.text.startswith(f"图{i}") for i in range(1, 8))),
        encoding="utf-8",
    )


def update_reviewer(path: Path) -> None:
    doc = Document(path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace("oracle and selected gain", "observed audit-best and validation-selected gain")
                cell.text = cell.text.replace("oracle gain", "observed audit-best gain")
                cell.text = cell.text.replace("oracle", "observed audit-best")
    doc.save(path)


def update_workbook(path: Path) -> None:
    workbook = load_workbook(path)
    replacements = {
        "oracle_validation_rank": "observed_audit_best_validation_rank",
        "heldout_oracle_variant": "heldout_audit_best_variant",
        "heldout_oracle_regret": "heldout_audit_best_regret",
        "oracle_gap": "audit_best_gap",
        "oracle_match": "audit_best_match",
        "oracle_opportunity_gain_mean": "observed_audit_best_opportunity_gain_mean",
        "oracle_opportunity_gain_low": "observed_audit_best_opportunity_gain_low",
        "oracle_opportunity_gain_high": "observed_audit_best_opportunity_gain_high",
        "oracle_candidate": "observed_audit_best_candidate",
        "oracle_utility": "observed_audit_best_utility",
        "oracle_opportunity_gain": "observed_audit_best_opportunity_gain",
    }
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    value = cell.value
                    for old, new in replacements.items():
                        value = value.replace(old, new)
                    value = value.replace("oracle", "observed_audit_best")
                    cell.value = value
    definitions = {
        "S28 Endpoint summary": "Observed audit-best opportunity gain is the finite outer-audit maximum utility minus the shared-anchor utility.",
        "S30 Selection units": "Observed audit-best denotes the best candidate within the finite audited registry, not a population optimum.",
        "S31 Paired pool effects": "Paired effects compare finite observed audit-best or validation-selected gains under matched seeds and folds.",
    }
    for sheet_name, note in definitions.items():
        if sheet_name in workbook.sheetnames:
            workbook[sheet_name]["A1"].comment = Comment(note, "OpenAI Codex")
    workbook.save(path)


def write_public_source_tables() -> None:
    exports = TARGET / "figure_source_data"
    rename = {
        "oracle_opportunity_gain_mean": "observed_audit_best_opportunity_gain_mean",
        "oracle_opportunity_gain_low": "observed_audit_best_opportunity_gain_low",
        "oracle_opportunity_gain_high": "observed_audit_best_opportunity_gain_high",
        "oracle_candidate": "observed_audit_best_candidate",
        "oracle_utility": "observed_audit_best_utility",
        "oracle_opportunity_gain": "observed_audit_best_opportunity_gain",
    }
    for source_name, target_name in [
        ("equal_size_endpoint_summary.csv", "Table_S28_endpoint_summary_source.csv"),
        ("equal_size_selection_units.csv", "Table_S30_selection_units_source.csv"),
        ("equal_size_paired_pool_effects.csv", "Table_S31_paired_pool_effects_source.csv"),
    ]:
        frame = pd.read_csv(ANALYSIS / source_name).rename(columns=rename)
        for column in frame.select_dtypes(include="object").columns:
            frame[column] = frame[column].str.replace("oracle_opportunity_gain", "observed_audit_best_opportunity_gain", regex=False)
        frame.to_csv(exports / target_name, index=False)


def main() -> None:
    shutil.copytree(SOURCE, TARGET, dirs_exist_ok=True,
                    ignore=shutil.ignore_patterns("main_figures", "figure_source_data", "tracked_unpacked", "rendered"))
    for i in range(1, 7):
        descriptive = next(FIG.glob(f"Figure_{i}_*.pdf"))
        for suffix in ("pdf", "svg"):
            src = descriptive.with_suffix(f".{suffix}")
            shutil.copy2(src, FIG / f"Figure{i}.{suffix}")
        png = descriptive.with_suffix(".png")
        shutil.copy2(png, FIG / f"Figure{i}_600dpi.png")
    for suffix in ("pdf", "svg"):
        shutil.copy2(FIG / f"Figure_7_equal_size_registry_composition.{suffix}", FIG / f"Figure7.{suffix}")
    shutil.copy2(FIG / "Figure_7_equal_size_registry_composition_600dpi.png", FIG / "Figure7_600dpi.png")

    english = TARGET / EN_NAME
    chinese = next(path for path in TARGET.glob("*.docx") if any(ord(char) > 127 for char in path.name))
    update_english(english)
    update_chinese(chinese)
    update_reviewer(TARGET / REVIEWER)
    workbook = TARGET / "supplementary" / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S31.xlsx"
    update_workbook(workbook)
    for old in (TARGET / "supplementary").glob("Additional_file_2_Machine_readable_Supplementary_Tables_S1-S*.xlsx"):
        if old != workbook:
            old.unlink()
    write_public_source_tables()

    en_doc = Document(english)
    zh_doc = Document(chinese)
    en_captions = [p.text for p in en_doc.paragraphs if p.text.startswith("Figure ")]
    zh_captions = [p.text for p in zh_doc.paragraphs if p.text.startswith("图") and "." in p.text[:5]]
    (TARGET / "English_Figure_Titles_and_Legends.txt").write_text("\n\n".join(en_captions), encoding="utf-8")
    (TARGET / "Chinese_Figure_Titles_and_Legends.txt").write_text("\n\n".join(zh_captions), encoding="utf-8")
    paragraph_by_prefix(en_doc, "Methods:")
    (TARGET / "Revised_Abstract.txt").write_text("\n".join(en_doc.paragraphs[i].text for i in (2, 3, 4, 5)), encoding="utf-8")
    print(TARGET)


if __name__ == "__main__":
    main()
