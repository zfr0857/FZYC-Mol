from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(r"D:\fzyc\output\paper33_final_minor_revision_20260718")
FILES = {
    "english": ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_FINAL_CLEAN.docx",
    "chinese": ROOT / "候选池扩张与模型选择损失_中文终稿.docx",
}


def xml_counts(path: Path) -> dict[str, int]:
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
        return {
            "native_math_objects": len(re.findall(r"<(?:m|ns\d+):oMath(?:\s|>)", xml)),
            "drawing_objects": len(re.findall(r"<w:drawing(?:\s|>)", xml)),
        }


def inspect(name: str, path: Path) -> dict:
    doc = Document(path)
    texts = [p.text.strip() for p in doc.paragraphs]
    full = "\n".join(texts)
    abstract_lines = []
    for paragraph in texts:
        if paragraph.startswith(("Background:", "Methods:", "Results:", "Conclusions:", "Scientific Contribution:")):
            abstract_lines.append(paragraph.split(":", 1)[1])
    declaration_heads = [
        "Ethics approval and consent to participate", "Consent for publication",
        "Availability of data and materials", "Competing interests", "Funding",
        "Authors' contributions", "Acknowledgements",
    ]
    result = {
        "file": str(path),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "abstract_words": len(re.findall(r"\b[\w’'-]+\b", " ".join(abstract_lines))) if name == "english" else None,
        "old_270_present": "270 outer" in full or "270个" in full,
        "old_three_endpoint_k16_present": "3 endpoints; K = 16 and 32" in full or "3个端点；K = 16" in full,
        "new_primary_exposure_present": "1,080" in full and "64,616.35" in full,
        "supplement_tables_S36_present": "S1–S36" in full,
        "supplement_figures_S21_present": "S1–S21" in full,
        "abstract_14_18_present": any("14 of 18" in x or "14/18" in x for x in texts[:20]),
        "abstract_10_18_present": any("10 of 18" in x or "10/18" in x for x in texts[:20]),
    }
    result.update(xml_counts(path))
    if name == "english":
        result["declaration_headings_present"] = {h: any(t == h for t in texts) for h in declaration_heads}
        result["author_confirmation_count"] = full.count("Author confirmation required before submission")
    else:
        result["declarations_present"] = "声明" in texts
        result["author_confirmation_count"] = full.count("投稿前须由作者确认")
    return result


def main() -> None:
    report = {name: inspect(name, path) for name, path in FILES.items()}
    out = ROOT / "Manuscript_structural_QC.json"
    out.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
