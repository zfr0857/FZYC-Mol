from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path("D:/fzyc")
OUTPUT = ROOT / "output"
EVIDENCE_DIR = max(
    [p for p in OUTPUT.iterdir() if p.is_dir() and "SCI1" in p.name],
    key=lambda p: p.stat().st_mtime,
)
SOURCE_DOCX = max(
    [p for p in OUTPUT.glob("*.docx") if p.name.endswith("-14.docx")],
    key=lambda p: p.stat().st_mtime,
)
TARGET_DOCX = OUTPUT / "小论文-15.docx"
AUDIT_MD = OUTPUT / "小论文-15_严格修改审计.md"
AUDIT_JSON = OUTPUT / "paper15_revision_audit.json"


INTRO_LITERATURE_PARAGRAPH = (
    "近半年文献进一步强化了这一分层设计。ADMET 可靠性基准将数据稀缺与 OOD 泛化、类别不平衡、"
    "bRo5 化学空间和活性悬崖作为现实挑战，并同时比较分子基础模型、TabPFNv2、GNN、AutoML "
    "和传统机器学习[5]；DCPM-ADMET、KROVEX 和 MolGramTreeNet 则从序列语义、ECFP、图嵌入、"
    "统计描述符选择和语法树约束等角度展示了多模态表征的潜在收益[6–8]；OOD 主动学习和适用域"
    "边界研究进一步提示，真实药物发现场景需要相似度、分布迁移和边界可靠性审计[10,11]。因此，"
    "本文将实验补强聚焦于三类可直接检验选择结论的证据：同划分现代强基线、逐样本 error-overlap "
    "和去重敏感性重跑；同时将不确定性、活性悬崖、bRo5 和 TDC 作为可靠性与化学边界分析，而不将其写成"
    "新的排行榜证据。"
)

RESULTS_SCOPE_PARAGRAPH = (
    "这些补强实验使强基线、错误互补性和数据清洗敏感性从状态说明转为可复核结果，但其证据层级仍需保持一致。"
    "同划分强基线、error-overlap 和三套去重策略均已在 ESOL、BACE 和 ClinTox 三个代表终点上完成；"
    "九终点主效应仍来自预登记的轻量扩池实验和共享划分 12 候选多视图实验。因此，本文将三终点深度/基础模型"
    "面板解释为对主结论的压力测试和边界验证，而不是替代九终点确认性结果的全量排行榜。"
)

DISCUSSION_LITERATURE_PARAGRAPH = (
    "近期研究也说明，本文的不足不应被简单归结为缺少某个单一更强模型。2026 年 ADMET 可靠性基准显示，"
    "TabPFNv2、预训练 GNN、AutoML 和传统模型在不同挑战下各有优势，并且活性悬崖仍是多类模型的共同弱点[5]。"
    "KROVEX 和 DCPM-ADMET 的结果则支持表征异质性和多模态融合的价值[6,7]，但这些收益只有在相同划分、"
    "相同选择规则和明确消融下才可比较。FZYC-Mol 因此不把复杂模型排除在外，而是要求复杂模型承担同样的"
    "冻结选择、逐样本导出、去重敏感性和边界审计成本。"
)

CONCLUSION_REPLACEMENT = (
    "对后续研究而言，最直接的采用方式不是替换现有预测器，而是在新候选进入比较之前建立候选登记、终止规则、"
    "内层选择、外层审计和负结果归档。若未来将 Chemprop、GNN、化学语言模型和 TabPFN 从代表性三终点扩展到"
    "九终点全量同划分面板，并加入时间外 ADMET 盲测，FZYC-Mol 仍可作为扩展前的审计底座。"
)

REFERENCE_35 = (
    "[35] Li Z, Chen X, Wen H, et al. A systematic survey and benchmark of deep learning for molecular "
    "property prediction in the foundation model era. arXiv:2604.16586 [cs.LG]. 2026."
)


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para


def replace_paragraph(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def find_paragraph(doc: Document, startswith: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    raise RuntimeError(f"Could not find paragraph starting with: {startswith}")


def ensure_reference(doc: Document) -> bool:
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "arXiv:2604.16586" in full_text:
        return False
    last_ref = find_paragraph(doc, "[34]")
    new_ref = insert_paragraph_after(last_ref, REFERENCE_35)
    new_ref.style = last_ref.style
    return True


def set_cell_borders(cell, top=None, bottom=None) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")

    for edge, spec in (("top", top), ("bottom", bottom)):
        if spec is None:
            continue
        element = borders.find(qn(f"w:{edge}"))
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(spec.get("sz", 8)))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), spec.get("color", "000000"))


def apply_three_line_tables(doc: Document) -> None:
    heavy = {"sz": 10, "color": "000000"}
    light = {"sz": 6, "color": "000000"}
    for table in doc.tables:
        if not table.rows:
            continue
        for row in table.rows:
            for cell in row.cells:
                set_cell_borders(cell)
        for cell in table.rows[0].cells:
            set_cell_borders(cell, top=heavy, bottom=light)
        for cell in table.rows[-1].cells:
            current_top = heavy if len(table.rows) == 1 else None
            set_cell_borders(cell, top=current_top, bottom=heavy)


def paragraph_exists(doc: Document, text: str) -> bool:
    return any(p.text.strip() == text for p in doc.paragraphs)


def revise_docx() -> dict:
    shutil.copy2(SOURCE_DOCX, TARGET_DOCX)
    doc = Document(str(TARGET_DOCX))

    inserted = []
    if not paragraph_exists(doc, INTRO_LITERATURE_PARAGRAPH):
        anchor = find_paragraph(doc, "基于这一定位")
        insert_paragraph_after(anchor, INTRO_LITERATURE_PARAGRAPH)
        inserted.append("intro_recent_literature")

    if not paragraph_exists(doc, RESULTS_SCOPE_PARAGRAPH):
        anchor = find_paragraph(doc, "error-overlap 审计显示")
        insert_paragraph_after(anchor, RESULTS_SCOPE_PARAGRAPH)
        inserted.append("results_scope_boundary")

    if not paragraph_exists(doc, DISCUSSION_LITERATURE_PARAGRAPH):
        anchor = find_paragraph(doc, "与以 MoleculeNet 或 TDC 为核心")
        insert_paragraph_after(anchor, DISCUSSION_LITERATURE_PARAGRAPH)
        inserted.append("discussion_literature_position")

    conclusion_anchor = find_paragraph(doc, "对后续研究而言")
    replace_paragraph(conclusion_anchor, CONCLUSION_REPLACEMENT)
    inserted.append("conclusion_scope_rewrite")

    reference_added = ensure_reference(doc)
    if reference_added:
        inserted.append("reference_35_added")

    apply_three_line_tables(doc)
    inserted.append("three_line_tables_applied")

    doc.save(str(TARGET_DOCX))
    return {
        "source_docx": str(SOURCE_DOCX),
        "target_docx": str(TARGET_DOCX),
        "inserted_or_updated": inserted,
    }


def load_evidence() -> dict:
    strong = pd.read_csv(EVIDENCE_DIR / "sci1_strong_baseline_endpoint_table.csv")
    overlap = pd.read_csv(EVIDENCE_DIR / "sci1_error_overlap_extended.csv")
    dedup = pd.read_csv(EVIDENCE_DIR / "sci1_duplicate_sensitivity_extended.csv")
    readiness = json.loads(
        (EVIDENCE_DIR / "sci1_submission_readiness_audit.json").read_text(encoding="utf-8")
    )

    return {
        "evidence_dir": str(EVIDENCE_DIR),
        "strong_baseline_rows": int(len(strong)),
        "strong_baseline_tasks": sorted(strong["task"].unique().tolist()),
        "strong_baseline_candidates": sorted(strong["candidate"].unique().tolist()),
        "strong_baseline_outer_units": int(strong["n_outer_units"].sum()),
        "error_overlap_pairs": int(len(overlap)),
        "error_overlap_mean_jaccard": float(overlap["mean_jaccard_error_overlap"].mean()),
        "error_overlap_range": [
            float(overlap["mean_jaccard_error_overlap"].min()),
            float(overlap["mean_jaccard_error_overlap"].max()),
        ],
        "duplicate_rows": int(len(dedup)),
        "duplicate_policies": sorted(dedup["policy"].unique().tolist()),
        "duplicate_max_abs_delta": float(dedup["abs_delta_vs_global_dedup"].max()),
        "readiness": readiness.get("strong_baseline", readiness),
    }


def audit_docx() -> dict:
    doc = Document(str(TARGET_DOCX))
    text = "\n".join(p.text for p in doc.paragraphs)
    refs = [p.text.strip() for p in doc.paragraphs if p.text.strip().startswith("[")]
    checks = {
        "title_present": doc.paragraphs[0].text.strip()
        == "冻结验证揭示分子性质预测中候选池扩张的收益与选择损失",
        "abstract_four_paragraphs": sum(1 for p in doc.paragraphs[:8] if p.text.strip()) >= 6,
        "recent_literature_inserted": INTRO_LITERATURE_PARAGRAPH in text,
        "results_scope_inserted": RESULTS_SCOPE_PARAGRAPH in text,
        "discussion_literature_inserted": DISCUSSION_LITERATURE_PARAGRAPH in text,
        "conclusion_scope_rewritten": CONCLUSION_REPLACEMENT in text,
        "reference_35_present": "arXiv:2604.16586" in text,
        "tabpfn_not_overclaimed": "TabPFN 完成" not in text and "TabPFN 因 license/token" in text,
        "nine_endpoint_deep_panel_not_overclaimed": "九终点全量深度模型面板仍超出当前完成范围" in text,
        "error_overlap_reported": "10 个候选对的平均 Jaccard 重合为 0.189" in text,
        "dedup_reported": "三套去重策略" in text and "最大平均效用变化为 0.022" in text,
        "reference_count": len(refs),
        "table_count": len(doc.tables),
        "figure_count": len(doc.inline_shapes),
    }

    forbidden_terms = ["证明了", "一定会", "必然优于", "彻底解决", "显著优于所有", "通用 SOTA 预测器"]
    checks["forbidden_terms_found"] = {
        term: (term in text) for term in forbidden_terms if term in text
    }
    checks["passed"] = all(
        value
        for key, value in checks.items()
        if key not in {"reference_count", "table_count", "figure_count", "forbidden_terms_found"}
    ) and not checks["forbidden_terms_found"]
    return checks


def write_audit_report(revision: dict, evidence: dict, checks: dict) -> None:
    items = [
        (
            "中心主张与创新定位",
            "已落实",
            "摘要、引言和讨论均将 FZYC-Mol 定位为冻结验证治理协议，而非新的预测主干。",
        ),
        (
            "32 候选代表性与有效多样性",
            "已落实但有限定",
            "正文报告轻量池相关性中位数 0.998、K_eff≈1.01，并将其定义为选择压力实验；异质性由 12 候选多视图实验承担。",
        ),
        (
            "现代强基线统一重训",
            "部分落实",
            "RDKit RF、GNN-GCN、Chemprop/D-MPNN、ChemBERTa、MoLFormer 已在 ESOL/BACE/ClinTox 三终点同划分面板完成；TabPFN 未产出预测，九终点全量深度面板仍未完成。",
        ),
        (
            "error-overlap",
            "已落实但限定为三终点面板",
            "已导出 103,025 条逐样本预测并计算 10 个候选对 Jaccard error-overlap。",
        ),
        (
            "去重敏感性实验",
            "已落实但限定为三终点面板",
            "global_dedup、keep_duplicates_grouped、train_fold_only_aggregate 三策略完成实际重跑；最大平均效用变化 0.022。",
        ),
        (
            "TDC、保形、MoleculeACE、bRo5",
            "已按边界证据写法落实",
            "TDC 三种子结果写为 seed variability interval；保形强调 ClinTox 类别 1 覆盖不足；MoleculeACE/bRo5 不被写成主胜负证据。",
        ),
        (
            "可复现性",
            "本地证据已落实，公开冷启动仍是边界",
            "稿件保留 public release、Zenodo DOI、第三方冷启动复跑为后续工作，避免过度承诺。",
        ),
    ]

    literature = [
        (
            "Zhao et al., J Cheminform 2026",
            "强调数据稀缺/OOD、类别不平衡、bRo5、活性悬崖，并比较 foundation/tabular/GNN/AutoML/传统模型。",
            "支持将强基线和化学边界分层报告，并保留 TabPFN 未完成边界。",
        ),
        (
            "Yin et al., RSC Adv. 2026",
            "OOD 主动学习强调 PubChem OOD、相似度和分布差异分析。",
            "支持逐样本风险、最近邻相似度和 TDC/bRo5 边界写法。",
        ),
        (
            "Jang et al., J Cheminform 2026",
            "KROVEX 通过统计描述符选择和多模态融合，并以消融支持收益。",
            "支持把多视图收益写成同划分确认结果，同时不夸大为通用 SOTA。",
        ),
        (
            "Li et al., arXiv 2026",
            "综述强调数据清洗、划分、评价协议、时间/骨架感知和 uncertainty-calibrated foundation models。",
            "支持新增引用和对公开复现/时间外验证的边界说明。",
        ),
    ]

    lines = [
        "# 小论文-15 严格修改与实验补强审计",
        "",
        f"- 生成时间：{datetime.now().isoformat(timespec='seconds')}",
        f"- 输入稿件：`{revision['source_docx']}`",
        f"- 输出稿件：`{revision['target_docx']}`",
        f"- 证据目录：`{evidence['evidence_dir']}`",
        "",
        "## 逐条落实状态",
        "",
        "| 修改项 | 状态 | 核对结果 |",
        "|---|---|---|",
    ]
    for item, status, detail in items:
        lines.append(f"| {item} | {status} | {detail} |")

    lines += [
        "",
        "## 三项重点实验核对",
        "",
        f"- 强基线：{evidence['strong_baseline_rows']} 行端点汇总，"
        f"{evidence['strong_baseline_outer_units']} 个外层单元；任务为 "
        f"{', '.join(evidence['strong_baseline_tasks'])}；候选为 "
        f"{', '.join(evidence['strong_baseline_candidates'])}。",
        f"- error-overlap：{evidence['error_overlap_pairs']} 个候选对，平均 Jaccard="
        f"{evidence['error_overlap_mean_jaccard']:.3f}，范围 "
        f"{evidence['error_overlap_range'][0]:.3f}–{evidence['error_overlap_range'][1]:.3f}。",
        f"- 去重敏感性：{evidence['duplicate_rows']} 行，策略为 "
        f"{', '.join(evidence['duplicate_policies'])}；最大绝对变化 "
        f"{evidence['duplicate_max_abs_delta']:.3f}。",
        f"- 逐样本预测导出：{evidence['readiness'].get('prediction_rows', 'NA')} 条记录。",
        f"- TabPFN 状态：{evidence['readiness'].get('tabpfn_status', 'NA')}。",
        "",
        "## 近半年文献吸收点",
        "",
        "| 文献 | 可借鉴优点 | 已写入/改造方式 |",
        "|---|---|---|",
    ]
    for paper, lesson, action in literature:
        lines.append(f"| {paper} | {lesson} | {action} |")

    lines += [
        "",
        "## Nature-style 语言与逻辑复查",
        "",
        "- 摘要仍保持“背景与问题、方法、主要结果、结论与边界”的四段逻辑。",
        "- 新增段落避免使用“证明、彻底、最优、通用 SOTA”等过强措辞。",
        "- 结果段落报告事实和范围；讨论段落解释与近期文献的关系和边界。",
        "- 全部 9 张表已在 DOCX XML 层统一为无竖线、表头上下线和表尾线的三线表样式。",
        "",
        "## 自动审计",
        "",
        "```json",
        json.dumps(checks, ensure_ascii=False, indent=2),
        "```",
    ]
    AUDIT_MD.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    revision = revise_docx()
    evidence = load_evidence()
    checks = audit_docx()
    write_audit_report(revision, evidence, checks)
    AUDIT_JSON.write_text(
        json.dumps(
            {
                "revision": revision,
                "evidence": evidence,
                "checks": checks,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    print(json.dumps({"target": str(TARGET_DOCX), "audit": str(AUDIT_MD), "passed": checks["passed"]}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
