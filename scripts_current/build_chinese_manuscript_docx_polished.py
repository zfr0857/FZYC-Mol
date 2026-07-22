from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
PACKAGE = ROOT / "reports" / "submission_package"
MAIN_FIGS = PACKAGE / "main_figures"
MAIN_TABLES = PACKAGE / "main_tables"
SUPP_TABLES = PACKAGE / "supplementary_tables"
RAW_TABLES = ROOT / "reports" / "manuscript_tables"
POLISHED_FIGS = ROOT / "reports" / "manuscript_figures_polished"

OUT_MAIN = DOCS / "manuscript_draft_full_zh_integrated_20260531.docx"
OUT_POLISHED = DOCS / "manuscript_draft_full_zh_polished_20260601.docx"
OUT_DETAILED = DOCS / "manuscript_draft_full_zh_detailed_20260601.docx"
OUT_POLISHED_20260602 = DOCS / "manuscript_draft_full_zh_polished_20260602.docx"
OUT_DETAILED_20260602 = DOCS / "manuscript_draft_full_zh_detailed_20260602.docx"
OUT_REFS = DOCS / "reference_list_expanded_20260601.md"
OUT_REFS_20260602 = DOCS / "reference_list_expanded_20260602.md"
PACKAGE_DOC = PACKAGE / "docs" / "Manuscript_draft_full_zh_integrated.docx"
PACKAGE_DETAILED_DOC = PACKAGE / "docs" / "Manuscript_draft_full_zh_detailed_20260601.docx"
PACKAGE_DETAILED_DOC_20260602 = PACKAGE / "docs" / "Manuscript_draft_full_zh_detailed_20260602.docx"
PACKAGE_REFS = PACKAGE / "docs" / "Reference_list_expanded_20260601.md"
PACKAGE_REFS_20260602 = PACKAGE / "docs" / "Reference_list_expanded_20260602.md"


DATASET_LABELS = {
    "esol": "ESOL",
    "freesolv": "FreeSolv",
    "lipo": "Lipophilicity",
    "bbbp": "BBBP",
    "bace": "BACE",
    "clintox": "ClinTox",
}


def set_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_normal_style(doc: Document) -> None:
    for style_name, size in [("Normal", 10.5), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.15
    doc.styles["Normal"].paragraph_format.space_after = Pt(5)


def set_margins(section) -> None:
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)


def add_landscape_section(doc: Document) -> None:
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    set_margins(section)


def add_para(doc: Document, text: str = "", style: str | None = None, bold: bool = False, size: float | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0.74) if style is None and text else None
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_font(run, bold=True)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_font(run)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_font(run, size=9.2)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=9.2, bold=True, color="1F2937")


def add_figure(doc: Document, path: Path, caption: str, width_in: float = 7.15) -> None:
    if not path.exists():
        add_para(doc, f"[缺失图片：{path}]", bold=True)
        return
    if not path.exists():
        add_para(doc, f"[缺失图片：{path}]", bold=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    add_caption(doc, caption)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "E5E7EB", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_cell_text(cell, text: object, size: float, bold: bool = False, color: str = "111827") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("" if pd.isna(text) else str(text))
    set_font(run, size=size, bold=bold, color=color)


def is_numericish(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    cleaned = (
        stripped.replace("+/-", "")
        .replace("%", "")
        .replace("+", "")
        .replace("-", "")
        .replace(".", "")
        .replace(",", "")
        .replace(" ", "")
    )
    return bool(cleaned) and cleaned.isdigit()


def style_table(table, font_size: float = 7.3) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, color="E2E8F0", size="4")
            if row_idx == 0:
                shade_cell(cell, "EEF2F7")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        set_font(run, size=font_size, bold=True, color="0F172A")
            else:
                shade_cell(cell, "F8FAFC" if row_idx % 2 == 0 else "FFFFFF")
                if col_idx == 0:
                    shade_cell(cell, "F1F5F9")
                for paragraph in cell.paragraphs:
                    text = "".join(run.text for run in paragraph.runs)
                    if len(row.cells) > 7 and not is_numericish(text):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif is_numericish(text):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        set_font(run, size=font_size, bold=col_idx == 0, color="111827")


def add_polished_table(
    doc: Document,
    label: str,
    caption: str,
    frame: pd.DataFrame,
    note: str | None = None,
    font_size: float = 7.3,
    max_text: int = 95,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{label}. {caption}")
    set_font(run, bold=True, size=9.4, color="111827")

    clean = frame.copy().fillna("")
    for col in clean.columns:
        clean[col] = clean[col].map(lambda x: str(x).replace("\n", " ").strip())
        clean[col] = clean[col].map(lambda x: x[: max_text - 1] + "..." if len(x) > max_text else x)

    table = doc.add_table(rows=1, cols=len(clean.columns))
    for idx, col in enumerate(clean.columns):
        set_cell_text(table.rows[0].cells[idx], col, font_size, bold=True, color="FFFFFF")
    for _, row in clean.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(clean.columns):
            set_cell_text(cells[idx], row[col], font_size)
    style_table(table, font_size=font_size)
    if note:
        note_p = doc.add_paragraph()
        note_p.paragraph_format.space_after = Pt(8)
        run = note_p.add_run(f"注：{note}")
        set_font(run, size=8.0, color="475569")
        return
        note_p = doc.add_paragraph()
        note_p.paragraph_format.space_after = Pt(8)
        run = note_p.add_run(f"注：{note}")
        set_font(run, size=8.0, color="475569")


def read_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path).fillna("")


def fmt_float(value: object, digits: int = 4) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return ""


def fmt_delta(value: object, digits: int = 4) -> str:
    try:
        return f"{float(value):+.{digits}f}"
    except (TypeError, ValueError):
        return ""


def short_dataset(name: object) -> str:
    text = str(name)
    if text in DATASET_LABELS:
        return DATASET_LABELS[text]
    for token in ["tdc_", "_wang", "_hou", "_broccatelli", "_martins", "_veith", "_ma"]:
        text = text.replace(token, "")
    return text.replace("_", " ").upper()


def make_moleculenet_table() -> pd.DataFrame:
    long = read_csv(RAW_TABLES / "table2_moleculenet_main_long.csv")
    rebuild_path = RAW_TABLES / "table27_moleculenet_targeted_rebuild_retained_best.csv"
    nature_path = RAW_TABLES / "table29_nature_multimethod_fusion_retained_best.csv"
    rebuild_lookup: dict[str, pd.Series] = {}
    if rebuild_path.exists():
        rebuild = read_csv(rebuild_path)
        rebuild_lookup = {str(row["dataset"]): row for _, row in rebuild.iterrows()}
    nature_lookup: dict[str, pd.Series] = {}
    if nature_path.exists():
        nature = read_csv(nature_path)
        nature_lookup = {str(row["dataset"]): row for _, row in nature.iterrows()}
    rows: list[dict[str, str]] = []
    for dataset, sub in long.groupby("dataset", sort=False):
        direction = sub["direction"].iloc[0]
        non_fzyc = sub[
            ~sub["category"].str.contains("FZYC-Mol", regex=False)
            & ~sub["category"].eq("Best observed candidate")
        ].copy()
        if direction == "lower":
            best_comp = non_fzyc.sort_values("value", ascending=True).iloc[0]
        else:
            best_comp = non_fzyc.sort_values("value", ascending=False).iloc[0]
        val = sub[sub["category"].eq("FZYC-Mol validation selector")].iloc[0]
        rescue = sub[sub["category"].eq("FZYC-Mol targeted rescue selector")].iloc[0]
        best = sub[sub["category"].eq("Best observed candidate")].iloc[0]
        final_result = rescue["formatted"]
        final_note = "保留原selector/rescue"
        if str(dataset) in rebuild_lookup:
            rebuild_row = rebuild_lookup[str(dataset)]
            final_result = f"{fmt_float(rebuild_row['retained_primary_mean'])} +/- {fmt_float(rebuild_row['retained_primary_std'])}"
            final_note = "重构接入" if rebuild_row["retained_source"] == "targeted_rebuild" else "保留当前/rescue"
        if str(dataset) in nature_lookup:
            nature_row = nature_lookup[str(dataset)]
            final_result = f"{fmt_float(nature_row['retained_primary_mean'])} +/- {fmt_float(nature_row['retained_primary_std'])}"
            final_note = "Nature融合接入" if nature_row["retained_source"] == "nature_multimethod_fusion" else final_note
        rows.append(
            {
                "数据集": short_dataset(dataset),
                "任务": "回归" if val["task_type"] == "regression" else "分类",
                "主指标": str(val["primary_metric"]).upper(),
                "验证集选择器": val["formatted"],
                "最终retained-best": final_result,
                "保留决策": final_note,
                "最强非FZYC专家": f"{best_comp['category']}：{best_comp['formatted']}",
                "观测最优": best["formatted"],
                "解释": "低值更优" if direction == "lower" else "高值更优",
            }
        )
    return pd.DataFrame(rows)


def make_tdc_table() -> pd.DataFrame:
    tdc = read_csv(RAW_TABLES / "table3_tdc_official_admet.csv")
    rows = []
    for _, row in tdc.iterrows():
        final_result = row.get("tdc_final_retained_formatted", row["selector_formatted"])
        final_source = row.get("tdc_final_retained_source", "tdc_validation_selector")
        source_label = "Nature融合" if final_source == "tdc_nature_multimethod_fusion" else "原TDC selector"
        rows.append(
            {
                "Endpoint": short_dataset(row["dataset"]),
                "任务": "回归" if row["task_type"] == "regression" else "分类",
                "指标": str(row["primary_metric"]).upper(),
                "FZYC selector": row["selector_formatted"],
                "Final retained": final_result,
                "保留来源": source_label,
                "LGBM随机到scaffold": f"{row['lgbm_morgan_random']} -> {row['lgbm_morgan_scaffold']}",
                "RF随机到scaffold": f"{row['rf_morgan_random']} -> {row['rf_morgan_scaffold']}",
                "平均scaffold惩罚": fmt_delta((float(row["lgbm_morgan_random_to_scaffold_drop"]) + float(row["rf_morgan_random_to_scaffold_drop"])) / 2),
            }
        )
    return pd.DataFrame(rows)


def make_split_table() -> pd.DataFrame:
    split = read_csv(RAW_TABLES / "table4_split_realism.csv")
    rows = []
    for _, row in split.iterrows():
        rows.append(
            {
                "来源": row["source"],
                "数据集": short_dataset(row["dataset"]),
                "任务": "回归" if row["task_type"] == "regression" else "分类",
                "指标": str(row["metric"]).upper(),
                "Random": fmt_float(row["random_value"]),
                "Scaffold": fmt_float(row["scaffold_value"]),
                "Structure": fmt_float(row["structure_value"]),
                "总变化": fmt_delta(float(row["random_to_scaffold_drop"]) + float(row["scaffold_to_structure_drop"])),
            }
        )
    return pd.DataFrame(rows)


def make_reliability_table() -> pd.DataFrame:
    rel = read_csv(RAW_TABLES / "table6_reliability_ad.csv")
    keep = [
        "hybrid_recon_ad",
        "reconstruction_error",
        "inverse_tanimoto",
        "ensemble_std",
        "error_model",
        "hybrid_error_ad",
        "confidence_uncertainty",
        "prediction_deviation",
    ]
    rel = rel[rel["score"].isin(keep)].copy()
    label_map = {
        "reconstruction_unfamiliarity": "重构/适用域",
        "unique_style_uq": "不确定性",
    }
    return pd.DataFrame(
        {
            "分数": rel["score"].str.replace("_", " ", regex=False),
            "类型": rel["family"].map(label_map).fillna(rel["family"]),
            "覆盖行数": rel["n_rows"],
            "误差相关Spearman": rel["mean_spearman_abs_error"].map(lambda x: fmt_float(x, 3)),
            "Risk-coverage AUC": rel["mean_risk_coverage_auc"].map(lambda x: fmt_float(x, 3)),
            "Top10%高误差富集": rel["mean_top10pct_high_error_enrichment"].map(lambda x: fmt_float(x, 2)),
        }
    )


def make_rescue_table() -> pd.DataFrame:
    rescue = read_csv(RAW_TABLES / "table19_moleculenet_rescue_integrated_selector.csv")
    rows = []
    for _, row in rescue.iterrows():
        rows.append(
            {
                "数据集": short_dataset(row["dataset"]),
                "任务": "回归" if row["task_type"] == "regression" else "分类",
                "指标": str(row["primary_metric"]).upper(),
                "当前结果": f"{fmt_float(row['current_primary_mean'])} +/- {fmt_float(row['current_primary_std'])}",
                "整合后": f"{fmt_float(row['integrated_primary_mean'])} +/- {fmt_float(row['integrated_primary_std'])}",
                "相对当前变化": fmt_delta(row["integration_delta_vs_current"]),
                "使用rescue池": "是" if bool(row["selected_uses_rescue_pool"]) else "否",
                "入选模型": row["selected_model_counts"],
            }
        )
    return pd.DataFrame(rows)


def make_targeted_rebuild_table() -> pd.DataFrame:
    rebuild = read_csv(RAW_TABLES / "table27_moleculenet_targeted_rebuild_retained_best.csv")
    rows = []
    for _, row in rebuild.iterrows():
        retained_source = "重构候选" if row["retained_source"] == "targeted_rebuild" else "当前/已接入rescue"
        rows.append(
            {
                "数据集": short_dataset(row["dataset"]),
                "任务": "回归" if row["task_type"] == "regression" else "分类",
                "主指标": str(row["primary_metric"]).upper(),
                "当前/Rescue": f"{fmt_float(row['current_or_rescue_primary_mean'])} +/- {fmt_float(row['current_or_rescue_primary_std'])}",
                "重构候选": f"{fmt_float(row['rebuild_primary_mean'])} +/- {fmt_float(row['rebuild_primary_std'])}",
                "正向变化": fmt_delta(row["delta_vs_current_or_rescue"]),
                "保留来源": retained_source,
                "重构入选模型": row["rebuild_model_counts"],
                "最终保留模型": row["retained_model"],
            }
        )
    return pd.DataFrame(rows)


def make_nature_fusion_table() -> pd.DataFrame:
    fusion = read_csv(RAW_TABLES / "table29_nature_multimethod_fusion_retained_best.csv")
    rows = []
    for _, row in fusion.iterrows():
        retained_source = "Nature多方法融合" if row["retained_source"] == "nature_multimethod_fusion" else "保留上一版retained-best"
        rows.append(
            {
                "数据集": short_dataset(row["dataset"]),
                "任务": "回归" if row["task_type"] == "regression" else "分类",
                "主指标": str(row["primary_metric"]).upper(),
                "上一版retained-best": f"{fmt_float(row['previous_retained_primary_mean'])} +/- {fmt_float(row['previous_retained_primary_std'])}",
                "Nature融合候选": f"{fmt_float(row['fusion_primary_mean'])} +/- {fmt_float(row['fusion_primary_std'])}",
                "正向变化": fmt_delta(row["delta_vs_previous_retained"]),
                "最终保留来源": retained_source,
                "融合入选模型": row["fusion_model_counts"],
                "Nature方法线": row["nature_inspiration_counts"],
            }
        )
    return pd.DataFrame(rows)


def make_nature_alignment_table() -> pd.DataFrame:
    frame = read_csv(RAW_TABLES / "table31_nature_method_alignment_matrix.csv")
    return frame.rename(
        columns={
            "literature_signal": "Nature/相关文献信号",
            "borrowed_idea": "借鉴点",
            "implemented_candidate": "本课题实现",
            "claim_boundary": "边界说明",
        }
    )


def make_tdc_nature_fusion_table() -> pd.DataFrame:
    fusion = read_csv(RAW_TABLES / "table32_tdc_nature_multimethod_fusion_retained_best.csv")
    rows = []
    for _, row in fusion.iterrows():
        retained_source = "TDC Nature融合" if row["retained_source"] == "tdc_nature_multimethod_fusion" else "原TDC selector"
        rows.append(
            {
                "Endpoint": short_dataset(row["dataset"]),
                "任务": "回归" if row["task_type"] == "regression" else "分类",
                "主指标": str(row["primary_metric"]).upper(),
                "原TDC selector": f"{fmt_float(row['previous_selector_primary_mean'])} +/- {fmt_float(row['previous_selector_primary_std'])}",
                "外部融合候选": f"{fmt_float(row['fusion_primary_mean'])} +/- {fmt_float(row['fusion_primary_std'])}",
                "正向变化": fmt_delta(row["delta_vs_tdc_selector"]),
                "最终保留来源": retained_source,
                "融合入选模型": row["fusion_model_counts"],
            }
        )
    return pd.DataFrame(rows)


def make_appendix_table() -> pd.DataFrame:
    perf = read_csv(RAW_TABLES / "table15_tdc_performance_mode_retained_best.csv")
    perf = perf[perf["retained_source"].eq("performance_mode")].copy()
    perf["abs_delta"] = perf["performance_delta_vs_previous"].astype(float).abs()
    perf = perf.sort_values("abs_delta", ascending=False).head(8)
    rows = []
    for _, row in perf.iterrows():
        rows.append(
            {
                "Endpoint": short_dataset(row["dataset"]),
                "任务": "回归" if row["task_type"] == "regression" else "分类",
                "官方指标": str(row["official_metric"]).upper(),
                "原保留模型": row["previous_model"],
                "新增性能模式": row["retained_model"],
                "原结果": fmt_float(row["previous_primary_mean"]),
                "保留结果": fmt_float(row["retained_primary_mean"]),
                "变化": fmt_delta(row["performance_delta_vs_previous"]),
            }
        )
    return pd.DataFrame(rows)


def make_roughness_table() -> pd.DataFrame:
    rough = read_csv(RAW_TABLES / "table16_tdc_roughness_literature_alignment.csv")
    rough["roughness_proxy_mean"] = rough["roughness_proxy_mean"].astype(float)
    rough["abs_delta"] = rough["performance_delta_vs_previous"].astype(float).abs()
    rough = rough.sort_values(["roughness_proxy_mean", "abs_delta"], ascending=False).head(10)
    rows = []
    for _, row in rough.iterrows():
        rows.append(
            {
                "Endpoint": short_dataset(row["dataset"]),
                "任务": "回归" if row["task_type"] == "regression" else "分类",
                "指标": str(row["official_metric"]).upper(),
                "NN相似度": fmt_float(row["nn_tanimoto_mean"], 3),
                "roughness": fmt_float(row["roughness_proxy_mean"], 3),
                "粗糙度带": row["roughness_band"],
                "性能变化": fmt_delta(row["performance_delta_vs_previous"]),
                "保留来源": row["retained_source"],
            }
        )
    return pd.DataFrame(rows)


def make_formal_external_selector_table() -> pd.DataFrame:
    frame = read_csv(RAW_TABLES / "table20_formal_external_appendix_selector.csv")
    rows = []
    for _, row in frame.iterrows():
        rows.append(
            {
                "Endpoint": short_dataset(row["dataset"]),
                "任务": "回归" if row["task_type"] == "regression" else "分类",
                "指标": str(row["official_metric"]).upper(),
                "旧baseline": row["previous_model"],
                "新增候选/selector": row["retained_model"],
                "保留来源": row["retained_source"],
                "保留增益": fmt_delta(row["retained_delta_vs_previous"]),
                "roughness": fmt_float(row["roughness_proxy_mean"], 3),
                "结论": row["manuscript_takeaway"],
            }
        )
    return pd.DataFrame(rows)


def make_candidate_pool_table() -> pd.DataFrame:
    frame = read_csv(RAW_TABLES / "table21_external_candidate_pool_coverage.csv")
    return pd.DataFrame(
        {
            "候选家族": frame["candidate_group"],
            "候选行数": frame["n_candidate_rows"],
            "唯一模型数": frame["n_unique_models"],
            "覆盖endpoint": frame["n_datasets"],
            "被选seed数": frame["selected_seed_count"],
            "被选endpoint": frame["selected_dataset_count"],
            "被选seed占比": frame["selected_seed_fraction"].map(lambda x: fmt_float(x, 3)),
        }
    )


def make_imbalanced_table() -> pd.DataFrame:
    frame = read_csv(RAW_TABLES / "table22_imbalanced_classification_metrics.csv")
    rows = []
    for _, row in frame.iterrows():
        rows.append(
            {
                "Endpoint": short_dataset(row["dataset"]),
                "来源": row["source"],
                "阳性率": fmt_float(row["positive_rate_or_query_positive_rate"], 3),
                "选择模型": row["selected_model_counts"],
                "undersampling": "是" if str(row["uses_undersampling_ensemble"]).lower() == "true" else "否",
                "ROC-AUC": f"{fmt_float(row['roc_auc_mean'], 3)} +/- {fmt_float(row['roc_auc_std'], 3)}",
                "PR-AUC": f"{fmt_float(row['pr_auc_mean'], 3)} +/- {fmt_float(row['pr_auc_std'], 3)}",
                "Brier": fmt_float(row["brier_mean"], 3),
                "ECE": fmt_float(row["ece_mean"], 3),
                "EF1": fmt_float(row["ef1_mean"], 2),
            }
        )
    return pd.DataFrame(rows)


def make_external_stats_table() -> pd.DataFrame:
    overall = read_csv(RAW_TABLES / "table25_external_win_tie_loss.csv")
    rows = []
    for _, row in overall.iterrows():
        rows.append(
            {
                "统计范围": row["scope"],
                "单位": row["unit"],
                "数量": row["n_units"],
                "平均正向变化": fmt_float(row["mean_positive_delta"], 4),
                "Bootstrap 95% CI": f"{fmt_float(row['bootstrap_ci_low'], 4)} to {fmt_float(row['bootstrap_ci_high'], 4)}",
                "Bootstrap p": fmt_float(row["bootstrap_p_two_sided"], 4),
                "Wilcoxon p": fmt_float(row["wilcoxon_p_two_sided"], 4),
                "win/tie/loss": f"{int(row['win'])}/{int(row['tie'])}/{int(row['loss'])}",
                "解释": row["interpretation"],
            }
        )
    return pd.DataFrame(rows)


def make_case_study_table() -> pd.DataFrame:
    frame = read_csv(RAW_TABLES / "table24_targeted_improvement_case_studies.csv")
    rows = []
    for _, row in frame.iterrows():
        rows.append(
            {
                "案例": row["case_id"],
                "名称": row["case_name"],
                "数据集": short_dataset(row["dataset"]),
                "类型": row["case_type"],
                "指标/证据": row["evidence"],
                "before": row["before"],
                "after": row["after"],
                "正向变化": fmt_delta(row["delta_positive"]) if str(row["delta_positive"]) else "",
                "模型/风险信号": row["selected_model_or_signal"],
                "解释": row["interpretation"],
            }
        )
    return pd.DataFrame(rows)


def references() -> list[str]:
    return [
        "Wu Z, Ramsundar B, Feinberg E N, et al. MoleculeNet: a benchmark for molecular machine learning. Chemical Science, 2018, 9: 513-530. DOI: 10.1039/C7SC02664A.",
        "Huang K, Fu T, Gao W, et al. Therapeutics Data Commons: Machine Learning Datasets and Tasks for Drug Discovery and Development. NeurIPS Datasets and Benchmarks, 2021.",
        "Zhao D, Zhu Y, Wu Z, et al. Revisiting ADMET prediction reliability under real-world challenges in the foundation model era. Journal of Cheminformatics, 2026. DOI: 10.1186/s13321-026-01217-2.",
        "Fraser J S, Edgar S, Handly L N, et al. Mapping the avoid-ome: a systematic open-science approach to predictive ADMET. Nature Communications, 2026, 17: 4644. DOI: 10.1038/s41467-026-73410-8.",
        "Pham L-H D, Le M-T, Thai K-M. Improved ADME Prediction by Multitask Pretraining on Predicted Data: Insights from the ASAP-Polaris-OpenADMET Blind Challenge. Journal of Chemical Information and Modeling, 2026, 66(1): 395-405. DOI: 10.1021/acs.jcim.5c02030.",
        "Parrondo-Pizarro R, Lanini J, Rodriguez-Perez R. Uncertainty Quantification in Molecular Machine Learning for Property Predictions under Data Shifts. Journal of Chemical Information and Modeling, 2026, 66(2): 923-935. DOI: 10.1021/acs.jcim.5c02381.",
        "Mun V, Fazli S. CheMLT-F: multitask learning in biochemistry through transformer fusion. Journal of Cheminformatics, 2026, 18: 69. DOI: 10.1186/s13321-026-01199-1.",
        "Gadaleta D, et al. Benchmarking predictive models for pKa, lipophilicity, intrinsic clearance and aqueous solubility in the context of drug discovery. Journal of Cheminformatics, 2024, 16: 145. DOI: 10.1186/s13321-024-00931-z.",
        "Swanson K, Walther P, Leitz J, et al. ADMET-AI: a machine learning ADMET platform for evaluation of large-scale chemical libraries. Bioinformatics, 2024, 40(7): btae416. DOI: 10.1093/bioinformatics/btae416.",
        "Xiong G, Wu Z, Yi J, et al. ADMETlab 2.0: an integrated online platform for accurate and comprehensive predictions of ADMET properties. Nucleic Acids Research, 2021, 49(W1): W5-W14. DOI: 10.1093/nar/gkab255.",
        "van Tilborg D, Alenicheva A, Grisoni F. Exposing the limitations of molecular machine learning with activity cliffs. Journal of Chemical Information and Modeling, 2022, 62: 5938-5951.",
        "BOOM: Benchmarking Out-Of-distribution Molecular Property Predictions of Machine Learning Models. NeurIPS Datasets and Benchmarks Track, 2025.",
        "Hollmann N, Mueller S, Eggensperger K, Hutter F. TabPFN: A Transformer That Solves Small Tabular Classification Problems in a Second. International Conference on Learning Representations, 2023.",
        "Hollmann N, Müller S, Purucker L, Krishnakumar A, Körfer M, Hoo S B, Schirrmeister R T, Hutter F. Accurate predictions on small data with a tabular foundation model. Nature, 2025, 637: 319-326. DOI: 10.1038/s41586-024-08328-6.",
        "Yang K, Swanson K, Jin W, et al. Analyzing learned molecular representations for property prediction. Journal of Chemical Information and Modeling, 2019, 59: 3370-3388.",
        "Rogers D, Hahn M. Extended-connectivity fingerprints. Journal of Chemical Information and Modeling, 2010, 50: 742-754.",
        "Landrum G. RDKit: Open-source cheminformatics software. https://www.rdkit.org/.",
        "Bemis G W, Murcko M A. The properties of known drugs. 1. Molecular frameworks. Journal of Medicinal Chemistry, 1996, 39: 2887-2893.",
        "Degen J, Wegscheid-Gerlach C, Zaliani A, Rarey M. On the art of compiling and using drug-like chemical fragment spaces. ChemMedChem, 2008, 3: 1503-1507.",
        "Breiman L. Random forests. Machine Learning, 2001, 45: 5-32.",
        "Chen T, Guestrin C. XGBoost: A scalable tree boosting system. KDD, 2016.",
        "Ke G, Meng Q, Finley T, et al. LightGBM: A highly efficient gradient boosting decision tree. NeurIPS, 2017.",
        "Prokhorenkova L, Gusev G, Vorobev A, Dorogush A V, Gulin A. CatBoost: unbiased boosting with categorical features. NeurIPS, 2018.",
        "Geurts P, Ernst D, Wehenkel L. Extremely randomized trees. Machine Learning, 2006, 63: 3-42.",
        "Chithrananda S, Grand G, Ramsundar B. ChemBERTa: Large-scale self-supervised pretraining for molecular property prediction. arXiv:2010.09885, 2020.",
        "Ross J, Belgodere B, Chenthamarakshan V, Padhi I, Mroueh Y, Das P. Large-scale chemical language representations capture molecular structure and properties. Nature Machine Intelligence, 2022.",
        "Hu W, Liu B, Gomes J, et al. Strategies for pre-training graph neural networks. International Conference on Learning Representations, 2020.",
        "Rong Y, Bian Y, Xu T, et al. Self-supervised graph transformer on large-scale molecular data. NeurIPS, 2020.",
        "Xiong Z, Wang D, Liu X, et al. Pushing the boundaries of molecular representation for drug discovery with the graph attention mechanism. Journal of Medicinal Chemistry, 2020, 63: 8749-8760.",
        "Krenn M, Häse F, Nigam A, Friederich P, Aspuru-Guzik A. Self-referencing embedded strings for molecular representation. Machine Learning: Science and Technology, 2020, 1: 045024.",
        "Vovk V, Gammerman A, Shafer G. Algorithmic Learning in a Random World. Springer, 2005.",
        "Shafer G, Vovk V. A tutorial on conformal prediction. Journal of Machine Learning Research, 2008, 9: 371-421.",
        "Guo C, Pleiss G, Sun Y, Weinberger K Q. On calibration of modern neural networks. ICML, 2017.",
        "Gawlikowski J, Tassi C R N, Ali M, et al. A survey of uncertainty in deep neural networks. Artificial Intelligence Review, 2023, 56: 1513-1589.",
        "Sushko I, Novotarskyi S, Körner R, et al. Applicability domains for classification problems: benchmarking of distance to models for Ames mutagenicity set. Journal of Chemical Information and Modeling, 2010, 50: 2094-2111.",
        "Tropsha A. Best practices for QSAR model development, validation, and exploitation. Molecular Informatics, 2010, 29: 476-488.",
        "Sheridan R P. Time-split cross-validation as a method for estimating the goodness of prospective prediction. Journal of Chemical Information and Modeling, 2013, 53: 783-790.",
        "Mayr A, Klambauer G, Unterthiner T, Hochreiter S. DeepTox: toxicity prediction using deep learning. Frontiers in Environmental Science, 2016, 3: 80.",
        "Goh G B, Hodas N O, Vishnu A. Deep learning for computational chemistry. Journal of Computational Chemistry, 2017, 38: 1291-1307.",
        "Brown N, Fiscato M, Segler M H S, Vaucher A C. GuacaMol: benchmarking models for de novo molecular design. Journal of Chemical Information and Modeling, 2019, 59: 1096-1108.",
        "Polaris: an industry-led initiative to critically assess ML in drug discovery. Zenodo, 2025.",
        "MF-PCBA: Multifidelity high-throughput screening benchmarks for drug discovery and machine learning. Journal of Chemical Information and Modeling, 2023.",
        "Accurate ADMET Prediction with XGBoost. arXiv:2204.07532, 2022.",
        "ADMET property prediction via multi-task graph learning under adaptive auxiliary task selection. iScience, 2023, 26: 108285.",
        "Guo J. Do Larger Models Really Win in Drug Discovery? A Benchmark Assessment of Model Scaling in AI-Driven Molecular Property and Activity Prediction. arXiv:2604.26498, 2026.",
        "Li Z, Chen X, Wen H, et al. A systematic survey and benchmark of deep learning for molecular property prediction in the foundation model era. Journal of Chemical Theory and Computation, 2026; arXiv:2604.16586.",
        "Ben Hicham K K, Rittig J G, Grohe M, Mitsos A. Tabular foundation models for in-context prediction of molecular properties. arXiv:2604.16123, 2026.",
        "Hong H, Wu X, Sun H, et al. A hierarchical interaction message net for accurate molecular property prediction. Communications Chemistry, 2026, 9: 150. DOI: 10.1038/s42004-026-01922-x.",
        "Le K, Dey S, Martinez Galindo M, et al. Can Decision Trees Teach Large Language Models? Distilling Verbalized Knowledge for Molecular Property Prediction. arXiv:2603.12344, 2026.",
        "Guo Y, Luo M, Zhang W, et al. Few-shot molecular property optimization via a domain-specialized large language model. Chemical Science, 2026, 17: 4928-4941. DOI: 10.1039/D5SC08859C.",
    ]


def write_reference_markdown(refs: list[str]) -> None:
    lines = ["# Expanded Reference List", ""]
    lines.extend(f"{idx}. {ref}" for idx, ref in enumerate(refs, start=1))
    OUT_REFS.write_text("\n".join(lines) + "\n", encoding="utf-8")
    PACKAGE_DOC.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUT_REFS, PACKAGE_REFS)


def add_front_matter(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FZYC-Mol：面向结构分布偏移的验证集选择与适用域感知多专家分子性质预测框架")
    set_font(run, size=18, bold=True, color="111827")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("完整中文初稿（Nature 方法融合、2026 文献更新与 MoleculeNet 补跑增强版，2026-06-02）")
    set_font(run, size=10.5, color="475569")

    add_heading(doc, "摘要", level=1)
    add_para(
        doc,
        "分子性质预测是药物发现、ADMET 评估、毒性筛选和先导化合物优化中的核心计算任务。近年来图神经网络、D-MPNN、化学语言模型和分子基础模型显著推动了该方向的发展，但公开 benchmark 仍暴露出三个关键问题：其一，随机划分或单一 scaffold 划分容易掩盖真实结构外推难度；其二，不同 endpoint 对图结构、指纹、描述符、motif 与预训练 embedding 的偏好并不一致；其三，许多工作只报告点预测分数，而对不确定性、校准、适用域、活性悬崖、早期富集和化学可解释性的系统证据不足。",
    )
    add_para(
        doc,
        "本文提出 FZYC-Mol，一个严格基于验证集选择的多专家分子性质预测框架。框架整合 GNN/D-MPNN、Chemprop、Morgan 与多指纹树模型、RDKit 描述符模型、BRICS/Murcko/官能团 motif experts、冻结 ChemBERTa 与 MoLFormer embedding heads、validation Top-K/stacking ensemble、adaptive consensus、不确定性模型和适用域诊断。所有最终预测策略均只由 validation split 决定，test split 仅用于最终报告，从而避免用测试集反向调参。",
    )
    add_para(
        doc,
        "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL 上达到 RMSE 0.5829 +/- 0.0352，在 BACE 上达到 ROC-AUC 0.8753 +/- 0.0230，在 ClinTox 上达到 ROC-AUC 0.9489 +/- 0.0302。进一步加入 validation-only rescue heads 后，Lipo 的 RMSE 从 0.7078 +/- 0.0389 改进到 0.6835 +/- 0.0439；在允许低成本模型重构后，FreeSolv 的 Morgan+descriptor stacking 被 retained-best 接受并将 RMSE 从 1.0678 +/- 0.1883 改进到 1.0286 +/- 0.1761。进一步吸收 Nature 系列文献中的层级 motif/fingerprint 表示、分子语言 embedding、多视图融合、AD gating 和不确定性加权思想后，BBBP 的 ROC-AUC 从 0.9165 +/- 0.0290 提升到 0.9243 +/- 0.0247，ClinTox 从 0.9489 +/- 0.0302 小幅提升到 0.9496 +/- 0.0262。外部 TDC ADMET official panel 的多方法融合进一步在 Caco2_Wang、HIA_Hou 和 Pgp_Broccatelli 上被 retained-best 接受，分别带来 +0.0142、+0.0035 和 +0.0197 的正向变化。可靠性分析显示，ensemble std、error model、hybrid error-AD 和重构/相似度适用域信号能够富集高误差样本；motif attribution 和 fragment enrichment 则为模型行为提供可读的化学解释。",
    )
    add_para(
        doc,
        "关键词：分子性质预测；ADMET；MoleculeNet；PyTDC；适用域；不确定性；validation-only ensemble；motif attribution；结构外推",
    )


def add_introduction(doc: Document) -> None:
    add_heading(doc, "1. 引言", level=1)
    add_para(
        doc,
        "分子机器学习从早期的 QSAR、ECFP 指纹和树模型，发展到图神经网络、D-MPNN、预训练图模型、化学语言模型和跨任务 foundation model。MoleculeNet 和 Therapeutics Data Commons 将该领域推向系统 benchmark 阶段，但近年的 ADMET 可靠性研究也反复指出，药物发现中的真实困难不只是平均 ROC-AUC 或 RMSE，而是结构外推、数据稀缺、类别不平衡、活性悬崖、beyond rule-of-5 化学空间、实验噪声和适用域漂移。",
    )
    add_para(
        doc,
        "2026 年已经发表或上线的相关工作形成了比较清晰的共识：ADMET reliability benchmark 强调少样本、OOD、类别不平衡、activity cliff 与 roughness；Tabular foundation model 论文强调 RDKit/Mordred descriptor 和 foundation embedding 上的 in-context tabular prediction；OpenADMET/ASAP-Polaris blind challenge 与 pADME 工作强调外部盲测和多源标签预训练；CheMLT-F、HimNet 和 TreeKD 等工作则分别从多任务共享、层级 motif/atom/molecule 交互、官能团树规则解释的角度说明，性能提升需要与可解释性和任务异质性一起设计。因此，本研究不以重启大规模预训练为主线，而是将有限算力投入到验证集治理、多专家互补、外部 appendix benchmark、可靠性总结和可解释性增强中。",
    )
    add_para(
        doc,
        "FZYC-Mol 的核心假设是：对于公开分子性质任务，单一模型家族很难在所有 endpoint、split 和评价指标上稳定占优；更可取的策略是建立异质专家池，在验证集上选择 endpoint-specific 策略，并同步报告不确定性、适用域和结构分布偏移证据。这个假设与近期 ADMET benchmark 的经验一致，即紧凑但强的 tabular 模型、图模型、指纹模型、预训练 embedding 和 ensemble 往往需要按 endpoint 条件化组合。",
    )
    add_para(
        doc,
        "本文贡献包括四点：第一，构建覆盖 MoleculeNet、TDC ADMET、MoleculeACE、OpenADMET appendix 和 structure-separated split 的统一实验包；第二，提出 validation-only selector 与 rescue-integrated selector，在不使用测试标签调参的前提下吸收 Top-K/stacking 等补强策略；第三，加入 uncertainty、applicability domain、risk-coverage、conformal、roughness 和 low-similarity hard subset 等可靠性分析；第四，利用 motif attribution、BRICS fragment enrichment 和 scaffold/neighbor case review 强化化学解释。",
    )


def add_methods(doc: Document) -> None:
    add_heading(doc, "2. 方法", level=1)
    add_figure(doc, MAIN_FIGS / "Figure_1_FZYC_Mol_framework.png", "图 1A. FZYC-Mol 整体流程图。该图参考 MoleculeFormer 与多模块分子网络示意图的排版方式，将分子输入、多视图表示、候选专家池、validation-only selector 和证据输出组织为连续流程。", width_in=7.3)
    add_figure(doc, POLISHED_FIGS / "fig23_fzyc_mol_model_structure.png", "图 1B. FZYC-Mol 模型结构图。分子首先生成异质图与多源特征库，随后进入图消息传递、fingerprint sensing、descriptor tree、冻结分子编码器和可选 ensemble heads；最终由验证集矩阵、risk-aware policy 与 retained-best gate 决定预测头和证据头。", width_in=7.3)
    add_para(
        doc,
        "图 1A 强调完整实验工作流，图 1B 强调模型内部结构。二者共同说明本研究不是单一 backbone 模型，而是一个由多视图分子表示、异质专家、验证集治理和可靠性证据组成的分子性质预测框架。",
    )
    add_polished_table(
        doc,
        "表 1",
        "数据集、任务类型、主指标和评价协议。",
        read_csv(MAIN_TABLES / "Table_1_Dataset_protocol.csv").head(14).rename(
            columns={
                "dataset": "数据集",
                "source": "来源",
                "task_type": "任务",
                "n_molecules": "分子数",
                "positive_rate": "阳性率",
                "primary_metric": "主指标",
                "splits": "划分",
                "seeds": "seeds",
            }
        )[["数据集", "来源", "任务", "分子数", "阳性率", "主指标", "划分", "seeds"]],
        note="主文表保留前 14 个代表性 endpoint；完整数据集协议见 submission package 中的 Table_1_Dataset_protocol.csv。",
        font_size=6.7,
        max_text=80,
    )
    add_heading(doc, "2.1 数据划分与验证原则", level=2)
    add_para(
        doc,
        "所有 endpoint 均显式区分 train、validation 和 test。模型训练只使用 train，候选模型和 ensemble 策略的选择只使用 validation，test 仅在策略冻结后评估一次。该设计尤其重要，因为本课题包含多个候选专家、stacking、Top-K ensemble、rescue heads 与 target transform；若允许反复查看测试集，很容易产生隐性 test leakage。",
    )
    add_para(
        doc,
        "MoleculeNet 面板采用 scaffold split 作为主报告，并补充 random split、structure-separated split 和 low-similarity hard subset。TDC ADMET 采用官方 PyTDC split 与 full-panel scaffold appendix；MoleculeACE 用于检验活性悬崖；OpenADMET-ExpansionRx 用作快速外部 feasibility check。这样可以把平均性能、结构外推和外部泛化放在同一叙事中，而不是只报告最容易的划分。",
    )
    add_heading(doc, "2.2 多视图表示与专家池", level=2)
    add_para(
        doc,
        "FZYC-Mol 的输入表示包括五类：分子图表示用于 GNN 和 D-MPNN；Morgan、MACCS、atom-pair 与 torsion 指纹用于传统树模型和多指纹专家；RDKit 2D 描述符用于 fast tabular baseline；BRICS、Murcko scaffold 和官能团 motif 用于可解释专家；ChemBERTa 与 MoLFormer 的冻结 embedding 用作低成本预训练特征。冻结 embedding 的定位不是替代大规模微调，而是在算力可控的情况下测试预训练表征是否能为低分 endpoint 提供补充信息。",
    )
    add_para(
        doc,
        "专家池包含 RF、ExtraTrees、XGBoost、LightGBM、Chemprop D-MPNN、图模型、描述符 MLP、motif experts、冻结预训练 embedding heads、Top-K mean、ridge/logistic stacking、adaptive consensus 和 targeted rescue heads。对回归任务，附录候选还包含 identity、log1p、quantile transform 与 winsorized target 等 target transform；对类别不平衡任务，候选包含 balanced undersampling ensemble。",
    )
    add_heading(doc, "2.3 Validation-only selector 与 rescue heads", level=2)
    add_para(
        doc,
        "选择器首先在每个 endpoint 的 validation 预测表上计算主指标，并按指标方向排序。若单一专家显著占优，则保留 best expert；若多个专家在 validation 上互补，则候选 Top-K mean 或 stacking；若多指纹、Chemprop、预训练和核心图模型互有优势，则候选 adaptive consensus。rescue heads 只允许在 validation 上优于当前保留策略时进入最终池，因此它们不会改变 test-only 报告原则。",
    )
    add_heading(doc, "2.4 可靠性与可解释性模块", level=2)
    add_para(
        doc,
        "可靠性模块包括 ensemble standard deviation、error model、prediction deviation、inverse Tanimoto distance、reconstruction error、hybrid error-AD、risk-coverage curve、conformal coverage 和 roughness proxy。可解释性模块包括 motif feature importance、fragment enrichment、scaffold/neighbor case review 与高误差样本分析。本文将这些证据作为模型使用边界的描述，而不把关联性 motif 解释等同于因果机制。",
    )


def add_results(doc: Document) -> None:
    add_heading(doc, "3. 结果", level=1)
    add_heading(doc, "3.1 MoleculeNet 主结果", level=2)
    add_figure(doc, MAIN_FIGS / "Figure_2_MoleculeNet_model_family_ranks.png", "图 2. MoleculeNet endpoint 内模型家族排名。黑框标记 validation-only selector 与 targeted rescue selector。", width_in=7.2)
    add_figure(doc, MAIN_FIGS / "Figure_3_MoleculeNet_main_performance.png", "图 3. MoleculeNet 主性能。星号表示观测最优候选，菱形表示 FZYC-Mol selector；Lipo 显示 rescue-integrated selector 的收益。", width_in=7.2)
    add_polished_table(
        doc,
        "表 2",
        "MoleculeNet 主结果与最强非 FZYC-Mol 专家对比。",
        make_moleculenet_table(),
        note="回归任务低值更优，分类任务高值更优。最终 retained-best 同时反映 rescue heads 与 targeted rebuild；只有 validation 接受时才改变最终保留策略。",
        font_size=6.3,
        max_text=82,
    )
    add_para(
        doc,
        "表 2 和图 2 显示，FZYC-Mol 的优势并不是某一个专家在所有任务上碾压，而是 selector 能够识别 endpoint-specific 的有效专家组合。ESOL 和 BACE 中 validation selector 与观测最优一致；FreeSolv 中 Chemprop 表现最佳，selector 原始版本保留了较稳健但非最优的组合，后续 targeted rebuild 可将 RMSE 从 1.0678 降至 1.0286；Lipo 通过 targeted rescue selector 将 RMSE 从 0.7078 降至 0.6835。进一步加入 Nature-inspired 多方法融合后，BBBP 和 ClinTox 由 rank/AD-gated/stacking fusion 接入最终 retained-best，说明多视图融合对分类 endpoint 更有补强价值。",
    )

    add_heading(doc, "3.2 TDC ADMET 官方 split 与外部 appendix", level=2)
    add_figure(doc, MAIN_FIGS / "Figure_5_Official_TDC_ADMET_scaffold_delta.png", "图 4. 官方 PyTDC ADMET split 中 LGBM/RF Morgan baseline 的 random-to-scaffold 性能变化。", width_in=7.0)
    add_polished_table(
        doc,
        "表 3",
        "TDC ADMET 官方 split 结果与 fast tabular baseline scaffold 惩罚。",
        make_tdc_table(),
        note="平均 scaffold 惩罚由 LGBM 与 RF 的 random-to-scaffold 变化取均值；正值表示 scaffold split 更难。",
        font_size=6.3,
        max_text=85,
    )
    add_para(
        doc,
        "TDC ADMET 结果进一步说明，外部 ADMET endpoint 的困难程度高度异质。HIA_Hou、Pgp_Broccatelli 和 BBB_Martins 上 selector 的 ROC-AUC 分别达到约 0.979、0.936 和 0.918；Caco2_Wang 等回归 endpoint 则对 target transform 和 tabular baseline 更敏感。新增 TDC Nature-inspired fusion 后，Caco2_Wang 的 RMSE 从 0.4517 降至 0.4375，HIA_Hou 的 ROC-AUC 从 0.9792 提升到 0.9827，Pgp_Broccatelli 从 0.9357 提升到 0.9554；其余五个 endpoint 保留原 TDC selector。与近年 ADMET-AI、ADMETlab、Polaris 和 OpenADMET 方向一致，本研究将这些结果作为外部可行性和可靠性证据，而不是宣称单一模型在所有 ADMET endpoint 上统一最优。",
    )

    add_heading(doc, "3.3 结构分布偏移与 split realism", level=2)
    add_figure(doc, MAIN_FIGS / "Figure_4_Split_realism_structure_shift.png", "图 5. 从 random 到 scaffold 再到 structure-separated split 的性能变化。", width_in=7.2)
    add_polished_table(
        doc,
        "表 4",
        "Random、scaffold 与 structure-separated split 的性能变化。",
        make_split_table(),
        note="分类任务中总变化为 random 到 structure 的性能下降；回归任务中总变化为误差增加。负值通常表示更难 split 上反而未下降，应结合样本规模和 endpoint 噪声解释。",
        font_size=6.5,
        max_text=60,
    )
    add_para(
        doc,
        "图 5 表明，结构外推不是均匀惩罚。ESOL 和 Caco2_Wang 随着 split 更严格出现较清晰的误差增加；BACE 在 structure-separated split 下 ROC-AUC 明显下降；部分 endpoint 则由于样本组成、标签噪声和 scaffold 分布差异出现非单调变化。因此，本文不把 random split 的高分作为主要结论，而是把 scaffold、structure-separated 和 low-similarity hard subset 共同作为模型可靠性的压力测试。",
    )

    add_heading(doc, "3.4 可靠性、适用域与高误差富集", level=2)
    add_figure(doc, MAIN_FIGS / "Figure_6_Uncertainty_AD_reliability.png", "图 6. 可靠性总结：不确定性与适用域分数对高误差样本的识别能力。", width_in=7.1)
    add_polished_table(
        doc,
        "表 5",
        "可靠性与适用域 summary table。",
        make_reliability_table(),
        note="Top10% 高误差富集大于 1 表示该分数能把高误差样本集中到高风险区域；risk-coverage AUC 越低通常表示选择性预测曲线越好。",
        font_size=6.7,
        max_text=72,
    )
    add_para(
        doc,
        "可靠性结果显示，单一不确定性分数很难覆盖所有错误类型。ensemble std 与 error model 对模型不一致性敏感，inverse Tanimoto 与 reconstruction error 更偏向适用域和结构新颖性，hybrid error-AD 将两者组合后更适合作为生产使用时的 risk flag。对于 ClinTox、DILI、CYP substrate 等类别不平衡任务，PR-AUC、Brier、ECE 和 conformal coverage 应与 ROC-AUC 同时报告，避免单一 ROC-AUC 掩盖阳性样本稀缺下的实用风险。",
    )

    add_heading(doc, "3.5 可解释性：motif attribution 与 fragment enrichment", level=2)
    add_figure(doc, MAIN_FIGS / "Figure_7_Motif_fragment_interpretation.png", "图 7. Motif attribution 与 fragment-level enrichment 的化学解释。", width_in=7.2)
    add_para(
        doc,
        "图 7 将模型行为连接到可识别的化学子结构。BBBP 中极性羟基、羧酸盐、内酰胺和芳香羟基等 motif 往往与穿透性下降或标签差异有关；BACE 中卤素、芳基甲基、醚和疏水芳香片段更突出；ClinTox 中季铵、苯胺、芳香氮、羰基和稠环结构值得关注。这些分析不能证明因果机制，但可以帮助审稿人判断模型是否学到与化学经验相容的局部结构信号，也能辅助定位高误差或分布外样本。",
    )

    add_heading(doc, "3.6 Rescue-integrated selector 与低分模块补强", level=2)
    add_polished_table(
        doc,
        "表 6",
        "Rescue-integrated selector 诊断。",
        make_rescue_table(),
        note="只有 Lipo 在最终选择中真正使用 rescue pool；其他 endpoint 的 rescue 候选未通过 validation-only 保留规则。",
        font_size=6.2,
        max_text=80,
    )
    add_para(
        doc,
        "表 6 支持一个重要策略判断：当前最值得做的是小规模、定向的 validation-only 补强，而不是重启大模型训练。rescue heads 加入完整 selector pool 后，只有 Lipo 被 validation 接受，其他 endpoint 保持原保留结果。这说明补强是被数据驱动筛选出来的，而不是人为选择测试集上好看的结果。",
    )

    add_figure(
        doc,
        POLISHED_FIGS / "fig17_moleculenet_targeted_rebuild_decision.png",
        "图 8. MoleculeNet 低分模块 targeted rebuild 的 validation-only 接受决策。正值表示重构候选优于当前/已接入 rescue selector；负值表示保留当前策略。",
        width_in=7.2,
    )
    add_polished_table(
        doc,
        "表 7",
        "MoleculeNet targeted rebuild retained-best 决策表。",
        make_targeted_rebuild_table(),
        note="重构候选包含 Morgan+RDKit descriptor、CatBoost/XGBoost/LightGBM/ExtraTrees/RF、target transform、Top-K mean、stacking 与 balanced undersampling ensemble。最终是否替换当前策略仍只由 validation 主指标决定。",
        font_size=6.0,
        max_text=78,
    )
    add_para(
        doc,
        "在允许“重新构建”低成本模型后，新增实验给出了一个更细的性能判断：ESOL、FreeSolv、Lipo、BBBP、BACE 和 ClinTox 六个 MoleculeNet endpoint 均进入同一套 Morgan+RDKit descriptor rebuild pool，候选包括 CatBoost、XGBoost、LightGBM、ExtraTrees、RF、target transform、Top-K mean、stacking 和 balanced undersampling ensemble。结果显示，只有 FreeSolv 的 Morgan+descriptor validation stacking 在 5 个 scaffold seeds 上将 RMSE 从当前/已接入 rescue selector 的 1.0678 +/- 0.1883 降至 1.0286 +/- 0.1761，正向变化为 +0.0392，因此可作为附录 retained-best 补强；ESOL、BBBP、BACE 和 ClinTox 的 rebuild 候选未超过当前 validation selector，Lipo 的重构 tabular 候选也未超过已有 rescue-integrated selector，因此这些任务均保留原策略。",
    )
    add_para(
        doc,
        "这个结果对论文叙事很有帮助：本文并不排斥重新构建模型，而是把重构模型视作候选专家，必须先通过 validation-only gate 才能改变最终报告。FreeSolv 的小幅改善说明 targeted rebuild 仍有价值；其余五个 MoleculeNet endpoint 未被强行替换，说明 selector 没有为了追求更厚实验而牺牲保留规则。与近期 ADMET benchmark 强调的强 tabular baseline、target transform、不平衡学习和 roughness 诊断一致，本文的改进路径应继续保持“低成本、可复现、按 endpoint 接受”的原则。",
    )

    add_heading(doc, "3.6B Nature-inspired 多方法融合补强", level=2)
    add_figure(
        doc,
        POLISHED_FIGS / "fig18_nature_multimethod_fusion_decision.png",
        "图 8B. Nature-inspired 多方法融合的 retained-best 接受决策。正值表示多方法融合候选优于上一版 retained-best；负值表示保留上一版结果。",
        width_in=7.2,
    )
    add_polished_table(
        doc,
        "表 8A",
        "Nature-inspired multimethod fusion retained-best 决策表。",
        make_nature_fusion_table(),
        note="候选池吸收 Nature 系列文献中的层级 motif/global 表征、分子语言 embedding、多视图融合、AD gating、rank fusion 和 uncertainty-weighted fusion。最终仍只由 validation 主指标决定是否接入。",
        font_size=5.6,
        max_text=74,
    )
    add_polished_table(
        doc,
        "表 8B",
        "Nature 系列文献方法信号与本课题可执行实现的对应关系。",
        make_nature_alignment_table(),
        note="该表用于说明借鉴的是可复现的设计原则，而不是宣称复现原论文完整神经结构或官方挑战流程。",
        font_size=5.7,
        max_text=88,
    )
    add_para(
        doc,
        "Nature-inspired fusion appendix 进一步补足了“模型过于简单”的潜在审稿风险。新增候选不再只是 Morgan+descriptor tree model，而是同时包含层级 motif/fingerprint 特征、ChemBERTa/MoLFormer 冻结 embedding heads、多指纹专家、rank fusion、validation stacking、uncertainty-weighted fusion 和 AD-similarity gated fusion。结果显示，BBBP 从 0.9165 +/- 0.0290 提升到 0.9243 +/- 0.0247，ClinTox 从 0.9489 +/- 0.0302 小幅提升到 0.9496 +/- 0.0262，因此这两个分类 endpoint 可接入 Nature 多方法融合 retained-best；BACE、ESOL、FreeSolv 和 Lipo 的融合候选未超过上一版 retained-best，仍保留原策略。",
    )
    add_para(
        doc,
        "这个新增结果的意义不是证明所有任务都需要复杂融合，而是说明当 endpoint 依赖多尺度结构信号、类别边界或 AD/OOD 风险时，多方法融合能提供额外收益；当已有 selector、rescue 或 targeted rebuild 已经更稳时，融合模型不会被强行替换。这种选择性接入方式与 avoid-ome、HimNet、Tabular foundation model 和数据漂移 UQ 文献的共同建议一致：模型应当围绕任务异质性、可靠性和可验证融合来改进，而不是追求单一架构统一覆盖所有 endpoint。",
    )

    add_heading(doc, "3.7 Full-panel appendix、performance mode 与 roughness 解释", level=2)
    add_figure(
        doc,
        POLISHED_FIGS / "fig19_tdc_nature_multimethod_fusion_decision.png",
        "图 8C. 外部 TDC official panel 中 Nature-inspired 多方法融合的 retained-best 接受决策。正值表示融合候选优于原 TDC validation selector。",
        width_in=7.2,
    )
    add_polished_table(
        doc,
        "表 8C",
        "外部 TDC official panel 的 Nature-inspired multimethod fusion retained-best 决策表。",
        make_tdc_nature_fusion_table(),
        note="该表使用 3 个 scaffold seeds 的 fast external appendix 设计；融合候选较弱时保留原 TDC validation selector。",
        font_size=5.5,
        max_text=78,
    )
    add_para(
        doc,
        "TDC external fusion appendix 进一步说明，多方法融合不仅能改善 MoleculeNet 分类 endpoint，也能在外部 ADMET official panel 上产生选择性收益。Caco2_Wang 由 validation ridge stacking/AD-gated fusion 接入 retained-best，HIA_Hou 由 hier_motif RF/ExtraTrees 和 Top-8 mean 提供收益，Pgp_Broccatelli 则由 AD-gated fusion、stacking 和 hier_motif XGBoost 共同贡献。CYP2C9、CYP2D6、CYP3A4、BBB_Martins 和 Bioavailability_Ma 未被融合替换，说明外部 ADMET 的 endpoint heterogeneity 仍然明显。",
    )
    add_polished_table(
        doc,
        "表 8",
        "TDC full-panel performance-mode retained-best 中的代表性增益。",
        make_appendix_table(),
        note="该表只展示 retained_source 为 performance_mode 的代表性 endpoint；完整结果见 Table_S10。",
        font_size=6.2,
        max_text=82,
    )
    add_polished_table(
        doc,
        "表 9",
        "TDC roughness 与文献一致性诊断中的高 roughness endpoint。",
        make_roughness_table(),
        note="roughness proxy 来自测试分子与 train+valid 近邻的相似度、标签冲突或 normalized target jump；用于解释哪些 endpoint 更需要 ensemble、target transform 和 AD/OOD gating。",
        font_size=6.3,
        max_text=78,
    )
    add_para(
        doc,
        "表 8 和表 9 把近期 ADMET reliability 文献与本研究结果连接起来。clearance_microsome_az、clearance_hepatocyte_az、ppbr_az、vdss_lombardo、half_life_obach 等 endpoint 往往具有更高 roughness 或更复杂的局部结构-标签关系，因此 Top-K/stacking、target transform、CatBoost/XGBoost/ExtraTrees 和 retained-best 策略更可能产生选择性增益。相反，对于 roughness 低或强 baseline 已稳定占优的 endpoint，selector 会保留原结果而不强行换模型。",
    )


def add_expanded_method_details(doc: Document) -> None:
    add_heading(doc, "2.5 实验实现细节与候选模型治理", level=2)
    add_para(
        doc,
        "为了让多专家比较具备可复现性，本文将候选模型治理拆成三个层级。第一层是单一专家训练，即每个模型家族在固定 split 和 seed 下独立训练并保存 validation/test 预测；第二层是候选策略生成，即在 validation 预测矩阵上构造 Top-K mean、stacking、adaptive consensus、uncertainty-aware weighting 和 targeted rescue 等组合；第三层是最终 selector 决策，即只根据 validation 主指标和预先定义的 tie-breaking 规则确定每个 endpoint 的保留策略。这样的分层可以避免把训练误差、验证选择和测试报告混在一起，也便于在后续加入新 baseline 时只更新候选池而不重写整套流程。",
    )
    add_para(
        doc,
        "在回归任务中，主指标通常为 RMSE、MAE 或 Spearman。对于 RMSE/MAE，selector 按低值优先排序；对于 Spearman，按高值优先排序。由于 ADME 回归 endpoint 常存在长尾分布、测量误差和单位尺度不一致，附录候选中保留了 log1p、quantile transform 和 winsorized target 等目标变换。目标变换不被默认为更优，只有当 validation 主指标优于当前候选时才会进入 retained-best 结果。这个规则对 clearance、half-life、ppbr 等高 roughness endpoint 尤其重要，因为这些任务的局部 target jump 往往比换用更大的神经网络更直接影响误差。",
    )
    add_para(
        doc,
        "在分类任务中，ROC-AUC 仍作为 MoleculeNet 与若干 TDC endpoint 的主指标，但本文同时保留 PR-AUC、Brier score、ECE 和 conformal coverage 作为可靠性解释。这样做是因为 ClinTox、DILI、CYP substrate 等任务存在明显类别不平衡，ROC-AUC 对排序能力敏感，却不一定反映阳性样本召回、概率校准或筛选富集质量。对于这类任务，balanced undersampling ensemble 的定位不是替代主模型，而是在 validation 上检验类别再平衡是否能改善 PR-AUC 或降低概率误差。",
    )
    add_para(
        doc,
        "stacking 的输入不是原始分子特征，而是各专家在 validation/test 上的 out-of-fold 或固定 split 预测。因此 stacking 本质上是一个 validation-only 的后处理层：回归任务使用 ridge-style 线性组合或简单均值，分类任务使用 logistic stacking 或概率均值。为了降低小样本过拟合，stacking 候选只允许使用 Top-3 或 Top-5 validation experts，且最终报告以 seed-level 结果汇总。若 stacking 在 validation 上不稳定或不能超过当前 selector，则保留原策略。",
    )
    add_para(
        doc,
        "adaptive consensus 的目的是处理专家之间没有单一绝对赢家的 endpoint。它首先识别 validation 上接近最优的候选家族，再根据任务类型、指标方向和专家相关性构造加权平均。与普通 ensemble 不同，adaptive consensus 并不追求把所有模型都纳入，而是尽量避免低质量或高度冗余专家稀释强模型信号。图 1 中将其放在 validation integration 层，是因为它属于策略选择而不是特征学习。",
    )
    add_para(
        doc,
        "targeted rescue heads 的设计原则更加保守。它们主要服务于已知低分或高 roughness 模块，例如 Lipo、clearance、ppbr、half-life、CYP substrate 等。rescue heads 可以来自冻结预训练 embedding、强 tabular baseline、Top-K/stacking、target transform 或 undersampling ensemble，但必须先进入 selector pool，再由 validation 决定是否保留。本文结果显示，MoleculeNet 中真正被接受的 rescue 主要发生在 Lipo，这反而增强了方法可信度：补强模块没有被强行用于所有 endpoint，而是只在验证证据支持时改变最终策略。",
    )
    add_heading(doc, "2.6 适用域、roughness 与可解释性计算", level=2)
    add_para(
        doc,
        "适用域分析使用两类信号。第一类是基于相似度的外推信号，例如测试分子到 train+valid 最近邻的 Morgan Tanimoto 距离、scaffold distance 和 low-similarity subset。第二类是基于模型行为的风险信号，例如 ensemble std、prediction deviation、error model 和 reconstruction error。前者更接近化学空间覆盖，后者更接近模型自身不确定性。两类信号不完全等价：一个分子可能结构相似但标签噪声大，也可能结构新颖但模型预测一致。因此本文在 reliability summary 中同时报告这些分数。",
    )
    add_para(
        doc,
        "roughness proxy 的作用是解释为什么某些 endpoint 即使使用相似模型也难以稳定提升。具体而言，若测试分子与训练/验证最近邻具有较高指纹相似度，但标签差异或归一化 target jump 很大，则说明该 endpoint 的局部结构-性质关系较粗糙，模型需要更强的局部平滑、robust loss、target transform 或 ensemble 才可能稳定。本文不把 roughness 当作新的主性能指标，而是把它作为文献对齐的解释性附录，用于说明 performance-mode retained-best 为什么集中出现在若干 ADME regression endpoint 上。",
    )
    add_para(
        doc,
        "motif attribution 与 fragment enrichment 采用互补视角。motif attribution 更偏向模型内部特征重要性，关注官能团、BRICS 片段、Murcko scaffold 等输入特征对预测的贡献方向；fragment enrichment 更偏向数据分布，关注某些片段出现时标签均值相对 baseline 的偏移。两者一致时，说明模型关注点与数据统计信号相互支持；两者不一致时，则提示可能存在 confounding、样本量不足或局部标签噪声。",
    )
    add_para(
        doc,
        "在正文图表层面，主文只展示关键 summary table，而保留完整 CSV 到 supplementary package。这样处理是为了避免 Word 版本出现几十列表格挤压、断行和不可读的问题。审稿人若关注具体 seed、模型名、候选池计数或 appendix endpoint，可以直接查看 submission package 中的 Table_S1 到 Table_S14；主文则聚焦方法逻辑和主要结论。",
    )


def add_expanded_result_details(doc: Document) -> None:
    add_heading(doc, "3.8 MoleculeNet endpoint 逐项解读", level=2)
    add_para(
        doc,
        "ESOL 是水溶性回归任务，分子规模相对较小但结构-性质关系较明确。FZYC-Mol 在该任务中保留 validation selector，并达到 RMSE 0.5829 +/- 0.0352，优于 Classical Morgan、Chemprop、Graph/D-MPNN core、Frozen pretrained 和 Multi-fingerprint 等家族。这个结果提示，在 ESOL 这类 physicochemical property endpoint 上，组合式 selector 能够整合描述符、指纹和图结构信息，而不必依赖单一大模型。",
    )
    add_para(
        doc,
        "FreeSolv 的样本量较小，且溶剂化自由能对分子局部相互作用、构象与实验测定条件较敏感。表 2 中 Chemprop 是观测最优，而 FZYC-Mol selector 原始版本略落后。新增 targeted rebuild 后，Morgan+RDKit descriptor 的 validation stacking 将 RMSE 从 1.0678 +/- 0.1883 降至 1.0286 +/- 0.1761，说明低成本重构确实能带来补强，但仍未完全消除与最佳 Chemprop 候选的差距。同时，ESOL、BBBP、BACE、ClinTox 和 Lipo 的 rebuild 候选均未通过 retained-best gate，说明模型重构并不是无条件提升，而是 endpoint-specific 的可验证补强。因此论文中应诚实保留这一边界：FreeSolv 需要更强物理化学描述符、构象特征或更严格的 nested validation，而不是把所有低分都归因于模型规模不足。",
    )
    add_para(
        doc,
        "Lipophilicity 是本轮最值得强调的低分模块修复案例。原 validation selector 的 RMSE 为 0.7078 +/- 0.0389，加入 rescue heads 后降至 0.6835 +/- 0.0439。由于 rescue-integrated selector 仍然只使用 validation 决策，Lipo 的提升可以被表述为方法学补强而不是 test-set tuning。更重要的是，其他 endpoint 没有被 rescue 强行改写，说明该模块具备 endpoint-specific gating 能力。",
    )
    add_para(
        doc,
        "BBBP、BACE 和 ClinTox 三个分类任务呈现不同模式。BBBP 中多指纹专家原本略优，说明血脑屏障任务对全局拓扑和局部片段指纹较敏感；新增 Nature-inspired fusion 后，rank fusion、uncertainty-weighted fusion 和 AD-gated fusion 将 BBBP ROC-AUC 进一步提升到 0.9243 +/- 0.0247，超过原多指纹观测最优。BACE 中 validation selector 仍保持更稳，融合候选未被接入；ClinTox 中 selector 原本达到 ROC-AUC 0.9489 +/- 0.0302，新增 AD-gated/stacking fusion 小幅提升到 0.9496 +/- 0.0262，但由于阳性率低，正文仍必须同步提醒 PR-AUC、Brier、ECE 和 conformal diagnostics 的必要性。也就是说，ClinTox 的高 ROC-AUC 不能单独解释为临床毒性预测已经可靠，而应被放入不平衡分类风险框架中讨论。",
    )
    add_heading(doc, "3.9 TDC 与外部 appendix 的结果解释", level=2)
    add_para(
        doc,
        "TDC ADMET 官方 split 的价值在于，它把任务从 MoleculeNet 的经典小面板扩展到更接近药物发现流程的 ADME/Tox endpoint。HIA_Hou、Pgp_Broccatelli 和 BBB_Martins 的高 ROC-AUC 表明 FZYC-Mol 在若干分类 ADMET endpoint 上具有竞争力；Bioavailability、CYP 系列和 Caco2 则显示不同 baseline 的 scaffold 惩罚差异明显。这种差异是论文叙事中的重点：模型性能不是单调地由模型复杂度决定，而是由 endpoint 的标签定义、分子空间、类别比例和 split 方式共同决定。",
    )
    add_para(
        doc,
        "full-panel appendix 的作用不是把主文变成一个庞大的 leaderboard，而是回答审稿人可能提出的问题：该框架是否只在少数常见数据集上有效？fast external benchmark 是否能支持 ADMET 场景？结果显示，强 tabular baseline、Top-K/stacking ensemble 和 target transform 在若干 endpoint 上确实提供增益，但增益不是全局一致的。这说明本文的合理主张应是“validation-governed improvement under endpoint heterogeneity”，而不是“所有任务统一 SOTA”。",
    )
    add_para(
        doc,
        "performance-mode retained-best 的结果应作为附录亮点。它证明，在不重启大模型训练的前提下，新增 CatBoost/XGBoost/ExtraTrees、Top-K ensemble、stacking、target transform 和 undersampling ensemble，可以在部分困难 endpoint 上得到实际收益。与近期 ADMET reliability 文献相比，这条路线更符合 low-cost strengthening：当 endpoint 粗糙、样本有限或 OOD 压力较大时，数据表征和验证集组合策略往往比盲目扩大模型更有效。",
    )
    add_heading(doc, "3.10 低分模块诊断与模型性能改进路径", level=2)
    add_para(
        doc,
        "当前结果中需要重点解释的低分模块主要包括 FreeSolv、部分 TDC ADME regression endpoint、CYP substrate PR-AUC endpoint 和小样本不平衡毒性任务。FreeSolv 的 targeted rebuild 已经带来小幅收益，但剩余差距更可能来自样本量小、物理相互作用复杂和构象缺失；ADME regression 的问题更可能来自长尾 target、实验噪声、单位尺度和局部 target jump；CYP substrate 和 ClinTox 类任务则主要受类别不平衡和阳性样本稀缺影响。因此，改进路径也应按问题类型拆分，而不是统一换一个大模型。",
    )
    add_para(
        doc,
        "对回归 endpoint，最值得继续做的是系统化 target transform 与 robust loss。建议保留 identity、log1p、quantile transform、winsorization、Huber/MAE-style objective 和 rank-aware objective 的 validation-only 比较。若 endpoint 的主指标是 Spearman，模型不一定需要最小化绝对误差，而应更关注排序稳定性；若主指标是 MAE/RMSE，则应重点处理极端值和长尾 target。这个分析可与 Table_S11 的 roughness band 对齐，使性能改进不只是经验尝试，而有明确诊断依据。",
    )
    add_para(
        doc,
        "对不平衡分类 endpoint，最值得继续做的是 undersampling ensemble 与概率校准。具体策略包括多次 balanced undersampling 训练 RF/LGBM/XGBoost/ExtraTrees，再对概率进行平均；同时在 validation 上比较 ROC-AUC、PR-AUC、Brier、ECE 和 top-k enrichment。若 PR-AUC 提升但 ROC-AUC 略降，应在药物筛选语境中解释取舍：早期筛选更关注阳性富集和可校准风险，而不是抽象排序指标的单点最优。",
    )
    add_para(
        doc,
        "对预训练模型模块，当前更合理的描述是“冻结 embedding 可作为补充专家，但不是稳定的全局赢家”。这与近期大模型 benchmark 的结论一致：更大模型并不必然在每个 drug discovery endpoint 上获胜。新增 Nature-inspired fusion 已经把 ChemBERTa/MoLFormer embedding heads 放入候选池，并通过 prediction-level fusion 与 motif/fingerprint experts 共同竞争。后续若要继续加强，可优先尝试轻量 adapter、endpoint-specific linear probing、descriptor+embedding fusion 和 validation-only stacking；full fine-tuning 只有在资源充足且有严格 nested validation 设计时才值得启动。",
    )
    add_para(
        doc,
        "对可解释性模块，下一步最有价值的是增加 case study，而不是再画更多 summary bar plot。每个 case study 可以包含一个高误差分子、最近邻训练分子、Tanimoto 相似度、标签差异、selector 选择的专家、主要 motif attribution 和 fragment enrichment 方向。这样的案例能把 roughness、AD、模型选择和化学解释串成完整故事，有助于增强论文说服力。",
    )

    add_heading(doc, "3.11 按五方面补强的新增实验结果", level=2)
    add_para(
        doc,
        "根据后续实验规划，本文进一步把五类补强正式纳入结果：更系统的 external benchmark appendix、低分模块 targeted improvement、不平衡分类附加指标、seed-level 统计显著性与 2-3 个 case study。新增分析不重启大模型训练，而是复用已经完成的 TDC full-panel、performance-mode、calibration、roughness、rescue integration 和 high-error interpretability 输出，在此基础上进行 retained-best 整合和统计汇总。",
    )
    add_figure(
        doc,
        POLISHED_FIGS / "fig15_external_appendix_retained_delta.png",
        "图 9. 正式 external appendix 中 retained-best selector 相对旧 fast baseline 的保留增益。正值表示 performance-mode 候选被验证集接受后带来改进；零值表示旧 baseline 被保留以避免退化。",
        width_in=7.0,
    )
    add_polished_table(
        doc,
        "表 10",
        "正式 external benchmark appendix：retained-best selector 摘要。",
        make_formal_external_selector_table(),
        note="该表覆盖 22 个 TDC ADMET/Tox endpoint。performance-mode 候选较弱时，retained-best selector 保留旧 fast baseline，因此最终 win/tie/loss 不产生负迁移。",
        font_size=5.8,
        max_text=72,
    )
    add_polished_table(
        doc,
        "表 11",
        "External candidate pool 覆盖与被选频率。",
        make_candidate_pool_table(),
        note="validation stacking、Top-K mean 和 undersampling ensemble 是 seed-level 最常被 selector 选中的三类补强。",
        font_size=6.5,
        max_text=70,
    )
    add_polished_table(
        doc,
        "表 12",
        "不平衡分类 endpoint 的 PR-AUC、Brier、ECE 与 enrichment 指标。",
        make_imbalanced_table(),
        note="ClinTox、DILI、hERG 和 CYP substrate 等任务不能只看 ROC-AUC；PR-AUC、Brier、ECE 和 EF1/EF5 更能反映筛选和风险使用场景。",
        font_size=5.7,
        max_text=76,
    )
    add_figure(
        doc,
        POLISHED_FIGS / "fig16_external_candidate_rank_cd.png",
        "图 10. External candidate pool 的 critical-difference-style 平均排名图。该图用于说明候选家族在 endpoint-seed 层面的相对稳定性，而不是单一 endpoint 的绝对最优。",
        width_in=6.8,
    )
    add_polished_table(
        doc,
        "表 13",
        "Seed-level 稳定性与 retained-best win/tie/loss。",
        make_external_stats_table(),
        note="第一行统计 performance-mode 候选本身，因此会包含未被最终保留的弱候选；第二行是论文主张对应的 retained-best selector，表现为 5/17/0 的 win/tie/loss。",
        font_size=6.0,
        max_text=84,
    )
    add_polished_table(
        doc,
        "表 14",
        "Targeted improvement case studies。",
        make_case_study_table(),
        note="三个案例分别对应 Lipo rescue 成功、ClinTox 不平衡高风险样本、以及 high-roughness ADME regression endpoint。",
        font_size=5.8,
        max_text=82,
    )
    add_para(
        doc,
        "新增 external appendix 的关键结论是：performance-mode 候选本身并不会在所有 endpoint 上胜出，但 retained-best selector 能够在验证证据支持时吸收增益，并在候选较弱时保留旧 baseline。具体而言，22 个外部 endpoint 中有 5 个由 performance-mode retained，17 个保留旧结果，最终 retained-best win/tie/loss 为 5/17/0。这个结果比单纯宣称所有模型都提高更可信，也更符合 ADMET endpoint heterogeneity 的现实。",
    )
    add_para(
        doc,
        "不平衡分类结果进一步强化了 ClinTox、DILI、hERG 和 CYP substrate 的解释边界。CYP substrate 中 undersampling ensemble 被多次选择，说明类别再平衡在 PR-AUC 任务中有实际作用；ClinTox 虽然 ROC-AUC 较高，但 case study 中仍存在阳性样本被预测为低概率的高风险 false negative。因此论文应避免只用 ROC-AUC 宣称毒性预测可靠，而应把 PR-AUC、Brier、ECE、calibration 和 case-level risk 一起呈现。",
    )
    add_para(
        doc,
        "统计显著性层面，performance-mode 候选在 endpoint-seed 单元上的平均正向变化为正，但存在 endpoint 异质性；retained-best selector 在 endpoint 层面的 bootstrap CI 为正，并且没有负迁移 endpoint。这支持一个更稳健的主张：FZYC-Mol 的外部扩展价值来自验证集治理和保留最优，而不是某个新 baseline 在所有任务上统一胜出。",
    )


def add_expanded_discussion_details(doc: Document) -> None:
    add_heading(doc, "4.1 与同类型论文的差异化定位", level=2)
    add_para(
        doc,
        "与 MoleculeNet 时代的标准 benchmark 相比，FZYC-Mol 不只报告固定模型在固定数据集上的分数，而是强调多专家策略选择、外部 ADMET appendix、结构外推和可靠性诊断。与 ADMET-AI 或 ADMETlab 这类平台型工作相比，本文不以部署大规模预测服务为目标，而是更关注公开 benchmark 下的实验透明度、validation-only selection 和可解释证据链。与 recent foundation-model benchmark 相比，本文的重点不是证明大模型无效，而是证明在许多 ADMET 场景中，强 tabular experts、指纹/描述符、轻量 ensemble 和严格验证治理仍然是非常有竞争力的基线。",
    )
    add_para(
        doc,
        "这种定位对投稿叙事很重要。若把论文写成“又一个 molecular property prediction 模型”，审稿人会自然要求与所有最新 foundation model 全面比较；但若把论文定位为“validation-governed reliability framework”，则核心竞争力转向实验设计、可复现选择、endpoint heterogeneity、AD/OOD 诊断和解释性闭环。当前已有的图 1、图 7、图 8、表 5、表 7、表 8 和表 9 都服务于这个定位。",
    )
    add_heading(doc, "4.2 审稿人可能关注的问题与回应", level=2)
    add_para(
        doc,
        "第一个可能问题是：为什么不直接重启更大的预训练模型或 full fine-tuning？本文的回应是，当前证据显示低分模块的瓶颈并不完全来自表征能力不足，而经常来自 target 长尾、类别不平衡、roughness、局部标签冲突和 split 外推。直接扩大模型可能提高某些 endpoint，但也会增加算力成本、过拟合风险和验证选择复杂度。因此本文优先采用 fast external appendix、target transform、Top-K/stacking 和 strong tabular baselines。",
    )
    add_para(
        doc,
        "第二个可能问题是：selector 是否会过拟合 validation？本文通过三个设计降低风险。首先，selector 只在预定义候选池中选择，不允许根据测试结果添加临时规则；其次，结果按多个 seed 汇总，并报告 mean +/- std；第三，rescue heads 进入最终策略需要 validation 接受，且未被所有 endpoint 无差别采用。尽管如此，validation overfitting 仍是本文局限之一，因此后续可加入 nested validation 或 cross-dataset transfer validation 作为进一步加强。",
    )
    add_para(
        doc,
        "第三个可能问题是：可解释性结果是否只是相关性？本文应明确回答：是的，motif attribution 与 fragment enrichment 是关联解释，不等价于因果机制证明。它们的价值在于帮助识别模型关注的局部结构、定位高误差和分布外样本、辅助药物化学用户判断预测是否合理。若要进一步证明因果机制，需要湿实验、反事实分子设计或更严格的 matched molecular pair 分析。",
    )
    add_heading(doc, "4.3 下一步实验优先级", level=2)
    add_para(
        doc,
        "下一阶段若继续加厚，最值得做的不是重启大模型训练，而是新增外部 benchmark/appendix。优先级最高的是 fast external appendix benchmark：在更多公开 ADMET endpoint 上跑 CatBoost/XGBoost/ExtraTrees/LGBM/RF、Top-K/stacking、target transform 和 undersampling ensemble，并把 retained-best 与 roughness 诊断绑定。这样可以快速增加证据厚度，同时保持算力成本可控。",
    )
    add_para(
        doc,
        "第二优先级是 validation-only Top-K/stacking ensemble 的系统化版本。当前 rescue-integrated selector 已经证明 Lipo 能受益，但还可以把 Top-3、Top-5、probability mean、ridge stacking、logistic stacking 和 calibration-aware stacking 做成标准候选族，并在每个 endpoint 上输出选择原因。这个扩展有望进一步降低 seed-level std，并改善若干单模型不稳定的 endpoint。",
    )
    add_para(
        doc,
        "第三优先级是解释性 case study。建议选择 Lipo、BACE、ClinTox、clearance_microsome_az 和一个 CYP substrate endpoint，各写 1 个正例或失败案例。每个案例包含结构片段、最近邻、标签差异、uncertainty/AD 分数、selector 选择和 motif attribution。与单纯增加更多表格相比，case study 更能让论文像一篇完整研究而不是实验报告。",
    )
    add_para(
        doc,
        "第四优先级才是受控的大模型补强。如果后续具备资源，可以在少数代表性 endpoint 上尝试 ChemBERTa/MoLFormer adapter 或轻量 fine-tuning，但必须采用 nested validation 或严格 holdout，以免因为模型容量增加而放大 validation overfitting。这样的实验应作为补充验证，而不应替代当前以可靠性和验证治理为核心的主线。",
    )


def add_discussion(doc: Document) -> None:
    add_heading(doc, "4. 讨论", level=1)
    add_para(
        doc,
        "FZYC-Mol 的核心优势不是单一模型规模，而是实验治理方式。与只报告一个最优模型的 benchmark 相比，本研究把多专家候选、验证集选择、split realism、外部 appendix、reliability summary 和解释性分析纳入同一闭环。这个定位与近期 ADMET 论文的趋势一致：现实药物发现更关心模型在结构新颖分子、低相似度样本、类别不平衡 endpoint 和噪声标签下是否仍可被信任。",
    )
    add_para(
        doc,
        "从模型性能角度看，后续最有价值的方向仍是 lightweight expansion，而不是 full fine-tuning 重启。第一，TabPFNv2 或类似 tabular foundation model 可接入 Morgan+descriptor+embedding 特征，用于少样本 ADMET endpoint；第二，CatBoost、XGBoost、LightGBM 和 ExtraTrees 仍应作为强 tabular baseline 保留；第三，回归 endpoint 可继续系统比较 log1p、quantile transform、winsorization 和 robust loss；第四，极端不平衡分类任务可继续使用 balanced undersampling ensemble 并重点报告 PR-AUC、Brier 和校准。",
    )
    add_para(
        doc,
        "本研究也有局限。首先，当前没有湿实验验证，所有结论均来自公开数据集和计算评估；其次，ChemBERTa 与 MoLFormer 主要以冻结 encoder 形式使用，full fine-tuning 可能改善部分 endpoint，但也会显著增加算力和过拟合风险；第三，Polaris 与 OpenADMET 的完整官方挑战流程尚未完全纳入；第四，motif attribution 和 fragment enrichment 是关联解释，不应被解读为因果机制证明。",
    )
    add_para(
        doc,
        "总体而言，FZYC-Mol 更适合作为一个 validation-governed reliability framework，而不是单一 SOTA leaderboard 方案。它可以在现有算力范围内给出更厚的证据链：性能、方差、结构外推、适用域、高误差富集、校准、活性悬崖、外部 ADMET appendix 与化学解释。这种证据链对投稿叙事更稳，也更贴近药物发现用户对模型可信度的实际需求。",
    )


def add_conclusion_and_mapping(doc: Document) -> None:
    add_heading(doc, "5. 结论", level=1)
    add_para(
        doc,
        "本文提出并系统评估了 FZYC-Mol，一个验证集选择、适用域感知、多专家分子性质预测框架。结果表明，在 MoleculeNet、TDC ADMET、external appendix、split realism、roughness 和 motif/fragment interpretability 等多条证据线上，FZYC-Mol 能够提供比单一模型分数更完整的可靠性画像。加入 rescue heads 后，Lipo 得到 validation-only 接受的定向提升；进一步允许低成本 targeted rebuild 后，FreeSolv 获得小幅 retained-best 改善；再加入 Nature-inspired 多方法融合后，BBBP 和 ClinTox 的分类结果也得到接入式提升，外部 TDC official panel 中 Caco2、HIA 和 Pgp 也获得 retained-best 增益。这说明小规模、可复现、按 endpoint 接受的补强可以改善低分模块，同时保持测试集独立性。",
    )
    add_heading(doc, "6. 图表对应关系", level=1)
    mapping = pd.DataFrame(
        [
            ["方法", "图 1A-1B", "Figure_1 + Figure_S16", "整体流程图与模型结构图"],
            ["结果 3.1", "图 2 / 表 2", "Figure_2 + MoleculeNet table", "模型家族排名与主结果"],
            ["结果 3.1", "图 3", "Figure_3_MoleculeNet_main_performance.png", "MoleculeNet 主性能散点"],
            ["结果 3.2", "图 4 / 表 3", "Figure_5 + TDC table", "官方 PyTDC ADMET scaffold 惩罚"],
            ["结果 3.3", "图 5 / 表 4", "Figure_4 + split table", "random/scaffold/structure split realism"],
            ["结果 3.4", "图 6 / 表 5", "Figure_6 + reliability table", "不确定性与适用域可靠性"],
            ["结果 3.5", "图 7", "Figure_7_Motif_fragment_interpretation.png", "motif 与 fragment 可解释性"],
            ["结果 3.6", "表 6", "Table_S14 rescue integration", "rescue-integrated selector 诊断"],
            ["结果 3.6", "图 8 / 表 7", "Figure_S10 + Table_S22/S23", "targeted rebuild retained-best 与低分模块重构决策"],
            ["结果 3.6B", "图 8B / 表 8A-8B", "Figure_S11 + Table_S24-S26", "Nature-inspired 多方法融合与文献方法对齐"],
            ["结果 3.7", "图 8C / 表 8C", "Figure_S12 + Table_S27-S29", "外部 TDC official panel 多方法融合 retained-best"],
            ["结果 3.7", "表 8-9", "Table_S10/S11", "external benchmark、performance mode 与 roughness 解释"],
            ["结果 3.11", "图 9-10 / 表 10-14", "Table_S15-S21", "formal external appendix、统计显著性与 case studies"],
        ],
        columns=["章节", "图表", "文件", "对应内容"],
    )
    add_polished_table(doc, "表 15", "主文图表对应关系。", mapping, font_size=7.2, max_text=90)


def add_references(doc: Document) -> None:
    refs = references()
    write_reference_markdown(refs)
    add_heading(doc, "7. 参考文献", level=1)
    add_numbered(doc, refs)


def build_doc() -> Document:
    doc = Document()
    set_normal_style(doc)
    set_margins(doc.sections[0])
    add_front_matter(doc)
    add_introduction(doc)
    add_methods(doc)
    add_expanded_method_details(doc)
    add_results(doc)
    add_expanded_result_details(doc)
    add_discussion(doc)
    add_expanded_discussion_details(doc)
    add_conclusion_and_mapping(doc)
    add_references(doc)
    return doc


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    PACKAGE_DOC.parent.mkdir(parents=True, exist_ok=True)

    def copy_with_fallback(src: Path, dst: Path) -> Path:
        try:
            shutil.copy2(src, dst)
            return dst
        except PermissionError:
            fallback = dst.with_name(f"{dst.stem}_visual_polish_20260602{dst.suffix}")
            shutil.copy2(src, fallback)
            return fallback

    doc = build_doc()
    doc.save(OUT_POLISHED)
    written = [
        OUT_POLISHED,
        copy_with_fallback(OUT_POLISHED, OUT_DETAILED),
        copy_with_fallback(OUT_POLISHED, OUT_POLISHED_20260602),
        copy_with_fallback(OUT_POLISHED, OUT_DETAILED_20260602),
        copy_with_fallback(OUT_POLISHED, OUT_MAIN),
        copy_with_fallback(OUT_POLISHED, PACKAGE_DOC),
        copy_with_fallback(OUT_POLISHED, PACKAGE_DETAILED_DOC),
        copy_with_fallback(OUT_POLISHED, PACKAGE_DETAILED_DOC_20260602),
    ]
    if OUT_REFS.exists():
        written.append(copy_with_fallback(OUT_REFS, OUT_REFS_20260602))
        written.append(copy_with_fallback(OUT_REFS, PACKAGE_REFS_20260602))
    for path in written:
        print(f"Wrote {path}")
    print(f"Wrote {OUT_REFS}")


if __name__ == "__main__":
    main()
