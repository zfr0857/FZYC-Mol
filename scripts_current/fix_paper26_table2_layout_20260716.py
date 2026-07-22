from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PATH = Path(r"D:\fzyc\output\paper26_split_regime_transfer_revision_20260716\Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript.docx")


def border(cell, edge: str, value: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    element = borders.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        borders.append(element)
    element.set(qn("w:val"), value)
    if value == "single":
        element.set(qn("w:sz"), "8")
        element.set(qn("w:color"), "000000")


def main() -> None:
    doc = Document(PATH)
    table = doc.tables[1]
    widths = [1.18, 1.05, 2.28, 1.18, 1.08]
    values = ["Split-regime audit", "32 Morgan", "3 endpoints; 5 seeds; scaffold vs similarity clusters", "Transfer", "5,760 fits"]
    for cell, value in zip(table.rows[-1].cells, values):
        cell.text = value
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for column, (cell, width) in enumerate(zip(row.cells, widths)):
            cell.width = Inches(width)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(7.5)
                    run.bold = row_index == 0
            border(cell, "top", "nil")
            border(cell, "bottom", "nil")
            if row_index == 0:
                border(cell, "top", "single")
                border(cell, "bottom", "single")
            if row_index == len(table.rows) - 1:
                border(cell, "bottom", "single")
    doc.save(PATH)


if __name__ == "__main__":
    main()
