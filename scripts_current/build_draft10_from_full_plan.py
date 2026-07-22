from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def set_text(paragraph: Paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def block_element(block):
    return block._p if isinstance(block, Paragraph) else block._tbl


def insert_paragraph_after(block, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    block_element(block).addnext(new_p)
    p = Paragraph(new_p, block._parent)
    if style:
        p.style = style
    p.add_run(text)
    return p


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        element = borders.find(qn("w:" + edge_name))
        if element is None:
            element = OxmlElement("w:" + edge_name)
            borders.append(element)
        if edge is None:
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), edge.get("val", "single"))
            element.set(qn("w:sz"), str(edge.get("sz", 8)))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), edge.get("color", "000000"))


def make_three_line(table: Table) -> None:
    line = {"val": "single", "sz": 8, "color": "000000"}
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_borders(cell, top=None, bottom=None, left=None, right=None)
            if ri == 0:
                set_cell_borders(cell, top=line, bottom=line, left=None, right=None)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
            elif ri == len(table.rows) - 1:
                set_cell_borders(cell, top=None, bottom=line, left=None, right=None)


def remove_extra_rows(table: Table, target_rows: int) -> None:
    while len(table.rows) > target_rows:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)


def find_para(doc: Document, predicate) -> Paragraph:
    for p in doc.paragraphs:
        if predicate(p.text.strip()):
            return p
    raise ValueError("paragraph not found")


def update_experiment_matrix(doc: Document) -> None:
    data = [
        ["实验", "证据层级", "核心问题", "数据/模型", "输出指标", "论文位置与实现模块"],
        ["A", "核心", "是否与 2025-2026 年强基线公平同划分比较", "MoleculeNet 6 终点与 TDC ADMET 关键终点；TabPFNv2-RDKit、AutoGluon-RDKit、XGBoost-RDKit/Mordred/MorganCount、Chemprop、KPGT representation、FZYC selector", "validation metric、test metric、rank、selected candidate、regret", "3.5、4.1、4.5；baselines_strong.py"],
        ["B", "核心", "OOD 难度是否形成 Random/Scaffold/Perimeter 梯度", "BBBP、hERG、Mutagenicity、HLM、Oral Bioavailability、Caco2、Half-Life、VDss；Morgan fingerprint max-distance perimeter split", "性能下降曲线、nearest-neighbor Tanimoto、低相似度分层、selector regret、rank stability", "3.1、4.3；data_splits.py"],
        ["C", "核心", "ROC-AUC 是否掩盖阳性召回与校准不足", "ClinTox、Tox21 NR ER、CYP2C9 Substrate、CYP2D6 Inhibition、hERG；class weight、oversampling、downsampling、downsampling ensemble、threshold moving", "ROC-AUC、PR-AUC、Precision、Recall、F1、MCC、Balanced accuracy、Brier、ECE、Recall at fixed Precision", "3.5、4.4；imbalance_panel.py"],
        ["D", "扩展", "普通小分子之外的 bRo5 化学空间是否仍可审计", "CycPept PAMPA、LinPept NonFouling、LinPept CellPen；XGBoost、TabPFNv2、AutoGluon、KPGT representation、FZYC selector", "性能、AD 覆盖率、失败案例、是否被选择器拒绝", "4.2、5.3；bro5_panel.py"],
        ["E", "核心", "活性悬崖是否从案例升级为全量任务证据", "MoleculeACE 30 任务；XGBoost-MorganCount、KPGT representation、FZYC selector、可运行 MOLMCL/KPGT 版本", "R2、RMSE、MAE、cliff-pair RMSE、gap Spearman、cliff recall、top cliff error", "4.6、4.7；moleculeace_cliff_audit.py"],
        ["F", "核心", "roughness 是否能解释选择器失败风险", "MoleculeNet、TDC、MoleculeACE、bRo5；ROGI/MODI/SARI 或可复现代理", "roughness 与 validation-test Spearman、Top1 match、Top3 hit、regret、optimism gap、高误差 AUROC 的相关性", "3.6、4.4；roughness_diagnostics.py"],
        ["G", "核心", "负结果是否透明进入候选接受/拒绝审计", "所有 dataset-seed 候选池；验证集最佳、最终保留、测试集观测最佳、拒绝原因", "candidate_registry.csv、selection_audit.csv、negative_results.csv", "4.5、5.4；selector_audit.py"],
        ["H", "扩展", "可靠性是否超越单一性能分数", "分类校准、Brier、ECE、reliability diagram、split conformal coverage、risk-coverage curve", "覆盖率、平均区间宽度、分类集合大小、risk-coverage AUC", "4.4、4.6；aggregate_results.py"],
    ]
    for table in doc.tables:
        if table.rows and table.cell(0, 0).text.strip() in {"模块", "实验"} and "证据" in table.cell(0, 1).text:
            remove_extra_rows(table, len(data))
            while len(table.rows) < len(data):
                table.add_row()
            for r, row in enumerate(data):
                for c, value in enumerate(row):
                    table.cell(r, c).text = value
            make_three_line(table)
            return
    raise ValueError("experiment matrix table not found")


def main() -> None:
    desktop = Path.home() / "Desktop"
    source = max([p for p in desktop.glob("FZYC-Mol_*.docx") if "-9" in p.name], key=lambda p: p.stat().st_mtime)
    output = desktop / "FZYC-Mol_初稿-10.docx"
    doc = Document(str(source))

    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("分子性质预测是药物发现"):
            set_text(
                p,
                "分子性质预测是药物发现、ADMET 评估和毒性风险筛查中的关键计算环节。随机划分下的平均 ROC-AUC 或 RMSE 难以充分反映模型在新骨架、少样本终点、不平衡标签、bRo5 化学空间、活性悬崖、实验噪声和适用域漂移条件下的可靠性。受近期真实挑战 ADMET 基准研究启发 [3]，本文提出 FZYC-Mol，一种由验证集治理的适用域感知分子性质预测框架。该框架不依赖单一大型主干模型，而是把多视图表示、强基线专家、Random/Scaffold/Perimeter 划分、类别不平衡策略、bRo5 外部压力、MoleculeACE 活性悬崖、粗糙度诊断、校准和适用域证据纳入冻结候选池；最终策略仅由验证集决定，测试集只在策略冻结后用于一次性评估。",
            )
        elif text.startswith("本文贡献可概括为五点"):
            set_text(
                p,
                "本文贡献可概括为五点。第一，围绕 MoleculeNet、TDC ADMET、bRo5 肽类压力任务、MoleculeACE 30 任务和 Random/Scaffold/Perimeter 划分构建统一评测流程，并明确训练集、验证集和测试集的使用边界。第二，将 TabPFNv2、AutoGluon、KPGT representation、XGBoost-RDKit/Mordred/MorganCount、Chemprop、树模型、冻结表征、采样策略、目标变换和融合策略统一登记为预注册候选，而不是在测试集后追加模型。第三，提出可审计的 accept/reject/retain 治理机制，使候选策略必须先通过验证集证据才能进入主结果；未通过候选作为负结果保留。第四，报告 validation-test rank correlation、Top1 match、Top3 hit、regret、optimism gap、不平衡分类、保形预测、risk-coverage、MoleculeACE cliff-pair 指标和 ROGI/MODI/SARI 粗糙度诊断，避免只依据单一 ROC-AUC 或 RMSE 下结论。第五，通过样本级失败案例、基序/片段解释和稳定命名的 CSV 输出，将模型性能、可靠性边界、负结果和可复现实验记录连接起来。",
            )
        elif text.startswith("为使补充实验与 Zhao et al. 2026"):
            set_text(
                p,
                "为使补充实验与 Zhao et al. 2026 的真实挑战框架 [3] 形成直接对照，本文把新增实验组织为“挑战维度—候选池—验证门控—风险输出”的矩阵，而不是把每个模型作为孤立追加结果。实验 A-C 和 E-G 构成支撑主线结论的核心证据，实验 D 与 H 用于加强真实药物发现场景和可靠性输出。所有模块均遵守测试集冻结原则：候选是否进入最终保留结果，只由验证集或预先定义的审计规则决定。",
            )
        elif text.startswith("表 1A."):
            set_text(p, "表 1A. 新增真实挑战实验矩阵、输出指标与正文落点。")
        elif text.startswith("为避免粗糙度分析停留在附属解释层面"):
            set_text(
                p,
                "为避免粗糙度分析停留在附属解释层面，本文将 ROGI、MODI 和 SARI 明确作为选择器风险诊断，而不是直接作为模型加权项。ROGI 用于刻画性质差异与化学距离之间的不平滑程度，MODI 用于估计局部邻域标签混杂，SARI 用于度量相似分子对中的性质突变。本文主要检验这些指标与 validation-test Spearman、Top1 match、Top3 hit、selector regret、optimism gap、高误差 AUROC 和 MoleculeACE cliff-pair error 的关系；若相关性不稳定，则作为负结果报告，而不写入最终门控规则。",
            )
        elif text.startswith("现代强基线同划分对照被作为"):
            set_text(
                p,
                "最新强基线同划分对照被作为 4.1 的独立证据块。该面板不只比较 CatBoost、XGBoost、LightGBM、ExtraTrees、RF 和 Chemprop，也把 TabPFNv2-RDKit、AutoGluon-RDKit、XGBoost-RDKit/Mordred/MorganCount 以及可获得的 KPGT representation 接入同一候选登记。每个 dataset-seed 同时输出 validation metric、test metric、rank、selected candidate 和 regret。其目标不是证明 FZYC-Mol 在每个终点击败所有强基线，而是检验当现代强基线进入同一验证集候选池后，最终保留策略是否仍能透明地接受、拒绝或保留旧基线。",
            )
        elif text.startswith("为补齐 Zhao et al."):
            set_text(
                p,
                "bRo5 外部压力测试用于补齐真实复杂化学空间。CycPept PAMPA、LinPept NonFouling 和 LinPept CellPen 分别代表环肽通透性、线性肽非特异性吸附和细胞穿透性；它们的分子量、构象柔性和低相似度外推难度均高于常规小分子面板。XGBoost-RDKit、XGBoost-MorganCount、TabPFNv2-RDKit、AutoGluon-RDKit、KPGT representation 和 FZYC selector 在同一 split 下比较；若某些候选未被验证集接受，结果仍作为边界场景审计报告，而不是被删去以维持性能叙事。",
            )
        elif text.startswith("Perimeter/max-min distance split"):
            set_text(
                p,
                "Random/Scaffold/Perimeter 三类划分被用于形成明确的 OOD 难度梯度。Perimeter split 以 Morgan fingerprint 距离选择最远化学空间作为测试集，使 random-to-scaffold-to-perimeter 的性能下降可以与对照论文 [3] 直接比较。该分析同时报告 nearest-neighbor Tanimoto、低相似度分层、selector regret 和 rank stability。若某个终点在 Perimeter split 下出现非单调变化，正文优先解释样本组成、标签噪声和骨架分布，而不是简单归因为模型失败或模型胜出。",
            )
        elif text.startswith("类别不平衡专项被放入可靠性小节"):
            set_text(
                p,
                "类别不平衡专项被放入可靠性小节，而不是作为单独的性能表。ClinTox、Tox21 NR ER、CYP2C9 Substrate、CYP2D6 Inhibition 和 hERG 同时报告 ROC-AUC、PR-AUC、Precision、Recall、F1、MCC、Balanced accuracy、Brier、ECE 和 Recall at fixed Precision，并比较 class weight、oversampling、downsampling、downsampling ensemble 与 threshold moving。该设置更接近筛选决策场景：模型必须说明在高精度阈值下还能保留多少阳性检出能力，以及概率输出是否可用于风险分层。",
            )
        elif text.startswith("计算成本与可复现性审计被纳入候选登记"):
            set_text(
                p,
                "候选接受/拒绝审计被纳入候选登记。每个 dataset-seed 记录候选池数量、验证集最佳、最终保留、测试集观测最佳、是否接入、拒绝原因、训练时间、推理时间、硬件需求和失败率。该表与性能表分开呈现，用于回答 FZYC-Mol 是否只是堆叠算力的问题；若某个高性能候选的时间成本或失败率显著高于轻量候选，则其进入最终保留结果需要更强的验证集证据。",
            )
        elif text.startswith("为形成完整的验证闭环"):
            set_text(
                p,
                "为形成完整的验证闭环，本文将关键补充证据纳入主文逻辑：先审计验证集选择是否存在排名偏差，再用配对统计和系统消融界定各模块贡献，随后在 Random/Scaffold/Perimeter 梯度、低相似度、bRo5 肽类压力、MoleculeACE 30 任务活性悬崖、不平衡分类、ROGI/MODI/SARI 粗糙度诊断、候选接受/拒绝审计、校准/保形预测和失败案例中检查适用边界。这些分析不改变测试集一次性报告原则，而是将“为什么接受、保留或拒绝某个候选”写成可追溯的证据链。",
            )
        elif text.startswith("MoleculeACE 配对结果与低相似度分析互相补充"):
            set_text(
                p,
                "MoleculeACE 配对结果与低相似度分析互相补充，并在本研究中从代表性案例扩展为 30 任务全量审计。主指标包括 R2、RMSE、MAE、cliff-pair RMSE、预测差异与真实差异的 gap Spearman、cliff recall 和 top cliff error；比较对象包括 XGBoost-MorganCount、KPGT representation、可运行的 MOLMCL/KPGT 版本、FZYC selector 以及 KPGT+XGBoost ensemble。已有代表性配对结果显示，活性悬崖目标候选在 51 个 seed 配对中的总 RMSE 平均正向变化为 0.0069，悬崖子集 RMSE 平均正向变化为 0.0056，但标准差较大；预测差异与真实差异的平均 Spearman 约为 0.252，部分任务接近零或为负。因此，本文将该结果解释为悬崖风险识别和候选治理的补充证据，而不是对活性悬崖预测已经解决的证明。",
            )
        elif text.startswith("不平衡分类与保形预测的补充结果改善了"):
            set_text(
                p,
                "校准、保形预测与风险覆盖补充结果用于支撑可靠性而非单纯性能。分类任务报告 reliability diagram、Brier、ECE、risk-coverage curve 和 split conformal coverage；回归任务报告覆盖率、平均区间宽度、预测区间随目标覆盖率变化的权衡以及 risk-coverage AUC。保形预测已在 80%/90%/95% 目标覆盖率下完成：分类平均经验覆盖率为 0.814/0.918/0.956，回归平均经验覆盖率为 0.823/0.925/0.962；区间宽度和集合大小随覆盖目标增加而上升，符合保形预测的风险-信息量权衡。",
            )
        elif text.startswith("样本级决策卡作为"):
            set_text(
                p,
                "样本级可靠性输出用于把预测结果转化为药物化学用户可检查的证据单元。每张决策卡同时给出预测值或类别概率、保形区间或预测集、AD 距离、风险分位、最近邻结构及标签差异、片段解释、是否拒用和拒用原因。该输出不参与模型选择，但能把 FZYC-Mol 的可靠性主张从平均分数延伸到具体分子层面的可审查判断。",
            )
        elif text.startswith("Zhao et al. 2026 将"):
            set_text(
                p,
                "Zhao et al. 2026 将 ADMET 可靠性组织为四类真实挑战：数据稀缺/OOD、类别不平衡、bRo5 化学空间和活性悬崖 [3]。该工作提供了一个高水平基准模板，强调 TabPFNv2、AutoGluon、KPGT、Uni-Mol、GEM、GNN 和传统机器学习在不同场景下的边界。与这种大规模 benchmark 不同，FZYC-Mol 的核心不是重新证明哪类模型总体最强，而是在候选池进入冻结测试之前，给出一个可审计的 accept/reject/retain 治理机制，并把负结果、roughness 风险和样本级失败案例纳入同一证据链。",
            )
        elif text.startswith("这种差异化定位决定了本文的补强方向"):
            set_text(
                p,
                "这种差异化定位决定了本文的补强方向。实验 A-H 将对照论文中的强基线、Perimeter split、类别不平衡、bRo5、活性悬崖、粗糙度、候选接受/拒绝审计和校准/保形预测转化为 FZYC-Mol 的候选治理压力测试。由此，本文不是把复杂模块直接写入主结论，而是让每个模块在同一验证证据下被接受、保留或拒绝。该写法比单纯宣称性能提升更稳健，也更符合当前结果中“选择性增益多于普遍胜出”的事实。",
            )
        elif text.startswith("本研究仍存在若干局限"):
            set_text(
                p,
                "本研究仍存在若干局限。首先，验证-测试排名审计和 9 个代表性终点的 3×3 nested validation 表明，验证集治理可以减少测试集事后选择，但内外层验证尚未覆盖全部 MoleculeNet、TDC、bRo5 和 MoleculeACE 终点，不能保证所有任务均达到测试最优；因此，小幅增益应与 regret、optimism gap、Top3 hit 和负结果共同解释。其次，TabPFNv2、AutoGluon、KPGT representation、bRo5 和 MoleculeACE 30 任务仍需在完整同划分结果中继续扩展，未完成的候选不能写成已取得性能提升。第三，bRo5、活性悬崖和低相似度样本仍是主要失败边界，FZYC-Mol 更适合识别和审计这些风险，而不是宣称已解决它们。第四，基序归因、片段富集和粗糙度相关性仍属于关联证据，不能替代因果机制或湿实验验证。",
            )
        elif text.startswith("进一步验证应优先扩展真实挑战矩阵"):
            set_text(
                p,
                "进一步验证应优先扩展真实挑战矩阵，而不是简单堆叠更多模型。更有价值的路线是在更多公开 ADMET、bRo5 和活性悬崖终点上系统评估 XGBoost、CatBoost、RandomForest、ExtraTrees、Chemprop、TabPFNv2、AutoGluon、KPGT/GEM/Uni-Mol 表征、目标变换、采样策略和 Top-K 集成，并把最终保留决策与 roughness、低相似度、Perimeter split 和 validation-test rank instability 关联起来。",
            )
        elif text.startswith("本文提出并系统评估了 FZYC-Mol"):
            set_text(
                p,
                "本文提出并系统评估了 FZYC-Mol，一种由验证集治理、适用域感知的多专家分子性质预测框架。结果表明，在 MoleculeNet、TDC ADMET、bRo5/肽类压力测试、Random/Scaffold/Perimeter 划分、低相似度子集、MoleculeACE 30 任务活性悬崖、ROGI/MODI/SARI 粗糙度诊断、候选接受/拒绝审计、校准/保形预测和基序/片段解释等多条证据线上，FZYC-Mol 能够提供比单一模型分数更完整的可靠性画像。其核心价值不是保证每个终点取得测试最优，而是在固定候选池内透明地接受、保留、拒绝和审计候选策略；当强基线、补救头、低成本重构或多方法融合通过验证集门控时，它们可以进入最终保留结果，当证据不足时则作为负结果保留。",
            )
        elif text.startswith("未来工作可沿三条路线推进"):
            set_text(
                p,
                "未来工作可沿三条路线推进。第一，继续扩展强基线同划分、bRo5、MoleculeACE 和盲测风格验证，以检验选择器在更多真实挑战下的稳定性。第二，在少量代表性低分任务上尝试受控适配器或轻量微调，但仍需使用嵌套验证限制过拟合。第三，进一步加强样本级可靠性输出，把基序、最近邻、不确定性、适用域、roughness 和实验标签噪声联系起来，为后续湿实验或专家审查提供更清晰的优先级。",
            )

    update_experiment_matrix(doc)

    table_caption = find_para(doc, lambda t: t.startswith("表 1A."))
    insert_paragraph_after(
        table_caption,
        "实现层面，新增 experiment_update 运行层以保证补充实验不污染既有结果。该层包含 data_splits.py、baselines_strong.py、imbalance_panel.py、bro5_panel.py、moleculeace_cliff_audit.py、roughness_diagnostics.py、selector_audit.py 和 aggregate_results.py，并统一导出 split_indices.csv、strong_baselines_metrics.csv、imbalance_panel_metrics.csv、bro5_data_status.csv、moleculeace_30_task_summary.csv、roughness_selector_risk_correlation.csv、selection_audit.csv、negative_results.csv 及 output_field_dictionary.csv。若 bRo5 数据或可选依赖缺失，模块输出 missing_data_report.csv 或 missing_dependency_report.csv，而不是静默跳过。",
    )

    doc.save(str(output))
    print("source=", source)
    print("output=", output)
    print("paragraphs=", len([p for p in doc.paragraphs if p.text.strip()]))
    print("tables=", len(doc.tables))
    print("images=", len(doc.inline_shapes))


if __name__ == "__main__":
    main()
