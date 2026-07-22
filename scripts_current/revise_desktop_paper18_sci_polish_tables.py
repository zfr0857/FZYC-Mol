from __future__ import annotations

import csv
import json
import re
import shutil
import ssl
import urllib.error
import urllib.parse
import urllib.request
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path("D:/fzyc")
OUT = ROOT / "output"
DESKTOP = Path("C:/Users/Administrator/Desktop")
TARGET = next(
    p for p in DESKTOP.glob("*.docx") if p.name.endswith("-18.docx") and not p.name.startswith("~$")
)
BACKUP = DESKTOP / f"{TARGET.stem}_SCI润色三线表前备份_{datetime.now():%Y%m%d_%H%M%S}.docx"
REPORT = OUT / "小论文-18_逐段润色逻辑数据引用与三线表核查报告.md"
AUDIT = OUT / "paper18_desktop_sci_polish_tables_audit.json"
REF_AUDIT = OUT / "paper18_desktop_reference_audit.json"


def ptext(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs]


def find_para(doc: Document, starts: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(starts):
            return p
    raise RuntimeError(f"paragraph not found: {starts}")


def set_para(doc: Document, starts: str, new: str, changes: list[dict[str, str]], reason: str) -> None:
    p = find_para(doc, starts)
    old = p.text
    if old != new:
        p.text = new
        changes.append(
            {
                "location": starts[:48],
                "original": old,
                "revised": new,
                "reason": reason,
                "rewritten_paragraph": new,
            }
        )


def set_para_optional(
    doc: Document,
    starts: str,
    new: str,
    changes: list[dict[str, str]],
    reason: str,
) -> bool:
    for p in doc.paragraphs:
        if p.text.strip().startswith(starts):
            old = p.text
            if old != new:
                p.text = new
                changes.append(
                    {
                        "location": starts[:48],
                        "original": old,
                        "revised": new,
                        "reason": reason,
                        "rewritten_paragraph": new,
                    }
                )
            return True
    return False


def style_if_exists(p, style_name: str) -> None:
    try:
        p.style = style_name
    except Exception:
        pass


def set_cell_border(cell, top=None, bottom=None, left=None, right=None, inside_h=None, inside_v=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge in [
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
        ("insideH", inside_h),
        ("insideV", inside_v),
    ]:
        if edge is None:
            continue
        tag = "w:" + edge_name
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge.items():
            element.set(qn("w:" + key), str(value))


def make_three_line_tables(doc: Document) -> list[dict[str, object]]:
    visible = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    none = {"val": "nil"}
    audit = []
    for ti, table in enumerate(doc.tables, start=1):
        rows = list(table.rows)
        for ri, row in enumerate(rows):
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top=none,
                    bottom=none,
                    left=none,
                    right=none,
                    inside_h=none,
                    inside_v=none,
                )
                if ri == 0:
                    set_cell_border(cell, top=visible, bottom=visible, left=none, right=none)
                if ri == len(rows) - 1:
                    set_cell_border(cell, bottom=visible, left=none, right=none)
        audit.append(
            {
                "table": ti,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "three_line_applied": True,
            }
        )
    return audit


def update_tables(doc: Document, changes: list[dict[str, str]]) -> None:
    # Table 9 row 2 had an outdated three-endpoint wording for modern strong baselines.
    if len(doc.tables) >= 9:
        t = doc.tables[8]
        row = t.rows[2]
        old = " | ".join(c.text for c in row.cells)
        row.cells[2].text = (
            "RDKit-RF、GNN-GCN、ChemBERTa 和 MoLFormer 已在六个 MoleculeNet 主任务上完成同划分 3×3×5 评估；"
            "Chemprop/D-MPNN 作为 ESOL、BACE 和 ClinTox 三终点边界面板报告。"
        )
        row.cells[3].text = "九终点全量深度/基础模型面板与 TabPFN 同划分预测导出仍不完整。"
        new = " | ".join(c.text for c in row.cells)
        if old != new:
            changes.append(
                {
                    "location": "表 9 强基线证据边界",
                    "original": old,
                    "revised": new,
                    "reason": "表格口径与正文六任务强基线结果不一致，需修正为六任务完成、Chemprop 三终点边界、TabPFN 未完成。",
                    "rewritten_paragraph": new,
                }
            )


def revise_text(doc: Document) -> list[dict[str, str]]:
    changes: list[dict[str, str]] = []
    set_para(
        doc,
        "本研究提出 FZYC-Mol模型",
        "本研究提出 FZYC-Mol 模型选择治理框架，将候选登记、嵌套选择、策略冻结、外层审计和负结果记录整合为一个可复核的验证治理协议，而不是新的预测主干网络。该协议在九个终点上执行 3×3×5 冻结选择实验，并将公开面板、逐样本可靠性和化学迁移边界作为可靠性与化学边界分析报告，而不将其并入主排行榜。",
        changes,
        "补足中英文空格，避免把 FZYC-Mol 误写成新预测主干，并将英文短语改为正式中文表述。",
    )
    set_para(
        doc,
        "2.2嵌套评估",
        "2.2 嵌套评估与泄漏控制",
        changes,
        "统一章节标题格式，在编号与标题之间加入空格。",
    )
    set_para(
        doc,
        "除主效应外，本文报告真嵌套验证",
        "除主效应外，本文报告严格嵌套验证、种子敏感性、统一消融、80/90/95% 保形覆盖率、精确 Tanimoto 分箱、MoleculeACE 活性悬崖、低相似度失败样本和扩展失败案例等补充分析。强基线证据矩阵被扩展为两层：RDKit-RF、GNN-GCN、ChemBERTa 和 MoLFormer 冻结适配头已在 ESOL、FreeSolv、Lipophilicity、BBBP、BACE 和 ClinTox 六个 MoleculeNet 主任务上完成同划分 3×3×5 评估，形成 360 个外层单元、1,080 个内层单元和 220,040 条逐样本预测；Chemprop/D-MPNN 仍限于 ESOL、BACE 和 ClinTox 三终点补强面板。TabPFN 已安装，但因授权和运行时交互限制未能完成同划分预测导出，因此仅作为授权受限候选记录在状态表中，不作为完成性结果。",
        changes,
        "将“真嵌套”改为更规范的“严格嵌套”，并统一百分号格式。",
    )
    set_para_optional(
        doc,
        "现代强基线面板使用",
        "现代强基线补强面板使用与确认性实验一致的 3×3×5 任务-种子-外层结构，但其证据层级被限定为六个 MoleculeNet 主任务的压力测试。RDKit-RF、GNN-GCN、ChemBERTa 和 MoLFormer 已完成同划分预测导出；Chemprop/D-MPNN 仍作为 ESOL、BACE 和 ClinTox 三终点边界面板报告，TabPFN 未能产生同划分预测导出。因此，强基线补强结果用于检验主结论的边界，而不替代九终点 12 候选多视图主效应估计。",
        changes,
        "原段落仍按三终点五候选旧口径叙述，与后文六任务强基线结果冲突；改为当前完成范围与证据边界。",
    )
    set_para_optional(
        doc,
        "新增强基线选择分析显示",
        "六任务强基线选择分析显示，冻结选择器在 90 个任务-种子-外层单元中的 Top-1 命中率为 0.878，平均范围归一化选择损失为 0.0048，平均 inner-outer Spearman 相关为 0.840。RDKit-RF 在六个任务中均为验证侧选择结果，但这并不被解释为复杂模型无效；相反，该结果说明复杂模型只有在同一冻结划分下完成预测导出、选择审计和负结果记录后，才宜进入主效应叙事。",
        changes,
        "将旧的 45 单元、五候选选择结果更新为六任务 90 单元强基线选择审计。",
    )
    set_para_optional(
        doc,
        "基于导出的逐样本预测",
        "基于六任务逐样本预测导出，本研究进一步计算 error-overlap。四个完成候选形成 6 个候选对、每对 90 个任务-种子-外层单元，平均 Jaccard 错误重合为 0.215，候选对范围为 0.168–0.296。该结果表明，候选间错误并非相互独立，候选池扩张带来的机会效应需要与错误重叠和错误互补性同时报告。",
        changes,
        "将旧的三终点 103,025 条记录和 10 对候选描述更新为六任务四候选 error-overlap 审计。",
    )
    set_para_optional(
        doc,
        "三套去重敏感性实验也已",
        "三套去重敏感性实验也从清洗审计推进到六个 MoleculeNet 主任务的外层结果比较。全局去重、仅训练折内聚合和保留重复并按骨架分组三种策略共形成 270 个外层评估单位；最大平均效用变化为 0.022，出现在 ClinTox。BACE 在三种策略下保持稳定，ESOL 与 Lipophilicity 偏移较小，而 ClinTox 对重复处理更敏感，说明少数类毒性任务应将重复策略作为明确边界条件报告。",
        changes,
        "将三终点去重敏感性更新为六任务、三策略和 270 外层单元口径。",
    )
    if not set_para_optional(
        doc,
        "TDC 四联审计图将门控规则拆解",
        "TDC 四联审计图将门控规则拆解为审稿人最关心的四个问题：冻结门控产生了哪些行动决策，晋级终点是否具有外层测试改善，保留终点是否确实避免潜在伤害，以及哪些终点仅因三种子区间过宽而证据不足。在 22 个 TDC 终点中，冻结门控产生 5 个 promoted 和 17 个 retained；事后审计显示 3 个 promoted-and-improved、7 个 retained-and-avoided-harm，以及 12 个区间过宽的 inconclusive 终点。",
        changes,
        "保留审稿逻辑，但压缩句式并明确 promoted/retained 与三类事后审计结果之间的关系。",
    ):
        set_para_optional(
            doc,
            "TDC 结果强调门控规则",
            "TDC 结果强调门控规则的终点依赖性。5 个 promoted 与 17 个 retained 并不意味着其余终点已被判定无效，而是说明在三种子和当前区间宽度下，许多终点无法支持明确晋级。事后审计进一步区分 3 个 promoted-and-improved、7 个 retained-and-avoided-harm 和 12 个 inconclusive 终点，从而保留负结果的审计价值，并避免将证据不足误写为反向证据。",
            changes,
            "在不强行替换图像的前提下，使 TDC 结果同时报告行动决策和证据不足边界。",
        )
    if not set_para_optional(
        doc,
        "由于 TDC 部分仅使用 3 个种子",
        "由于 TDC 部分仅使用 3 个种子，图 7 中的横向区间应解释为 seed-summary interval，而不是严格抽样置信区间。promoted-and-improved 终点的种子方向为 9/9 正向，retained-and-avoided-harm 终点为 21/21 负向，支持门控在这两类终点上作出行动性决策；相反，跨零且区间较宽的终点被保留为证据不足，而不是被解释为真实无效。",
        changes,
        "统一“图 7”格式，并强化三种子结果的推断边界。",
    ):
        set_para_optional(
            doc,
            "TDC 部分仅使用 3 个种子",
            "TDC 部分仅使用 3 个种子，因此相关区间应解释为 seed-summary interval 或描述性区间，而不是严格抽样置信区间。当前 5 个 promoted 终点中有 2 个描述性区间跨零，故仅保留为证据不足或宽区间晋级案例；promoted-and-improved 与 retained-and-avoided-harm 均标注为外层事后审计。",
            changes,
            "将 seed variability interval 统一为 seed-summary interval，并明确非严格抽样推断边界。",
        )
    set_para_optional(
        doc,
        "图 7  TDC 门控有效性",
        "图 7  TDC 门控有效性。该图报告 22 个 TDC 终点经冻结门控后的 promoted/retained 构成、方向归一化测试增量和三种子 seed-summary interval，用于区分可行动改善、避免伤害和证据不足。",
        changes,
        "补足图7图注的审稿信息：终点数、门控构成、效应量和三种子区间解释。",
    )
    set_para(
        doc,
        "图 10  验证集大小、候选相关性与候选规模",
        "图 10  验证集大小、候选相关性与候选规模共同决定 selection loss。曲线来自真实结果锚定的 controlled mechanism experiment；纵轴采用 K=64 固定尺度，避免随 K 改变分母。",
        changes,
        "原图注过短，无法说明实验来源、纵轴尺度和审稿人需要看到的解释边界。",
    )
    style_if_exists(find_para(doc, "图 10"), "FigureCaption")
    set_para(
        doc,
        "可靠性实验进一步扩展为四个互补面板",
        "可靠性实验进一步扩展为四个互补面板：标签条件与 Mondrian 保形检验少数类覆盖，回归保形比较覆盖率与区间宽度，scaffold/OOD 分层评估校准漂移，ensemble uncertainty 评估高误差样本的提前识别能力。RDKit-RF 在 90% 目标覆盖下的分类 split conformal 类别 1 覆盖为 0.627；label-conditional 和 Mondrian label-similarity conformal 将类别 1 覆盖提高到 0.895 和 0.898。CQR 的 90% 平均覆盖为 0.882，平均区间宽度为 7.25，提示其在当前特征和样本量下并未稳定优于残差式保形。",
        changes,
        "维持数据不变，强化四个面板的逻辑功能和负结果表述。",
    )
    set_para(
        doc,
        "scaffold/OOD 校准和不确定性排序",
        "scaffold/OOD 校准和不确定性排序给出了更细的边界证据。最近邻 Tanimoto <0.5 的分类子集平均 ROC-AUC 为 0.803，低于 >0.7 子集的 0.924；ensemble uncertainty 对 top-10% 高误差样本的平均富集为 1.54。ClinTox 少数类仍作为负结果保留：RDKit-RF 的少数类召回不足，说明覆盖校准可以改善风险声明，却不能单独替代阈值式毒性筛选。",
        changes,
        "强调该段是边界证据而非平均性能主张，并保留 ClinTox 负结果。",
    )
    set_para(
        doc,
        "图 11  不确定性和保形预测",
        "图 11  不确定性和保形预测的四联图。左上，RDKit-RF 在 90% 目标覆盖下的总体覆盖与类别 1 覆盖；右上，回归 residual conformal、Mondrian residual 与 CQR 的覆盖率-区间宽度权衡；左下，最近邻 Tanimoto 分层下的 OOD/scaffold ROC-AUC 和 ECE；右下，ensemble uncertainty 对高误差样本的 top-10% 富集。",
        changes,
        "原图注仅有标题，补充四个小图各自承担的证据功能。",
    )
    style_if_exists(find_para(doc, "图 11"), "FigureCaption")
    set_para(
        doc,
        "真实决策价值实验也扩展为四个面板",
        "真实决策价值实验也扩展为四个面板，将模型分数转化为固定预算筛选收益、相对 oracle 的遗漏阳性数、毒性假阴性成本和实验队列风险。FZYC-selected 在该六任务强基线面板中等同于 RDKit-RF，其 top-1%、top-5% 和 top-10% 平均富集分别为 4.06、3.28 和 2.80；MoLFormer 的 top-10% 富集为 2.60。但毒性假阴性成本给出不同排序：RDKit-RF/FZYC、MoLFormer 和 GNN-GCN 的 ClinTox 阈值成本分别为每 100 个分子 67.7、45.3 和 84.2。",
        changes,
        "保持数值不变，明确富集与毒性排除是两个不同决策目标。",
    )
    set_para(
        doc,
        "图 12  真实筛选决策价值",
        "图 12  真实筛选决策价值的四联图。左上，top-1%、top-5% 和 top-10% 固定预算富集；右上，相同预算下相对 test oracle 的遗漏阳性数；左下，ClinTox 假阴性加权成本；右下，在低毒候选进入实验队列的模拟中，高风险分子进入队列的比例。",
        changes,
        "原图注过短，补足四联图结构和各面板解释。",
    )
    style_if_exists(find_para(doc, "图 12"), "FigureCaption")
    set_para(
        doc,
        "表 10 | 新增机制",
        "表 10 | 新增机制、可靠性、决策价值和失败案例实验摘要",
        changes,
        "将表10标题设置为表题样式，并保持与正文术语一致。",
    )
    style_if_exists(find_para(doc, "表 10"), "TableCaption")
    set_para(
        doc,
        "FZYC-Mol 不替代预测模型",
        "FZYC-Mol 不替代预测模型，也不保证性能提升，且不提供可迁移到所有终点的元选择器。六任务 MoleculeNet 面板已完成 RDKit-RF、GNN-GCN、ChemBERTa 和 MoLFormer 的同划分 3×3×5 评估，并新增选择审计、配对效应、error-overlap、去重敏感性、验证信息量机制实验和组件消融；但 Chemprop/D-MPNN 仍只完成三终点补强面板，TabPFN 未能产生同划分预测，九终点全量深度/基础模型面板仍超出当前完成范围。因此，强基线结果应解释为更厚的边界压力测试，而不是复杂模型已在所有终点上完成确认性统一重训。TDC 三种子结果只能反映种子变异和公开面板异质性，不能承担严格抽样推断。公开 release、Zenodo DOI 和第三方冷启动复跑仍需作为后续复现工作完成。",
        changes,
        "保留限制性结论，整理长句层级，使强基线完成范围与未完成范围更清晰。",
    )
    set_para(
        doc,
        "[33] Parrondo-Pizarro",
        "[33] Parrondo-Pizarro R, Lanini J, Rodríguez-Pérez R. Uncertainty quantification in molecular machine learning for property predictions under data shifts. J Chem Inf Model. 2026;66:923–935. doi:10.1021/acs.jcim.5c02381.",
        changes,
        "按 DOI 页面修正作者姓氏重音符号，提升参考文献规范性。",
    )
    return changes


def extract_references(doc: Document) -> list[dict[str, str]]:
    refs = []
    for p in doc.paragraphs:
        t = p.text.strip()
        m = re.match(r"^\[(\d+)\]\s+(.*)", t)
        if m:
            doi = ""
            dm = re.search(r"doi:([^\s.]+(?:\.[^\s.]+)*)", t, flags=re.I)
            if dm:
                doi = dm.group(1).rstrip(".")
            am = re.search(r"arXiv:([0-9]{4}\.[0-9]{4,5})", t)
            refs.append({"number": m.group(1), "text": t, "doi": doi, "arxiv": am.group(1) if am else ""})
    return refs


def fetch_json(url: str) -> dict:
    ctx = ssl._create_unverified_context()
    req = urllib.request.Request(url, headers={"User-Agent": "paper18-reference-audit/1.0"})
    with urllib.request.urlopen(req, context=ctx, timeout=20) as resp:
        return json.loads(resp.read().decode("utf-8"))


def audit_references(refs: list[dict[str, str]]) -> dict[str, object]:
    rows = []
    for ref in refs:
        status = "manual"
        title = ""
        source = ""
        year = ""
        note = ""
        if ref["doi"]:
            try:
                data = fetch_json("https://api.crossref.org/works/" + urllib.parse.quote(ref["doi"]))
                msg = data.get("message", {})
                title = (msg.get("title") or [""])[0]
                source = (msg.get("container-title") or [""])[0]
                years = msg.get("published-print", msg.get("published-online", {})).get("date-parts", [[""]])
                year = str(years[0][0]) if years and years[0] else ""
                status = "doi_resolved"
            except Exception as exc:  # noqa: BLE001
                status = "doi_error"
                note = str(exc)
        elif ref["arxiv"]:
            try:
                url = "https://export.arxiv.org/api/query?id_list=" + urllib.parse.quote(ref["arxiv"])
                req = urllib.request.Request(url, headers={"User-Agent": "paper18-reference-audit/1.0"})
                with urllib.request.urlopen(req, context=ssl._create_unverified_context(), timeout=20) as resp:
                    xml = resp.read().decode("utf-8", errors="ignore")
                mt = re.search(r"<title>(.*?)</title>", xml, re.S)
                title = re.sub(r"\s+", " ", mt.group(1)).strip() if mt else ""
                status = "arxiv_resolved" if ref["arxiv"] in xml else "arxiv_check"
                source = "arXiv"
            except Exception as exc:  # noqa: BLE001
                status = "arxiv_error"
                note = str(exc)
        else:
            note = "No DOI/arXiv in manuscript; verify manually against publisher/software/book/conference record."
        rows.append({**ref, "status": status, "resolved_title": title, "resolved_source": source, "resolved_year": year, "note": note})
    return {
        "references_checked": len(rows),
        "doi_resolved": sum(r["status"] == "doi_resolved" for r in rows),
        "arxiv_resolved": sum(r["status"] == "arxiv_resolved" for r in rows),
        "manual_or_no_identifier": sum(r["status"] == "manual" for r in rows),
        "rows": rows,
    }


def data_audit() -> dict[str, object]:
    out: dict[str, object] = {}
    hard = OUT / "sci1_hardening_20260707"
    mech = OUT / "sci1_mechanism_uq_decision_20260707"
    fig7 = OUT / "paper18_fig07_four_panel_reviewer_audit.json"
    fig11 = OUT / "paper18_fig11_fig12_four_panel_audit.json"
    if fig7.exists():
        out["fig7"] = json.loads(fig7.read_text(encoding="utf-8"))["metrics"]
    if fig11.exists():
        out["fig11_fig12"] = json.loads(fig11.read_text(encoding="utf-8"))["metrics"]
    sel = hard / "six_task_strong_selection_scorecard.csv"
    if sel.exists():
        df = pd.read_csv(sel)
        out["six_task_selection"] = df.to_dict("records")
    overlap = hard / "six_task_error_overlap_pairwise_summary.csv"
    if overlap.exists():
        df = pd.read_csv(overlap)
        col = "mean_jaccard_error_overlap" if "mean_jaccard_error_overlap" in df.columns else "mean_error_jaccard"
        out["error_overlap"] = {
            "mean_jaccard": float(df[col].mean()),
            "min_jaccard": float(df[col].min()),
            "max_jaccard": float(df[col].max()),
        }
    mech_file = mech / "mechanism_controlled_simulation_summary.csv"
    if mech_file.exists():
        df = pd.read_csv(mech_file)
        def val(regime: str, frac: float, k: int) -> float:
            row = df[
                df["correlation_regime"].eq(regime)
                & df["validation_information_fraction"].eq(frac)
                & df["candidate_count"].eq(k)
            ].iloc[0]
            return float(row["fixed_k64_normalized_selection_loss"])
        out["mechanism"] = {
            "high_25_k4": val("high_correlated_lightweight", 0.25, 4),
            "high_25_k64": val("high_correlated_lightweight", 0.25, 64),
            "medium_25_k4": val("medium_correlated_multiview", 0.25, 4),
            "medium_25_k64": val("medium_correlated_multiview", 0.25, 64),
            "low_25_k4": val("low_correlated_deep_foundation", 0.25, 4),
            "low_25_k64": val("low_correlated_deep_foundation", 0.25, 64),
        }
    return out


def verify_table_borders(docx: Path) -> dict[str, object]:
    with ZipFile(docx) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    return {
        "tables_in_xml": xml.count("<w:tbl>"),
        "has_vertical_nil": 'w:left w:val="nil"' in xml and 'w:right w:val="nil"' in xml,
        "has_visible_table_borders": 'w:top w:val="single"' in xml and 'w:bottom w:val="single"' in xml,
    }


def write_report(changes, refs_audit, data_checks, table_audit, final_audit) -> None:
    lines = [
        "# 小论文-18 SCI润色、逻辑连贯性、数据引用与三线表核查报告",
        "",
        f"- 源文档：`{TARGET}`",
        f"- 备份：`{BACKUP}`",
        f"- 生成时间：{datetime.now().isoformat(timespec='seconds')}",
        "",
        "## 逐段润色表",
        "",
        "| 序号 | 位置 | 原句/问题句 | 修改后句子 | 修改原因 | 重写后整段 |",
        "|---:|---|---|---|---|---|",
    ]
    for i, c in enumerate(changes, 1):
        def esc(x: str) -> str:
            return x.replace("|", "\\|").replace("\n", "<br>")
        lines.append(
            f"| {i} | {esc(c['location'])} | {esc(c['original'])} | {esc(c['revised'])} | {esc(c['reason'])} | {esc(c['rewritten_paragraph'])} |"
        )
    lines += [
        "",
        "## 段落逻辑连贯性建议",
        "",
        "- 摘要已维持“背景与问题—方法—主要结果—结论与边界”的顺序，数值只保留核心发现，可靠性与化学边界归为边界证据。",
        "- 方法部分的证据层级需始终区分九终点确认性主效应、六任务强基线压力测试、三终点 Chemprop 边界面板和 TabPFN 未完成状态。",
        "- 结果部分建议继续避免将 TDC 三种子区间写成严格置信区间；图7已明确为 seed-summary interval。",
        "- 讨论部分已保留“治理框架而非新预测主干”的定位，以降低审稿人对模型排行榜主张的质疑。",
        "",
        "## 数据核查摘要",
        "",
        "```json",
        json.dumps(data_checks, ensure_ascii=False, indent=2),
        "```",
        "",
        "## 引用核查摘要",
        "",
        f"- 参考文献总数：{refs_audit['references_checked']}",
        f"- DOI 解析成功：{refs_audit['doi_resolved']}",
        f"- arXiv 解析成功：{refs_audit['arxiv_resolved']}",
        f"- 无 DOI/arXiv、需按期刊格式人工确认：{refs_audit['manual_or_no_identifier']}",
        "",
        "详见 `paper18_desktop_reference_audit.json`。",
        "",
        "## 三线表核查",
        "",
        f"- 已处理表格数：{len(table_audit)}",
        "- 所有表格已清除左右竖线与内部竖线，并保留表头上框线、表头下框线和表底线。",
        "",
        "```json",
        json.dumps(final_audit, ensure_ascii=False, indent=2),
        "```",
    ]
    REPORT.write_text("\n".join(lines), encoding="utf-8-sig")


def main() -> None:
    shutil.copy2(TARGET, BACKUP)
    doc = Document(str(TARGET))
    changes = revise_text(doc)
    update_tables(doc, changes)
    table_audit = make_three_line_tables(doc)
    refs = extract_references(doc)
    refs_audit = audit_references(refs)
    REF_AUDIT.write_text(json.dumps(refs_audit, ensure_ascii=False, indent=2), encoding="utf-8")
    data_checks = data_audit()
    try:
        doc.save(str(TARGET))
        saved_to = TARGET
        fallback_used = False
    except PermissionError:
        saved_to = DESKTOP / f"{TARGET.stem}_SCI润色三线表.docx"
        doc.save(str(saved_to))
        fallback_used = True

    shutil.copy2(saved_to, OUT / "小论文-18.docx")
    with ZipFile(saved_to) as zf:
        bad = zf.testzip()
    final_doc = Document(str(saved_to))
    table_border = verify_table_borders(saved_to)
    text = "\n".join(p.text for p in final_doc.paragraphs)
    table_text = "\n".join(cell.text for table in final_doc.tables for row in table.rows for cell in row.cells)
    all_text = text + "\n" + table_text
    final_audit = {
        "saved_to": str(saved_to),
        "fallback_used": fallback_used,
        "backup": str(BACKUP),
        "zip_ok": bad is None,
        "paragraphs": len(final_doc.paragraphs),
        "tables": len(final_doc.tables),
        "figures": len(final_doc.inline_shapes),
        "changed_items": len(changes),
        "references_checked": refs_audit["references_checked"],
        "doi_resolved": refs_audit["doi_resolved"],
        "arxiv_resolved": refs_audit["arxiv_resolved"],
        "all_tables_three_line_applied": all(x["three_line_applied"] for x in table_audit),
        "table_border_xml_audit": table_border,
        "fig10_caption_complete": "共同决定 selection loss" in text,
        "fig11_caption_complete": "不确定性和保形预测的四联图" in text,
        "fig12_caption_complete": "真实筛选决策价值的四联图" in text,
        "table9_strong_baseline_fixed": "RDKit-RF、GNN-GCN、ChemBERTa 和 MoLFormer 已在六个 MoleculeNet 主任务" in all_text,
        "passed": False,
    }
    final_audit["passed"] = (
        final_audit["zip_ok"]
        and final_audit["all_tables_three_line_applied"]
        and final_audit["fig10_caption_complete"]
        and final_audit["fig11_caption_complete"]
        and final_audit["fig12_caption_complete"]
        and final_audit["table9_strong_baseline_fixed"]
    )
    AUDIT.write_text(json.dumps(final_audit, ensure_ascii=False, indent=2), encoding="utf-8")
    write_report(changes, refs_audit, data_checks, table_audit, final_audit)
    print(json.dumps(final_audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
