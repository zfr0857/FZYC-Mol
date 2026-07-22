from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
TABLE_DIR = ROOT / "reports" / "manuscript_tables"
FIG_DIR = ROOT / "reports" / "manuscript_figures_polished"
TABLE_PREVIEW_DIR = ROOT / "reports" / "manuscript_tables_polished"

BASE_DOCX = DOCS / "manuscript_draft_full_zh_detailed_20260602_visual_polish_20260602.docx"
BASE_MD = DOCS / "manuscript_draft_full_zh_integrated_20260531.md"
OUT_DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_20260603.docx"
OUT_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_20260603.md"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def normalize_styles(doc: Document) -> None:
    for style_name, size in [("Normal", 10.5), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.18
    doc.styles["Normal"].paragraph_format.space_after = Pt(5)


def set_margins(section) -> None:
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)


def add_landscape_section(doc: Document) -> None:
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    set_margins(section)


def add_portrait_section(doc: Document) -> None:
    section = doc.add_section()
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width
    set_margins(section)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    set_run_font(run, bold=True)


def add_para(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color="111827")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=9.0, bold=True, color="334155")


def add_figure(doc: Document, path: Path, caption: str, width: float = 7.0) -> None:
    if not path.exists():
        add_para(doc, f"[缺少图片：{path}]", first_line=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "DDE6ED") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), color)


def set_cell_text(cell, text: object, size: float = 7.0, bold: bool = False, color: str = "111827", align: str = "center") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT, "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    run = p.add_run("" if pd.isna(text) else str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc: Document, label: str, caption: str, frame: pd.DataFrame, note: str = "", font_size: float = 6.6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{label}. {caption}")
    set_run_font(run, size=9.2, bold=True, color="111827")

    table = doc.add_table(rows=1, cols=len(frame.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for c, col in enumerate(frame.columns):
        cell = table.rows[0].cells[c]
        shade_cell(cell, "14532D")
        set_cell_border(cell)
        set_cell_text(cell, col, font_size, bold=True, color="FFFFFF")
    for r, (_, row) in enumerate(frame.iterrows(), start=1):
        cells = table.add_row().cells
        for c, col in enumerate(frame.columns):
            cell = cells[c]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            shade_cell(cell, "F8FAFC" if r % 2 == 0 else "FFFFFF")
            if c == 0:
                shade_cell(cell, "ECFDF5")
            text = row[col]
            align = "left" if col in {"主要结论", "文件", "说明", "入选模型", "候选来源", "结果定位"} else "center"
            color = "111827"
            bold = c == 0
            if str(text).startswith("+"):
                color = "166534"
                bold = True
            elif str(text).startswith("-"):
                color = "B91C1C"
                bold = True
            set_cell_text(cell, text, font_size, bold=bold, color=color, align=align)
    if note:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f"注：{note}")
        set_run_font(run, size=8.0, color="64748B")


def read_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path).fillna("")


def fmt_float(value: object, digits: int = 4) -> str:
    try:
        value = float(value)
        return f"{value:.{digits}f}"
    except (TypeError, ValueError):
        return ""


def fmt_delta(value: object, digits: int = 4) -> str:
    try:
        value = float(value)
        return f"{value:+.{digits}f}"
    except (TypeError, ValueError):
        return ""


def build_experiment_index() -> pd.DataFrame:
    rows = [
        ["核心协议", "Table 1", "table1_dataset_protocol.csv", "MoleculeNet、TDC、MoleculeACE、OpenADMET 等数据协议；明确 split、seed、primary metric。"],
        ["MoleculeNet 主结果", "Figure 2-3 / Table 2", "table2_moleculenet_main*.csv", "ESOL、BACE、ClinTox、Lipo 等主结果；Lipo rescue 已接入，FreeSolv 保留为 remaining limitation。"],
        ["TDC official split", "Figure 5 / Table 3", "table3_tdc_official_admet*.csv", "Caco2、HIA、Pgp 等外部 ADMET official split；展示 random-to-scaffold 惩罚。"],
        ["Split realism", "Figure 4 / Table 4", "table4_split_realism.csv", "random、scaffold、structure-separated split 对比，用于证明结构外推压力。"],
        ["Ablation/significance", "Figure S3 / Table 5", "table5_ablation_significance.csv", "验证 selector 相对单专家和消融池的 win/tie/loss 与显著性证据。"],
        ["Reliability/AD", "Figure 6 / Table 6 / Table S6", "table6_reliability_ad.csv; table11_reliability_summary.csv", "ensemble std、error model、inverse Tanimoto、reconstruction AD 等用于定位高误差风险。"],
        ["Calibration/conformal", "Figure S2/S5", "table9_risk_calibrated_selector.csv", "Brier、ECE、conformal coverage 与 risk-calibrated selector 支持可靠性叙事。"],
        ["MoleculeACE/cliffs", "Figure S7 / Table S4-S5", "table10_moleculeace_cliff_objective_selector*.csv", "活性悬崖和早期富集补充 ROC-AUC/RMSE 之外的筛选相关指标。"],
        ["OpenADMET fast appendix", "Figure S8 / Table S8", "table13_openadmet_expansionrx_fast_external_benchmark.csv", "快速外部可行性检查，不替代主 benchmark。"],
        ["TDC full-panel appendix", "Table S9-S11", "table14-16_*.csv", "18 个 ADME + 4 个 toxicity endpoint；包含 performance mode 与 roughness/literature alignment。"],
        ["Pretrained/rescue heads", "Table S13-S14", "table18-19_*.csv", "冻结 ChemBERTa/MoLFormer embedding rescue heads；Lipo 被 validation-only selector 接受。"],
        ["Targeted rebuild", "Figure S10 / Table S22-S23", "table27-28_*.csv", "低分 MoleculeNet 模块定向重构；FreeSolv 小幅 retained-best 改善，其他 endpoint 保留旧结果。"],
        ["Nature-style MoleculeNet fusion", "Figure S11 / Table S24-S26", "table29-31_*.csv", "融合 tabular、embedding、AD-gated、rank fusion 等方法；BBBP/ClinTox 有接入收益。"],
        ["TDC Nature-style fusion", "Figure S12 / Table S27-S29", "table32-34_*.csv", "TDC official panel 多方法融合；Caco2、HIA、Pgp 等 endpoint 获得 retained-best 改善。"],
        ["3D-lite/roughness experts", "Figure S13 / Table S30-S32", "table35-37_*.csv", "3D-lite 和 roughness-weighted regression 未被 validation gate 接入；oracle 显示候选池仍有未捕获信号。"],
        ["Selector strategy audit", "Figure S14 / Table S33-S35", "table38-40_*.csv", "固定 selector policy 审计，risk-adjusted positive endpoint pool 最多，但平均 delta 仍为负。"],
        ["Formal fixed selector", "Figure 8 / Table 8 / Table S36-S38", "table8_formal_fixed_selector_main.csv; table41-43_*.csv", "预先固定 risk_adjusted_lambda_0.5，32 个 endpoint-metric 中 10 个正向 promotion。"],
    ]
    return pd.DataFrame(rows, columns=["实验模块", "图表编号", "文件", "主要结论"])


def build_policy_summary() -> pd.DataFrame:
    summary = read_csv(TABLE_DIR / "table43_formal_fixed_selector_policy_summary.csv")
    rows = []
    for _, row in summary.iterrows():
        policy = "risk-adjusted λ=0.5" if str(row["fixed_policy"]) == "risk_adjusted_lambda_0.5" else "stability tie-breaker"
        rows.append(
            {
                "固定策略": policy,
                "正向/总数": f"{int(row['n_positive_vs_current'])}/{int(row['n_endpoint_metrics'])}",
                "负向数": int(row["n_negative_vs_current"]),
                "正向率": f"{float(row['positive_rate']) * 100:.1f}%",
                "平均 Δ": fmt_delta(row["mean_delta_vs_current"]),
                "最大正向 Δ": fmt_delta(row["largest_positive_delta"]),
                "最大负向 Δ": fmt_delta(row["largest_negative_delta"]),
            }
        )
    return pd.DataFrame(rows)


def build_roughness_summary() -> pd.DataFrame:
    table = read_csv(TABLE_DIR / "table35_3d_roughness_regression_retained_best.csv")
    keep = ["source", "dataset", "primary_metric", "new_primary_mean", "previous_primary_mean", "delta_vs_previous", "retained_source"]
    out = table[keep].copy()
    out = out.rename(
        columns={
            "source": "来源",
            "dataset": "Endpoint",
            "primary_metric": "指标",
            "new_primary_mean": "3D/roughness 候选",
            "previous_primary_mean": "当前 retained",
            "delta_vs_previous": "Δ",
            "retained_source": "结论",
        }
    )
    out["3D/roughness 候选"] = out["3D/roughness 候选"].map(fmt_float)
    out["当前 retained"] = out["当前 retained"].map(fmt_float)
    out["Δ"] = out["Δ"].map(fmt_delta)
    out["结论"] = out["结论"].map(lambda x: "保留当前" if str(x) == "previous_retained" else str(x))
    return out


def build_promoted_table() -> pd.DataFrame:
    path = TABLE_DIR / "table8_formal_fixed_selector_main.csv"
    if path.exists():
        return read_csv(path)
    table = read_csv(TABLE_DIR / "table42_formal_risk_adjusted_integration_retained_best.csv")
    table = table[table["promotion_status"].eq("promoted_by_fixed_policy")].copy()
    return table


def append_integration_update(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "8. 2026-06-03 全量整合更新：selector 改进、负结果与主表重排", level=1)
    add_para(
        doc,
        "本节将 2026-06-03 之前分散在初稿、实验报告和附录更新中的结果统一接入同一份完整中文初稿。此前 2026-06-02 详细版已经包含 MoleculeNet 主面板、TDC official split、split realism、uncertainty/AD、motif attribution、rescue heads、targeted rebuild、Nature-style multimethod fusion 和 formal external appendix。本次新增的核心内容是：3D-lite/roughness-weighted regression 的负结果、selector strategy audit、预先固定的 formal fixed-selector integration，以及重新排版后的主图和主表。",
    )
    add_para(
        doc,
        "整合后的论文叙事应更加清楚：FZYC-Mol 不是继续无边界堆叠 candidate 的模型集合，而是一个 validation-governed reliability framework。性能提升来自预注册的候选治理、endpoint-specific retained-best gate、selector robustness、reliability/AD 诊断和针对低分模块的定向 rescue。对于没有通过 validation gate 的候选，例如 3D-lite/roughness experts，需要作为负结果保留，以证明论文没有用测试集事后挑选策略。",
    )

    add_heading(doc, "8.1 全部实验模块索引", level=2)
    add_landscape_section(doc)
    add_table(
        doc,
        "整合表 A",
        "截至 2026-06-03 的全部实验结果与论文图表对应关系。",
        build_experiment_index(),
        note="该索引用于替代分散文件查找；完整 CSV、PNG/SVG 和报告均已纳入 submission package。",
        font_size=6.0,
    )
    add_portrait_section(doc)

    add_heading(doc, "8.2 精排主图与主表更新", level=2)
    add_para(
        doc,
        "上一版 20260603 快速稿内容偏短，主要原因是它只展示了 formal fixed-selector integration。现在保留 20260602 详细初稿的完整正文，并把新版 Figure 8 与 Table 8 接入主文，同时提供 MoleculeNet 和 TDC official split 的精排主表预览。主文中展示 retained 后的结果变化；完整固定策略的负向审计仍保留在 Tables S36-S38。",
    )
    add_figure(
        doc,
        FIG_DIR / "fig22_formal_fixed_selector_integration.png",
        "图 8. Formal fixed-selector integration。risk-adjusted λ=0.5 为主策略，stability tie-breaker 为敏感性分析。图中显示固定策略正向 endpoint 数与 retained-best promotion。",
        width=6.9,
    )
    add_figure(
        doc,
        TABLE_PREVIEW_DIR / "table2_moleculenet_main_polished.png",
        "精排表预览 1. MoleculeNet 主结果。该表保留主文可读字段；完整 seed-level 与候选池细节见 Table S1、S22-S26。",
        width=7.2,
    )
    add_figure(
        doc,
        TABLE_PREVIEW_DIR / "table3_tdc_official_admet_polished.png",
        "精排表预览 2. TDC official split 摘要。该表同时显示 retained 结果和平均 scaffold 惩罚。",
        width=7.2,
    )
    add_figure(
        doc,
        TABLE_PREVIEW_DIR / "table8_formal_fixed_selector_main_polished.png",
        "精排表预览 3. Formal fixed-selector rescue summary。该表只列出预先固定 risk-adjusted selector 的正向 promotion。",
        width=7.2,
    )

    add_heading(doc, "8.3 Formal fixed-selector integration 的正式结论", level=2)
    add_table(
        doc,
        "整合表 B",
        "固定 selector policy 的正负结果汇总。",
        build_policy_summary(),
        note="risk-adjusted λ=0.5 是主策略；stability tie-breaker 仅作为敏感性分析。平均 Δ 为负，说明该策略不能全局替换主 selector。",
        font_size=7.0,
    )
    add_para(
        doc,
        "正式 integration 合并了已有候选池，并以 canonical endpoint 和 primary metric 重新组织候选。主策略固定为 risk_adjusted_lambda_0.5，不再按 endpoint 使用测试集挑选不同 policy。结果显示，该策略在 32 个 endpoint-metric 中有 10 个相对当前 retained-best 为正向，最明显的收益集中在 half-life Obach、clearance microsome AZ、clearance hepatocyte AZ、CYP2C9 substrate、VDss、Pgp、BBBP 和 ClinTox。",
    )
    add_table(
        doc,
        "整合表 C",
        "Formal risk-adjusted selector 带来的 retained-best promotion。",
        build_promoted_table(),
        note="Δ 已按指标方向统一为正值代表更优；完整正负固定策略审计见 Table S36-S38。",
        font_size=6.4,
    )
    add_para(
        doc,
        "这些提升的共同特点是 endpoint 本身存在较强验证噪声、类别不平衡、长尾回归或 roughness 压力。换言之，本轮结果支持的是 selector-level performance rescue，而不是新增大模型或继续扩大 candidate pool。论文中应将其写成 targeted rescue 和 selector robustness，而非全局 SOTA 替代。",
    )

    add_heading(doc, "8.4 3D-lite 与 roughness-weighted regression 的负结果", level=2)
    add_table(
        doc,
        "整合表 D",
        "3D-lite/roughness-weighted regression gate 的 retained-best 结果。",
        build_roughness_summary(),
        note="所有 endpoint 均保留当前 retained-best；oracle audit 只用于诊断候选池中是否存在未被 validation gate 捕捉的信号。",
        font_size=6.5,
    )
    add_para(
        doc,
        "3D-lite 和 roughness-weighted tree 是合理的下一步尝试，因为它们贴合 FreeSolv、clearance、half-life 等低分回归模块。然而正式 validation gate 没有接入任何新 endpoint，这一负结果很有价值：它说明继续堆 candidate 并不能自然提高主结果，真正瓶颈在于 selector 如何处理验证噪声、roughness 和 endpoint-specific uncertainty。该结论直接支撑转向 fixed selector integration 的决策。",
    )

    add_heading(doc, "8.5 写作层面的合并建议", level=2)
    add_para(
        doc,
        "最终投稿版本建议采用“三层结果”结构。第一层是主结果：MoleculeNet、TDC official split、split realism、reliability/AD、motif interpretability 和 Figure 8/Table 8。第二层是强附录：formal external appendix、TDC performance-mode、rescue-integrated selector、Nature-style fusion、selector audit。第三层是诊断附录：3D-lite/roughness negative result、oracle audit、roughness/literature alignment、candidate family details。这样既能展示实验厚度，又不会让主文被几十张原始 CSV 表淹没。",
    )
    add_para(
        doc,
        "表格呈现方面，主文只放经过精排的紧凑表；完整候选级数据保留在 supplementary tables。对于 fixed selector，主文 Table 8 只显示 promotion；但正文必须明确说明完整固定策略有 22 个负向 endpoint-metric，完整审计在 Table S36-S38。这个写法比只报告正向结果更稳，也更符合高水平审稿对选择偏差的敏感性。",
    )


def make_markdown() -> None:
    base = BASE_MD.read_text(encoding="utf-8", errors="replace") if BASE_MD.exists() else ""
    additions = [
        "# 2026-06-03 全量整合更新",
        "",
        "本文件在 2025-05-31 中文整合初稿基础上，追加 2026-06-02/2026-06-03 的全部新增实验与写作更新。",
        "",
        "## 全部实验模块索引",
        "",
        build_experiment_index().to_markdown(index=False),
        "",
        "## Formal fixed-selector policy summary",
        "",
        build_policy_summary().to_markdown(index=False),
        "",
        "## Formal risk-adjusted selector promotion",
        "",
        build_promoted_table().to_markdown(index=False),
        "",
        "## 3D-lite/roughness negative gate",
        "",
        build_roughness_summary().to_markdown(index=False),
        "",
    ]
    for path in [
        DOCS / "recent_literature_competitive_improvement_20260602.md",
        DOCS / "nature_literature_method_fusion_update_20260602.md",
        DOCS / "three_d_roughness_regression_update_20260603.md",
        DOCS / "selector_strategy_audit_update_20260603.md",
        DOCS / "formal_fixed_selector_integration_update_20260603.md",
        DOCS / "manuscript_draft_full_zh_polished_20260603.md",
    ]:
        if path.exists():
            additions.extend(["", f"## 附加整合来源：{path.name}", "", path.read_text(encoding="utf-8", errors="replace")])
    OUT_MD.write_text(base.rstrip() + "\n\n" + "\n".join(additions) + "\n", encoding="utf-8")


def main() -> None:
    if not BASE_DOCX.exists():
        raise FileNotFoundError(BASE_DOCX)
    shutil.copy2(BASE_DOCX, OUT_DOCX)
    doc = Document(OUT_DOCX)
    normalize_styles(doc)
    for section in doc.sections:
        set_margins(section)
    append_integration_update(doc)
    doc.save(OUT_DOCX)
    make_markdown()
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {OUT_MD}")


if __name__ == "__main__":
    main()
