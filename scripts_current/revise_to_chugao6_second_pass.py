from __future__ import annotations

import math
import re
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "reviewer_revision_20260607"
FIG_DIR = REPORT_DIR / "figures"
FIG_DIR.mkdir(parents=True, exist_ok=True)


def find_source_docx() -> Path:
    candidates = list(Path(r"C:/Users/Administrator/Desktop").rglob("FZYC-Mol_*.docx"))
    matches = [p for p in candidates if p.name.endswith("-5.docx")]
    if not matches:
        raise FileNotFoundError("Could not locate the current -5 manuscript.")
    return matches[0]


def set_run_font(run, size: float | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)


def set_para_text(paragraph, text: str, size: float | None = None) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size)


def insert_paragraph_after(paragraph, text: str = "", size: float | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.paragraphs[-1]
    # Re-find the inserted paragraph by object identity in document order.
    for para in paragraph._parent.paragraphs:
        if para._p is new_p:
            new_para = para
            break
    if text:
        set_para_text(new_para, text, size)
    return new_para


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def add_picture_to_paragraph(paragraph, image_path: Path, width_inches: float = 6.7) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def save_figure(fig, base: Path) -> None:
    fig.savefig(base.with_suffix(".png"), dpi=600, bbox_inches="tight", facecolor="white")
    fig.savefig(base.with_suffix(".pdf"), bbox_inches="tight", facecolor="white")
    fig.savefig(base.with_suffix(".svg"), bbox_inches="tight", facecolor="white")
    fig.savefig(base.with_suffix(".tiff"), dpi=600, bbox_inches="tight", facecolor="white")


def draw_box(ax, xy, wh, text, fc="#F6F8FA", ec="#2F3A45", lw=1.25, fontsize=9.5):
    x, y = xy
    w, h = wh
    box = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.018,rounding_size=0.025",
        linewidth=lw,
        facecolor=fc,
        edgecolor=ec,
    )
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, color="#1F2933")
    return box


def arrow(ax, start, end, color="#57606A"):
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=13,
            linewidth=1.2,
            color=color,
            shrinkA=8,
            shrinkB=8,
        )
    )


def generate_figure_1() -> Path:
    plt.rcParams.update({"font.family": "Arial", "font.size": 9.5, "axes.linewidth": 0.8})
    fig, ax = plt.subplots(figsize=(7.2, 3.15))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    palette = ["#E8F1FA", "#F5EFE6", "#EAF5EC", "#F3ECF8", "#F7F1D7", "#E9F4F4"]
    labels = [
        "Endpoint\nregistry",
        "Frozen\nsplits",
        "Molecular\nviews",
        "Candidate\nexperts",
        "Validation\nselection",
        "Locked\ntest audit",
    ]
    xs = [0.035, 0.19, 0.345, 0.50, 0.655, 0.81]
    for i, (x, label) in enumerate(zip(xs, labels)):
        draw_box(ax, (x, 0.58), (0.135, 0.20), label, fc=palette[i], fontsize=8.6)
        if i < len(xs) - 1:
            arrow(ax, (x + 0.135, 0.68), (xs[i + 1], 0.68))

    lower = [
        ("Data lock", "No test-set edits"),
        ("Split realism", "Scaffold audits\nLow-sim bins"),
        ("Reproducibility", "Seed-level CSVs\nSource data"),
    ]
    x0 = 0.12
    for i, (title, body) in enumerate(lower):
        x = x0 + i * 0.28
        draw_box(ax, (x, 0.23), (0.20, 0.15), f"{title}\n{body}", fc="#FFFFFF", ec="#A7B0BA", fontsize=7.3)
        arrow(ax, (x + 0.11, 0.58), (x + 0.11, 0.38), color="#8A949E")

    ax.text(0.05, 0.90, "FZYC-Mol workflow", fontsize=13, weight="bold", color="#111827")
    ax.text(
        0.05,
        0.84,
        "Predefined data, candidates and validation rules are fixed before the final test audit.",
        fontsize=8.8,
        color="#4B5563",
    )
    base = FIG_DIR / "fig1_workflow_only"
    save_figure(fig, base)
    plt.close(fig)
    return base.with_suffix(".png")


def generate_figure_2() -> Path:
    plt.rcParams.update({"font.family": "Arial", "font.size": 9.5, "axes.linewidth": 0.8})
    fig, ax = plt.subplots(figsize=(7.2, 3.85))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    ax.text(0.05, 0.92, "Selector, gate and output evidence", fontsize=13, weight="bold", color="#111827")
    ax.text(
        0.05,
        0.86,
        "Evidence modules decide whether a candidate is accepted, rejected or retained as a negative result.",
        fontsize=8.8,
        color="#4B5563",
    )

    draw_box(ax, (0.06, 0.58), (0.23, 0.20), "Validation selector\nrank, regret,\noptimism gap", fc="#E8F1FA")
    draw_box(ax, (0.39, 0.58), (0.23, 0.20), "AD/UQ gate\nTanimoto,\nensemble spread", fc="#EAF5EC")
    draw_box(ax, (0.72, 0.58), (0.21, 0.20), "Output decision\naccept, retain,\nor reject", fc="#F5EFE6")
    arrow(ax, (0.29, 0.66), (0.39, 0.66))
    arrow(ax, (0.62, 0.66), (0.72, 0.66))
    arrow(ax, (0.825, 0.58), (0.825, 0.46), color="#8A949E")

    evidence = [
        ("Performance", "ROC/PR\nRMSE/MAE\npaired seeds"),
        ("Reliability", "Calibration\nrisk-coverage\nconformal coverage"),
        ("Chemistry", "Motif/fragment\nsupport, effect\np/FDR"),
        ("Boundary", "Failure cases\nunaccepted\ncandidates"),
    ]
    xs = [0.06, 0.29, 0.52, 0.75]
    for x, (title, body) in zip(xs, evidence):
        draw_box(ax, (x, 0.14), (0.18, 0.22), f"{title}\n{body}", fc="#FFFFFF", ec="#A7B0BA", fontsize=7.8)
        arrow(ax, (x + 0.09, 0.43), (x + 0.09, 0.36), color="#C0C7CF")

    ax.plot([0.15, 0.85], [0.43, 0.43], color="#C7CDD3", lw=1.2)
    ax.text(0.50, 0.465, "frozen evidence package", ha="center", va="center", fontsize=8.5, color="#4B5563")
    base = FIG_DIR / "fig2_selector_gate_output_evidence"
    save_figure(fig, base)
    plt.close(fig)
    return base.with_suffix(".png")


def generate_risk_coverage_figure() -> Path:
    expanded = pd.read_csv(ROOT / "reports" / "uncertainty_ad_expanded" / "risk_coverage.csv")
    tdc = pd.read_csv(ROOT / "reports" / "uncertainty_ad_tdc_admet" / "risk_coverage.csv")
    wanted = [
        ("bbbp", "BBBP", expanded),
        ("clintox", "ClinTox", expanded),
        ("tdc_caco2_wang", "Caco2_Wang", tdc),
        ("tdc_pgp_broccatelli", "Pgp_Broccatelli", tdc),
    ]
    fig, axes = plt.subplots(2, 2, figsize=(7.2, 5.0), sharex=True)
    colors = ["#386CB0", "#DD8452", "#55A868", "#C44E52"]
    for ax, (dataset, title, frame), color in zip(axes.ravel(), wanted, colors):
        sub = frame[frame["dataset"].eq(dataset)].copy()
        metric = sub["risk_metric"].iloc[0]
        agg = sub.groupby("coverage")["risk"].agg(["mean", "std"]).reset_index()
        x = agg["coverage"].to_numpy()
        y = agg["mean"].to_numpy()
        sd = agg["std"].fillna(0).to_numpy()
        ax.plot(x, y, marker="o", color=color, linewidth=1.7, markersize=4)
        ax.fill_between(x, np.maximum(0, y - sd), y + sd, color=color, alpha=0.14, linewidth=0)
        ax.set_title(title, fontsize=10.5, weight="bold")
        ax.grid(True, color="#E5E7EB", linewidth=0.6, alpha=0.8)
        ax.set_ylabel("Error rate" if metric == "error_rate" else metric.upper())
        ax.spines[["top", "right"]].set_visible(False)
    for ax in axes[-1]:
        ax.set_xlabel("Coverage retained")
    fig.suptitle("Risk-coverage curves for representative reliability endpoints", fontsize=13, weight="bold", y=0.995)
    fig.tight_layout(rect=[0, 0, 1, 0.96])
    base = FIG_DIR / "fig11_risk_coverage_bbbp_clintox_caco2_pgp"
    save_figure(fig, base)
    plt.close(fig)
    return base.with_suffix(".png")


def cell_text(cell) -> str:
    return re.sub(r"\s+", " ", cell.text.strip())


def table_matrix(table) -> list[list[str]]:
    return [[cell_text(cell) for cell in row.cells] for row in table.rows]


def joined(row: list[str], cols: list[int]) -> str:
    parts = [row[i] for i in cols if i < len(row) and row[i]]
    return "; ".join(parts)


def simple_compress(rows: list[list[str]], headers: list[str], groups: list[list[int]]) -> tuple[list[str], list[list[str]]]:
    out_rows = []
    for row in rows[1:]:
        out_rows.append([joined(row, group) for group in groups])
    return headers, out_rows


CLINTOX_P80 = "0.370 ± 0.132"
CLINTOX_P90 = "P≥0.90: 0.232 ± 0.144"


def compress_table_16(rows: list[list[str]]) -> tuple[list[str], list[list[str]]]:
    headers = ["Endpoint/来源", "阳性率/模型", "ROC/PR", "Brier/ECE", "Recall@P≥0.80", "富集/说明"]
    out = []
    for row in rows[1:]:
        endpoint = row[0]
        fixed = f"{CLINTOX_P80}; {CLINTOX_P90}" if endpoint.lower() == "clintox" else "未计算"
        out.append(
            [
                joined(row, [0, 1]),
                joined(row, [2, 3, 4]),
                joined(row, [5, 6]),
                joined(row, [7, 8]),
                fixed,
                f"EF1 {row[9]}" if len(row) > 9 and row[9] else "",
            ]
        )
    return headers, out


def compress_table_30(rows: list[list[str]]) -> tuple[list[str], list[list[str]]]:
    headers = ["数据集/阳性率", "ROC-AUC/PR-AUC", "Brier/ECE", "Recall@P≥0.80", "EF1/EF5"]
    out = []
    for row in rows[1:]:
        endpoint = row[0].lower()
        fixed = f"{CLINTOX_P80}; {CLINTOX_P90}" if endpoint == "clintox" else "未计算"
        out.append([joined(row, [0, 1]), joined(row, [2, 3]), joined(row, [4, 5]), fixed, joined(row, [6, 7])])
    return headers, out


def compressed_table_content(index: int, rows: list[list[str]]) -> tuple[list[str], list[list[str]]] | None:
    configs: dict[int, tuple[list[str], list[list[int]]]] = {
        0: (["数据集", "来源/任务", "n/阳性率", "主指标", "划分", "seeds"], [[0], [1, 2], [3, 4], [5], [6], [7]]),
        2: (["数据集", "任务/指标", "验证选择器", "最终保留", "最强对照/观测最优", "解释"], [[0], [1, 2], [3], [4, 5], [6, 7], [8]]),
        3: (["数据集", "任务/指标", "当前", "整合后", "变化", "补救/模型"], [[0], [1, 2], [3], [4], [5], [6, 7]]),
        4: (["数据集", "任务/指标", "当前/Rescue", "重构候选", "变化", "保留/模型", "最终模型"], [[0], [1, 2], [3], [4], [5], [6, 7], [8]]),
        5: (["数据集", "任务/指标", "原策略", "融合候选", "变化", "保留/模型", "文献信号"], [[0], [1, 2], [3], [4], [5], [6, 7], [8]]),
        7: (["Endpoint", "任务/指标", "FZYC/保留", "随机到scaffold", "平均惩罚"], [[0], [1, 2], [3, 4, 5], [6, 7], [8]]),
        8: (["Endpoint", "任务/指标", "原策略", "外部融合候选", "变化", "来源/模型"], [[0], [1, 2], [3], [4], [5], [6, 7]]),
        9: (["Endpoint", "任务/指标", "原模型/候选", "原结果/保留", "变化", "说明"], [[0], [1, 2], [3, 4], [5, 6], [7], []]),
        10: (["Endpoint", "任务/指标", "NN/粗糙度", "粗糙度带", "性能变化", "来源"], [[0], [1, 2], [3, 4], [5], [6], [7]]),
        11: (["Endpoint", "任务/指标", "旧基线/候选", "保留来源", "增益", "粗糙度/结论"], [[0], [1, 2], [3, 4], [5], [6], [7, 8]]),
        13: (["来源/数据集", "任务/指标", "Random", "Scaffold", "Structure", "总变化"], [[0, 1], [2, 3], [4], [5], [6], [7]]),
        18: (["Endpoint", "任务/指标", "当前/固定策略", "Δ", "来源/模型"], [[0], [1, 2], [3, 4], [5], [6, 7]]),
        21: (["候选池", "规模/覆盖", "中位候选", "排名一致性", "风险单元"], [[0], [1, 2, 3], [4], [5, 6], [7]]),
        22: (["统计范围", "样本/数量", "平均变化/CI", "p值", "win/tie/loss/解释"], [[0], [1, 2], [3, 4], [5, 6], [7, 8]]),
        25: (["来源/n", "Spearman", "Top-1/Top-3", "负相关", "gap/regret"], [[0, 1], [2], [3, 4], [5], [6, 7]]),
        27: (["来源/任务/bin", "n/相似度", "不确定性", "高误差富集", "性能/校准"], [[0, 1, 2], [3, 4], [5], [6], [7]]),
        29: (["任务/seeds", "cliff pairs/Δy", "gap Spearman", "方向准确率", "gap MAE/相似度"], [[0, 1], [2, 3], [4], [5], [6, 7]]),
        31: (["任务/覆盖", "n/经验覆盖", "集合/单例", "区间宽度"], [[0, 1], [2, 3, 4], [5, 6], [7]]),
        32: (["案例/数据集/类型", "指标", "修改前/后", "Δ", "解释"], [[0, 1, 2], [3], [4, 5], [6], [7]]),
        34: (["案例/名称", "数据集/类型", "指标/证据", "修改前/后/Δ", "模型/风险信号", "解释"], [[0, 1], [2, 3], [4], [5, 6, 7], [8], [9]]),
    }
    if index == 16:
        return compress_table_16(rows)
    if index == 30:
        return compress_table_30(rows)
    if index in configs:
        headers, groups = configs[index]
        return simple_compress(rows, headers, groups)
    return None


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        if edge_data is None:
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), edge_data.get("val", "single"))
            element.set(qn("w:sz"), str(edge_data.get("sz", 8)))
            element.set(qn("w:space"), str(edge_data.get("space", 0)))
            element.set(qn("w:color"), edge_data.get("color", "000000"))


def apply_three_line_style(table) -> None:
    border = {"val": "single", "sz": 8, "space": 0, "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, 8.0)
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=border, bottom=border)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=border)


def replace_table(doc: Document, old_table, headers: list[str], rows: list[list[str]]):
    new_table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for j, header in enumerate(headers):
        new_table.cell(0, j).text = header
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            new_table.cell(i, j).text = value
    apply_three_line_style(new_table)
    old_table._tbl.addprevious(new_table._tbl)
    old_table._element.getparent().remove(old_table._element)
    return new_table


def update_and_compress_tables(doc: Document) -> None:
    original_tables = list(doc.tables)
    for index, table in enumerate(original_tables):
        if len(table.columns) <= 7:
            continue
        rows = table_matrix(table)
        content = compressed_table_content(index, rows)
        if content is not None:
            headers, out_rows = content
            replace_table(doc, table, headers, out_rows)
    for table in doc.tables:
        apply_three_line_style(table)


def insert_motif_fdr_table(doc: Document) -> None:
    if any("基序/片段富集的支持度" in p.text for p in doc.paragraphs):
        return
    summary = pd.read_csv(REPORT_DIR / "motif_fragment_support_fdr_summary.csv")
    summary = summary[["dataset", "kind", "feature", "n", "feature_positive_rate", "delta_positive_rate", "p_value", "fdr_q"]]
    summary = summary.head(12)
    target = next(p for p in doc.paragraphs if p.text.strip().startswith("基序/片段解释性分析"))
    caption = insert_paragraph_after(target, "表 24. 基序/片段富集的支持度、效应量与 FDR 证据。", size=10)
    table = doc.add_table(rows=1 + len(summary), cols=7)
    headers = ["数据集", "片段/骨架", "n", "阳性率/基线差", "p值", "FDR q", "解释边界"]
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
    for i, row in enumerate(summary.itertuples(index=False), start=1):
        effect = f"{row.feature_positive_rate:.3f}; Δ={row.delta_positive_rate:+.3f}"
        q_note = "FDR支持" if row.fdr_q <= 0.10 else "探索性"
        values = [
            str(row.dataset),
            f"{row.kind}: {row.feature}",
            str(int(row.n)),
            effect,
            f"{row.p_value:.2e}",
            f"{row.fdr_q:.2e}",
            q_note,
        ]
        for j, value in enumerate(values):
            table.cell(i, j).text = value
    apply_three_line_style(table)
    caption._p.addnext(table._tbl)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("表 24. 定向改进案例"):
            set_para_text(paragraph, "表 25. 定向改进案例。", size=10)
            break


def update_text_and_figures(doc: Document, fig1: Path, fig2: Path, fig11: Path) -> None:
    add_picture_to_paragraph(doc.paragraphs[30], fig1)
    set_para_text(doc.paragraphs[31], "图 1. FZYC-Mol 的冻结实验工作流：数据登记、划分锁定、视图构建、候选专家池、验证集选择和锁定测试报告。")
    add_picture_to_paragraph(doc.paragraphs[32], fig2)
    set_para_text(doc.paragraphs[33], "图 2. FZYC-Mol 的选择器、适用域门控和输出证据：验证排名、风险门控、性能/校准、风险覆盖、片段统计和失败案例共同决定最终保留。")
    set_para_text(
        doc.paragraphs[34],
        "图 1 仅展示实验工作流，强调所有数据划分、候选池和选择规则在测试集报告前锁定；图 2 则展示模型治理层，说明选择器、适用域门控、可靠性输出和化学解释如何共同形成接受、保留或拒绝候选的证据链。",
    )

    add_picture_to_paragraph(doc.paragraphs[154], fig11)
    set_para_text(doc.paragraphs[155], "图 11. BBBP、ClinTox、Caco2_Wang 和 Pgp_Broccatelli 的风险-覆盖曲线。")
    set_para_text(
        doc.paragraphs[158],
        "可靠性结果显示，单一不确定性分数难以覆盖所有错误类型。图 11 明确给出 BBBP、ClinTox、Caco2_Wang 和 Pgp_Broccatelli 的 risk-coverage 曲线：当低风险样本优先保留时，BBBP、ClinTox 和 Pgp_Broccatelli 的错误率随覆盖率下降而降低，Caco2_Wang 的 RMSE 曲线则提示回归外推收益较温和。集成标准差和错误模型对模型不一致性敏感，反向 Tanimoto 距离与重构误差更偏向适用域和结构新颖性，错误-适用域混合指标将两者组合后更适合作为实际使用时的风险标记。",
    )

    clin_para = doc.paragraphs[163]
    if "固定精度召回" not in clin_para.text:
        insert_paragraph_after(
            clin_para,
            f"针对审稿人点名的 ClinTox 固定精度审计，本文在与表 S17 一致的 consensus_strict_core_multifp fast 预测上计算 recall at fixed precision：在 precision≥0.80 时，五个 seed 的召回率为 {CLINTOX_P80}；在更严格的 precision≥0.90 时为 {CLINTOX_P90}。因此，ClinTox 的高 ROC-AUC 不应被解释为阳性样本均被充分检出，而应与 PR-AUC、Brier、ECE、fixed-precision recall 和样本级假阴性案例共同报告。",
            size=10,
        )

    motif_para = next(p for p in doc.paragraphs if p.text.strip().startswith("基序/片段解释性分析"))
    set_para_text(
        motif_para,
        "基序/片段解释性分析将模型行为连接到可识别的化学子结构。本文不再只描述候选片段，而是报告最小支持度、阳性率差异、Fisher 精确检验 p 值和 Benjamini-Hochberg FDR q 值。BBBP 中 N-连接片段和羰基/内酰胺相关 BRICS 片段与穿透性下降显著相关；BACE 中若干疏水芳香/卤代片段与阳性标签富集相关；ClinTox 中哌嗪/含氮片段、芳香片段和羰基片段显示阳性富集。该证据仍应解释为关联性解释，而不是因果机制证明。",
    )

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("本研究也有明确局限"):
            set_para_text(
                paragraph,
                "本研究也有明确局限。首先，验证-测试排名审计和 nested validation 共同表明，验证集治理可以避免测试集事后挑选，但仍存在验证集选择偏差，不能保证测试最优；因此所有小幅增益都需要与 regret、optimism gap、Top-3 命中和负结果一起解释。第二，收益具有明显终点异质性，BBBP、ClinTox、HIA、Pgp 等终点的增益较小，FreeSolv 仍落后于观测最佳 Chemprop 候选。第三，基序归因和片段富集已补充 support、effect size、p 值和 FDR q 值，但这些统计仍属于关联证据，不应被解读为未经湿实验验证的因果机制。第四，当前没有湿实验验证，ChemBERTa 与 MoLFormer 主要以冻结编码器形式使用，Polaris 与 OpenADMET 的完整官方挑战流程尚未完全纳入。",
            )
            break


def update_verified_references(doc: Document) -> None:
    fixes = {
        "[16]": "[16] Hong H, Wu X, Sun H, et al. A hierarchical interaction message net for accurate molecular property prediction. Communications Chemistry, 2026, 9: 150. DOI: 10.1038/s42004-026-01922-x.",
        "[44]": "[44] Liu H, Zhu B, Nie S, et al. Advancing ADMET prediction through multiscale fragment-aware pretraining with MSformer-ADMET. Briefings in Bioinformatics, 2025, 26(5): bbaf506. DOI: 10.1093/bib/bbaf506.",
        "[49]": "[49] Kamuntavicius G, Paquet T, Bastas O, et al. Benchmarking ML in ADMET predictions: the practical impact of feature representations in ligand-based models. Journal of Cheminformatics, 2025, 17: 108. DOI: 10.1186/s13321-025-01041-0.",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for prefix, replacement in fixes.items():
            if text.startswith(prefix):
                set_para_text(paragraph, replacement)


def verify_layout(doc_path: Path) -> pd.DataFrame:
    doc = Document(str(doc_path))
    rows = []
    for i, table in enumerate(doc.tables):
        rows.append({"table_index": i, "rows": len(table.rows), "cols": len(table.columns)})
    return pd.DataFrame(rows)


def main() -> None:
    fig1 = generate_figure_1()
    fig2 = generate_figure_2()
    fig11 = generate_risk_coverage_figure()

    src = find_source_docx()
    out = src.with_name(src.name.replace("-5.docx", "-6.docx"))
    doc = Document(str(src))

    update_text_and_figures(doc, fig1, fig2, fig11)
    update_and_compress_tables(doc)
    insert_motif_fdr_table(doc)
    update_verified_references(doc)
    doc.save(str(out))

    layout = verify_layout(out)
    layout.to_csv(REPORT_DIR / "chugao6_table_layout_audit.csv", index=False)
    wide = layout[layout["cols"] > 7]
    report = [
        "# 初稿-6 second-pass revision audit",
        "",
        f"- Source: `{src}`",
        f"- Output: `{out}`",
        f"- Figure 1: `{fig1}`",
        f"- Figure 2: `{fig2}`",
        f"- Figure 11: `{fig11}`",
        f"- Tables with >7 columns after compression: {len(wide)}",
        f"- ClinTox Recall@P≥0.80: {CLINTOX_P80}; {CLINTOX_P90}",
        f"- Motif/FDR source: `{REPORT_DIR / 'motif_fragment_support_fdr_summary.csv'}`",
        "- 2025/2026 references checked online; refs [16], [44], [49] were format-normalized.",
    ]
    if not wide.empty:
        report.append("")
        report.append(wide.to_markdown(index=False))
    (REPORT_DIR / "chugao6_revision_audit.md").write_text("\n".join(report), encoding="utf-8")
    print(out)
    print(f"wide_tables={len(wide)}")


if __name__ == "__main__":
    main()
