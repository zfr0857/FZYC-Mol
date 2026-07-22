from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "小论文-2.docx"
OUTPUT = ROOT / "output" / "小论文-3.docx"
FIG = ROOT / "output" / "小论文-3_图表包" / "figures"


def find_prefix(doc: Document, prefix: str):
    matches = [p for p in doc.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting with {prefix!r}; found {len(matches)}")
    return matches[0]


def replace_text(paragraph, text: str) -> None:
    run_properties = copy.deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs and paragraph.runs[0]._r.rPr is not None else None
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


def replace_prefix(doc: Document, prefix: str, text: str) -> None:
    replace_text(find_prefix(doc, prefix), text)


def insert_after(paragraph, text: str):
    doc = paragraph._parent
    created = doc.add_paragraph(text, style=paragraph.style)
    paragraph._p.addnext(created._p)
    return created


def replace_image_before_caption(doc: Document, caption_prefix: str, image_path: Path) -> None:
    caption = find_prefix(doc, caption_prefix)
    paragraphs = list(doc.paragraphs)
    index = next(i for i, p in enumerate(paragraphs) if p._p is caption._p)
    drawing = next(p for p in reversed(paragraphs[:index]) if p._p.xpath(".//w:drawing"))
    for child in list(drawing._p):
        if child.tag != qn("w:pPr"):
            drawing._p.remove(child)
    drawing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    drawing.paragraph_format.first_line_indent = Cm(0)
    drawing.add_run().add_picture(str(image_path), width=Cm(16.2))


def build() -> Path:
    doc = Document(SOURCE)

    replace_prefix(doc, "FZYC-Mol：", "FZYC-Mol：轻量候选池扩张下分子性质模型选择的验证治理与风险审计")
    replace_prefix(
        doc,
        "分子性质预测常在有限验证信息上比较",
        "分子性质预测常在有限验证信息上比较不断扩张的模型候选，但候选增加既可能提高可达到的预测上界，也可能放大验证排序的不确定性。FZYC-Mol 将冻结候选登记、验证侧选择、外层审计和可靠性证据组织为可追踪的治理协议。对 9 个终点开展的随机顺序、随机子集和家族平衡审计显示，随机子集从 4 扩至 32 个候选时，机会校正命中率由 0.723 降至 0.142，全 32 池固定分母遗憾由 0.090 增至 0.214。在 3 外层×3 内层×5 重复的确认性嵌套验证中，机会校正命中率由 0.881 降至 0.240，固定分母遗憾由 0.049 增至 0.171。随机排序负对照在各池规模的机会校正命中率均约为 0，而真实流程仍显著高于机会水平。仅由 one-SE 集合大小、内折选择频率和内折波动构成的验证侧选择风险与外层遗憾呈正相关（Spearman ρ=0.235，端点 bootstrap 95% CI 0.012–0.378）；低风险一半单元的平均遗憾为 0.098，高风险一半为 0.148。留一端点迁移中，validation-best 在 9 次训练端点选择中均胜出，但仅在 7 个留出端点达到端点内最优。结果形成了候选扩张、排序失真、风险预警与接受/保留/拒绝决策的闭环，同时表明冻结规则不能保证普遍性能提升。确认性结论限于公开离线基准中的轻量 Morgan-512 候选池，不外推至尚未在相同外层划分统一重训的图模型和化学语言模型。",
    )
    contributions = [
        ("FZYC-Mol 以冻结候选登记", "FZYC-Mol 以冻结候选登记、随机化候选池和重复嵌套选择审计评价分子性质模型选择，并用全 32 池固定分母遗憾与机会校正排序保真度区分真实选择信号和候选数带来的机械机会变化。"),
        ("候选池规模与组成通过重复", "随机排序负对照在保持验证效用和外层效用边际分布的同时破坏二者对应关系，为扩池后的排序信息衰减提供显式零假设。"),
        ("版本化决策卡把终点级选择", "验证侧选择风险把 one-SE 歧义、内折选择频率和内折波动连接到外层遗憾与风险-覆盖曲线；留一端点迁移进一步检验治理规则的跨端点边界。"),
    ]
    for prefix, text in contributions:
        replace_prefix(doc, prefix, text)
    replace_prefix(doc, "关键词：", "关键词：分子性质预测；模型选择治理；候选池扩张；选择风险；随机排序负对照；固定分母遗憾；重复嵌套验证")

    replace_prefix(
        doc,
        "近期研究已分别覆盖现实挑战",
        "近期研究已分别覆盖现实挑战、强表示和可靠性估计。Zhao 等系统比较了稀缺、分布外、类别不平衡、bRo5 和活性悬崖场景[5]；DCPM-ADMET、KROVEX 与 MolGramTreeNet 进一步发展了预训练、指纹、描述符和图表示融合[6-8]。OOD 表征学习、主动学习、适用域边界和不确定性基准也在快速推进[9-13]，数据漂移研究进一步显示化学距离与模型分歧提供互补的不确定性信号[33]。这些工作说明困难场景、多模态融合、适用域和不确定性本身均不能作为本文的首创主张。",
    )
    replace_prefix(doc, "本研究聚焦一个可检验问题", "本研究聚焦一个可检验问题：在验证信息保持不变时，候选池扩张何时降低排序保真度，以及这种风险能否在读取外层标签之前被验证侧诊断量识别。FZYC-Mol 因而是冻结登记、验证侧选择、风险预警和外层审计组成的治理协议，而不是新的主干预测网络。")
    replace_prefix(doc, "分析明确区分三类证据", "分析明确区分四类证据：冻结验证流程给出可实施选择；测试事后最优只定义遗憾参照；随机排序负对照定义无排序信息时的零基线；未通过门控、运行失败或跨端点不能迁移的策略用于界定边界。该分层防止将探索性收益或测试集机会转写为确认性结论。")
    replace_prefix(doc, "主要假设不是候选增加必然降低", "主要假设不是候选增加必然降低点预测性能，而是搜索自由度扩大后，验证排序与外层效用的一致性会下降。本文以三种随机化候选池控制和置换负对照检验排序信号，以 3×3×5 重复嵌套验证确认遗憾变化，以验证侧选择风险和留一端点迁移连接预警、保留与拒绝决策，并用 AutoGluon、TDC、风险-覆盖、标签条件保形和 MoleculeACE 检验结论边界。")

    anchor = find_prefix(doc, "AutoGluon-Tabular 使用相同")
    p1 = insert_after(anchor, "2.5.1 选择风险、随机排序负对照与跨端点迁移")
    p2 = insert_after(p1, "选择风险仅由验证侧可观测量构成。对每个终点-重复-外层单元和池规模 K，歧义分量定义为 (|A¹ˢᴱ|−1)/(K−1)，不稳定分量定义为 1−selection frequency，波动分量定义为内折效用标准差在同一终点和 K 内的百分位。三项取等权均值，形成 0–1 的 selection-risk score；外层固定分母遗憾仅在分数冻结后用于审计。风险-覆盖曲线按该分数从低到高保留外层单元。")
    p3 = insert_after(p2, "随机排序负对照在每个冻结外层单元内独立置换内层验证效用与外层效用的候选对应关系，同时保持两侧边际分布、候选数和外层效用范围不变。K=4/8/16/32 各进行 1,000 次置换，并重新计算机会校正 Top-3 命中率与全 32 池固定分母遗憾。该分析检验真实排序是否优于无对应信息的选择器，而不产生新的候选或重新训练模型。")
    insert_after(p3, "治理规则迁移采用 leave-one-endpoint-out（LOEO）设计。每次留出一个终点，仅用其余八个终点的冻结外层遗憾选择平均遗憾最低的规则，再在留出终点评价。留出终点内的最优规则仅作为事后 oracle 参照，不参与规则选择。")

    captions = {
        "图 1 |": "图 1 | FZYC-Mol 总体工作流。任务协议、分子视图、专家池、验证治理和证据输出在最终测试前形成单向冻结流程；底部列出外部验证、校准/不平衡、Top-K 压力、稳健性、困难子集和负结果等实验层。该图为概念图，不构成性能结果。",
        "图 2 |": "图 2 | FZYC-Mol 模块结构。多视图分子输入进入图、树模型和冻结 embedding 专家；验证选择器、选择风险和风险感知融合共同形成接受、保留或拒绝决策。历史候选仅在共享冻结划分时进入确认性池，虚线框表示当前未满足该条件的候选。",
        "图 3 |": "图 3 | 随机化候选池与随机排序负对照。a，三种随机化模式的全 32 池固定分母遗憾。b，真实重复嵌套选择与 1,000 次置换负对照的机会校正 Top-3 命中率。c，三种随机化模式的 MRR。d，真实流程与置换负对照的固定分母遗憾。区间按终点聚类或置换分布给出。",
        "图 4 |": "图 4 | 重复嵌套选择与验证侧风险。a,b，3×3×5 重复嵌套验证的固定分母遗憾和机会校正命中率。c，仅用 one-SE 歧义、内折选择频率和内折波动构造的选择风险与外层遗憾。d，按选择风险从低到高保留外层单元的风险-覆盖曲线。测试标签仅用于分数冻结后的外层评价。",
        "图 6 |": "图 6 | TDC 门控有效性审计。a，按 gate_category 汇总的晋级与保留终点数。b，22 个终点的方向归一化测试变化及三随机种子区间；蓝色为晋级，灰色为保留。宽置信区间单独标记为 inconclusive。",
        "图 7 |": "图 7 | 预测风险与选择风险的双层选择性审计。a-c，BBBP、ClinTox 和 Caco2 的逐样本风险-覆盖曲线，真实误差排序仅定义 oracle 下界。d，按验证侧 selection-risk score 保留外层选择单元后的平均固定遗憾。两类风险的统计单位不同，不作直接数值比较。",
        "图 8 |": "图 8 | 标签条件保形审计。a，分类总体及类别条件覆盖。b，回归经验覆盖。c，分类预测集合大小与 pooled fallback 率。d，回归区间宽度及按训练标签 SD 标准化的宽度。标称覆盖为 80%、90% 和 95%。",
        "图 10 |": "图 10 | 治理、候选组成与跨端点迁移。a，完整 32 候选池上的四种治理规则。b，冻结 one-SE 下的候选家族移除。c，终点×规则遗憾矩阵，橙点表示端点内事后最优。d，LOEO 训练端点选择的规则与留出端点 oracle 的遗憾差距。",
    }
    for prefix, text in captions.items():
        replace_prefix(doc, prefix, text)

    negative_anchor = find_prefix(doc, "上述随机化分析使用既有冻结预测")
    n1 = insert_after(negative_anchor, "随机排序负对照进一步表明，机会校正指标具有正确零基线。K=4/8/16/32 时，置换后的平均机会校正命中率为 0.005/−0.002/−0.002/0.000，而真实重复嵌套流程为 0.881/0.550/0.334/0.240。")
    insert_after(n1, "真实流程的固定分母遗憾为 0.049/0.131/0.139/0.171，低于置换负对照的 0.452/0.387/0.319/0.314。候选池扩大后仍保留高于随机的排序信息，但真实与零基线的间隔随 K 缩小。")

    risk_anchor = find_prefix(doc, "AutoGluon 相对固定单模型")
    r1 = insert_after(risk_anchor, "验证侧选择风险在 540 个外层单元中与固定分母遗憾呈正相关（Spearman ρ=0.235，端点 bootstrap 95% CI 0.012–0.378；图 4c）。该相关为诊断性而非因果性，但其输入在外层标签读取前即可获得。")
    insert_after(r1, "按选择风险中位数二分时，低风险一半单元的平均遗憾为 0.098，高风险一半为 0.148；风险-覆盖曲线在保留低风险单元时保持较低遗憾（图 4d）。分池分析的区间较宽，因此正文只主张总体预警价值，不声称每个 K 均独立显著。")

    transfer_anchor = find_prefix(doc, "两类消融必须分开解释")
    t1 = insert_after(transfer_anchor, "LOEO 迁移提供了更严格的跨端点检验。其余八个端点在 9 次留一分析中均选择 validation-best，但该规则只在 7 个留出端点达到端点内最低遗憾；FreeSolv 偏好 one-SE 低方差，HIA 偏好 one-SE 低成本。")
    insert_after(t1, "LOEO 规则的平均留出遗憾为 0.129，端点内 oracle 为 0.121。0.009 的平均差距不支持在当前九终点样本上学习复杂元选择器，却确认了单一治理规则仍存在端点特异失败。")

    replace_prefix(doc, "主文数字、表和更新后的图", "主文数字、表和图 1–10 均由冻结 source data 或概念图脚本重建；新增选择风险、置换负对照和 LOEO 迁移分别保存逐单元、逐置换和逐留出端点结果。完整文件清单、SHA-256、冷启动日志和发布缺口见补充复现材料。")

    discussion_updates = [
        ("随机化候选池控制与 3×3×5", "随机化候选池、置换负对照与 3×3×5 重复嵌套验证共同表明，候选池扩张削弱了选择可信度，但并未把真实排序降至随机水平。机会校正命中率、MRR、固定分母遗憾和置换零基线给出了同向证据，因此结论既不依赖固定 Top-3 的机械机会变化，也不等同于候选越多性能必然越差。"),
        ("候选规模与候选组成均影响", "候选规模与候选组成共同影响效应大小。三种随机化模式在 K<32 时存在差异而在完整池中收敛，说明家族组成会改变局部难度；置换负对照则显示，候选池扩大还会缩小真实排序相对于无信息选择器的优势。"),
        ("没有一种治理规则在所有终点占优", "没有一种治理规则在所有终点占优。validation-best 的平均遗憾最低，LOEO 迁移也在每次训练端点集合中选择该规则，但 FreeSolv 与 HIA 的留出结果仍偏好不同的保守规则。规则冻结的价值是限制事后自由度并显式保存失败，而不是保证最低遗憾。"),
        ("可靠性结果也需要条件化解释", "可靠性结果需要区分预测风险与选择风险。逐样本 E-AURC、适用域和标签条件保形回答单个预测何时应降置信；验证侧 selection-risk score 回答一次候选选择何时更可能产生较大外层遗憾。二者使用不同统计单位，却共同服务于接受、保留和拒绝决策。"),
        ("MoleculeACE 的 17 个官方任务", "MoleculeACE、bRo5 与低相似度结果把选择失败定位到化学边界，近期 OOD 与数据漂移研究则说明距离和模型分歧通常只提供互补而非完备的风险证据[9-13,33]。因此，本文把这些信号接入治理闭环，而不将其表述为新的不确定性方法或部署充分条件。"),
        ("本研究的首要范围限制", "本研究的首要范围限制是确认性候选池仅含轻量 Morgan-512 模型。历史 Chemprop、GNN、ChemBERTa 和 MoLFormer 预测未共享冻结外层划分，不能被事后拼接为异质嵌套证据。其次，选择风险相关性为中等且分池区间较宽，LOEO 仅覆盖九个端点；此外，AutoGluon 只覆盖固定 CPU 树模型集合，尚无独立 ADMET 时间外盲测。"),
        ("综上，候选增加可提高", "综上，候选增加可提高可达到的预测上界，也会增加验证侧选择自由度。实践中应在测试前冻结候选登记和选择目标，以置换或其他负对照校准排序指标，按终点报告外层遗憾与选择风险，并把无法运行、未晋级和跨端点不能迁移的结果作为证据的一部分。"),
    ]
    for prefix, text in discussion_updates:
        replace_prefix(doc, prefix, text)

    replace_prefix(doc, "FZYC-Mol 将冻结候选登记、随机化", "FZYC-Mol 将冻结候选登记、随机化候选池、置换负对照、验证侧选择风险和外层审计组织为可复核的轻量候选治理流程。随机子集控制和 3×3×5 重复嵌套验证均显示，扩池伴随机会校正排序质量下降和固定分母遗憾上升；置换结果同时确认真实流程仍保留高于随机的排序信息。")
    replace_prefix(doc, "治理规则和候选家族消融未发现", "验证侧选择风险能够在外层标签读取前富集高遗憾单元；LOEO 迁移则表明 validation-best 虽在 7/9 个留出端点达到端点内最优，仍不能覆盖 FreeSolv 与 HIA 的端点特异偏好。治理规则和候选组成因此需要与终点级边界共同报告。")
    replace_prefix(doc, "最稳健的结论是：", "最稳健的结论是：候选池扩大增加了验证排序与选择的不确定性，冻结规则、负对照、选择风险和外层审计可以暴露并管理这种风险，但不能保证普遍性能提升。该结论限于公开离线基准和冻结的轻量 Morgan-512 候选池，不支持对异质重型模型、时间外 ADMET 数据或临床部署的外推。")

    replace_prefix(doc, "补充材料逐项提供", "补充材料逐项提供：数据清洗流与删除 reason；完整候选登记；100 次随机顺序/子集/家族平衡清单；3×3×5 重复嵌套逐折排名、遗憾和稳定性；1,000 次随机排序负对照；验证侧选择风险逐单元结果；LOEO 治理迁移；AutoGluon 三预算；TDC 门控；风险-覆盖、标签条件保形、MoleculeACE 和候选家族消融。")
    replace_prefix(doc, "更新主图均配套 source data", "图 1–10 均配套 PNG/SVG、source data 和生成脚本。选择闭环数值由 selection_closure_values.json 汇总；manuscript_values.json 继续承担既有结果的零差异重建。reproducibility_manifest.json 与 SHA256SUMS 用于校验文件，但公开仓库 release 和 Zenodo 归档仍待作者完成。")

    reference_anchor = find_prefix(doc, "[32]")
    insert_after(reference_anchor, "[33] Parrondo-Pizarro R, Lanini J, Rodriguez-Perez R. Uncertainty quantification in molecular machine learning for property predictions under data shifts. J Chem Inf Model. 2026;66:923-935. doi:10.1021/acs.jcim.5c02381.")

    figure_files = {
        "图 1 |": "fig01_overall_workflow.png",
        "图 2 |": "fig02_model_structure.png",
        "图 3 |": "fig03_candidate_pool_and_null.png",
        "图 4 |": "fig04_nested_selection_risk.png",
        "图 5 |": "fig05_moleculenet_rank_audit.png",
        "图 6 |": "fig06_tdc_gate_audit.png",
        "图 7 |": "fig07_prediction_and_selection_risk.png",
        "图 8 |": "fig08_conformal_audit.png",
        "图 9 |": "fig09_chemical_boundaries.png",
        "图 10 |": "fig10_governance_transfer.png",
    }
    for caption, filename in figure_files.items():
        replace_image_before_caption(doc, caption, FIG / filename)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
