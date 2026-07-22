from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path("D:/fzyc")
OUT = ROOT / "output"
SOURCE = max([p for p in OUT.glob("*.docx") if p.name.endswith("-15.docx")], key=lambda p: p.stat().st_mtime)
TARGET = OUT / "小论文-16.docx"
REPORT = OUT / "小论文-16_一区审稿人视角不足与修订说明.md"
AUDIT = OUT / "paper16_q1_reviewer_audit.json"


REPLACEMENTS = {
    "冻结验证揭示分子性质预测中候选池扩张的收益与选择损失": (
        "冻结验证量化分子性质预测中候选池扩张的收益与选择损失"
    ),
    "分子性质预测的最终模型通常从持续扩张的候选池中选出。": (
        "分子性质预测通常在不断扩大的候选模型、表征和调参方案中选择最终模型。"
        "候选扩张可以提高可达到的性能上界，但在验证信息固定时也会重复消耗排序信号，"
        "使外层测试结果同时受到真实表征收益和模型选择偏差的影响。"
    ),
    "本研究提出 FZYC-Mol（Frozen validation governance for molecular model selection）": (
        "本研究提出 FZYC-Mol（Frozen validation governance for molecular model selection），"
        "将候选登记、嵌套选择、策略冻结、外层审计和负结果记录整合为一个验证治理协议，"
        "而非新的预测主干网络。该协议在九个终点上执行 3 外层×3 内层×5 重复的冻结选择实验，"
        "并将公开面板、逐样本可靠性和化学迁移边界归入 reliability and chemical-boundary analyses"
        "（可靠性与化学边界分析），而非主排行榜证据。"
    ),
    "在确认性候选池扩张实验中，K=32 相对 K=4": (
        "在确认性候选池扩张实验中，K=32 相对 K=4 使完整池范围归一化选择损失增加 "
        "0.122（端点聚类 95% CI 0.072–0.175；精确 P=0.0078；Holm P=0.039），"
        "机会校正 Top-3 命中率下降 0.642。共享冻结划分的 12 候选多视图实验显示，"
        "validation-best 相对 Morgan-only 的实际兑现效用增益为 0.343（0.210–0.483；9/9 终点）。"
        "跨终点元风险仅支持探索性预警：严格留一端点验证的高遗憾 AUC 为 0.648，"
        "保留预测风险最低 50% 单元时平均遗憾降低 0.034（95% CI 0.020–0.047），"
        "但九个终点不足以建立通用元选择器。"
    ),
    "这些结果表明，候选池扩张在分子性质预测中同时提高可达到上界并增加选择损失。": (
        "这些结果表明，候选池扩张在分子性质预测中同时提高可达到上界并增加选择损失。"
        "冻结治理使这一权衡可被审计，并将可靠性与化学边界从平均性能叙事中分离出来。"
        "本文结论不外推到受授权限制的 TabPFN、九终点全量深度模型面板或时间外前瞻验证。"
    ),
    "基于这一定位，本研究将文献对照中的差异转化为可检查的扩展分析": (
        "近年及近半年分子机器学习研究共同强调，强基线、数据划分、OOD 泛化、类别不平衡、"
        "活性悬崖、bRo5 化学空间和不确定性评估需要在同一证据框架下报告[1,2,5–11,21–23,33–35]。"
        "据此，本文将文献要求落实为三类补强证据：代表性同划分强基线、逐样本 error-overlap "
        "和去重敏感性重跑。保形可靠性、活性悬崖、bRo5 与 TDC 分析用于界定适用边界，"
        "而非形成新的总排行榜。"
    ),
    "近半年文献进一步强化了这一分层设计。": (
        "这一分层也限定了本文的中心问题：在验证信息固定的条件下，候选池扩张是否会同时提高可达到性能上界、"
        "降低验证排序保真度并增加模型选择损失。围绕该问题，候选池扩张、随机排序负对照、信号恢复正对照、"
        "端点层配对推断和共享划分多视图验证构成确认性主证据；AutoGluon、跨端点元风险、TDC、逐样本风险、"
        "保形预测、MoleculeACE 和 bRo5 构成次要或探索性边界证据。"
    ),
    "本研究的中心主张是：": "",
    "除主效应外，本研究还报告真嵌套验证、种子敏感性": (
        "除主效应外，本文报告真嵌套验证、种子敏感性、统一消融、80/90/95 保形覆盖率、"
        "精确 Tanimoto 分箱、MoleculeACE 活性悬崖、低相似度失败样本和扩展失败案例等补充分析。"
        "强基线证据矩阵被限定为代表性压力测试：ESOL、BACE 和 ClinTox 分别覆盖回归、常规分类和"
        "少数类毒性终点；GNN、RDKit-RF、ChemBERTa/MoLFormer 冻结适配头完成完整 3 外层×3 内层×5 "
        "种子的内外层评估，Chemprop/D-MPNN 完成相同外层 folds 和 seeds 下的外层确认性训练以及全部 "
        "3 个 inner folds 的确认性重训。TabPFN 已安装，但因授权和运行时交互限制未能完成同划分预测导出；"
        "因此仅作为授权受限候选记录在状态表中，不作为完成性结果。"
    ),
    "error-overlap 审计显示，五个候选的错误集并非完全重合": (
        "error-overlap 审计显示，五个候选的错误集并非完全重合：10 个候选对的平均 Jaccard 重合为 "
        "0.189，范围为 0.057–0.449。去重敏感性分析也已从清洗审计推进到实际重跑：三个代表终点、"
        "三套策略和 135 个外层单元的最大平均效用变化为 0.022。这些结果支持一个审慎的比较原则："
        "当代复杂模型只有在同一冻结划分下完成预测导出、选择审计和负结果记录后，才宜进入主文效应叙事。"
    ),
    "FZYC-Mol 不替代预测模型，也不必然带来性能提升": (
        "FZYC-Mol 不替代预测模型，也不保证性能提升，且不提供可迁移到所有终点的元选择器。"
        "代表性三终点面板已完成 GNN、ChemBERTa/MoLFormer 冻结适配头、RDKit-RF 和 Chemprop/D-MPNN "
        "的 3×3×5 同划分评估，并新增验证选择、配对效应和 error-overlap 审计；但 TabPFN 未能产生"
        "同划分预测，九终点全量深度模型面板仍超出当前完成范围。因此，强基线结果应解释为边界压力测试，"
        "不能被解读为复杂模型已在所有终点上完成确认性统一重训。TDC 三种子结果只能反映种子变异和公开面板异质性，"
        "不能承担严格抽样推断。公开 release、Zenodo DOI 和第三方冷启动复跑仍需作为后续复现工作完成。"
    ),
    "近期研究也说明，本文的不足不应被简单归结为缺少某个单一更强模型。": (
        "近期研究提示，本文的主要证据缺口在于比较范围和前瞻复现，而非单一模型缺席。"
        "2026 年 ADMET 可靠性基准显示，TabPFNv2、预训练 GNN、AutoML 和传统模型在不同挑战下各有优势，"
        "并且活性悬崖仍是多类模型的共同弱点[5]。KROVEX 和 DCPM-ADMET 的结果支持表征异质性和多模态融合的价值[6,7]，"
        "但这些收益只有在相同划分、相同选择规则和明确消融下才可比较。FZYC-Mol 因此不排除复杂模型，"
        "而是要求复杂模型承担同样的冻结选择、逐样本导出、去重敏感性和边界审计成本。"
    ),
    "这些结果同时界定了后续扩展的优先级：第一": (
        "后续扩展应优先完成三项工作：将强基线从三个代表终点扩展到九终点全量面板；"
        "在解决 TabPFN 授权或运行环境限制后纳入同划分比较；完成公开 release、Zenodo DOI 和独立冷启动复跑。"
        "这些边界已在正文中明确限定，以避免将代表性扩展分析写成已完成的全量确证。"
    ),
    "对后续研究而言，最直接的采用方式不是替换现有预测器": (
        "对后续研究而言，FZYC-Mol 的直接用途是在新候选进入比较之前建立候选登记、终止规则、内层选择、"
        "外层审计和负结果归档。若未来将 Chemprop、GNN、化学语言模型和 TabPFN 从代表性三终点扩展到"
        "九终点全量同划分面板，并加入时间外 ADMET 盲测，该框架仍可作为扩展前的审计底座。"
    ),
}


def replace_paragraph(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def apply_revisions() -> dict:
    shutil.copy2(SOURCE, TARGET)
    doc = Document(str(TARGET))
    hits = {}

    for prefix, replacement in REPLACEMENTS.items():
        matched = []
        for idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip().startswith(prefix):
                matched.append(idx)
        hits[prefix] = matched
        if not matched:
            continue
        for idx in matched:
            if replacement:
                replace_paragraph(doc.paragraphs[idx], replacement)
            else:
                replace_paragraph(doc.paragraphs[idx], "")

    doc.save(str(TARGET))
    return hits


def audit_docx(hits: dict) -> dict:
    with ZipFile(TARGET) as zf:
        bad = zf.testzip()
    doc = Document(str(TARGET))
    text = "\n".join(p.text for p in doc.paragraphs)
    forbidden = ["揭示分子性质预测", "本研究的中心主张是：", "不应被简单归结", "复杂模型必须", "license/token"]
    checks = {
        "source": str(SOURCE),
        "target": str(TARGET),
        "zip_ok": bad is None,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "figures": len(doc.inline_shapes),
        "references": sum(1 for p in doc.paragraphs if p.text.strip().startswith("[")),
        "title_rewritten": doc.paragraphs[0].text.strip()
        == "冻结验证量化分子性质预测中候选池扩张的收益与选择损失",
        "abstract_boundary": "本文结论不外推到受授权限制的 TabPFN" in text,
        "intro_condensed": "近年及近半年分子机器学习研究共同强调" in text,
        "q1_scope_language": "代表性压力测试" in text and "比较范围和前瞻复现" in text,
        "tabpfn_formal_boundary": "因授权和运行时交互限制未能完成同划分预测导出" in text,
        "forbidden_hits": {term: term in text for term in forbidden if term in text},
        "replacement_hits": hits,
    }
    checks["passed"] = (
        checks["zip_ok"]
        and checks["title_rewritten"]
        and checks["abstract_boundary"]
        and checks["intro_condensed"]
        and checks["q1_scope_language"]
        and checks["tabpfn_formal_boundary"]
        and not checks["forbidden_hits"]
    )
    return checks


def write_report(checks: dict) -> None:
    lines = [
        "# 小论文-16：一区审稿人视角不足与修订说明",
        "",
        f"- 生成时间：{datetime.now().isoformat(timespec='seconds')}",
        f"- 输入稿件：`{SOURCE}`",
        f"- 输出稿件：`{TARGET}`",
        "",
        "## 一区审稿人可能提出的主要不足",
        "",
        "| 类别 | 审稿人可能质疑 | 本轮处理 |",
        "|---|---|---|",
        "| 主张范围 | 标题和摘要容易让人以为覆盖所有分子预测模型，但强基线仍是三终点代表性面板。 | 将标题由“揭示”改为“量化”，摘要和讨论明确不外推到 TabPFN、九终点全量深度模型面板或时间外验证。 |",
        "| 实验证据 | 现代强基线已有实质补强，但仍不是九终点全量确认性重训。 | 将强基线统一称为“代表性压力测试”，避免写成全量排行榜。 |",
        "| TabPFN 边界 | TabPFN 未能产出预测，若写得含糊会被认为选择性报告。 | 将“license/token”口语化表达改为“授权和运行时交互限制”，并说明不作为完成性结果。 |",
        "| 叙事聚焦 | 引言中近期文献和补强实验罗列较多，容易稀释核心问题。 | 合并文献定位段，收束到三类补强证据和一个中心问题。 |",
        "| 结果口吻 | “复杂模型必须”等表达带有规范性，结果段应主要报告观察。 | 改为“审慎的比较原则”，降低审稿人对过度主张的反感。 |",
        "| 讨论边界 | 局限段较长且重复，容易显得防御性。 | 将证据缺口概括为“比较范围和前瞻复现”，并列出后续优先级。 |",
        "| 写作风格 | 部分句子口语化或混用英文运行术语。 | 替换“license/token”等表达，减少口语化和过强措辞。 |",
        "",
        "## 仍然无法仅靠文字完全解决的缺口",
        "",
        "- 九终点全量 Chemprop/GNN/ChemBERTa/MoLFormer/TabPFN 同划分重训仍未完成。",
        "- TabPFN 仍没有同划分逐样本预测导出。",
        "- 时间外前瞻 ADMET 盲测、公开 release、Zenodo DOI 和第三方冷启动复跑仍是后续复现工作。",
        "- 图件复杂度和部分主图 panel 数量若目标期刊严格限制，仍建议在投稿排版前再做一次图版压缩。",
        "",
        "## 自动审计",
        "",
        "```json",
        json.dumps({k: v for k, v in checks.items() if k != "replacement_hits"}, ensure_ascii=False, indent=2),
        "```",
    ]
    REPORT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    hits = apply_revisions()
    checks = audit_docx(hits)
    AUDIT.write_text(json.dumps(checks, ensure_ascii=False, indent=2), encoding="utf-8")
    write_report(checks)
    print(json.dumps({"target": str(TARGET), "report": str(REPORT), "passed": checks["passed"]}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
