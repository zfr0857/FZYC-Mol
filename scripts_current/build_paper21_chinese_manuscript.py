from __future__ import annotations

import json
import os
from pathlib import Path
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from paper25_manuscript_content import (
    CHINESE_DISCUSSION_SECTIONS,
    CHINESE_INTRODUCTION_PARAGRAPHS,
    render_chinese_abstract,
)
from paper25_docx_svg import embed_svg_figures

ROOT=Path("D:/fzyc")
NEW=Path(os.environ.get("FZYC_ANALYSIS_OUT", ROOT/"output"/"paper21_final_reanalysis_20260713"))
MINOR=Path(os.environ.get("FZYC_MINOR_OUT", ROOT/"output"/"paper23_minor_revision_20260713"))
MASTER=MINOR/"Minor_revision_master_results_and_verification.xlsx"
ENG=Path(os.environ.get("FZYC_MANUSCRIPT_OUT", ROOT/"output"/"Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript.docx"))
OUT=Path(os.environ.get("FZYC_CHINESE_OUT", ROOT/"output"/"候选池扩张与模型选择损失_中文完整论文.docx"))
FIG=Path(os.environ.get("FZYC_FIG_OUT", MINOR/"main_figures"))
DISPLAY={"bace":"BACE","bbbp":"BBBP","clintox":"ClinTox","esol":"ESOL","freesolv":"FreeSolv","lipo":"Lipophilicity","tdc_caco2_wang":"Caco2","tdc_hia_hou":"HIA","tdc_pgp_broccatelli":"P-gp"}
TRANSFORM_LABELS={"raw":"原始矩阵","row_centred":"逐行中心化矩阵","fixed_reference_relative":"固定参照差值矩阵","within_unit_rank":"单元内秩矩阵"}

def font_run(r, size=12, bold=False, heading=False):
    r.bold=bold; r.font.size=Pt(size); r.font.name="Times New Roman"
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"黑体" if heading else "宋体")

def addp(doc,text="",style=None,bold=False,align=None):
    size=16 if style=="Title" else 15 if style=="Heading 1" else 13 if style=="Heading 2" else 12
    p=doc.add_paragraph(style=style); r=p.add_run(text); font_run(r,size,bold,style and style.startswith("Heading"))
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(0)
    if align is not None: p.alignment=align
    return p

def border(cell, **edges):
    tcPr=cell._tc.get_or_add_tcPr(); b=tcPr.first_child_found_in("w:tcBorders")
    if b is None: b=OxmlElement("w:tcBorders"); tcPr.append(b)
    for edge,attrs in edges.items():
        el=b.find(qn("w:"+edge))
        if el is None: el=OxmlElement("w:"+edge); b.append(el)
        for k,v in attrs.items(): el.set(qn("w:"+k),str(v))

def three_line(t):
    none={"val":"nil"}; line={"val":"single","sz":"8","color":"000000"}
    for i,row in enumerate(t.rows):
        for c in row.cells:
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            border(c,top=none,bottom=none,left=none,right=none,insideH=none,insideV=none)
            if i==0: border(c,top=line,bottom=line)
            if i==len(t.rows)-1: border(c,bottom=line)
            for p in c.paragraphs:
                p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(0)
                for r in p.runs: font_run(r,8.5,i==0)

def table(doc, headers, rows):
    t=doc.add_table(rows=1,cols=len(headers))
    for c,x in zip(t.rows[0].cells,headers): c.text=str(x)
    for row in rows:
        cells=t.add_row().cells
        for c,x in zip(cells,row): c.text=str(x)
    three_line(t); return t

def fig(doc,name,caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(FIG/name),width=Inches(6.7))
    p=addp(doc,caption); p.style="Caption"

def page_number(section):
    p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); p._p.append(fld)

def zh_list(values):
    return "、".join(values) if values else "无终点"

def effective_rank_text(frame):
    parts=[]
    for mode in ["raw","row_centred","fixed_reference_relative","within_unit_rank"]:
        row=frame.loc[frame.transformation.eq(mode)].iloc[0]
        parts.append(f"{TRANSFORM_LABELS[mode]}{row.entropy_rank_median:.2f}（IQR {row.entropy_rank_q25:.2f}–{row.entropy_rank_q75:.2f}；范围{row.entropy_rank_min:.2f}–{row.entropy_rank_max:.2f}）")
    return "；".join(parts)

def main():
    master=pd.read_excel(MASTER,sheet_name="Master result table")
    ranking=pd.read_excel(MASTER,sheet_name="Ranking main")
    effective=pd.read_excel(MASTER,sheet_name="Effective rank")
    cross=pd.read_excel(MASTER,sheet_name="Cross-fitted")
    matched_endpoint=pd.read_excel(MASTER,sheet_name="Matched K3")
    support=pd.read_excel(MASTER,sheet_name="Chemical support")
    null_summary=pd.read_csv(MINOR/"mechanism_permutation_null_summary.csv")
    mechanism=json.loads((MINOR/"mechanism_calibration_audit.json").read_text(encoding="utf-8"))
    rank4=ranking.loc[ranking.candidate_count.eq(4)].iloc[0]
    rank32=ranking.loc[ranking.candidate_count.eq(32)].iloc[0]
    class_pos=[DISPLAY[t] for t in cross.loc[cross.task_type.eq("classification") & cross.cross_fitted_effect.gt(0),"task"]]
    class_neg=[DISPLAY[t] for t in cross.loc[cross.task_type.eq("classification") & cross.cross_fitted_effect.lt(0),"task"]]
    reg_pos=[DISPLAY[t] for t in cross.loc[cross.task_type.eq("regression") & cross.cross_fitted_effect.gt(0),"task"]]
    reg_neg=[DISPLAY[t] for t in cross.loc[cross.task_type.eq("regression") & cross.cross_fitted_effect.lt(0),"task"]]
    class_excludes=[DISPLAY[r.task] for r in cross[cross.task_type.eq("classification")].itertuples() if r.split_seed_bootstrap95_low_cross_fitted>0 or r.split_seed_bootstrap95_high_cross_fitted<0]
    reg_excludes=[DISPLAY[r.task] for r in cross[cross.task_type.eq("regression")].itertuples() if r.split_seed_bootstrap95_low_cross_fitted>0 or r.split_seed_bootstrap95_high_cross_fitted<0]
    endpoint_medians=matched_endpoint.rename(columns={"endpoint_median_gain":"selected_model_gain_median"})
    matched_positive=[DISPLAY[t] for t in endpoint_medians.loc[endpoint_medians.selected_model_gain_median.gt(0),"task"]]
    matched_negative=[DISPLAY[t] for t in endpoint_medians.loc[endpoint_medians.selected_model_gain_median.lt(0),"task"]]
    subset_positive=matched_endpoint.set_index("task").positive_subset_proportion
    support_medians=support.set_index(["task_type","tanimoto_bin"]).selected_performance_median
    diversity_text=effective_rank_text(effective)
    positive_count=int(master.loc[master.metric.eq("positive cross-fitted direction count"),"estimate"].iloc[0])
    negative_count=int(master.loc[master.metric.eq("negative cross-fitted direction count"),"estimate"].iloc[0])
    abstract=render_chinese_abstract(endpoint_count=cross.task.nunique(),positive_count=positive_count,negative_count=negative_count)
    d=Document(); sec=d.sections[0]; sec.top_margin=Cm(2.54); sec.bottom_margin=Cm(2.54); sec.left_margin=Cm(2.54); sec.right_margin=Cm(2.54); page_number(sec)
    styles=d.styles
    for s in ["Normal","Caption"]:
        styles[s].font.name="Times New Roman"; styles[s]._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体"); styles[s].font.size=Pt(12)
    for s,n in [("Title",16),("Heading 1",15),("Heading 2",13)]:
        styles[s].font.name="Times New Roman"; styles[s]._element.rPr.rFonts.set(qn("w:eastAsia"),"黑体"); styles[s].font.size=Pt(n); styles[s].font.bold=True
    addp(d,"分子性质预测中候选池扩张、验证排序失真与模型选择损失：一项回顾性嵌套审计研究","Title",True,WD_ALIGN_PARAGRAPH.CENTER)
    addp(d,"摘要","Heading 1")
    for text in abstract.values(): addp(d,text)
    addp(d,"关键词：分子性质预测；候选池扩张；模型选择；嵌套交叉验证；效用模式多样性；验证排序保真度；化学支持范围")

    sections=[
    ("1 引言",CHINESE_INTRODUCTION_PARAGRAPHS),
    ("2 方法",[]),
    ("2.1 研究设计与证据层级",[
      "本研究是对既有分子性质实验的回顾性审计。9个主要终点、候选顺序、K值、模型seed、划分逻辑、选择规则、结局和排除条件在所报告分析中保持不变。审计规范是在原始外层结果产生后重建的，因此不属于前瞻性预注册研究（图1）。",
      "证据分为四层：32候选近重复传统模型池为主要审计；组成重采样和已知真值模拟用于机制解释；12候选多视图实验用于评估异质表征在相同公开终点上的可兑现收益；四模型逐样本预测、保形预测和化学边界分析用于可靠性限定。这些分析不合并为单一排行榜。"]),
    ("2.2 数据集与分子标准化",[
      "主要终点包括ESOL、FreeSolv、Lipophilicity、BBBP、BACE、ClinTox、Caco2 Wang、HIA Hou和P-gp Broccatelli。ESOL、FreeSolv、Lipophilicity和Caco2使用RMSE，其余分类终点使用ROC-AUC作为主要选择效用。PR-AUC、Brier分数、期望校准误差和少数类召回仅用于可靠性分析。",
      "分子使用RDKit解析并按端点级固定规则标准化；保留最大有效片段，执行化学合理的电荷规范化，并保留立体信息。插补、缩放及其他拟合变换仅在相应训练折上估计。分类重复结构仅在标签一致时合并，冲突标签删除；回归重复标签取平均并记录重复数。ClinTox清洗后含1,376个结构，其中58个为阳性。数据集摘要见表1。"]),
    ("2.3 候选登记表与计算暴露",[
      "受控扩张登记表含32个共享Morgan-512表征的候选。K=4、8、16和32为预先固定的候选顺序前缀，依次加入正则化线性、袋装树和多类提升模型。顺序未根据外层表现重排。",
      "每个候选在不同K下使用相同预处理、内层折和重拟合逻辑，但总搜索暴露随K增加。日志记录17,280次拟合和4,437.95候选拟合秒。该登记表用于研究相关候选的重复选择，不代表所有图模型或基础模型的完整搜索空间。候选登记和计算暴露见表2。"]),
    ("2.4 重复嵌套骨架评估",[
      "划分seed为11、23、37、53和71。分类与回归终点均生成5套真正不同的骨架划分；每套含3个外层骨架折，并在每个外层训练分区内生成3个内层骨架折。Bemis–Murcko骨架组保持完整，分配使用seed相关的随机并列处理并尽量平衡样本数，且不依据模型性能选择划分。",
      "验证最优选择器按平均内层效用排序候选，分类效用为ROC-AUC，回归效用为负RMSE；并列时按登记顺序决定。候选在完整外层训练分区重拟合后仅在外层折评估一次。完整划分清单记录样本数、骨架数、标签均值/标准差/范围、split hash及无跨折骨架重叠检查。"]),
    ("2.5 审计差距与排序估计量",[
      "对外层单元u和候选池C_K，内层平均效用记为V_uj，外层效用记为A_uj。验证选择候选为s_u=argmax_j V_uj，有限外层最佳候选为b_u=argmax_j A_uj。原始选择损失L_u(K)=A_u,b_u−A_u,s_u；分类中为ROC-AUC差，回归中对应RMSE增加。",
      "Top-3恢复使用CAHit@3=(Hit@3−3/K)/(1−3/K)进行机会校正。MRR以随机排序期望H_K/K校正。排序指标先在每个划分seed内平均3个外层折，再在每个终点内平均5个seed，最后以9个终点的中位数和IQR作为主汇总；同时报告NDCG、Spearman、Kendall和外层最佳候选的验证秩百分位。分类与回归始终分轴报告，不合并不同单位。",
      "置换负对照在每个锁定外层单元内随机抽取候选秩，执行5,000次重复，并采用与真实分析完全相同的折—seed—终点聚合。分级正对照使用135个审计单元中的4,320条锁定外层候选效用，注入0、0.10、0.25、0.50、0.75和1.0六级验证—审计信号；除完全信号外，每个配置模拟500次。对照仅校准指标零点和信号恢复，不改变真实候选选择。"]),
    ("2.6 有效多样性估计",[
      "对每个终点和K构建候选效用矩阵，分别分析原始效用、逐行中心化效用、固定参照差值效用和单元内秩。逐行中心化用于去除共同审计单元难度；固定参照和秩矩阵用于检验候选相对模式。",
      "报告经验相关谱、Ledoit–Wolf收缩相关谱、谱熵有效秩、参与率有效秩和非对角相关中位数。使用固定seed 20260713执行5,000次层级bootstrap，并报告Monte Carlo稳定性、留一seed范围、留一fold范围及预先指定参照敏感性。跨终点主汇总采用中位数、IQR和范围，不把9个终点视为总体随机样本，也不提供总体推断性置信区间。"]),
    ("2.7 候选组成控制",[
      "随机顺序、随机子集和家族平衡控制用于判断K趋势是否由特定前缀组成驱动。每种模式在终点内汇总，随机化seed仅为重采样装置，不作为独立终点。K=32时各模式均等于完整候选池。"]),
    ("2.8 交叉拟合审计参照",[
      "对每个终点、K和留出seed，在其他4个seed上计算候选平均外层效用并固定最高者作为参照，再在留出seed的3个折上评估。交叉拟合差距为参照效用减验证所选效用；同一单元差距使用该折内最大候选效用。该参照降低了同一折上定义和评估最大值的循环性，但仍复用同一公开终点，不能视为外部验证。"]),
    ("2.9 有限审计赢家乐观偏差模拟",[
      "在所有候选真实标准化效用相同的情形下，模拟相关高斯审计误差。K取4、8、16、32，有效审计量取25、50、100、200，候选相关取0、0.5、0.9，每个配置30,000次重复。赢家乐观偏差定义为最大观测效用减相应候选真实效用。该模拟仅说明有限最大值的机械性质，不用于校正真实终点。"]),
    ("2.10 等规模多视图压力测试",[
      "12候选登记表由Morgan-512、MACCS、RDKit2D及其拼接表征与3类学习器交叉组成。利用候选级内外层效用，完整枚举C(12,3)=220个K=3子集，无需重新训练。全部子集属于同一登记表的组成敏感性分析，并非220次独立实验。",
      "子集采用互斥组成分类：Morgan-only参照、仅单一表征、仅单一学习器、无拼接的表征平衡、含拼接的表征平衡及混合不平衡。每个终点报告中位数、IQR、分布范围和优于Morgan-only的子集比例，不计算把220个重叠子集当作独立样本的P值。"]),
    ("2.11 有限表征基线分析",[
      "六终点四模型面板包括RDKit-RF、GCN、ChemBERTa冻结嵌入线性探针和MoLFormer冻结嵌入线性探针。各模型训练预算并不等同，因此该面板仅用于共享划分下的选择行为、预测相关和错误互补性，不用于比较充分调优的现代架构。"]),
    ("2.12 可靠性与化学边界分析",[
      "逐样本预测按最大训练集Morgan Tanimoto划分为低支持(<0.5)、中支持(0.5至<0.7)和高支持(≥0.7)。每个区间比较所选模型和交叉拟合参照性能、错误Jaccard重叠、模型分歧、不确定性—错误相关和假阴性富集。",
      "外层精确Murcko骨架按设计均未见，因此采用预先固定的骨架指纹最大相似度0.5阈值区分‘已见或相关骨架’与‘新颖骨架’。四模型面板还计算预测矩阵和误差矩阵的相关性、谱熵有效秩和参与率有效秩，不外推至32候选登记表。"]),
    ("2.13 统计推断",[
      "各终点的K=32减K=4效应先在每个不同划分seed内平均3个折，再对5个seed均值进行10,000次聚类bootstrap。分类ROC-AUC损失与回归RMSE损失构成按任务分层的共同主要证据；两者保留各自单位，不计算跨任务平均效应。排序指标、同一单元差距、有效多样性和等规模组成效应为关键互补证据。220个K=3子集结果是登记子集的分布，不视为220次独立实验。"]),
    ]
    for h,pars in sections:
        addp(d,h,"Heading 1" if h in ["1 引言","2 方法"] else "Heading 2")
        for x in pars: addp(d,x)

    # Main tables use the same audited values as the English manuscript.
    addp(d,"表1 主要终点与评价指标。单位：ESOL为log mol/L；FreeSolv为kcal/mol；Lipophilicity为实验logD；Caco2 Wang为数据集给定对数渗透性量纲。分类终点报告阳性类别n（%）。",bold=True)
    eng=Document(ENG)
    rows1=[[c.text for c in row.cells] for row in eng.tables[0].rows[1:]]
    table1_terms={" to ":"至"," hydration free energy ":" 水合自由能 "," experimental logD":" 实验logD"," positive ":" 阳性 "," dataset-provided log-permeability scale":" 数据集给定对数渗透性量纲"}
    for row in rows1:
        for old,new in table1_terms.items(): row[2]=row[2].replace(old,new)
    table(d,["终点","分析n","类别比例或目标范围","主要指标"],rows1)
    addp(d,"表2 审计组成与计算暴露",bold=True)
    rows2=[
      ["受控前缀审计","32个Morgan候选","5组seeded scaffold partitions；3个外层折 × 3个内层折","扩张审计","17,280次候选拟合"],
      ["机制校准对照","32个Morgan候选","5,000次置换；6级信号 × 4个K","零点与信号恢复","不增加模型拟合"],
      ["组成控制","32个Morgan候选","每个模式、K和划分seed各100个子集seed","组成敏感性","不增加拟合"],
      ["等规模多视图","12个多视图候选","5组seeded scaffold partitions；穷举C(12,3)","组成效应","6,480次候选拟合"],
      ["4模型面板","4个表征候选","可追溯的外层预测面板","可靠性边界","360个候选–折单元"],
    ]
    table(d,["分析","登记表","重采样设计","用途","暴露"],rows2)
    fig(d,"Figure_1_retrospective_nested_audit_architecture.png","图1 候选池扩张与模型选择损失的回顾性嵌套审计框架。中央模块将5个种子划分中3个内层骨架折的候选排序，与3个外层折的一次性审计严格分离，并禁止外层标签反馈。左侧输入已登记分子终点和预设候选前缀；周围分支汇总矩阵依赖多样性、机会校正排序、同单元与交叉拟合选择差距、候选组成对照、等规模多视图压力测试、有限审计赢家乐观偏差、化学支持边界以及四模型预测和误差可靠性，最终汇入可审计证据图谱，而非单一模型排行榜。")

    results=[
    ("3 结果",[]),
    ("3.1 有效多样性取决于矩阵构造",[f"K=32时，9个终点的Ledoit–Wolf谱熵有效秩主汇总为{diversity_text}。原始效用保留共同审计单元难度，逐行中心化移除共同水平位移，固定参照估计依赖参照，单元内秩保留顺序但丢失效用间距。参与率有效秩和非对角相关中位数见补充表S6；不把任何一种变换视为唯一真实的候选数。","留一seed、留一fold和预设参照分析均表明终点估计对矩阵构造和参照敏感。逐行中心化同时引入每行和为零的约束；收缩稳定相关谱，但不会增加独立审计单元（图2；补充表S6–S7）。"]),
    ("3.2 机会校正排序保真度随K变化",[f"按照折内到seed、seed内到终点、再跨9个终点的预定汇总，端点中位CAHit@3由K=4时的{rank4.chance_adjusted_hit_median:.3f}（IQR {rank4.chance_adjusted_hit_q25:.3f}–{rank4.chance_adjusted_hit_q75:.3f}）变为K=32时的{rank32.chance_adjusted_hit_median:.3f}（{rank32.chance_adjusted_hit_q25:.3f}–{rank32.chance_adjusted_hit_q75:.3f}）；标准化MRR增益由{rank4.normalized_mrr_gain_median:.3f}（{rank4.normalized_mrr_gain_q25:.3f}–{rank4.normalized_mrr_gain_q75:.3f}）变为{rank32.normalized_mrr_gain_median:.3f}（{rank32.normalized_mrr_gain_q25:.3f}–{rank32.normalized_mrr_gain_q75:.3f}）。标准化MRR使用随机排序期望H_K/K。NDCG、Spearman、Kendall和秩百分位提供互补证据。",f"置换对照回到校正零点，四个K下的观测CAHit@3终点中位数均高于95%置换包络（最大单侧置换P={null_summary.loc[null_summary.metric.eq('chance_adjusted_hit'),'one_sided_p_observed_le_null'].max():.4f}）。在K=4、8、16和32下，注入信号增强均使CAHit@3和标准化MRR单调提高、固定范围选择损失单调降低；零信号CAHit@3最大绝对偏差为{mechanism['max_null_signal_abs_cahit']:.3f}，完全信号时所有K的中位选择损失均为{mechanism['max_perfect_signal_selection_loss']:.3f}（图3A–B）。"]),
    ("3.3 审计差距分解区分机会与兑现",["K=32时，分类终点平均有限审计最佳收益为0.0311 ROC-AUC，所选模型收益为0.0173，增量审计差距为0.0138；K=4时分别为0.0028、−0.0014和0.0042。回归K=32时对应RMSE收益为0.8078、0.7378和0.0700，K=4时为0.7619、0.7442和0.0178。不同单位分开报告（图4A–B）。"]),
    ("3.4 交叉拟合效应在终点间异质",[f"K=32相对K=4的交叉拟合效应在9个终点中的6个为正（{zh_list(class_pos + reg_pos)}），在3个为负（{zh_list(class_neg + reg_neg)}）。划分seed区间在分类ROC-AUC损失层的{zh_list(class_excludes)}和回归RMSE损失层的{zh_list(reg_excludes)}排除零；FreeSolv为显著负向。两个任务层是共同主要证据，按各自单位分开报告（图3C；表3）。"]),
    ("3.5 有限审计最大值包含赢家乐观偏差",["在候选相关为0.9、有效审计量为50且候选真实效用相同的模拟中，平均赢家乐观偏差由K=4时的0.046个标准差增至K=32时的0.092个标准差。K增加、有效审计量减小或候选相关降低时，有限最大值的乐观偏差增大（图4C–D）。"]),
    ("3.6 等规模多视图收益具有终点依赖性",[f"完整220个K=3子集中，终点中位收益在{zh_list(matched_positive)}为正，在{zh_list(matched_negative)}为负。各终点正平均收益子集比例介于{subset_positive.min():.3f}和{subset_positive.max():.3f}。这些比例描述固定登记表中的组成空间，不把重叠子集数量解释为统计功效（图5；补充表S9–S10）。"]),
    ("3.7 表征错误相关但具有部分互补性",["四模型面板的预测相关和高错误Jaccard重叠均明显高于零，预测谱熵有效秩低于名义4。候选错误既非独立也非完全相同。由于训练预算不等，这些结果仅说明固定配置下的错误结构，不能形成现代架构性能排序（图6A）。"]),
    ("3.8 化学边界处可靠性下降",[f"分类终点在低、中、高Tanimoto支持下的审计单元中位ROC-AUC分别为{support_medians.loc[('classification','<0.5')]:.3f}、{support_medians.loc[('classification','0.5-0.7')]:.3f}和{support_medians.loc[('classification','>0.7')]:.3f}；回归中位RMSE分别为{support_medians.loc[('regression','<0.5')]:.3f}、{support_medians.loc[('regression','0.5-0.7')]:.3f}和{support_medians.loc[('regression','>0.7')]:.3f}。新颖骨架以已见或相关骨架为reference = 1报告错误重叠、分歧和高错误富集相对变化（图6B–C）。","ClinTox面板将ROC-AUC、PR-AUC、少数类召回、假阴性率、条件覆盖率和预测集大小分区显示，不把不同可靠性量合并为单一性能排名（图6D）。"]),
    ]
    for h,pars in results:
        addp(d,h,"Heading 1" if h=="3 结果" else "Heading 2")
        for x in pars: addp(d,x)
        if h.startswith("3.1"): fig(d,"Figure_2_candidate_diversity_after_adjustment.png","图2 矩阵依赖的候选多样性。A比较原始、逐行中心化、固定参照差值和单元内秩矩阵的Ledoit–Wolf谱熵秩；实线为终点中位数，阴影为IQR，灰色点线为名义K。原始效用保留共同审计单元难度，逐行中心化去除共同水平位移。B在K=32比较谱熵秩与参与率秩，并给出45°参考线。C为校正后的候选相关中位数。D显示K=32的固定参照差值终点估计；实线和虚线分别为留一seed与留一fold范围，空心点为预设参照。")
        if h.startswith("3.2"): fig(d,"Figure_3_chance_adjusted_ranking_and_selection_gaps.png","图3 机会校正排序校准与交叉拟合差距。A显示CAHit@3和标准化MRR增益的终点中位数与IQR；灰带为相同聚合下5,000次随机秩置换形成的联合95%包络。B显示六级注入验证—审计信号下、四种候选数量的CAHit@3正对照恢复。C按任务分层显示分类ROC-AUC损失和回归RMSE损失及split-seed bootstrap 95%区间；实心点表示区间排除零，空心点表示区间跨零。D为登记前缀、随机顺序、随机子集和家族平衡控制；K=32时各模式按设计等于完整登记表。")
        if h.startswith("3.3"): fig(d,"Figure_4_selection_gap_and_winner_optimism.png","图4 选择差距分解与赢家乐观偏差。A和B比较分类与回归的观测审计最佳收益和验证所选收益。C使用终点内标准化效应，在一个按任务分层的连续森林图中展示全部终点；原始ROC-AUC与RMSE效应另行列示。D显示候选真实效用相同时的有限审计赢家乐观偏差。")
        if h.startswith("3.4"):
            disp={"bace":"BACE","bbbp":"BBBP","clintox":"ClinTox","esol":"ESOL","freesolv":"FreeSolv","lipo":"Lipophilicity","tdc_caco2_wang":"Caco2","tdc_hia_hou":"HIA","tdc_pgp_broccatelli":"P-gp"}
            rows=[]
            for _,r in cross.iterrows():
                metric="ROC-AUC损失" if r.task_type=="classification" else "RMSE损失"
                lower=r.split_seed_bootstrap95_low_cross_fitted; upper=r.split_seed_bootstrap95_high_cross_fitted
                ci=f"{lower:.4f}至{upper:.4f}"
                if lower>0: interpretation="K=32损失较大"
                elif upper<0: interpretation="K=32损失较小"
                elif r.cross_fitted_effect>=0: interpretation="不确定；点估计为正"
                else: interpretation="不确定；点估计为负"
                rows.append([disp[r.task],metric,f"{r.cross_fitted_effect:.4f}",ci,interpretation])
            addp(d,"表3 按任务分层共同主要证据中的终点特异性K=32减K=4交叉拟合效应",bold=True); table(d,["终点","指标","交叉拟合效应","95% CI","解释"],rows)
        if h.startswith("3.6"): fig(d,"Figure_5_matched_size_multiview_composition.png","图5 等规模表征组成效应。A展示全部220个登记K=3子集的终点特异性分布。B汇总互斥的组成类别。C展示K=3、6、9和12阶梯。D使用终点MAD标准化效应，在一个连续森林图中展示全部终点；原始ROC-AUC收益和RMSE降低另行列示。")
        if h.startswith("3.8"): fig(d,"Figure_6_prediction_errors_across_chemical_support.png","图6 化学支持边界下的预测可靠性。A在对称四模型矩阵中合并预测相关与高错误Jaccard重叠。B将分类ROC-AUC、回归RMSE、分类假阴性率以及分类和回归高错误富集整合为一张Tanimoto支持风险矩阵；格内文字为自然尺度中位数，颜色仅表示该指标行内的不利方向。C在对数比值轴上报告新颖骨架相对于已见或相关骨架的错误重叠、模型分歧、高错误富集和假阴性富集。D在一张对齐点图中比较四个固定配置模型的ROC-AUC、PR-AUC、少数类召回、假阴性率、条件覆盖和预测集大小；预测集大小仅为显示而映射到1–2尺度。")

    discussions=[
    ("4 讨论",[]),
    *CHINESE_DISCUSSION_SECTIONS,
    ("5 结论",[f"在9个分子性质审计中，名义候选数不能唯一描述有效搜索多样性，矩阵构造会显著改变有效秩估计。候选池扩张伴随机会校正排序减弱和异质的交叉拟合选择差距；效应在{zh_list(class_pos + reg_pos)}为正、在{zh_list(class_neg + reg_neg)}为负。等规模多视图比较显示表征组成收益依赖终点，逐样本分析进一步将可靠性限制定位到低化学支持和新颖骨架。", "分子模型选择研究应同时报告候选资格、名义K、矩阵依赖的效用模式多样性、机会校正排序、同一单元和交叉拟合差距、划分唯一性、计算暴露、失败候选和化学支持边界。该审计框架提高证据透明度，但不构成通用选择器或部署级筛选系统。"]),
    ]
    for h,pars in discussions:
        addp(d,h,"Heading 1" if h in ["4 讨论","5 结论"] else "Heading 2")
        for x in pars: addp(d,x)

    addp(d,"缩写表","Heading 1"); addp(d,"BACE，β-分泌酶1；BBBP，血脑屏障通透性；bRo5，超出五规则；CAHit@3，机会校正Top-3命中；CQR，保形化分位数回归；ECE，期望校准误差；HIA，人体肠道吸收；MRR，平均倒数秩；NDCG，标准化折损累计增益；P-gp，P-糖蛋白；PR-AUC，精确率–召回率曲线下面积；RMSE，均方根误差；ROC-AUC，受试者工作特征曲线下面积。")
    addp(d,"声明","Heading 1")
    for h,x in [("伦理批准与参与同意","不适用。本研究使用公开分子数据集，不涉及人类参与者、人类数据或动物实验。"),("发表同意","不适用。"),("数据与材料可用性","公开数据来源见补充表S1；派生逐折表、源文件哈希和分析代码随提交包提供。")]: addp(d,h,"Heading 2"); addp(d,x)
    author_declarations = [
        ("利益冲突", os.environ.get("FZYC_COMPETING_INTERESTS", "提交前需作者确认；未推断任何利益冲突声明。")),
        ("基金资助", os.environ.get("FZYC_FUNDING", "提交前需作者确认；未推断资助来源、项目编号或资助方角色。")),
        ("作者贡献", os.environ.get("FZYC_AUTHOR_CONTRIBUTIONS", "提交前需提供经核实的作者姓名缩写与CRediT贡献角色；未推断作者贡献。")),
        ("致谢", os.environ.get("FZYC_ACKNOWLEDGEMENTS", "提交前需作者确认或提供经核实的致谢文字；未推断致谢内容。")),
    ]
    for h,x in author_declarations: addp(d,h,"Heading 2"); addp(d,x)
    addp(d,"补充信息","Heading 1"); addp(d,"Additional file 1：补充方法与结果；Additional file 2：机器可读补充表S1–S20；Additional file 3：补充图S1–S17。")
    addp(d,"参考文献","Heading 1")
    # Preserve the exact numbered reference list from the audited English manuscript.
    started=False
    for p in eng.paragraphs:
        if p.text.strip()=="References": started=True; continue
        if started and p.text.strip(): addp(d,p.text.strip())
    d.core_properties.title="分子性质预测中候选池扩张、验证排序失真与模型选择损失：一项回顾性嵌套审计研究"
    d.save(OUT); embed_svg_figures(OUT, FIG); print(OUT)

if __name__=="__main__": main()
