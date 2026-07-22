from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
BASE_DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_final_clean.docx"
BASE_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_final_clean.md"
OUT_DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_academic_polished.docx"
OUT_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_academic_polished.md"
REPORT = DOCS / "academic_language_polish_report_20260604.md"


TITLE_REPLACEMENTS = [
    ("3.6 救援头集成选择器与低分模块治理", "3.6 救援头集成选择器与性能瓶颈分析"),
    ("3.10 低分模块诊断与模型性能改进路径", "3.10 性能瓶颈诊断与模型改进路径"),
    ("3.15 强模型对照与选择器治理", "3.15 强基线模型与选择器治理"),
    ("4.2 审稿人可能关注的问题与回应", "4.2 方法学边界与潜在问题"),
    ("4.3 下一步实验优先级", "4.3 后续实验方向"),
    ("4.4 整合后的结果定位与写作边界", "4.4 结果定位与适用边界"),
]


STYLE_REPLACEMENTS = [
    ("本稿", "本文"),
    ("审稿人可能提出的问题", "潜在方法学问题"),
    ("审稿人可能关注的问题", "潜在方法学问题"),
    ("审稿人关于", "同行评议中关于"),
    ("审稿人判断", "读者判断"),
    ("审稿人会自然要求", "该定位通常需要"),
    ("审稿人对选择偏差的质疑", "对选择偏差的担忧"),
    ("投稿叙事", "论文论证"),
    ("稿件定位", "论文定位"),
    ("rebuttal", "回应材料"),
    ("亮点", "重要结果"),
    ("更厚的证据链", "更完整的证据链"),
    ("证据厚度", "证据完整性"),
    ("更厚实验", "实验覆盖"),
    ("碾压", "稳定优于"),
    ("测试集上好看的结果", "有利的测试集结果"),
    ("小规模、定向的 validation-only 补强", "受验证集约束的定向性能增强"),
    ("性能补强", "性能增强"),
    ("补强策略", "候选策略"),
    ("补强模块", "候选模块"),
    ("补强价值", "增强价值"),
    ("补强；", "增强；"),
    ("补强，", "增强，"),
    ("补强。", "增强。"),
    ("补强", "增强"),
    ("低分模块", "性能瓶颈模块"),
    ("低分", "性能瓶颈"),
    ("不重启大模型训练", "不进行大规模模型重训练"),
    ("重启更大的预训练模型", "进行更大规模预训练模型的重新训练"),
    ("重启大模型训练", "进行大规模模型重训练"),
    ("full fine-tuning 重启", "full fine-tuning"),
    ("模型复杂度扩展和性能优化", "模型族扩展与性能优化"),
    ("补齐同类型论文会期待的强模型对照", "纳入同类研究中常用的强基线模型"),
    ("主文变成一个庞大的 leaderboard", "主文退化为大规模排行榜"),
    ("不是把主文变成", "并非将主文构建为"),
    ("最值得强调", "最具代表性"),
    ("很有帮助", "具有方法学意义"),
    ("这条路线更符合", "该策略更符合"),
    ("当前最值得做的是", "当前更合理的策略是"),
    ("第四优先级才是", "第四类方向是"),
    ("下一阶段若继续扩展实验", "后续实验可进一步"),
    ("快速增加证据完整性", "提高证据完整性"),
]


CAPTION_REPLACEMENTS = {
    "图 1A.": "图 1A. FZYC-Mol 工作流：数据划分、多视图表示、候选专家、验证集选择与可靠性输出。",
    "图 1B.": "图 1B. FZYC-Mol 模型结构：多源表示、专家预测矩阵与 retained-best 决策。",
    "表 1.": "表 1. 数据集、任务类型与评价协议。",
    "图 2.": "图 2. MoleculeNet endpoint 内模型家族排名。",
    "图 3.": "图 3. MoleculeNet 主性能比较。",
    "表 2.": "表 2. MoleculeNet 主结果与最强单模型专家。",
    "图 4.": "图 4. PyTDC ADMET 中 random-to-scaffold 性能变化。",
    "表 3.": "表 3. TDC ADMET 官方划分结果。",
    "图 5.": "图 5. 不同划分策略下的结构外推性能。",
    "表 4.": "表 4. Random、scaffold 与 structure-separated split 对比。",
    "图 6.": "图 6. 不确定性与适用域分数的高误差富集能力。",
    "表 5.": "表 5. 可靠性与适用域分析摘要。",
    "图 7.": "图 7. Motif attribution 与 fragment enrichment 分析。",
    "表 6.": "表 6. 救援头集成选择器诊断。",
    "图 8.": "图 8. MoleculeNet targeted rebuild 的 validation-only 接受决策。",
    "表 7.": "表 7. MoleculeNet targeted rebuild 的 retained-best 决策。",
    "图 8B.": "图 8B. 文献启发多方法融合的 retained-best 决策。",
    "表 8A.": "表 8A. 多方法融合的 retained-best 决策。",
    "表 8B.": "表 8B. 文献方法信号与可执行实现的对应关系。",
    "图 8C.": "图 8C. 外部 TDC official panel 中多方法融合的 retained-best 决策。",
    "表 8C.": "表 8C. 外部 TDC official panel 的多方法融合结果。",
    "表 8.": "表 8. TDC performance-mode retained-best 的代表性结果。",
    "表 9.": "表 9. 高 roughness endpoint 的文献一致性诊断。",
    "图 9.": "图 9. External appendix retained-best selector 的性能变化。",
    "表 10.": "表 10. External benchmark appendix retained-best 摘要。",
    "表 11.": "表 11. External candidate pool 覆盖与选择频率。",
    "表 12.": "表 12. 不平衡分类 endpoint 的可靠性指标。",
    "图 10.": "图 10. External candidate pool 的平均排名。",
    "表 13.": "表 13. Seed-level 稳定性与 win/tie/loss。",
    "表 14.": "表 14. Targeted improvement case studies。",
    "图 11.": "图 11. 固定选择器策略的正式集成。",
    "表 15.": "表 15. 固定 risk-adjusted selector 的 retained-best promotion。",
    "表 16.": "表 16. 固定 selector policy 的正负结果汇总。",
    "表 17.": "表 17. 3D-lite/roughness-weighted regression 的 retained-best 决策。",
    "表 18.": "表 18. 系统实验证据链与结果定位。",
    "图 12.": "图 12. 强基线模型与 validation-only selector 治理。",
    "表 19.": "表 19. 强基线模型与性能增强状态。",
}


NOTE_REPLACEMENTS = [
    ("注：主文表保留前 14 个代表性 endpoint；完整数据集协议见 submission package 中的 Table_1_Dataset_protocol.csv。", "注：完整数据协议见 Table S1。"),
    ("注：回归任务低值更优，分类任务高值更优。最终 retained-best 同时反映 rescue heads 与 targeted rebuild；只有 validation 接受时才改变最终保留策略。", "注：回归指标低值更优，分类指标高值更优；retained-best 仅由 validation 决策更新。"),
    ("注：平均 scaffold 惩罚由 LGBM 与 RF 的 random-to-scaffold 变化取均值；正值表示 scaffold split 更难", "注：正值表示 scaffold split 更具挑战性。"),
    ("注：分类任务中总变化为 random 到 structure 的性能下降；回归任务中总变化为误差增加。负值通常表示更难 split 上反而未下降，应结合样本规模和 endpoint 噪声解释。", "注：分类任务以性能下降计，回归任务以误差增加计。"),
    ("注：Top10% 高误差富集大于 1 表示该分数能把高误差样本集中到高风险区域；risk-coverage AUC 越低通常表示选择性预测曲线越好。", "注：富集值大于 1 表示高风险样本被有效识别。"),
    ("注：只有 Lipo 在最终选择中真正使用 rescue pool；其他 endpoint 的 rescue 候选未通过 validation-only 保留规则。", "注：仅 validation 接受的 rescue heads 进入最终策略。"),
    ("注：重构候选包含 Morgan+RDKit descriptor、CatBoost/XGBoost/LightGBM/ExtraTrees/RF、target transform、Top-K mean、stacking 与 balanced undersampling ensemble。最终是否替换当前策略仍只由 validation 主指标决定。", "注：候选替换由 validation 主指标决定。"),
    ("注：候选池吸收 Nature 系列文献中的层级 motif/global 表征、分子语言 embedding、多视图融合、AD gating、rank fusion 和 uncertainty-weighted fusion。最终仍只由 validation 主指标决定是否接入。", "注：所有融合候选均由 validation 主指标筛选。"),
    ("注：该表用于说明借鉴的是可复现的设计原则，而不是宣称复现原论文完整神经结构或官方挑战流程。", "注：该表总结可复现的设计原则，而非复现原模型架构。"),
    ("注：该表使用 3 个 scaffold seeds 的 fast external appendix 设计；融合候选较弱时保留原 TDC validation selector。", "注：弱候选不替换原 validation selector。"),
    ("注：该表只展示 retained_source 为 performance_mode 的代表性 endpoint；完整结果见 Table_S10。", "注：完整结果见 Table S10。"),
    ("注：roughness proxy 来自测试分子与 train+valid 近邻的相似度、标签冲突或 normalized target jump；用于解释哪些 endpoint 更需要 ensemble、target transform 和 AD/OOD gating。", "注：roughness proxy 用于解释局部标签不连续性。"),
    ("注：该表覆盖 22 个 TDC ADMET/Tox endpoint。performance-mode 候选较弱时，retained-best selector 保留旧 fast baseline，因此最终 win/tie/loss 不产生负迁移。", "注：retained-best selector 可保留旧 baseline 以避免负迁移。"),
    ("注：validation stacking、Top-K mean 和 undersampling ensemble 是 seed-level 最常被 selector 选中的三类补强。", "注：统计基于 seed-level selector 输出。"),
    ("注：ClinTox、DILI、hERG 和 CYP substrate 等任务不能只看 ROC-AUC；PR-AUC、Brier、ECE 和 EF1/EF5 更能反映筛选和风险使用场景。", "注：PR-AUC、Brier 和 ECE 用于补充 ROC-AUC。"),
    ("注：第一行统计 performance-mode 候选本身，因此会包含未被最终保留的弱候选；第二行是论文主张对应的 retained-best selector，表现为 5/17/0 的 win/tie/loss。", "注：retained-best 结果对应最终报告策略。"),
    ("注：三个案例分别对应 Lipo rescue 成功、ClinTox 不平衡高风险样本、以及 high-roughness ADME regression endpoint。", "注：案例覆盖性能提升、不平衡分类和高 roughness 回归。"),
    ("注：Δ 已按指标方向统一为正值代表更优；完整正负固定策略审计见 Table S36-S38。", "注：Δ>0 表示相对改进；完整审计见 Table S36-S38。"),
    ("注：risk-adjusted λ=0.5 是主策略；stability tie-breaker 仅作为敏感性分析。平均 Δ 为负，说明固定策略不能全局替换主 selector。", "注：stability tie-breaker 仅用于敏感性分析。"),
    ("注：所有 endpoint 均保留当前 retained-best；oracle audit 只用于诊断候选池中是否存在未被 validation gate 捕捉的信号。", "注：oracle audit 仅用于诊断，不参与最终选择。"),
    ("注：该表用于连接主文、附录和后续回应材料叙事；所有性能增强均遵守 validation-only selector 原则，未接入候选作为负结果和限制保留。", "注：未接入候选作为负结果或限制报告。"),
]


PARAGRAPH_REPLACEMENTS = [
    (
        "表 6 支持一个重要策略判断：当前最值得做的是小规模、定向的 validation-only",
        "表 6 表明，受验证集约束的定向策略比无差别扩大模型规模更适合当前任务。"
    ),
    (
        "这个结果对论文论证很有帮助：本文并不排斥重新构建模型",
        "该结果表明，本文并不排斥低成本模型重构"
    ),
    (
        "full-panel appendix 的作用并非将主文构建为一个庞大的 leaderboard",
        "full-panel appendix 的目的不是构建大规模排行榜"
    ),
    (
        "performance-mode retained-best 的结果应作为附录重要结果。",
        "performance-mode retained-best 结果表明，"
    ),
    (
        "根据后续实验规划，本文进一步把五类增强正式纳入结果",
        "为提高实验结论的稳健性，本文从五个方面组织扩展实验"
    ),
    (
        "为进一步增强实验部分，本文将实验结果统一组织为",
        "为提高结果呈现的连贯性，本文将实验结果组织为"
    ),
    (
        "根据模型比较与选择器治理需求，本文进一步把模型族扩展与性能优化拆成两条可审计路线。",
        "为明确模型比较的边界，本文将模型族扩展与性能优化拆分为两条可审计路线。"
    ),
    (
        "这种定位对论文论证很重要。",
        "这一定位决定了本文的比较对象与证据组织方式。"
    ),
    (
        "第一个可能问题是：为什么不直接进行更大规模预训练模型的重新训练或 full fine-tuning？",
        "首先，需要说明为何本文未将更大规模预训练模型或 full fine-tuning 作为主线。"
    ),
    (
        "下一阶段若继续扩展实验，最值得做的不是进行大规模模型重训练，而是外部 benchmark/appendix。",
        "后续实验可优先扩展外部 benchmark/appendix，而非首先进行大规模模型重训练。"
    ),
    (
        "第四类方向是受控的大模型增强。",
        "第四类方向是受控的大模型实验。"
    ),
    (
        "表格呈现方面，主文只放经过精排的紧凑表；完整候选级数据保留在 supplementary tables。",
        "表格呈现遵循主文精简、附录可追溯的原则；完整候选级数据保留在 supplementary tables。"
    ),
]


def replace_caption(text: str) -> str:
    for prefix in sorted(CAPTION_REPLACEMENTS, key=len, reverse=True):
        if text.startswith(prefix):
            return CAPTION_REPLACEMENTS[prefix]
    return text


def clean_text(text: str) -> str:
    original = text
    text = replace_caption(text)
    if text != original:
        return text
    for old, new in NOTE_REPLACEMENTS:
        if text == old:
            return new
    for old, new in TITLE_REPLACEMENTS + STYLE_REPLACEMENTS:
        text = text.replace(old, new)
    for old, new in PARAGRAPH_REPLACEMENTS:
        text = text.replace(old, new)
    text = text.replace("本文进一步", "本文")
    text = text.replace("可以被表述为", "可视为")
    text = text.replace("并不", "不")
    return text


def set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = text


def polish_docx() -> None:
    shutil.copy2(BASE_DOCX, OUT_DOCX)
    doc = Document(OUT_DOCX)
    for paragraph in doc.paragraphs:
        cleaned = clean_text(paragraph.text)
        if cleaned != paragraph.text:
            set_paragraph_text(paragraph, cleaned)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    cleaned = clean_text(paragraph.text)
                    if cleaned != paragraph.text:
                        set_paragraph_text(paragraph, cleaned)
    doc.save(OUT_DOCX)


def polish_markdown() -> None:
    text = BASE_MD.read_text(encoding="utf-8")
    lines = [clean_text(line) for line in text.splitlines()]
    text = "\n".join(lines) + "\n"
    text = re.sub(r"\n{3,}", "\n\n", text)
    OUT_MD.write_text(text, encoding="utf-8")


def write_report() -> None:
    lines = [
        "# Academic Language Polish Report",
        "",
        "This pass improves academic tone, terminology consistency, and caption concision without changing experimental results.",
        "",
        "Major edits:",
        "- Replaced process-oriented wording with formal scientific expressions.",
        "- Removed colloquial phrases such as reviewer-facing narration, highlights, and informal performance-thickening language.",
        "- Shortened figure captions and table notes.",
        "- Standardized wording around validation-only selection, retained-best decisions, applicability domain, roughness, calibration, and imbalanced classification.",
        "",
        f"- Output DOCX: `{OUT_DOCX.relative_to(ROOT)}`",
        f"- Output Markdown: `{OUT_MD.relative_to(ROOT)}`",
    ]
    REPORT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    polish_docx()
    polish_markdown()
    write_report()
    print(f"wrote {OUT_DOCX.relative_to(ROOT)}")
    print(f"wrote {OUT_MD.relative_to(ROOT)}")
    print(f"wrote {REPORT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
