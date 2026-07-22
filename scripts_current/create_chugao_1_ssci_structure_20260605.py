from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from restructure_manuscript_integrated_results_20260604 import (
    add_image,
    add_table,
    set_font,
    setup_doc,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SRC_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260605_nine_chapter_structure.md"
OUT_MD = DOCS / "初稿-1.md"
OUT_DOCX = DOCS / "初稿-1.docx"


HEADINGS = [
    "# 摘要",
    "# 第一章 Introduction（引言）",
    "# 第二章 Literature Review（文献回顾）",
    "# 第三章 Methodology（研究方法/数据收集方法）",
    "# 第四章 Results（数据结果）",
    "# 第五章 Discussion（讨论）",
    "# 第六章 Conclusion（总结/展望）",
    "# 第七章 Acknowledgement（致谢）",
    "# 第八章 Declaration of Conflicting Interests / Funding（利益冲突声明/基金）",
    "# 第九章 References（参考文献）",
]


def chapter(text: str, heading: str) -> str:
    start = text.index(heading) + len(heading)
    next_positions = [text.index(h, start) for h in HEADINGS if h != heading and h in text[start:]]
    end = min(next_positions) if next_positions else len(text)
    return text[start:end].strip()


def split_level2(body: str) -> tuple[str, list[tuple[str, str]]]:
    matches = list(re.finditer(r"^##\s+(.+)$", body, flags=re.M))
    if not matches:
        return body.strip(), []
    preamble = body[: matches[0].start()].strip()
    sections: list[tuple[str, str]] = []
    for idx, match in enumerate(matches):
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(body)
        sections.append((match.group(1).strip(), body[match.end() : end].strip()))
    return preamble, sections


def take_section(sections: list[tuple[str, str]], prefix: str) -> str:
    for heading, content in sections:
        if heading.startswith(prefix):
            return content.strip()
    return ""


def replace_keywords(abstract_body: str) -> str:
    return re.sub(
        r"关键词：.*",
        "关键词：分子性质预测；ADMET；适用域；验证集选择；多专家模型",
        abstract_body,
    )


def split_intro(intro_body: str) -> str:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", intro_body) if p.strip()]
    problem = paragraphs[1] if len(paragraphs) > 1 else ""
    solution = paragraphs[2] if len(paragraphs) > 2 else ""
    objective = (
        "本文的目标是在不依赖单一大规模 backbone 的前提下，建立一个可复现、可审计、适用域感知的多专家分子性质预测框架。"
        "研究假设是：如果所有候选模型、融合策略和 rescue heads 均通过 validation-only 规则进入 retained-best 决策，"
        "则框架能够在保持测试集独立性的同时，对不同 endpoint 的性能、可靠性和解释边界给出更稳定的判断。"
    )
    return "\n\n".join(
        [
            "# 第一章 Introduction（引言）",
            "本章介绍研究背景、当前分子性质预测面临的问题、已有解决方案以及本文的研究目标和基本假设。",
            "## 1.1 当前研究的问题",
            problem,
            "当前问题可以概括为三点：第一，随机划分下的平均分数难以代表真实药物发现中的结构外推难度；第二，ADMET endpoint 常受到标签噪声、类别不平衡和局部 activity cliff 的影响；第三，复杂模型、强 tabular baseline 与冻结预训练表征之间的优劣并不固定，需要在统一验证规则下比较。",
            "## 1.2 当前的解决方案",
            solution,
            "已有解决方案包括 scaffold split、external benchmark、uncertainty quantification、applicability-domain analysis、target transform、balanced undersampling ensemble、Top-K/stacking ensemble 以及基于图模型或分子语言模型的预训练表征。这些方法各有价值，但如果缺少统一的验证集治理流程，新增候选模型容易造成事后选择偏差。",
            "## 1.3 本文的目标和假设",
            objective,
        ]
    )


def split_literature_review(lit_body: str) -> str:
    _preamble, sections = split_level2(lit_body)
    s21 = take_section(sections, "2.1")
    s22 = take_section(sections, "2.2")
    s23 = take_section(sections, "2.3")
    s24 = take_section(sections, "2.4")
    s25 = take_section(sections, "2.5")
    return "\n\n".join(
        [
            "# 第二章 Literature Review（文献回顾）",
            "本章从理论背景、已有研究、研究空白和研究目的四个方面回顾相关工作，为后续方法设计和结果解释提供文献依据。",
            "## 2.1 介绍理论背景或模型",
            s21,
            s23,
            "## 2.2 介绍已有研究",
            s22,
            s24,
            "## 2.3 研究空白",
            s25,
            "综合已有研究，当前空白主要体现在：第一，许多论文仍以单一分数比较模型，缺少验证集治理和负结果保留；第二，外部 ADMET、结构偏移、校准和不平衡分类常被分散报告，难以形成统一证据链；第三，强 tabular baseline、冻结 embedding、target transform 和可解释性模块往往没有在同一候选池中接受同一规则审查。",
            "## 2.4 研究目的",
            "本文旨在构建一个以 validation-only selector 为核心的分子性质预测框架，将多视图表示、强基线模型、图模型、冻结预训练表征、融合策略、适用域分析和可解释性证据整合到统一流程中。研究目的不是提出无条件替代所有模型的新 backbone，而是检验在结构分布偏移和 ADMET 异质性场景下，验证集治理能否更稳健地选择可靠候选。",
        ]
    )


def split_table1_block(preamble: str) -> tuple[str, str]:
    lines = preamble.splitlines()
    start = next((i for i, line in enumerate(lines) if line.startswith("表 1.")), None)
    if start is None:
        return preamble.strip(), ""
    return "\n".join(lines[:start]).strip(), "\n".join(lines[start:]).strip()


def split_methodology(method_body: str) -> str:
    preamble, sections = split_level2(method_body)
    design_block, data_block = split_table1_block(preamble)
    old31 = take_section(sections, "3.1")
    old32 = take_section(sections, "3.2")
    old33 = take_section(sections, "3.3")
    old34 = take_section(sections, "3.4")
    old35 = take_section(sections, "3.5")
    old36 = take_section(sections, "3.6")
    return "\n\n".join(
        [
            "# 第三章 Methodology（研究方法/数据收集方法）",
            "本章说明研究设计、分析方法和数据来源。全部实验遵循 train/validation/test 分离原则，候选模型和融合策略只在验证集上选择，测试集仅用于策略冻结后的最终评估。",
            "## 3.1 研究的设计",
            design_block,
            old31,
            old33,
            "## 3.2 分析方法",
            old32,
            old34,
            old35,
            old36,
            "## 3.3 数据",
            "本研究使用公开分子性质预测数据集和公开 ADMET benchmark，包括 MoleculeNet、TDC ADMET、MoleculeACE、OpenADMET-ExpansionRx feasibility appendix 以及结构外推相关划分。数据均以 SMILES、任务标签、划分索引和 endpoint 元信息为基础组织。",
            data_block,
        ]
    )


def demote_results_headings(results_body: str) -> str:
    counter = {"n": 0}

    def repl(match: re.Match[str]) -> str:
        counter["n"] += 1
        title = re.sub(r"^4\.\d+\s*", "", match.group(1)).strip()
        return f"### 4.1.{counter['n']} {title}"

    return re.sub(r"^##\s+(.+)$", repl, results_body, flags=re.M)


def split_results(results_body: str) -> str:
    factual = demote_results_headings(results_body)
    analysis = (
        "从结果含义看，FZYC-Mol 的主要价值不在于某一个专家模型在所有任务上统一获胜，而在于 selector 能够在 endpoint-specific 条件下保留更合适的候选。"
        "MoleculeNet 结果显示，Lipo 和 FreeSolv 的性能瓶颈可以通过 targeted rescue 或低成本 rebuild 得到部分缓解；TDC ADMET 结果说明，外部 endpoint 的改善更依赖任务本身的标签分布、roughness 和结构外推难度。"
        "可靠性分析进一步表明，ROC-AUC、RMSE 或 MAE 只能描述平均性能，不能替代校准、PR-AUC、适用域、不确定性和案例级解释。"
    )
    return "\n\n".join(
        [
            "# 第四章 Results（数据结果）",
            "本章先按照图表和指标列出实验事实，再分析这些结果对模型性能、可靠性和适用边界的含义。",
            "## 4.1 讲述事实（没有分析）",
            "以下内容按照任务面板列出主要数据结果。图表和数值用于呈现实验事实，详细解释在 4.2 和第五章讨论中展开。",
            factual,
            "## 4.2 分析事实：分析结果的含义",
            analysis,
            "总体而言，positive retained-best 结果支持 validation-only selector 的可行性；未被接入的候选结果同样具有意义，因为它们说明新增模型并不应自动覆盖既有最优策略。这样的结果组织可以降低选择性报告风险，也更适合解释 ADMET endpoint 的异质性。",
        ]
    )


def split_discussion(discussion_body: str) -> str:
    _preamble, sections = split_level2(discussion_body)
    s51 = take_section(sections, "5.1")
    s52 = take_section(sections, "5.2")
    s53 = take_section(sections, "5.3")
    s54 = take_section(sections, "5.4")
    return "\n\n".join(
        [
            "# 第五章 Discussion（讨论）",
            "本章讨论实验结论与观察现象之间的联系，并进一步说明这些现象与分子机器学习理论、ADMET 可靠性和适用域分析之间的关系。",
            "## 5.1 讨论结论与现象之间的联系",
            s51,
            "## 5.2 讨论现象与理论之间的因果关系",
            s52,
            s53,
            "## 5.3 结合数据-现象-理论得出研究结论",
            s54,
            "结合数据、现象和理论可以得到本文的核心判断：在分子性质预测中，模型规模并不是唯一决定因素；验证集治理、结构外推评估、可靠性指标和可解释性证据共同决定模型是否适合进入最终报告。",
        ]
    )


def split_conclusion(conclusion_body: str) -> str:
    parts = conclusion_body.split("未来工作可沿三条路线推进。", 1)
    summary = parts[0].strip()
    future = "未来工作可沿三条路线推进。" + parts[1].strip() if len(parts) > 1 else ""
    limitation = (
        "本研究仍存在不足。首先，当前实验均基于公开数据集和计算评估，缺少湿实验验证；其次，部分 foundation model 或 TabPFNv2 候选受到授权、缓存和计算资源限制，尚未形成完整正式对照；第三，motif attribution 与 fragment enrichment 主要提供关联性解释，不能直接证明因果机制。"
    )
    return "\n\n".join(
        [
            "# 第六章 Conclusion（总结/展望）",
            "## 6.1 总结本研究",
            summary,
            "## 6.2 指出不足，讨论原因",
            limitation,
            "这些不足的原因主要来自数据可得性、公开 benchmark 与真实药物发现流程之间的差距，以及大模型实验在算力和复现条件上的成本。本文因此将部分未完成或未接入结果保留为 appendix 或 limitation，而不将其包装为确定性结论。",
            "## 6.3 展望日后的研究",
            future,
        ]
    )


def split_acknowledgement(ack_body: str) -> str:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", ack_body) if p.strip()]
    people = paragraphs[0] if paragraphs else ""
    funding = paragraphs[1] if len(paragraphs) > 1 else "正式投稿前，可根据实际情况补充具体基金、平台或计算资源支持。"
    return "\n\n".join(
        [
            "# 第七章 Acknowledgement（致谢）",
            "## 7.1 致谢相关人员的贡献",
            people,
            "## 7.2 表示本研究的基金支持（可选）",
            funding,
        ]
    )


def split_declarations(decl_body: str) -> str:
    conflict = re.search(r"Declaration of Conflicting Interests：(.+?)(?:\n\n|$)", decl_body, flags=re.S)
    funding = re.search(r"Funding：(.+?)(?:\n\n|$)", decl_body, flags=re.S)
    conflict_text = conflict.group(1).strip() if conflict else "作者声明，目前未发现与本文研究内容直接相关的商业、财务或个人利益冲突。"
    funding_text = funding.group(1).strip() if funding else "当前稿件未提供具体基金项目信息。正式投稿前应补充基金名称、项目编号和资助机构。"
    return "\n\n".join(
        [
            "# 第八章 Declaration of Conflicting Interests / Funding（利益冲突声明/基金）",
            "## 8.1 声明本研究没有与任何单位或个人有利益冲突",
            conflict_text,
            "## 8.2 项目基金的支持",
            funding_text,
        ]
    )


def build_markdown() -> str:
    text = SRC_MD.read_text(encoding="utf-8")
    title = text[: text.index("# 摘要")].strip()
    title = re.sub(r"完整中文初稿（.+?）", "初稿-1（SSCI通用框架版，2026-06-05）", title)
    abstract = replace_keywords(chapter(text, "# 摘要"))
    refs = chapter(text, "# 第九章 References（参考文献）")
    refs = refs.replace(
        "[7] Gadaleta D, et al. Benchmarking predictive models for pKa, lipophilicity, intrinsic clearance and aqueous solubility in the context of drug discovery. Journal of Cheminformatics, 2024, 16: 145. DOI: 10.1186/s13321-024-00931-z.",
        "[7] Gadaleta D, et al. Comprehensive benchmarking of computational tools for predicting toxicokinetic and physicochemical properties of chemicals. Journal of Cheminformatics, 2024, 16: 145. DOI: 10.1186/s13321-024-00931-z.",
    )
    parts = [
        title,
        "# 摘要",
        abstract,
        split_intro(chapter(text, "# 第一章 Introduction（引言）")),
        split_literature_review(chapter(text, "# 第二章 Literature Review（文献回顾）")),
        split_methodology(chapter(text, "# 第三章 Methodology（研究方法/数据收集方法）")),
        split_results(chapter(text, "# 第四章 Results（数据结果）")),
        split_discussion(chapter(text, "# 第五章 Discussion（讨论）")),
        split_conclusion(chapter(text, "# 第六章 Conclusion（总结/展望）")),
        split_acknowledgement(chapter(text, "# 第七章 Acknowledgement（致谢）")),
        split_declarations(chapter(text, "# 第八章 Declaration of Conflicting Interests / Funding（利益冲突声明/基金）")),
        "# 第九章 References（参考文献）",
        "## 9.1 通篇引用的文献，按照出现顺序列出",
        refs,
    ]
    markdown = "\n\n".join(part.strip() for part in parts if part.strip()) + "\n"
    markdown = re.sub(r"(?i)^#.*high\s*lights.*\n.*?(?=^#|\Z)", "", markdown, flags=re.M | re.S)
    markdown = markdown.replace("亮点", "")
    markdown = markdown.replace("reports/manuscript_figures_hires", "reports/manuscript_figures_nature_style")
    return markdown


def add_paragraph(doc: Document, text: str, index: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    size = 10.5
    bold = False
    color = "111827"
    if index == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 16
        bold = True
    elif index == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 10.5
        color = "475569"
    elif re.match(r"^图\s*\d+\.", text):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 9.0
        bold = True
        color = "334155"
    elif re.match(r"^表\s*\d+\.", text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 9.2
        bold = True
    elif text.startswith("注："):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 8.2
        color = "64748B"
    elif re.match(r"^\[\d+\]\s+", text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 9.0
    else:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def markdown_to_docx(markdown: str) -> None:
    doc = Document()
    setup_doc(doc)
    lines = markdown.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line:
            idx += 1
            continue
        if line.startswith("!["):
            add_image(doc, line)
            idx += 1
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            p = doc.add_heading("", level=min(level, 3))
            run = p.add_run(text)
            set_font(run, size={1: 16, 2: 13.5, 3: 12}.get(level, 11), bold=True)
            idx += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            if len(table_lines) >= 2:
                add_table(doc, table_lines)
            continue
        add_paragraph(doc, line, idx)
        idx += 1
    doc.save(OUT_DOCX)


def main() -> None:
    markdown = build_markdown()
    OUT_MD.write_text(markdown, encoding="utf-8")
    markdown_to_docx(markdown)
    print(f"wrote {OUT_MD.relative_to(ROOT)}")
    print(f"wrote {OUT_DOCX.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
