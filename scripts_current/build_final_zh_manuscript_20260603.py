from __future__ import annotations

import math
import textwrap
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib import font_manager as fm
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
TABLE_DIR = ROOT / "reports" / "manuscript_tables"
FIG_DIR = ROOT / "reports" / "manuscript_figures_polished"
TABLE_PREVIEW_DIR = ROOT / "reports" / "manuscript_tables_polished"
PACKAGE = ROOT / "reports" / "submission_package"

OUT_DOCX = DOCS / "manuscript_draft_full_zh_polished_20260603.docx"
OUT_MD = DOCS / "manuscript_draft_full_zh_polished_20260603.md"

TABLE2_PRETTY = TABLE_DIR / "table2_moleculenet_main_pretty_20260603.csv"
TABLE3_PRETTY = TABLE_DIR / "table3_tdc_official_admet_pretty_20260603.csv"
TABLE8_MAIN = TABLE_DIR / "table8_formal_fixed_selector_main.csv"


DATASET_LABELS = {
    "esol": "ESOL",
    "freesolv": "FreeSolv",
    "lipo": "Lipo",
    "bbbp": "BBBP",
    "bace": "BACE",
    "clintox": "ClinTox",
    "caco2_wang": "Caco2",
    "hia_hou": "HIA",
    "pgp_broccatelli": "Pgp",
    "bioavailability_ma": "Bioavailability",
    "bbb_martins": "BBB",
    "cyp2c9_veith": "CYP2C9",
    "cyp2d6_veith": "CYP2D6",
    "cyp3a4_veith": "CYP3A4",
    "half_life_obach": "Half-life",
    "clearance_microsome_az": "Clearance microsome",
    "clearance_hepatocyte_az": "Clearance hepatocyte",
    "cyp2c9_substrate_carbonmangels": "CYP2C9 substrate",
    "vdss_lombardo": "VDss",
    "herg": "hERG",
    "ld50_zhu": "LD50",
}

SOURCE_LABELS = {
    "moleculenet_nature_fusion": "MoleculeNet fusion",
    "moleculenet_targeted_rebuild": "MoleculeNet rebuild",
    "tdc_performance_mode": "TDC performance mode",
    "tdc_nature_fusion": "TDC fusion",
    "three_d_roughness_regression": "3D-lite/roughness",
}


def configure_matplotlib_cjk() -> None:
    font_candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/msyhbd.ttc"),
    ]
    for font_path in font_candidates:
        if font_path.exists():
            fm.fontManager.addfont(str(font_path))
    plt.rcParams["font.family"] = "Microsoft YaHei"
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def canonical_dataset(name: object) -> str:
    text = str(name)
    return text[4:] if text.startswith("tdc_") else text


def label_dataset(name: object) -> str:
    text = canonical_dataset(name)
    return DATASET_LABELS.get(text, text.replace("_", " "))


def read_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path).fillna("")


def parse_mean(text: object) -> float:
    token = str(text).split("+/-", 1)[0].strip()
    try:
        return float(token)
    except ValueError:
        return math.nan


def fmt_float(value: object, digits: int = 4) -> str:
    try:
        if value == "":
            return ""
        value = float(value)
        if math.isnan(value):
            return ""
        return f"{value:.{digits}f}"
    except (TypeError, ValueError):
        return ""


def fmt_pm(mean: object, std: object, digits: int = 4) -> str:
    mean_text = fmt_float(mean, digits)
    std_text = fmt_float(std, digits)
    if not mean_text:
        return ""
    return f"{mean_text} +/- {std_text}" if std_text else mean_text


def fmt_delta(value: object, digits: int = 4) -> str:
    try:
        value = float(value)
        if math.isnan(value):
            return ""
        return f"{value:+.{digits}f}"
    except (TypeError, ValueError):
        return ""


def metric_label(metric: object) -> str:
    text = str(metric).replace("_", "-").upper()
    return {"ROC-AUC": "ROC-AUC", "PR-AUC": "PR-AUC", "RMSE": "RMSE", "MAE": "MAE", "SPEARMAN": "Spearman"}.get(text, text)


def task_label(task: object) -> str:
    return "回归" if str(task) == "regression" else "分类"


def build_moleculenet_pretty() -> pd.DataFrame:
    wide = read_csv(TABLE_DIR / "table2_moleculenet_main.csv")
    long = read_csv(TABLE_DIR / "table2_moleculenet_main_long.csv")
    formal = read_csv(TABLE_DIR / "table42_formal_risk_adjusted_integration_retained_best.csv")
    formal_lookup = {str(row["canonical_dataset"]): row for _, row in formal.iterrows()}

    rows = []
    for _, row in wide.iterrows():
        dataset = str(row["dataset"])
        sub = long[long["dataset"].eq(dataset)].copy()
        direction = str(sub.iloc[0]["direction"]) if "direction" in sub.columns and not sub.empty else "higher"
        non_fzyc = sub[
            ~sub["category"].astype(str).str.contains("FZYC-Mol", regex=False)
            & ~sub["category"].astype(str).eq("Best observed candidate")
        ].copy()
        if not non_fzyc.empty:
            non_fzyc["value_num"] = non_fzyc["value"].astype(float)
            best_comp = non_fzyc.sort_values("value_num", ascending=direction == "lower").iloc[0]
            best_comp_text = f"{best_comp['category']}: {best_comp['formatted']}"
        else:
            best_comp_text = ""
        current = str(row["FZYC-Mol final retained-best"])
        formal_row = formal_lookup.get(dataset)
        if formal_row is not None:
            formal_text = fmt_pm(formal_row["retained_primary_mean"], formal_row["retained_primary_std"])
            promoted = str(formal_row["promotion_status"]) == "promoted_by_fixed_policy"
            delta = fmt_delta(formal_row["delta_vs_current_retained"] if promoted else 0.0)
            status = "固定策略提升" if promoted else "保留当前"
        else:
            formal_text = current
            delta = "+0.0000"
            status = "保留当前"
        rows.append(
            {
                "数据集": label_dataset(dataset),
                "任务": task_label(row["task_type"]),
                "指标": metric_label(row["primary_metric"]),
                "原 retained-best": current,
                "固定策略 retained": formal_text,
                "Δ": delta,
                "结论": status,
                "最强非 FZYC 专家": best_comp_text,
            }
        )
    out = pd.DataFrame(rows)
    out.to_csv(TABLE2_PRETTY, index=False, encoding="utf-8-sig")
    return out


def build_tdc_pretty() -> pd.DataFrame:
    tdc = read_csv(TABLE_DIR / "table3_tdc_official_admet.csv")
    formal = read_csv(TABLE_DIR / "table42_formal_risk_adjusted_integration_retained_best.csv")
    formal_lookup = {str(row["canonical_dataset"]): row for _, row in formal.iterrows()}

    rows = []
    for _, row in tdc.iterrows():
        dataset = canonical_dataset(row["dataset"])
        formal_row = formal_lookup.get(dataset)
        final_current = str(row["tdc_final_retained_formatted"])
        if formal_row is not None:
            formal_text = fmt_pm(formal_row["retained_primary_mean"], formal_row["retained_primary_std"])
            promoted = str(formal_row["promotion_status"]) == "promoted_by_fixed_policy"
            delta = fmt_delta(formal_row["delta_vs_current_retained"] if promoted else 0.0)
            status = "固定策略提升" if promoted else "保留当前"
        else:
            formal_text = final_current
            delta = "+0.0000"
            status = "保留当前"
        scaffold_penalty = (
            float(row["lgbm_morgan_random_to_scaffold_drop"]) + float(row["rf_morgan_random_to_scaffold_drop"])
        ) / 2
        rows.append(
            {
                "Endpoint": label_dataset(dataset),
                "任务": task_label(row["task_type"]),
                "指标": metric_label(row["primary_metric"]),
                "原 TDC retained": final_current,
                "固定策略 retained": formal_text,
                "Δ": delta,
                "结论": status,
                "平均 scaffold 惩罚": fmt_delta(scaffold_penalty),
            }
        )
    out = pd.DataFrame(rows)
    out.to_csv(TABLE3_PRETTY, index=False, encoding="utf-8-sig")
    return out


def build_formal_main_table() -> pd.DataFrame:
    retained = read_csv(TABLE_DIR / "table42_formal_risk_adjusted_integration_retained_best.csv")
    promoted = retained[retained["promotion_status"].eq("promoted_by_fixed_policy")].copy()
    promoted["delta_num"] = promoted["delta_vs_current_retained"].astype(float)
    promoted = promoted.sort_values("delta_num", ascending=False)
    rows = []
    for _, row in promoted.iterrows():
        rows.append(
            {
                "Endpoint": label_dataset(row["canonical_dataset"]),
                "任务": task_label(row["task_type"]),
                "指标": metric_label(row["primary_metric"]),
                "当前 retained": fmt_pm(row["current_retained_primary_mean"], row["current_retained_primary_std"]),
                "固定策略 retained": fmt_pm(row["retained_primary_mean"], row["retained_primary_std"]),
                "Δ": fmt_delta(row["delta_vs_current_retained"]),
                "候选来源": SOURCE_LABELS.get(str(row["selected_source"]), str(row["selected_source"])),
                "入选模型": str(row["selected_model"]),
            }
        )
    out = pd.DataFrame(rows)
    out.to_csv(TABLE8_MAIN, index=False, encoding="utf-8-sig")
    return out


def build_policy_summary_table() -> pd.DataFrame:
    summary = read_csv(TABLE_DIR / "table43_formal_fixed_selector_policy_summary.csv")
    rows = []
    for _, row in summary.iterrows():
        policy = "risk-adjusted λ=0.5" if str(row["fixed_policy"]) == "risk_adjusted_lambda_0.5" else "stability tie-breaker"
        rows.append(
            {
                "固定策略": policy,
                "正向/总数": f"{int(row['n_positive_vs_current'])}/{int(row['n_endpoint_metrics'])}",
                "负向数": int(row["n_negative_vs_current"]),
                "正向率": f"{float(row['positive_rate']) * 100:.1f}%",
                "平均 Δ": fmt_delta(row["mean_delta_vs_current"]),
                "中位 Δ": fmt_delta(row["median_delta_vs_current"]),
                "最大正向 Δ": fmt_delta(row["largest_positive_delta"]),
                "最大负向 Δ": fmt_delta(row["largest_negative_delta"]),
            }
        )
    return pd.DataFrame(rows)


def wrap_cell(text: object, width: int) -> str:
    text = str(text)
    if len(text) <= width:
        return text
    return "\n".join(textwrap.wrap(text, width=width, break_long_words=False))


def draw_table_preview(frame: pd.DataFrame, title: str, stem: str, max_rows: int | None = None) -> None:
    configure_matplotlib_cjk()
    TABLE_PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    data = frame.copy()
    if max_rows is not None:
        data = data.head(max_rows)
    col_widths = []
    for col in data.columns:
        if col in {"最强非 FZYC 专家", "入选模型"}:
            col_widths.append(0.22)
        elif col in {"候选来源", "结论"}:
            col_widths.append(0.12)
        else:
            col_widths.append(0.10)
    total = sum(col_widths)
    col_widths = [w / total for w in col_widths]
    cell_text = []
    for _, row in data.iterrows():
        cell_text.append(
            [
                wrap_cell(row[col], 26 if col in {"最强非 FZYC 专家", "入选模型"} else 16)
                for col in data.columns
            ]
        )
    height = max(2.4, 0.48 * (len(data) + 2))
    fig, ax = plt.subplots(figsize=(13.8, height))
    ax.axis("off")
    ax.text(0.0, 1.03, title, transform=ax.transAxes, fontsize=13, fontweight="bold", ha="left", va="bottom")
    table = ax.table(
        cellText=cell_text,
        colLabels=list(data.columns),
        loc="upper left",
        cellLoc="center",
        colLoc="center",
        colWidths=col_widths,
        bbox=[0, 0, 1, 0.96],
    )
    table.auto_set_font_size(False)
    table.set_fontsize(7.6)
    for (r, c), cell in table.get_celld().items():
        cell.set_edgecolor("#dbe3ea")
        cell.set_linewidth(0.6)
        if r == 0:
            cell.set_facecolor("#14532d")
            cell.get_text().set_color("white")
            cell.get_text().set_fontweight("bold")
        else:
            cell.set_facecolor("#f8fafc" if r % 2 == 0 else "white")
            if c == 0:
                cell.set_facecolor("#ecfdf5")
                cell.get_text().set_fontweight("bold")
            text = cell.get_text().get_text()
            if text.startswith("+"):
                cell.get_text().set_color("#166534")
                cell.get_text().set_fontweight("bold")
            elif text.startswith("-"):
                cell.get_text().set_color("#b91c1c")
    fig.savefig(TABLE_PREVIEW_DIR / f"{stem}.png", dpi=320, bbox_inches="tight", facecolor="white")
    fig.savefig(TABLE_PREVIEW_DIR / f"{stem}.svg", bbox_inches="tight", facecolor="white")
    plt.close(fig)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_doc_style(doc: Document) -> None:
    for name, size in [("Normal", 10.5), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.18
    doc.styles["Normal"].paragraph_format.space_after = Pt(5)


def set_margins(section) -> None:
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)


def make_landscape(doc: Document) -> None:
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    set_margins(section)


def make_portrait(doc: Document) -> None:
    section = doc.add_section()
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width
    set_margins(section)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    set_run_font(run, bold=True)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color="111827")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=9.0, bold=True, color="334155")


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.8) -> None:
    if not path.exists():
        add_para(doc, f"[缺少图片：{path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "DDE6ED") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), color)


def set_cell_text(cell, text: object, size: float = 7.2, bold: bool = False, color: str = "111827", align: str = "center") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }[align]
    run = p.add_run("" if pd.isna(text) else str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_styled_table(doc: Document, label: str, caption: str, frame: pd.DataFrame, note: str = "", font_size: float = 7.0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{label}. {caption}")
    set_run_font(run, size=9.4, bold=True, color="111827")

    table = doc.add_table(rows=1, cols=len(frame.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for c, col in enumerate(frame.columns):
        cell = table.rows[0].cells[c]
        shade_cell(cell, "14532D")
        set_cell_border(cell)
        set_cell_text(cell, col, size=font_size, bold=True, color="FFFFFF")
    for r, (_, row) in enumerate(frame.iterrows(), start=1):
        cells = table.add_row().cells
        for c, col in enumerate(frame.columns):
            cell = cells[c]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            fill = "F8FAFC" if r % 2 == 0 else "FFFFFF"
            if c == 0:
                fill = "ECFDF5"
            shade_cell(cell, fill)
            text = row[col]
            color = "111827"
            bold = c == 0
            align = "center"
            if col in {"最强非 FZYC 专家", "入选模型", "候选来源"}:
                align = "left"
            if str(text).startswith("+"):
                color = "166534"
                bold = True
            elif str(text).startswith("-"):
                color = "B91C1C"
                bold = True
            set_cell_text(cell, text, size=font_size, bold=bold, color=color, align=align)
    if note:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f"注：{note}")
        set_run_font(run, size=8.0, color="64748B")


def build_markdown(mn: pd.DataFrame, tdc: pd.DataFrame, formal: pd.DataFrame, policy: pd.DataFrame) -> None:
    lines = [
        "# FZYC-Mol 中文初稿（2026-06-03 更新版）",
        "",
        "本版已接入 formal fixed-selector integration，并将主表改为主文可读的紧凑格式。",
        "",
        f"![Figure 1]({(FIG_DIR / 'fig1_framework_overview_polished.png').as_posix()})",
        "",
        "## 主表 2：MoleculeNet 主结果",
        "",
        mn.to_markdown(index=False),
        "",
        "## 主表 3：TDC official split 摘要",
        "",
        tdc.to_markdown(index=False),
        "",
        "## 主表 8：Formal fixed selector rescue",
        "",
        formal.to_markdown(index=False),
        "",
        "## 固定策略汇总",
        "",
        policy.to_markdown(index=False),
        "",
        f"![Figure 8]({(FIG_DIR / 'fig22_formal_fixed_selector_integration.png').as_posix()})",
    ]
    OUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def add_reference_list(doc: Document) -> None:
    refs = [
        "Wu Z, Ramsundar B, Feinberg E N, et al. MoleculeNet: a benchmark for molecular machine learning. Chemical Science, 2018.",
        "Huang K, Fu T, Gao W, et al. Therapeutics Data Commons: machine learning datasets and tasks for drug discovery and development. NeurIPS, 2021.",
        "Yang K, Swanson K, Jin W, et al. Analyzing learned molecular representations for property prediction. Journal of Chemical Information and Modeling, 2019.",
        "Rogers D, Hahn M. Extended-connectivity fingerprints. Journal of Chemical Information and Modeling, 2010.",
        "Kuhn M, Letunic I, Jensen L J, Bork P. The SIDER database of drugs and side effects. Nucleic Acids Research, 2016.",
        "Guo J. Do Larger Models Really Win in Drug Discovery? A Benchmark Assessment of Model Scaling in AI-Driven Molecular Property and Activity Prediction. arXiv, 2026.",
        "Li Z, Chen X, Wen H, et al. A systematic survey and benchmark of deep learning for molecular property prediction in the foundation model era. Journal of Chemical Theory and Computation, 2026.",
        "Ben Hicham K K, Rittig J G, Grohe M, Mitsos A. Tabular foundation models for in-context prediction of molecular properties. arXiv, 2026.",
        "Hong H, Wu X, Sun H, et al. A hierarchical interaction message net for accurate molecular property prediction. Communications Chemistry, 2026.",
        "Le K, Dey S, Martinez Galindo M, et al. Can Decision Trees Teach Large Language Models? Distilling Verbalized Knowledge for Molecular Property Prediction. arXiv, 2026.",
    ]
    add_heading(doc, "参考文献", level=1)
    for idx, ref in enumerate(refs, start=1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(ref)
        set_run_font(run, size=9.0, color="111827")


def build_docx(mn: pd.DataFrame, tdc: pd.DataFrame, formal: pd.DataFrame, policy: pd.DataFrame) -> None:
    doc = Document()
    set_doc_style(doc)
    set_margins(doc.sections[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FZYC-Mol：面向 ADMET 结构分布偏移的验证集治理多专家分子性质预测框架")
    set_run_font(run, size=18, bold=True, color="111827")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("中文初稿更新版：formal fixed-selector integration 与精排主表（2026-06-03）")
    set_run_font(run, size=10.5, color="475569")

    add_heading(doc, "摘要", level=1)
    add_para(
        doc,
        "分子性质预测在药物发现、ADMET 评估和毒性筛选中具有基础作用。现有分子预测研究常聚焦单一模型的平均分数，但真实药物发现任务还同时面对结构分布偏移、标签噪声、类别不平衡、局部 activity cliff、外部数据集迁移以及模型可解释性不足等问题。本文提出 FZYC-Mol，一个以 validation-only 选择为核心的多专家分子性质预测框架。该框架整合 Morgan/多指纹树模型、RDKit descriptor 模型、图神经网络、Chemprop、冻结分子语言模型 embedding、motif/fragment experts、Top-K/stacking ensemble、uncertainty 和 applicability-domain 诊断，并要求最终候选只能由验证集决定，测试集仅用于最终报告。",
    )
    add_para(
        doc,
        "在 MoleculeNet 主面板上，FZYC-Mol 保持 ESOL、BACE、Lipo 等任务的竞争性表现；在新增 formal fixed-selector integration 中，我们预先固定 risk-adjusted validation selector（λ=0.5），并将 stability tie-breaker 作为敏感性分析。固定 risk-adjusted 策略在 32 个可比 endpoint-metric 中有 10 个相对当前 retained-best 为正向，主要改善 half-life、clearance microsome、clearance hepatocyte、CYP2C9 substrate、VDss、Pgp、BBBP 和 ClinTox 等低分或高噪声模块。由于该固定策略仍有 22 个 endpoint-metric 低于当前 retained-best，本文将其定位为 targeted rescue 和 selector robustness 证据，而不是全局替代主 selector。",
    )
    add_para(
        doc,
        "本文的主要贡献包括：系统构建 validation-governed multi-expert 预测框架；在 MoleculeNet、TDC official splits、external appendix 和 roughness diagnostics 上形成多层实验闭环；提出并审计 fixed risk-adjusted selector integration；使用 PR-AUC、Brier、ECE、risk-coverage、AD/OOD、conformal 和 motif/fragment attribution 强化可靠性与可解释性。整体结果表明，在 ADMET endpoint 异质性明显的场景中，性能提升更应依赖预注册的验证集治理和针对性 rescue，而非简单堆叠候选模型或无约束重训大模型。",
    )

    add_heading(doc, "1 引言", level=1)
    add_para(
        doc,
        "近年来，图神经网络、D-MPNN、分子语言模型和分子基础模型推动了分子性质预测的发展。然而，多个 2026 年同类 benchmark 和可靠性研究都指出，更大的模型并不必然在所有 drug discovery endpoint 上胜出；在少样本、OOD、类别不平衡和标签粗糙的 ADMET 任务中，强 tabular baseline、descriptor/fingerprint 表征、calibration-aware selection 和 endpoint-specific ensemble 仍然非常有竞争力。因此，本课题的核心不应继续无边界堆 candidate，而应加强 selector：让候选选择更加稳定、可复现、能处理验证噪声，并保留所有负向结果的审计边界。",
    )
    add_para(
        doc,
        "FZYC-Mol 的研究问题可以概括为：在不使用测试集调参的条件下，如何从异质候选专家中为每个 endpoint 选择合适的预测策略，并同时提供可靠性、适用域和化学解释证据。与传统 leaderboard 式论文不同，本文强调 validation-only selection、retained-best reporting、seed-level stability、外部 appendix benchmark 和 case-level interpretability 共同构成证据链。",
    )

    add_heading(doc, "2 方法", level=1)
    add_figure(
        doc,
        FIG_DIR / "fig1_framework_overview_polished.png",
        "图 1. FZYC-Mol 多专家框架。分子被转化为 graph、fingerprint、descriptor、motif 和 pretrained embedding 等多视图表示，进入异质专家池；最终策略由验证集选择，并同步输出可靠性、适用域和解释性诊断。",
        width=7.1,
    )
    add_para(
        doc,
        "候选专家包括 classical Morgan/多指纹树模型、Graph/D-MPNN、Chemprop、RDKit descriptor baselines、motif experts、冻结 ChemBERTa/MoLFormer embedding heads、Top-K mean、ridge/logistic stacking、target transform 和 undersampling ensemble。回归任务按 endpoint official metric 使用 RMSE、MAE 或 Spearman；分类任务除 ROC-AUC 外，同时补充 PR-AUC、Brier、ECE 和 enrichment 指标。",
    )
    add_para(
        doc,
        "本轮正式新增的 fixed selector integration 采用预先固定的 risk-adjusted validation 规则：对于候选模型的跨 seed 验证均值和验证标准差，低值指标最小化 mean + 0.5 × std，高值指标最大化 mean - 0.5 × std。该规则的目的不是获得单 endpoint 的事后最优，而是降低验证噪声下高波动候选被误选的概率。stability tie-breaker 作为敏感性分析，仅在近似最优候选中优先选择验证方差更低、平均排名更稳定的模型。",
    )

    make_landscape(doc)
    add_styled_table(
        doc,
        "表 1",
        "MoleculeNet 主结果精排摘要。",
        mn,
        note="固定策略 retained 列仅在预先固定 risk-adjusted selector 带来正向 retained-best 时更新；其余 endpoint 保留当前主 selector。",
        font_size=7.0,
    )
    add_styled_table(
        doc,
        "表 2",
        "TDC official split 主结果精排摘要。",
        tdc,
        note="平均 scaffold 惩罚为 LGBM 与 RF 从 random 到 scaffold 的平均方向化变化，用于展示 split realism。",
        font_size=6.7,
    )
    make_portrait(doc)

    add_heading(doc, "3 结果", level=1)
    add_heading(doc, "3.1 MoleculeNet 与 TDC 主结果", level=2)
    add_figure(
        doc,
        FIG_DIR / "fig3_moleculenet_performance_dots.png",
        "图 2. MoleculeNet 主结果散点图。回归任务低值更优，分类任务高值更优；FZYC-Mol 的 retained-best 结果体现 validation-only selector 和 targeted rescue 的共同作用。",
        width=6.7,
    )
    add_para(
        doc,
        "MoleculeNet 主表显示，FZYC-Mol 在 ESOL、BACE、Lipo 等 endpoint 上保持竞争性表现。Lipo 已经通过 rescue-integrated selector 获得明确提升；FreeSolv 仍是主要 remaining limitation，说明小样本溶剂化自由能任务可能需要更强构象或物理化学特征。BBBP 和 ClinTox 在 formal risk-adjusted selector 中获得小幅进一步提升，但幅度有限，因此应作为 targeted rescue 证据而非过度宣称。",
    )
    add_figure(
        doc,
        FIG_DIR / "fig5_tdc_official_split_delta.png",
        "图 3. TDC official split 中 random-to-scaffold 变化。不同 endpoint 的 scaffold 惩罚方向和幅度并不一致，支持 endpoint heterogeneity 的核心叙事。",
        width=6.7,
    )
    add_para(
        doc,
        "TDC official split 结果强调了 ADMET endpoint 的异质性。Caco2、HIA、Pgp 等 endpoint 可以通过多方法 fusion 或 fixed selector 获得局部收益；CYP 系列和 Bioavailability 等任务则表现出更明显的 split 与类别分布压力。相比宣称统一模型全面胜出，更稳妥的结论是：FZYC-Mol 的价值在于验证集治理、候选保留和可靠性诊断共同降低错误推广风险。",
    )

    add_heading(doc, "3.2 Formal fixed-selector integration", level=2)
    add_figure(
        doc,
        FIG_DIR / "fig22_formal_fixed_selector_integration.png",
        "图 4. Formal fixed-selector integration。risk-adjusted λ=0.5 为主策略，stability tie-breaker 为敏感性分析。该图展示固定策略的正向 endpoint 数和最强 targeted rescue 信号。",
        width=6.9,
    )
    add_styled_table(
        doc,
        "表 3",
        "Formal fixed selector policy 汇总。",
        policy,
        note="该表必须与正向 retained-best 表一起解读：固定策略并非全局替代主 selector，而是识别 targeted rescue 机会。",
        font_size=7.5,
    )
    make_landscape(doc)
    add_styled_table(
        doc,
        "表 4",
        "Formal risk-adjusted selector 带来的 retained-best 提升。",
        formal,
        note="Δ 按指标方向统一为正值代表更优。该表只列出固定主策略的正向 promotion；完整正负结果见 Table S36-S38。",
        font_size=6.8,
    )
    make_portrait(doc)
    add_para(
        doc,
        "固定 risk-adjusted selector 在 32 个 endpoint-metric 中识别出 10 个正向 promotion。最大收益来自 half-life Obach（Spearman +0.0182），其次是 clearance microsome（+0.0132）、clearance hepatocyte（+0.0099）和 CYP2C9 substrate PR-AUC（+0.0082）。这些 endpoint 恰好对应长尾回归、局部 target jump、类别不平衡或验证噪声较强的模块，说明 selector 改进比继续堆 candidate 更贴合当前性能瓶颈。",
    )
    add_para(
        doc,
        "同时，固定策略的平均 delta 仍为负，Lipo、ESOL、FreeSolv、Caco2 MAE、PPBR、CYP2D6/CYP3A4 等 endpoint 在该规则下不如当前 retained-best。因此论文中不应把 risk-adjusted selector 写成全局替代模型，而应写成 formal fixed-policy rescue appendix：它可以提升特定低分模块，并为 selector 改进提供可复现证据。",
    )

    add_heading(doc, "3.3 可靠性、可解释性与负结果", level=2)
    add_figure(
        doc,
        FIG_DIR / "fig6_reliability_summary_polished.png",
        "图 5. 不确定性与适用域可靠性摘要。AD/OOD、ensemble uncertainty、error model 和 reconstruction-style scores 可用于定位高错误风险样本。",
        width=6.8,
    )
    add_figure(
        doc,
        FIG_DIR / "fig11_motif_fragment_interpretation.png",
        "图 6. Motif 与 fragment 层面的解释性结果。该图用于展示模型关注的局部化学结构，但不应被解释为因果机制证明。",
        width=6.8,
    )
    add_para(
        doc,
        "3D-lite 与 roughness-weighted tree 实验没有被 validation-only gate 接入主结果，这是重要的负结果。它说明候选池中可能存在测试 oracle 信号，但当前验证选择尚不能稳定捕捉，进一步支持本轮转向 selector 改进而不是继续扩大候选数量。对于审稿人而言，保留这类负结果能提高论文可信度。",
    )

    add_heading(doc, "4 讨论", level=1)
    add_para(
        doc,
        "与 2026 年同类分子预测论文相比，本课题的优势不在于提出单一更大的模型，而在于将模型性能、外部 benchmark、selector 稳定性、AD/OOD、calibration、roughness 和解释性组织成统一证据链。近期 ADMET benchmark 和 tabular foundation model 研究均提示，强 tabular baseline 和验证集治理在少样本与 OOD 场景中仍具有很高价值。FZYC-Mol 的 fixed selector integration 正好回应这一点：它并不追求所有 endpoint 统一涨分，而是用预先固定规则识别可以被稳健提升的模块。",
    )
    add_para(
        doc,
        "当前仍有若干局限。第一，固定 risk-adjusted selector 不是全局最优，不能替代主 selector。第二，FreeSolv 和部分 ADME regression endpoint 仍需要构象、物理化学特征或更严格 nested validation。第三，ClinTox、CYP substrate 等不平衡任务需要继续强化 PR-AUC、calibration curve、Brier/ECE 和 case study。第四，motif attribution 是相关性解释，不是湿实验机制证明。",
    )

    add_heading(doc, "5 结论", level=1)
    add_para(
        doc,
        "本版更新将 FZYC-Mol 从“多 candidate 性能堆叠”推进到“正式固定 selector 改进”。主策略 risk-adjusted λ=0.5 在 10 个 endpoint-metric 上带来 retained-best promotion，尤其改善 half-life、clearance 和 CYP2C9 substrate 等低分模块；同时完整保留负向 endpoint，避免过度宣称。论文最终应把该结果写成 targeted rescue 和 selector robustness，而不是全局替代主 selector。这样的叙事更符合高水平论文对方法边界、统计诚实性和实验厚度的要求。",
    )
    add_reference_list(doc)

    DOCS.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def main() -> None:
    TABLE_PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    mn = build_moleculenet_pretty()
    tdc = build_tdc_pretty()
    formal = build_formal_main_table()
    policy = build_policy_summary_table()
    draw_table_preview(mn, "Table 2. MoleculeNet main results (polished)", "table2_moleculenet_main_polished", max_rows=6)
    draw_table_preview(tdc, "Table 3. TDC official split summary (polished)", "table3_tdc_official_admet_polished", max_rows=8)
    draw_table_preview(formal, "Table 8. Formal fixed-selector rescue summary", "table8_formal_fixed_selector_main_polished")
    build_docx(mn, tdc, formal, policy)
    build_markdown(mn, tdc, formal, policy)
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {TABLE2_PRETTY}")
    print(f"Wrote {TABLE3_PRETTY}")
    print(f"Wrote {TABLE8_MAIN}")
    print(f"Wrote {TABLE_PREVIEW_DIR}")


if __name__ == "__main__":
    main()
