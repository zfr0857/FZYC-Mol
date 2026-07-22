from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIG_DIR = ROOT / "reports" / "manuscript_figures_polished"
HIRES_FIG_DIR = ROOT / "reports" / "manuscript_figures_hires"

BASE_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_academic_polished.md"
OUT_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_academic_integrated_ordered_3line.md"
OUT_DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_academic_integrated_ordered_3line.docx"
REPORT = DOCS / "manuscript_results_integration_restructure_report_20260604.md"


FIGURES = {
    "fig1": HIRES_FIG_DIR / "fig1_framework_overview_polished.png",
    "fig2": HIRES_FIG_DIR / "fig2_moleculenet_rank_heatmap_polished.png",
    "fig3": HIRES_FIG_DIR / "fig3_moleculenet_performance_dots.png",
    "fig4": HIRES_FIG_DIR / "fig5_tdc_official_split_delta.png",
    "fig5": HIRES_FIG_DIR / "fig4_split_realism_polished.png",
    "fig6": HIRES_FIG_DIR / "fig6_reliability_summary_polished.png",
    "fig7": HIRES_FIG_DIR / "fig11_motif_fragment_interpretation.png",
    "fig8": HIRES_FIG_DIR / "fig17_moleculenet_targeted_rebuild_decision.png",
    "fig8b": HIRES_FIG_DIR / "fig18_nature_multimethod_fusion_decision.png",
    "fig8c": HIRES_FIG_DIR / "fig19_tdc_nature_multimethod_fusion_decision.png",
    "fig9": HIRES_FIG_DIR / "fig15_external_appendix_retained_delta.png",
    "fig10": HIRES_FIG_DIR / "fig16_external_candidate_rank_cd.png",
    "fig11": HIRES_FIG_DIR / "fig22_formal_fixed_selector_integration.png",
    "fig12": HIRES_FIG_DIR / "fig24_strong_baseline_selector_governance.png",
    "fig23": HIRES_FIG_DIR / "fig23_fzyc_mol_model_structure.png",
    "fig20": HIRES_FIG_DIR / "fig20_3d_roughness_regression_gate.png",
}


REFERENCE_LIST = [
    "Wu Z, Ramsundar B, Feinberg E N, et al. MoleculeNet: a benchmark for molecular machine learning. Chemical Science, 2018, 9: 513-530. DOI: 10.1039/C7SC02664A.",
    "Huang K, Fu T, Gao W, et al. Therapeutics Data Commons: machine learning datasets and tasks for drug discovery and development. Proceedings of NeurIPS Datasets and Benchmarks, 2021.",
    "Zhao D, Zhu Y, Wu Z, et al. Revisiting ADMET prediction reliability under real-world challenges in the foundation model era. Journal of Cheminformatics, 2026. DOI: 10.1186/s13321-026-01217-2.",
    "Fraser J S, Edgar S, Handly L N, et al. Mapping the avoid-ome: a systematic open-science approach to predictive ADMET. Nature Communications, 2026, 17: 4644. DOI: 10.1038/s41467-026-73410-8.",
    "Dinh Pham L-H, Le M-T, Thai K-M. Improved ADME prediction by multitask pretraining on predicted data: insights from the ASAP-Polaris-OpenADMET blind challenge. Journal of Chemical Information and Modeling, 2026, 66: 395-405. DOI: 10.1021/acs.jcim.5c02030.",
    "Parrondo-Pizarro R, Lanini J, Rodriguez-Perez R. Uncertainty quantification in molecular machine learning for property predictions under data shifts. Journal of Chemical Information and Modeling, 2026, 66: 923-935. DOI: 10.1021/acs.jcim.5c02381.",
    "Gadaleta D, et al. Benchmarking predictive models for pKa, lipophilicity, intrinsic clearance and aqueous solubility in the context of drug discovery. Journal of Cheminformatics, 2024, 16: 145. DOI: 10.1186/s13321-024-00931-z.",
    "Swanson K, Walther P, Leitz J, et al. ADMET-AI: a machine learning ADMET platform for evaluation of large-scale chemical libraries. Bioinformatics, 2024, 40: btae416. DOI: 10.1093/bioinformatics/btae416.",
    "Xiong G, Wu Z, Yi J, et al. ADMETlab 2.0: an integrated online platform for accurate and comprehensive predictions of ADMET properties. Nucleic Acids Research, 2021, 49: W5-W14. DOI: 10.1093/nar/gkab255.",
    "van Tilborg D, Alenicheva A, Grisoni F. Exposing the limitations of molecular machine learning with activity cliffs. Journal of Chemical Information and Modeling, 2022, 62: 5938-5951.",
    "Hollmann N, Mueller S, Eggensperger K, Hutter F. TabPFN: a transformer that solves small tabular classification problems in a second. International Conference on Learning Representations, 2023.",
    "Hollmann N, Müller S, Purucker L, Krishnakumar A, Körfer M, Hoo S B, Schirrmeister R T, Hutter F. Accurate predictions on small data with a tabular foundation model. Nature, 2025, 637: 319-326. DOI: 10.1038/s41586-024-08328-6.",
    "Ben Hicham K K, Rittig J G, Grohe M, Mitsos A. Tabular foundation models for in-context prediction of molecular properties. arXiv:2604.16123, 2026.",
    "Guo J. Do larger models really win in drug discovery? A benchmark assessment of model scaling in AI-driven molecular property and activity prediction. arXiv:2604.26498, 2026.",
    "Li Z, Chen X, Wen H, et al. A systematic survey and benchmark of deep learning for molecular property prediction in the foundation model era. arXiv:2604.16586, 2026.",
    "Hong H, Wu X, Sun H, et al. A hierarchical interaction message net for accurate molecular property prediction. Communications Chemistry, 2026. DOI: 10.1038/s42004-026-01922-x.",
    "Guo Y, Luo M, Zhang W, et al. Few-shot molecular property optimization via a domain-specialized large language model. Chemical Science, 2026, 17: 4928-4941. DOI: 10.1039/D5SC08859C.",
    "Yang K, Swanson K, Jin W, et al. Analyzing learned molecular representations for property prediction. Journal of Chemical Information and Modeling, 2019, 59: 3370-3388.",
    "Rogers D, Hahn M. Extended-connectivity fingerprints. Journal of Chemical Information and Modeling, 2010, 50: 742-754.",
    "Landrum G. RDKit: open-source cheminformatics software. https://www.rdkit.org/.",
    "Bemis G W, Murcko M A. The properties of known drugs. 1. Molecular frameworks. Journal of Medicinal Chemistry, 1996, 39: 2887-2893.",
    "Degen J, Wegscheid-Gerlach C, Zaliani A, Rarey M. On the art of compiling and using drug-like chemical fragment spaces. ChemMedChem, 2008, 3: 1503-1507.",
    "Breiman L. Random forests. Machine Learning, 2001, 45: 5-32.",
    "Chen T, Guestrin C. XGBoost: a scalable tree boosting system. Proceedings of KDD, 2016.",
    "Ke G, Meng Q, Finley T, et al. LightGBM: a highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems, 2017.",
    "Prokhorenkova L, Gusev G, Vorobev A, Dorogush A V, Gulin A. CatBoost: unbiased boosting with categorical features. Advances in Neural Information Processing Systems, 2018.",
    "Geurts P, Ernst D, Wehenkel L. Extremely randomized trees. Machine Learning, 2006, 63: 3-42.",
    "Chithrananda S, Grand G, Ramsundar B. ChemBERTa: large-scale self-supervised pretraining for molecular property prediction. arXiv:2010.09885, 2020.",
    "Ross J, Belgodere B, Chenthamarakshan V, Padhi I, Mroueh Y, Das P. Large-scale chemical language representations capture molecular structure and properties. Nature Machine Intelligence, 2022, 4: 1256-1264. DOI: 10.1038/s42256-022-00580-7.",
    "Hu W, Liu B, Gomes J, et al. Strategies for pre-training graph neural networks. International Conference on Learning Representations, 2020.",
    "Rong Y, Bian Y, Xu T, et al. Self-supervised graph transformer on large-scale molecular data. Advances in Neural Information Processing Systems, 2020.",
    "Xiong Z, Wang D, Liu X, et al. Pushing the boundaries of molecular representation for drug discovery with graph attention mechanism. Journal of Medicinal Chemistry, 2020, 63: 8749-8760.",
    "Vovk V, Gammerman A, Shafer G. Algorithmic Learning in a Random World. Springer, 2005.",
    "Shafer G, Vovk V. A tutorial on conformal prediction. Journal of Machine Learning Research, 2008, 9: 371-421.",
    "Guo C, Pleiss G, Sun Y, Weinberger K Q. On calibration of modern neural networks. Proceedings of ICML, 2017.",
    "Gawlikowski J, Tassi C R N, Ali M, et al. A survey of uncertainty in deep neural networks. Artificial Intelligence Review, 2023, 56: 1513-1589.",
    "Sushko I, Novotarskyi S, Körner R, et al. Applicability domains for classification problems: benchmarking of distance to models for Ames mutagenicity set. Journal of Chemical Information and Modeling, 2010, 50: 2094-2111.",
    "Tropsha A. Best practices for QSAR model development, validation, and exploitation. Molecular Informatics, 2010, 29: 476-488.",
    "Sheridan R P. Time-split cross-validation as a method for estimating prospective prediction performance. Journal of Chemical Information and Modeling, 2013, 53: 783-790.",
    "Mayr A, Klambauer G, Unterthiner T, Hochreiter S. DeepTox: toxicity prediction using deep learning. Frontiers in Environmental Science, 2016, 3: 80.",
]


def read_base() -> str:
    text = BASE_MD.read_text(encoding="utf-8")
    text = text.replace(
        "完整中文初稿（Nature 方法融合、2026 文献更新与 MoleculeNet 补跑增强版，2026-06-02）",
        "完整中文初稿（实验融合整改版，2026-06-04）",
    )
    text = text.replace(
        "在同一套验证集规则决定是否接入主结果。这样既能回应",
        "在同一套验证集规则决定是否接入主结果。后续扩展的 rescue、rebuild、fusion、external appendix、fixed selector 和 3D-lite/roughness 结果均被整合到相应科学问题中，而不是作为独立追加实验处理。这样既能回应",
    )
    text = text.replace(
        "ClinTox、DILI、hERG 和 CYP substrate 因此被单独纳入不平衡分类增强。",
        "ClinTox、DILI、hERG 和 CYP substrate 因此在不平衡分类增强中重点报告。",
    )
    text = text.replace(
        "加入 rescue heads 后，Lipo 得到 validation-only 接受的定向提升；进一步允许低成本 targeted rebuild 后，FreeSolv 获得小幅 retained-best 改善；再加入 Nature-inspired 多方法融合后，BBBP 和 ClinTox 的分类结果也得到接入式提升，外部 TDC official panel 中 Caco2、HIA 和 Pgp 也获得 retained-best 增益。",
        "在同一 validation-only 候选池中，rescue heads、低成本 targeted rebuild 和 Nature-inspired 多方法融合分别在 Lipo、FreeSolv、BBBP/ClinTox 以及外部 TDC official panel 的 Caco2、HIA 和 Pgp 上形成 retained-best 增益。",
    )
    return text


def split_level2_results(text: str) -> dict[str, str]:
    start = text.index("# 3. 结果")
    end = text.index("# 4. 讨论")
    result = text[start:end]
    matches = list(re.finditer(r"(?m)^## 3\.[^\n]+", result))
    sections: dict[str, str] = {}
    for idx, match in enumerate(matches):
        sec_start = match.start()
        sec_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(result)
        heading = match.group(0).strip()
        sections[heading] = result[sec_start:sec_end].strip()
    return sections


def strip_heading(section: str) -> str:
    lines = section.splitlines()
    return "\n".join(lines[1:]).strip()


def extract_block(section: str, start_prefix: str, end_prefix: str | None = None) -> str:
    lines = section.splitlines()
    start = next((i for i, line in enumerate(lines) if line.startswith(start_prefix)), None)
    if start is None:
        return ""
    if end_prefix is None:
        end = len(lines)
    else:
        end = next((i for i in range(start + 1, len(lines)) if lines[i].startswith(end_prefix)), len(lines))
    return "\n".join(lines[start:end]).strip()


def extract_paragraph(section: str, prefix: str) -> str:
    paragraphs = [p.strip() for p in section.split("\n\n") if p.strip()]
    return next((p for p in paragraphs if p.startswith(prefix)), "")


def image_markdown(key: str, alt: str) -> str:
    path = FIGURES[key]
    return f"![{alt}]({path.as_posix()})" if path.exists() else ""


def clean_fragment(text: str) -> str:
    replacements = {
        "后续 targeted rebuild 可将 RMSE": "同一候选池中的 targeted rebuild 将 RMSE",
        "Nature-inspired fusion appendix 进一步补足了“模型过于简单”的潜在审稿风险。": "文献启发的多方法融合被纳入同一 MoleculeNet 候选池，用于检验层级 motif/fingerprint、冻结 embedding、AD gating 和 rank/stacking fusion 是否能在标准任务上带来可验证收益。",
        "TDC external fusion appendix 进一步说明": "外部 TDC official panel 中的融合结果进一步说明",
        "external appendix 的关键结论是：": "外部 ADMET retained-best 分析表明：",
        "为提高实验结论的稳健性，本文从五个方面组织扩展实验：更系统的 external benchmark appendix、性能瓶颈模块 targeted improvement、不平衡分类附加指标、seed-level 统计显著性与 2-3 个 case study。这些分析不进行大规模模型重训练，而是复用已经完成的 TDC full-panel、performance-mode、calibration、roughness、rescue integration 和 high-error interpretability 输出，在此基础上进行 retained-best 整合和统计汇总。": "",
        "为避免无边界扩展候选模型，本文将 selector 改进纳入正式结果部分。": "在完成候选池扩展后，本文进一步用固定 selector policy 审计候选选择过程。",
        "为明确模型比较的边界，本文将模型族扩展与性能优化拆分为两条可审计路线。": "模型族扩展与性能优化在本文中被统一纳入 selector governance，而非作为独立追加实验。",
        "3D-lite descriptors 和 roughness-weighted tree 是针对性能瓶颈回归模块的合理尝试，尤其贴合 FreeSolv、clearance、half-life、PPBR 等 endpoint。然而 validation-only gate 没有在任何测试 endpoint 上正式接入新候选。这个负结果应保留在主文或强附录中，因为它说明性能瓶颈不是单纯缺少 candidate，而是 selector 如何处理验证噪声、roughness 和 endpoint-specific uncertainty。": "3D-lite descriptors 和 roughness-weighted tree 被放入性能瓶颈与 roughness 诊断框架中，用于解释同一类高 roughness 回归瓶颈。它们尤其贴合 FreeSolv、clearance、half-life、PPBR 等 endpoint，但 validation-only gate 没有在任何测试 endpoint 上正式接入新候选。该负结果说明，性能瓶颈不只来自候选模型不足，也来自验证噪声、roughness 与 endpoint-specific uncertainty。",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def build_integrated_markdown() -> str:
    text = read_base()
    prefix = text[: text.index("# 3. 结果")].rstrip()
    discussion = text[text.index("# 4. 讨论") :].strip()
    sections = split_level2_results(text)

    s31 = sections["## 3.1 MoleculeNet 主结果"]
    s32 = sections["## 3.2 TDC ADMET 官方划分与外部附录"]
    s33 = sections["## 3.3 结构分布偏移与划分真实性"]
    s34 = sections["## 3.4 可靠性、适用域与高误差富集"]
    s35 = sections["## 3.5 可解释性：motif attribution 与 fragment enrichment"]
    s36 = sections["## 3.6 救援头集成选择器与性能瓶颈分析"]
    s36b = sections["## 3.6B 文献启发的多方法融合验证"]
    s37 = sections["## 3.7 全面板附录、性能模式与粗糙度解释"]
    s38 = sections["## 3.8 MoleculeNet endpoint 逐项解读"]
    s39 = sections["## 3.9 TDC 与外部 appendix 的结果解释"]
    s310 = sections["## 3.10 性能瓶颈诊断与模型改进路径"]
    s311 = sections["## 3.11 五类实验证据与稳健性分析"]
    s312 = sections["## 3.12 固定选择器策略的正式集成"]
    s313 = sections["## 3.13 3D-lite 与粗糙度加权回归的负结果"]
    s314 = sections["## 3.14 系统实验证据链"]
    s315 = sections["## 3.15 强基线模型与选择器治理"]

    external_blocks = "\n\n".join(
        block
        for block in [
            image_markdown("fig9", "External appendix retained-best selector"),
            extract_block(s311, "图 9.", "表 12."),
            extract_paragraph(s311, "外部 ADMET retained-best"),
        ]
        if block
    )
    imbalance_block = "\n\n".join(
        block
        for block in [
            extract_block(s311, "表 12.", "图 10."),
            extract_paragraph(s311, "不平衡分类结果"),
        ]
        if block
    )
    stability_block = "\n\n".join(
        block
        for block in [
            image_markdown("fig10", "External candidate pool average ranks"),
            extract_block(s311, "图 10.", "表 14."),
            extract_paragraph(s311, "统计显著性层面"),
        ]
        if block
    )
    case_block = extract_block(s311, "表 14.")

    results_parts = [
        "# 3. 结果",
        "",
        "本节不再按照实验加入的时间顺序展开，而是按照论文需要回答的科学问题组织证据。MoleculeNet、TDC ADMET、结构外推、可靠性、selector 审计和可解释性分别构成六条主线；rescue heads、targeted rebuild、Nature-inspired fusion、performance-mode appendix、fixed selector、3D-lite/roughness 和强基线结果均作为相应主线的证据嵌入解释。",
        "",
        "## 3.1 标准任务性能：MoleculeNet 主结果、rescue heads 与多方法融合",
        "",
        image_markdown("fig2", "MoleculeNet model-family rank heatmap"),
        "",
        image_markdown("fig3", "MoleculeNet main performance"),
        "",
        clean_fragment(strip_heading(s31)),
        "",
        clean_fragment(strip_heading(s36)),
        "",
        image_markdown("fig8", "MoleculeNet targeted rebuild decision"),
        "",
        clean_fragment(strip_heading(s36b)),
        "",
        image_markdown("fig8b", "Nature-inspired multimethod fusion decision"),
        "",
        clean_fragment(strip_heading(s38)),
        "",
        "上述结果在同一 MoleculeNet 结果主线下共同说明：FZYC-Mol 的性能提升来自 endpoint-specific 候选治理，而不是后续实验的孤立堆叠。Lipo、FreeSolv、BBBP 和 ClinTox 分别对应 rescue、rebuild、multi-view fusion 和 AD-gated/stacking fusion 的有效场景；未被接受的候选则作为负结果保留，用于界定模型边界。",
        "",
        "## 3.2 外部 ADMET 泛化：TDC official split、full-panel appendix 与 performance-mode",
        "",
        image_markdown("fig4", "TDC official split delta"),
        "",
        clean_fragment(strip_heading(s32)),
        "",
        image_markdown("fig8c", "TDC Nature-inspired fusion decision"),
        "",
        clean_fragment(strip_heading(s37)),
        "",
        clean_fragment(strip_heading(s39)),
        "",
        external_blocks,
        "",
        "因此，TDC 与 external appendix 不应被写成时间顺序上的补充 benchmark，而应作为外部泛化证据：official panel 证明 selector 在常见 ADMET endpoint 上可用，full-panel appendix 证明 retained-best gate 能吸收 performance-mode 的选择性增益，同时在候选较弱时避免负迁移。",
        "",
        "## 3.3 结构分布偏移与划分真实性",
        "",
        image_markdown("fig5", "Split realism and structure shift"),
        "",
        clean_fragment(strip_heading(s33)),
        "",
        "这一小节与 3.2 共同支撑外推结论：性能并不只由模型家族决定，还受到 scaffold 迁移、低相似度样本和 endpoint 标签定义的共同影响。因此，random split、scaffold split、structure-separated split 和 external panel 应被作为同一可靠性检验框架中的不同压力层级。",
        "",
        "## 3.4 可靠性、校准、不平衡分类与 roughness 诊断",
        "",
        image_markdown("fig6", "Reliability and applicability-domain summary"),
        "",
        clean_fragment(strip_heading(s34)),
        "",
        imbalance_block,
        "",
        clean_fragment(strip_heading(s310)),
        "",
        image_markdown("fig20", "3D-lite and roughness gate"),
        "",
        clean_fragment(strip_heading(s313)),
        "",
        "将不确定性、calibration、不平衡分类、roughness 和 3D-lite 负结果放在同一小节中，可以使各类诊断共同回答同一个问题：模型在何种样本、endpoint 和标签结构下不应被过度信任。",
        "",
        "## 3.5 Selector governance、强基线与统计稳定性",
        "",
        image_markdown("fig11", "Formal fixed selector integration"),
        "",
        clean_fragment(strip_heading(s312)),
        "",
        clean_fragment(strip_heading(s315)),
        "",
        stability_block,
        "",
        clean_fragment(strip_heading(s314)),
        "",
        "这一组织方式把 CatBoost/XGBoost/LightGBM/ExtraTrees/RF、Chemprop D-MPNN、冻结 embedding、TabPFNv2 通道、Top-K/stacking、target transform 和 risk-adjusted selector 统一为 selector governance 的候选与审计对象。它们不是彼此割裂的补充实验，而是同一 validation-only 决策框架中的不同候选来源和稳健性检验。",
        "",
        "## 3.6 可解释性与案例分析",
        "",
        image_markdown("fig7", "Motif and fragment interpretation"),
        "",
        clean_fragment(strip_heading(s35)),
        "",
        case_block,
        "",
        "案例分析用于连接性能、可靠性和化学解释：Lipo 展示 endpoint-level rescue 如何被 validation 接受，ClinTox 展示高 ROC-AUC 任务中仍需关注高风险 false negative，high-roughness ADME regression 则展示 roughness、最近邻标签差异和 Top-K/stacking 候选之间的联系。这样写可以让可解释性服务于主结果，并与性能和可靠性证据形成同一论证链条。",
    ]

    new_results = "\n".join(part for part in results_parts if part is not None)
    discussion = discussion.replace("## 4.3 后续实验方向", "## 4.3 局限与进一步验证")
    discussion = discussion.replace(
        "整合后的论文论证建议采用三层结构。第一层是主结果：MoleculeNet、TDC official split、split realism、reliability/AD、motif interpretability，以及 formal fixed-selector integration。第二层是强附录：formal external appendix、performance-mode、rescue-integrated selector、Nature-style fusion 和 selector strategy audit。第三层是诊断附录：3D-lite/roughness negative result、oracle audit、roughness/literature alignment 和 candidate family details。",
        "修订后的结果呈现采用主线融合结构：MoleculeNet 性能、TDC 外部泛化、split realism、reliability/AD、selector governance 和 motif/case interpretability 分别承载相应证据。formal external appendix、performance-mode、rescue-integrated selector、Nature-style fusion、3D-lite/roughness negative result 和 selector strategy audit 均嵌入上述主线中解释，而不再作为按时间追加的独立实验块。",
    )
    integrated = "\n\n".join([prefix, new_results, discussion]).replace("\n\n\n", "\n\n")
    integrated = integrated.replace(
        "ClinTox 的高 ROC-AUC 不能单独解释为临床毒性预测已经可靠",
        "ClinTox 的高 ROC-AUC 不能仅凭该指标解释为临床毒性预测已经可靠",
    )
    integrated = integrated.replace("图 11. 固定选择器策略的正式集成。", "图 11. Risk-adjusted selector 的固定策略审计。")
    integrated = integrated.replace("表 18. 系统实验证据链与结果定位。", "表 18. 主线式证据整合与结果定位。")
    return integrated + "\n"


def normalize_method_figures(markdown: str) -> str:
    method_figures = "\n\n".join(
        [
            image_markdown("fig1", "FZYC-Mol framework overview"),
            "图 1. FZYC-Mol 整体工作流：数据划分、多视图表示、候选专家、验证集选择与可靠性输出。",
            image_markdown("fig23", "FZYC-Mol model structure"),
            "图 2. FZYC-Mol 模型结构：多源表示、专家预测矩阵与 retained-best 决策。",
            "图 1 和图 2 分别强调实验工作流与模型内部结构。二者共同说明，本研究的主体不是单一 backbone，而是由多视图表示、异质专家、验证集治理、适用域判断和可解释性证据组成的预测框架。",
        ]
    )
    pattern = (
        r"图 1A\.[^\n]+\n\n"
        r"图 1B\.[^\n]+\n\n"
        r"图 1A[^\n]+\n"
    )
    return re.sub(pattern, method_figures + "\n", markdown, count=1)


def neutralize_old_crossrefs(markdown: str) -> str:
    replacements = {
        "表 2 和图 2 显示": "MoleculeNet 主结果显示",
        "表 6 表明": "rescue-integrated selector 结果表明",
        "表 8 和表 9 把": "performance-mode 与 roughness 诊断把",
        "图 5 表明": "结构外推结果表明",
        "图 7 将": "motif/fragment 解释性分析将",
        "主文表 15 只显示 promotion": "主文固定 selector 摘要表只显示 promotion",
        "表 15 只显示 promotion": "固定 selector 摘要表只显示 promotion",
    }
    for old, new in replacements.items():
        markdown = markdown.replace(old, new)
    return markdown


def renumber_display_items(markdown: str) -> str:
    fig_counter = 0
    table_counter = 0
    lines = []
    fig_pattern = re.compile(r"^图\s*\d+[A-Z]?\.")
    table_pattern = re.compile(r"^表\s*\d+[A-Z]?\.")
    for line in markdown.splitlines():
        if fig_pattern.match(line):
            fig_counter += 1
            line = fig_pattern.sub(f"图 {fig_counter}.", line, count=1)
        elif table_pattern.match(line):
            table_counter += 1
            line = table_pattern.sub(f"表 {table_counter}.", line, count=1)
        lines.append(line)
    return "\n".join(lines) + "\n"


def add_missing_figure_captions(markdown: str) -> str:
    image = image_markdown("fig20", "3D-lite and roughness gate")
    if image and image in markdown and "3D-lite/roughness-weighted regression 的 validation gate" not in markdown:
        markdown = markdown.replace(
            image + "\n\n3D-lite descriptors",
            image + "\n\n图 20. 3D-lite/roughness-weighted regression 的 validation gate。\n\n3D-lite descriptors",
            1,
        )
    return markdown


def relocate_image_before_caption(markdown: str, filename: str, caption_fragment: str) -> str:
    lines = markdown.splitlines()
    image_idx = next((idx for idx, line in enumerate(lines) if line.startswith("![") and filename in line), None)
    if image_idx is None:
        return markdown
    image_line = lines.pop(image_idx)
    caption_idx = next(
        (idx for idx, line in enumerate(lines) if line.startswith("图 ") and caption_fragment in line),
        None,
    )
    if caption_idx is None:
        lines.insert(image_idx, image_line)
        return "\n".join(lines) + "\n"
    if caption_idx > 0 and lines[caption_idx - 1] == image_line:
        lines.insert(image_idx, image_line)
        return "\n".join(lines) + "\n"
    insert = [image_line, ""]
    lines[caption_idx:caption_idx] = insert
    return "\n".join(lines) + "\n"


def relocate_figures_near_captions(markdown: str) -> str:
    moves = [
        ("fig3_moleculenet_performance_dots.png", "MoleculeNet 主性能比较"),
        ("fig17_moleculenet_targeted_rebuild_decision.png", "MoleculeNet targeted rebuild"),
        ("fig18_nature_multimethod_fusion_decision.png", "文献启发多方法融合"),
        ("fig20_3d_roughness_regression_gate.png", "3D-lite/roughness-weighted regression 的 validation gate"),
    ]
    for filename, fragment in moves:
        markdown = relocate_image_before_caption(markdown, filename, fragment)
    return re.sub(r"\n{3,}", "\n\n", markdown)


def replace_references(markdown: str) -> str:
    match = re.search(r"(?m)^# 6\. 参考文献\s*$", markdown)
    if not match:
        return markdown
    refs = "\n\n".join(f"[{idx}] {ref}" for idx, ref in enumerate(REFERENCE_LIST, start=1))
    return markdown[: match.end()] + "\n\n" + refs + "\n"


def finalize_markdown(markdown: str) -> str:
    markdown = normalize_method_figures(markdown)
    markdown = neutralize_old_crossrefs(markdown)
    markdown = add_missing_figure_captions(markdown)
    markdown = renumber_display_items(markdown)
    markdown = relocate_figures_near_captions(markdown)
    markdown = replace_references(markdown)
    markdown = markdown.replace("![图 12. 强模型对照与 selector 治理。]", "![Strong baseline selector governance]")
    markdown = markdown.replace("reports/manuscript_figures_polished", "reports/manuscript_figures_hires")
    markdown = re.sub(r"\n{3,}", "\n\n", markdown)
    return markdown.strip() + "\n"


def set_font(run, size: float = 10.5, bold: bool = False, color: str = "111827") -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def setup_doc(doc: Document) -> None:
    for style_name, size in [("Normal", 10.5), ("Heading 1", 16), ("Heading 2", 13.5), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
    for section in doc.sections:
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(1.65)
        section.right_margin = Cm(1.65)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, top: bool = False, bottom: bool = False) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        active = (edge == "top" and top) or (edge == "bottom" and bottom)
        tag.set(qn("w:val"), "single" if active else "nil")
        tag.set(qn("w:sz"), "8" if active else "0")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "111827" if active else "FFFFFF")


def set_cell_text(
    cell,
    text: str,
    size: float,
    bold: bool = False,
    color: str = "111827",
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def parse_table_line(line: str, expected_cols: int | None = None) -> list[str]:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    if expected_cols and len(cells) > expected_cols:
        cells = cells[: expected_cols - 1] + [" | ".join(cells[expected_cols - 1 :])]
    if expected_cols and len(cells) < expected_cols:
        cells += [""] * (expected_cols - len(cells))
    return cells


def add_table(doc: Document, lines: list[str]) -> None:
    header = parse_table_line(lines[0])
    rows = [parse_table_line(line, len(header)) for line in lines[2:] if line.strip().startswith("|")]
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    font_size = 5.8 if len(header) >= 8 else 6.8
    for idx, col in enumerate(header):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "FFFFFF")
        set_cell_borders(cell, top=True, bottom=True)
        set_cell_text(cell, col, font_size, bold=True, color="111827")
    for r_idx, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for c_idx, value in enumerate(row):
            cell = cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade_cell(cell, "FFFFFF")
            set_cell_borders(cell, bottom=r_idx == len(rows))
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 22 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cell, value, font_size, bold=False, align=align)
    doc.add_paragraph()


def add_image(doc: Document, line: str) -> None:
    match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line.strip())
    if not match:
        return
    path = Path(match.group(1))
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    width = Inches(6.7)
    if "fig1_" in path.name or "fig23_" in path.name:
        width = Inches(7.0)
    run.add_picture(str(path), width=width)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    size = 10.5
    bold = False
    color = "111827"
    if re.match(r"^图\s+\d+\.", text):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 9.0
        bold = True
        color = "334155"
    elif re.match(r"^表\s+\d+\.", text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 9.2
        bold = True
    elif text.startswith("注："):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 8.2
        color = "64748B"
    elif re.match(r"^\[\d+\]\s+", text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 9.0
    else:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def markdown_to_docx(markdown: str) -> None:
    doc = Document()
    setup_doc(doc)
    lines = markdown.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line:
            idx += 1
            continue
        if line.startswith("!["):
            add_image(doc, line)
            idx += 1
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            p = doc.add_heading("", level=min(level, 3))
            run = p.add_run(text)
            set_font(run, size={1: 16, 2: 13.5, 3: 12}.get(level, 11), bold=True)
            idx += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            if len(table_lines) >= 2:
                add_table(doc, table_lines)
            continue
        add_paragraph(doc, line)
        idx += 1
    doc.save(OUT_DOCX)


def write_report() -> None:
    REPORT.write_text(
        "\n".join(
            [
                "# Manuscript Results Integration Restructure Report",
                "",
                "This revision reorganizes distributed experimental evidence into the main scientific result threads.",
                "",
                "Main changes:",
                "- MoleculeNet section now integrates rescue heads, targeted rebuild, Nature-inspired fusion, and endpoint interpretation.",
                "- TDC section now integrates official split, external fusion, full-panel appendix, performance-mode retained-best, and external candidate coverage.",
                "- Reliability section now integrates uncertainty/AD, imbalanced classification metrics, roughness diagnostics, and 3D-lite/roughness negative results.",
                "- Selector governance section now integrates fixed selector policy, strong baselines, seed-level statistics, and win/tie/loss analysis.",
                "- Interpretability section now integrates motif/fragment analysis with case studies.",
                "- Figure and table captions are renumbered by order of appearance.",
                "- DOCX tables use a three-line academic table style.",
                "- References are rewritten as a numbered, cleaned list with verified DOI/arXiv information where available.",
                "",
                f"- Output DOCX: `{OUT_DOCX.relative_to(ROOT)}`",
                f"- Output Markdown: `{OUT_MD.relative_to(ROOT)}`",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def main() -> None:
    markdown = finalize_markdown(build_integrated_markdown())
    OUT_MD.write_text(markdown, encoding="utf-8")
    markdown_to_docx(markdown)
    write_report()
    print(f"wrote {OUT_MD.relative_to(ROOT)}")
    print(f"wrote {OUT_DOCX.relative_to(ROOT)}")
    print(f"wrote {REPORT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
