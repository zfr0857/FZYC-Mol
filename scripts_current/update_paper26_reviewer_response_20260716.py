from pathlib import Path

from docx import Document
from docx.shared import Pt


PATH = Path(r"D:\fzyc\output\paper26_split_regime_transfer_revision_20260716\Reviewer_concern_Response_Location.docx")


def main() -> None:
    doc = Document(PATH)
    table = doc.tables[0]
    for row in table.rows:
        if row.cells[0].text.startswith("Realign Discussion sections"):
            row.cells[0].text = "Realign Discussion sections 4.1-4.8"
            row.cells[1].text = (
                "Reorganized the Discussion so that each heading addresses nominal versus effective diversity, matrix construction, "
                "chance-adjusted ranking, cross-fitting, matched-size multiview evidence, chemical-support reliability, split-regime "
                "transfer and limitations."
            )
            row.cells[2].text = "Discussion 4.1-4.8"
    additions = [
        (
            "Test whether the main audit is an artefact of scaffold splitting",
            "Added a fully nested transfer audit on ClinTox, BACE and ESOL. The same 32-candidate registry, K prefixes, five seeds and 3 x 3 "
            "design were retrained with intact Morgan-512 Tanimoto connected components at threshold 0.70. All splits were group disjoint, "
            "maximum cross-fold similarity was 0.699 and no random fallback occurred. CAHit@3 K32-K4 changes and cross-fitted gap directions "
            "agreed across regimes in all three endpoints, while magnitudes and interval exclusion remained endpoint dependent.",
            "Methods 2.13; Results 3.9; Discussion 4.7-4.8; Tables S23-S26; Figure S18",
        ),
        (
            "Avoid overgeneralizing from a homogeneous Morgan registry",
            "Retained the controlled near-duplicate registry as the primary stress test rather than presenting it as representative of modern "
            "architectures. The matched-size multiview and four-model panels remain boundary analyses, and the revised Abstract, Scientific "
            "Contribution, Discussion and Limitations explicitly restrict the claim to direction transport in three endpoints rather than a "
            "universal selector or model-family law.",
            "Abstract; Scientific Contribution; Discussion 4.5, 4.7-4.8; Conclusions",
        ),
    ]
    existing = {row.cells[0].text for row in table.rows}
    for values in additions:
        if values[0] in existing:
            continue
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8.5)
                    run.bold = row_index == 0
    doc.save(PATH)


if __name__ == "__main__":
    main()
