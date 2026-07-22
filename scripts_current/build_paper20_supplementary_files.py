from __future__ import annotations

import hashlib
import json
import re
import sys
from pathlib import Path

import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from matplotlib.backends.backend_pdf import PdfPages


ROOT = Path("D:/fzyc")
sys.path.insert(0, str(ROOT / "scripts"))
from build_paper19_manuscript import add_body, add_heading, configure_document  # noqa: E402

OUT = ROOT / "output"
CORE = OUT / "paper20_candidate_pool_audit_20260712"
MAJOR = OUT / "paper19_major_revision_20260712"
PREV = OUT / "paper19_rejection_driven_experiments_20260712"
REV = OUT / "paper19_jcheminformatics_revision_20260712"
HARD = OUT / "sci1_hardening_20260707"
UQ = OUT / "sci1_mechanism_uq_decision_20260707"
BASE = ROOT / "results" / "nested_selection" / "repeated_nested"
SUPP = CORE / "supplementary"
FILE1 = SUPP / "Additional_file_1_Supplementary_Methods_and_Results.docx"
FILE2 = SUPP / "Additional_file_2_Supplementary_Tables_S1-S17.xlsx"
FILE3 = SUPP / "Additional_file_3_Supplementary_Figures_S1-S14.pdf"
DISPLAY = {
    "bace": "BACE", "bbbp": "BBBP", "clintox": "ClinTox", "esol": "ESOL", "freesolv": "FreeSolv",
    "lipo": "Lipophilicity", "tdc_caco2_wang": "Caco2", "tdc_hia_hou": "HIA", "tdc_pgp_broccatelli": "P-gp",
}


def latest(name: str) -> Path | None:
    paths = [p for p in OUT.rglob(name) if "代码包" not in str(p) and "artifacts" not in str(p)]
    return max(paths, key=lambda p: p.stat().st_mtime) if paths else None


def read_latest(name: str) -> pd.DataFrame:
    path = latest(name)
    if path is None:
        return pd.DataFrame({"status": ["requires source data"], "requested_file": [name]})
    frame = pd.read_csv(path)
    frame.insert(0, "source_file", str(path))
    return frame


def load_base(name: str) -> pd.DataFrame:
    frames = []
    for seed in [11, 23, 37, 53, 71]:
        frame = pd.read_csv(BASE / f"seed_{seed}" / name).rename(columns={"dataset": "task"})
        frame.insert(0, "seed", seed)
        frames.append(frame)
    return pd.concat(frames, ignore_index=True)


def sha256(path: Path) -> str:
    digest = hashlib.sha256(path.read_bytes())
    return digest.hexdigest()


def build_tables() -> dict[str, pd.DataFrame]:
    dataset = pd.read_csv(REV / "dataset_characteristics.csv")
    registry = pd.read_csv(PREV / "paper19_candidate_registry.csv")
    compute = pd.read_csv(PREV / "paper19_compute_budget.csv")
    outer = load_base("outer_candidate_scores.csv")
    split_manifest = outer.groupby(["task", "task_type", "seed", "outer_fold", "outer_split_type"], as_index=False).agg(
        n_candidates=("candidate_order", "nunique"), n_test=("n_test", "max")
    )
    diversity = pd.read_csv(CORE / "utility_pattern_diversity_endpoint.csv")
    composition = pd.read_csv(REV / "candidate_composition_controls.csv")
    cross = pd.read_csv(CORE / "cross_fitted_reference_units.csv")
    multiview = pd.read_csv(MAJOR / "multiview_absolute_paired_units.csv")
    matched = pd.read_csv(CORE / "matched_k_multiview_summary.csv")
    representation = pd.read_csv(HARD / "six_task_strong_endpoint_table.csv")
    overlap = pd.read_csv(HARD / "six_task_error_overlap_pairwise_detail.csv")
    rep_combined = pd.concat([
        representation.assign(record_type="endpoint_performance"),
        overlap.assign(record_type="prediction_error_overlap"),
    ], ignore_index=True, sort=False)
    dedup = pd.read_csv(HARD / "six_task_duplicate_sensitivity_summary.csv")
    conformal = pd.read_csv(UQ / "conformal_crossfold_summary.csv").assign(record_type="classification_conformal")
    cqr = pd.read_csv(UQ / "cqr_regression_summary.csv").assign(record_type="regression_cqr")
    conformal_combined = pd.concat([conformal, cqr], ignore_index=True, sort=False)
    tdc = read_latest("fig06_tdc_22_endpoint_raw_table.csv")
    moleculeace = read_latest("fig09_moleculeace_task_summary.csv")
    bro5 = read_latest("fig09_bro5_cycpept_pampa_compact_summary.csv")
    runtime = pd.read_csv(HARD / "six_task_strong_baseline_runtime_status.csv")
    unavailable = pd.DataFrame([
        {"record_type": "missing_source", "candidate_or_item": "multiview candidate-level predictions", "status": "requires source data", "reason": "prediction exports absent"},
        {"record_type": "missing_source", "candidate_or_item": "multiview split-assignment hash", "status": "requires source data", "reason": "manifest records split generator but not assignment hash"},
        {"record_type": "historical_document", "candidate_or_item": "小论文-18(4).docx", "status": "requires source data", "reason": "exact named file not found"},
    ])
    failed = pd.concat([runtime.assign(record_type="representation_runtime"), unavailable], ignore_index=True, sort=False)
    tables = {
        "S1 Dataset provenance": dataset,
        "S2 Candidate registry": registry,
        "S3 Config compute": compute,
        "S4 Split manifest": split_manifest,
        "S5 Fold utilities": outer,
        "S6 Effective diversity": diversity,
        "S7 Composition controls": composition,
        "S8 Audit references": cross,
        "S9 Multiview effects": multiview,
        "S10 Matched-K": matched,
        "S11 Rep predictions": rep_combined,
        "S12 Dedup sensitivity": dedup,
        "S13 Conformal": conformal_combined,
        "S14 TDC endpoints": tdc,
        "S15 MoleculeACE": moleculeace,
        "S16 bRo5": bro5,
        "S17 Failed unavailable": failed,
    }
    return tables


def write_tables(tables: dict[str, pd.DataFrame]) -> None:
    SUPP.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(FILE2, engine="openpyxl") as writer:
        for sheet, frame in tables.items():
            frame.to_excel(writer, sheet_name=sheet[:31], index=False)
            ws = writer.book[sheet[:31]]
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions
            for cell in ws[1]:
                cell.font = cell.font.copy(bold=True)
            for col in ws.columns:
                sample = [str(c.value) if c.value is not None else "" for c in list(col)[:200]]
                ws.column_dimensions[col[0].column_letter].width = min(max(max(map(len, sample), default=8) + 2, 10), 40)


def build_file1(tables: dict[str, pd.DataFrame]) -> None:
    doc = Document(); configure_document(doc)
    add_heading(doc, "Supplementary Methods and Results", 1)
    add_body(doc, "This file preserves secondary analyses and extended implementation detail while the main article remains focused on candidate-pool expansion, validation-ranking distortion, cross-fitted audit gaps and matched-size multiview effects. Exact numerical records are provided in Additional file 2; regenerated diagnostic figures are provided in Additional file 3.")
    sections = [
        ("S1 Extended nested-audit methods", "Candidate eligibility, split manifests, fold-level utilities, computational exposure and failed-candidate status were retained at the finest available resolution. The primary outer matrix contained 15 seed-by-fold rows per endpoint. Common-difficulty adjustment, cross-fitted reference construction and matched-K pool definitions are described in the main Methods and reproduced in Tables S4-S10."),
        ("S2 TDC endpoint analysis", f"The TDC source table contains {len(tables['S14 TDC endpoints'])} endpoint rows. These results retain official endpoint direction, prior comparator, performance-run source and retained model. They are contextual reliability evidence because the endpoint panel and evaluation budgets differ from the nine-endpoint expansion audit."),
        ("S3 Automated machine-learning budgets", "Historical 30-, 300- and 1,800-second automated machine-learning outputs were retained only where machine-readable files were available. They are not used to claim superiority over automated search because candidate budgets, feature preparation and hardware exposure are not equivalent to the controlled prefix registry."),
        ("S4 Leave-one-endpoint-out meta-risk", "The leave-one-endpoint-out selector is treated as an exploratory meta-risk analysis. It does not define a generally applicable molecular selector because nine endpoints are insufficient to identify stable cross-task decision rules and because endpoint meta-features were derived retrospectively."),
        ("S5 Conformal prediction and CQR", f"Table S13 combines {len(tables['S13 Conformal'])} classification-conformal and regression-CQR summary rows. Label-conditional and similarity-conditioned methods, ensemble uncertainty and CQR were evaluated under existing folds. Coverage is descriptive for the studied endpoints and does not guarantee screening utility under new chemical shifts."),
        ("S6 MoleculeACE activity cliffs", f"The MoleculeACE table contains {len(tables['S15 MoleculeACE'])} task summaries. Gap correlation, direction accuracy and pair counts were retained. Activity-cliff pairs expose local property discontinuity and therefore complement, rather than validate, aggregate scaffold performance."),
        ("S7 Beyond-rule-of-five boundaries", f"The bRo5 table contains {len(tables['S16 bRo5'])} split summaries. Random, scaffold and perimeter splits are reported separately because perimeter migration changes both chemical support and task difficulty."),
        ("S8 Deduplication sensitivity", f"Table S12 contains {len(tables['S12 Dedup sensitivity'])} duplicate-policy summaries. Global deduplication, training-fold aggregation and retained-duplicate strategies were compared where source predictions existed. These results assess data-processing sensitivity and are not added to the primary candidate-expansion effect."),
        ("S9 Representation errors and failure cases", "Prediction-level error overlap was available for RDKit-RF, a GCN and frozen ChemBERTa and MoLFormer probes. Jaccard overlap was incomplete, indicating complementary errors, but unequal optimization budgets preclude a modern-model leaderboard. Low-similarity, extreme-label, activity-cliff, bRo5 and ClinTox minority-class failures are retained as distinct categories."),
        ("S10 Source limitations", "The multiview directory contains candidate registries, inner utilities, outer candidate scores and selected-policy details. It does not contain molecule-level prediction exports or a split-assignment hash. Consequently, score-level and key-level hashes were recomputed, whereas prediction-hash verification remains unavailable. The exact historical file named 小论文-18(4).docx was also unavailable; historical CSVs were used only as supplementary evidence and did not overwrite the current main manuscript."),
        ("S11 Extended interpretation", "The supplementary analyses support a restrained interpretation. Model-selection loss depends on candidate correlation, validation information and endpoint chemistry. Reliability depends additionally on class imbalance, scaffold novelty, similarity and label discontinuity. These axes should be reported jointly but not collapsed into one score."),
    ]
    for heading, text in sections:
        add_heading(doc, heading, 2); add_body(doc, text)
    add_heading(doc, "Supplementary table directory", 1)
    for i, (sheet, frame) in enumerate(tables.items(), start=1):
        add_body(doc, f"Table S{i}. {sheet.split(' ', 1)[1]}. {len(frame)} rows in Additional file 2.")
    add_heading(doc, "Supplementary figure directory", 1)
    names = ["Dataset cleaning flow", "Candidate correlation matrices", "Eigenvalue spectra", "Raw Top-1 and raw MRR", "Permutation calibration", "Signal-recovery control", "Full winner-optimism grid", "Fold-level audit gaps", "Multiview seed audit", "TDC endpoint panel", "Conformal and risk-coverage curves", "MoleculeACE results", "bRo5 migration results", "Low-similarity and ClinTox failure cases"]
    for i, name in enumerate(names, start=1): add_body(doc, f"Figure S{i}. {name}.")
    doc.save(FILE1)


def page_title(fig: plt.Figure, text: str) -> None:
    fig.suptitle(text, x=0.06, ha="left", fontsize=12, fontweight="bold")


def build_figures(tables: dict[str, pd.DataFrame]) -> None:
    mpl.rcParams.update({"font.family": "Arial", "font.size": 8.5, "axes.spines.top": False, "axes.spines.right": False, "pdf.fonttype": 42})
    with PdfPages(FILE3) as pdf:
        d = tables["S1 Dataset provenance"]
        fig, ax = plt.subplots(figsize=(7.2, 5.2)); page_title(fig, "Figure S1. Dataset cleaning flow")
        y=np.arange(len(d)); ax.barh(y,d["raw_n"],color="#AEB5BD",label="Raw"); ax.barh(y,d["analysis_n"],color="#376FA8",label="Analysis"); ax.set(yticks=y,yticklabels=d["display_name"],xlabel="Molecules"); ax.legend(frameon=False); fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        outer = tables["S5 Fold utilities"]; mats=[]
        for task in outer.task.unique():
            m=outer.loc[outer.task.eq(task)].pivot_table(index=["seed","outer_fold"],columns="candidate_order",values="outer_utility")
            mats.append(np.nan_to_num(np.corrcoef(m.to_numpy(float),rowvar=False)))
        corr=np.mean(mats,axis=0)
        fig, ax=plt.subplots(figsize=(7.2,5.2)); page_title(fig,"Figure S2. Candidate correlation matrices"); im=ax.imshow(corr,vmin=-1,vmax=1,cmap="RdBu_r"); ax.set(xlabel="Candidate order",ylabel="Candidate order"); fig.colorbar(im,ax=ax,label="Mean outer-utility correlation"); fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        eig=np.array([np.sort(np.linalg.eigvalsh(m))[::-1] for m in mats]); mean=eig.mean(axis=0)
        fig,ax=plt.subplots(figsize=(7.2,5.2)); page_title(fig,"Figure S3. Eigenvalue spectra"); ax.plot(range(1,len(mean)+1),mean,marker="o",ms=3); ax.set(xlabel="Ordered component",ylabel="Mean correlation eigenvalue",yscale="log"); fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        rank=pd.read_csv(MAJOR/"chance_adjusted_ranking_summary.csv").sort_values("candidate_count")
        fig,ax=plt.subplots(figsize=(7.2,5.2)); page_title(fig,"Figure S4. Raw Top-1 and raw MRR"); ax.plot(rank.candidate_count,rank.mean_top1_hit,marker="o",label="Top-1"); ax.plot(rank.candidate_count,rank.mean_mrr,marker="s",label="MRR"); ax.set(xlabel="K",ylabel="Raw metric",xticks=[4,8,16,32],ylim=(0,1)); ax.legend(frameon=False); fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        signal=read_latest("fig11_signal_recovery_summary.csv")
        null=signal.loc[signal.get("signal_correlation",pd.Series(dtype=float)).eq(0)] if "signal_correlation" in signal else pd.DataFrame()
        fig,ax=plt.subplots(figsize=(7.2,5.2)); page_title(fig,"Figure S5. Permutation calibration");
        if len(null): ax.errorbar(null.pool_size,null.chance_adjusted_hit_mean,yerr=[null.chance_adjusted_hit_mean-null.chance_adjusted_hit_ci95_low,null.chance_adjusted_hit_ci95_high-null.chance_adjusted_hit_mean],marker="o",capsize=2)
        else: ax.text(0.5,0.5,"requires source data",ha="center",va="center")
        ax.axhline(0,color="black",lw=.8); ax.set(xlabel="K",ylabel="Chance-adjusted Hit under null"); fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        fig,ax=plt.subplots(figsize=(7.2,5.2)); page_title(fig,"Figure S6. Signal-recovery control")
        if "signal_correlation" in signal:
            for k,g in signal.groupby("pool_size"): ax.plot(g.signal_correlation,g.chance_adjusted_hit_mean,marker="o",label=f"K={k}")
            ax.legend(frameon=False,ncol=2)
        else: ax.text(.5,.5,"requires source data",ha="center")
        ax.set(xlabel="Injected validation-audit correlation",ylabel="Chance-adjusted Hit"); fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        sim=pd.read_csv(PREV/"paper19_oracle_extreme_value_simulation.csv"); s=sim.loc[sim.truth_scenario.eq("equal_truth") & sim.pairwise_candidate_correlation.eq(0.9)]; grid=s.pivot_table(index="effective_audit_sample_size",columns="candidate_count",values="mean_observed_oracle_optimism").sort_index(ascending=False)
        fig,ax=plt.subplots(figsize=(7.2,5.2)); page_title(fig,"Figure S7. Full winner-optimism grid"); im=ax.imshow(grid,aspect="auto",cmap="YlOrRd"); ax.set(xticks=range(len(grid.columns)),xticklabels=grid.columns,yticks=range(len(grid.index)),yticklabels=grid.index,xlabel="K",ylabel="Effective audit n"); fig.colorbar(im,ax=ax,label="Mean optimism"); fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        gaps=pd.read_csv(CORE/"audit_gap_decomposition_units.csv"); g=gaps.loc[gaps.candidate_count.eq(32)]
        fig,ax=plt.subplots(figsize=(7.2,5.2)); page_title(fig,"Figure S8. Fold-level audit gaps"); data=[g.loc[g.task.eq(t),"incremental_observed_audit_gap"] for t in sorted(g.task.unique())]; ax.boxplot(data,labels=[DISPLAY.get(t,t) for t in sorted(g.task.unique())],showfliers=False); ax.set(ylabel="Incremental observed audit gap"); plt.setp(ax.get_xticklabels(),rotation=45,ha="right"); fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        mv=pd.read_csv(CORE/"matched_k_multiview_units.csv"); q=mv.loc[mv.pool_name.eq("ladder_K12")]
        fig,ax=plt.subplots(figsize=(7.2,5.2)); page_title(fig,"Figure S9. Multiview seed audit")
        for task,grp in q.groupby("task"): ax.plot(sorted(grp.seed.unique()),grp.groupby("seed").gain_vs_morgan_k3.mean(),marker="o",label=DISPLAY.get(task,task))
        ax.set(xlabel="Seed",ylabel="Mean paired gain"); ax.legend(frameon=False,ncol=3,fontsize=6.5); fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        tdc=tables["S14 TDC endpoints"]
        fig,ax=plt.subplots(figsize=(7.2,5.2)); page_title(fig,"Figure S10. TDC endpoint panel")
        if "performance_delta_vs_previous" in tdc:
            q=tdc.sort_values("performance_delta_vs_previous"); ax.barh(range(len(q)),q.performance_delta_vs_previous,color=np.where(q.performance_delta_vs_previous>=0,"#376FA8","#B54D4D")); ax.set(yticks=range(len(q)),yticklabels=q.dataset, xlabel="Direction-normalized delta versus previous")
        else: ax.text(.5,.5,"requires source data",ha="center")
        fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        risk=read_latest("fig07_risk_coverage_curves.csv")
        fig,axes=plt.subplots(1,2,figsize=(7.2,5.2)); page_title(fig,"Figure S11. Conformal and risk-coverage curves")
        conf=pd.read_csv(UQ/"conformal_crossfold_summary.csv"); q=conf.groupby("method",as_index=False).agg(coverage=("mean_class_1_coverage","mean")) if "mean_class_1_coverage" in conf else pd.DataFrame()
        if len(q): axes[0].bar(range(len(q)),q.coverage,color="#2D8A7E"); axes[0].set(xticks=range(len(q)),xticklabels=q.method,ylabel="Mean minority coverage"); plt.setp(axes[0].get_xticklabels(),rotation=45,ha="right")
        if "coverage" in risk:
            for dataset,g in risk.groupby("dataset"): axes[1].plot(g.coverage,g.observed_risk,alpha=.4)
        axes[1].set(xlabel="Coverage",ylabel="Observed risk"); fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        ma=tables["S15 MoleculeACE"]
        fig,ax=plt.subplots(figsize=(7.2,5.2)); page_title(fig,"Figure S12. MoleculeACE results")
        if "direction_accuracy" in ma:
            q=ma.sort_values("direction_accuracy"); ax.barh(range(len(q)),q.direction_accuracy,color="#D07A34"); ax.set(yticks=range(len(q)),yticklabels=q.task,xlabel="Cliff-pair direction accuracy",xlim=(0,1))
        else: ax.text(.5,.5,"requires source data",ha="center")
        fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        br=tables["S16 bRo5"]
        fig,ax=plt.subplots(figsize=(7.2,5.2)); page_title(fig,"Figure S13. bRo5 migration results")
        if "test_RMSE" in br:
            values=br.test_RMSE.astype(str).str.extract(r"([0-9.]+)")[0].astype(float); ax.bar(br.split,values,color=["#376FA8","#D07A34","#B54D4D"][:len(br)]); ax.set(ylabel="Test RMSE")
        else: ax.text(.5,.5,"requires source data",ha="center")
        fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)

        simi=pd.read_csv(UQ/"calibration_ood_scaffold_summary.csv"); cli=pd.read_csv(UQ/"clintox_minority_negative_result.csv").drop_duplicates("candidate")
        fig,axes=plt.subplots(1,2,figsize=(7.2,5.2)); page_title(fig,"Figure S14. Low-similarity and ClinTox failure cases")
        q=simi.groupby("tanimoto_bin",as_index=False).mean(numeric_only=True); order=[x for x in ["<0.5","0.5-0.7",">0.7"] if x in q.tanimoto_bin.values]; q=q.set_index("tanimoto_bin").loc[order]; axes[0].plot(order,q.mean_roc_auc,marker="o"); axes[0].set(xlabel="Maximum train-set Tanimoto",ylabel="Mean ROC-AUC",ylim=(0,1))
        axes[1].barh(range(len(cli)),cli.minority_recall,color="#376FA8"); axes[1].set(yticks=range(len(cli)),yticklabels=cli.candidate,xlabel="ClinTox minority recall",xlim=(0,1)); fig.tight_layout(rect=[0,0,1,0.95]); pdf.savefig(fig); plt.close(fig)


def main() -> None:
    SUPP.mkdir(parents=True, exist_ok=True)
    tables = build_tables()
    write_tables(tables)
    build_file1(tables)
    build_figures(tables)
    audit = {
        "additional_file_1": {"path": str(FILE1), "sha256": sha256(FILE1)},
        "additional_file_2": {"path": str(FILE2), "sha256": sha256(FILE2), "sheets": list(tables)},
        "additional_file_3": {"path": str(FILE3), "sha256": sha256(FILE3), "figures": 14},
    }
    (SUPP/"supplementary_audit.json").write_text(json.dumps(audit,ensure_ascii=False,indent=2),encoding="utf-8")
    print(json.dumps(audit,ensure_ascii=False,indent=2))


if __name__ == "__main__":
    main()
