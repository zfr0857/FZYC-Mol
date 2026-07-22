from __future__ import annotations

import json
import re
import os
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from paper25_manuscript_content import (
    DISCUSSION_SECTIONS,
    INTRODUCTION_PARAGRAPHS,
    render_abstract,
)
from paper25_docx_svg import embed_svg_figures


ROOT=Path("D:/fzyc")
SRC=ROOT/"output"/"小论文-20_Journal_of_Cheminformatics_主文.docx"
NEW=Path(os.environ.get("FZYC_ANALYSIS_OUT", ROOT/"output"/"paper21_final_reanalysis_20260713"))
MINOR=Path(os.environ.get("FZYC_MINOR_OUT", ROOT/"output"/"paper23_minor_revision_20260713"))
MASTER=MINOR/"Minor_revision_master_results_and_verification.xlsx"
OUT=Path(os.environ.get("FZYC_MANUSCRIPT_OUT", ROOT/"output"/"Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript.docx"))
FIG=Path(os.environ.get("FZYC_FIG_OUT", NEW/"main_figures"))
DISPLAY={"bace":"BACE","bbbp":"BBBP","clintox":"ClinTox","esol":"ESOL","freesolv":"FreeSolv","lipo":"Lipophilicity","tdc_caco2_wang":"Caco2","tdc_hia_hou":"HIA","tdc_pgp_broccatelli":"P-gp"}
TRANSFORM_LABELS={"raw":"raw","row_centred":"row-centred","fixed_reference_relative":"fixed-reference-relative","within_unit_rank":"within-unit-rank"}


def find(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix): return p
    raise KeyError(prefix)


def set_text(p, text):
    p.clear(); r=p.add_run(text); r.font.name="Times New Roman"; r.font.size=Pt(11)


def delete_paragraph(p):
    p._element.getparent().remove(p._element)


def set_after_heading(doc, heading, text):
    paragraphs = doc.paragraphs
    for i, paragraph in enumerate(paragraphs[:-1]):
        if paragraph.text.strip() == heading:
            set_text(paragraphs[i + 1], text)
            return
    raise KeyError(heading)


def replace_between(doc, start_prefix, end_prefix, items):
    start=find(doc,start_prefix); end=find(doc,end_prefix)
    node=start._element.getnext()
    while node is not end._element:
        following=node.getnext(); node.getparent().remove(node); node=following
    for text,style in items:
        paragraph=end.insert_paragraph_before(style=style)
        set_text(paragraph,text)


def page_field(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=paragraph.add_run(); fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); run._r.append(fld)


def set_cell_border(cell, **edges):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in("w:tcBorders")
    if borders is None: borders=OxmlElement("w:tcBorders"); tcPr.append(borders)
    for edge,attrs in edges.items():
        tag="w:"+edge; el=borders.find(qn(tag))
        if el is None: el=OxmlElement(tag); borders.append(el)
        for k,v in attrs.items(): el.set(qn("w:"+k),str(v))


def three_line(table):
    none={"val":"nil"}; line={"val":"single","sz":"8","color":"000000"}
    for i,row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell,top=none,bottom=none,left=none,right=none,insideH=none,insideV=none)
            if i==0: set_cell_border(cell,top=line,bottom=line)
            if i==len(table.rows)-1: set_cell_border(cell,bottom=line)
            for p in cell.paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
                for r in p.runs: r.font.name="Times New Roman"; r.font.size=Pt(8.5); r.bold=(i==0)


def insert_figure(caption_p, path, width):
    p=caption_p.insert_paragraph_before(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path),width=Inches(width))


def effect_summary(frame, task_type):
    q=frame[frame.task_type.eq(task_type)].copy()
    positive=[DISPLAY[t] for t in q.loc[q.cross_fitted_effect.gt(0),"task"]]
    negative=[DISPLAY[t] for t in q.loc[q.cross_fitted_effect.lt(0),"task"]]
    excludes=[DISPLAY[r.task] for r in q.itertuples() if r.split_seed_bootstrap95_low_cross_fitted>0 or r.split_seed_bootstrap95_high_cross_fitted<0]
    return positive,negative,excludes


def effective_rank_text(frame):
    parts=[]
    for mode in ["raw","row_centred","fixed_reference_relative","within_unit_rank"]:
        row=frame.loc[frame.transformation.eq(mode)].iloc[0]
        parts.append(
            f"{TRANSFORM_LABELS[mode]} {row.entropy_rank_median:.2f} "
            f"(IQR {row.entropy_rank_q25:.2f}–{row.entropy_rank_q75:.2f}; "
            f"range {row.entropy_rank_min:.2f}–{row.entropy_rank_max:.2f})"
        )
    return "; ".join(parts)


def main():
    doc=Document(SRC)
    master=pd.read_excel(MASTER,sheet_name="Master result table")
    ranking=pd.read_excel(MASTER,sheet_name="Ranking main")
    effective=pd.read_excel(MASTER,sheet_name="Effective rank")
    cross=pd.read_excel(MASTER,sheet_name="Cross-fitted")
    matched_endpoint=pd.read_excel(MASTER,sheet_name="Matched K3")
    support=pd.read_excel(MASTER,sheet_name="Chemical support")
    null_summary=pd.read_csv(MINOR/"mechanism_permutation_null_summary.csv")
    signal_summary=pd.read_csv(MINOR/"mechanism_signal_recovery_summary.csv")
    mechanism=json.loads((MINOR/"mechanism_calibration_audit.json").read_text(encoding="utf-8"))
    rank4=ranking.loc[ranking.candidate_count.eq(4)].iloc[0]
    rank8=ranking.loc[ranking.candidate_count.eq(8)].iloc[0]
    rank16=ranking.loc[ranking.candidate_count.eq(16)].iloc[0]
    rank32=ranking.loc[ranking.candidate_count.eq(32)].iloc[0]
    class_pos,class_neg,class_excludes=effect_summary(cross,"classification")
    reg_pos,reg_neg,reg_excludes=effect_summary(cross,"regression")
    same_positive=[DISPLAY[t] for t in cross.loc[cross.same_unit_effect.gt(0),"task"]]
    same_negative=[DISPLAY[t] for t in cross.loc[cross.same_unit_effect.lt(0),"task"]]
    endpoint_medians=matched_endpoint.rename(columns={"endpoint_median_gain":"selected_model_gain_median"})
    matched_positive=[DISPLAY[t] for t in endpoint_medians.loc[endpoint_medians.selected_model_gain_median.gt(0),"task"]]
    matched_negative=[DISPLAY[t] for t in endpoint_medians.loc[endpoint_medians.selected_model_gain_median.lt(0),"task"]]
    subset_positive=matched_endpoint.set_index("task").positive_subset_proportion
    support_medians=support.set_index(["task_type","tanimoto_bin"]).selected_performance_median
    diversity_text=effective_rank_text(effective)
    positive_count=int(master.loc[master.metric.eq("positive cross-fitted direction count"),"estimate"].iloc[0])
    negative_count=int(master.loc[master.metric.eq("negative cross-fitted direction count"),"estimate"].iloc[0])
    abstract=render_abstract(endpoint_count=cross.task.nunique(),positive_count=positive_count,negative_count=negative_count)
    replacements={
      "Background:": abstract["Background:"],
      "Methods: We performed": abstract["Methods:"],
      "Results: At K = 32": abstract["Results:"],
      "Conclusions: Within the evaluated design": abstract["Conclusions:"],
      "Scientific Contribution:": abstract["Scientific Contribution:"],
      "We ask a limited question": "We ask a limited question: under a retrospectively locked repeated nested scaffold evaluation, how did registered candidate-pool expansion relate to matrix-dependent utility-pattern diversity, chance-adjusted validation-ranking fidelity and endpoint-specific selection loss? The study contributes three elements: separation of nominal K from utility-pattern diversity; calibration of adjusted ranking with permutation and graded signal-recovery controls, coupled to cross-fitted selection gaps; and matched-size representation analyses bounded by chemical support. This hierarchy does not turn the audit into a new predictor, universal selector, prospective preregistration or independent external validation study.",
      "The study was a retrospective frozen audit": "The study was a retrospective audit of completed molecular-property experiments. The nine primary endpoints, candidate order, K values, model seeds, split logic, selection rules, outcomes and exclusions were held fixed for the analyses reported here. Because the audit specification was reconstructed after the original outer outcomes existed, the design was not prospectively preregistered. Frozen denotes unchanged eligibility and analysis rules during this audit; it does not imply an untouched prospective cohort.",
      "The outer folds served": "Outer folds served as audit sets for frozen inner-selection decisions. Outer labels did not influence candidate eligibility, hyperparameters, fold-specific preprocessing or tie breaking within a completed unit. The same outer results were subsequently used to quantify the largest observed candidate utility, ranking agreement and selection loss; the resulting maximum is therefore an observed finite-audit best rather than a population bound. The multiview and representation panels reused public endpoint definitions and are not independent external validations.",
      "The present revision treated": "The candidate-pool audit was the primary inferential line. TDC, automated machine learning, conformal prediction, activity-cliff, beyond-rule-of-five and duplicate-sensitivity analyses were retained as supplementary reliability or boundary evidence and were not presented as independent confirmation.",
      "For every endpoint, five seeds": "Five seeded scaffold partitions were generated independently for both classification and regression endpoints using seeds 11, 23, 37, 53 and 71. Each partition contained three outer scaffold folds and three inner scaffold folds within every outer-training partition. Bemis–Murcko scaffold groups were kept intact; group allocation used seed-dependent random tie breaking while balancing sample counts, and no model performance informed allocation. The split manifest records sample and scaffold counts, target mean, standard deviation and range, split hashes and the absence of cross-fold scaffold overlap.",
      "Outer folds from the same seed": "Outer folds share training observations and are paired audit units rather than independent biological replications. For every endpoint, fold effects were first averaged within each of five distinct split seeds before split-seed resampling. The five seeds define repeated scaffold partitions; the three folds within a seed do not constitute three independent replications.",
      "Uncertainty was estimated with a hierarchical bootstrap": "Uncertainty in endpoint-specific effective diversity was estimated with 5,000 hierarchical bootstrap replicates using fixed seed 20260713. Split-seed blocks were sampled with replacement and, within each sampled block, three outer-fold rows were sampled with replacement; the shrinkage correlation and effective ranks were recomputed. Across-endpoint summaries are medians, IQRs and ranges, without a population-level confidence interval because the nine endpoints were not treated as a random sample from a molecular-task population. Sensitivity analyses omitted each seed and fold and repeated fixed-reference-relative estimates with predefined references; no reference was chosen from outer performance.",
      "The multiview registry crossed": "The multiview registry crossed four representations (Morgan-512, MACCS, RDKit2D and their concatenation) with three learners. The matched-size analysis exhaustively enumerated all C(12,3) = 220 three-candidate subsets without retraining, using stored candidate-level inner and outer utilities. These overlapping subsets are composition-sensitivity contrasts within one registry, not independent experiments; the endpoint remained the interpretation unit and no subset-level P values were calculated. Subsets were assigned to mutually exclusive classes: Morgan-only reference, single-representation only, single-learner only, representation-balanced without concatenation, representation-balanced with concatenation, or mixed unbalanced. We report medians, IQRs, distribution ranges and the within-endpoint proportion outperforming Morgan-only.",
      "For every paired unit": "For every paired unit we retained the endpoint, split seed, outer fold, registered candidates, selected candidate and paired inner and outer utilities. Deterministic split hashes and SHA-256 source hashes were recorded in the machine-readable supplement. Prediction-level quantities were calculated only for registries with traceable candidate-level prediction exports.",
      "The rank of the observed outer-best": "The rank of the observed outer-best candidate within the validation ordering generated several fidelity estimands. Top-3 recovery was chance adjusted as CAHit@3 = (Hit@3 − 3/K)/(1 − 3/K), and MRR was normalized against its random-order expectation H_K/K. A 5,000-replicate permutation control drew random candidate ranks within each locked outer unit and repeated the same fold-to-seed-to-endpoint aggregation as the empirical analysis. A graded positive control used the 4,320 locked outer-candidate utilities from 135 audit units and injected validation–audit signal levels of 0, 0.10, 0.25, 0.50, 0.75 and 1.0; each non-perfect cell used 500 simulations. These controls calibrated metric zero points and signal recovery and were not used to alter empirical candidate selections.",
      "Per-molecule outer predictions were available": "Per-molecule outer predictions were available for the rerun 32-candidate regression audit and for the completed four-model representation panel. Prediction correlations, high-error Jaccard overlap, Ledoit–Wolf entropy rank and participation-ratio rank were calculated only where the requisite candidate-level predictions were traceable. Prediction-level diversity is therefore reported as a boundary analysis and is not treated as equivalent to utility-pattern diversity.",
      "Ensemble uncertainty was assessed": "Ensemble uncertainty was assessed by its association with absolute error, risk-coverage behaviour and high-error enrichment. Test molecules were stratified by maximum training-set Morgan Tanimoto similarity into low (<0.5), intermediate (0.5 to <0.7) and high (at least 0.7) support. We compared selected and cross-fitted-reference performance, candidate error overlap, model disagreement, uncertainty-error association and false-negative enrichment within each stratum. Exact held-out Murcko scaffolds were unseen by construction; related versus novel scaffolds were therefore defined by a prespecified maximum scaffold-fingerprint similarity threshold of 0.5. This derived relatedness definition is reported explicitly.",
      "For endpoint-specific K = 32": "For endpoint-specific K = 32 minus K = 4 contrasts, fold differences were averaged within each of five distinct split seeds and resampled 10,000 times. Classification cross-fitted ROC-AUC loss and regression cross-fitted RMSE loss were prespecified task-stratified co-primary strata. Their units were not combined, no cross-task average effect was calculated, and endpoint effects with 95% intervals were interpreted by direction and heterogeneity. Ranking measures, same-unit gaps, effective diversity and matched-size effects were key complementary evidence. No endpoint-aggregation P values or multiplicity-adjusted confirmatory family was used.",
      "For multiview paired gains": "Matched-size subset results are distributional summaries over all 220 registered subsets, not 220 independent experiments. Medians, interquartile ranges, 95% distribution ranges and endpoint-level composition frequencies are reported. Prediction-level support and scaffold analyses use fold-level medians and retain classification and regression on separate numerical axes.",
      "The first analysis asked": f"Effective diversity depended strongly on utility-matrix construction. At K = 32, endpoint-specific Ledoit–Wolf entropy-rank point estimates summarized across the nine endpoints were {diversity_text}. Raw utilities retained common audit-unit difficulty; row centring removed common level shifts; fixed-reference estimates depended on the reference; and within-unit ranks preserved order while discarding utility spacing. The corresponding participation-ratio ranks and correlations are reported in Additional file 2: Table S6. No transformation is treated as a unique true candidate count.",
      "Leave-one-seed and leave-one-fold": "Leave-one-seed, leave-one-fold and predefined-reference analyses retained a substantial nominal-effective gap but changed endpoint magnitudes. Monte Carlo checkpoints through 5,000 replicates showed that central estimates were more stable than some tail limits. Effective diversity is therefore reported as an estimator- and matrix-dependent audit quantity rather than a single intrinsic candidate count (Additional file 2: Tables S6-S7).",
      "This analysis asked whether": f"This analysis asked whether the validation ordering recovered the outer-best candidate beyond the opportunity expected from K alone. Under the predefined fold-to-seed-to-endpoint aggregation, endpoint-median CAHit@3 was {rank4.chance_adjusted_hit_median:.3f} (IQR {rank4.chance_adjusted_hit_q25:.3f}–{rank4.chance_adjusted_hit_q75:.3f}) at K = 4, {rank8.chance_adjusted_hit_median:.3f} ({rank8.chance_adjusted_hit_q25:.3f}–{rank8.chance_adjusted_hit_q75:.3f}) at K = 8, {rank16.chance_adjusted_hit_median:.3f} ({rank16.chance_adjusted_hit_q25:.3f}–{rank16.chance_adjusted_hit_q75:.3f}) at K = 16 and {rank32.chance_adjusted_hit_median:.3f} ({rank32.chance_adjusted_hit_q25:.3f}–{rank32.chance_adjusted_hit_q75:.3f}) at K = 32. The overall loss of top-rank fidelity was nonmonotonic between K = 16 and K = 32 and cannot be attributed solely to the changing 3/K random baseline.",
      "Normalized MRR gain showed": f"Endpoint-median normalized MRR gain, using H_K/K as the random-order expectation, was {rank4.normalized_mrr_gain_median:.3f} (IQR {rank4.normalized_mrr_gain_q25:.3f}–{rank4.normalized_mrr_gain_q75:.3f}) at K = 4, {rank8.normalized_mrr_gain_median:.3f} ({rank8.normalized_mrr_gain_q25:.3f}–{rank8.normalized_mrr_gain_q75:.3f}) at K = 8, {rank16.normalized_mrr_gain_median:.3f} ({rank16.normalized_mrr_gain_q25:.3f}–{rank16.normalized_mrr_gain_q75:.3f}) at K = 16 and {rank32.normalized_mrr_gain_median:.3f} ({rank32.normalized_mrr_gain_q25:.3f}–{rank32.normalized_mrr_gain_q75:.3f}) at K = 32. From K = 4 to K = 32, endpoint-median NDCG changed from {rank4.ndcg_median:.3f} to {rank32.ndcg_median:.3f}, Spearman correlation from {rank4.spearman_median:.3f} to {rank32.spearman_median:.3f}, Kendall correlation from {rank4.kendall_median:.3f} to {rank32.kendall_median:.3f} and rank percentile from {rank4.rank_percentile_median:.3f} to {rank32.rank_percentile_median:.3f}.",
      "The opportunity-adjusted measures": f"The permutation control was centred near the adjusted zero, and observed endpoint-median CAHit@3 exceeded its 95% permutation envelope at all four K values (maximum one-sided permutation P = {null_summary.loc[null_summary.metric.eq('chance_adjusted_hit'),'one_sided_p_observed_le_null'].max():.4f}). Across K = 4, 8, 16 and 32, increasing injected validation–audit signal monotonically increased median CAHit@3 and normalized MRR gain and monotonically reduced fixed-range selection loss. The maximum absolute zero-signal CAHit@3 was {mechanism['max_null_signal_abs_cahit']:.3f}, and perfect signal reduced median selection loss to {mechanism['max_perfect_signal_selection_loss']:.3f} at every K. Thus the adjusted measures retained information above chance while exhibiting the intended null and recovery behaviour (Figure 3A–B).",
      "Separating reference selection": f"Separating reference selection from held-out-seed evaluation attenuated same-unit maxima. Cross-fitted K = 32 minus K = 4 effects were positive in six of nine endpoints ({', '.join(class_pos + reg_pos)}) and negative in three ({', '.join(class_neg + reg_neg)}). Split-seed intervals excluded zero for {', '.join(class_excludes)} in the classification ROC-AUC-loss stratum and for {', '.join(reg_excludes)} in the regression RMSE-loss stratum; FreeSolv was negative. Endpoint estimates and intervals are reported without pooling the two numerical scales.",
      "The same-unit effect was positive": f"The same-unit effect was positive for {', '.join(same_positive) if same_positive else 'no endpoint'} and negative for {', '.join(same_negative) if same_negative else 'no endpoint'}. It exceeded the cross-fitted effect in most endpoints, consistent with additional optimism when the reference is the maximum from the evaluated unit. Endpoint exceptions are retained and prevent a deterministic or universal interpretation.",
      "The matched K = 3 analysis changed": f"Exhaustive enumeration of all 220 K = 3 subsets refined the multiview interpretation. The endpoint median of subset-level selected-model gains versus Morgan K = 3 was positive for {', '.join(matched_positive) if matched_positive else 'no endpoint'} and negative for {', '.join(matched_negative) if matched_negative else 'no endpoint'}. The within-endpoint proportion of subsets with positive mean gain ranged from {subset_positive.min():.3f} to {subset_positive.max():.3f}. These proportions describe the registered composition space and do not turn 220 overlapping subsets into independent replications.",
      "All representation pairs at K = 6": f"The K = 3, 6, 9 and 12 ladder retained endpoint heterogeneity, with negative endpoint-level composition medians preserved for {', '.join(matched_negative) if matched_negative else 'no endpoint'}. Because classification ROC-AUC gains and regression RMSE reductions have different units, they are shown on separate axes and are not combined into a continuous cross-task effect.",
      "The full multiview-pool gains remained": "The full multiview-pool and K = 3, 6, 9 and 12 ladder analyses retained endpoint heterogeneity under five distinct scaffold partitions for classification and regression. Composition-class summaries are mutually exclusive, while the full set of 220 subset identities and their overlapping candidate memberships remains machine readable. Complete subset and ladder results are in Additional file 2: Tables S9-S10.",
      "This limited panel asked": "The limited four-model panel provided 90 endpoint-seed-fold prediction sets across six endpoints. Prediction correlations and error overlaps were substantial but incomplete; mean Ledoit-Wolf prediction entropy rank was below the nominal four models. These quantities characterize the completed fixed configurations and cannot rank fully optimized modern architectures.",
      "Pairwise error-overlap Jaccard": "Across chemical-support strata, pairwise high-error Jaccard overlap was non-zero and model disagreement increased at weaker support in several endpoints. Prediction-level complementarity therefore existed, but it was not equivalent to independent candidate opportunity. Unequal model-training budgets further restrict any architecture-level comparison.",
      "Prediction-level error exports": "The support-stratified audit retained one traceable selected-candidate result per split-seed outer fold and avoided plotting duplicate selected and reference points. Candidate performance, error overlap and disagreement were summarized within prespecified support strata; no support stratum was used to choose a candidate.",
      "Chemical similarity stratification": f"Chemical-support stratification showed a boundary signal. Median classification ROC-AUC across audit units was {support_medians.loc[('classification','<0.5')]:.3f} below Tanimoto 0.5, {support_medians.loc[('classification','0.5-0.7')]:.3f} from 0.5 to <0.7 and {support_medians.loc[('classification','>0.7')]:.3f} at or above 0.7. Median regression RMSE was {support_medians.loc[('regression','<0.5')]:.3f}, {support_medians.loc[('regression','0.5-0.7')]:.3f} and {support_medians.loc[('regression','>0.7')]:.3f} across the same strata. Novel-scaffold relative changes in error overlap, disagreement and high-error enrichment were reported against seen-or-related scaffolds as reference = 1.",
      "ClinTox remained the strongest negative result": "ClinTox remained the strongest negative result in the traceable four-model panel. Discrimination, minority recall, false-negative rate, label-conditional conformal coverage and mean prediction-set size are shown as distinct quantities rather than collapsed into one ranking. The observed minority-safety limitation prevents interpreting a favourable ROC-AUC as deployment readiness.",
      "ClinTox was treated": "ClinTox was treated as a prespecified negative reliability example because only 58 positive standardized structures were available. ROC-AUC, PR-AUC, minority recall, false-negative rate, label-conditional conformal coverage and mean prediction-set size were reported as distinct quantities. Coverage addresses prediction-set validity, whereas recall and false-negative rate address screening safety; none establishes deployment-level toxicology performance.",
      "ClinTox illustrates": "ClinTox illustrates why discrimination, coverage and screening safety must be separated. ROC-AUC can remain favourable while PR-AUC, minority recall and false-negative rate expose rare-class limitations. Conformal coverage indicates whether prediction sets contain the true label at the target frequency; it does not ensure that a single threshold retrieves toxic molecules at an acceptable rate. The present ClinTox results do not support deployment as a toxicity filter.",
      "ClinTox remained the principal negative result": "",
      "The primary audit shows": "The primary audit shows that nominal K and matrix-dependent utility-pattern diversity are different properties of a model search. At K = 32, median shrinkage entropy rank ranged from 2.98 for raw utilities to 27.14 for within-unit ranks. The low raw rank should not be interpreted as evidence that the registry contained only three independent predictive behaviours; it also reflects common audit-unit difficulty retained by the raw utility matrix.",
      "Effective rank itself is an estimate": "Effective diversity is not a fixed attribute of a named registry. At K = 32, each endpoint matrix contains 15 outer audit rows, so the empirical candidate correlation matrix is rank deficient. Shrinkage, hierarchical resampling and omission analyses stabilize and expose sensitivity, but they do not create independent audit units. Row centring removes a common audit-unit level shift while imposing a row-sum-zero constraint; fixed-reference estimates depend on the reference; within-unit ranks discard utility spacing; and shrinkage stabilizes the spectrum without adding data.",
      "Common outer-unit difficulty": "Candidate count, learner-family count, representation count, compute exposure and utility-pattern diversity are not interchangeable. The registered prefixes increased candidate and compute exposure mechanically, whereas effective diversity depended on how candidate utilities varied across audit units. The audit therefore treats K as search exposure and the matrix-derived ranks as conditional descriptions of observed utility patterns.",
      "Raw Top-1 recovery": "Raw Top-1 recovery and MRR become mechanically harder as K grows. Under the predefined fold-to-seed-to-endpoint aggregation, endpoint-median CAHit@3 changed from " + f"{rank4.chance_adjusted_hit_median:.3f} to {rank32.chance_adjusted_hit_median:.3f}" + ", while endpoint-median normalized MRR gain changed from " + f"{rank4.normalized_mrr_gain_median:.3f} to {rank32.normalized_mrr_gain_median:.3f}" + ". Normalized MRR used the random-order expectation H_K/K. NDCG, Spearman, Kendall and rank percentile provide complementary views of global ordering, so the interpretation rests on their joint pattern rather than one cut-off.",
      "The eight-endpoint increase": "Weaker chance-adjusted validation ranking accompanied endpoint-dependent selection gaps but did not determine their sign. Positive and negative endpoint effects are both retained, and classification and regression remain on their native ROC-AUC-loss and RMSE-loss scales. The data support a risk pattern within these endpoints, not a universal penalty from candidate expansion.",
      "No single ranking statistic": "CAHit@3 asks whether the eventual outer best entered a short validation list; normalized MRR gives graded credit to its rank; NDCG evaluates the ordered registry; and Spearman and Kendall describe broader reordering. Their chance or opportunity adjustments prevent the decline in raw winner recovery from being interpreted as evidence by itself.",
      "Cross-fitting sharpened": "The ranking results are complementary to the task-stratified co-primary cross-fitted effects. They are mechanistically compatible with selection loss, but they do not prove that ranking degradation caused every endpoint effect.",
      "The multiview stress test": "Cross-fitted references chosen without the held-out split seed attenuated same-unit maxima in most endpoints. This comparison reduces same-unit maximum reuse but still draws on the same retrospective public endpoints; it is a cross-fitted sensitivity reference, not an external validation cohort.",
      "The matched K = 3 alternatives": "Finite-audit maxima also contain winner optimism because a maximum is taken over noisy outer estimates. The equal-truth simulation shows how K, candidate correlation and effective audit size influence that optimism. These simulations explain why an observed audit best is an opportunity benchmark rather than a population-optimal bound.",
      "Nested cross-validation separates": "Classification ROC-AUC loss and regression RMSE loss form task-stratified co-primary evidence strata. They are reported on separate axes and interpreted without a cross-task average effect. Split-seed intervals quantify repeated-partition variation, but only nine endpoints and five seeds limit precision and generalization.",
      "The outer folds are also reused": "The three outer folds within a seed share training observations, and all candidates are evaluated on paired folds. Fold rows are therefore paired audit units rather than independent biological experiments. Averaging within seed before resampling preserves this dependence more faithfully than treating 15 folds as independent.",
      "The zero-width multiview intervals": "Cross-fitting did not eliminate finite-maximum optimism, public-dataset reuse or model-registry conditioning. Same-unit and cross-fitted effects should be reported together because their separation reveals how much of the observed opportunity gap depends on reusing the same audit unit to define its reference maximum.",
      "Molecular benchmark reports": "The matched-size multiview analysis changed the interpretation from 'more candidates are better' to an endpoint-dependent composition claim. Exhaustive K = 3 enumeration held subset size fixed while changing representation and learner composition; the 220 subsets overlap extensively and therefore map sensitivity within one registry rather than supply 220 replications.",
      "Ranking fidelity should also": "Mutually exclusive composition classes prevent concatenation, balance, learner and representation labels from being treated as independent groups. Endpoint-specific distributions preserve the negative and weak-gain cases, including BACE, and avoid combining classification ROC-AUC gain with regression RMSE reduction on one continuous scale.",
      "Compute accounting is part": "The K = 3, 6, 9 and 12 ladder still combines representation breadth, learner interactions and additional selection opportunities. Candidate eligibility, completed fits and compute exposure must therefore accompany the gain estimates; matched size controls one dimension of opportunity but does not equalize every training budget.",
      "A compact benchmark report": "The complete candidate ledger and subset membership are retained in machine-readable supplementary tables. This keeps the main interpretation at the endpoint level while preserving every candidate, split seed, fold, status and composition contrast for audit.",
      "The leave-one-seed-out reference": "The leave-one-seed-out reference is cross-fitted only at the repeated-partition level and still reuses public endpoints; it is not external validation. Post-hoc reruns generated all-candidate molecule-level exports for the classification 32-candidate registry. Small stochastic refit drift was detected (maximum absolute outer-candidate metric difference 0.0010, with two one-standard-error selection changes), so the locked primary metrics remained the source of record and the new exports were not used to revise primary estimates. Complete candidate-level prediction files remained unavailable across the full multiview registry; consequently, prediction-level diversity was not evaluated for the matched-subset analysis. The matched K = 3 analysis is exhaustive only for the 12-candidate registry, and its 220 overlapping subsets remain conditional on fixed candidates and compute exposure.",
      "Within nine repeated nested": "Within the studied endpoints, candidate-pool expansion was accompanied by weaker chance-adjusted validation ranking and heterogeneous model-selection gaps. Estimated candidate diversity depended strongly on whether common audit-unit difficulty and utility-level shifts were removed. Cross-fitted references attenuated same-unit maxima, while matched-size analyses showed that representation composition could produce endpoint-dependent gains.",
      "Molecular model-selection studies should": "Molecular model-selection studies should jointly report candidate eligibility, nominal K, matrix-dependent utility-pattern diversity, chance-adjusted ranking fidelity, endpoint-specific same-unit and cross-fitted gaps, split uniqueness, computational exposure, failed candidates and chemical-support boundaries. These quantities support transparent audit interpretation but do not define a universal selector or a deployment-ready screening system.",
    }
    for prefix,text in replacements.items(): set_text(find(doc,prefix),text)

    # Explicit first citations keep every main figure and table connected to the narrative.
    citations={
      "The study was a retrospective audit":" (Figure 1).",
      "Cleaning affected endpoints differently":" Dataset-level summaries are provided in Table 1.",
      "Per-candidate opportunity was held constant":" Candidate registries and recorded exposure are summarized in Table 2.",
      "Leave-one-seed, leave-one-fold":" (Figure 2; Additional file 2: Tables S6-S7).",
      "The decomposition shows":" (Figure 4A-B; Additional file 2: Table S8).",
      "Separating reference selection":" (Figure 3C; Table 3).",
      "At candidate correlation 0.9":" (Figure 4C-D).",
      "The full multiview-pool and":" (Figure 5; Additional file 2: Tables S9-S10).",
      "The support-stratified audit":" (Figure 6A-B).",
      "ClinTox remained the strongest negative result":" (Figure 6C-D).",
    }
    for prefix,suffix in citations.items():
        p=find(doc,prefix)
        # Remove superseded table-only endings before adding the complete citation once.
        text=re.sub(r"\s*\(Additional file 2: Tables? S[0-9-]+\)\.?$","",p.text)
        if suffix.strip() not in text: set_text(p,text.rstrip()+suffix)

    headings={"3.1 Nominal expansion":"3.1 Effective diversity depended on matrix construction",
      "3.3 Audit-gap":"3.3 Expansion increased observed opportunity more than selected gain",
      "3.4 Cross-fitted effects remained":"3.4 Cross-fitted effects were heterogeneous across endpoints",
      "3.6 Matched-size":"3.6 Matched-size representation effects were endpoint dependent",
      "3.7 Representation errors":"3.7 Prediction errors were correlated but partially complementary",
      "3.8 Reliability":"3.8 Reliability weakened at chemical-support boundaries",
      "4.1 Nominal K":"4.1 Nominal K did not define candidate diversity",
      "4.2 Chance-adjusted":"4.2 Matrix construction changed effective-rank interpretation",
      "4.3 Matched K":"4.3 Chance-adjusted ranking degradation accompanied selection gaps",
      "4.4 Same-unit":"4.4 Cross-fitting attenuated same-unit maxima",
      "4.5 Implications":"4.5 Matched-size analyses refined the multiview interpretation",
      "4.6 Reliability remains":"4.6 Reliability remained conditional on chemical support"}
    for prefix,text in headings.items(): set_text(find(doc,prefix),text)

    replace_between(doc,"1 Introduction","2 Methods",[(text,"Normal") for text in INTRODUCTION_PARAGRAPHS])
    discussion_items=[]
    for heading,paragraphs in DISCUSSION_SECTIONS:
        discussion_items.append((heading,"Heading 2"))
        discussion_items.extend((text,"Normal") for text in paragraphs)
    replace_between(doc,"4 Discussion","5 Conclusions",discussion_items)

    # Crossref-verified title correction; DOI and all remaining metadata are unchanged.
    set_text(find(doc,"30. Sheridan RP."),"30. Sheridan RP. Time-split cross-validation as a method for estimating the goodness of prospective prediction. J Chem Inf Model. 2013;53:783–790. doi:10.1021/ci400084k.")

    captions={
      "Figure 1.":"Figure 1. Retrospective nested audit of candidate-pool expansion and model-selection loss. The central module separates candidate ranking in three inner scaffold folds from one-shot auditing in three outer folds across five seeded partitions, with no outer-label feedback. Registered molecular endpoints and prespecified candidate prefixes enter from the left. Surrounding branches summarize matrix-dependent diversity, chance-adjusted ranking, same-unit and cross-fitted selection gaps, candidate-composition controls, equal-size multiview stress tests, finite-audit winner optimism, chemical-support boundaries and four-model prediction/error reliability. These streams converge on an auditable evidence map rather than a single model leaderboard.",
      "Figure 2.":"Figure 2. Matrix-dependent candidate diversity. (A) Ledoit–Wolf entropy rank across raw, row-centred, fixed-reference-relative and within-unit-rank matrices; solid lines are endpoint medians, ribbons are endpoint IQRs and the grey dotted line is nominal K. Raw utilities retain common audit-unit difficulty, whereas row centring removes common level shifts. (B) Entropy and participation-ratio concordance at K = 32 with a 45° reference. (C) Median candidate correlation after adjustment. (D) Fixed-reference-relative endpoint estimates at K = 32; solid and dashed lines are leave-one-seed and leave-one-fold ranges, respectively, and hollow points show predefined references.",
      "Figure 3.":"Figure 3. Chance-adjusted ranking calibration and cross-fitted gaps. (A) Endpoint-median CAHit@3 and normalized MRR gain with endpoint IQRs; the grey band is the joint 95% envelope from 5,000 random-rank permutations under the same aggregation. (B) Positive-control recovery of endpoint-median CAHit@3 across six injected validation–audit signal levels and four candidate counts. (C) Task-stratified cross-fitted classification ROC-AUC loss and regression RMSE loss with split-seed bootstrap 95% intervals; filled markers exclude zero, hollow markers include zero, and colour denotes task type rather than significance. (D) Registered-prefix, random-order, random-subset and family-balanced controls; all modes equal the complete registry at K = 32 by design.",
      "Figure 4.":"Figure 4. Selection-gap decomposition and winner optimism. Panels A and B compare observed audit-best and validation-selected gains for classification and regression. Panel C presents all endpoints in one task-stratified forest plot using within-endpoint normalized effects; raw ROC-AUC and RMSE effects are listed separately. Panel D shows finite-audit winner optimism under equal candidate truth.",
      "Figure 5.":"Figure 5. Matched-size representation-composition effects. Panel A shows endpoint-specific distributions across all 220 registered K = 3 subsets. Panel B summarizes mutually exclusive composition classes. Panel C shows the K = 3, 6, 9 and 12 ladder. Panel D presents all endpoints in one integrated forest plot using endpoint-MAD-normalized effects, with raw ROC-AUC gains and RMSE reductions listed separately.",
      "Figure 6.":"Figure 6. Prediction reliability across chemical-support boundaries. Panel A combines prediction correlations and high-error Jaccard overlaps in a symmetric four-model matrix. Panel B consolidates classification ROC-AUC, regression RMSE, classification false-negative rate and classification/regression high-error enrichment into one Tanimoto-support risk matrix; cell text gives natural-scale medians and colour encodes only the within-row adverse direction. Panel C reports novel-scaffold relative changes in error overlap, model disagreement, high-error enrichment and false-negative enrichment on a log-ratio axis. Panel D compares all four fixed-configuration models in one aligned dot plot across ROC-AUC, PR-AUC, minority recall, false-negative rate, conditional coverage and prediction-set size; the set-size values are mapped from their 1-2 scale only for display.",
    }
    for prefix,text in captions.items(): set_text(find(doc,prefix),text)

    # Remove old embedded drawings, then insert current 600-dpi figures before captions.
    for p in list(doc.paragraphs):
        if p._element.xpath('.//w:drawing'): delete_paragraph(p)
    paths=["Figure_1_retrospective_nested_audit_architecture.png","Figure_2_candidate_diversity_after_adjustment.png","Figure_3_chance_adjusted_ranking_and_selection_gaps.png","Figure_4_selection_gap_and_winner_optimism.png","Figure_5_matched_size_multiview_composition.png","Figure_6_prediction_errors_across_chemical_support.png"]
    for i,name in enumerate(paths,1): insert_figure(find(doc,f"Figure {i}."),FIG/name,6.7)

    # Keep the three main tables compact and move complete registries to the supplement.
    table1=doc.tables[0]
    table1.rows[0].cells[1].text="Analysis n"
    for row in table1.rows[1:]:
        endpoint=row.cells[0].text.strip()
        if endpoint=="ESOL": row.cells[2].text="-11.60 to 1.58"
        elif endpoint=="FreeSolv": row.cells[2].text="-25.47 to 3.43"
        elif endpoint=="Lipophilicity": row.cells[2].text="-1.50 to 4.50"
        elif endpoint.startswith("Caco2"): row.cells[2].text="-7.76 to -3.51"
    set_text(find(doc,"Table 1."),"Table 1. Datasets and primary endpoint metrics. Units: ESOL, log mol/L; FreeSolv, kcal/mol; Lipophilicity, experimental logD; Caco2 Wang, the dataset-provided log-permeability scale. Classification rows report positive-class n (%).")
    table2=doc.tables[1]
    table2_rows=[
        ["Analysis","Registry","Resampling design","Purpose","Exposure"],
        ["Controlled prefix audit","32 Morgan candidates","5 seeded scaffold partitions; 3 outer × 3 inner folds","Expansion audit","17,280 candidate fits"],
        ["Mechanism controls","32 Morgan candidates","5,000 permutations; 6 signal levels × 4 K values","Null and signal recovery","No model fitting"],
        ["Composition controls","32 Morgan candidates","100 subset seeds per mode, K and split seed","Composition sensitivity","No additional fitting"],
        ["Matched multiview","12 multiview candidates","5 seeded scaffold partitions; exhaustive C(12,3)","Composition effects","6,480 candidate fits"],
        ["Four-model panel","4 representation candidates","Traceable outer prediction panel","Reliability boundary","360 candidate–fold units"],
    ]
    while len(table2.rows)<len(table2_rows): table2.add_row()
    for row,values in zip(table2.rows,table2_rows):
        for cell,value in zip(row.cells,values): cell.text=value
    set_text(find(doc,"Table 3."),"Table 3. Endpoint-specific cross-fitted K = 32 minus K = 4 effects in the task-stratified co-primary evidence strata.")

    # Replace Table 3 with the primary cross-fitted endpoint results.
    old=doc.tables[2]; old_el=old._element; parent=old_el.getparent(); pos=parent.index(old_el); parent.remove(old_el)
    table=doc.add_table(rows=1,cols=5); table.style="Table Grid"
    hdr=["Endpoint","Metric","Cross-fitted effect","95% CI","Interpretation"]
    for c,t in zip(table.rows[0].cells,hdr): c.text=t
    for _,r in cross.iterrows():
        cells=table.add_row().cells; metric="ROC-AUC loss" if r.task_type=="classification" else "RMSE loss"
        lower=r.split_seed_bootstrap95_low_cross_fitted; upper=r.split_seed_bootstrap95_high_cross_fitted
        uncertainty=f"{lower:.4f} to {upper:.4f}"
        if lower>0:
            interpretation="Greater loss at K = 32"
        elif upper<0:
            interpretation="Lower loss at K = 32"
        elif r.cross_fitted_effect>=0:
            interpretation="Uncertain; point estimate positive"
        else:
            interpretation="Uncertain; point estimate negative"
        vals=[DISPLAY[r.task],metric,f"{r.cross_fitted_effect:.4f}",uncertainty,interpretation]
        for c,t in zip(cells,vals): c.text=t
    parent.insert(pos,table._element)
    for t in doc.tables: three_line(t)

    set_text(find(doc,"Additional file 2."),"Additional file 2. Machine-readable Supplementary Tables (Tables S1-S22; XLSX).")
    set_text(find(doc,"Additional file 3."),"Additional file 3. Supplementary Figures (Figures S1-S17; PDF).")
    for i,p in enumerate(doc.paragraphs):
        if p.text.strip()=="Ethics approval and consent to participate" and i+1<len(doc.paragraphs):
            set_text(doc.paragraphs[i+1],"Not applicable. This study used publicly available molecular datasets and involved no human participants, human data or animal experiments.")
        elif p.text.strip()=="Consent for publication" and i+1<len(doc.paragraphs):
            set_text(doc.paragraphs[i+1],"Not applicable.")

    # Verified statements can be injected at build time; transparent placeholders remain otherwise.
    declarations = {
        "Competing interests": os.environ.get(
            "FZYC_COMPETING_INTERESTS",
            "Author confirmation required before submission; no competing-interest statement has been inferred.",
        ),
        "Funding": os.environ.get(
            "FZYC_FUNDING",
            "Author confirmation required before submission; no funding source, grant number or funder role has been inferred.",
        ),
        "Authors' contributions": os.environ.get(
            "FZYC_AUTHOR_CONTRIBUTIONS",
            "Verified author initials and CRediT roles are required before submission; contribution roles have not been inferred.",
        ),
        "Acknowledgements": os.environ.get(
            "FZYC_ACKNOWLEDGEMENTS",
            "Author confirmation or verified acknowledgement text is required before submission; no acknowledgement has been inferred.",
        ),
    }
    for heading, statement in declarations.items():
        set_after_heading(doc, heading, statement)

    # Journal layout: editable text, double spacing, continuous line numbering, automatic pages, no manual breaks.
    for p in doc.paragraphs:
        p.paragraph_format.page_break_before=False
        for br in p._element.xpath('.//w:br[@w:type="page"]'): br.getparent().remove(br)
        if not p.style.name.startswith("Heading") and p.style.name!="Figure Caption": p.paragraph_format.line_spacing=2.0
        p.paragraph_format.space_after=Pt(0 if p.style.name=="Normal" else 6)
        for r in p.runs:
            r.font.name="Times New Roman"; r._element.rPr.rFonts.set(qn("w:eastAsia"),"Times New Roman")
    for section in doc.sections:
        sectPr=section._sectPr
        for old_ln in sectPr.findall(qn("w:lnNumType")): sectPr.remove(old_ln)
        ln=OxmlElement("w:lnNumType"); ln.set(qn("w:countBy"),"1"); ln.set(qn("w:restart"),"continuous"); sectPr.append(ln)
        footer=section.footer; footer.is_linked_to_previous=False
        for p in footer.paragraphs: p.clear()
        page_field(footer.paragraphs[0])
    doc.core_properties.title="Candidate-pool expansion, validation-ranking distortion and model-selection loss in molecular property prediction: a retrospective nested audit"
    doc.save(OUT)
    embed_svg_figures(OUT, FIG)
    print(OUT)


if __name__=="__main__": main()
