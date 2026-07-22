from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\fzyc\output\paper32_equation_table_format_20260718")
SOURCES = [
    (
        ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_equations_tables_large_Figure7.docx",
        ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_final_unified_format.docx",
        "en",
    ),
    (
        ROOT / "Chinese_manuscript_equations_tables_large_Figure7.docx",
        ROOT / "Chinese_manuscript_final_unified_format.docx",
        "zh",
    ),
]


def set_style_font(style, latin: str, east_asia: str, size: float, bold: bool | None = None) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)


def set_run_font(run, latin: str, east_asia: str, size: float, bold: bool | None = None) -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)


def has_math(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//m:oMath"))


def has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def ensure_style(document: Document, name: str, base: str = "Normal"):
    if name in document.styles:
        return document.styles[name]
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = document.styles[base]
    return style


def configure_styles(document: Document, language: str) -> dict[str, object]:
    latin = "Times New Roman"
    body_east = "宋体" if language == "zh" else latin
    heading_east = "黑体" if language == "zh" else latin
    normal = document.styles["Normal"]
    set_style_font(normal, latin, body_east, 11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25 if language == "zh" else 1.15
    normal.paragraph_format.space_after = Pt(3 if language == "en" else 0)

    title = ensure_style(document, "Title")
    set_style_font(title, latin, heading_east, 16, True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_with_next = True

    for style_name, size, before, after in [
        ("Heading 1", 14, 12, 6),
        ("Heading 2", 12, 9, 4),
        ("Heading 3", 11, 6, 3),
    ]:
        style = ensure_style(document, style_name)
        set_style_font(style, latin, heading_east, size, True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for caption_name in ["Caption", "Figure Caption"]:
        style = ensure_style(document, caption_name)
        set_style_font(style, latin, body_east, 9)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_together = True
    return {"latin": latin, "body_east": body_east, "heading_east": heading_east}


def format_document(source: Path, destination: Path, language: str) -> dict[str, object]:
    document = Document(source)
    fonts = configure_styles(document, language)
    latin = fonts["latin"]
    body_east = fonts["body_east"]
    heading_east = fonts["heading_east"]

    for section in document.sections:
        section.top_margin = Cm(2.25)
        section.bottom_margin = Cm(2.25)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)

    in_references = False
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else "Normal"
        if index == 0:
            paragraph.style = document.styles["Title"]
            style_name = "Title"
        if text in {"References", "参考文献"}:
            in_references = True

        if has_drawing(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.keep_with_next = True
            continue
        if has_math(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.keep_together = True
            continue

        is_caption = style_name in {"Caption", "Figure Caption"} or text.startswith(("Figure ", "Table ", "图", "表"))
        is_heading = style_name.startswith("Heading")
        if is_caption:
            paragraph.style = document.styles["Figure Caption" if text.startswith(("Figure ", "图")) else "Caption"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = None
            size = 9
            east = body_east
        elif style_name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            size = 16
            east = heading_east
        elif is_heading:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = None
            size = {"Heading 1": 14, "Heading 2": 12, "Heading 3": 11}.get(style_name, 11)
            east = heading_east
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.paragraph_format.line_spacing = 1.25 if language == "zh" else 1.15
            paragraph.paragraph_format.space_after = Pt(3 if language == "en" else 0)
            if in_references and text:
                paragraph.paragraph_format.left_indent = Cm(0.63)
                paragraph.paragraph_format.first_line_indent = Cm(-0.63)
            elif language == "zh" and index > 7 and text and not text.startswith(("关键词", "Keywords")):
                paragraph.paragraph_format.left_indent = None
                paragraph.paragraph_format.first_line_indent = Cm(0.74)
            else:
                paragraph.paragraph_format.left_indent = None
                paragraph.paragraph_format.first_line_indent = None
            size = 11
            east = body_east
        for run in paragraph.runs:
            set_run_font(run, latin, east, size, True if is_heading or style_name == "Title" else None)

    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        set_run_font(run, latin, body_east, 9, True if row_index == 0 else None)

    document.save(destination)
    with ZipFile(destination) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    return {
        "source": str(source),
        "output": str(destination),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "native_equations": xml.count("<m:oMath>"),
        "language": language,
    }


def main() -> None:
    results = [format_document(source, destination, language) for source, destination, language in SOURCES]
    (ROOT / "Unified_manuscript_formatting_audit.json").write_text(
        json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(results, ensure_ascii=False))


if __name__ == "__main__":
    main()
