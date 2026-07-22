from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


OUT = Path(r"D:\fzyc\output\paper33_final_minor_revision_20260718")
EN_BASE = Path(r"D:\fzyc\output\paper32_equation_table_format_20260718\Candidate_pool_expansion_Journal_of_Cheminformatics_final_unified_format.docx")
ZH_BASE = Path(r"C:\Users\Administrator\Desktop\Chinese_manuscript_final_Times_New_Roman_figures.docx")
EN_OUT = OUT / "Candidate_pool_expansion_Journal_of_Cheminformatics_FINAL_CLEAN.docx"
ZH_OUT = OUT / "候选池扩张与模型选择损失_中文终稿.docx"
OUT.mkdir(parents=True, exist_ok=True)


def find(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(prefix)


def replace_prefix(doc: Document, prefix: str, text: str) -> None:
    paragraph = find(doc, prefix)
    paragraph.text = text


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def replace_section(doc: Document, start: str, end: str, blocks: list[tuple[str, str]]) -> None:
    start_p = find(doc, start)
    end_p = find(doc, end)
    paragraphs = doc.paragraphs
    i = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is start_p._p)
    j = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is end_p._p)
    preserved = []
    for paragraph in paragraphs[i + 1 : j]:
        has_drawing = bool(paragraph._p.xpath(".//w:drawing"))
        style_name = paragraph.style.name if paragraph.style is not None else ""
        is_caption = style_name == "Caption" or paragraph.text.strip().startswith(("Figure ", "图"))
        if has_drawing or is_caption:
            preserved.append(paragraph._p)
        else:
            remove_paragraph(paragraph)
    for style, text in blocks:
        end_p.insert_paragraph_before(text, style=style)
    # Move the retained figure paragraphs and captions after the revised prose,
    # preserving their original order and relationships.
    for element in preserved:
        end_p._p.addprevious(element)


def border(parent, edge: str, val: str, size: str = "8") -> None:
    node = parent.find(qn(f"w:{edge}"))
    if node is None:
        node = OxmlElement(f"w:{edge}")
        parent.append(node)
    node.set(qn("w:val"), val)
    node.set(qn("w:sz"), size)
    node.set(qn("w:color"), "000000")


def three_line(table) -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge, val, size in [
        ("top", "single", "12"),
        ("left", "nil", "0"),
        ("bottom", "single", "12"),
        ("right", "nil", "0"),
        ("insideH", "nil", "0"),
        ("insideV", "nil", "0"),
    ]:
        border(borders, edge, val, size)
    tbl_pr.append(borders)
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        border(tc_borders, "bottom", "single", "8")
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.bold = True
                    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                    rfonts.set(qn("w:ascii"), "Times New Roman")
                    rfonts.set(qn("w:hAnsi"), "Times New Roman")
                    rfonts.set(qn("w:eastAsia"), "宋体")


def notation_table(doc: Document, before_prefix: str, language: str) -> None:
    before = find(doc, before_prefix)
    table = doc.add_table(rows=1, cols=2)
    headers = ["Symbol", "Definition"] if language == "en" else ["符号", "定义"]
    rows_en = [
        ("e", "endpoint"), ("s", "split seed"), ("f", "outer fold"),
        ("u = (s, f)", "outer audit unit"), ("j", "candidate"),
        ("p", "registry composition"), ("K", "candidate count"),
        ("V(u,j)", "inner-validation utility"), ("A(u,j)", "outer-audit utility"),
    ]
    rows_zh = [
        ("e", "终点"), ("s", "切分种子"), ("f", "外层折"),
        ("u = (s, f)", "外层审计单元"), ("j", "候选"),
        ("p", "候选池组成"), ("K", "候选数量"),
        ("V(u,j)", "内层验证效用"), ("A(u,j)", "外层审计效用"),
    ]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in (rows_en if language == "en" else rows_zh):
        cells = table.add_row().cells
        cells[0].text, cells[1].text = values
    three_line(table)
    before._p.addprevious(table._tbl)


EN_ABSTRACT = {
    "Background:": "Background: Molecular property-prediction studies increasingly compare correlated registries of representations, learners and tuning variants. Candidate expansion can introduce complementary chemical information while increasing repeated selection pressure on finite validation data.",
    "Methods:": "Methods: We conducted a retrospective, task-stratified nested audit across nine public endpoints and a prespecified six-endpoint composition intervention. The expanded intervention included component, downstream-budget, anchor, normalization, selection-stability and split-mechanism sensitivity analyses.",
    "Results:": "Results: The composition intervention covered six prespecified endpoints. At K = 32, modern augmentation produced positive validation-selected gains relative to the homogeneous registry across all six endpoints, while component, compute-matched, stability and split analyses showed that realised benefit remained conditional on validation alignment and registry composition.",
    "Conclusions:": "Conclusions: Candidate expansion creates both opportunity and selection pressure. Matrix-dependent diversity is descriptive rather than causal, and net value depends on complementarity, validation alignment and bounded downstream cost within the evaluated endpoints, registries and split mechanisms.",
    "Scientific Contribution:": "Scientific Contribution: This study jointly audits candidate count, registry composition, validation-ranking distortion, cross-fitted gaps, selection stability and bounded downstream compute in molecular property prediction. It provides an auditable domain-specific decomposition rather than a new universal statistical law, general-purpose selector or modern-architecture leaderboard.",
}

ZH_ABSTRACT = {
    "背景：": "背景：分子性质预测研究日益比较由表征、学习器与调参变体构成的相关候选池。候选扩张可引入互补化学信息，同时增加有限验证数据上的重复选择压力。",
    "方法：": "方法：本研究在九个公开端点开展回顾性、任务分层的嵌套审计，并加入预设的六终点候选池组成干预。扩展干预包含组件、下游预算、锚点、归一化、选择稳定性和切分机制敏感性分析。",
    "结果：": "结果：组成干预覆盖六个预设端点。K = 32 时，现代增强相对同质候选池的验证选择增益在六个端点均为正；组件、计算匹配、稳定性与切分分析表明，实际收益仍取决于验证一致性和候选池组成。",
    "结论：": "结论：候选扩张同时创造机会并增加选择压力。矩阵依赖的多样性具有描述性而非因果性；在已评估的端点、候选池和切分机制内，净价值取决于互补性、验证一致性与受限下游成本。",
    "科学贡献：": "科学贡献：本研究联合审计分子性质预测中的候选数量、候选池组成、验证排序失真、交叉拟合差距、选择稳定性与受限下游计算。本文提供领域特异的可审计分解，不提出新的普适统计定律、通用选择器或现代架构排行榜。",
}


def update_table_2(doc: Document, language: str) -> None:
    table = doc.tables[1]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    if language == "en":
        headers = ["Audit component", "Registry", "Evaluation design", "Recorded exposure"]
        rows = [
            ["Controlled prefix audit", "32 Morgan candidates", "9 endpoints; K = 4, 8, 16 and 32; 5 seeds; 3 outer × 3 inner", "17,280 candidate fits; 4,437.95 candidate-fit s"],
            ["Calibration and resampling controls", "Stored 32-candidate results", "5,000 permutations; 100 resamples per mode/K/seed", "No additional model fitting"],
            ["Matched multiview audit", "12 multiview candidates", "9 endpoints; 5 seeds; all C(12,3) registered subsets", "6,480 candidate fits"],
            ["Split-regime transfer audit", "32 Morgan candidates", "3 endpoints; seeded scaffold vs Tanimoto-component splits", "5,760 candidate fits"],
            ["Expanded registry-composition intervention", "Three locked candidate registries", "6 endpoints; K = 4, 8, 16 and 32; component, budget, anchor, normalization, stability and split controls", "1,080 primary outer units; 64,616.35 downstream s recorded; total fit count not reconstructed"],
        ]
    else:
        headers = ["审计组成", "候选池", "评价设计", "记录的计算暴露"]
        rows = [
            ["受控前缀审计", "32个Morgan候选", "9端点；K=4/8/16/32；5种子；3外层×3内层", "17,280次候选拟合；4,437.95候选拟合秒"],
            ["校准与重采样控制", "已存储的32候选结果", "5,000次置换；每模式/K/种子100次重采样", "无新增模型拟合"],
            ["匹配多表征审计", "12个多表征候选", "9端点；5种子；全部C(12,3)登记子集", "6,480次候选拟合"],
            ["切分机制迁移审计", "32个Morgan候选", "3端点；种子化骨架与Tanimoto分量切分", "5,760次候选拟合"],
            ["扩展候选池组成干预", "三类锁定候选池", "6端点；K=4/8/16/32；组件/预算/锚点/归一化/稳定性/切分控制", "1,080主要外层单元；64,616.35下游秒；总拟合数未重构"],
        ]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    three_line(table)


def localize_zh_tables(doc: Document) -> None:
    table1, _, table3 = doc.tables
    headers = ["终点", "样本量", "类别平衡/目标范围", "指标"]
    for cell, text in zip(table1.rows[0].cells, headers):
        cell.text = text
    for row in table1.rows[1:]:
        value = row.cells[2].text.replace(" positive", " 阳性").replace(" to ", "至")
        row.cells[2].text = re.sub(r"(?<![A-Za-z])-(?=\d)", "−", value)
    table3_headers = ["终点", "效应（95%区间）", "方向"]
    for cell, text in zip(table3.rows[0].cells, table3_headers):
        cell.text = text
    replacements = {
        "Classification: ROC-AUC loss": "分类：ROC-AUC损失",
        "Regression: RMSE loss": "回归：RMSE损失",
        "Uncertain": "不确定",
        "Greater loss": "损失增加",
        "Lower loss": "损失降低",
    }
    for row in table3.rows[1:]:
        for cell in row.cells:
            value = cell.text
            for old, new in replacements.items():
                value = value.replace(old, new)
            cell.text = re.sub(r"(?<![A-Za-z])-(?=\d)", "−", value)
    three_line(table1)
    three_line(table3)
    note = doc.add_paragraph("注：分类终点以阳性例数（比例）表示类别平衡；回归终点按原始数据集的目标单位报告范围。")
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(3)
    for run in note.runs:
        run.font.name = "宋体"
        run.font.size = Pt(9)
    table1._tbl.addnext(note._p)


def build(language: str, source: Path, destination: Path) -> None:
    shutil.copy2(source, destination)
    doc = Document(destination)
    abstracts = EN_ABSTRACT if language == "en" else ZH_ABSTRACT
    for prefix, text in abstracts.items():
        replace_prefix(doc, prefix, text)

    if language == "en":
        replace_prefix(doc, "Under a retrospectively locked", (
            "We therefore conducted a retrospectively locked, task-stratified audit rather than a model leaderboard. The evidence comprised a nine-endpoint primary audit and a six-endpoint expanded composition intervention. Locked registries were evaluated at K = 4, 8, 16 and 32, with component ablations, equal-downstream-budget sensitivity, selection-stability analysis, anchor and normalization sensitivity, and three-endpoint split-regime transport. These layers address distinct questions and are interpreted only within the evaluated endpoints, registries and split mechanisms."
        ))
        replace_section(doc, "2.1 ", "2.2 ", [
            ("Normal", "The study was a retrospective audit of completed molecular-property experiments. Endpoints, candidate order, K values, seeds, split logic, selection rules, outcomes and exclusions were held fixed for the analyses reported here. Because the audit specification was reconstructed after the original outer outcomes existed, the study was not prospectively preregistered; frozen denotes unchanged eligibility and analysis rules during this audit, not an untouched prospective cohort (Figure 1)."),
            ("Normal", "Evidence was organized into five levels. Level 1 was the nine-endpoint controlled prefix audit. Level 2 comprised permutation calibration, composition controls and finite-audit simulation. Level 3 was the matched-size twelve-candidate multiview audit. Level 4 was the six-endpoint expanded registry-composition intervention with component, budget, anchor, normalization, stability and split-mechanism sensitivity analyses. Level 5 comprised prediction-level reliability and chemical-support boundary evidence."),
            ("Normal", "The levels answer different questions and were not merged into a single leaderboard. Outer folds audited frozen inner-selection decisions; subsequent outer-best, ranking and selection-loss decompositions are retrospective finite-audit quantities rather than population bounds. The multiview, composition and prediction panels reuse public endpoint definitions and do not constitute independent external validation."),
        ])
        replace_prefix(doc, "Per-candidate opportunity was held", (
            "Per-candidate opportunity was held constant across K: the same preprocessing, inner folds and refit logic applied whenever a candidate was eligible, while total search exposure increased with the registry. The controlled audit recorded 17,280 candidate fits and 4,437.95 candidate-fit seconds; calibration and composition controls reused stored results. The expanded intervention contained 1,080 primary outer units across six endpoints, three registries, four K values, five seeds and three outer folds, with 64,616.35 downstream fit/predict seconds recorded at unit level. Total expanded-intervention fit count was not reconstructed. Exposure measures are analysis specific and are summarized in Table 2; complete resources are in Tables S2–S3 and S28–S36."
        ))
        replace_prefix(doc, "Search budgets differed", "Search budgets differed across candidates, so the limited four-model panel was restricted to selector behaviour and error correlation under shared splits. D-MPNN was not included in that prediction-level panel; a separately locked one-epoch D-MPNN candidate was included only in the expanded registry-composition intervention. TabPFN was not included because a complete shared-split audit was unavailable.")
        replace_section(doc, "2.14 ", "2.15 ", [
            ("Heading 3", "Endpoint and registry design"),
            ("Normal", "Six endpoints were fixed from task coverage before expanded outcomes were read: ClinTox, BACE and BBBP for classification, and ESOL, Lipophilicity and Caco2_Wang for regression. Homogeneous Morgan, classical multiview and modern-augmented exact-prefix registries were evaluated at K = 4, 8, 16 and 32 with seeds 11, 23, 37, 53 and 71 and identical three-outer by three-inner folds. The shared Morgan linear candidate occupied the first locked position; candidate configuration and order did not change across eligible prefixes."),
            ("Heading 3", "Component ablations"),
            ("Normal", "Fixed-K ablations compared classical multiview, classical plus ChemBERTa, classical plus MoLFormer, classical plus D-MPNN and the full modern-augmented registry at K = 16 and 32. Each addition replaced a prespecified classical candidate so that K remained exact. The separately locked one-epoch D-MPNN belonged only to this expanded intervention."),
            ("Heading 3", "Equal-downstream-budget analysis"),
            ("Normal", "Each locked prefix was truncated at the endpoint- and K-specific median downstream time of the classical multiview registry; outer performance never determined retention. Recorded time includes inner and outer downstream fit/predict time but excludes model acquisition, encoder pretraining and cached embedding extraction. This is a bounded downstream-compute sensitivity, not an end-to-end efficiency comparison."),
            ("Heading 3", "Anchor, normalization, selection-stability and split-regime sensitivity"),
            ("Normal", "Anchor sensitivity used the shared Morgan linear model, a fixed Morgan random forest and the predefined registry-median candidate. Raw gains, endpoint-MAD normalization and paired homogeneous-audit-best normalization were reported. Candidate, representation and learner-family frequencies, normalized selection entropy, modal proportion, leave-one-seed-out agreement and adjacent-fold switches quantified stability. ClinTox, BACE and ESOL completed the three-registry loop under seeded scaffold and Tanimoto-component splits with registries, seeds, folds, rules and metrics held fixed."),
        ])
        replace_section(doc, "2.16 ", "3 Results", [
            ("Normal", "For endpoint-specific K = 32 minus K = 4 contrasts, fold differences were averaged within each of five split seeds and the seed blocks were resampled 10,000 times. Classification cross-fitted ROC-AUC loss and regression cross-fitted RMSE loss were prespecified task-stratified co-primary strata; raw effects were not pooled, no cross-task average was calculated and no endpoint-aggregation P value or confirmatory multiplicity family was used."),
            ("Normal", "Endpoint–pool–K cells are not independent. Fold effects were averaged within seed, seed-block uncertainty was retained and direction counts were interpreted descriptively. Component and budget comparisons were paired within endpoint, seed and fold. Matched-size results summarize 220 overlapping registered subsets rather than independent experiments; support and scaffold analyses retain classification and regression on separate numerical axes."),
        ])
        replace_section(doc, "3.1 ", "3.2 ", [
            ("Normal", "At K = 32, the endpoint medians of Ledoit–Wolf entropy rank were 2.98 for raw utilities, 24.23 after row centring, 5.86 for fixed-reference-relative utilities and 27.14 for within-unit ranks. These matrices retain different information: common audit difficulty, level-shift removal, dependence on a prespecified reference or ordering without utility spacing, respectively. No transformation is treated as a unique true candidate count (Figure 2; Tables S6–S7)."),
            ("Normal", "Estimator, seed, fold and reference sensitivities changed magnitude while retaining a nominal–effective gap. Full IQRs, ranges, participation-ratio ranks and correlation summaries are reported in Table S6 rather than repeated in the main text."),
        ])
        replace_section(doc, "3.2 ", "3.3 ", [
            ("Normal", "Ranking fidelity declined overall as K expanded, although the K = 16 to K = 32 pattern was not monotonic. Endpoint-median CAHit@3 changed from 1.000 at K = 4 to 0.264 at K = 32, and normalized MRR gain changed from 0.861 to 0.157; NDCG and rank correlations provided concordant but non-identical views of reordering."),
            ("Normal", "The permutation negative control was centred near the adjusted zero. In the graded positive control, increasing injected validation–audit signal increased CAHit@3 and normalized MRR and reduced fixed-range selection loss, establishing face validity of the metrics without proving a biochemical mechanism (Figure 3A–B)."),
        ])
        replace_section(doc, "3.3 ", "3.4 ", [
            ("Normal", "For classification, mean observed audit-best gain increased from 0.0028 to 0.0311 ROC-AUC between K = 4 and K = 32, while selected-model gain increased from −0.0014 to 0.0173. For regression, observed audit-best gain increased from 0.7619 to 0.8078 RMSE units, while selected-model gain changed from 0.7442 to 0.7378; classification and regression scales were not pooled."),
            ("Normal", "Across both task strata, expansion increased observed opportunity faster than realised gain. Near-zero audit-best denominators made realization ratios unstable, so they remain secondary machine-readable quantities (Figure 4A–B; Table S8)."),
        ])
        replace_section(doc, "3.4 ", "3.5 ", [
            ("Normal", "Cross-fitted K = 32 minus K = 4 effects were positive in six endpoints and negative in three. Five split-seed intervals excluded zero; FreeSolv was the negative exception. Table 3 retains all endpoint estimates, intervals and uncertain directions without pooling classification and regression scales (Figure 3C; Figure 4C; Table S8)."),
            ("Normal", "The leave-one-seed reference attenuates the circularity of a same-unit maximum but remains conditional on the same public endpoints and split generator. Positive and negative results are therefore endpoint-qualified sensitivity evidence rather than external validation."),
        ])
        replace_section(doc, "3.6 ", "3.7 ", [
            ("Normal", "The matched-size multiview audit retained substantial endpoint heterogeneity: composition distributions included positive, weak and negative effects, and larger K ladders did not remove those differences. The 220 K = 3 subsets overlap in candidate membership and are distributional contrasts within one registry, not independent experiments."),
            ("Normal", "Classification ROC-AUC gains and regression RMSE reductions remain on separate axes. Complete subset identities, IQRs, ranges and ladder results are provided in Tables S9–S10 (Figure 5)."),
        ])
        replace_section(doc, "3.9 ", "3.10 ", [
            ("Normal", "Across ClinTox, BACE and ESOL, the prespecified CAHit@3 change and cross-fitted-gap directions transferred from seeded scaffold partitions to Tanimoto-component splits, but effect magnitudes and several individual contrasts changed. Effective-rank ordering was partly retained while absolute values remained split dependent."),
            ("Normal", "This three-endpoint transport analysis does not establish external validation, time-split robustness or a universal model-selection law. Complete intervals, splitter feasibility checks, missing cells and task-specific effects are retained in Tables S24–S26 and Figure S19."),
        ])
        replace_section(doc, "3.10 ", "4 Discussion", [
            ("Normal", "Across the six prespecified endpoints and three registries, CAHit@3 decreased from K = 4 to K = 32 in 14 of 18 endpoint–pool cells, whereas the cross-fitted gap increased in 10. At K = 32, both multiview and modern augmentation produced positive validation-selected gains relative to the homogeneous registry across all six endpoints; raw classification and regression utilities were not pooled."),
            ("Normal", "Modern components were heterogeneous. The full modern registry exceeded the corresponding classical multiview selected gain in 6 of 12 endpoint–K cells, while component-specific results differed and are reported in Table S33 and Figure S17."),
            ("Normal", "Under the locked classical-time threshold, equal-budget selected gain was at least the equal-K value in 17 of 36 endpoint–pool–K cells. Stability and selection-frequency results qualify this downstream-compute sensitivity and are reported in Table S36 and Figures S20–S21."),
            ("Normal", "Anchor and normalization choices changed scale more than the overall direction pattern, and 54 of 72 prespecified composition-by-K metric contrasts retained direction across the two evaluated split mechanisms. Missing normalization cells were retained; complete sensitivity results are in Tables S34–S35 and Figures S18–S19."),
        ])
        replace_section(doc, "4.8 ", "4.9 ", [
            ("Normal", "Registry composition changes the opportunity available to a selector, whereas finite validation determines how much of that opportunity is realised. Candidate expansion can therefore raise observed audit-best utility and selection pressure simultaneously. Matrix-dependent effective diversity complements nominal K by describing relative independence of candidate utilities across audit units, but it is descriptive and does not itself cause or determine selection loss."),
            ("Normal", "Modern candidates were not a homogeneous class: frozen language-model representations, the separately locked one-epoch D-MPNN and the full modern registry affected endpoints differently. This heterogeneity explains why modern augmentation can improve the opportunity frontier without guaranteeing that the inner selector captures every gain."),
            ("Normal", "Equal-budget conclusions depend on the prespecified prefix order because an expensive early candidate can consume the downstream budget before later complementary candidates become eligible. Anchors and normalizations affected numerical scale more than direction. These findings concern bounded downstream fit/predict exposure and do not imply end-to-end architecture efficiency within or beyond the evaluated endpoints, registries and split mechanisms."),
        ])
        replace_section(doc, "4.9 ", "5 Conclusions", [
            ("Normal", "The nine-endpoint primary audit and six-endpoint composition intervention are retrospective evaluations of public datasets, not prospective or independent external validation. Each endpoint supplies 15 outer audit units; shrinkage stabilizes covariance estimates but does not create independent information."),
            ("Normal", "Modern candidates were limited to frozen representation probes and a separately locked one-epoch D-MPNN. Equal-budget results depend on the evaluated hardware, measured downstream fit/predict time and registry order, and exclude acquisition and pretraining cost. Tanimoto-component transport covered only three endpoints and one similarity rule."),
            ("Normal", "Endpoint–pool–K and subset cells reuse endpoints and folds and are not independent. Near-zero normalization denominators and unavailable cells were retained as missing. Direction counts and descriptive associations were not promoted to population-level confirmatory tests."),
        ])
        replace_section(doc, "5 Conclusions", "Supplementary Information", [
            ("Normal", "Candidate expansion changed finite-audit opportunity, validation-realised gain, ranking fidelity, selection stability and bounded downstream cost. Matrix-dependent effective diversity supplements nominal K by describing the relative independence of candidate utilities across audit units, but it cannot by itself determine selection loss."),
            ("Normal", "Net value depends on whether added candidates provide complementary chemical information that finite inner validation can identify at acceptable downstream cost. This conclusion is limited to the evaluated endpoints, registries and split mechanisms and does not establish a universal candidate-expansion law."),
        ])
        replace_prefix(doc, "Additional file 1", "Additional file 1. Supplementary Methods and Results.")
        replace_prefix(doc, "Additional file 2", "Additional file 2. Machine-readable Supplementary Tables, Tables S1–S36.")
        replace_prefix(doc, "Additional file 3", "Additional file 3. Supplementary Figures, Figures S1–S21.")
        replace_prefix(doc, "Public dataset provenance", "Public dataset provenance is listed in Additional file 2, Table S1. Frozen registries, fold-level utilities, split manifests, source hashes, timing records, Figure 7 source data, equation-to-code mappings and analysis scripts are included in the submission package. No public GitHub or Zenodo archival status is claimed.")
        for heading in ["Competing interests", "Funding", "Authors' contributions", "Acknowledgements"]:
            h = find(doc, heading)
            paragraphs = doc.paragraphs
            index = next(i for i, p in enumerate(paragraphs) if p._p is h._p)
            paragraphs[index + 1].text = "Author confirmation required before submission; no information was inferred in this revision."
        update_table_2(doc, "en")
        notation_table(doc, "Selection and ranking estimands", "en")
    else:
        replace_prefix(doc, "本研究提出", "因此，本研究开展回顾性锁定、任务分层的审计，而非模型排行榜。证据包括九终点主审计和六终点扩展组成干预；三类已登记候选池在K = 4、8、16、32下评价，并完成组件消融、等下游预算敏感性、选择稳定性、锚点与归一化敏感性以及三终点切分机制迁移。各层级回答不同问题，所有结论均限定于已评估的端点、候选池和切分机制。")
        replace_section(doc, "2.1 ", "2.2 ", [
            ("Normal", "本研究是对已完成分子性质实验的回顾性审计。端点、候选顺序、K值、种子、切分逻辑、选择规则、结局与排除规则在本次分析中保持固定。由于审计规范在原始外层结果产生后重建，本研究并非前瞻性预注册；“冻结”仅表示审计期间资格与分析规则不变，不代表未被查看的前瞻性队列（图1）。"),
            ("Normal", "证据分为五层。第一层为九终点受控前缀审计；第二层为置换校准、组成控制和有限审计模拟；第三层为匹配规模的十二候选多表征审计；第四层为六终点扩展候选池组成干预及组件、预算、锚点、归一化、稳定性和切分机制敏感性分析；第五层为预测层可靠性与化学支持边界证据。"),
            ("Normal", "不同层级回答不同问题，不合并为单一排行榜。外层折仅审计冻结的内层选择决策；随后计算的外层最佳、排序与选择损失均为回顾性有限审计量，而非总体上界。多表征、组成及预测面板复用公开端点定义，不构成独立外部验证。"),
        ])
        replace_prefix(doc, "每个候选在不同K", "候选在不同K下沿用相同预处理、内层折和重拟合逻辑，但候选池扩张增加总搜索暴露。受控审计记录17,280次候选拟合和4,437.95候选拟合秒，校准与组成控制复用已存结果。扩展干预在六个端点、三类候选池、四个K、五个种子和三个外层折下包含1,080个主要外层单元，并记录64,616.35下游拟合/预测秒；总拟合数未重构。各分析暴露量不可直接互换，表2给出摘要，完整资源见表S2–S3及S28–S36。")
        replace_prefix(doc, "六终点四模型面板", "六终点四模型面板包括RDKit-RF、GCN、ChemBERTa冻结嵌入线性探针和MoLFormer冻结嵌入线性探针。各模型训练预算并不等同，因此该面板仅用于共享切分下的选择行为与误差结构。该预测层面板未包含D-MPNN；另行锁定的一轮D-MPNN候选仅用于扩展候选池组成干预。由于缺少完整共享切分审计，TabPFN未纳入。")
        replace_section(doc, "2.14 ", "2.15 ", [
            ("Heading 3", "端点与候选池设计"),
            ("Normal", "在读取扩展结果前按任务覆盖固定六个端点：分类为ClinTox、BACE、BBBP，回归为ESOL、Lipophilicity、Caco2_Wang。同质Morgan、多表征经典和现代增强三类精确前缀候选池均在K = 4、8、16、32下，采用种子11、23、37、53、71以及一致的3外层×3内层折评价。共享Morgan线性候选固定在首位，候选配置与顺序不随前缀改变。"),
            ("Heading 3", "组件消融"),
            ("Normal", "固定K消融在K = 16和32下比较多表征经典、经典+ChemBERTa、经典+MoLFormer、经典+D-MPNN及完整现代增强候选池。每次加入现代组件均替换一个预设经典候选，保持K严格相同。另行锁定的一轮D-MPNN仅属于本扩展干预。"),
            ("Heading 3", "等下游预算分析"),
            ("Normal", "各锁定前缀按端点与K使用多表征经典候选池的中位下游时间阈值截断，外层性能从不参与保留决策。记录时间仅包含内外层下游拟合与预测，不含模型获取、编码器预训练或缓存嵌入提取，因此该分析是受限下游计算敏感性而非端到端效率比较。"),
            ("Heading 3", "锚点、归一化、选择稳定性与切分机制敏感性"),
            ("Normal", "锚点包括共享Morgan线性候选、固定Morgan随机森林和预设候选池中位候选；报告原始增益、端点MAD归一化及配对同质审计最优归一化。候选、表征和学习器家族频率、标准化选择熵、众数比例、留一种子一致性及相邻折切换用于刻画稳定性。ClinTox、BACE和ESOL在种子化骨架划分与Tanimoto分量切分下完成三候选池闭环，候选池、种子、折、规则和指标均保持固定。"),
        ])
        replace_section(doc, "2.16 ", "3 ", [
            ("Normal", "端点特异的K = 32减K = 4效应先在五个切分种子内平均外层折，再对种子区组进行10,000次重采样。分类交叉拟合ROC-AUC损失和回归交叉拟合RMSE损失为预设任务分层共同主要层；原始效应不合并，不计算跨任务平均，也不报告端点汇总P值或确认性多重性家族。"),
            ("Normal", "“端点×候选池×K”单元并不独立。折效应在种子内平均，保留种子区组不确定性，方向计数仅作描述。组件与预算比较在端点、种子和折内配对。匹配规模结果是220个重叠已登记子集的分布摘要，分类与回归始终保留在不同数值轴。"),
        ])
        replace_section(doc, "3.1 ", "3.2 ", [
            ("Normal", "K = 32时，九个端点的Ledoit–Wolf谱熵秩中位数分别为：原始效用2.98、逐行中心化24.23、固定参照相对效用5.86、单元内秩27.14。四种矩阵分别保留共同审计难度、去除水平位移、依赖预设参照或仅保留顺序；任何一种都不被视为唯一真实候选数（图2；表S6–S7）。"),
            ("Normal", "估计器、种子、折和参照敏感性改变数值大小，但均保留名义数量与有效秩之间的差距。完整IQR、范围、参与率秩与相关性汇总移至表S6。"),
        ])
        replace_section(doc, "3.2 ", "3.3 ", [
            ("Normal", "随K扩张，排序保真度总体下降，但K = 16至32并非单调。终点CAHit@3中位数由K = 4时的1.000变为K = 32时的0.264，标准化MRR增益由0.861变为0.157；NDCG和秩相关提供一致但不完全相同的重排信息。"),
            ("Normal", "置换负对照位于校正零点附近。分级正对照中，注入验证—审计信号增强时，CAHit@3和标准化MRR提高，固定范围选择损失降低，说明指标具有预期的表面效度，但不能据此证明真实生物化学机制（图3A–B）。"),
        ])
        replace_section(doc, "3.3 ", "3.4 ", [
            ("Normal", "分类任务中，K = 4至32的平均观测审计最优增益由0.0028增至0.0311 ROC-AUC，选择模型增益由−0.0014增至0.0173。回归任务中，观测审计最优增益由0.7619增至0.8078 RMSE单位，选择模型增益由0.7442变为0.7378；两类任务尺度不合并。"),
            ("Normal", "在两类任务中，观测机会的增加均快于实际兑现增益。审计最优增益接近零时兑现比例不稳定，因此仅作为机器可读次要结果（图4A–B；表S8）。"),
        ])
        replace_section(doc, "3.4 ", "3.5 ", [
            ("Normal", "交叉拟合K = 32减K = 4效应在六个端点为正、三个端点为负；五个切分种子区间排除零，FreeSolv为负向例外。表3保留全部端点估计、区间和不确定方向，不合并分类与回归尺度（图3C；图4C；表S8）。"),
            ("Normal", "留一种子参照可减弱同一单元最大值的循环性，但仍依赖相同公开端点和切分生成器。因此，正负结果属于端点限定的敏感性证据，而非外部验证。"),
        ])
        replace_section(doc, "3.6 ", "3.7 ", [
            ("Normal", "匹配规模多表征审计保留明显终点异质性：组成分布同时出现正、弱和负效应，更大K阶梯亦未消除差异。220个K = 3子集共享候选成员，是同一候选池中的分布对比而非独立实验。"),
            ("Normal", "分类ROC-AUC增益与回归RMSE降低保持在不同轴上。完整子集身份、IQR、范围和阶梯结果见表S9–S10（图5）。"),
        ])
        replace_section(doc, "3.9 ", "3.10 ", [
            ("Normal", "在ClinTox、BACE和ESOL中，预设CAHit@3变化与交叉拟合差距方向从种子化骨架划分迁移至Tanimoto分量切分，但效应幅度及若干单独对比发生改变；有效秩顺序部分保留，绝对值仍依赖切分机制。"),
            ("Normal", "该三终点迁移分析不构成外部验证、时间切分鲁棒性或普适模型选择定律。完整区间、切分可行性、缺失单元和任务特异效应见表S24–S26及图S19。"),
        ])
        replace_section(doc, "3.10 ", "4 ", [
            ("Normal", "在六个预设端点和三类候选池中，CAHit@3从K = 4到32在14/18个“端点×候选池”单元下降，交叉拟合差距在10个单元增加。K = 32时，多表征与现代增强相对同质候选池的验证选择增益在六个端点均为正；分类和回归原始效用不合并。"),
            ("Normal", "现代组件具有异质性。完整现代候选池在6/12个“端点×K”单元超过对应多表征经典选择增益，其他组件特异结果见表S33和图S17。"),
            ("Normal", "在锁定的经典时间阈值下，等预算选择增益在17/36个“端点×候选池×K”单元不低于等K结果。稳定性与选择频率结果用于限定该下游计算敏感性，见表S36及图S20–S21。"),
            ("Normal", "锚点与归一化对数值尺度的影响大于总体方向；两种切分机制下，54/72个预设“组成×K”指标对比方向一致。缺失归一化单元全部保留，完整敏感性结果见表S34–S35及图S18–S19。"),
        ])
        replace_section(doc, "4.8 ", "4.9 ", [
            ("Normal", "候选池组成改变选择器可利用的机会，有限验证决定机会能否兑现。候选扩张因而可能同时提高观测审计最优效用并增加选择压力。矩阵依赖的有效多样性补充了名义K，可描述候选效用在审计单元中的相对独立变化，但其本身不构成因果惩罚，也不能决定选择损失。"),
            ("Normal", "现代候选并非同质类别：冻结语言模型表征、另行锁定的一轮D-MPNN及完整现代候选池对各端点的影响不同。因此，现代增强可提高机会前沿，却不能保证内层选择器兑现全部增益。"),
            ("Normal", "等预算结果依赖预设前缀顺序，因为较昂贵的早期候选可能在后续互补候选进入前耗尽预算。锚点与归一化对效应尺度的影响大于方向。上述结论仅涉及受限下游拟合/预测暴露，不代表端到端架构效率，并限定于已评估的端点、候选池和切分机制。"),
        ])
        replace_section(doc, "4.9 ", "5 ", [
            ("Normal", "九终点主审计和六终点组成干预均为公开数据上的回顾性评价，不是前瞻性或独立外部验证。每个端点仅提供15个外层审计单元；收缩可稳定协方差估计，但不能创造独立信息。"),
            ("Normal", "现代候选仅包括冻结表征探针和另行锁定的一轮D-MPNN。等预算结果依赖当前硬件、下游拟合/预测时间与候选顺序，且不含获取和预训练成本；Tanimoto分量迁移仅覆盖三个端点和一种相似度规则。"),
            ("Normal", "“端点×候选池×K”及子集单元复用端点与折，并不独立。近零归一化分母和不可用单元均保留为缺失；方向计数与描述性关联未提升为总体确认性检验。"),
        ])
        replace_section(doc, "5 ", "补充信息", [
            ("Normal", "候选扩张改变有限审计机会、验证兑现增益、排序保真度、选择稳定性与受限下游成本。矩阵依赖的有效多样性补充了名义K，可描述候选效用在审计单元中的相对独立变化，但其本身不能决定选择损失。"),
            ("Normal", "净价值取决于新增候选能否提供有限内层验证可识别的互补化学信息，并保持可接受的下游成本。该结论仅限于已评估的端点、候选池和切分机制，不构成普适候选扩张定律。"),
        ])
        replace_prefix(doc, "Additional file 1", "Additional file 1：补充方法与结果。Additional file 2：机器可读补充表S1–S36。Additional file 3：补充图S1–S21。")
        update_table_2(doc, "zh")
        localize_zh_tables(doc)
        notation_table(doc, "选择与排序估计量", "zh")
        # Use consistent Chinese terminology in the remaining unchanged text.
        for paragraph in doc.paragraphs:
            original = paragraph.text
            revised = original.replace("seeded scaffold partitions", "种子化骨架划分（seeded scaffold partition）")
            revised = revised.replace("source hashes", "源文件哈希")
            revised = revised.replace("scaffold split", "骨架切分")
            revised = revised.replace("Tanimoto 连通分量 split", "Tanimoto连通分量切分")
            # Reassigning paragraph.text removes drawings and OMML; change only
            # genuinely affected plain-text paragraphs.
            has_drawing = bool(paragraph._p.xpath(".//w:drawing"))
            has_math = bool(paragraph._p.xpath(".//m:oMath | .//m:oMathPara"))
            if revised != original and not has_drawing and not has_math:
                paragraph.text = revised

        ref = find(doc, "参考文献")
        declaration_blocks = [
            ("Heading 1", "声明"),
            ("Heading 2", "伦理批准与参与同意"),
            ("Normal", "不适用。本研究使用公开分子数据集，不涉及人类参与者、人类数据或动物实验。"),
            ("Heading 2", "发表同意"), ("Normal", "不适用。"),
            ("Heading 2", "数据与材料可获得性"),
            ("Normal", "公开数据来源见Additional file 2的表S1。提交包包含冻结候选池、逐折效用、切分清单、源文件哈希、时间记录、Figure 7源数据、公式—代码映射及分析脚本；本文不声称GitHub或Zenodo永久归档已经完成。"),
            ("Heading 2", "竞争性利益"), ("Normal", "投稿前需作者确认；本次修订未推断相关信息。"),
            ("Heading 2", "基金资助"), ("Normal", "投稿前需作者确认；本次修订未推断相关信息。"),
            ("Heading 2", "作者贡献"), ("Normal", "投稿前需依据真实作者首字母与CRediT角色补充；本次修订未推断作者信息。"),
            ("Heading 2", "致谢"), ("Normal", "投稿前需作者确认；本次修订未推断相关信息。"),
        ]
        for style, text in declaration_blocks:
            ref.insert_paragraph_before(text, style=style)

    doc.save(destination)


def main() -> None:
    build("en", EN_BASE, EN_OUT)
    build("zh", ZH_BASE, ZH_OUT)
    audit = {
        "status": "clean_manuscripts_built",
        "english": str(EN_OUT),
        "chinese": str(ZH_OUT),
        "expanded_primary_outer_units": 1080,
        "expanded_recorded_downstream_seconds": 64616.35152543575,
        "expanded_fit_count": "not reconstructed",
        "author_declaration_fields": "confirmation required; no information inferred",
    }
    (OUT / "Paper33_manuscript_build_audit.json").write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(audit, ensure_ascii=False))


if __name__ == "__main__":
    main()
