from __future__ import annotations

import csv
import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import fitz
from PIL import Image
from docx import Document
from openpyxl import load_workbook


ROOT = Path(r"D:\fzyc\output\paper28_pre_submission_minor_revision_20260717")
FIGDIR = ROOT / "main_figures"
SUPP = ROOT / "supplementary"
EN = ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript.docx"
TRACKED = ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript_tracked_changes.docx"
ZH = ROOT / "候选池扩张与模型选择损失_中文完整论文.docx"
REVIEWER = ROOT / "Reviewer_concern_Response_Location.docx"
SUPP_DOC = SUPP / "Additional_file_1_Supplementary_Methods_and_Results.docx"
BOOK = SUPP / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S31.xlsx"

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def doc_text(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def xml_package_ok(path: Path) -> tuple[bool, int, list[str]]:
    errors: list[str] = []
    xml_count = 0
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if name.lower().endswith((".xml", ".rels")):
                xml_count += 1
                try:
                    ET.fromstring(zf.read(name))
                except Exception as exc:  # pragma: no cover - diagnostic branch
                    errors.append(f"{name}: {exc}")
    return not errors, xml_count, errors


def check_figures() -> dict:
    rows = []
    all_ok = True
    for n in range(1, 8):
        png = FIGDIR / f"Figure{n}_600dpi.png"
        svg = FIGDIR / f"Figure{n}.svg"
        pdf = FIGDIR / f"Figure{n}.pdf"
        exists = all(p.exists() and p.stat().st_size > 0 for p in (png, svg, pdf))
        mode = None
        width = height = 0
        svg_text = ""
        svg_min_size = None
        svg_bad_fonts: list[str] = []
        pdf_fonts: set[str] = set()
        pdf_bad_fonts: list[str] = []
        if exists:
            with Image.open(png) as im:
                mode = im.mode
                width, height = im.size
            svg_text = svg.read_text(encoding="utf-8")
            sizes = [float(x) for x in re.findall(r"font-size:\s*([0-9.]+)px", svg_text)]
            if not sizes:
                sizes = [float(x) for x in re.findall(r"font-size=['\"]([0-9.]+)", svg_text)]
            svg_min_size = min(sizes) if sizes else None
            families = set(re.findall(r"font-family:\s*([^;\"]+)", svg_text))
            svg_bad_fonts = sorted(
                f for f in families
                if "Times New Roman" not in f and "TimesNewRoman" not in f
            )
            with fitz.open(pdf) as document:
                for page in document:
                    for font in page.get_fonts(full=True):
                        pdf_fonts.add(str(font[3]))
            pdf_bad_fonts = sorted(
                f for f in pdf_fonts
                if "TimesNewRoman" not in f.replace(" ", "")
            )
        ok = (
            exists
            and mode == "RGB"
            and width >= 3000
            and height >= 2000
            and width <= 4700
            and "<text" in svg_text
            and svg_min_size is not None
            and svg_min_size >= 7.5
            and not svg_bad_fonts
            and bool(pdf_fonts)
            and not pdf_bad_fonts
        )
        all_ok &= ok
        rows.append({
            "figure": n,
            "ok": ok,
            "pixels": f"{width}x{height}",
            "mode": mode,
            "svg_min_font_pt": svg_min_size,
            "svg_bad_fonts": svg_bad_fonts,
            "pdf_fonts": sorted(pdf_fonts),
            "pdf_bad_fonts": pdf_bad_fonts,
        })
    return {"ok": all_ok, "rows": rows}


def check_documents() -> dict:
    paths = [EN, TRACKED, ZH, REVIEWER, SUPP_DOC]
    package_rows = []
    all_packages_ok = True
    for path in paths + [BOOK]:
        ok, count, errors = xml_package_ok(path)
        all_packages_ok &= ok
        package_rows.append({"file": path.name, "ok": ok, "xml_parts": count, "errors": errors})

    en_text = doc_text(EN)
    zh_text = doc_text(ZH)
    reviewer_text = doc_text(REVIEWER)
    supp_text = doc_text(SUPP_DOC)
    clean_oracle_counts = {
        "english": len(re.findall(r"\boracle\b", en_text, flags=re.I)),
        "chinese": len(re.findall(r"\boracle\b", zh_text, flags=re.I)),
        "reviewer": len(re.findall(r"\boracle\b", reviewer_text, flags=re.I)),
        "supplement": len(re.findall(r"\boracle\b", supp_text, flags=re.I)),
    }
    phrases = {
        "observed audit-best": "observed audit-best" in en_text.lower(),
        "finite audit maximum": "finite audit maximum" in en_text.lower(),
        "not 18 independent experiments": "not 18 independent experiments" in en_text.lower(),
        "downstream cost scope": "downstream nested fitting and prediction" in en_text.lower(),
        "useful complementarity qualification": "useful complementarity" in en_text.lower(),
        "not modern leaderboard": "does not constitute a modern-model leaderboard" in en_text.lower(),
        "Chinese observed audit-best definition": "有限外层审计最佳值" in zh_text or "观测审计最佳" in zh_text,
    }

    layout_rows = []
    layout_ok = True
    for path in (EN, ZH):
        doc = Document(path)
        section = doc.sections[0]
        printable = section.page_width - section.left_margin - section.right_margin
        widths = [shape.width for shape in doc.inline_shapes]
        ok = len(widths) == 7 and all(w <= printable for w in widths)
        layout_ok &= ok
        layout_rows.append({
            "file": path.name,
            "inline_shapes": len(widths),
            "printable_width_emu": printable,
            "maximum_shape_width_emu": max(widths) if widths else 0,
            "ok": ok,
        })

    with zipfile.ZipFile(TRACKED) as zf:
        document_root = ET.fromstring(zf.read("word/document.xml"))
        settings_root = ET.fromstring(zf.read("word/settings.xml"))
    insertions = len(document_root.findall(f".//{{{NS_W}}}ins"))
    deletions = len(document_root.findall(f".//{{{NS_W}}}del"))
    track_enabled = settings_root.find(f".//{{{NS_W}}}trackRevisions") is not None
    tracked_ok = track_enabled and insertions > 0 and deletions > 0

    return {
        "ok": (
            all_packages_ok
            and all(v == 0 for v in clean_oracle_counts.values())
            and all(phrases.values())
            and layout_ok
            and tracked_ok
        ),
        "ooxml_packages": package_rows,
        "clean_oracle_counts": clean_oracle_counts,
        "required_phrases": phrases,
        "layout": layout_rows,
        "tracked_changes": {
            "ok": tracked_ok,
            "track_revisions_enabled": track_enabled,
            "insertions": insertions,
            "deletions": deletions,
        },
    }


def check_workbook() -> dict:
    wb = load_workbook(BOOK, read_only=False, data_only=False)
    sheets = wb.sheetnames
    oracle_hits: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                value = cell.value
                if isinstance(value, str) and re.search(r"\boracle\b", value, flags=re.I):
                    oracle_hits.append(f"{ws.title}!{cell.coordinate}")
                if cell.comment and re.search(r"\boracle\b", cell.comment.text, flags=re.I):
                    oracle_hits.append(f"{ws.title}!{cell.coordinate}:comment")
    wb.close()
    return {
        "ok": len(sheets) == 31 and not oracle_hits,
        "sheet_count": len(sheets),
        "first_sheet": sheets[0] if sheets else None,
        "last_sheet": sheets[-1] if sheets else None,
        "oracle_hits": oracle_hits,
    }


def check_qc_and_definitions() -> dict:
    qc = json.loads((ROOT / "Final_minor_revision_QC_audit.json").read_text(encoding="utf-8"))
    definition = json.loads(
        (ROOT / "figure_source_data" / "Figure_7_definition_audit.json").read_text(encoding="utf-8")
    )
    with (ROOT / "figure_source_data" / "Figure_7_normalization_denominator_audit.csv").open(
        encoding="utf-8-sig", newline=""
    ) as handle:
        denom_rows = list(csv.DictReader(handle))
    denominators = []
    for row in denom_rows:
        for key, value in row.items():
            if key and value and "denominator" in key.lower():
                try:
                    denominators.append(float(value))
                except ValueError:
                    pass
    positive = bool(denominators) and min(denominators) > 0
    qc_ok = qc.get("status") == "complete" and all(qc.get("checks", {}).values())
    definition_ok = (
        definition.get("all_denominators_positive") is True
        and float(definition.get("minimum_denominator", 0)) > 0
        and "raw Ledoit-Wolf entropy effective rank" in definition.get("relative_entropy_rank_definition", "")
        and "downstream" in definition.get("cost_scope", "").lower()
    )
    return {
        "ok": qc_ok and definition_ok and positive,
        "qc_status": qc.get("status"),
        "qc_checks": qc.get("checks"),
        "definition": definition,
        "denominator_rows": len(denom_rows),
        "minimum_denominator_recomputed": min(denominators) if denominators else None,
        "all_recomputed_denominators_positive": positive,
    }


def main() -> None:
    checks = {
        "figures": check_figures(),
        "documents": check_documents(),
        "workbook": check_workbook(),
        "qc_and_definitions": check_qc_and_definitions(),
    }
    overall = all(item["ok"] for item in checks.values())
    audit = {
        "status": "complete" if overall else "failed",
        "root": str(ROOT),
        "checks": checks,
        "author_confirmation_still_required": [
            "Competing interests",
            "Funding",
            "Authors' contributions",
            "Acknowledgements",
        ],
    }
    (ROOT / "Final_package_verification_audit.json").write_text(
        json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    with (ROOT / "Final_package_verification_checklist.csv").open(
        "w", encoding="utf-8-sig", newline=""
    ) as handle:
        writer = csv.writer(handle)
        writer.writerow(["check", "status", "detail"])
        for name, result in checks.items():
            writer.writerow([name, "PASS" if result["ok"] else "FAIL", json.dumps(result, ensure_ascii=False)])
        writer.writerow([
            "author declarations",
            "AUTHOR CONFIRMATION REQUIRED",
            "Competing interests; Funding; Authors' contributions; Acknowledgements",
        ])
    print(json.dumps({
        "status": audit["status"],
        "figure_ok": checks["figures"]["ok"],
        "document_ok": checks["documents"]["ok"],
        "workbook_ok": checks["workbook"]["ok"],
        "qc_and_definitions_ok": checks["qc_and_definitions"]["ok"],
        "tracked_changes": checks["documents"]["tracked_changes"],
        "minimum_denominator": checks["qc_and_definitions"]["minimum_denominator_recomputed"],
    }, ensure_ascii=False, indent=2))
    if not overall:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
