from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path("D:/fzyc")
OUT = ROOT / "output"
HARD = OUT / "sci1_hardening_20260707"
SOURCE = max([p for p in OUT.glob("*.docx") if p.name.endswith("-16.docx")], key=lambda p: p.stat().st_mtime)
TARGET = OUT / "小论文-17.docx"
REPORT = OUT / "小论文-17_SCI1实验补强报告.md"
AUDIT_JSON = OUT / "paper17_sci1_hardening_audit.json"


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para


def replace_para(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def find_para(doc: Document, startswith: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {startswith}")


def load_values() -> dict[str, object]:
    audit = json.loads((HARD / "sci1_hardening_audit.json").read_text(encoding="utf-8"))
    overlap = pd.read_csv(HARD / "six_task_error_overlap_pairwise_summary.csv")
    dup = pd.read_csv(HARD / "six_task_duplicate_sensitivity_summary.csv")
    val = pd.read_csv(HARD / "validation_information_sensitivity_summary.csv")
    comp = pd.read_csv(HARD / "component_policy_ablation_summary.csv")
    full_val = val[val["variant"].eq("full_multiview")].sort_values("validation_information_fraction")
    comp_full = comp[comp["variant"].eq("full_multiview")]
    return {
        "audit": audit,
        "overlap_mean": float(overlap["mean_jaccard_error_overlap"].mean()),
        "overlap_min": float(overlap["mean_jaccard_error_overlap"].min()),
        "overlap_max": float(overlap["mean_jaccard_error_overlap"].max()),
        "duplicate_max": float(dup["abs_delta_vs_global_dedup"].max()),
        "duplicate_max_task": str(dup.sort_values("abs_delta_vs_global_dedup", ascending=False).iloc[0]["task"]),
        "full_loss_025": float(full_val[full_val["validation_information_fraction"].eq(0.25)]["mean_range_normalized_selection_loss"].iloc[0]),
        "full_loss_100": float(full_val[full_val["validation_information_fraction"].eq(1.00)]["mean_range_normalized_selection_loss"].iloc[0]),
        "full_hit_025": float(full_val[full_val["validation_information_fraction"].eq(0.25)]["top1_hit_rate"].iloc[0]),
        "full_hit_100": float(full_val[full_val["validation_information_fraction"].eq(1.00)]["top1_hit_rate"].iloc[0]),
        "full_vb_regret": float(comp_full[comp_full["policy"].eq("validation_best")]["mean_normalized_regret"].iloc[0]),
        "full_fixed_regret": float(comp_full[comp_full["policy"].eq("fixed_morgan_rf")]["mean_normalized_regret"].iloc[0]),
        "full_risk_regret": float(comp_full[comp_full["policy"].eq("risk_adjusted")]["mean_normalized_regret"].iloc[0]),
        "full_one_se_regret": float(comp_full[comp_full["policy"].eq("one_se_stable")]["mean_normalized_regret"].iloc[0]),
    }


def revise_docx(values: dict[str, object]) -> None:
    shutil.copy2(SOURCE, TARGET)
    doc = Document(str(TARGET))
    audit = values["audit"]

    replace_para(
        find_para(doc, "除主效应外，本文报告真嵌套验证"),
        (
            "除主效应外，本文报告真嵌套验证、种子敏感性、统一消融、80/90/95 保形覆盖率、"
            "精确 Tanimoto 分箱、MoleculeACE 活性悬崖、低相似度失败样本和扩展失败案例等补充分析。"
            "强基线证据矩阵被扩展为两层：RDKit-RF、GNN-GCN、ChemBERTa 和 MoLFormer 冻结适配头已在 "
            "ESOL、FreeSolv、Lipophilicity、BBBP、BACE 和 ClinTox 六个 MoleculeNet 主任务上完成同划分 "
            "3 外层×3 内层×5 种子评估，形成 360 个外层单元、1,080 个内层单元和 220,040 条逐样本预测；"
            "Chemprop/D-MPNN 仍限定为 ESOL、BACE 和 ClinTox 三终点补强面板。TabPFN 已安装，但因授权和运行时"
            "交互限制未能完成同划分预测导出；因此仅作为授权受限候选记录在状态表中，不作为完成性结果。"
        ),
    )

    replace_para(
        find_para(doc, "为使文献对照可复核，本研究整理了六类证据边界"),
        (
            "为使文献对照可复核，本研究整理了六类证据边界，覆盖标准基准、现代强基线、模型选择偏倚、"
            "预测可靠性、化学边界和可复现性。新增 SCI1 补强实验将现代强基线从三终点代表面板扩展到六个 "
            "MoleculeNet 主任务：RDKit-RF、GNN-GCN、ChemBERTa 和 MoLFormer 在同一冻结划分下共产生 "
            f"{audit['six_task_outer_rows']} 个外层单元、{audit['six_task_inner_rows']} 个内层单元和 "
            f"{audit['six_task_prediction_rows']:,} 条逐样本预测。冻结选择器在 "
            f"{audit['six_task_selection_units']} 个选择单元中的 Top-1 命中率为 {audit['six_task_top1_hit_rate']:.3f}，"
            f"平均范围归一化选择损失为 {audit['six_task_mean_range_normalized_selection_loss']:.4f}。"
        ),
    )

    replace_para(
        find_para(doc, "error-overlap 审计显示，五个候选的错误集并非完全重合"),
        (
            "error-overlap 审计已扩展到六任务强基线面板。四个完成候选形成 6 个候选对、每对 90 个任务-种子-外层"
            f"单元，平均 Jaccard 错误重合为 {values['overlap_mean']:.3f}，候选对范围为 "
            f"{values['overlap_min']:.3f}–{values['overlap_max']:.3f}。去重敏感性也从三终点扩展到六个 MoleculeNet "
            f"主任务、三套策略和 {audit['duplicate_detail_rows']} 个外层单元；最大平均效用变化为 "
            f"{values['duplicate_max']:.3f}（{values['duplicate_max_task']}）。这些结果支持一个审慎的比较原则："
            "当代复杂模型只有在同一冻结划分下完成预测导出、选择审计、去重敏感性和负结果记录后，才宜进入主文效应叙事。"
        ),
    )

    anchor = find_para(doc, "这些补强实验使强基线、错误互补性和数据清洗敏感性")
    mechanism = (
        "新增验证信息量机制实验进一步支持主命题。在九终点 12 候选多视图池中，模拟验证信息量从 25% 增加到 100% 时，"
        f"full-multiview 的平均范围归一化选择损失由 {values['full_loss_025']:.3f} 降至 "
        f"{values['full_loss_100']:.3f}，Top-1 命中率由 {values['full_hit_025']:.3f} 升至 "
        f"{values['full_hit_100']:.3f}。组件消融显示，full-multiview validation-best 的平均归一化遗憾为 "
        f"{values['full_vb_regret']:.3f}，低于固定 Morgan-RF 的 {values['full_fixed_regret']:.3f}，"
        f"并接近 risk-adjusted 的 {values['full_risk_regret']:.3f} 和 one-SE stable 的 {values['full_one_se_regret']:.3f}。"
    )
    if not any(p.text.strip() == mechanism for p in doc.paragraphs):
        insert_after(anchor, mechanism)

    replace_para(
        find_para(doc, "FZYC-Mol 不替代预测模型，也不保证性能提升"),
        (
            "FZYC-Mol 不替代预测模型，也不保证性能提升，且不提供可迁移到所有终点的元选择器。"
            "六任务 MoleculeNet 面板已完成 RDKit-RF、GNN-GCN、ChemBERTa 和 MoLFormer 的同划分 3×3×5 评估，"
            "并新增选择审计、配对效应、error-overlap、去重敏感性、验证信息量机制实验和组件消融；"
            "但 Chemprop/D-MPNN 仍只完成三终点补强面板，TabPFN 未能产生同划分预测，九终点全量深度/基础模型面板"
            "仍超出当前完成范围。因此，强基线结果应解释为更厚的边界压力测试，不能被解读为复杂模型已在所有终点上"
            "完成确认性统一重训。TDC 三种子结果只能反映种子变异和公开面板异质性，不能承担严格抽样推断。"
            "公开 release、Zenodo DOI 和第三方冷启动复跑仍需作为后续复现工作完成。"
        ),
    )

    replace_para(
        find_para(doc, "后续扩展应优先完成三项工作"),
        (
            "后续扩展应优先完成三项工作：将 Chemprop/D-MPNN、TabPFN 和更多深度/基础模型从当前三终点或六任务层级"
            "扩展到九终点全量同划分面板；加入真正时间外或外部 ADMET 盲测以替代公开面板边界证据；完成公开 release、"
            "Zenodo DOI 和独立冷启动复跑。这些边界已在正文中明确限定，以避免将代表性扩展分析写成已完成的全量确证。"
        ),
    )

    doc.save(str(TARGET))


def write_report(values: dict[str, object], checks: dict[str, object]) -> None:
    audit = values["audit"]
    lines = [
        "# 小论文-17 SCI1 实验补强报告",
        "",
        f"- 生成时间：{datetime.now().isoformat(timespec='seconds')}",
        f"- 输入稿件：`{SOURCE}`",
        f"- 输出稿件：`{TARGET}`",
        f"- 新实验目录：`{HARD}`",
        "",
        "## 已完成的新实验",
        "",
        "| 实验 | 新增规模 | 关键结果 |",
        "|---|---:|---|",
        f"| 六任务现代强基线 | {audit['six_task_outer_rows']} 外层、{audit['six_task_inner_rows']} 内层、{audit['six_task_prediction_rows']:,} 条预测 | Top-1 命中率 {audit['six_task_top1_hit_rate']:.3f}；平均范围归一化选择损失 {audit['six_task_mean_range_normalized_selection_loss']:.4f} |",
        f"| error-overlap | 6 个候选对，每对 90 个单元 | 平均 Jaccard {values['overlap_mean']:.3f}，范围 {values['overlap_min']:.3f}–{values['overlap_max']:.3f} |",
        f"| 去重敏感性 | {audit['duplicate_detail_rows']} 个外层单元 | 最大平均效用变化 {values['duplicate_max']:.3f}（{values['duplicate_max_task']}） |",
        f"| 验证信息量机制实验 | {audit['validation_sensitivity_rows']} 个模拟单元 | full-multiview loss {values['full_loss_025']:.3f}→{values['full_loss_100']:.3f}；Top-1 {values['full_hit_025']:.3f}→{values['full_hit_100']:.3f} |",
        f"| 组件/规则消融 | {audit['component_policy_rows']} 条策略汇总 | validation-best regret {values['full_vb_regret']:.3f}；fixed Morgan-RF {values['full_fixed_regret']:.3f} |",
        f"| OOD/time 边界汇总 | {audit['ood_boundary_rows']} 条边界证据 | TDC、bRo5 time split、MoleculeACE 被统一为边界证据 |",
        "",
        "## 仍保留的 SCI1 风险",
        "",
        "- Chemprop/D-MPNN 尚未从三终点扩展到六任务或九终点全量同划分面板。",
        "- TabPFN 未能完成同划分预测导出。",
        "- TDC/bRo5/MoleculeACE 仍是公开边界证据，尚不能替代真正时间外或外部盲测。",
        "- 公开 release、Zenodo DOI 和第三方冷启动复跑仍需在投稿前补齐。",
        "",
        "## 自动审计",
        "",
        "```json",
        json.dumps(checks, ensure_ascii=False, indent=2),
        "```",
    ]
    REPORT.write_text("\n".join(lines), encoding="utf-8")


def audit_docx() -> dict[str, object]:
    with ZipFile(TARGET) as zf:
        bad = zf.testzip()
    doc = Document(str(TARGET))
    text = "\n".join(p.text for p in doc.paragraphs)
    checks = {
        "zip_ok": bad is None,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "figures": len(doc.inline_shapes),
        "references": sum(1 for p in doc.paragraphs if p.text.strip().startswith("[")),
        "six_task_strong_written": "220,040 条逐样本预测" in text and "360 个外层单元" in text,
        "six_task_overlap_written": "6 个候选对、每对 90 个" in text,
        "duplicate_written": "270 个外层单元" in text,
        "validation_sensitivity_written": "验证信息量从 25% 增加到 100%" in text,
        "chemprop_boundary_written": "Chemprop/D-MPNN 仍只完成三终点补强面板" in text,
        "tabpfn_boundary_written": "TabPFN 未能产生同划分预测" in text,
    }
    checks["passed"] = all(v for k, v in checks.items() if isinstance(v, bool))
    return checks


def main() -> None:
    values = load_values()
    revise_docx(values)
    checks = audit_docx()
    AUDIT_JSON.write_text(
        json.dumps({"values": values, "checks": checks}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    write_report(values, checks)
    print(json.dumps({"target": str(TARGET), "report": str(REPORT), "passed": checks["passed"]}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
