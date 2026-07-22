from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path("D:/fzyc")
SOURCE = ROOT / "output" / "小论文-20_修订分配与审稿回复.md"
TARGET = ROOT / "output" / "小论文-20_修订分配与审稿回复.docx"


def add_table(doc: Document, block: list[str]) -> None:
    rows = [[x.strip() for x in line.strip().strip("|").split("|")] for line in block]
    rows = [rows[0]] + rows[2:]
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = value.replace("`", "")
            for paragraph in cells[j].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"; run.font.size = Pt(8); run.bold = i == 0


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.2); section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1]):
            block = [lines[i], lines[i + 1]]; i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            add_table(doc, block); continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(line, style="List Number")
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:]); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line:
            p = doc.add_paragraph(line.replace("`", ""))
        else:
            i += 1; continue
        for run in p.runs:
            run.font.name = "Arial"; run.font.size = Pt(10)
        i += 1
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
