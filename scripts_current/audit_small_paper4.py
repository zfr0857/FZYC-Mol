from __future__ import annotations

import hashlib
import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "小论文-4.docx"
TRACKED = ROOT / "output" / "小论文-4_修订痕迹.docx"
PACKAGE = ROOT / "output" / "小论文-4_图表包"
AUDIT = ROOT / "results" / "audits" / "small_paper_4_audit.json"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def xml_errors(path: Path) -> list[str]:
    errors = []
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            errors.append(f"bad_member:{bad}")
        for name in archive.namelist():
            if name.endswith(".xml"):
                try:
                    ElementTree.fromstring(archive.read(name))
                except Exception as exc:
                    errors.append(f"{name}:{type(exc).__name__}")
    return errors


def audit() -> Path:
    doc = Document(DOCX)
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    expected = [
        "K=32 相对 K=4 的配对固定分母遗憾增加 0.122",
        "端点×池规模内置换 P=0.953",
        "MAE 为 0.112，低于常数基线 0.123",
        "12 候选共享划分重训",
        "相对 Morgan-only 选择的配对效用增益为 0.343",
        "6,480 次拟合",
        "图 1–11 均配套 PNG/SVG",
        "表 8 | 审稿核心补实验与共享划分多视图结果",
    ]
    stale = [
        "低风险一半单元的平均遗憾为 0.098，高风险一半为 0.148",
        "确认性候选池仅含轻量 Morgan-512 模型",
        "图 1–10 均配套 PNG/SVG",
        "主文数字、表和图 1–10",
    ]
    captions = [p.text.strip() for p in doc.paragraphs if re.match(r"^图\s*\d+\s*\|", p.text.strip())]
    table_captions = [p.text.strip() for p in doc.paragraphs if re.match(r"^表\s*\d+\s*\|", p.text.strip())]
    svg = sorted((PACKAGE / "figures").glob("fig*.svg"))
    png = sorted((PACKAGE / "figures").glob("fig*.png"))
    svg_stems = {p.stem for p in svg}
    png_stems = {p.stem for p in png}
    figure_checks = {}
    for path in png:
        with Image.open(path) as image:
            figure_checks[path.stem] = {
                "width_px": image.width,
                "height_px": image.height,
                "bytes": path.stat().st_size,
                "svg_editable_text": "<text" in (PACKAGE / "figures" / f"{path.stem}.svg").read_text(encoding="utf-8"),
            }
    shape_cm = [{"width": shape.width / 360000, "height": shape.height / 360000} for shape in doc.inline_shapes]
    with zipfile.ZipFile(TRACKED) as archive:
        tracked_xml = archive.read("word/document.xml")
        settings = archive.read("word/settings.xml")
    tracked = {
        "insertions": len(re.findall(br"<w:ins(?:\s|>)", tracked_xml)),
        "deletions": len(re.findall(br"<w:del(?:\s|>)", tracked_xml)),
        "tracking_enabled": b"trackRevisions" in settings,
    }
    report = {
        "docx": {"path": str(DOCX), "sha256": sha256(DOCX), "bytes": DOCX.stat().st_size},
        "tracked_docx": {"path": str(TRACKED), "sha256": sha256(TRACKED), "bytes": TRACKED.stat().st_size},
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "equation_paragraphs": sum(p.style and p.style.name == "Equation" for p in doc.paragraphs),
        "caption_count": len(captions),
        "caption_numbers": [int(re.search(r"\d+", caption).group()) for caption in captions],
        "table_caption_numbers": [int(re.search(r"\d+", caption).group()) for caption in table_captions],
        "expected_text": {value: value in text for value in expected},
        "stale_text": {value: value in text for value in stale},
        "clean_xml_errors": xml_errors(DOCX),
        "tracked_xml_errors": xml_errors(TRACKED),
        "tracked_changes": tracked,
        "figure_counts": {"svg": len(svg), "png": len(png), "matching_stems": svg_stems == png_stems},
        "figure_checks": figure_checks,
        "inline_shape_cm": shape_cm,
        "source_data_files": len(list((PACKAGE / "source_data").glob("*.csv"))),
    }
    report["passed"] = all(
        [
            report["tables"] == 10,
            report["inline_shapes"] == 11,
            report["equation_paragraphs"] == 17,
            report["caption_count"] == 11,
            report["caption_numbers"] == list(range(1, 12)),
            report["table_caption_numbers"] == list(range(1, 11)),
            all(report["expected_text"].values()),
            not any(report["stale_text"].values()),
            not report["clean_xml_errors"],
            not report["tracked_xml_errors"],
            report["tracked_changes"]["insertions"] > 0,
            report["tracked_changes"]["deletions"] > 0,
            report["tracked_changes"]["tracking_enabled"],
            report["figure_counts"]["svg"] == 11,
            report["figure_counts"]["png"] == 11,
            report["figure_counts"]["matching_stems"],
            all(item["svg_editable_text"] for item in figure_checks.values()),
            all(item["width_px"] >= 2000 and item["bytes"] >= 50000 for item in figure_checks.values()),
            all(item["width"] <= 16.21 and item["height"] <= 13.0 for item in shape_cm),
            report["source_data_files"] >= 40,
        ]
    )
    AUDIT.parent.mkdir(parents=True, exist_ok=True)
    AUDIT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    if not report["passed"]:
        raise SystemExit(json.dumps(report, ensure_ascii=False, indent=2))
    return AUDIT


if __name__ == "__main__":
    print(audit())
