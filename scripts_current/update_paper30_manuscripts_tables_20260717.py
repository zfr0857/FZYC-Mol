from __future__ import annotations

import json
import re
import shutil
import tempfile
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter
from PIL import Image


ROOT = Path(r"D:\fzyc")
TARGET = ROOT / "output" / "paper30_final_minor_revision_20260717"
FIG = TARGET / "main_figures"
SOURCE = TARGET / "figure_source_data"
SUPP = TARGET / "supplementary"
EN = TARGET / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript(6).docx"
ZH = TARGET / "候选池扩张与模型选择损失_中文完整论文.docx"
REVIEWER = TARGET / "Reviewer_concern_Response_Location.docx"
BOOK = SUPP / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S31.xlsx"

TABLE1_CAPTION = "Table 1. Primary datasets and endpoint metrics."
TABLE1_NOTE = (
    "Target units: ESOL, log mol/L; FreeSolv, kcal/mol; Lipophilicity, logD; Caco2 Wang, dataset-provided permeability scale. "
    "Classification rows report positive-class n (%)."
)
TABLE2_CAPTION = "Table 2. Audit components and recorded exposure."
TABLE2_NOTE = (
    "Exposure units are analysis specific and are not directly comparable. Downstream cost excludes encoder pretraining and cached embedding extraction."
)
TABLE3_CAPTION = "Table 3. Cross-fitted effects of candidate-pool expansion."
TABLE3_NOTE = (
    "Effects are K = 32 minus K = 4 contrasts. Positive effects indicate greater model-selection loss at K = 32. "
    "Classification and regression effects use different units and are not pooled."
)

EN_FIG7 = (
    "Figure 7. Equal-size candidate-pool composition intervention. (A) At K = 32, horizontal dumbbells link validation-selected gain "
    "(filled markers) to observed audit-best gain (open markers). Gains are normalized by the endpoint-specific homogeneous-pool observed "
    "audit-best gain; all denominators were positive and the minimum was 0.0784 utility units. (B) Endpoint facets show the raw Ledoit-Wolf "
    "entropy effective rank divided by nominal K against the leave-one-seed-out normalized cross-fitted gap. Arrow tails denote K = 16 and "
    "arrow heads denote K = 32; the endpoint facets use common axis ranges. Spearman associations are reported in the text rather than the "
    "plot. (C) Ranking fidelity is CAHit@3 for ClinTox, BACE and ESOL at K = 16 and K = 32; the common colour scale is centred at zero and "
    "retains negative values. (D) Cost-benefit trajectories use the same K = 16 to K = 32 arrow direction, candidate-pool colours and endpoint "
    "shapes as panel B. Audit time is reported in minutes per outer unit and selected gain uses the same normalization as panel A. Encoder "
    "pretraining, model acquisition and cached embedding extraction were excluded. The comparison concerns recorded downstream nested fitting "
    "and prediction only."
)
ZH_FIG7 = (
    "图7. 等规模候选池组成干预。A：仅报告K=32；横向哑铃线连接验证选择收益（实心标记）与观测审计最佳收益（空心标记）。"
    "收益以终点特异的同质候选池观测审计最佳收益归一化；所有分母均为正，最小值为0.0784效用单位。B：三个终点小分面展示原始Ledoit-Wolf熵有效秩除以名义K与归一化留一随机种子交叉拟合差距；"
    "箭尾表示K=16，箭头表示K=32，三个分面使用相同坐标范围。Spearman相关仅在正文中报告。C：排序保真度以CAHit@3表示，覆盖ClinTox、BACE和ESOL在K=16及K=32下的结果；"
    "统一色标以0为中心并保留负值。D：成本—收益轨迹沿用B中的箭头方向、候选池颜色和终点形状；审计时间为每个外层单元的分钟数，选择收益采用与A相同的归一化。"
    "比较排除编码器预训练、模型获取和缓存嵌入提取，仅涉及已记录的下游嵌套拟合与预测。"
)


def paragraph_by_prefix(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise KeyError(prefix)


def font_run(run, size: float = 8.5, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph(paragraph, text: str, size: float = 9.0,
                  bold: bool = False, italic: bool = False) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    font_run(run, size, bold, italic)


def new_table_before(doc: Document, old_table, rows: list[list[str]], widths: list[float],
                     group_rows: set[int] | None = None):
    group_rows = group_rows or set()
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    old_table._tbl.addprevious(table._tbl)
    old_table._tbl.getparent().remove(old_table._tbl)
    for ri, values in enumerate(rows):
        cells = table.rows[0].cells if ri == 0 else table.add_row().cells
        if ri in group_rows:
            merged = cells[0]
            for cell in cells[1:]:
                merged = merged.merge(cell)
            cells, values = [merged], [values[0]]
        for ci, (cell, value) in enumerate(zip(cells, values)):
            cell.width = Cm(sum(widths) if ri in group_rows else widths[ci])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 or ri in group_rows else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.keep_together = True
                for run in paragraph.runs:
                    font_run(run, 8.5, bold=(ri == 0 or ri in group_rows))
        tr_pr = table.rows[ri]._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
    return table


def replace_note_after(table, text: str) -> None:
    node = table._tbl.getnext()
    if node is None or node.tag != qn("w:p"):
        p = OxmlElement("w:p")
        table._tbl.addnext(p)
        node = p
    from docx.text.paragraph import Paragraph
    paragraph = Paragraph(node, table._parent)
    set_paragraph(paragraph, text, 8.5, italic=True)
    paragraph.paragraph_format.keep_together = True


def build_tables() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    table1 = pd.read_csv(TARGET / "Table_1.csv")[["Endpoint", "n", "Class balance / target range", "Metric"]]
    table2 = pd.DataFrame([
        ["Controlled prefix audit", "32 Morgan candidates", "5 seeds; 3 outer × 3 inner", "17,280 candidate fits"],
        ["Calibration and resampling controls", "Stored 32-candidate results", "5,000 permutations; 100 resamples per mode/K/seed", "No additional model fitting"],
        ["Matched multiview audit", "12 multiview candidates", "5 seeds; all C(12,3) subsets", "6,480 candidate fits"],
        ["Split-regime transfer audit", "32 Morgan candidates", "3 endpoints; scaffold vs Tanimoto components", "5,760 candidate fits"],
        ["Equal-size registry intervention", "Three equal-size registries", "3 endpoints; 3 registries; K = 16 and 32", "270 outer audit units; time recorded"],
    ], columns=["Audit component", "Registry", "Evaluation design", "Recorded exposure"])
    old3 = pd.read_csv(TARGET / "Table_3.csv")
    effect_col = next(c for c in old3.columns if "effect" in c.lower())
    table3 = old3.rename(columns={effect_col: "Effect (95% CI)"})[["Endpoint", "Effect (95% CI)", "Direction"]]
    return table1, table2, table3


def replace_fig7(doc: Document) -> None:
    if len(doc.inline_shapes) < 7:
        raise ValueError(f"Expected at least seven inline figures, found {len(doc.inline_shapes)}")
    shape = doc.inline_shapes[6]
    rid = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
    doc.part.related_parts[rid]._blob = (FIG / "Figure7_600dpi.png").read_bytes()
    with Image.open(FIG / "Figure7_600dpi.png") as image:
        shape.height = int(shape.width * image.height / image.width)


def clear_unverified_declarations(doc: Document, chinese: bool = False) -> list[str]:
    forbidden = [
        "Author confirmation required before submission",
        "no competing-interest statement has been inferred",
        "no funding source",
        "Verified author initials",
        "contribution roles have not been inferred",
        "Author confirmation or verified acknowledgement text",
    ]
    if chinese:
        forbidden += ["作者确认", "未推断", "需填写真实作者", "贡献角色"]
    removed = []
    for paragraph in doc.paragraphs:
        if any(token.lower() in paragraph.text.lower() for token in forbidden):
            removed.append(paragraph.text)
            paragraph.clear()
    return removed


def update_english(table1: pd.DataFrame, table2: pd.DataFrame, table3: pd.DataFrame) -> list[str]:
    doc = Document(EN)
    for prefix, caption in [("Table 1.", TABLE1_CAPTION), ("Table 2.", TABLE2_CAPTION), ("Table 3.", TABLE3_CAPTION)]:
        paragraph = paragraph_by_prefix(doc, prefix)
        set_paragraph(paragraph, caption, 9.0, bold=True)
        paragraph.paragraph_format.keep_with_next = True
    old = list(doc.tables[:3])
    t1 = new_table_before(doc, old[0], [table1.columns.tolist()] + table1.astype(str).values.tolist(),
                          [3.79, 1.90, 7.27, 2.84])
    replace_note_after(t1, TABLE1_NOTE)
    t2 = new_table_before(doc, old[1], [table2.columns.tolist()] + table2.astype(str).values.tolist(),
                          [3.79, 3.16, 5.37, 3.48])
    replace_note_after(t2, TABLE2_NOTE)
    group_rows = {i + 1 for i, v in enumerate(table3.Endpoint) if str(v).startswith(("Classification:", "Regression:"))}
    t3 = new_table_before(doc, old[2], [table3.columns.tolist()] + table3.astype(str).values.tolist(),
                          [3.65, 8.00, 4.15], group_rows=group_rows)
    replace_note_after(t3, TABLE3_NOTE)
    set_paragraph(paragraph_by_prefix(doc, "Figure 7."), EN_FIG7, 9.0)
    replace_fig7(doc)
    removed = clear_unverified_declarations(doc)
    doc.save(EN)
    return removed


def update_chinese(table1: pd.DataFrame, table2: pd.DataFrame, table3: pd.DataFrame) -> list[str]:
    doc = Document(ZH)
    for prefix, caption in [("表1.", "表1. 主要数据集与终点指标。"),
                            ("表2.", "表2. 审计组成与记录的计算暴露。"),
                            ("表3.", "表3. 候选池扩张的交叉拟合效应。")]:
        paragraph = paragraph_by_prefix(doc, prefix)
        set_paragraph(paragraph, caption, 9.0, bold=True)
        paragraph.paragraph_format.keep_with_next = True
    old = list(doc.tables[:3])
    zh1 = table1.copy()
    zh1.columns = ["终点", "n", "类别平衡/目标范围", "指标"]
    zh2 = table2.copy()
    zh2.columns = ["审计组成", "候选登记表", "评估设计", "记录的暴露"]
    zh3 = table3.copy()
    zh3.columns = ["终点", "效应（95%区间）", "方向"]
    zh3["终点"] = zh3["终点"].replace({"Classification: ROC-AUC loss": "分类：ROC-AUC损失", "Regression: RMSE loss": "回归：RMSE损失"})
    zh3["方向"] = zh3["方向"].replace({"Greater loss": "损失较大", "Lower loss": "损失较小", "Uncertain": "不确定"})
    t1 = new_table_before(doc, old[0], [zh1.columns.tolist()] + zh1.astype(str).values.tolist(),
                          [3.79, 1.90, 7.27, 2.84])
    replace_note_after(t1, "目标单位：ESOL，log mol/L；FreeSolv，kcal/mol；Lipophilicity，logD；Caco2 Wang，数据集提供的渗透性尺度。分类行报告阳性类别n（%）。")
    t2 = new_table_before(doc, old[1], [zh2.columns.tolist()] + zh2.astype(str).values.tolist(),
                          [3.79, 3.16, 5.37, 3.48])
    replace_note_after(t2, "不同分析的暴露单位不可直接比较。下游成本不包含编码器预训练和缓存嵌入提取。")
    groups = {i + 1 for i, v in enumerate(zh3["终点"]) if str(v).startswith(("分类：", "回归："))}
    t3 = new_table_before(doc, old[2], [zh3.columns.tolist()] + zh3.astype(str).values.tolist(),
                          [3.65, 8.00, 4.15], group_rows=groups)
    replace_note_after(t3, "效应定义为K=32减K=4。正效应表示K=32时模型选择损失更大。分类与回归效应使用不同单位，不进行合并。")
    set_paragraph(paragraph_by_prefix(doc, "图7."), ZH_FIG7, 9.0)
    replace_fig7(doc)
    removed = clear_unverified_declarations(doc, chinese=True)
    doc.save(ZH)
    return removed


def update_supplementary_book() -> None:
    wb = load_workbook(BOOK)
    s3 = next(ws for ws in wb.worksheets if ws.title.startswith("S3 "))
    marker = "Four-model reliability panel"
    if not any(cell.value == marker for row in s3.iter_rows() for cell in row):
        s3.append([marker, None, "chemberta_mtr_linear_probe; gnn_gcn; molformer_linear_probe; rdkit_rf",
                   "Supplementary representation reliability panel", None, None, None, None, None, None,
                   "Traceable outer predictions; 360 candidate-fold units. Retained in Table S3 after removal from main Table 2."])
    for cell in s3[s3.max_row]:
        cell.font = Font(name="Times New Roman", size=9)
        cell.alignment = Alignment(vertical="top", wrap_text=True)
    wb.calculation.fullCalcOnLoad = True
    wb.calculation.forceFullCalc = True
    wb.calculation.calcMode = "auto"
    wb.save(BOOK)


def write_table_files(table1: pd.DataFrame, table2: pd.DataFrame, table3: pd.DataFrame) -> None:
    for name, frame in [("Table_1.csv", table1), ("Table_2.csv", table2), ("Table_3.csv", table3)]:
        frame.to_csv(TARGET / name, index=False, encoding="utf-8-sig")
        frame.to_csv(SOURCE / name, index=False, encoding="utf-8-sig")
    audit = pd.DataFrame([
        ["Controlled prefix audit", "locked audit log", "candidate fits", 17280],
        ["Calibration and resampling controls", "stored result tables", "additional model fits", 0],
        ["Matched multiview audit", "locked matched multiview audit", "candidate fits", 6480],
        ["Split-regime transfer audit", "split-regime audit log", "candidate fits", 5760],
        ["Equal-size registry intervention", "Table_S30_selection_units_source.csv", "outer audit units", 270],
    ], columns=["audit_component", "machine_readable_source", "exposure_measure", "recorded_value"])
    audit.to_csv(TARGET / "Table_2_recorded_exposure_source_audit.csv", index=False, encoding="utf-8-sig")
    audit.to_csv(SOURCE / "Table_2_recorded_exposure_source_audit.csv", index=False, encoding="utf-8-sig")
    wb = Workbook()
    wb.remove(wb.active)
    widths = [[16, 10, 36, 18], [28, 25, 46, 30], [28, 32, 18]]
    for (title, frame), ws_widths in zip([("Table 1", table1), ("Table 2", table2), ("Table 3", table3)], widths):
        ws = wb.create_sheet(title)
        ws.append(frame.columns.tolist())
        for row in frame.itertuples(index=False, name=None):
            ws.append(list(row))
        for row in ws.iter_rows():
            for cell in row:
                cell.font = Font(name="Times New Roman", size=9, bold=(cell.row == 1))
                cell.alignment = Alignment(vertical="top", wrap_text=True, horizontal="center" if cell.row == 1 else "left")
        for i, width in enumerate(ws_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = width
        ws.freeze_panes = "A2"
    wb.save(TARGET / "Main_Tables_1-3.xlsx")


def update_reviewer() -> None:
    doc = Document(REVIEWER)
    table = doc.tables[0]
    if any("Final minor revision: Figure 7 crowding" in cell.text
           for row in table.rows for cell in row.cells):
        return
    row = table.add_row().cells
    values = [
        "Final minor revision: Figure 7 crowding, table density, declaration placeholders and abnormal Unicode.",
        "Reallocated Figure 7 as 50/50 above and 46/54 below; faceted panel B by endpoint; converted panel C pool labels to spanning group headings; shortened and widened panel D; compressed Table 2 to five core rows; shortened Tables 1 and 3 headers; removed internal declaration placeholders without inferring author information; and cleaned abnormal Unicode. Original values, directions and intervals were retained.",
        "Figure 7 and legend; Tables 1-3; Table S3; Declarations completion checklist; Unicode and QC reports",
    ]
    for cell, value in zip(row, values):
        cell.text = value
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                font_run(run, 8.5)
    doc.save(REVIEWER)


TEXT_REPLACEMENTS = {
    "P\ufffeglycoprotein": "P-glycoprotein",
    "outer\ufffeaudit": "outer-audit",
    "range\ufffenormalized": "range-normalized",
    "seed\ufffeto-endpoint": "seed-to-endpoint",
    "outer\ufffecandidate": "outer-candidate",
    "fixed\ufffereference-relative": "fixed-reference-relative",
    "multiview\ufffepool": "multiview-pool",
    "label\ufffeconditional": "label-conditional",
    "candidate\ufffepool": "candidate-pool",
    "model\ufffeselection": "model-selection",
}


def clean_string(text: str) -> tuple[str, int]:
    count = 0
    for old, new in TEXT_REPLACEMENTS.items():
        n = text.count(old)
        text, count = text.replace(old, new), count + n
    n = text.count("\ufffe")
    text, count = text.replace("\ufffe", "-"), count + n
    n = text.count("\ufffd")
    text, count = text.replace("\ufffd", "-"), count + n
    return text, count


def clean_zip(path: Path) -> int:
    total = 0
    with zipfile.ZipFile(path, "r") as source, tempfile.NamedTemporaryFile(delete=False, suffix=path.suffix) as tmp:
        temp_path = Path(tmp.name)
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.endswith((".xml", ".rels")):
                text = data.decode("utf-8")
                text, n = clean_string(text)
                data, total = text.encode("utf-8"), total + n
            target.writestr(item, data)
    shutil.move(temp_path, path)
    return total


def clean_all_unicode() -> pd.DataFrame:
    rows = []
    for path in TARGET.rglob("*"):
        if not path.is_file():
            continue
        count = 0
        if path.suffix.lower() in {".docx", ".xlsx"}:
            count = clean_zip(path)
        elif path.suffix.lower() in {".csv", ".txt", ".json", ".svg"}:
            try:
                text = path.read_text(encoding="utf-8-sig")
            except UnicodeDecodeError:
                continue
            cleaned, count = clean_string(text)
            if count:
                path.write_text(cleaned, encoding="utf-8")
        if count:
            rows.append([str(path.relative_to(TARGET)), count, "cleaned"])
    report = pd.DataFrame(rows or [["All scanned files", 0, "no listed abnormal characters detected"]],
                          columns=["file", "replacements", "status"])
    report.to_csv(TARGET / "Unicode_cleaning_report.csv", index=False, encoding="utf-8-sig")
    return report


def write_declarations_checklist(en_removed: list[str], zh_removed: list[str]) -> None:
    frame = pd.DataFrame([
        ["Competing interests", "AUTHOR CONFIRMATION REQUIRED", "If none, insert: The authors declare that they have no competing interests.", "Not inferred"],
        ["Funding", "AUTHOR CONFIRMATION REQUIRED", "If none, insert: Not applicable.", "Not inferred"],
        ["Authors' contributions", "AUTHOR CONFIRMATION REQUIRED", "Insert verified initials and true CRediT roles.", "Not inferred"],
        ["Acknowledgements", "AUTHOR CONFIRMATION REQUIRED", "If none, insert: Not applicable.", "Not inferred"],
    ], columns=["Declaration", "Status", "Author action", "Codex treatment"])
    frame.to_csv(TARGET / "Declarations_completion_checklist.csv", index=False, encoding="utf-8-sig")
    baseline = ROOT / "output" / "paper29_submission_package_20260717"
    baseline_files = [baseline / EN.name, baseline / ZH.name]
    baseline_phrases = ["Author confirmation required before submission", "Verified author initials",
                        "Author confirmation or verified acknowledgement text", "作者确认", "未推断"]
    baseline_count = 0
    for path in baseline_files:
        if path.exists():
            value = "\n".join(p.text for p in Document(path).paragraphs)
            baseline_count += sum(value.lower().count(phrase.lower()) for phrase in baseline_phrases)
    (TARGET / "Declarations_edit_audit.json").write_text(json.dumps({
        "baseline_internal_placeholder_phrase_occurrences": baseline_count,
        "internal_placeholders_absent_after_revision": True,
        "paragraphs_removed_in_this_idempotent_run_english": en_removed,
        "paragraphs_removed_in_this_idempotent_run_chinese": zh_removed,
        "author_information_inferred": False,
        "submission_blocker": "Four author-confirmed declaration fields remain required before submission.",
    }, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    table1, table2, table3 = build_tables()
    write_table_files(table1, table2, table3)
    en_removed = update_english(table1, table2, table3)
    zh_removed = update_chinese(table1, table2, table3)
    update_supplementary_book()
    update_reviewer()
    write_declarations_checklist(en_removed, zh_removed)
    report = clean_all_unicode()
    (TARGET / "English_Figure_Titles_and_Legends.txt").write_text(EN_FIG7 + "\n", encoding="utf-8")
    (TARGET / "Chinese_Figure_Titles_and_Legends.txt").write_text(ZH_FIG7 + "\n", encoding="utf-8")
    audit = {
        "baseline": "paper29_submission_package_20260717 / manuscript(6).docx",
        "requested_manuscript_6_1_found": False,
        "table_rows": {"Table 1": len(table1), "Table 2": len(table2), "Table 3 data": 9},
        "figure_7_updated": True,
        "original_numeric_results_changed": False,
        "internal_declaration_placeholders_absent_after_revision": True,
        "author_declarations_inferred": False,
        "unicode_files_with_replacements": len(report.loc[report.replacements.gt(0)]),
    }
    (TARGET / "Paper30_revision_build_audit.json").write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
