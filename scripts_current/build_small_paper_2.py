# -*- coding: utf-8 -*-
from __future__ import annotations

import copy
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(r"C:\Users\Administrator\Desktop\修改\初稿-10.docx")
OUTPUT = ROOT / "output" / "小论文-2.docx"
VALUES = ROOT / "results" / "manuscript_values.json"


def find_prefix(doc: Document, prefix: str):
    matches = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting with {prefix!r}; found {len(matches)}")
    return matches[0]


def replace_text(paragraph, text: str) -> None:
    run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = copy.deepcopy(paragraph.runs[0]._r.rPr)
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


def replace_prefix(doc: Document, prefix: str, text: str) -> None:
    replace_text(find_prefix(doc, prefix), text)


def delete_prefix(doc: Document, prefix: str) -> None:
    paragraph = find_prefix(doc, prefix)
    paragraph._element.getparent().remove(paragraph._element)


def replace_image_before_caption(doc: Document, caption_prefix: str, image_path: Path) -> None:
    caption = find_prefix(doc, caption_prefix)
    paragraphs = list(doc.paragraphs)
    index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    drawing = next(paragraph for paragraph in reversed(paragraphs[:index]) if paragraph._p.xpath(".//w:drawing"))
    for child in list(drawing._p):
        if child.tag != qn("w:pPr"):
            drawing._p.remove(child)
    drawing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    drawing.paragraph_format.first_line_indent = Cm(0)
    drawing.add_run().add_picture(str(image_path), width=Cm(16.2))


def set_cell(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def set_table(table, matrix: list[list[str]]) -> None:
    while len(table.rows) > len(matrix):
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < len(matrix):
        table.add_row()
    for row, values in zip(table.rows, matrix, strict=True):
        for cell, value in zip(row.cells, values, strict=True):
            set_cell(cell, value)


def f(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}"


def build() -> Path:
    values = json.loads(VALUES.read_text(encoding="utf-8"))
    doc = Document(SOURCE)
    candidate = values["candidate_pool"]["random_subset"]
    repeated = values["repeated_nested"]
    stable = values["repeated_stability"]
    auto = values["autogluon_budget"]
    gate = values["tdc_gate"]
    risk = values["risk_coverage"]
    conformal = values["conformal"]
    ablation = values["ablation"]
    cleaning = values["data_cleaning"]

    replace_prefix(doc, "FZYC-Mol：", "FZYC-Mol：轻量候选池扩张下分子性质模型选择的验证治理")
    replace_prefix(
        doc,
        "分子性质预测常在同一验证集上比较",
        "分子性质预测常在有限验证信息上比较不断扩张的模型候选，但固定 Top-k 命中率会随候选数机械下降，池内动态范围归一化也会掩盖跨规模选择损失。FZYC-Mol 将冻结候选登记、验证侧选择和外层审计组织为可追踪的治理协议。对 9 个终点开展的随机顺序、随机子集和家族平衡审计显示，随机子集从 4 扩至 32 个候选时，机会校正命中率由 "
        f"{f(candidate['4']['chance_adjusted_hit_mean'])} 降至 {f(candidate['32']['chance_adjusted_hit_mean'])}，全 32 池固定分母归一化遗憾由 {f(candidate['4']['fixed_regret_mean'])} 增至 {f(candidate['32']['fixed_regret_mean'])}。"
        f"在 3 外层×3 内层×5 重复的确认性嵌套验证中，机会校正命中率由 {f(repeated['4']['chance_adjusted_hit']['mean'])} 降至 {f(repeated['32']['chance_adjusted_hit']['mean'])}，固定分母遗憾由 {f(repeated['4']['fixed_normalized_regret']['mean'])} 增至 {f(repeated['32']['fixed_normalized_regret']['mean'])}，主导候选比例由 {f(stable['4']['modal_selection_rate_mean'])} 降至 {f(stable['32']['modal_selection_rate_mean'])}。"
        "验证最优、one-SE 与候选家族移除消融均未产生跨终点一致的最低遗憾。AutoGluon 在 30/300/1800 s 三个预算下相对验证最优均为 7/0/2 个终点胜/平/负；TDC 门控为 5 个晋级、17 个保留。结果表明，冻结规则和外层审计能够暴露选择风险，但不能保证普遍性能提升。确认性结论限于公开离线基准中的轻量 Morgan-512 候选池，不外推至尚未在相同外层划分统一重训的图模型和化学语言模型。",
    )
    contribution_prefixes = [
        "本研究的创新不建立",
        "（1）将候选池规模设为受控变量",
        "（2）建立冻结登记",
    ]
    contribution_texts = [
        "FZYC-Mol 以冻结候选登记和重复嵌套选择审计评价分子性质模型选择，并以全 32 池固定分母测试遗憾、机会校正排序保真度和选择稳定性作为主要证据。",
        "候选池规模与组成通过重复的分层子集和顺序扰动进行控制，从而分离搜索空间扩张、登记顺序与模型家族组成的影响。",
        "版本化决策卡把终点级选择与适用域、校准和保形证据连接到可自动重建的 source data，同时将确认性主张限定于冻结的轻量 Morgan-512 候选池。",
    ]
    for prefix, text in zip(contribution_prefixes, contribution_texts, strict=True):
        replace_prefix(doc, prefix, text)
    delete_prefix(doc, "（3）在同一外层单元")
    delete_prefix(doc, "（4）把低相似度")
    replace_prefix(doc, "关键词：", "关键词：分子性质预测；模型选择治理；候选池扩张；机会校正排名；固定分母遗憾；重复嵌套验证；保形预测")

    replace_prefix(
        doc,
        "本研究据此提出四个相互连接的问题",
        "本研究聚焦一个可检验问题：在验证信息保持不变时，候选池规模和组成的扩张是否降低排序保真度并增加外层选择遗憾。FZYC-Mol 因而是冻结登记、验证侧选择和外层审计组成的治理协议，而不是新的主干预测网络。",
    )
    replace_prefix(
        doc,
        "主要假设不是候选增加必然降低",
        "主要假设不是候选增加必然降低点预测性能，而是搜索自由度扩大后，验证排序与外层效用的一致性会下降。本文以三种随机化候选池控制作为组成与顺序审计，以 3×3×5 重复嵌套验证作为确认性检验，并用多预算 AutoGluon、TDC 门控、风险-覆盖、标签条件保形和 MoleculeACE 流程检验结论边界。",
    )

    replace_prefix(
        doc,
        "所有 SMILES 经 RDKit 解析",
        f"SMILES 先经 RDKit 解析和 Cleanup，再选择最大分子片段、在可行时中和电荷并生成带立体信息的 canonical SMILES。分类重复结构仅在标签一致时合并；冲突标签组整体排除。回归重复结构按均值聚合并记录重复数。14 个主流程端点的审计共输入 {cleaning['input_count']:,} 条记录，输出 {cleaning['output_count']:,} 条唯一结构；15 条无效 SMILES、{cleaning['duplicate_consistent_merged']} 条一致重复和 {cleaning['duplicate_conflict_excluded']} 条冲突重复均保留逐行 reason。",
    )
    replace_prefix(
        doc,
        "每次运行保存 split index",
        "每次运行保存 split index、标准化 SMILES、标签、预测、候选配置和软件版本。运行身份由 config hash、data hash、split hash、seed、code hash 和 prediction hash 联合确定；运行时间和最终指标不参与去重，因此不会因保留较快运行而系统性压低计算成本。历史文件缺少上述身份字段时标记为 unverified，而不按指标相同静默合并。",
    )
    replace_prefix(
        doc,
        "候选专家使用四类分子表示",
        "历史探索包括 Morgan 指纹、RDKit 描述符、图模型、Chemprop、ChemBERTa、MoLFormer 和预测融合。确认性重复嵌套实验仅使用 32 个可在统一外层划分稳定复跑的轻量 Morgan-512 候选；历史重型候选预测与冻结的 11/23/37/53/71 外层划分不兼容，因此不被拼接为异质确认性候选池。",
    )
    replace_prefix(
        doc,
        "候选登记表至少记录",
        "候选登记表记录 candidate_id、family、representation、eligible、status、registry_order、complexity_level 和 config_hash。确认性轻量登记包含 linear、bagging 和 boosting 家族的 32 个固定变体；任何 failed、missing_data 或 rejected 状态均保留在日志中，不由测试结果触发替换。",
    )
    replace_prefix(doc, "注：AutoGluon-Tabular 使用 Morgan-512", "注：AutoGluon-Tabular 与轻量候选使用相同 Morgan-512 特征和外层划分，分别设置每折 30/300/1800 s 时间上限；逐折记录实际耗时、模型数和峰值 RSS。测试集事后最优只用于评价。")

    replace_prefix(doc, "Stabₜ(a) =", "\tStabₜ(a) = S⁻¹∑ₛ𝟙(aₜₛ* = a)\t(5)")
    replace_prefix(doc, "aₜ* = LexMin", "\taₜ* = LexMinₐ∈Aₜ¹ˢᴱ(−Stabₜₐ, SDₜₐ, Calₜₐ, Costₐ, IDₐ)\t(7)")
    replace_prefix(
        doc,
        "词典序依次比较波动",
        "one-SE 集合内依次最大化 selection frequency，再最小化 fold variability、calibration loss、compute cost 和稳定 candidate_id。测试标签不进入集合构建、排序、阈值、风险权重或平局处理；仅在候选冻结后计算下列审计量。",
    )
    replace_prefix(doc, "R̃ₜₛ =", "\tR̃ₜₛ⁽³²⁾ = Rₜₛ/[maxₐ∈A³²uₜₐₛ⁽ᵗᵉˢᵗ⁾ − minₐ∈A³²uₜₐₛ⁽ᵗᵉˢᵗ⁾]\t(9)")
    replace_prefix(doc, "Oₜₛ =", "\tR̃ₜₛ⁽ᵈʸⁿ⁾ = Rₜₛ/[maxₐ∈Aᴷuₜₐₛ⁽ᵗᵉˢᵗ⁾ − minₐ∈Aᴷuₜₐₛ⁽ᵗᵉˢᵗ⁾]\t(10)")
    replace_prefix(doc, "Hit@kₜₛ =", "\tCAHit@3ₜₛ = [Hit@3ₜₛ − 3/K]/[1 − 3/K],   K > 3\t(11)")
    replace_prefix(
        doc,
        "式 (8)-(11) 分别给出",
        "式 (8)-(11) 分别给出绝对测试遗憾、全 32 池固定分母遗憾、仅作敏感性分析的动态分母遗憾和机会校正 Top-3 命中率。跨 K 比较同时报告 Top-25% 命中、MRR、排名百分位、NDCG、Spearman 和 Kendall；测试事后最优仅定义审计参照。",
    )
    replace_prefix(doc, "qα =", "\tqα,y = Quantile⌈(ny+1)(1−α)⌉({sᵢ: i∈cal, yᵢ=y})\t(15)")
    replace_prefix(doc, "Cαʳᵉᵍ", "\tCαʳᵉᵍ(x) = [ŷ(x)−qα, ŷ(x)+qα]\t(16)")
    replace_prefix(doc, "Cαᶜˡˢ", "\tCαᶜˡˢ(x) = {y: s(x,y) ≤ qα,y}\t(17)")
    replace_prefix(
        doc,
        "本文在 α=0.20",
        "本文在 α=0.20、0.10 和 0.05 下评价 80%、90% 和 95% 标称覆盖。分类采用标签条件阈值；某类别校准样本不足时回退到 pooled 阈值并记录 fallback_reason。回归同时报告原尺度宽度、训练标签 SD/IQR 标准化宽度和 interval score，避免跨任务直接平均原尺度区间。",
    )

    replace_prefix(
        doc,
        "核心嵌套实验覆盖 BBBP",
        "核心嵌套实验覆盖 BBBP、BACE、ClinTox、ESOL、FreeSolv、Lipophilicity、Caco2、HIA 和 Pgp。每个终点使用 3 个外层骨架折、3 个内层骨架折和 5 个预先登记的划分种子（11、23、37、53、71），形成每个池规模 135 个终点-重复-外层单元。",
    )
    replace_prefix(
        doc,
        "回顾性候选池压力分析保留",
        "候选池控制固定 K=4/8/16/32，并对 random-order、random-subset 和 family-balanced 三种模式各生成 100 个可复现子集或顺序。random-subset 从完整登记直接抽取 K 项；family-balanced 尽量保持家族比例；固定基线强制纳入时，其余 K−1 项仍按冻结随机方案抽取。",
    )
    replace_prefix(
        doc,
        "跨任务测试遗憾按同一终点",
        "绝对遗憾由子集内测试事后最优效用减去被冻结选候选效用。主分析的归一化分母始终取同一外层单元完整 32 候选的测试效用范围；当前 K 池内动态范围只作为敏感性指标。所有选择规则只读取内层验证汇总。",
    )
    replace_prefix(
        doc,
        "统计单位为终点-外层折",
        "统计以 endpoint 为主要聚类单位。分层 bootstrap 先有放回抽取 9 个终点，再在所抽终点内重采样重复-外层单元，共 5,000 次；正文报告终点等权估计、终点四分位距和 95% bootstrap 区间，不把 135 个外层单元解释为 135 个独立数据集。",
    )
    replace_prefix(
        doc,
        "AutoGluon-Tabular 对照使用相同",
        "AutoGluon-Tabular 使用相同 Morgan-512 特征、相同外层骨架划分和独立内层骨架调优折，候选限定为 LightGBM、CatBoost、随机森林和极端随机树。每个外层折分别设置 30、300 和 1,800 s 上限，记录实际拟合时间、排行榜模型数和 Linux ru_maxrss；模型完成全部固定配置后允许提前结束。",
    )

    replace_prefix(doc, "选择性预测按风险从低到高", "选择性预测按风险从低到高保留样本。覆盖率 c 下，分类风险为保留样本的错误率，回归风险为 RMSE。AURC 在 c=0.1–1.0 上积分；E-AURC 为冻结风险曲线面积减去按真实误差排序得到的 oracle 风险下界面积。该下界仅用于评价风险排序，不参与模型、阈值或权重选择。")
    replace_prefix(doc, "分类保形结果报告经验覆盖率", "分类保形结果报告总体与类别条件覆盖、平均集合大小、singleton rate、empty-set rate 和 pooled fallback；回归报告经验覆盖、SD/IQR 标准化宽度和 interval score。总体覆盖不能替代稀有类别、低相似度或高粗糙度子群的条件审计。")
    replace_prefix(doc, "MoleculeNet 结果报告 5 个随机种子", "MoleculeNet 结果报告 5 个随机种子的均值与标准差；TDC 和 bRo5 报告 3 个随机种子。候选池与重复嵌套的主不确定性由 endpoint-clustered hierarchical bootstrap 给出；候选行、外层折和随机种子均不被当作独立生物学重复。")
    replace_prefix(doc, "计算成本记录拟合时间", "计算成本记录实际拟合时间、预测时间、模型数、候选数和峰值内存。AutoGluon 的时间上限不是强制耗时；若固定模型集合提前完成，则同时报告上限和实际耗时。固定治理策略不依据终点测试结果选择最有利规则。")
    replace_prefix(doc, "复现包应包含数据下载", "复现包包含 LICENSE、requirements.lock、environment.yml、Dockerfile、持续集成配置、数据下载与清洗脚本、候选登记、逐折结果、source data、SHA256SUMS 和最小冷启动脚本。当前已完成本机冷启动和分析级重建；Docker 第三方冷启动、公开仓库 release 和 Zenodo DOI 仍待外部发布，因此不声称通过全部发布门禁。")
    replace_prefix(doc, "本文依照 Journal of Cheminformatics", "本文将主要创新限定为选择治理、受控候选池审计和可追踪结果重建；表示融合、适用域、不确定性和困难场景均视为既有方法或压力载体。由于异质重型候选尚未在相同外层划分统一重训，标题、摘要和讨论均明确限定为轻量候选。")
    replace_prefix(doc, "为限制多任务、多候选和多指标", "为限制多任务、多候选和多指标带来的叙事自由度，主要主张预先限定为机会校正排名保真度、全 32 池固定分母遗憾、重复嵌套选择稳定性和测试标签隔离；AutoGluon、TDC、风险-覆盖、保形与 MoleculeACE 为边界和对照证据。")
    replace_prefix(doc, "每项主张与统计单位一一对应", "每项主张与统计单位一一对应：候选池随机化与重复嵌套均以 endpoint 为主要聚类单位；TDC 晋级以终点为单位；风险与保形保留终点-随机种子单元；MoleculeACE 纳入以官方任务为单位。候选行、外层折和分子对不互相替代为独立重复。")
    replace_prefix(doc, "摘要、结果、表格、图注和结论", "摘要、结果、表格、图注和结论均使用同一冻结 source data。TDC 统一写为 5 promoted/17 retained，并同时报告门控错误类型；总体保形覆盖与类别条件覆盖并列。任何结果均不自动推及异质重型候选、时间外数据或低相似度子群。")

    c4, c32 = candidate["4"], candidate["32"]
    replace_prefix(doc, "3.1 候选池扩张降低", "3.1 随机化候选池控制确认规模效应")
    result_31 = [
        f"随机化候选池控制排除了固定登记顺序这一单一解释。random-subset 模式下，全 32 池固定分母遗憾从 K=4 的 {f(c4['fixed_regret_mean'])}（95% CI {f(c4['fixed_regret_ci95_low'])}–{f(c4['fixed_regret_ci95_high'])}）升至 K=32 的 {f(c32['fixed_regret_mean'])}（{f(c32['fixed_regret_ci95_low'])}–{f(c32['fixed_regret_ci95_high'])}）。",
        f"跨 K 的排序质量同向下降：random-subset 的机会校正命中率由 {f(c4['chance_adjusted_hit_mean'])} 降至 {f(c32['chance_adjusted_hit_mean'])}，MRR 由 {f(c4['mrr_mean'])} 降至 {f(c32['mrr_mean'])}。因此，扩池信号不是固定 Top-3 机会率下降的机械结果。",
        f"random-order 在 K=4/8/16/32 的固定分母遗憾为 {f(values['candidate_pool']['random_order']['4']['fixed_regret_mean'])}/{f(values['candidate_pool']['random_order']['8']['fixed_regret_mean'])}/{f(values['candidate_pool']['random_order']['16']['fixed_regret_mean'])}/{f(values['candidate_pool']['random_order']['32']['fixed_regret_mean'])}；family-balanced 对应 {f(values['candidate_pool']['family_balanced']['4']['fixed_regret_mean'])}/{f(values['candidate_pool']['family_balanced']['8']['fixed_regret_mean'])}/{f(values['candidate_pool']['family_balanced']['16']['fixed_regret_mean'])}/{f(values['candidate_pool']['family_balanced']['32']['fixed_regret_mean'])}。三种控制均在大池中保留更高选择损失。",
        "K=32 时三种模式等同于完整候选池，因此其结果一致；K<32 时不同组成产生可见波动，说明候选家族影响效应大小，但不能消除规模相关趋势。",
        "动态池内分母遗憾仅作敏感性分析。正文的跨规模比较均使用完整 32 候选的固定分母，避免后加入极端候选同时扩大分母并人为压低归一化遗憾。",
        "上述随机化分析使用既有冻结预测进行组成和顺序审计；确认性证据来自下一节的 3×3×5 重复嵌套重训。",
    ]
    for prefix, text in zip(
        ["候选池扩张持续削弱", "测试遗憾随候选数", "在 32 候选池中", "风险调整策略的平均优势", "终点分解显示", "乐观偏差的方向"],
        result_31,
        strict=True,
    ):
        replace_prefix(doc, prefix, text)
    replace_prefix(doc, "图 3 |", "图 3 | 随机化候选池控制。a，random-order、random-subset 和 family-balanced 三种模式的全 32 池固定分母遗憾；点和区间由终点级分层 bootstrap 给出。b，机会校正 Top-3 命中率。c，MRR。每个终点、模式和 K 使用 100 个冻结子集或顺序；K=32 时三种模式均为完整登记。")
    replace_image_before_caption(doc, "图 3 |", ROOT / "results/figures/candidate_pool_control.png")

    r4, r8, r16, r32 = (repeated[key] for key in ("4", "8", "16", "32"))
    replace_prefix(doc, "3.2 32 候选嵌套", "3.2 重复嵌套验证确认排序与选择不确定性")
    result_32 = [
        f"5 个预注册划分种子均完成 3 外层×3 内层嵌套重训，共得到每个池规模 135 个外层单元。全 32 池固定分母遗憾在 K=4/8/16/32 时为 {f(r4['fixed_normalized_regret']['mean'])}/{f(r8['fixed_normalized_regret']['mean'])}/{f(r16['fixed_normalized_regret']['mean'])}/{f(r32['fixed_normalized_regret']['mean'])}。",
        f"机会校正命中率由 {f(r4['chance_adjusted_hit']['mean'])}（95% CI {f(r4['chance_adjusted_hit']['ci95_low'])}–{f(r4['chance_adjusted_hit']['ci95_high'])}）降至 {f(r32['chance_adjusted_hit']['mean'])}（{f(r32['chance_adjusted_hit']['ci95_low'])}–{f(r32['chance_adjusted_hit']['ci95_high'])}）；MRR 同期由 {f(r4['mrr']['mean'])} 降至 {f(r32['mrr']['mean'])}。",
        f"选择稳定性也随扩池下降。终点内 15 个重复-外层单元的主导候选比例均值由 {f(stable['4']['modal_selection_rate_mean'])} 降至 {f(stable['32']['modal_selection_rate_mean'])}，one-SE 集合的平均成对 Jaccard 由 {f(stable['4']['pairwise_jaccard_mean'])} 降至 {f(stable['32']['pairwise_jaccard_mean'])}。",
        f"多预算 AutoGluon 结果在 30/300/1800 s 上限下均形成 27 个外层结果，相对 32 候选验证最优策略均为 7/0/2 个终点胜/平/负。实际总拟合时间分别为 {auto['30']['actual_fit_seconds_total']:.1f}/{auto['300']['actual_fit_seconds_total']:.1f}/{auto['1800']['actual_fit_seconds_total']:.1f} s；300 与 1800 s 预算在固定模型集合完成后给出相同终点级胜负。",
        "AutoGluon 相对固定单模型为 9/0/0，相对 one-SE 为 8/0/1，相对风险调整为 7/0/2。该性能优势支持把 FZYC-Mol 的价值限定为可追踪选择和偏差审计，而不是优于 AutoML 的点预测器。",
    ]
    for prefix, text in zip(
        ["九个终点均完成 3 外层", "验证集最优策略的平均归一化", "验证集最优策略的平均外层", "AutoGluon-Tabular CPU", "原四候选嵌套面板"],
        result_32,
        strict=True,
    ):
        replace_prefix(doc, prefix, text)
    replace_prefix(doc, "表 3 |", "表 3 | 轻量候选池 3×3×5 重复嵌套审计")
    replace_prefix(doc, "注：共 9 个终点、27 个外层", "注：共 9 个终点、5 个重复和 135 个外层测试单元/池规模。遗憾分母固定为同一外层单元完整 32 候选的测试效用范围；区间以终点为主要聚类单位。")
    replace_prefix(doc, "图 4 |", "图 4 | 3×3×5 重复嵌套候选池审计。a，全 32 池固定分母归一化外层测试遗憾及终点级 95% bootstrap 区间。b，机会校正命中率与 MRR。c，终点内主导候选比例和归一化选择熵。测试标签仅用于冻结后的外层评价。")
    replace_image_before_caption(doc, "图 4 |", ROOT / "results/figures/repeated_nested_control.png")

    replace_prefix(doc, "TDC 的 22 个终点采用统一冻结逻辑", f"TDC 的 22 个终点采用统一冻结门控，得到 {gate['promoted']} 个晋级和 {gate['retained']} 个保留。事后门控审计进一步区分 3 个 promoted-and-improved、7 个 retained-and-avoided-harm 和 12 个因置信区间过宽而 inconclusive 的终点；不存在 promoted-but-harmed 类别。")
    replace_prefix(doc, "其余 17 个终点中", "17 个保留终点不再被写成“0 个下降”。其中 7 个终点的保留避免了测试伤害，10 个保留终点和 2 个晋级终点因三种子置信区间过宽而证据不足。门控状态与事后类别均完整保留。")
    replace_prefix(doc, "图 6 |", "图 6 | TDC 门控有效性审计。横轴为按 gate_category 分类的终点数，颜色区分晋级与保留。22 个终点均进入审计；宽置信区间单独标记为 inconclusive，不与改善或伤害合并。")
    replace_image_before_caption(doc, "图 6 |", ROOT / "results/figures/tdc_gate_audit.png")

    class80, class90, class95 = (conformal[f"classification_{level}"] for level in ("0.80", "0.90", "0.95"))
    reg80, reg90, reg95 = (conformal[f"regression_{level}"] for level in ("0.80", "0.90", "0.95"))
    result_35 = [
        f"85 个终点-随机种子风险单元均按正确损失定义重建。分类 AURC 和 E-AURC 的中位数为 {f(risk['classification']['median_aurc'])} 和 {f(risk['classification']['median_e_aurc'])}；回归对应 {f(risk['regression']['median_aurc'])} 和 {f(risk['regression']['median_e_aurc'])}。所有 E-AURC 均为非负。",
        "oracle 曲线按真实逐样本误差从低到高排序，因此是风险下界而不是性能上界。分类纵轴为错误率，回归纵轴为 RMSE；低覆盖区的小样本波动被原样保留。",
        f"标签条件分类保形在 80%/90%/95% 目标下的总体覆盖为 {f(class80['mean_coverage'])}/{f(class90['mean_coverage'])}/{f(class95['mean_coverage'])}。类别 0 覆盖为 {f(class80['mean_class_0_coverage'])}/{f(class90['mean_class_0_coverage'])}/{f(class95['mean_class_0_coverage'])}，类别 1 仅为 {f(class80['mean_class_1_coverage'])}/{f(class90['mean_class_1_coverage'])}/{f(class95['mean_class_1_coverage'])}。",
        f"分类每个标称水平有 {class80['fallback_count']} 个单元触发 pooled fallback，主要来自 ClinTox 少数类校准样本不足。总体覆盖接近标称值并不意味着少数类条件覆盖充分。",
        f"回归在 80%/90%/95% 目标下的覆盖为 {f(reg80['mean_coverage'])}/{f(reg90['mean_coverage'])}/{f(reg95['mean_coverage'])}，按训练标签 SD 标准化的平均区间宽度为 {f(reg80['mean_normalized_width_sd'])}/{f(reg90['mean_normalized_width_sd'])}/{f(reg95['mean_normalized_width_sd'])}。",
        "覆盖增加伴随集合或区间扩大。决策卡因此同时保存类别条件覆盖、fallback、标准化宽度、相似度和风险分位数，而不以单一总体覆盖判定样本可靠。",
        "低相似度、高局部粗糙度和类别稀少分别对应不同失败机制；风险排序、适用域和保形预测只提供降置信证据，不能替代外部实验验证。",
    ]
    for prefix, text in zip(
        ["冻结风险分数在 BBBP", "风险-覆盖曲线并非处处", "保形预测的总体覆盖", "严格 Tanimoto 分层显示", "平均性能、边际覆盖", "分类保形集合的效率", "风险识别在分类和回归间"],
        result_35,
        strict=True,
    ):
        replace_prefix(doc, prefix, text)
    replace_prefix(doc, "图 7 |", "图 7 | 风险-覆盖审计。蓝线为冻结风险排序，绿虚线为按真实误差排序得到的 oracle 风险下界，灰线为随机拒用。AURC 在覆盖率 0.1–1.0 上积分，E-AURC 为冻结曲线面积减去 oracle 下界面积；分类风险为错误率，回归风险为 RMSE。")
    replace_prefix(doc, "图 8 |", "图 8 | 标签条件保形审计。a，分类总体及类别条件覆盖；少数类校准不足时记录 pooled fallback。b，回归经验覆盖与按训练标签 SD 标准化的区间宽度。点为终点-随机种子汇总，标称覆盖为 80%、90% 和 95%。")
    replace_image_before_caption(doc, "图 8 |", ROOT / "results/figures/conformal_corrected.png")

    replace_prefix(doc, "MoleculeACE 的 17 个可用任务", "MoleculeACE 官方可用配置共 17 个，本流程全部纳入且无失败或排除任务。纳入表逐任务记录训练/测试分子数、唯一分子数及 cliff molecule 数；共享分子和悬崖对不被当作独立任务重复。17 个任务×3 随机种子形成 51 个验证选择测试单元，平均 RMSE 为 0.711，活性悬崖子集 RMSE 为 0.813。")
    replace_prefix(doc, "不同任务的差异相关性变化较大", "任务级差异相关性具有明显区间异质性：高相似分子对的平均 Spearman 为 0.252，任务-随机种子范围为 −0.018 至 0.661，方向准确率为 0.750。方向准确率高于随机水平不等同于悬崖幅度预测准确，代表性分子对仅作失败案例。")

    gov = ablation["governance_rule"]
    fam = ablation["candidate_family_removal"]
    replace_prefix(doc, "3.7 统一消融", "3.7 治理规则与候选家族消融给出不同结论")
    result_37 = [
        f"治理规则消融只改变选择规则，不改变完整 32 候选组成。验证最优、冻结 one-SE、one-SE 低方差和 one-SE 低成本的固定分母遗憾分别为 {f(gov['validation_best']['mean_fixed_regret'])}/{f(gov['frozen_one_se_governance']['mean_fixed_regret'])}/{f(gov['one_se_low_variance']['mean_fixed_regret'])}/{f(gov['one_se_low_cost']['mean_fixed_regret'])}。",
        "验证最优在均值上最低，但其终点级区间与保守规则重叠；因此结果不支持把 one-SE、稳定性或成本惩罚写成普遍降低遗憾的机制。治理规则的作用是预先约束选择自由度并暴露负结果。",
        f"候选家族移除保持冻结 one-SE 规则不变。完整池、移除 bagging、移除 boosting 和移除 linear 的遗憾为 {f(fam['full_pool']['mean_fixed_regret'])}/{f(fam['remove_bagging']['mean_fixed_regret'])}/{f(fam['remove_boosting']['mean_fixed_regret'])}/{f(fam['remove_linear']['mean_fixed_regret'])}。移除任一候选家族均未产生可外推的固定预测器优势。",
        "两类消融必须分开解释：治理消融回答选择规则是否稳健，家族移除回答候选组成是否改变可选上界。完整系统不再被描述为必须包含所有模块的固定预测器。",
    ]
    for prefix, text in zip(
        ["统一消融未支持", "融合和补救头的作用", "最佳单模型和简单均值", "轻量预训练适配器"],
        result_37,
        strict=True,
    ):
        replace_prefix(doc, prefix, text)
    replace_prefix(doc, "表 7 |", "表 7 | 治理规则与候选家族移除消融")
    replace_prefix(doc, "注：Δ 表示相对完整系统", "注：治理消融仅改变选择规则；家族移除仅改变候选组成。均值和 95% 区间均按终点级分层 bootstrap 计算。")
    replace_prefix(doc, "图 10 |", "图 10 | 治理与组成消融分离。a，在完整 32 候选池上比较冻结 one-SE、验证最优、one-SE 低方差和 one-SE 低成本。b，在冻结 one-SE 规则下移除 bagging、boosting 或 linear 家族。纵轴均为全 32 池固定分母遗憾，误差条为终点级 95% bootstrap 区间。")
    replace_image_before_caption(doc, "图 10 |", ROOT / "results/figures/ablation_separation.png")

    replace_prefix(doc, "证据一致性审计区分了", "证据一致性审计确认，重复嵌套 540 行结果、三种候选池控制、TDC 22 终点、85 个风险单元、90 个保形单元和 AutoGluon 三预算共 81 个外层结果均可由结构化 CSV 重建。异质重型候选的统一外层重训和独立 ADMET 时间外盲测仍未完成。")
    replace_prefix(doc, "3.9 结果包一致性", "3.9 Source data 自动重建")
    replace_prefix(doc, "主文结果均连接到结构化文件", "主文数字、表和更新后的图 3、4、6、8、10 均由 source data 和 manuscript_values.json 自动生成；一致性验证的 difference_count 为 0。完整文件清单、SHA-256、冷启动日志和发布缺口见补充复现材料。")
    for prefix in ["文件级审计确认", "图形复现包包含", "当前已达到核心实验", "后续冷启动复跑应"]:
        delete_prefix(doc, prefix)

    discussion_texts = [
        "随机化候选池控制与 3×3×5 重复嵌套验证共同表明，候选池扩张首先削弱选择可信度，而不是简单决定最高分。机会校正命中率、MRR、固定分母遗憾和主导候选比例给出了同向证据，因此结论不依赖固定 Top-3 的机械机会效应。",
        "候选规模与候选组成均影响效应大小。random-order、random-subset 和 family-balanced 的结果在 K<32 时不同，但在完整池中收敛；这说明家族组成会改变局部难度，却不足以解释全部扩池信号。",
        "没有一种治理规则在所有终点占优。验证最优的平均遗憾低于三个保守变体，但区间重叠且终点异质性明显；移除 bagging、boosting 或 linear 也可能降低平均遗憾。冻结规则的价值因此是约束叙事自由度和保存负结果，而不是保证最低遗憾。",
        "AutoGluon 在三个预算上相对验证最优均为 7/0/2，且 300 与 1,800 s 上限没有改变终点级胜负。该结果直接承认自动化树模型的点预测优势，并把 FZYC-Mol 的贡献限定为候选登记、选择偏差审计和可追踪决策，而非性能优于 AutoML。",
        "可靠性结果也需要条件化解释。E-AURC 使用 oracle 风险下界后保持非负；标签条件保形揭示少数类覆盖低于总体覆盖；回归标准化宽度随标称覆盖增加。适用域、风险和保形输出应被视为互补的降置信证据。",
        "MoleculeACE 的 17 个官方任务全部纳入，但任务间差异相关性变化大；bRo5 和低相似度结果同样揭示结构外推压力。上述化学边界支持谨慎使用，却不能把离线基准转换为实验或临床部署证据。",
        "本研究的首要范围限制是确认性候选池仅含轻量 Morgan-512 模型。历史 Chemprop、GNN、ChemBERTa 和 MoLFormer 预测未共享冻结的外层划分，不能被事后拼接为异质嵌套证据。第二，AutoGluon 仅覆盖固定 CPU 树模型集合；第三，尚无独立 ADMET 时间外盲测。",
        "代码包已包含许可证、环境锁定、CI 配置、Dockerfile、校验和和本机冷启动结果，但公开仓库 release、Zenodo DOI 与真正第三方空环境复跑仍待作者执行。因此本文可表述为分析级可重建，而不是已通过全部开放发布门禁。",
        "综上，候选增加可提高可达到的预测上界，也会增加验证侧选择自由度。实践中应在测试前冻结候选登记和选择目标，按终点报告外层遗憾与稳定性，并把无法运行、未晋级和被保留的结果作为证据的一部分。",
    ]
    discussion_prefixes = [
        "回顾性审计与 32 候选",
        "FZYC-Mol 的四项贡献",
        "近期文献核验进一步",
        "治理规则不能消除",
        "可靠性分析同样需要",
        "化学边界结果支持",
        "本研究有四项主要限制",
        "最优先的后续工作",
        "统计单位决定结论强度",
    ]
    for prefix, text in zip(discussion_prefixes, discussion_texts, strict=True):
        replace_prefix(doc, prefix, text)
    for prefix in ["不同选择目标存在", "在实际筛选中", "因此，文稿的新颖性", "结果报告的顺序", "第三方应能够"]:
        delete_prefix(doc, prefix)

    replace_prefix(doc, "FZYC-Mol 将受控候选池扩张", "FZYC-Mol 将冻结候选登记、随机化候选池控制、验证侧选择和外层审计组织为可复核的轻量候选治理流程。随机子集控制和 3×3×5 重复嵌套验证均显示，扩池伴随机会校正排序质量下降、固定分母遗憾上升和选择稳定性降低。")
    replace_prefix(doc, "保留既有结果和负结果", "治理规则和候选家族消融未发现跨终点通用的最低遗憾策略；AutoGluon 在三个预算上均相对验证最优取得 7/0/2 个终点胜/平/负。TDC 的 17 个保留终点、少数类保形覆盖不足和所有未晋级结果均被保留。")
    replace_prefix(doc, "本文最稳健的结论是", "最稳健的结论是：扩大候选池会增加验证排序与选择的不确定性，冻结规则和外层审计能够暴露这种风险，但不能保证普遍性能提升。该结论限于公开离线基准和冻结的轻量 Morgan-512 候选池，不支持对异质重型模型、时间外 ADMET 数据或临床部署的外推。")

    replace_prefix(doc, "本文使用 MoleculeNet", "本文使用 MoleculeNet、TDC、MoleculeACE 和公共 bRo5 数据。处理数据、划分索引、候选登记、逐样本预测、source data 和 SHA-256 清单已在本地复现包整理；公开仓库地址和 Zenodo DOI 尚待作者发布：[公开仓库 URL]；[Zenodo DOI]。受原始许可证限制的数据仅提供下载脚本与校验值。")
    replace_prefix(doc, "数据清洗、32 候选嵌套重训", "数据清洗、3×3×5 重复嵌套、候选池随机化、AutoGluon 多预算、可靠性和绘图脚本已形成可运行工作流，并提供 MIT LICENSE、requirements.lock、environment.yml、Dockerfile 和 CI 配置。公开 release 标签尚待作者补充：[代码仓库 URL 与 release 标签]。")
    replace_prefix(doc, "[请作者补充基金名称", "[请作者补充基金名称、编号及资助方角色；若无资助，请明确声明“本研究未获得专项资助”。]")
    replace_prefix(doc, "[请作者按照 CRediT", "[请作者按照 CRediT 补充概念构思、方法学、软件、验证、形式分析、数据整理、初稿撰写、审阅与编辑和监督等贡献。]")
    replace_prefix(doc, "[请作者补充致谢", "生成式人工智能使用说明：本研究使用 Codex 辅助代码实现、文稿结构整理和数值一致性检查；所有实验设计、结果解释、引用和最终文字均需由作者核验并承担责任。其他致谢请作者补充。")

    replace_prefix(doc, "补充材料与图表源数据逐项对应", "补充材料逐项提供：数据清洗流与删除 reason；完整候选登记；100 次随机顺序/子集/家族平衡清单；3×3×5 重复嵌套逐折排名、遗憾和稳定性；AutoGluon 30/300/1800 s 逐折性能、耗时、模型数和内存；TDC 门控混淆类别；风险-覆盖、标签条件保形、MoleculeACE 纳入流程；治理与候选家族消融。")
    replace_prefix(doc, "每张图均配套 CSV", "更新主图均配套 source data 和生成脚本。manuscript_values.json 由结构化 CSV 自动生成并通过零差异重建；reproducibility_manifest.json 与 SHA256SUMS 用于校验文件，但公开仓库 release 和 Zenodo 归档仍待作者完成。")

    set_cell(doc.tables[0].rows[3].cells[3], "3×3×5\n135 个外层单元/池规模")
    set_cell(doc.tables[0].rows[3].cells[4], "机会校正排名；固定分母遗憾；稳定性")
    for row in doc.tables[1].rows:
        if row.cells[0].text.strip() == "AutoGluon-Tabular":
            set_cell(row.cells[1], "相同 Morgan-512；30/300/1800 s")
            set_cell(row.cells[4], "三预算均完成")
            set_cell(row.cells[5], "性能、稳定性与成本对照")

    table3 = [["候选数", "机会校正命中率", "MRR", "固定分母遗憾（95% CI）", "主导候选比例", "归一化选择熵"]]
    for key in ("4", "8", "16", "32"):
        table3.append(
            [
                key,
                f(repeated[key]["chance_adjusted_hit"]["mean"]),
                f(repeated[key]["mrr"]["mean"]),
                f"{f(repeated[key]['fixed_normalized_regret']['mean'])} ({f(repeated[key]['fixed_normalized_regret']['ci95_low'])}–{f(repeated[key]['fixed_normalized_regret']['ci95_high'])})",
                f(stable[key]["modal_selection_rate_mean"]),
                f(stable[key]["normalized_entropy_mean"]),
            ]
        )
    set_table(doc.tables[2], table3)

    table7 = [["变体", "消融类别", "候选数", "平均固定遗憾", "95% CI", "改变内容"]]
    variant_rows = [
        ("冻结 one-SE", "治理规则", "32", gov["frozen_one_se_governance"], "选择规则"),
        ("验证最优", "治理规则", "32", gov["validation_best"], "选择规则"),
        ("one-SE 低方差", "治理规则", "32", gov["one_se_low_variance"], "选择规则"),
        ("one-SE 低成本", "治理规则", "32", gov["one_se_low_cost"], "选择规则"),
        ("完整候选池", "候选组成", "32", fam["full_pool"], "候选组成"),
        ("移除 bagging", "候选组成", "24", fam["remove_bagging"], "候选组成"),
        ("移除 boosting", "候选组成", "12", fam["remove_boosting"], "候选组成"),
        ("移除 linear", "候选组成", "28", fam["remove_linear"], "候选组成"),
    ]
    for name, kind, count, stats, changed in variant_rows:
        table7.append([name, kind, count, f(stats["mean_fixed_regret"]), f"{f(stats['ci95_low'])}–{f(stats['ci95_high'])}", changed])
    set_table(doc.tables[6], table7)

    doc.core_properties.title = "FZYC-Mol：轻量候选池扩张下分子性质模型选择的验证治理"
    doc.core_properties.subject = "候选池控制、重复嵌套选择审计与可靠性证据"
    doc.core_properties.comments = "小论文-2；正文数值由 results/manuscript_values.json 自动生成。"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
