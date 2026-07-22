from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
BASE_DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_normative.docx"
BASE_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_normative.md"
OUT_DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_final_clean.docx"
OUT_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_final_clean.md"


REPLACEMENTS = [
    ("3.6 Rescue-integrated selector 与低分模块补强", "3.6 Rescue-integrated selector 与低分模块治理"),
    ("3.6 Rescue-integrated selector 与低分模块治理", "3.6 救援头集成选择器与低分模块治理"),
    ("3.6B Nature-inspired 多方法融合补强", "3.6B 文献启发的多方法融合验证"),
    ("3.2 TDC ADMET 官方 split 与外部 appendix", "3.2 TDC ADMET 官方划分与外部附录"),
    ("3.3 结构分布偏移与 split realism", "3.3 结构分布偏移与划分真实性"),
    ("3.7 Full-panel appendix、performance mode 与 roughness 解释", "3.7 全面板附录、性能模式与粗糙度解释"),
    ("3.11 按五方面补强的新增实验结果", "3.11 五类实验证据与稳健性分析"),
    ("3.12 Formal fixed-selector integration", "3.12 固定选择器策略的正式集成"),
    ("3.13 3D-lite 与 roughness-weighted regression 的负结果", "3.13 3D-lite 与粗糙度加权回归的负结果"),
    ("3.14 实验加厚后的证据链", "3.14 系统实验证据链"),
    ("3.15 强模型对照与性能补强更新", "3.15 强模型对照与选择器治理"),
    ("7. 参考文献", "6. 参考文献"),
    ("表 6. Rescue-integrated selector 诊断。", "表 6. 救援头集成选择器诊断。"),
    ("图 11. Formal fixed-selector integration。", "图 11. 固定选择器策略的正式集成。"),
    ("表 15. Formal risk-adjusted selector 带来的 retained-best promotion。", "表 15. 固定 risk-adjusted selector 带来的 retained-best promotion。"),
    ("强模型对照与 selector 治理更新", "强模型对照与 selector 治理"),
    ("新增候选/selector", "候选/selector"),
    ("新增性能模式", "性能模式候选"),
    ("新增候选必须先通过验证集证据", "候选策略必须先通过验证集证据"),
    ("新增 rescue head", "rescue head"),
    ("新增 TDC Nature-inspired fusion 后", "TDC Nature-inspired fusion 接入后"),
    ("新增实验给出了", "该组实验给出了"),
    ("新增候选不再只是", "候选池不再只是"),
    ("这个新增结果", "该结果"),
    ("新增 targeted rebuild 后", "targeted rebuild 接入后"),
    ("新增 Nature-inspired fusion 后", "Nature-inspired fusion 接入后"),
    ("新增 AD-gated/stacking fusion", "AD-gated/stacking fusion"),
    ("新增 CatBoost/XGBoost/ExtraTrees、Top-K ensemble、stacking、target transform 和 undersampling ensemble", "引入 CatBoost/XGBoost/ExtraTrees、Top-K ensemble、stacking、target transform 和 undersampling ensemble"),
    ("新增 Nature-inspired fusion 已经", "Nature-inspired fusion 已经"),
    ("新增分析不重启", "这些分析不重启"),
    ("新增 external appendix", "external appendix"),
    ("新增实验从“事后挑结果”改成", "候选扩展从“事后挑结果”改成"),
    ("本稿将新增实验统一组织为", "本稿将实验结果统一组织为"),
    ("新增候选没有被正式接入", "候选没有被正式接入"),
    ("新增 TabPFNv2、树模型全集、Chemprop/D-MPNN 和冻结分子语言模型表征后", "TabPFNv2、树模型全集、Chemprop/D-MPNN 和冻结分子语言模型表征进入候选池后"),
    ("新增外部 benchmark/appendix", "外部 benchmark/appendix"),
    ("进一步把“增加模型复杂度”和“冲性能”拆成两条可审计路线", "进一步把模型复杂度扩展和性能优化拆成两条可审计路线"),
    ("根据最新实验方向", "根据模型比较与选择器治理需求"),
    ("继续加厚", "继续扩展实验"),
    ("补强策略", "候选策略"),
    ("补强模块", "候选模块"),
]


def clean_text(text: str) -> str:
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    text = text.replace("新增", "引入的")
    return text


def set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = text


def element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter() if node.tag == qn("w:t"))


def remove_docx_section(doc: Document, start_prefix: str, end_prefix: str) -> None:
    body = doc._body._element
    children = list(body)
    start = None
    end = None
    for idx, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = element_text(child).strip()
        if start is None and text.startswith(start_prefix):
            start = idx
        elif start is not None and text.startswith(end_prefix):
            end = idx
            break
    if start is None or end is None:
        raise RuntimeError(f"Could not find section {start_prefix!r} -> {end_prefix!r}.")
    for child in children[start:end]:
        body.remove(child)


def clean_docx() -> None:
    shutil.copy2(BASE_DOCX, OUT_DOCX)
    doc = Document(OUT_DOCX)
    for paragraph in doc.paragraphs:
        cleaned = clean_text(paragraph.text)
        if cleaned != paragraph.text:
            set_paragraph_text(paragraph, cleaned)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    cleaned = clean_text(paragraph.text)
                    if cleaned != paragraph.text:
                        set_paragraph_text(paragraph, cleaned)
    remove_docx_section(doc, "6. 图表对应关系", "6. 参考文献")
    doc.save(OUT_DOCX)


def clean_markdown() -> None:
    text = BASE_MD.read_text(encoding="utf-8")
    text = clean_text(text)
    text = re.sub(r"\n# 6\. 图表对应关系\n.*?(?=\n# 6\. 参考文献\n)", "\n", text, flags=re.S)
    OUT_MD.write_text(text, encoding="utf-8")


def main() -> None:
    clean_docx()
    clean_markdown()
    print(f"wrote {OUT_DOCX.relative_to(ROOT)}")
    print(f"wrote {OUT_MD.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
