from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    path = Path(r"C:\Users\Administrator\Desktop\修改\初稿-15.docx")
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    print(f"Document: {path}")
    print(f"Paragraphs: {len(paras)}; tables: {len(doc.tables)}; images: {len(doc.inline_shapes)}")
    terms = [
        "未跑",
        "已跑",
        "没跑",
        "审稿人要求",
        "按照第一个文件",
        "当前草稿",
        "Codex",
        "完美",
        "绝对",
        "虚构",
        "伪造",
        "Figure X",
        "Table X",
    ]
    print("\nTerm audit:")
    for term in terms:
        hits = [(idx + 1, text) for idx, text in enumerate(paras) if term in text]
        print(f"- {term}: {len(hits)}")
        for idx, text in hits[:3]:
            pos = text.find(term)
            print(f"  P{idx}: {text[max(0, pos-80):pos+120]}")
    print("\nKey updated paragraphs:")
    for key in ["bRo5", "MoleculeACE", "FreeSolv", "Full、best single", "fixed precision", "3 x 3 nested", "轻量大模型适配器"]:
        print(f"\n## {key}")
        count = 0
        for idx, text in enumerate(paras, 1):
            if key in text:
                print(f"P{idx}: {text[:700]}")
                count += 1
                if count >= 3:
                    break
    print("\nTable widths:")
    for i, table in enumerate(doc.tables, 1):
        cols = len(table.columns) if table.rows else 0
        print(f"T{i:02d}: {len(table.rows)}x{cols}")


if __name__ == "__main__":
    main()
