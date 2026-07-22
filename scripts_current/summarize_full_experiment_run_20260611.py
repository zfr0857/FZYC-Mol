from __future__ import annotations

import csv
import shutil
from collections import Counter
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "full_missing_experiment_run_20260611"
OUT.mkdir(parents=True, exist_ok=True)
DESKTOP = Path.home() / "Desktop" / "修改"
SOURCE_DOCX = DESKTOP / "初稿-15.docx"
UPDATED_DOCX = DESKTOP / "初稿-16_补充实验整理版.docx"
REPORT_DOCX = DESKTOP / "补充实验全跑整理报告_20260611.docx"


def fmt(value: object, digits: int = 3) -> str:
    try:
        if pd.isna(value):
            return ""
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def mean_pm(series: pd.Series) -> str:
    series = pd.to_numeric(series, errors="coerce").dropna()
    if series.empty:
        return ""
    return f"{series.mean():.3f} ± {series.std(ddof=1):.3f}" if len(series) > 1 else f"{series.iloc[0]:.3f}"


def compact_path(path: str | Path) -> str:
    p = Path(path)
    try:
        return str(p.relative_to(ROOT)).replace("\\", "/")
    except Exception:
        return str(p).replace("\\", "/")


def load_bro5_summary() -> tuple[pd.DataFrame, list[list[str]], str]:
    path = ROOT / "reports" / "bro5_cycpept_pampa_20260611" / "validation_selected_results.csv"
    df = pd.read_csv(path)
    rows = []
    for split, group in df.groupby("split", sort=False):
        rows.append(
            [
                split,
                str(len(group)),
                mean_pm(group["test_rmse"]),
                mean_pm(group["test_mae"]),
                mean_pm(group["test_spearman"]),
                "; ".join(f"{k}:{v}" for k, v in Counter(group["model"]).items()),
            ]
        )
    summary = pd.DataFrame(rows, columns=["split", "n_seed", "test_RMSE", "test_MAE", "test_Spearman", "selected_models"])
    summary.to_csv(OUT / "bro5_cycpept_pampa_compact_summary.csv", index=False, encoding="utf-8-sig")
    n = int(df["n_train"].iloc[0] + df["n_valid"].iloc[0] + df["n_test"].iloc[0])
    line = (
        f"CycPept-PAMPA n={n}; 4 split x 3 seeds. "
        f"Random RMSE {summary.loc[summary.split.eq('random'), 'test_RMSE'].iloc[0]}, "
        f"scaffold RMSE {summary.loc[summary.split.eq('scaffold'), 'test_RMSE'].iloc[0]}, "
        f"perimeter RMSE {summary.loc[summary.split.eq('perimeter'), 'test_RMSE'].iloc[0]}, "
        f"time RMSE {summary.loc[summary.split.eq('time'), 'test_RMSE'].iloc[0]}."
    )
    return df, rows, line


def load_adapter_summary() -> tuple[pd.DataFrame, str]:
    selected = pd.read_csv(ROOT / "reports" / "pretrained_lightweight_adapter_20260611" / "adapter_validation_selected.csv")
    n = len(selected)
    n_mlp = int(selected["adapter_type"].eq("mlp_adapter").sum())
    mean_delta = selected["delta_vs_linear_test_positive"].mean()
    by_dataset = (
        selected.groupby("dataset")
        .agg(n=("seed", "count"), mlp_selected=("adapter_type", lambda s: int((s == "mlp_adapter").sum())), mean_delta=("delta_vs_linear_test_positive", "mean"))
        .reset_index()
    )
    by_dataset.to_csv(OUT / "adapter_compact_summary.csv", index=False, encoding="utf-8-sig")
    line = f"Lightweight adapters were selected in {n_mlp}/{n} dataset-encoder-seed units; mean test delta versus linear probe was {mean_delta:.3f} in the task-specific positive direction."
    return selected, line


def load_moleculeace_summary() -> tuple[pd.DataFrame, str]:
    raw = pd.read_csv(ROOT / "reports" / "moleculeace_full_available_20260611" / "metrics_raw.csv")
    selected = raw[(raw["split"].eq("test")) & (raw["selected_by_validation"].eq(1))].copy()
    selected.to_csv(OUT / "moleculeace_full_available_selected_rows.csv", index=False, encoding="utf-8-sig")
    n_tasks = selected["task"].nunique()
    n_rows = len(selected)
    line = (
        f"MoleculeACE available panel: {n_tasks} Hugging Face configs x 3 seeds = {n_rows} validation-selected test rows; "
        f"mean RMSE {selected['rmse'].mean():.3f}, mean cliff RMSE {selected['cliff_rmse'].mean():.3f}."
    )
    return selected, line


def load_other_lines() -> dict[str, str]:
    nested = pd.read_csv(ROOT / "reports" / "remaining_missing_experiments_20260606" / "true_nested_validation" / "true_nested_validation_summary.csv")
    conformal = pd.read_csv(ROOT / "reports" / "remaining_missing_experiments_20260606" / "conformal_80_90_95_summary.csv")
    ablation = pd.read_csv(ROOT / "reports" / "remaining_missing_experiments_20260606" / "unified_ablation_matrix_summary.csv")
    gap = pd.read_csv(ROOT / "reports" / "remaining_missing_experiments_20260606" / "moleculeace_gap_correlation_summary.csv")
    three_d = pd.read_csv(ROOT / "reports" / "three_d_roughness_regression_experts_20260603" / "retained_best_summary.csv")
    return {
        "nested": f"True nested validation completed on {len(nested)} representative endpoints with 3 outer folds each.",
        "conformal": "Conformal prediction completed at 80/90/95% target coverage for classification and regression tasks.",
        "ablation": f"Unified ablation matrix completed with {len(ablation)} task-type/category rows; no table exceeds 7 columns in the manuscript audit.",
        "gap": f"MoleculeACE gap-correlation audit completed: mean gap Spearman {gap['gap_spearman'].mean():.3f}, mean direction accuracy {gap['direction_accuracy'].mean():.3f}.",
        "three_d": f"3D-lite/roughness regression gate verified across {len(three_d)} retained-best rows.",
    }


def build_completion_matrix(bro5_line: str, adapter_line: str, mace_line: str, other: dict[str, str]) -> pd.DataFrame:
    rows = [
        {
            "module": "bRo5 CycPept-PAMPA",
            "status": "completed_partial",
            "scope": "CycPeptMPDB/CycPeptMP PAMPA subset; random/scaffold/perimeter/time splits; 3 seeds",
            "key_result": bro5_line,
            "evidence": "reports/bro5_cycpept_pampa_20260611/",
            "manuscript_use": "Report as available bRo5 stress test; keep LinPept entries as missing_data.",
        },
        {
            "module": "True nested validation",
            "status": "completed",
            "scope": "9 representative MoleculeNet/TDC endpoints; 3 outer x 3 inner",
            "key_result": other["nested"],
            "evidence": "reports/remaining_missing_experiments_20260606/true_nested_validation/",
            "manuscript_use": "Use as selector-bias diagnostic, not replacement for frozen test results.",
        },
        {
            "module": "Unified ablation matrix",
            "status": "completed",
            "scope": "Full, best single, simple mean, no selector, no fusion, no AD gate, no uncertainty, no motif/fingerprint, no rescue head",
            "key_result": other["ablation"],
            "evidence": "reports/remaining_missing_experiments_20260606/unified_ablation_matrix_summary.csv",
            "manuscript_use": "Report endpoint-dependent module contribution and negative results.",
        },
        {
            "module": "Exact Tanimoto bins",
            "status": "completed",
            "scope": ">0.7, 0.5-0.7 and <0.5 mutually exclusive bins",
            "key_result": "Performance, uncertainty, calibration proxies and risk enrichment were generated for the low-similarity bins.",
            "evidence": "reports/remaining_missing_experiments_20260606/exact_tanimoto_bins_summary.csv",
            "manuscript_use": "Use as strict low-similarity/OOD evidence.",
        },
        {
            "module": "Conformal 80/90/95",
            "status": "completed",
            "scope": "Classification and regression conformal outputs",
            "key_result": other["conformal"],
            "evidence": "reports/remaining_missing_experiments_20260606/conformal_80_90_95_summary.csv",
            "manuscript_use": "Report coverage and information-width trade-off.",
        },
        {
            "module": "MoleculeACE full available panel",
            "status": "completed_available",
            "scope": "17 available Hugging Face configs; 3 seeds; balanced cliff selection objective",
            "key_result": mace_line,
            "evidence": "reports/moleculeace_full_available_20260611/",
            "manuscript_use": "State full available panel, not 30-task claim.",
        },
        {
            "module": "MoleculeACE gap correlation and cases",
            "status": "completed",
            "scope": "Prediction-gap versus true-gap correlation and representative cliff pairs",
            "key_result": other["gap"],
            "evidence": "reports/remaining_missing_experiments_20260606/moleculeace_gap_correlation_summary.csv",
            "manuscript_use": "Use as activity-cliff risk evidence, not causal mechanism.",
        },
        {
            "module": "Lightweight pretrained adapter",
            "status": "completed_controlled",
            "scope": "ChemBERTa and MoLFormer frozen embeddings; linear probe vs one-hidden-layer MLP adapter; 6 datasets x 2 encoders x 3 seeds",
            "key_result": adapter_line,
            "evidence": "reports/pretrained_lightweight_adapter_20260611/",
            "manuscript_use": "Report as controlled supplement; do not claim full fine-tuning.",
        },
        {
            "module": "3D-lite and physical descriptors",
            "status": "completed_verified",
            "scope": "MoleculeNet regression and TDC ADME regression endpoints",
            "key_result": other["three_d"],
            "evidence": "reports/three_d_roughness_regression_experts_20260603/",
            "manuscript_use": "Keep as endpoint-dependent rescue/negative-result evidence.",
        },
        {
            "module": "Failure-case evidence chain",
            "status": "completed_expanded",
            "scope": "Low-similarity OOD cases and MoleculeACE cliff-pair cases",
            "key_result": "Extended failure cases table and molecule-pair figures were generated.",
            "evidence": "reports/remaining_missing_experiments_20260606/extended_failure_cases.csv",
            "manuscript_use": "Use to connect performance, AD and chemical interpretation.",
        },
        {
            "module": "LinPept NonFouling / LinPept CellPen",
            "status": "missing_data",
            "scope": "No local or directly verified public CSV/TSV with required SMILES/y fields in this run",
            "key_result": "Not run; no performance claim should be made.",
            "evidence": "reports/experiment_update/bro5_data_status.csv",
            "manuscript_use": "State as missing_data/future validation only.",
        },
    ]
    df = pd.DataFrame(rows)
    df.to_csv(OUT / "experiment_completion_matrix_20260611.csv", index=False, encoding="utf-8-sig")
    return df


def add_docx_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    apply_three_line_table(table)


def set_border(element, edge: str, val: str, size: str = "8", color: str = "000000") -> None:
    tag = "w:tblBorders" if element.tag.endswith("tblPr") else "w:tcBorders"
    borders = element.find(qn(tag))
    if borders is None:
        borders = OxmlElement(tag)
        element.append(borders)
    edge_el = borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        borders.append(edge_el)
    edge_el.set(qn("w:val"), val)
    if val != "nil":
        edge_el.set(qn("w:sz"), size)
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)


def apply_three_line_table(table) -> None:
    tbl_pr = table._tbl.tblPr
    for edge in ("left", "right", "insideH", "insideV"):
        set_border(tbl_pr, edge, "nil")
    set_border(tbl_pr, "top", "single", "12")
    set_border(tbl_pr, "bottom", "single", "12")
    if table.rows:
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for edge in ("left", "right", "top"):
                set_border(tc_pr, edge, "nil")
            set_border(tc_pr, "bottom", "single", "8")
        for row in table.rows[1:]:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                for edge in ("left", "right", "top", "bottom"):
                    set_border(tc_pr, edge, "nil")


def write_markdown_report(matrix: pd.DataFrame, bro5_rows: list[list[str]]) -> None:
    lines = [
        "# Full Missing-Experiment Run Summary, 2026-06-11",
        "",
        "All runnable missing experiments were rerun or newly added. Experiments without verified local/public data are retained as missing_data and are not converted into performance claims.",
        "",
        "## Completion matrix",
        "",
        matrix.to_markdown(index=False),
        "",
        "## CycPept-PAMPA bRo5 compact results",
        "",
        pd.DataFrame(bro5_rows, columns=["split", "n_seed", "test_RMSE", "test_MAE", "test_Spearman", "selected_models"]).to_markdown(index=False),
        "",
    ]
    (OUT / "full_missing_experiment_run_summary.md").write_text("\n".join(lines), encoding="utf-8")


def write_docx_report(matrix: pd.DataFrame, bro5_rows: list[list[str]]) -> None:
    doc = Document()
    doc.add_heading("补充实验全跑整理报告", level=1)
    doc.add_paragraph("日期：2026-06-11")
    doc.add_paragraph(
        "本报告汇总初稿-15 中被标注为未完成、待验证或仅有数据状态的实验模块。已能运行的模块均保留原始输出；无可核验数据源的模块仅记录为 missing_data，不写作性能结论。"
    )
    doc.add_heading("完成度矩阵", level=2)
    add_docx_table(
        doc,
        ["模块", "状态", "范围", "关键结果", "证据文件", "论文处理"],
        matrix[["module", "status", "scope", "key_result", "evidence", "manuscript_use"]].values.tolist(),
    )
    doc.add_heading("bRo5 CycPept-PAMPA 紧凑结果", level=2)
    add_docx_table(doc, ["划分", "seed数", "测试RMSE", "测试MAE", "测试Spearman", "验证选择模型"], bro5_rows)
    doc.add_heading("写入论文时的边界", level=2)
    doc.add_paragraph(
        "CycPept-PAMPA 可以写成一个已完成的 bRo5 公开压力测试；LinPept NonFouling 和 LinPept CellPen 仍应写为 missing_data。MoleculeACE 可写为当前 Hugging Face 数据源可获得的 17 个任务全量运行，不能写成 30 任务全量完成。轻量 adapter 仅为冻结表征上的 MLP adapter/linear probe 对照，不能写成大模型 full fine-tuning。"
    )
    doc.save(REPORT_DOCX)
    shutil.copy2(REPORT_DOCX, OUT / REPORT_DOCX.name)


def clear_paragraph(paragraph: Paragraph) -> None:
    if hasattr(paragraph, "clear"):
        paragraph.clear()
    else:
        for run in paragraph.runs:
            run.text = ""


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)


def replace_text(doc: Document, replacements: dict[str, str]) -> list[str]:
    changed = []
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new = text
        for old, value in replacements.items():
            new = new.replace(old, value)
        if new != text:
            set_paragraph_text(paragraph, new)
            changed.append(text[:60])
    return changed


def insert_before(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    set_paragraph_text(new_para, text)
    return new_para


def insert_table_before(paragraph: Paragraph, headers: list[str], rows: list[list[str]]) -> None:
    tbl = OxmlElement("w:tbl")
    paragraph._p.addprevious(tbl)
    table = paragraph._parent.add_table(rows=1, cols=len(headers), width=Inches(6.5))
    tbl.addprevious(table._tbl)
    tbl.getparent().remove(tbl)
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    apply_three_line_table(table)


def update_manuscript(matrix: pd.DataFrame, bro5_rows: list[list[str]], bro5_line: str, adapter_line: str, mace_line: str) -> None:
    doc = Document(SOURCE_DOCX)
    replacements = {
        "bRo5 模块在当前版本中仅作为数据状态审计与后续运行接口，不构成已完成性能评估。": "bRo5 模块已新增 CycPept-PAMPA 公开子集压力测试，但 LinPept NonFouling 和 LinPept CellPen 仍为 missing_data，因此该模块不能表述为完整 bRo5 性能评估。",
        "bRo5 当前仅完成数据状态审计，不作为已完成性能结果。": "bRo5 当前新增 CycPept-PAMPA 公开子集结果；LinPept NonFouling 和 LinPept CellPen 仍保留为 missing_data，因此不作为完整 bRo5 全量结论。",
        "bRo5 三个预设数据集在当前审计中均为 missing_data，不能作为已完成性能评估；MoleculeACE 当前可核验结果覆盖 17 个任务和 51 个 seed 配对，不能表述为完整 30 任务全量结论。": "bRo5 审计已新增 CycPept-PAMPA 公开子集实验，覆盖 random、scaffold、perimeter 和 time split；但 LinPept NonFouling 和 LinPept CellPen 仍为 missing_data，因此不能表述为完整 bRo5 全量结论。MoleculeACE 当前可访问数据源包含 17 个任务，本文已完成这 17 个可用任务的 51 个 seed 配对，不能表述为 30 任务全量结论。",
    }
    replace_text(doc, replacements)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "4.7 可解释性与案例分析":
            insert_before(paragraph, "4.6.1 补充实验补齐后的证据状态", style="Heading 3")
            insert_before(
                paragraph,
                "根据补充实验清单，本文进一步补齐了能够在当前数据和依赖条件下运行的关键模块。新增 bRo5 CycPept-PAMPA 实验覆盖 random、scaffold、perimeter 和 time split，验证集选择的测试结果显示："
                + bro5_line
                + " 轻量分子大模型适配器在冻结 ChemBERTa/MoLFormer 表征上完成 linear probe 与 MLP adapter 对照，"
                + adapter_line
                + " MoleculeACE 当前可访问的 17 个任务均已按 3 seeds 运行，"
                + mace_line
                + " 这些结果扩大了证据链，但仍保持边界：LinPept 两项、官方盲测提交和大规模 full fine-tuning 尚不具备完整可核验证据。",
            )
            compact_rows = matrix[["module", "status", "key_result", "evidence", "manuscript_use"]].values.tolist()
            insert_table_before(paragraph, ["模块", "状态", "关键结果", "证据文件", "论文处理"], compact_rows)
            break
    doc.save(UPDATED_DOCX)
    shutil.copy2(UPDATED_DOCX, OUT / UPDATED_DOCX.name)


def main() -> None:
    _, bro5_rows, bro5_line = load_bro5_summary()
    _, adapter_line = load_adapter_summary()
    _, mace_line = load_moleculeace_summary()
    other = load_other_lines()
    matrix = build_completion_matrix(bro5_line, adapter_line, mace_line, other)
    write_markdown_report(matrix, bro5_rows)
    write_docx_report(matrix, bro5_rows)
    update_manuscript(matrix, bro5_rows, bro5_line, adapter_line, mace_line)
    print(f"Wrote {OUT / 'experiment_completion_matrix_20260611.csv'}")
    print(f"Wrote {OUT / 'full_missing_experiment_run_summary.md'}")
    print(f"Wrote {REPORT_DOCX}")
    print(f"Wrote {UPDATED_DOCX}")


if __name__ == "__main__":
    main()
