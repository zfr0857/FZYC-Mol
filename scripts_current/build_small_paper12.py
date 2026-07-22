from __future__ import annotations

import json
import re
import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
SRC = OUT / "小论文-11.docx"
DOCX = OUT / "小论文-12.docx"
EXP = OUT / "小论文-12_严格补实验"
AUDIT = ROOT / "results" / "audits" / "small_paper_12_audit.json"


TITLE = "候选池扩张增加分子性质预测中的选择损失"


def set_para_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    p.paragraph_format.line_spacing = 1.15


def insert_after(p, text: str, style: str | None = None):
    new = p.insert_paragraph_before("")
    p._p.addprevious(new._p)
    p._p.addnext(new._p)
    if style:
        new.style = style
    set_para_text(new, text)
    new.paragraph_format.first_line_indent = Cm(0.74) if style != "Heading 2" else Cm(0)
    new.paragraph_format.space_after = Pt(4)
    return new


def replace_first_containing(doc: Document, needle: str, text: str) -> None:
    for p in doc.paragraphs:
        if needle in p.text:
            set_para_text(p, text)
            return
    raise ValueError(f"needle not found: {needle}")


def insert_after_first_containing(doc: Document, needle: str, additions: list[str]) -> None:
    for p in doc.paragraphs:
        if needle in p.text:
            last = p
            for text in additions:
                last = insert_after(last, text)
            return
    raise ValueError(f"needle not found: {needle}")


def replace_all_literals(doc: Document, replacements: dict[str, str]) -> None:
    for p in doc.paragraphs:
        text = p.text
        updated = text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != text:
            set_para_text(p, updated)


def xml_errors(path: Path) -> list[str]:
    errors = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if name.endswith(".xml"):
                try:
                    ElementTree.fromstring(zf.read(name))
                except Exception as exc:
                    errors.append(f"{name}: {exc}")
    return errors


def metrics() -> dict[str, object]:
    combined = json.loads((EXP / "paper12_combined_gap_audit.json").read_text(encoding="utf-8"))
    strict = json.loads((EXP / "paper12_strict_gap_experiments_audit.json").read_text(encoding="utf-8"))
    chem = json.loads((EXP / "chemprop_outer_audit.json").read_text(encoding="utf-8"))
    chem_inner = json.loads((EXP / "chemprop_inner_audit.json").read_text(encoding="utf-8"))
    core = json.loads((EXP / "paper12_core_addon_audit.json").read_text(encoding="utf-8"))
    strong = pd.read_csv(EXP / "combined_strong_baseline_summary.csv")
    overlap = pd.read_csv(EXP / "combined_strong_error_overlap_pairwise_summary.csv")
    dup = pd.read_csv(EXP / "duplicate_sensitivity_summary.csv")
    selection = pd.read_csv(EXP / "combined_strong_selection_summary.csv")
    selected_counts = pd.read_csv(EXP / "combined_strong_selected_candidate_counts.csv")
    paired = pd.read_csv(EXP / "combined_strong_candidate_paired_effects.csv")
    status = pd.read_csv(EXP / "strong_baseline_runtime_status.csv")
    chem_status = pd.read_csv(EXP / "chemprop_outer_runtime_status.csv")

    def row(candidate: str, task_type: str):
        sub = strong[(strong["candidate"].eq(candidate)) & (strong["task_type"].eq(task_type))]
        return sub.iloc[0].to_dict() if len(sub) else {}

    def selection_row(task: str):
        sub = selection[selection["task"].eq(task)]
        return sub.iloc[0].to_dict() if len(sub) else {}

    def paired_row(candidate: str):
        sub = paired[(paired["candidate"].eq(candidate)) & (paired["task"].eq("__overall__"))]
        return sub.iloc[0].to_dict() if len(sub) else {}

    max_overlap = overlap.sort_values("mean_jaccard_error_overlap", ascending=False).iloc[0].to_dict()
    min_overlap = overlap.sort_values("mean_jaccard_error_overlap", ascending=True).iloc[0].to_dict()
    max_dup = dup.loc[dup["delta_vs_global_dedup"].abs().idxmax()].to_dict()
    tab = status[status["candidate"].eq("tabpfn_rdkit")].iloc[0].to_dict()
    overall_selection = selection_row("__overall__")
    rdkit_selected = selected_counts[selected_counts["selected_candidate"].eq("rdkit_rf")]["n_selected_units"].sum()
    return {
        "combined_prediction_rows": combined["combined_prediction_rows"],
        "combined_candidates": combined["combined_candidates"],
        "combined_overlap_pairs": combined["combined_strong_error_overlap_pairs"],
        "combined_overlap_mean": combined["combined_strong_error_overlap_mean"],
        "combined_inner_rows": core["combined_inner_score_rows"],
        "combined_outer_rows": core["combined_outer_score_rows"],
        "strict_inner_rows": strict["strong_inner_rows"],
        "strict_outer_rows": strict["strong_outer_rows"],
        "strict_prediction_rows": strict["strong_prediction_rows"],
        "duplicate_rows": strict["duplicate_sensitivity_rows"],
        "duplicate_max_abs_delta": strict["duplicate_max_abs_delta_vs_global"],
        "chemprop_rows": chem["rows"],
        "chemprop_prediction_rows": chem["prediction_rows"],
        "chemprop_inner_rows": chem_inner["inner_rows"],
        "chemprop_inner_prediction_rows": chem_inner["prediction_rows"],
        "tabpfn_status": tab["status"],
        "tabpfn_reason": tab["reason"],
        "chemprop_status": chem_status.iloc[0]["status"],
        "selection_units": core["strong_selection_units"],
        "selection_top1_hit_rate": overall_selection.get("outer_top1_hit_rate"),
        "selection_loss": overall_selection.get("mean_selection_loss"),
        "selection_norm_loss": overall_selection.get("mean_range_normalized_selection_loss"),
        "selection_spearman": overall_selection.get("mean_inner_outer_spearman"),
        "rdkit_selected_units": int(rdkit_selected),
        "molformer_delta_vs_rdkit": paired_row("molformer_linear_probe").get("mean_delta_vs_rdkit"),
        "chemberta_delta_vs_rdkit": paired_row("chemberta_mtr_linear_probe").get("mean_delta_vs_rdkit"),
        "chemprop_delta_vs_rdkit": paired_row("chemprop_dmpnn").get("mean_delta_vs_rdkit"),
        "gnn_delta_vs_rdkit": paired_row("gnn_gcn").get("mean_delta_vs_rdkit"),
        "molformer_win_rate_vs_rdkit": paired_row("molformer_linear_probe").get("win_rate_vs_rdkit"),
        "rdkit_cls_auc": row("rdkit_rf", "classification").get("roc_auc_mean"),
        "molformer_cls_auc": row("molformer_linear_probe", "classification").get("roc_auc_mean"),
        "chemberta_cls_auc": row("chemberta_mtr_linear_probe", "classification").get("roc_auc_mean"),
        "chemprop_cls_auc": row("chemprop_dmpnn", "classification").get("roc_auc_mean"),
        "gnn_cls_auc": row("gnn_gcn", "classification").get("roc_auc_mean"),
        "rdkit_esol_rmse": row("rdkit_rf", "regression").get("rmse_mean"),
        "chemberta_esol_rmse": row("chemberta_mtr_linear_probe", "regression").get("rmse_mean"),
        "molformer_esol_rmse": row("molformer_linear_probe", "regression").get("rmse_mean"),
        "chemprop_esol_rmse": row("chemprop_dmpnn", "regression").get("rmse_mean"),
        "gnn_esol_rmse": row("gnn_gcn", "regression").get("rmse_mean"),
        "max_overlap_pair": f"{max_overlap['candidate_a']} vs {max_overlap['candidate_b']}",
        "max_overlap": max_overlap["mean_jaccard_error_overlap"],
        "min_overlap_pair": f"{min_overlap['candidate_a']} vs {min_overlap['candidate_b']}",
        "min_overlap": min_overlap["mean_jaccard_error_overlap"],
        "max_dup_task": max_dup["task"],
        "max_dup_policy": max_dup["policy"],
        "max_dup_delta": max_dup["delta_vs_global_dedup"],
    }


def revise() -> dict[str, object]:
    m = metrics()
    shutil.copy2(SRC, DOCX)
    doc = Document(DOCX)
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "小论文-12：标题与摘要按 Nature 风格重写"
    set_para_text(doc.paragraphs[0], TITLE)

    abstract_parts = [
        (
            "分子性质预测中的最终模型通常从不断扩张的候选池中选出。"
            "这一做法可以提高可达到性能上界，但也会在固定验证信息上反复消费排序信号，使最终外层表现混合真实表征收益和模型选择偏差。"
        ),
        (
            "本文提出 FZYC-Mol（Frozen validation governance for molecular model selection），将候选登记、嵌套选择、策略冻结、外层审计和负结果记录整合为一个验证治理协议，而不是新的预测主干网络。"
            "我们在九个终点上执行 3 外层×3 内层×5 重复的冻结选择实验，并以 reliability and chemical-boundary analyses（可靠性与化学边界分析）界定公开面板、逐样本可靠性和化学迁移边界。"
        ),
        (
            "在确认性候选池扩张实验中，K=32 相对 K=4 的完整池范围归一化选择损失增加 0.122（端点聚类 95% CI 0.072–0.175；精确 P=0.0078；Holm P=0.039），机会校正 Top-3 命中率下降 0.642。"
            "在共享冻结划分的 12 候选多视图实验中，validation-best 相对 Morgan-only 的实际兑现效用增益为 0.343（0.210–0.483；9/9 终点）。"
            "跨终点元风险仅作为探索性预警：严格留一端点验证得到高遗憾 AUC=0.648，保留预测风险最低 50% 单元时平均遗憾降低 0.034（95% CI 0.020–0.047），但九个终点不足以建立通用元选择器。"
        ),
        (
            "这些结果表明，候选池扩张在分子性质预测中同时带来更高上界和更大的选择损失。"
            "冻结治理能够使这一权衡可审计，并把可靠性与化学边界从平均性能叙事中分离出来；其结论不应外推到受授权限制的 TabPFN、九终点全量深度模型面板或时间外前瞻验证。"
        ),
    ]
    replace_first_containing(doc, "分子性质预测常在有限验证信息上", abstract_parts[0])
    insert_after_first_containing(doc, abstract_parts[0][:24], abstract_parts[1:])

    replace_first_containing(
        doc,
        "为回应名义候选数与有效候选数的区分",
        (
            "为区分名义候选数与有效候选数，本文增加候选有效多样性审计。"
            "32 候选轻量池的内层验证效用两两相关中位数为 0.998，特征值有效候选数 K_eff=1.01；外层测试效用相关中位数为 0.999，K_eff=1.01。"
            "因此，32 候选池应解释为高度相关轻量候选的受控选择压力实验，而不是代表 GNN、D-MPNN、Transformer 和融合模型共同扩张的真实异质池。"
        ),
    )

    replace_first_containing(
        doc,
        "真实异质性的确认由 12 候选多视图实验承担",
        (
            "九终点真实异质性效应由 12 候选多视图实验估计：Morgan-512、MACCS、RDKit2D 与拼接多视图分别配对线性模型、随机森林和 LightGBM，并在九个终点、相同外层折、相同内层折和相同种子上完成 6,480 次冻结重训。"
            "Chemprop/D-MPNN、GNN-GCN、ChemBERTa 和 MoLFormer 的同划分结果作为 ESOL、BACE 和 ClinTox 三终点强基线面板单独报告；因此，它们不并入九终点 12 候选主效应估计。"
        ),
    )

    replace_first_containing(
        doc,
        "为把相关实验做厚",
        (
            "为检验表示组合对选择收益的影响，本文将异质池拆分为 Morgan-only、Morgan+MACCS、Morgan+MACCS+RDKit2D 和完整 12 候选多视图四级阶梯。"
            "在 135 个任务-种子-外层折单位中，完整异质池相对 Morgan-only 的 G_attain=0.303，L_select=0.015，G_realized=0.296，中位 η=1.000；对应候选有效多样性、逐终点效应和完整候选登记均作为 source data 输出。"
        ),
    )

    replace_first_containing(
        doc,
        "3.4 共享冻结划分的多视图候选兑现可达到收益",
        "3.4 共享冻结划分确认多视图收益并审计强基线",
    )

    replace_first_containing(
        doc,
        "这一规则看似保守",
        (
            "历史探索中的 GNN、Chemprop、ChemBERTa、MoLFormer 和更复杂融合为候选来源提供背景，但只有在统一外层划分上重训并完成登记后才可进入确认性效应估计。"
            "这一约束较为保守，但可防止历史探索中表现最好的结果被事后选入主文。"
            "换言之，FZYC-Mol 的治理逻辑并不排斥复杂模型，而是要求复杂模型承担与轻量候选相同的审计成本。"
        ),
    )

    replace_first_containing(
        doc,
        "赢家频率",
        (
            "每个选择单元提取 16 个不读取外层标签的验证特征，包括池规模、任务类型、验证效用离散度、前两名标准化间隔、折间排名一致性、胜出频率、候选家族数、one-SE 大小和原风险分量。"
            "最外层完整留出一个终点；其余八个终点内部再次 LOEO，并按 MAE 在固定 Ridge、浅层随机森林和强正则 HistGradientBoosting 中选模。"
            "高遗憾阈值、模型类型和全部参数均由训练端点确定。"
        ),
    )

    replace_first_containing(
        doc,
        "规则本身没有跨终点普适赢家",
        (
            "治理消融表明，规则本身没有跨终点普适最优解。"
            "validation-best 在完整多视图池中表现较好，但 one-SE、低方差、低成本和 LOEO 规则都呈现条件性收益或失败。"
            "该结果使 FZYC-Mol 的定位更接近审计框架：它帮助研究者知道何时、为什么、以何种成本选择某个规则，而不是替研究者预设唯一最优规则。"
        ),
    )

    replace_first_containing(
        doc,
        "结果部分按证据梯度组织",
        (
            "结果部分按证据等级组织。"
            "前两节检验候选池扩张是否造成选择压力；第三节用正负对照和跨端点风险检验指标解释；第四节确认多视图候选在共享划分下是否兑现收益。"
            "随后各节依次给出公开面板、逐样本可靠性、化学迁移、治理消融和负结果库存，使主证据和边界证据在结构上保持区分。"
        ),
    )

    replace_first_containing(
        doc,
        "多视图确认实验补上了旧稿最容易被质疑的比较条件",
        (
            "多视图确认实验处理了候选比较中的关键可比性问题。"
            "所有候选在相同外层折、内层折和种子上重训，避免把历史表现最好的模型与新候选直接比较。"
            "结果显示，拼接多视图不仅提高了 test oracle 的可达到上界，也在 validation-best 选择下兑现了配对效用增益；因此，多视图收益不是单纯事后 oracle 现象。"
        ),
    )

    replace_first_containing(
        doc,
        "TDC 结果强调门控规则的终点依赖性",
        (
            "TDC 结果强调门控规则的终点依赖性。"
            "5 个 promoted 与 17 个 retained 并不意味着其余终点被证明无效，而是说明在三种子和当前区间宽度下，许多终点无法支持明确晋级。"
            "这种处理保留了负结果的审计价值，也避免把证据不足误写为反向证据。"
        ),
    )

    replace_first_containing(
        doc,
        "候选扩张不是简单的",
        (
            "候选扩张不能概括为单调收益或单调损害。"
            "在本文的轻量 32 候选池中，完整池 oracle 上界随 K 增加而上升，但验证选择未能等比例兑现这一上界，选择损失同步增加。"
            "新增有效多样性分析显示，32 个轻量候选的内层效用相关性中位数为 0.998，有效候选数约为 1.11，说明该实验主要检验候选数量和超参数自由度带来的选择压力，而不是声称覆盖真实药物发现中的全部模型异质性。"
            "因此，多视图共享划分实验被单独作为异质候选确认，而不是与轻量扩池混成一个总排行榜。"
        ),
    )

    replace_first_containing(
        doc,
        "同时补齐真嵌套验证",
        (
            "除主效应外，本文还报告真嵌套验证、种子敏感性、统一消融、80/90/95 保形覆盖率、精确 Tanimoto 分箱、MoleculeACE 活性悬崖、低相似度失败样本和扩展失败案例等 10 张补充实验表。"
            "强基线证据矩阵扩展为代表性同划分执行：ESOL、BACE 和 ClinTox 分别覆盖回归、常规分类和少数类毒性终点；GNN、RDKit-RF、ChemBERTa/MoLFormer 冻结适配头完成完整 3 外层×3 内层×5 种子的内外层评估，Chemprop/D-MPNN 完成相同外层 folds 和 seeds 下的外层确认性训练以及全部 3 个 inner folds 的确认性重训。"
            "TabPFN 已安装，但因 license/token 交互导致本地运行时不可用；因此，仅作为授权受限候选记录在状态表中，不作为完成性结果。"
        ),
    )

    insert_after_first_containing(
        doc,
        "完整多视图 test oracle 相对 Morgan-only oracle",
        [
            (
                f"现代强基线面板使用与确认性实验一致的 3×3×5 任务-种子-外层结构，合并后包含 {m['combined_inner_rows']} 条内层评分和 {m['combined_outer_rows']} 条外层评分。"
                f"在分类终点上，RDKit-RF 的平均 ROC-AUC 为 {m['rdkit_cls_auc']:.3f}，MoLFormer、ChemBERTa、Chemprop/D-MPNN 和 GNN-GCN 分别为 {m['molformer_cls_auc']:.3f}、{m['chemberta_cls_auc']:.3f}、{m['chemprop_cls_auc']:.3f} 和 {m['gnn_cls_auc']:.3f}。"
                f"在 ESOL 回归上，对应 RMSE 分别为 {m['rdkit_esol_rmse']:.3f}、{m['chemberta_esol_rmse']:.3f}、{m['molformer_esol_rmse']:.3f}、{m['chemprop_esol_rmse']:.3f} 和 {m['gnn_esol_rmse']:.3f}。"
                "这些结果不支持将模型复杂度预设为跨终点优势来源，而支持将复杂模型登记为需要同划分审计的候选来源。"
            ),
            (
                f"新增强基线选择分析显示，五候选池在 {m['selection_units']} 个 task-seed-fold 单元中由验证均值选择 RDKit-RF {m['rdkit_selected_units']} 次，外层 oracle top-1 命中率为 {m['selection_top1_hit_rate']:.3f}，平均 range-normalized selection loss 为 {m['selection_norm_loss']:.4f}，inner-outer Spearman 相关均值为 {m['selection_spearman']:.3f}。"
                f"以 RDKit-RF 为配对参照，MoLFormer、ChemBERTa、Chemprop/D-MPNN 和 GNN-GCN 的平均效用差分别为 {m['molformer_delta_vs_rdkit']:.3f}、{m['chemberta_delta_vs_rdkit']:.3f}、{m['chemprop_delta_vs_rdkit']:.3f} 和 {m['gnn_delta_vs_rdkit']:.3f}；MoLFormer 仅在 {m['molformer_win_rate_vs_rdkit']:.3f} 的外层单位超过 RDKit-RF。"
                "该分析将强基线比较从汇总性能扩展到验证选择、可达到上界和配对效应的共同审计。"
            ),
            (
                f"基于导出的逐样本预测，本文直接计算强基线 error-overlap。合并强基线预测池包含 {m['combined_prediction_rows']} 条测试样本-候选记录，覆盖 {len(m['combined_candidates'])} 个可运行候选和 {m['combined_overlap_pairs']} 个候选对。"
                f"平均 Jaccard error-overlap 为 {m['combined_overlap_mean']:.3f}；最高重叠来自 {m['max_overlap_pair']}（{m['max_overlap']:.3f}），最低重叠来自 {m['min_overlap_pair']}（{m['min_overlap']:.3f}）。"
                "因此，候选间错误并非相互独立，增加候选数带来的机会效应需要与错误重叠同时报告。"
            ),
            (
                f"三套去重敏感性实验也已从清洗审计推进到外层结果比较。全局去重、仅训练折内聚合和保留重复并按骨架分组三种策略共形成 {m['duplicate_rows']} 个外层评估单位。"
                f"最大绝对偏移出现在 {m['max_dup_task']} 的 {m['max_dup_policy']} 策略，外层效用相对全局去重改变 {m['max_dup_delta']:.3f}。"
                "BACE 在三种策略下保持不变，ESOL 偏移很小，而 ClinTox 对保留重复更敏感，说明少数类毒性任务应把重复处理作为边界条件报告。"
            ),
        ],
    )

    replace_first_containing(
        doc,
        "FZYC-Mol 不替代预测模型",
        (
            "FZYC-Mol 不替代预测模型，不保证性能提升，也不提供可迁移到所有终点的元选择器。本文已在代表性三终点面板中完成 GNN、ChemBERTa/MoLFormer 冻结适配头、RDKit-RF 和 Chemprop/D-MPNN 的完整 3×3×5 同划分评估，并新增验证选择、配对效应和 error-overlap 审计；"
            "但 TabPFN 因 license/token 交互限制未能产生预测，九终点全量深度模型面板也仍超出当前完成范围。"
            "因此，强基线结果应解释为代表性边界测试，而不能被解读为复杂模型已在所有终点上完成确认性统一重训。TDC 三种子结果只能反映种子变异和公开面板异质性，不能承担严格抽样推断。公开 release、Zenodo DOI 和第三方冷启动复跑仍需在正式投稿前完成。"
        ),
    )

    replace_all_literals(
        doc,
        {
            "完整 32 池完整池范围归一化选择损失": "32 候选完整池范围归一化选择损失",
            "历史最佳模型": "历史表现最好的模型",
            "这样做的目的不是声称所有外部复现障碍已经消失，而是保证本地分析级复现可以区分": "这样做的目的不是声称所有外部复现障碍已经消失，而是使本地分析级复现能够区分",
            "泄漏审计覆盖完全相同的 SMILES": "泄漏审计覆盖相同 SMILES",
            "接近最佳候选": "接近验证最优候选",
            "FZYC-Mol 不替代预测模型，不保证性能提升，也不提供可迁移到所有终点的元选择器。": "FZYC-Mol 不替代预测模型，也不必然带来性能提升，且不提供可迁移到所有终点的元选择器。",
            "本文的唯一主命题": "本文的中心主张",
            "在稿件中承担不同角色": "在本文中承担不同角色",
            "这样的分工避免把所有数据集混成一个泛化声明，也让每个实验的证据等级更加清楚。": "这种分工避免把所有数据集混成一个泛化声明，并使每个实验的证据等级更明确。",
            "加厚后的保形明细": "扩展后的保形明细",
            "单纯事后 oracle 现象": "单纯的事后 oracle 现象",
            "主文只报告": "主文仅报告",
            "放回常见分子性质预测语境": "置于常见分子性质预测语境",
            "真实项目前瞻验证": "实际项目前瞻验证",
            "等权 selection-risk 的负验证说明": "等权 selection-risk 的负验证表明",
            "10 张补充实验表": "10 个补充实验表",
        },
    )

    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.keep_with_next = True
    doc.save(DOCX)
    return m


def audit(m: dict[str, object]) -> None:
    doc = Document(DOCX)
    text = "\n".join(p.text for p in doc.paragraphs)
    fig_caps = [p.text.strip() for p in doc.paragraphs if re.match(r"^图\s*\d+\b", p.text.strip())]
    tab_caps = [p.text.strip() for p in doc.paragraphs if re.match(r"^表\s*\d+\s*\|", p.text.strip())]
    report = {
        "docx": str(DOCX),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "figure_caption_numbers": [int(re.search(r"\d+", x).group()) for x in fig_caps],
        "table_caption_numbers": [int(re.search(r"\d+", x).group()) for x in tab_caps],
        "has_title": TITLE in text,
        "has_strong_panel": "现代强基线面板使用与确认性实验一致的 3×3×5" in text,
        "has_error_overlap": "error-overlap" in text and "Jaccard" in text,
        "has_duplicate_sensitivity": "三套去重敏感性实验" in text,
        "has_tabpfn_boundary": "TabPFN" in text and "license/token" in text,
        "has_chemprop_inner_completed": "Chemprop/D-MPNN 完成" in text and "全部 3 个 inner folds" in text,
        "has_core_selection_addon": "range-normalized selection loss" in text and "配对效应" in text,
        "old_unfinished_deep_phrase": "Chemprop、GNN、ChemBERTa 和 MoLFormer 等强模型尚未在本文统一外层划分上完成确认性重训" in text,
        "old_chemprop_inner_unfinished_phrase": "Chemprop 尚未在全部三个 inner folds 上重复训练" in text,
        "xml_errors": xml_errors(DOCX),
        "metrics": m,
    }
    report["passed"] = all(
        [
            report["has_title"],
            report["has_strong_panel"],
            report["has_error_overlap"],
            report["has_duplicate_sensitivity"],
            report["has_tabpfn_boundary"],
            report["has_chemprop_inner_completed"],
            report["has_core_selection_addon"],
            not report["old_unfinished_deep_phrase"],
            not report["old_chemprop_inner_unfinished_phrase"],
            not report["xml_errors"],
        ]
    )
    AUDIT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    if not report["passed"]:
        raise SystemExit(json.dumps(report, ensure_ascii=False, indent=2))


def main() -> None:
    m = revise()
    audit(m)
    print(DOCX)
    print(AUDIT)


if __name__ == "__main__":
    main()
