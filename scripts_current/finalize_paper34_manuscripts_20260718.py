from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(r"D:\fzyc\output\paper34_submission_ready_20260718")
EN = ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_FINAL_CLEAN.docx"
ZH = next(path for path in ROOT.glob("*.docx") if not path.name.startswith("Candidate"))


def set_text(paragraph: Paragraph, text: str) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "SimSun")


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def normal_after(doc: Document, heading_prefix: str, ordinal: int = 1) -> Paragraph:
    heading = find_paragraph(doc, heading_prefix)
    seen = 0
    passed = False
    for paragraph in doc.paragraphs:
        if paragraph._p is heading._p:
            passed = True
            continue
        if passed and paragraph.style.name.startswith("Heading"):
            break
        if passed and paragraph.style.name == "Normal" and paragraph.text.strip():
            seen += 1
            if seen == ordinal:
                return paragraph
    raise RuntimeError(f"Normal paragraph {ordinal} after {heading_prefix} not found")


def insert_paragraph_before(element, parent, text: str, style: str) -> Paragraph:
    node = OxmlElement("w:p")
    element.addprevious(node)
    paragraph = Paragraph(node, parent)
    paragraph.style = style
    set_text(paragraph, text)
    return paragraph


def table_by_header(doc: Document, first: str, second: str | None = None) -> Table:
    for table in doc.tables:
        head = [cell.text.strip() for cell in table.rows[0].cells]
        if head and head[0] == first and (second is None or len(head) > 1 and second in head[1]):
            return table
    raise RuntimeError(f"Table not found: {first} / {second}")


def prevent_table_split(table: Table) -> None:
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = row_index < len(table.rows) - 1


def prepare_caption(paragraph: Paragraph, text: str) -> None:
    set_text(paragraph, text)
    paragraph.style = "Caption"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def place_table_before_heading(table: Table, caption: Paragraph, heading: Paragraph, note: Paragraph | None = None) -> None:
    heading._p.addprevious(caption._p)
    heading._p.addprevious(table._tbl)
    if note is not None:
        heading._p.addprevious(note._p)


def place_table_after_paragraph(table: Table, caption: Paragraph, paragraph: Paragraph) -> None:
    paragraph._p.addnext(table._tbl)
    paragraph._p.addnext(caption._p)


def add_notation_title(doc: Document, table: Table, title: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == title:
            paragraph._p.getparent().remove(paragraph._p)
    insert_paragraph_before(table._tbl, table._parent, title, "Heading 3")


def keep_figure_with_caption(doc: Document, caption_prefix: str) -> None:
    caption = find_paragraph(doc, caption_prefix)
    caption.paragraph_format.keep_together = True
    previous = caption._p.getprevious()
    while previous is not None and previous.tag != qn("w:p"):
        previous = previous.getprevious()
    if previous is not None:
        Paragraph(previous, caption._parent).paragraph_format.keep_with_next = True


def replace_after_heading(doc: Document, heading_prefix: str, replacements: list[str]) -> None:
    heading = find_paragraph(doc, heading_prefix)
    targets: list[Paragraph] = []
    passed = False
    for paragraph in doc.paragraphs:
        if paragraph._p is heading._p:
            passed = True
            continue
        if passed and paragraph.style.name.startswith("Heading"):
            break
        if passed and paragraph.style.name == "Normal" and paragraph.text.strip():
            targets.append(paragraph)
    if len(targets) < len(replacements):
        raise RuntimeError(f"Insufficient paragraphs after {heading_prefix}")
    for paragraph, text in zip(targets, replacements):
        set_text(paragraph, text)
    for paragraph in targets[len(replacements):]:
        paragraph._p.getparent().remove(paragraph._p)


def remove_second_conclusion_paragraph(doc: Document, heading_prefix: str) -> None:
    second = normal_after(doc, heading_prefix, 2)
    second._p.getparent().remove(second._p)


def format_all_captions(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.style.name in {"Caption", "Figure Caption"}:
            paragraph.paragraph_format.keep_together = True
            if paragraph.style.name == "Caption":
                paragraph.paragraph_format.keep_with_next = True


def revise_english(path: Path) -> dict:
    doc = Document(path)
    set_text(doc.paragraphs[4], "Results: The composition intervention covered six prespecified endpoints. Modern augmentation yielded positive validation-selected gains relative to the homogeneous registry across all six endpoints, while ranking and cross-fitted effects remained heterogeneous.")

    set_text(find_paragraph(doc, "Equations (1)"), "Equations (1)–(14) define the candidate decision, selection loss, chance-adjusted ranking, cross-fitted comparison, matrix transformations, effective ranks, normalized gains and selection stability used below.")
    p = find_paragraph(doc, "Here r_u")
    set_text(p, p.text.rstrip() + " Equations (1)–(4) define the candidate decision, selection loss and chance-adjusted ranking quantities.")
    p = find_paragraph(doc, "U_(-s)")
    set_text(p, p.text.rstrip() + " Equations (5)–(6) define the cross-fitted reference and comparison.")
    p = find_paragraph(doc, "X is the outer-utility matrix")
    set_text(p, p.text.rstrip() + " Equations (7)–(10) define the matrix transformations and effective ranks.")
    p = find_paragraph(doc, "The label hom denotes")
    set_text(p, p.text.rstrip() + " Equations (11)–(13) define the composition contrasts and normalized gains.")
    p = find_paragraph(doc, "Here π_j")
    set_text(p, p.text.rstrip() + " Equation (14) defines normalized selection entropy.")

    notation = table_by_header(doc, "Symbol")
    add_notation_title(doc, notation, "Mathematical notation used in the audit estimands")

    table1 = table_by_header(doc, "Endpoint", "n")
    table2 = table_by_header(doc, "Audit component")
    table3 = table_by_header(doc, "Endpoint", "Effect")
    cap1 = find_paragraph(doc, "Table 1.")
    cap2 = find_paragraph(doc, "Table 2.")
    cap3 = find_paragraph(doc, "Table 3.")
    prepare_caption(cap1, "Table 1. Primary datasets and endpoint metrics.")
    prepare_caption(cap2, "Table 2. Audit components and recorded computational exposure.")
    prepare_caption(cap3, "Table 3. Cross-fitted effects of candidate-pool expansion.")
    note = find_paragraph(doc, "Target units:")
    place_table_before_heading(table1, cap1, find_paragraph(doc, "2.3 Candidate"), note)
    place_table_before_heading(table2, cap2, find_paragraph(doc, "2.4 Repeated"))
    place_table_after_paragraph(table3, cap3, normal_after(doc, "3.4 Cross-fitted", 1))
    for table in (table1, table2, table3, notation):
        prevent_table_split(table)

    p = find_paragraph(doc, "At candidate correlation 0.9")
    set_text(p, p.text.replace("(Figure 4C-D)", "(Figure 4D)").replace("(Figure 4C–D)", "(Figure 4D)"))

    replace_after_heading(doc, "3.10 Expanded", [
        "Across the six prespecified endpoints and three registries, CAHit@3 decreased from K = 4 to K = 32 in most endpoint–pool combinations, whereas cross-fitted gaps remained heterogeneous. At K = 32, multiview and modern augmentation produced positive validation-selected gains relative to the homogeneous registry across all six endpoints; raw classification and regression utilities were not pooled. Complete counts are reported in Table S33 and Figure S16.",
        "Modern components were heterogeneous: frozen language-model representations, the separately locked one-epoch D-MPNN and the full modern registry did not contribute uniformly across endpoints or candidate-pool sizes. Component-level paired results are reported in Table S33 and Figure S17.",
        "Equal-budget truncation and selection-stability analyses showed that realised benefit depended on the locked prefix order and finite-validation alignment. Complete equal-K, equal-budget, selection-frequency and entropy results are reported in Table S36 and Figures S20–S21.",
        "Anchor and normalization choices changed numerical scale more than the overall direction pattern, and most prespecified directions transported across the two evaluated split mechanisms although magnitudes changed. Missing cells and complete sensitivity results are retained in Tables S34–S35 and Figures S18–S19.",
    ])

    set_text(find_paragraph(doc, "Figure 7."), "Figure 7. Expanded equal-size candidate-pool composition intervention. (A) At K = 32, horizontal segments connect paired homogeneous-normalized validation-selected gain (filled squares) and observed finite-audit opportunity (open circles) for six prespecified endpoints and three pool compositions; classification and regression occupy separate row blocks. (B) The K = 4, 8, 16 and 32 ladders show mean paired homogeneous-normalized selected gain and cross-fitted gap in vertically aligned panels, with classification and regression displayed in separate x-axis bands. (C) CAHit@3 is shown for every endpoint–pool–K cell using four candidate-pool-size columns; narrow strips identify H, MV and M, and selection entropy is reported in Figure S20. (D) Equal-K and equal-downstream-budget trajectories relate measured downstream audit time per outer unit to paired homogeneous-normalized selected gain; the step line marks the empirical Pareto frontier. Uncertainty used five seed blocks after averaging the three outer folds within seed. Downstream time excludes pretrained-encoder acquisition, pretraining and cached embedding extraction.")
    keep_figure_with_caption(doc, "Figure 7.")

    set_text(normal_after(doc, "5 Conclusions", 1), "Within the evaluated endpoints, registries and split mechanisms, candidate expansion altered finite-audit opportunity, validation-realised gain, ranking fidelity, selection stability and bounded downstream cost. Matrix-dependent effective diversity complemented nominal K by describing relative candidate-utility movement across audit units, but did not determine selection loss.")
    remove_second_conclusion_paragraph(doc, "5 Conclusions")

    if not any(p.text.startswith("Additional file 4.") for p in doc.paragraphs):
        declarations = find_paragraph(doc, "List of abbreviations")
        insert_paragraph_before(declarations._p, declarations._parent, "Additional file 4. Code and reproducibility package.", "Normal")
    set_text(find_paragraph(doc, "Public dataset provenance"), "Public dataset provenance is listed in Additional file 2, Table S1. Complete source code, data-download and cleaning scripts, pinned environment files, software versions, candidate registries, split manifests, fold/seed/candidate exports, figure and table source data, exclusion logs, and quick and training-level reproduction entry points are provided in Additional file 4. The corresponding package version is paper34-final-minor-revision-20260718 and is distributed under the MIT License. No public GitHub, Zenodo release or DOI is claimed; access is through the journal supplementary files.")
    format_all_captions(doc)
    doc.save(path)
    return {"tables": 4, "equations": 14, "results_3_10_paragraphs": 4}


def revise_chinese(path: Path) -> dict:
    doc = Document(path)
    set_text(doc.paragraphs[4], "结果：组成干预覆盖六个预设端点。现代增强相对同质候选池的验证选择增益在六个端点均为正，而排序与交叉拟合效应仍具有异质性。")

    set_text(find_paragraph(doc, "公式（1）"), "公式（1）–（14）依次定义下文使用的候选决策、选择损失、机会校正排序、交叉拟合对比、矩阵变换、有效秩、归一化增益和选择稳定性。")
    p = find_paragraph(doc, "其中，r_u")
    set_text(p, p.text.rstrip() + " 公式（1）–（4）定义候选决策、选择损失和机会校正排序量。")
    p = find_paragraph(doc, "U_(-s)")
    set_text(p, p.text.rstrip() + " 公式（5）–（6）定义交叉拟合参照与对比。")
    p = find_paragraph(doc, "X 为外层效用矩阵")
    set_text(p, p.text.rstrip() + " 公式（7）–（10）定义矩阵变换与有效秩。")
    p = find_paragraph(doc, "hom表示")
    set_text(p, p.text.rstrip() + " 公式（11）–（13）定义组成对比与归一化增益。")
    p = find_paragraph(doc, "此处 π_j")
    set_text(p, p.text.rstrip() + " 公式（14）定义标准化选择熵。")

    notation = table_by_header(doc, "符号")
    add_notation_title(doc, notation, "审计估计量使用的数学符号")
    table1 = table_by_header(doc, "终点", "样本量")
    table2 = table_by_header(doc, "审计组成")
    table3 = table_by_header(doc, "终点", "效应")
    cap1 = find_paragraph(doc, "表1.")
    cap2 = find_paragraph(doc, "表2.")
    cap3 = find_paragraph(doc, "表3.")
    prepare_caption(cap1, "表1. 主要数据集与终点指标。")
    prepare_caption(cap2, "表2. 审计组件与记录的计算暴露。")
    prepare_caption(cap3, "表3. 候选池扩张的交叉拟合效应。")
    note = find_paragraph(doc, "注：分类终点")
    place_table_before_heading(table1, cap1, find_paragraph(doc, "2.3 候选"), note)
    place_table_before_heading(table2, cap2, find_paragraph(doc, "2.4 重复"))
    place_table_after_paragraph(table3, cap3, normal_after(doc, "3.4 交叉拟合", 1))
    for table in (table1, table2, table3, notation):
        prevent_table_split(table)

    p = normal_after(doc, "3.5 有限审计", 1)
    set_text(p, p.text.replace("（图4C–D）", "（图4D）").replace("（图4C-D）", "（图4D）"))

    replace_after_heading(doc, "3.10 扩展", [
        "在六个预设端点和三类候选池中，CAHit@3从K = 4到K = 32在多数“端点×候选池”组合下降，而交叉拟合差距仍具有异质性。K = 32时，多表征与现代增强相对同质候选池的验证选择增益在六个端点均为正；分类与回归原始效用不合并。完整计数见表S33和图S16。",
        "现代组件具有异质性：冻结语言模型表征、另行锁定的一轮D-MPNN与完整现代候选池在不同端点和候选规模下并未产生一致贡献。组件层面的配对结果见表S33和图S17。",
        "等预算截断与选择稳定性分析表明，实际收益取决于锁定的候选前缀顺序及有限验证一致性。完整的等K、等预算、选择频率与熵结果见表S36和图S20–S21。",
        "锚点与归一化选择对数值尺度的影响大于总体方向模式；多数预设方向在两种已评估切分机制间保持迁移，但效应幅度发生改变。缺失单元与完整敏感性结果保留于表S34–S35和图S18–S19。",
    ])

    set_text(find_paragraph(doc, "图 7"), "图7 扩展后的等规模候选池组成干预。A，K = 32时，横线连接六个预设端点和三类候选池的配对同质归一化验证选择增益（实心方形）与有限审计机会（空心圆）；分类与回归分区排列。B，K = 4、8、16和32阶梯在上下对齐面板中展示配对同质归一化选择增益与交叉拟合差距，分类和回归位于独立横轴区段。C，四个候选规模列展示全部“端点×候选池×K”单元的CAHit@3；窄色条标识H、MV和M，选择熵移至图S20。D，等K与等下游预算轨迹展示单个外层单元的下游审计时间与配对同质归一化选择增益，阶梯线表示经验Pareto前沿。不确定性以五个种子为区组，并先在种子内平均三个外层折。下游时间不含预训练编码器获取、预训练与缓存嵌入提取。")
    keep_figure_with_caption(doc, "图7")

    set_text(normal_after(doc, "5 结论", 1), "在已评估的端点、候选池和切分机制中，候选扩张改变了有限审计机会、验证兑现增益、排序保真度、选择稳定性和受限下游成本。矩阵依赖的有效多样性补充了名义K，可描述候选效用在审计单元中的相对独立变化，但其本身不能决定选择损失。")
    remove_second_conclusion_paragraph(doc, "5 结论")

    set_text(find_paragraph(doc, "Additional file 1"), "Additional file 1：补充方法与结果。Additional file 2：机器可读补充表S1–S36。Additional file 3：补充图S1–S21。Additional file 4：代码与可复现性包。")
    set_text(find_paragraph(doc, "公开数据来源"), "公开数据来源见Additional file 2的表S1。完整源代码、数据下载与清洗脚本、锁定环境文件、软件版本、候选登记表、切分清单、逐折/逐种子/逐候选导出、图表源数据、排除日志以及快速复现和训练级复现入口均见Additional file 4。对应版本为paper34-final-minor-revision-20260718，并按MIT License分发。本文不声称GitHub、Zenodo公开版本或DOI已经完成；访问方式为期刊补充文件。")
    format_all_captions(doc)
    doc.save(path)
    return {"tables": 4, "equations": 14, "results_3_10_paragraphs": 4}


def main() -> None:
    report = {"english": revise_english(EN), "chinese": revise_chinese(ZH)}
    (ROOT / "Paper34_manuscript_revision_audit.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
