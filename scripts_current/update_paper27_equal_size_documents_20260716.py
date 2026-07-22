from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "paper26_split_regime_transfer_revision_20260716"
BASE = ROOT / "output" / "paper27_equal_size_registry_composition_20260716"
TARGET = ROOT / "output" / "paper27_equal_size_registry_composition_revision_20260716"
FIGURE = BASE / "figures" / "Figure_7_equal_size_registry_composition_600dpi.png"
EN_NAME = "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript.docx"
TRACKED_NAME = "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript_tracked_changes.docx"


EN_METHODS = (
    "Methods: We conducted a retrospective nested audit across nine public molecular-property endpoints using repeated seeded scaffold "
    "partitions. Candidate registries were evaluated through matrix-dependent diversity estimators, chance-adjusted ranking measures, "
    "calibration controls, cross-fitted references and matched-size analyses. A split-regime transfer audit retrained the 32-candidate "
    "registry under Tanimoto-component separation. An equal-size composition intervention then compared Morgan-only, classical multiview "
    "and modern-augmented registries at K = 16 and 32 on ClinTox, BACE and ESOL under identical nested folds and a shared anchor."
)
EN_RESULTS = (
    "Results: Effective diversity and chance-adjusted ranking remained matrix and endpoint dependent. In the split-regime audit, the "
    "K = 32 minus K = 4 CAHit@3 direction remained negative and the cross-fitted gap direction positive under both split regimes for all "
    "three endpoints, although magnitudes changed. At matched K = 32, classical multiview and modern-augmented pools increased both oracle "
    "opportunity and selected-model gain relative to the homogeneous Morgan pool in all three endpoints. Modern augmentation had the largest "
    "raw effective rank at every endpoint, but relative diversity was not monotonically associated with cross-fitted gap across the 18 "
    "endpoint-pool-K cells (Spearman rho = -0.211). ESOL further showed higher selected gain despite lower CAHit@3, separating ranking "
    "recovery from realised utility. Modern augmentation required substantially more downstream audit time and had lower gain per audit hour."
)
EN_CONCLUSIONS = (
    "Conclusions: Nominal K and matrix-relative effective diversity were both incomplete descriptions of molecular model-selection risk. "
    "Registry expansion was valuable when added candidates supplied complementary information that finite validation ranking could realise; "
    "diversity alone did not imply degradation. Molecular benchmarks should jointly report opportunity, realised gain, chance-adjusted "
    "ranking, cross-fitted gaps, registry composition, computational exposure, split sensitivity and chemical-support boundaries."
)
EN_CONTRIBUTION = (
    "Scientific Contribution: This study provides an auditable joint decomposition of nominal candidate-pool size, matrix-dependent utility-pattern "
    "diversity, chance-adjusted ranking, cross-fitted selection gaps and realised model-selection gain. A split-regime transfer audit separates "
    "direction transport from invariant effect size, while an equal-size registry-composition intervention shows that useful complementarity and "
    "selection alignment, rather than K or effective diversity alone, govern net benefit. The contribution concerns molecular model-selection "
    "practice and does not introduce a universal statistical law, a new predictor or a fully optimized modern-architecture benchmark."
)


def find(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise KeyError(prefix)


def set_text(paragraph, text: str) -> None:
    paragraph.text = text


def fit_inline_figures_to_page(doc: Document) -> None:
    """Keep embedded figures inside the narrowest printable section."""
    printable = min(section.page_width - section.left_margin - section.right_margin for section in doc.sections)
    target = min(printable, Inches(6.3))
    for shape in doc.inline_shapes:
        if shape.width > target:
            ratio = target / shape.width
            shape.width = target
            shape.height = int(shape.height * ratio)


def insert_after(paragraph, text: str = "", style: str | None = None):
    new = OxmlElement("w:p")
    paragraph._p.addnext(new)
    from docx.text.paragraph import Paragraph
    result = Paragraph(new, paragraph._parent)
    if style:
        try:
            result.style = style
        except KeyError:
            result.style = "Caption" if style == "Figure Caption" else "Normal"
    if text:
        result.add_run(text)
    return result


def insert_before(paragraph, text: str = "", style: str | None = None):
    new = OxmlElement("w:p")
    paragraph._p.addprevious(new)
    from docx.text.paragraph import Paragraph
    result = Paragraph(new, paragraph._parent)
    if style:
        try:
            result.style = style
        except KeyError:
            result.style = "Caption" if style == "Figure Caption" else "Normal"
    if text:
        result.add_run(text)
    return result


def insert_figure_after(paragraph, image: Path, caption: str):
    picture = insert_after(paragraph)
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(image), width=Inches(6.55))
    cap = insert_after(picture, caption, "Figure Caption")
    return cap


def update_english(path: Path) -> None:
    doc = Document(path)
    set_text(find(doc, "Methods:"), EN_METHODS)
    set_text(find(doc, "Results:"), EN_RESULTS)
    set_text(find(doc, "Conclusions:"), EN_CONCLUSIONS)
    set_text(find(doc, "Scientific Contribution:"), EN_CONTRIBUTION)
    set_text(find(doc, "Keywords:"), (
        "Keywords: molecular property prediction; candidate-pool expansion; registry composition; model selection; nested cross-validation; "
        "utility-pattern diversity; ranking fidelity; audit gap"
    ))
    set_text(find(doc, "Under a retrospectively locked"), (
        "Under a retrospectively locked repeated nested scaffold evaluation, how did candidate-pool expansion relate to matrix-dependent "
        "utility-pattern diversity, chance-adjusted ranking fidelity, cross-fitted selection gaps and representation-composition effects; did "
        "the principal audit directions persist under a stricter structure-separated split; and, at fixed K, did homogeneous, classical "
        "multiview and modern-augmented registries create different opportunity, realised gain and computational exposure? We address this "
        "limited question through calibration controls, split transfer and a matched-budget registry-composition intervention."
    ))

    stat = find(doc, "2.14 Statistical inference")
    insert_before(stat, "2.14 Equal-size registry-composition intervention", "Heading 2")
    insert_before(stat, (
        "ClinTox, BACE and ESOL were reused with seeds 11, 23, 37, 53 and 71 and the same three outer by three inner scaffold folds. "
        "Three registries had exact nested prefixes of K = 16 and 32 and shared a Morgan-512 linear anchor. The homogeneous registry added "
        "Morgan-only learner and tuning variants. The classical multiview registry balanced Morgan-512, MACCS, RDKit 2D descriptors and their "
        "concatenation across eight lightweight learner variants. The modern-augmented registry interleaved 16 classical candidates with frozen "
        "ChemBERTa-MTR, ChemBERTa-MLM and MoLFormer embeddings fitted by nested downstream heads and the locked one-epoch D-MPNN candidate. "
        "The modern panel was a composition intervention, not a fully optimized architecture ranking."
    ), "Normal")
    insert_before(stat, (
        "For every endpoint, seed, outer fold, pool and K, opportunity gain was the outer oracle utility minus the shared-anchor utility, and "
        "realised gain was the validation-selected utility minus the same anchor. We also calculated CAHit@3, leave-one-seed-out cross-fitted "
        "selection gaps, Ledoit-Wolf matrix-relative effective diversity under four transformations, complete downstream audit time and selected "
        "gain per audit hour. Encoder pretraining and cached embedding-extraction costs were unavailable and excluded; therefore cost comparisons "
        "refer only to observed nested downstream fitting and D-MPNN fit/predict time."
    ), "Normal")
    set_text(stat, "2.15 Statistical inference")

    discussion = find(doc, "4 Discussion")
    h = insert_before(discussion, "3.10 Equal-size composition changed opportunity, realisation and cost", "Heading 2")
    p = insert_after(h, (
        "At K = 32, classical multiview versus homogeneous Morgan pools increased mean oracle opportunity gain from 0.125 to 0.188 in "
        "ClinTox, 0.079 to 0.084 in BACE and 1.419 to 2.260 utility units in ESOL. The corresponding selected-model gains increased from "
        "0.103 to 0.172, 0.072 to 0.078 and 1.325 to 2.179. Modern-augmented values were similar or slightly higher for selected gain: "
        "0.176, 0.079 and 2.191. All K = 32 oracle- and selected-gain differences versus the homogeneous pool were positive in all five seed "
        "means for both composition interventions (Additional file 2: Tables S28 and S31)."
    ), "Normal")
    p = insert_after(p, (
        "Raw K = 32 entropy ranks were 10.88, 2.90 and 1.87 for homogeneous ClinTox, BACE and ESOL pools; 12.40, 6.09 and 9.07 for "
        "classical multiview pools; and 13.96, 7.92 and 11.47 for modern-augmented pools. Nevertheless, relative entropy rank was not "
        "monotonically associated with same-unit or cross-fitted selection gap across the 18 endpoint-pool-K cells (Spearman rho = -0.008 and "
        "-0.211). At K = 32 the modern cross-fitted gap was smaller than the homogeneous gap for ClinTox (0.0048 versus 0.0092) and ESOL "
        "(0.0105 versus 0.0578), but slightly larger for BACE (0.0035 versus 0.0025)."
    ), "Normal")
    p = insert_after(p, (
        "Ranking fidelity and realised utility also separated. ESOL CAHit@3 was 0.338, 0.044 and -0.030 for homogeneous, classical multiview "
        "and modern-augmented K = 32 pools, even though the latter two had substantially higher selected gain. Modern augmentation consumed "
        "139-187 seconds of downstream audit time per outer unit versus 18-41 seconds for classical multiview and 21-31 seconds for homogeneous "
        "pools. Its selected gain per audit hour was lower than the classical multiview value in all three endpoints, even before encoder "
        "pretraining cost was counted (Figure 7; Tables S28-S30)."
    ), "Normal")
    caption = (
        "Figure 7. Equal-size candidate-pool composition intervention. A, K = 32 oracle and validation-selected gains normalized to the "
        "homogeneous-pool oracle gain for display; raw endpoint-scale values are retained in Table S28. B, matrix-relative raw entropy rank "
        "against the leave-one-seed-out cross-fitted selection gap normalized to the homogeneous oracle gain. C, chance-adjusted Hit@3. D, "
        "downstream audit time against selected gain relative to the homogeneous pool. Filled and open markers denote K = 16 and K = 32; "
        "marker shape denotes endpoint. Pretraining and cached embedding-extraction costs are excluded."
    )
    insert_figure_after(p, FIGURE, caption)

    limitations = find(doc, "4.8 Limitations")
    h = insert_before(limitations, "4.8 Registry composition, not diversity alone, governed net value", "Heading 2")
    p = insert_after(h, (
        "Holding K and folds fixed directly addressed the concern that a Morgan-dominated registry might determine the primary result. Both "
        "classical and modern additions raised the attainable ceiling and usually allowed the validation selector to realise most of that gain. "
        "The absence of a positive monotonic diversity-gap association shows that effective rank is a structural audit descriptor, not a causal "
        "penalty parameter. Complementarity is useful only when inner validation ranking can identify it with sufficient stability."
    ), "Normal")
    insert_after(p, (
        "This result sharpens the expansion story. Nominal K measures eligibility, effective diversity describes independent utility movement, "
        "oracle gain measures available chemical opportunity, selected gain measures realisation and compute records its price. The classical "
        "multiview pool occupied the most favourable cost-benefit region here; frozen language-model representations and the lightweight D-MPNN "
        "expanded the modern boundary but were not a substitute for a compute-matched, fully tuned architecture benchmark."
    ), "Normal")
    set_text(limitations, "4.9 Limitations")
    limitation_text = find(doc, "The primary registry was intentionally")
    limitation_text.text += (
        " The equal-size intervention covered three endpoints, 18 endpoint-pool-K cells and one modern composition; its frozen encoders, "
        "lightweight downstream heads and one-epoch D-MPNN do not represent exhaustive contemporary architecture tuning. Observed cost excludes "
        "encoder pretraining and cached embedding extraction, and effective-rank associations at this cell count are descriptive."
    )

    set_text(find(doc, "Within the studied endpoints"), (
        "Within the studied endpoints, candidate-pool expansion was accompanied by weaker chance-adjusted ranking and heterogeneous selection "
        "gaps, and selected audit directions transferred conditionally to stricter structure separation. The equal-size intervention further "
        "showed that classical multiview and modern-augmented registries could increase both opportunity and realised utility without a monotonic "
        "increase in selection gap. Nominal K and effective diversity therefore described exposure and structure, but useful complementarity and "
        "validation alignment governed net benefit."
    ))
    set_text(find(doc, "Molecular model-selection studies should"), EN_CONCLUSIONS.replace("Conclusions: ", ""))
    set_text(find(doc, "Additional file 1."), "Additional file 1. Supplementary Methods and Results, including split-regime and equal-size composition audits.")
    set_text(find(doc, "Additional file 2."), "Additional file 2. Machine-readable Supplementary Tables (Tables S1-S31; XLSX).")
    set_text(find(doc, "Public dataset provenance"), (
        "Public dataset provenance is listed in Additional file 2: Table S1. Derived fold-level tables, split manifests, source hashes, "
        "split-regime and equal-size registry-composition outputs, timing records and analysis code are supplied in the submission package."
    ))
    fit_inline_figures_to_page(doc)
    doc.save(path)


def update_chinese(path: Path) -> None:
    doc = Document(path)
    set_text(find(doc, "方法："), "方法：采用重复嵌套骨架切分审计九个公开分子性质终点，并结合矩阵依赖有效多样性、机会校正排序、校准对照、交叉拟合参照和等规模分析。切分机制迁移审计在ClinTox、BACE和ESOL上采用Tanimoto连通分量切分。进一步在相同种子、折分、K=16/32和共享锚点下比较Morgan同质池、经典多表征池和现代增强池。")
    set_text(find(doc, "结果："), "结果：切分机制迁移中，三个终点的K=32减K=4 CAHit@3方向均保持为负，交叉拟合差距方向均为正，但效应大小改变。固定K=32时，经典多表征池和现代增强池在三个终点均提高oracle机会收益和实际选择收益。现代池在三个终点均具有最高原始有效秩，但18个终点-池-K单元中相对多样性与交叉拟合差距不呈单调关系（Spearman ρ=-0.211）。ESOL在实际收益提高的同时CAHit@3下降，说明排序恢复与实际效用并不等价。现代池下游审计时间明显增加，单位审计时间收益较低。")
    set_text(find(doc, "结论："), "结论：名义K和矩阵相对有效多样性均不足以单独描述分子模型选择风险。只有当新增候选提供可被有限验证排序识别的互补信息时，候选池扩张才转化为净收益。应联合报告机会收益、实际收益、机会校正排序、交叉拟合差距、注册表组成、计算暴露、切分敏感性和化学支持边界。")

    stat = find(doc, "2.14 统计推断")
    insert_before(stat, "2.14 等规模候选池组成干预", "Heading 2")
    insert_before(stat, "在ClinTox、BACE和ESOL上保持种子11、23、37、53、71及相同3×3嵌套骨架折。三类候选池均包含严格嵌套的K=16和K=32前缀，并共享Morgan-512线性锚点：同质池仅加入Morgan学习器与调参变体；经典多表征池平衡Morgan、MACCS、RDKit 2D描述符及其拼接表征；现代增强池将16个经典候选与冻结ChemBERTa-MTR、ChemBERTa-MLM、MoLFormer嵌入的嵌套下游头及锁定的一轮D-MPNN交错组合。现代池用于组成干预，不用于完整现代架构排名。", "Normal")
    insert_before(stat, "相对于共享锚点计算oracle机会收益和验证选择的实际收益，并报告CAHit@3、留一随机种子交叉拟合差距、四种矩阵变换下的Ledoit-Wolf相对有效秩、完整下游审计时间及单位审计小时收益。预训练和缓存嵌入提取成本未计入。", "Normal")
    set_text(stat, "2.15 统计推断")

    discussion = find(doc, "4 讨论")
    h = insert_before(discussion, "3.10 等规模组成改变机会、兑现与成本", "Heading 2")
    p = insert_after(h, "在K=32时，经典多表征池相对于同质Morgan池将ClinTox、BACE和ESOL的平均oracle机会收益分别从0.125、0.079和1.419提高至0.188、0.084和2.260；实际选择收益分别从0.103、0.072和1.325提高至0.172、0.078和2.179。现代增强池的实际收益分别为0.176、0.079和2.191。两种组成干预相对于同质池的K=32机会与实际收益差在五个随机种子均为正。", "Normal")
    p = insert_after(p, "K=32原始熵秩在同质池中分别为10.88、2.90和1.87，在经典多表征池中为12.40、6.09和9.07，在现代增强池中为13.96、7.92和11.47。然而18个终点-池-K单元的相对熵秩与同单元或交叉拟合选择差距的Spearman相关仅为-0.008和-0.211。ESOL现代池CAHit@3为-0.030，但实际收益高于同质池，进一步表明排序命中与效用兑现并不等价。", "Normal")
    p = insert_after(p, "现代池每个外层单元消耗139–187秒下游审计时间，经典多表征池为18–41秒，同质池为21–31秒；即使不计预训练成本，现代池在三个终点的单位审计小时收益均低于经典多表征池。", "Normal")
    caption = "图7. 等规模候选池组成干预。A：K=32的oracle机会收益与实际选择收益；B：矩阵相对有效多样性与交叉拟合选择差距；C：机会校正Hit@3；D：下游审计成本—收益前沿。实心和空心标记分别表示K=16和K=32，形状表示终点。图中跨终点收益进行了归一化，原始终点尺度数值见表S28。"
    insert_figure_after(p, FIGURE, caption)

    limitations = find(doc, "4.8 局限性")
    h = insert_before(limitations, "4.8 注册表组成而非多样性本身决定净收益", "Heading 2")
    p = insert_after(h, "固定K与折分后，经典和现代候选均提高了性能上限，并通常允许验证选择器兑现大部分收益。有效多样性与差距缺乏正向单调关系，说明有效秩是结构审计量而不是因果惩罚参数；互补性只有在内层排序能够稳定识别时才有价值。", "Normal")
    insert_after(p, "因此，候选资格数量、独立效用运动、可获得机会、实际兑现和计算价格需要分开报告。本研究中经典多表征池位于更有利的成本—收益区域；冻结语言模型嵌入和轻量D-MPNN扩展了现代模型边界，但不能替代计算预算匹配的完整现代架构基准。", "Normal")
    set_text(limitations, "4.9 局限性")
    lim = find(doc, "主要登记表有意包含")
    lim.text += " 等规模干预仅覆盖三个终点、18个终点-池-K单元和一种现代组成；冻结编码器、轻量下游头及一轮D-MPNN不代表完整现代架构调优。成本未包括编码器预训练和缓存嵌入提取，有效秩关联仅作描述性解释。"
    set_text(find(doc, "在9个分子性质审计中"), "在本研究终点中，候选池扩张伴随机会校正排序减弱和异质选择差距，且部分审计方向可条件性迁移至更严格结构分离。等规模组成干预进一步表明，经典多表征和现代增强池能够同时提高机会与实际效用，而不必单调增加选择差距。名义K和有效多样性描述暴露与结构，但有用互补性及其与验证排序的一致性决定净收益。")
    set_text(find(doc, "分子模型选择研究应"), "分子模型选择研究应联合报告候选资格、名义K、矩阵依赖有效多样性、机会与实际收益、机会校正排序、交叉拟合差距、注册表组成、计算暴露、切分敏感性和化学支持边界。")
    set_text(find(doc, "Additional file 1："), "Additional file 1：补充方法与结果（含切分机制和等规模组成审计）；Additional file 2：机器可读补充表S1–S31；Additional file 3：补充图S1–S18。")
    fit_inline_figures_to_page(doc)
    doc.save(path)


def add_sheet(workbook, name: str, frame: pd.DataFrame) -> None:
    if name in workbook.sheetnames:
        del workbook[name]
    ws = workbook.create_sheet(name)
    ws.append(list(frame.columns))
    for row in frame.itertuples(index=False, name=None):
        ws.append([None if pd.isna(value) else value for value in row])
    fill = PatternFill("solid", fgColor="D9EAF2")
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for column in ws.columns:
        letter = column[0].column_letter
        width = min(max(max(len(str(cell.value or "")) for cell in column) + 2, 10), 34)
        ws.column_dimensions[letter].width = width


def update_supplementary_workbook(path: Path) -> None:
    wb = load_workbook(path)
    add_sheet(wb, "S27 Registry composition", pd.read_csv(BASE / "equal_size_candidate_registry.csv"))
    add_sheet(wb, "S28 Endpoint summary", pd.read_csv(BASE / "equal_size_endpoint_summary.csv"))
    add_sheet(wb, "S29 Effective diversity", pd.read_csv(BASE / "equal_size_effective_diversity.csv"))
    add_sheet(wb, "S30 Selection units", pd.read_csv(BASE / "equal_size_selection_units.csv"))
    add_sheet(wb, "S31 Paired pool effects", pd.read_csv(BASE / "equal_size_paired_pool_effects.csv"))
    wb.save(path)


def update_supplementary_doc(path: Path) -> None:
    doc = Document(path)
    limitations = find(doc, "S15 Source limitations")
    h = insert_before(limitations, "S15 Equal-size registry-composition intervention", "Heading 2")
    p = insert_after(h, (
        "Three exact K = 16/32 registries shared one Morgan-linear anchor and the same ClinTox, BACE and ESOL nested scaffold units. "
        "The homogeneous registry used Morgan-only learner/tuning variants; the classical multiview registry balanced four representations "
        "across eight learner variants; and the modern-augmented registry interleaved classical candidates with frozen ChemBERTa-MTR, "
        "ChemBERTa-MLM and MoLFormer downstream heads plus the locked one-epoch D-MPNN. Candidate identities and ordering are in Table S27."
    ), "Normal")
    insert_after(p, (
        "Tables S28-S31 retain endpoint-scale opportunity and realised gains, CAHit@3, same-unit and leave-one-seed-out cross-fitted gaps, "
        "four effective-rank transformations, complete selection units, observed downstream time and paired seed-direction counts. Encoder "
        "pretraining and cached embedding extraction were excluded. The 18-cell diversity-gap association was descriptive and did not support "
        "a monotonic penalty interpretation."
    ), "Normal")
    set_text(limitations, "S16 Source limitations")
    directory = find(doc, "Table S26.")
    p = insert_after(directory, "Table S27. Equal-size candidate registry, ordering, representation and source.", "Normal")
    p = insert_after(p, "Table S28. Endpoint-level opportunity, realised gain, ranking, cross-fitted gap and downstream cost summary.", "Normal")
    p = insert_after(p, "Table S29. Matrix-dependent effective diversity for all endpoint-pool-K combinations.", "Normal")
    p = insert_after(p, "Table S30. Complete outer-unit equal-size selection audit.", "Normal")
    insert_after(p, "Table S31. Paired composition effects relative to the homogeneous Morgan registry.", "Normal")
    doc.save(path)


def update_reviewer_response(path: Path) -> None:
    doc = Document(path)
    table = doc.tables[0]
    cells = table.add_row().cells
    texts = [
        "The 32-candidate registry is dominated by Morgan-based classical learners, so the conclusions may reflect one homogeneous registry rather than contemporary molecular machine learning.",
        "Added a fully nested equal-size registry-composition intervention at K=16/32 on ClinTox, BACE and ESOL. Morgan-only, classical multiview and modern-augmented pools share the same anchor, seeds, folds and candidate count. Frozen ChemBERTa/MoLFormer heads and the locked D-MPNN expand the architecture boundary. Both heterogeneous pools increased oracle and selected gain, but effective diversity was not monotonically related to selection gap and modern gains had substantially higher downstream cost. Claims are therefore framed around useful complementarity and realisation, not diversity as a universal penalty.",
        "Methods 2.14; Results 3.10; Discussion 4.8-4.9; Figure 7; Additional file 1 S15-S16; Tables S27-S31."
    ]
    for cell, text in zip(cells, texts):
        cell.text = text
    doc.save(path)


def main() -> None:
    if TARGET.exists():
        raise FileExistsError(TARGET)
    shutil.copytree(SOURCE, TARGET, ignore=shutil.ignore_patterns("tracked_unpacked", "rendered"))
    (TARGET / "main_figures").mkdir(exist_ok=True)
    for suffix in (".pdf", ".svg"):
        shutil.copy2(BASE / "figures" / f"Figure_7_equal_size_registry_composition{suffix}", TARGET / "main_figures")
    shutil.copy2(FIGURE, TARGET / "main_figures")

    update_english(TARGET / EN_NAME)
    chinese = next(path for path in TARGET.glob("*.docx") if any(ord(char) > 127 for char in path.name))
    update_chinese(chinese)
    supplementary = TARGET / "supplementary"
    old_workbook = supplementary / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S26.xlsx"
    new_workbook = supplementary / "Additional_file_2_Machine_readable_Supplementary_Tables_S1-S31.xlsx"
    shutil.copy2(old_workbook, new_workbook)
    update_supplementary_workbook(new_workbook)
    update_supplementary_doc(supplementary / "Additional_file_1_Supplementary_Methods_and_Results.docx")
    update_reviewer_response(TARGET / "Reviewer_concern_Response_Location.docx")
    for name, text in {
        "New_Abstract.txt": "\n".join([EN_METHODS, EN_RESULTS, EN_CONCLUSIONS]),
        "Revised_Abstract.txt": "\n".join([EN_METHODS, EN_RESULTS, EN_CONCLUSIONS]),
        "New_Scientific_Contribution.txt": EN_CONTRIBUTION,
        "Scientific_Contribution.txt": EN_CONTRIBUTION,
    }.items():
        (TARGET / name).write_text(text, encoding="utf-8")
    print(TARGET)


if __name__ == "__main__":
    main()
