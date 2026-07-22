from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "paper15_revision_audit"
OUT.mkdir(parents=True, exist_ok=True)

REVIEW_DOCX = Path(
    r"C:\Users\Administrator\Documents\WeChat Files\wxid_561hsyfdwabz22\FileStorage\File\2026-06\FZYC-Mol_修改意见.docx"
)
MANUSCRIPT = ROOT / "output" / "\u5c0f\u8bba\u6587-14.docx"

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
}


def para_text(p) -> str:
    return "".join(run.text for run in p.runs).strip()


def docx_visible_text(path: Path) -> dict:
    doc = Document(path)
    paras = []
    for i, p in enumerate(doc.paragraphs, start=1):
        text = para_text(p)
        if text:
            paras.append({"index": i, "style": p.style.name if p.style else "", "text": text})
    tables = []
    for ti, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"index": ti, "rows": rows})
    return {"paragraphs": paras, "tables": tables}


def xml_texts(path: Path, member: str) -> list[str]:
    try:
        with zipfile.ZipFile(path) as z:
            data = z.read(member)
    except KeyError:
        return []
    root = ET.fromstring(data)
    texts = []
    for node in root.iter():
        if node.tag in {
            f"{{{NS['w']}}}t",
            f"{{{NS['w']}}}delText",
            f"{{{NS['w']}}}instrText",
        }:
            if node.text:
                texts.append(node.text)
    return texts


def comments(path: Path) -> list[str]:
    try:
        with zipfile.ZipFile(path) as z:
            data = z.read("word/comments.xml")
    except KeyError:
        return []
    root = ET.fromstring(data)
    out = []
    for c in root.findall("w:comment", NS):
        cid = c.attrib.get(f"{{{NS['w']}}}id", "")
        author = c.attrib.get(f"{{{NS['w']}}}author", "")
        texts = [n.text for n in c.iter(f"{{{NS['w']}}}t") if n.text]
        out.append(f"[comment {cid} {author}] {''.join(texts).strip()}")
    return out


def drawing_count(path: Path) -> int:
    try:
        with zipfile.ZipFile(path) as z:
            data = z.read("word/document.xml")
    except KeyError:
        return 0
    return data.count(b"<a:blip")


def extract_references(text: str) -> list[str]:
    lines = [ln.strip() for ln in text.splitlines()]
    start = None
    for i, line in enumerate(lines):
        if line.lower() in {"references", "reference"} or line.startswith("References"):
            start = i + 1
    if start is None:
        return []
    refs = []
    buff = ""
    for line in lines[start:]:
        if not line:
            continue
        if re.match(r"^\d+[\.\)]\s+", line) and buff:
            refs.append(buff.strip())
            buff = line
        else:
            buff = (buff + " " + line).strip()
    if buff:
        refs.append(buff.strip())
    return refs


def main() -> None:
    review = docx_visible_text(REVIEW_DOCX)
    manuscript = docx_visible_text(MANUSCRIPT)
    review_all_text = "\n".join(p["text"] for p in review["paragraphs"])
    manuscript_all_text = "\n".join(p["text"] for p in manuscript["paragraphs"])

    for name, obj in [("review_comments", review), ("manuscript14", manuscript)]:
        (OUT / f"{name}.json").write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")
        text_lines = []
        for p in obj["paragraphs"]:
            text_lines.append(f"P{p['index']:04d} [{p['style']}] {p['text']}")
        for t in obj["tables"]:
            text_lines.append(f"\n[TABLE {t['index']}]")
            for row in t["rows"]:
                text_lines.append(" | ".join(row))
        (OUT / f"{name}.txt").write_text("\n".join(text_lines), encoding="utf-8")

    review_comments = comments(REVIEW_DOCX)
    manuscript_comments = comments(MANUSCRIPT)
    (OUT / "docx_comments.json").write_text(
        json.dumps({"review_docx": review_comments, "manuscript": manuscript_comments}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    summary = {
        "review_docx": str(REVIEW_DOCX),
        "manuscript": str(MANUSCRIPT),
        "review_paragraphs": len(review["paragraphs"]),
        "review_tables": len(review["tables"]),
        "review_comments": len(review_comments),
        "manuscript_paragraphs": len(manuscript["paragraphs"]),
        "manuscript_tables": len(manuscript["tables"]),
        "manuscript_figures": drawing_count(MANUSCRIPT),
        "references_detected": len(extract_references(manuscript_all_text)),
        "review_xml_text_chars": len("".join(xml_texts(REVIEW_DOCX, "word/document.xml"))),
        "manuscript_xml_text_chars": len("".join(xml_texts(MANUSCRIPT, "word/document.xml"))),
    }
    (OUT / "extraction_summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
