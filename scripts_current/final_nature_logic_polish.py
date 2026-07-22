from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def set_text(paragraph, text: str) -> None:
    """Replace paragraph text while preserving its paragraph style."""
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.add_run(text)
    return inserted


def main() -> None:
    desktop = Path.home() / "Desktop"
    source = max(
        desktop.glob("FZYC-Mol_初稿-8_Nature全稿语言逻辑修订.docx"),
        key=lambda p: p.stat().st_mtime,
    )
    output = desktop / "FZYC-Mol_初稿-8_Nature全稿语言逻辑终审版.docx"
    doc = Document(str(source))

    replaced = []

    for p in doc.paragraphs:
        text = p.text.strip()

        if text.startswith("分类任务保留 ROC-AUC 作为 MoleculeNet"):
            set_text(
                p,
                "分类任务保留 ROC-AUC 作为 MoleculeNet 和若干 TDC 终点的主指标，同时报告 PR-AUC、Brier 分数、ECE、MCC、平衡准确率、风险-覆盖曲线和保形覆盖率。该指标组合用于降低类别不平衡任务中的解释风险：ROC-AUC 较高并不必然意味着阳性样本召回、概率校准或筛选富集可靠。ClinTox、DILI、hERG 和 CYP 底物任务因此在不平衡分类增强中重点报告。",
            )
            replaced.append("classification_metric_rationale")

        elif (
            "MoleculeNet 时代" in text
            and "ADMET-AI" in text
            and "非常有竞争力" in text
        ):
            set_text(
                p,
                "与 MoleculeNet 时代的标准基准研究相比，FZYC-Mol 不只报告固定模型在固定数据集上的分数，而是强调多专家策略选择、外部 ADMET 评估、结构外推和可靠性诊断。与 ADMET-AI 或 ADMETlab 这类平台型工作相比，本文不以部署大规模预测服务为目标，而是更关注公开基准下的实验透明度、仅基于验证集的选择和可解释证据链。与近期基础模型评测相比，本文的重点不是证明大模型无效，而是说明在许多 ADMET 场景中，强表格专家、指纹/描述符、轻量集成和严格验证治理仍可构成具有竞争力且可解释的基线。",
            )
            replaced.append("discussion_positioning_baseline")

        elif text.startswith("这种定位也与近年高水平论文"):
            set_text(
                p,
                "这种定位也与近年高水平论文的证据组织方式相一致。SCAGE、MotiL 和 MSformer-ADMET 都把结构单元、基序或片段作为理解模型行为的重要入口 [43-45]；多通道层级表示、HimNet 和多模态融合研究则从不同角度说明，原子、基序、指纹、构象和全局分子信息之间存在互补性 [16,46,50]。这些工作通常依赖新的预训练任务、复杂图结构或更高算力投入。FZYC-Mol 的差异在于把类似思想落实为低成本、同划分、可审计的候选池治理：层级基序、冻结表征或多方法融合只有在通过验证集门控后才进入最终策略；证据不足的候选则作为负结果保留。这一处理方式使本文既能吸收最新模型设计的启发，也能避免把未经验证的复杂模块直接写入主结论。",
            )
            replaced.append("discussion_recent_work")

        elif text.startswith("此外，OmniMol"):
            set_text(
                p,
                "此外，OmniMol、ADMETlab 3.0 和近期 ADMET 表征基准提示，真实 ADMET 建模往往同时面对不完整标注、终点间相关性、平台化决策支持和跨数据集迁移问题 [47-49]。这为 FZYC-Mol 提供了两个自然扩展方向：一是把当前 TDC full-panel appendix 扩展为更系统的多任务外部附录，报告属性间相关性、缺失标注处理和候选选择频率；二是把可靠性输出组织为面向药物化学用户的决策层，包括适用域、校准、风险分位、基序解释和相似邻域证据。因此，FZYC-Mol 的价值不只在于判定哪个模型分数最高，更在于说明候选策略在什么边界内可用、为什么可用以及何时应被拒绝。",
            )
            replaced.append("discussion_omnimol")

        elif "写成单一新模型" in text:
            set_text(
                p,
                "这一定位也决定了本文的比较对象与证据组织方式。FZYC-Mol 的主要主张不是引入一个孤立的新模型，而是提出一套由验证集治理的可靠性框架；核心竞争力因此转向实验设计、可复现选择、终点异质性、适用域诊断和解释性闭环。图 1、图 7、图 8、表 5、表 7、表 8 和表 9 均服务于这一主线。",
            )
            replaced.append("discussion_not_single_model")

        elif text.startswith("针对审稿人点名的 ClinTox"):
            set_text(
                p,
                "为补充 ClinTox 在高精度筛选条件下的召回证据，本文在与表 S17 一致的 consensus_strict_core_multifp fast 预测上计算 recall at fixed precision：在 precision≥0.80 时，五个 seed 的召回率为 0.370 ± 0.132；在更严格的 precision≥0.90 时，召回率为 0.232 ± 0.144。因此，ClinTox 的结论不只依赖 ROC-AUC 或 PR-AUC，还同时报告了高精度阈值下可保留的阳性检出能力。",
            )
            replaced.append("clintox_fixed_precision")

        elif text.startswith("本研究仍存在明确局限。验证-测试排名审计"):
            delete_paragraph(p)
            replaced.append("floating_limitations_removed")

        elif text.startswith("因此，FZYC-Mol 更适合作为一种"):
            set_text(
                p,
                "总体而言，FZYC-Mol 更适合作为一种由验证集治理的可靠性框架，而不是单一追求最高分的排行榜方案。它在现有算力范围内提供了更完整的证据链，包括性能、方差、结构外推、适用域、高误差富集、校准、活性悬崖、外部 ADMET 附录和化学解释。这种证据链使论文主张更稳健，也更贴近药物发现用户对模型可信度的实际需求。",
            )
            replaced.append("discussion_transition")

        elif text.startswith("这种定位也决定了主文表格的角色"):
            set_text(
                p,
                "因此，主文表格仅保留能够直接支撑结论的摘要证据，完整候选级和 seed 级结果则置于补充材料。固定选择器的主文摘要表展示正向提升，同时正文保留 22 个负向终点-指标单元并指向完整审计表。该写法降低了选择偏差风险，也使负结果成为方法边界的一部分。",
            )
            replaced.append("discussion_table_role")

        elif text.startswith("案例层面的解释性分析仍需扩展。当前稿件"):
            set_text(
                p,
                "案例层面的解释性分析仍需扩展。本研究已纳入 Lipophilicity 补救、ClinTox 高风险假阴性和高粗糙度 ADME 回归案例；后续可扩展 BACE、clearance_microsome_az 和 CYP 底物终点，并在同一样本级审计表中连接结构片段、最近邻、标签差异、不确定性和基序归因。",
            )
            replaced.append("case_explanation_scope")

    # Ensure the limitations section starts with limitations before future work.
    paras = list(doc.paragraphs)
    for idx, p in enumerate(paras):
        if p.text.strip() == "5.3 局限与进一步验证":
            following = [q.text.strip() for q in paras[idx + 1 : idx + 4]]
            if not any("本研究仍存在" in q for q in following):
                insert_after(
                    p,
                    "本研究仍存在若干局限。首先，验证-测试排名审计和 9 个代表性终点的 3×3 nested validation 表明，验证集治理可以减少测试集事后选择，但内外层验证尚未覆盖全部 MoleculeNet 与 TDC 终点，不能保证所有任务均达到测试最优；因此，小幅增益应与 regret、optimism gap、Top-3 命中和负结果共同解释。其次，收益具有明显终点异质性，BBBP、ClinTox、HIA 和 Pgp 等终点的增益较小，FreeSolv 仍落后于观测最佳 Chemprop 候选。第三，基序归因和片段富集虽已补充 support、effect size、p 值和 FDR q 值，但仍属于关联证据，不能替代因果机制或湿实验验证。第四，ChemBERTa 与 MoLFormer 主要以冻结编码器形式使用，Polaris 与 OpenADMET 的完整官方挑战流程仍需进一步扩展。",
                )
                replaced.append("limitations_inserted")
            break

    doc.save(str(output))
    print("source=", source)
    print("output=", output)
    print("replaced=", ",".join(replaced))


if __name__ == "__main__":
    main()
