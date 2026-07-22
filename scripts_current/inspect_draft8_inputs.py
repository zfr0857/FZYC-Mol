# -*- coding: utf-8 -*-
from __future__ import annotations

from collections import Counter
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "draft8_inspection"
REPORT_DIR.mkdir(parents=True, exist_ok=True)

CHECKLIST = Path(r"C:\Users\Administrator\Downloads\FZYC-Mol_Journal_of_Cheminformatics_修改清单.docx")
MANUSCRIPT = Path(r"C:\Users\Administrator\Desktop\修改\初稿-7.docx")


def inspect_docx(path: Path, output_name: str) -> None:
    doc = Document(path)
    lines: list[str] = [
        f"# {path.name}",
        "",
        f"- path: `{path}`",
        f"- paragraphs: {len(doc.paragraphs)}",
        f"- tables: {len(doc.tables)}",
        f"- inline_shapes: {len(doc.inline_shapes)}",
        f"- sections: {len(doc.sections)}",
        "",
        "## Style summary",
    ]
    counts = Counter((p.style.name if p.style else "<none>") for p in doc.paragraphs)
    for style, count in counts.most_common():
        lines.append(f"- {style}: {count}")

    lines.extend(["", "## Paragraphs"])
    for i, p in enumerate(doc.paragraphs):
        text = p.text.replace("\t", " <TAB> ").replace("\n", " / ").strip()
        style = p.style.name if p.style else "<none>"
        lines.append(f"P{i:04d} [{style}] {text}")

    lines.extend(["", "## Tables"])
    for ti, table in enumerate(doc.tables):
        lines.append("")
        lines.append(f"### Table {ti} ({len(table.rows)} x {len(table.columns)})")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.replace("\n", " / ").replace("\t", " ").strip() for cell in row.cells]
            lines.append(f"R{ri:03d}: " + " | ".join(cells))

    lines.extend(["", "## Headings"])
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ""
        if "Heading" in style or "标题" in style or style in {"Title", "Subtitle"}:
            lines.append(f"P{i:04d} [{style}] {p.text.strip()}")

    (REPORT_DIR / output_name).write_text("\n".join(lines), encoding="utf-8-sig")


def main() -> None:
    for path in (CHECKLIST, MANUSCRIPT):
        if not path.exists():
            raise FileNotFoundError(path)
    inspect_docx(CHECKLIST, "checklist.md")
    inspect_docx(MANUSCRIPT, "manuscript.md")
    print(REPORT_DIR)


if __name__ == "__main__":
    main()
