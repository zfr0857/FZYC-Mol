from __future__ import annotations

import json
import shutil
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "output" / "paper30_submission_package_20260717"
DATA = ROOT / "output" / "paper31_expanded_intervention_20260717"
PACKAGE = ROOT / "output" / "paper31_submission_package_20260717"
EN_BASE = BASE / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript(6).docx"
ZH_BASE = BASE / "候选池扩张与模型选择损失_中文完整论文.docx"
EN_OUT = PACKAGE / "Candidate_pool_expansion_Journal_of_Cheminformatics_manuscript(7).docx"
ZH_OUT = PACKAGE / "候选池扩张与模型选择损失_中文完整论文(7).docx"


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_paragraph(doc: Document, predicate):
    for paragraph in doc.paragraphs:
        if predicate(paragraph):
            return paragraph
    raise ValueError("Required paragraph was not found")


def replace_section(doc: Document, start_prefix: str, end_prefix: str, heading: str, blocks: list[tuple[str, str]]) -> None:
    start = find_paragraph(doc, lambda p: p.text.strip().startswith(start_prefix))
    end = find_paragraph(doc, lambda p: p.text.strip().startswith(end_prefix))
    paragraphs = doc.paragraphs
    start_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is start._p)
    end_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is end._p)
    for paragraph in paragraphs[start_index + 1:end_index]:
        remove_paragraph(paragraph)
    start.text = heading
    available_styles = {style.name for style in doc.styles}
    for style, text in blocks:
        resolved_style = style
        if resolved_style not in available_styles:
            resolved_style = "Caption" if style == "Figure Caption" and "Caption" in available_styles else "Normal"
        inserted = end.insert_paragraph_before(text, style=resolved_style)
        if text == "[[FIGURE_7_VECTOR]]":
            inserted.alignment = WD_ALIGN_PARAGRAPH.CENTER


def replace_exact_prefix(doc: Document, prefix: str, text: str) -> None:
    paragraph = find_paragraph(doc, lambda p: p.text.strip().startswith(prefix))
    paragraph.text = text


def equation_blocks(language: str) -> list[tuple[str, str]]:
    if language == "en":
        groups = [
            ("Selection and ranking estimands", range(1, 9)),
            ("Cross-fitted reference", range(9, 12)),
            ("Matrix transformations and effective rank", range(12, 20)),
            ("Composition, normalization and downstream cost", range(20, 26)),
            ("Selection stability", range(26, 28)),
        ]
        group_notes = {
            "Selection and ranking estimands": "Here r_u is the inner-validation rank of the outer finite-audit best candidate, q is a rank cutoff (q = 3), I is the indicator function and H_K is the Kth harmonic number.",
            "Cross-fitted reference": "U_(-s) contains all outer units from seeds other than held-out seed s; S = 5 and F = 3 are the numbers of repeat seeds and outer folds.",
            "Matrix transformations and effective rank": "X is the outer-utility matrix; S in Equation 16 is the sample covariance (not the seed count), T is the scaled-identity target, lambda_i are non-negative eigenvalues of the shrinkage correlation matrix and p_i are their unit-sum proportions.",
            "Composition, normalization and downstream cost": "The label hom denotes the homogeneous Morgan pool, and T_down is measured inner-plus-outer downstream fit/predict wall time.",
            "Selection stability": "Here q_j is the proportion of repeated outer audit units in which candidate j is selected; it is distinct from the rank cutoff q in Equation 5.",
        }
        intro = (
            "Let u = (s, f) index an outer audit unit formed by repeat seed s and outer fold f; "
            "V(u,j) is the mean inner-validation utility and A(u,j) is outer-audit utility for candidate j. "
            "The eligible registry prefix C(p,K), candidate order, anchor a and reference j0 were fixed before outer outcomes were inspected. "
            "All unit-level quantities were averaged over folds within seed before five-seed block aggregation."
        )
        tail = (
            "We used epsilon = 1e-12. Ledoit–Wolf shrinkage used the scaled-identity target and the analytic coefficient implemented in scikit-learn. "
            "The shrinkage covariance was rescaled to a correlation matrix R before eigenanalysis. "
            "Correlation eigenvalues below zero only because of floating-point error were clipped to zero, and 0 log 0 was defined as zero. "
            "Equations 22 and 24 divide within paired outer units before fold-to-seed-to-endpoint aggregation; units with an absolute denominator at or below epsilon were reported as missing."
        )
    else:
        groups = [
            ("选择与排序估计量", range(1, 9)),
            ("交叉拟合参照", range(9, 12)),
            ("矩阵变换与有效秩", range(12, 20)),
            ("候选池组成、归一化与下游成本", range(20, 26)),
            ("选择稳定性", range(26, 28)),
        ]
        group_notes = {
            "选择与排序估计量": "其中，r_u 为外层有限审计最佳候选在内层验证排序中的名次，q 为排序截断值（q = 3），I 为示性函数，H_K 为第 K 个调和数。",
            "交叉拟合参照": "U_(-s) 表示除留出种子 s 外的所有外层单元；S = 5 与 F = 3 分别为重复种子数和每个种子的外层折数。",
            "矩阵变换与有效秩": "X 为外层效用矩阵；公式（16）中的 S 为样本协方差（不是种子数），T 为缩放单位阵目标；lambda_i 为收缩相关矩阵的非负特征值，p_i 为其和为 1 的比例。",
            "候选池组成、归一化与下游成本": "hom 表示同质 Morgan 候选池，T_down 为内外层下游拟合与预测的实测墙钟时间。",
            "选择稳定性": "此处 q_j 为候选 j 在重复外层审计单元中被选中的比例，与公式（5）的排序截断 q 不同。",
        }
        intro = (
            "令 u = (s, f) 表示由重复种子 s 与外层折 f 构成的外层审计单元；V(u,j) 为候选 j 的平均内层验证效用，A(u,j) 为其外层审计效用。"
            "候选池前缀 C(p,K)、候选顺序、锚点 a 与参照候选 j0 均在查看外层结果前固定。所有单元层统计量先在种子内对外层折求平均，再以 5 个种子为区组汇总。"
        )
        tail = (
            "数值稳定常数 epsilon 取 1e-12。Ledoit–Wolf 收缩采用缩放单位阵目标及 scikit-learn 的解析收缩系数估计；"
            "在特征分解前，将收缩协方差重新缩放为相关矩阵 R；"
            "仅因浮点误差产生的微小负特征值裁剪为 0，并约定 0 log 0 = 0。公式（22）和（24）先在配对外层审计单元内相除，再按折—种子—端点汇总；"
            "绝对分母不超过 epsilon 的单元记为缺失并单独报告。"
        )
    blocks: list[tuple[str, str]] = [("Normal", intro)]
    for title, numbers in groups:
        blocks.append(("Heading 3", title))
        blocks.append(("Normal", group_notes[title]))
        for number in numbers:
            blocks.append(("Normal", f"[[EQ_{number}]]"))
    blocks.append(("Normal", tail))
    return blocks


def metrics() -> dict[str, object]:
    units = pd.read_csv(DATA / "Paper31_selection_units.csv")
    summary = pd.read_csv(DATA / "Paper31_endpoint_pool_K_summary.csv")
    ablation = pd.read_csv(DATA / "Paper31_component_ablation_summary.csv")
    budget = pd.read_csv(DATA / "Paper31_equal_budget_units.csv")
    stability = pd.read_csv(DATA / "Paper31_selection_stability.csv")
    direction = pd.read_csv(
        DATA / "composition_split_loop" / "Paper31_composition_split_direction_concordance.csv"
    )
    anchor_direction = pd.read_csv(DATA / "Paper31_anchor_normalization_direction_concordance.csv")
    primary = summary.loc[
        summary.design.eq("equal_K") & summary.anchor_scheme.eq("shared_morgan_linear")
    ]
    k4 = primary.loc[primary.candidate_count.eq(4)].set_index(["pool", "task"])
    k32 = primary.loc[primary.candidate_count.eq(32)].set_index(["pool", "task"])
    cahit_delta = k32.chance_adjusted_hit3_mean - k4.chance_adjusted_hit3_mean
    gap_delta = k32.cross_fitted_selection_gap_mean - k4.cross_fitted_selection_gap_mean

    hom32 = k32.loc["Homogeneous Morgan"]
    modern32 = k32.loc["Modern-augmented"]
    classical32 = k32.loc["Classical multiview"]
    modern_vs_hom = modern32.selected_model_gain_mean - hom32.selected_model_gain_mean
    classical_vs_hom = classical32.selected_model_gain_mean - hom32.selected_model_gain_mean

    base = ablation.loc[ablation.pool.eq("Classical multiview")].set_index(["task", "candidate_count"])
    ablation_counts = {}
    for pool in ["+ChemBERTa", "+MoLFormer", "+D-MPNN", "Full modern-augmented"]:
        comp = ablation.loc[ablation.pool.eq(pool)].set_index(["task", "candidate_count"])
        delta = comp.selected_model_gain_mean - base.selected_model_gain_mean
        ablation_counts[pool] = int((delta > 0).sum())

    equal_k_units = units.loc[
        units.design.eq("equal_K") & units.anchor_scheme.eq("shared_morgan_linear")
        & units.candidate_count.isin([16, 32])
    ].set_index(["pool", "task", "seed", "outer_fold", "candidate_count"])
    budget_indexed = budget.set_index(["pool", "task", "seed", "outer_fold", "candidate_count"])
    aligned = budget_indexed.join(
        equal_k_units[["selected_model_gain"]].rename(columns={"selected_model_gain": "equal_k_gain"}),
        how="left",
    )
    budget_delta = aligned.selected_model_gain - aligned.equal_k_gain
    budget_endpoint = budget_delta.groupby(["pool", "task", "candidate_count"]).mean()

    primary_units = units.loc[
        units.design.eq("equal_K") & units.anchor_scheme.eq("shared_morgan_linear")
        & units.candidate_count.eq(32)
    ]
    denom_unique = primary_units[[
        "task", "seed", "outer_fold", "paired_homogeneous_best_gain", "homogeneous_denominator_valid"
    ]].drop_duplicates(["task", "seed", "outer_fold"])

    st32 = stability.loc[stability.candidate_count.eq(32)]
    return {
        "cahit_negative_cells": int((cahit_delta < 0).sum()),
        "gap_positive_cells": int((gap_delta > 0).sum()),
        "total_pool_endpoint_cells": int(len(cahit_delta)),
        "modern_vs_hom_positive": int((modern_vs_hom > 0).sum()),
        "classical_vs_hom_positive": int((classical_vs_hom > 0).sum()),
        "endpoint_count": 6,
        "ablation_counts": ablation_counts,
        "ablation_total": int(len(base)),
        "budget_nonnegative_cells": int((budget_endpoint >= 0).sum()),
        "budget_total_cells": int(len(budget_endpoint)),
        "modern_budget_mean_candidates": float(
            budget.loc[budget.pool.eq("Modern-augmented")].eligible_candidate_count.mean()
        ),
        "split_same_direction": int(direction.same_direction.sum()),
        "split_total": int(len(direction)),
        "median_entropy_k32": {
            pool: float(st32.loc[st32.pool.eq(pool)].candidate_selection_entropy_normalized.median())
            for pool in ["Homogeneous Morgan", "Classical multiview", "Modern-augmented"]
        },
        "median_seed_agreement_k32": {
            pool: float(st32.loc[st32.pool.eq(pool)].leave_one_seed_out_agreement.median())
            for pool in ["Homogeneous Morgan", "Classical multiview", "Modern-augmented"]
        },
        "invalid_denominators_k32": int((~denom_unique.homogeneous_denominator_valid).sum()),
        "denominator_units_k32": int(len(denom_unique)),
        "minimum_positive_denominator_k32": float(
            denom_unique.loc[denom_unique.homogeneous_denominator_valid, "paired_homogeneous_best_gain"].abs().min()
        ),
        "anchor_direction_min": float(anchor_direction.direction_concordance.min()),
        "anchor_direction_max": float(anchor_direction.direction_concordance.max()),
        "anchor_invalid_summary_cells": int(anchor_direction.invalid_summary_cells.sum()),
    }


def english_blocks(m: dict[str, object]) -> dict[str, list[tuple[str, str]]]:
    ab = m["ablation_counts"]
    methods = [
        ("Normal", (
            "Six endpoints were fixed from task coverage before expanded outcomes were read: ClinTox, BACE and BBBP for classification, and ESOL, "
            "Lipophilicity and Caco2_Wang for regression. Three exact-prefix registries—homogeneous Morgan, classical multiview and modern-augmented—were evaluated at K = 4, 8, 16 and 32 with seeds 11, 23, 37, 53 and 71 and identical three-outer by three-inner folds. "
            "The shared Morgan linear candidate occupied the first locked position. Candidate configuration and order did not change across eligible prefixes."
        )),
        ("Normal", (
            "Fixed-K component ablations compared classical multiview, classical plus ChemBERTa, classical plus MoLFormer, classical plus D-MPNN and the full modern-augmented registry at K = 16 and 32. "
            "Each addition replaced a prespecified classical candidate so that K remained exact. Equal-downstream-budget analyses truncated each locked registry prefix at the endpoint- and K-specific median downstream time of the classical multiview registry; outer performance never determined retention. "
            "Observed downstream time includes inner and outer fit/predict time but excludes encoder pretraining, model acquisition and cached embedding extraction."
        )),
        ("Normal", (
            "Anchor sensitivity used the shared Morgan linear model, a fixed Morgan random forest and the predefined registry-median candidate. Raw gains, endpoint-MAD normalization and paired homogeneous-audit-best normalization were reported. "
            "Selection stability comprised candidate, representation and learner-family frequencies, selection entropy, modal proportion, leave-one-seed-out agreement and adjacent-fold switches. "
            "ClinTox, BACE and ESOL additionally completed the full three-composition loop under seeded scaffold and Tanimoto-component splits while holding registries, seeds, folds, rules and metrics fixed."
        )),
    ]
    results = [
        ("Normal", (
            f"The expanded intervention completed six endpoints, three candidate-pool compositions and four locked K values. From K = 4 to K = 32, CAHit@3 decreased in {m['cahit_negative_cells']} of {m['total_pool_endpoint_cells']} endpoint–pool cells, whereas the cross-fitted gap increased in {m['gap_positive_cells']} cells. "
            f"At K = 32, classical multiview and modern augmentation increased validation-selected gain relative to the homogeneous registry in {m['classical_vs_hom_positive']} and {m['modern_vs_hom_positive']} of six endpoints, respectively. "
            "These counts describe direction across prespecified endpoints; classification ROC-AUC utility and regression negative-RMSE utility were not pooled on their raw scales."
        )),
        ("Normal", (
            f"At fixed K = 16 and 32, selected gain exceeded the corresponding classical multiview value in {ab['+ChemBERTa']}, {ab['+MoLFormer']}, {ab['+D-MPNN']} and {ab['Full modern-augmented']} of {m['ablation_total']} endpoint–K cells for the ChemBERTa, MoLFormer, D-MPNN and full-modern registries. "
            "The component results therefore separated architecture identity from the effect of merely enlarging K (Figure S17; Table S33)."
        )),
        ("Normal", (
            f"Under the locked classical-time budget, equal-budget selected gain was at least the equal-K value in {m['budget_nonnegative_cells']} of {m['budget_total_cells']} endpoint–pool–K cells. "
            f"Because D-MPNN occupied the second modern position and dominated downstream time, the modern registry retained a mean of {m['modern_budget_mean_candidates']:.2f} candidates under this strict prefix rule. "
            "This result is a downstream-compute sensitivity, not an end-to-end efficiency comparison (Figure S21; Tables S33 and S36)."
        )),
        ("Normal", (
            f"At K = 32, median normalized selection entropy was {m['median_entropy_k32']['Homogeneous Morgan']:.3f}, {m['median_entropy_k32']['Classical multiview']:.3f} and {m['median_entropy_k32']['Modern-augmented']:.3f} for homogeneous, classical multiview and modern-augmented registries; corresponding median leave-one-seed-out agreement was {m['median_seed_agreement_k32']['Homogeneous Morgan']:.3f}, {m['median_seed_agreement_k32']['Classical multiview']:.3f} and {m['median_seed_agreement_k32']['Modern-augmented']:.3f}. "
            f"Paired homogeneous normalization had {m['invalid_denominators_k32']} near-zero denominators among {m['denominator_units_k32']} endpoint outer units at K = 32; valid denominators were no smaller than {m['minimum_positive_denominator_k32']:.4g}. "
            f"Across the nine predefined anchor–normalization combinations, selected-gain direction concordance with the primary raw-scale reference ranged from {m['anchor_direction_min']:.3f} to {m['anchor_direction_max']:.3f}; {m['anchor_invalid_summary_cells']} undefined summary cells were retained as missing. "
            "Raw, endpoint-MAD and alternative-anchor results are reported in Figure S18 and Table S34."
        )),
        ("Normal", (
            f"Across ClinTox, BACE and ESOL, {m['split_same_direction']} of {m['split_total']} prespecified composition-by-K metric contrasts had the same direction under scaffold and Tanimoto-component splits. "
            "Magnitudes and several individual contrasts changed, so this analysis supports composition-effect transport across two evaluated split regimes rather than external validation or universal robustness (Figure S19; Table S35)."
        )),
        ("Normal", "[[FIGURE_7_VECTOR]]"),
        ("Figure Caption", (
            "Figure 7. Expanded equal-size candidate-pool composition intervention. (A) At K = 32, horizontal segments connect paired homogeneous-normalized validation-selected gain (filled squares) and observed finite-audit opportunity (open circles) for six prespecified endpoints and three pool compositions; classification and regression occupy separate row blocks. "
            "(B) The K = 4, 8, 16 and 32 ladders show mean paired homogeneous-normalized selected gain (solid) and cross-fitted gap (dashed), with classification and regression displayed in separate x-axis bands. "
            "(C) CAHit@3 is shown for every endpoint–pool–K cell; the right strip reports normalized selection entropy at K = 32. Negative CAHit@3 values are retained. "
            "(D) Equal-K and equal-downstream-budget trajectories relate measured downstream audit time per outer unit to paired homogeneous-normalized selected gain; the step line marks the empirical Pareto frontier. "
            "Uncertainty used five seed blocks after averaging the three outer folds within seed. Downstream time excludes pretrained-encoder acquisition, pretraining and cached embedding extraction."
        )),
    ]
    discussion = [
        ("Normal", (
            "The six-endpoint intervention changes the interpretation from a three-dataset boundary check to a prespecified composition experiment spanning rare-class, regular classification, small regression, large regression and permeability settings. "
            "Expansion raised finite audited opportunity in many cells, but the realised gain depended on whether inner validation could rank the added candidates consistently. Effective diversity, ranking fidelity and selected utility therefore describe different parts of the audit."
        )),
        ("Normal", (
            "The component ablations show why the modern panel should not be treated as one indivisible category. Frozen chemical-language-model embeddings and the one-epoch D-MPNN changed opportunity, selection stability and downstream time differently across endpoints. "
            "The strict equal-budget prefix analysis further shows that a computationally expensive candidate can consume the budget before later complementary candidates become eligible. This is a property of the prespecified order and bounded downstream accounting, not evidence that any architecture is intrinsically inefficient."
        )),
        ("Normal", (
            "Anchor and normalization sensitivity constrains the strongest claim. Paired homogeneous normalization is interpretable only when its finite-audit opportunity denominator is away from zero; near-zero units were not regularized into large ratios. "
            "The raw and endpoint-MAD results, together with the fixed Morgan-RF and registry-median anchors, distinguish directional conclusions from scale-specific presentation."
        )),
    ]
    limitations = [
        ("Normal", (
            "The primary nine-endpoint audit and the expanded six-endpoint composition intervention remain retrospective evaluations of public data. Effective rank at each endpoint used 15 outer units, so shrinkage stabilizes but does not create independent information. "
            "The split-regime composition loop covered only ClinTox, BACE and ESOL and one Morgan-derived Tanimoto threshold; it was not temporal, target-aware or prospective external validation."
        )),
        ("Normal", (
            "The modern candidates were frozen ChemBERTa and MoLFormer representations with lightweight nested heads plus a locked one-epoch D-MPNN. They do not form an exhaustive or compute-optimized modern molecular-learning benchmark. "
            "Equal-budget results depend on measured hardware-specific downstream time and the prespecified prefix order, while encoder pretraining and acquisition cost were excluded. Paired homogeneous normalization was undefined when the homogeneous finite-audit opportunity was at most 1e-12; all such units are disclosed rather than imputed."
        )),
        ("Normal", (
            "Finally, component-ablation and endpoint–pool–K cells reuse endpoints and folds and are not independent experiments. Inference therefore retained seed blocks and endpoint-specific reporting; cell counts and descriptive associations were not promoted to confirmatory population-level tests."
        )),
    ]
    conclusions = [
        ("Normal", (
            "Across six prespecified molecular-property endpoints, candidate-pool expansion simultaneously changed finite audited opportunity, validation-realised gain, ranking fidelity, selection stability and downstream cost. Effective diversity rather than nominal K better described independent utility movement, but diversity alone did not determine selection loss."
        )),
        ("Normal", (
            "The strengthened audit supports a bounded practical conclusion: registry expansion is useful when added candidates contribute complementary chemical information that finite inner validation can identify at an acceptable downstream budget. "
            "Robust reporting should therefore combine locked composition and K ladders, chance-adjusted ranking, cross-fitted gaps, matrix-relative diversity, selection frequencies, anchor/normalization sensitivity, compute-matched analysis and split-mechanism sensitivity."
        )),
    ]
    return {"methods": methods, "results": results, "discussion": discussion, "limitations": limitations, "conclusions": conclusions}


def chinese_blocks(m: dict[str, object]) -> dict[str, list[tuple[str, str]]]:
    ab = m["ablation_counts"]
    methods = [
        ("Normal", "在读取扩展结果前，依据任务覆盖预先固定 6 个端点：分类任务为 ClinTox、BACE 和 BBBP，回归任务为 ESOL、Lipophilicity 与 Caco2_Wang。三类精确前缀候选池——同质 Morgan、多表征经典与现代增强——均在 K = 4、8、16、32 下，使用种子 11、23、37、53、71 及完全一致的 3 外层×3 内层折进行评估。共享 Morgan 线性模型固定在首位；候选配置及顺序不随 K 改变。"),
        ("Normal", "固定 K 的组件消融在 K = 16 和 32 下比较多表征经典、经典+ChemBERTa、经典+MoLFormer、经典+D-MPNN 与完整现代增强候选池。每次加入现代组件时均用预先规定的经典候选进行替换，保持 K 严格相同。等下游预算分析按端点与 K 使用多表征经典候选池的中位下游时间阈值，并依锁定顺序截断；外层表现从不参与保留决策。下游时间仅包含内外层拟合与预测，不包含编码器预训练、模型获取及缓存嵌入提取。"),
        ("Normal", "锚点敏感性预先设定共享 Morgan 线性模型、固定 Morgan 随机森林及候选表中位位置模型，并分别报告原始增益、端点 MAD 归一化及配对同质审计最优归一化。选择稳定性包括候选、表征与学习器家族频率、选择熵、众数比例、留一种子一致性和相邻外层折切换次数。ClinTox、BACE 与 ESOL 进一步在 scaffold split 与 Tanimoto 连通分量 split 下完成三类候选池闭环，候选表、种子、折、选择规则和指标均保持不变。"),
    ]
    results = [
        ("Normal", f"扩展干预完整覆盖 6 个端点、3 类候选池与 4 个锁定 K。从 K = 4 到 K = 32，CAHit@3 在 {m['cahit_negative_cells']}/{m['total_pool_endpoint_cells']} 个“端点×候选池”单元下降，交叉拟合差距在 {m['gap_positive_cells']} 个单元增加。K = 32 时，多表征经典与现代增强相对同质候选池的验证选择增益分别在 {m['classical_vs_hom_positive']} 和 {m['modern_vs_hom_positive']} 个端点为正。上述计数仅描述预设端点中的方向；分类 ROC-AUC 效用与回归负 RMSE 效用未在原始尺度上合并。"),
        ("Normal", f"在固定 K = 16 和 32 下，ChemBERTa、MoLFormer、D-MPNN 与完整现代增强候选池的选择增益分别在 {ab['+ChemBERTa']}、{ab['+MoLFormer']}、{ab['+D-MPNN']} 和 {ab['Full modern-augmented']}/{m['ablation_total']} 个“端点×K”单元超过相应多表征经典候选池，从而将具体组件作用与单纯增加 K 区分开（图 S17；表 S33）。"),
        ("Normal", f"在锁定的经典候选池时间预算下，等预算选择增益在 {m['budget_nonnegative_cells']}/{m['budget_total_cells']} 个“端点×候选池×K”单元不低于等 K 结果。由于 D-MPNN 位于现代候选表第二位且占据主要下游时间，严格前缀规则下现代候选池平均仅保留 {m['modern_budget_mean_candidates']:.2f} 个候选。该结果只表示下游计算敏感性，不能解释为端到端效率比较（图 S21；表 S33、S36）。"),
        ("Normal", f"K = 32 时，同质 Morgan、多表征经典和现代增强候选池的标准化选择熵中位数分别为 {m['median_entropy_k32']['Homogeneous Morgan']:.3f}、{m['median_entropy_k32']['Classical multiview']:.3f} 和 {m['median_entropy_k32']['Modern-augmented']:.3f}；对应留一种子一致性中位数为 {m['median_seed_agreement_k32']['Homogeneous Morgan']:.3f}、{m['median_seed_agreement_k32']['Classical multiview']:.3f} 和 {m['median_seed_agreement_k32']['Modern-augmented']:.3f}。K = 32 的 {m['denominator_units_k32']} 个端点外层单元中，有 {m['invalid_denominators_k32']} 个同质归一化分母接近 0；有效分母最小为 {m['minimum_positive_denominator_k32']:.4g}。九种预设“锚点×归一化”组合相对主原始尺度参照的选择增益方向一致性介于 {m['anchor_direction_min']:.3f}–{m['anchor_direction_max']:.3f}，并保留 {m['anchor_invalid_summary_cells']} 个无定义汇总单元为缺失。完整结果见图 S18 和表 S34。"),
        ("Normal", f"在 ClinTox、BACE 和 ESOL 中，{m['split_same_direction']}/{m['split_total']} 个预设“组成×K”指标对比在 scaffold split 与 Tanimoto 连通分量 split 下方向一致；但效应大小及若干单独对比发生变化。因此，该分析仅支持两种已评估切分机制之间的组成效应迁移，不构成外部验证或普适鲁棒性证据（图 S19；表 S35）。"),
        ("Normal", "[[FIGURE_7_VECTOR]]"),
        ("Figure Caption", "图 7｜扩展后的等规模候选池组成干预。A，K = 32 时，横线连接 6 个预设端点与 3 类候选池的配对同质归一化验证选择增益（实心方形）和有限审计机会（空心圆）；分类与回归分区排列。B，K = 4、8、16、32 阶梯显示配对同质归一化选择增益（实线）和交叉拟合差距（虚线），分类与回归位于独立横轴区段。C，热图显示全部“端点×候选池×K”单元的 CAHit@3，右侧窄列为 K = 32 的标准化选择熵；负 CAHit@3 保留。D，等 K 与等下游预算轨迹展示单个外层单元的下游审计时间与配对同质归一化选择增益，阶梯线表示经验 Pareto 前沿。不确定性以 5 个种子为区组，并先在种子内平均 3 个外层折。下游时间不含预训练编码器获取、预训练与缓存嵌入提取。"),
    ]
    discussion = [
        ("Normal", "六端点干预使证据从三数据集边界检查扩展为覆盖稀有类别、常规分类、小样本回归、大样本回归与渗透性任务的预设组成实验。候选池扩张在许多单元提高有限审计机会，但能否形成实际选择增益，取决于内层验证能否稳定识别新增候选。有效多样性、排序保真度与实际选择效用因此对应审计链条中的不同环节。"),
        ("Normal", "组件消融表明，现代候选池不能被视为不可分割的单一类别。冻结的化学语言模型表征与一轮 D-MPNN 对不同端点的机会、选择稳定性和下游时间影响不同。严格等预算前缀分析还显示，一个计算昂贵的早期候选可能在后续互补候选进入前耗尽预算。这是预设顺序与受限下游成本核算的结果，并不证明某种架构具有内在低效率。"),
        ("Normal", "锚点与归一化敏感性进一步限制了最强结论：只有当有限审计机会分母远离 0 时，配对同质归一化才具有稳定解释。本文未将近零分母正则化为巨大比值，而是将其记为缺失；原始尺度、端点 MAD、固定 Morgan-RF 与候选表中位锚点共同区分方向性结论与特定尺度展示。"),
    ]
    limitations = [
        ("Normal", "主审计的 9 个端点与扩展组成干预的 6 个端点仍属于公共数据上的回顾性评估。每个端点的有效秩仅基于 15 个外层单元；收缩估计能够提高数值稳定性，但不能创造新的独立信息。组成×切分闭环仅覆盖 ClinTox、BACE、ESOL 与一个 Morgan-Tanimoto 阈值，不能替代时间切分、靶点感知切分或前瞻性外部验证。"),
        ("Normal", "现代候选包括冻结 ChemBERTa/MoLFormer 表征的轻量嵌套预测头及锁定的一轮 D-MPNN，并非穷尽或计算充分优化的现代分子学习基准。等预算结论依赖当前硬件上的下游时间与预设前缀顺序，且不包含编码器预训练和获取成本。配对同质归一化在有限审计机会不超过 1e-12 时无定义，所有此类单元均公开而未插补。"),
        ("Normal", "组件消融及“端点×候选池×K”单元重复使用相同端点与折，不能视为相互独立的实验。因此，推断始终保留种子区组与端点层报告；单元计数和描述性关联未提升为总体确认性检验。"),
    ]
    conclusions = [
        ("Normal", "在 6 个预设分子性质端点中，候选池扩张同时改变有限审计机会、验证可实现增益、排序保真度、选择稳定性与下游成本。相较名义 K，有效多样性更接近独立效用变化，但多样性本身不能决定选择损失。"),
        ("Normal", "强化后的审计支持一个边界明确的实践结论：只有当新增候选提供有限内层验证能够识别的互补化学信息，且下游预算可接受时，候选池扩张才具有净价值。分子机器学习基准应联合报告锁定的组成与 K 阶梯、机会校正排序、交叉拟合差距、矩阵相对多样性、选择频率、锚点/归一化敏感性、计算匹配分析及切分机制敏感性。"),
    ]
    return {"methods": methods, "results": results, "discussion": discussion, "limitations": limitations, "conclusions": conclusions}


def update_document(source: Path, destination: Path, language: str, m: dict[str, object]) -> None:
    shutil.copy2(source, destination)
    doc = Document(destination)
    blocks = english_blocks(m) if language == "en" else chinese_blocks(m)
    if language == "en":
        replace_exact_prefix(doc, "Methods:", (
            "Methods: We conducted a retrospective nested audit across nine public endpoints and a prespecified six-endpoint composition intervention. "
            "Three locked registries were evaluated at K = 4, 8, 16 and 32 with repeated 3×3 nested scaffold folds. Component ablations, equal-downstream-budget truncation, three predefined anchors, three normalizations, selection-stability measures and a three-endpoint Tanimoto-component loop were added."
        ))
        replace_exact_prefix(doc, "Results:", (
            f"Results: Across 18 endpoint–pool cells, CAHit@3 decreased from K = 4 to 32 in {m['cahit_negative_cells']} cells and the cross-fitted gap increased in {m['gap_positive_cells']}. "
            f"At K = 32, modern augmentation increased validation-selected gain relative to the homogeneous registry in {m['modern_vs_hom_positive']} of six endpoints. Component, budget, anchor and split-regime analyses showed that useful complementarity, selection stability and downstream exposure—not nominal K alone—governed realised benefit."
        ))
        replace_exact_prefix(doc, "Conclusions:", (
            "Conclusions: Candidate-pool expansion raises both finite audited opportunity and selection pressure. Effective diversity is informative but not causal; net value depends on whether finite validation can identify complementary candidates at an acceptable downstream budget."
        ))
        replace_exact_prefix(doc, "Scientific Contribution:", (
            "Scientific Contribution: This study provides an auditable molecular-model-selection decomposition that jointly intervenes on candidate count and composition, evaluates component and compute controls, quantifies selection stability, tests predefined anchors and normalizations, and transports composition effects across two split mechanisms. It does not claim a new universal statistical law or a fully optimized architecture leaderboard."
        ))
        replace_exact_prefix(doc, "Public dataset provenance", (
            "Public dataset provenance is listed in Additional file 2: Table S1. The submission package supplies frozen registries, fold-level inner and outer utilities, split manifests, source hashes, Figure 7 and Figures S16–S21 as editable SVG/PDF and 600-dpi PNG, Tables S32–S36, timing records, native equation-to-code mappings and all analysis scripts required to reproduce the expanded intervention."
        ))
        replace_section(doc, "2.14 ", "2.15 ", "2.14 Expanded equal-size registry-composition intervention", blocks["methods"])
        statistical = find_paragraph(doc, lambda p: p.text.strip().startswith("2.15 Statistical inference"))
        statistical.insert_paragraph_before("2.15 Mathematical estimands and matrix-relative diversity", style="Heading 2")
        for style, text in equation_blocks("en"):
            statistical.insert_paragraph_before(text, style=style)
        statistical.text = "2.16 Statistical inference"
        replace_section(doc, "3.10 ", "4 Discussion", "3.10 Expanded composition, stability and compute controls", blocks["results"])
        replace_section(doc, "4.8 ", "4.9 ", "4.8 Registry composition, selection stability and bounded compute", blocks["discussion"])
        replace_section(doc, "4.9 ", "5 Conclusions", "4.9 Limitations", blocks["limitations"])
        replace_section(doc, "5 Conclusions", "Supplementary Information", "5 Conclusions", blocks["conclusions"])
    else:
        # Chinese manuscript retains the same numbered section prefixes and formula numbering.
        replace_exact_prefix(doc, "方法：", "方法：在九个公开端点的重复嵌套审计基础上，加入预先固定的六端点候选池组成干预。三类锁定候选池在 K = 4、8、16、32 下采用重复 3×3 嵌套骨架切分，并完成组件消融、等下游预算截断、三个预设锚点、三种归一化、选择稳定性及三端点 Tanimoto 连通分量闭环。")
        replace_exact_prefix(doc, "结果：", f"结果：18 个‘端点×候选池’单元中，CAHit@3 从 K = 4 到 32 在 {m['cahit_negative_cells']} 个单元下降，交叉拟合差距在 {m['gap_positive_cells']} 个单元增加。K = 32 时，现代增强相对同质候选池的验证选择增益在 {m['modern_vs_hom_positive']}/6 个端点为正。组件、预算、锚点与切分机制分析表明，实际收益由互补信息、选择稳定性与下游计算暴露共同决定，而非仅由名义 K 决定。")
        replace_exact_prefix(doc, "结论：", "结论：候选池扩张同时提高有限审计机会并增加选择压力。有效多样性具有描述价值，但不是选择损失的因果参数；净收益取决于有限验证能否在可接受下游预算内识别互补候选。")
        replace_exact_prefix(doc, "科学贡献：", "科学贡献：本研究在分子模型选择中联合干预候选数量与组成，完成组件和计算控制，量化选择稳定性，检验预设锚点与归一化，并评估组成效应在两种切分机制间的迁移。本文不声称新的普适统计定律，也不构成充分优化的现代架构排行榜。")
        replace_exact_prefix(doc, "公开数据来源", "公开数据来源见补充表 S1。提交包提供冻结候选表、内外层逐折效用、切分清单、源文件哈希、Figure 7 与图 S16–S21 的可编辑 SVG/PDF 和 600 dpi PNG、表 S32–S36、时间记录、原生公式—代码映射及复现扩展干预所需的全部分析脚本。")
        replace_section(doc, "2.14 ", "2.15 ", "2.14 扩展的等规模候选池组成干预", blocks["methods"])
        statistical = find_paragraph(doc, lambda p: p.text.strip().startswith("2.15 "))
        statistical.insert_paragraph_before("2.15 数学估计量与矩阵相对多样性", style="Heading 2")
        for style, text in equation_blocks("zh"):
            statistical.insert_paragraph_before(text, style=style)
        statistical.text = "2.16 统计推断"
        replace_section(doc, "3.10 ", "4 ", "3.10 扩展的组成、稳定性与计算控制", blocks["results"])
        replace_section(doc, "4.8 ", "4.9 ", "4.8 候选池组成、选择稳定性与受限计算", blocks["discussion"])
        replace_section(doc, "4.9 ", "5 ", "4.9 局限性", blocks["limitations"])
        replace_section(doc, "5 ", "补充信息", "5 结论", blocks["conclusions"])
    doc.save(destination)


def build_reviewer_response_matrix() -> None:
    rows = [
        ("Equal-size evidence covered only three endpoints and K=16/32", "Six prespecified endpoints; K=4/8/16/32 exact prefixes", "Methods 2.14; Results 3.10; Figure 7; Figure S16; Table S32"),
        ("Modern panel was a boundary check rather than a component intervention", "Fixed-K ChemBERTa, MoLFormer, D-MPNN and full-modern ablations", "Methods 2.14; Figure S17; Table S33"),
        ("Compute exposure could explain composition effects", "Locked equal-downstream-budget prefix sensitivity", "Figure 7D; Figure S21; Tables S33 and S36"),
        ("Anchor and normalization choices could determine direction", "Three predefined anchors and raw/MAD/paired-homogeneous normalizations", "Methods 2.14-2.15; Figure S18; Table S34"),
        ("Composition effects were not tested under strict structure separation", "Full three-pool loop under scaffold and Tanimoto-component splits", "Results 3.10; Figure S19; Table S35"),
        ("Effective diversity may not be usable by validation", "Candidate/representation/family frequencies, entropy, seed agreement and fold switches", "Figure 7C; Figure S20; Table S36"),
        ("Core estimands lacked numbered mathematical definitions", "Twenty-seven native Word equations with explicit aggregation and epsilon", "Methods 2.15; equation notation and code-mapping workbook"),
        ("Endpoint-pool-K cells risk pseudoreplication", "Outer folds averaged within seed; five seeds retained as bootstrap blocks", "Methods 2.16; all expanded summary tables"),
    ]
    doc = Document()
    doc.add_heading("Paper31 reviewer-concern response and evidence-location matrix", level=1)
    doc.add_paragraph(
        "This matrix records where each strengthened analysis is implemented. It is an internal submission aid and does not convert sensitivity analyses into confirmatory claims."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["Reviewer concern", "Added control", "Evidence location"]):
        cell.text = text
    for concern, control, location in rows:
        cells = table.add_row().cells
        cells[0].text = concern
        cells[1].text = control
        cells[2].text = location
    doc.save(PACKAGE / "Reviewer_concern_Response_Location_Paper31.docx")
    pd.DataFrame(rows, columns=["reviewer_concern", "added_control", "evidence_location"]).to_csv(
        DATA / "Paper31_reviewer_concern_response_matrix.csv", index=False, encoding="utf-8-sig"
    )


def main() -> None:
    PACKAGE.mkdir(parents=True, exist_ok=True)
    m = metrics()
    update_document(EN_BASE, EN_OUT, "en", m)
    update_document(ZH_BASE, ZH_OUT, "zh", m)
    build_reviewer_response_matrix()
    audit = {
        "status": "text_updated_equation_placeholders_pending_native_OMath_insertion",
        "baseline_note": "requested manuscript(6)(2).docx was not present; Paper30 manuscript(6).docx was the latest discoverable English source",
        "english": str(EN_OUT), "chinese": str(ZH_OUT), "metrics": m,
    }
    (DATA / "Paper31_manuscript_build_audit.json").write_text(
        json.dumps(audit, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    print(EN_OUT)
    print(ZH_OUT)


if __name__ == "__main__":
    main()
