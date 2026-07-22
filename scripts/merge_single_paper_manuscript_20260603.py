from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from merge_complete_integrated_manuscript_20260603 import (
    BASE_DOCX,
    BASE_MD,
    DOCS,
    FIG_DIR,
    OUT_DOCX,
    OUT_MD,
    TABLE_PREVIEW_DIR,
    add_figure,
    add_heading,
    add_para,
    add_table,
    build_experiment_index,
    build_policy_summary,
    build_promoted_table,
    build_roughness_summary,
    normalize_styles,
    set_margins,
)


OUT_SINGLE_DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260603.docx"
OUT_SINGLE_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260603.md"


def build_strengthened_experiment_table() -> pd.DataFrame:
    rows = [
        [
            "主性能基线",
            "MoleculeNet、TDC official、MoleculeACE",
            "ROC-AUC、RMSE、MAE、Spearman、PR-AUC",
            "以 retained-best gate 报告最终结果，同时保留 seed-level mean/std 和 split 惩罚，避免只展示单次最优数值。",
            "主文结果与主表",
        ],
        [
            "外部 benchmark appendix",
            "OpenADMET fast subset、TDC full-panel appendix",
            "endpoint-level primary metric、retention source、AD/OOD proxy",
            "证明 selector 在外部 ADMET 任务上可以稳定挑选强候选；目标是外推证据而不是无边界重跑大模型。",
            "强附录；主文摘要引用",
        ],
        [
            "性能模式补强",
            "Top-K mean、stacking、target transform、树模型专家",
            "validation-only Δ、win/tie/loss、retained-best promotion",
            "补强低分模块时先看 validation gate；Lipo 与部分 TDC endpoint 有接入收益，FreeSolv 仍作为 remaining limitation。",
            "低分模块 subsection",
        ],
        [
            "不平衡分类与校准",
            "ClinTox、DILI、hERG、CYP substrate",
            "PR-AUC、Brier、ECE、coverage、risk-coverage",
            "把 ROC-AUC 之外的可靠性指标纳入证据链，降低类别不平衡导致的虚高风险。",
            "可靠性结果与附表",
        ],
        [
            "selector 固定策略审计",
            "risk_adjusted_lambda_0.5、stability tie-breaker",
            "正向数、平均 Δ、最大正/负向 Δ",
            "预先固定 policy 后再集成；10/32 endpoint-metric 正向，但平均 Δ 仍为负，因此定位为 targeted rescue 而非全局替代。",
            "主文表 15-16；附表 S36-S38",
        ],
        [
            "负结果与瓶颈诊断",
            "3D-lite descriptors、roughness-weighted regression",
            "validation gate、oracle audit、roughness proxy",
            "新增候选没有被正式接入，说明当前瓶颈更多来自验证噪声、局部 target jump 和 selector uncertainty。",
            "主文表 17；限制讨论",
        ],
        [
            "可解释性与案例",
            "motif attribution、fragment enrichment、uncertainty cases",
            "片段贡献、邻域相似度、不确定性/误差一致性",
            "用 Lipo rescue、ClinTox 高风险样本和 high-roughness ADME regression 案例连接性能、可靠性和化学解释。",
            "主文案例；补充图",
        ],
    ]
    return pd.DataFrame(rows, columns=["实验层", "覆盖任务", "核心指标", "主要结论", "结果定位"])


def build_single_paper_experiment_index() -> pd.DataFrame:
    frame = build_experiment_index().copy()
    mask = frame["实验模块"].eq("Formal fixed selector")
    frame.loc[mask, "图表编号"] = "图 11 / 表 15-16 / Table S36-S38"
    frame.loc[mask, "主要结论"] = "预先固定 risk_adjusted_lambda_0.5，32 个 endpoint-metric 中 10 个正向 promotion；完整负向审计保留在附表。"
    extra = pd.DataFrame(
        [
            {
                "实验模块": "3D-lite/roughness negative gate",
                "图表编号": "表 17 / Figure S13 / Table S30-S32",
                "文件": "table35-37_*.csv",
                "主要结论": "3D-lite 和 roughness-weighted regression 未通过 validation-only gate，作为低分回归瓶颈与负结果诊断保留。",
            },
            {
                "实验模块": "Experiment-thickening evidence chain",
                "图表编号": "表 18",
                "文件": "manuscript section 3.14",
                "主要结论": "将主性能、外部 benchmark、性能模式、校准/不平衡、selector audit、负结果和解释性案例统一成可审稿的证据链。",
            },
        ]
    )
    return pd.concat([frame, extra], ignore_index=True)


def find_paragraph(doc: Document, prefix: str):
    for para in doc.paragraphs:
        if para.text.strip().startswith(prefix):
            return para
    raise ValueError(f"paragraph starting with {prefix!r} not found")


def move_appended_blocks_before(doc: Document, marker_text: str, target_para) -> None:
    body = doc._body._element
    children = list(body)
    marker = None
    for child in children:
        if child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.iter() if node.tag == qn("w:t"))
            if marker_text in text:
                marker = child
                break
    if marker is None:
        raise ValueError(f"marker {marker_text!r} not found")
    moving = []
    found = False
    for child in list(body):
        if child is marker:
            found = True
            continue
        if found and child.tag != qn("w:sectPr"):
            moving.append(child)
    body.remove(marker)
    for child in moving:
        body.remove(child)
    target = target_para._p
    for child in moving:
        target.addprevious(child)


def relabel_existing_captions(doc: Document) -> None:
    replacements = {
        "表 15. 主文图表对应关系。": "表 19. 主文图表对应关系。",
    }
    for para in doc.paragraphs:
        text = para.text.strip()
        if text not in replacements:
            continue
        new_text = replacements[text]
        if para.runs:
            for run in para.runs:
                run.text = ""
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)


def append_results_blocks(doc: Document) -> None:
    marker = "__MOVE_RESULTS_20260603__"
    doc.add_paragraph(marker)
    add_heading(doc, "3.12 Formal fixed-selector integration", level=2)
    add_para(
        doc,
        "为避免继续无边界堆叠 candidate，本轮将 selector 改进正式接入结果部分。我们预先固定 risk_adjusted_lambda_0.5 作为主策略，并把 stability tie-breaker 作为敏感性分析；不允许按 endpoint 使用测试集挑选不同 policy。该设计把新增实验从“事后挑结果”改成“预注册固定规则后的审计”，更符合高水平论文对模型选择边界的要求。",
    )
    add_figure(
        doc,
        FIG_DIR / "fig22_formal_fixed_selector_integration.png",
        "图 11. Formal fixed-selector integration。risk-adjusted λ=0.5 为主策略，stability tie-breaker 为敏感性分析；图中展示固定策略正向 endpoint-metric 数和 retained-best promotion。",
        width=6.9,
    )
    add_table(
        doc,
        "表 15",
        "Formal risk-adjusted selector 带来的 retained-best promotion。",
        build_promoted_table(),
        note="Δ 已按指标方向统一为正值代表更优；完整正负固定策略审计见 Table S36-S38。",
        font_size=6.3,
    )
    add_table(
        doc,
        "表 16",
        "固定 selector policy 的正负结果汇总。",
        build_policy_summary(),
        note="risk-adjusted λ=0.5 是主策略；stability tie-breaker 仅作为敏感性分析。平均 Δ 为负，说明固定策略不能全局替换主 selector。",
        font_size=6.8,
    )
    add_para(
        doc,
        "正式 integration 后，risk-adjusted λ=0.5 在 32 个 endpoint-metric 中有 10 个相对当前 retained-best 为正向，最明显的收益来自 half-life Obach、clearance microsome AZ、clearance hepatocyte AZ 和 CYP2C9 substrate。与此同时，该策略仍有 22 个 endpoint-metric 低于当前 retained-best，因此主文应把它定位为 targeted rescue 和 selector robustness，而不是全局替代主 selector。",
    )

    add_heading(doc, "3.13 3D-lite 与 roughness-weighted regression 的负结果", level=2)
    add_para(
        doc,
        "3D-lite descriptors 和 roughness-weighted tree 是针对低分回归模块的合理尝试，尤其贴合 FreeSolv、clearance、half-life、PPBR 等 endpoint。然而 validation-only gate 没有在任何测试 endpoint 上正式接入新候选。这个负结果应保留在主文或强附录中，因为它说明性能瓶颈并不是单纯缺少 candidate，而是 selector 如何处理验证噪声、roughness 和 endpoint-specific uncertainty。",
    )
    add_table(
        doc,
        "表 17",
        "3D-lite/roughness-weighted regression gate 的 retained-best 结果。",
        build_roughness_summary(),
        note="所有 endpoint 均保留当前 retained-best；oracle audit 只用于诊断候选池中是否存在未被 validation gate 捕捉的信号。",
        font_size=6.4,
    )
    add_para(
        doc,
        "与 formal fixed-selector integration 放在一起解读时，3D-lite/roughness 的负结果反而加强了论文结论：在当前实验条件下，更值得改进的是 selector 的稳定性和不确定性处理，而不是继续扩大候选模型数量。",
    )

    add_heading(doc, "3.14 实验加厚后的证据链", level=2)
    add_para(
        doc,
        "为进一步增强实验部分，本稿将新增实验统一组织为“主性能、外部外推、性能模式补强、可靠性/校准、固定 selector 审计、负结果诊断和可解释性案例”七类证据。这样的写法可以避免实验结果显得零散：主文先回答 FZYC-Mol 是否在标准分子性质预测任务上有效，再回答这种有效性是否能在更真实的 ADMET 外推、不平衡分类和低相似度样本中保持，最后用负结果与案例分析说明哪些模块仍然是后续工作瓶颈。",
    )
    add_para(
        doc,
        "从模型性能角度看，当前最有价值的改进不是继续无差别增加 candidate 数量，而是把候选模型纳入可复现的 selector governance：所有 rescue heads、Top-K/stacking ensemble、target transform、tree baselines、3D-lite descriptors 和 roughness-weighted regression 都必须先通过 validation-only gate。未通过 gate 的模块仍保留为诊断证据，而不是被删除；通过 gate 的模块才进入 retained-best 主结果。这一点对于回应审稿人关于测试集选择偏差、过拟合和负结果透明度的质疑尤其关键。",
    )
    add_table(
        doc,
        "表 18",
        "实验加厚后的证据链与稿件定位。",
        build_strengthened_experiment_table(),
        note="该表用于连接主文、附录和后续 rebuttal 叙事；所有性能补强均遵守 validation-only selector 原则，未接入候选作为负结果和限制保留。",
        font_size=5.7,
    )
    move_appended_blocks_before(doc, marker, find_paragraph(doc, "4. 讨论"))


def append_discussion_blocks(doc: Document) -> None:
    marker = "__MOVE_DISCUSSION_20260603__"
    doc.add_paragraph(marker)
    add_heading(doc, "4.4 整合后的结果定位与写作边界", level=2)
    add_para(
        doc,
        "整合后的投稿叙事建议采用三层结构。第一层是主结果：MoleculeNet、TDC official split、split realism、reliability/AD、motif interpretability，以及 formal fixed-selector integration。第二层是强附录：formal external appendix、performance-mode、rescue-integrated selector、Nature-style fusion 和 selector strategy audit。第三层是诊断附录：3D-lite/roughness negative result、oracle audit、roughness/literature alignment 和 candidate family details。",
    )
    add_para(
        doc,
        "表格呈现方面，主文只放经过精排的紧凑表；完整候选级数据保留在 supplementary tables。对于 fixed selector，主文表 15 只显示 promotion，但正文必须明确说明完整固定策略仍有 22 个负向 endpoint-metric，完整审计在 Table S36-S38。这个写法比只报告正向结果更稳，也更能抵御审稿人对选择偏差的质疑。",
    )
    move_appended_blocks_before(doc, marker, find_paragraph(doc, "5. 结论"))


def append_figure_index_blocks(doc: Document) -> None:
    marker = "__MOVE_INDEX_20260603__"
    doc.add_paragraph(marker)
    add_heading(doc, "6.1 全部实验模块与图表索引", level=2)
    add_table(
        doc,
        "表 20",
        "截至 2026-06-03 的全部实验结果与论文图表对应关系。",
        build_single_paper_experiment_index(),
        note="该索引用于替代分散文件查找；完整 CSV、PNG/SVG 和报告均已纳入 submission package。",
        font_size=5.8,
    )
    add_figure(
        doc,
        TABLE_PREVIEW_DIR / "table2_moleculenet_main_polished.png",
        "精排表预览 1. MoleculeNet 主结果；完整 seed-level 与候选池细节见 Table S1、S22-S26。",
        width=7.2,
    )
    add_figure(
        doc,
        TABLE_PREVIEW_DIR / "table3_tdc_official_admet_polished.png",
        "精排表预览 2. TDC official split 摘要；该表同时显示 retained 结果和平均 scaffold 惩罚。",
        width=7.2,
    )
    add_figure(
        doc,
        TABLE_PREVIEW_DIR / "table8_formal_fixed_selector_main_polished.png",
        "精排表预览 3. Formal fixed-selector rescue summary；该表只列出预先固定 risk-adjusted selector 的正向 promotion。",
        width=7.2,
    )
    move_appended_blocks_before(doc, marker, find_paragraph(doc, "7. 参考文献"))


def iter_block_items(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def table_to_markdown(table: Table) -> str:
    rows = [[cell.text.replace("\n", " ").strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    header = rows[0]
    sep = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def docx_to_markdown(docx_path: Path, out_path: Path) -> None:
    doc = Document(docx_path)
    lines: list[str] = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            style = block.style.name if block.style is not None else ""
            if style.startswith("Heading 1"):
                lines.extend(["", f"# {text}", ""])
            elif style.startswith("Heading 2"):
                lines.extend(["", f"## {text}", ""])
            elif style.startswith("Heading 3"):
                lines.extend(["", f"### {text}", ""])
            else:
                lines.extend([text, ""])
        else:
            markdown = table_to_markdown(block)
            if markdown:
                lines.extend([markdown, ""])
    out_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def integrate_markdown() -> None:
    docx_to_markdown(OUT_SINGLE_DOCX, OUT_SINGLE_MD)
    try:
        shutil.copy2(OUT_SINGLE_MD, OUT_MD)
    except PermissionError:
        print(f"Skipped markdown overwrite because the file is open: {OUT_MD}")


def main() -> None:
    shutil.copy2(BASE_DOCX, OUT_SINGLE_DOCX)
    doc = Document(OUT_SINGLE_DOCX)
    normalize_styles(doc)
    for section in doc.sections:
        set_margins(section)
    relabel_existing_captions(doc)
    append_results_blocks(doc)
    append_discussion_blocks(doc)
    append_figure_index_blocks(doc)
    doc.save(OUT_SINGLE_DOCX)
    try:
        shutil.copy2(OUT_SINGLE_DOCX, OUT_DOCX)
    except PermissionError:
        print(f"Skipped overwrite because the file is open: {OUT_DOCX}")
    integrate_markdown()
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {OUT_SINGLE_DOCX}")
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_SINGLE_MD}")


if __name__ == "__main__":
    main()
