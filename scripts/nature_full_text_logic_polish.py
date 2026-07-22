from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE_NAME = "FZYC-Mol_初稿-7.docx"
OUTPUT_NAME = "FZYC-Mol_初稿-8_Nature全稿语言逻辑修订.docx"


def set_run_font(run, latin="Times New Roman", east_asia="宋体", size=10.5, bold=False):
    run.font.name = latin
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def replace_text(paragraph, text):
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    run = paragraph.add_run(text)
    set_run_font(run)


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_before(paragraph, text, style_name="Normal"):
    new_p = paragraph.insert_paragraph_before("")
    new_p.style = style_name
    run = new_p.add_run(text)
    set_run_font(run)
    return new_p


def first_para(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p.text.strip()):
            return p
    return None


def replace_by_start(doc, start, new_text):
    p = first_para(doc, lambda t: t.startswith(start))
    if p is None:
        raise RuntimeError(f"Paragraph not found: {start}")
    replace_text(p, new_text)


def main():
    source = Path.home() / "Desktop" / SOURCE_NAME
    if not source.exists():
        matches = sorted((Path.home() / "Desktop").glob("FZYC-Mol_初稿-7*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
        if not matches:
            raise FileNotFoundError("Cannot find 初稿-7 on Desktop.")
        source = matches[0]

    doc = Document(str(source))


    replace_by_start(
        doc,
        "分子性质预测是药物发现、ADMET 评估和毒性风险筛查中",
        "分子性质预测是药物发现、ADMET 评估和毒性风险筛查中的关键计算环节。随机划分下的平均 ROC-AUC 或 RMSE 难以充分反映模型在新骨架、少样本终点、不平衡标签、活性悬崖、实验噪声和适用域漂移条件下的可靠性。为解决这一问题，本文提出 FZYC-Mol，一种由验证集治理的适用域感知分子性质预测框架。该框架不依赖单一大型主干模型，而是把多视图表示、强基线专家、目标变换、融合候选、补救头、校准和适用域证据纳入冻结候选池；最终策略仅由验证集决定，测试集只在策略冻结后用于一次性评估。",
    )
    replace_by_start(
        doc,
        "在 MoleculeNet 主面板中",
        "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别达到 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302。验证集接受的定向补救改善了 Lipophilicity，低成本重构改善了 FreeSolv，但 FreeSolv 仍落后于观测最佳 Chemprop 候选，因此被保留为边界案例。多视图融合和适用域门控在 BBBP、ClinTox 以及外部 TDC ADMET 的 Caco2_Wang、HIA_Hou 和 Pgp_Broccatelli 上产生选择性增益；22 个外部终点的最终保留结果为 win/tie/loss = 5/17/0。结果表明，FZYC-Mol 更适合被界定为终点依赖的选择、拒绝和审计流程，而不是保证所有任务统一提升的单一模型。",
    )
    replace_by_start(
        doc,
        "可靠性分析显示",
        "可靠性分析显示，风险分数对分类错误具有较好的识别能力，但对回归高误差样本的识别能力为有限至中等；对应中位 AUROC 分别为 0.788 和 0.652。保形预测在 MoleculeNet 分类任务的 80%/90%/95% 目标覆盖率下达到平均经验覆盖率 0.814/0.918/0.956，在回归任务下达到 0.823/0.925/0.962。验证-测试排名审计显示，跨 200 个 dataset-seed 候选池的中位 Spearman 相关为 0.667，测试最佳落入验证 Top-3 的比例为 0.295，Top-1 一致比例为 0.135。上述结果说明，验证集治理可以减少测试集事后选择，但不能保证测试最优；所有增益仍需结合 regret、optimism gap、嵌套验证和负结果审计解释。基序归因和片段富集仅作为关联性化学解释，不应被视为因果机制证明。",
    )


    replace_by_start(
        doc,
        "进一步看近年高水平分子机器学习论文",
        "近期分子机器学习研究更强调把模型设计与可解释化学单元、数据划分压力、任务异质性和部署边界联系起来。MotiL 将分子基序学习作为预训练目标，强调 scaffold 与 motif-level 信息在性质预测中的作用 [45]；多通道层级表示研究则把分子结构层级、通道级预训练和任务特异聚合结合起来，以应对 activity cliffs 等困难场景 [46]；OmniMol 从超图视角处理不完整标注数据，说明端点缺失和标签不完备本身就是 ADMET 建模的重要挑战 [47]。这些研究共同提示，FZYC-Mol 的评价重点应放在验证集治理、任务边界和可解释证据链上，而不是把候选模型数量本身作为创新。",
    )
    replace_by_start(
        doc,
        "近期分子性质预测与 ADMET benchmark 论文的写作趋势也提示",
        "近期分子性质预测与 ADMET benchmark 研究也提示，正文应避免堆叠过多候选细节，而应围绕数据划分、主指标、独立测试、不确定性、适用域和失败边界组织证据。层级交互模型、ADMET 表征基准、分子性质预测综述和多模态毒性预测研究通常把紧凑主表、seed 级补充表、清晰图注和负结果说明结合起来。基于这一原则，本文主表仅保留直接支撑结论的摘要信息，候选级、seed 级、超参数级和完整指标级结果统一放入补充材料。",
    )


    replace_by_start(
        doc,
        "为符合 Nature 系列期刊对公式可编辑性",
        "为使方法定义可复核、可引用和可重现，本节仅保留直接支撑 FZYC-Mol 方法学主张的编号公式。正文引用编号公式时采用“式 (n)”的形式，通用评价指标和统计检验细节放入 Supplementary Methods。设 t 表示预测终点，a 表示候选策略，x_i 和 y_i 分别表示第 i 个分子的结构输入和观测标签，所有候选池、阈值、容差和打破平局规则均在测试集评估前冻结。",
    )

    formula_replacements = [
        ("在 equation (1) 中", "在式 (1) 中，训练集、验证集和测试集互不重叠。式 (2) 定义候选专家可使用的多视图分子表示。式 (3) 中，M_t 为预先指定的主指标；当指标越大越优时 dir_t=1，当误差类指标越小越优时 dir_t=-1。该定义把 ROC-AUC、PR-AUC、RMSE 和 MAE 转换到统一的验证效用方向，同时保留原始指标的报告方式。"),
        ("Equations (4)-(6)", "式 (4)-(6) 给出由验证集治理的选择器。系统先识别验证效用最高的候选，再把落入 ε_t 容差范围的近似并列候选保留下来，最后按验证风险、模型复杂度和跨 seed 稳定性确定最终策略。若没有触发近似并列或稳定性规则，式 (6) 退化为式 (4) 中的验证集最佳候选。"),
        ("Equations (7)-(10)", "式 (7)-(10) 用于区分冻结选择器和事后测试集最佳候选。Regret_t 衡量冻结策略与测试集观测最佳之间的差距，OptGap_t 衡量验证集乐观偏差。排名审计比较验证集和测试集候选排序，式 (10) 则定义 nested validation，其中外层测试折不参与内层候选选择。"),
        ("Equations (11)-(16)", "式 (11)-(16) 定义多专家融合和适用域输出。B_t 表示通过验证集保留的融合候选集合，u_{i,a} 表示候选专家 a 在样本 i 上的不确定性，τ 控制验证效用对权重的影响强度。最近邻 Tanimoto 分数用于估计化学空间支持度；落在适用域之外的样本被路由到更保守的保留专家，或被标记为高风险样本。"),
        ("Equation (17) 是", "式 (17) 是 Fig. 11 中 risk-coverage 曲线的计算基础。式 (18) 用于报告分类任务的概率校准。式 (19) 用于 ClinTox fixed-precision audit，因为在严重类别不平衡条件下，高 ROC-AUC 可能掩盖阳性类别召回不足的问题。"),
        ("Equations (20) 和", "式 (20) 和式 (21) 分别定义基于校准集的 split-conformal 回归区间和分类集合。式 (22) 固定低相似度分析中的三个互斥 Tanimoto 分层。式 (23) 定义 MoleculeACE 活性悬崖分子对及预测差异与真实差异的相关性。式 (24) 报告片段层面的效应量和 Benjamini-Hochberg FDR 校正证据，因此基序和片段结果被解释为关联性化学证据，而不是因果机制证明。"),
        ("综上，equations (1)-(24)", "综上，式 (1)-(24) 共同定义了冻结数据边界、仅基于验证集的选择器、最终保留门控、排名审计、nested validation、融合、适用域距离、risk-coverage、校准、conformal prediction、低相似度分层、活性悬崖审计和片段富集统计。这些公式也明确区分了用于模型选择的量和策略冻结后才报告的审计量。"),
    ]
    for start, replacement in formula_replacements:
        replace_by_start(doc, start, replacement)


    replace_by_start(
        doc,
        "本节按审稿问题重新组织证据链",
        "本节按照证据链组织结果。首先报告 MoleculeNet 和外部 ADMET 的基线与最终保留性能，并明确小幅增益和 FreeSolv 边界。随后评估验证集选择偏差，包括验证-测试排名、Top-1/Top-3、regret、optimism gap 和 nested validation。接着展示 OOD、低相似度三档和 MoleculeACE 活性悬崖结果，并报告系统消融、固定选择器和负结果。最后，结合校准、risk-coverage、保形预测、化学解释和失败案例说明哪些预测不应被过度信任。",
    )

    # Remove duplicated paragraphs that repeated earlier Results statements inside the interpretability section.
    duplicate_starts = [
        "不平衡分类结果进一步强化了 ClinTox、DILI",
        "统计稳定性层面，性能增强候选在终点-seed",
    ]
    for start in duplicate_starts:
        seen = False
        for p in list(doc.paragraphs):
            text = p.text.strip()
            if text.startswith(start):
                if not seen:
                    seen = True
                else:
                    delete_paragraph(p)

    replace_by_start(
        doc,
        "外部附录的关键结论是",
        "外部附录显示，性能增强候选并不会在所有终点上稳定胜出，但最终保留选择器能够在验证证据支持时吸收增益，并在候选较弱时保留旧基线。具体而言，22 个外部终点中有 5 个由性能增强候选保留，17 个保留旧结果，最终保留的 win/tie/loss 为 5/17/0。该结果支持更克制的结论：FZYC-Mol 的优势来自验证集治理下的选择性改进，而不是所有模型或所有终点的统一提升。",
    )


    replace_by_start(
        doc,
        "本研究的主要信息是",
        "本研究表明，分子性质预测模型的价值不应只由单一排行榜分数决定，还应由选择过程、结构外推能力、适用域边界和失败模式共同界定。FZYC-Mol 的主要贡献不在于扩大模型规模，而在于把多专家候选、验证集选择、划分真实性、外部 ADMET 评估、可靠性分析和化学解释纳入同一证据链。该定位与近期 ADMET 研究趋势一致：真实药物发现更关心模型在结构新颖分子、低相似度样本、不平衡毒性标签和噪声终点下是否仍可被信任。",
    )
    replace_by_start(
        doc,
        "本研究也有明确局限",
        "本研究仍存在明确局限。验证-测试排名审计和 nested validation 表明，验证集治理可以减少测试集事后选择，但仍存在验证集选择偏差，不能保证测试最优；因此，小幅增益应与 regret、optimism gap、Top-3 命中和负结果共同解释。收益也具有明显终点异质性，BBBP、ClinTox、HIA 和 Pgp 等终点的增益较小，FreeSolv 仍落后于观测最佳 Chemprop 候选。此外，基序归因和片段富集虽已补充 support、effect size、p 值和 FDR q 值，但仍属于关联证据，不应被解读为未经湿实验验证的因果机制。当前稿件也尚未纳入湿实验验证，ChemBERTa 与 MoLFormer 主要以冻结编码器形式使用，Polaris 与 OpenADMET 的完整官方挑战流程仍需进一步扩展。",
    )
    replace_by_start(
        doc,
        "首先，需要说明为何本文未将",
        "本研究未将更大规模预训练模型或全量微调作为主线，是因为当前证据显示，性能瓶颈并不完全来自表征能力不足。目标值长尾、类别不平衡、局部结构-性质关系粗糙、标签冲突和结构外推同样会限制模型表现。单纯扩大模型规模可能提高某些终点分数，但也会增加算力成本、过拟合风险和验证选择复杂度。因此，本文优先采用快速外部评估、目标变换、Top-K 集成和强表格基线，并把更大模型实验保留为受控扩展方向。",
    )
    replace_by_start(
        doc,
        "第二个可能问题是",
        "验证集过拟合是该框架的核心风险。本文通过五个设计暴露并降低这一风险：选择器只在预定义候选池中选择，不允许根据测试结果添加临时规则；结果按多个随机种子汇总，并报告均值、标准差和配对统计；补救头进入最终策略需要通过验证集接受，且不会被所有终点无差别采用；排名审计、Top-3 命中、optimism gap 和负结果表将低相关候选池明确写为选择风险；9 个代表性终点的 3×3 nested validation 用于补充检查内外层划分下的稳定性。更大规模的跨数据集迁移验证仍是后续工作重点。",
    )
    replace_by_start(
        doc,
        "第三个可能问题是",
        "可解释性结果也应受到边界约束。本文将基序归因与片段富集界定为关联解释，而非因果机制证明。它们的价值在于帮助识别模型关注的局部结构、定位高误差和分布外样本，并辅助药物化学用户判断预测是否合理。进一步证明因果机制仍需要湿实验、反事实分子设计或更严格的匹配分子对分析。",
    )
    replace_by_start(
        doc,
        "进一步验证仍应优先扩展",
        "进一步验证应优先扩展外部评估附录，而不是简单堆叠更多模型。更有价值的路线是在更多公开 ADMET 终点上系统评估 CatBoost、XGBoost、ExtraTrees、LightGBM、RF、Top-K 集成、目标变换和欠采样集成，并把最终保留决策与粗糙度、低相似度和验证-测试排名一致性关联起来。",
    )
    replace_by_start(
        doc,
        "第二优先级是把目前已有",
        "候选族也需要进一步预注册和标准化。当前 Top-K、堆叠、风险调整和稳定性打破平局规则已显示一定价值，但仍应在嵌套验证或外部时间划分上复核。Lipophilicity 补救和 MoleculeACE 悬崖目标候选的平均收益较小，不能被过度解释为普适改进。",
    )
    replace_by_start(
        doc,
        "第三优先级是继续加强",
        "案例层面的解释性分析仍需扩展。当前稿件已纳入 Lipophilicity 补救、ClinTox 高风险假阴性和高粗糙度 ADME 回归案例；后续可扩展 BACE、clearance_microsome_az 和 CYP 底物终点，并在同一样本级审计表中连接结构片段、最近邻、标签差异、不确定性和基序归因。",
    )
    replace_by_start(
        doc,
        "第四类方向是受控的大模型实验",
        "大模型实验应作为受控补充验证开展。在少数代表性终点上尝试 ChemBERTa/MoLFormer 适配器或轻量微调时，应采用嵌套验证或严格留出测试，以免模型容量增加放大验证集过拟合。这类实验可以检验冻结表征是否限制了当前结果，但不应替代以可靠性和验证治理为核心的主线。",
    )
    replace_by_start(
        doc,
        "结果呈现采用主线融合结构",
        "从适用边界看，FZYC-Mol 更适合作为候选策略治理和可靠性审计框架，而不是自动输出最高分模型的黑箱系统。MoleculeNet 性能、TDC 外部泛化、划分真实性、可靠性/适用域、选择器治理、补充实验闭环和基序/案例解释分别承担不同证据功能。外部附录、补救选择器、多方法融合、3D-lite/粗糙度负结果、MoleculeACE 活性悬崖和选择器策略审计共同说明，该框架的价值来自可追踪的接受与拒绝机制。",
    )
    replace_by_start(
        doc,
        "表格呈现遵循主文精简",
        "这种定位也决定了主文表格的角色。主表应呈现能够直接支撑结论的摘要证据，完整候选级和 seed 级结果应保留在补充材料中。固定选择器的主文摘要表展示正向提升，同时正文保留 22 个负向终点-指标单元并指向完整审计表。该写法能降低选择偏差风险，也使负结果成为方法边界的一部分。",
    )

    # Data availability: remove duplicated heading-like paragraph and merge the statement.
    replace_by_start(
        doc,
        "数据和代码可用性：",
        "数据和代码可用性：本文使用的公开数据集可通过原始平台获得。与本文结果对应的 split seeds、候选登记 CSV、验证/测试预测表、统计检验脚本、图表 source data、环境文件和主表生成命令应在正式投稿前整理至 GitHub/Zenodo 或期刊认可的数据仓库；当前草稿尚未填入正式链接，不应虚构 accession number。补充材料中的表 S1-Sxx 需在投稿前与正文表格编号统一核对。",
    )
    for p in list(doc.paragraphs):
        if p.text.strip() == "数据与代码可用性" or p.text.strip().startswith("本文使用的公开数据集、划分脚本"):
            delete_paragraph(p)

    output = source.parent / OUTPUT_NAME
    doc.save(str(output))

    # Audit after saving.
    revised = Document(str(output))
    all_text = "\n".join(p.text for p in revised.paragraphs)
    red_flags = [s for s in ["审稿", "点名", "Nature 系列", "本节按审稿", "写成单一新模型"] if s in all_text]
    duplicates = []
    seen = {}
    for i, p in enumerate(revised.paragraphs):
        txt = " ".join(p.text.strip().split())
        if not txt:
            continue
        if txt in seen:
            duplicates.append((seen[txt] + 1, i + 1, txt[:80]))
        else:
            seen[txt] = i
    formula_nums = []
    for p in revised.paragraphs:
        txt = p.text.strip()
        if txt.endswith(")") and "\t" in p.text:
            tail = txt.rsplit("(", 1)[-1].rstrip(")")
            if tail.isdigit():
                formula_nums.append(int(tail))
    wide_tables = [(i + 1, len(t.columns)) for i, t in enumerate(revised.tables) if len(t.columns) > 7]
    print(f"source={source}")
    print(f"output={output}")
    print(f"paragraphs={len(revised.paragraphs)} tables={len(revised.tables)} images={len(revised.inline_shapes)}")
    print(f"red_flags={red_flags}")
    print(f"duplicates={duplicates[:5]}")
    print(f"formula_numbers={formula_nums}")
    print(f"wide_tables={wide_tables}")


if __name__ == "__main__":
    main()
