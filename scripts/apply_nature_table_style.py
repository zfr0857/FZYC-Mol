from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]


def _get_or_add(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    specs = {
        "top": ("single", "12", "1F2937"),
        "bottom": ("single", "12", "1F2937"),
        "insideH": ("single", "4", "D0D5DD"),
        "left": ("nil", "0", "FFFFFF"),
        "right": ("nil", "0", "FFFFFF"),
        "insideV": ("nil", "0", "FFFFFF"),
    }
    for edge, (val, size, color) in specs.items():
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    tbl_pr.append(borders)


def set_table_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    margins = _get_or_add(tbl_pr, "w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), "85")
        node.set(qn("w:type"), "dxa")


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:cantSplit"))
    if node is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:tblHeader"))
    if node is None:
        node = OxmlElement("w:tblHeader")
        tr_pr.append(node)
    node.set(qn("w:val"), "true")


def format_run(run, *, size: float, bold: bool = False, color: str = "1F2937") -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def polish_tables(doc: Document) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table)
        set_table_cell_margins(table)
        if table.rows:
            repeat_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            keep_row_together(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if row_index == 0:
                    shade(cell, "F2F4F7")
                else:
                    shade(cell, "FFFFFF")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        format_run(run, size=8.2 if row_index else 8.5, bold=(row_index == 0))

    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name == "TableCaption":
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                format_run(run, size=9.2, bold=True, color="1F4E79")
        elif paragraph.style and paragraph.style.name == "TableNote":
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(5)
            for run in paragraph.runs:
                format_run(run, size=7.8, color="667085")


def main() -> None:
    source = Path(sys.argv[1]) if len(sys.argv) > 1 else ROOT / "output" / "小论文-7.docx"
    target = Path(sys.argv[2]) if len(sys.argv) > 2 else ROOT / "output" / "小论文-6.docx"
    doc = Document(source)
    before = "\n".join(p.text for p in doc.paragraphs)
    before += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    polish_tables(doc)
    after = "\n".join(p.text for p in doc.paragraphs)
    after += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    if before != after:
        raise RuntimeError("table styling changed document text; aborting")
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    print(f"Nature-style tables applied to {len(doc.tables)} tables: {target}")


if __name__ == "__main__":
    main()
