from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(r"D:\fzyc\output\paper35_submission_ready_20260718")
EN = ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_FINAL_CLEAN.docx"
ZH_CANDIDATES = [
    path for path in ROOT.glob("*.docx")
    if not path.name.startswith(("Candidate", "~$"))
]
ZH = next(
    (path for path in ZH_CANDIDATES if "Figure4更新版" in path.name),
    ZH_CANDIDATES[0],
)
TRACK = ROOT / "Candidate_pool_expansion_Journal_of_Cheminformatics_FINAL_TRACK_CHANGES.docx"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def xml_counts(path: Path) -> dict[str, int]:
    with ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
    return {
        "omml": len(re.findall(rb"<m:oMath(?:\s|>)", xml)),
        "drawings": xml.count(b"<w:drawing"),
        "insertions": xml.count(b"<w:ins"),
        "deletions": xml.count(b"<w:del"),
    }


def main() -> None:
    table = json.loads((ROOT / "Table_placement_report.json").read_text(encoding="utf-8"))
    equations = json.loads((ROOT / "Equation_rendering_report.json").read_text(encoding="utf-8"))
    crossrefs = json.loads((ROOT / "Cross_reference_report.json").read_text(encoding="utf-8"))
    figure7 = json.loads((ROOT / "Figure7_QC_report.json").read_text(encoding="utf-8"))
    figure4 = json.loads((ROOT / "reports" / "Figure4_QC_report.json").read_text(encoding="utf-8"))
    summary = json.loads((ROOT / "Paper35_final_report_summary.json").read_text(encoding="utf-8"))
    unicode_report = json.loads((ROOT / "Unicode_cleaning_report.json").read_text(encoding="utf-8"))
    references = json.loads((ROOT / "reports" / "Reference_verification_report.json").read_text(encoding="utf-8"))
    en_text = "\n".join(paragraph.text for paragraph in Document(EN).paragraphs)
    zh_text = "\n".join(paragraph.text for paragraph in Document(ZH).paragraphs)
    placeholders = [
        text for text in [
            "Author confirmation required before submission; no information was inferred in this revision.",
            "投稿前需作者确认",
            "本次修订未推断",
        ]
        if text in en_text or text in zh_text
    ]
    track_counts = xml_counts(TRACK)
    checks = {
        "english_clean_exists": EN.exists(),
        "english_track_changes_exists": TRACK.exists(),
        "chinese_manuscript_exists": ZH.exists(),
        "word_native_equations": xml_counts(EN)["omml"] == xml_counts(ZH)["omml"] == 14,
        "seven_main_figures_embedded": xml_counts(EN)["drawings"] == xml_counts(ZH)["drawings"] == 7,
        "track_changes_markup_present": track_counts["insertions"] > 0 and track_counts["deletions"] > 0,
        "table_placement": table["status"] == "pass",
        "equation_rendering": equations["status"] == "pass",
        "cross_references": crossrefs["status"] == "pass",
        "figure7_qc": figure7["overall_status"] == "pass",
        "figure4_qc": figure4["status"] == "pass",
        "unicode_cleaning": unicode_report["status"] == "pass",
        "reference_verification_36_of_36": references["reference_count"] == references["verified"] == 36 and references["unresolved"] == 0,
        "code_and_data": summary["code_data"] == "pass",
        "english_chinese_consistency": summary["english_chinese_consistency"] == "pass",
        "additional_file_4_zip": (ROOT / "supplementary" / "Additional_file_4_Code_and_reproducibility_package.zip").exists(),
        "dockerfile_present": (ROOT / "supplementary" / "Additional_file_4_Code_and_reproducibility_package" / "Dockerfile").exists(),
        "declarations_complete": not placeholders and summary["declarations"] == "pass",
    }
    technical_pass = all(value for key, value in checks.items() if key != "declarations_complete")
    status = "pass" if technical_pass and checks["declarations_complete"] else "blocked_author_declarations" if technical_pass else "fail"
    report = {
        "status": status,
        "checks": checks,
        "blocking_items": [
            "Verified author initials and CRediT roles",
            "Confirmation of competing-interests statement",
            "Confirmation of funding statement",
            "Confirmation of acknowledgements statement",
        ] if not checks["declarations_complete"] else [],
        "internal_placeholder_strings_detected": placeholders,
        "track_changes_counts": track_counts,
    }
    (ROOT / "Final_submission_validation_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    (ROOT / "Final_submission_validation_report.txt").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    manifest_exclude = {"Final_output_manifest.json", "SHA256SUMS.txt"}
    files = sorted(
        path for path in ROOT.rglob("*")
        if path.is_file()
        and path.name not in manifest_exclude
        and not path.name.startswith("~$")
    )
    manifest = {
        "package": ROOT.name,
        "version": "paper35-final-minor-revision-20260718",
        "validation_status": status,
        "file_count": len(files),
        "files": [{"path": path.relative_to(ROOT).as_posix(), "bytes": path.stat().st_size, "sha256": sha256(path)} for path in files],
    }
    (ROOT / "Final_output_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    hashes = [f"{sha256(path)}  {path.relative_to(ROOT).as_posix()}" for path in files]
    hashes.append(f"{sha256(ROOT / 'Final_output_manifest.json')}  Final_output_manifest.json")
    (ROOT / "SHA256SUMS.txt").write_text("\n".join(hashes) + "\n", encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
