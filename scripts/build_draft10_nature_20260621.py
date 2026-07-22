# -*- coding: utf-8 -*-
from __future__ import annotations

import copy
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

import polish_draft9_nature_20260621 as d9


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "初稿-9.docx"
OUTPUT = ROOT / "output" / "初稿-10.docx"
DESKTOP_OUTPUT = Path(r"C:\Users\Administrator\Desktop\修改\初稿-10.docx")

EXPANDED = ROOT / "reports" / "draft10_core_experiments_20260621" / "expanded_nested"
AUTOGLUON = ROOT / "reports" / "draft10_core_experiments_20260621" / "autogluon_nested"
FIGURE4 = ROOT / "output" / "初稿-10_图表与源数据" / "figures" / "fig04_expanded_nested_candidate_pool.png"


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    matches = [p for p in doc.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting with {prefix!r}, found {len(matches)}")
    return matches[0]


def replace_prefix(doc: Document, prefix: str, text: str) -> Paragraph:
    paragraph = find_paragraph(doc, prefix)
    d9.replace_paragraph_text(paragraph, text)
    return paragraph


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_after(paragraph: Paragraph, text: str = "", style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    created = Paragraph(new_p, paragraph._parent)
    created.style = style
    if text:
        d9.replace_paragraph_text(created, text)
    return created


def paragraph_index(paragraphs: list[Paragraph], target: Paragraph) -> int:
    return next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is target._p)


def set_formula(paragraph: Paragraph, formula: str, number: int) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.tab_stops.clear_all()
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(8.1), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(16.2), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.add_run("\t")
    math_run = paragraph.add_run(formula)
    math_run.font.name = "Cambria Math"
    math_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Cambria Math")
    math_run.font.size = Pt(10.5)
    paragraph.add_run("\t")
    number_run = paragraph.add_run(f"({number})")
    d9.set_run_font(number_run, 10.5)


def rebuild_formula_section(doc: Document) -> None:
    start = find_paragraph(doc, "2.4 验证治理选择器")
    end = find_paragraph(doc, "2.5 嵌套验证")
    paragraphs = list(doc.paragraphs)
    start_idx = paragraph_index(paragraphs, start)
    end_idx = paragraph_index(paragraphs, end)
    for paragraph in paragraphs[start_idx + 1 : end_idx]:
        delete_paragraph(paragraph)
    # Risk-coverage equations were previously repeated in section 2.6; the
    # rebuilt section below is the single mathematical definition block.
    for paragraph in list(doc.paragraphs):
        if paragraph.style and paragraph.style.name == "Equation":
            delete_paragraph(paragraph)

    content: list[tuple[str, str | tuple[str, int]]] = [
        (
            "Normal",
            "设 t 表示任务，a 表示候选，s 表示内层重复或外层折，r∈{val,test} 表示评价分区。"
            "为避免长推导打断方法叙述，主文只保留决定候选冻结、选择偏差和样本级可靠性的定义；"
            "ROC-AUC、RMSE、Brier 分数与 FDR 等通用统计量列入补充方法。",
        ),
        ("Normal", "方向一致的验证效用与单标准误集合定义为："),
        ("Equation", ("uₜₐₛ⁽ʳ⁾ = dₜmₜₐₛ⁽ʳ⁾,   dₜ ∈ {−1,+1}", 1)),
        ("Equation", ("μₜₐ = S⁻¹∑ₛuₜₐₛ⁽ᵛᵃˡ⁾", 2)),
        ("Equation", ("SEₜₐ = SD(uₜₐ·⁽ᵛᵃˡ⁾)/√S", 3)),
        ("Equation", ("Aₜ¹ˢᴱ = {a: μₜₐ ≥ μₜ* − SEₜ*}", 4)),
        (
            "Normal",
            "其中 dₜ 统一指标方向，a* 为验证均值最高的候选。式 (4) 只把与最优候选差异不超过一个标准误的候选送入平局处理，"
            "避免把抽样波动解释为确定性优势。稳定性、风险调整分数和最终冻结规则分别为：",
        ),
        ("Equation", ("Stabₜ(a) = S⁻¹∑ₛ𝟙(aₜₛ* = a)", 5)),
        ("Equation", ("qₜₐ = μₜₐ − λSD(uₜₐ·⁽ᵛᵃˡ⁾),   λ = 0.5", 6)),
        ("Equation", ("aₜ* = LexMinₐ∈Aₜ¹ˢᴱ(SDₜₐ, Calₜₐ, Costₐ, IDₐ)", 7)),
        (
            "Normal",
            "词典序依次比较波动、校准损失、计算成本和 candidate_id；风险调整策略作为独立固定对照，不替代式 (7)。"
            "外层测试标签仅用于下列审计量：",
        ),
        ("Equation", ("Rₜₛ = maxₐuₜₐₛ⁽ᵗᵉˢᵗ⁾ − uₜₐₜ*ₛ⁽ᵗᵉˢᵗ⁾", 8)),
        ("Equation", ("R̃ₜₛ = Rₜₛ/[maxₐuₜₐₛ⁽ᵗᵉˢᵗ⁾ − minₐuₜₐₛ⁽ᵗᵉˢᵗ⁾]", 9)),
        ("Equation", ("Oₜₛ = uₜₐₜ*ₛ⁽ᵛᵃˡ⁾ − uₜₐₜ*ₛ⁽ᵗᵉˢᵗ⁾", 10)),
        ("Equation", ("Hit@kₜₛ = 𝟙[aₜₛᵒʳᵃᶜˡᵉ ∈ Top-kₜₛ⁽ᵛᵃˡ⁾]", 11)),
        (
            "Normal",
            "式 (8)-(11) 分别给出测试遗憾、归一化遗憾、乐观偏差和排序命中率。测试集事后最优候选只定义评价上界，"
            "不得反向修改候选池、λ 或平局规则。样本级风险与选择性预测定义为：",
        ),
        ("Equation", ("rᵢ = ∑ⱼλⱼzᵢⱼ,   λⱼ ≥ 0,   ∑ⱼλⱼ = 1", 12)),
        ("Equation", ("I(c) = {i: rank(rᵢ) ≤ ⌈cN⌉}", 13)),
        ("Equation", ("Risk(c) = |I(c)|⁻¹∑ᵢ∈I(c)ℓ(yᵢ,ŷᵢ)", 14)),
        (
            "Normal",
            "zᵢⱼ 表示标准化的模型分歧、交叉拟合误差分数、1−Tanimoto 相似度或重构误差；权重仅在训练/验证侧拟合。"
            "保形阈值及回归区间、分类集合写为：",
        ),
        ("Equation", ("qα = Quantile⌈(ncal+1)(1−α)⌉({sᵢ}ᵢ∈cal)", 15)),
        ("Equation", ("Cαʳᵉᵍ(x) = [ŷ(x)−qα, ŷ(x)+qα]", 16)),
        ("Equation", ("Cαᶜˡˢ(x) = {y: s(x,y) ≤ qα}", 17)),
        (
            "Normal",
            "本文在 α=0.20、0.10 和 0.05 下分别评价 80%、90% 和 95% 标称覆盖。公式只规定可复核的决策接口；"
            "低相似度三档、固定精度召回、活性悬崖分子对和片段富集的实现细节与阈值均在补充方法和图表源数据中保留。",
        ),
    ]

    anchor = start
    for style, payload in content:
        paragraph = insert_after(anchor, style=style)
        if style == "Equation":
            formula, number = payload  # type: ignore[misc]
            set_formula(paragraph, formula, number)
        else:
            d9.replace_paragraph_text(paragraph, str(payload))
        anchor = paragraph


def resize_table(table, n_rows: int) -> None:
    while len(table.rows) > n_rows:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < n_rows:
        table.add_row()


def replace_figure(paragraph: Paragraph, image_path: Path) -> None:
    drawings = paragraph._p.xpath(".//w:drawing")
    if not drawings:
        raise RuntimeError("Target figure paragraph contains no drawing")
    rel_ids = {
        blip.get(qn("r:embed"))
        for blip in paragraph._p.xpath(".//a:blip")
        if blip.get(qn("r:embed"))
    }
    for rel_id in rel_ids:
        paragraph.part.drop_rel(rel_id)
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.add_run().add_picture(str(image_path), width=Cm(16.2))


def rebuild_tables(doc: Document, prospective: pd.DataFrame, stability: pd.DataFrame) -> None:
    matrices = copy.deepcopy(d9.TABLE_MATRICES)
    captions = copy.deepcopy(d9.TABLE_CAPTIONS)
    notes = copy.deepcopy(d9.TABLE_NOTES)
    widths = copy.deepcopy(d9.TABLE_WIDTH_WEIGHTS)

    matrices[1][3] = [
        "32 候选嵌套验证",
        "9 个终点；4/8/16/32 候选",
        "3 外层×3 内层\n骨架分组",
        "27 个外层单元",
        "遗憾/命中率/稳定性",
        "选择可靠性主检验",
    ]

    matrices[2] = [
        ["策略", "选择依据", "跨折形式", "测试标签用途", "完成状态", "证据解释"],
        ["固定单模型", "登记首个基线", "固定", "仅评价", "已完成", "低搜索自由度对照"],
        ["验证集最优", "内层验证均值", "允许切换", "仅评价", "已完成", "朴素选择对照"],
        ["单标准误+稳定性", "均值容差、波动与成本", "固定规则", "仅评价", "已完成", "保守冻结规则"],
        ["风险调整", "验证均值−0.5×标准差", "固定规则", "仅评价", "已完成", "全局固定 λ"],
        ["Top-K/堆叠", "内层折外预测融合", "按折冻结", "仅评价", "已完成", "候选族对照"],
        ["AutoGluon-Tabular", "独立内层骨架调优", "外层一次评价", "仅评价", "已完成", "CPU 树模型 AutoML 对照"],
        ["随机期望", "候选等概率期望", "随机期望", "仅评价", "已完成", "搜索下界"],
        ["测试集事后最优", "测试效用最大", "事后计算", "仅作上界", "已完成", "不得用于晋级"],
    ]
    captions[2] = "候选选择策略、AutoML 对照及测试标签使用边界"
    notes[2] = (
        "注：AutoGluon-Tabular 使用 Morgan-512、LightGBM/CatBoost/随机森林/极端随机树、每个外层折 30 s 调优预算和全量重拟合；"
        "测试集事后最优只用于评价。"
    )

    rows = [["候选数", "Top-3 命中率", "验证最优遗憾", "单标准误遗憾", "风险调整遗憾", "验证最优稳定性"]]
    for pool_size in (4, 8, 16, 32):
        p = prospective[prospective.pool_size == pool_size].set_index("policy")
        top3 = p.loc["validation_best", "top3_hit_rate"]
        stab = stability[
            (stability.pool_size == pool_size) & (stability.policy == "validation_best")
        ].modal_selection_rate.mean()
        rows.append(
            [
                str(pool_size),
                f"{top3:.3f}",
                f"{p.loc['validation_best', 'normalized_regret_mean']:.3f}",
                f"{p.loc['one_se_stable', 'normalized_regret_mean']:.3f}",
                f"{p.loc['risk_adjusted', 'normalized_regret_mean']:.3f}",
                f"{stab:.3f}",
            ]
        )
    matrices[3] = rows
    captions[3] = "32 候选前瞻性冻结协议的 3×3 嵌套压力结果"
    notes[3] = (
        "注：共 9 个终点、27 个外层测试单元；候选数为同一冻结登记的前 4/8/16/32 项。"
        "遗憾按各外层候选效用范围归一化；稳定性为各终点外层折主导候选比例的均值。"
    )
    widths[3] = [0.8, 1.25, 1.2, 1.2, 1.2, 1.35]

    matrices[9] = [
        ["研究方向", "近期代表工作", "已实现内容", "本研究的差异", "主张边界"],
        ["现实挑战基准", "Zhao 等[5]", "稀缺/OOD/不平衡/bRo5/悬崖", "将场景用于选择器压力审计", "不主张挑战集合首创"],
        ["预训练与多模态", "Zhang、Jang 等[24-26]", "预训练、指纹与图表示融合", "不以新表示作为核心贡献", "不主张融合首创"],
        ["OOD、适用域与不确定性", "Wen、Yin、Uchibori 等[27-31]", "OOD 学习、域边界与 UQ", "证据写入冻结选择决策", "不主张 AD/UQ 首创"],
        ["分子对与活性悬崖", "Fralish 和 Reker[32]", "分子对差异学习与优化", "用于定位选择失败条件", "不主张分子对分析首创"],
        ["候选治理", "本研究", "4/8/16/32 受控扩池与外层审计", "遗憾、命中、稳定性与三态门控", "结论限于公开离线基准"],
    ]
    captions[9] = "近期相关研究与四项创新主张的边界"
    notes[9] = "注：检索截至 2026 年 6 月；“未发现完整覆盖”不等同于证明绝对首次。"
    widths[9] = [1.15, 1.65, 2.1, 2.25, 1.55]

    d9.TABLE_MATRICES = matrices
    d9.TABLE_CAPTIONS = captions
    d9.TABLE_NOTES = notes
    d9.TABLE_WIDTH_WEIGHTS = widths

    caption_paragraphs = [p for p in doc.paragraphs if p.style and p.style.name == "TableCaption"]
    note_paragraphs = [p for p in doc.paragraphs if p.style and p.style.name == "TableNote"]
    if len(caption_paragraphs) != 9 or len(note_paragraphs) != 9 or len(doc.tables) != 9:
        raise RuntimeError("Expected nine main tables with captions and notes")

    for table_no, table in enumerate(doc.tables, start=1):
        resize_table(table, len(matrices[table_no]))
        d9.format_caption(caption_paragraphs[table_no - 1], f"表 {table_no} | ", captions[table_no])
        d9.format_table(table, table_no)
        d9.format_note(note_paragraphs[table_no - 1], notes[table_no])


def load_results() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, str]:
    policy_detail = pd.read_csv(EXPANDED / "policy_detail.csv")
    policy_summary = pd.read_csv(EXPANDED / "policy_summary.csv")
    stability = pd.read_csv(EXPANDED / "selection_stability.csv")
    prospective = policy_summary.copy()

    ag_detail = pd.read_csv(AUTOGLUON / "outer_results.csv")
    ag_summary = pd.read_csv(AUTOGLUON / "summary.csv")
    if len(ag_detail) != 27 or ag_detail["dataset"].nunique() != 9:
        raise RuntimeError("AutoGluon baseline is not complete for 9 endpoints x 3 outer folds")
    fzyc = policy_detail[
        (policy_detail.pool_size == 32) & (policy_detail.policy == "validation_best")
    ][["dataset", "outer_fold", "outer_utility"]].rename(columns={"outer_utility": "fzyc_utility"})
    paired = ag_detail[["dataset", "outer_fold", "outer_utility"]].rename(
        columns={"outer_utility": "autogluon_utility"}
    ).merge(fzyc, on=["dataset", "outer_fold"], validate="one_to_one")
    endpoint_delta = paired.groupby("dataset", as_index=False).agg(
        autogluon_utility=("autogluon_utility", "mean"),
        fzyc_utility=("fzyc_utility", "mean"),
    )
    endpoint_delta["delta"] = endpoint_delta.autogluon_utility - endpoint_delta.fzyc_utility
    wins = int((endpoint_delta.delta > 1e-9).sum())
    losses = int((endpoint_delta.delta < -1e-9).sum())
    ties = 9 - wins - losses
    classification = ag_summary[ag_summary.task_type == "classification"].primary_mean
    regression = ag_summary[ag_summary.task_type == "regression"].primary_mean
    ag_text = (
        f"AutoGluon-Tabular CPU 树模型对照在 5 个分类终点上的平均外层 ROC-AUC 范围为 "
        f"{classification.min():.3f}-{classification.max():.3f}，在 4 个回归终点上的 RMSE 范围为 "
        f"{regression.min():.3f}-{regression.max():.3f}。按终点平均外层效用与 32 候选验证最优策略配对比较，"
        f"AutoGluon 为 {wins}/{ties}/{losses} 个胜/平/负。该对照采用单个内层骨架调优折和固定时间预算，"
        "用于检验结论是否依赖手工候选，不作为 FZYC-Mol 治理协议的替代。"
    )
    return prospective, stability, ag_detail, ag_summary, ag_text


def append_recent_references(doc: Document) -> None:
    references = [
        "[24] Zhang L, Zeng Y, Qi Y, et al. DCPM-ADMET: fusion of dual-component pre-trained model and molecular fingerprints to enhance drug ADMET properties prediction. J Cheminform. 2026. doi:10.1186/s13321-026-01244-z.",
        "[25] Jang Y, Lee J, Jeong K, Kim J. Multimodal graph fusion with statistically guided parsimonious descriptor selection for molecular property prediction. J Cheminform. 2026;18:18. doi:10.1186/s13321-025-01140-y.",
        "[26] Zhang Y, Liu W, Zhao H, et al. MolGramTreeNet: a multimodal molecular property prediction model via grammar tree-constrained molecular representation. iScience. 2026;29:114928. doi:10.1016/j.isci.2026.114928.",
        "[27] Wen X, Liu H, Long W, Wei S, Zhu R. Consistent semantic representation learning for out-of-distribution molecular property prediction. Brief Bioinform. 2025;26:bbaf147. doi:10.1093/bib/bbaf147.",
        "[28] Yin T, Gao P, Panapitiya G, Saldanha EG. Out-of-distribution evaluation of active learning pipelines for molecular property prediction. RSC Adv. 2026;16:5281-5295. doi:10.1039/D5RA08055J.",
        "[29] Uchibori Y, Kaneko H. Generation of molecules near the applicability domain boundaries of property prediction models. J Chem Inf Model. 2026. doi:10.1021/acs.jcim.5c03220.",
        "[30] Kim JY, Vlachos DG. Distance-aware molecular property prediction in nonlinear structure-property space. J Chem Inf Model. 2025;65:6744-6756. doi:10.1021/acs.jcim.5c01037.",
        "[31] Tang H, Yue T, Li Y. Assessing uncertainty in machine learning for polymer property prediction: a benchmark study. J Chem Inf Model. 2025;65:6585-6598. doi:10.1021/acs.jcim.5c00550.",
        "[32] Fralish Z, Reker D. Pairwise learning for molecular property prediction and optimization. Front Drug Discov. 2026;6:1859068. doi:10.3389/fddsv.2026.1859068.",
    ]
    existing = "\n".join(p.text for p in doc.paragraphs)
    if "[24] Zhang L" in existing:
        return
    for text in references:
        paragraph = doc.add_paragraph(style="Normal")
        d9.replace_paragraph_text(paragraph, text)
        paragraph.paragraph_format.left_indent = Cm(0.63)
        paragraph.paragraph_format.first_line_indent = Cm(-0.63)
        paragraph.paragraph_format.space_after = Pt(2)


def expand_citation(content: str) -> list[int]:
    values: list[int] = []
    for token in content.split(","):
        token = token.strip()
        if not token:
            continue
        if "-" in token:
            start, end = (int(x) for x in token.split("-", 1))
            values.extend(range(start, end + 1))
        else:
            values.append(int(token))
    return values


def compress_citation(values: list[int]) -> str:
    ordered = sorted(set(values))
    chunks: list[str] = []
    start = previous = ordered[0]
    for value in ordered[1:]:
        if value == previous + 1:
            previous = value
            continue
        chunks.append(str(start) if start == previous else f"{start}-{previous}")
        start = previous = value
    chunks.append(str(start) if start == previous else f"{start}-{previous}")
    return ",".join(chunks)


def renumber_references(doc: Document) -> None:
    citation_pattern = re.compile(r"\[([0-9]+(?:\s*[-,]\s*[0-9]+)*)\]")

    reference_paragraphs = [p for p in doc.paragraphs if re.match(r"^\[\d+\]", p.text.strip())]
    reference_text: dict[int, str] = {}
    for paragraph in reference_paragraphs:
        match = re.match(r"^\[(\d+)\]\s*(.*)$", paragraph.text.strip())
        if not match:
            raise RuntimeError(f"Malformed reference: {paragraph.text}")
        reference_text[int(match.group(1))] = match.group(2)

    # Number references by their first appearance in the main text. All
    # references are cited in prose; table citations are rewritten below but
    # do not determine the first-appearance order.
    reference_elements = {p._p for p in reference_paragraphs}
    order: list[int] = []
    for paragraph in doc.paragraphs:
        if paragraph._p in reference_elements:
            continue
        for match in citation_pattern.finditer(paragraph.text):
            for value in expand_citation(match.group(1).replace(" ", "")):
                if value not in order:
                    order.append(value)
    order.extend(sorted(set(reference_text) - set(order)))
    if set(order) != set(reference_text):
        raise RuntimeError("Citation and reference sets differ")
    mapping = {old_number: new_number for new_number, old_number in enumerate(order, start=1)}

    def rewrite(text: str) -> str:
        def repl(match: re.Match[str]) -> str:
            old_values = expand_citation(match.group(1).replace(" ", ""))
            new_values = [mapping[value] for value in old_values]
            return f"[{compress_citation(new_values)}]"

        return citation_pattern.sub(repl, text)

    for paragraph in doc.paragraphs:
        if paragraph._p in reference_elements:
            continue
        new_text = rewrite(paragraph.text)
        if new_text != paragraph.text:
            d9.replace_paragraph_text(paragraph, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text = rewrite(paragraph.text)
                    if new_text != paragraph.text:
                        if paragraph.runs:
                            paragraph.runs[0].text = new_text
                            for run in paragraph.runs[1:]:
                                run.text = ""
                        else:
                            paragraph.add_run(new_text)

    for paragraph in reference_paragraphs:
        delete_paragraph(paragraph)
    for old_number in order:
        new_number = mapping[old_number]
        paragraph = doc.add_paragraph(style="Normal")
        d9.replace_paragraph_text(paragraph, f"[{new_number}] {reference_text[old_number]}")
        paragraph.paragraph_format.left_indent = Cm(0.63)
        paragraph.paragraph_format.first_line_indent = Cm(-0.63)
        paragraph.paragraph_format.space_after = Pt(2)


def add_contributions(doc: Document) -> None:
    paragraph = find_paragraph(doc, "FZYC-Mol 将候选登记")
    d9.replace_paragraph_text(
        paragraph,
        "本研究的创新不建立在新的主干网络、融合表示或适用域算法上，而由以下四项相互约束的证据构成。",
    )
    contributions = [
        "（1）将候选池规模设为受控变量。在同一冻结登记中依次扩展 4、8、16 和 32 个候选，并以 Top-k 命中率、测试遗憾、乐观偏差和选择稳定性直接评价“能否选对”，而非只比较最终分数。",
        "（2）建立冻结登记与接受、保留、拒绝三态治理。指标方向、候选顺序、单标准误容差、稳定性平局规则和风险惩罚在外层测试前固定，测试集事后最优只保留为审计上界。",
        "（3）在同一外层单元联合评价选择可信度与预测可信度。选择层报告遗憾、命中与稳定性，预测层报告适用域、固定精度召回、校准、风险-覆盖和 80%/90%/95% 保形覆盖。",
        "（4）把低相似度、活性悬崖、bRo5 和类别不平衡转化为选择器压力测试。困难场景用于定位候选何时应晋级、保留或拒绝，而不被表述为新的数据集或普遍性能优势。",
    ]
    anchor = paragraph
    for text in contributions:
        anchor = insert_after(anchor, text, "Normal")


def update_manuscript(doc: Document, prospective: pd.DataFrame, stability: pd.DataFrame, ag_text: str) -> None:
    p32 = prospective[prospective.pool_size == 32].set_index("policy")
    top3_values = {
        int(k): float(v)
        for k, v in prospective[prospective.policy == "validation_best"]
        .set_index("pool_size")["top3_hit_rate"]
        .items()
    }
    stab_values = {
        pool: stability[(stability.pool_size == pool) & (stability.policy == "validation_best")]
        .modal_selection_rate.mean()
        for pool in (4, 8, 16, 32)
    }

    abstract = (
        "分子性质预测常在同一验证集上比较不断扩张的表示、模型和融合候选，反复查询可能使验证集演变为隐性开发集。"
        "我们提出 FZYC-Mol，一套由冻结候选登记、内外层隔离、单标准误容差、稳定性与风险调整规则以及证据化决策卡组成的选择治理框架。"
        "评估覆盖 6 个 MoleculeNet 任务、22 个 TDC ADMET 终点、17 个 MoleculeACE 活性悬崖任务及 CycPept-PAMPA 和 LinPept bRo5 数据。"
        "回顾性审计中，候选数由 4 增至 32 时，测试集事后最优进入验证前 3 名的比例由 0.897 降至 0.333。"
        "在 9 个终点、32 个冻结轻量候选的 3×3 嵌套重训中，该比例由 0.926 降至 0.222，验证集最优策略的外层选择稳定性由 0.926 降至 0.444；"
        "其归一化测试遗憾在 4/8/16/32 候选时为 0.167/0.184/0.151/0.165，未随扩池单调改善。"
        "在 32 候选时，固定风险调整与验证集最优策略的遗憾分别为 0.165 和 0.165，说明保守规则提高部分终点稳定性但不保证降低平均遗憾。"
        "TDC 冻结策略得到 5 个晋级、17 个保留、0 个下降；分类和回归保形预测在 80%/90%/95% 标称水平下的平均覆盖率分别为 "
        "0.814/0.918/0.956 和 0.823/0.925/0.962。结果表明，候选池扩张会增加验证排序不确定性；冻结协议、外层审计和负结果保留可界定选择风险，"
        "但证据限于公开离线基准，不支持普遍最优或临床部署主张。"
    )
    replace_prefix(doc, "分子性质预测通常通过", abstract)

    add_contributions(doc)
    replace_prefix(
        doc,
        "近期 Zhao",
        "近期研究已分别覆盖现实挑战、强表示和可靠性估计。Zhao 等系统比较了稀缺、分布外、类别不平衡、bRo5 和活性悬崖场景[5]；"
        "DCPM-ADMET、KROVEX 与 MolGramTreeNet 进一步发展了预训练、指纹、描述符和图表示融合[24-26]。"
        "OOD 表征学习、主动学习、适用域边界和不确定性基准也在快速推进[27-31]，分子对学习则直接建模局部性质差异[32]。"
        "这些工作说明困难场景、多模态融合、适用域和不确定性本身均不能作为本文的首创主张。",
    )
    replace_prefix(
        doc,
        "本研究据此提出三个问题",
        "本研究据此提出四个相互连接的问题：候选池按预定顺序扩张时，验证排序命中和外层选择稳定性如何变化；"
        "预先冻结的单标准误、稳定性和风险调整规则能否降低选择损失；选择可信度能否与适用域、校准和保形覆盖在同一外层单元联合审计；"
        "低相似度、活性悬崖、bRo5 和不平衡场景能否定位晋级、保留或拒绝的边界。FZYC-Mol 因而是选择治理协议，而不是新的主干预测网络。",
    )
    replace_prefix(
        doc,
        "主要假设不是候选增加",
        "主要假设不是候选增加必然降低预测性能，而是固定验证信息下的搜索自由度扩张会增加排序不确定性，使测试遗憾不再保证单调下降。"
        "我们以回顾性候选池压力为发现性证据，以 32 候选的 3×3 嵌套重训为确认性主检验，再用 AutoGluon、TDC、MoleculeACE、bRo5、"
        "低相似度分层和保形预测检验结论是否跨选择器与化学边界保持一致。",
    )

    rebuild_formula_section(doc)
    replace_prefix(
        doc,
        "真正嵌套验证覆盖",
        "核心嵌套实验覆盖 BBBP、BACE、ClinTox、ESOL、FreeSolv、Lipophilicity、Caco2、HIA 和 Pgp 九个终点。"
        "每个终点采用 3 个外层骨架分组折和 3 个内层骨架分组折；32 个预先登记的 Morgan-512 轻量候选覆盖线性模型、随机森林、"
        "极端随机树、梯度提升、LightGBM、XGBoost 和 CatBoost 的固定超参数变体。候选池依登记顺序截取前 4、8、16 和 32 项，"
        "所有候选均在每个外层训练集重拟合，外层测试标签只用于计算性能、测试遗憾和事后上界。",
    )
    replace_prefix(
        doc,
        "候选池压力分析复用",
        "回顾性候选池压力分析保留为独立发现性证据：它复用历史验证与留出测试预测，并按相同登记顺序比较 4/8/16/32 项。"
        "前瞻性冻结协议的嵌套重训则从原始特征重新拟合全部 32 个轻量候选。两类分析分别回答历史结果是否存在扩池信号，以及该信号在严格内外层隔离后能否复现。",
    )
    replace_prefix(
        doc,
        "跨任务测试遗憾按",
        "跨任务测试遗憾按同一终点、外层折和候选池中的测试效用范围归一化；分母为零的单元不进入归一化汇总。"
        "每个池规模比较固定单模型、验证集最优、单标准误加稳定性、风险调整和测试集事后最优上界。Top-3 命中率以测试事后最优候选是否进入内层验证前 3 名定义。",
    )
    replace_prefix(
        doc,
        "选择器比较以终点-随机种子",
        "统计单位为终点-外层折而非候选行。选择稳定性先在每个终点内计算 3 个外层折中主导候选的比例，再跨 9 个终点平均；"
        "策略遗憾对 27 个外层测试单元汇总。由于同一终点内折并非独立数据集，正文以效应量、分布和终点异质性为主，不把 27 个折直接解释为 27 个独立生物学重复。",
    )
    method_anchor = find_paragraph(doc, "统计单位为终点-外层折")
    insert_after(
        method_anchor,
        "AutoGluon-Tabular 对照使用相同 Morgan-512 特征和外层骨架分组折；每个外层训练集再划出一个独立内层骨架调优折，"
        "候选限定为 LightGBM、CatBoost、随机森林和极端随机树，时间预算为每个外层折 30 s，随后以 refit_full 在外层训练数据上重拟合[22]。"
        "该对照衡量自动化树模型选择，而非完整深度学习 AutoML 或 FZYC-Mol 治理规则。",
        "Normal",
    )

    replace_prefix(
        doc,
        "本文依照 Journal of Cheminformatics",
        "本文依照 Journal of Cheminformatics 的可重复性要求[6]，将主要创新限定为四项选择治理贡献，并把表示融合、适用域、不确定性和困难场景明确列为既有方法或压力载体。"
        "由于缺少独立且带时间戳的 ADMET 盲测集，CycPept-PAMPA 时间划分只被解释为时间迁移压力，不替代前瞻性外部验证。",
    )
    replace_prefix(
        doc,
        "适用域（applicability domain",
        "适用域（applicability domain, AD）由测试分子与训练集最近邻的 Tanimoto 相似度、集成分歧、描述符距离和重构误差共同表征[20]。"
        "相似度阈值不用于删除测试样本，只用于条件性能和风险提示。样本被严格划分为 >0.7、0.5-0.7 和 <0.5 三档，并分别统计性能、不确定性、高误差率、误差富集、校准和风险-误差相关。",
    )
    replace_prefix(
        doc,
        "分类概率在独立校准集上",
        "分类概率在独立校准集上比较未校准、Platt 或温度缩放以及等距回归等方法，并依据 Brier 分数和 ECE 选择校准器[17-19]。"
        "分类保形预测采用标签条件非一致性分数，回归采用绝对残差分位数，标称覆盖率设为 80%、90% 和 95%。校准集与模型选择验证集相互独立，避免同一数据同时决定候选和覆盖阈值。",
    )
    replace_prefix(
        doc,
        "bRo5 评估比较随机划分",
        "bRo5 评估比较随机划分与更严格的骨架、外缘和时间划分[21]。CycPept-PAMPA 用于渗透性回归；LinPept CellPen 与 NonFouling 用于分类、概率校准和固定精度召回。"
        "鉴于这些数据与近期 ADMET 可靠性研究存在来源重叠[5]，本文将其作为选择器和适用域的压力测试，而非新的数据资源贡献。",
    )
    replace_prefix(
        doc,
        "MoleculeNet 结果报告 5 个随机种子",
        "MoleculeNet 结果报告 5 个随机种子的均值±标准差；TDC 和 bRo5 多数结果报告 3 个随机种子。嵌套验证或配对终点差异在 n=3 时以 t=4.303 计算 95% 置信区间，"
        "跨终点胜/平/负以终点为单位[23]。片段多重检验采用 FDR 控制；模型家族比较不以单个未经校正的 P 值宣称显著优越。",
    )
    replace_prefix(
        doc,
        "为限制多任务、多候选",
        "为限制多任务、多候选和多指标带来的叙事自由度，主要主张预先限定为候选池扩张下的 Top-3 命中率、测试遗憾、乐观偏差、选择稳定性和内外层隔离；"
        "次要主张涉及 MoleculeNet/TDC 冻结性能、AutoGluon 对照、风险-覆盖和保形覆盖；MoleculeACE 分子对、片段富集、3D-lite 和轻量适配器属于探索性证据。",
    )
    replace_prefix(
        doc,
        "摘要、结果、表格、图注和结论使用同一证据边界",
        "摘要、结果、表格、图注和结论使用同一证据边界。回顾性与嵌套重训共同支持“排序命中下降”和“测试遗憾非单调”，"
        "但 32 候选时风险调整与验证最优遗憾近乎相同，因此不支持“保守规则普遍降低外层遗憾”。TDC 的 5/17/0 支持保守晋级，"
        "保形总体覆盖支持边际有效性，二者均不自动推及时间外或低相似度子群。",
    )

    replace_prefix(doc, "3.2 真正嵌套验证", "3.2 32 候选嵌套重训复现选择困难")
    replace_prefix(
        doc,
        "九个终点均完成 3 外层",
        "九个终点均完成 3 外层×3 内层、32 个冻结轻量候选的嵌套重训，所有外层划分均保持骨架分组。"
        f"候选池由 4、8、16 扩至 32 时，测试集事后最优候选进入内层验证前 3 名的比例为 "
        f"{top3_values[4]:.3f}、{top3_values[8]:.3f}、{top3_values[16]:.3f} 和 {top3_values[32]:.3f}（表 3，图 4b）。"
        "该趋势与回顾性审计的 0.897 至 0.333 同向，并在严格外层隔离下下降得更明显。",
    )
    replace_prefix(
        doc,
        "嵌套面板仅包含 4 个轻量候选",
        "验证集最优策略的平均归一化测试遗憾在 4/8/16/32 候选时为 0.167/0.184/0.151/0.165，呈非单调变化。"
        f"32 候选时，单标准误、风险调整和验证最优策略的遗憾分别为 "
        f"{p32.loc['one_se_stable', 'normalized_regret_mean']:.3f}、{p32.loc['risk_adjusted', 'normalized_regret_mean']:.3f} 和 "
        f"{p32.loc['validation_best', 'normalized_regret_mean']:.3f}（图 4a）。风险调整未降低该规模的平均遗憾，这一负结果被原样保留。",
    )
    replace_prefix(
        doc,
        "随机种子嵌套审计覆盖",
        f"验证集最优策略的平均外层选择稳定性由 {stab_values[4]:.3f} 降至 {stab_values[8]:.3f}、"
        f"{stab_values[16]:.3f} 和 {stab_values[32]:.3f}（图 4c）。32 候选时，风险调整在 BBBP、HIA 等终点降低遗憾，"
        "但在 ClinTox、FreeSolv 和 Caco2 等终点恶化；Pgp 则由验证最优获益。终点异质性说明固定惩罚不能取代逐终点外层审计。",
    )
    replace_prefix(doc, "嵌套验证证实，外层隔离", ag_text)
    ag_paragraph = find_paragraph(doc, "AutoGluon-Tabular CPU 树模型对照")
    insert_after(
        ag_paragraph,
        "原四候选嵌套面板的绝对性能仍作为可追溯参照保留：BBBP、BACE、ClinTox、HIA 和 Pgp 的外层 ROC-AUC 分别为 "
        "0.900、0.895、0.793、0.917 和 0.938；ESOL、FreeSolv、Lipophilicity 和 Caco2 的外层 RMSE 分别为 "
        "1.153、2.071、0.859 和 0.462。该面板检验早期轻量登记，不能与扩展候选池的选择遗憾合并解释。",
        "Normal",
    )

    figure4_caption = (
        "图 4 | 32 候选的嵌套候选池压力。a，3×3 嵌套重训下不同选择策略的归一化外层测试遗憾；点为 27 个终点-外层折单元均值，误差条为 95% 区间。"
        "b，回顾性审计与嵌套重训的 Top-3 命中率。c，验证最优、单标准误和风险调整策略跨外层折的选择稳定性。"
        "d，32 候选时 9 个终点的策略遗憾，显示固定惩罚的终点异质性。测试集事后最优只用于评价。"
    )
    caption = replace_prefix(doc, "图 4 |", figure4_caption)
    paragraphs = list(doc.paragraphs)
    image_paragraph = paragraphs[paragraph_index(paragraphs, caption) - 1]
    replace_figure(image_paragraph, FIGURE4)

    replace_prefix(
        doc,
        "证据一致性审计区分了",
        "证据一致性审计区分了 FreeSolv 定向重建、Lipophilicity 补救头和预测融合；ClinTox 主结果统一为 ROC-AUC 0.950±0.026；"
        "TDC 的 5/17/0、32 候选嵌套结果和 AutoGluon 27 个外层单元均可由结构化 CSV 重建。独立 ADMET 时间外盲测仍未完成，未被表述为完成结果。",
    )
    replace_prefix(
        doc,
        "主文结果均连接到结构化文件",
        "主文结果均连接到结构化文件：图 3 对应回顾性候选池压力，图 4 对应 32 候选嵌套重训，图 5 对应 MoleculeNet、ClinTox 固定精度召回和排序审计，"
        "图 6 对应 TDC 22 终点；图 7-10 分别对应逐样本风险、保形预测与相似度分层、MoleculeACE 与 bRo5、统一消融。AutoGluon 的逐折排行榜、预测和运行清单单独归档。",
    )
    replace_prefix(
        doc,
        "当前已达到分析级复现",
        "当前已达到核心实验的分析级复现：32 候选嵌套重训、AutoGluon 外层对照、风险-覆盖、保形预测、相似度分层和化学压力测试均保存逐折或逐样本结果。"
        "尚未完成的是全部历史图模型与预训练候选的统一外层重训、独立 ADMET 时间外盲测和第三方冷启动复跑，因此本文不表述为全候选端到端复现。",
    )

    replace_prefix(
        doc,
        "本研究表明，候选池扩张首先改变",
        "回顾性审计与 32 候选嵌套重训共同表明，候选池扩张首先改变模型选择的可信度，而非简单决定最高分能否提高。"
        "两类分析的 Top-3 命中率均随 4/8/16/32 候选扩张下降；嵌套重训同时显示验证最优遗憾非单调、选择稳定性由 0.926 降至 0.444。"
        "这与固定验证信息下搜索自由度吸收抽样噪声的理论预期一致[3,4]。",
    )
    replace_prefix(
        doc,
        "FZYC-Mol 的贡献是一组",
        "FZYC-Mol 的四项贡献形成一条闭环：受控扩池量化选择难度，冻结三态治理约束晋级自由度，选择层与预测层可靠性联合记录，"
        "化学压力场景定位失效条件。单标准误、嵌套验证、适用域和保形预测均有统计或化学信息学先例；本文的贡献在于其组合、外层量化和可审计决策接口，而非重新命名既有方法。",
    )
    replace_prefix(
        doc,
        "与 Zhao 等的同刊研究相比",
        "近期文献核验进一步收紧了新颖性边界。Zhao 等已覆盖四类现实挑战和 AutoML[5]，预训练与多模态融合已有 DCPM-ADMET、KROVEX 和 MolGramTreeNet[24-26]，"
        "OOD、适用域和不确定性也有专门方法与基准[27-31]。本文可辩护的差异不是这些模块本身，而是把候选池规模作为受控变量，并将上述证据连接到测试遗憾、冻结晋级和三态决策。",
    )
    replace_prefix(
        doc,
        "治理规则不能消除全部选择误差",
        "治理规则不能消除全部选择误差。32 候选时，风险调整与验证最优的平均归一化遗憾分别为 0.165 和 0.165，近乎相同；"
        "单标准误策略为 0.201。风险调整提高了部分终点的稳定性，却在 ClinTox 等终点产生更大遗憾。该结果反驳了“更保守必然更优”的简单叙事，也说明外层审计必须保留。",
    )
    replace_prefix(
        doc,
        "本研究有四项主要限制",
        "本研究有四项主要限制。第一，32 候选嵌套池由 CPU 可复跑的轻量模型构成，不等于把全部历史图模型和预训练模型在每个外层折重训。"
        "第二，AutoGluon 对照采用单个内层骨架调优折和固定时间预算，不代表所有 AutoML 配置。第三，CycPept 时间划分不是独立 ADMET 时间外盲测。"
        "第四，公开仓库、许可证、环境锁定和 Zenodo DOI 尚待完成，当前证据包仍需第三方冷启动复核。",
    )
    replace_prefix(
        doc,
        "最优先的后续工作是在冻结候选版本后",
        "最优先的后续工作是获得带明确采集时间且标签定义兼容的独立 ADMET 数据，在冻结候选版本后进行一次性盲测；其次是在统一计算预算下把 Chemprop、"
        "图模型和冻结预训练表示纳入外层重训。只有时间外盲测、冷启动复跑和永久归档完成，当前选择治理证据才能进一步支持前瞻性应用主张。",
    )
    replace_prefix(
        doc,
        "因此，文稿的新颖性取决于",
        "因此，文稿的新颖性取决于以选择过程为对象的受控证据，而非模型清单长度。本次截至 2026 年 6 月的检索未发现同时覆盖受控扩池、"
        "冻结三态治理、选择与预测双层可靠性以及化学场景化失败定位的研究，但这一检索结论不能证明绝对首次；正文相应避免“首次”“唯一”和“普遍优于”等表述。",
    )

    replace_prefix(
        doc,
        "FZYC-Mol 将候选池扩张",
        "FZYC-Mol 将受控候选池扩张、冻结选择、外层审计和可靠性输出组织为可复核流程。回顾性审计和 32 候选嵌套重训均显示 Top-3 命中率随候选池扩大下降；"
        "嵌套重训进一步显示验证最优遗憾非单调、选择稳定性降低，且固定风险调整并不在所有规模或终点占优。",
    )
    replace_prefix(
        doc,
        "负结果是治理流程的组成部分",
        "保留既有结果和负结果使证据链闭合：MoleculeNet/TDC 给出冻结性能，AutoGluon 检验手工候选依赖，风险-覆盖与保形预测界定样本级可信度，"
        "MoleculeACE、bRo5、低相似度和失败案例定位化学边界；TDC 的 17 个保留终点和 32 候选风险调整负结果均未被事后删除。",
    )
    replace_prefix(
        doc,
        "本文最稳健的结论是",
        "本文最稳健的结论是：扩大候选池会增加验证排序不确定性，冻结规则和外层审计能够暴露选择风险，但没有单一保守策略保证最低遗憾。"
        "在独立 ADMET 时间外盲测、全部重型候选统一外层重训和永久开放归档完成前，证据范围限定于公开离线基准，不支持普遍最优或临床部署主张。",
    )
    replace_prefix(
        doc,
        "数据清洗、候选训练、嵌套验证",
        "数据清洗、32 候选嵌套重训、AutoGluon-Tabular 对照、候选池压力、统计分析和 Python 绘图脚本已形成可运行工作流。"
        "公开发行版仍应补充环境锁定文件、许可证、命令行入口和持续集成测试；仓库地址与发布标签尚待作者补充：[代码仓库地址与发布标签]。",
    )
    replace_prefix(
        doc,
        "补充材料与图表源数据逐项对应",
        "补充材料与图表源数据逐项对应：表 S1 记录数据版本与清洗；表 S2 记录完整候选登记、超参数和软件版本；表 S3 提供 MoleculeNet/TDC 逐任务结果；"
        "表 S4 提供回顾性与 32 候选嵌套的遗憾、乐观偏差、Top-k 和稳定性；表 S5 提供 AutoGluon 逐折排行榜、校准、AURC/E-AURC、保形预测和相似度分层；"
        "表 S6 提供 MoleculeACE 与 bRo5；表 S7 记录失败候选、拒绝原因和计算成本。",
    )

    rebuild_tables(doc, prospective, stability)
    append_recent_references(doc)
    renumber_references(doc)


def final_format(doc: Document) -> None:
    d9.configure_styles(doc)
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else "Normal"
        if style_name == "Equation":
            paragraph.paragraph_format.widow_control = True
            continue
        if style_name == "Normal":
            paragraph.paragraph_format.widow_control = True
            if re.match(r"^\[\d+\]", paragraph.text.strip()):
                paragraph.paragraph_format.left_indent = Cm(0.63)
                paragraph.paragraph_format.first_line_indent = Cm(-0.63)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    d9.set_run_font(run, 9.0)
            else:
                for run in paragraph.runs:
                    d9.set_run_font(run, 10.5)
        elif style_name.startswith("Heading"):
            size = 14 if style_name == "Heading 1" else 12 if style_name == "Heading 2" else 10.5
            for run in paragraph.runs:
                d9.set_run_font(run, size, bold=True, east_asia="黑体")
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)


def validate(doc: Document) -> None:
    text = "\n".join(p.text for p in doc.paragraphs)
    stale = [
        "AutoGluon 选择器未产生可核验输出",
        "完整大候选池尚未",
        "AutoGluon、完整大候选池",
        "将科学贡献限制为三句话",
    ]
    found = [item for item in stale if item in text]
    if found:
        raise RuntimeError(f"Stale claims remain: {found}")
    equation_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == "Equation")
    figure_count = sum(bool(p._p.xpath(".//w:drawing")) for p in doc.paragraphs)
    if equation_count != 17 or len(doc.tables) != 9 or figure_count != 10:
        raise RuntimeError(
            f"Unexpected structure: equations={equation_count}, tables={len(doc.tables)}, figures={figure_count}"
        )
    refs = [p.text for p in doc.paragraphs if re.match(r"^\[\d+\]", p.text.strip())]
    if len(refs) != 32:
        raise RuntimeError(f"Expected 32 references, found {len(refs)}")


def main() -> None:
    for path in (SOURCE, FIGURE4, AUTOGLUON / "outer_results.csv", AUTOGLUON / "summary.csv"):
        if not path.exists():
            raise FileNotFoundError(path)
    prospective, stability, _, _, ag_text = load_results()
    doc = Document(SOURCE)
    update_manuscript(doc, prospective, stability, ag_text)
    final_format(doc)
    validate(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    DESKTOP_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    doc.save(DESKTOP_OUTPUT)
    print(OUTPUT)
    print(DESKTOP_OUTPUT)


if __name__ == "__main__":
    main()
