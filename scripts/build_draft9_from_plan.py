from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def set_text(paragraph: Paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def block_element(block):
    return block._p if isinstance(block, Paragraph) else block._tbl


def insert_paragraph_after(block, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    block_element(block).addnext(new_p)
    parent = block._parent
    paragraph = Paragraph(new_p, parent)
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge_name, edge in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():
        tag = "w:" + edge_name
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        if edge is None:
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), edge.get("val", "single"))
            element.set(qn("w:sz"), str(edge.get("sz", 8)))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), edge.get("color", "000000"))


def make_three_line(table: Table) -> None:
    line = {"val": "single", "sz": 8, "color": "000000"}
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_borders(cell, top=None, bottom=None, left=None, right=None)
            if ri == 0:
                set_cell_borders(cell, top=line, bottom=line, left=None, right=None)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
            elif ri == len(table.rows) - 1:
                set_cell_borders(cell, top=None, bottom=line, left=None, right=None)


def insert_table_after(doc: Document, block, data: list[list[str]]) -> Table:
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    make_three_line(table)
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    block_element(block).addnext(tbl)
    return table


def find_para(doc: Document, predicate) -> Paragraph:
    for p in doc.paragraphs:
        if predicate(p.text.strip()):
            return p
    raise ValueError("paragraph not found")


def insert_sequence(doc: Document, anchor, items):
    current = anchor
    for item in items:
        if isinstance(item, str):
            current = insert_paragraph_after(current, item)
        elif isinstance(item, list):
            current = insert_table_after(doc, current, item)
        else:
            raise TypeError(item)
    return current


def main() -> None:
    desktop = Path.home() / "Desktop"
    source = max(
        desktop.glob("FZYC-Mol_初稿-8_Nature全稿语言逻辑终审版.docx"),
        key=lambda p: p.stat().st_mtime,
    )
    output = desktop / "FZYC-Mol_初稿-9.docx"
    doc = Document(str(source))

    # Abstract and introduction: align the central claim with Zhao et al. 2026
    # without turning the manuscript into a simple performance-winner claim.
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("分子性质预测是药物发现"):
            set_text(
                p,
                "分子性质预测是药物发现、ADMET 评估和毒性风险筛查中的关键计算环节。随机划分下的平均 ROC-AUC 或 RMSE 难以充分反映模型在新骨架、少样本终点、不平衡标签、bRo5 化学空间、活性悬崖、实验噪声和适用域漂移条件下的可靠性。受近期真实挑战 ADMET 基准研究启发 [3]，本文提出 FZYC-Mol，一种由验证集治理的适用域感知分子性质预测框架。该框架不依赖单一大型主干模型，而是把多视图表示、强基线专家、目标变换、融合候选、补救头、校准、粗糙度诊断和适用域证据纳入冻结候选池；最终策略仅由验证集决定，测试集只在策略冻结后用于一次性评估。",
            )
        elif text.startswith("本文贡献可概括为五点"):
            set_text(
                p,
                "本文贡献可概括为五点。第一，围绕 MoleculeNet、TDC ADMET、MoleculeACE、bRo5 肽类压力任务和结构外推划分构建统一评测流程，并明确训练集、验证集和测试集的使用边界。第二，将 TabPFNv2/AutoGluon/KPGT representation 等现代强基线、树模型、Chemprop、冻结表征、目标变换、采样策略和跨模态融合统一登记为预注册候选，而不是在测试集后追加模型。第三，提出可审计的 accept/reject/retain 治理机制，使候选策略必须先通过验证集证据才能进入主结果；未通过候选作为负结果保留。第四，报告验证-测试排名、代表性嵌套验证、random/scaffold/perimeter 梯度、低相似度分层、MoleculeACE 活性悬崖、不平衡分类、保形预测、校准、风险-覆盖曲线和 ROGI/MODI/SARI 粗糙度诊断，避免只依据单一 ROC-AUC 或 RMSE 下结论。第五，通过基序归因、片段富集、失败案例和计算成本审计，将模型性能、可靠性边界、负结果和可复现实验记录连接起来。",
            )

    # Methods: add the cross-paper challenge matrix.
    anchor = find_para(doc, lambda t: t.startswith("为保证多专家比较可复现"))
    insert_sequence(
        doc,
        anchor,
        [
            "为使补充实验与 Zhao et al. 2026 的真实挑战框架 [3] 形成直接对照，本文把新增实验组织为“挑战维度—候选池—验证门控—风险输出”的矩阵，而不是把每个模型作为孤立追加结果。M1、M2、M3、M5 和 M6 构成支撑主线结论的核心证据；M4、M7 和 M8 用于加强与真实药物发现场景的可比性和可复现性；M9 则提供样本级决策输出。所有模块均遵守测试集冻结原则：候选是否进入最终保留结果，只由验证集或预先定义的审计规则决定。",
            "表 1A. 对照参考论文驱动的补充实验矩阵与正文落点。",
            [
                ["模块", "证据层级", "挑战维度", "主要数据集", "候选/指标", "正文落点"],
                ["M1", "核心", "现代强基线同划分对照", "MoleculeNet 6 终点；ADMET 小样本/OOD 终点", "TabPFNv2、AutoGluon、XGBoost-RDKit/Mordred/MorganCount、Chemprop、FZYC selector；ROC-AUC、PR-AUC、RMSE/MAE/R2、regret、rank correlation", "3.5、4.1、4.5；Supplementary Note 3"],
                ["M2", "核心", "Random/scaffold/perimeter OOD 梯度", "BBBP、hERG、Mutagenicity、HLM、Oral Bioavailability、Caco2、Half-Life、VDss", "各 split 性能下降、最近邻 Tanimoto、selector regret、低相似度分层", "3.1、4.3；Supplementary Note 4"],
                ["M3", "核心", "类别不平衡决策场景", "ClinTox、Tox21-NR-ER、CYP2C9 substrate、CYP2D6 inhibition、hERG", "weighted、oversampling、downsampling、undersampling ensemble、calibrated ensemble；PR-AUC、Recall@Precision、F1、MCC、Brier、ECE", "3.5、4.4；Supplementary Note 5"],
                ["M4", "扩展", "bRo5/肽类外部压力", "CycPept-PAMPA、LinPept-NonFouling、LinPept-CellPen；可选 Macrocycle-PAMPA", "KPGT/GEM/Uni-Mol 表征、XGBoost、TabPFNv2、FZYC selector；R2/RMSE、ROC-AUC/PR-AUC、AD 分层", "4.2、5.3；Supplementary Note 6"],
                ["M5", "核心", "MoleculeACE 全量活性悬崖", "MoleculeACE 30 任务", "XGBoost-MorganCount、KPGT representation、MOLMCL/KPGT 可运行版本、FZYC selector；cliff-pair RMSE、gap Spearman、cliff recall", "4.6；Supplementary Note 6"],
                ["M6", "核心", "ROGI/MODI/SARI 选择风险", "MoleculeNet、TDC 回归、MoleculeACE、bRo5", "粗糙度与性能、rank correlation、regret、高误差率、AD 失败率的相关性", "3.6、4.4；Supplementary Note 7"],
                ["M7", "扩展", "跨模态/正交优势集成", "Caco2、Half-Life、VDss、CycPept-PAMPA、MoleculeACE 代表任务", "KPGT+GEM、KPGT+XGBoost、TabPFNv2+XGBoost、FZYC gated ensemble；融合前后性能、risk-coverage、AD 分层", "4.1、4.2、4.6"],
                ["M8", "扩展", "计算成本与可复现审计", "所有正式候选模型", "训练/推理时间、硬件、候选数量、失败率、time-performance Pareto", "4.5；数据和代码可用性"],
                ["M9", "样本级增强", "样本级可靠性决策卡", "每个终点 20–50 个代表分子", "预测、置信区间、AD 距离、风险分位、最近邻、片段解释、拒用标记", "4.7；Supplementary Note 8"],
            ],
        ],
    )

    anchor = find_para(doc, lambda t: t.startswith("粗糙度代理指标用于解释"))
    insert_sequence(
        doc,
        anchor,
        [
            "为避免粗糙度分析停留在附属解释层面，本文将 ROGI、MODI 和 SARI 明确作为选择器风险诊断，而不是直接作为模型加权项。ROGI 用于刻画性质差异与化学距离之间的不平滑程度，MODI 用于估计局部邻域标签混杂，SARI 用于度量相似分子对中的性质突变。本文主要检验这些指标与验证-测试 rank correlation、selector regret、高误差分位、AD 失败率和 MoleculeACE cliff-pair error 的关系；若相关性不稳定，则作为负结果报告，而不写入最终门控规则。",
        ],
    )


    anchor = find_para(doc, lambda t: t.startswith("MoleculeNet 主结果显示"))
    insert_sequence(
        doc,
        anchor,
        [
            "现代强基线同划分对照被作为 4.1 的独立证据块。该面板不只比较 CatBoost、XGBoost、LightGBM、ExtraTrees、RF 和 Chemprop，也把 TabPFNv2-RDKit、AutoGluon-RDKit、XGBoost-RDKit/Mordred/MorganCount 以及可获得的 KPGT representation 接入同一候选登记。其目标不是证明 FZYC-Mol 在每个终点击败所有强基线，而是检验当现代强基线进入同一验证集候选池后，最终保留策略是否仍能透明地接受、拒绝或保留旧基线。",
        ],
    )

    anchor = find_para(doc, lambda t: t.startswith("TDC ADMET 结果进一步说明"))
    insert_sequence(
        doc,
        anchor,
        [
            "为补齐 Zhao et al. [3] 中最有辨识度的真实挑战之一，外部 ADMET 结果需与 bRo5/肽类压力测试并列解释。CycPept-PAMPA、LinPept-NonFouling 和 LinPept-CellPen 分别代表环肽通透性、线性肽非特异性吸附和细胞穿透性；它们的分子量、构象柔性和低相似度外推难度均高于常规小分子面板。若 KPGT、GEM/Uni-Mol 或 TabPFNv2 在这些任务中未被验证集接受，结果仍应作为边界场景审计报告，而不是被删去以维持性能叙事。",
        ],
    )

    anchor = find_para(doc, lambda t: t.startswith("这一小节与 3.2 共同支撑外推结论"))
    insert_sequence(
        doc,
        anchor,
        [
            "Perimeter/max-min distance split 被加入为最严格的 OOD 梯度，使 random→scaffold→perimeter 的性能下降可以与对照论文 [3] 直接比较。该分析同时报告测试分子到训练/验证集合的最近邻 Tanimoto、低相似度分层性能和 selector regret。若某个终点在 perimeter split 下出现非单调变化，正文需优先解释样本组成、标签噪声和骨架分布，而不是简单归因为模型失败或模型胜出。",
        ],
    )

    anchor = find_para(doc, lambda t: t.startswith("可靠性结果显示"))
    insert_sequence(
        doc,
        anchor,
        [
            "类别不平衡专项被放入可靠性小节，而不是作为单独的性能表。ClinTox、Tox21-NR-ER、CYP2C9 substrate、CYP2D6 inhibition 和 hERG 同时报告 ROC-AUC、PR-AUC、Recall@Precision≥0.8/0.9、F1、MCC、Brier 和 ECE，并比较 weighted loss、oversampling、single downsampling、undersampling ensemble 与 calibrated ensemble。该设置更接近筛选决策场景：模型必须说明在高精度阈值下还能保留多少阳性检出能力，以及概率输出是否可用于风险分层。",
            "粗糙度诊断在本节承担选择风险解释功能。ROGI/MODI/SARI 与 risk decile enrichment、AD 失败率、MoleculeACE cliff-pair error 和验证-测试排名一致性一起报告。若粗糙度只与高误差样本相关、但不能稳定改善最终选择器，则应把它界定为低成本风险提示，而不是新的性能增强模块。",
        ],
    )

    anchor = find_para(doc, lambda t: t.startswith("当前同划分模型登记覆盖"))
    insert_sequence(
        doc,
        anchor,
        [
            "计算成本与可复现性审计被纳入候选登记。每个正式候选需记录训练时间、推理时间、CPU/GPU 或显存需求、随机种子、失败率、调参范围和默认参数。该表与性能表分开呈现，用于回答 FZYC-Mol 是否只是堆叠算力的问题；若某个高性能候选的时间成本或失败率显著高于轻量候选，则其进入最终保留结果需要更强的验证集证据。",
        ],
    )

    anchor = find_para(doc, lambda t: t.startswith("为形成完整的验证闭环"))
    set_text(
        anchor,
        "为形成完整的验证闭环，本文将关键补充证据纳入主文逻辑：先审计验证集选择是否存在排名偏差，再用配对统计和系统消融界定各模块贡献，随后在 random/scaffold/perimeter 梯度、低相似度、bRo5 肽类压力、MoleculeACE 全量活性悬崖、不平衡分类、ROGI/MODI/SARI 粗糙度诊断、计算成本和失败案例中检查适用边界。这些分析不改变测试集一次性报告原则，而是将“为什么接受、保留或拒绝某个候选”写成可追溯的证据链。",
    )

    anchor = find_para(doc, lambda t: t.startswith("MoleculeACE 配对结果与低相似度分析互相补充"))
    set_text(
        anchor,
        "MoleculeACE 配对结果与低相似度分析互相补充，并在本研究中从代表性案例扩展为 30 任务全量审计。主指标包括平均 R2/RMSE、cliff-pair RMSE、预测差异与真实差异的 gap Spearman、cliff recall 和 top cliff error；比较对象包括 XGBoost-MorganCount、KPGT representation、可运行的 MOLMCL/KPGT 版本、FZYC selector 以及 KPGT+XGBoost ensemble。已有代表性配对结果显示，活性悬崖目标候选在 51 个 seed 配对中的总 RMSE 平均正向变化为 0.0069，悬崖子集 RMSE 平均正向变化为 0.0056，但标准差较大；预测差异与真实差异的平均 Spearman 约为 0.252，部分任务接近零或为负。因此，本文将该结果解释为悬崖风险识别和候选治理的补充证据，而不是对活性悬崖预测已经解决的证明。",
    )

    anchor = find_para(doc, lambda t: t.startswith("不平衡分类与保形预测的补充结果改善了"))
    set_text(
        anchor,
        "不平衡分类与保形预测的补充结果改善了 ClinTox、DILI、hERG 和 CYP 底物等任务的解释。ROC-AUC 仍作为标准指标，但主文同时报告 PR-AUC、Recall@Precision≥0.8/0.9、F1、MCC、Brier、ECE、富集指标和样本级风险案例，并明确不同采样策略是否被验证集接受。保形预测已在 80%/90%/95% 目标覆盖率下完成：分类平均经验覆盖率为 0.814/0.918/0.956，回归平均经验覆盖率为 0.823/0.925/0.962；区间宽度和集合大小随覆盖目标增加而上升，符合保形预测的风险-信息量权衡。",
    )

    anchor = find_para(doc, lambda t: t.startswith("案例分析用于连接性能"))
    insert_sequence(
        doc,
        anchor,
        [
            "样本级决策卡作为可选增强，用于把预测结果转化为药物化学用户可检查的证据单元。每张决策卡应同时给出预测值或类别概率、保形区间或预测集、AD 距离、风险分位、最近邻结构及标签差异、片段解释、是否拒用和拒用原因。该输出不参与模型选择，但能把 FZYC-Mol 的可靠性主张从平均分数延伸到具体分子层面的可审查判断。",
        ],
    )

    # Discussion: make the comparison with Zhao et al. explicit and bounded.
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("与 MoleculeNet 时代的标准基准研究相比"):
            set_text(
                p,
                "Zhao et al. 2026 将 ADMET 可靠性组织为四类真实挑战：数据稀缺/OOD、类别不平衡、bRo5 化学空间和活性悬崖 [3]。该工作提供了一个高水平基准模板，强调 TabPFNv2、KPGT、Uni-Mol、AutoGluon、GNN 和传统机器学习在不同场景下的边界。与这种大规模 benchmark 不同，FZYC-Mol 的核心不是重新证明哪类模型总体最强，而是在候选池进入冻结测试之前，给出一个可审计的 accept/reject/retain 治理机制。",
            )
        elif text.startswith("这种定位也与近年高水平论文的证据组织方式相一致"):
            set_text(
                p,
                "这种差异化定位决定了本文的补强方向。M1–M8 将对照论文中的强基线、perimeter split、类别不平衡、bRo5、活性悬崖、粗糙度、跨模态集成和计算成本转化为 FZYC-Mol 的候选治理压力测试；M9 则把可靠性输出落到样本级决策卡。由此，本文不是把复杂模块直接写入主结论，而是让每个模块在同一验证证据下被接受、保留或拒绝。该写法比单纯宣称性能提升更稳健，也更符合当前结果中“选择性增益多于普遍胜出”的事实。",
            )
        elif text.startswith("此外，OmniMol"):
            set_text(
                p,
                "此外，OmniMol、ADMETlab 3.0 和近期 ADMET 表征基准提示，真实 ADMET 建模往往同时面对不完整标注、终点间相关性、平台化决策支持和跨数据集迁移问题 [47-49]。因此，FZYC-Mol 的价值不只在于判定哪个模型分数最高，更在于说明候选策略在什么边界内可用、为什么可用以及何时应被拒绝。外部 ADMET 附录、bRo5 压力测试、MoleculeACE 活性悬崖、ROGI/MODI/SARI 诊断和计算成本审计共同服务于这一点。",
            )
        elif text.startswith("本研究仍存在若干局限"):
            set_text(
                p,
                "本研究仍存在若干局限。首先，验证-测试排名审计和 9 个代表性终点的 3×3 nested validation 表明，验证集治理可以减少测试集事后选择，但内外层验证尚未覆盖全部 MoleculeNet、TDC、bRo5 和 MoleculeACE 终点，不能保证所有任务均达到测试最优；因此，小幅增益应与 regret、optimism gap、Top-3 命中和负结果共同解释。其次，收益具有明显终点异质性，BBBP、ClinTox、HIA 和 Pgp 等终点的增益较小，FreeSolv 仍落后于观测最佳 Chemprop 候选。第三，bRo5、MoleculeACE 和低相似度样本仍是主要失败边界，FZYC-Mol 更适合识别和审计这些风险，而不是宣称已解决它们。第四，基序归因、片段富集和粗糙度相关性仍属于关联证据，不能替代因果机制或湿实验验证。第五，ChemBERTa、MoLFormer、KPGT、GEM 和 Uni-Mol 的更充分微调仍需受控扩展，Polaris 与 OpenADMET 的完整官方挑战流程也需进一步验证。",
            )
        elif text.startswith("进一步验证应优先扩展外部评估附录"):
            set_text(
                p,
                "进一步验证应优先扩展真实挑战矩阵，而不是简单堆叠更多模型。更有价值的路线是在更多公开 ADMET、bRo5 和活性悬崖终点上系统评估 CatBoost、XGBoost、ExtraTrees、LightGBM、RF、TabPFNv2、AutoGluon、KPGT/GEM/Uni-Mol 表征、Top-K 集成、目标变换和欠采样集成，并把最终保留决策与粗糙度、低相似度、perimeter split 和验证-测试排名一致性关联起来。",
            )


    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("本文提出并系统评估了 FZYC-Mol"):
            set_text(
                p,
                "本文提出并系统评估了 FZYC-Mol，一种由验证集治理、适用域感知的多专家分子性质预测框架。结果表明，在 MoleculeNet、TDC ADMET、bRo5/肽类压力测试、划分真实性、低相似度子集、MoleculeACE 活性悬崖、粗糙度诊断、计算成本审计和基序/片段解释等多条证据线上，FZYC-Mol 能够提供比单一模型分数更完整的可靠性画像。其核心价值不是保证每个终点取得测试最优，而是在固定候选池内透明地接受、保留、拒绝和审计候选策略；当补救头、低成本重构、强基线或多方法融合通过验证集门控时，它们可以进入最终保留结果，当证据不足时则作为负结果保留。",
            )
        elif text.startswith("未来工作可沿三条路线推进"):
            set_text(
                p,
                "未来工作可沿三条路线推进。第一，继续扩展外部 ADMET、bRo5、MoleculeACE 和盲测风格验证，以检验选择器在更多真实挑战下的稳定性。第二，在少量代表性低分任务上尝试受控适配器或轻量微调，但仍需使用嵌套验证限制过拟合。第三，进一步加强样本级决策卡，把基序、最近邻、不确定性、适用域、粗糙度和实验标签噪声联系起来，为后续湿实验或专家审查提供更清晰的优先级。",
            )

    doc.save(str(output))
    print("source=", source)
    print("output=", output)
    print("paragraphs=", len([p for p in doc.paragraphs if p.text.strip()]))
    print("tables=", len(doc.tables))
    print("images=", len(doc.inline_shapes))


if __name__ == "__main__":
    main()
