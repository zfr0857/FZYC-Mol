from __future__ import annotations

import json
import zipfile
from io import BytesIO
from pathlib import Path

from docx import Document
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "work" / "small_paper5_source_audit"
SOURCES = {
    "draft7": Path(r"C:\Users\Administrator\Desktop\修改\初稿-7.docx"),
    "smallpaper4": ROOT / "output" / "小论文-4.docx",
    "fusion_spec": Path(r"C:\Users\Administrator\Downloads\FZYC-Mol_双文档全量融合说明书_内容实验结果图表版.docx"),
}


def extract(name: str, path: Path) -> dict[str, object]:
    doc = Document(path)
    lines = [f"# {name}", "", f"Source: `{path}`", "", "## Paragraphs", ""]
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.replace("\u000b", " ").strip()
        if text:
            lines.append(f"P{index:04d} [{paragraph.style.name}] {text}")
    lines.extend(["", "## Tables", ""])
    for table_index, table in enumerate(doc.tables):
        lines.append(f"### Table {table_index + 1} [{table.style.name if table.style else 'None'}]")
        lines.append("")
        for row in table.rows:
            cells = [" ".join(cell.text.split()).replace("|", "∣") for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append("")
    media = []
    with zipfile.ZipFile(path) as archive:
        for member in sorted(n for n in archive.namelist() if n.startswith("word/media/")):
            data = archive.read(member)
            record = {"member": member, "bytes": len(data)}
            try:
                image = Image.open(BytesIO(data))
                record.update({"width": image.width, "height": image.height, "format": image.format})
            except Exception:
                record.update({"width": None, "height": None, "format": None})
            media.append(record)
    (OUT / f"{name}.md").write_text("\n".join(lines), encoding="utf-8")
    summary = {
        "source": str(path),
        "bytes": path.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "sections": len(doc.sections),
        "media": media,
        "headings": [
            {"index": i, "style": p.style.name, "text": p.text.strip()}
            for i, p in enumerate(doc.paragraphs)
            if p.text.strip() and p.style.name.startswith("Heading")
        ],
        "figure_captions": [
            {"index": i, "style": p.style.name, "text": p.text.strip()}
            for i, p in enumerate(doc.paragraphs)
            if p.text.strip().startswith(("图 ", "Figure ", "Fig. "))
        ],
        "table_captions": [
            {"index": i, "style": p.style.name, "text": p.text.strip()}
            for i, p in enumerate(doc.paragraphs)
            if p.text.strip().startswith(("表 ", "Table "))
        ],
    }
    return summary


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    summaries = {name: extract(name, path) for name, path in SOURCES.items()}
    (OUT / "inventory.json").write_text(json.dumps(summaries, ensure_ascii=False, indent=2), encoding="utf-8")
    print(OUT)


if __name__ == "__main__":
    main()
