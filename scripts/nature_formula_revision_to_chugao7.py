from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE_NEEDLE = "近期文献格式规范修订版"
OUTPUT_NAME = "FZYC-Mol_初稿-7.docx"


def find_source() -> Path:
    desktop = Path.home() / "Desktop"
    candidates = sorted(
        desktop.glob("FZYC-Mol_*.docx"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    for path in candidates:
        if SOURCE_NEEDLE in path.name:
            return path
    if not candidates:
        raise FileNotFoundError("No FZYC-Mol docx file was found on Desktop.")
    return candidates[0]


def set_run_font(run, latin="Times New Roman", east_asia="宋体", size=10.5, bold=False):
    run.font.name = latin
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def style_normal(paragraph, text, bold=False):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(21)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.15
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)


def style_subhead(paragraph, text):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(3)
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=True)


def style_formula(paragraph, formula, number):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.space_before = Pt(2)
    fmt.space_after = Pt(2)
    fmt.line_spacing = 1
    try:
        fmt.tab_stops.clear_all()
    except Exception:
        pass
    fmt.tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    fmt.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.clear()
    run = paragraph.add_run(f"\t{formula}\t({number})")
    set_run_font(run, latin="Cambria Math", east_asia="Cambria Math", size=9.5)


def style_heading2(paragraph, text):
    paragraph.style = "Heading 2"
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=12, bold=True)


def build_formula_block():
    block = [
        ("heading2", "3.7 数学定义、公式编号与验证集治理"),
        (
            "normal",
            "为符合 Nature 系列期刊对公式可编辑性和可追踪性的要求，本节仅保留直接支撑 FZYC-Mol 方法学主张的编号公式。正文引用编号公式时采用 equation (n) 的形式，通用评价指标和统计检验细节放入 Supplementary Methods。设 t 表示预测终点，a 表示候选策略，x_i 和 y_i 分别表示第 i 个分子的结构输入和观测标签，所有候选池、阈值、容差和打破平局规则均在测试集评估前冻结。",
        ),
        ("subhead", "数据边界、特征视图与候选池。"),
        (
            "formula",
            "D_t = D_t^{tr} ∪ D_t^{val} ∪ D_t^{te},  D_t^r = {(x_i,y_i)}_{i=1}^{n_r},  r ∈ {tr,val,te}",
            1,
        ),
        (
            "formula",
            "ϕ_i = [ϕ_i^{fp}, ϕ_i^{desc}, ϕ_i^{graph}, ϕ_i^{pre}],  A_t = {a_1,...,a_{m_t}}",
            2,
        ),
        (
            "formula",
            "ẑ_{i,a} = f_{t,a}(ϕ_i),  S_t^r(a) = dir_t · M_t({y_i,ẑ_{i,a}}_{i∈D_t^r})",
            3,
        ),
        (
            "normal",
            "在 equation (1) 中，训练集、验证集和测试集互不重叠。Equation (2) 定义候选专家可使用的多视图分子表示。Equation (3) 中，M_t 为预先指定的主指标；当指标越大越优时 dir_t=1，当误差类指标越小越优时 dir_t=-1。该写法把 ROC-AUC、PR-AUC、RMSE 和 MAE 转换到统一的验证效用方向，同时保留原始指标的报告方式。",
        ),
        ("subhead", "验证集选择器、容差集合与最终保留。"),
        (
            "formula",
            "a_t^{best} = argmax_{a∈A_t} S_t^{val}(a)",
            4,
        ),
        (
            "formula",
            "A_t(ε_t) = {a∈A_t : S_t^{val}(a) ≥ max_{b∈A_t} S_t^{val}(b) - ε_t}",
            5,
        ),
        (
            "formula",
            "a_t^* = argmin_{a∈A_t(ε_t)} [η_1R_t^{val}(a) + η_2C(a) - η_3Stab_t(a)]",
            6,
        ),
        (
            "normal",
            "Equations (4)-(6) 给出由验证集治理的选择器。系统先识别验证效用最高的候选，再把落入 ε_t 容差范围的近似并列候选保留下来，最后按验证风险、模型复杂度和跨 seed 稳定性确定最终策略。若没有触发近似并列或稳定性规则，equation (6) 退化为 equation (4) 中的验证集最佳候选。",
        ),
        ("subhead", "验证偏差、排名审计与嵌套验证。"),
        (
            "formula",
            "a_t^{†} = argmax_{a∈A_t} S_t^{te}(a),  Regret_t = S_t^{te}(a_t^{†}) - S_t^{te}(a_t^*)",
            7,
        ),
        (
            "formula",
            "OptGap_t = S_t^{val}(a_t^*) - S_t^{te}(a_t^*)",
            8,
        ),
        (
            "formula",
            "ρ_{rank,t} = Spearman(rank_val(A_t), rank_test(A_t)),  Hit@K_t = I(a_t^{†}∈TopK_val(A_t))",
            9,
        ),
        (
            "formula",
            "Perf_{t,o}^{outer} = S_{t,o}^{outer-test}(argmax_{a∈A_t} K^{-1}∑_{k=1}^{K}S_{t,o,k}^{inner-val}(a))",
            10,
        ),
        (
            "normal",
            "Equations (7)-(10) 用于区分冻结选择器和事后测试集最佳候选。Regret_t 衡量冻结策略与测试集观测最佳之间的差距，OptGap_t 衡量验证集乐观偏差。排名审计比较验证集和测试集候选排序，equation (10) 则定义 nested validation，其中外层测试折不参与内层候选选择。",
        ),
        ("subhead", "多专家融合、适用域门控与样本级风险。"),
        (
            "formula",
            "ẑ_i^{mean} = K^{-1}∑_{a∈TopK_val(A_t)}ẑ_{i,a}",
            11,
        ),
        (
            "formula",
            "w_{i,a} = [exp(S_t^{val}(a)/τ)/(u_{i,a}+ε)] / ∑_{b∈B_t}[exp(S_t^{val}(b)/τ)/(u_{i,b}+ε)]",
            12,
        ),
        (
            "formula",
            "ẑ_i^{uw} = ∑_{a∈B_t}w_{i,a}ẑ_{i,a}",
            13,
        ),
        (
            "formula",
            "T(m_i,m_j)=|m_i∧m_j|/|m_i∨m_j|,  s_i^{NN}=max_{j∈D_t^{tr}∪D_t^{val}}T(m_i,m_j),  d_i^{AD}=1-s_i^{NN}",
            14,
        ),
        (
            "formula",
            "ẑ_i^{AD}=I(s_i^{NN}≥δ_AD)ẑ_i^{ens}+I(s_i^{NN}<δ_AD)ẑ_i^{safe}",
            15,
        ),
        (
            "formula",
            "r_i = λp_i^{err} + (1-λ)d_i^{AD}",
            16,
        ),
        (
            "normal",
            "Equations (11)-(16) 定义多专家融合和适用域输出。B_t 表示通过验证集保留的融合候选集合，u_{i,a} 表示候选专家 a 在样本 i 上的不确定性，τ 控制验证效用对权重的影响强度。最近邻 Tanimoto 分数用于估计化学空间支持度；落在适用域之外的样本被路由到更保守的保留专家，或被标记为高风险样本。",
        ),
        ("subhead", "风险覆盖、校准与固定精确率召回。"),
        (
            "formula",
            "Coverage(q)=n^{-1}∑_{i=1}^{n}I(r_i≤Q_q(r)),  Risk(q)=mean(e_i | r_i≤Q_q(r))",
            17,
        ),
        (
            "formula",
            "Brier=n^{-1}∑_{i=1}^{n}(p_i-y_i)^2,  ECE=∑_{b=1}^{B}(|B_b|/n)|acc(B_b)-conf(B_b)|",
            18,
        ),
        (
            "formula",
            "Recall@P≥π = max_{γ:Precision(γ)≥π} TP(γ)/(TP(γ)+FN(γ))",
            19,
        ),
        (
            "normal",
            "Equation (17) 是 Fig. 11 中 risk-coverage 曲线的计算基础。Equation (18) 用于报告分类任务的概率校准。Equation (19) 用于 ClinTox fixed-precision audit，因为在严重类别不平衡条件下，高 ROC-AUC 可能掩盖阳性类别召回不足的问题。",
        ),
        ("subhead", "保形预测、低相似度分层、活性悬崖与片段统计。"),
        (
            "formula",
            "q_α^{reg}=Quantile_cal(|y-ẑ|,1-α),  I_α(x)=[ẑ(x)-q_α^{reg}, ẑ(x)+q_α^{reg}]",
            20,
        ),
        (
            "formula",
            "q_α^{cls}=Quantile_cal(1-p_y,1-α),  C_α(x)={c : 1-p_c(x)≤q_α^{cls}}",
            21,
        ),
        (
            "formula",
            "bin_i∈{high,mid,low};  high:s_i^{NN}>0.7,  mid:0.5≤s_i^{NN}≤0.7,  low:s_i^{NN}<0.5",
            22,
        ),
        (
            "formula",
            "cliff(i,j)=I(T(m_i,m_j)≥θ_s and |y_i-y_j|≥θ_y),  ρ_gap=Spearman(|y_i-y_j|,|ẑ_i-ẑ_j|)",
            23,
        ),
        (
            "formula",
            "Δ_g=mean(y_i|g∈x_i)-mean(y_i|g∉x_i),  q_g=BH({p_g})",
            24,
        ),
        (
            "normal",
            "Equations (20) 和 (21) 分别定义基于校准集的 split-conformal 回归区间和分类集合。Equation (22) 固定低相似度分析中的三个互斥 Tanimoto 分层。Equation (23) 定义 MoleculeACE 活性悬崖分子对及预测差异与真实差异的相关性。Equation (24) 报告片段层面的效应量和 Benjamini-Hochberg FDR 校正证据，因此基序和片段结果被解释为关联性化学证据，而不是因果机制证明。",
        ),
        (
            "normal",
            "综上，equations (1)-(24) 共同定义了冻结数据边界、仅基于验证集的选择器、最终保留门控、排名审计、nested validation、融合、适用域距离、risk-coverage、校准、conformal prediction、低相似度分层、活性悬崖审计和片段富集统计。这些公式也明确区分了用于模型选择的量和策略冻结后才报告的审计量。",
        ),
    ]
    return block


def replace_formula_section(doc: Document):
    paragraphs = doc.paragraphs
    start = None
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip().startswith("3.7 "):
            start = idx
            break
    if start is None:
        raise RuntimeError("Section 3.7 was not found.")

    end = None
    for idx in range(start + 1, len(paragraphs)):
        text = paragraphs[idx].text.strip()
        style = paragraphs[idx].style.name if paragraphs[idx].style else ""
        if text == "结果" or (style.startswith("Heading 1") and idx > start):
            end = idx
            break
    if end is None:
        raise RuntimeError("The end of section 3.7 was not found.")

    old_elements = [p._element for p in paragraphs[start:end]]
    anchor = paragraphs[start]
    for kind, text, *maybe_num in build_formula_block():
        new_para = anchor.insert_paragraph_before("")
        if kind == "heading2":
            style_heading2(new_para, text)
        elif kind == "subhead":
            style_subhead(new_para, text)
        elif kind == "formula":
            style_formula(new_para, text, maybe_num[0])
        else:
            style_normal(new_para, text)

    for element in old_elements:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def audit(doc: Document):
    formula_numbers = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.endswith(")") and "\t" in paragraph.text:
            tail = text.rsplit("(", 1)[-1].rstrip(")")
            if tail.isdigit():
                formula_numbers.append(int(tail))
    wide_tables = [(i + 1, len(t.columns)) for i, t in enumerate(doc.tables) if len(t.columns) > 7]
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "formula_numbers": formula_numbers,
        "wide_tables": wide_tables,
    }


def main():
    source = find_source()
    output = source.parent / OUTPUT_NAME
    doc = Document(str(source))
    replace_formula_section(doc)
    doc.save(str(output))

    revised = Document(str(output))
    result = audit(revised)
    print(f"source={source}")
    print(f"output={output}")
    print(f"paragraphs={result['paragraphs']} tables={result['tables']} inline_shapes={result['inline_shapes']}")
    print(f"formula_numbers={result['formula_numbers']}")
    print(f"wide_tables={result['wide_tables']}")


if __name__ == "__main__":
    main()
