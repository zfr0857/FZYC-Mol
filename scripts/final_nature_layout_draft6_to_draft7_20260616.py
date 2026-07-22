# -*- coding: utf-8 -*-
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports" / "draft7_final_nature_layout_20260616"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def locate_source() -> Path:
    candidates = [p for p in (Path.home() / "Desktop").rglob("初稿-6.docx")]
    if not candidates:
        raise FileNotFoundError("Could not locate 初稿-6.docx.")
    return max(candidates, key=lambda p: p.stat().st_mtime)


SRC_DOCX = locate_source()
DEST_DOCX = SRC_DOCX.parent / "初稿-7.docx"
REPORT_DOCX = SRC_DOCX.parent / "初稿-7_Nature终排QA报告.docx"


def get_or_add_ppr(paragraph):
    return paragraph._p.get_or_add_pPr()


def remove_child(parent, tag: str) -> None:
    for child in list(parent):
        if child.tag == qn(tag):
            parent.remove(child)


def ensure_onoff(parent, tag: str, enabled: bool = True, single: bool = True) -> None:
    if single:
        remove_child(parent, tag)
    if enabled:
        parent.append(OxmlElement(tag))


def set_paragraph_font(paragraph, size_pt: float = 10.5, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold


def set_border(element, edge: str, val: str, size: str = "8", color: str = "000000") -> None:
    tag = "w:tblBorders" if element.tag.endswith("tblPr") else "w:tcBorders"
    borders = element.find(qn(tag))
    if borders is None:
        borders = OxmlElement(tag)
        element.append(borders)
    edge_el = borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        borders.append(edge_el)
    edge_el.set(qn("w:val"), val)
    edge_el.set(qn("w:sz"), size)
    edge_el.set(qn("w:space"), "0")
    edge_el.set(qn("w:color"), color)


def set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))


def set_cell_margins(table, top=45, start=65, bottom=45, end=65) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for edge, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = cell_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_width(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")


def apply_three_line_table(table) -> None:
    tbl_pr = table._tbl.tblPr
    for edge in ["top", "bottom"]:
        set_border(tbl_pr, edge, "single", "12")
    for edge in ["left", "right", "insideH", "insideV"]:
        set_border(tbl_pr, edge, "nil")

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for edge in ["left", "right", "insideH", "insideV"]:
                set_border(tc_pr, edge, "nil")
    if table.rows:
        for cell in table.rows[0].cells:
            set_border(cell._tc.get_or_add_tcPr(), "bottom", "single", "8")


def set_row_properties(row, header: bool) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    ensure_onoff(tr_pr, "w:cantSplit", True, single=True)
    ensure_onoff(tr_pr, "w:tblHeader", header, single=True)


def set_table_cell_text_style(cell, size_pt: float, header: bool = False, center: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if (header or center) else WD_ALIGN_PARAGRAPH.LEFT
        ppr = get_or_add_ppr(paragraph)
        ensure_onoff(ppr, "w:widowControl", True, single=True)
        ensure_onoff(ppr, "w:keepLines", True, single=True)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(size_pt)
            run.bold = header


def set_row_values(table, row_idx: int, values: list[str]) -> None:
    cells = table.rows[row_idx].cells
    for i, value in enumerate(values):
        cells[i].text = value


def remove_manual_break_and_empty_paragraphs(doc: Document) -> int:
    removed = 0
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        has_draw = "w:drawing" in paragraph._p.xml or "w:pict" in paragraph._p.xml
        has_page_break = 'w:type="page"' in paragraph._p.xml
        if (has_page_break or text == "") and not has_draw:
            paragraph._element.getparent().remove(paragraph._element)
            removed += 1
    return removed


def update_table_content(doc: Document) -> list[str]:
    changes: list[str] = []
    if len(doc.tables) >= 1:
        t = doc.tables[0]
        rows = [
            ["ESOL", "MoleculeNet；回归", "1117", "RMSE", "Scaffold/random/结构分离/低相似度；5 个种子"],
            ["FreeSolv", "MoleculeNet；回归", "642", "RMSE", "Scaffold/random/结构分离/低相似度；5 个种子"],
            ["Lipophilicity", "MoleculeNet；回归", "4200", "RMSE", "Scaffold/random/结构分离/低相似度；5 个种子"],
            ["BBBP", "MoleculeNet；分类", "1975；0.76", "ROC-AUC", "Scaffold/random/结构分离/低相似度；5 个种子"],
            ["BACE", "MoleculeNet；分类", "1513；0.4567", "ROC-AUC", "Scaffold/random/结构分离/低相似度；5 个种子"],
            ["ClinTox", "MoleculeNet；分类", "1461；0.0705", "ROC-AUC", "Scaffold/random/结构分离/低相似度；5 个种子"],
            ["TDC ADMET", "外部 ADMET", "578-13130", "RMSE / ROC-AUC", "PyTDC 官方划分；scaffold 审计"],
            ["MoleculeACE", "活性悬崖", "可用任务子集", "RMSE / R2 / 悬崖指标", "悬崖对与粗糙度诊断"],
            ["bRo5", "规则五以外压力测试", "CycPept-PAMPA；LinPept", "RMSE / ROC-AUC / PR-AUC", "公开数据外推与适用域审计；非盲测"],
        ]
        for idx, row in enumerate(rows, start=1):
            set_row_values(t, idx, row)
        changes.append("表1：统一中文分隔符，替换 AD 审计为适用域审计。")

    if len(doc.tables) >= 2:
        t = doc.tables[1]
        rows = [
            ["ESOL", "RMSE", "0.5829 ± 0.0352", "0.5829 ± 0.0352", "多指纹 0.6352；最优 0.5829", "低值优"],
            ["FreeSolv", "RMSE", "1.0678 ± 0.1883", "1.0286 ± 0.1761", "Chemprop 0.9518；最优 0.9518", "边界案例"],
            ["Lipophilicity", "RMSE", "0.7078 ± 0.0389", "0.6835 ± 0.0439", "D-MPNN 0.7456；最优 0.6835", "低值优"],
            ["BBBP", "ROC-AUC", "0.9165 ± 0.0290", "0.9243 ± 0.0247", "多指纹 0.9215；最优 0.9243", "高值优"],
            ["BACE", "ROC-AUC", "0.8753 ± 0.0230", "0.8753 ± 0.0230", "D-MPNN 0.8571；最优 0.8753", "高值优"],
            ["ClinTox", "ROC-AUC", "0.9489 ± 0.0302", "0.9496 ± 0.0262", "多指纹 0.9479；最优 0.9496", "高值优"],
        ]
        for idx, row in enumerate(rows, start=1):
            set_row_values(t, idx, row)
        changes.append("表2：去除未定义缩写 Multi-fp/best，改为中文可读表达。")

    if len(doc.tables) >= 3:
        t = doc.tables[2]
        set_row_values(t, 1, ["TDC 多方法融合", "Caco2、HIA、Pgp 出现选择性增益", "多视图和适用域门控对部分外部终点有效"])
        changes.append("表3：将 AD 门控改为适用域门控，减少歧义。")

    if len(doc.tables) >= 5:
        t = doc.tables[4]
        set_row_values(t, 2, ["bRo5 化学空间", "公共压力测试已纳入", "仅支持公开数据外推与适用域审计；非独立盲测"])
        changes.append("表5：统一 bRo5 行边界表述，避免 AD 缩写歧义。")

    return changes


def final_table_layout(doc: Document) -> list[str]:
    widths = {
        0: [0.72, 1.22, 1.38, 1.10, 1.95],
        1: [0.72, 0.72, 1.05, 1.05, 1.85, 0.70],
        2: [1.40, 2.10, 2.70],
        3: [1.85, 1.65, 2.70],
        4: [1.45, 2.00, 2.75],
        5: [1.30, 4.90],
    }
    notes: list[str] = []
    for ti, table in enumerate(doc.tables):
        set_table_fixed_width(table)
        set_cell_margins(table)
        apply_three_line_table(table)
        for ri, row in enumerate(table.rows):
            set_row_properties(row, header=(ri == 0))
            for ci, cell in enumerate(row.cells):
                if ti in widths and ci < len(widths[ti]):
                    set_cell_width(cell, widths[ti][ci])
                size = 7.4 if ti == 1 else 8.4
                center = ci in {0, 1, 2, 3} and ti in {0, 1, 3}
                set_table_cell_text_style(cell, size, header=(ri == 0), center=center)
        notes.append(f"表{ti + 1}：三线表、固定宽度、表头重复、行不跨页拆分、紧凑字号。")
    return notes


def final_paragraph_layout(doc: Document) -> list[str]:
    notes: list[str] = []
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        has_draw = "w:drawing" in paragraph._p.xml or "w:pict" in paragraph._p.xml
        ppr = get_or_add_ppr(paragraph)
        ensure_onoff(ppr, "w:widowControl", True, single=True)

        if has_draw:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            notes.append(f"图片段落 {i}：与图题同页。")
            continue

        if not text:
            continue

        if i == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            set_paragraph_font(paragraph, 15, bold=True)
            continue

        if text in {"Research Article", "作者：XXX，XXX，XXX", "单位：XXX大学 / XXX学院 / XXX实验室", "通讯作者：XXX，Email：XXX@XXX.com", "稿件类型：Research Article"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.keep_with_next = True
            set_paragraph_font(paragraph, 10.5, bold=False)
            continue

        if text.startswith("图"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.keep_together = True
            set_paragraph_font(paragraph, 9.0, bold=False)
            continue

        if text.startswith("表"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            set_paragraph_font(paragraph, 9.2, bold=False)
            continue

        if text.startswith("注："):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.keep_together = True
            set_paragraph_font(paragraph, 8.8, bold=False)
            continue

        is_heading = text in {"摘要", "1 引言", "2 材料与方法", "3 结果", "4 讨论", "5 结论", "List of abbreviations", "Declarations", "References"} or (
            len(text) > 2 and text[0].isdigit() and " " in text[:5]
        )
        if is_heading:
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            set_paragraph_font(paragraph, 12 if text[1:2] == " " or text in {"摘要", "Declarations", "References"} else 11, bold=True)
            continue

        if text.startswith("关键词："):
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.keep_together = True
            set_paragraph_font(paragraph, 10.0, bold=False)
            continue

        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(21)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.keep_together = False
        set_paragraph_font(paragraph, 10.5, bold=False)
    return notes


def resize_images_to_text_width(doc: Document) -> list[str]:
    notes: list[str] = []
    section = doc.sections[0]
    text_width = section.page_width - section.left_margin - section.right_margin
    max_width = min(text_width, Inches(7.05))
    for idx, shape in enumerate(doc.inline_shapes, 1):
        if shape.width > max_width:
            ratio = max_width / shape.width
            old_w, old_h = shape.width, shape.height
            shape.width = int(shape.width * ratio)
            shape.height = int(shape.height * ratio)
            notes.append(f"图{idx}：宽度 {old_w / 914400:.2f} in 调整为 {shape.width / 914400:.2f} in。")
        else:
            notes.append(f"图{idx}：宽度 {shape.width / 914400:.2f} in，未超出版心。")
    return notes


def set_page_setup(doc: Document) -> list[str]:
    notes: list[str] = []
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(0.60)
        section.right_margin = Inches(0.60)
        section.top_margin = Inches(0.70)
        section.bottom_margin = Inches(0.70)
    notes.append("页面：A4 纵向，左右 0.60 in、上下 0.70 in，兼顾宽表与正文可读性。")
    return notes


def collect_qa(doc: Document) -> list[str]:
    text = "\n".join(p.text for p in doc.paragraphs)
    def is_cjk(char: str) -> bool:
        return 0x4E00 <= ord(char) <= 0x9FFF

    qa: list[str] = []
    qa.append(f"段落数：{len(doc.paragraphs)}；表格数：{len(doc.tables)}；图片数：{len(doc.inline_shapes)}。")
    qa.append(f"最大表格列数：{max((len(t.columns) for t in doc.tables), default=0)}。")
    qa.append(f"长表格单元格（>90 字）：{sum(1 for t in doc.tables for r in t.rows for c in r.cells if len(c.text.strip()) > 90)}。")
    qa.append(f"禁止跨页拆分表格行：{sum(1 for t in doc.tables for r in t.rows if r._tr.get_or_add_trPr().find(qn('w:cantSplit')) is not None)}。")
    qa.append(f"重复表头行：{sum(1 for t in doc.tables for r in t.rows if r._tr.get_or_add_trPr().find(qn('w:tblHeader')) is not None)}。")
    qa.append(f"图片段落数量：{sum(1 for p in doc.paragraphs if 'w:drawing' in p._p.xml or 'w:pict' in p._p.xml)}；图题数量：{sum(1 for p in doc.paragraphs if p.text.strip().startswith('图'))}。")
    qa.append(f"表题数量：{sum(1 for p in doc.paragraphs if p.text.strip().startswith('表'))}；表注数量：{sum(1 for p in doc.paragraphs if p.text.strip().startswith('注：'))}。")
    text_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    qa.append(f"图片超出版心：{sum(1 for s in doc.inline_shapes if s.width > text_width)}。")
    qa.append(f"正文中文字符：{sum(1 for c in text.split('References')[0] if is_cjk(c))}。")
    for term in ["首创", "革命性", "证明了", "数据状态审计", "6 投稿前自我审查", "弱 基线模型", "backbone", "class weight", "validation metric", "AD 门控", "AD 审计"]:
        qa.append(f"风险词/旧表述检查 {term}：{text.count(term)} 处。")
    return qa


def xml_counts(path: Path) -> list[str]:
    from zipfile import ZipFile

    with ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    manual_page_breaks = xml.count('w:type="page"')
    return [
        f"XML 手动分页符：{manual_page_breaks}。",
        f"XML lastRenderedPageBreak：{xml.count('lastRenderedPageBreak')}。",
        f"XML keepNext：{xml.count('<w:keepNext')}。",
        f"XML keepLines：{xml.count('<w:keepLines')}。",
        f"XML cantSplit：{xml.count('<w:cantSplit')}。",
        f"XML tblHeader：{xml.count('<w:tblHeader')}。",
    ]


def build_report(doc: Document, content_changes: list[str], table_notes: list[str], figure_notes: list[str], para_notes: list[str], page_notes: list[str], qa: list[str], xml_qa: list[str]) -> None:
    report = Document()
    report.add_heading("初稿-7 Nature 终排 QA 报告", level=1)
    report.add_paragraph(f"源文档：{SRC_DOCX}")
    report.add_paragraph(f"输出文档：{DEST_DOCX}")
    report.add_heading("Nature 终排原则", level=2)
    report.add_paragraph(
        "本轮按 nature-polishing 的布局原则处理 Word 稿：避免强制分页造成松散页面，绑定图/表与题注，保持主文表短而清晰，禁止表格行跨页拆分，统一三线表、字号、行距、页边距和孤行控制。"
    )
    report.add_heading("表格内容修订", level=2)
    for item in content_changes:
        report.add_paragraph(item)
    report.add_heading("表格版式修订", level=2)
    for item in table_notes:
        report.add_paragraph(item)
    report.add_heading("页面与图文绑定", level=2)
    for item in page_notes + figure_notes[:5]:
        report.add_paragraph(item)
    report.add_paragraph(f"图片段落绑定记录共 {len(para_notes)} 项，已在 QA 中复核。")
    report.add_heading("QA 结果", level=2)
    for item in qa + xml_qa:
        report.add_paragraph(item)
    report.save(REPORT_DOCX)


def main() -> None:
    doc = Document(SRC_DOCX)
    removed = remove_manual_break_and_empty_paragraphs(doc)
    content_changes = update_table_content(doc)
    page_notes = set_page_setup(doc)
    figure_notes = resize_images_to_text_width(doc)
    table_notes = final_table_layout(doc)
    para_notes = final_paragraph_layout(doc)
    qa = collect_qa(doc)
    qa.insert(0, f"删除手动分页/空段落：{removed} 个。")

    doc.save(DEST_DOCX)
    shutil.copy2(DEST_DOCX, OUT_DIR / DEST_DOCX.name)
    xml_qa = xml_counts(DEST_DOCX)
    build_report(doc, content_changes, table_notes, figure_notes, para_notes, page_notes, qa, xml_qa)
    shutil.copy2(REPORT_DOCX, OUT_DIR / REPORT_DOCX.name)

    print(f"Wrote {DEST_DOCX}")
    print(f"Wrote {REPORT_DOCX}")
    for item in qa + xml_qa:
        print(item)


if __name__ == "__main__":
    main()
