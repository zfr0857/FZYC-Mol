from __future__ import annotations

import csv
import sys
from pathlib import Path

from docx import Document


def find_file(base: Path, patterns: list[str], required: list[str]) -> Path:
    matches: list[Path] = []
    for pattern in patterns:
        matches.extend(base.glob(pattern))
        matches.extend(base.rglob(pattern))
    filtered = []
    for path in matches:
        name = path.name
        if all(token in name for token in required) and path.suffix.lower() == ".docx":
            filtered.append(path)
    if not filtered:
        raise FileNotFoundError(f"No matching file under {base}: {required}")
    return max(filtered, key=lambda p: p.stat().st_mtime)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")

    draft = find_file(
        Path(r"C:\Users\Administrator\Desktop"),
        ["*.docx"],
        ["14"],
    )
    doc = Document(draft)

    out_dir = Path("work/draft15_audit")
    out_dir.mkdir(parents=True, exist_ok=True)

    paras = []
    for i, paragraph in enumerate(doc.paragraphs, 1):
        text = " ".join(paragraph.text.split())
        if text:
            paras.append((i, paragraph.style.name if paragraph.style else "", text))

    with (out_dir / "draft14_paragraph_index.csv").open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["index", "style", "text"])
        writer.writerows(paras)

    with (out_dir / "draft14_table_index.csv").open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["table", "rows", "cols", "first_row", "first_col"])
        for idx, table in enumerate(doc.tables, 1):
            rows = len(table.rows)
            cols = len(table.columns) if rows else 0
            first_row = " | ".join(cell.text.strip().replace("\n", " / ") for cell in table.rows[0].cells) if rows else ""
            first_col = " | ".join(row.cells[0].text.strip().replace("\n", " / ") for row in table.rows[:8]) if rows and cols else ""
            writer.writerow([idx, rows, cols, first_row, first_col])

    terms = [
        "bRo5",
        "MoleculeACE",
        "FreeSolv",
        "ClinTox",
        "Caco2",
        "Pgp",
        "nested",
        "ablation",
        "selector",
        "Tanimoto",
        "risk-coverage",
        "conformal",
        "external",
        "blind",
        "\u6d88\u878d",
        "\u4f4e\u76f8\u4f3c",
        "\u5d4c\u5957",
        "\u76f2\u6d4b",
        "\u5916\u90e8",
        "\u5c40\u9650",
        "\u5927\u6a21\u578b",
        "3D",
    ]
    print(f"Draft: {draft}")
    print(f"Paragraphs: {len(paras)}; tables: {len(doc.tables)}; images: {len(doc.inline_shapes)}")
    print("\nMatched paragraphs:")
    for i, style, text in paras:
        if any(term.lower() in text.lower() for term in terms):
            print(f"P{i:04d} [{style}] {text[:700]}")

    print("\nTables:")
    for idx, table in enumerate(doc.tables, 1):
        rows = len(table.rows)
        cols = len(table.columns) if rows else 0
        first = " | ".join(cell.text.strip().replace("\n", " / ") for cell in table.rows[0].cells) if rows else ""
        print(f"T{idx:02d}: {rows}x{cols} :: {first[:300]}")


if __name__ == "__main__":
    main()
