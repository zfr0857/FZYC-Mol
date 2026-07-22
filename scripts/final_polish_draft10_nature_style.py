# -*- coding: utf-8 -*-
"""Final Nature-leaning style polish for FZYC-Mol draft 10.

The script keeps the original document intact and writes a polished copy.
It only edits manuscript prose/table text that reads like internal process
notes, reviewer-response language, or informal wording.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


DESKTOP = Path(r"C:\Users\Administrator\Desktop")
SRC = max(DESKTOP.glob("FZYC-Mol_*-10.docx"), key=lambda p: p.stat().st_mtime)
OUT = DESKTOP / "FZYC-Mol_初稿-10_Nature终审润色版.docx"


PARAGRAPH_REPLACEMENTS_BY_PREFIX: dict[str, str] = {
    "为提高稿件的学术可读性，正文采用": (
        "结果呈现遵循紧凑、可追溯原则：主表只保留数据集、任务类型、评价指标、最终保留结果、"
        "最强对照和结论解释；宽表、候选矩阵、seed-level 差值、bootstrap 置信区间、Wilcoxon 检验、"
        "消融和完整失败单元均在补充表中给出。所有图表按首次出现顺序编号，并在图注或表注中说明"
        "指标方向、误差线含义、选择规则和负结果解释。"
    ),
    "近期分子性质预测与 ADMET benchmark 研究也提示": (
        "近期分子性质预测与 ADMET benchmark 研究通常将证据围绕数据划分、主指标、独立测试、"
        "不确定性、适用域和失败边界组织，而不是在正文堆叠过多候选细节。层级交互模型、"
        "ADMET 表征基准、分子性质预测综述和多模态毒性预测研究通常把紧凑主表、seed 级补充表、"
        "清晰图注和负结果说明结合起来。基于这一原则，本文主表仅保留直接支撑结论的摘要信息，"
        "候选级、seed 级、超参数级和完整指标级结果统一放入补充材料。"
    ),
    "为使补充实验与 Zhao et al. 2026 的真实挑战框架": (
        "为与 Zhao et al. 2026 的真实挑战框架 [3] 保持可比，本文将扩展实验组织为"
        "“挑战维度—候选池—验证门控—风险输出”的矩阵。该矩阵把强基线、OOD 划分、不平衡分类、"
        "bRo5 化学空间、活性悬崖、粗糙度、候选接受/拒绝审计和校准/保形预测纳入同一验证集治理流程，"
        "而不是将各模型作为相互孤立的追加结果。实验 A-C 和 E-G 构成支撑主线结论的核心证据，"
        "实验 D 与 H 用于加强真实药物发现场景和可靠性输出。所有模块均遵守测试集冻结原则："
        "候选是否进入最终保留结果，只由验证集或预先定义的审计规则决定。"
    ),
    "表 1A. 新增真实挑战实验矩阵": "表 1A. 真实挑战实验矩阵、输出指标与正文位置。",
    "实现层面，新增 experiment_update 运行层": (
        "为保证扩展实验可复现，本文将扩展实验与既有结果解耦，并要求所有运行过程输出稳定命名的 CSV、"
        "字段字典、缺失数据报告和缺失依赖报告。该设计使强基线、OOD 划分、不平衡策略、bRo5 压力测试、"
        "MoleculeACE 审计、粗糙度诊断和候选接受/拒绝记录能够从同一结果目录追溯到数据划分、随机种子、"
        "候选模型和最终表格。具体脚本、配置和字段说明见代码可用性部分及 Supplementary Methods。"
    ),
    "为提高可读性和可追溯性，主文图表只保留": (
        "图表呈现遵循可读性和可追溯性原则：正文只保留能直接回答核心问题的总结图和紧凑表格；"
        "完整随机种子级、候选级、指标级和终点级 CSV 进入补充材料包。每张正文图表均说明其回答的问题、"
        "指标方向、是否为最终保留结果，以及未接入候选作为负结果保留的依据。"
    ),
    "最新强基线同划分对照被作为": (
        "现代强基线同划分对照被作为 4.1 的独立证据块。该面板不只比较 CatBoost、XGBoost、"
        "LightGBM、ExtraTrees、RF 和 Chemprop，也把 TabPFNv2-RDKit、AutoGluon-RDKit、"
        "XGBoost-RDKit/Mordred/MorganCount 以及可获得的 KPGT representation 接入同一候选登记。"
        "每个 dataset-seed 同时输出 validation metric、test metric、rank、selected candidate 和 regret。"
        "其目标不是证明 FZYC-Mol 在每个终点击败所有强基线，而是检验当现代强基线进入同一验证集候选池后，"
        "最终保留策略是否仍能透明地接受、拒绝或保留旧基线。"
    ),
    "FreeSolv 的样本量较小": (
        "FreeSolv 的样本量较小，且溶剂化自由能对分子局部相互作用、构象与实验测定条件较敏感。"
        "表 2 中 Chemprop 是观测最优，而 FZYC-Mol 选择器原始版本略落后。定向重构接入后，"
        "Morgan 指纹与 RDKit 描述符的验证集堆叠将 RMSE 从 1.0678 ± 0.1883 降至 1.0286 ± 0.1761，"
        "表明低成本重构具有有限但可量化的增强作用，但仍未完全消除与最佳 Chemprop 候选的差距。"
        "同时，ESOL、BBBP、BACE、ClinTox 和 Lipophilicity 的重构候选均未通过最终保留门控，"
        "说明模型重构不是无条件提升，而是按终点接受的可验证增强。这一边界被保留在结果解释中："
        "FreeSolv 可能需要更强物理化学描述符、构象特征或更严格的嵌套验证，而不能把所有性能瓶颈都归因于模型规模不足。"
    ),
    "bRo5 外部压力测试用于补齐": (
        "bRo5 外部压力测试用于覆盖更复杂的化学空间。CycPept PAMPA、LinPept NonFouling 和 "
        "LinPept CellPen 分别代表环肽通透性、线性肽非特异性吸附和细胞穿透性；它们的分子量、"
        "构象柔性和低相似度外推难度均高于常规小分子面板。XGBoost-RDKit、XGBoost-MorganCount、"
        "TabPFNv2-RDKit、AutoGluon-RDKit、KPGT representation 和 FZYC selector 被纳入同划分候选池，"
        "并同步报告 RMSE/MAE、low-similarity bin、AD failure rate 和 risk enrichment。"
        "该实验不被解释为完整肽类 ADMET 建模，而是用于检验 FZYC-Mol 在更复杂化学空间中的候选治理和失败边界。"
    ),
    "全终点附录的作用不是把主文退化为大规模排行榜": (
        "全终点附录用于检验该框架能否从少数常见数据集泛化到更广泛的 ADMET 场景，而不是构建新的大规模排行榜。"
        "结果显示，强表格基线、Top-K/堆叠集成和目标变换在若干终点上提供可测量增益，但增益并不具有全局一致性。"
        "因此，本文的合理主张是“终点异质性下由验证集治理的选择性改进”，而非所有任务上的统一最优性声明。"
    ),
    "类别不平衡专项被放入可靠性小节": (
        "类别不平衡专项被放入可靠性小节，而不是作为单独的性能表。ClinTox、Tox21 NR ER、"
        "CYP2C9 Substrate、CYP2D6 Inhibition 和 hERG 同时报告 ROC-AUC、PR-AUC、Precision、"
        "Recall、F1、MCC、Balanced accuracy、Brier、ECE 和 Recall at fixed Precision，并比较 "
        "class weight、oversampling、downsampling、downsampling ensemble 与 threshold moving。"
        "该设置更接近筛选决策场景：模型还需量化在高精度阈值下可保留的阳性检出能力，"
        "并说明概率输出是否可用于风险分层。"
    ),
    "因此，TDC 与外部附录不应被写成": (
        "因此，TDC 与外部附录不作为时间顺序上的补充基准，而作为外部泛化证据：官方面板支持选择器在"
        "常见 ADMET 终点上的可用性，全终点附录显示最终保留门控能够吸收性能增强候选的选择性增益，"
        "同时在候选较弱时避免负迁移。"
    ),
    "Random/Scaffold/Perimeter 三类划分被用于形成明确的 OOD 难度梯度": (
        "Random/Scaffold/Perimeter 三类划分被用于形成明确的 OOD 难度梯度。Perimeter split 以 "
        "Morgan fingerprint 距离选择最远化学空间作为测试集，使 random-to-scaffold-to-perimeter "
        "的性能下降可以与 Zhao et al. [3] 的真实挑战设置直接比较。该分析同时报告 nearest-neighbor "
        "Tanimoto、低相似度分层、selector regret 和 rank stability。若某个终点在 Perimeter split 下出现"
        "非单调变化，正文优先解释样本组成、标签噪声和骨架分布，而不是简单归因为模型失败或模型胜出。"
    ),
    "探索性强基线实验未触发主结果更新": (
        "探索性强基线实验未触发主结果更新。MoleculeNet 强基线探索在 FreeSolv、Lipophilicity 和 ClinTox "
        "上分别选择堆叠或排序融合候选，但其性能未超过已有最终保留结果，因此仅作为附录证据报告。"
        "FreeSolv、Lipophilicity 和 ClinTox 的探索结果分别为 1.4417、0.7901 和 0.9186，"
        "对应最终保留/参考结果为 1.0286、0.6835 和 0.9489。该结果说明，候选扩展必须受到验证集门控约束，"
        "不能仅因模型类别较新或复杂度更高而进入最终保留结果。"
    ),
    "候选接受/拒绝审计被纳入候选登记": (
        "候选接受/拒绝审计被纳入候选登记。每个 dataset-seed 记录候选池数量、验证集最佳、最终保留、"
        "测试集观测最佳、是否接入、拒绝原因、训练时间、推理时间、硬件需求和失败率。该表与性能表分开呈现，"
        "用于评估 FZYC-Mol 是否主要依赖计算成本堆叠；若某个高性能候选的时间成本或失败率显著高于轻量候选，"
        "则其进入最终保留结果需要更强的验证集证据。"
    ),
    "注：该审计只用于评估验证集选择风险": (
        "注：该审计只用于评估验证集选择风险，不参与最终模型选择。总体中位 Spearman 为 0.661，"
        "说明验证集排名与测试集排名存在中等一致性；但验证集第一名等于测试集第一名的比例较低，"
        "提示验证集选择不应被解释为测试最优保证，而应解释为可冻结、可审计且仍有残余风险的模型治理流程。"
    ),
    "为提高结果呈现的连贯性": (
        "本文将结果证据组织为主性能、外部外推、性能增强、可靠性/校准、固定选择器审计、负结果诊断和"
        "可解释性案例七类。该组织方式首先评估 FZYC-Mol 在标准分子性质预测任务中的有效性，随后检验其在"
        "ADMET 外推、不平衡分类和低相似度样本中的稳健性，最后通过负结果与案例分析界定后续改进空间。"
    ),
    "从模型性能角度看，最有价值的改进不是继续无差别增加候选数量": (
        "从模型性能角度看，最有价值的改进不是继续无差别增加候选数量，而是把候选模型纳入可复现的选择器治理："
        "所有补救头、Top-K/堆叠集成、目标变换、树模型基线、3D-lite 描述符和粗糙度加权回归均需先通过"
        "验证集门控。未通过门控的模块仍保留为诊断证据；通过门控的模块才进入最终保留主结果。这样的设计"
        "直接暴露选择偏差、过拟合和负结果透明度，而不是将失败候选从证据链中删除。"
    ),
    "为形成完整的验证闭环": (
        "为形成完整的验证闭环，本文将关键扩展证据纳入结果逻辑：先审计验证集选择是否存在排名偏差，再用"
        "配对统计和系统消融界定各模块贡献，随后在 Random/Scaffold/Perimeter 梯度、低相似度、bRo5 肽类压力、"
        "MoleculeACE 30 任务活性悬崖、不平衡分类、ROGI/MODI/SARI 粗糙度诊断、候选接受/拒绝审计、"
        "校准/保形预测和失败案例中检查适用边界。这些分析不改变测试集一次性报告原则，而是将"
        "“为什么接受、保留或拒绝某个候选”写成可追溯的证据链。"
    ),
    "补充表 S1. 补充实验清单": "补充表 S1. 扩展实验证据在正文中的定位。",
    "系统消融与配对统计进一步限定了模块贡献": (
        "系统消融与配对统计进一步限定了模块贡献。完整选择器相对单一 Chemprop、多指纹、冻结预训练表征和"
        "核心家族候选均保持净正向，但 no_chemprop 和 no_pretrained 的负向或不稳定结果说明，候选家族数量"
        "增加本身并不保证性能提升；只有经过验证门控的候选才能进入最终保留策略。"
    ),
    "这种差异化定位决定了本文的补强方向": (
        "这种差异化定位决定了本文的补强方向。实验 A-H 将近期真实挑战研究强调的强基线、Perimeter split、"
        "类别不平衡、bRo5、活性悬崖、粗糙度、候选接受/拒绝审计和校准/保形预测转化为 FZYC-Mol 的候选治理压力测试。"
        "由此，复杂模块不被直接纳入主结论，而是在统一验证证据下被接受、保留或拒绝。这种证据组织与当前结果中"
        "“选择性增益多于普遍胜出”的事实一致。"
    ),
    "总体而言，FZYC-Mol 更适合作为一种由验证集治理的可靠性框架": (
        "总体而言，FZYC-Mol 更适合作为一种由验证集治理的可靠性框架，而不是单一追求最高分的排行榜方案。"
        "它在现有算力范围内提供了更完整的证据链，包括性能、方差、结构外推、适用域、高误差富集、校准、"
        "活性悬崖、外部 ADMET 附录和化学解释。这种证据链使结论更克制，也更贴近药物发现用户对模型可信度的实际需求。"
    ),
    "本研究仍存在若干局限": (
        "本研究仍存在若干局限。首先，验证-测试排名审计和 9 个代表性终点的 3×3 nested validation 表明，"
        "验证集治理可以减少测试集事后选择，但内外层验证尚未覆盖全部 MoleculeNet、TDC、bRo5 和 MoleculeACE 终点，"
        "不能保证所有任务均达到测试最优；因此，小幅增益应与 regret、optimism gap、Top3 hit 和负结果共同解释。"
        "其次，TabPFNv2、AutoGluon、KPGT representation、bRo5 和 MoleculeACE 30 任务仍需在完整同划分结果中继续扩展，"
        "尚未完成完整同划分验证的候选仅作为候选接口或后续验证方向，不作为已取得性能提升的证据。第三，"
        "bRo5、活性悬崖和低相似度样本仍是主要失败边界，FZYC-Mol 更适合识别和审计这些风险，而不是宣称已解决它们。"
        "第四，基序归因、片段富集和粗糙度相关性仍属于关联证据，不能替代因果机制或湿实验验证。"
    ),
    "因此，主文表格仅保留能够直接支撑结论的摘要证据": (
        "因此，主文表格仅保留能够直接支撑结论的摘要证据，完整候选级和 seed 级结果则置于补充材料。"
        "固定选择器的主文摘要表展示正向提升，同时正文保留 22 个负向终点-指标单元并指向完整审计表。"
        "这种呈现方式降低了选择偏差风险，也使负结果成为方法边界的一部分。"
    ),
    "利益冲突、基金和数据代码可用性声明": "利益冲突、基金、数据与代码可用性声明",
    "基金声明：本研究的具体资助信息需由作者": (
        "基金声明：本研究的资助信息将在投稿版本中由作者按期刊要求核定；若无专项资助，"
        "将声明“本研究未获得特定资助”。"
    ),
    "数据和代码可用性：本文使用的公开数据集可通过原始平台获得": (
        "数据和代码可用性：本文使用的公开数据集可通过 MoleculeNet、Therapeutics Data Commons、"
        "MoleculeACE 及相应原始平台获得。与本文结果对应的 split seeds、候选登记表、验证/测试预测、"
        "统计检验脚本、图表 source data、环境文件和表格生成脚本将与接收版本同步存档于 GitHub/Zenodo "
        "或期刊认可的数据仓库。仓库冻结后将在本文中补充永久链接和 accession number。"
    ),
}


INLINE_REPLACEMENTS: dict[str, str] = {
    "较好的识别能力": "较强的识别能力",
    "必须": "需",
    "未完成同划分正式运行的通道仅作为候选接口，不进入主文数值比较": "尚未完成同划分正式运行的通道仅作为候选接口，不进入正文数值比较",
    "无新增湿实验，也未声称完成官方盲测挑战": "未引入新的湿实验，也未声称完成官方盲测挑战",
    "新增 Top-3 命中率": "纳入 Top-3 命中率",
    "未完成同划分正式数值纳入": "尚未完成同划分正式数值纳入",
    "已补齐": "已纳入",
}


def set_paragraph_text(paragraph, text: str) -> None:
    # Preserve paragraph-level style while replacing run contents.
    paragraph.text = text


def polish_paragraphs(paragraphs) -> tuple[int, list[str]]:
    changed = 0
    missed: list[str] = []
    used = set()

    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        for prefix, replacement in PARAGRAPH_REPLACEMENTS_BY_PREFIX.items():
            if text.startswith(prefix):
                set_paragraph_text(paragraph, replacement)
                changed += 1
                used.add(prefix)
                break

    for paragraph in paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in INLINE_REPLACEMENTS.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            set_paragraph_text(paragraph, new_text)
            changed += 1

    for prefix in PARAGRAPH_REPLACEMENTS_BY_PREFIX:
        if prefix not in used:
            missed.append(prefix)
    return changed, missed


def polish_tables(doc: Document) -> int:
    changed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    new_text = text
                    for old, new in INLINE_REPLACEMENTS.items():
                        new_text = new_text.replace(old, new)
                    if new_text != text:
                        set_paragraph_text(paragraph, new_text)
                        changed += 1
    return changed


def polish_experiment_matrix(doc: Document) -> int:
    """Remove code-file names from the main challenge-experiment matrix."""
    if len(doc.tables) < 3:
        return 0
    table = doc.tables[2]
    rows = table.rows
    if len(rows) < 9 or len(rows[0].cells) < 6:
        return 0

    updates = {
        (1, 6): "正文位置与补充材料",
        (2, 6): "3.5、4.1、4.5；Supplementary Tables",
        (3, 6): "3.1、4.3；Supplementary Methods",
        (4, 6): "3.5、4.4；Supplementary Tables",
        (5, 6): "4.2、5.3；Supplementary Data",
        (6, 6): "4.6、4.7；Supplementary Tables",
        (7, 6): "3.6、4.4；Supplementary Methods",
        (8, 5): "候选池规模、接受/拒绝状态、拒绝原因、训练/推理成本、负结果摘要",
        (8, 6): "4.5、5.4；Supplementary Data",
        (9, 6): "4.4、4.6；Supplementary Methods",
    }
    changed = 0
    for (r, c), text in updates.items():
        paragraph = rows[r - 1].cells[c - 1].paragraphs[0]
        if paragraph.text != text:
            set_paragraph_text(paragraph, text)
            changed += 1
    return changed


def main() -> None:
    doc = Document(SRC)
    paragraph_changes, missed = polish_paragraphs(doc.paragraphs)
    table_changes = polish_tables(doc)
    matrix_changes = polish_experiment_matrix(doc)
    doc.save(OUT)
    print(f"source={SRC}")
    print(f"output={OUT}")
    print(f"paragraph_changes={paragraph_changes}")
    print(f"table_changes={table_changes}")
    print(f"matrix_changes={matrix_changes}")
    if missed:
        print("missed_prefixes=")
        for item in missed:
            print(f"- {item}")


if __name__ == "__main__":
    main()
