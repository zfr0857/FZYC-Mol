from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from restructure_manuscript_integrated_results_20260604 import (
    add_image,
    add_table,
    set_font,
    setup_doc,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SRC_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_academic_integrated_ordered_3line.md"
OUT_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260605_nine_chapter_structure.md"
OUT_DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260605_nine_chapter_structure.docx"
REPORT = DOCS / "manuscript_nine_chapter_restructure_report_20260605.md"


RECENT_WORK_PARAGRAPH = (
    "近几个月的相关工作为本文写法提供了清晰参照。ADMET reliability benchmark 将数据稀缺、OOD、类别不平衡、bRo5 和 activity cliffs "
    "作为核心挑战，并指出 TabPFNv2、undersampling ensemble 和 roughness index 在不同场景下各有价值。OpenADMET avoid-ome 研究强调，"
    "ADMET 失败来自一组共享的机制性反靶标和性质空间，需要开放数据、blind challenge 和适用域问题的系统研究。Tabular foundation model "
    "论文显示，RDKit/Mordred descriptors 与冻结 foundation embedding 结合 TabPFN 可以成为低成本强对照。HimNet、LGSM 和 CHAMP 等多尺度模型则强调 "
    "atom、motif、fingerprint、3D/geometric 和 global context 的层级融合与化学解释。"
)


def between(text: str, start: str, end: str | None = None) -> str:
    start_pos = text.index(start) + len(start)
    end_pos = text.index(end, start_pos) if end else len(text)
    return text[start_pos:end_pos].strip()


def renumber_subsections(body: str, old: int, new: int) -> str:
    return re.sub(rf"^##\s+{old}\.", f"## {new}.", body, flags=re.M)


def build_literature_review() -> str:
    return """# 第二章 Literature Review（文献回顾）

本章围绕分子性质预测、ADMET 可靠性、强基线模型和可解释性评价四条线索梳理相关研究。与只列举已有模型不同，本章重点说明现有文献如何界定问题难度、实验划分和模型可信度，从而为 FZYC-Mol 的方法设计提供依据。

## 2.1 分子性质预测 benchmark 与评价范式

MoleculeNet 和 Therapeutics Data Commons 为分子机器学习提供了标准化任务集合，使不同模型能够在统一数据集和指标下比较 [1,2]。这些 benchmark 的贡献不仅在于提供分数表，更在于推动研究者关注任务类型、数据规模、标签噪声和化学结构分布。早期工作常以随机划分下的平均 ROC-AUC 或 RMSE 作为主要依据，但药物发现应用更关心模型能否外推到新 scaffold、新系列或低相似度候选分子。Activity-cliff 研究进一步表明，即使分子结构高度相似，性质标签也可能出现突变，因此单纯报告平均性能容易掩盖局部高风险区域 [10]。

近年来，split realism 已成为分子性质预测论文中的关键审查点。Scaffold split、time split、structure-separated split 和 low-similarity hard subset 能够更接近真实 lead optimization 或 hit expansion 场景 [39]。这类研究提示，评价协议本身就是方法学的一部分；若模型只在简单随机划分上表现优越，却无法在结构外推场景中保持稳定，其药物发现价值应被谨慎解释。

## 2.2 ADMET 可靠性与外部验证研究

ADMET 任务具有 endpoint 异质性强、标签来源复杂、阳性比例不均衡和实验噪声高等特点。近期 ADMET reliability benchmark 将数据稀缺、OOD、类别不平衡、beyond-rule-of-5 化学空间和 activity cliffs 作为影响模型可信度的核心因素 [3]。OpenADMET avoid-ome 研究进一步把 ADMET 失败理解为跨靶点、跨性质空间的系统性风险，强调开放数据、blind challenge 和适用域报告对模型落地的重要性 [4]。ASAP-Polaris-OpenADMET blind challenge 则显示，利用预测数据进行多任务预训练可以改善部分 ADME endpoint，但其收益仍依赖任务、划分和外部验证设置 [5]。

这些工作共同说明，ADMET 论文不能只追求单一 leaderboard 提升，而应同时报告外部 benchmark、校准质量、不确定性、适用域和失败案例。ADMET-AI 与 ADMETlab 2.0 等平台型工作展示了大规模预测系统的实用价值 [8,9]，但对研究型论文而言，更重要的是说明模型何时可靠、何时应被拒用、何时需要额外实验验证。

## 2.3 分子表示、强基线与 foundation model

分子表示方法大体可分为传统指纹/描述符、图神经网络、消息传递模型和预训练分子语言模型。ECFP、MACCS、RDKit descriptors 与树模型长期以来仍是强基线，尤其适用于样本量有限、特征维度适中或 endpoint 噪声较高的任务 [19,23-27]。Chemprop D-MPNN 证明了有向消息传递在分子性质预测中的竞争力 [18]；ChemBERTa、MoLFormer 和图预训练模型则尝试从大规模未标注分子中学习可迁移表征 [28-31]。

近年的 TabPFN 和 tabular foundation model 研究显示，小样本表格预测不一定需要大规模分子专用 fine-tuning；在 RDKit/Mordred descriptors、fingerprint 或冻结 embedding 特征上，TabPFN 类模型可能形成低成本强对照 [11-13]。因此，本研究将 CatBoost、XGBoost、LightGBM、ExtraTrees、RF、Chemprop、冻结 ChemBERTa/MoLFormer、Top-K/stacking 和 TabPFN 代码通道纳入同一候选池，而不是把某一类模型预设为唯一主线。

## 2.4 不确定性、适用域与可解释性

不确定性量化、校准和 conformal prediction 为分子机器学习提供了性能之外的可信度证据 [33-36]。在药物发现中，模型不仅要给出数值预测，还要提示哪些分子处于低相似度区域、哪些样本存在高 ensemble disagreement、哪些 endpoint 可能受局部 target jump 或实验噪声影响。传统 QSAR 适用域研究已经指出，距离模型、相似度阈值和外推边界应与预测结果同时报告 [35,38]。

可解释性方面，motif attribution、fragment enrichment、scaffold case review 和 nearest-neighbor analysis 能够把模型行为连接到可读的化学片段。需要强调的是，这些分析通常提供关联性证据，而不是直接证明因果机制。对本文而言，可解释性模块的核心作用是帮助定位高误差区域、理解 selector 的选择边界，并为后续实验验证提供优先级。

## 2.5 本研究与已有工作的关系

综合上述文献，FZYC-Mol 的研究空白不在于提出一个更大的 backbone，而在于构建可审计的 validation-governed model selection framework。该框架把强 tabular baseline、图模型、冻结 embedding、target transform、undersampling ensemble、Top-K/stacking、AD gating、roughness 诊断和可解释性分析放入同一验证集治理流程。这样可以避免根据测试集事后挑选模型，也能把正结果、负结果和待复现实验放在统一证据链中解释。"""


def build_acknowledgement() -> str:
    return """# 第七章 Acknowledgement（致谢）

作者感谢 MoleculeNet、Therapeutics Data Commons、MoleculeACE、OpenADMET/Polaris 相关公开资源以及 RDKit、scikit-learn、XGBoost、LightGBM、CatBoost、Chemprop、DeepChem 等开源软件社区对本研究的支持。上述数据集和工具为分子性质预测、ADMET benchmark、适用域分析和可解释性评估提供了可复现基础。

作者同时感谢前期讨论中对模型性能、外部 benchmark、可靠性指标、图表规范和论文结构提出建议的同事与同行。正式投稿前，可根据实际情况补充具体人员、平台、计算资源或实验室支持信息。"""


def build_declarations() -> str:
    return """# 第八章 Declaration of Conflicting Interests / Funding（利益冲突声明/基金）

Declaration of Conflicting Interests：作者声明，目前未发现与本文研究内容直接相关的商业、财务或个人利益冲突。若后续存在专利申请、企业合作、咨询关系、软件授权或数据访问限制，应在正式投稿前据实补充。

Funding：当前稿件未提供具体基金项目信息。正式投稿前应补充基金名称、项目编号和资助机构；若本研究未获得特定外部资助，可在英文稿中写明：The authors received no specific funding for this work.

Data and Code Availability：本文使用的公开数据集、候选模型结果、图表和补充表格应随投稿材料提供可追溯路径。若目标期刊要求开放代码仓库，建议在正式投稿前整理训练脚本、随机种子、环境文件和结果表生成流程。"""


def fix_references(refs: str) -> str:
    refs = refs.strip()
    refs = re.sub(
        r"\[7\]\s+Gadaleta D, et al\..*?DOI: 10\.1186/s13321-024-00931-z\.",
        "[7] Gadaleta D, et al. Comprehensive benchmarking of computational tools for predicting toxicokinetic and physicochemical properties of chemicals. Journal of Cheminformatics, 2024, 16: 145. DOI: 10.1186/s13321-024-00931-z.",
        refs,
        flags=re.S,
    )
    return refs


def build_markdown() -> str:
    source = SRC_MD.read_text(encoding="utf-8")
    front = source[: source.index("# 1. 引言")].strip()
    front = front.replace(
        "完整中文初稿（实验融合整改版，2026-06-04）",
        "完整中文初稿（九章结构版，2026-06-05）",
    )

    intro = between(source, "# 1. 引言", "# 2. 方法")
    intro = intro.replace(RECENT_WORK_PARAGRAPH + "\n\n", "")
    intro = (
        "本章介绍分子性质预测在药物发现和 ADMET 评估中的研究背景，明确本文关注的结构外推、验证集治理和模型可靠性问题，并概述本文贡献。\n\n"
        + intro.strip()
    )

    methodology = renumber_subsections(between(source, "# 2. 方法", "# 3. 结果"), old=2, new=3)
    methodology = (
        "本章说明数据来源、数据划分、候选模型构建、验证集选择、可靠性评估和可解释性分析流程。为避免实验结果因事后挑选而产生偏差，所有候选策略均在 validation split 上完成选择，test split 仅用于冻结策略后的最终报告。\n\n"
        + methodology.strip()
    )

    results = renumber_subsections(between(source, "# 3. 结果", "# 4. 讨论"), old=3, new=4)
    results = (
        "本章按照科学问题而非实验加入时间组织结果：首先报告 MoleculeNet 主面板与 targeted rescue，其次呈现外部 ADMET 泛化，再分析结构分布偏移、可靠性、selector governance 和可解释性案例。\n\n"
        + results.strip()
    )

    discussion = renumber_subsections(between(source, "# 4. 讨论", "# 5. 结论"), old=4, new=5)
    conclusion = between(source, "# 5. 结论", "# 6. 参考文献")
    conclusion = (
        conclusion.strip()
        + "\n\n未来工作可沿三条路线推进。第一，继续扩展外部 ADMET appendix 和 blind-challenge 风格验证，以检验 selector 在更多真实 endpoint 上的稳定性。第二，在少量代表性低分任务上尝试受控的 adapter 或 lightweight fine-tuning，但仍需使用 nested validation 限制过拟合。第三，进一步加强 case-level interpretability，把 motif、nearest-neighbor、uncertainty 和实验标签噪声联系起来，为后续湿实验或专家审查提供更清晰的优先级。"
    )

    refs = fix_references(between(source, "# 6. 参考文献"))

    parts = [
        front,
        "# 第一章 Introduction（引言）",
        intro,
        build_literature_review(),
        "# 第三章 Methodology（研究方法/数据收集方法）",
        methodology,
        "# 第四章 Results（数据结果）",
        results,
        "# 第五章 Discussion（讨论）",
        discussion,
        "# 第六章 Conclusion（总结/展望）",
        conclusion,
        build_acknowledgement(),
        build_declarations(),
        "# 第九章 References（参考文献）",
        refs,
    ]
    return "\n\n".join(part.strip() for part in parts if part.strip()) + "\n"


def add_paragraph(doc: Document, text: str, line_index: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    size = 10.5
    bold = False
    color = "111827"
    if line_index == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 16
        bold = True
    elif line_index == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 10.5
        color = "475569"
    elif re.match(r"^图\s*\d+\.", text):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 9.0
        bold = True
        color = "334155"
    elif re.match(r"^表\s*\d+\.", text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 9.2
        bold = True
    elif text.startswith("注："):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 8.2
        color = "64748B"
    elif re.match(r"^\[\d+\]\s+", text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 9.0
    elif not text.startswith("|"):
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def markdown_to_docx(markdown: str) -> None:
    doc = Document()
    setup_doc(doc)
    lines = markdown.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line:
            idx += 1
            continue
        if line.startswith("!["):
            add_image(doc, line)
            idx += 1
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            p = doc.add_heading("", level=min(level, 3))
            run = p.add_run(text)
            set_font(run, size={1: 16, 2: 13.5, 3: 12}.get(level, 11), bold=True)
            idx += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            if len(table_lines) >= 2:
                add_table(doc, table_lines)
            continue
        add_paragraph(doc, line, idx)
        idx += 1
    doc.save(OUT_DOCX)


def write_report(markdown: str) -> None:
    figures = len(re.findall(r"^!\[", markdown, flags=re.M))
    tables = len(re.findall(r"^表\s*\d+\.", markdown, flags=re.M))
    REPORT.write_text(
        "\n".join(
            [
                "# Nine-Chapter Manuscript Restructure Report",
                "",
                "This revision reorganizes the Chinese manuscript according to the requested chapter sequence.",
                "",
                "Main changes:",
                "- Preserved the abstract before the numbered chapters.",
                "- Replaced the original six-section body with nine formal chapters.",
                "- Added a standalone Literature Review chapter with benchmark, ADMET reliability, strong-baseline, uncertainty, applicability-domain, and interpretability context.",
                "- Moved Methodology, Results, Discussion, and Conclusion to chapters 3-6 while preserving the original evidence flow.",
                "- Added Acknowledgement and Declaration of Conflicting Interests / Funding sections.",
                "- Corrected the title of reference [7] according to its DOI metadata.",
                f"- Figures retained: {figures}.",
                f"- Table captions retained: {tables}.",
                "",
                f"Markdown: `{OUT_MD.relative_to(ROOT)}`",
                f"DOCX: `{OUT_DOCX.relative_to(ROOT)}`",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def main() -> None:
    markdown = build_markdown()
    OUT_MD.write_text(markdown, encoding="utf-8")
    markdown_to_docx(markdown)
    write_report(markdown)
    print(f"wrote {OUT_MD.relative_to(ROOT)}")
    print(f"wrote {OUT_DOCX.relative_to(ROOT)}")
    print(f"wrote {REPORT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
