from __future__ import annotations

import json
import re
import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
SRC = OUT / "\u5c0f\u8bba\u6587-13.docx"
DST = OUT / "\u5c0f\u8bba\u6587-13_Nature\u683c\u5f0f\u7ec8\u5ba1.docx"
AUDIT = OUT / "paper13_nature_format_audit.json"
TEXT = OUT / "paper13_nature_format_text.txt"


ZH_FONT = "\u5b8b\u4f53"
EN_FONT = "Times New Roman"


def has_drawing_or_break(p) -> bool:
    xml = p._p.xml
    return "<w:drawing" in xml or "<w:pict" in xml or "<w:br" in xml


def remove_paragraph(p) -> None:
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = EN_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), ZH_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style_paragraph(p, *, role: str = "body") -> None:
    text = p.text.strip()
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)

    if role == "title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
        pf.space_after = Pt(12)
        for r in p.runs:
            set_run_font(r, 16, True)
        return

    if role == "heading1":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        for r in p.runs:
            set_run_font(r, 12, True)
        return

    if role == "heading2":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(8)
        pf.space_after = Pt(4)
        for r in p.runs:
            set_run_font(r, 10.5, True)
        return

    if role == "heading3":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(3)
        for r in p.runs:
            set_run_font(r, 10, True)
        return

    if role == "caption":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.line_spacing = 1.0
        pf.space_before = Pt(4)
        pf.space_after = Pt(3)
        for r in p.runs:
            set_run_font(r, 9, False)
        return

    if role == "keywords":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.space_after = Pt(8)
        for r in p.runs:
            set_run_font(r, 10.5, False)
        return

    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0.74) if text else Cm(0)
    for r in p.runs:
        set_run_font(r, 10.5, False)


def set_cell_border(cell, *, top=None, bottom=None, left="nil", right="nil") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    def set_edge(edge: str, spec) -> None:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        if spec is None:
            el.set(qn("w:val"), "nil")
            return
        if spec == "nil":
            el.set(qn("w:val"), "nil")
            return
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(spec.get("sz", 6)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")

    set_edge("top", top)
    set_edge("bottom", bottom)
    set_edge("left", left)
    set_edge("right", right)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    def edge(name: str, val: str, sz: int = 6) -> None:
        el = borders.find(qn(f"w:{name}"))
        if el is None:
            el = OxmlElement(f"w:{name}")
            borders.append(el)
        el.set(qn("w:val"), val)
        if val != "nil":
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")

    edge("top", "single", 8)
    edge("bottom", "single", 8)
    edge("left", "nil")
    edge("right", "nil")
    edge("insideH", "nil")
    edge("insideV", "nil")


def remove_cell_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in list(tc_pr.findall(qn("w:shd"))):
        tc_pr.remove(shd)


def set_cell_margin(cell, margin_twips: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "bottom", "left", "right"):
        el = tc_mar.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(margin_twips))
        el.set(qn("w:type"), "dxa")


def format_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for style_name in ("Normal Table", "Table Normal"):
        try:
            table.style = style_name
            break
        except Exception:
            continue
    set_table_borders(table)

    n_rows = len(table.rows)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            remove_cell_shading(cell)
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            top = {"sz": 8} if r_idx == 0 else None
            bottom = {"sz": 6} if r_idx == 0 else ({"sz": 8} if r_idx == n_rows - 1 else None)
            set_cell_border(cell, top=top, bottom=bottom)
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, 8.5, r_idx == 0)


def replace_text_preserve_single_run(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    style_paragraph(p, role="body")


def polish_repeated_methods(doc: Document) -> None:
    replacements = {
        "\u9664\u4e3b\u6548\u5e94\u5916\uff0c\u672c\u6587\u8fd8\u62a5\u544a\u771f\u5d4c\u5957\u9a8c\u8bc1\u3001\u79cd\u5b50\u654f\u611f\u6027\u3001\u7edf\u4e00\u6d88\u878d\u300180/90/95 \u4fdd\u5f62\u8986\u76d6\u7387\u3001\u7cbe\u786e Tanimoto \u5206\u7bb1\u3001MoleculeACE \u6d3b\u6027\u60ac\u5d16\u3001\u4f4e\u76f8\u4f3c\u5ea6\u5931\u8d25\u6837\u672c\u548c\u6269\u5c55\u5931\u8d25\u6848\u4f8b\u7b49 10 \u4e2a\u8865\u5145\u5b9e\u9a8c\u3002":
        "\u9664\u4e3b\u6548\u5e94\u5916\uff0c\u672c\u6587\u8fd8\u62a5\u544a\u771f\u5d4c\u5957\u9a8c\u8bc1\u3001\u79cd\u5b50\u654f\u611f\u6027\u3001\u7edf\u4e00\u6d88\u878d\u300180/90/95 \u4fdd\u5f62\u8986\u76d6\u7387\u3001\u7cbe\u786e Tanimoto \u5206\u7bb1\u3001MoleculeACE \u6d3b\u6027\u60ac\u5d16\u3001\u4f4e\u76f8\u4f3c\u5ea6\u5931\u8d25\u6837\u672c\u548c\u6269\u5c55\u5931\u8d25\u6848\u4f8b\u7b49\u5341\u7c7b\u8fb9\u754c\u5206\u6790\u3002",
        "\u5386\u53f2\u63a2\u7d22\u4e2d\u7684 GNN\u3001Chemprop\u3001ChemBERTa\u3001MoLFormer \u548c\u66f4\u590d\u6742\u878d\u5408\u4e3a\u5019\u9009\u6765\u6e90\u63d0\u4f9b\u80cc\u666f\uff0c\u4f46\u53ea\u6709\u5728\u7edf\u4e00\u5916\u5c42\u5212\u5206\u4e0a\u91cd\u8bad\u5e76\u5b8c\u6210\u767b\u8bb0\u540e\u624d\u53ef\u8fdb\u5165\u786e\u8ba4\u6027\u6548\u5e94\u4f30\u8ba1\u3002\u8fd9\u4e00\u7ea6\u675f\u8f83\u4e3a\u4fdd\u5b88\uff0c\u4f46\u53ef\u9632\u6b62\u5386\u53f2\u63a2\u7d22\u4e2d\u8868\u73b0\u6700\u597d\u7684\u7ed3\u679c\u88ab\u4e8b\u540e\u9009\u5165\u4e3b\u6587\u3002\u6362\u8a00\u4e4b\uff0cFZYC-Mol \u7684\u6cbb\u7406\u903b\u8f91\u5e76\u4e0d\u6392\u65a5\u590d\u6742\u6a21\u578b\uff0c\u800c\u662f\u8981\u6c42\u590d\u6742\u6a21\u578b\u627f\u62c5\u4e0e\u8f7b\u91cf\u5019\u9009\u76f8\u540c\u7684\u5ba1\u8ba1\u6210\u672c\u3002":
        "\u5386\u53f2\u5019\u9009\u6765\u6e90\u7528\u4e8e\u754c\u5b9a\u641c\u7d22\u7a7a\u95f4\uff0c\u800c\u4e0d\u76f4\u63a5\u53c2\u4e0e\u786e\u8ba4\u6027\u6548\u5e94\u4f30\u8ba1\u3002GNN\u3001Chemprop\u3001ChemBERTa\u3001MoLFormer\u3001\u6307\u7eb9\u3001\u4e8c\u7ef4\u63cf\u8ff0\u7b26\u3001\u7247\u6bb5/\u9aa8\u67b6\u548c\u9884\u6d4b\u878d\u5408\u53ea\u6709\u5728\u7edf\u4e00\u5916\u5c42\u5212\u5206\u4e0a\u91cd\u8bad\u3001\u767b\u8bb0\u5e76\u8f93\u51fa\u8d1f\u7ed3\u679c\u540e\uff0c\u624d\u53ef\u8fdb\u5165\u4e3b\u6548\u5e94\u6bd4\u8f83\u3002",
        "\u8fd9\u4f7f\u672c\u6587\u7684\u5b9a\u4f4d\u66f4\u63a5\u8fd1\u6a21\u578b\u8bc4\u4f30\u548c\u9a8c\u8bc1\u6cbb\u7406\uff0c\u800c\u4e0d\u662f\u53c8\u4e00\u4e2a\u65b0\u7684\u9884\u6d4b\u5668\u6392\u540d\u3002":
        "\u8fd9\u4f7f\u672c\u6587\u7684\u5b9a\u4f4d\u66f4\u63a5\u8fd1\u6a21\u578b\u8bc4\u4f30\u548c\u9a8c\u8bc1\u6cbb\u7406\uff0c\u800c\u4e0d\u662f\u65b0\u7684\u9884\u6d4b\u5668\u6392\u5e8f\u7814\u7a76\u3002",
    }
    remove_prefix = "\u5386\u53f2\u63a2\u7d22\u5305\u542b\u5206\u5b50\u56fe\u3001Chemprop\u3001ChemBERTa\u3001MoLFormer\u3001\u6307\u7eb9\u3001\u4e8c\u7ef4\u63cf\u8ff0\u7b26\u3001\u7247\u6bb5/\u9aa8\u67b6\u548c\u9884\u6d4b\u878d\u5408\u3002"

    for p in list(doc.paragraphs):
        t = p.text.strip()
        if t in replacements:
            replace_text_preserve_single_run(p, replacements[t])
        elif t.startswith(remove_prefix):
            remove_paragraph(p)

    substring_replacements = {
        "\u9a8c\u8bc1\u6700\u4f18\u5019\u9009": "\u9a8c\u8bc1\u6548\u7528\u6700\u9ad8\u7684\u5019\u9009",
        "\u63a5\u8fd1\u9a8c\u8bc1\u6700\u4f18\u5019\u9009": "\u63a5\u8fd1\u6700\u9ad8\u9a8c\u8bc1\u6548\u7528\u7684\u5019\u9009",
        "\u88ab\u8bc1\u660e\u65e0\u6548": "\u5df2\u88ab\u5224\u5b9a\u65e0\u6548",
        "\u6ca1\u6709\u8de8\u7ec8\u70b9\u666e\u9002\u6700\u4f18\u89e3": "\u6ca1\u6709\u5f62\u6210\u8de8\u7ec8\u70b9\u666e\u9002\u89c4\u5219",
        "\u552f\u4e00\u6700\u4f18\u89c4\u5219": "\u552f\u4e00\u6b63\u786e\u89c4\u5219",
        "\u5386\u53f2\u6700\u4f18\u5206\u6570": "\u5386\u53f2\u6700\u9ad8\u5206\u6570",
        "\u6700\u4f18\u89c4\u5219": "\u8868\u73b0\u6700\u7a33\u5b9a\u7684\u89c4\u5219",
        "\u672c\u6587\u7684\u4e0d\u8db3\u662f\u5e76\u672a\u63d0\u4f9b": "\u672c\u6587\u7684\u8303\u56f4\u5e76\u4e0d\u5305\u62ec\u63d0\u4f9b",
        "\u800c\u4e0d\u662f\u53c8\u4e00\u4e2a\u65b0\u7684\u9884\u6d4b\u5668\u6392\u540d": "\u800c\u4e0d\u662f\u65b0\u7684\u9884\u6d4b\u5668\u6392\u5e8f\u7814\u7a76",
    }
    for p in doc.paragraphs:
        t = p.text
        updated = t
        for old, new in substring_replacements.items():
            updated = updated.replace(old, new)
        if updated != t:
            replace_text_preserve_single_run(p, updated)


def format_document() -> None:
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    polish_repeated_methods(doc)

    for p in list(doc.paragraphs):
        if not p.text.strip() and not has_drawing_or_break(p):
            remove_paragraph(p)

    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if i == 0:
            try:
                p.style = "Title"
            except Exception:
                pass
            style_paragraph(p, role="title")
        elif t.startswith("\u8868 "):
            try:
                p.style = "TableCaption"
            except Exception:
                pass
            style_paragraph(p, role="caption")
        elif t.startswith("\u56fe "):
            try:
                p.style = "FigureCaption"
            except Exception:
                pass
            style_paragraph(p, role="caption")
        elif t.startswith("\u5173\u952e\u8bcd"):
            style_paragraph(p, role="keywords")
        elif (
            t in {"\u6458\u8981", "\u53c2\u8003\u6587\u732e"}
            or re.match(r"^[1-5]\s+(\u5f15\u8a00|\u6750\u6599\u4e0e\u65b9\u6cd5|\u7ed3\u679c|\u8ba8\u8bba|\u7ed3\u8bba)$", t)
        ):
            try:
                p.style = "Heading 1"
            except Exception:
                pass
            style_paragraph(p, role="heading1")
        elif re.match(r"^\d+\.\d+\.\d+", t):
            try:
                p.style = "Heading 3"
            except Exception:
                pass
            style_paragraph(p, role="heading3")
        elif p.style.name in {"Heading 2", "Heading 3"} or re.match(r"^\d+\.\d+\s", t):
            try:
                p.style = "Heading 2"
            except Exception:
                pass
            style_paragraph(p, role="heading2")
        elif has_drawing_or_break(p):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(4)
        else:
            style_paragraph(p, role="body")

    for table in doc.tables:
        format_table(table)

    doc.core_properties.title = "\u5019\u9009\u6c60\u6269\u5f20\u589e\u52a0\u5206\u5b50\u6027\u8d28\u9884\u6d4b\u4e2d\u7684\u9009\u62e9\u635f\u5931"
    doc.core_properties.subject = "\u5c0f\u8bba\u6587-13 Nature \u683c\u5f0f\u4e09\u7ebf\u8868\u7ec8\u5ba1\u7248"
    doc.save(DST)


def xml_errors(path: Path) -> list[str]:
    errors: list[str] = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if name.endswith(".xml"):
                try:
                    ElementTree.fromstring(zf.read(name))
                except Exception as exc:
                    errors.append(f"{name}: {exc}")
    return errors


def extract_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def table_border_audit(table) -> dict[str, object]:
    findings: list[str] = []
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            vals = {}
            if borders is not None:
                for edge in ("top", "bottom", "left", "right"):
                    el = borders.find(qn(f"w:{edge}"))
                    vals[edge] = el.get(qn("w:val")) if el is not None else None
            if vals.get("left") not in {"nil", None} or vals.get("right") not in {"nil", None}:
                findings.append(f"cell {r_idx},{c_idx} has vertical border")
            if r_idx == 0 and vals.get("top") != "single":
                findings.append(f"cell {r_idx},{c_idx} missing top rule")
            if r_idx == 0 and vals.get("bottom") != "single":
                findings.append(f"cell {r_idx},{c_idx} missing header rule")
            if r_idx == len(table.rows) - 1 and vals.get("bottom") != "single":
                findings.append(f"cell {r_idx},{c_idx} missing bottom rule")
    return {"passed": not findings, "findings": findings[:20]}


def audit() -> dict[str, object]:
    doc = Document(DST)
    text = extract_text(doc)
    TEXT.write_text(text, encoding="utf-8")

    abstract = []
    capture = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "\u6458\u8981":
            capture = True
            continue
        if capture and t.startswith("\u5173\u952e\u8bcd"):
            break
        if capture and t:
            abstract.append(t)

    colloquial_terms = [
        "\u505a\u539a",
        "\u8865\u539a",
        "\u8865\u5f3a",
        "\u6295\u7a3f\u524d",
        "\u4e0d\u8db3\u70b9",
        "\u4e0d\u8db3\u662f",
        "\u770b\u8d77\u6765",
        "\u8001\u5b9e\u8bf4",
        "\u5f88\u725b",
        "\u53c8\u4e00\u4e2a",
        "SCI 1",
    ]
    overclaim_terms = ["\u8bc1\u660e", "\u6bcb\u5eb8\u7f6e\u7591", "\u6700\u4f18", "\u9996\u6b21", "\u7a81\u7834\u6027"]
    table_audits = [table_border_audit(t) for t in doc.tables]
    figure_nums = []
    table_nums = []
    for p in doc.paragraphs:
        t = p.text.strip()
        m = re.match(r"^\u56fe\s*(\d+)", t)
        if m:
            figure_nums.append(int(m.group(1)))
        m = re.match(r"^\u8868\s*(\d+)", t)
        if m:
            table_nums.append(int(m.group(1)))

    styles = {}
    heading1_texts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        styles.setdefault(p.style.name, 0)
        styles[p.style.name] += 1
        if p.style.name == "Heading 1":
            heading1_texts.append(t)

    allowed_h1 = {
        "\u6458\u8981",
        "1 \u5f15\u8a00",
        "2 \u6750\u6599\u4e0e\u65b9\u6cd5",
        "3 \u7ed3\u679c",
        "4 \u8ba8\u8bba",
        "5 \u7ed3\u8bba",
        "\u53c2\u8003\u6587\u732e",
    }

    audit_data = {
        "docx": str(DST),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "styles_used": styles,
        "heading1_texts": heading1_texts,
        "heading1_only_section_titles": set(heading1_texts).issubset(allowed_h1),
        "abstract_paragraphs": len(abstract),
        "abstract_secondary_names_absent": not any(
            any(x in p for x in ["MoleculeNet", "TDC", "MoleculeACE", "bRo5", "\u4fdd\u5f62"])
            for p in abstract
        ),
        "table_caption_numbers": table_nums,
        "figure_caption_numbers": figure_nums,
        "table_styles": [t.style.name if t.style else None for t in doc.tables],
        "all_tables_three_line": all(item["passed"] for item in table_audits),
        "table_border_audits": table_audits,
        "colloquial_hits": {term: text.count(term) for term in colloquial_terms if term in text},
        "overclaim_hits": {term: text.count(term) for term in overclaim_terms if term in text},
        "has_tabpfn_boundary": "TabPFN" in text and "\u6388\u6743\u6216\u8fd0\u884c\u73af\u5883\u9650\u5236" in text,
        "has_error_overlap": "error-overlap" in text and "0.189" in text,
        "has_duplicate_sensitivity": "\u53bb\u91cd\u654f\u611f\u6027" in text and "0.022" in text,
        "xml_errors": xml_errors(DST),
    }
    audit_data["passed"] = (
        audit_data["abstract_paragraphs"] == 4
        and audit_data["abstract_secondary_names_absent"]
        and audit_data["table_caption_numbers"] == list(range(1, 10))
        and audit_data["figure_caption_numbers"] == list(range(1, 10))
        and audit_data["heading1_only_section_titles"]
        and all(style == "Normal Table" for style in audit_data["table_styles"])
        and audit_data["all_tables_three_line"]
        and not audit_data["colloquial_hits"]
        and not audit_data["overclaim_hits"]
        and audit_data["has_tabpfn_boundary"]
        and audit_data["has_error_overlap"]
        and audit_data["has_duplicate_sensitivity"]
        and not audit_data["xml_errors"]
    )
    AUDIT.write_text(json.dumps(audit_data, ensure_ascii=False, indent=2), encoding="utf-8")
    return audit_data


def main() -> None:
    format_document()
    print(json.dumps(audit(), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
