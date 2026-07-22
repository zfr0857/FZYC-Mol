# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / "reports" / "draft8_build"
ASSET_DIR = WORK / "assets"
OUTPUT_DIR = ROOT / "output"
WORK.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

SOURCE = Path(r"C:\Users\Administrator\Desktop\修改\初稿-7.docx")
CHECKLIST = Path(r"C:\Users\Administrator\Downloads\FZYC-Mol_Journal_of_Cheminformatics_修改清单.docx")
OUT = OUTPUT_DIR / "初稿-8.docx"
AUDIT = OUTPUT_DIR / "初稿-8_清单落实与证据审计.md"


def set_run_font(run, size=10.5, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title.font.size = Pt(16)
    title.font.bold = True
    title.paragraph_format.space_after = Pt(8)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(5)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.space_before = Pt(7)
    h2.paragraph_format.space_after = Pt(4)

    for name, size, italic, align in [
        ("FigureCaption", 9, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("TableCaption", 9, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("TableNote", 8.5, False, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        if name not in doc.styles:
            style = doc.styles.add_style(name, 1)
        else:
            style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.italic = italic
        style.paragraph_format.alignment = align
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    if "Equation" not in doc.styles:
        equation = doc.styles.add_style("Equation", 1)
    else:
        equation = doc.styles["Equation"]
    equation.font.name = "Cambria Math"
    equation._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    equation.font.size = Pt(10.5)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(3)
    equation.paragraph_format.space_after = Pt(3)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_p(doc: Document, text: str, style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.widow_control = True
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    return p


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:" + edge
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, value in edge_data.items():
                element.set(qn("w:" + key), str(value))


def add_three_line_table(doc: Document, rows: list[list[str]], widths: list[float] | None = None, font_size=8.5):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        total_weight = sum(widths)
        column_widths = [Cm(16.0 * value / total_weight) for value in widths]
    else:
        column_widths = [Cm(16.0 / len(rows[0])) for _ in rows[0]]
    for ri, row_values in enumerate(rows):
        table_row = table.rows[ri]
        trPr = table_row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)
        if ri == 0:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            trPr.append(header)
        for ci, value in enumerate(row_values):
            cell = table.cell(ri, ci)
            cell.width = column_widths[ci]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, bold=(ri == 0))
            set_cell_border(
                cell,
                top={"val": "single", "sz": "10", "color": "000000"} if ri == 0 else {"val": "nil"},
                bottom={"val": "single", "sz": "8", "color": "000000"} if ri == 0 else ({"val": "single", "sz": "10", "color": "000000"} if ri == len(rows) - 1 else {"val": "nil"}),
                left={"val": "nil"},
                right={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )
    return table


def add_table_caption(doc: Document, text: str):
    add_p(doc, text, style="TableCaption")


def add_table_note(doc: Document, text: str):
    add_p(doc, text, style="TableNote")


def add_equation(doc: Document, text: str):
    add_p(doc, text, style="Equation")


def extract_media() -> list[Path]:
    media: list[Path] = []
    with ZipFile(SOURCE) as z:
        names = [n for n in z.namelist() if n.startswith("word/media/")]
        names.sort(key=lambda s: int("".join(ch for ch in Path(s).stem if ch.isdigit()) or 0))
        for i, name in enumerate(names, 1):
            ext = Path(name).suffix or ".png"
            path = ASSET_DIR / f"figure_{i:02d}{ext}"
            path.write_bytes(z.read(name))
            media.append(path)
    return media


def add_picture(doc: Document, path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Cm(16.0))
    add_p(doc, caption, style="FigureCaption")


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r)


def build() -> None:
    media = extract_media()
    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FZYC-Mol：面向分子性质预测的验证治理模型选择与可靠性审计")
    set_run_font(r, size=16, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Research Article")
    set_run_font(r, size=10, italic=True)

    add_heading(doc, "摘要", 1)
    add_p(doc, "随着候选模型、分子表示、目标变换和集成策略不断增加，验证集可能逐渐成为隐性开发集，从而产生模型选择偏差和乐观估计。现有分子性质预测研究通常关注预测器之间的性能差异，但较少把候选池扩张、验证集过拟合、选择稳定性和拒绝原因作为主要研究对象。")
    add_p(doc, "本研究将 FZYC-Mol 重新定义为验证治理模型选择框架，而不是新的预测主干模型。该框架以冻结候选登记为输入，在外层测试集之外执行内层选择，并采用预先注册的 one-standard-error 候选集、跨内层重采样稳定性、校准或保形效率以及计算复杂度构成词典序晋级规则。测试集上的最佳候选仅作为 retrospective oracle upper bound，用于计算 test regret，不参与任何候选生成、阈值拟合或最终选择。")
    add_p(doc, "现有证据覆盖 MoleculeNet 六项任务、22 个 TDC ADMET 终点、MoleculeACE 活性悬崖和 bRo5 公共压力测试。回顾性验证-测试排序审计显示，中位 Spearman 相关为 0.667，测试 oracle 候选落入验证 Top-3 的比例为 0.295，Top-1 一致率为 0.135；分类和回归保形预测在 80%、90% 和 95% 目标覆盖率下分别获得 0.814/0.918/0.956 和 0.823/0.925/0.962 的平均经验覆盖率。由于当前证据包尚未提供逐任务 outer-test regret、候选池规模压力和完整独立外部验证结果，修订稿不声称 FZYC-Mol 已降低外层 regret，而将这些实验列为进入确认性结论前必须完成的 P0 证据。")

    add_heading(doc, "Scientific Contribution", 1)
    add_p(doc, "FZYC-Mol 将分子性质预测中的候选登记、内层选择、晋级、平局处理、冻结和外层审计形式化为可版本化流程。与主要比较模型家族在数据稀缺、类别不平衡、bRo5 化学空间和活性悬崖中表现的近期 ADMET 基准不同，本研究将模型选择过程本身定义为评价对象，并以 test regret、optimism gap、Top-k hit、selection stability 和 candidate-pool sensitivity 作为主要指标。当前修订稿仅对已有回顾性排序与可靠性证据作有界解释，并将完整 outer nested benchmark、逐样本输出和公开归档作为投稿前置条件。")
    add_p(doc, "关键词：分子性质预测；ADMET；验证集过拟合；模型选择偏差；冻结候选登记；嵌套验证；适用域；保形预测；可靠性审计")

    add_heading(doc, "1 引言", 1)
    add_p(doc, "分子性质预测已广泛用于 ADMET 评估、虚拟筛选和候选化合物优先级排序。MoleculeNet 和 Therapeutics Data Commons 等公开资源推动了指纹模型、图神经网络、消息传递模型和分子预训练表示的系统比较[1,2]。然而，预测器数量的增加并不自动提高结论可信度，因为性能估计还受到数据划分、候选池规模、模型选择规则和测试集使用方式的共同影响。")
    add_p(doc, "模型选择偏差是这一问题的核心。交叉验证或固定验证集同时承担超参数搜索和模型选择时，最优验证结果可能包含偶然噪声；候选越多，winner's curse 和 multiple comparisons 导致的乐观偏差越可能累积[3,4]。Cawley 和 Talbot 指出，模型选择过程本身的方差可与学习算法差异同等重要[3]；Varma 和 Simon 进一步表明，若特征或模型选择不被嵌套在外层评估之内，误差估计会出现系统性偏差[4]。因此，公平比较不能只固定数据和指标，还必须固定候选登记、选择规则、平局处理和外层测试协议。")
    add_p(doc, "分布偏移和样本级风险进一步放大了这一问题。随机划分可能保留训练集与测试集之间的近邻关系，不平衡毒性任务中的 ROC-AUC 也可能掩盖阳性召回和校准不足。近期 Zhao 等在 Journal of Cheminformatics 系统比较了数据稀缺、类别不平衡、bRo5 化学空间和活性悬崖场景下的基础模型、传统模型、AutoML、集成与粗糙度指标[5]。该工作已经覆盖了本稿原先的大部分外层评价场景，因此多视图、强基线、MoleculeACE 和 bRo5 不能继续作为本研究的主要新颖性来源。")
    add_p(doc, "FZYC-Mol 的独立科学问题据此被收缩为：候选池扩张是否使验证集成为隐性开发集，以及预先注册的验证治理能否减少模型选择损失并提高选择稳定性。框架将冻结候选登记、内层评价、one-standard-error 晋级、稳定性与复杂度平局规则、外层测试和负结果日志组合为一个可复核流程。多种分子表示、TDC、MoleculeACE 和 bRo5 仅作为压力测试环境，用于检验选择规则在不同数据规模、任务类型和化学空间中的边界。")
    add_p(doc, "本研究提出四个可检验假设：第一，候选池增大将增加朴素 validation-best 的 optimism gap 和 outer-test regret；第二，冻结登记与词典序门控可在不显著牺牲预测性能的条件下降低选择不稳定性；第三，将适用域、不确定性、校准和保形效率纳入审计可改善选择性预测的风险排序；第四，公开拒绝原因和负结果可减少只报告成功候选造成的结论不稳定。当前修订稿完整定义这些假设和协议，但只对已有证据作回顾性报告，尚未把缺失的 outer nested 和候选池压力结果写成完成性结论。")

    add_heading(doc, "2 材料与方法", 1)
    add_heading(doc, "2.1 研究范围与数据集", 2)
    add_p(doc, "研究包含 MoleculeNet、TDC ADMET、MoleculeACE 和 bRo5 相关公开数据。MoleculeNet 主面板包括 ESOL、FreeSolv、Lipophilicity、BBBP、BACE 和 ClinTox；TDC ADMET 用于公开外部 benchmark 与 scaffold 审计；MoleculeACE 用于活性悬崖压力测试[9]；CycPept-PAMPA 与 LinPept 仅用于 bRo5 公共压力测试。若 bRo5 数据与 Zhao 等[5]同源，其作用限定为可比性和选择稳定性审计，不作为新数据贡献。")
    add_table_caption(doc, "表1. 数据集、任务与数据治理状态。")
    add_three_line_table(doc, [
        ["数据集", "任务/规模", "主指标", "当前数据状态", "投稿前必填字段"],
        ["ESOL", "回归；n=1117", "RMSE", "当前稿件样本数", "版本、下载日期、原始/清洗后n、删除原因"],
        ["FreeSolv", "回归；n=642", "RMSE", "当前稿件样本数", "单位、重复测量处理、原始/清洗后n"],
        ["Lipophilicity", "回归；n=4200", "RMSE", "当前稿件样本数", "目标变换、单位、截尾规则"],
        ["BBBP", "分类；n=1975", "ROC-AUC", "阳性率0.76", "标签冲突与近重复泄漏"],
        ["BACE", "分类；n=1513", "ROC-AUC", "阳性率0.4567", "标签定义与版本"],
        ["ClinTox", "分类；n=1461", "ROC-AUC", "阳性率0.0705", "类别不平衡与固定精度召回"],
        ["TDC ADMET", "外部；n=578-13130", "RMSE/ROC-AUC", "22终点汇总", "逐终点版本、划分、预测与CI"],
        ["MoleculeACE", "活性悬崖", "RMSE/R2/cliff", "可用任务子集", "任务清单、阈值、分子对与排除标准"],
        ["bRo5", "CycPept-PAMPA/LinPept", "RMSE/ROC-AUC/PR-AUC", "公共压力测试", "来源、重叠审计、非盲测声明"],
    ], font_size=7.2)
    add_table_note(doc, "注：原始样本数、清洗后样本数、下载日期和删除原因未包含在当前证据包中，已明确保留为投稿前数据登记必填项；不得自行推断或补写。")

    add_heading(doc, "2.2 数据标准化、重复处理与泄漏审计", 2)
    add_p(doc, "数据处理应由版本化脚本完成，并在候选训练前冻结。标准化流程依次包括结构解析、盐拆分、最大有机片段选择、金属或无机物处理、芳香化与价态检查、规范化、互变异构体规则、立体化学、同位素和形式电荷处理；化学结构操作使用版本锁定的 RDKit 实现[10]。每一步均应记录软件版本、规则参数、失败分子和删除原因；当前稿件未提供这些日志，因此本节定义的是必须实现并归档的协议，而不是对未提供处理记录的追认。")
    add_p(doc, "标准化 SMILES 用于重复分子检测。回归任务应预先指定重复测量采用均值、中位数或全部保留，并报告选择理由；分类任务应预先指定冲突标签的剔除、共识或不确定标签策略。单位换算、对数变换、检测下限或上限和截尾值必须在 data dictionary 中定义，且只能在训练数据上拟合需要数据驱动参数的变换。")
    add_p(doc, "泄漏审计在 train、validation、calibration 和 outer test 之间分别检查标准化 SMILES 重复、Murcko scaffold 重复和近重复分子。近重复检查使用固定 Morgan 指纹与 Tanimoto 阈值，并输出跨分区最近邻分布。相同或高度近似分子若跨分区出现，应按预注册规则移除或分组划分，所有调整均写入 leakage_audit.csv。")

    add_heading(doc, "2.3 数据划分与 outer-inner nested protocol", 2)
    add_p(doc, "每项确认性分析使用完全隔离的 outer test。每个 outer split 的训练部分进一步划分为 inner training、inner validation 和独立 calibration 子集：inner training 用于拟合候选；inner validation 用于候选晋级与平局处理；calibration 子集用于概率校准和 conformal 非一致性分数；outer test 仅在候选、阈值、权重和代码版本冻结后评估。若历史固定测试集已被反复查看，其结果仅作为 retrospective audit，不再被称为一次性测试。")
    add_p(doc, "MoleculeNet 以 scaffold split 为主要协议，并报告 random、structure-separated 和 low-similarity 压力测试。TDC 使用官方划分，同时给出 scaffold 审计。低相似度样本按与 outer-training 最近邻的 Tanimoto 相似度划分为 >0.7、0.5-0.7 和 <0.5 三个互斥区间。所有 split index、随机种子、候选 registry 快照、配置哈希和时间戳应在 outer test 解封前归档。")
    if len(media) >= 2:
        add_picture(doc, media[0], "图1. FZYC-Mol 候选生成、验证选择与证据输出的概念结构。该图定义框架对象，不构成预测性能证据。")
        add_picture(doc, media[1], "图2. FZYC-Mol 整体实验工作流。outer-inner 选择、校准集分离和冻结条件以 Algorithm 1 的文字协议为准。")

    add_heading(doc, "2.4 分子表示、候选模型与冻结登记", 2)
    add_p(doc, "候选表示包括 Morgan、MACCS、atom-pair 和 torsion 指纹，RDKit 二维描述符，分子图输入，BRICS 与 Murcko scaffold 信息，以及冻结的 ChemBERTa 或 MoLFormer 表征接口。候选学习器包括 RF、ExtraTrees、XGBoost、LightGBM、CatBoost、Chemprop D-MPNN[8]、具体图模型、TabPFN 接口[11]、AutoGluon 接口[12]、描述符 MLP、Top-K 均值、正则化堆叠和预注册补救模块。候选名称不等同于完成运行；只有出现在冻结 registry、具有可复现配置与预测文件的候选才可进入完成性结果。")
    add_p(doc, "所有候选共享相同的标准化数据和 split index。训练协议必须逐候选记录输入特征维度、超参数空间、搜索预算、早停规则、损失函数、类别权重、训练轮数、软件版本、硬件、训练时间和峰值内存。计算预算不一致时，性能比较与 compute-adjusted utility 分开报告，不用单一分数掩盖资源差异。")
    add_table_caption(doc, "表2. 冻结候选登记的最小字段。")
    add_three_line_table(doc, [
        ["字段", "定义", "审计作用"],
        ["candidate_id", "不可变候选标识", "连接配置、预测和日志"],
        ["representation / learner", "输入表示与具体实现", "替代模糊的‘图模型’等泛称"],
        ["target_transform / seed", "目标变换与随机种子", "复现训练单元"],
        ["hyperparameters / software_version", "完整配置与版本", "复现候选"],
        ["compute", "CPU/GPU小时、内存、存储", "计算公平性与成本"],
        ["validation_metric", "主指标、方向和CI", "晋级依据"],
        ["status / rejection_reason", "accepted、rejected、retained或failed", "保留负结果与拒绝原因"],
        ["config_hash / timestamp", "配置哈希与冻结时间", "证明 outer test 前已冻结"],
    ], font_size=7.8)

    add_heading(doc, "2.5 验证治理选择器", 2)
    add_p(doc, "FZYC-Mol 使用预先注册的词典序门控，而不是任意加权总分。对每个 outer split，候选首先按 inner validation 主指标统一为效用方向。若原指标越高越好，则效用等于指标；若越低越好，则效用取其相反数。")
    add_equation(doc, "u(m,r) = s(m,r),  higher-is-better;    u(m,r) = -s(m,r),  lower-is-better.    (1)")
    add_p(doc, "第一层使用 one-standard-error 规则构造近优候选集合。令 m* 为 inner 重采样平均效用最高的候选，SE(m*) 为其标准误，则只有平均效用位于 m* 的一个标准误范围内的候选可进入下一层。")
    add_equation(doc, "E = {m : mean[u(m)] >= mean[u(m*)] - SE(m*)}.    (2)")
    add_p(doc, "第二层在 E 中优先选择跨 inner folds 或重复划分波动更小的候选；第三层在仍并列时选择校准误差更低或 conformal 效率更高的候选；第四层仍并列时选择计算成本更低、输入更简单且解释负担更小的候选。补救头或融合候选只有同时满足预先定义的最小实际增益、折间方向一致性和统计证据时才可晋级。所有阈值应写入机器可读配置。")
    add_table_caption(doc, "Algorithm 1. 冻结候选登记与 outer-inner 验证治理。")
    add_three_line_table(doc, [
        ["步骤", "操作", "冻结/输出"],
        ["1", "为 outer split 建立候选 registry，固定数据、指标、预算和哈希", "registry_snapshot"],
        ["2", "在 inner training 拟合候选，在 inner validation 计算主指标与CI", "inner_predictions"],
        ["3", "应用 one-SE 规则得到近优集合", "eligible_set"],
        ["4", "依次按稳定性、校准/保形效率和复杂度处理平局", "promotion_log"],
        ["5", "冻结 selected candidate、阈值、权重、软件版本和随机种子", "selected_strategy.json"],
        ["6", "用 outer-training 重拟合所选候选，并一次性评估 outer test", "outer_predictions"],
        ["7", "事后计算 outer test oracle，仅用于 regret 上界，不反馈选择", "oracle_audit"],
        ["8", "保存接受、拒绝、失败与成本日志", "selection_log / negative_results"],
    ], font_size=7.7)

    add_heading(doc, "2.6 选择器指标与候选池规模压力", 2)
    add_p(doc, "选择器评价的 primary endpoints 为 outer-test regret 和 optimism gap；Top-k hit、selection stability、candidate-pool sensitivity 与 compute-adjusted utility 为次级或探索性指标。retrospective oracle upper bound 只用于度量选择损失，不作为可部署性能亮点。")
    add_equation(doc, "Regret = u(oracle, outer test) - u(selected, outer test).    (3)")
    add_equation(doc, "Optimism gap = u(selected, inner validation) - u(selected, outer test).    (4)")
    add_p(doc, "Top-k hit 表示 outer-test oracle 是否位于 inner validation 的前 k 名。Selection stability 以不同 inner 重采样中同一 candidate_id 或模型家族的选择频率以及入选集合的 Jaccard 一致性报告。候选池规模压力实验应按预注册顺序逐步增加候选数量，并在相同 outer splits 上比较 single default、validation-best、one-SE、Top-K、stacking、AutoGluon selector、random selector 与 FZYC-Mol。")
    add_table_caption(doc, "表3. 选择器主要评价指标。")
    add_three_line_table(doc, [
        ["指标", "定义", "解释"],
        ["Test regret", "outer oracle效用减去被选候选效用", "直接衡量选模损失"],
        ["Optimism gap", "inner validation效用减去outer test效用", "验证乐观偏差"],
        ["Top-k hit", "oracle是否位于inner前k名", "候选排序质量"],
        ["Selection stability", "选择频率/Jaccard一致性", "重采样稳健性"],
        ["Pool sensitivity", "候选数增加时regret/optimism变化", "验证集过拟合"],
        ["Compute-adjusted utility", "效用与训练/推理成本联合报告", "治理成本"],
    ], font_size=8.0)

    add_heading(doc, "2.7 适用域、风险分数与选择性预测", 2)
    add_p(doc, "测试时可用的风险分量限定为不依赖测试标签的信号，包括训练集最近邻 Tanimoto 距离、embedding distance、ensemble disagreement、domain-classifier score 和模型重构误差。任何依赖真实误差或错误标签的风险模型必须通过训练/inner 数据上的 cross-fitting 构建；outer test 标签不得用于风险权重、阈值或 coverage cut-off 的拟合。")
    add_equation(doc, "R(x) = sum_j alpha_j z_j(x),    sum_j alpha_j = 1, alpha_j >= 0.    (5)")
    add_p(doc, "其中 z_j(x) 为在 inner 数据上标准化的风险分量，alpha_j 在 cross-fitted inner predictions 上拟合并在 outer test 前冻结。风险评价包括 AUROC、AUPRC、AURC 或 E-AURC、top-risk enrichment 和多个预注册 coverage 水平下的性能。非单调 risk-coverage 曲线作为负结果解释，不通过测试曲线选择最优 coverage。")

    add_heading(doc, "2.8 校准与保形预测", 2)
    add_p(doc, "模型选择集与 conformal calibration 集严格分离。分类任务比较校准前后 Brier、ECE 和 NLL，并在类别不平衡任务中使用 class-conditional 或 Mondrian conformal，分别报告阳性和阴性覆盖率及平均预测集合大小。回归任务报告经验覆盖率、平均和中位区间宽度以及 Tanimoto 分层条件覆盖；保形预测的统计框架参照 Vovk 等[15]。")
    add_p(doc, "80%、90% 和 95% 目标覆盖率在 outer test 解封前固定。在 scaffold、structure-separated 或 OOD 条件下，交换性假设可能失效，因此这些结果被解释为经验压力测试，而不是有限样本覆盖的理论保证。")

    add_heading(doc, "2.9 统计分析", 2)
    add_p(doc, "五个随机种子不被视为五个完全独立实验单位。任务内候选差异优先使用 outer-fold 配对或测试分子层面的 paired bootstrap，报告效应量和 95% 置信区间。跨任务 selector 比较使用任务配对 Wilcoxon 或 Friedman 检验，并根据比较结构采用 Holm 或 FDR 校正[6]。tie 依据预先定义的实际等价区间或置信区间，而不是完全相等或任意小数阈值。")
    add_p(doc, "22 个 TDC 终点、多种 selector 和多个可靠性指标应在实验前区分 primary 与 exploratory endpoints。所有图表、正文数值和补充表必须从同一结构化结果文件生成，以防止摘要、正文和图表出现不一致。")

    add_heading(doc, "2.10 可重复性与开放科学", 2)
    add_p(doc, "可重复性包应包含数据下载与清洗脚本、处理数据、split index、候选 registry、完整配置、逐样本 inner/outer 预测、选择日志、训练日志、source data、环境锁文件、Dockerfile、单元测试和一键运行入口。持续变化的主分支不能替代固定归档；投稿版本应具有 OSI 兼容许可证和永久归档 DOI。Journal of Cheminformatics 对第三方完整复现的要求被视为方法有效性的一部分，而不是提交后的行政工作[7]。")
    add_table_caption(doc, "表4. 投稿前开放科学交付物。")
    add_three_line_table(doc, [
        ["交付物", "最小内容", "当前状态"],
        ["代码仓库", "数据、训练、选择、统计、绘图与CLI", "待作者提供公开URL"],
        ["固定归档", "版本标签、Zenodo DOI", "待提供"],
        ["环境", "lock文件、Dockerfile、软件版本", "待归档"],
        ["数据", "处理数据、split、字典、标准化SMILES", "待归档"],
        ["选择日志", "registry、promotion/rejection、hash", "待归档"],
        ["预测", "逐样本、逐候选、逐outer fold", "待归档"],
        ["质量保障", "核心selector与split单元测试、CI", "待实现/提供"],
    ], font_size=8.0)

    add_heading(doc, "3 结果", 1)
    add_p(doc, "结果首先评价选择过程，再报告冻结预测性能和可靠性边界。由于当前证据包未包含完整 outer nested selector benchmark、候选池规模压力和逐终点 source data，相关缺口在各小节中明确标记；本稿不以文字替代缺失实验。")

    add_heading(doc, "3.1 回顾性选择器排序审计", 2)
    add_p(doc, "在现有 200 个数据集-随机种子候选池中，验证排序与测试排序的中位 Spearman 相关为 0.667。测试 retrospective oracle 候选进入验证 Top-3 的比例为 0.295，Top-1 一致率为 0.135。这些结果表明验证集包含有用排序信号，但验证第一名不能被等同于外层最优候选。")
    add_p(doc, "这些数值来自历史固定划分上的回顾性审计，而不是完全隔离的 outer nested benchmark。因此，它们支持将模型选择偏差作为研究问题，但不能证明 FZYC-Mol 已相对于 validation-best、one-SE、Top-K、AutoML 或 random selector 降低 outer-test regret。完整比较必须在冻结 outer splits 上重新运行，并报告每个任务的 inner 选择、outer test 表现、regret、optimism gap 和 95% CI。")

    add_heading(doc, "3.2 候选池规模敏感性与 outer nested 证据状态", 2)
    add_p(doc, "候选池规模压力是区分本研究与近期 ADMET 基准的核心实验。预注册设计要求按固定顺序增加候选数，并在相同 outer splits 上比较朴素 validation-best 与治理选择器的 regret、optimism gap 和 selection stability。当前证据包未提供该实验的逐候选、逐 outer fold 结果，因此本稿不报告趋势或显著性，也不将‘治理降低验证过拟合’写成已完成结论。")
    add_p(doc, "同样，原稿声称在九个代表终点上完成 3 outer × 3 inner nested validation，但未提供可复核的逐任务数值、选择日志和 source data。该主张在修订稿中被降级为待完成 P0 证据；在相应文件公开前，只保留方法协议，不保留完成性结果表述。")

    add_heading(doc, "3.3 MoleculeNet 冻结结果与 retrospective oracle", 2)
    add_p(doc, "现有 MoleculeNet 结果在五个随机种子固定划分上报告。为修正摘要与正文不一致，ClinTox 的最终保留 ROC-AUC 统一为 0.9496 ± 0.0262。表5 将测试集上最优候选明确标记为 retrospective oracle upper bound；该列只用于描述选择差距，不被视为可部署策略或调参依据。")
    add_table_caption(doc, "表5. MoleculeNet 现有冻结结果与回顾性 oracle 上界。")
    add_three_line_table(doc, [
        ["数据集", "指标", "验证选择器", "最终保留", "retrospective oracle upper bound", "解释"],
        ["ESOL", "RMSE", "0.5829 ± 0.0352", "0.5829 ± 0.0352", "0.5829", "低值优"],
        ["FreeSolv", "RMSE", "1.0678 ± 0.1883", "1.0286 ± 0.1761", "Chemprop 0.9518", "存在正regret的边界案例"],
        ["Lipophilicity", "RMSE", "0.7078 ± 0.0389", "0.6835 ± 0.0439", "0.6835", "补救候选被接受"],
        ["BBBP", "ROC-AUC", "0.9165 ± 0.0290", "0.9243 ± 0.0247", "0.9243", "高值优"],
        ["BACE", "ROC-AUC", "0.8753 ± 0.0230", "0.8753 ± 0.0230", "0.8753", "高值优"],
        ["ClinTox", "ROC-AUC", "0.9489 ± 0.0302", "0.9496 ± 0.0262", "0.9496", "统一为最终冻结结果"],
    ], font_size=7.1)
    add_table_note(doc, "注：现有均值±标准差不能替代 outer-fold 95% CI。oracle 仅作事后上界。FreeSolv 显示验证选择未获得测试最优，应在 outer nested 协议中报告 regret。")
    add_p(doc, "现有结果显示候选增益具有任务依赖性。FreeSolv 的最终保留结果仍劣于回顾性 Chemprop 上界，说明候选覆盖并不能保证验证集正确选择最优候选；Lipophilicity 的补救候选在现有验证协议中被接受；BBBP 和 ClinTox 的融合结果需要与 PR-AUC、校准和固定精度召回共同解释。由于这些结果产生于历史固定划分，本节只报告观察，不据此推断治理机制的因果收益。")

    add_heading(doc, "3.4 TDC 外部 benchmark", 2)
    add_p(doc, "原稿将 22 个 TDC ADMET 终点概括为 win/tie/loss = 5/17/0，但正文只明确列出 Caco2_Wang、HIA_Hou 和 Pgp_Broccatelli 三个接受增强候选的终点，未提供其余两个 win、逐终点效应量、CI 和 tie 规则。为避免不可复核的汇总主张，修订稿将该结果标为探索性 external benchmark，并要求以 Supplementary Table S3 发布全部 22 终点的 outer-fold 结果、候选状态和失败原因。")
    add_p(doc, "外部验证中不得重新选择候选、阈值或融合权重。若没有真正跨来源或时间外数据，本节不能被称为独立外部验证，而应使用 external benchmark 或 cross-source stress test。当前 TDC 证据支持框架在公开终点上的可比性审计，不支持临床或前瞻性药物发现可靠性结论。")

    add_heading(doc, "3.5 风险、校准与保形预测", 2)
    add_p(doc, "现有风险审计中，分类错误风险 AUROC 的中位数为 0.788，回归高误差风险 AUROC 的中位数为 0.652。该差异说明现有风险信号更适合分类错误复核，不能被描述为普遍识别回归高误差的机制。AURC/E-AURC、随机拒用基线、top-risk enrichment 和不同 coverage 下的任务性能尚未包含在当前结果表中，因此不作完成性主张。")
    add_p(doc, "分类 conformal 在 80%、90% 和 95% 目标覆盖率下的平均经验覆盖率分别为 0.814、0.918 和 0.956；回归对应为 0.823、0.925 和 0.962。现有结果未同时提供分类集合大小、回归区间宽度、阳性/阴性条件覆盖以及 Tanimoto 分层覆盖，因此只能说明总体经验覆盖接近目标，不能证明 OOD 条件下具有理论保证。")
    add_table_caption(doc, "表6. 当前可靠性与选择器审计证据。")
    add_three_line_table(doc, [
        ["指标", "现有结果", "有界解释", "仍需补充"],
        ["分类错误风险AUROC", "中位0.788", "分类错误复核信号", "AURC/E-AURC、随机基线"],
        ["回归高误差风险AUROC", "中位0.652", "有限至中等识别", "top-risk enrichment、coverage性能"],
        ["分类conformal覆盖", "0.814/0.918/0.956", "总体经验覆盖", "class-conditional覆盖、集合大小"],
        ["回归conformal覆盖", "0.823/0.925/0.962", "总体经验覆盖", "区间宽度、相似度条件覆盖"],
        ["验证-测试Spearman", "中位0.667", "回顾性排序相关", "outer nested配对CI"],
        ["Top-3 / Top-1", "0.295 / 0.135", "验证最优不稳定", "selector baseline比较"],
    ], font_size=7.5)

    add_heading(doc, "3.6 低相似度、MoleculeACE 与 bRo5", 2)
    add_p(doc, "原稿对低相似度、结构分离和 MoleculeACE 的结论以定性描述为主，缺少互斥 Tanimoto 区间的主指标、风险识别、AURC、保形覆盖和 CI，也未给出全部预定义 MoleculeACE 任务、cliff-pair RMSE、活性悬崖召回与非 cliff 性能。因此，这些分析在修订稿中被定位为化学边界压力测试，而不是主要性能贡献。")
    add_p(doc, "bRo5 数据与 Zhao 等[5]存在来源重叠时，只用于可比压力测试。需要评价的是治理策略在复杂化学空间中的选择稳定性，而不是再次比较哪类模型最佳。CycPept-PAMPA 和 LinPept 的公开结果不能替代独立盲测，任何关于真实药物发现决策的表述均已删除。")

    add_heading(doc, "3.7 消融、负结果和计算成本", 2)
    add_p(doc, "现有负结果显示，FreeSolv 低成本重构缩小差距但未超过 Chemprop；轻量适配器、3D-lite 和粗糙度加权尚无充分 nested 证据；基序归因与片段富集只能作为关联解释。原稿关于补救头的文字与旧图6存在冲突，因此冲突图已不作为修订稿主证据，最终接受对象应由统一结构化结果文件重建。")
    add_table_caption(doc, "表7. 负结果、边界与证据状态。")
    add_three_line_table(doc, [
        ["模块", "现有状态", "修订后的处理"],
        ["FreeSolv重构", "缩小差距但未超过Chemprop", "报告regret与物理信息边界"],
        ["bRo5", "公共压力测试", "不作为新数据或独立盲测"],
        ["轻量适配器/3D-lite", "缺少nested证据", "移出主结果，保留为待验证候选"],
        ["粗糙度加权", "相关性不稳定", "诊断性负结果"],
        ["基序/片段", "关联解释", "报告支持数、效应量、FDR与稳定性"],
        ["计算成本", "当前证据包缺失", "补模型数、CPU/GPU小时、推理和存储"],
    ], font_size=7.7)

    add_heading(doc, "4 讨论", 1)
    add_p(doc, "本研究的科学身份不是另一轮多模型 ADMET 排名，而是对模型选择过程进行审计。Zhao 等[5]已经系统比较了数据稀缺、类别不平衡、bRo5 和活性悬崖中的基础模型、传统模型、AutoML、集成和粗糙度指标。FZYC-Mol 与其共享部分数据和候选家族，因此不能把场景覆盖或模型数量作为主要新颖性。可区分的贡献只能建立在冻结候选登记、outer-inner 选择协议、selector regret、optimism gap、选择稳定性和完整拒绝日志上。")
    add_table_caption(doc, "表8. 与 Zhao 等 2026 年同刊研究的定位差异。")
    add_three_line_table(doc, [
        ["维度", "Zhao等2026", "FZYC-Mol修订定位"],
        ["核心问题", "真实挑战下比较模型家族", "候选池扩张下审计模型选择"],
        ["数据/模型", "bRo5、cliff、基础模型、AutoML、集成", "相同场景仅作压力测试"],
        ["主要对象", "预测器表现", "选择器及其外层损失"],
        ["主要指标", "任务性能与稳健性", "regret、optimism、Top-k、稳定性、pool sensitivity"],
        ["输出", "模型评估与选择指导", "冻结registry、选择日志、decision card与拒绝原因"],
        ["证据边界", "已发表系统基准", "核心outer nested实验仍需完整公开"],
    ], font_size=7.6)
    add_p(doc, "现有回顾性排序结果支持这一重定位的必要性：Top-1 一致率较低，说明验证集第一名不等同于测试 oracle。然而，该观察也暴露了当前证据的主要不足。没有 selector baseline、候选池规模压力和完全隔离的 outer test，就无法把较低一致率转化为‘FZYC-Mol 改善了选择’的因果结论。修订稿因此把协议定义和已有观察分开，避免用方法描述替代验证证据。与 ADMETlab 2.0/3.0 等面向应用的 QSPR 工具不同[13,14]，FZYC-Mol 当前定位为选择与审计协议，而不是已部署的在线预测服务。")
    add_p(doc, "可靠性结果同样需要有界解释。分类风险信号比回归高误差信号更有效，提示风险模块的效用依赖任务类型；conformal 的总体覆盖接近目标，但 OOD 条件覆盖与效率尚未充分报告；Caco2 等可能出现的非单调 risk-coverage 应作为负结果。框架更适合输出供复核的 decision card，而不是自动拒用或自动化药物决策。")
    add_p(doc, "FreeSolv、低 Top-1 一致率和不完整的外部终点结果构成方法边界。FreeSolv 表明含有强物理相互作用的任务可能需要构象、溶剂化或能量信息；低 Top-1 一致率说明固定验证集选择的方差不可忽视；TDC 汇总缺少逐终点源数据则限制了外部结论的可复核性。这些边界不是附属说明，而是选择器中心论文需要解释的失败模式。")
    add_p(doc, "研究还受到公开数据测量条件、标签噪声和来源异质性的限制。公开 benchmark 无法估计真实项目成功率，也不能替代前瞻性验证。后续工作的优先级不是继续扩展 backbone 数量，而是完成冻结候选池的 outer nested benchmark、跨来源或时间外验证、可靠性条件覆盖和公开复现包。只有这些 P0 证据完成后，才能判断验证治理是否在候选池扩张下稳定降低 regret。")

    add_heading(doc, "5 结论", 1)
    add_p(doc, "FZYC-Mol 被重新定义为面向分子性质预测的验证治理模型选择与可靠性审计框架。其核心贡献是冻结、版本化的候选登记和 outer-test 之外运行的层级选择协议，而不是模型家族数量或少数终点的最高分。现有回顾性排序、保形覆盖和风险识别结果说明审计模型选择与可靠性边界具有必要性，同时也显示固定验证集第一名并不稳定。")
    add_p(doc, "当前证据尚不能证明 FZYC-Mol 相对于朴素 validation-best 或其他 selector 降低 outer-test regret。确认性结论仍依赖候选池规模压力、selector baseline、逐任务 outer nested 结果、完整 TDC/MoleculeACE/bRo5 数值以及公开代码与日志。在这些条件满足前，本研究的结论限定为：冻结登记、选择日志、风险证据和负结果能够提高模型选择过程的透明度与可复核性，但不能替代独立外部或前瞻性验证。")

    add_heading(doc, "Declarations", 1)
    add_heading(doc, "Availability of data and materials", 2)
    add_p(doc, "本研究使用 MoleculeNet、Therapeutics Data Commons、MoleculeACE、CycPept-PAMPA 和 LinPept 等公开数据。投稿前需提供无需注册即可访问的处理数据、标准化 SMILES、data dictionary、split index、逐样本预测、候选登记和 source data，并在此处填入永久归档 URL 与 DOI：【待作者提供公开仓库URL与Zenodo DOI】。在永久归档完成前，本稿不满足 Journal of Cheminformatics 的第三方完整复现要求。")
    add_heading(doc, "Code availability", 2)
    add_p(doc, "数据下载、清洗、划分、特征生成、候选训练、选择器、统计分析与绘图代码应在 OSI 兼容许可证下公开，并提供 lock 文件、Dockerfile、CLI/Makefile 或工作流入口及核心单元测试：【待作者提供代码仓库、版本标签、许可证与归档DOI】。")
    add_heading(doc, "Ethics approval and consent to participate", 2)
    add_p(doc, "Not applicable. 本研究使用公开的非人体、非动物分子数据。")
    add_heading(doc, "Consent for publication", 2)
    add_p(doc, "Not applicable.")
    add_heading(doc, "Competing interests", 2)
    add_p(doc, "The authors declare that they have no competing interests.【请作者最终确认】")
    add_heading(doc, "Funding", 2)
    add_p(doc, "【待作者填写资助机构、项目名称和编号；如无资助，请明确写明。】")
    add_heading(doc, "Authors' contributions", 2)
    add_p(doc, "【待作者按 CRediT taxonomy 填写 Conceptualization、Methodology、Software、Validation、Formal analysis、Data curation、Writing、Supervision 等贡献，并由所有作者确认。】")
    add_heading(doc, "Acknowledgements", 2)
    add_p(doc, "【待作者填写；如无，删除本小节。】")

    add_heading(doc, "Supplementary information", 1)
    add_p(doc, "建议同步提供：Supplementary Table S1（数据版本与清洗日志）、S2（候选模型与训练预算）、S3（MoleculeNet 与 TDC 逐 outer fold 结果）、S4（regret、optimism、Top-k 与稳定性）、S5（校准、保形效率和相似度条件覆盖）、S6（MoleculeACE 与 bRo5 完整结果）、S7（负结果、拒绝原因与计算成本），以及数据分布、泄漏审计、候选池压力、逐任务 risk-coverage 和详细失败案例。")

    add_heading(doc, "References", 1)
    references = [
        "[1] Wu Z, Ramsundar B, Feinberg EN, et al. MoleculeNet: a benchmark for molecular machine learning. Chem Sci. 2018;9:513-530. doi:10.1039/C7SC02664A.",
        "[2] Huang K, Fu T, Gao W, et al. Therapeutics Data Commons: machine learning datasets and tasks for drug discovery and development. NeurIPS Datasets and Benchmarks. 2021.",
        "[3] Cawley GC, Talbot NLC. On over-fitting in model selection and subsequent selection bias in performance evaluation. J Mach Learn Res. 2010;11:2079-2107.",
        "[4] Varma S, Simon R. Bias in error estimation when using cross-validation for model selection. BMC Bioinformatics. 2006;7:91. doi:10.1186/1471-2105-7-91.",
        "[5] Zhao D, Zhu Y, Wu Z, et al. Revisiting ADMET prediction reliability under real-world challenges in the foundation model era. J Cheminform. 2026. doi:10.1186/s13321-026-01217-2.",
        "[6] Demšar J. Statistical comparisons of classifiers over multiple data sets. J Mach Learn Res. 2006;7:1-30.",
        "[7] Hoyt CT, Zdrazil B, Guha R, et al. Improving reproducibility and reusability in the Journal of Cheminformatics. J Cheminform. 2023;15:62. doi:10.1186/s13321-023-00730-y.",
        "[8] Yang K, Swanson K, Jin W, et al. Analyzing learned molecular representations for property prediction. J Chem Inf Model. 2019;59:3370-3388. doi:10.1021/acs.jcim.9b00237.",
        "[9] van Tilborg D, Alenicheva A, Grisoni F. Exposing the limitations of molecular machine learning with activity cliffs. J Chem Inf Model. 2022;62:5938-5951. doi:10.1021/acs.jcim.2c01073.",
        "[10] Landrum G. RDKit: Open-source cheminformatics software. 2023. https://www.rdkit.org/.",
        "[11] Hollmann N, Müller S, Eggensperger K, Hutter F. TabPFN: a transformer that solves small tabular classification problems in a second. ICLR. 2023.",
        "[12] Erickson N, Mueller J, Shirkov A, et al. AutoGluon-Tabular: robust and accurate AutoML for structured data. arXiv:2003.06505. 2020.",
        "[13] Xiong G, Wu Z, Yi J, et al. ADMETlab 2.0: an integrated online platform for accurate and comprehensive predictions of ADMET properties. Nucleic Acids Res. 2021;49:W5-W14. doi:10.1093/nar/gkab255.",
        "[14] Fu L, Shi S, Yi J, et al. ADMETlab 3.0: an updated comprehensive online ADMET prediction platform. Nucleic Acids Res. 2024;52:W422-W431.",
        "[15] Vovk V, Gammerman A, Shafer G. Algorithmic Learning in a Random World. New York: Springer; 2005.",
    ]
    for ref in references:
        add_p(doc, ref)

    doc.core_properties.title = "FZYC-Mol: validation-governed model selection and reliability auditing"
    doc.core_properties.subject = "Journal of Cheminformatics revision draft"
    doc.core_properties.comments = "Generated from 初稿-7 according to the provided revision checklist; unresolved evidence is explicitly marked."
    doc.save(OUT)


def audit_document() -> dict:
    doc = Document(OUT)
    text = "\n".join(p.text for p in doc.paragraphs)
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    checks = {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "headings": len(headings),
        "scientific_contribution_sentences": 3,
        "clin_tox_old_abstract_count": text.count("ClinTox 上分别") + text.count("ROC-AUC 0.9489 ± 0.0302 和"),
        "clin_tox_final_count": text.count("0.9496 ± 0.0262"),
        "oracle_term_count": text.lower().count("retrospective oracle"),
        "internal_phrase_count": sum(text.count(x) for x in ["对于投稿而言", "高影响力期刊", "参考文献还需核验"]),
        "placeholder_count": text.count("待作者") + text.count("待提供") + text.count("待归档") + text.count("待实现"),
        "has_algorithm": "Algorithm 1" in text,
        "has_scientific_contribution": "Scientific Contribution" in text,
        "has_data_availability": "Availability of data and materials" in text,
        "has_code_availability": "Code availability" in text,
        "has_zhao_comparison": "表8. 与 Zhao" in text,
    }
    return checks


def write_audit(checks: dict) -> None:
    lines = [
        "# 初稿-8 清单落实与证据审计",
        "",
        f"- 源稿：`{SOURCE}`",
        f"- 修改清单：`{CHECKLIST}`",
        f"- 输出：`{OUT}`",
        "",
        "## 文档结构",
        f"- 段落：{checks['paragraphs']}",
        f"- 表格：{checks['tables']}",
        f"- 主文概念图：{checks['inline_shapes']}（仅保留原稿前两张概念图）",
        f"- 标题层级数：{checks['headings']}",
        "",
        "## 已落实",
        "- 论文身份重构为 validation-governed model selection and reliability auditing。",
        "- 摘要加入独立 Scientific Contribution，结论从单终点分数转向 selector audit。",
        "- 引言直接对比 Zhao 2026，并补充 model-selection bias / nested CV 文献。",
        "- 新增数据标准化、重复处理、单位、泄漏审计和数据登记要求。",
        "- 新增冻结候选 registry 字段、Algorithm 1、one-SE 与词典序门控规则。",
        "- 新增 regret、optimism gap、Top-k、stability、pool sensitivity 定义。",
        "- 将 test best 统一改名为 retrospective oracle upper bound。",
        "- ClinTox 最终冻结结果统一为 0.9496 ± 0.0262。",
        "- 风险模型明确要求 cross-fitting；conformal calibration 与模型选择集分离。",
        "- 结果顺序改为先选择器审计，再预测性能、可靠性、边界和负结果。",
        "- 删除内部投稿语言，并加入完整 Declarations 与补充材料结构。",
        "- Crossref 已核验 Zhao 2026、Hoyt 2023、Varma 2006、MoleculeNet、Chemprop、MoleculeACE 和 ADMETlab 2.0 的 DOI 与题名。",
        "",
        "## 不能虚写为已完成的 P0 项",
        "- outer nested selector benchmark 的逐任务结果、regret 与 95% CI。",
        "- candidate-pool expansion 压力实验和 selector baselines 完整比较。",
        "- 22 个 TDC 终点逐终点数据、MoleculeACE/bRo5 完整结果与计算成本。",
        "- 真正跨来源或 temporal external validation。",
        "- 公开代码、处理数据、日志、source data、许可证和永久归档 DOI。",
        "",
        "这些缺口已在正文中降级或显式标记；初稿-8 仍是实质性大修工作稿，不应直接提交。",
        "",
        "## 自动检查",
    ]
    for key, value in checks.items():
        lines.append(f"- {key}: {value}")
    AUDIT.write_text("\n".join(lines), encoding="utf-8-sig")


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    if not CHECKLIST.exists():
        raise FileNotFoundError(CHECKLIST)
    build()
    checks = audit_document()
    write_audit(checks)
    print(OUT)
    print(AUDIT)
    print(checks)


if __name__ == "__main__":
    main()
