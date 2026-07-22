from __future__ import annotations

import csv
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


def find_file(base: Path, required_tokens: list[str]) -> Path:
    matches: list[Path] = []
    for path in base.rglob("*.docx"):
        if all(token in path.name for token in required_tokens):
            matches.append(path)
    if not matches:
        raise FileNotFoundError(f"No .docx matching {required_tokens} under {base}")
    return max(matches, key=lambda p: p.stat().st_mtime)


def clear_paragraph(paragraph: Paragraph) -> None:
    if hasattr(paragraph, "clear"):
        paragraph.clear()
        return
    for run in paragraph.runs:
        run.text = ""


def set_text(paragraph: Paragraph, text: str, bold: bool = False) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)


def replace_by_prefix(doc: Document, prefix: str, new_text: str) -> bool:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            set_text(paragraph, new_text)
            return True
    return False


def insert_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    set_text(new_paragraph, text)
    return new_paragraph


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    paragraph.paragraph_format.space_after = Pt(0)


def ensure_table_size(table, rows: int, cols: int) -> None:
    while len(table.rows) < rows:
        table.add_row()
    while len(table.rows) > rows:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)
    if len(table.columns) != cols:
        raise ValueError(f"Expected {cols} columns, got {len(table.columns)}")


def set_border(element, edge: str, val: str, size: str = "8", color: str = "000000") -> None:
    tag = "w:tblBorders" if element.tag.endswith("tblPr") else "w:tcBorders"
    borders = element.find(qn(tag))
    if borders is None:
        borders = OxmlElement(tag)
        element.append(borders)
    edge_el = borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        borders.append(edge_el)
    edge_el.set(qn("w:val"), val)
    if val != "nil":
        edge_el.set(qn("w:sz"), size)
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)


def apply_three_line_table(table) -> None:
    tbl_pr = table._tbl.tblPr
    for edge in ("left", "right", "insideH", "insideV"):
        set_border(tbl_pr, edge, "nil")
    set_border(tbl_pr, "top", "single", size="12")
    set_border(tbl_pr, "bottom", "single", size="12")

    if not table.rows:
        return
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        for edge in ("left", "right", "top"):
            set_border(tc_pr, edge, "nil")
        set_border(tc_pr, "bottom", "single", size="8")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in table.rows[1:]:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for edge in ("left", "right", "top", "bottom"):
                set_border(tc_pr, edge, "nil")


def update_experiment_matrix(table) -> list[list[str]]:
    rows = [
        ["实验模块", "证据层级", "核心问题", "数据/候选", "输出指标", "论文位置"],
        [
            "强基线同划分比较",
            "主线证据",
            "结论是否依赖弱 baseline",
            "MoleculeNet、TDC ADMET；CatBoost、XGBoost、LightGBM、ExtraTrees、RF、Chemprop、TabPFNv2、AutoGluon、KPGT 表征",
            "主指标、配对差值、win/tie/loss、regret、候选接受状态",
            "3.5、4.1、4.5；完整表入补充材料",
        ],
        [
            "FreeSolv 定向补救",
            "边界案例",
            "溶剂化相关任务是否需要物理代理特征",
            "Morgan/RDKit、目标变换、Top-K/stacking、3D-lite 候选与高误差样本审计",
            "RMSE、MAE、seed 级差值、最近邻标签差异、补救接受/拒绝",
            "4.1、4.7、5.3",
        ],
        [
            "bRo5 化学空间",
            "数据状态/待验证",
            "规则五以外空间能否形成可核验证据",
            "CycPept PAMPA、LinPept NonFouling、LinPept CellPen 或可公开替代数据",
            "数据可用性、字段缺失、样本量、适用域覆盖；数据未就绪时不报告性能",
            "3.5、4.6、5.3",
        ],
        [
            "MoleculeACE 活性悬崖",
            "困难结构证据",
            "相似结构标签跳变是否被平均指标掩盖",
            "当前可核验任务子集，后续扩展到更多可运行任务；cliff-aware 候选与 KPGT/专用模型对照",
            "RMSE、MAE、R2、cliff-pair RMSE、gap Spearman、cliff recall、案例对",
            "4.6、4.7",
        ],
        [
            "OOD、Perimeter 与低相似度压力测试",
            "外推边界",
            "模型在新骨架和化学空间边缘是否稳定",
            "Random、Scaffold、Structure-separated、Perimeter、低相似度互斥三档",
            "性能下降、校准、风险富集、最近邻 Tanimoto、样本量不确定性",
            "4.3、4.6",
        ],
        [
            "不平衡 ADMET 分类",
            "筛选实用性",
            "高 ROC-AUC 是否伴随阳性召回和概率可信度",
            "ClinTox、DILI、hERG、CYP 相关终点；class weight、采样、阈值移动、focal loss 候选",
            "ROC-AUC、PR-AUC、Brier、ECE、MCC、balanced accuracy、Recall@fixed Precision",
            "4.4、4.6",
        ],
        [
            "嵌套验证与选择器偏差审计",
            "选择风险诊断",
            "候选池是否放大验证集过拟合",
            "代表性 MoleculeNet/TDC 终点；3 outer x 3 inner，关键任务可扩展",
            "outer metric、candidate switch rate、regret、Top1/Top3、validation-test Spearman",
            "3.5、4.5、5.3",
        ],
        [
            "完整消融",
            "模块贡献",
            "选择器、融合、AD gate、uncertainty、rescue head、表示模块是否必要",
            "Full、best single、simple mean、w/o selector、w/o fusion、w/o AD gate、w/o uncertainty、w/o motif/fingerprint、w/o rescue head",
            "相对 Full 的变化、任务依赖性、可靠性变化、负结果",
            "4.6；完整矩阵入补充材料",
        ],
        [
            "轻量分子大模型适配器",
            "受控扩展",
            "冻结表征是否限制困难终点",
            "ChemBERTa、MoLFormer、图预训练表征；linear probe、adapter、LoRA 或小规模微调",
            "验证接受状态、低相似度收益、计算成本、过拟合风险",
            "3.5、5.3；未完成者仅作候选方向",
        ],
        [
            "3D-lite 与物理描述符",
            "受控补救",
            "2D 表征是否不足以描述构象或物理相互作用",
            "ETKDG、MMFF/UFF、体积、表面积、形状、构象失败标记",
            "终点级增益、构象失败率、高误差案例、是否被验证门控接受",
            "4.1、4.6、5.3",
        ],
        [
            "外部盲测风格验证/时间切分",
            "迁移边界",
            "内部划分结果能否迁移到独立来源或时间后数据",
            "TDC 外部终点、公开挑战数据、版本或时间切分；外部测试前冻结候选池",
            "外部指标、训练-测试重叠检查、适用域覆盖、外部失败案例",
            "4.2、5.3；无时间戳时写为 holdout/后续验证",
        ],
        [
            "样本级解释与失败证据链",
            "解释边界",
            "模型为什么错、哪些分子不宜过度信任",
            "FreeSolv 高误差、ClinTox 假阴性、低相似度 OOD、MoleculeACE cliff pairs",
            "最近邻、Murcko scaffold、BRICS/官能团、风险分数、标签差异、p/FDR 或探索性标注",
            "4.7、补充案例表",
        ],
    ]
    ensure_table_size(table, len(rows), 6)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            set_cell(table.rows[r_idx].cells[c_idx], value, bold=(r_idx == 0))
    return rows


def polish_document(doc: Document) -> list[str]:
    changes: list[str] = []
    replacements = [
        (
            "分子性质预测是药物发现、ADMET 评估及毒性风险筛查中的关键计算步骤。",
            "分子性质预测已成为药物发现、ADMET 评估和毒性风险筛查中的基础计算环节，但随机划分下的平均 ROC-AUC 或 RMSE 难以充分反映模型在新骨架外推、低相似度分子、不平衡毒性标签、规则五以外化学空间和活性悬崖中的可靠性。本文提出 FZYC-Mol，一种由验证集治理驱动的适用域感知分子性质预测框架。该框架将多视图表示、强基线专家、Random/Scaffold/Perimeter 划分、类别不平衡策略、MoleculeACE 可用任务子集、粗糙度诊断、校准和适用域证据纳入同一冻结候选池；候选接受、拒绝或保留仅由验证集证据决定，测试集在策略冻结后只用于一次性最终评估。bRo5 模块在当前版本中仅作为数据状态审计与后续运行接口，不构成已完成性能评估。",
            "abstract_scope",
        ),
        (
            "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别达到 RMSE",
            "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别达到 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302。验证集接受的定向补救改善了 Lipophilicity；低成本重构缩小了 FreeSolv 与当前选择器之间的差距，但仍未超过观测最佳 Chemprop 候选，因此 FreeSolv 被保留为物理相互作用相关任务的边界案例。多视图融合与适用域门控在 BBBP、ClinTox 以及外部 TDC ADMET 任务 Caco2_Wang、HIA_Hou 和 Pgp_Broccatelli 上产生选择性增益；22 个外部终点的最终保留结果为 win/tie/loss = 5/17/0。上述结果将 FZYC-Mol 界定为终点依赖的选择、拒绝与审计流程，而非保证所有任务统一提升的单一模型。",
            "abstract_results",
        ),
        (
            "可靠性分析显示，风险分数对分类错误具有较强的识别能力",
            "可靠性分析显示，风险分数对分类错误的识别能力强于对回归高误差样本的识别能力；两类任务的中位 AUROC 分别为 0.788 和 0.652。保形预测在 MoleculeNet 分类任务 80%/90%/95% 目标覆盖率下的平均经验覆盖率为 0.814/0.918/0.956，在回归任务下为 0.823/0.925/0.962。验证-测试排名审计显示，跨 200 个 dataset-seed 候选池的中位 Spearman 相关为 0.667，测试最佳候选落入验证 Top-3 的比例为 0.295，Top-1 一致比例为 0.135。这些结果表明，验证集治理可以减少测试集事后选择风险，但不能保证获得测试集最优结果；所有小幅增益均需结合 regret、optimism gap、嵌套验证和负结果审计解释。基序归因与片段富集仅作为关联性化学解释，不作为因果机制证据。",
            "abstract_reliability",
        ),
        (
            "本文贡献可归纳为五个方面。",
            "本文贡献可归纳为五个相互约束的证据层面。第一，围绕 MoleculeNet、TDC ADMET、bRo5 数据状态审计、MoleculeACE 可用任务子集和 Random/Scaffold/Perimeter 划分构建统一评测流程，并明确训练集、验证集和测试集的使用边界。第二，将 TabPFNv2、AutoGluon、KPGT representation、XGBoost-RDKit/Mordred/MorganCount、Chemprop、树模型、冻结表征、采样策略、目标变换和融合策略预先登记为候选，而非依据测试集表现临时追加。第三，提出 accept/reject/retain 治理机制，使补救头、融合、适用域门控和强基线必须先通过验证集证据审查，方可进入最终保留结果。第四，报告 validation-test rank correlation、Top1 match、Top3 hit、regret、optimism gap、不平衡分类、保形预测、risk-coverage、MoleculeACE cliff-pair 指标和 ROGI/MODI/SARI 粗糙度诊断，避免仅依据单一 ROC-AUC 或 RMSE 下结论。第五，将数据缺失、候选失败、低相似度失败和活性悬崖失败写入限制与补充材料，使负结果成为证据链的一部分。",
            "contribution",
        ),
        (
            "为与 Zhao et al. 2026 的真实挑战框架",
            "为与近期真实挑战 ADMET 框架保持可比，本文将补充实验组织为“挑战维度—候选池—验证门控—风险输出”的矩阵。该矩阵覆盖强基线同划分、FreeSolv 边界补救、bRo5 数据状态、MoleculeACE 活性悬崖、OOD/Perimeter/低相似度压力测试、不平衡分类、嵌套验证、完整消融、轻量大模型适配器、3D-lite 物理代理、外部 holdout 与样本级失败案例。所有模块均遵守同一冻结原则：候选是否进入最终保留结果，只能由验证集或预定义审计规则决定；数据未就绪或证据不足的模块被标注为缺失、负结果或后续验证，而不是写作已完成性能提升。",
            "method_matrix_intro",
        ),
        (
            "为保证扩展实验可复现，本文将扩展实验与既有结果解耦",
            "为保证扩展实验可复现，本文将候选登记、训练运行、验证选择、测试评估和负结果记录解耦保存。每个模块均输出稳定命名的候选登记表、字段字典、缺失数据报告、缺失依赖报告和最终决策表；强基线、OOD 划分、不平衡策略、bRo5 数据状态审计、MoleculeACE 可用子集、粗糙度诊断和候选接受/拒绝记录均可追溯到数据划分、随机种子、候选模型和最终表格。具体脚本、配置和字段说明见代码可用性部分及 Supplementary Methods。",
            "method_reproducibility",
        ),
        (
            "分类任务保留 ROC-AUC 作为 MoleculeNet 和若干 TDC 终点的主指标",
            "分类任务保留 ROC-AUC 作为 MoleculeNet 和若干 TDC 终点的主指标，同时报告 PR-AUC、Brier 分数、ECE、MCC、平衡准确率、风险-覆盖曲线、保形覆盖率和固定 precision 条件下的 recall。该指标组合用于降低类别不平衡任务中的解释偏差：ROC-AUC 较高并不必然意味着阳性样本召回、概率校准或筛选富集可靠。ClinTox、DILI、hERG 和 CYP 底物任务因此在不平衡分类增强中重点报告；所有阈值均在验证集上设定，并在测试集上一次性评估。",
            "classification_metrics",
        ),
        (
            "定向补救头服务于已知性能瓶颈或高粗糙度模块",
            "定向补救头服务于已知性能瓶颈或高粗糙度模块，例如 Lipophilicity、FreeSolv、clearance、ppbr、half-life 和 CYP 底物任务。补救头可以来自冻结表征、强表格基线、Top-K/堆叠集成、目标变换、欠采样集成或 3D-lite 物理代理，但需先进入冻结候选池，再由验证集决定是否保留。当前结果显示，Lipophilicity 的定向补救通过验证集门控进入最终策略；FreeSolv 的低成本重构缩小误差但仍保留为边界案例；其余候选若未通过门控，则作为负结果而非主结论处理。",
            "rescue_head",
        ),
        (
            "本节按照证据链组织结果。",
            "本节按证据链组织结果。首先报告 MoleculeNet 和外部 ADMET 的基线与最终保留性能，并明确小幅增益和 FreeSolv 边界。随后评估验证集选择偏差，包括 validation-test ranking、Top-1/Top-3、regret、optimism gap 和代表性 nested validation。接着展示 OOD、Perimeter、低相似度三档和 MoleculeACE 活性悬崖结果，并报告系统消融、固定选择器和负结果。最后，结合校准、risk-coverage、保形预测、化学解释和失败案例说明哪些预测不宜被过度信任。",
            "results_overview",
        ),
        (
            "MoleculeNet 主结果显示，FZYC-Mol 的价值并非来自某一专家在所有任务上的稳定胜出",
            "MoleculeNet 主结果显示，FZYC-Mol 的价值并不来自某一专家在所有任务上的稳定胜出，而来自验证集选择器对不同终点候选策略的可审计接受与拒绝。ESOL 和 BACE 中，验证集选择器与观测最优一致；Lipophilicity 中，定向补救选择器将 RMSE 从 0.7078 降至 0.6835；FreeSolv 中，低成本重构将 RMSE 降至 1.0286 ± 0.1761，但仍未超过观测最佳 Chemprop 候选，因此被保留为选择器稳健性和物理表征不足的边界案例。BBBP 和 ClinTox 的多方法融合候选进入最终保留策略，但增益幅度较小，需结合 seed 级差值、置信区间和选择器审计解释。",
            "moleculenet_results",
        ),
        (
            "为形成完整的验证闭环，本文将关键扩展证据纳入结果逻辑",
            "为形成完整的验证闭环，本文将补充实验转化为结果中的压力测试，而不是作为独立清单堆叠。验证集选择偏差首先由排名审计、Top-3 命中、regret、optimism gap 和 nested validation 诊断；模块贡献随后由 Full、best single、simple mean、w/o selector、w/o fusion、w/o AD gate、w/o uncertainty、w/o motif/fingerprint 和 w/o rescue head 等消融矩阵限定；结构外推由 Random/Scaffold/Structure/Perimeter 和低相似度互斥三档呈现；活性悬崖、bRo5 数据状态、不平衡分类、校准/保形预测和样本级案例用于界定适用边界。上述分析不改变测试集一次性报告原则，只说明为何接受、保留或拒绝候选。",
            "closure_intro",
        ),
        (
            "系统消融与配对统计进一步限定了模块贡献。",
            "系统消融与配对统计进一步限定了模块贡献。完整选择器相对单一 Chemprop、多指纹、冻结预训练表征和核心家族候选总体保持净正向，但 no_chemprop、no_pretrained 或去除特定表示后的负向结果说明，候选家族数量增加本身并不保证性能提升。适用域门控、不确定性加权、融合和补救头的收益具有终点依赖性；未产生稳定增益的模块被保留为负结果，而不是被用于扩大主结论。",
            "ablation_results",
        ),
        (
            "低相似度和结构分离结果提醒，平均性能并不能等同于新骨架外推能力。",
            "低相似度和结构分离结果显示，平均性能不能等同于新骨架外推能力。按最近邻 Tanimoto 相似度划分的高、中、低三档中，低相似度样本通常伴随更高误差、更差校准或更高风险富集；结构分离划分中 FreeSolv、Pgp 和 BACE 等任务出现明显划分惩罚。因此，本文将适用域、risk-coverage、最近邻证据和样本量不确定性与性能分数同时报告，避免把随机划分结果推广到所有外推场景。",
            "low_similarity_results",
        ),
        (
            "MoleculeACE 配对结果与低相似度分析互相补充。",
            "MoleculeACE 配对结果与低相似度分析互相补充。当前可核验结果覆盖 17 个 MoleculeACE 任务和 51 个 seed 配对，而不是完整 30 任务全量审计。主指标包括 R2、RMSE、MAE、cliff-pair RMSE、预测差异与真实差异的 gap Spearman、cliff recall 和 top cliff error；比较对象包括 XGBoost-MorganCount、KPGT representation、可运行的 MOLMCL/KPGT 版本、FZYC selector 以及 KPGT+XGBoost ensemble。代表性配对结果显示，总 RMSE 平均正向变化为 0.0069，悬崖子集 RMSE 平均正向变化为 0.0056，但标准差较大；gap Spearman 平均约为 0.252，部分任务接近零或为负。因此，该模块被解释为活性悬崖风险识别和候选治理的补充证据，而非证明活性悬崖预测已经解决。",
            "moleculeace_results",
        ),
        (
            "最后，负结果和失败案例被明确保留。",
            "最后，负结果和失败案例被明确保留。ClinTox 高风险假阴性表明，高 ROC-AUC 任务仍需样本级风险证据；FreeSolv 低相似度高误差样本提示适用域外推和显式物理相互作用仍是短板；MoleculeACE 活性悬崖失败案例显示，相似分子的真实性质差异仍可能被低估；bRo5 当前仅能提供数据状态审计，不能作为性能结论；3D-lite 和粗糙度加权在 oracle 条件下偶有潜在收益，但验证集门控未稳定接入最终策略。这些负结果共同界定 FZYC-Mol 的使用边界，也避免论文只报告正向结果。",
            "negative_results",
        ),
        (
            "基序/片段解释性分析将模型行为连接到可识别的化学子结构。",
            "基序/片段解释性分析将模型行为连接到可识别的化学子结构，但其作用仅限于关联性解释。本文报告最小支持度、阳性率差异、Fisher 精确检验 p 值和 Benjamini-Hochberg FDR q 值：BBBP 中 N-连接片段和羰基/内酰胺相关 BRICS 片段与穿透性下降相关；BACE 中若干疏水芳香/卤代片段与阳性标签富集相关；ClinTox 中哌嗪/含氮片段、芳香片段和羰基片段显示阳性富集。上述证据不能替代因果机制验证，也不应外推为湿实验机制结论。",
            "motif_results",
        ),
        (
            "这种差异化定位决定了本文的补强方向。",
            "这种差异化定位决定了本文的补强方向。强基线同划分、FreeSolv 补救、不平衡分类、OOD/Perimeter 低相似度压力测试、MoleculeACE、nested validation、完整消融、bRo5 数据状态审计、3D-lite/大模型适配器候选、外部 holdout 和样本级案例共同构成 FZYC-Mol 的治理压力测试。复杂模块并不直接扩大主结论，而是在统一验证证据下被接受、保留、拒绝或标注为数据未就绪。这种证据组织与当前结果中“选择性增益多于普遍胜出”的事实一致。",
            "discussion_positioning",
        ),
        (
            "本研究仍存在若干局限。",
            "本研究仍存在若干局限。首先，验证-测试排名审计和 9 个代表性终点的 3 x 3 nested validation 表明，验证集治理能够减少测试集事后选择，但内外层验证尚未覆盖全部 MoleculeNet、TDC、bRo5 和 MoleculeACE 终点，不能保证所有任务均达到测试最优；小幅增益应与 regret、optimism gap、Top3 hit 和负结果共同解释。其次，bRo5 三个预设数据集在当前审计中均为 missing_data，不能作为已完成性能评估；MoleculeACE 当前可核验结果覆盖 17 个任务和 51 个 seed 配对，不能表述为完整 30 任务全量结论。第三，TabPFNv2、AutoGluon、KPGT representation、轻量大模型适配器、3D-lite、bRo5 和外部时间切分仍需完整同划分验证，尚未完成的候选仅作为候选接口或后续验证方向，不作为已取得性能提升的证据。第四，bRo5、活性悬崖和低相似度样本仍是主要失败边界，FZYC-Mol 更适合识别和审计这些风险，而非宣称已解决它们。第五，基序归因、片段富集和粗糙度相关性仍属于关联证据，不能替代因果机制验证或湿实验验证。",
            "limitations",
        ),
        (
            "进一步验证应优先扩展真实挑战矩阵",
            "进一步验证应优先扩展真实挑战矩阵，而不是简单堆叠更多模型。更有价值的路线是在更多公开 ADMET、bRo5 和活性悬崖终点上系统评估 XGBoost、CatBoost、RandomForest、ExtraTrees、Chemprop、TabPFNv2、AutoGluon、KPGT/GEM/Uni-Mol 表征、目标变换、采样策略、Top-K 集成、3D-lite 和轻量适配器，并将最终保留决策与 roughness、低相似度、Perimeter split、外部 holdout 和 validation-test rank instability 关联起来。",
            "further_validation",
        ),
        (
            "本文提出并系统评估了 FZYC-Mol",
            "本文提出并评估了 FZYC-Mol，一种由验证集治理、适用域感知的多专家分子性质预测框架。结果表明，在 MoleculeNet、TDC ADMET、Random/Scaffold/Perimeter 划分、低相似度子集、MoleculeACE 可用任务子集、ROGI/MODI/SARI 粗糙度诊断、候选接受/拒绝审计、校准/保形预测和基序/片段解释等证据线上，FZYC-Mol 能够提供比单一模型分数更完整的可靠性画像。bRo5 当前仅完成数据状态审计，不作为已完成性能结果。该框架的核心价值不是保证每个终点取得测试最优，而是在固定候选池内透明地接受、保留、拒绝和审计候选策略；当强基线、补救头、低成本重构或多方法融合通过验证集门控时，它们可以进入最终保留结果，当证据不足或数据未就绪时则作为负结果或待验证模块保留。",
            "conclusion",
        ),
        (
            "未来工作可沿三条路线推进。",
            "未来工作可沿三条路线推进。第一，继续扩展强基线同划分、bRo5、MoleculeACE、外部 holdout 和盲测风格验证，以检验选择器在更多真实挑战下的稳定性。第二，在少量代表性低分任务上开展受控适配器、轻量微调或 3D-lite 物理代理实验，并使用嵌套验证限制过拟合。第三，进一步加强样本级可靠性输出，将基序、最近邻、不确定性、适用域、roughness 和实验标签噪声联系起来，为后续湿实验或专家审查提供更清晰的优先级。",
            "outlook",
        ),
    ]
    for prefix, text, label in replacements:
        if replace_by_prefix(doc, prefix, text):
            changes.append(label)
        else:
            changes.append(f"missing:{label}")

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("为与近期真实挑战 ADMET 框架保持可比"):
            insert_after(
                paragraph,
                "按照补充实验方案，表中的每一项均被限定为一个明确的证据功能：已具备结果的模块报告性能、稳定性和负结果；仅具备接口或数据状态的模块只报告可用性、失败原因和后续运行条件。这样处理可以把论文从“增加模型数量”转向“呈现选择过程的可审计性”，同时避免将尚未完成的实验写成已获得的正向结论。",
            )
            changes.append("inserted_method_boundary_paragraph")
            break

    cautious_replacements = {
        "使补救头、融合、适用域门控和强基线必须先通过验证集证据审查": "使补救头、融合、适用域门控和强基线需先通过验证集证据审查",
        "多尺度结构与片段解释能改善部分任务，但必须报告其适用边界": "多尺度结构与片段解释能改善部分任务，但需要同步报告其适用边界",
        "若单一专家显著占优": "若单一专家在验证集上稳定占优",
        "标签差异或归一化目标值跳变显著": "标签差异或归一化目标值跳变较大",
        "复杂融合的收益具有显著的终点依赖性": "复杂融合的收益具有清晰的终点依赖性",
        "失败率显著高于轻量候选": "失败率高于轻量候选",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in cautious_replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            set_text(paragraph, new_text)
            changes.append("cautious_style_replacement")

    if len(doc.tables) >= 3:
        matrix_rows = update_experiment_matrix(doc.tables[2])
        changes.append("updated_experiment_matrix_table")
        out_dir = Path("work/draft15_audit")
        out_dir.mkdir(parents=True, exist_ok=True)
        with (out_dir / "draft15_plan_mapping.csv").open("w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerows(matrix_rows)

    return changes


def apply_document_style(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            if run.font.size is None:
                run.font.size = Pt(10.5)
    for table in doc.tables:
        apply_three_line_table(table)


def audit_document(doc: Document, output_dir: Path, changes: list[str]) -> None:
    forbidden_terms = [
        "未跑",
        "已跑",
        "没跑",
        "审稿人要求",
        "按照第一个文件",
        "当前草稿",
        "AI",
        "Codex",
        "完美",
        "绝对",
    ]
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    flags = []
    for term in forbidden_terms:
        count = sum(text.count(term) for text in texts)
        if count:
            flags.append((term, count))

    table_dims = [(i + 1, len(t.rows), len(t.columns) if t.rows else 0) for i, t in enumerate(doc.tables)]
    long_tables = [item for item in table_dims if item[2] > 7]

    output_dir.mkdir(parents=True, exist_ok=True)
    with (output_dir / "draft15_revision_audit.txt").open("w", encoding="utf-8") as f:
        f.write("Draft 15 revision audit\n")
        f.write("=======================\n")
        f.write("Changes:\n")
        for change in changes:
            f.write(f"- {change}\n")
        f.write("\nTable dimensions:\n")
        for idx, rows, cols in table_dims:
            f.write(f"- T{idx:02d}: {rows} x {cols}\n")
        f.write("\nLong tables (>7 columns):\n")
        if long_tables:
            for idx, rows, cols in long_tables:
                f.write(f"- T{idx:02d}: {rows} x {cols}\n")
        else:
            f.write("- none\n")
        f.write("\nFlagged terms:\n")
        if flags:
            for term, count in flags:
                f.write(f"- {term}: {count}\n")
        else:
            f.write("- none\n")


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    input_path = find_file(Path(r"C:\Users\Administrator\Desktop"), ["初稿", "14"])
    output_path = Path(r"C:\Users\Administrator\Desktop\修改\初稿-15.docx")
    doc = Document(input_path)
    changes = polish_document(doc)
    apply_document_style(doc)
    doc.save(output_path)
    audit_document(doc, Path("work/draft15_audit"), changes)
    print(f"Input: {input_path}")
    print(f"Output: {output_path}")
    print(f"Changes: {len(changes)}")
    for change in changes:
        print(f"- {change}")


if __name__ == "__main__":
    main()
