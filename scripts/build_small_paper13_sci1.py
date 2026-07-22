from __future__ import annotations

import json
import re
import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
SRC = OUT / "\u5c0f\u8bba\u6587-12_\u7ec8\u5ba1.docx"
DOCX = OUT / "\u5c0f\u8bba\u6587-13.docx"
PACK = OUT / "\u5c0f\u8bba\u6587-13_SCI1\u8865\u5f3a\u8bc1\u636e"
AUDIT = ROOT / "results" / "audits" / "small_paper_13_sci1_audit.json"
TEXT_AUDIT = OUT / "paper13_sci1_fulltext_audit.txt"


def set_run_font(run, size: float | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
    if size is not None:
        run.font.size = Pt(size)


def set_para_text(p, text: str, size: float | None = None) -> None:
    for r in p.runs:
        r.text = ""
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    set_run_font(run, size=size)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)


def insert_before(anchor, text: str, style: str | None = None, first_line: bool = True):
    p = anchor.insert_paragraph_before("")
    if style:
        p.style = style
    set_para_text(p, text)
    if first_line and style not in {"Heading 1", "Heading 2"}:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def paragraph_texts(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def replace_literals(doc: Document, replacements: dict[str, str]) -> None:
    for p in doc.paragraphs:
        text = p.text
        updated = text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != text:
            set_para_text(p, updated)


def find_para(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"paragraph not found: {text}")


def find_first_startswith(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError(f"paragraph prefix not found: {prefix}")


def add_gap_table_before(doc: Document, anchor) -> None:
    rows = [
        [
            "\u6807\u51c6\u57fa\u51c6\u4e0e\u4efb\u52a1\u53ef\u6bd4\u6027",
            "[1,2]",
            "\u4e5d\u7ec8\u70b9\u51bb\u7ed3 3\u00d73\u00d75 \u8bc4\u4f30\uff1bTDC \u95e8\u63a7\u4e0e source data \u8f93\u51fa",
            "\u4e0d\u58f0\u79f0 MoleculeNet/TDC \u5168\u699c\u5355\u4f18\u52bf\uff1b\u65f6\u95f4\u5916\u524d\u77bb\u9a8c\u8bc1\u5c1a\u672a\u5b8c\u6210",
        ],
        [
            "\u73b0\u4ee3\u5f3a\u57fa\u7ebf\u4e0e\u57fa\u7840\u6a21\u578b",
            "[21\u201323,34]",
            "RDKit RF\u3001GNN-GCN\u3001Chemprop/D-MPNN\u3001ChemBERTa\u548c MoLFormer \u5728 ESOL/BACE/ClinTox \u540c\u5212\u5206\u9762\u677f\u5b8c\u6210",
            "\u4e5d\u7ec8\u70b9\u5168\u91cf\u6df1\u5ea6\u9762\u677f\u4e0e TabPFN \u4ecd\u4e0d\u5b8c\u6574",
        ],
        [
            "\u6a21\u578b\u9009\u62e9\u504f\u501a",
            "[3,4]",
            "K \u6269\u5f20\u3001\u968f\u673a\u6392\u5e8f\u8d1f\u5bf9\u7167\u3001\u4fe1\u53f7\u6062\u590d\u6b63\u5bf9\u7167\u548c\u51bb\u7ed3\u51b3\u7b56\u5361\u5f62\u6210\u4e3b\u8bc1\u636e",
            "\u7b2c\u4e09\u65b9\u51b7\u542f\u52a8\u590d\u8dd1\u4e0e DOI \u53d1\u5e03\u5c1a\u672a\u5b8c\u6210",
        ],
        [
            "\u9884\u6d4b\u53ef\u9760\u6027",
            "[26,27,33]",
            "\u9010\u6837\u672c\u98ce\u9669\u66f2\u7ebf\u3001E-AURC\u3001\u5206\u88c2\u4fdd\u5f62\u8986\u76d6\u548c\u6807\u7b7e\u6761\u4ef6\u8986\u76d6",
            "\u4e0d\u662f\u5168\u90e8\u4e0d\u786e\u5b9a\u6027\u4f30\u8ba1\u5668\u7684\u5bf9\u7167\u57fa\u51c6",
        ],
        [
            "\u5316\u5b66\u8fb9\u754c\u4e0e\u8fc1\u79fb",
            "[29]",
            "MoleculeACE \u6210\u5bf9\u5dee\u503c\u3001bRo5 \u5207\u5206\u5bf9\u7167\u548c\u6700\u8fd1\u90bb Tanimoto \u5206\u5c42",
            "\u4ec5\u4f5c\u6b21\u8981\u8fb9\u754c\u8bc1\u636e\uff1b\u4e0d\u7b49\u540c\u4e8e\u5b8c\u6574\u90e8\u7f72\u9a8c\u8bc1",
        ],
        [
            "\u53ef\u590d\u73b0\u4e0e\u8d1f\u7ed3\u679c",
            "[32]",
            "\u751f\u6210 source data\u3001XML \u5ba1\u8ba1\u3001runtime status\u3001\u53bb\u91cd\u654f\u611f\u6027\u548c error-overlap \u660e\u7ec6",
            "\u6b63\u5f0f\u4ee3\u7801\u53d1\u5e03\u3001Zenodo DOI \u548c\u72ec\u7acb\u590d\u8dd1\u4ecd\u9700\u5b8c\u6210",
        ],
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header = ["\u53ef\u6bd4\u6027\u95ee\u9898", "\u5bf9\u7167\u6587\u732e", "\u672c\u6587\u8bc1\u636e", "\u4ecd\u4fdd\u7559\u7684\u8fb9\u754c"]
    for cell, text in zip(table.rows[0].cells, header):
        cell.text = text
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=8)
    anchor._p.addprevious(table._tbl)


def update_references(doc: Document) -> None:
    refs = paragraph_texts(doc)
    if any("A systematic study of key elements underlying molecular property prediction" in t for t in refs):
        return
    p = doc.add_paragraph()
    set_para_text(
        p,
        "[34] Deng J, Yang Z, Wang H, Ojima I, Samaras D, Wang F. A systematic study of key elements underlying molecular property prediction. Nat Commun. 2023;14:6395. doi:10.1038/s41467-023-41948-6.",
    )
    p.paragraph_format.first_line_indent = Cm(0)


def insert_sci1_content(doc: Document) -> None:
    intro_anchor = find_first_startswith(doc, "\u672c\u6587\u7684\u4e2d\u5fc3\u4e3b\u5f20\u662f")
    insert_before(
        intro_anchor,
        "\u8fd1\u5e74\u9ad8\u5f71\u54cd\u5206\u5b50\u673a\u5668\u5b66\u4e60\u7814\u7a76\u5df2\u5f62\u6210\u4e09\u7c7b\u53ef\u6bd4\u6027\u8981\u6c42\u3002MoleculeNet \u548c TDC \u5f3a\u8c03\u6570\u636e\u96c6\u3001\u5212\u5206\u3001\u6307\u6807\u548c\u516c\u5f00\u4efb\u52a1\u5b9a\u4e49\u7684\u7edf\u4e00\uff1bChemprop/D-MPNN\u3001ChemBERTa \u548c MoLFormer \u7b49\u5de5\u4f5c\u4f7f\u5f3a\u57fa\u7ebf\u4ece\u6811\u6a21\u578b\u6269\u5c55\u5230\u56fe\u6a21\u578b\u548c\u5316\u5b66\u8bed\u8a00\u6a21\u578b\uff1b\u5927\u89c4\u6a21\u7ecf\u9a8c\u7814\u7a76\u5219\u8868\u660e\uff0c\u56fa\u5b9a\u8868\u5f81\u548c\u7ecf\u8c03\u6574\u7684\u4f20\u7edf\u6a21\u578b\u5728\u8bb8\u591a\u5206\u5b50\u4efb\u52a1\u4e0a\u4ecd\u5177\u6709\u5f3a\u7ade\u4e89\u529b[1,2,21\u201323,34]\u3002\u56e0\u6b64\uff0c\u82e5\u5c06\u672c\u7814\u7a76\u8868\u8ff0\u4e3a\u65b0\u7684\u901a\u7528 SOTA \u9884\u6d4b\u5668\uff0c\u5176\u7ec8\u70b9\u6570\u3001\u65f6\u95f4\u5916\u9a8c\u8bc1\u548c\u5168\u91cf\u57fa\u7840\u6a21\u578b\u9762\u677f\u5747\u4e0d\u8db3\uff1b\u66f4\u51c6\u786e\u7684\u5b9a\u4f4d\u662f\u5efa\u7acb\u4e00\u4e2a\u53ef\u5ba1\u8ba1\u7684\u5019\u9009\u6c60\u9009\u62e9\u6cbb\u7406\u534f\u8bae\u3002",
    )
    insert_before(
        intro_anchor,
        "\u56f4\u7ed5\u8fd9\u4e00\u5b9a\u4f4d\uff0c\u672c\u6587\u5c06\u6587\u732e\u5bf9\u7167\u4e2d\u7684\u5dee\u5f02\u8f6c\u5316\u4e3a\u53ef\u68c0\u67e5\u7684\u6269\u5c55\u5206\u6790\uff1a\u5728\u4e09\u4e2a\u4ee3\u8868\u7ec8\u70b9\u4e0a\u7edf\u4e00\u91cd\u8bad RDKit RF\u3001GNN-GCN\u3001Chemprop/D-MPNN\u3001ChemBERTa \u548c MoLFormer\uff1b\u8f93\u51fa\u9010\u6837\u672c\u9884\u6d4b\u5e76\u8ba1\u7b97 error-overlap\uff1b\u5b8c\u6210\u4e09\u5957\u53bb\u91cd\u654f\u611f\u6027\u5206\u6790\uff1b\u540c\u65f6\u628a\u4fdd\u5f62\u53ef\u9760\u6027\u3001MoleculeACE\u3001bRo5 \u548c TDC \u4f5c\u4e3a\u6b21\u8981\u8fb9\u754c\u8bc1\u636e\u800c\u975e\u4e3b\u8981\u80dc\u8d1f\u8bc1\u636e\u3002\u8fd9\u79cd\u5199\u6cd5\u4f7f\u7a3f\u4ef6\u4e0e\u9ad8\u5f71\u54cd\u57fa\u51c6\u5de5\u4f5c\u7684\u6807\u51c6\u5bf9\u9f50\uff0c\u4e5f\u907f\u514d\u628a\u5c1a\u672a\u5b8c\u6210\u7684 TabPFN \u548c\u4e5d\u7ec8\u70b9\u5168\u91cf\u6df1\u5ea6\u9762\u677f\u5199\u6210\u5df2\u5b8c\u6210\u7ed3\u8bba\u3002",
    )

    discussion_anchor = find_para(doc, "4 \u8ba8\u8bba")
    insert_before(discussion_anchor, "3.10 \u4e0e\u5f53\u524d\u57fa\u51c6\u5b9e\u8df5\u7684\u5bf9\u7167\u548c\u8bc1\u636e\u8fb9\u754c", style="Heading 2", first_line=False)
    insert_before(
        discussion_anchor,
        "\u4e3a\u4f7f\u6587\u732e\u5bf9\u7167\u53ef\u590d\u6838\uff0c\u672c\u6587\u6574\u7406\u4e86\u516d\u7c7b\u8bc1\u636e\u8fb9\u754c\uff0c\u8986\u76d6\u6807\u51c6\u57fa\u51c6\u3001\u73b0\u4ee3\u5f3a\u57fa\u7ebf\u3001\u6a21\u578b\u9009\u62e9\u504f\u501a\u3001\u9884\u6d4b\u53ef\u9760\u6027\u3001\u5316\u5b66\u8fb9\u754c\u548c\u53ef\u590d\u73b0\u6027\u3002\u5f3a\u57fa\u7ebf\u8bc1\u636e\u4e0d\u518d\u505c\u7559\u4e8e\u8fd0\u884c\u72b6\u6001\u8bf4\u660e\uff1a\u4ee3\u8868\u6027\u9762\u677f\u5305\u542b 225 \u4e2a\u5916\u5c42\u5355\u5143\u3001675 \u4e2a\u5185\u5c42\u5355\u5143\u548c 103,025 \u884c\u9010\u6837\u672c\u9884\u6d4b\u3002\u5728\u8be5\u9762\u677f\u4e2d\uff0cRDKit RF \u88ab\u51bb\u7ed3\u9009\u62e9\u5668\u5728 45/45 \u4e2a\u9009\u62e9\u5355\u5143\u4e2d\u9009\u4e2d\uff0c\u5916\u5c42 Top-1 \u547d\u4e2d\u7387\u4e3a 0.956\uff0c\u5e73\u5747\u8303\u56f4\u5f52\u4e00\u5316\u9009\u62e9\u635f\u5931\u4e3a 0.0017\u3002",
    )
    insert_before(
        discussion_anchor,
        "error-overlap \u5ba1\u8ba1\u663e\u793a\uff0c\u4e94\u4e2a\u5019\u9009\u7684\u9519\u8bef\u96c6\u5e76\u975e\u5b8c\u5168\u91cd\u5408\uff1a10 \u4e2a\u5019\u9009\u5bf9\u7684\u5e73\u5747 Jaccard \u91cd\u5408\u4e3a 0.189\uff0c\u8303\u56f4\u4e3a 0.057\u20130.449\u3002\u53bb\u91cd\u654f\u611f\u6027\u5206\u6790\u4e5f\u5df2\u4ece\u6e05\u6d17\u5ba1\u8ba1\u63a8\u8fdb\u5230\u5b9e\u9645\u91cd\u8dd1\uff1a\u4e09\u4e2a\u4ee3\u8868\u7ec8\u70b9\u3001\u4e09\u5957\u7b56\u7565\u548c 135 \u4e2a\u5916\u5c42\u5355\u5143\u7684\u6700\u5927\u5e73\u5747\u6548\u7528\u53d8\u5316\u4e3a 0.022\u3002\u8fd9\u4e9b\u7ed3\u679c\u652f\u6301\u4e00\u4e2a\u4fdd\u5b88\u7ed3\u8bba\uff1a\u5f53\u4ee3\u590d\u6742\u6a21\u578b\u5fc5\u987b\u5728\u540c\u4e00\u51bb\u7ed3\u5212\u5206\u4e0b\u627f\u62c5\u5ba1\u8ba1\u6210\u672c\uff0c\u800c\u4e0d\u80fd\u4ec5\u4f9d\u8d56\u5386\u53f2\u6700\u4f18\u5206\u6570\u8fdb\u5165\u4e3b\u6587\u53d9\u4e8b\u3002",
    )
    insert_before(
        discussion_anchor,
        "\u8868 9 | \u5f53\u524d\u9ad8\u5f71\u54cd\u5206\u5b50\u673a\u5668\u5b66\u4e60\u7814\u7a76\u8981\u6c42\u4e0e\u672c\u6587\u8bc1\u636e\u5bf9\u7167",
        first_line=False,
    )
    add_gap_table_before(doc, discussion_anchor)

    conclusion_anchor = find_para(doc, "5 \u7ed3\u8bba")
    insert_before(conclusion_anchor, "4.5 \u4e0e\u8fd1\u671f\u9ad8\u5f71\u54cd\u7814\u7a76\u7684\u5b9a\u4f4d\u5dee\u5f02", style="Heading 2", first_line=False)
    insert_before(
        conclusion_anchor,
        "\u4e0e\u4ee5 MoleculeNet \u6216 TDC \u4e3a\u6838\u5fc3\u7684\u57fa\u51c6\u8bba\u6587\u76f8\u6bd4\uff0c\u672c\u6587\u7684\u4e0d\u8db3\u662f\u5e76\u672a\u63d0\u4f9b\u66f4\u5927\u89c4\u6a21\u7684\u516c\u5f00\u4efb\u52a1\u6536\u96c6\u6216\u6392\u884c\u699c\uff1b\u4e0e Chemprop\u3001ChemBERTa \u6216 MoLFormer \u7b49\u65b9\u6cd5\u8bba\u6587\u76f8\u6bd4\uff0c\u672c\u6587\u4e5f\u6ca1\u6709\u63d0\u51fa\u66f4\u5f3a\u7684\u8868\u5f81\u5b66\u4e60\u4e3b\u5e72\u3002\u672c\u6587\u7684\u4f18\u52bf\u5728\u53e6\u4e00\u4e2a\u5c42\u9762\uff1a\u5b83\u628a\u5019\u9009\u6c60\u6269\u5f20\u3001\u9a8c\u8bc1\u6392\u5e8f\u5931\u771f\u3001\u5916\u5c42\u9009\u62e9\u635f\u5931\u548c\u8d1f\u7ed3\u679c\u4fdd\u7559\u5199\u6210\u53ef\u590d\u6838\u534f\u8bae\u3002\u8fd9\u4f7f\u672c\u6587\u7684\u5b9a\u4f4d\u66f4\u63a5\u8fd1\u6a21\u578b\u8bc4\u4f30\u548c\u9a8c\u8bc1\u6cbb\u7406\uff0c\u800c\u4e0d\u662f\u53c8\u4e00\u4e2a\u65b0\u7684\u9884\u6d4b\u5668\u6392\u540d\u3002",
    )
    insert_before(
        conclusion_anchor,
        "\u8fd9\u4e9b\u7ed3\u679c\u4e5f\u754c\u5b9a\u4e86\u540e\u7eed\u6269\u5c55\u7684\u4f18\u5148\u7ea7\uff1a\u7b2c\u4e00\uff0c\u5728\u8d44\u6e90\u5141\u8bb8\u65f6\u5c06\u5f3a\u57fa\u7ebf\u4ece\u4e09\u4e2a\u4ee3\u8868\u7ec8\u70b9\u6269\u5c55\u5230\u4e5d\u7ec8\u70b9\u5168\u91cf\u9762\u677f\uff1b\u7b2c\u4e8c\uff0c\u89e3\u51b3 TabPFN \u6388\u6743\u6216\u8fd0\u884c\u73af\u5883\u9650\u5236\u540e\u518d\u7eb3\u5165\u540c\u5212\u5206\u6bd4\u8f83\uff1b\u7b2c\u4e09\uff0c\u5b8c\u6210\u516c\u5f00 release\u3001Zenodo DOI \u548c\u72ec\u7acb\u51b7\u542f\u52a8\u590d\u8dd1\u3002\u8fd9\u4e9b\u8fb9\u754c\u5df2\u5728\u6b63\u6587\u4e2d\u660e\u786e\u9650\u5b9a\uff0c\u4ee5\u907f\u514d\u628a\u4ee3\u8868\u6027\u6269\u5c55\u5206\u6790\u5199\u6210\u5df2\u5b8c\u6210\u7684\u5168\u91cf\u786e\u8bc1\u3002",
    )


def xml_errors(path: Path) -> list[str]:
    errors: list[str] = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if name.endswith(".xml"):
                try:
                    ElementTree.fromstring(zf.read(name))
                except Exception as exc:
                    errors.append(f"{name}: {exc}")
    return errors


def extract_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def audit_doc(path: Path) -> dict[str, object]:
    doc = Document(path)
    text = extract_text(doc)
    TEXT_AUDIT.write_text(text, encoding="utf-8")
    abstract = []
    capture = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "\u6458\u8981":
            capture = True
            continue
        if capture and t.startswith("\u5173\u952e\u8bcd"):
            break
        if capture and t:
            abstract.append(t)

    table_caption_numbers = []
    for p in doc.paragraphs:
        m = re.match(r"\u8868\s*(\d+)", p.text.strip())
        if m:
            table_caption_numbers.append(int(m.group(1)))
    figure_caption_numbers = []
    for p in doc.paragraphs:
        m = re.match(r"\u56fe\s*(\d+)", p.text.strip())
        if m:
            figure_caption_numbers.append(int(m.group(1)))

    colloquial = [
        "\u505a\u539a",
        "\u8865\u539a",
        "\u8865\u5f3a",
        "\u6295\u7a3f\u524d",
        "\u4e0d\u8db3\u70b9",
        "SCI 1",
        "\u5f88\u725b",
        "\u8d8a\u591a\u8d8a\u597d",
        "\u8d8a\u591a\u8d8a\u5dee",
        "\u8001\u5b9e\u8bf4",
        "\u7a3f\u4ef6\u4e2d",
        "\u770b\u8d77\u6765",
    ]
    abstract_forbidden = ["MoleculeNet", "TDC", "MoleculeACE", "bRo5", "\u4fdd\u5f62"]
    required = [
        "K=32 \u76f8\u5bf9 K=4",
        "\u5b9e\u9645\u5151\u73b0\u6548\u7528\u589e\u76ca",
        "\u8de8\u7ec8\u70b9\u5143\u98ce\u9669",
    ]
    audit = {
        "docx": str(path),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "abstract_paragraphs": len(abstract),
        "abstract_required_hits": {item: any(item in p for p in abstract) for item in required},
        "abstract_forbidden_hits": {item: any(item in p for p in abstract) for item in abstract_forbidden},
        "has_sci1_gap_section": "\u5bf9\u7167\u548c\u8bc1\u636e\u8fb9\u754c" in text,
        "has_positioning_section": "\u5b9a\u4f4d\u5dee\u5f02" in text,
        "has_table9_caption": 9 in table_caption_numbers,
        "has_new_reference_34": "[34] Deng J" in text,
        "has_tabpfn_boundary": "TabPFN \u4ecd\u4e0d\u5b8c\u6574" in text
        or "TabPFN \u6388\u6743" in text,
        "figure_caption_numbers": figure_caption_numbers,
        "table_caption_numbers": table_caption_numbers,
        "colloquial_hits": {item: text.count(item) for item in colloquial if item in text},
        "xml_errors": xml_errors(path),
    }
    audit["passed"] = (
        audit["abstract_paragraphs"] == 4
        and all(audit["abstract_required_hits"].values())
        and not any(audit["abstract_forbidden_hits"].values())
        and audit["has_sci1_gap_section"]
        and audit["has_positioning_section"]
        and audit["has_table9_caption"]
        and audit["has_new_reference_34"]
        and audit["has_tabpfn_boundary"]
        and not audit["colloquial_hits"]
        and not audit["xml_errors"]
    )
    return audit


def main() -> None:
    if not PACK.exists():
        raise SystemExit(f"evidence pack missing: {PACK}")
    shutil.copy2(SRC, DOCX)
    doc = Document(DOCX)
    insert_sci1_content(doc)
    replace_literals(
        doc,
        {
            "\u4ecd\u9700\u5728\u6b63\u5f0f\u6295\u7a3f\u524d\u5b8c\u6210": "\u4ecd\u9700\u4f5c\u4e3a\u540e\u7eed\u590d\u73b0\u5de5\u4f5c\u5b8c\u6210",
        },
    )
    update_references(doc)
    doc.save(DOCX)
    audit = audit_doc(DOCX)
    AUDIT.parent.mkdir(parents=True, exist_ok=True)
    AUDIT.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
