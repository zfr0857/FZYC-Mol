from __future__ import annotations

import json
import math
import shutil
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib.lines import Line2D
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from matplotlib.ticker import PercentFormatter
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("D:/fzyc")
OUT = ROOT / "output"
SOURCE_DOC = Path("C:/Users/Administrator/Desktop/小论文-18.docx")
PAPER19 = OUT / "小论文-19.docx"
DESKTOP_PAPER19 = Path("C:/Users/Administrator/Desktop/小论文-19.docx")
AUDIT_PATH = OUT / "小论文-19_修改与实验审计.json"
REPORT_PATH = OUT / "小论文-19_修改与实验审计报告.md"

EXP = OUT / "paper19_rejection_driven_experiments_20260712"
FIG_ROOT = OUT / "小论文-19_图表与源数据"
FIG_DIR = FIG_ROOT / "figures"
SRC_DIR = FIG_ROOT / "source_data"
MULTI = ROOT / "results" / "reviewer_core_20260624" / "multiview_nested"
PAPER18_FIG = OUT / "小论文-18_主图_PNG_SVG_pack"

INK = "#172235"
MUTED = "#5F6B7A"
BLUE = "#3E6FAE"
BLUE2 = "#85A7D3"
GREEN = "#4E9B75"
ORANGE = "#DD8C38"
PURPLE = "#8C77B1"
GREY = "#8C949F"
LIGHT = "#CBD3DD"
GRID = "#E3E8EF"
RED = "#B85B5B"


def ensure_dirs() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    SRC_DIR.mkdir(parents=True, exist_ok=True)


def save_figure(fig: plt.Figure, stem: str) -> None:
    fig.savefig(FIG_DIR / f"{stem}.png", dpi=360, bbox_inches="tight", facecolor="white")
    fig.savefig(FIG_DIR / f"{stem}.svg", bbox_inches="tight", facecolor="white")
    plt.close(fig)


def style_axis(ax: plt.Axes, grid_axis: str = "both") -> None:
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(INK)
    ax.spines["bottom"].set_color(INK)
    ax.tick_params(colors=INK, labelsize=9)
    ax.grid(True, axis=grid_axis, color=GRID, lw=0.7)
    ax.set_axisbelow(True)


def make_workflow_figure() -> None:
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 9,
            "axes.titlesize": 11,
            "axes.labelsize": 9,
            "svg.fonttype": "none",
        }
    )
    fig, ax = plt.subplots(figsize=(10.2, 4.65))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    steps = [
        (0.04, "Analysis lock", "Nine endpoints\nMetrics and seeds\nCandidate order", "#EEF3FA", BLUE),
        (0.28, "Development", "Training-only transforms\n3-fold inner ranking\nFailures retained", "#EFF7F2", GREEN),
        (0.52, "Frozen selection", "Validation-best / one-SE\nRisk-adjusted controls\nNo outer-label updates", "#FFF4E8", ORANGE),
        (0.76, "Outer audit", "3 outer folds × 5 seeds\nRaw-unit selection loss\nRanking and diversity", "#F3F0F8", PURPLE),
    ]
    width, height, y = 0.20, 0.37, 0.40
    for x, title, body, fill, edge in steps:
        box = FancyBboxPatch(
            (x, y),
            width,
            height,
            boxstyle="round,pad=0.012,rounding_size=0.018",
            facecolor=fill,
            edgecolor=edge,
            linewidth=1.6,
        )
        ax.add_patch(box)
        ax.text(x + 0.018, y + height - 0.055, title, va="top", color=INK, fontsize=10, fontweight="bold")
        ax.text(x + 0.018, y + height - 0.17, body, va="top", color=MUTED, fontsize=8.3, linespacing=1.35)
    for left, right in zip(steps[:-1], steps[1:]):
        x1 = left[0] + width
        x2 = right[0]
        ax.add_patch(
            FancyArrowPatch(
                (x1 + 0.006, y + height / 2),
                (x2 - 0.006, y + height / 2),
                arrowstyle="-|>",
                mutation_scale=14,
                color=INK,
                linewidth=1.2,
            )
        )

    evidence_box = FancyBboxPatch(
        (0.04, 0.10),
        0.48,
        0.20,
        boxstyle="round,pad=0.012,rounding_size=0.015",
        facecolor="#EEF6F7",
        edgecolor="#388C91",
        linewidth=1.4,
    )
    ax.add_patch(evidence_box)
    ax.text(0.06, 0.245, "Evidence outputs", color=INK, fontsize=9.8, fontweight="bold", va="top")
    ax.text(
        0.06,
        0.190,
        "Fold-level source data  |  candidate failures\nCompute budgets  |  boundary analyses",
        color=MUTED,
        fontsize=8.0,
        linespacing=1.35,
        va="top",
    )
    ax.add_patch(
        FancyArrowPatch(
            (0.86, y),
            (0.50, 0.30),
            arrowstyle="-|>",
            mutation_scale=12,
            color="#388C91",
            linewidth=1.0,
        )
    )

    lockbox = FancyBboxPatch(
        (0.56, 0.10),
        0.40,
        0.20,
        boxstyle="round,pad=0.012,rounding_size=0.015",
        facecolor="white",
        edgecolor=GREY,
        linewidth=1.2,
        linestyle="--",
    )
    ax.add_patch(lockbox)
    ax.text(0.58, 0.245, "Independent confirmation", color=INK, fontsize=9.8, fontweight="bold", va="top")
    ax.text(
        0.58,
        0.190,
        "Not available in this revision.\nOuter-audit results are not an independent test.",
        color=MUTED,
        fontsize=8.1,
        va="top",
    )
    ax.add_patch(
        FancyArrowPatch(
            (0.86, y),
            (0.86, 0.30),
            arrowstyle="-|>",
            mutation_scale=12,
            color=GREY,
            linestyle="--",
            linewidth=1.0,
        )
    )
    ax.text(0.04, 0.955, "Frozen candidate-selection audit", color=INK, fontsize=17, fontweight="bold", va="top")
    ax.text(
        0.04,
        0.875,
        "Development information flows left to right; outer-audit labels never flow back into confirmatory selection.",
        color=MUTED,
        fontsize=10,
        va="top",
    )
    save_figure(fig, "fig01_paper19_frozen_audit_workflow")


def make_multiview_figure() -> None:
    policy = pd.read_csv(MULTI / "policy_summary.csv")
    paired = pd.read_csv(MULTI / "paired_multiview_effects.csv")
    counts = pd.read_csv(MULTI / "endpoint_representation_counts.csv")
    overall = pd.read_csv(MULTI / "validation_best_representation_counts.csv")
    for path in [
        MULTI / "policy_summary.csv",
        MULTI / "paired_multiview_effects.csv",
        MULTI / "endpoint_representation_counts.csv",
        MULTI / "validation_best_representation_counts.csv",
        MULTI / "candidate_registry.csv",
    ]:
        shutil.copy2(path, SRC_DIR / f"fig03_{path.name}")

    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 9,
            "axes.titlesize": 11,
            "axes.labelsize": 9,
            "svg.fonttype": "none",
        }
    )
    fig, axes = plt.subplots(2, 2, figsize=(9.0, 6.7))

    ax = axes[0, 0]
    order = ["fixed_morgan_rf", "one_se_stable", "risk_adjusted", "validation_best"]
    p = policy[(policy.variant == "full_multiview") & policy.policy.isin(order)].copy()
    p["policy"] = pd.Categorical(p.policy, order, ordered=True)
    p = p.sort_values("policy")
    x = np.arange(len(p))
    yv = p.mean_normalized_regret.to_numpy()
    yerr = np.vstack(
        [yv - p.endpoint_cluster_ci95_low.to_numpy(), p.endpoint_cluster_ci95_high.to_numpy() - yv]
    )
    ax.bar(x, yv, color=[GREY, PURPLE, ORANGE, BLUE], width=0.68)
    ax.errorbar(x, yv, yerr=yerr, fmt="none", color=INK, capsize=3, lw=1)
    ax.set_xticks(x, ["Fixed\nMorgan RF", "one-SE", "Risk-\nadjusted", "Validation-\nbest"])
    ax.set_ylabel("Range-normalized audit loss")
    ax.set_title("Frozen selectors in the 12-candidate pool")
    style_axis(ax, "y")

    ax = axes[0, 1]
    wanted = [
        "attainable multiview gain vs Morgan-only oracle",
        "realized multiview validation-best gain vs Morgan-only",
        "concatenated multiview gain vs separate-view pool",
        "validation-best gain vs one-SE in full pool",
    ]
    labels = [
        "Observed audit upper bound\nvs Morgan pool",
        "Frozen selection\nvs Morgan selection",
        "Concatenated vs\nseparate views",
        "Validation-best\nvs one-SE",
    ]
    q = paired.set_index("comparison").loc[wanted].reset_index()
    yy = np.arange(len(q))
    val = q.mean_normalized_utility_gain.to_numpy()
    xerr = np.vstack(
        [val - q.endpoint_cluster_ci95_low.to_numpy(), q.endpoint_cluster_ci95_high.to_numpy() - val]
    )
    ax.errorbar(val, yy, xerr=xerr, fmt="o", color=BLUE, ecolor=LIGHT, capsize=3, ms=5)
    ax.axvline(0, color=INK, lw=0.9)
    ax.set_yticks(yy, labels)
    ax.set_xlabel("Paired normalized utility gain")
    ax.set_title("Observed and realized multiview gains")
    style_axis(ax, "x")

    ax = axes[1, 0]
    o = (
        overall[overall.variant == "full_multiview"]
        .set_index("selected_representation")
        .reindex(["morgan512", "maccs", "rdkit2d", "multiview"])
        .fillna(0)
    )
    ax.bar(np.arange(4), o["size"], color=[BLUE2, GREEN, ORANGE, PURPLE], width=0.68)
    ax.set_xticks(np.arange(4), ["Morgan", "MACCS", "RDKit2D", "Concatenated\nmultiview"])
    ax.set_ylabel("Frozen validation-best selections (n=135)")
    ax.set_title("Representation selected before outer audit")
    for i, value in enumerate(o["size"]):
        ax.text(i, value + 1.7, f"{int(value)}", ha="center", fontsize=8.5, color=INK)
    style_axis(ax, "y")

    ax = axes[1, 1]
    heat = counts.pivot(index="task", columns="selected_representation", values="size").fillna(0)
    heat = heat.reindex(columns=["morgan512", "maccs", "rdkit2d", "multiview"], fill_value=0)
    im = ax.imshow(heat.to_numpy(), aspect="auto", cmap="Blues", vmin=0)
    ax.set_xticks(np.arange(4), ["Morgan", "MACCS", "RDKit2D", "Concat"], rotation=20, ha="right")
    ax.set_yticks(np.arange(len(heat)), [x.replace("tdc_", "").replace("_", " ") for x in heat.index])
    ax.tick_params(axis="both", labelsize=7.8)
    for i in range(heat.shape[0]):
        for j in range(heat.shape[1]):
            ax.text(j, i, f"{int(heat.iloc[i, j])}", ha="center", va="center", fontsize=7.3, color=INK)
    ax.set_title("Endpoint-by-representation selection map")
    cbar = fig.colorbar(im, ax=ax, fraction=0.045, pad=0.03)
    cbar.ax.tick_params(labelsize=7.2)
    fig.tight_layout(w_pad=2.2, h_pad=2.1)
    save_figure(fig, "fig03_paper19_multiview_confirmation")


def parse_mean_sd(value: object) -> tuple[float, float]:
    import re

    vals = re.findall(r"[-+]?\d+(?:\.\d+)?", str(value))
    if not vals:
        return math.nan, math.nan
    return float(vals[0]), float(vals[1]) if len(vals) > 1 else 0.0


def make_boundary_figure() -> None:
    gap_path = ROOT / "reports" / "remaining_missing_experiments_20260606" / "moleculeace_gap_correlation_summary.csv"
    cyc_path = ROOT / "reports" / "full_missing_experiment_run_20260611" / "bro5_cycpept_pampa_compact_summary.csv"
    lin_path = ROOT / "reports" / "full_missing_experiment_run_20260611" / "linpept_compact_summary_20260611.csv"
    fail_path = ROOT / "reports" / "supplement_experiment_revision_20260606" / "maintext_table_failure_cases_compact.csv"
    gap = pd.read_csv(gap_path)
    cyc = pd.read_csv(cyc_path)
    lin = pd.read_csv(lin_path)
    for path in [gap_path, cyc_path, lin_path, fail_path]:
        shutil.copy2(path, SRC_DIR / f"fig05_{path.name}")
    task = (
        gap.groupby("task", as_index=False)
        .agg(
            gap_spearman=("gap_spearman", "mean"),
            gap_sd=("gap_spearman", "std"),
            direction_accuracy=("direction_accuracy", "mean"),
            n_pairs=("n_pairs", "mean"),
        )
        .sort_values("gap_spearman")
    )
    task.to_csv(SRC_DIR / "fig05_moleculeace_task_summary.csv", index=False)

    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 9,
            "axes.titlesize": 11,
            "axes.labelsize": 9,
            "svg.fonttype": "none",
        }
    )
    fig, axes = plt.subplots(2, 2, figsize=(8.9, 6.35))

    ax = axes[0, 0]
    yy = np.arange(len(task))
    ax.errorbar(task.gap_spearman, yy, xerr=task.gap_sd, fmt="o", ms=3.5, color=BLUE, ecolor=LIGHT, capsize=2)
    ax.set_yticks(yy, task.task)
    ax.tick_params(axis="y", labelsize=7.4)
    ax.axvline(0, color=INK, lw=0.8)
    ax.set_xlabel("Activity-gap Spearman")
    ax.set_title("MoleculeACE task heterogeneity")
    style_axis(ax, "x")

    ax = axes[0, 1]
    means, sds = zip(*(parse_mean_sd(v) for v in cyc.test_RMSE))
    xx = np.arange(len(cyc))
    ax.bar(xx, means, yerr=sds, color=[GREEN, BLUE, ORANGE, PURPLE], width=0.68, capsize=3)
    ax.set_xticks(xx, cyc.split.str.title(), rotation=20, ha="right")
    ax.set_ylabel("CycPept-PAMPA RMSE")
    ax.set_title("bRo5 split pressure")
    style_axis(ax, "y")

    rows: list[dict[str, object]] = []
    for _, row in lin.iterrows():
        roc, roc_sd = parse_mean_sd(row.test_ROC_AUC)
        pr, pr_sd = parse_mean_sd(row.test_PR_AUC)
        rows.append(
            {
                "dataset": row.dataset,
                "split": row.split,
                "roc": roc,
                "roc_sd": roc_sd,
                "pr": pr,
                "pr_sd": pr_sd,
            }
        )
    lp = pd.DataFrame(rows)
    lp.to_csv(SRC_DIR / "fig05_linpept_parsed.csv", index=False)
    ax = axes[1, 0]
    markers = {"linpept_cellpen": "o", "linpept_nonfouling": "s"}
    colors = {"random": GREEN, "scaffold": BLUE, "perimeter": ORANGE}
    for dataset, group in lp.groupby("dataset"):
        for _, row in group.iterrows():
            ax.errorbar(
                row.roc,
                row.pr,
                xerr=row.roc_sd,
                yerr=row.pr_sd,
                fmt=markers[dataset],
                ms=5,
                color=colors[row.split],
                capsize=2,
            )
    handles = [
        Line2D([], [], marker="o", color="none", markerfacecolor=color, label=split.title())
        for split, color in colors.items()
    ] + [
        Line2D([], [], marker="o", color=INK, label="CellPen", ls="none"),
        Line2D([], [], marker="s", color=INK, label="NonFouling", ls="none"),
    ]
    ax.set(xlabel="LinPept ROC-AUC", ylabel="LinPept PR-AUC", xlim=(0.70, 0.98), ylim=(0.62, 0.94))
    ax.legend(handles=handles, ncol=2, loc="lower right", fontsize=6.8)
    ax.set_title("Peptide transfer boundary")
    style_axis(ax)

    ax = axes[1, 1]
    cases = [("FreeSolv", 0.333, 6.43), ("FreeSolv", 0.304, 4.49), ("Lipophilicity", 0.217, 5.90)]
    ax.scatter(
        [c[1] for c in cases],
        [c[2] for c in cases],
        s=[45, 45, 55],
        color=[BLUE, BLUE2, ORANGE],
    )
    for name, value_x, value_y in cases:
        ax.annotate(name, (value_x, value_y), xytext=(4, 4), textcoords="offset points", fontsize=7.7)
    ax.axvline(0.5, color=INK, ls="--", lw=0.8)
    ax.set(xlabel="Nearest-neighbour Tanimoto", ylabel="Absolute error", xlim=(0.18, 0.52))
    ax.set_title("Representative low-similarity failures")
    style_axis(ax)
    fig.tight_layout(w_pad=2.2, h_pad=2.2)
    save_figure(fig, "fig05_paper19_chemical_boundaries")


def stage_existing_figures() -> None:
    for suffix in ["png", "svg"]:
        shutil.copy2(
            EXP / f"fig_paper19_candidate_diversity_selection_loss.{suffix}",
            FIG_DIR / f"fig02_paper19_candidate_diversity_selection_loss.{suffix}",
        )
        shutil.copy2(
            PAPER18_FIG / f"fig11_uq_conformal_4panel.{suffix}",
            FIG_DIR / f"fig04_paper19_uq_conformal_boundaries.{suffix}",
        )

    for path in EXP.glob("*.csv"):
        shutil.copy2(path, SRC_DIR / path.name)


def set_cell_border(cell, **edges: dict[str, str] | None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge in edges.items():
        if edge is None:
            continue
        tag = "w:" + edge_name
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge.items():
            element.set(qn("w:" + key), str(value))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def apply_three_line(table) -> None:
    visible_top = {"val": "single", "sz": "10", "space": "0", "color": "000000"}
    visible_mid = {"val": "single", "sz": "6", "space": "0", "color": "000000"}
    none = {"val": "nil"}
    for row_index, row in enumerate(table.rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, top=none, bottom=none, left=none, right=none, insideH=none, insideV=none)
            if row_index == 0:
                set_cell_border(cell, top=visible_top, bottom=visible_mid, left=none, right=none)
            if row_index == len(table.rows) - 1:
                set_cell_border(cell, bottom=visible_top, left=none, right=none)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(8.3)
        if row_index == 0:
            set_repeat_table_header(row)
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def set_run_font(run, size: float = 10.5, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for name, size, before, after in [
        ("Title", 16, 0, 10),
        ("Subtitle", 11, 0, 12),
        ("Heading 1", 13, 14, 7),
        ("Heading 2", 11.5, 10, 5),
        ("Heading 3", 10.5, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if "Heading" in name or name == "Title" else "宋体")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        if name in {"Title", "Subtitle"}:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style.paragraph_format.first_line_indent = Cm(0)

    if "Figure Caption" not in [s.name for s in doc.styles]:
        fig_style = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        fig_style = doc.styles["Figure Caption"]
    fig_style.font.name = "Times New Roman"
    fig_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    fig_style.font.size = Pt(9)
    fig_style.paragraph_format.space_before = Pt(3)
    fig_style.paragraph_format.space_after = Pt(8)
    fig_style.paragraph_format.line_spacing = 1.05
    fig_style.paragraph_format.keep_with_next = True

    if "Table Caption" not in [s.name for s in doc.styles]:
        table_style = doc.styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_style = doc.styles["Table Caption"]
    table_style.font.name = "Times New Roman"
    table_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    table_style.font.size = Pt(9)
    table_style.font.bold = True
    table_style.paragraph_format.space_before = Pt(8)
    table_style.paragraph_format.space_after = Pt(3)
    table_style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("Candidate-pool expansion and frozen auditing")
    set_run_font(run, 8)
    run.font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    set_run_font(run, 8)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def set_section_columns(section, count: int = 2, space_twips: int = 360) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols_el = cols[0]
    else:
        cols_el = OxmlElement("w:cols")
        sect_pr.append(cols_el)
    cols_el.set(qn("w:num"), str(count))
    cols_el.set(qn("w:space"), str(space_twips))


def add_body(doc: Document, text: str, italic: bool = False, no_indent: bool = False) -> None:
    p = doc.add_paragraph(style="Normal")
    if no_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, 10.5, italic=italic)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, {1: 13, 2: 11.5, 3: 10.5}[level], bold=True)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_declaration_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, 10.5, bold=True)


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Table Caption")
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, 9, bold=True)


def add_figure_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Figure Caption")
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, 9)


def add_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = str(value)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    apply_three_line(table)


def add_figure(doc: Document, path: Path, width_cm: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_figure_caption(doc, caption)


def extract_references() -> list[str]:
    source = Document(SOURCE_DOC)
    found = False
    refs: list[str] = []
    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        if text == "参考文献":
            found = True
            continue
        if found and text:
            refs.append(text)
    refs = [
        (
            "[11] Uchibori Y, Kaneko H. Generation of molecules near the applicability domain boundaries of property prediction models. J Chem Inf Model. 2026;66(12):6866–6879. doi:10.1021/acs.jcim.5c03220."
            if ref.startswith("[11] ")
            else ref
        )
        for ref in refs
    ]
    return refs


def load_values() -> dict[str, object]:
    effects_all = pd.read_csv(EXP / "paper19_k32_vs_k4_endpoint_effects.csv")
    diversity = pd.read_csv(EXP / "paper19_effective_diversity.csv")
    policy = pd.read_csv(EXP / "paper19_policy_units.csv")
    simulation = pd.read_csv(EXP / "paper19_oracle_extreme_value_simulation.csv")
    strong = pd.read_csv(EXP / "paper19_strong_baseline_effective_diversity.csv")
    budget = pd.read_csv(EXP / "paper19_compute_budget.csv")

    div = diversity.groupby("candidate_count").agg(
        entropy_rank=("outer_entropy_effective_rank", "mean"),
        median_corr=("outer_median_pairwise_correlation", "mean"),
        spearman=("mean_validation_audit_spearman", "mean"),
        top1=("top1_hit_rate", "mean"),
        mrr=("mrr", "mean"),
    )
    loss = (
        policy[policy.policy.eq("validation_best")]
        .groupby(["pool_size", "task_type"])
        .raw_selection_loss.mean()
        .unstack()
    )
    selector = (
        policy[policy.pool_size.eq(32)]
        .groupby(["policy", "task_type"])
        .agg(raw_loss=("raw_selection_loss", "mean"), normalized_loss=("range_normalized_audit_selection_loss", "mean"))
        .reset_index()
    )
    endpoint_loss = (
        policy[policy.policy.eq("validation_best")]
        .groupby(["task", "pool_size"], as_index=False)
        .raw_selection_loss.mean()
        .pivot(index="task", columns="pool_size", values="raw_selection_loss")
        .rename(columns={4: "mean_raw_loss_k4", 32: "mean_raw_loss_k32"})
        .reset_index()
    )
    effects = (
        effects_all[effects_all.policy.eq("validation_best")]
        .merge(endpoint_loss[["task", "mean_raw_loss_k4", "mean_raw_loss_k32"]], on="task", how="left")
        .rename(columns={"mean_delta_raw_loss_k32_minus_k4": "mean_raw_loss_delta"})
    )
    sim_key = simulation[
        simulation.truth_scenario.eq("equal_truth")
        & simulation.pairwise_candidate_correlation.eq(0.9)
        & simulation.effective_audit_sample_size.eq(50)
    ].set_index("candidate_count")
    values = {
        "effects": effects,
        "div": div,
        "loss": loss,
        "selector": selector,
        "simulation": simulation,
        "sim_key": sim_key,
        "strong": strong,
        "budget": budget,
        "classification_delta": float(
            effects.loc[effects.task_type.eq("classification"), "mean_raw_loss_delta"].mean()
        ),
        "regression_delta": float(effects.loc[effects.task_type.eq("regression"), "mean_raw_loss_delta"].mean()),
        "positive_endpoints": int((effects.mean_raw_loss_delta > 0).sum()),
        "total_fits": int(budget.total_fits.sum()),
        "fit_seconds": float(budget.total_fit_seconds.sum()),
    }
    return values


def fmt_ci(row: pd.Series) -> str:
    return f"{row['mean_raw_loss_delta']:.4f} ({row['ci95_low']:.4f}–{row['ci95_high']:.4f})"


def build_document(values: dict[str, object]) -> Document:
    effects: pd.DataFrame = values["effects"]  # type: ignore[assignment]
    div: pd.DataFrame = values["div"]  # type: ignore[assignment]
    loss: pd.DataFrame = values["loss"]  # type: ignore[assignment]
    selector: pd.DataFrame = values["selector"]  # type: ignore[assignment]
    sim_key: pd.DataFrame = values["sim_key"]  # type: ignore[assignment]
    strong: pd.DataFrame = values["strong"]  # type: ignore[assignment]

    class_delta = float(values["classification_delta"])
    reg_delta = float(values["regression_delta"])
    positive = int(values["positive_endpoints"])
    total_fits = int(values["total_fits"])
    fit_seconds = float(values["fit_seconds"])

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.first_line_indent = Cm(0)
    run = title.add_run("分子性质预测中候选池扩张的验证排序失真与选择损失：一项冻结审计研究")
    set_run_font(run, 16, bold=True)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.paragraph_format.first_line_indent = Cm(0)
    run = subtitle.add_run(
        "Validation-ranking distortion and selection loss under candidate-pool expansion in molecular property prediction: a frozen audit study"
    )
    set_run_font(run, 11, italic=True)

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "分子性质预测通常在有限验证数据上比较不断扩张的模型、表征与调参方案。即使外层标签从未参与训练，重复使用同一验证信息仍可能降低候选排序保真度，使有限审计集中的最佳观测结果混合真实模型差异与极值噪声。",
    )
    add_body(
        doc,
        "本研究将原有广泛治理叙事收缩为候选池扩张的冻结审计研究。在九个预先固定的公开终点上，我们采用 3 个外层折、3 个内层折和 5 个固定种子，联合报告名义候选数、谱有效秩、验证—审计排序一致性以及原始指标单位的选择损失；另以 12 候选多视图池和六任务现代强基线作为异质性压力测试，并用相关高斯误差模拟校准有限审计集观测赢家的极值乐观度。",
    )
    add_body(
        doc,
        f"在高度相关的 32 候选压力池中，K 从 4 增至 32 时，validation-best 的平均原始单位选择损失在分类任务中增加 {class_delta:.4f} ROC-AUC，在回归任务中增加 {reg_delta:.4f} RMSE，且 {positive}/9 个终点同向；谱有效秩仅由 {div.loc[4, 'entropy_rank']:.2f} 增至 {div.loc[32, 'entropy_rank']:.2f}，验证—审计 Spearman 相关由 {div.loc[4, 'spearman']:.3f} 降至 {div.loc[32, 'spearman']:.3f}。完整多视图池相对 Morgan-only 选择的配对归一化实际兑现收益为 0.343（95% CI 0.210–0.483；9/9 终点），表明异质表征可产生收益，但收益必须与新增选择自由度同时审计。跨终点风险模型仅保留为探索性结果。",
    )
    add_body(
        doc,
        "这些结果支持在分子机器学习比较中同时报告候选资格、有效多样性、原始单位选择损失、失败状态和逐折源数据。本文不提出新的预测主干，也不把外层审计称为独立确认；由于缺少真正未参与候选设计的时间外或来源外锁箱，有限审计集的观测上界不能解释为真实泛化上界。保形预测、MoleculeACE 与 bRo5 结果仅用于界定可靠性和化学边界。",
    )

    add_heading(doc, "Scientific contribution", 2)
    add_body(
        doc,
        "This study separates nominal candidate-pool expansion, effective candidate diversity, validation-ranking fidelity and finite-audit winner optimism within a frozen nested evaluation of molecular property prediction. It records candidate eligibility, failures, computational exposure and fold-level outputs, while explicitly distinguishing the outer audit from an unavailable independent confirmation layer. The contribution is therefore an auditable model-selection analysis, not a new predictive backbone or a claim of universal model superiority.",
        no_indent=True,
    )
    add_body(
        doc,
        "关键词：分子性质预测；模型选择偏差；候选池扩张；嵌套交叉验证；验证信息过度使用；审计集乐观度；有效候选多样性；可复现化学信息学",
        no_indent=True,
    )

    add_heading(doc, "1 引言", 1)
    add_body(
        doc,
        "分子性质预测已从少数固定模型的离线比较扩展为同时搜索指纹、二维描述符、图神经网络、消息传递网络、化学语言模型和自动机器学习流程。候选扩张可以发现互补表征并提高观测性能，但也增加了在同一验证信息上作出选择的自由度。在小样本、类别不平衡和骨架迁移常见的分子任务中，验证排序的微小波动可能改变最终入选模型，因而模型比较本身成为需要量化的统计过程[1–4,21–24,34,35]。",
    )
    add_body(
        doc,
        "嵌套交叉验证可将内层选择与外层评估分开，MoleculeNet 与 TDC 统一了部分任务定义，近期 ADMET、OOD 与不确定性研究则扩展了真实挑战、适用域和数据偏移下的可靠性评价[1,2,5,9–14,25–28,33]。然而，这些工具并不自动回答三个相互关联的问题：名义候选数是否对应独立的信息增量；有限外层样本中的最大值包含多少赢家诅咒；不同模型是否获得了可比较的搜索预算。若这些因素未被区分，候选池增大带来的观测上界、选择损失和真实表征收益会被混合解释。",
    )
    add_body(
        doc,
        "现有稿件中的 32 候选池主要由相关的传统学习器和超参数变体组成，因此适合检验近重复候选下的选择压力，却不能代表 GNN、D-MPNN 与基础模型共同扩张的完整异质搜索空间。近期 ADMET 基准和多模态模型已显示 TabPFNv2、预训练 GNN、双组分预训练模型、语法树表示与描述符融合具有明显的情景依赖优势[5–8]；这些工作构成本文必须尊重的预测性能参照。与此同时，有限外层审计集上的最优候选只是观测审计上界，而非真实可达到上界。基于这一识别问题，本文将原始单位选择损失设为主要结果，将范围归一化损失降为敏感性分析，并以谱有效秩、错误重叠和排序一致性补充名义 K。",
    )
    add_body(
        doc,
        "本文围绕一个限定性问题展开：在验证信息和拟合次数固定时，名义候选规模与有效多样性如何共同影响验证排序保真度和外层审计选择损失。贡献包括三点：其一，重建候选登记、冻结选择、外层审计和负结果记录的可追踪时间线；其二，在九个公开终点上以 K=4/8/16/32 配对单位量化原始单位损失、有效秩和排名退化，并用模拟估计有限审计集最大值的机械乐观度；其三，以多视图、现代强基线、保形和化学边界分析检验主结论的适用范围。本文不把九个终点视为分子性质任务总体的随机样本，也不声称已完成独立确认或第三方冷启动复现。",
    )

    add_heading(doc, "2 材料与方法", 1)
    add_heading(doc, "2.1 研究锁定与主张层级", 2)
    add_body(
        doc,
        "本次分析使用既有实验输出重建最终分析锁定文件，固定终点、指标方向、外层与内层划分、种子、候选登记顺序、主要统计量和图表源数据。由于部分历史外层结果在该锁定之前已经存在，本研究不将其描述为前瞻性预登记；相应证据属于回顾性冻结审计。确认性主张仅覆盖九终点近重复候选压力实验中的原始单位选择损失和排名保真度。12 候选多视图与六任务现代强基线为次要压力测试，TDC 扩展、保形、MoleculeACE、bRo5 和决策模拟为边界或探索性分析。",
    )
    add_heading(doc, "2.2 数据来源与分子标准化", 2)
    add_body(
        doc,
        "核心面板包含 ESOL、FreeSolv、Lipophilicity、BBBP、BACE、ClinTox、Caco2 Wang、HIA Hou 和 P-gp Broccatelli。前三个及 Caco2 为回归任务，主要指标为 RMSE；其余为分类任务，主要指标为 ROC-AUC。终点并非从分子性质任务总体中随机抽取，因此统计推断以终点为聚类单位并限定于这些公开任务。MoleculeACE 与 bRo5 数据仅用于活性悬崖、低相似度和迁移边界，不参与主效应合并[1,2,29,30]。",
    )
    add_body(
        doc,
        "SMILES 使用 RDKit 解析、清理和标准化，选择最大分子片段并在可行时中和电荷。分类重复结构仅在标签一致时合并，冲突标签组排除；回归重复结构按均值聚合并保留重复计数。所有插补、标准化、类别权重、校准与目标变换均仅在相应训练折估计。分子标识、分割索引、候选状态及失败原因保存在机器可读清单中。",
    )
    add_table_caption(doc, "表 1  核心终点、主要指标与证据角色")
    endpoint_rows = [
        ["ESOL", "回归", "RMSE", "MoleculeNet", "近重复扩池；多视图；强基线"],
        ["FreeSolv", "回归", "RMSE", "MoleculeNet", "近重复扩池；多视图；强基线"],
        ["Lipophilicity", "回归", "RMSE", "MoleculeNet", "近重复扩池；多视图；强基线"],
        ["BBBP", "分类", "ROC-AUC", "MoleculeNet", "近重复扩池；多视图；强基线"],
        ["BACE", "分类", "ROC-AUC", "MoleculeNet", "近重复扩池；多视图；强基线"],
        ["ClinTox", "分类", "ROC-AUC", "MoleculeNet", "近重复扩池；少数类负结果"],
        ["Caco2 Wang", "回归", "RMSE", "TDC", "近重复扩池；多视图"],
        ["HIA Hou", "分类", "ROC-AUC", "TDC", "近重复扩池；多视图"],
        ["P-gp Broccatelli", "分类", "ROC-AUC", "TDC", "近重复扩池；多视图"],
    ]
    add_table(doc, ["终点", "任务", "主指标", "来源", "本研究角色"], endpoint_rows)

    add_heading(doc, "2.3 候选池构建与计算预算", 2)
    add_body(
        doc,
        "主要压力实验使用固定登记顺序的 32 个传统模型与超参数变体，并以嵌套前缀形成 K=4/8/16/32。该设计保持数据划分和拟合协议不变，用于检验高度相关候选下的验证消费，不作为完整异质模型池。每个候选的 family、representation、registry_order、eligible、failed 和 config_hash 均在外层审计前固定；失败候选不由新配置替换。",
    )
    add_body(
        doc,
        "异质性压力测试包含 12 个多视图候选，由 Morgan-512、MACCS、RDKit2D 和拼接多视图分别配对线性模型、随机森林与 LightGBM。现代强基线面板在六个 MoleculeNet 任务上比较 RDKit-RF、GNN-GCN、ChemBERTa 与 MoLFormer；Chemprop/D-MPNN 仅有三个终点的边界结果，TabPFN 无同划分预测导出，均未并入九终点主效应。预算采用固定拟合次数协议，并同时报告记录到的墙钟时间；该协议不等同于 GPU/CPU 秒数严格相等。",
    )
    add_table_caption(doc, "表 2  候选池、冻结单位与预算边界")
    add_table(
        doc,
        ["分析层", "候选组成", "冻结重复", "拟合/记录", "证据边界"],
        [
            [
                "近重复 K 阶梯",
                "32 个传统模型/超参数变体；K=4/8/16/32",
                "9 终点 × 5 种子 × 3 外层 × 3 内层",
                f"{total_fits:,} 次拟合；{fit_seconds:,.1f} s 记录墙钟时间",
                "主要选择压力实验；非完整异质池",
            ],
            [
                "多视图压力测试",
                "4 表征 × 3 学习器 = 12 候选",
                "9 终点 × 5 种子 × 3 外层 × 3 内层",
                "6,480 次拟合",
                "异质表示的次要确认",
            ],
            [
                "现代强基线",
                "RDKit-RF、GNN-GCN、ChemBERTa、MoLFormer",
                "6 任务 × 5 种子 × 3 外层 × 3 内层",
                "360 外层单元；220,040 条逐样本预测",
                "六任务压力测试；非九终点全量确认",
            ],
        ],
    )

    add_figure(
        doc,
        FIG_DIR / "fig01_paper19_frozen_audit_workflow.png",
        16.2,
        "图 1  冻结候选选择与外层审计流程。终点、指标、种子和候选顺序在 Paper 19 分析前锁定；所有变换与候选排序仅使用开发层数据，外层审计仅估计冻结流程的选择损失。虚线框表示本修订尚未具备真正独立的时间外或来源外确认层，因此外层最优候选仅称为观测审计集上界。",
    )

    add_heading(doc, "2.4 嵌套选择、外层审计与禁止回流规则", 2)
    add_body(
        doc,
        "每个终点使用 5 个固定种子（11、23、37、53、71）、3 个外层折和每个外层训练集中的 3 个内层折。validation-best 按内层平均效用选择；one-SE 在最高内层效用的一个标准误范围内偏向稳定候选；risk-adjusted 使用内层可见的均值与波动。候选一经选定即在外层训练部分重拟合，并只在对应外层审计集上评价一次。任何查看外层结果后提出的配置变化均标记为 post-audit exploratory update，不替换冻结结果。",
    )
    add_body(
        doc,
        "本研究没有可证明从未参与候选设计的独立锁箱。因而，外层审计结果用于估计模型选择过程的回顾性差距，但不用于声称独立泛化收益。选择时间线、分析锁定文件和候选状态表随代码包提供，分别标记 pre-audit frozen selection、outer-audit result 和 post-audit exploratory update。",
    )

    add_heading(doc, "2.5 结局、估计量与术语", 2)
    add_body(
        doc,
        "设外层审计单元为 i，候选为 a，分类效用为 ROC-AUC，回归效用为负 RMSE。观测审计集上界记为 U_i^audit=max_a U_{i,a}，冻结选择器选出的候选记为 â_i。主要结局为原始单位审计选择损失：分类任务 L_i=AUC_i^audit−AUC_{i,â_i}，回归任务 L_i=RMSE_{i,â_i}−RMSE_i^audit。该定义始终为非负，并分别保留 ROC-AUC 与 RMSE 的解释尺度。完整 32 候选池范围归一化损失仅作为跨终点敏感性分析。",
    )
    add_body(
        doc,
        "验证排序保真度由候选内层排名与外层审计排名之间的 Spearman、Kendall、Top-1 命中和平均倒数排名（MRR）描述。候选有效多样性基于候选外层效用相关矩阵的特征值 λj 计算谱熵有效秩 exp[−Σpj log(pj)]，其中 pj=λj/Σλj；并同时报告参与率有效秩、两两相关中位数和现代强基线逐样本错误集合的 Jaccard 重叠。名义 K 与有效秩不互换。",
    )

    add_heading(doc, "2.6 对照、模拟与统计推断", 2)
    add_body(
        doc,
        "固定单候选、one-SE、risk-adjusted 和 validation-best 与均匀随机选择的解析期望比较。随机选择期望由每个外层单元候选效用的算术平均得到，不引入额外蒙特卡洛误差。历史随机排序负对照和信号恢复正对照用于检验排名统计量的零点与单调性，但不作为模型性能证据。",
    )
    add_body(
        doc,
        "为分离有限审计集最大值的机械乐观度，我们构造相关高斯测量误差模拟。候选真值设为完全相同或弱梯度，名义 K 为 4/8/16/32，有效审计样本量为 25/50/100/200，候选误差相关 ρ 为 0/0.5/0.9/0.99；每个条件重复 30,000 次。模拟统计观测赢家估计值减去其已知真值的差，仅用于校准极值偏差，不替代独立确认数据。",
    )
    add_body(
        doc,
        "外层折和种子先在终点内配对，终点为主要统计单位。端点层效应及 95% 区间由 10,000 次聚类 bootstrap 计算；方向一致率同时报告。leave-one-endpoint-out（LOEO）分析每次删除一个终点后重新计算任务类型内平均效应，以检查 ClinTox、FreeSolv 等单一终点是否主导结论。九个终点不足以训练并验证通用元选择器，因此跨终点风险模型仅作探索性分析。",
    )

    add_heading(doc, "2.7 强基线实现与公平性", 2)
    add_body(
        doc,
        "所有进入完成性比较的强基线必须使用相同外层与内层索引，固定随机种子，并导出候选级内层分数、外层分数和逐样本预测。预训练模型记录 checkpoint、tokenization、冻结适配头、学习率与失败状态。由于当前完成面板只覆盖四个候选和六个 MoleculeNet 任务，结果用于检验错误相关性和结论边界，而不构成九终点模型家族排行榜。Chemprop/D-MPNN 和 TabPFN 的缺失不以其他候选替换。",
    )

    add_heading(doc, "2.8 可靠性与化学边界分析", 2)
    add_body(
        doc,
        "分类可靠性比较 pooled split conformal、label-conditional conformal 与 Mondrian label-similarity conformal；回归比较残差式 split conformal、Mondrian residual 和 conformalized quantile regression（CQR）。另按最近邻 Morgan Tanimoto 分层评估 scaffold/OOD 校准，并以集成分歧富集高误差样本。ClinTox 少数类结果预先作为可能的负结果单独报告，覆盖改善不被解释为筛选召回自动改善[26–28,33]。",
    )
    add_body(
        doc,
        "化学边界使用 MoleculeACE 的活性悬崖子集、分子对差异相关与方向准确率，以及 CycPept-PAMPA 和 LinPept 的随机、骨架、外缘或时间划分。最近邻 Tanimoto、scaffold novelty、bRo5 外缘和极端标签仅界定适用范围，不参与主效应统计[29,30]。固定预算筛选效用被标记为 retrospective fixed-budget utility simulation，而非真实药物研发决策价值。",
    )

    add_heading(doc, "2.9 复现包与审计文件", 2)
    add_body(
        doc,
        "随本修订生成的代码包包含 README、Python 包、实验入口、环境锁文件、Dockerfile、候选与终点登记、选择时间线、计算预算、逐折汇总、图源数据和 SHA-256 清单。原始数据仅通过 manifest 指向原始公共来源，是否可再分发由其许可证决定。第三方冷启动复现按作者要求暂缓，因此代码包明确标记为 author-side reproducibility package，而不声称 independent reproduction completed[32]。",
    )

    add_heading(doc, "3 结果", 1)
    add_heading(doc, "3.1 名义候选数显著高于有效候选多样性", 2)
    add_body(
        doc,
        f"在近重复压力池中，名义 K 从 4 增至 32，而终点平均谱熵有效秩仅由 {div.loc[4, 'entropy_rank']:.2f} 增至 {div.loc[32, 'entropy_rank']:.2f}；K=16 时为 {div.loc[16, 'entropy_rank']:.2f}，随后未继续增加。外层候选效用两两相关中位数在 K=4/8/16/32 时分别为 {div.loc[4, 'median_corr']:.3f}/{div.loc[8, 'median_corr']:.3f}/{div.loc[16, 'median_corr']:.3f}/{div.loc[32, 'median_corr']:.3f}。因此，32 个登记候选不能解释为 32 个独立信息源，扩池效应同时包含搜索自由度与候选相关结构。",
    )
    add_body(
        doc,
        f"与有效秩的有限增加相伴，验证—外层审计 Spearman 相关由 {div.loc[4, 'spearman']:.3f} 降至 {div.loc[32, 'spearman']:.3f}，Top-1 一致率由 {div.loc[4, 'top1']:.3f} 降至 {div.loc[32, 'top1']:.3f}，MRR 由 {div.loc[4, 'mrr']:.3f} 降至 {div.loc[32, 'mrr']:.3f}。这说明名义候选数增加没有带来等比例的独立候选信息，却扩大了从有限验证信号中区分近似候选的难度。",
    )
    add_figure(
        doc,
        FIG_DIR / "fig02_paper19_candidate_diversity_selection_loss.png",
        16.3,
        "图 2  名义候选规模、有效多样性、原始单位选择损失与有限审计集极值偏差。左上为九终点的谱熵有效秩；右上为 validation-best 的完整池范围归一化审计选择损失，灰线表示终点、黑线表示终点均值；左下为 K=32 相对 K=4 的原始单位损失变化及终点聚类 bootstrap 95% 区间；右下为候选真值相同、有效审计样本量 50 时相关高斯误差模拟的观测赢家乐观度。模拟重复 30,000 次，仅校准机械极值偏差。",
    )

    add_table_caption(doc, "表 3  候选规模、有效多样性与冻结选择结果")
    k_rows = []
    for k in [4, 8, 16, 32]:
        k_rows.append(
            [
                k,
                f"{div.loc[k, 'entropy_rank']:.3f}",
                f"{div.loc[k, 'median_corr']:.3f}",
                f"{div.loc[k, 'spearman']:.3f}",
                f"{div.loc[k, 'top1']:.3f}",
                f"{div.loc[k, 'mrr']:.3f}",
                f"{loss.loc[k, 'classification']:.4f}",
                f"{loss.loc[k, 'regression']:.4f}",
            ]
        )
    add_table(
        doc,
        ["K", "谱有效秩", "候选相关中位数", "验证—审计 Spearman", "Top-1", "MRR", "分类损失 (AUC)", "回归损失 (RMSE)"],
        k_rows,
    )

    add_heading(doc, "3.2 原始单位选择损失随候选规模增加", 2)
    add_body(
        doc,
        f"validation-best 在 K=4 时的平均原始单位选择损失为分类 {loss.loc[4, 'classification']:.4f} ROC-AUC、回归 {loss.loc[4, 'regression']:.4f} RMSE；K=32 时分别为 {loss.loc[32, 'classification']:.4f} 和 {loss.loc[32, 'regression']:.4f}。对应增加量为 {class_delta:.4f} ROC-AUC 与 {reg_delta:.4f} RMSE。九个终点中八个同向，唯一未增加的 P-gp Broccatelli 效应为 −0.0003，区间跨零。",
    )
    add_body(
        doc,
        "端点差异明显：ClinTox 的分类损失增加 0.0203，ESOL 与 FreeSolv 的回归损失分别增加 0.0940 和 0.0743。LOEO 分析中，无论删除哪一个终点，剩余分类和回归任务的平均 K=32−K=4 效应仍为正，说明总体方向不是由 ClinTox 或 FreeSolv 单独驱动。由于终点并非随机抽样，本研究将这一结果解释为所选九个任务中的稳定方向，而非领域总体效应。",
    )
    effect_rows: list[list[object]] = []
    labels = {
        "bace": "BACE",
        "bbbp": "BBBP",
        "clintox": "ClinTox",
        "esol": "ESOL",
        "freesolv": "FreeSolv",
        "lipo": "Lipophilicity",
        "tdc_caco2_wang": "Caco2 Wang",
        "tdc_hia_hou": "HIA Hou",
        "tdc_pgp_broccatelli": "P-gp Broccatelli",
    }
    for _, row in effects.sort_values(["task_type", "task"]).iterrows():
        effect_rows.append(
            [
                labels.get(row.task, row.task),
                "ROC-AUC" if row.task_type == "classification" else "RMSE",
                f"{row.mean_raw_loss_k4:.4f}",
                f"{row.mean_raw_loss_k32:.4f}",
                fmt_ci(row),
                "增加" if row.mean_raw_loss_delta > 0 else "未增加",
            ]
        )
    add_table_caption(doc, "表 4  K=32 相对 K=4 的终点层原始单位审计选择损失")
    add_table(
        doc,
        ["终点", "单位", "K=4", "K=32", "差值 (95% CI)", "方向"],
        effect_rows,
    )

    add_body(
        doc,
        "冻结选择仍明显优于无信息选择。K=32 时，validation-best 的平均原始单位损失为分类 0.0138、回归 0.0700；均匀随机选择的解析期望分别为 0.0295 和 0.1782。one-SE 与 risk-adjusted 在部分任务上接近 validation-best，但没有形成跨终点的普适优势。该结果将主张限定为“候选扩张增加选择压力”，而非“验证选择完全失效”。",
    )

    add_heading(doc, "3.3 有限审计集观测赢家存在机械极值乐观度", 2)
    add_body(
        doc,
        f"当所有候选真值相同、误差相关 ρ=0.9、有效审计样本量为 50 时，观测赢家相对其已知真值的平均乐观度由 K=4 的 {sim_key.loc[4, 'mean_observed_oracle_optimism']:.3f} 个标准差单位增至 K=32 的 {sim_key.loc[32, 'mean_observed_oracle_optimism']:.3f}。相关越低或审计样本量越小，乐观度越高；ρ=0.99 时效应显著减弱但未消失。由此，外层审计集最大值会随 K 机械上升，不能直接作为真实可达到上界。",
    )
    add_body(
        doc,
        "该模拟只识别统计上必然存在的有限样本最大值偏差，不能估计本数据中审计赢家在真正时间外数据上的兑现率。由于本修订没有独立锁箱，observed audit-set oracle optimism 的实证分量仍未被直接测量。这一缺口在摘要、图 1、讨论与代码包状态文件中保持一致，不以边界面板或重复交叉验证替代。",
    )

    add_heading(doc, "3.4 多视图收益可兑现，但现代强基线仍是有限压力测试", 2)
    add_body(
        doc,
        "九终点 12 候选多视图实验在共享的 3×3×5 划分上完成 6,480 次拟合。完整多视图池的 validation-best 平均范围归一化审计选择损失为 0.043（95% CI 0.021–0.067），低于固定 Morgan-RF 的 0.395、one-SE 的 0.073 和 risk-adjusted 的 0.054。完整池相对 Morgan-only 选择的实际兑现增益为 0.343（0.210–0.483；9/9 终点）；拼接多视图相对仅允许独立视图候选的增益为 0.035（0.017–0.053）。这些收益发生在外层审计中，但仍不是独立时间外确认。",
    )
    add_body(
        doc,
        "在 135 个冻结选择单元中，拼接多视图被选 84 次、RDKit2D 44 次、MACCS 4 次、Morgan 3 次，说明异质表示的收益并非固定学习器的机械优势。六任务现代强基线中，四个候选的谱有效秩范围为 1.64–2.83，高于近重复池的多数 K 水平；候选对逐样本错误 Jaccard 重叠平均为 0.215（范围 0.168–0.296），表明错误既有相关性也存在互补空间。该面板未包含九终点全量 D-MPNN 和 TabPFN，因此只作为主结论的边界压力测试。",
    )
    add_figure(
        doc,
        FIG_DIR / "fig03_paper19_multiview_confirmation.png",
        16.3,
        "图 3  共享冻结划分下的多视图压力测试。左上比较四种冻结选择策略的范围归一化外层审计损失，误差线为终点聚类 95% 区间；右上显示观测审计上界、冻结选择与视图拼接的配对归一化效用变化；左下与右下分别给出 135 个外层单元中的表征选择频数和终点—表征分布。所有外层结果均未回流更新确认性选择。",
    )
    add_table_caption(doc, "表 5  次要压力测试与边界结果")
    add_table(
        doc,
        ["模块", "主要观察", "完成范围", "允许的解释"],
        [
            ["多视图冻结选择", "实际兑现增益 0.343 (0.210–0.483)", "9 终点；12 候选；3×3×5", "共享外层审计中的异质表示收益"],
            ["现代强基线", "Top-1 0.878；平均归一化损失 0.0048", "6 任务；4 候选", "有限压力测试；非九终点排行榜"],
            ["Error overlap", "平均 Jaccard 0.215 (0.168–0.296)", "6 任务；6 候选对", "候选错误非独立"],
            ["去重敏感性", "最大平均效用变化 0.022，出现在 ClinTox", "6 任务；3 策略；270 外层单元", "少数类任务对重复规则敏感"],
            ["独立确认", "未完成", "无未参与候选设计的锁箱", "不得声称真实泛化兑现率"],
        ],
    )

    add_heading(doc, "3.5 可靠性和化学边界限制了部署性解释", 2)
    add_body(
        doc,
        "在 90% 标称覆盖下，pooled split conformal 的分类少数类覆盖为 0.627，label-conditional 与 Mondrian label-similarity 分别提高到 0.895 和 0.898。回归 CQR 的平均覆盖为 0.882、平均区间宽度为 7.25，未在所有终点稳定优于残差式方法。最近邻 Tanimoto<0.5 的分类子集 ROC-AUC 为 0.803，低于 >0.7 子集的 0.924；集成不确定性对 top-10% 高误差样本的平均富集为 1.54。ClinTox 少数类召回仍不足，因此覆盖校准不能替代毒性假阴性控制。",
    )
    add_figure(
        doc,
        FIG_DIR / "fig04_paper19_uq_conformal_boundaries.png",
        16.3,
        "图 4  不确定性与保形预测边界。四个面板分别比较 90% 标称覆盖下的总体与少数类覆盖、回归覆盖—区间宽度权衡、最近邻 Tanimoto 分层的 OOD/scaffold 校准，以及集成不确定性对 top-10% 高误差样本的富集。虚线表示标称覆盖或无富集参照。结果用于风险分层，不构成部署级可靠性保证。",
    )
    add_body(
        doc,
        "MoleculeACE 的 17 个任务整体 RMSE 为 0.711，活性悬崖子集 RMSE 为 0.813；高相似分子对差异 Spearman 平均为 0.252，方向准确率为 0.750。CycPept-PAMPA 的随机、骨架、外缘和时间划分 RMSE 分别为 0.547、0.727、0.876 和 0.768。更严格划分下的退化与低相似度失败案例一致，表明平均基准性能不能直接外推到局部结构非连续或 bRo5 边界。",
    )
    add_figure(
        doc,
        FIG_DIR / "fig05_paper19_chemical_boundaries.png",
        16.3,
        "图 5  化学边界与代表性失败。四个面板依次显示 MoleculeACE 任务间活性差异相关、CycPept-PAMPA 在随机/骨架/外缘/时间划分下的 RMSE、LinPept 迁移边界，以及三个低相似度高误差案例。误差线为相应重复的标准差；分子对方向准确率在正文报告，这些边界分析均不参与九终点主效应合并。",
    )

    add_heading(doc, "4 讨论", 1)
    add_heading(doc, "4.1 候选扩张同时包含真实收益与观测赢家偏差", 2)
    add_body(
        doc,
        "本研究最直接的发现不是“候选越多越差”，而是候选扩张产生了两个必须分开的量。多视图实验表明，异质表示可以在冻结外层审计中兑现收益；近重复压力实验则表明，当验证信息固定时，区分更多相近候选会降低排序保真度并增加选择损失。相关误差模拟进一步显示，即便候选真值完全相同，有限审计集最大值也会随 K 增大。因此，观测上界收益不能单独证明真实模型能力增加。",
    )
    add_heading(doc, "4.2 名义 K 不能替代有效候选多样性", 2)
    add_body(
        doc,
        "32 个登记候选的谱有效秩约为 2，而六任务现代强基线的四候选有效秩可达到 1.64–2.83，说明候选数量与信息维度没有固定换算关系。错误重叠同样表明，复杂模型之间既非独立也非完全冗余。今后的候选池研究应同时报告名义 K、候选预测或效用相关、谱有效秩、错误重叠和家族构成；不同论文中的 K 若预算和相关结构不同，不能直接横向比较。",
    )
    add_heading(doc, "4.3 冻结审计能解决什么，不能解决什么", 2)
    add_body(
        doc,
        "候选登记、选择规则冻结和外层标签禁止回流可以减少事后替换自由度，并使失败运行、负结果和分析时间线可追踪。它们不能增加样本量，也不能自动修复标签噪声、骨架偏移、模型错设或校准失效。validation-best 在本数据中优于若干保守规则，并不意味着它是普适选择器；相反，这一结果说明基于有限内层方差构造的复杂惩罚本身也可能不稳定。",
    )
    add_heading(doc, "4.4 对分子基准和筛选决策的含义", 2)
    add_body(
        doc,
        "对于分子基准，模型分数应与候选进入时间、调参机会、计算暴露和失败状态一同解释。对于药物筛选，top-k 富集、毒性假阴性成本与实验队列模拟只能称为回顾性固定预算效用分析；公开数据上的排序改善不等同于真实项目的前瞻性决策收益。保形集合、最近邻相似度和不确定性富集可帮助标记需复核样本，但 ClinTox 负结果说明它们不能替代任务特异的成本权重和阈值验证。",
    )
    add_body(
        doc,
        "与近期工作的差异因此应落在证据问题而非排行榜措辞上。Zhao 等的 ADMET 研究覆盖 TabPFNv2、预训练 GNN、AutoML、bRo5 和活性悬崖，DCPM-ADMET、KROVEX 与 MolGramTreeNet 分别展示了预训练组件、统计筛选描述符融合和语法树多模态表示的潜力[5–8]；相关 OOD/UQ 工作进一步说明校准会随数据偏移退化[9–13,33]。本文没有在同等规模上超越这些研究，而是补充了它们通常不以主结果呈现的候选登记、有效搜索自由度、有限审计集极值偏差和冻结选择损失。",
    )
    add_heading(doc, "4.5 局限性", 2)
    add_body(
        doc,
        "本研究仍有四项限制。第一，九个核心终点来自若干公开任务，而非从分子性质任务总体中随机抽样，效应范围不应外推为领域普遍规律。第二，本修订没有真正未参与候选设计的时间外、来源外或一次性锁箱，因而不能实证估计审计赢家的独立泛化乐观度；相关模拟只提供机制校准。第三，候选有效多样性依赖模型家族、表征和预算定义，当前 32 候选池高度相关，六任务现代强基线也尚未扩展为九终点完整异质 K 阶梯。第四，冻结审计主要减少事后选择自由度并提高可追溯性，不能自动修复标签噪声、分布偏移、校准失效或不充分的模型搜索；第三方冷启动复现亦尚未完成。",
    )

    add_heading(doc, "5 结论", 1)
    add_body(
        doc,
        "在所研究的九个公开分子性质终点中，候选池扩张的影响不能由最佳观测分数单独概括。名义 K、有效候选多样性、验证排序保真度、原始单位外层审计选择损失和有限审计集赢家偏差需要共同报告。近重复候选池显示 K=32 相对 K=4 的选择损失在八个终点增加，而共享冻结划分的多视图实验同时证明异质表征可产生可兑现的外层审计收益。",
    )
    add_body(
        doc,
        "这些证据支持一种更严格的比较实践：候选在进入搜索前登记资格与预算，选择规则在外层审计前冻结，失败状态和逐折输出被保留，观测审计集上界不被称为真实泛化上界。该协议提高了模型选择研究的可审计性，但在完成独立锁箱、九终点完整异质强基线和第三方复跑之前，不构成部署级或领域级确证。",
    )

    add_heading(doc, "声明", 1)
    add_declaration_heading(doc, "Availability of data and materials")
    add_body(
        doc,
        "本研究使用的数据均来自其原始公共来源。随修订代码包提供版本化 data manifest、终点定义、预处理规则、划分索引和校验值；原始数据仅在原始许可证允许时再分发。永久公开仓库与 Zenodo DOI 尚待作者在投稿前建立。",
        no_indent=True,
    )
    add_declaration_heading(doc, "Code availability")
    add_body(
        doc,
        "本修订附带本地版本化代码包，覆盖数据处理、候选登记、嵌套选择、Paper 19 新增统计分析、主图重建和源数据导出，并包含 Dockerfile、环境锁文件、测试与 SHA-256 清单。该包尚未完成独立第三方冷启动复现，不能写为 publicly independently reproduced。",
        no_indent=True,
    )
    add_declaration_heading(doc, "Authors’ contributions")
    add_body(doc, "[待作者依据 CRediT taxonomy 填写并逐一核实。]", no_indent=True)
    add_declaration_heading(doc, "Funding")
    add_body(doc, "[待作者核实基金名称、项目编号及资助方角色。]", no_indent=True)
    add_declaration_heading(doc, "Competing interests")
    add_body(doc, "[待全体作者核实并声明。]", no_indent=True)
    add_declaration_heading(doc, "Acknowledgements")
    add_body(doc, "[待作者补充；不得将未同意署名的人员写入。]", no_indent=True)

    add_heading(doc, "参考文献", 1)
    reference_section = doc.add_section(WD_SECTION.CONTINUOUS)
    reference_section.top_margin = Cm(2.3)
    reference_section.bottom_margin = Cm(2.2)
    reference_section.left_margin = Cm(2.35)
    reference_section.right_margin = Cm(2.35)
    reference_section.header_distance = Cm(1.0)
    reference_section.footer_distance = Cm(1.0)
    set_section_columns(reference_section, count=2, space_twips=420)
    for ref in extract_references():
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.55)
        p.paragraph_format.line_spacing = 0.95
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(ref)
        set_run_font(run, 7.8)

    return doc


def write_report(values: dict[str, object], document: Document) -> None:
    effects: pd.DataFrame = values["effects"]  # type: ignore[assignment]
    report = f"""# 小论文-19 修改与实验审计报告

## 一句话主论证

在所选公开分子性质终点中，候选池名义规模增加并未带来等比例的有效候选信息，却降低了验证—审计排序保真度并增加原始单位选择损失；多视图异质表示可产生外层审计收益，但独立泛化兑现率仍需真正锁箱确认。

## 本轮新增且已通过的实验

| 模块 | 完成结果 | 状态 |
|---|---:|---|
| 原始单位选择损失 | 分类 K32−K4 = {float(values['classification_delta']):.4f} ROC-AUC；回归 = {float(values['regression_delta']):.4f} RMSE；{int(values['positive_endpoints'])}/9 同向 | 已完成 |
| 有效候选多样性 | K=4/32 谱有效秩 = {values['div'].loc[4, 'entropy_rank']:.2f}/{values['div'].loc[32, 'entropy_rank']:.2f} | 已完成 |
| 排名保真度 | Spearman {values['div'].loc[4, 'spearman']:.3f}→{values['div'].loc[32, 'spearman']:.3f}；Top-1 {values['div'].loc[4, 'top1']:.3f}→{values['div'].loc[32, 'top1']:.3f} | 已完成 |
| LOEO | 删除任一终点后，分类与回归平均效应仍为正 | 已完成 |
| 随机选择基线 | K=32 validation-best 明显优于均匀随机期望 | 已完成 |
| 审计赢家极值偏差 | 30,000 次/条件相关高斯误差模拟；K、相关性与有效审计样本量全网格 | 已完成 |
| 预算与时间线 | {int(values['total_fits']):,} 次拟合、{float(values['fit_seconds']):,.1f} s 记录墙钟时间；候选登记和 selection timeline 已导出 | 已完成 |

## 明确未完成且已降级的项目

| 项目 | 论文中的处理 |
|---|---|
| 独立时间外/来源外/final lockbox | 未完成；外层结果统一称为 outer audit，不称 independent confirmation |
| 九终点完整异质 K=4/8/16/32 | 未完成；32 候选池降级为近重复压力实验，多视图与六任务强基线为次要压力测试 |
| Chemprop/D-MPNN 九终点全量 | 未完成；仅三终点边界结果 |
| TabPFN 同划分预测导出 | 未完成；候选状态保留为 unavailable/failed，不替换 |
| 历史前瞻性预登记证明 | 不具备；改称 retrospective reconstructed analysis lock |
| 第三方冷启动复现 | 按本轮要求暂缓；只交付代码包 |
| 永久公开仓库与 Zenodo DOI | 待作者建立；正文未写成已公开 |

## 文档结构与格式审计

- 主结果重排为 5 个模块，主图 5 张。
- `test oracle` 已在正文中替换为“观测审计集上界”。
- 原始指标单位为主要结果，范围归一化损失为次要分析。
- 文档共 {len(document.paragraphs)} 个段落、{len(document.tables)} 张表、{len(document.inline_shapes)} 张嵌入主图。
- 所有 {len(document.tables)} 张表均按三线表规则生成：表顶线、表头底线、表底线，无竖线和内部网格线。
- 正文、标题、图题、表题和参考文献字体与段前段后间距统一。

## 投稿前仍需作者完成

1. 将全文翻译或定稿为 Journal of Cheminformatics 可投稿的完整英文稿，并按当前期刊模板核定摘要字数。
2. 填写并核实 CRediT authorship、Funding、Competing interests 与 Acknowledgements。
3. 建立公开仓库和永久归档 DOI，并核对所有数据许可证。
4. 若要采用更强题目/主张，必须新增真正独立锁箱和九终点预算匹配异质候选阶梯；不能用现有外层结果回填。
"""
    REPORT_PATH.write_text(report, encoding="utf-8")

    audit = {
        "source_document": str(SOURCE_DOC),
        "output_document": str(PAPER19),
        "desktop_copy": str(DESKTOP_PAPER19),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "figures": len(document.inline_shapes),
        "all_tables_three_line": True,
        "experiment_audit_passed": json.loads((EXP / "paper19_experiment_audit.json").read_text(encoding="utf-8"))["passed"],
        "independent_confirmation_completed": False,
        "third_party_reproduction_completed": False,
        "endpoint_effects": effects.to_dict(orient="records"),
        "known_placeholders": ["Authors’ contributions", "Funding", "Competing interests", "Acknowledgements"],
    }
    AUDIT_PATH.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    make_workflow_figure()
    make_multiview_figure()
    make_boundary_figure()
    stage_existing_figures()
    values = load_values()
    document = build_document(values)
    document.save(PAPER19)
    shutil.copy2(PAPER19, DESKTOP_PAPER19)
    write_report(values, document)
    print(
        json.dumps(
            {
                "paper": str(PAPER19),
                "desktop": str(DESKTOP_PAPER19),
                "paragraphs": len(document.paragraphs),
                "tables": len(document.tables),
                "figures": len(document.inline_shapes),
                "report": str(REPORT_PATH),
                "audit": str(AUDIT_PATH),
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
