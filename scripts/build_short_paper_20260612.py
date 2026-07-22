from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]


def locate_source_docx() -> Path:
    desktop = Path.home() / "Desktop"
    candidates = [
        p
        for p in desktop.rglob("*.docx")
        if p.stat().st_size > 1_000_000 and "18" in p.name and "Nature" in p.name
    ]
    if not candidates:
        raise FileNotFoundError("Could not locate the draft-18 Nature-final manuscript.")
    return max(candidates, key=lambda p: p.stat().st_mtime)


SOURCE_DOCX = locate_source_docx()
OUT_DOCX = SOURCE_DOCX.parent / "小论文-1.docx"
OUT_REPORT = ROOT / "reports" / "full_missing_experiment_run_20260611" / "short_paper_1_build_report.txt"


def set_doc_style(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)


def set_border(element, edge: str, val: str, size: str = "8", color: str = "000000") -> None:
    tag = "w:tblBorders" if element.tag.endswith("tblPr") else "w:tcBorders"
    borders = element.find(qn(tag))
    if borders is None:
        borders = OxmlElement(tag)
        element.append(borders)
    edge_el = borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        borders.append(edge_el)
    edge_el.set(qn("w:val"), val)
    if val != "nil":
        edge_el.set(qn("w:sz"), size)
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)


def apply_three_line_table(table) -> None:
    tbl_pr = table._tbl.tblPr
    for edge in ("left", "right", "insideH", "insideV"):
        set_border(tbl_pr, edge, "nil")
    set_border(tbl_pr, "top", "single", "12")
    set_border(tbl_pr, "bottom", "single", "12")
    if table.rows:
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for edge in ("left", "right", "top"):
                set_border(tc_pr, edge, "nil")
            set_border(tc_pr, "bottom", "single", "8")
        for row in table.rows[1:]:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                for edge in ("left", "right", "top", "bottom"):
                    set_border(tc_pr, edge, "nil")


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    add_caption(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    apply_three_line_table(table)


def add_picture_if_exists(doc: Document, path: Path, width: float = 6.2) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))


def pm(mean: float, sd: float, digits: int = 3) -> str:
    return f"{mean:.{digits}f} ± {sd:.{digits}f}"


def load_stress_tables() -> tuple[list[list[str]], list[list[str]]]:
    cyc = pd.read_csv(ROOT / "reports" / "bro5_cycpept_pampa_20260611" / "validation_selected_results.csv")
    cyc_rows = []
    for split in ["random", "scaffold", "perimeter", "time"]:
        g = cyc[cyc["split"].eq(split)]
        cyc_rows.append(
            [
                "CycPept-PAMPA",
                split,
                str(len(g)),
                pm(g["test_rmse"].mean(), g["test_rmse"].std()),
                pm(g["test_mae"].mean(), g["test_mae"].std()),
                f"{g['test_spearman'].mean():.3f}",
            ]
        )

    lin = pd.read_csv(
        ROOT / "reports" / "full_missing_experiment_run_20260611" / "linpept_compact_summary_20260611.csv",
        encoding="utf-8-sig",
    )
    lin_rows = []
    labels = {"linpept_cellpen": "LinPept CellPen", "linpept_nonfouling": "LinPept NonFouling"}
    for _, row in lin.iterrows():
        lin_rows.append(
            [
                labels.get(row["dataset"], row["dataset"]),
                row["split"],
                str(int(row["n_seed"])),
                row["test_ROC_AUC"],
                row["test_PR_AUC"],
                row["recall_at_precision_0.80"],
            ]
        )
    return cyc_rows, lin_rows


def selected_references() -> list[str]:
    refs_path = ROOT / "reports" / "full_missing_experiment_run_20260611" / "draft18_references.txt"
    refs = refs_path.read_text(encoding="utf-8").splitlines()
    keep = {1, 2, 3, 4, 5, 6, 8, 9, 10, 12, 18, 19, 20, 23, 24, 25, 26, 27, 28, 29, 33, 34, 35, 37, 38, 39, 41, 43, 45, 48, 49}
    out = []
    for ref in refs:
        if ref.startswith("["):
            number = int(ref.split("]", 1)[0][1:])
            if number in keep:
                out.append(ref)
    return out


def build_doc() -> None:
    cyc_rows, lin_rows = load_stress_tables()
    doc = Document()
    set_doc_style(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("FZYC-Mol：验证集治理驱动的适用域感知分子性质预测框架")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("作者信息待补充").italic = True

    doc.add_heading("摘要", level=1)
    add_paragraph(
        doc,
        "分子性质预测已成为药物发现和 ADMET 风险筛查的基础工具，但随机划分下的平均 ROC-AUC 或 RMSE 往往难以揭示新骨架外推、低相似度分子、不平衡毒性标签、bRo5 化学空间和活性悬崖中的可靠性风险。本文提出 FZYC-Mol，一种由验证集治理驱动的适用域感知多专家框架。该框架在测试集读取前冻结候选池，将强表格基线、图模型、冻结分子表征、目标变换、融合策略、适用域门控、校准、保形预测、粗糙度诊断和可解释性证据纳入同一 accept/reject/retain 流程。MoleculeNet 主面板显示，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别达到 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302；外部 TDC ADMET 面板中 22 个终点的最终保留结果为 win/tie/loss = 5/17/0。可靠性审计显示，分类错误风险分数的中位 AUROC 为 0.788，回归高误差检出中位 AUROC 为 0.652；验证-测试排名中位 Spearman 为 0.667。bRo5 与 MoleculeACE 压力测试进一步界定了外推边界。FZYC-Mol 的核心价值不是构建单一最高分模型，而是提供可冻结、可复核、保留负结果的分子性质预测证据链。"
    )
    add_paragraph(
        doc,
        "关键词：分子性质预测；ADMET；适用域；验证集治理；不确定性；活性悬崖；bRo5"
    )

    doc.add_heading("引言", level=1)
    add_paragraph(
        doc,
        "分子机器学习模型已广泛用于早期筛选、ADMET 评估和毒性预警。MoleculeNet、Therapeutics Data Commons 和近期真实挑战 ADMET benchmark 使模型比较更标准化，但也暴露出一个更基本的问题：在结构外推、少样本终点、标签噪声、不平衡分类和活性悬崖区域，平均性能分数并不足以说明模型是否可被信任。"
    )
    add_paragraph(
        doc,
        "已有研究分别从强基线模型、深度图模型、预训练表征、不确定性、保形预测、适用域和外部挑战数据集等角度推进这一问题。然而，这些证据常被分散呈现：模型候选、验证选择、测试报告、负结果和样本级失败案例之间缺少统一的治理规则。对于药物发现应用而言，更关键的问题不是某一模型是否在单个排行榜上获胜，而是候选策略何时被接受、何时应被拒绝，以及这种判断是否在测试集读取前已经固定。"
    )
    add_paragraph(
        doc,
        "本文将原完整稿件压缩为一版小论文，保留 FZYC-Mol 的核心论证：在固定候选池内，用验证集证据治理多专家模型接入，并将性能、适用域、校准、不确定性、结构外推、活性悬崖和负结果组织为同一证据链。"
    )

    doc.add_heading("结果", level=1)
    doc.add_heading("FZYC-Mol 将候选选择转化为可审计治理流程", level=2)
    add_paragraph(
        doc,
        "FZYC-Mol 将分子性质预测拆分为三个层级：单专家训练、验证集候选生成和最终冻结选择。候选专家包括 RF、ExtraTrees、XGBoost、LightGBM、CatBoost、Chemprop D-MPNN、图模型、描述符 MLP、冻结 ChemBERTa/MoLFormer 表征头、Top-K 均值、堆叠集成、自适应共识、不确定性加权、适用域门控和定向补救头。所有候选均在验证集上比较，测试集仅用于策略冻结后的最终评估。"
    )
    add_picture_if_exists(doc, ROOT / "reports" / "reviewer_revision_20260607" / "figures" / "fig1_workflow_only.png")
    add_caption(
        doc,
        "图 1. FZYC-Mol 工作流。数据划分、多视图表示、候选专家、验证集选择和可靠性输出在测试集读取前完成冻结。"
    )

    add_table(
        doc,
        "表 1. 小论文保留的核心证据层。",
        ["证据层", "核心问题", "保留数据", "主指标", "结论边界"],
        [
            ["MoleculeNet 主面板", "标准分子性质任务中是否有效", "ESOL、FreeSolv、Lipophilicity、BBBP、BACE、ClinTox", "RMSE / ROC-AUC", "报告终点依赖收益，不宣称全局最优"],
            ["外部 ADMET", "是否可迁移到官方外部划分", "TDC ADMET 与外部附录", "ROC-AUC / MAE / RMSE", "22 个终点 win/tie/loss = 5/17/0"],
            ["可靠性审计", "预测何时可信", "风险分数、保形预测、校准、适用域", "AUROC / coverage / Brier", "分类风险识别强于回归高误差检出"],
            ["结构外推", "随机划分是否高估泛化", "Random/Scaffold/Perimeter 与低相似度 bins", "性能下降、Tanimoto", "低相似度样本需单独解释"],
            ["bRo5 与活性悬崖", "困难化学空间是否暴露边界", "CycPept-PAMPA、LinPept、MoleculeACE", "RMSE / ROC-AUC / cliff RMSE", "作为压力测试，不替代盲测"],
        ],
    )

    doc.add_heading("标准任务和外部 ADMET 呈现选择性收益", level=2)
    add_paragraph(
        doc,
        "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别达到 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302。定向补救改善 Lipophilicity；FreeSolv 的 Morgan+descriptor 验证集堆叠将 RMSE 从 1.0678 ± 0.1883 降至 1.0286 ± 0.1761，但仍未超过观测最佳 Chemprop 候选，因此被保留为物理相互作用相关任务的边界案例。"
    )
    add_paragraph(
        doc,
        "外部 TDC ADMET 面板显示，多方法融合与适用域门控在 Caco2_Wang、HIA_Hou 和 Pgp_Broccatelli 等任务上产生选择性增益；22 个外部终点的最终保留结果为 win/tie/loss = 5/17/0。该结果支持的主张是“验证集治理下的选择性改进”，而非所有终点的统一提升。"
    )
    add_picture_if_exists(doc, ROOT / "reports" / "reviewer_revision_20260607" / "figures" / "fig2_selector_gate_output_evidence.png")
    add_caption(
        doc,
        "图 2. 验证集选择器、门控与证据输出。候选接入由验证集证据决定，未通过门控的候选作为负结果或补充结果保留。"
    )

    add_table(
        doc,
        "表 2. 小论文主结果与可靠性摘要。",
        ["模块", "任务/范围", "主要结果", "审计指标", "小论文解释"],
        [
            ["MoleculeNet", "ESOL / BACE / ClinTox", "0.5829 ± 0.0352 RMSE；0.8753 ± 0.0230 ROC-AUC；0.9489 ± 0.0302 ROC-AUC", "seed-level 均值±标准差", "标准任务上形成可审计主面板"],
            ["定向补救", "Lipophilicity / FreeSolv", "Lipophilicity 被验证集接受；FreeSolv 降至 1.0286 ± 0.1761 RMSE", "验证集门控", "补救是终点依赖增强"],
            ["外部 ADMET", "22 个 TDC/外部终点", "win/tie/loss = 5/17/0", "最终保留策略", "选择性收益，无负迁移"],
            ["风险分数", "分类/回归错误检出", "中位 AUROC 0.788 / 0.652", "risk-coverage / enrichment", "分类错误识别更稳定"],
            ["保形预测", "80/90/95% 覆盖", "分类 0.814/0.918/0.956；回归 0.823/0.925/0.962", "经验覆盖率", "可提供覆盖率受控输出"],
            ["选择器审计", "200 个 dataset-seed 候选池", "中位 Spearman 0.667；Top-3 hit 0.295；Top-1 match 0.135", "rank audit", "降低事后选择风险但非测试最优保证"],
        ],
    )

    doc.add_heading("可靠性分析界定预测的可用边界", level=2)
    add_paragraph(
        doc,
        "FZYC-Mol 的可靠性输出包括集成标准差、预测偏差、反向 Tanimoto 距离、重构误差、风险-覆盖曲线、保形覆盖率、校准和粗糙度代理。分类任务以预测错误为风险检出目标，回归任务以绝对误差最高的 20% 样本为高误差目标。结果显示，风险分数对分类错误的识别更稳定，而回归高误差检出仍受终点粗糙度和标签噪声限制。"
    )
    add_picture_if_exists(doc, ROOT / "reports" / "reviewer_revision_20260607" / "figures" / "fig11_risk_coverage_bbbp_clintox_caco2_pgp.png")
    add_caption(
        doc,
        "图 3. 代表性风险-覆盖曲线。低风险 retained subset 通常具有更低错误率，但回归高误差检出弱于分类错误检出。"
    )

    doc.add_heading("bRo5 与活性悬崖压力测试保留为边界证据", level=2)
    add_paragraph(
        doc,
        "bRo5 模块现包含 CycPept-PAMPA 回归任务和 LinPept CellPen/NonFouling 分类任务。CycPept-PAMPA 在 random、scaffold、perimeter 和 time split 下的 RMSE 分别为 0.547 ± 0.021、0.727 ± 0.009、0.876 ± 0.012 和 0.768 ± 0.013，显示外推难度随划分增强而上升。LinPept CellPen 在 random、scaffold 和 perimeter split 下的 ROC-AUC 分别为 0.937 ± 0.020、0.894 ± 0.029 和 0.859 ± 0.005；LinPept NonFouling 对应结果为 0.766 ± 0.012、0.765 ± 0.004 和 0.761 ± 0.000。"
    )
    add_paragraph(
        doc,
        "MoleculeACE 当前可核验结果覆盖 17 个任务和 51 个 seed 配对，验证集选择后的平均 RMSE 为 0.711，平均 cliff subset RMSE 为 0.813。gap Spearman 平均约为 0.252，部分任务接近零或为负，说明活性悬崖仍是主要失败边界。"
    )
    add_table(
        doc,
        "表 3. bRo5 公开压力测试摘要。",
        ["数据集", "split", "n_seed", "RMSE / ROC-AUC", "MAE / PR-AUC", "Spearman / Recall@P0.80"],
        cyc_rows + lin_rows,
    )

    doc.add_heading("讨论", level=1)
    add_paragraph(
        doc,
        "这版小论文将 FZYC-Mol 定位为可靠性治理框架，而不是单一追求最高分的模型。其主要贡献在于把候选池扩展、验证集选择、测试集冻结、适用域判断、风险输出和负结果保留组织到同一逻辑链中。对于药物发现用户，这种证据链比单点分数更接近实际决策需求。"
    )
    add_paragraph(
        doc,
        "结果也提示两个边界。第一，验证集治理降低了测试集事后选择风险，但验证集排名与测试集排名并不完全一致，因此小幅增益必须结合 rank audit、regret、optimism gap 和 nested validation 解释。第二，bRo5、低相似度分子、粗糙度较高的 ADME 回归和 MoleculeACE 活性悬崖仍是主要失败区域，当前结果支持风险识别和候选治理，而不支持“已解决外推问题”的结论。"
    )
    add_paragraph(
        doc,
        "与大规模 ADMET benchmark 和在线平台不同，FZYC-Mol 的目标不是替代平台化预测系统，而是在公开数据和同划分比较中给出可复核的模型接入规则。后续更适合开展盲测式外部验证、时间划分验证、更多 bRo5 终点和受控轻量适配器实验，而不是简单增加候选模型数量。"
    )

    doc.add_heading("方法概要", level=1)
    add_paragraph(
        doc,
        "每个终点被划分为训练集、验证集和测试集。训练集用于拟合单专家；验证集用于生成候选预测矩阵、融合候选、适用域门控、补救头和最终保留决策；测试集仅在策略冻结后评估一次。对任一候选 c 和终点 t，验证效用 U(c,t) 由官方主指标按方向转换为正向量，最终保留候选由 c* = argmax U(c,t) 及预定义平局、复杂度和风险规则确定。"
    )
    add_paragraph(
        doc,
        "候选池包括传统指纹/描述符树模型、图模型、Chemprop、冻结预训练表征、目标变换、Top-K 均值、堆叠集成、不确定性加权、适用域门控和定向补救头。分类任务报告 ROC-AUC、PR-AUC、Brier、ECE、MCC、固定精度召回和风险富集；回归任务报告 RMSE、MAE、Spearman、R2、高误差富集和保形区间宽度。完整 seed-level、candidate-level、metric-level 和 endpoint-level 输出保留在补充材料和 CSV 中。"
    )

    doc.add_heading("局限", level=1)
    add_paragraph(
        doc,
        "本小论文为压缩稿，保留核心证据而不展开全部候选矩阵。当前 nested validation 覆盖代表性终点而非全部任务；MoleculeACE 结果限定于当前可访问的 17 个任务；LinPept 与 CycPept-PAMPA 为公开数据压力测试，其外推范围不同于官方盲测或未见外部队列；基序归因和片段富集属于关联性解释，不能替代因果机制验证或湿实验验证。"
    )

    doc.add_heading("数据和代码可用性", level=1)
    add_paragraph(
        doc,
        "本文使用 MoleculeNet、Therapeutics Data Commons、MoleculeACE、Benchmark-ADMET-2025 及相应原始平台的公开数据。与小论文结果对应的 split seeds、候选登记表、验证/测试预测、统计检验脚本、图表 source data、环境文件和表格生成脚本将与投稿或接收版本同步存档于 GitHub/Zenodo 或期刊认可的数据仓库。当前稿件尚未分配永久 DOI；仓库冻结后应补充永久链接、版本号和 accession number。"
    )

    doc.add_heading("利益冲突与基金声明", level=1)
    add_paragraph(doc, "作者声明，目前未发现与本文研究内容直接相关的商业、财务或个人利益冲突。基金信息将在投稿版本中由作者按期刊要求核定。")

    doc.add_heading("参考文献", level=1)
    for ref in selected_references():
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        for run in p.runs:
            run.font.size = Pt(8.5)

    doc.save(OUT_DOCX)
    OUT_REPORT.write_text(
        f"Built short paper from {SOURCE_DOCX}\nOutput: {OUT_DOCX}\nFigures: workflow, selector/gate, risk coverage\nTables: evidence map, main results, bRo5 stress tests\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    build_doc()
    print(f"Wrote {OUT_DOCX}")
