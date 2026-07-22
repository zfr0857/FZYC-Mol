from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "小论文-3.docx"
OUTPUT = ROOT / "output" / "小论文-4.docx"
FIGURE = ROOT / "output" / "小论文-4_图表包" / "figures" / "fig11_reviewer_core_closure.png"


def find_prefix(doc: Document, prefix: str):
    matches = [p for p in doc.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting with {prefix!r}; found {len(matches)}")
    return matches[0]


def replace_text(paragraph, text: str) -> None:
    rpr = copy.deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs and paragraph.runs[0]._r.rPr is not None else None
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if rpr is not None:
        run._r.insert(0, rpr)


def replace_prefix(doc: Document, prefix: str, text: str) -> None:
    replace_text(find_prefix(doc, prefix), text)


def insert_after(paragraph, text: str = "", style=None):
    doc = paragraph._parent
    created = doc.add_paragraph(text, style=style or paragraph.style)
    paragraph._p.addnext(created._p)
    return created


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        from docx.oxml import OxmlElement

        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def format_table(table) -> None:
    table.style = "Normal Table"
    widths = [Cm(5.0), Cm(3.0), Cm(3.4), Cm(2.4), Cm(2.2)]
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.width = widths[col_index]
            cell.vertical_alignment = 1
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Cm(0)
                paragraph.paragraph_format.space_before = Cm(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    if row_index == 0:
                        run.bold = True
            if row_index == 0:
                shade_cell(cell, "DDEBF7")


def add_core_table(doc: Document, after_paragraph):
    caption = insert_after(after_paragraph, "表 8 | 审稿核心补实验与共享划分多视图结果", style="TableCaption")
    table = doc.add_table(rows=1, cols=5)
    headers = ["实验或对照", "效应估计", "端点聚类 95% CI", "端点方向", "精确 P"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    rows = [
        ["K=32 相对 K=4 的固定遗憾变化", "+0.122", "0.072–0.175", "8/9 增加", "0.0078；Holm 0.039"],
        ["完整多视图池：固定 Morgan RF 遗憾", "0.395", "0.253–0.545", "—", "—"],
        ["完整多视图池：one-SE 遗憾", "0.073", "0.043–0.104", "—", "—"],
        ["完整多视图池：风险调整遗憾", "0.054", "0.028–0.081", "—", "—"],
        ["完整多视图池：validation-best 遗憾", "0.043", "0.021–0.067", "—", "—"],
        ["多视图 validation-best 相对 Morgan-only 的兑现效用增益", "+0.343", "0.210–0.483", "9/9 增加", "0.0039"],
        ["拼接多视图相对分离视图池的效用增益", "+0.035", "0.017–0.053", "9/9 增加", "0.0039"],
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    format_table(table)
    caption._p.addnext(table._tbl)
    note = doc.add_paragraph(
        "注：遗憾与效用增益均以同一外层单元完整候选效用范围归一化。配对扩池检验以九个终点均值作 2⁹ 次精确符号翻转；多视图效应同样以终点为配对单位。测试 oracle 仅定义可达到上界，不参与选择。",
        style="TableNote",
    )
    table._tbl.addnext(note._p)
    return note


def add_figure(doc: Document, after_paragraph):
    drawing = doc.add_paragraph()
    drawing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    drawing.paragraph_format.first_line_indent = Cm(0)
    drawing.add_run().add_picture(str(FIGURE), width=Cm(16.2))
    after_paragraph._p.addnext(drawing._p)
    caption = doc.add_paragraph(
        "图 11 | 审稿核心闭环实验。a，九个终点中 K=32 相对 K=4 的配对固定分母遗憾变化。b，K=32 下从零信号到完全信号的正对照恢复曲线；蓝线为机会校正 Top-3 命中，橙线为固定分母遗憾。c，原等权 selection-risk score 的遗憾四分位；端点×池规模内置换未支持独立关联。d，严格嵌套留一端点元风险模型在 50% 保留覆盖下的遗憾变化，阴影为总体端点聚类 95% CI。e，共享冻结划分上 12 个多视图候选的治理规则遗憾。f，多视图可达到增益、实际选择增益及规则间配对增益；误差条为端点聚类 95% CI。",
        style="FigureCaption",
    )
    drawing._p.addnext(caption._p)
    return caption


def build() -> Path:
    doc = Document(SOURCE)
    replace_prefix(doc, "FZYC-Mol：", "FZYC-Mol：候选扩张与多视图异质性下分子性质模型选择的验证治理")
    replace_prefix(
        doc,
        "分子性质预测常在有限验证信息上比较",
        "分子性质预测需要在有限验证信息上选择不断扩张且表征异质的候选，但候选增加既提高可达到上界，也放大选择偏差。FZYC-Mol 将冻结候选登记、验证侧选择、外层审计和可靠性证据组织为可追踪治理协议。九个终点的 3 外层×3 内层×5 重复确认性实验中，K=32 相对 K=4 的配对固定分母遗憾增加 0.122（端点聚类 95% CI 0.072–0.175；精确 P=0.0078，九检验 Holm P=0.039），机会校正命中率下降 0.642（0.442–0.780）。从零信号到完全信号的 500 次/单元正对照在全部池规模上单调恢复命中率、MRR 和遗憾，零信号机会校正偏差绝对值不超过 0.004。原等权 selection-risk score 的总体相关虽为 0.235，但端点×池规模内置换 P=0.953，留一端点门控区间跨零，故降级为描述性诊断。严格嵌套的跨端点元风险模型仅使用其他端点的验证证据，在完全留出端点上将 MAE 从 0.123 降至 0.112，Spearman 为 0.313，高遗憾 AUC 为 0.648；50% 风险门控使平均遗憾降低 0.034（0.020–0.047），九个端点中八个改善。为检验表征异质性，Morgan-512、MACCS、RDKit2D 和拼接多视图分别配对线性、随机森林和 LightGBM，在相同冻结划分上完成 6,480 次拟合。完整 12 候选池的 validation-best 遗憾为 0.043（0.021–0.067），相对 Morgan-only 选择的配对效用增益为 0.343（0.210–0.483），九个端点方向一致。结果形成候选扩张、统计校准、跨端点预警和异质候选兑现的闭环；结论仍不外推至未在相同划分重训的图模型、化学语言模型或真实时间外盲测。",
    )
    contribution_updates = [
        (
            "FZYC-Mol 以冻结候选登记、随机化候选池",
            "FZYC-Mol 以冻结候选登记、随机化候选池和 3×3×5 重复嵌套审计评价模型选择；端点配对效应量、精确符号翻转和 Holm 校正确认 K=32 相对 K=4 的遗憾增加不是少数折驱动。",
        ),
        (
            "随机排序负对照在保持验证效用",
            "随机排序负对照给出零信息基线，连续信号注入正对照进一步证明机会校正命中、MRR 和固定分母遗憾能随可恢复排序信息单调响应。",
        ),
        (
            "验证侧选择风险把 one-SE",
            "风险证据被分为探索性等权分数和严格嵌套跨端点元风险：前者的负验证被完整保留，后者在完全留出端点上检验可迁移门控；共享划分多视图重训则检验治理能否兑现异质候选上界。",
        ),
    ]
    for prefix, text in contribution_updates:
        replace_prefix(doc, prefix, text)
    replace_prefix(
        doc,
        "关键词：",
        "关键词：分子性质预测；模型选择治理；候选池扩张；多视图候选；嵌套验证；跨端点风险；信号恢复；固定分母遗憾",
    )

    method_anchor = find_prefix(doc, "治理规则迁移采用 leave-one-endpoint-out")
    m1 = insert_after(method_anchor, "2.5.2 配对推断与信号恢复正对照")
    m2 = insert_after(
        m1,
        "扩池效应以同一终点-重复-外层折的 K=4 与 K=8/16/32 结果配对。先在终点内对 15 个外层单元求平均，再对九个终点执行全部 2⁹ 种符号翻转，获得双侧精确 P；固定遗憾、机会校正命中和 MRR 的三种池比较共九项检验，使用 Holm 法控制家族错误率。区间由 10,000 次端点聚类 bootstrap 给出。",
    )
    m3 = insert_after(
        m2,
        "正对照在每个冻结外层单元的真实候选效用分布上构造合成验证分数 s·z(u_outer)+√(1−s²)·ε，其中 s=0/0.10/0.25/0.50/0.75/1.00，ε 为标准正态噪声。除 s=1 外，每个单元和池规模重复 500 次。该模拟只验证指标的零点与检出能力，不作为模型性能证据。",
    )
    m4 = insert_after(m3, "2.5.3 严格嵌套的跨端点元风险")
    m5 = insert_after(
        m4,
        "每个选择单元提取 16 个不读取外层标签的验证证据，包括池规模、任务类型、候选验证效用离散度、前两名标准化间隔、最佳候选与总体内折波动、折间候选排序一致性、折赢家频率与家族数、one-SE 大小及三项原风险分量。最外层每次完整留出一个终点；在其余八个终点内部再次 leave-one-endpoint-out，按 MAE 在固定 Ridge、浅层随机森林和强正则 HistGradientBoosting 中选模，再拟合八个终点并预测完全留出的终点。",
    )
    m6 = insert_after(
        m5,
        "元风险以留出端点 MAE、Spearman、高遗憾 AUC 和 50% 风险门控评价。高遗憾阈值、模型类型和全部参数均只由训练端点确定；门控效应以留出端点内 retained−all 遗憾定义，并用端点聚类 bootstrap 汇总。原等权分数同时进行端点×池规模内置换，以区分跨层级相关与层内预警能力。",
    )
    m7 = insert_after(m6, "2.5.4 共享冻结划分的多视图候选确认")
    insert_after(
        m7,
        "Morgan-512、MACCS、RDKit2D 与三者拼接多视图分别配对线性模型、80 树随机森林和 80 轮 LightGBM，形成 12 个预登记候选。九个终点使用与主确认实验相同的 11/23/37/53/71 种子、3 外层和 3 内层划分；缺失值插补和标准化仅在各训练折拟合，数值不稳定的 RDKit Ipc 在查看标签前固定排除。每个候选完成 3 个内折拟合和一次外层重拟合，共 12×9×5×3×4=6,480 次拟合。比较完整池、Morgan-only、双指纹池和不含拼接视图池，并分别评价 fixed Morgan RF、validation-best、one-SE、风险调整和测试 oracle。",
    )

    negative_anchor = find_prefix(doc, "真实流程的固定分母遗憾为 0.049")
    n1 = insert_after(
        negative_anchor,
        "端点配对推断确认该变化不是独立折伪重复造成。K=32 相对 K=4 的平均固定分母遗憾增加 0.122（端点聚类 95% CI 0.072–0.175），九个终点中八个方向为正，精确 P=0.0078；九项池规模×指标检验经 Holm 校正后 P=0.039。机会校正命中率的配对变化为 −0.642（−0.780 至 −0.442），同样有八个终点下降。",
    )
    insert_after(
        n1,
        "信号恢复正对照给出互补验证。K=4/8/16/32 的机会校正命中率和 MRR 均随注入信号单调增加，固定分母遗憾均单调降低；零信号时各 K 的机会校正命中绝对偏差不超过 0.004，完全信号时遗憾为 0（图 11b）。因此，主指标既具有正确零点，也能恢复逐级增强的排序信息。",
    )

    replace_prefix(
        doc,
        "验证侧选择风险在 540 个外层单元中",
        "原等权 selection-risk score 在 540 个单元中的总体 Spearman 为 0.235（端点 bootstrap 95% CI 0.019–0.378），且遗憾由最低风险四分位的 0.076 增至最高四分位的 0.169。然而，在端点×池规模内置换后单侧 P=0.953，说明总体相关主要包含端点和 K 的层级差异，不能据此主张层内前瞻门控。",
    )
    replace_prefix(
        doc,
        "按选择风险中位数二分时",
        "严格留一端点时，等权分数的校准 MAE 为 0.120，常数基线为 0.123，Spearman 为 0.099，高遗憾 AUC 为 0.514；以其他端点中位风险门控的平均遗憾变化为 −0.009（95% CI −0.038 至 0.020），仅 5/9 个端点改善。因此该分数降级为描述性诊断，不再作为已验证门控。",
    )
    risk_anchor = find_prefix(doc, "严格留一端点时，等权分数")
    insert_after(
        risk_anchor,
        "嵌套跨端点元风险在完全留出端点上给出更强但仍有限的前瞻证据：MAE 为 0.112，低于常数基线 0.123；Spearman 为 0.313，高遗憾 AUC 为 0.648。按预测风险保留最低 50% 单元时，平均遗憾降低 0.034（端点聚类 95% CI 0.020–0.047），8/9 个端点改善（图 11d）。九次外层留出中，内层模型选择六次选 Ridge、三次选随机森林。",
    )

    caption10 = find_prefix(doc, "图 10 |")
    section = insert_after(caption10, "3.8 共享冻结划分的多视图候选确认", style="Heading 2")
    s1 = insert_after(
        section,
        "九个终点、五个重复和三个外层折均完成 12 候选共享划分重训，共 135 个外层单元、6,480 次拟合，无缺失候选或折失败。完整池 validation-best 的平均标准化遗憾为 0.043（端点聚类 95% CI 0.021–0.067），低于风险调整的 0.054（0.028–0.081）、one-SE 的 0.073（0.043–0.104）和固定 Morgan RF 的 0.395（0.253–0.545）。",
        style="Normal",
    )
    s2 = insert_after(
        s1,
        "完整多视图 test oracle 相对 Morgan-only oracle 的可达到效用增益为 0.347（0.228–0.472），九个终点方向一致；validation-best 实际兑现的配对增益为 0.343（0.210–0.483），同样为 9/9 个终点改善，精确 P=0.0039。两者接近说明验证选择器在该轻量异质池中兑现了大部分可达到表征增益。",
        style="Normal",
    )
    s3 = insert_after(
        s2,
        "拼接多视图相对仅允许各视图独立候选的验证选择增益为 0.035（0.017–0.053；9/9 端点，P=0.0039）。135 个 validation-best 决策中，拼接多视图被选择 84 次，RDKit2D 44 次，MACCS 4 次，Morgan-512 3 次；表征优势并不等价于某个固定模型普适占优。",
        style="Normal",
    )
    s4 = insert_after(
        s3,
        "在完整池中，validation-best 相对 one-SE 的配对效用增益为 0.030（0.014–0.048；8/9 端点，P=0.0117）；相对风险调整为 0.011（0.003–0.021），但端点符号翻转 P=0.0625。后者应解释为小幅均值优势而非确认性普适差异。",
        style="Normal",
    )
    note = add_core_table(doc, s4)
    add_figure(doc, note)

    replace_prefix(doc, "3.8 失败案例", "3.9 失败案例、片段统计与证据完整性")
    replace_prefix(doc, "3.9 Source data", "3.10 Source data 自动重建")
    replace_prefix(doc, "表 8 | 代表性失败", "表 9 | 代表性失败案例及其可审计风险信号")
    replace_prefix(doc, "表 9 | 近期相关", "表 10 | 近期相关研究与四项创新主张的边界")
    replace_prefix(
        doc,
        "证据一致性审计确认，重复嵌套 540 行结果",
        "证据一致性审计确认，主轻量池 540 个选择单元、12 候选多视图池 135 个外层单元和 6,480 次拟合、三种候选池控制、5,000 次/分量置换与 bootstrap、TDC 22 终点、85 个风险单元、90 个保形单元及 AutoGluon 三预算均可由结构化 CSV 重建。异质多视图确认已完成；图神经网络、化学语言模型的共享外层重训和独立 ADMET 时间外盲测仍未完成。",
    )
    replace_prefix(
        doc,
        "主文数字、表和图 1–10",
        "主文数字、表和图 1–11 均由冻结 source data 或概念图脚本重建；新增端点配对推断、信号恢复、严格嵌套元风险和共享划分多视图实验分别保存逐单元、逐端点、逐候选和逐策略结果。完整文件清单、SHA-256、冷启动日志和发布缺口见补充复现材料。",
    )

    replace_prefix(
        doc,
        "随机化候选池、置换负对照与 3×3×5",
        "随机化候选池、置换负对照、端点配对精确检验与 3×3×5 重复嵌套验证共同表明，候选池扩张削弱选择可信度，但并未把真实排序降至随机水平。K=32 相对 K=4 的遗憾增加在八个终点方向一致并经多重校正保留；连续信号恢复又证明指标能从正确零点单调响应。因此核心结论不依赖独立折伪重复、固定 Top-3 的机械机会变化或单一指标。",
    )
    replace_prefix(
        doc,
        "候选规模与候选组成共同影响",
        "候选规模、模型家族和分子表征共同影响选择难度。共享划分多视图实验把历史异质候选的不可比性替换为真正同折重训：可达到上界与 validation-best 实际增益在九个终点上方向一致，说明治理可以兑现轻量多视图增益；但这仍不等同于图模型或化学语言模型已被确认。",
    )
    replace_prefix(
        doc,
        "没有一种治理规则在所有终点占优",
        "治理规则的结论变得更具体。32 候选 Morgan 池的 LOEO 仍显示 FreeSolv 与 HIA 的端点特异偏好；12 候选多视图池中 validation-best 相对 one-SE 的配对增益得到支持，而相对风险调整的端点符号检验未跨越 0.05。规则冻结的价值是限制事后自由度并保存失败，而不是预设保守规则必然更优。",
    )
    replace_prefix(
        doc,
        "可靠性结果需要区分预测风险与选择风险",
        "可靠性结果需要区分预测风险、描述性选择风险和可迁移元风险。原等权 selection-risk score 在总体四分位上呈剂量趋势，却未通过层内置换和留一端点门控，因此已明确降级；嵌套元风险只使用其他端点训练，获得中等区分和 8/9 端点门控改善，但九端点样本仍不足以声称普适部署。逐样本 E-AURC、适用域和标签条件保形则回答不同统计单位上的预测可靠性。",
    )
    replace_prefix(
        doc,
        "本研究的首要范围限制",
        "本研究的首要范围限制已从单一 Morgan-512 表征缩小为轻量候选范围：确认性证据现覆盖 Morgan、MACCS、RDKit2D 和拼接多视图及三类传统学习器，但尚未在相同外层划分统一重训 Chemprop、GNN、ChemBERTa 和 MoLFormer。其次，元风险只覆盖九个端点，且 P-gp 是唯一未从 50% 门控获益的端点；回归任务的 GroupKFold 在五个种子下保持相同骨架组，仅模型随机性改变。此外，仍无独立 ADMET 时间外盲测。",
    )
    replace_prefix(
        doc,
        "综上，候选增加可提高",
        "综上，候选增加可以提高可达到上界并在多视图共享划分中产生可兑现增益，也会增加验证侧选择自由度。审稿充分的证据链应同时包含冻结登记、端点配对推断、零/正对照、严格留出风险验证、同折异质候选重训和完整负结果；仅报告最优测试分数或总体风险相关均不足以证明选择可靠。",
    )

    replace_prefix(
        doc,
        "FZYC-Mol 将冻结候选登记、随机化候选池",
        "FZYC-Mol 将冻结候选登记、随机化候选池、端点配对推断、零/正对照、跨端点元风险和外层审计组织为可复核的模型选择治理流程。重复嵌套与精确检验确认扩池增加遗憾；信号恢复证明指标可校准；共享划分多视图重训则表明 validation-best 能兑现轻量异质候选的可达到增益。",
    )
    replace_prefix(
        doc,
        "验证侧选择风险能够在外层标签读取前",
        "原等权选择风险只保留为描述性信号；严格嵌套元风险在完全留出端点上取得 MAE 0.112、AUC 0.648，并在 8/9 个端点降低 50% 门控遗憾。治理规则仍具有端点依赖性，validation-best 相对风险调整的多视图配对优势也未通过端点符号 0.05 阈值。",
    )
    replace_prefix(
        doc,
        "最稳健的结论是：",
        "最稳健的结论是：候选扩张同时增加可达到上界和选择不确定性；冻结治理、配对推断、负/正对照、严格留出风险验证及同折异质候选重训可以暴露并部分管理这种风险，但不能保证普遍性能提升。结论限于公开离线基准和轻量多视图候选，不支持对未重训深度模型、时间外 ADMET 或临床部署外推。",
    )

    replace_prefix(
        doc,
        "补充材料逐项提供",
        "补充材料逐项提供：数据清洗与候选登记；100 次随机候选控制；3×3×5 重复嵌套逐折结果；端点配对精确检验与 Holm 校正；1,000 次随机排序负对照；六级信号恢复正对照；原风险分数的层内置换与 LOEO 负验证；严格嵌套元风险；12 候选共享划分多视图的 6,480 次拟合；TDC、风险-覆盖、保形预测、MoleculeACE、bRo5 和完整负结果。",
    )
    replace_prefix(
        doc,
        "图 1–10 均配套",
        "图 1–11 均配套 PNG/SVG、source data 和生成脚本。reviewer_core_values.json、cross_endpoint_meta_risk_summary.json 与 multiview_values.json 汇总新增结果；selection_closure_values.json 和 manuscript_values.json 继续承担既有结果的零差异重建。公开仓库 release、第三方空环境复跑和 Zenodo 归档仍待作者完成。",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
