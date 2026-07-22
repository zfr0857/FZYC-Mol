from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
TABLE_DIR = ROOT / "reports" / "manuscript_tables"
BASE_DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604.docx"
BASE_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604.md"
OUT_DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_normative.docx"
OUT_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604_normative.md"
REPORT = DOCS / "manuscript_frontmatter_methods_normative_review_20260604.md"
TABLE51 = TABLE_DIR / "table51_recent_paper_writing_alignment.csv"
TABLE52 = TABLE_DIR / "table52_frontmatter_methods_revision_matrix.csv"


ABSTRACT_P1 = (
    "分子性质预测是药物发现、ADMET 评估和毒性风险筛查中的基础计算环节。"
    "近几个月的 ADMET reliability benchmark、OpenADMET avoid-ome 研究、tabular foundation model 评估以及多尺度分子表示论文共同显示，"
    "当前领域的核心问题已经从单一 leaderboard 分数转向更真实的模型使用边界：结构外推、数据稀缺、类别不平衡、activity cliffs、"
    "beyond-rule-of-5 化学空间、实验标签噪声、概率校准和适用域漂移。"
    "因此，一篇面向高水平投稿的分子性质预测论文不仅需要报告 ROC-AUC 或 RMSE，还应说明模型在何种 split、何种 endpoint、何种可靠性条件下可被信任。"
)

ABSTRACT_P2 = (
    "本文提出 FZYC-Mol，一个 validation-governed 多专家分子性质预测框架。"
    "框架将实验流程显式拆分为四层：多视图分子表示、候选专家训练、验证集策略选择和测试集冻结报告。"
    "候选池覆盖 Morgan/MACCS/atom-pair/torsion 指纹、RDKit 2D 描述符、BRICS/Murcko/官能团 motif、"
    "LightGBM/XGBoost/CatBoost/RF/ExtraTrees、Chemprop D-MPNN、图模型、冻结 ChemBERTa/MoLFormer embedding heads、"
    "Top-K mean、ridge/logistic stacking、target transform、balanced undersampling ensemble 以及适用域/不确定性证据头。"
    "所有最终 retained-best 策略均只由 validation split 决定，test split 仅在策略冻结后用于一次性报告。"
)

ABSTRACT_P3 = (
    "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别达到 RMSE 0.5829 +/- 0.0352、"
    "ROC-AUC 0.8753 +/- 0.0230 和 ROC-AUC 0.9489 +/- 0.0302。"
    "validation-only rescue heads 使 Lipo 的 RMSE 从 0.7078 +/- 0.0389 改进到 0.6835 +/- 0.0439；"
    "低成本 targeted rebuild 使 FreeSolv 的 retained-best RMSE 从 1.0678 +/- 0.1883 改进到 1.0286 +/- 0.1761。"
    "进一步引入层级 motif/fingerprint 表示、冻结分子语言 embedding、多视图融合、AD gating 与不确定性加权后，"
    "BBBP 和 ClinTox 获得小幅 retained-best 改善；外部 TDC ADMET official panel 中 Caco2_Wang、HIA_Hou 和 Pgp_Broccatelli 也被 validation gate 接受。"
)

ABSTRACT_P4 = (
    "可靠性分析显示，ensemble std、error model、hybrid error-AD、Tanimoto 适用域和 reconstruction error 能够富集高误差样本；"
    "motif attribution 与 fragment enrichment 为 BBBP、BACE、ClinTox 和 ADME 回归任务提供了可读的化学解释。"
    "同时，本文保留未被 validation 接受的 3D-lite、roughness-weighted regression 和 TabPFNv2 未授权 pilot 作为负结果或待复现实验，"
    "以明确方法边界。整体而言，FZYC-Mol 的贡献不是提出单一更大的 backbone，而是提供一套可复现、可审计、适用域感知的多专家选择范式。"
)


REPLACEMENTS = {
    2: ABSTRACT_P1,
    3: ABSTRACT_P2,
    4: ABSTRACT_P3,
    7: (
        "分子机器学习已经从 QSAR、ECFP 指纹和树模型扩展到图神经网络、D-MPNN、分子语言模型和跨任务 foundation model。"
        "然而，近期 benchmark 的共同结论是：药物发现中的实际预测难度不能由随机划分下的平均 ROC-AUC 或 RMSE 充分刻画。"
        "当评价场景转向 scaffold expansion、低相似度分子、少样本 endpoint、类别不平衡毒性任务和 activity-cliff 区域时，"
        "许多复杂模型与传统 descriptor/fingerprint 模型之间的优势会发生变化。因此，模型比较必须同时报告 split realism、endpoint 异质性和可靠性证据。"
    ),
    8: (
        "近几个月的相关工作为本文写法提供了清晰参照。ADMET reliability benchmark 将数据稀缺、OOD、类别不平衡、bRo5 和 activity cliffs 作为核心挑战，"
        "并指出 TabPFNv2、undersampling ensemble 和 roughness index 在不同场景下各有价值。OpenADMET avoid-ome 研究强调，ADMET 失败来自一组共享的机制性反靶标和性质空间，"
        "需要开放数据、blind challenge 和适用域问题的系统研究。Tabular foundation model 论文显示，RDKit/Mordred descriptors 与冻结 foundation embedding 结合 TabPFN 可以成为低成本强对照。"
        "HimNet、LGSM 和 CHAMP 等多尺度模型则强调 atom、motif、fingerprint、3D/geometric 和 global context 的层级融合与化学解释。"
    ),
    9: (
        "这些研究共同提示：若本文只把 FZYC-Mol 写成一个新的 ensemble 模型，会削弱其真正价值。"
        "更规范的定位应是 validation-governed model selection framework：在预定义候选池中比较紧凑 tabular 模型、图模型、冻结 embedding、target transform 和 calibration/reliability heads，"
        "并用同一套验证集规则决定是否接入主结果。这样既能回应“为什么不直接训练更大模型”的问题，也能解释为什么某些 pilot 结果应保留在 appendix 而不覆盖 retained-best。"
    ),
    10: (
        "本文贡献可概括为五点。第一，构建覆盖 MoleculeNet、TDC ADMET、MoleculeACE、OpenADMET appendix 和 structure-separated split 的统一实验包。"
        "第二，提出 validation-only selector、rescue-integrated selector 和 retained-best gate，使新增候选必须先通过验证集证据。"
        "第三，补齐 CatBoost/XGBoost/LightGBM/ExtraTrees/RF、Chemprop D-MPNN、冻结 ChemBERTa/MoLFormer、TabPFNv2 代码通道、Top-K/stacking 和 target-transform 等同类型论文会期待的强对照。"
        "第四，加入 uncertainty、applicability domain、risk-coverage、calibration、conformal、roughness 和 low-similarity hard subset 等可靠性分析。"
        "第五，通过 motif attribution、fragment enrichment 和 case-study 设计，把模型性能与可解释性边界连接起来。"
    ),
    13: (
        "图 1A. FZYC-Mol 整体流程图。该图按近期多模块分子预测论文的规范，将输入分子、数据划分、多视图表示、候选专家池、validation-only selector、retained-best gate 和可靠性/解释性输出放在同一条闭环流程中；图中不把候选专家误写为最终模型，而是强调验证集治理。"
    ),
    15: (
        "图 1B. FZYC-Mol 模型结构图。结构图展示 atom/bond/motif/fingerprint/descriptor/embedding 多源表示如何形成候选专家，以及预测矩阵如何进入 Top-K、stacking、risk-aware policy 和 retained-best gate；该图用于解释方法边界，而不是宣称复现某一篇深度模型的完整架构。"
    ),
    16: (
        "图 1A 强调实验工作流，图 1B 强调模型内部结构。二者共同说明本研究的主体不是单一 backbone，而是一个由多视图表示、异质专家、验证集治理、适用域判断和可解释性证据组成的预测框架。"
    ),
    20: (
        "所有 endpoint 均显式区分 train、validation 和 test。训练阶段只使用 train；候选专家、target transform、Top-K/stacking、rescue heads 和 selector policy 的选择只使用 validation；test 仅在策略冻结后评估一次。"
        "这一规则是全文最重要的方法学边界，因为候选池不断扩展时，若允许根据 test 结果回改候选或 tie-breaker，就会产生隐性 test leakage。"
    ),
    21: (
        "数据划分按问题难度分层报告。MoleculeNet 以 scaffold split 为主，并补充 random split、structure-separated split 和 low-similarity hard subset；"
        "TDC ADMET 使用官方 PyTDC split 和 full-panel scaffold appendix；MoleculeACE 用于 activity-cliff 检验；OpenADMET-ExpansionRx 作为快速外部 feasibility check。"
        "这种组织方式与近期 benchmark 对 split realism 和 blind challenge 的强调一致，避免只在最容易的随机划分上展示模型优势。"
    ),
    23: (
        "FZYC-Mol 的输入表示分为五类。第一类是分子图表示，用于 GNN、D-MPNN 和 Chemprop 类模型；第二类是 Morgan、MACCS、atom-pair 与 torsion 指纹，用于传统树模型和多指纹专家；"
        "第三类是 RDKit 2D 描述符，用于 fast tabular baseline、TabPFN 代码通道和 target-transform regression；第四类是 BRICS、Murcko scaffold 和官能团 motif，用于可解释专家与 fragment enrichment；"
        "第五类是冻结 ChemBERTa 与 MoLFormer embedding，用于低成本测试预训练表征是否能补充低分 endpoint。"
    ),
    24: (
        "候选专家池包括 RF、ExtraTrees、XGBoost、LightGBM、CatBoost、Chemprop D-MPNN、图模型、描述符 MLP、motif experts、冻结预训练 embedding heads、"
        "Top-K mean、ridge/logistic stacking、adaptive consensus 和 targeted rescue heads。回归任务额外包含 identity、log1p、quantile-normal、winsorized target 等目标变换；"
        "类别不平衡任务额外包含 balanced undersampling ensemble 和 PR-AUC/Brier/ECE/MCC/balanced accuracy 评估。TabPFNv2 已接入代码路径，但在本机授权/缓存未准备前不报告数值。"
    ),
    26: (
        "selector 首先在每个 endpoint 的 validation 预测表上计算官方主指标，并按指标方向排序。若单一专家显著占优，则保留 best expert；若多个专家在 validation 上互补，则候选 Top-K mean、stacking 或 uncertainty-weighted fusion；"
        "若新增 rescue head、target transform 或 strong baseline 未超过当前 retained-best，则进入附录而不改变主结果。该规则使 positive result、negative result 和 pending result 都有明确位置。"
    ),
    28: (
        "可靠性模块包括 ensemble standard deviation、error model、prediction deviation、inverse Tanimoto distance、reconstruction error、hybrid error-AD、risk-coverage curve、conformal coverage、calibration 和 roughness proxy。"
        "可解释性模块包括 motif feature importance、fragment enrichment、scaffold/neighbor case review 与高误差样本分析。本文只把这些证据解释为模型使用边界和化学相关性，不把关联性 motif 解释等同于因果机制。"
    ),
    30: (
        "为保证多专家比较可复现，本文将候选模型治理拆成三个层级。第一层是单一专家训练：每个模型家族在固定 split 和 seed 下独立训练并保存 validation/test 预测。"
        "第二层是候选策略生成：仅在 validation 预测矩阵上构造 Top-K mean、stacking、adaptive consensus、uncertainty-aware weighting、AD gating 和 targeted rescue。"
        "第三层是最终 selector 决策：只根据 validation 主指标和预先定义的 tie-breaking/risk-adjusted 规则确定 retained-best。"
    ),
    31: (
        "回归任务的主指标为 RMSE、MAE 或 Spearman。RMSE/MAE 按低值优先，Spearman 按高值优先。考虑到 clearance、half-life、ppbr、VDss 和 Caco2 等 ADME endpoint 常具有长尾分布、测量误差和局部 target jump，"
        "本文将 log1p、quantile-normal、winsorized target 和 robust loss 作为候选，而不是默认策略。只有当这些变换在 validation 上优于当前候选时，才允许进入 retained-best。"
    ),
    32: (
        "分类任务保留 ROC-AUC 作为 MoleculeNet 和若干 TDC endpoint 的主指标，同时报告 PR-AUC、Brier score、ECE、MCC、balanced accuracy、risk-coverage 和 conformal coverage。"
        "这种写法是为了回应类别不平衡任务中的常见审稿质疑：ROC-AUC 高并不一定意味着阳性样本召回、概率校准或筛选富集可靠。ClinTox、DILI、hERG 和 CYP substrate 因此被单独纳入不平衡分类补强。"
    ),
    33: (
        "stacking 被定义为 validation-only prediction-level candidate，而不是新的原始特征学习模型。回归任务使用 ridge-style 线性组合或均值，分类任务使用 logistic stacking 或概率均值。"
        "为降低小样本过拟合，stacking 只使用 Top-3 或 Top-5 validation experts，并在 seed-level 上汇总；若 validation 上不稳定或不能超过当前 selector，则保留原策略。"
    ),
    34: (
        "adaptive consensus 用于处理专家之间没有单一绝对赢家的 endpoint。它先识别 validation 上接近最优的候选家族，再根据任务类型、指标方向和专家相关性构造加权平均。"
        "与简单堆叠所有模型不同，adaptive consensus 的目标是避免低质量或高度冗余专家稀释强模型信号。"
    ),
    35: (
        "targeted rescue heads 服务于已知低分或高 roughness 模块，例如 Lipo、FreeSolv、clearance、ppbr、half-life 和 CYP substrate。"
        "rescue heads 可以来自冻结 embedding、强 tabular baseline、Top-K/stacking、target transform 或 undersampling ensemble，但必须先进入 selector pool，再由 validation 决定是否保留。"
        "本文结果显示，MoleculeNet 中真正被接受的 rescue 主要发生在 Lipo；这一负选择同样重要，因为它说明补强模块没有被强行用于所有 endpoint。"
    ),
    37: (
        "适用域分析使用两类信号。第一类是化学空间外推信号，包括测试分子到 train+valid 最近邻的 Morgan Tanimoto 距离、scaffold distance 和 low-similarity subset；"
        "第二类是模型行为风险信号，包括 ensemble std、prediction deviation、error model 和 reconstruction error。两类信号互补：结构相似的分子仍可能因标签噪声或 activity cliff 出错，结构新颖的分子也可能因模型一致性较高而风险较低。"
    ),
    38: (
        "roughness proxy 用于解释 endpoint 的局部结构-性质关系是否平滑。若测试分子与最近邻高度相似但标签差异或归一化 target jump 很大，则说明该 endpoint 具有高 roughness，"
        "仅增大模型容量不一定有效，反而需要 target transform、robust loss、局部邻域诊断或 ensemble 稳定化。本文不把 roughness 作为主性能指标，而是作为低分 ADME regression 和 activity-cliff 任务的解释性证据。"
    ),
    39: (
        "motif attribution 与 fragment enrichment 提供互补解释。前者面向模型内部特征重要性，关注官能团、BRICS 片段和 Murcko scaffold 对预测的贡献方向；"
        "后者面向数据分布，关注某些片段出现时标签均值相对 baseline 的偏移。两者一致时，说明模型关注点与数据统计信号相互支持；不一致时，则提示 confounding、样本量不足或局部标签噪声。"
    ),
    40: (
        "图表呈现遵循“主文总结、附录追溯”的原则。主文只保留可读的 summary figure 和 compact table；完整 seed-level、candidate-level、metric-level 和 endpoint-level CSV 进入 supplementary package。"
        "因此，正文图表说明需要明确每张图回答的问题、指标方向、是否为 retained-best，以及未接入候选为何作为负结果保留。"
    ),
}


def set_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(10.5)


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    set_text(new_para, text)
    return new_para


def replace_docx() -> None:
    shutil.copy2(BASE_DOCX, OUT_DOCX)
    doc = Document(OUT_DOCX)
    for idx, text in REPLACEMENTS.items():
        set_text(doc.paragraphs[idx], text)
    insert_after(doc.paragraphs[4], ABSTRACT_P4)
    doc.save(OUT_DOCX)


def replace_markdown() -> None:
    text = BASE_MD.read_text(encoding="utf-8")
    doc = Document(BASE_DOCX)
    pairs = [(doc.paragraphs[idx].text, new) for idx, new in REPLACEMENTS.items()]
    for old, new in pairs:
        if old:
            text = text.replace(old, new)
    text = text.replace(ABSTRACT_P3, ABSTRACT_P3 + "\n\n" + ABSTRACT_P4)
    OUT_MD.write_text(text, encoding="utf-8")


def build_tables() -> None:
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    alignment = pd.DataFrame(
        [
            {
                "recent_paper_signal": "ADMET reliability benchmark, Journal of Cheminformatics 2026",
                "writing_norm": "把 OOD、少样本、类别不平衡、bRo5、activity cliffs 和 roughness 写成问题定义，而不是结果后解释。",
                "revision_action": "摘要和引言首段改为 real-world reliability framing；方法中新增 roughness/imbalance 指标边界。",
            },
            {
                "recent_paper_signal": "OpenADMET avoid-ome, Nature Communications 2026",
                "writing_norm": "强调 ADMET 是共享机制性失败空间，突出适用域、开放数据、blind challenge 和模型使用边界。",
                "revision_action": "引言增加 avoid-ome/外部 appendix 叙事，方法中强化 AD/OOD 与 retained-best gate。",
            },
            {
                "recent_paper_signal": "Tabular foundation models for molecular properties, arXiv 2026",
                "writing_norm": "把 TabPFN + descriptors/embeddings 写作低成本强 baseline，而不是和大模型 fine-tuning 混在一起。",
                "revision_action": "摘要和方法中明确 TabPFNv2 已接入代码路径但未授权前不报告数值。",
            },
            {
                "recent_paper_signal": "Do Larger Models Really Win in Drug Discovery?, arXiv 2026",
                "writing_norm": "强调模型尺度不是唯一决定因素，split 与任务类型决定模型家族优势。",
                "revision_action": "引言把本研究定位为 validation-governed model selection framework，而非单一更大模型。",
            },
            {
                "recent_paper_signal": "HimNet / LGSM / CHAMP 2026",
                "writing_norm": "图和方法描述应突出 atom-motif-global 多尺度表示、跨视图融合与可解释性验证。",
                "revision_action": "图 1A/1B caption 和方法 2.2 改为多层表示+候选专家+验证治理的结构化描述。",
            },
        ]
    )
    matrix = pd.DataFrame(
        [
            {"section": "摘要", "issue_before": "偏实验汇总，结果段过长。", "revision": "改成背景/方法/结果/结论边界四段式。"},
            {"section": "引言", "issue_before": "文献信号较多但逻辑层次略散。", "revision": "按 benchmark challenge -> recent consensus -> gap -> contributions 组织。"},
            {"section": "方法 2.1", "issue_before": "validation-only 原则已有，但与 test leakage 风险关系可更明确。", "revision": "明确 train/validation/test 三层边界和候选扩展时的泄漏风险。"},
            {"section": "方法 2.2", "issue_before": "专家池列表完整但缺少分类层级。", "revision": "按图、指纹、描述符、motif、冻结 embedding 五类表示重写。"},
            {"section": "方法 2.3-2.5", "issue_before": "selector、stacking、rescue、target transform 混在一起。", "revision": "拆成 candidate generation、validation gate、retained-best rule 和 appendix negative result。"},
            {"section": "方法 2.6", "issue_before": "可靠性/解释性定义充分，但与近期文献 framing 可更直接。", "revision": "把 AD、roughness、calibration、motif/fragment 的用途写成模型使用边界。"},
            {"section": "图表", "issue_before": "部分图表说明像文件索引。", "revision": "caption 增加回答的问题、指标方向和是否影响 retained-best。"},
        ]
    )
    alignment.to_csv(TABLE51, index=False)
    matrix.to_csv(TABLE52, index=False)


def write_report() -> None:
    lines = [
        "# 中文初稿前文与方法学规范化审查（2026-06-04）",
        "",
        "本次审查参考近几个月 ADMET reliability、OpenADMET avoid-ome、Tabular foundation model、模型尺度 benchmark、HimNet/LGSM/CHAMP 等相关论文的写法。核心结论是：当前稿件内容已经很厚，但摘要、引言和方法部分需要从“实验记录式汇总”改成“问题定义-方法边界-验证治理-结果证据”的论文式结构。",
        "",
        "## 已完成修订",
        "",
        "- 摘要改为四段式：背景与缺口、方法框架、主要结果、可靠性与边界。",
        "- 引言改为四步逻辑：领域背景、近期共识、本文定位、贡献列表。",
        "- 方法 2.1-2.6 重写为更标准的实验协议：数据划分、表示层、候选池、validation gate、retained-best rule、可靠性与解释性边界。",
        "- 图 1A/1B caption 改为说明图回答的问题和方法边界，而不是只描述画面元素。",
        "- 新增两张审查表：Table S46 最近论文写法对齐；Table S47 前文/方法修订矩阵。",
        "",
        "## 输出文件",
        "",
        f"- 修订 docx：`{OUT_DOCX.relative_to(ROOT)}`",
        f"- 修订 markdown：`{OUT_MD.relative_to(ROOT)}`",
        f"- 写法对齐表：`{TABLE51.relative_to(ROOT)}`",
        f"- 修订矩阵：`{TABLE52.relative_to(ROOT)}`",
        "",
    ]
    REPORT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    replace_docx()
    replace_markdown()
    build_tables()
    write_report()
    print(f"wrote {OUT_DOCX.relative_to(ROOT)}")
    print(f"wrote {OUT_MD.relative_to(ROOT)}")
    print(f"wrote {TABLE51.relative_to(ROOT)}")
    print(f"wrote {TABLE52.relative_to(ROOT)}")
    print(f"wrote {REPORT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
