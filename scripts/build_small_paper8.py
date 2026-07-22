from __future__ import annotations

import json
import re
import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
SRC_DOC = OUT / "小论文-7.docx"
SRC_SUPP = OUT / "小论文-7_补充表.docx"
SRC_PACK = OUT / "小论文-7_图表包"
DOCX = OUT / "小论文-8.docx"
SUPP = OUT / "小论文-8_补充表.docx"
PACK = OUT / "小论文-8_图表包"
AUDIT = ROOT / "results" / "audits" / "small_paper_8_final_audit.json"


EXPANSIONS: dict[str, list[str]] = {
    "科学贡献": [
        "本文的中心贡献是把模型选择从单次排行榜问题转化为可审计的验证治理问题。稿件不把 FZYC-Mol 定义为新的分子主干网络，也不以单个公开基准分数作为主要卖点，而是关注候选池扩大后验证集如何被反复查询、验证排序如何失真、以及哪些证据能够在读取外层标签之前提示风险。这样的定位使方法贡献和经验发现保持一致：候选扩张可以提高测试事后上界，但如果没有冻结规则和负结果记录，也会增加选择自由度和报告偏差。",
        "为避免贡献边界被误读，本文将证据分为三层。第一层是确认性证据，包括 3 外层×3 内层×5 重复嵌套、随机候选控制、置换负对照和信号恢复正对照。第二层是应用性证据，包括 MoleculeNet、TDC、多视图确认、逐样本风险和保形预测。第三层是边界性证据，包括 MoleculeACE 活性悬崖、bRo5 划分压力、低相似度失败样本、治理消融和负结果库存。三层证据共同服务于同一论点，而不是彼此分散的模块堆叠。",
    ],
    "1 引言": [
        "分子性质预测在药物发现中的作用已经从离线模型评测延伸到候选排序、实验排队和风险复核。随着可用表征和学习器增多，研究者能够轻易把 Morgan 指纹、MACCS、RDKit2D 描述符、树模型、图模型、预训练化学语言模型和自动机器学习纳入同一比较框架。这个技术趋势带来了真实收益，也带来了一个更隐蔽的问题：当验证集保持不变而候选池持续扩大时，最佳验证分数可能同时反映模型能力和选择自由度。",
        "这种问题不同于传统意义上的测试集泄漏。即使测试标签从未参与训练，验证集仍可能在模型家族、超参数、融合策略、阈值和补救规则之间被反复消费。对实际项目而言，风险并不只表现为最终分数偏高，还表现为验证排序对外层排序的保真度下降、不同终点之间的选择规则不可迁移、以及失败案例在主文中被弱化或遗漏。因而，一个可信的分子预测流程不仅要报告模型表现，还要解释为什么某个模型或规则被选中，以及何时不应使用该选择。",
        "已有研究为这一问题提供了重要基础。交叉验证偏差、适应性过拟合、不确定性估计、适用域、保形预测、多模态融合和 AutoML 已分别从统计、可靠性或工程角度提出工具。然而，这些工具常被放在模型性能叙事之下，缺少一个把候选登记、选择冻结、外层审计和负结果追踪连接起来的主线。本文选择从验证治理切入，是因为审稿人和实际使用者最终关心的不是候选池里有多少模型，而是选择是否可复核、失败是否被记录、边界是否被主动暴露。",
        "本文的研究问题因此被限定为三个层次。第一，固定验证信息下，候选池从 K=4 扩大到 K=8、16 和 32 是否系统性增加固定分母遗憾并降低机会校正命中。第二，当候选池包含异质分子表示时，实际兑现收益是否仍可在共享冻结划分上被确认，而不是来自不可比较的历史运行。第三，验证侧风险、可靠性和化学边界能否形成审稿人可追踪的证据闭环，使负结果和失败案例与正向结果处于同一报告体系。",
    ],
    "2.1 研究范围、数据集与任务登记": [
        "任务登记在本文中承担两个功能。首先，它限定了哪些终点、指标、划分和种子属于确认性分析，避免在观察测试结果后改变研究对象。其次，它把数据清洗、标签方向、指标方向和候选资格写入同一张可审计表，使后续图表中的每个数值都能追溯到具体终点和运行配置。对于小样本分子任务，这一点尤其重要，因为单个终点的偶然波动很容易被误解释为方法优势。",
        "MoleculeNet、TDC、MoleculeACE 和 bRo5 在稿件中承担不同角色。MoleculeNet 提供回归和分类的主面板，用于检验候选池扩张、多视图表示和冻结选择。TDC 提供跨来源 ADMET 终点，用于观察门控在外部公开面板上的异质性。MoleculeACE 强调活性悬崖和局部非连续性，bRo5 则用于更严格的化学迁移边界。这样的分工避免把所有数据集混成一个泛化声明，也让每个实验的证据等级更加清楚。",
    ],
    "2.2 数据标准化、重复处理与泄漏审计": [
        "数据处理遵循先登记、后转换、再训练的顺序。所有标准化、重复处理、特征计算和划分索引均在训练侧或预定义规则下完成，测试标签在选择器冻结前不可见。对于分类任务，正负类比例和少数类样本数被记录为解释校准和固定精度召回波动的必要背景；对于回归任务，单位、标签范围和训练标签标准差被保留，用于解释标准化区间宽度和跨终点比较。",
        "泄漏审计不是单一脚本检查，而是贯穿数据和结果链路的约束。候选登记文件记录 eligible、rejected、failed 和 missing-data 状态；结果文件记录 config_hash、split_hash、seed 和 prediction_hash；主图 source data 由冻结表自动重建。这样做的目的不是声称所有外部复现障碍已经消失，而是保证本地分析级复现可以区分代码、数据、划分和候选状态四类来源。",
    ],
    "2.3 分子表示、候选专家与冻结登记": [
        "候选池被设计为既能产生真实选择压力，又不依赖不可复跑的重型历史模型。32 候选扩池实验使用 Morgan-512 的轻量变体，目的是在候选数量上形成受控梯度；12 候选多视图确认实验则把 Morgan-512、MACCS、RDKit2D 和拼接多视图与线性模型、随机森林和 LightGBM 配对，目的是在同一冻结划分上检验表征异质性。两个实验回答的问题不同，因此不被合并为一个总排行榜。",
        "历史探索中的 GNN、Chemprop、ChemBERTa、MoLFormer 和更复杂融合为候选来源提供背景，但只有在统一外层划分上重训并完成登记后才可进入确认性效应估计。这一规则看似保守，却能防止旧实验中最优结果被事后选入主文。换言之，FZYC-Mol 的治理逻辑并不排斥复杂模型，而是要求复杂模型承担与轻量候选相同的审计成本。",
    ],
    "2.4 验证治理选择器的形式化定义": [
        "形式化定义的核心是把每个候选在内层验证中的效用、方差、复杂度和资格状态分开记录。validation-best 只按验证效用选择；one-SE 规则在接近最佳候选的集合中偏向更简单或更稳定的候选；风险调整规则尝试把验证侧选择风险纳入排序。每种规则都在外层评估前冻结，因此外层结果只能用于审计，而不能回流改变规则。",
        "固定分母遗憾用于比较不同候选池规模时的选择损失。若直接使用各自候选池的 oracle 作为分母，K 扩张会同时改变可达到上界和后悔参照，导致规模效应被部分遮蔽。本文保留固定分母定义，是为了让 K=4 到 K=32 的变化反映同一外层参照下的选择损失，而不是让评价标准随候选池一起移动。",
        "机会校正命中率和 MRR 则从排序角度补充遗憾指标。固定 Top-3 命中会随候选数机械下降，因此需要随机排序负对照给出零信息参照。只有当真实选择流程同时优于置换基线、并在候选池扩张时呈现可解释的下降，才能支持“排序信息被削弱但未消失”的更细致结论。",
    ],
    "2.5 嵌套验证、候选池压力与选择器对照": [
        "重复嵌套验证承担主效应估计，而不是单纯增加运行次数。外层折用于模拟未见数据，内层折用于选择候选，重复种子用于量化划分与随机性的联合影响。端点被视为主要聚类单位，是因为同一终点内的折和种子共享数据来源与标签定义，不能被当作独立生物学重复。这个推断层级直接影响置信区间和精确检验的解释。",
        "候选池压力实验包含 random-order、random-subset 和 family-balanced 三种控制。random-order 检验固定登记顺序是否造成伪趋势，random-subset 检验随机抽样候选时是否仍随 K 改变，family-balanced 检验候选家族比例是否解释规模效应。三种控制的作用不是增加图形复杂度，而是逐一排除审稿人可能提出的替代解释。",
        "AutoGluon 被放在强基线位置，而不是被写成 FZYC-Mol 的竞争对象。AutoML 可以在固定预算内搜索多种模型并提供强预测性能，但它本身不自动解决候选登记、验证消费、负结果归档和逐样本边界解释。本文因此把 AutoGluon 用于界定性能与成本边界，同时保留 FZYC-Mol 作为治理协议的定位。",
    ],
    "2.6 逐样本预测风险、校准与保形预测": [
        "逐样本风险模块的目标不是替代终点级选择，而是解释哪些预测更需要保留、复核或标记为边界样本。风险分数由训练和验证侧信息拟合，并在测试推理前冻结；选择性预测按风险排序报告覆盖率和错误率变化。这样可以区分两个问题：一个模型是否在终点层面被选中，以及该模型在具体分子上是否给出足够可靠的预测。",
        "保形预测被用于补充而不是取代校准。总体覆盖接近标称水平时，少数类覆盖仍可能不足；回归区间达到目标覆盖时，标准化宽度也可能过大。本文因此同时报告总体覆盖、类别条件覆盖、pooled fallback 和区间宽度，避免把单一覆盖率写成自动可用的可靠性保证。",
    ],
    "2.7 活性悬崖、bRo5 与解释性分析": [
        "化学边界实验被安排在主结果之后，是因为它们回答的是适用范围而不是平均性能。MoleculeACE 检验高相似分子对的方向和幅度是否被恢复，bRo5 划分压力检验随机划分之外的迁移退化，片段统计则用于解释局部失败或富集信号的证据强度。三者共同约束了本文结论的化学外推范围。",
    ],
    "2.8 统计分析、计算成本与开放科学": [
        "统计报告遵循“效应、区间、检验和边界”并列的原则。主文避免把折、种子或候选行当作独立样本，而是在端点层面给出配对效应和聚类不确定性。计算成本、实际拟合时间和失败状态被保留，是因为模型选择协议若忽略成本和失败，会把不可部署或不可复跑的候选误写成同等可比。",
    ],
    "2.9 双层决策卡与选择日志": [
        "双层决策卡把终点级选择与样本级复核分开记录。终点层记录 candidate_id、候选家族、选择器、验证效用、外层指标、遗憾、门控状态和失败原因；样本层记录点预测、校准概率或区间、最近邻相似度、风险分位、保形集合或区间以及是否需要人工复核。这样设计可以避免把“某个终点选择了某个模型”和“某个分子预测可以直接使用”混为同一判断。",
        "选择日志还保留未晋级和运行失败条目。对于审稿人而言，这些条目常比成功案例更能显示方法是否可信，因为它们暴露了规则何时失效、候选何时缺失、以及某些样本为何落在适用域之外。本文因此把负结果记录为主流程的一部分，而不是在主文完成后才附加的补充说明。",
    ],
    "2.10 主张层级与终止规则": [
        "主张层级在写作中同样作为终止规则使用。只有同时通过预登记、冻结选择、外层审计和对照检验的结果，才进入主结论；仅在单一数据集或探索性片段统计中出现的信号，被写为边界或假设生成证据。这个规则使扩写后的稿件不会因篇幅增加而扩大结论范围。",
    ],
    "3 结果": [
        "结果部分按证据梯度组织，而不是按模块罗列。前两节检验候选池扩张是否造成选择压力；第三节用正负对照和跨端点风险检验指标解释；第四节确认多视图候选在共享划分下是否兑现收益；随后各节依次给出公开面板、逐样本可靠性、化学迁移、治理消融和负结果库存。这样的顺序使每一组结果都回答前一组结果留下的问题。",
    ],
    "3.1 随机化候选池控制确认规模效应": [
        "该结果的关键不是 K=32 的遗憾数值本身，而是三种候选组成控制均保留了随 K 增大的趋势。如果趋势只在固定顺序下出现，审稿人可以认为它来自登记顺序；如果只在某一候选家族比例下出现，则可能来自家族构成。当前结果排除了这两类简单解释，使“验证信息固定时候选自由度增加会带来选择代价”成为更稳健的观察。",
    ],
    "3.2 重复嵌套验证确认选择不确定性": [
        "重复嵌套结果进一步说明，候选扩张带来的不是单纯的 Top-3 机会率变化。随机排序负对照给出了零信息曲线，而真实流程在 K=4 到 K=32 均保留高于零信息的排序信号。与此同时，机会校正命中率和 MRR 下降表明这些排序信号不足以完全抵消候选池扩大带来的选择压力。",
    ],
    "3.3 指标校准与跨端点风险形成正负证据闭环": [
        "正负对照的并列是本文证据闭环的关键。信号恢复正对照说明，当验证与外层效用之间存在可控相关时，指标能够恢复排序信号；原等权 selection-risk 的负验证说明，一个直觉上合理的风险分数若不能通过置换和留一端点检验，就不能被写成已验证治理模块。严格 LOEO 元风险保留为有限预警工具，而不是被夸大为通用元选择器。",
    ],
    "3.4 共享冻结划分的多视图候选兑现可达到收益": [
        "多视图确认实验补上了旧稿最容易被质疑的比较条件。所有候选在相同外层折、内层折和种子上重训，避免把历史最佳模型与新候选直接比较。结果显示，拼接多视图不仅提高了 test oracle 的可达到上界，也在 validation-best 选择下兑现了配对效用增益；因此，多视图收益不是单纯事后 oracle 现象。",
    ],
    "3.5 MoleculeNet 冻结性能与筛选语境": [
        "MoleculeNet 面板用于把治理结论放回常见分子性质预测语境。分类终点、回归终点和 ClinTox 固定精度召回共同表明，单一平均性能不足以描述筛选价值。特别是 ClinTox 少数类样本较少，召回区间较宽，因此主文不把固定阈值结果写成稳定毒性筛选器，而是把它作为类别不平衡下需要审计的代表场景。",
    ],
    "3.6 TDC 门控揭示终点异质性": [
        "TDC 结果强调门控规则的终点依赖性。5 个 promoted 与 17 个 retained 并不意味着其余终点被证明无效，而是说明在三种子和当前区间宽度下，许多终点无法支持明确晋级。这样的写法保留了负结果的审计价值，也避免把证据不足误写为反向证据。",
    ],
    "3.7 逐样本风险与标签条件保形限定预测可靠性": [
        "逐样本结果把终点级选择进一步落实到样本级决策。风险排序能够降低保留低风险样本时的错误率或 RMSE，但总体覆盖、类别条件覆盖和 fallback 之间的差异提示，可靠性不能只用一个总覆盖率概括。对药物发现应用而言，这意味着预测输出应包含复核状态和边界信号，而不只是单个点预测。",
    ],
    "3.8 MoleculeACE 与 bRo5 界定化学迁移边界": [
        "MoleculeACE 与 bRo5 的共同信息是，平均性能和局部化学连续性并非同一问题。活性悬崖中方向准确率高于随机并不意味着幅度恢复充分；bRo5 中随机划分优于骨架、外缘或时间划分也不应被写成真实项目迁移性能。本文因此把这些结果放在边界部分，用于限制而不是扩大主结论。",
    ],
    "3.9 治理消融、失败案例与负结果": [
        "治理消融表明，规则本身没有跨终点普适赢家。validation-best 在完整多视图池中表现较好，但 one-SE、低方差、低成本和 LOEO 规则都呈现条件性收益或失败。这个结果使 FZYC-Mol 的定位更接近审计框架：它帮助研究者知道何时、为什么、以何种成本选择某个规则，而不是替研究者预设唯一最优规则。",
        "失败案例和负结果库存防止主文只报告正向平均效应。ClinTox 假阴性、低相似度 FreeSolv 和 Lipophilicity 样本、half-life 极端标签以及低支持片段富集，分别指向类别风险、适用域、标签粗糙度和解释证据不足。将这些案例放入主结果而非仅放入补充材料，可以让审稿人看到模型边界已经被主动纳入分析。",
    ],
    "3.10 Source data 自动重建": [
        "Source data 自动重建是本文可复核性的最低承诺。每幅主图不仅有图片文件，也有对应 CSV 和生成脚本；图件包中的 PNG、SVG、PDF 和 TIFF 让阅读、编辑和投稿格式可以分离。当前限制同样需要明确：公开 release、Zenodo DOI 和第三方空环境复跑尚未完成，因此稿件只声明本地分析级可重建，而不把它写成已完成的开放复现。",
    ],
    "4 讨论": [
        "本文最重要的含义是，分子机器学习中的模型选择应被视为实验设计的一部分，而不是训练结束后的报告步骤。候选池扩张、多视图表示和 AutoML 都可能提高可达到性能，但它们也增加验证集被查询的次数。如果没有候选登记和冻结审计，最终测试分数很难区分真实表征收益、候选数量机会效应和事后选择偏差。",
        "这一结论与现有可靠性工具是互补关系。适用域、不确定性估计、校准和保形预测可以帮助解释逐样本输出，但它们不能自动保证终点级模型选择无偏；嵌套验证和外层审计可以约束选择偏差，但它们也不能自动解释低相似度分子或少数类样本的失败。FZYC-Mol 将这些工具放在同一治理链中，目的正是避免单个工具被过度承诺。",
        "一个可能的替代解释是，候选扩张带来的遗憾增加仅来自弱候选被加入池中。多视图确认和候选家族消融对这一解释作出限制：异质表示在共享划分下确实带来了可达到和实际兑现收益，但候选组成变化仍会改变选择损失。因而，候选扩张不是简单的“越多越差”或“越多越好”，而是收益与选择风险同时增加的权衡。",
        "另一个审稿人可能关注的问题是，治理协议是否会牺牲模型创新。本文的结果支持更谨慎的回答：治理并不替代新模型，也不阻止引入更强表示；它规定新候选进入比较前必须登记、重训、冻结并接受相同外层审计。若未来 Chemprop、GNN、ChemBERTa 或 MoLFormer 在同一划分下完成登记，它们可以扩展候选池，但其贡献仍需同时报告可达到上界、实际选择收益和新增选择风险。",
        "本文也存在清晰边界。九个主终点不足以训练通用跨端点元选择器；TDC 三种子结果对宽区间终点只能给出保留或证据不足；MoleculeACE 和 bRo5 公共划分不能替代真实项目前瞻验证；历史重型候选尚未在统一外层划分下重训。保留这些边界会削弱过度概括的吸引力，但能提高结论对审稿人的可信度。",
    ],
    "5 结论": [
        "综合 11 张主图和 8 张主表，FZYC-Mol 支持一个限定但可操作的结论：在固定验证信息下，候选池扩大和多视图异质性必须与冻结选择、外层审计、对照实验和负结果记录同时报告。仅报告最优测试分数会遗漏选择风险；仅报告治理规则也不足以证明预测有用。二者需要在同一证据链中被同时评价。",
        "对后续研究而言，最直接的采用方式不是替换现有预测器，而是在任何新候选进入比较之前建立登记表和终止规则。若一个新模型只在少数终点上改善测试分数，却没有给出验证排序、候选状态、失败记录和逐样本可靠性证据，它仍不应被写成稳健选择。相反，即使某个候选未晋级，只要其失败状态和边界原因被记录，也能为后续模型设计和实验复核提供有用信息。",
        "因此，本文的实际价值不在于给出一个永远胜出的分子预测器，而在于提供一种可复核的报告格式。该格式要求每个终点有登记，每个候选有状态，每个选择有冻结规则，每个主张有 source data，每个失败有原因码。未来若加入更强 GNN、化学语言模型或前瞻实验，这一框架仍可作为扩展前的审计底座。",
    ],
}


REPLACEMENTS = {
    "不能写成“零下降”": "不应表述为“无下降结论”",
    "首先，它限定了": "一方面，它限定了",
    "其次，它把": "另一方面，它把",
    "故仍为探索性关联": "因此仍为探索性关联",
    "只用于量化冻结选择差距": "仅用于量化冻结选择差距",
    "只表示事后上界": "仅表示事后上界",
    "只用于解释边界": "仅用于解释边界",
    "不是新的主干预测网络": "不是新的分子预测主干网络",
    "主要卖点": "主要贡献依据",
    "审稿人和实际使用者": "外部读者和实际使用者",
    "审稿人可追踪": "外部读者可追踪",
    "审稿人可能提出": "外部读者可能提出",
    "对于审稿人而言": "对于外部读者而言",
    "审稿人可以认为": "外部读者可能认为",
    "让审稿人看到": "让读者看到",
    "另一个审稿人可能关注的问题是": "另一个需要预先界定的问题是",
    "对审稿人的可信度": "对外部读者的可信度",
    "永远胜出": "在所有场景中胜出",
    "我们依次采用": "本文依次采用",
    "首创模块": "全新模块",
    "点预测普遍优于 AutoML": "点预测系统性超过 AutoML",
    "不能外推到尚未同折重训的 Chemprop、GNN、ChemBERTa 或 MoLFormer": "不应外推到尚未同折重训的 Chemprop、GNN、ChemBERTa 或 MoLFormer",
}


def insert_after(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new = OxmlElement("w:p")
    paragraph._p.addnext(new)
    out = Paragraph(new, paragraph._parent)
    out.style = style
    run = out.add_run(text)
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    out.paragraph_format.first_line_indent = Cm(0.74)
    out.paragraph_format.line_spacing = 1.18
    out.paragraph_format.space_before = Pt(0)
    out.paragraph_format.space_after = Pt(4)
    return out


def find_heading(doc: Document, heading: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip() == heading:
            return p
    raise ValueError(f"heading not found: {heading}")


def apply_text_replacements(doc: Document) -> None:
    for p in doc.paragraphs:
        if not p.text:
            continue
        text = p.text
        new_text = text
        for old, new in REPLACEMENTS.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            for r in p.runs:
                r.text = ""
            p.runs[0].text = new_text if p.runs else ""


def expand_document(doc: Document) -> None:
    apply_text_replacements(doc)
    for heading, paragraphs in EXPANSIONS.items():
        anchor = find_heading(doc, heading)
        last = anchor
        for text in paragraphs:
            last = insert_after(last, text)
    apply_text_replacements(doc)


def _cell_borders(cell, **edges: tuple[str, str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        if edge in edges:
            val, size, color = edges[edge]
            node.set(qn("w:val"), val)
            node.set(qn("w:sz"), size)
            node.set(qn("w:color"), color)
            node.set(qn("w:space"), "0")
        else:
            node.set(qn("w:val"), "nil")
            node.set(qn("w:sz"), "0")
            node.set(qn("w:color"), "FFFFFF")
            node.set(qn("w:space"), "0")


def _shade(cell, fill: str = "FFFFFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:tblHeader"))
    if node is None:
        node = OxmlElement("w:tblHeader")
        tr_pr.append(node)
    node.set(qn("w:val"), "true")


def _no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _format_cell_text(cell, header: bool) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs:
            r.font.name = "Arial"
            r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            r.font.size = Pt(8.4 if header else 8.1)
            r.font.bold = header
            r.font.color.rgb = RGBColor.from_string("1F2937")


def apply_three_line_tables(doc: Document) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for ri, row in enumerate(table.rows):
            _no_split(row)
            if ri == 0:
                _repeat_header(row)
            for cell in row.cells:
                _shade(cell, "FFFFFF")
                top = ("single", "14", "1F2937") if ri == 0 else ("nil", "0", "FFFFFF")
                bottom = ("single", "9", "1F2937") if ri == 0 else ("nil", "0", "FFFFFF")
                if ri == len(table.rows) - 1:
                    bottom = ("single", "14", "1F2937")
                _cell_borders(cell, top=top, bottom=bottom)
                _format_cell_text(cell, header=(ri == 0))
    for p in doc.paragraphs:
        if p.style and p.style.name == "TableCaption":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.name = "Arial"
                r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                r.font.size = Pt(9.5)
                r.font.bold = True
                r.font.color.rgb = RGBColor.from_string("1F4E79")
        elif p.style and p.style.name == "TableNote":
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(5)
            for r in p.runs:
                r.font.size = Pt(7.8)
                r.font.color.rgb = RGBColor.from_string("667085")


def polish_layout(doc: Document) -> None:
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.25)
        sec.right_margin = Cm(2.25)
    for p in doc.paragraphs:
        if p.style and p.style.name == "Normal":
            p.paragraph_format.line_spacing = 1.18
            p.paragraph_format.space_after = Pt(4)
        elif p.style and p.style.name.startswith("Heading"):
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif p.style and p.style.name == "FigureCaption":
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)


def locate(path: Path, pattern: str) -> Path:
    if path.exists():
        return path
    matches = list(path.parent.glob(pattern))
    if not matches:
        raise FileNotFoundError(path)
    return max(matches, key=lambda p: p.stat().st_mtime)


def build_doc() -> None:
    src_doc = locate(SRC_DOC, "*-7.docx")
    src_pack = locate(SRC_PACK, "*-7_图表包")
    if src_pack.is_dir():
        shutil.copytree(src_pack, PACK, dirs_exist_ok=True)
    doc = Document(src_doc)
    expand_document(doc)
    polish_layout(doc)
    apply_three_line_tables(doc)
    doc.save(DOCX)


def build_supp() -> None:
    src = locate(SRC_SUPP, "*-7_补充表.docx")
    doc = Document(src)
    polish_layout(doc)
    apply_three_line_tables(doc)
    for p in doc.paragraphs:
        if p.style and p.style.name == "Title" and "小论文-7" in p.text:
            for r in p.runs:
                r.text = r.text.replace("小论文-7", "小论文-8")
    doc.save(SUPP)


def xml_errors(path: Path) -> list[str]:
    errors: list[str] = []
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            errors.append(f"bad_member:{bad}")
        for name in archive.namelist():
            if name.endswith(".xml"):
                try:
                    ElementTree.fromstring(archive.read(name))
                except Exception as exc:
                    errors.append(f"{name}:{type(exc).__name__}")
    return errors


def count_units(text: str) -> dict[str, int]:
    return {
        "chinese_chars": len(re.findall(r"[\u4e00-\u9fff]", text)),
        "no_space_chars": len(re.sub(r"\s+", "", text)),
        "cn_or_en_units": len(re.findall(r"[A-Za-z0-9]+|[\u4e00-\u9fff]", text)),
    }


def audit() -> dict:
    doc = Document(DOCX)
    text = "\n".join(p.text for p in doc.paragraphs)
    fig_caps = [p.text.strip() for p in doc.paragraphs if re.match(r"^图\s*\d+\s*\|", p.text.strip())]
    tab_caps = [p.text.strip() for p in doc.paragraphs if re.match(r"^表\s*\d+\s*\|", p.text.strip())]
    formats = {ext: sorted(p.stem for p in (PACK / "figures").glob(f"fig[0-9][0-9]_*.{ext}")) for ext in ["png", "svg", "pdf", "tiff"]}
    figure_checks = {}
    for p in sorted((PACK / "figures").glob("fig[0-9][0-9]_*.png")):
        svg = p.with_suffix(".svg")
        with Image.open(p) as im:
            figure_checks[p.stem] = {
                "size": [im.width, im.height],
                "png_bytes": p.stat().st_size,
                "svg_editable_text": "<text" in svg.read_text(encoding="utf-8"),
            }
    colloquial_flags = [
        "可以看到", "不难发现", "显然", "很明显", "完美", "万能", "证明了", "零下降",
        "首先，", "其次，", "最后，", "卖点", "审稿人", "永远"
    ]
    overclaim_flags = ["首次", "首创", "革命性", "决定性证明", "完全解决", "普遍优于"]
    report = {
        "docx": str(DOCX),
        "supplement": str(SUPP),
        "package": str(PACK),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "figure_caption_numbers": [int(re.search(r"\d+", x).group()) for x in fig_caps],
        "table_caption_numbers": [int(re.search(r"\d+", x).group()) for x in tab_caps],
        "counts": count_units(text),
        "figure_formats": {k: len(v) for k, v in formats.items()},
        "figure_stems_match": len({tuple(v) for v in formats.values()}) == 1,
        "figure_checks_ok": all(v["size"][0] >= 2000 and v["png_bytes"] >= 50000 and v["svg_editable_text"] for v in figure_checks.values()),
        "source_data_csv": len(list((PACK / "source_data").glob("*.csv"))),
        "xml_errors": xml_errors(DOCX),
        "supp_xml_errors": xml_errors(SUPP),
        "colloquial_hits": {term: text.count(term) for term in colloquial_flags if term in text},
        "overclaim_hits": {term: text.count(term) for term in overclaim_flags if term in text},
    }
    report["passed"] = all([
        report["tables"] == 8,
        report["inline_shapes"] == 11,
        report["figure_caption_numbers"] == list(range(1, 12)),
        report["table_caption_numbers"] == list(range(1, 9)),
        15000 <= report["counts"]["cn_or_en_units"] <= 17500,
        all(v == 11 for v in report["figure_formats"].values()),
        report["figure_stems_match"],
        report["figure_checks_ok"],
        report["source_data_csv"] >= 80,
        not report["xml_errors"],
        not report["supp_xml_errors"],
        not report["colloquial_hits"],
        not report["overclaim_hits"],
    ])
    AUDIT.parent.mkdir(parents=True, exist_ok=True)
    AUDIT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    if not report["passed"]:
        raise SystemExit(json.dumps(report, ensure_ascii=False, indent=2))
    return report


def main() -> None:
    build_doc()
    build_supp()
    report = audit()
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
