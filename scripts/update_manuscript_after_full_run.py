from __future__ import annotations

import shutil
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from scripts.update_manuscript_supplement import (  # noqa: E402
    fill_table,
    fmt,
    insert_para,
    insert_table,
    replace_paragraph,
    set_run_font,
    style_paragraph,
)


OUT = ROOT / "reports" / "remaining_missing_experiments_20260606"
NESTED = OUT / "true_nested_validation"
FIG = OUT / "figures"


def source_docx() -> Path:
    base = Path.home() / "Desktop" / "\u4fee\u6539"
    preferred = base / "FZYC-Mol_\u521d\u7a3f4_\u8865\u5145\u5b9e\u9a8c\u4fee\u8ba2.docx"
    if preferred.exists():
        return preferred
    candidates = sorted(base.glob("FZYC-Mol*\u521d\u7a3f4*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not candidates:
        raise FileNotFoundError("Cannot find FZYC-Mol draft 4 under Desktop/修改.")
    return candidates[0]


def target_docx(src: Path) -> Path:
    return src.with_name("FZYC-Mol_\u521d\u7a3f4_\u8865\u5145\u5b9e\u9a8c\u5168\u8dd1\u4fee\u8ba2.docx")


def compact(text: object, n: int = 76) -> str:
    if text is None or pd.isna(text):
        return ""
    value = str(text).replace("\n", " ").strip()
    return value if len(value) <= n else value[: n - 1] + "..."


def pct_text(x: object) -> str:
    if x is None or pd.isna(x):
        return ""
    return f"{100 * float(x):.1f}%"


def model_counts(text: str) -> str:
    mapping = {
        "rf_balanced": "RF",
        "extratrees_balanced": "ExtraTrees",
        "logreg_l2_c0.3": "LR C0.3",
        "logreg_l2_c1": "LR C1",
        "ridge_alpha1": "Ridge α1",
        "ridge_alpha10": "Ridge α10",
        "rf_reg": "RF",
        "extratrees_reg": "ExtraTrees",
    }
    out = str(text)
    for old, new in mapping.items():
        out = out.replace(old, new)
    return out


def category_label(value: str) -> str:
    labels = {
        "full": "Full",
        "best_single": "Best single",
        "simple_mean": "Simple mean",
        "no_validation_selector_fixed_morgan": "Best single/no selector",
        "without_selector": "w/o selector",
        "without_fusion": "w/o fusion",
        "without_ad_gate": "w/o AD gate",
        "without_uncertainty_weighting": "w/o uncertainty weighting",
        "without_hier_motif_multifp": "w/o motif/fingerprint",
        "without_rescue_head": "w/o rescue head",
        "without_rescue_head_current": "w/o rescue head (accepted endpoints)",
    }
    return labels.get(value, value)


def source_label(value: str) -> str:
    labels = {
        "moleculenet": "MoleculeNet",
        "tdc": "TDC",
        "moleculenet_multimethod": "MoleculeNet multimethod",
        "moleculenet_selector_ablation": "MoleculeNet selector ablation",
        "tdc_multimethod": "TDC multimethod",
        "tdc_selector_ablation": "TDC selector ablation",
    }
    return labels.get(value, value)


def task_type_label(value: str) -> str:
    return "\u5206\u7c7b" if value == "classification" else "\u56de\u5f52"


def build_true_nested_rows() -> list[list[str]]:
    df = pd.read_csv(NESTED / "true_nested_validation_summary.csv")
    rows: list[list[str]] = []
    for r in df.itertuples(index=False):
        if r.task_type == "classification":
            metric = f"ROC-AUC {fmt(r.roc_auc_mean)} +/- {fmt(r.roc_auc_sd)}"
            aux = f"PR-AUC {fmt(r.pr_auc_mean)}; Brier {fmt(r.brier_mean)}"
        else:
            metric = f"RMSE {fmt(r.rmse_mean)} +/- {fmt(r.rmse_sd)}"
            aux = f"MAE {fmt(r.mae_mean)}; Spearman {fmt(r.spearman_mean)}"
        rows.append(
            [
                r.dataset,
                task_type_label(r.task_type),
                str(int(r.n_outer)),
                model_counts(r.selected_candidate_count),
                metric,
                aux,
            ]
        )
    return rows


def build_seed_nested_rows() -> list[list[str]]:
    df = pd.read_csv(OUT / "nested_seed_validation_summary.csv")
    rows: list[list[str]] = []
    for r in df.itertuples(index=False):
        rows.append(
            [
                r.dataset,
                source_label(r.source),
                task_type_label(r.task_type),
                str(int(r.n_outer)),
                fmt(r.selected_test_mean),
                fmt(r.median_regret_vs_test_oracle),
                str(int(r.top_model_switches)),
            ]
        )
    return rows


def build_ablation_rows() -> list[list[str]]:
    df = pd.read_csv(OUT / "unified_ablation_matrix_summary.csv")
    order = {
        "full": 0,
        "best_single": 1,
        "simple_mean": 2,
        "no_validation_selector_fixed_morgan": 3,
        "without_selector": 4,
        "without_fusion": 5,
        "without_ad_gate": 6,
        "without_uncertainty_weighting": 7,
        "without_hier_motif_multifp": 8,
        "without_rescue_head": 9,
        "without_rescue_head_current": 10,
    }
    df["sort_key"] = df["category"].map(order).fillna(99)
    df = df.sort_values(["sort_key", "task_type"])
    return [
        [
            category_label(r.category),
            task_type_label(r.task_type),
            str(int(r.n_units)),
            fmt(r.mean_test),
            fmt(r.mean_delta_vs_full_positive),
            pct_text(r.positive_fraction_vs_full),
        ]
        for r in df.itertuples(index=False)
    ]


def build_tanimoto_rows() -> list[list[str]]:
    df = pd.read_csv(OUT / "exact_tanimoto_bins_summary.csv")
    bin_order = {">0.7": 0, "0.5-0.7": 1, "<0.5": 2}
    df["sort_key"] = df["similarity_bin"].map(bin_order).fillna(9)
    df = df.sort_values(["source", "task_type", "sort_key"])
    rows: list[list[str]] = []
    for r in df.itertuples(index=False):
        if r.task_type == "classification":
            metric = f"ROC {fmt(r.roc_auc)}; PR {fmt(r.pr_auc)}; ECE {fmt(r.ece)}"
        else:
            metric = f"RMSE {fmt(r.rmse)}; MAE {fmt(r.mae)}; ρ {fmt(r.spearman)}"
        rows.append(
            [
                source_label(r.source),
                task_type_label(r.task_type),
                r.similarity_bin,
                fmt(r.n, 1),
                fmt(r.mean_similarity),
                fmt(r.high_error_enrichment),
                metric,
            ]
        )
    return rows


def build_conformal_rows() -> list[list[str]]:
    df = pd.read_csv(OUT / "conformal_80_90_95_summary.csv")
    rows: list[list[str]] = []
    for r in df.sort_values(["task_type", "target_coverage"]).itertuples(index=False):
        if r.task_type == "classification":
            size = f"set {fmt(r.avg_set_size_mean)}"
            risk = f"singleton {pct_text(r.singleton_rate_mean)}; empty {pct_text(r.empty_rate_mean)}"
        else:
            size = f"width {fmt(r.mean_width_mean)}"
            risk = f"median AE {fmt(r.median_abs_error_mean)}"
        rows.append(
            [
                task_type_label(r.task_type),
                pct_text(r.target_coverage),
                str(int(r.n)),
                fmt(r.coverage_mean),
                fmt(r.coverage_median),
                size,
                risk,
            ]
        )
    return rows


def build_moleculeace_rows() -> list[list[str]]:
    df = pd.read_csv(OUT / "moleculeace_gap_correlation_summary.csv")
    agg = (
        df.groupby("task", as_index=False)
        .agg(
            seeds=("seed", "count"),
            cliff_pairs=("n_cliff_pairs", "mean"),
            cutoff=("cliff_abs_delta_cutoff", "mean"),
            gap_spearman=("gap_spearman", "mean"),
            direction_accuracy=("direction_accuracy", "mean"),
            mean_gap_abs_error=("mean_gap_abs_error", "mean"),
            mean_similarity=("mean_similarity", "mean"),
        )
        .sort_values("task")
    )
    return [
        [
            r.task,
            str(int(r.seeds)),
            fmt(r.cliff_pairs, 1),
            fmt(r.cutoff),
            fmt(r.gap_spearman),
            pct_text(r.direction_accuracy),
            fmt(r.mean_gap_abs_error),
        ]
        for r in agg.itertuples(index=False)
    ]


def build_failure_rows() -> list[list[str]]:
    df = pd.read_csv(OUT / "extended_failure_cases.csv")
    rows: list[list[str]] = []
    for r in df.itertuples(index=False):
        if pd.isna(r.y_true):
            evidence = ""
        else:
            evidence = f"true {fmt(r.y_true)}; pred {fmt(r.y_pred)}; risk/sim {fmt(r.risk_or_similarity)}"
        rows.append(
            [
                str(r.case_id),
                compact(r.case_type, 34),
                str(r.dataset),
                evidence,
                compact(r.interpretation, 82),
            ]
        )
    return rows


def update_tracker_table(doc: Document) -> None:
    replacements = {
        "P0-1": (
            "\u5df2\u8865\u8dd1",
            "\u589e\u52a0\u771f\u5b9e 3x3 nested validation\uff089 \u4e2a\u4ee3\u8868\u7ec8\u70b9\uff09\uff0c\u5e76\u4fdd\u7559 seed-nested selector audit\u3001Top-3\u3001optimism gap \u548c regret\u3002",
        ),
        "P0-3": (
            "\u5df2\u8865\u8dd1",
            "\u8865\u9f50 Full / best single / simple mean / w/o selector / w/o fusion / w/o AD gate / w/o uncertainty weighting / w/o motif-fingerprint / w/o rescue head \u7edf\u4e00\u77e9\u9635\u3002",
        ),
        "P0-4": (
            "\u5df2\u8865\u8dd1",
            "\u6309 >0.7\u30010.5-0.7\u3001<0.5 \u4e09\u4e2a\u4e92\u65a5 Tanimoto bin \u8f93\u51fa\u6027\u80fd\u3001\u6821\u51c6\u3001\u4e0d\u786e\u5b9a\u6027\u548c\u98ce\u9669\u5bcc\u96c6\u3002",
        ),
        "P0-5": (
            "\u5df2\u8865\u8dd1",
            "\u589e\u52a0\u9884\u6d4b\u5dee\u5f02 vs \u771f\u5b9e\u5dee\u5f02\u76f8\u5173\u6027\uff0c\u5e76\u8f93\u51fa\u4ee3\u8868\u6027 cliff pair \u6848\u4f8b\u56fe\u3002",
        ),
        "P1-2": (
            "\u5df2\u8865\u8dd1",
            "\u8865\u9f50 80%/90%/95% \u76ee\u6807\u8986\u76d6\u7387\uff0c\u5206\u7c7b\u5e73\u5747\u5b9e\u6d4b\u8986\u76d6 0.814/0.918/0.956\uff0c\u56de\u5f52\u4e3a 0.823/0.925/0.962\u3002",
        ),
        "P1-3": (
            "\u5df2\u6269\u5c55",
            "\u539f 3 \u7c7b\u6848\u4f8b\u5916\uff0c\u65b0\u589e\u4f4e\u76f8\u4f3c\u5ea6\u5931\u8d25\u6848\u4f8b\u548c MoleculeACE \u6d3b\u6027\u60ac\u5d16\u5931\u8d25\u6848\u4f8b\u3002",
        ),
    }
    for table in doc.tables:
        for row in table.rows:
            if not row.cells:
                continue
            key = row.cells[0].text.strip()
            if key in replacements and len(row.cells) >= 4:
                row.cells[2].text = replacements[key][0]
                row.cells[3].text = replacements[key][1]
                for cell in (row.cells[2], row.cells[3]):
                    for p in cell.paragraphs:
                        style_paragraph(p, size=8)


def insert_picture(anchor, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        return
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(5.8))
    insert_para(anchor, caption, size=8)


def update_core_text(doc: Document) -> None:
    replace_paragraph(
        doc,
        "\u53ef\u9760\u6027\u5206\u6790\u663e\u793a",
        "\u53ef\u9760\u6027\u5206\u6790\u663e\u793a\uff0c\u96c6\u6210\u6807\u51c6\u5dee\u3001\u9519\u8bef\u6a21\u578b\u3001\u9519\u8bef-\u9002\u7528\u57df\u6df7\u5408\u6307\u6807\u3001Tanimoto \u9002\u7528\u57df\u548c\u91cd\u6784\u8bef\u5dee\u80fd\u591f\u8bc6\u522b\u9ad8\u8bef\u5dee\u6837\u672c\uff1b\u98ce\u9669\u5206\u6570\u5bf9\u5206\u7c7b\u9519\u8bef\u7684\u4e2d\u4f4d AUROC \u4e3a 0.788\uff0c\u5bf9\u56de\u5f52\u9ad8\u8bef\u5dee\u6837\u672c\u7684\u4e2d\u4f4d AUROC \u4e3a 0.652\u3002\u4fdd\u5f62\u9884\u6d4b\u5728 80%/90%/95% \u76ee\u6807\u4e0b\u7684\u5e73\u5747\u7ecf\u9a8c\u8986\u76d6\u7387\u4e3a 0.814/0.918/0.956\uff08\u5206\u7c7b\uff09\u548c 0.823/0.925/0.962\uff08\u56de\u5f52\uff09\u3002\u65b0\u589e\u771f\u5b9e 3x3 nested validation \u5728 9 \u4e2a\u4ee3\u8868\u7ec8\u70b9\u4e0a\u68c0\u9a8c\u4e86\u5916\u5c42\u8bc4\u4f30\u4e0e\u5185\u5c42\u5019\u9009\u9009\u62e9\u7684\u9694\u79bb\uff1bseed-nested selector audit \u5219\u8868\u660e\uff0c\u8de8 15 \u4e2a\u6570\u636e\u96c6-\u5019\u9009\u6c60\u7ec4\u5408\u7684\u4e2d\u4f4d test-oracle regret \u603b\u4f53\u8f83\u5c0f\uff0c\u4f46\u4ecd\u5b58\u5728\u7ec8\u70b9\u4f9d\u8d56\u7684\u6a21\u578b\u5207\u6362\u3002\u57fa\u5e8f\u5f52\u56e0\u3001\u7247\u6bb5\u5bcc\u96c6\u3001\u4e25\u683c Tanimoto \u4e09\u6863\u5206\u5c42\u548c MoleculeACE \u5dee\u5f02\u76f8\u5173\u5206\u6790\u4e00\u8d77\u754c\u5b9a\u4e86 FZYC-Mol \u7684\u4f7f\u7528\u8fb9\u754c\u3002",
    )
    replace_paragraph(
        doc,
        "\u6839\u636e\u8865\u5145\u5b9e\u9a8c\u6e05\u5355",
        "\u6839\u636e\u8865\u5145\u5b9e\u9a8c\u6e05\u5355\uff0c\u672c\u8f6e\u4fee\u8ba2\u5c06 P0/P1 \u8bc1\u636e\u4ece\u8865\u5145\u6750\u6599\u62c9\u56de\u4e3b\u6587\u903b\u8f91\uff1a\u5148\u5ba1\u8ba1\u9a8c\u8bc1\u96c6\u9009\u62e9\u662f\u5426\u5b58\u5728\u6392\u540d\u504f\u5dee\uff0c\u518d\u7528\u771f\u5b9e 3x3 nested validation\u3001seed-nested selector audit\u3001\u914d\u5bf9\u7edf\u8ba1\u548c\u7edf\u4e00\u6d88\u878d\u77e9\u9635\u754c\u5b9a\u54ea\u4e9b\u6a21\u5757\u771f\u6b63\u6709\u6548\uff0c\u6700\u540e\u5728\u4e25\u683c Tanimoto \u4e09\u6863\u3001\u7ed3\u6784\u5206\u79bb\u3001MoleculeACE \u6d3b\u6027\u60ac\u5d16\u3001\u4e0d\u5e73\u8861\u5206\u7c7b\u300180/90/95 \u4fdd\u5f62\u8986\u76d6\u548c\u5931\u8d25\u6848\u4f8b\u4e2d\u68c0\u67e5\u9002\u7528\u8fb9\u754c\u3002\u8fd9\u4e9b\u8865\u5145\u5b9e\u9a8c\u4e0d\u6539\u53d8\u6d4b\u8bd5\u96c6\u4e00\u6b21\u6027\u62a5\u544a\u539f\u5219\uff0c\u800c\u662f\u5c06\u201c\u4e3a\u4ec0\u4e48\u63a5\u53d7\u6216\u62d2\u7edd\u67d0\u4e2a\u5019\u9009\u201d\u5199\u6210\u53ef\u8ffd\u6eaf\u7684\u8bc1\u636e\u94fe\u3002",
    )
    replace_paragraph(
        doc,
        "\u6ce8\uff1aP1 \u4e2d\u7684\u4fdd\u5f62\u9884\u6d4b",
        "\u6ce8\uff1aP1 \u4e2d\u7684\u4fdd\u5f62\u9884\u6d4b\u5df2\u8865\u8dd1 80%/90%/95% \u4e09\u4e2a\u8986\u76d6\u7387\uff0c\u5e76\u5728\u8865\u5145\u8868 S15 \u4e2d\u7edf\u4e00\u62a5\u544a\u3002",
    )
    replace_paragraph(
        doc,
        "\u4f4e\u76f8\u4f3c\u5ea6\u548c\u7ed3\u6784\u5206\u79bb\u7ed3\u679c\u63d0\u9192",
        "\u4f4e\u76f8\u4f3c\u5ea6\u548c\u7ed3\u6784\u5206\u79bb\u7ed3\u679c\u63d0\u9192\uff0c\u5e73\u5747\u6027\u80fd\u5e76\u4e0d\u80fd\u7b49\u540c\u4e8e\u65b0\u9aa8\u67b6\u5916\u63a8\u80fd\u529b\u3002\u65b0\u589e\u4e25\u683c\u4e92\u65a5 Tanimoto \u4e09\u6863\u663e\u793a\uff0c<0.5 \u533a\u95f4\u5728 MoleculeNet \u5206\u7c7b\u3001MoleculeNet \u56de\u5f52\u3001TDC \u5206\u7c7b\u548c TDC \u56de\u5f52\u4e2d\u5747\u51fa\u73b0\u9ad8\u8bef\u5dee\u5bcc\u96c6\uff0c\u5bcc\u96c6\u500d\u6570\u5206\u522b\u4e3a 1.346\u30011.431\u30011.250 \u548c 1.070\u3002\u8fd9\u4e00\u7ed3\u679c\u652f\u6301\u5c06\u9002\u7528\u57df\u3001\u98ce\u9669-\u8986\u76d6\u66f2\u7ebf\u3001\u6700\u8fd1\u90bb\u8bc1\u636e\u548c\u6027\u80fd\u5206\u6570\u540c\u65f6\u62a5\u544a\u3002",
    )
    replace_paragraph(
        doc,
        "MoleculeACE \u914d\u5bf9\u7ed3\u679c\u4e0e\u4f4e\u76f8\u4f3c\u5ea6\u5206\u6790\u4e92\u76f8\u8865\u5145",
        "MoleculeACE \u914d\u5bf9\u7ed3\u679c\u4e0e\u4f4e\u76f8\u4f3c\u5ea6\u5206\u6790\u4e92\u76f8\u8865\u5145\u3002\u6d3b\u6027\u60ac\u5d16\u76ee\u6807\u5019\u9009\u5728 51 \u4e2a seed \u914d\u5bf9\u4e2d\u7684\u603b RMSE \u5e73\u5747\u6b63\u5411\u53d8\u5316\u4e3a 0.0069\uff0c\u60ac\u5d16\u5b50\u96c6 RMSE \u5e73\u5747\u6b63\u5411\u53d8\u5316\u4e3a 0.0056\u3002\u65b0\u589e\u5dee\u5f02\u76f8\u5173\u5206\u6790\u8fdb\u4e00\u6b65\u663e\u793a\uff0c\u9884\u6d4b\u5dee\u5f02\u4e0e\u771f\u5b9e\u5dee\u5f02\u7684 Spearman \u76f8\u5173\u6027\u5177\u6709\u660e\u663e\u4efb\u52a1\u4f9d\u8d56\u6027\uff0c\u4f46\u591a\u6570\u4efb\u52a1\u7684\u65b9\u5411\u5224\u522b\u51c6\u786e\u7387\u9ad8\u4e8e 0.70\u3002\u56e0\u6b64\uff0c\u672c\u6587\u5c06\u8be5\u7ed3\u679c\u89e3\u91ca\u4e3a\u60ac\u5d16\u98ce\u9669\u8bc6\u522b\u548c\u5019\u9009\u6cbb\u7406\u7684\u8865\u5145\u8bc1\u636e\uff0c\u800c\u4e0d\u662f\u5bf9\u6570\u503c\u60ac\u5d16\u9884\u6d4b\u5df2\u7ecf\u89e3\u51b3\u7684\u8bc1\u660e\u3002",
    )
    replace_paragraph(
        doc,
        "\u4e0d\u5e73\u8861\u5206\u7c7b\u4e0e\u4fdd\u5f62\u9884\u6d4b\u7684\u8865\u5145\u7ed3\u679c",
        "\u4e0d\u5e73\u8861\u5206\u7c7b\u4e0e\u4fdd\u5f62\u9884\u6d4b\u7684\u8865\u5145\u7ed3\u679c\u6539\u5584\u4e86 ClinTox\u3001DILI\u3001hERG \u548c CYP \u5e95\u7269\u7b49\u4efb\u52a1\u7684\u89e3\u91ca\u3002ROC-AUC \u4ecd\u4f5c\u4e3a\u6807\u51c6\u6307\u6807\uff0c\u4f46\u4e3b\u6587\u540c\u65f6\u62a5\u544a PR-AUC\u3001Brier\u3001ECE \u548c\u5bcc\u96c6\u6307\u6807\uff0c\u5e76\u5c06 80%/90%/95% \u7ecf\u9a8c\u8986\u76d6\u7387\u4f5c\u4e3a\u591a\u5c42\u4e0d\u786e\u5b9a\u6027\u8fb9\u754c\u8bc1\u636e\u3002\u8865\u8dd1\u540e\uff0c\u5206\u7c7b\u4efb\u52a1\u7684\u5e73\u5747\u5b9e\u6d4b\u8986\u76d6\u7387\u4e3a 0.814/0.918/0.956\uff0c\u56de\u5f52\u4efb\u52a1\u4e3a 0.823/0.925/0.962\uff0c\u4e0e\u76ee\u6807\u8986\u76d6\u57fa\u672c\u4e00\u81f4\u3002",
    )
    replace_paragraph(
        doc,
        "\u8865\u5145\u8868 S8. \u4fdd\u5f62\u9884\u6d4b 90%",
        "\u8865\u5145\u8868 S8. \u4fdd\u5f62\u9884\u6d4b 90% \u76ee\u6807\u8986\u76d6\u7387\u6458\u8981\uff08\u539f\u8868\uff09\u3002",
    )
    replace_paragraph(
        doc,
        "\u6700\u540e\uff0c\u8d1f\u7ed3\u679c\u548c\u5931\u8d25\u6848\u4f8b\u88ab\u660e\u786e\u4fdd\u7559",
        "\u6700\u540e\uff0c\u8d1f\u7ed3\u679c\u548c\u5931\u8d25\u6848\u4f8b\u88ab\u660e\u786e\u4fdd\u7559\u3002ClinTox \u9ad8\u98ce\u9669\u5047\u9634\u6027\u8868\u660e\uff0c\u9ad8 ROC-AUC \u4efb\u52a1\u4ecd\u9700\u6837\u672c\u7ea7\u98ce\u9669\u8bc1\u636e\uff1b3D-lite \u548c\u7c97\u7cd9\u5ea6\u52a0\u6743\u5728 oracle \u6761\u4ef6\u4e0b\u5076\u6709\u6f5c\u5728\u6536\u76ca\uff0c\u4f46\u9a8c\u8bc1\u96c6\u95e8\u63a7\u672a\u5c06\u5176\u63a5\u5165\u6700\u7ec8\u7b56\u7565\u3002\u672c\u8f6e\u65b0\u589e\u4f4e\u76f8\u4f3c\u5ea6\u5931\u8d25\u548c MoleculeACE \u6d3b\u6027\u60ac\u5d16\u5931\u8d25\u6848\u4f8b\uff0c\u4f7f\u5931\u8d25\u7c7b\u522b\u4ece 3 \u7c7b\u6269\u5c55\u5230 5 \u7c7b\u3002\u8fd9\u4e9b\u7ed3\u679c\u5171\u540c\u754c\u5b9a FZYC-Mol \u7684\u4f7f\u7528\u8fb9\u754c\uff0c\u4e5f\u907f\u514d\u8bba\u6587\u53ea\u62a5\u544a\u6b63\u7ed3\u679c\u3002",
    )
    replace_paragraph(
        doc,
        "\u7b2c\u4e8c\u4e2a\u53ef\u80fd\u95ee\u9898\u662f",
        "\u7b2c\u4e8c\u4e2a\u53ef\u80fd\u95ee\u9898\u662f\uff0c\u9009\u62e9\u5668\u662f\u5426\u4f1a\u8fc7\u62df\u5408\u9a8c\u8bc1\u96c6\u3002\u672c\u6587\u901a\u8fc7\u4e94\u4e2a\u8bbe\u8ba1\u964d\u4f4e\u8fd9\u4e00\u98ce\u9669\u3002\u9996\u5148\uff0c\u9009\u62e9\u5668\u53ea\u5728\u9884\u5b9a\u4e49\u5019\u9009\u6c60\u4e2d\u9009\u62e9\uff0c\u4e0d\u5141\u8bb8\u6839\u636e\u6d4b\u8bd5\u7ed3\u679c\u6dfb\u52a0\u4e34\u65f6\u89c4\u5219\u3002\u5176\u6b21\uff0c\u7ed3\u679c\u6309\u591a\u4e2a\u968f\u673a\u79cd\u5b50\u6c47\u603b\uff0c\u5e76\u62a5\u544a\u5747\u503c\u3001\u6807\u51c6\u5dee\u548c\u914d\u5bf9\u7edf\u8ba1\u3002\u7b2c\u4e09\uff0c\u8865\u6551\u5934\u8fdb\u5165\u6700\u7ec8\u7b56\u7565\u9700\u8981\u901a\u8fc7\u9a8c\u8bc1\u96c6\u63a5\u53d7\uff0c\u4e14\u6ca1\u6709\u88ab\u6240\u6709\u7ec8\u70b9\u65e0\u5dee\u522b\u91c7\u7528\u3002\u7b2c\u56db\uff0c\u672c\u8f6e\u589e\u52a0\u6392\u540d\u5ba1\u8ba1\u3001Top-3 \u547d\u4e2d\u3001optimism gap \u548c\u8d1f\u7ed3\u679c\u8868\uff0c\u5c06\u4f4e\u76f8\u5173\u5019\u9009\u6c60\u660e\u786e\u5199\u4e3a\u9009\u62e9\u98ce\u9669\u800c\u975e\u80dc\u5229\u8bc1\u636e\u3002\u7b2c\u4e94\uff0c\u65b0\u589e\u771f\u5b9e 3x3 nested validation \u5728 9 \u4e2a\u4ee3\u8868\u7ec8\u70b9\u4e0a\u5c06\u5185\u5c42\u5019\u9009\u9009\u62e9\u4e0e\u5916\u5c42\u8bc4\u4f30\u5206\u79bb\uff0c\u7528\u4e8e\u5ba1\u8ba1\u800c\u4e0d\u7528\u4e8e\u6539\u5199\u4e3b\u6d4b\u8bd5\u96c6\u7ed3\u679c\u3002",
    )
    replace_paragraph(
        doc,
        "\u7b2c\u4e8c\u4f18\u5148\u7ea7\u662f\u628a\u76ee\u524d\u5df2\u6709\u7684 Top-K",
        "\u7b2c\u4e8c\u4f18\u5148\u7ea7\u662f\u628a\u76ee\u524d\u5df2\u6709\u7684 Top-K\u3001\u5806\u53e0\u3001\u98ce\u9669\u8c03\u6574\u548c\u7a33\u5b9a\u6027\u6253\u7834\u5e73\u5c40\u89c4\u5219\u6539\u6210\u66f4\u4e25\u683c\u7684\u9884\u6ce8\u518c\u5019\u9009\u65cf\u3002\u5f53\u524d\u8865\u6551\u6574\u5408\u9009\u62e9\u5668\u5df2\u663e\u793a Lipophilicity \u80fd\u53d7\u76ca\uff0cMoleculeACE \u60ac\u5d16\u76ee\u6807\u5019\u9009\u4e5f\u6709\u5c0f\u5e45\u5e73\u5747\u6536\u76ca\uff1b\u4f46\u771f\u5b9e nested validation \u5c1a\u53ea\u8986\u76d6\u4ee3\u8868\u7ec8\u70b9\u548c\u8f7b\u91cf\u5019\u9009\u6c60\uff0c\u540e\u7eed\u4ecd\u5e94\u5728\u5916\u90e8\u65f6\u95f4\u5212\u5206\u6216\u66f4\u5927\u89c4\u6a21\u76f2\u6d4b\u4e0a\u590d\u6838\u3002",
    )


def insert_full_run_section(doc: Document) -> None:
    anchor = next((p for p in doc.paragraphs if p.text.strip().startswith("4.7 ")), None)
    if anchor is None:
        raise RuntimeError("Cannot find section 4.7 anchor.")

    insert_para(anchor, "4.6.1 \u8865\u8dd1\u5b8c\u6210\u540e\u7684\u65b0\u589e\u9a8c\u8bc1", size=11, bold=True)
    insert_para(
        anchor,
        "\u672c\u8f6e\u8865\u8dd1\u628a\u524d\u7248\u6807\u8bb0\u4e3a\u201c\u90e8\u5206\u5b8c\u6210\u201d\u6216\u201c\u672a\u8dd1\u9f50\u201d\u7684\u9879\u76ee\u5168\u90e8\u8f6c\u5316\u4e3a\u53ef\u5ba1\u8ba1\u8f93\u51fa\u3002\u5176\u4e2d\uff0c\u771f\u5b9e 3x3 nested validation \u7528\u4ee3\u8868\u7ec8\u70b9\u68c0\u9a8c\u6d41\u7a0b\u9694\u79bb\uff1bseed-nested selector audit \u7528\u73b0\u6709\u5019\u9009\u6c60\u68c0\u67e5\u8de8 seed \u9009\u62e9\u7a33\u5b9a\u6027\uff1b\u7edf\u4e00\u6d88\u878d\u77e9\u9635\u3001\u4e25\u683c Tanimoto \u4e09\u6863\u3001\u4fdd\u5f62 80/90/95\u3001MoleculeACE \u5dee\u5f02\u76f8\u5173\u548c\u6269\u5c55\u5931\u8d25\u6848\u4f8b\u5206\u522b\u5bf9\u5e94\u5ba1\u7a3f\u610f\u89c1\u4e2d\u6700\u5bb9\u6613\u88ab\u8ffd\u95ee\u7684\u51e0\u4e2a\u8bc1\u636e\u7f3a\u53e3\u3002",
        size=10,
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S11. \u771f\u5b9e 3x3 nested validation \u4ee3\u8868\u7ec8\u70b9\u6458\u8981\u3002",
        ["\u6570\u636e\u96c6", "\u7c7b\u578b", "\u5916\u5c42\u6298", "\u5185\u5c42\u9009\u62e9", "\u5916\u5c42\u4e3b\u6307\u6807", "\u8f85\u52a9\u6307\u6807"],
        build_true_nested_rows(),
        "\u6ce8\uff1a\u8be5 nested validation \u4ec5\u7528\u4e8e\u5ba1\u8ba1\u6d41\u7a0b\u9694\u79bb\uff0c\u4e0d\u7528\u4e8e\u66ff\u6362\u4e3b\u6587\u4e00\u6b21\u6027\u6d4b\u8bd5\u96c6\u7ed3\u679c\u3002",
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S12. Seed-nested selector audit \u6458\u8981\u3002",
        ["\u6570\u636e\u96c6", "\u5019\u9009\u6c60", "\u7c7b\u578b", "\u5916\u5c42 seed", "\u9009\u4e2d\u6d4b\u8bd5\u5747\u503c", "test-oracle regret", "\u6a21\u578b\u5207\u6362"],
        build_seed_nested_rows(),
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S13. \u7edf\u4e00\u7cfb\u7edf\u6d88\u878d\u77e9\u9635\u3002",
        ["\u77e9\u9635\u9879", "\u7c7b\u578b", "n", "\u65b9\u5411\u7edf\u4e00\u5747\u503c", "\u76f8\u5bf9 Full \u0394", "\u4f18\u4e8e Full"],
        build_ablation_rows(),
        "\u6ce8\uff1a\u56de\u5f52\u4efb\u52a1\u5df2\u8f6c\u6210\u6b63\u5411\u6307\u6807\u540e\u6c47\u603b\uff1b\u56e0\u6b64 \u0394 \u4e3a\u6b63\u8868\u793a\u8f83 Full \u6709\u6539\u5584\u3002",
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S14. \u4e25\u683c\u4e92\u65a5 Tanimoto \u4e09\u6863\u5206\u5c42\u3002",
        ["\u6765\u6e90", "\u7c7b\u578b", "bin", "n", "\u5747\u503c\u76f8\u4f3c\u5ea6", "\u9ad8\u8bef\u5dee\u5bcc\u96c6", "\u4e3b\u8981\u6027\u80fd/\u6821\u51c6"],
        build_tanimoto_rows(),
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S15. \u4fdd\u5f62\u9884\u6d4b 80%/90%/95% \u8986\u76d6\u7387\u6458\u8981\u3002",
        ["\u7c7b\u578b", "\u76ee\u6807", "n", "\u5e73\u5747\u8986\u76d6", "\u4e2d\u4f4d\u8986\u76d6", "\u96c6\u5408/\u533a\u95f4", "\u98ce\u9669\u6458\u8981"],
        build_conformal_rows(),
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S16. MoleculeACE \u9884\u6d4b\u5dee\u5f02 vs \u771f\u5b9e\u5dee\u5f02\u76f8\u5173\u6027\u3002",
        ["\u4efb\u52a1", "seed", "\u60ac\u5d16\u5bf9", "\u9608\u503c", "gap Spearman", "\u65b9\u5411\u51c6\u786e\u7387", "gap MAE"],
        build_moleculeace_rows(),
    )
    insert_picture(
        anchor,
        FIG / "moleculeace_cliff_pair_cases.png",
        "\u8865\u5145\u56fe S1. MoleculeACE \u4ee3\u8868\u6027 cliff pair \u6848\u4f8b\u56fe\u3002\u56fe\u4e2d\u5c55\u793a\u9ad8\u76f8\u4f3c\u5ea6\u4f46\u771f\u5b9e\u6d3b\u6027\u5dee\u5f02\u8f83\u5927\u7684\u5206\u5b50\u5bf9\uff0c\u5e76\u6807\u6ce8\u9884\u6d4b gap \u4e0e\u771f\u5b9e gap \u7684\u504f\u79bb\u3002",
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S17. \u6269\u5c55\u5931\u8d25\u6848\u4f8b\u6458\u8981\u3002",
        ["\u6848\u4f8b", "\u7c7b\u578b", "\u6570\u636e\u96c6", "\u6570\u503c\u8bc1\u636e", "\u89e3\u91ca"],
        build_failure_rows(),
        "\u6ce8\uff1a\u5b8c\u6574 SMILES \u548c\u914d\u5bf9\u5206\u5b50\u4fe1\u606f\u89c1 extended_failure_cases.csv \u548c moleculeace_cliff_pair_cases.csv\u3002",
    )


def write_markdown_summary(path: Path) -> None:
    conformal = pd.read_csv(OUT / "conformal_80_90_95_summary.csv")
    true_nested = pd.read_csv(NESTED / "true_nested_validation_summary.csv")
    tanimoto = pd.read_csv(OUT / "exact_tanimoto_bins_summary.csv")
    ablation = pd.read_csv(OUT / "unified_ablation_matrix_summary.csv")
    nested = pd.read_csv(OUT / "nested_seed_validation_summary.csv")
    moleculeace = pd.read_csv(OUT / "moleculeace_gap_correlation_summary.csv")
    failure = pd.read_csv(OUT / "extended_failure_cases.csv")

    cls_cov = conformal[conformal.task_type == "classification"].sort_values("target_coverage")
    reg_cov = conformal[conformal.task_type == "regression"].sort_values("target_coverage")
    low = tanimoto[tanimoto.similarity_bin == "<0.5"][["source", "task_type", "high_error_enrichment"]]
    ab_drop = ablation[ablation.category == "no_validation_selector_fixed_morgan"]
    mol_agg = moleculeace.groupby("task")["gap_spearman"].mean().describe()

    lines = [
        "# Remaining Experiments Full-Run Summary",
        "",
        "- True nested validation: 9 representative endpoints, 3 outer folds and 3 inner folds per endpoint.",
        f"- Nested classification ROC-AUC examples: BBBP {fmt(true_nested.loc[true_nested.dataset == 'bbbp', 'roc_auc_mean'].iloc[0])}, BACE {fmt(true_nested.loc[true_nested.dataset == 'bace', 'roc_auc_mean'].iloc[0])}, Pgp {fmt(true_nested.loc[true_nested.dataset == 'tdc_pgp_broccatelli', 'roc_auc_mean'].iloc[0])}.",
        f"- Nested regression RMSE examples: Lipo {fmt(true_nested.loc[true_nested.dataset == 'lipo', 'rmse_mean'].iloc[0])}, Caco2 {fmt(true_nested.loc[true_nested.dataset == 'tdc_caco2_wang', 'rmse_mean'].iloc[0])}.",
        f"- Seed-nested selector audit: {len(nested)} dataset-pool summaries.",
        f"- Unified ablation: fixed Morgan/no-selector drops by {fmt(abs(ab_drop[ab_drop.task_type == 'classification']['mean_delta_vs_full_positive'].iloc[0]))} in classification and {fmt(abs(ab_drop[ab_drop.task_type == 'regression']['mean_delta_vs_full_positive'].iloc[0]))} in regression positive-direction metric.",
        f"- Exact Tanimoto <0.5 high-error enrichment: " + "; ".join(f"{r.source}/{r.task_type} {fmt(r.high_error_enrichment)}" for r in low.itertuples(index=False)),
        f"- Conformal classification coverage means at 80/90/95%: {'/'.join(fmt(x) for x in cls_cov.coverage_mean)}.",
        f"- Conformal regression coverage means at 80/90/95%: {'/'.join(fmt(x) for x in reg_cov.coverage_mean)}.",
        f"- MoleculeACE gap Spearman across tasks: mean {fmt(mol_agg['mean'])}, min {fmt(mol_agg['min'])}, max {fmt(mol_agg['max'])}.",
        f"- Failure cases expanded to {len(failure)} rows, adding low-similarity and MoleculeACE cliff failures.",
        "",
        "Key files:",
        f"- {NESTED / 'true_nested_validation_summary.csv'}",
        f"- {OUT / 'unified_ablation_matrix_summary.csv'}",
        f"- {OUT / 'exact_tanimoto_bins_summary.csv'}",
        f"- {OUT / 'conformal_80_90_95_summary.csv'}",
        f"- {OUT / 'moleculeace_gap_correlation_summary.csv'}",
        f"- {OUT / 'extended_failure_cases.csv'}",
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    src = source_docx()
    out = target_docx(src)
    shutil.copy2(src, out)
    doc = Document(out)
    update_core_text(doc)
    update_tracker_table(doc)
    insert_full_run_section(doc)
    doc.save(out)
    write_markdown_summary(OUT / "full_run_summary.md")
    print(out)


if __name__ == "__main__":
    main()
