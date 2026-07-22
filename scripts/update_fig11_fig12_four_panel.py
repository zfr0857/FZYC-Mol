from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches


ROOT = Path("D:/fzyc")
OUT = ROOT / "output"
EXP = OUT / "sci1_mechanism_uq_decision_20260707"
SCRIPT_OUT = OUT / "paper18_fig11_fig12_four_panel"
AUDIT = OUT / "paper18_fig11_fig12_four_panel_audit.json"


COLORS = {
    "fzyc_selected": "#3A6EA5",
    "rdkit_rf": "#6E7A86",
    "molformer_linear_probe": "#4C9A74",
    "chemberta_mtr_linear_probe": "#B86B5E",
    "gnn_gcn": "#9A7BB8",
    "split": "#6E7A86",
    "label": "#3A6EA5",
    "mondrian": "#4C9A74",
    "cqr": "#B86B5E",
}

CAND_LABELS = {
    "fzyc_selected": "FZYC-selected",
    "rdkit_rf": "RDKit-RF",
    "molformer_linear_probe": "MoLFormer",
    "chemberta_mtr_linear_probe": "ChemBERTa",
    "gnn_gcn": "GNN-GCN",
}

METHOD_LABELS = {
    "split_conformal": "Split",
    "label_conditional_conformal": "Label-conditional",
    "mondrian_label_similarity_conformal": "Mondrian",
    "split_conformal_residual": "Residual split",
    "mondrian_similarity_residual": "Mondrian residual",
    "conformalized_quantile_regression_morgan": "CQR",
}


def docx_path() -> Path:
    matches = [p for p in OUT.glob("*.docx") if p.name.endswith("-18.docx")]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paper-18 docx, found {matches}")
    return matches[0]


def setup_plot() -> None:
    plt.rcParams.update(
        {
            "font.family": "Arial",
            "font.size": 11,
            "axes.titlesize": 13,
            "axes.labelsize": 12,
            "xtick.labelsize": 10,
            "ytick.labelsize": 10,
            "legend.fontsize": 9,
            "axes.linewidth": 0.9,
            "savefig.dpi": 600,
            "svg.fonttype": "none",
        }
    )


def savefig(fig: plt.Figure, stem: str) -> tuple[Path, Path]:
    png = SCRIPT_OUT / f"{stem}.png"
    svg = SCRIPT_OUT / f"{stem}.svg"
    fig.savefig(png, bbox_inches="tight", facecolor="white")
    fig.savefig(svg, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return png, svg


def load_frames() -> dict[str, pd.DataFrame]:
    files = {
        "conf": "conformal_crossfold_summary.csv",
        "cqr": "cqr_regression_summary.csv",
        "cal": "calibration_ood_scaffold_summary.csv",
        "ens": "ensemble_uncertainty_summary.csv",
        "clin": "clintox_minority_negative_result.csv",
        "dec": "decision_enrichment_summary.csv",
        "tox": "toxicity_false_negative_cost.csv",
        "queue": "toxicity_queue_simulation.csv",
        "fail": "failure_case_category_summary.csv",
    }
    return {key: pd.read_csv(EXP / name) for key, name in files.items()}


def fig11_source(d: dict[str, pd.DataFrame]) -> pd.DataFrame:
    conf90 = d["conf"][d["conf"]["alpha"].eq(0.10)].copy()
    rdkit_class = conf90[
        conf90["candidate"].eq("rdkit_rf") & conf90["task_type"].eq("classification")
    ]
    panel_a = (
        rdkit_class.groupby("method", as_index=False)
        .agg(
            overall_coverage=("mean_coverage", "mean"),
            class0_coverage=("mean_class_0_coverage", "mean"),
            class1_coverage=("mean_class_1_coverage", "mean"),
            set_size=("mean_set_size", "mean"),
        )
        .assign(panel="minority_class_coverage")
    )

    residual = conf90[
        conf90["candidate"].eq("rdkit_rf") & conf90["task_type"].eq("regression")
    ][
        [
            "task",
            "method",
            "target_coverage",
            "mean_coverage",
            "mean_interval_width",
        ]
    ].copy()
    cqr = d["cqr"][d["cqr"]["alpha"].eq(0.10)][
        ["task", "method", "target_coverage", "mean_coverage", "mean_interval_width"]
    ].copy()
    panel_b = pd.concat([residual, cqr], ignore_index=True).assign(panel="regression_width_coverage")

    panel_c = d["cal"].copy().assign(panel="ood_scaffold_calibration")

    panel_d = d["ens"][~d["ens"]["task"].eq("__overall__")].copy()
    panel_d = panel_d.assign(panel="ensemble_uncertainty_triage")

    rows = []
    for frame in [panel_a, panel_b, panel_c, panel_d]:
        rows.extend(frame.to_dict("records"))
    src = pd.DataFrame(rows)
    src.to_csv(SCRIPT_OUT / "fig11_uq_conformal_4panel_source.csv", index=False)
    return src


def draw_fig11(d: dict[str, pd.DataFrame]) -> tuple[Path, Path, dict[str, float]]:
    fig11_source(d)
    fig, axes = plt.subplots(2, 2, figsize=(12.2, 8.6))
    ax1, ax2, ax3, ax4 = axes.ravel()

    conf90 = d["conf"][d["conf"]["alpha"].eq(0.10)].copy()
    rdkit_class = conf90[
        conf90["candidate"].eq("rdkit_rf") & conf90["task_type"].eq("classification")
    ]
    cov = (
        rdkit_class.groupby("method")[["mean_coverage", "mean_class_1_coverage"]]
        .mean()
        .reindex(
            [
                "split_conformal",
                "label_conditional_conformal",
                "mondrian_label_similarity_conformal",
            ]
        )
    )
    x = np.arange(len(cov))
    width = 0.36
    ax1.bar(
        x - width / 2,
        cov["mean_coverage"],
        width,
        label="Overall",
        color="#B7BEC7",
        edgecolor="#1F2937",
        linewidth=0.4,
    )
    ax1.bar(
        x + width / 2,
        cov["mean_class_1_coverage"],
        width,
        label="Class 1",
        color=COLORS["label"],
        edgecolor="#1F2937",
        linewidth=0.4,
    )
    ax1.axhline(0.90, color="#1F2937", lw=1.0, ls="--")
    ax1.set_ylim(0, 1.04)
    ax1.set_ylabel("Coverage")
    ax1.set_title("Class-conditional coverage")
    ax1.set_xticks(x, ["Split", "Label-cond.", "Mondrian"], rotation=15, ha="right")
    ax1.legend(frameon=False, loc="lower right")

    residual = conf90[
        conf90["candidate"].eq("rdkit_rf") & conf90["task_type"].eq("regression")
    ].copy()
    residual = residual[
        residual["method"].isin(["split_conformal_residual", "mondrian_similarity_residual"])
    ]
    cqr = d["cqr"][d["cqr"]["alpha"].eq(0.10)].copy()
    reg = pd.concat(
        [
            residual[
                [
                    "task",
                    "method",
                    "mean_coverage",
                    "mean_interval_width",
                ]
            ],
            cqr[
                [
                    "task",
                    "method",
                    "mean_coverage",
                    "mean_interval_width",
                ]
            ],
        ],
        ignore_index=True,
    )
    for method, group in reg.groupby("method"):
        label = METHOD_LABELS.get(method, method)
        color = (
            COLORS["split"]
            if "split" in method
            else COLORS["mondrian"]
            if "mondrian" in method
            else COLORS["cqr"]
        )
        ax2.scatter(
            group["mean_interval_width"],
            group["mean_coverage"],
            s=52,
            color=color,
            edgecolor="#1F2937",
            linewidth=0.45,
            label=label,
            alpha=0.92,
        )
    ax2.axhline(0.90, color="#1F2937", lw=1.0, ls="--")
    ax2.set_xlabel("Mean interval width")
    ax2.set_ylabel("Coverage")
    ax2.set_title("Regression interval trade-off")
    ax2.set_ylim(0.76, 0.98)
    ax2.legend(frameon=False, loc="lower right")

    cal = d["cal"].copy()
    order = ["<0.5", "0.5-0.7", ">0.7"]
    cal_g = cal.groupby("tanimoto_bin")[["mean_roc_auc", "mean_ece"]].mean().reindex(order)
    x = np.arange(len(order))
    ax3.bar(
        x - 0.18,
        cal_g["mean_roc_auc"],
        0.36,
        label="ROC-AUC",
        color=COLORS["mondrian"],
        edgecolor="#1F2937",
        linewidth=0.4,
    )
    ax3.bar(
        x + 0.18,
        cal_g["mean_ece"],
        0.36,
        label="ECE",
        color=COLORS["cqr"],
        edgecolor="#1F2937",
        linewidth=0.4,
    )
    ax3.set_xticks(x, order)
    ax3.set_ylim(0, 1.0)
    ax3.set_xlabel("Nearest-neighbour Tanimoto")
    ax3.set_ylabel("Metric value")
    ax3.set_title("OOD/scaffold calibration")
    ax3.legend(frameon=False, loc="upper center", ncol=2)

    ens = d["ens"][~d["ens"]["task"].eq("__overall__")].copy()
    ens["task"] = ens["task"].str.upper()
    ens = ens.sort_values("mean_top10_high_error_enrichment")
    ax4.barh(
        ens["task"],
        ens["mean_top10_high_error_enrichment"],
        color="#8EA4C8",
        edgecolor="#1F2937",
        linewidth=0.4,
    )
    ax4.axvline(1.0, color="#1F2937", lw=1.0, ls="--")
    ax4.set_xlabel("Top-10% error enrichment")
    ax4.set_title("Ensemble uncertainty triage")
    ax4.set_xlim(0, max(2.7, ens["mean_top10_high_error_enrichment"].max() + 0.25))

    for ax in axes.ravel():
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.grid(axis="y", color="#E7EBF0", lw=0.8)
        ax.set_axisbelow(True)

    fig.tight_layout(w_pad=2.2, h_pad=2.0)
    png, svg = savefig(fig, "fig11_uq_conformal_4panel")

    rdkit = rdkit_class.groupby("method")[["mean_coverage", "mean_class_1_coverage"]].mean()
    metrics = {
        "split_class1_coverage": float(rdkit.loc["split_conformal", "mean_class_1_coverage"]),
        "label_class1_coverage": float(
            rdkit.loc["label_conditional_conformal", "mean_class_1_coverage"]
        ),
        "mondrian_class1_coverage": float(
            rdkit.loc["mondrian_label_similarity_conformal", "mean_class_1_coverage"]
        ),
        "cqr_90_mean_coverage": float(d["cqr"][d["cqr"]["alpha"].eq(0.10)]["mean_coverage"].mean()),
        "cqr_90_mean_width": float(
            d["cqr"][d["cqr"]["alpha"].eq(0.10)]["mean_interval_width"].mean()
        ),
        "low_similarity_auc": float(cal_g.loc["<0.5", "mean_roc_auc"]),
        "high_similarity_auc": float(cal_g.loc[">0.7", "mean_roc_auc"]),
        "ensemble_top10_enrichment": float(
            d["ens"].loc[d["ens"]["task"].eq("__overall__"), "mean_top10_high_error_enrichment"].iloc[0]
        ),
    }
    return png, svg, metrics


def fig12_source(d: dict[str, pd.DataFrame]) -> pd.DataFrame:
    dec = d["dec"].copy().assign(panel="enrichment_and_regret")
    tox = d["tox"].copy().assign(panel="toxicity_cost")
    queue = d["queue"].copy().assign(panel="queue_simulation")
    fail = d["fail"].copy().assign(panel="failure_taxonomy")
    src = pd.concat([dec, tox, queue, fail], ignore_index=True, sort=False)
    src.to_csv(SCRIPT_OUT / "fig12_decision_value_4panel_source.csv", index=False)
    return src


def draw_fig12(d: dict[str, pd.DataFrame]) -> tuple[Path, Path, dict[str, float]]:
    fig12_source(d)
    fig, axes = plt.subplots(2, 2, figsize=(12.2, 8.6))
    ax1, ax2, ax3, ax4 = axes.ravel()

    dec = d["dec"].copy()
    keep = [
        "fzyc_selected",
        "molformer_linear_probe",
        "chemberta_mtr_linear_probe",
        "gnn_gcn",
        "rdkit_rf",
    ]
    budgets = [0.01, 0.05, 0.10]
    for candidate in keep:
        g = (
            dec[dec["candidate"].eq(candidate)]
            .groupby("budget_fraction", as_index=False)["mean_enrichment"]
            .mean()
            .sort_values("budget_fraction")
        )
        ax1.plot(
            g["budget_fraction"] * 100,
            g["mean_enrichment"],
            marker="o",
            lw=2.0,
            color=COLORS.get(candidate, "#6E7A86"),
            label=CAND_LABELS[candidate],
        )
    ax1.set_xticks([1, 5, 10])
    ax1.set_xlabel("Screening budget (%)")
    ax1.set_ylabel("Mean enrichment")
    ax1.set_title("Top-budget enrichment")
    ax1.legend(frameon=False, ncol=2, loc="upper right")

    regret = (
        dec.groupby(["candidate", "budget_fraction"], as_index=False)["mean_regret_vs_oracle"].mean()
    )
    for candidate in keep:
        g = regret[regret["candidate"].eq(candidate)].sort_values("budget_fraction")
        ax2.plot(
            g["budget_fraction"] * 100,
            g["mean_regret_vs_oracle"],
            marker="o",
            lw=2.0,
            color=COLORS.get(candidate, "#6E7A86"),
            label=CAND_LABELS[candidate],
        )
    ax2.set_xticks([1, 5, 10])
    ax2.set_xlabel("Screening budget (%)")
    ax2.set_ylabel("Missed positives vs oracle")
    ax2.set_title("Fixed-budget regret")

    tox = d["tox"].copy()
    tox_g = (
        tox.groupby("candidate", as_index=False)
        .agg(
            cost_mean=("cost_per_100_molecules", "mean"),
            cost_sem=("cost_per_100_molecules", lambda x: x.std(ddof=1) / np.sqrt(len(x))),
        )
        .set_index("candidate")
        .reindex(keep)
        .reset_index()
    )
    x = np.arange(len(tox_g))
    ax3.bar(
        x,
        tox_g["cost_mean"],
        yerr=tox_g["cost_sem"],
        capsize=3,
        color=[COLORS.get(c, "#6E7A86") for c in tox_g["candidate"]],
        edgecolor="#1F2937",
        linewidth=0.4,
    )
    ax3.set_xticks(x, [CAND_LABELS[c] for c in tox_g["candidate"]], rotation=25, ha="right")
    ax3.set_ylabel("Cost per 100 molecules")
    ax3.set_title("Toxicity false-negative cost")

    queue = d["queue"].copy()
    queue_g = (
        queue.groupby(["candidate", "advance_budget_fraction"], as_index=False)[
            "toxic_advanced_rate"
        ].mean()
    )
    for candidate in keep:
        g = queue_g[queue_g["candidate"].eq(candidate)].sort_values("advance_budget_fraction")
        ax4.plot(
            g["advance_budget_fraction"] * 100,
            g["toxic_advanced_rate"] * 100,
            marker="o",
            lw=2.0,
            color=COLORS.get(candidate, "#6E7A86"),
            label=CAND_LABELS[candidate],
        )
    ax4.set_xticks([1, 5, 10])
    ax4.set_xlabel("Advanced queue budget (%)")
    ax4.set_ylabel("Toxic advanced rate (%)")
    ax4.set_title("Experimental queue simulation")

    for ax in axes.ravel():
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.grid(axis="y", color="#E7EBF0", lw=0.8)
        ax.set_axisbelow(True)

    fig.tight_layout(w_pad=2.2, h_pad=2.0)
    png, svg = savefig(fig, "fig12_decision_value_4panel")

    fzyc = dec[dec["candidate"].eq("fzyc_selected")].groupby("budget_fraction")[
        "mean_enrichment"
    ].mean()
    mol = dec[dec["candidate"].eq("molformer_linear_probe")].groupby("budget_fraction")[
        "mean_enrichment"
    ].mean()
    q = queue_g.set_index(["candidate", "advance_budget_fraction"])["toxic_advanced_rate"]
    metrics = {
        "fzyc_ef1": float(fzyc.loc[0.01]),
        "fzyc_ef5": float(fzyc.loc[0.05]),
        "fzyc_ef10": float(fzyc.loc[0.10]),
        "molformer_ef10": float(mol.loc[0.10]),
        "rdkit_tox_cost": float(tox_g.loc[tox_g["candidate"].eq("rdkit_rf"), "cost_mean"].iloc[0]),
        "molformer_tox_cost": float(
            tox_g.loc[tox_g["candidate"].eq("molformer_linear_probe"), "cost_mean"].iloc[0]
        ),
        "gnn_tox_cost": float(tox_g.loc[tox_g["candidate"].eq("gnn_gcn"), "cost_mean"].iloc[0]),
        "fzyc_queue_toxic_rate_10": float(q.loc[("fzyc_selected", 0.10)]),
        "molformer_queue_toxic_rate_10": float(q.loc[("molformer_linear_probe", 0.10)]),
        "failure_categories": int(d["fail"]["category"].nunique()),
        "failure_cases": int(d["fail"]["n_cases"].sum()),
    }
    return png, svg, metrics


def replace_text(doc: Document, old_start: str, new_text: str) -> None:
    for p in doc.paragraphs:
        if p.text.strip().startswith(old_start):
            p.text = new_text
            return
    raise RuntimeError(f"Paragraph not found: {old_start}")


def replace_image_before_caption(doc: Document, caption_start: str, image_path: Path, width: float) -> None:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(caption_start):
            image_p = doc.paragraphs[i - 1]
            image_p.clear()
            image_p.add_run().add_picture(str(image_path), width=Inches(width))
            return
    raise RuntimeError(f"Caption not found: {caption_start}")


def update_docx(fig11_png: Path, fig12_png: Path, m11: dict[str, float], m12: dict[str, float]) -> Path:
    target = docx_path()
    backup = OUT / f"{target.stem}_图11图12四联图前备份_{datetime.now():%Y%m%d_%H%M%S}.docx"
    shutil.copy2(target, backup)
    doc = Document(str(target))

    replace_text(
        doc,
        "可靠性实验从总体覆盖扩展",
        (
            "可靠性实验进一步扩展为四个互补面板：标签条件与 Mondrian 保形检验少数类覆盖，"
            "回归保形比较覆盖率与区间宽度，scaffold/OOD 分层评估校准漂移，ensemble uncertainty "
            "评估高误差样本的提前识别能力。RDKit-RF 在 90% 目标覆盖下的分类 split conformal "
            f"类别 1 覆盖为 {m11['split_class1_coverage']:.3f}；label-conditional 和 Mondrian "
            f"label-similarity conformal 将类别 1 覆盖提高到 {m11['label_class1_coverage']:.3f} "
            f"和 {m11['mondrian_class1_coverage']:.3f}。CQR 的 90% 平均覆盖为 "
            f"{m11['cqr_90_mean_coverage']:.3f}，平均区间宽度为 {m11['cqr_90_mean_width']:.2f}，"
            "提示其在当前特征和样本量下并未稳定优于残差式保形。"
        ),
    )
    replace_text(
        doc,
        "scaffold/OOD 校准进一步显示",
        (
            "scaffold/OOD 校准和不确定性排序给出了更细的边界证据。最近邻 Tanimoto <0.5 "
            f"的分类子集平均 ROC-AUC 为 {m11['low_similarity_auc']:.3f}，低于 >0.7 子集的 "
            f"{m11['high_similarity_auc']:.3f}；ensemble uncertainty 对 top-10% 高误差样本的"
            f"平均富集为 {m11['ensemble_top10_enrichment']:.2f}。ClinTox 少数类仍作为负结果保留："
            "RDKit-RF 的少数类召回不足，说明覆盖校准可以改善风险声明，却不能单独替代阈值式毒性筛选。"
        ),
    )
    replace_text(
        doc,
        "图 11",
        (
            "图 11  不确定性和保形预测的四联图。左上，RDKit-RF 在 90% 目标覆盖下的总体覆盖与类别 1 覆盖；"
            "右上，回归 residual conformal、Mondrian residual 与 CQR 的覆盖率-区间宽度权衡；"
            "左下，最近邻 Tanimoto 分层下的 OOD/scaffold ROC-AUC 和 ECE；右下，ensemble uncertainty "
            "对高误差样本的 top-10% 富集。"
        ),
    )
    replace_text(
        doc,
        "真实决策价值实验将模型分数转化",
        (
            "真实决策价值实验也扩展为四个面板，将模型分数转化为固定预算筛选收益、相对 oracle 的遗漏阳性数、"
            "毒性假阴性成本和实验队列风险。FZYC-selected 在该六任务强基线面板中等同于 RDKit-RF，"
            f"其 top-1%、top-5% 和 top-10% 平均富集分别为 {m12['fzyc_ef1']:.2f}、"
            f"{m12['fzyc_ef5']:.2f} 和 {m12['fzyc_ef10']:.2f}；MoLFormer 的 top-10% 富集为 "
            f"{m12['molformer_ef10']:.2f}。但毒性假阴性成本给出不同排序：RDKit-RF/FZYC、MoLFormer "
            f"和 GNN-GCN 的 ClinTox 阈值成本分别为每 100 个分子 {m12['rdkit_tox_cost']:.1f}、"
            f"{m12['molformer_tox_cost']:.1f} 和 {m12['gnn_tox_cost']:.1f}。"
        ),
    )
    replace_text(
        doc,
        "图 12",
        (
            "图 12  真实筛选决策价值的四联图。左上，top-1%、top-5% 和 top-10% 固定预算富集；"
            "右上，相同预算下相对 test oracle 的遗漏阳性数；左下，ClinTox 假阴性加权成本；"
            "右下，在低毒候选进入实验队列的模拟中，高风险分子进入队列的比例。"
        ),
    )

    replace_image_before_caption(doc, "图 11", fig11_png, 6.6)
    replace_image_before_caption(doc, "图 12", fig12_png, 6.6)
    doc.save(str(target))
    return backup


def audit_docx(target: Path) -> dict[str, object]:
    with ZipFile(target) as zf:
        bad = zf.testzip()
        media_count = sum(n.startswith("word/media/") for n in zf.namelist())
    doc = Document(str(target))
    text = "\n".join(p.text for p in doc.paragraphs)
    return {
        "target": str(target),
        "zip_ok": bad is None,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "figures": len(doc.inline_shapes),
        "media_count": media_count,
        "fig11_caption_updated": "不确定性和保形预测的四联图" in text,
        "fig12_caption_updated": "真实筛选决策价值的四联图" in text,
        "source_data_fig11": str(SCRIPT_OUT / "fig11_uq_conformal_4panel_source.csv"),
        "source_data_fig12": str(SCRIPT_OUT / "fig12_decision_value_4panel_source.csv"),
    }


def write_report(backup: Path, metrics: dict[str, float], audit: dict[str, object]) -> Path:
    report = OUT / "小论文-18_图11图12四联图加厚报告.md"
    lines = [
        "# 小论文-18 图11和图12四联图加厚报告",
        "",
        f"- 生成时间：{datetime.now().isoformat(timespec='seconds')}",
        f"- 更新文档：`{docx_path()}`",
        f"- 备份文档：`{backup}`",
        f"- 图表目录：`{SCRIPT_OUT}`",
        "",
        "## 图11新增内容",
        "",
        "- 将保形/UQ结果扩展为四个小图：少数类覆盖、回归覆盖-宽度权衡、OOD/scaffold校准、ensemble uncertainty高误差富集。",
        f"- split conformal 类别1覆盖 {metrics['fig11']['split_class1_coverage']:.3f}；label-conditional 和 Mondrian 提高到 {metrics['fig11']['label_class1_coverage']:.3f} 和 {metrics['fig11']['mondrian_class1_coverage']:.3f}。",
        f"- CQR 90%平均覆盖 {metrics['fig11']['cqr_90_mean_coverage']:.3f}，平均区间宽度 {metrics['fig11']['cqr_90_mean_width']:.2f}。",
        "",
        "## 图12新增内容",
        "",
        "- 将真实决策价值扩展为四个小图：固定预算富集、相对oracle遗漏阳性、毒性假阴性成本、实验队列模拟。",
        f"- FZYC/RDKit top-1/5/10%富集为 {metrics['fig12']['fzyc_ef1']:.2f}/{metrics['fig12']['fzyc_ef5']:.2f}/{metrics['fig12']['fzyc_ef10']:.2f}。",
        f"- 毒性成本为 RDKit/FZYC {metrics['fig12']['rdkit_tox_cost']:.1f}、MoLFormer {metrics['fig12']['molformer_tox_cost']:.1f}、GNN-GCN {metrics['fig12']['gnn_tox_cost']:.1f}/100 molecules。",
        "",
        "## 审计",
        "",
        "```json",
        json.dumps(audit, ensure_ascii=False, indent=2),
        "```",
    ]
    report.write_text("\n".join(lines), encoding="utf-8-sig")
    return report


def main() -> None:
    SCRIPT_OUT.mkdir(parents=True, exist_ok=True)
    setup_plot()
    data = load_frames()
    fig11_png, fig11_svg, m11 = draw_fig11(data)
    fig12_png, fig12_svg, m12 = draw_fig12(data)
    backup = update_docx(fig11_png, fig12_png, m11, m12)
    target = docx_path()
    audit = audit_docx(target)
    audit.update(
        {
            "fig11_png": str(fig11_png),
            "fig11_svg": str(fig11_svg),
            "fig12_png": str(fig12_png),
            "fig12_svg": str(fig12_svg),
            "backup": str(backup),
            "metrics": {"fig11": m11, "fig12": m12},
        }
    )
    audit["passed"] = (
        audit["zip_ok"]
        and audit["fig11_caption_updated"]
        and audit["fig12_caption_updated"]
        and Path(audit["source_data_fig11"]).exists()
        and Path(audit["source_data_fig12"]).exists()
    )
    AUDIT.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    report = write_report(backup, {"fig11": m11, "fig12": m12}, audit)
    print(
        json.dumps(
            {
                "docx": str(target),
                "report": str(report),
                "fig11_png": str(fig11_png),
                "fig12_png": str(fig12_png),
                "passed": audit["passed"],
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
