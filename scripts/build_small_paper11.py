from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree

import numpy as np
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
SRC = OUT / "小论文-8.docx"
DOCX = OUT / "小论文-11.docx"
ANALYSIS = OUT / "小论文-11_补充分析"
THICK_ANALYSIS = OUT / "小论文-11_加厚实验"
AUDIT = ROOT / "results" / "audits" / "small_paper_11_audit.json"
THICK_AUDIT = ROOT / "results" / "audits" / "small_paper_11_thickened_experiments.json"


TITLE = "分子性质预测中候选池扩张的选择风险：嵌套验证与冻结治理框架"

ABSTRACT = (
    "分子性质预测常在有限验证信息上从不断扩张的候选池中冻结最终模型；候选增加可能提高可达到性能上界，"
    "也可能降低验证排序对外层表现的保真度。FZYC-Mol（Frozen validation governance for molecular model selection）"
    "在本文中指候选登记、嵌套选择、策略冻结、外层审计和负结果记录组成的验证治理协议，而不是新的预测主干网络。"
    "在九个终点的 3 外层×3 内层×5 重复确认性实验中，K=32 相对 K=4 的完整池范围归一化选择损失增加 0.122"
    "（端点聚类 95% CI 0.072–0.175；精确 P=0.0078；Holm P=0.039），机会校正 Top-3 命中率下降 0.642。"
    "随机排序负对照和六级信号恢复正对照排除了固定 Top-3 机会率及指标失真的替代解释。共享冻结划分上的 12 个多视图候选"
    "完成 6,480 次拟合，validation-best 相对 Morgan-only 的配对效用增益为 0.343（0.210–0.483；9/9 终点）。"
    "严格嵌套的跨端点元风险仅作为探索性预警，在留出端点上将 MAE 从 0.123 降至 0.112，并在 8/9 个终点降低 50% 门控遗憾。"
    "TDC、逐样本风险、标签条件保形、MoleculeACE 和 bRo5 分析进一步界定可靠性和化学外推边界。结果表明，候选扩张同时带来收益上界和选择损失；"
    "冻结治理能够暴露并部分管理这一权衡，但不保证跨终点普遍性能提升，也不能替代统一划分下的深度模型重训或时间外前瞻验证。"
)

KEYWORDS = (
    "关键词：分子性质预测；模型选择风险；候选池扩张；冻结治理；嵌套验证；多视图候选；"
    "选择损失；保形预测；化学边界"
)

INTRO = [
    (
        "分子性质预测已从公开基准上的离线性能比较，进入候选排序、实验排队和风险复核等真实决策场景。"
        "随着 Morgan 指纹、MACCS、RDKit2D 描述符、树模型、图模型、D-MPNN、预训练化学语言模型和 AutoML 被纳入同一比较框架，"
        "研究者获得了更高的可达到性能上界，也增加了在同一验证信息上反复选择的自由度。"
    ),
    (
        "这一风险不同于直接测试集泄漏。即使测试标签从未参与训练，验证集仍可能在模型家族、超参数、融合策略、阈值和补救规则之间被持续消费。"
        "因此，最终测试分数可能混合真实表征收益、候选数量机会效应和事后选择偏差。对分子任务而言，小样本毒性终点、骨架分布变化、低相似度分子和活性悬崖又会进一步放大验证排序的不稳定性。"
    ),
    (
        "已有交叉验证偏差、nested CV、AutoML、适用域、不确定性估计、保形预测和多模态分子学习研究为这一问题提供了基础。"
        "但这些工具常被放在模型性能叙事之下，缺少一个同时记录候选资格、选择时点、外层审计、负结果和逐样本边界的证据链。"
        "本文的新意不在于单独发明这些组件，而在于把它们用于候选池扩张的受控实验，并形成可复核的分子性质预测模型选择日志。"
    ),
    (
        "本文的唯一主命题是：在验证信息固定的条件下，候选池扩张虽然能够提高可达到性能上界，但会降低验证排序保真度并增加模型选择损失；冻结式验证治理可以量化并部分管理这一权衡。"
        "围绕该命题，本文把候选池扩张、随机排序负对照、信号恢复正对照、端点层配对推断和共享划分多视图验证作为确认性主证据；"
        "AutoGluon、跨端点元风险、TDC、逐样本风险、保形预测、MoleculeACE 和 bRo5 则被定位为次要或探索性边界证据。"
    ),
]

DISCUSSION = [
    ("4.1 候选扩张为何同时增加上界与选择风险", "Heading 2"),
    (
        "候选扩张不是简单的“越多越好”或“越多越差”。在本文的轻量 32 候选池中，完整池 oracle 上界随 K 增加而上升，"
        "但验证选择未能等比例兑现这一上界，选择损失同步增加。新增有效多样性分析显示，32 个轻量候选的内层效用相关性中位数为 0.998，"
        "有效候选数约为 1.11，说明该实验主要检验候选数量和超参数自由度带来的选择压力，而不是声称覆盖真实药物发现中的全部模型异质性。"
        "因此，多视图共享划分实验被单独作为异质候选确认，而不是与轻量扩池混成一个总排行榜。"
    ),
    ("4.2 validation-best 为何优于部分保守规则", "Heading 2"),
    (
        "one-SE、风险调整、低方差和稳定性规则在本文中没有形成普适优势。一个合理解释是，分子任务的内层样本量有限，"
        "方差估计和复杂度惩罚本身可能带有噪声；过强的保守惩罚会放弃真实可兑现收益。validation-best 在多视图确认中取得最低平均遗憾，"
        "但这并不意味着它在所有终点上都是最优规则。冻结治理的作用是约束事后自由度并保存失败记录，而不是预设唯一正确的选择器。"
    ),
    ("4.3 分子任务为何更容易出现选择不稳定", "Heading 2"),
    (
        "分子性质预测的选择风险来自数据和化学空间两端。骨架划分会改变训练和测试分子的局部结构分布；ClinTox 等少数类毒性终点对召回、校准和保形覆盖更敏感；"
        "MoleculeACE 和低相似度案例显示，高相似分子对也可能存在幅度不连续；bRo5 外缘和时间划分进一步提示随机划分性能不能直接外推到更困难的化学迁移场景。"
        "这些结果支持把逐样本风险、最近邻相似度、保形集合或区间、失败原因码与终点级选择日志并列报告。"
    ),
    ("4.4 实际使用边界", "Heading 2"),
    (
        "FZYC-Mol 不替代预测模型，不保证性能提升，也不提供可迁移到所有终点的元选择器。Chemprop、GNN、ChemBERTa 和 MoLFormer 等强模型尚未在本文统一外层划分上完成确认性重训，"
        "因此只能作为候选来源和后续扩展方向，不能写入当前确认性效应估计。TDC 三种子结果只能反映种子变异和公开面板异质性，不能承担严格抽样推断。"
        "公开 release、Zenodo DOI 和第三方冷启动复跑仍需在正式投稿前完成，否则开放治理主张与交付物之间会存在缺口。"
    ),
]

CONCLUSION = [
    (
        "本文将 FZYC-Mol 严格限定为一种分子性质预测模型选择的冻结治理框架。九个终点的重复嵌套实验、随机排序负对照、信号恢复正对照和端点层配对推断共同支持一个限定性结论："
        "在固定验证信息下，候选池扩张会提高可达到上界，也会增加选择损失并降低验证排序保真度。共享划分多视图实验显示，异质表示可以在冻结选择下兑现收益，但其收益必须与新增选择自由度同时报告。"
    ),
    (
        "对后续研究而言，最直接的采用方式不是替换现有预测器，而是在新候选进入比较之前建立候选登记、终止规则、内层选择、外层审计和负结果归档。"
        "若未来加入统一划分重训的 Chemprop、GNN、化学语言模型或时间外 ADMET 盲测，这一框架仍可作为扩展前的审计底座。"
    ),
]


def set_para_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")


def insert_after(p, text: str, style: str | None = None):
    new = p.insert_paragraph_before("")
    p._p.addprevious(new._p)
    # The paragraph was inserted before p; move it after p.
    p._p.addnext(new._p)
    if style:
        new.style = style
    set_para_text(new, text)
    new.paragraph_format.first_line_indent = Cm(0.74) if style != "Heading 2" else Cm(0)
    new.paragraph_format.space_after = Pt(4)
    return new


def delete_paragraph(p) -> None:
    el = p._element
    el.getparent().remove(el)


def paragraph_has_drawing(p) -> bool:
    return bool(p._element.xpath(".//w:drawing"))


def find_para(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"paragraph not found: {text}")


def delete_between(doc: Document, start_text: str, end_text: str) -> None:
    paras = list(doc.paragraphs)
    start = next(i for i, p in enumerate(paras) if p.text.strip() == start_text)
    end = next(i for i, p in enumerate(paras) if p.text.strip() == end_text)
    for p in paras[start + 1 : end]:
        delete_paragraph(p)


def insert_block_after_heading(doc: Document, heading: str, block: list[str | tuple[str, str]]) -> None:
    anchor = find_para(doc, heading)
    last = anchor
    for item in block:
        if isinstance(item, tuple):
            text, style = item
        else:
            text, style = item, "Normal"
        last = insert_after(last, text, style)


def replace_section(doc: Document, start: str, end: str, block: list[str | tuple[str, str]]) -> None:
    delete_between(doc, start, end)
    insert_block_after_heading(doc, start, block)


def replace_first_containing(doc: Document, needle: str, text: str) -> None:
    for p in doc.paragraphs:
        if needle in p.text:
            set_para_text(p, text)
            return
    raise ValueError(f"needle not found: {needle}")


def insert_after_first_containing(doc: Document, needle: str, additions: list[str]) -> None:
    for p in doc.paragraphs:
        if needle in p.text:
            last = p
            for text in additions:
                last = insert_after(last, text)
            return
    raise ValueError(f"needle not found: {needle}")


def remove_figure(doc: Document, number: int) -> None:
    cap_re = re.compile(rf"^图\s*{number}\b")
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if cap_re.match(p.text.strip()):
            if i > 0 and paragraph_has_drawing(paras[i - 1]):
                delete_paragraph(paras[i - 1])
            delete_paragraph(p)


def compute_analysis() -> dict[str, float]:
    ANALYSIS.mkdir(parents=True, exist_ok=True)
    root = ROOT
    inns, outs = [], []
    for f in sorted((root / "results/nested_selection/repeated_nested").glob("seed_*/inner_scores.csv")):
        df = pd.read_csv(f)
        df["repeat_seed"] = int(f.parent.name.split("_")[-1])
        inns.append(df)
    for f in sorted((root / "results/nested_selection/repeated_nested").glob("seed_*/outer_candidate_scores.csv")):
        df = pd.read_csv(f)
        df["repeat_seed"] = int(f.parent.name.split("_")[-1])
        outs.append(df)
    inner = pd.concat(inns, ignore_index=True)
    outer = pd.concat(outs, ignore_index=True)

    def keff(df, value_col):
        idx = ["dataset", "repeat_seed", "outer_fold"] + (["inner_fold"] if "inner_fold" in df.columns else [])
        # Candidate names differ slightly between classification and regression
        # for the same registry slot. Align by the frozen registry order so the
        # effective diversity audit reflects the nominal K=32 pool.
        piv = df.pivot_table(index=idx, columns="candidate_order", values=value_col, aggfunc="mean")
        corr = piv.corr(min_periods=10).fillna(0)
        np.fill_diagonal(corr.values, 1)
        vals = np.clip(np.linalg.eigvalsh(corr.values), 1e-12, None)
        off = corr.values[np.triu_indices_from(corr.values, k=1)]
        return {
            "nominal_K": float(piv.shape[1]),
            "K_eff": float((vals.sum() ** 2) / (vals @ vals)),
            "median_corr": float(np.median(off)),
        }

    inner_k = keff(inner, "inner_utility")
    outer_k = keff(outer, "outer_utility")
    pd.DataFrame([{"measure": "inner-validation utility", **inner_k}, {"measure": "outer-test utility", **outer_k}]).to_csv(
        ANALYSIS / "candidate_effective_diversity.csv", index=False
    )

    regret = pd.read_csv(root / "results/candidate_pool/subset_regret_long.csv")
    ro = regret[(regret["mode"].eq("random_order")) & (regret["subset_seed"].eq(0))].copy()
    base = ro[ro.pool_size.eq(4)][
        ["endpoint", "repeat", "outer_fold", "oracle_test_utility", "selected_test_utility"]
    ].rename(columns={"oracle_test_utility": "oracle4", "selected_test_utility": "selected4"})
    merged = ro.merge(base, on=["endpoint", "repeat", "outer_fold"], how="left")
    merged["G_attain"] = merged["oracle_test_utility"] - merged["oracle4"]
    merged["L_select"] = merged["oracle_test_utility"] - merged["selected_test_utility"]
    merged["G_realized"] = merged["selected_test_utility"] - merged["selected4"]
    merged["eta"] = np.where(np.abs(merged["G_attain"]) > 1e-12, merged["G_realized"] / merged["G_attain"], np.nan)
    gain = (
        merged.groupby("pool_size")
        .agg(
            G_attain=("G_attain", "mean"),
            L_select=("L_select", "mean"),
            G_realized=("G_realized", "mean"),
            eta_median=("eta", "median"),
            fixed_regret=("fixed_normalized_regret", "mean"),
            dynamic_regret=("dynamic_normalized_regret", "mean"),
            full_range_median=("full_range", "median"),
        )
        .reset_index()
    )
    gain.to_csv(ANALYSIS / "gain_decomposition.csv", index=False)

    tdc = pd.read_csv(root / "results/source_data/tdc_gate_audit.csv")
    conf = pd.read_csv(root / "results/source_data/conformal_long.csv")
    conf_summary = (
        conf.groupby(["task_type", "target_coverage"])
        .agg(
            coverage_mean=("coverage", "mean"),
            class0=("class_0_coverage", "mean"),
            class1=("class_1_coverage", "mean"),
            avg_set_size=("avg_set_size", "mean"),
            empty_rate=("empty_rate", "mean"),
            fallback_count=("fallback_reason", lambda s: s.notna().sum()),
            n=("coverage", "size"),
        )
        .reset_index()
    )
    conf_summary.to_csv(ANALYSIS / "conformal_summary_recomputed.csv", index=False)

    if not THICK_AUDIT.exists() or not (THICK_ANALYSIS / "heterogeneous_staged_pool_summary.csv").exists():
        subprocess.run([sys.executable, str(ROOT / "scripts" / "thicken_paper11_experiments.py")], check=True)
    thick = json.loads(THICK_AUDIT.read_text(encoding="utf-8"))

    return {
        "inner_keff": inner_k["K_eff"],
        "inner_corr": inner_k["median_corr"],
        "outer_keff": outer_k["K_eff"],
        "outer_corr": outer_k["median_corr"],
        "g32_attain": float(gain.loc[gain.pool_size.eq(32), "G_attain"].iloc[0]),
        "g32_select": float(gain.loc[gain.pool_size.eq(32), "L_select"].iloc[0]),
        "g32_realized": float(gain.loc[gain.pool_size.eq(32), "G_realized"].iloc[0]),
        "g32_eta": float(gain.loc[gain.pool_size.eq(32), "eta_median"].iloc[0]),
        "tdc_promoted": int(tdc["promoted"].sum()),
        "tdc_cross_zero_promoted": int(((tdc["promoted"]) & (tdc["ci_low"] * tdc["ci_high"] <= 0)).sum()),
        "thick_multiview_K": int(thick["multiview_stage_K"]),
        "thick_multiview_units": int(thick["multiview_n_units"]),
        "thick_multiview_G_attain": float(thick["multiview_G_attain"]),
        "thick_multiview_L_select": float(thick["multiview_L_select"]),
        "thick_multiview_G_realized": float(thick["multiview_G_realized"]),
        "thick_multiview_eta": float(thick["multiview_eta"]),
        "thick_completion_tables": int(thick["copied_completion_tables"]),
        "thick_strong_baseline_rows": int(thick["strong_baseline_rows"]),
        "thick_conformal_rows": int(thick["conformal_rows"]),
        "thick_cleaning_events": int(thick["cleaning_events"]),
    }


def update_tables(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text
                if "95% 配对区间" in txt:
                    cell.text = txt.replace("95% 配对区间", "三种子描述性区间")
                if "固定分母遗憾（95% CI）" in txt:
                    cell.text = txt.replace("固定分母遗憾（95% CI）", "完整池范围归一化选择损失（95% CI）")
                if "固定分母遗憾" in txt and "完整池" not in txt:
                    cell.text = txt.replace("固定分母遗憾", "完整池范围归一化选择损失")


def revise_document(metrics: dict[str, float]) -> None:
    shutil.copy2(SRC, DOCX)
    doc = Document(DOCX)
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "根据修改意见修订：候选池扩张、冻结治理、候选有效多样性、收益分解与边界分析"

    set_para_text(doc.paragraphs[0], TITLE)
    replace_first_containing(doc, "分子性质预测已成为候选化合物排序", ABSTRACT)
    replace_first_containing(doc, "关键词：", KEYWORDS)

    replace_section(doc, "1 引言", "2 材料与方法", INTRO)

    heading_map = {
        "2.1 研究范围、数据集与任务登记": "2.1 Preregistered tasks and candidate pools（任务登记与候选池）",
        "2.2 数据标准化、重复处理与泄漏审计": "2.2 Nested evaluation and leakage control（嵌套评估与泄漏控制）",
        "2.3 分子表示、候选专家与冻结登记": "2.3 Model-selection strategies and primary outcomes（选择策略与主要结果）",
        "2.4 验证治理选择器的形式化定义": "2.4 Controls and statistical inference（对照与统计推断）",
        "2.5 嵌套验证、候选池压力与选择器对照": "2.5 Multi-view confirmation and strong baselines（多视图确认与强基线）",
        "2.6 逐样本预测风险、校准与保形预测": "2.6 Secondary reliability and chemical-boundary analyses（次要可靠性与化学边界）",
    }
    for old, new in heading_map.items():
        try:
            set_para_text(find_para(doc, old), new)
        except ValueError:
            pass
    for old in ["2.7 活性悬崖、bRo5 与解释性分析", "2.8 统计分析、计算成本与开放科学", "2.9 双层决策卡与选择日志", "2.10 主张层级与终止规则"]:
        try:
            p = find_para(doc, old)
            p.style = "Heading 3"
        except ValueError:
            pass

    insert_after_first_containing(
        doc,
        "32 候选扩池实验使用 Morgan-512 的轻量变体",
        [
            (
                f"为回应名义候选数与有效候选数的区分，本轮新增候选有效多样性审计。"
                f"32 候选轻量池的内层验证效用两两相关中位数为 {metrics['inner_corr']:.3f}，特征值有效候选数 K_eff={metrics['inner_keff']:.2f}；"
                f"外层测试效用相关中位数为 {metrics['outer_corr']:.3f}，K_eff={metrics['outer_keff']:.2f}。"
                "因此，32 候选池应解释为高度相关轻量候选的受控选择压力实验，而不是代表 GNN、D-MPNN、Transformer 和融合模型共同扩张的真实异质池。"
            ),
            (
                "真实异质性的确认由 12 候选多视图实验承担：Morgan-512、MACCS、RDKit2D 与拼接多视图分别配对线性模型、随机森林和 LightGBM，"
                "并在九个终点、相同外层折、相同内层折和相同种子上完成 6,480 次冻结重训。Chemprop、GNN、ChemBERTa 和 MoLFormer 尚未完成同折确认性重训，"
                "因此不进入当前确认性效应估计。"
            ),
            (
                f"为把相关实验做厚，本轮进一步把异质池拆成 Morgan-only、Morgan+MACCS、Morgan+MACCS+RDKit2D 和完整 12 候选多视图四级阶梯。"
                f"在 {metrics['thick_multiview_units']} 个任务-种子-外层折单位中，完整异质池相对 Morgan-only 的 G_attain={metrics['thick_multiview_G_attain']:.3f}，"
                f"L_select={metrics['thick_multiview_L_select']:.3f}，G_realized={metrics['thick_multiview_G_realized']:.3f}，中位 η={metrics['thick_multiview_eta']:.3f}；"
                "对应候选有效多样性、逐终点效应和完整候选登记均作为 source data 输出。"
            ),
            (
                f"同时补齐真嵌套验证、种子敏感性、统一消融、80/90/95 保形覆盖率、精确 Tanimoto 分箱、MoleculeACE 活性悬崖、低相似度失败样本和扩展失败案例等 {metrics['thick_completion_tables']} 张补充实验表。"
                f"强基线证据矩阵共 {metrics['thick_strong_baseline_rows']} 行：CatBoost、XGBoost、LightGBM、ExtraTrees 和 RF 作为同分割强树模型证据；"
                "Chemprop/D-MPNN 与 ChemBERTa/MoLFormer 冻结适配器仅作为历史或边界证据，不写入当前确认性效应估计。"
            ),
        ],
    )

    insert_after_first_containing(
        doc,
        "动态池内分母遗憾仅作敏感性分析",
        [
            (
                "本文同时报告四个收益量：可达到收益 G_attain(K)=U_oracle(K)-U_oracle(4)，选择损失 L_select(K)=U_oracle(K)-U_selected(K)，"
                "实际兑现收益 G_realized(K)=U_selected(K)-U_selected(4)，以及兑现比例 η(K)=G_realized(K)/G_attain(K)。"
            ),
            (
                f"在 fixed-prefix 轻量池中，K=32 的平均 G_attain={metrics['g32_attain']:.3f}，L_select={metrics['g32_select']:.3f}，"
                f"G_realized={metrics['g32_realized']:.3f}，中位 η={metrics['g32_eta']:.3f}。"
                "这一分解说明，候选扩张确实增加了理论上界，但验证选择只兑现了其中一部分。完整池效用范围的中位数及动态分母敏感性已写入补充分析表。"
            ),
        ],
    )

    insert_after_first_containing(
        doc,
        "扩池效应在同一终点-重复-外层折内配对",
        [
            "统计推断预先区分确认性、次要和探索性分析。唯一主要效应为 K=32 相对 K=4 的完整池范围归一化选择损失差；CAHit、MRR、NDCG、Spearman 和 Kendall 作为次要或敏感性指标。Holm 校正覆盖三种池比较与三类排序/遗憾指标构成的九项检验家族。"
        ],
    )

    insert_after_first_containing(
        doc,
        "保形预测被用于补充而不是取代校准",
        [
            "分类保形采用 split conformal 的类别条件非一致性阈值：真实类别概率的 1-p_y 作为分数，在每个类别内估计有限样本修正分位数；当某一类别校准样本低于预设最低数量时，回退到 pooled 阈值。预测集合包含所有满足 1-p_y≤q_y 的标签，并记录空集合、多标签集合、平均集合大小和 fallback 原因。回归任务采用 split conformal 绝对残差区间，并按训练标签标准差报告标准化宽度。",
            (
                f"加厚后的保形明细覆盖 {metrics['thick_conformal_rows']} 条 seed-candidate-target 记录，并逐终点报告目标覆盖率、最低覆盖率、类别条件覆盖、"
                "平均集合大小、空集合率、区间宽度和 fallback 次数。重复与清洗审计同步汇总为端点-动作-原因三层表，共记录 "
                f"{metrics['thick_cleaning_events']} 条标准化、合并或剔除事件，以便复核重复样本处理是否影响结论边界。"
            ),
        ],
    )

    insert_after_first_containing(
        doc,
        "TDC 结果强调门控规则的终点依赖性",
        [
            (
                "TDC 部分仅使用 3 个种子，因此相关区间应解释为 seed variability interval 或描述性区间，而不是严格抽样置信区间。"
                f"当前 {metrics['tdc_promoted']} 个 promoted 终点中有 {metrics['tdc_cross_zero_promoted']} 个描述性区间跨零，故只保留为证据不足或宽区间晋级案例；"
                "promoted-and-improved 与 retained-and-avoided-harm 均标注为外层事后审计。"
            )
        ],
    )

    replacements = {
        "固定分母遗憾": "完整池范围归一化选择损失",
        "95% 配对区间": "三种子描述性区间",
        "综合 11 张主图和 8 张主表": "综合确认性主结果、次要应用结果和边界分析",
        "图 1–11": "图 1–9",
        "11 张主图": "9 张主文图",
        "TDC 又显示晋级证据强度随终点变化": "TDC 显示晋级证据强度随终点变化，且三种子区间不能承担强抽样推断",
    }
    for p in doc.paragraphs:
        text = p.text
        new = text
        for old, val in replacements.items():
            new = new.replace(old, val)
        if new != text:
            set_para_text(p, new)

    replace_section(doc, "4 讨论", "5 结论", DISCUSSION)
    replace_section(doc, "5 结论", "参考文献", CONCLUSION)

    for n in (10, 11):
        remove_figure(doc, n)
    try:
        delete_between(doc, "3.10 Source data 自动重建", "4 讨论")
        delete_paragraph(find_para(doc, "3.10 Source data 自动重建"))
    except ValueError:
        pass

    update_tables(doc)
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.keep_with_next = True
        elif p.style and p.style.name in {"FigureCaption", "TableCaption"}:
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p.paragraph_format.line_spacing = 1.15
    doc.save(DOCX)


def xml_errors(path: Path) -> list[str]:
    errors = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if name.endswith(".xml"):
                try:
                    ElementTree.fromstring(zf.read(name))
                except Exception as exc:
                    errors.append(f"{name}: {exc}")
    return errors


def audit(metrics: dict[str, float]) -> None:
    doc = Document(DOCX)
    text = "\n".join(p.text for p in doc.paragraphs)
    thick_report = json.loads(THICK_AUDIT.read_text(encoding="utf-8")) if THICK_AUDIT.exists() else {}
    thick_outputs = [
        THICK_ANALYSIS / "heterogeneous_staged_pool_summary.csv",
        THICK_ANALYSIS / "heterogeneous_pool_effective_diversity.csv",
        THICK_ANALYSIS / "strong_baseline_evidence_matrix.csv",
        THICK_ANALYSIS / "conformal_endpoint_detail.csv",
        THICK_ANALYSIS / "duplicate_and_cleaning_audit_summary.csv",
        THICK_ANALYSIS / "existing_completion_experiment_manifest.csv",
        THICK_ANALYSIS / "lightweight32_candidate_registry_enriched.csv",
        THICK_ANALYSIS / "lightweight32_outer_utility_correlation_matrix.csv",
        THICK_ANALYSIS / "lightweight32_validation_rank_correlation_matrix.csv",
        THICK_ANALYSIS / "lightweight32_error_overlap_status.csv",
    ]
    fig_caps = [p.text.strip() for p in doc.paragraphs if re.match(r"^图\s*\d+\b", p.text.strip())]
    tab_caps = [p.text.strip() for p in doc.paragraphs if re.match(r"^表\s*\d+\s*\|", p.text.strip())]
    report = {
        "docx": str(DOCX),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "figure_caption_numbers": [int(re.search(r"\d+", x).group()) for x in fig_caps],
        "table_caption_numbers": [int(re.search(r"\d+", x).group()) for x in tab_caps],
        "has_new_title": TITLE in text,
        "has_unique_thesis": "唯一主命题" in text,
        "has_keff": "K_eff" in text,
        "has_gain_decomposition": "G_attain" in text and "G_realized" in text,
        "has_tdc_interval_caution": "seed variability interval" in text,
        "has_conformal_definition": "类别条件非一致性阈值" in text,
        "mentions_unfinished_deep_models": "尚未完成同折确认性重训" in text,
        "has_thickened_experiment_text": "为把相关实验做厚" in text and "强基线证据矩阵" in text,
        "thickened_experiments_passed": bool(thick_report.get("passed")),
        "thickened_outputs_exist": all(p.exists() for p in thick_outputs),
        "xml_errors": xml_errors(DOCX),
        "metrics": metrics,
    }
    report["passed"] = all(
        [
            report["has_new_title"],
            report["has_unique_thesis"],
            report["has_keff"],
            report["has_gain_decomposition"],
            report["has_tdc_interval_caution"],
            report["has_conformal_definition"],
            report["mentions_unfinished_deep_models"],
            report["has_thickened_experiment_text"],
            report["thickened_experiments_passed"],
            report["thickened_outputs_exist"],
            not report["xml_errors"],
        ]
    )
    AUDIT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    if not report["passed"]:
        raise SystemExit(json.dumps(report, ensure_ascii=False, indent=2))


def main() -> None:
    metrics = compute_analysis()
    revise_document(metrics)
    audit(metrics)
    print(DOCX)
    print(AUDIT)


if __name__ == "__main__":
    main()
