# -*- coding: utf-8 -*-
"""Compress the formula section to a Nature-style main-text set.

Starting from draft 11, this produces draft 12. The main text keeps only the
non-standard equations needed to define FZYC-Mol. Standard metric formulae and
long derivations are explicitly moved to Supplementary Methods in the prose.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


DESKTOP = Path(r"C:\Users\Administrator\Desktop")
SRC = max(DESKTOP.glob("FZYC-Mol_*-11.docx"), key=lambda p: p.stat().st_mtime)
OUT = DESKTOP / "FZYC-Mol_初稿-12.docx"


NEW_FORMULA_SECTION = [
    "3.7 核心数学定义",
    (
        "按照 Nature 系列论文的主文写法，本节仅保留定义 FZYC-Mol 所必需的核心公式。"
        "RMSE、MAE、ROC-AUC、PR-AUC、Brier、ECE、bootstrap、Wilcoxon、Benjamini-Hochberg FDR "
        "以及 split-conformal 具体实现均属于通用统计或评价细节，放入 Supplementary Methods。"
        "设 t 表示预测终点，a 表示候选策略，x_i 和 y_i 分别表示第 i 个分子的结构输入和观测标签。"
        "所有候选池、阈值、容差、复杂度惩罚和打破平局规则均在测试集评估前冻结。"
    ),
    "数据边界与候选池。",
    "D_t = D_t^{tr} ∪ D_t^{val} ∪ D_t^{te},  D_t^{tr} ∩ D_t^{val} ∩ D_t^{te} = ∅,  A_t = {a_1,...,a_{m_t}}	(1)",
    "ŷ_{i,a} = f_{t,a}(ϕ_i),  ϕ_i = [ϕ_i^{fp}, ϕ_i^{desc}, ϕ_i^{graph}, ϕ_i^{pre}]	(2)",
    (
        "式 (1) 固定训练、验证和测试边界；式 (2) 定义每个候选策略在多视图分子表示上的预测。"
        "该边界是防止测试集事后选择的核心前提。"
    ),
    "验证效用与冻结选择器。",
    "S_t^r(a) = dir_t · M_t({y_i, ŷ_{i,a}}_{i∈D_t^r})	(3)",
    "a_t^* = argmax_{a∈A_t(ε_t)} [S_t^{val}(a) - λ_R R_t^{val}(a) - λ_C C(a) + λ_S Stab_t(a)]	(4)",
    (
        "式 (3) 将不同方向的主指标转换为统一的验证效用，其中误差类指标取 dir_t=-1，"
        "越大越优的指标取 dir_t=1。式 (4) 定义最终保留策略：只有验证集效用落入预设容差 "
        "A_t(ε_t) 的候选才进入打破平局步骤，随后按验证风险 R_t^{val}、模型复杂度 C 和跨 seed 稳定性 "
        "Stab_t 决定最终策略。若没有触发容差或稳定性规则，式 (4) 退化为验证集最佳候选。"
    ),
    "选择偏差审计。",
    "Regret_t = max_{a∈A_t} S_t^{te}(a) - S_t^{te}(a_t^*),  OptGap_t = S_t^{val}(a_t^*) - S_t^{te}(a_t^*)	(5)",
    "ρ_{rank,t} = Spearman(rank_val(A_t), rank_test(A_t)),  Hit@K_t = I(argmax_a S_t^{te}(a) ∈ TopK_val(A_t))	(6)",
    (
        "式 (5) 和式 (6) 不参与模型选择，只用于审计验证集治理的残余风险。Regret 衡量冻结策略与测试集观测最佳之间的差距，"
        "OptGap 衡量验证集乐观偏差，ρ_rank 和 Hit@K 衡量验证排名与测试排名的一致性。"
        "Nested validation 将候选选择放入内层验证，外层测试折仅用于偏差诊断；其完整索引形式见 Supplementary Methods。"
    ),
    "融合、适用域与样本级风险。",
    "ŷ_i^{ens} = ∑_{a∈B_t} w_{i,a}ŷ_{i,a},  w_{i,a} ∝ exp(S_t^{val}(a)/τ)/(u_{i,a}+ε)	(7)",
    "d_i^{AD} = 1 - max_{j∈D_t^{tr}∪D_t^{val}} T(m_i,m_j),  r_i = λp_i^{err} + (1-λ)d_i^{AD}	(8)",
    (
        "式 (7) 定义验证集接受的融合候选，其中 B_t 为保留候选集合，u_{i,a} 为样本级不确定性，τ 控制验证效用权重。"
        "式 (8) 将化学空间距离和错误模型组合为样本级风险分数；d_i^{AD} 越高，表示该分子越远离训练/验证化学空间。"
    ),
    "风险覆盖与活性悬崖审计。",
    "Risk(q) = mean(e_i | r_i ≤ Q_q(r)),  Coverage(q) = n^{-1}∑_{i=1}^{n}I(r_i ≤ Q_q(r))	(9)",
    "cliff(i,j)=I(T(m_i,m_j)≥θ_s and |y_i-y_j|≥θ_y),  ρ_gap=Spearman(|y_i-y_j|, |ŷ_i-ŷ_j|)	(10)",
    (
        "式 (9) 是 risk-coverage 曲线的主文定义，用于检验低风险样本优先保留时错误是否下降。"
        "式 (10) 定义活性悬崖分子对及预测差异与真实差异的一致性。低相似度三档、conformal coverage、"
        "校准图、fixed-precision recall 和片段富集统计均在 Supplementary Methods 中给出完整公式与实现细节。"
    ),
    (
        "综上，式 (1)-(10) 只定义 FZYC-Mol 的数据边界、验证集治理、选择偏差审计、融合、适用域风险、"
        "risk-coverage 和活性悬崖审计。其他通用评价公式不在主文重复展开，以减少公式密度并保持结果叙述的连续性。"
    ),
]


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_para(paragraph, text: str) -> None:
    paragraph.text = text
    if "\t(" in text or text.strip().endswith(tuple(f"({i})" for i in range(1, 20))):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    doc = Document(SRC)
    start = None
    end = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("3.7 ") and ("数学" in text or "公式" in text):
            start = idx
        if start is not None and idx > start and text == "结果":
            end = idx
            break
    if start is None or end is None:
        raise RuntimeError("Could not locate formula section boundaries.")

    anchor = doc.paragraphs[start]
    for text in NEW_FORMULA_SECTION:
        p = anchor.insert_paragraph_before(text)
        if "\t(" in text:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Remove old formula section after inserting the replacement before it.
    for paragraph in list(doc.paragraphs)[start + len(NEW_FORMULA_SECTION) : end + len(NEW_FORMULA_SECTION)]:
        delete_paragraph(paragraph)

    doc.save(OUT)
    print(f"source={SRC}")
    print(f"output={OUT}")
    print(f"old_paragraph_range={start+1}-{end}")
    print(f"new_formula_count=10")


if __name__ == "__main__":
    main()
