from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]


def source_docx() -> Path:
    base = Path.home() / "Desktop" / "\u4fee\u6539"
    src = base / "FZYC-Mol_\u521d\u7a3f4_\u8865\u5145\u5b9e\u9a8c\u5168\u8dd1\u4fee\u8ba2.docx"
    if not src.exists():
        raise FileNotFoundError(src)
    return src


def target_docx(src: Path) -> Path:
    return src.with_name("FZYC-Mol_\u521d\u7a3f4_\u8865\u5145\u5b9e\u9a8c\u5168\u8dd1\u516c\u5f0f\u8be6\u7ec6\u4fee\u8ba2.docx")


def set_font(run, size: float = 10, bold: bool | None = None, font: str = "Times New Roman") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style(p, size: float = 10, bold: bool = False) -> None:
    for run in p.runs:
        set_font(run, size=size, bold=bold)


def insert_para(anchor, text: str, size: float = 10, bold: bool = False, align=None):
    p = anchor.insert_paragraph_before(text)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(3)
    style(p, size=size, bold=bold)
    return p


def insert_formula(anchor, formula: str, number: int) -> None:
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.15), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(formula)
    set_font(run, size=9.0, font="Cambria Math")
    run = p.add_run(f"\t({number})")
    set_font(run, size=9.0, font="Cambria Math")


def replace_paragraph(doc: Document, old_start: str, new_text: str) -> None:
    for p in doc.paragraphs:
        if p.text.strip().startswith(old_start):
            p.text = new_text
            style(p, size=10)
            return


def add_formula_refs(doc: Document) -> None:
    replace_paragraph(
        doc,
        "\u9009\u62e9\u5668\u9996\u5148\u5728\u6bcf\u4e2a\u7ec8\u70b9\u7684\u9a8c\u8bc1\u96c6\u9884\u6d4b\u8868\u4e0a\u8ba1\u7b97\u5b98\u65b9\u4e3b\u6307\u6807",
        "\u9009\u62e9\u5668\u9996\u5148\u5728\u6bcf\u4e2a\u7ec8\u70b9\u7684\u9a8c\u8bc1\u96c6\u9884\u6d4b\u8868\u4e0a\u8ba1\u7b97\u5b98\u65b9\u4e3b\u6307\u6807\uff0c\u5e76\u6309\u516c\u5f0f (3) \u5c06\u5206\u7c7b\u548c\u56de\u5f52\u6307\u6807\u7edf\u4e00\u4e3a\u6b63\u5411\u5f97\u5206\u3002\u6700\u7ec8\u5019\u9009\u7531\u516c\u5f0f (6)-(8) \u51b3\u5b9a\uff1b\u82e5\u591a\u4e2a\u4e13\u5bb6\u5728\u9a8c\u8bc1\u96c6\u4e0a\u4e92\u8865\uff0c\u5219\u5019\u9009 Top-K \u5747\u503c\u3001\u5806\u53e0\u96c6\u6210\u3001\u4e0d\u786e\u5b9a\u6027\u52a0\u6743\u878d\u5408\u6216\u76f8\u5173\u6027\u60e9\u7f5a\u5171\u8bc6\uff0c\u5bf9\u5e94\u516c\u5f0f (13)-(16)\u3002\u5019\u9009\u6c60\u3001\u4e3b\u6307\u6807\u3001\u5bb9\u5dee\u548c\u5e73\u5c40\u89c4\u5219\u5728\u8bfb\u53d6\u6d4b\u8bd5\u96c6\u524d\u56fa\u5b9a\uff1b\u672a\u8d85\u8fc7\u5f53\u524d\u6700\u7ec8\u4fdd\u7559\u7ed3\u679c\u7684\u8865\u6551\u5934\u3001\u76ee\u6807\u53d8\u6362\u6216\u5f3a\u57fa\u7ebf\u8fdb\u5165\u9644\u5f55\u800c\u4e0d\u6539\u53d8\u4e3b\u7ed3\u679c\u3002\u8be5\u89c4\u5219\u4f7f\u6b63\u7ed3\u679c\u3001\u8d1f\u7ed3\u679c\u548c\u672a\u63a5\u5165\u5019\u9009\u90fd\u6709\u660e\u786e\u4f4d\u7f6e\uff0c\u4e5f\u964d\u4f4e\u4e86\u6d4b\u8bd5\u96c6\u4e8b\u540e\u9009\u62e9\u504f\u5dee\u3002",
    )
    replace_paragraph(
        doc,
        "\u4e3a\u5ba1\u8ba1\u9a8c\u8bc1\u96c6\u9009\u62e9\u662f\u5426\u53ef\u80fd\u8fc7\u62df\u5408",
        "\u4e3a\u5ba1\u8ba1\u9a8c\u8bc1\u96c6\u9009\u62e9\u662f\u5426\u53ef\u80fd\u8fc7\u62df\u5408\uff0c\u672c\u6587\u989d\u5916\u8ba1\u7b97\u6bcf\u4e2a dataset-seed \u5019\u9009\u6c60\u4e2d\u9a8c\u8bc1\u96c6\u6392\u540d\u4e0e\u6d4b\u8bd5\u96c6\u6392\u540d\u7684 Spearman \u76f8\u5173\u6027\uff0c\u5e76\u8bb0\u5f55\u9a8c\u8bc1\u96c6\u7b2c\u4e00\u540d\u662f\u5426\u4e5f\u662f\u6d4b\u8bd5\u96c6\u7b2c\u4e00\u540d\u3001\u6d4b\u8bd5\u96c6\u7b2c\u4e00\u540d\u662f\u5426\u843d\u5165\u9a8c\u8bc1\u96c6 Top-3\uff0c\u4ee5\u53ca\u516c\u5f0f (9)-(12) \u5b9a\u4e49\u7684 regret\u3001optimism gap\u3001\u6392\u540d\u5ba1\u8ba1\u548c nested validation \u5916\u5c42\u8bc4\u4f30\u3002\u8be5\u5206\u6790\u4e0d\u53c2\u4e0e\u6a21\u578b\u9009\u62e9\uff0c\u53ea\u7528\u4e8e\u8bc4\u4f30\u9009\u62e9\u5668\u98ce\u9669\uff1a\u82e5\u9a8c\u8bc1-\u6d4b\u8bd5\u6392\u540d\u76f8\u5173\u6027\u8f83\u4f4e\u3001Top-3 \u547d\u4e2d\u4e0d\u8db3\u6216\u51fa\u73b0\u8d1f\u76f8\u5173\uff0c\u76f8\u5173\u7ed3\u679c\u5728\u6b63\u6587\u4e2d\u6309\u8fb9\u754c\u6216\u8d1f\u7ed3\u679c\u89e3\u91ca\uff0c\u800c\u4e0d\u4f5c\u4e3a\u65e0\u6761\u4ef6\u6027\u80fd\u63d0\u5347\u8bc1\u636e\u3002",
    )
    replace_paragraph(
        doc,
        "\u53ef\u9760\u6027\u6a21\u5757\u5305\u62ec\u96c6\u6210\u6807\u51c6\u5dee",
        "\u53ef\u9760\u6027\u6a21\u5757\u5305\u62ec\u96c6\u6210\u6807\u51c6\u5dee\u3001\u9519\u8bef\u6a21\u578b\u3001\u9884\u6d4b\u504f\u5dee\u3001\u53cd\u5411 Tanimoto \u8ddd\u79bb\u3001\u91cd\u6784\u8bef\u5dee\u3001\u9519\u8bef-\u9002\u7528\u57df\u6df7\u5408\u6307\u6807\u3001\u98ce\u9669-\u8986\u76d6\u66f2\u7ebf\u3001\u4fdd\u5f62\u8986\u76d6\u7387\u3001\u6821\u51c6\u548c\u7c97\u7cd9\u5ea6\u4ee3\u7406\u6307\u6807\u3002\u5176\u4e2d Tanimoto \u6700\u8fd1\u90bb\u76f8\u4f3c\u5ea6\u3001\u4e25\u683c\u5206\u5c42\u3001\u9002\u7528\u57df\u95e8\u63a7\u3001\u96c6\u6210\u4e0d\u786e\u5b9a\u6027\u3001\u9519\u8bef\u6807\u7b7e\u3001\u9519\u8bef\u6a21\u578b\u3001\u6df7\u5408\u98ce\u9669\u548c\u98ce\u9669-\u8986\u76d6\u5bcc\u96c6\u5206\u522b\u7531\u516c\u5f0f (17)-(25) \u5b9a\u4e49\u3002\u5206\u7c7b\u4efb\u52a1\u628a\u9884\u6d4b\u6807\u7b7e\u9519\u8bef\u4f5c\u4e3a\u9519\u8bef\u68c0\u51fa\u76ee\u6807\uff0c\u56de\u5f52\u4efb\u52a1\u628a\u7edd\u5bf9\u8bef\u5dee\u6700\u9ad8\u7684 20% \u6837\u672c\u4f5c\u4e3a\u9ad8\u8bef\u5dee\u76ee\u6807\uff0c\u5e76\u7528 AUROC\u3001AUPRC\u3001\u98ce\u9669-\u8986\u76d6\u66f2\u7ebf\u548c top-10% \u5bcc\u96c6\u5171\u540c\u8bc4\u4f30\u98ce\u9669\u5206\u6570\u3002\u53ef\u89e3\u91ca\u6027\u6a21\u5757\u5305\u62ec\u57fa\u5e8f\u7279\u5f81\u91cd\u8981\u6027\u3001\u7247\u6bb5\u5bcc\u96c6\u3001\u9aa8\u67b6/\u8fd1\u90bb\u6848\u4f8b\u590d\u6838\u4e0e\u9ad8\u8bef\u5dee\u6837\u672c\u5206\u6790\u3002\u672c\u6587\u53ea\u628a\u8fd9\u4e9b\u8bc1\u636e\u89e3\u91ca\u4e3a\u6a21\u578b\u4f7f\u7528\u8fb9\u754c\u548c\u5316\u5b66\u76f8\u5173\u6027\uff0c\u4e0d\u628a\u5173\u8054\u6027\u57fa\u5e8f\u89e3\u91ca\u7b49\u540c\u4e8e\u56e0\u679c\u673a\u5236\u3002",
    )
    replace_paragraph(
        doc,
        "\u5206\u7c7b\u4efb\u52a1\u4fdd\u7559 ROC-AUC \u4f5c\u4e3a MoleculeNet",
        "\u5206\u7c7b\u4efb\u52a1\u4fdd\u7559 ROC-AUC \u4f5c\u4e3a MoleculeNet \u548c\u82e5\u5e72 TDC \u7ec8\u70b9\u7684\u4e3b\u6307\u6807\uff0c\u540c\u65f6\u62a5\u544a PR-AUC\u3001Brier \u5206\u6570\u3001ECE\u3001MCC\u3001\u5e73\u8861\u51c6\u786e\u7387\u3001\u98ce\u9669-\u8986\u76d6\u66f2\u7ebf\u548c\u4fdd\u5f62\u8986\u76d6\u7387\u3002Brier/ECE \u548c 80%/90%/95% \u4fdd\u5f62\u9884\u6d4b\u7684\u5b9a\u4e49\u89c1\u516c\u5f0f (5)\u3001(26)-(28)\uff1b\u8fd9\u79cd\u5199\u6cd5\u662f\u4e3a\u4e86\u56de\u5e94\u7c7b\u522b\u4e0d\u5e73\u8861\u4efb\u52a1\u4e2d\u7684\u5e38\u89c1\u5ba1\u7a3f\u8d28\u7591\uff1aROC-AUC \u8f83\u9ad8\u5e76\u4e0d\u5fc5\u7136\u610f\u5473\u7740\u9633\u6027\u6837\u672c\u53ec\u56de\u3001\u6982\u7387\u6821\u51c6\u6216\u7b5b\u9009\u5bcc\u96c6\u53ef\u9760\u3002ClinTox\u3001DILI\u3001hERG \u548c CYP \u5e95\u7269\u4efb\u52a1\u56e0\u6b64\u5728\u4e0d\u5e73\u8861\u5206\u7c7b\u589e\u5f3a\u4e2d\u91cd\u70b9\u62a5\u544a\u3002",
    )
    replace_paragraph(
        doc,
        "\u7c97\u7cd9\u5ea6\u4ee3\u7406\u6307\u6807\u7528\u4e8e\u89e3\u91ca\u7ec8\u70b9\u7684\u5c40\u90e8\u7ed3\u6784-\u6027\u8d28\u5173\u7cfb\u662f\u5426\u5e73\u6ed1",
        "\u7c97\u7cd9\u5ea6\u4ee3\u7406\u6307\u6807\u7528\u4e8e\u89e3\u91ca\u7ec8\u70b9\u7684\u5c40\u90e8\u7ed3\u6784-\u6027\u8d28\u5173\u7cfb\u662f\u5426\u5e73\u6ed1\uff0c\u5177\u4f53\u6309\u516c\u5f0f (31) \u5c06\u6700\u8fd1\u90bb\u6807\u7b7e\u8df3\u53d8\u4e0e Tanimoto \u8ddd\u79bb\u8054\u7cfb\u8d77\u6765\u3002\u82e5\u6d4b\u8bd5\u5206\u5b50\u4e0e\u6700\u8fd1\u90bb\u9ad8\u5ea6\u76f8\u4f3c\u4f46\u6807\u7b7e\u5dee\u5f02\u6216\u5f52\u4e00\u5316\u76ee\u6807\u503c\u8df3\u53d8\u663e\u8457\uff0c\u5219\u8bf4\u660e\u8be5\u7ec8\u70b9\u5177\u6709\u9ad8\u7c97\u7cd9\u5ea6\u3002\u5355\u7eaf\u589e\u5927\u6a21\u578b\u5bb9\u91cf\u901a\u5e38\u4e0d\u8db3\u4ee5\u89e3\u51b3\u8be5\u95ee\u9898\uff0c\u9700\u8981\u76ee\u6807\u53d8\u6362\u3001\u7a33\u5065\u635f\u5931\u3001\u5c40\u90e8\u90bb\u57df\u8bca\u65ad\u6216\u96c6\u6210\u7a33\u5b9a\u5316\u3002\u672c\u6587\u4e0d\u628a\u7c97\u7cd9\u5ea6\u4f5c\u4e3a\u4e3b\u6027\u80fd\u6307\u6807\uff0c\u800c\u662f\u4f5c\u4e3a\u6027\u80fd\u74f6\u9888 ADME \u56de\u5f52\u548c\u6d3b\u6027\u60ac\u5d16\u4efb\u52a1\u7684\u89e3\u91ca\u6027\u8bc1\u636e\u3002",
    )
    replace_paragraph(
        doc,
        "\u57fa\u5e8f\u5f52\u56e0\u4e0e\u7247\u6bb5\u5bcc\u96c6\u63d0\u4f9b\u4e92\u8865\u89e3\u91ca",
        "\u57fa\u5e8f\u5f52\u56e0\u4e0e\u7247\u6bb5\u5bcc\u96c6\u63d0\u4f9b\u4e92\u8865\u89e3\u91ca\uff0c\u5176\u4e2d\u7247\u6bb5\u5bcc\u96c6\u6309\u516c\u5f0f (34) \u62a5\u544a\u6807\u7b7e\u504f\u79fb\u6216\u6bd4\u503c\u6bd4\u3002\u524d\u8005\u9762\u5411\u6a21\u578b\u5185\u90e8\u7279\u5f81\u91cd\u8981\u6027\uff0c\u5173\u6ce8\u5b98\u80fd\u56e2\u3001BRICS \u7247\u6bb5\u548c Murcko \u9aa8\u67b6\u5bf9\u9884\u6d4b\u7684\u8d21\u732e\u65b9\u5411\uff1b\u540e\u8005\u9762\u5411\u6570\u636e\u5206\u5e03\uff0c\u5173\u6ce8\u67d0\u4e9b\u7247\u6bb5\u51fa\u73b0\u65f6\u6807\u7b7e\u5747\u503c\u76f8\u5bf9\u57fa\u7ebf\u7684\u504f\u79fb\u3002\u4e24\u8005\u4e00\u81f4\u65f6\uff0c\u8bf4\u660e\u6a21\u578b\u5173\u6ce8\u70b9\u4e0e\u6570\u636e\u7edf\u8ba1\u4fe1\u53f7\u76f8\u4e92\u652f\u6301\uff1b\u4e0d\u4e00\u81f4\u65f6\uff0c\u5219\u63d0\u793a\u6df7\u6742\u56e0\u7d20\u3001\u6837\u672c\u91cf\u4e0d\u8db3\u6216\u5c40\u90e8\u6807\u7b7e\u566a\u58f0\u3002",
    )


def add_section(doc: Document) -> None:
    anchor = next((p for p in doc.paragraphs if p.text.strip() == "\u7ed3\u679c"), None)
    if anchor is None:
        raise RuntimeError("Cannot find Results heading.")

    insert_para(anchor, "3.7 \u6570\u5b66\u5b9a\u4e49\u4e0e\u8bc4\u4ef7\u516c\u5f0f", size=11, bold=True)
    insert_para(
        anchor,
        "\u4e3a\u4f7f FZYC-Mol \u7684\u9a8c\u8bc1\u96c6\u9009\u62e9\u3001\u5019\u9009\u878d\u5408\u3001\u9002\u7528\u57df\u95e8\u63a7\u3001\u4e0d\u786e\u5b9a\u6027\u3001\u4fdd\u5f62\u9884\u6d4b\u3001\u4f4e\u76f8\u4f3c\u5ea6\u5206\u5c42\u3001MoleculeACE \u6d3b\u6027\u60ac\u5d16\u548c\u57fa\u5e8f/\u7247\u6bb5\u89e3\u91ca\u53ef\u88ab\u590d\u6838\uff0c\u672c\u6587\u5c06\u6838\u5fc3\u8ba1\u7b97\u7edf\u4e00\u5199\u6210\u4ee5\u4e0b\u516c\u5f0f\u3002\u516c\u5f0f\u4e2d\u6240\u6709\u8d85\u53c2\u6570\u548c\u5bb9\u5dee\u5747\u53ea\u80fd\u7531\u8bad\u7ec3\u96c6\u6216\u9a8c\u8bc1\u96c6\u786e\u5b9a\uff0c\u4e0d\u5141\u8bb8\u4f7f\u7528\u6d4b\u8bd5\u96c6\u56de\u6539\u3002",
    )
    insert_para(anchor, "\u7b26\u53f7\u4e0e\u4e3b\u6307\u6807\u65b9\u5411\u3002", size=10, bold=True)
    insert_para(anchor, "\u5bf9\u4e8e\u7ec8\u70b9 t\uff0c\u6570\u636e\u88ab\u56fa\u5b9a\u5206\u4e3a\u8bad\u7ec3\u3001\u9a8c\u8bc1\u548c\u6d4b\u8bd5\u4e09\u90e8\u5206\uff0c\u5019\u9009\u4e13\u5bb6\u96c6\u4e3a A_t\u3002")
    insert_formula(anchor, "D_t = D_t^tr union D_t^val union D_t^te,  D_t^r = {(x_i,y_i)}_{i=1}^{n_r},  r in {tr,val,te}", 1)
    insert_formula(anchor, "Z_t^r = [z_hat_{i,a}^r]_{n_r x |A_t|},  z_hat_{i,a}^r = f_{t,a}(x_i)", 2)
    insert_para(anchor, "\u4e3a\u8ba9\u5206\u7c7b\u548c\u56de\u5f52\u4efb\u52a1\u5171\u7528\u540c\u4e00\u9009\u62e9\u5668\uff0c\u5148\u5c06\u4e3b\u6307\u6807\u7edf\u4e00\u6210\u6b63\u5411\u5f97\u5206\u3002")
    insert_formula(anchor, "S_t^r(a) = M_t(y^r,z_hat_a^r) if M_t is higher-better;  S_t^r(a) = -M_t(y^r,z_hat_a^r) if M_t is lower-better", 3)
    insert_formula(anchor, "RMSE = sqrt(n^{-1} sum_i (y_i-z_hat_i)^2),  MAE = n^{-1} sum_i |y_i-z_hat_i|,  rho = Spearman(y,z_hat)", 4)
    insert_formula(anchor, "Brier = n^{-1} sum_i (p_i-y_i)^2,  ECE = sum_{b=1}^B |B_b| n^{-1} |acc(B_b)-conf(B_b)|", 5)

    insert_para(anchor, "\u9a8c\u8bc1\u96c6\u9009\u62e9\u5668\u3001\u5bb9\u5dee\u548c\u5ba1\u8ba1\u3002", size=10, bold=True)
    insert_formula(anchor, "a_t^best = argmax_{a in A_t} S_t^val(a)", 6)
    insert_formula(anchor, "A_t(epsilon) = {a in A_t: S_t^val(a) >= max_{b in A_t} S_t^val(b) - epsilon_t}", 7)
    insert_formula(anchor, "a_t^* = argmin_{a in A_t(epsilon)} [eta_1 R_t^val(a) + eta_2 C(a) - eta_3 Stab_t(a)]", 8)
    insert_para(anchor, "\u5176\u4e2d R_t^val(a) \u4e3a\u9a8c\u8bc1\u96c6\u98ce\u9669\u6216\u6821\u51c6\u60e9\u7f5a\uff0cC(a) \u4e3a\u6a21\u578b\u590d\u6742\u5ea6\uff0cStab_t(a) \u4e3a\u8de8 seed \u7a33\u5b9a\u6027\u3002\u82e5\u4e0d\u542f\u7528\u5e73\u5c40\u60e9\u7f5a\uff0ceta_1=eta_2=eta_3=0\uff0c\u9009\u62e9\u9000\u5316\u4e3a\u516c\u5f0f (6)\u3002", size=9)
    insert_formula(anchor, "Regret_t = S_t^te(a_t^dagger) - S_t^te(a_t^*),  a_t^dagger = argmax_{a in A_t} S_t^te(a)", 9)
    insert_formula(anchor, "OptGap_t = S_t^val(a_t^*) - S_t^te(a_t^*)", 10)
    insert_formula(anchor, "rho_rank,t = Spearman(rank_val(A_t), rank_test(A_t)),  Hit@K_t = I(a_t^dagger in TopK_val(A_t))", 11)
    insert_formula(anchor, "a_{t,o}^inner = argmax_a mean_{k in I_o} S_{t,k}^{inner-val}(a),  Perf_{t,o}^{outer}=S_{t,o}^{outer-test}(a_{t,o}^inner)", 12)

    insert_para(anchor, "\u878d\u5408\u3001\u5806\u53e0\u4e0e\u5171\u8bc6\u6743\u91cd\u3002", size=10, bold=True)
    insert_formula(anchor, "T_K(t) = TopK_val(A_t),  z_hat_i^mean = K^{-1} sum_{a in T_K(t)} z_hat_{i,a}", 13)
    insert_formula(anchor, "beta_t^* = argmin_beta L_val(y, g(sum_{a in T_K} beta_a z_hat_a)) + lambda ||beta||_2^2,  z_hat_i^stack = g(sum_a beta_a^* z_hat_{i,a})", 14)
    insert_formula(anchor, "w_{i,a}^{unc} = exp(S_t^val(a)/tau)(u_{i,a}+epsilon)^{-1} / sum_b exp(S_t^val(b)/tau)(u_{i,b}+epsilon)^{-1}", 15)
    insert_formula(anchor, "w_a^{cons} = exp(S_t^val(a)/tau)/(1 + mean_{b != a} |corr(z_hat_a^val,z_hat_b^val)|),  z_hat_i^{ens}=sum_a w_a z_hat_{i,a}/sum_a w_a", 16)

    insert_para(anchor, "\u9002\u7528\u57df\u3001\u4e0d\u786e\u5b9a\u6027\u4e0e\u98ce\u9669-\u8986\u76d6\u3002", size=10, bold=True)
    insert_formula(anchor, "T(m_i,m_j) = |m_i AND m_j| / |m_i OR m_j|,  s_i^NN = max_{j in D^tr union D^val} T(m_i,m_j)", 17)
    insert_formula(anchor, "d_i^AD = 1 - s_i^NN,  bin_i = >0.7 if s_i^NN>0.7;  0.5-0.7 if 0.5<=s_i^NN<=0.7;  <0.5 if s_i^NN<0.5", 18)
    insert_formula(anchor, "g_i = I(s_i^NN >= delta_AD),  z_hat_i^AD = g_i z_hat_i^{ens} + (1-g_i) z_hat_i^{safe}", 19)
    insert_formula(anchor, "u_i^{ens} = sqrt((K-1)^{-1} sum_{a in T_K} (z_hat_{i,a} - mean_a z_hat_{i,a})^2)", 20)
    insert_formula(anchor, "e_i = I(y_hat_i != y_i) for classification;  e_i = I(|y_i-z_hat_i| >= q_{0.80}^{val}(|y-z_hat|)) for regression", 21)
    insert_formula(anchor, "p_i^{err} = h_phi(z_hat_i, u_i^{ens}, d_i^AD, |z_hat_i-z_hat_i^{NN}|, descriptors_i)", 22)
    insert_formula(anchor, "r_i = lambda p_i^{err} + (1-lambda)d_i^AD,  0 <= lambda <= 1", 23)
    insert_formula(anchor, "Coverage(q)=n^{-1} sum_i I(r_i <= Q_q(r)),  Risk(q)=sum_i e_i I(r_i <= Q_q(r))/sum_i I(r_i <= Q_q(r))", 24)
    insert_formula(anchor, "EF_q = P(e_i=1 | r_i in top q%) / P(e_i=1)", 25)

    insert_para(anchor, "\u4fdd\u5f62\u9884\u6d4b\u4e0e\u7ecf\u9a8c\u8986\u76d6\u3002", size=10, bold=True)
    insert_formula(anchor, "S_i^{cls}=1-p_hat_{i,y_i},  q_alpha^{cls}=Quantile_{ceil((n_cal+1)(1-alpha))/n_cal}({S_i^{cls}})", 26)
    insert_formula(anchor, "C_alpha(x)={c: 1-p_hat_c(x) <= q_alpha^{cls}},  S_i^{reg}=|y_i-z_hat_i|,  I_alpha(x)=[z_hat(x)-q_alpha^{reg}, z_hat(x)+q_alpha^{reg}]", 27)
    insert_formula(anchor, "Cov_alpha=n^{-1} sum_i I(y_i in C_alpha(x_i) or y_i in I_alpha(x_i)),  Size_alpha=n^{-1}sum_i |C_alpha(x_i)|,  Width_alpha=n^{-1}sum_i |I_alpha(x_i)|", 28)
    insert_para(anchor, "\u516c\u5f0f (26)-(28) \u4e2d\u7684\u5206\u4f4d\u6570\u53ea\u5728 calibration \u5b50\u96c6\u4e0a\u8ba1\u7b97\uff0c\u56e0\u6b64\u8986\u76d6\u7387\u53ea\u8868\u793a\u7ed9\u5b9a\u5212\u5206\u4e0b\u7684\u6709\u9650\u6837\u672c\u4fdd\u8bc1\u3002", size=9)

    insert_para(anchor, "\u6d3b\u6027\u60ac\u5d16\u3001\u7c97\u7cd9\u5ea6\u3001\u6d88\u878d\u548c\u53ef\u89e3\u91ca\u6027\u3002", size=10, bold=True)
    insert_formula(anchor, "Delta y_{ij}=|y_i-y_j|,  Delta z_hat_{ij}=|z_hat_i-z_hat_j|,  cliff(i,j)=I(T(m_i,m_j)>=theta_s and Delta y_{ij}>=theta_y)", 29)
    insert_formula(anchor, "rho_gap=Spearman({Delta y_{ij}},{Delta z_hat_{ij}}),  Acc_dir=n^{-1}sum_{ij} I(sign(y_i-y_j)=sign(z_hat_i-z_hat_j))", 30)
    insert_formula(anchor, "Rough_i = |y_i-y_{NN(i)}|/(1-s_i^NN+epsilon),  Rough_t = median_i(Rough_i)", 31)
    insert_formula(anchor, "Delta Q_{t,c}=S_t^te(c)-S_t^te(Full),  WinFrac_c=|{t: Delta Q_{t,c}>0}|/|T|", 32)
    insert_formula(anchor, "Accept_rescue(t)=I(S_t^val(rescue)-S_t^val(current)>epsilon_t and Risk_t^val(rescue)<=Risk_t^val(current)+gamma_t)", 33)
    insert_formula(anchor, "E_f^{reg}=mean(y_i | f in x_i)-mean(y_i | f notin x_i),  OR_f^{cls}=[P(y=1|f)/(1-P(y=1|f))]/[P(y=1|not f)/(1-P(y=1|not f))]", 34)
    insert_para(anchor, "\u8fd9\u4e9b\u516c\u5f0f\u4e0e\u8865\u5145\u8868 S11-S17 \u4e00\u4e00\u5bf9\u5e94\uff1a\u516c\u5f0f (12) \u5bf9\u5e94 nested validation\uff0c\u516c\u5f0f (18)\u3001(24)-(25) \u5bf9\u5e94\u4f4e\u76f8\u4f3c\u5ea6\u548c\u98ce\u9669\u5bcc\u96c6\uff0c\u516c\u5f0f (26)-(28) \u5bf9\u5e94 80%/90%/95% \u4fdd\u5f62\u9884\u6d4b\uff0c\u516c\u5f0f (29)-(31) \u5bf9\u5e94 MoleculeACE \u548c\u7c97\u7cd9\u5ea6\u8bca\u65ad\uff0c\u516c\u5f0f (32)-(33) \u5bf9\u5e94\u6d88\u878d\u4e0e\u8865\u6551\u5934\u63a5\u53d7\u89c4\u5219\u3002")


def main() -> None:
    src = source_docx()
    out = target_docx(src)
    shutil.copy2(src, out)
    doc = Document(out)
    add_formula_refs(doc)
    add_section(doc)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
