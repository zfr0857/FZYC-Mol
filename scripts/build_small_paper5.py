from __future__ import annotations

import csv
import json
import os
import shutil
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "output" / "小论文-4.docx"
OUTPUT = Path(os.environ.get("FZYC_DOCX_OUTPUT", str(ROOT / "output" / "小论文-5.docx")))
PACKAGE = Path(os.environ.get("FZYC_DOCX_PACKAGE", str(ROOT / "output" / "小论文-5_图表包")))
FIG = PACKAGE / "figures"
AUDIT = Path(os.environ.get("FZYC_FUSION_AUDIT", str(ROOT / "output" / "小论文-5_融合审计")))
SUPP = Path(os.environ.get("FZYC_SUPP_OUTPUT", str(ROOT / "output" / "小论文-5_补充表.docx")))
DOC_LABEL = OUTPUT.stem


FIGURES = {
    1: "fig01_overall_workflow.png",
    2: "fig02_candidate_pool_controls.png",
    3: "fig03_repeated_nested_selection.png",
    4: "fig04_metric_calibration_meta_risk.png",
    5: "fig05_multiview_confirmation.png",
    6: "fig06_moleculenet_decisions.png",
    7: "fig07_tdc_gate_audit.png",
    8: "fig08_prediction_reliability_conformal.png",
    9: "fig09_chemical_boundaries_decision_card.png",
    10: "fig10_governance_automl_transfer.png",
    11: "fig11_failures_negative_results_traceability.png",
}


def clear_body(doc: Document) -> None:
    body = doc._element.body
    sect = body.sectPr
    for child in list(body):
        if child is not sect:
            body.remove(child)


def ensure_styles(doc: Document) -> None:
    existing = {s.name for s in doc.styles}
    for name, size, bold, colour, keep_next in [
        ("FigureCaption", 9.0, False, "404040", False),
        ("TableCaption", 9.5, True, "1F4E79", True),
        ("TableNote", 8.3, False, "606060", False),
        ("Equation", 10.2, False, "202020", False),
    ]:
        if name not in existing:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[name]
        style.font.name = "Arial"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.keep_with_next = keep_next
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(3)


def set_repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trpr.append(node)


def shade(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def no_split(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:cantSplit"))


def safe_style(doc: Document, preferred: str, fallback: str = "Normal") -> str:
    return preferred if preferred in [s.name for s in doc.styles] else fallback


def add_p(doc: Document, text: str, style: str = "Normal", align=None):
    p = doc.add_paragraph(text, style=safe_style(doc, style))
    if align is not None:
        p.alignment = align
    return p


def add_h(doc: Document, text: str, level: int = 1):
    return add_p(doc, text, f"Heading {level}")


def add_equation(doc: Document, text: str):
    p = add_p(doc, text, "Equation")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_figure(doc: Document, number: int, caption: str, width_cm: float = 16.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(FIG / FIGURES[number]), width=Cm(width_cm))
    cap = add_p(doc, f"图 {number} | {caption}", "FigureCaption")
    cap.paragraph_format.first_line_indent = Cm(0)
    return cap


def add_table(doc: Document, number: int, title: str, headers: list[str], rows: list[list[str]], note: str = ""):
    cap = add_p(doc, f"表 {number} | {title}", "TableCaption")
    cap.paragraph_format.first_line_indent = Cm(0)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = safe_style(doc, "Normal Table", "Table Grid")
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = str(value)
        shade(cell, "DDEBF7")
    set_repeat_table_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = str(value)
        no_split(table.rows[-1])
    for r_i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.first_line_indent = Cm(0)
                for run in p.runs:
                    run.font.name = "Arial"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(8.2)
                    if r_i == 0:
                        run.bold = True
    if note:
        p = add_p(doc, f"注：{note}", "TableNote")
        p.paragraph_format.first_line_indent = Cm(0)
    return table


TABLES = {
    1: (
        "数据资源、划分方案及证据用途",
        ["数据资源", "任务与规模", "划分与重复", "主要输出", "证据用途"],
        [
            ["MoleculeNet", "3 回归+3 分类；n=642–4,200", "骨架划分；5 种子", "RMSE/ROC-AUC；校准", "冻结应用面板"],
            ["TDC ADMET", "18 ADME+4 毒性；22 终点", "官方骨架划分；3 种子", "官方指标；逐终点区间", "跨来源公开面板"],
            ["32 候选确认实验", "9 终点；K=4/8/16/32", "3 外层×3 内层×5 重复", "CAHit、MRR、固定遗憾", "候选扩张主检验"],
            ["12 候选多视图", "4 表示×3 学习器；9 终点", "共享 3×3×5 划分", "遗憾、可达到/兑现增益", "异质候选确认"],
            ["MoleculeACE", "17 回归任务", "公开配置；3 种子", "整体/悬崖 RMSE；差异相关", "局部非连续性边界"],
            ["公共 bRo5", "CycPept 7,334；LinPept 1,960/7,239", "随机/骨架/外缘/时间；3 种子", "RMSE、ROC-AUC、PR-AUC", "化学迁移压力"],
        ],
        "TDC 与 bRo5 均为公开压力面板，不等同于独立前瞻性盲测。",
    ),
    2: (
        "候选选择策略、测试边界与计算预算",
        ["策略", "验证侧依据", "测试标签用途", "状态", "解释"],
        [
            ["固定单模型", "预登记候选", "仅评价", "完成", "低自由度对照"],
            ["validation-best", "内层验证均值", "仅评价", "完成", "朴素验证选择"],
            ["one-SE+稳定性", "均值容差、波动与成本", "仅评价", "完成", "冻结保守规则"],
            ["风险调整", "验证均值−0.5×SD", "仅评价", "完成", "固定 λ 的负结果审计"],
            ["Top-K/堆叠", "内层折外预测", "仅评价", "完成", "融合候选"],
            ["AutoGluon", "相同 Morgan-512；30/300/1,800 s", "仅评价", "三预算完成", "性能与成本强对照"],
            ["随机期望", "候选等概率", "仅评价", "完成", "零信息参照"],
            ["test oracle", "测试效用最大", "只定义上界", "完成", "不得用于晋级"],
        ],
        "所有阈值、权重和平局规则均在外层测试前冻结。",
    ),
    3: (
        "轻量候选池 3×3×5 重复嵌套审计",
        ["候选数", "机会校正命中", "MRR", "固定分母遗憾（95% CI）", "主导候选比例"],
        [
            ["4", "0.881", "0.869", "0.049（0.015–0.090）", "0.859"],
            ["8", "0.550", "0.518", "0.131（0.086–0.191）", "0.474"],
            ["16", "0.334", "0.401", "0.139（0.090–0.195）", "0.430"],
            ["32", "0.240", "0.264", "0.171（0.113–0.237）", "0.326"],
        ],
        "每个池规模包含 9 个终点、5 个重复和 135 个外层单元；遗憾分母固定为同一外层单元完整 32 候选的测试效用范围。",
    ),
    4: (
        "MoleculeNet 冻结结果与 test oracle 差距",
        ["数据集", "主指标", "原验证选择器", "最终保留", "test oracle", "证据解释"],
        [
            ["ESOL", "RMSE", "0.5829±0.0352", "0.5829±0.0352", "0.5829±0.0352", "冻结结果与上界一致"],
            ["FreeSolv", "RMSE", "1.0678±0.1883", "1.0286±0.1761", "0.9518±0.1314", "仍有选择差距"],
            ["Lipophilicity", "RMSE", "0.7078±0.0389", "0.6835±0.0439", "0.6835±0.0439", "验证支持定向更新"],
            ["BBBP", "ROC-AUC", "0.9165±0.0290", "0.9243±0.0247", "0.9243±0.0247", "冻结结果与上界一致"],
            ["BACE", "ROC-AUC", "0.8753±0.0230", "0.8753±0.0230", "0.8753±0.0230", "冻结结果与上界一致"],
            ["ClinTox", "ROC-AUC", "0.9489±0.0302", "0.9496±0.0262", "0.9496±0.0262", "需并列少数类召回"],
        ],
        "test oracle 仅用于定义遗憾或上界，不参与策略替换。",
    ),
    5: (
        "TDC 冻结策略晋级的五个终点",
        ["终点", "官方指标", "原基线", "最终保留", "相对增益", "95% 配对区间"],
        [
            ["clearance_hepatocyte_az", "Spearman", "0.456", "0.459", "0.7%", "−15.1%–16.5%"],
            ["clearance_microsome_az", "Spearman", "0.486", "0.512", "5.5%", "−17.7%–28.7%"],
            ["ppbr_az", "MAE", "8.025", "6.874", "14.4%", "12.3%–16.4%"],
            ["vdss_lombardo", "Spearman", "0.532", "0.692", "30.0%", "14.1%–45.8%"],
            ["half_life_obach", "Spearman", "0.268", "0.448", "67.1%", "41.4%–92.8%"],
        ],
        "共有 5 promoted 与 17 retained；其余终点和门控错误类别见补表 S5。",
    ),
    6: (
        "MoleculeACE 与 bRo5 化学边界摘要",
        ["数据或模块", "范围", "主要结果", "辅助结果", "解释边界"],
        [
            ["MoleculeACE", "17 任务×3 种子", "整体 RMSE 0.711", "悬崖 RMSE 0.813", "悬崖子集误差更高"],
            ["高相似分子对", "51 任务-种子单元", "差异 Spearman 0.252", "方向准确率 0.750", "方向不等于幅度准确"],
            ["CycPept", "随机", "RMSE 0.547±0.021", "Spearman 0.761±0.028", "插值参照"],
            ["CycPept", "外缘", "RMSE 0.876±0.012", "Spearman 0.303±0.026", "结构外推压力最大"],
            ["LinPept CellPen", "外缘", "ROC-AUC 0.859±0.005", "PR-AUC 0.822±0.008", "相对随机划分下降"],
            ["LinPept NonFouling", "外缘", "ROC-AUC 0.761", "PR-AUC 0.698", "筛选富集有限"],
        ],
        "这些结果用于界定化学迁移和局部非连续性，不构成前瞻性部署验证。",
    ),
    7: (
        "统计校准、跨端点风险与多视图确认",
        ["实验或对照", "效应/结果", "95% CI", "方向", "P 值或结论"],
        [
            ["K=32 相对 K=4 固定遗憾", "+0.122", "0.072–0.175", "8/9 增加", "精确 P=0.0078；Holm=0.039"],
            ["K=32 相对 K=4 CAHit", "−0.642", "−0.780–−0.442", "8/9 下降", "机会校正"],
            ["等权 selection-risk", "Spearman 0.235", "层内置换 P=0.953", "LOEO 区间跨 0", "描述性负验证"],
            ["严格元风险", "MAE 0.112 vs 0.123", "门控 −0.034（−0.047–−0.020）", "8/9 改善", "AUC 0.648"],
            ["多视图 validation-best 遗憾", "0.043", "0.021–0.067", "—", "6,480 次拟合"],
            ["多视图 vs Morgan-only", "+0.343", "0.210–0.483", "9/9 增加", "精确 P=0.0039"],
            ["拼接 vs 分离视图池", "+0.035", "0.017–0.053", "9/9 增加", "精确 P=0.0039"],
        ],
        "信号恢复正对照在全部 K 上单调，零信号机会校正偏差绝对值小于 0.004，完全信号遗憾为 0。",
    ),
    8: (
        "代表性失败案例与双层决策卡字段",
        ["案例/字段", "层级", "观测", "风险或相似度", "可审计解释"],
        [
            ["ClinTox 假阴性", "分子", "真实 1；预测 0.273", "风险分位 0.913", "总体 AUC 不保证单样本召回"],
            ["Half-life 极端标签", "分子", "真实 820；预测 47.2", "粗糙度 0.413", "标签尺度与局部不连续并存"],
            ["FreeSolv 低相似度", "分子", "绝对误差 6.43", "Tanimoto 0.333", "边界案例，不估计发生率"],
            ["Lipophilicity 低相似度", "分子", "真实 2.50；预测 −3.40", "Tanimoto 0.217", "进入人工复核"],
            ["CHEMBL204_Ki", "分子对", "跨种子预测差异波动", "相似度 0.889", "关联证据，不推断机制"],
            ["终点级字段", "选择单元", "candidate_id、selector、验证/外层指标", "遗憾与 gate_state", "记录选择时点与负结果"],
            ["样本级字段", "测试分子", "点预测、保形集合/区间", "AD 相似度与风险分位", "accept/manual review/outside AD"],
        ],
        "案例按预定义高误差、高风险或低相似度条件选取；完整记录见补表 S9。",
    ),
}


REFERENCES = [
    "[1] Wu Z, Ramsundar B, Feinberg EN, et al. MoleculeNet: a benchmark for molecular machine learning. Chem Sci. 2018;9:513–530. doi:10.1039/C7SC02664A.",
    "[2] Huang K, Fu T, Gao W, et al. Therapeutics Data Commons: machine learning datasets and tasks for drug discovery and development. NeurIPS Datasets and Benchmarks. 2021.",
    "[3] Cawley GC, Talbot NLC. On over-fitting in model selection and subsequent selection bias in performance evaluation. J Mach Learn Res. 2010;11:2079–2107.",
    "[4] Varma S, Simon R. Bias in error estimation when using cross-validation for model selection. BMC Bioinformatics. 2006;7:91. doi:10.1186/1471-2105-7-91.",
    "[5] Zhao D, Zhu Y, Wu Z, et al. Revisiting ADMET prediction reliability under real-world challenges in the foundation model era. J Cheminform. 2026. doi:10.1186/s13321-026-01217-2.",
    "[6] Zhang L, Zeng Y, Qi Y, et al. DCPM-ADMET: fusion of dual-component pre-trained model and molecular fingerprints to enhance drug ADMET properties prediction. J Cheminform. 2026. doi:10.1186/s13321-026-01244-z.",
    "[7] Jang Y, Lee J, Jeong K, Kim J. Multimodal graph fusion with statistically guided parsimonious descriptor selection for molecular property prediction. J Cheminform. 2026;18:18. doi:10.1186/s13321-025-01140-y.",
    "[8] Zhang Y, Liu W, Zhao H, et al. MolGramTreeNet: a multimodal molecular property prediction model via grammar tree-constrained molecular representation. iScience. 2026;29:114928. doi:10.1016/j.isci.2026.114928.",
    "[9] Wen X, Liu H, Long W, Wei S, Zhu R. Consistent semantic representation learning for out-of-distribution molecular property prediction. Brief Bioinform. 2025;26:bbaf147. doi:10.1093/bib/bbaf147.",
    "[10] Yin T, Gao P, Panapitiya G, Saldanha EG. Out-of-distribution evaluation of active learning pipelines for molecular property prediction. RSC Adv. 2026;16:5281–5295. doi:10.1039/D5RA08055J.",
    "[11] Uchibori Y, Kaneko H. Generation of molecules near the applicability domain boundaries of property prediction models. J Chem Inf Model. 2026. doi:10.1021/acs.jcim.5c03220.",
    "[12] Kim JY, Vlachos DG. Distance-aware molecular property prediction in nonlinear structure-property space. J Chem Inf Model. 2025;65:6744–6756. doi:10.1021/acs.jcim.5c01037.",
    "[13] Tang H, Yue T, Li Y. Assessing uncertainty in machine learning for polymer property prediction: a benchmark study. J Chem Inf Model. 2025;65:6585–6598. doi:10.1021/acs.jcim.5c00550.",
    "[14] Fralish Z, Reker D. Pairwise learning for molecular property prediction and optimization. Front Drug Discov. 2026;6:1859068. doi:10.3389/fddsv.2026.1859068.",
    "[15] Landrum G. RDKit: open-source cheminformatics software. https://www.rdkit.org/.",
    "[16] Rogers D, Hahn M. Extended-connectivity fingerprints. J Chem Inf Model. 2010;50:742–754.",
    "[17] Breiman L. Random forests. Mach Learn. 2001;45:5–32.",
    "[18] Ke G, Meng Q, Finley T, et al. LightGBM: a highly efficient gradient boosting decision tree. Adv Neural Inf Process Syst. 2017.",
    "[19] Chen T, Guestrin C. XGBoost: a scalable tree boosting system. Proceedings of KDD. 2016.",
    "[20] Prokhorenkova L, Gusev G, Vorobev A, et al. CatBoost: unbiased boosting with categorical features. Adv Neural Inf Process Syst. 2018.",
    "[21] Yang K, Swanson K, Jin W, et al. Analyzing learned molecular representations for property prediction. J Chem Inf Model. 2019;59:3370–3388. doi:10.1021/acs.jcim.9b00237.",
    "[22] Chithrananda S, Grand G, Ramsundar B. ChemBERTa: large-scale self-supervised pretraining for molecular property prediction. arXiv:2010.09885. 2020.",
    "[23] Ross J, Belgodere B, Chenthamarakshan V, et al. Large-scale chemical language representations capture molecular structure and properties. Nat Mach Intell. 2022;4:1256–1264. doi:10.1038/s42256-022-00580-7.",
    "[24] Erickson N, Mueller J, Shirkov A, et al. AutoGluon-Tabular: robust and accurate AutoML for structured data. arXiv:2003.06505. 2020.",
    "[25] Tropsha A. Best practices for QSAR model development, validation, and exploitation. Mol Inform. 2010;29:476–488.",
    "[26] Vovk V, Gammerman A, Shafer G. Algorithmic Learning in a Random World. New York: Springer; 2005.",
    "[27] Shafer G, Vovk V. A tutorial on conformal prediction. J Mach Learn Res. 2008;9:371–421.",
    "[28] Guo C, Pleiss G, Sun Y, Weinberger KQ. On calibration of modern neural networks. Proceedings of ICML. 2017.",
    "[29] van Tilborg D, Alenicheva A, Grisoni F. Exposing the limitations of molecular machine learning with activity cliffs. J Chem Inf Model. 2022;62:5938–5951. doi:10.1021/acs.jcim.2c01073.",
    "[30] Sheridan RP. Time-split cross-validation as a method for estimating prospective prediction performance. J Chem Inf Model. 2013;53:783–790.",
    "[31] Demšar J. Statistical comparisons of classifiers over multiple data sets. J Mach Learn Res. 2006;7:1–30.",
    "[32] Hoyt CT, Zdrazil B, Guha R, et al. Improving reproducibility and reusability in the Journal of Cheminformatics. J Cheminform. 2023;15:62. doi:10.1186/s13321-023-00730-y.",
    "[33] Parrondo-Pizarro R, Lanini J, Rodriguez-Perez R. Uncertainty quantification in molecular machine learning for property predictions under data shifts. J Chem Inf Model. 2026;66:923–935. doi:10.1021/acs.jcim.5c02381.",
]


def add_manuscript(doc: Document) -> None:
    p = add_p(doc, "FZYC-Mol：候选扩张与多视图异质性下分子性质模型选择的验证治理", "Title", WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.first_line_indent = Cm(0)
    add_p(doc, "研究论文", "Normal", WD_ALIGN_PARAGRAPH.CENTER).paragraph_format.first_line_indent = Cm(0)

    add_h(doc, "摘要")
    add_p(doc, "分子性质预测已成为候选化合物排序、ADMET 预警和毒性筛查的重要计算环节，但模型与表示不断增加后，有限验证集同时承担选择和调优，可能使可达到性能上界与选择不确定性同步上升。FZYC-Mol 将候选登记、验证侧选择、策略冻结、外层审计和负结果记录组织为验证治理协议。九个终点的 3 外层×3 内层×5 重复实验中，K=32 相对 K=4 的配对固定分母遗憾增加 0.122（端点聚类 95% CI 0.072–0.175；精确 P=0.0078，Holm P=0.039），机会校正命中率下降 0.642。随机排序负对照和六级信号恢复正对照排除了固定 Top-3 机会率及指标失真的替代解释。原等权选择风险未通过层内置换（P=0.953）和留一端点门控，因而保留为负验证；严格嵌套的跨端点元风险将留出端点 MAE 从 0.123 降至 0.112，并在 8/9 个终点降低 50% 门控遗憾。共享冻结划分上的 12 个多视图候选完成 6,480 次拟合，validation-best 遗憾为 0.043（0.021–0.067），相对 Morgan-only 的配对效用增益为 0.343（0.210–0.483；9/9 端点）。MoleculeNet、TDC、逐样本风险、标签条件保形、MoleculeACE 和 bRo5 结果进一步界定了应用与化学边界。该证据链表明，候选扩张既能增加可达到上界，也会放大验证选择风险；冻结治理能够暴露并部分管理这一权衡，但不保证跨终点普遍增益。")
    add_h(doc, "科学贡献")
    add_p(doc, "FZYC-Mol 将候选池规模作为受控变量，通过三种候选组成随机化、3×3×5 重复嵌套、固定分母遗憾和端点配对精确推断识别验证选择损失。")
    add_p(doc, "随机排序负对照与连续信号注入正对照共同校准机会校正命中、MRR 和遗憾，原等权风险的失败也被完整保留。")
    add_p(doc, "严格留出端点的元风险检验了预警的有限迁移性，共享划分多视图重训则检验治理能否兑现异质候选收益。")
    add_p(doc, "关键词：分子性质预测；模型选择治理；候选池扩张；多视图候选；嵌套验证；跨端点风险；保形预测；化学边界")

    add_h(doc, "1 引言")
    add_p(doc, "溶解度、脂溶性、渗透性、毒性和药代动力学预测会影响药物发现早期的候选排序、实验排队与风险复核。公开基准上的单一 ROC-AUC 或 RMSE 可以比较模型，却不能说明被选模型是否稳定、某个预测是否处于适用域内，以及何时应转入实验或人工复核。随机划分、类别不平衡、低相似度分子、规则五以外化学空间和活性悬崖进一步扩大了评价指标与实际决策之间的距离[1,2,25,29]。")
    add_p(doc, "这一问题随候选空间扩张而加剧。Morgan 与 MACCS 指纹、二维描述符、树模型、图神经网络、D-MPNN、预训练化学语言模型和预测融合可以提供互补表征[16–23]，但它们也把少量算法比较转变为异质候选池中的多重决策。候选增加可能提高测试事后可达到上界，同时增加验证集被查询和适配的次数，因此“选择哪个模型”本身成为需要独立评价的方法学问题。")
    add_p(doc, "即使测试标签从未直接训练模型，持续依据同一验证集调整表示、超参数、融合、补救头和阈值，仍可能产生适应性过拟合。交叉验证同时承担调参与性能评价时会引入选择偏差[3,4]；在分子任务中，小样本、骨架分组和端点异质性使验证排序更易波动。仅报告最终测试分数无法区分真实表征增益、候选数增加带来的机会效应和事后选择。")
    add_p(doc, "近期研究已覆盖现实 ADMET 挑战、多模态融合、分布外学习、适用域、不确定性估计和分子对建模[5–14,33]。这些工作建立了困难场景和可靠性工具的必要性，但候选池扩张如何改变验证排序保真度，以及风险能否在读取外层标签前被诊断，仍缺少受控证据。因此，本文不把多视图、适用域或保形预测表述为首创模块，而将它们作为验证治理的收益与边界载体。")
    add_p(doc, "本文提出的核心问题是：在验证信息固定时，候选池从 K=4 扩展到 8、16 和 32 是否增加测试遗憾并降低排序保真度，以及这种风险能否仅凭验证侧证据在新终点上迁移。FZYC-Mol 在本文中指一套由候选登记、内层选择、策略冻结、外层审计和负结果记录组成的验证治理协议，不是新的主干预测网络。")
    add_p(doc, "为回答这一问题，我们依次采用随机化候选组成、重复嵌套重训、机会校正指标、固定分母遗憾、随机排序负对照、信号恢复正对照、精确符号翻转、严格跨端点元风险和共享划分多视图确认。随后以 MoleculeNet、AutoGluon 和 TDC 检查应用性能与终点异质性，并用逐样本风险、标签条件保形、低相似度、MoleculeACE 与 bRo5 界定预测和化学边界。")

    add_h(doc, "2 材料与方法")
    add_h(doc, "2.1 研究范围、数据集与任务登记", 2)
    add_p(doc, "主面板包括 ESOL、FreeSolv、Lipophilicity、BBBP、BACE 和 ClinTox 六个 MoleculeNet 任务。前三项为回归，主指标为 RMSE；后三项为分类，主指标为 ROC-AUC，并记录 PR-AUC、Brier 分数、期望校准误差和固定精度召回。外部 ADMET 面板采用 TDC 的 18 个 ADME 与 4 个毒性终点，遵循官方任务类型、指标和骨架划分。化学边界包含 17 个实际完成的 MoleculeACE 配置，以及 CycPept-PAMPA、LinPept CellPen 和 LinPept NonFouling 公共 bRo5 数据。")
    add_p(doc, "终点登记表在运行前固定数据版本、清洗前后样本数、标签定义与单位、指标方向、划分、种子和允许进入的候选家族。终点登记与候选登记相互独立，避免观察测试结果后扩大研究对象或搜索空间。")
    add_table(doc, 1, *TABLES[1])

    add_h(doc, "2.2 数据标准化、重复处理与泄漏审计", 2)
    add_p(doc, "SMILES 经 RDKit 解析与 Cleanup 后选择最大分子片段，在可行时中和电荷，并生成带立体信息的 canonical SMILES。分类重复结构仅在标签一致时合并，冲突组整体排除；回归重复结构按均值聚合并记录重复数。14 个主流程终点共输入 53,878 条记录，得到 53,522 个唯一结构；15 条无效 SMILES、186 条一致重复和 155 条冲突重复均保存逐行原因。")
    add_p(doc, "泄漏审计覆盖完全相同的 SMILES、Bemis–Murcko 骨架和 Morgan 最近邻相似度。所有插补、标准化、类别权重、目标变换和概率校准只在相应训练折估计。运行身份由 config hash、data hash、split hash、seed、code hash 和 prediction hash 联合确定；缺少身份字段的历史文件标为 unverified，不依据指标相同静默合并。")

    add_h(doc, "2.3 分子表示、候选专家与冻结登记", 2)
    add_p(doc, "历史探索包含分子图、Chemprop、ChemBERTa、MoLFormer、指纹、二维描述符、片段/骨架和预测融合。由于历史重型候选未在统一外层划分重训，它们只用于说明候选来源，不能进入确认性效应估计。32 候选扩池实验使用可稳定复跑的 Morgan-512 轻量变体；异质确认实验则将 Morgan-512、MACCS、RDKit2D 和三者拼接视图分别与线性模型、80 树随机森林和 80 轮 LightGBM 配对，形成 12 个预登记候选。")
    add_p(doc, "候选登记字段包括 candidate_id、family、representation、eligible、status、registry_order、complexity_level 和 config_hash。状态分为 eligible、rejected、failed 和 missing-data。测试结果较差、运行失败或数据缺失不会触发事后替换；所有状态均进入选择日志。")
    add_table(doc, 2, *TABLES[2])
    add_figure(doc, 1, "FZYC-Mol 总体工作流与证据层级。数据协议、泄漏控制、多视图候选、验证治理和可靠性输出在最终测试前形成单向冻结流程。历史重型候选以虚线边界标示，仅在共享划分重训后可进入确认性池；eligible、rejected、failed 和 missing-data 状态均被记录。该图为概念图。")

    add_h(doc, "2.4 验证治理选择器的形式化定义", 2)
    add_p(doc, "设 t 为任务，a 为候选，s 为内层重复或外层单元，r∈{val,test} 为评价分区，d_t 将指标统一为正向效用。验证均值、单标准误集合和风险调整分数定义为：")
    add_equation(doc, "uₜₐₛ⁽ʳ⁾ = dₜmₜₐₛ⁽ʳ⁾,   dₜ ∈ {−1,+1}                                                (1)")
    add_equation(doc, "μₜₐ = S⁻¹∑ₛuₜₐₛ⁽ᵛᵃˡ⁾,   SEₜₐ = SD(uₜₐ·⁽ᵛᵃˡ⁾)/√S                         (2)")
    add_equation(doc, "Aₜ¹ˢᴱ = {a: μₜₐ ≥ μₜ* − SEₜ*}                                                    (3)")
    add_equation(doc, "qₜₐ = μₜₐ − 0.5SD(uₜₐ·⁽ᵛᵃˡ⁾)                                                      (4)")
    add_p(doc, "one-SE 集合内依次最大化选择频率，再最小化折间波动、校准损失、计算成本和稳定 candidate_id。测试标签不参与集合构建、排序、阈值、权重或平局处理。外层绝对遗憾、完整 32 池固定分母遗憾和机会校正 Top-3 命中定义为：")
    add_equation(doc, "Rₜₛ = maxₐuₜₐₛ⁽ᵗᵉˢᵗ⁾ − uₜₐ*ₛ⁽ᵗᵉˢᵗ⁾                                            (5)")
    add_equation(doc, "R̃ₜₛ⁽³²⁾ = Rₜₛ/[maxₐ∈A³²uₜₐₛ⁽ᵗᵉˢᵗ⁾ − minₐ∈A³²uₜₐₛ⁽ᵗᵉˢᵗ⁾]       (6)")
    add_equation(doc, "CAHit@3ₜₛ = [Hit@3ₜₛ − 3/K]/[1 − 3/K],   K>3                              (7)")
    add_p(doc, "动态池内分母遗憾仅作敏感性分析。跨 K 比较同时报告 Top-25% 命中、MRR、排名百分位、NDCG、Spearman 和 Kendall，以防止结论依赖单一排名指标。")

    add_h(doc, "2.5 嵌套验证、候选池压力与选择器对照", 2)
    add_p(doc, "核心实验覆盖 BBBP、BACE、ClinTox、ESOL、FreeSolv、Lipophilicity、Caco2、HIA 和 P-gp。每个终点使用 3 个外层骨架折、3 个内层骨架折和 5 个预登记种子（11、23、37、53、71），形成每个池规模 135 个终点-重复-外层单元。K 固定为 4/8/16/32，并对 random-order、random-subset 和 family-balanced 三种模式各生成 100 个冻结顺序或子集。")
    add_h(doc, "2.5.1 选择风险、随机排序负对照与治理迁移", 3)
    add_p(doc, "等权 selection-risk 只由验证侧可见量构成，包括 one-SE 歧义、1−选择频率和内折波动百分位。随机排序负对照在每个冻结外层单元内置换验证效用与外层效用的候选对应关系，各 K 进行 1,000 次。治理规则迁移采用 leave-one-endpoint-out（LOEO）：仅用其余八个终点选择平均遗憾最低的规则，再评价留出终点。")
    add_h(doc, "2.5.2 配对推断与信号恢复正对照", 3)
    add_p(doc, "扩池效应在同一终点-重复-外层折内配对。先在终点内平均 15 个外层单元，再对九个终点执行全部 2⁹ 种符号翻转；固定遗憾、CAHit 和 MRR 的三种池比较共九项检验，采用 Holm 法控制家族错误率。区间由 10,000 次端点聚类 bootstrap 给出。")
    add_p(doc, "正对照在真实候选外层效用分布上构造合成验证分数 s·z(u_outer)+√(1−s²)·ε，其中 s=0、0.10、0.25、0.50、0.75 和 1.00，ε 为标准正态噪声。除完全信号外，每个外层单元和池规模重复 500 次。该模拟只校准指标零点与检出能力，不作为模型性能证据。")
    add_h(doc, "2.5.3 严格嵌套的跨端点元风险", 3)
    add_p(doc, "每个选择单元提取 16 个不读取外层标签的验证特征，包括池规模、任务类型、验证效用离散度、前两名标准化间隔、折间排名一致性、赢家频率、候选家族数、one-SE 大小和原风险分量。最外层完整留出一个终点；其余八个终点内部再次 LOEO，并按 MAE 在固定 Ridge、浅层随机森林和强正则 HistGradientBoosting 中选模。高遗憾阈值、模型类型和全部参数均由训练端点确定。")
    add_h(doc, "2.5.4 共享冻结划分的多视图候选确认", 3)
    add_p(doc, "12 个多视图候选使用与主确认实验相同的 5 个种子、3 外层和 3 内层划分。插补和标准化仅在训练折拟合，数值不稳定的 RDKit Ipc 在查看标签前排除。每个候选完成 3 个内折拟合和一次外层重拟合，共 12×9×5×3×4=6,480 次拟合。比较完整池、Morgan-only、双指纹池和不含拼接视图池，并评价 fixed Morgan RF、validation-best、one-SE、风险调整和 test oracle。")

    add_h(doc, "2.6 逐样本预测风险、校准与保形预测", 2)
    add_p(doc, "适用域由测试分子到训练集最近邻的 Morgan Tanimoto 相似度、集成分歧、描述符距离和重构误差表征。逐样本预测风险只在训练/验证侧拟合，测试推理前冻结。选择性预测按风险从低到高保留样本，分类风险为错误率，回归风险为 RMSE；E-AURC 定义为冻结风险曲线面积减去真实误差排序形成的 oracle 风险下界面积。")
    add_p(doc, "分类概率校准比较未校准、Platt 或温度缩放与等距回归，并在独立校准集上按 Brier 分数和 ECE 选择。分类保形采用标签条件阈值，回归采用绝对残差分位数，标称覆盖率为 80%、90% 和 95%。类别校准样本不足时回退到 pooled 阈值并记录 fallback_reason。回归同时报告训练标签 SD 标准化区间宽度。")

    add_h(doc, "2.7 活性悬崖、bRo5 与解释性分析", 2)
    add_p(doc, "MoleculeACE 除整体 RMSE 外，单独计算活性悬崖子集 RMSE，并比较高相似分子对的预测差异与真实差异 Spearman、Pearson、方向准确率和 gap MAE。CycPept-PAMPA 比较随机、骨架、外缘和时间划分；LinPept 比较随机、骨架和外缘划分。片段/基序分析同时报告最小支持数、效应量、Fisher exact P、Benjamini–Hochberg FDR 和跨种子稳定性，仅解释为关联。")

    add_h(doc, "2.8 统计分析、计算成本与开放科学", 2)
    add_p(doc, "MoleculeNet 报告 5 个随机种子的均值与标准差，TDC 和 bRo5 报告 3 个随机种子。候选池与重复嵌套的主不确定性以 endpoint 为聚类单位；候选行、外层折、种子或分子对不被当作独立生物学重复。AutoGluon 同时记录预算、实际拟合时间、模型数和峰值内存。复现包包含环境锁定、Dockerfile、持续集成配置、数据清洗、候选登记、逐折结果、source data 和 SHA-256 清单；公开 release、Zenodo DOI 和第三方空环境复跑尚待作者完成。")
    add_h(doc, "2.9 双层决策卡与选择日志", 2)
    add_p(doc, "终点级决策卡记录 candidate_id、selector、验证与外层指标、固定分母遗憾、晋级/保留/拒绝状态和失败类型。样本级决策卡记录点预测、AD 相似度、逐样本预测风险、校准状态、保形集合或区间、最近邻案例和 accept/manual review/outside AD 状态。两类风险使用不同统计单位，不作直接数值比较。")
    add_h(doc, "2.10 主张层级与终止规则", 2)
    add_p(doc, "确认性主张限定为机会校正排名保真度、完整 32 池固定分母遗憾、3×3×5 重复嵌套、严格留出风险验证和共享划分轻量多视图。MoleculeNet、AutoGluon、TDC、逐样本风险、保形、MoleculeACE 与 bRo5 为应用或边界证据。若新候选未改善预定验证指标，或改善位于 one-SE 容差内但稳定性、校准或成本更差，则终止晋级并记录负结果。")

    add_h(doc, "3 结果")
    add_h(doc, "3.1 随机化候选池控制确认规模效应", 2)
    add_p(doc, "random-subset 模式下，完整 32 池固定分母遗憾由 K=4 的 0.090（95% CI 0.041–0.154）升至 K=32 的 0.214（0.143–0.293），CAHit 由 0.723 降至 0.142，MRR 由 0.712 降至 0.223。random-order 的固定遗憾为 0.089/0.129/0.170/0.214，family-balanced 为 0.087/0.125/0.162/0.214。三种控制均保留规模相关趋势，因此固定登记顺序或候选家族比例不能单独解释该结果。")
    add_p(doc, "随机排序负对照给出了零信息参照。K=4/8/16/32 时，置换 CAHit 为 0.005/−0.002/−0.002/0.000，而真实重复嵌套流程为 0.881/0.550/0.334/0.240。真实固定遗憾仍低于置换值，表明扩大候选池削弱但未消除可用排序信息。")
    add_figure(doc, 2, "随机化候选池与置换负对照。三种候选组成控制显示完整 32 池固定分母遗憾与 MRR 随 K 变化；1,000 次候选对应置换定义零信息 CAHit 与遗憾基线。K=32 时三种组成模式等同。")

    add_h(doc, "3.2 重复嵌套验证确认选择不确定性", 2)
    add_p(doc, "5 个预登记种子均完成 3 外层×3 内层重训。K=4/8/16/32 的固定分母遗憾为 0.049/0.131/0.139/0.171，CAHit 为 0.881/0.550/0.334/0.240，MRR 为 0.869/0.518/0.401/0.264。主导候选比例由 0.859 降至 0.326，one-SE 集合成对 Jaccard 由 0.787 降至 0.418，显示候选增加同时降低排序保真度和选择稳定性。")
    add_p(doc, "端点配对推断避免把 135 个外层单元误作独立数据集。K=32 相对 K=4 的平均固定遗憾增加 0.122（95% CI 0.072–0.175），九个终点中八个同向，精确 P=0.0078；九项比较经 Holm 校正后 P=0.039。CAHit 的配对变化为 −0.642（−0.780–−0.442），同样有八个终点下降。")
    add_p(doc, "AutoGluon 在 30/300/1,800 s 上限下相对 32 候选 validation-best 均为 7/0/2 个终点胜/平/负，实际总拟合时间为 216.9/497.5/495.3 s。该强对照说明 FZYC-Mol 的贡献是可追踪选择与偏差审计，而不是点预测普遍优于 AutoML。")
    add_table(doc, 3, *TABLES[3])
    add_figure(doc, 3, "3×3×5 重复嵌套扩池效应。固定分母遗憾和机会校正命中随 K 变化，选择稳定性由主导候选比例与 one-SE 集合一致性表征。区间以 endpoint 为主要聚类单位。")

    add_h(doc, "3.3 指标校准与跨端点风险形成正负证据闭环", 2)
    add_p(doc, "六级信号恢复正对照在全部 K 上使 CAHit 和 MRR 随注入信号单调增加、固定遗憾单调降低。零信号时 CAHit 绝对偏差不超过 0.004，完全信号时遗憾为 0。主指标因此既具有正确零点，也能恢复逐级增强的排序信息。")
    add_p(doc, "原等权 selection-risk 的总体 Spearman 为 0.235，遗憾由最低风险四分位的 0.076 增至最高四分位的 0.169；但端点×K 层内置换 P=0.953。严格 LOEO 时，其高遗憾 AUC 为 0.514，50% 门控遗憾变化为 −0.009（95% CI −0.038–0.020），仅 5/9 个终点改善。该分数因此降级为描述性负验证。")
    add_p(doc, "严格嵌套元风险仅使用其他端点的验证特征，在完全留出端点上取得 MAE 0.112，相比常数基线 0.123 更低；Spearman 为 0.313，高遗憾 AUC 为 0.648。按预测风险保留最低 50% 单元时，平均遗憾降低 0.034（95% CI 0.020–0.047），8/9 个终点改善。该结果支持有限跨端点预警，但九个终点不足以建立通用元选择器。")
    add_figure(doc, 4, "指标校准与跨端点风险。a，K=32 相对 K=4 的端点配对固定遗憾变化。b，K=32 信号恢复正对照。c，原等权 selection-risk 的负验证。d，严格 LOEO 元风险在 50% 保留覆盖下的遗憾变化，绿色表示改善；浅绿色带为总体端点聚类 95% CI。")

    add_h(doc, "3.4 共享冻结划分的多视图候选兑现可达到收益", 2)
    add_p(doc, "九个终点、五个重复和三个外层折均完成 12 候选重训，共 135 个外层单元和 6,480 次拟合，无候选缺失或折失败。完整池 validation-best 的平均标准化遗憾为 0.043（95% CI 0.021–0.067），低于风险调整 0.054、one-SE 0.073 和固定 Morgan RF 0.395。")
    add_p(doc, "完整多视图 test oracle 相对 Morgan-only oracle 的可达到效用增益为 0.347，validation-best 实际兑现增益为 0.343（0.210–0.483；9/9 端点，P=0.0039）。拼接视图相对仅允许独立视图候选的增益为 0.035（0.017–0.053；9/9，P=0.0039）。135 次 validation-best 选择中，拼接视图被选 84 次、RDKit2D 44 次、MACCS 4 次、Morgan 3 次；表征收益不等同于固定学习器普适占优。")
    add_figure(doc, 5, "共享划分多视图确认。a，完整 12 候选池的治理规则遗憾。b，可达到和实际兑现的配对效用增益。c，135 个外层单元中 validation-best 选择的表示计数。d，终点×表示选择热图。误差条为端点聚类 95% CI。")

    add_h(doc, "3.5 MoleculeNet 冻结性能与筛选语境", 2)
    add_p(doc, "六个 MoleculeNet 任务的最终结果均可追溯到冻结候选。ESOL、FreeSolv 和 Lipophilicity 的 RMSE 分别为 0.5829±0.0352、1.0286±0.1761 和 0.6835±0.0439；BBBP、BACE 和 ClinTox 的 ROC-AUC 分别为 0.9243±0.0247、0.8753±0.0230 和 0.9496±0.0262。")
    add_p(doc, "FreeSolv 的 test oracle Chemprop 为 0.9518±0.1314，仅用于量化冻结选择差距。ClinTox 阳性率约 7%，在精度不低于 0.80 和 0.90 时的召回率分别为 0.588±0.168 和 0.491±0.195；每个测试划分仅含 5–14 个阳性样本，较大波动限制了固定阈值筛选。历史全候选审计的 Spearman 中位数为 0.667，Top-1 一致率为 0.135，测试 oracle 进入验证 Top-3 的比例为 0.295，这些值仅作为应用背景，不替代机会校正主结果。")
    add_table(doc, 4, *TABLES[4])
    add_figure(doc, 6, "MoleculeNet 冻结性能、ClinTox 固定精度召回与历史排序审计。分类与回归分面报告 5 种子均值和标准差；test oracle 只表示事后上界。ClinTox 召回并列展示少数类筛选边界，FreeSolv 保留冻结选择差距。")

    add_h(doc, "3.6 TDC 门控揭示终点异质性", 2)
    add_p(doc, "TDC 的 22 个终点经统一冻结门控后得到 5 promoted 和 17 retained。晋级终点为 clearance_hepatocyte_az、clearance_microsome_az、ppbr_az、vdss_lombardo 和 half_life_obach。事后审计进一步得到 3 个 promoted-and-improved、7 个 retained-and-avoided-harm 和 12 个 inconclusive 终点，没有 promoted-but-harmed。")
    add_p(doc, "前两个清除率终点的三种子区间跨零，ppbr、vdss 和 half-life 的区间保持为正。17 个 retained 终点不能写成“零下降”：其中 7 个保留避免了测试伤害，其余多因三种子区间过宽而证据不足。TDC 在本文中是跨来源公开面板，不是独立时间外盲测。")
    add_table(doc, 5, *TABLES[5])
    add_figure(doc, 7, "TDC 门控有效性。a，5 promoted 与 17 retained 的事后类别。b，22 个终点的方向归一化测试变化及三种子区间；蓝色为晋级，灰色为保留，宽区间标记为 inconclusive。")

    add_h(doc, "3.7 逐样本风险与标签条件保形限定预测可靠性", 2)
    add_p(doc, "85 个终点-随机种子风险单元按正确损失定义重建。分类 AURC/E-AURC 中位数为 0.060/0.044，回归为 0.465/0.245，且所有 E-AURC 非负。oracle 曲线按真实逐样本误差排序，因此是风险下界；分类纵轴为错误率，回归纵轴为 RMSE。")
    add_p(doc, "分类保形在 80%/90%/95% 目标下的总体覆盖为 0.856/0.932/0.956，类别 1 覆盖仅为 0.673/0.807/0.832。每个标称水平有 5 个单元触发 pooled fallback，主要来自 ClinTox 少数类校准样本不足。回归覆盖为 0.823/0.925/0.962，按训练标签 SD 标准化的区间宽度为 0.879/1.329/2.005。总体覆盖接近标称值不能替代类别条件或低相似度审计。")
    add_figure(doc, 8, "逐样本预测风险与标签条件保形。a–c，BBBP、ClinTox 和 Caco2 的风险-覆盖曲线，oracle 为风险下界。d，分类总体和类别条件覆盖。e，回归经验覆盖。f，标准化回归区间宽度与分类 pooled fallback。两类输出均不构成自动拒用阈值。")

    add_h(doc, "3.8 MoleculeACE 与 bRo5 界定化学迁移边界", 2)
    add_p(doc, "17 个 MoleculeACE 任务均完成三种子运行，形成 51 个任务-种子单元。整体 RMSE 为 0.711，活性悬崖子集 RMSE 为 0.813。高相似分子对差异 Spearman 平均为 0.252，任务-种子范围为 −0.018–0.661，方向准确率为 0.750。方向判断高于随机并不等于活性差异幅度已准确恢复。")
    add_p(doc, "CycPept-PAMPA 在随机、骨架、外缘和时间划分下的 RMSE 为 0.547±0.021、0.727±0.009、0.876±0.012 和 0.768±0.013。LinPept CellPen 在随机、骨架和外缘划分下的 ROC-AUC 为 0.937±0.020、0.894±0.029 和 0.859±0.005；NonFouling 外缘 PR-AUC 为 0.698。更严格划分下的退化表明随机划分可能高估迁移性能，但这些公开划分不等同于真实项目前瞻验证。")
    add_table(doc, 6, *TABLES[6])
    add_table(doc, 7, *TABLES[7])
    add_figure(doc, 9, "化学边界。a,b，MoleculeACE 任务异质性及方向与幅度的区别。c，CycPept-PAMPA 划分压力。d，LinPept 迁移边界。e，低相似度代表失败。案例只用于解释边界。")

    add_h(doc, "3.9 治理消融、失败案例与负结果", 2)
    add_p(doc, "只改变选择规则时，validation-best、冻结 one-SE、one-SE 低方差和 one-SE 低成本的固定遗憾为 0.129/0.171/0.174/0.185；只改变候选组成时，完整池及移除 bagging、boosting 或 linear 的遗憾为 0.171/0.147/0.131/0.150。两类消融回答不同问题，均不支持预设某个规则或家族普遍占优。LOEO 规则平均遗憾为 0.129，端点内 oracle 为 0.121；FreeSolv 和 HIA 保留端点特异失败。")
    add_p(doc, "代表性案例连接选择级与分子级边界。ClinTox 的真实阳性样本预测为 0.273、风险分位为 0.913；half-life 样本真实值为 820、预测约 47.2；FreeSolv 与 Lipophilicity 低相似度案例的 Tanimoto 为 0.333 和 0.217。BACE 带正电含氮桥环片段支持数为 62、效应 Δ=+0.527，BBBP 羧酸片段支持数为 76、Δ=−0.602；ClinTox N-甲基哌嗪片段支持数仅 5，故仍为探索性关联。")
    add_p(doc, "固定风险调整策略在 32 个终点-指标单元中有 10 个正向和 22 个负向，平均变化 −0.014；稳定性平局规则有 7 个正向和 25 个负向，平均变化 −0.015。未晋级、规则失效、运行失败和数据缺失均以独立负结果类型归档。")
    add_table(doc, 8, *TABLES[8])
    add_figure(doc, 10, "治理规则、候选家族与 AutoML 强基线。a，治理规则消融。b，候选家族移除。c，AutoGluon 三预算相对 validation-best 的胜负与实际拟合时间。d，留一端点规则迁移与端点内 oracle 对照。")
    add_figure(doc, 11, "负结果、失败案例与证据追踪。a，固定策略负结果。b，代表性失败边界信号。c，片段关联强度受支持样本数约束。d，可追踪证据库存。")

    add_h(doc, "3.10 Source data 自动重建", 2)
    add_p(doc, "主文 11 幅图和 8 张表均由冻结 source data 或概念图脚本重建。主轻量池包含 540 个选择单元，多视图池包含 135 个外层单元和 6,480 次拟合；候选池控制、置换、bootstrap、TDC 22 终点、85 个风险单元、90 个保形单元和 AutoGluon 三预算均可追溯到结构化文件。公开 release、Zenodo DOI 和独立第三方冷启动尚未完成，故当前表述限定为分析级可重建。")

    add_h(doc, "4 讨论")
    add_p(doc, "受控候选组成、3×3×5 重复嵌套、端点配对精确检验和正负对照共同支持一个限定性结论：当验证信息不增加时，候选扩张会降低排序保真度并增加选择损失，但真实排序仍优于随机。该结论不依赖固定 Top-3 的机械机会变化，也没有把外层折误作独立数据集。")
    add_p(doc, "候选规模、家族和分子表示共同决定这一权衡。共享划分多视图实验将旧稿中不可比较的历史异质候选替换为同折重训，显示可达到上界与 validation-best 实际增益在九个终点上方向一致。多视图候选因此值得加入，但收益必须与新增选择自由度同时评价；该结果仍不能外推到尚未同折重训的 Chemprop、GNN、ChemBERTa 或 MoLFormer。")
    add_p(doc, "治理规则本身没有普适赢家。validation-best 在完整多视图池中取得最低平均遗憾，但 32 候选 Morgan 池的 LOEO 仍显示端点特异偏好。冻结规则的作用是约束事后自由度、保存未晋级与失败记录，并使策略变化可追溯，而不是保证 one-SE、稳定性或风险调整在所有终点更优。")
    add_p(doc, "AutoGluon 在三个预算上均赢得多数终点，TDC 又显示晋级证据强度随终点变化。这些结果限制了方法身份：FZYC-Mol 是候选登记、选择偏差审计和决策记录协议，不是全面优于 AutoML 的点预测器。MoleculeNet 和 TDC 提供应用性能语境，不能替代确认性选择实验。")
    add_p(doc, "预测可靠性需要区分三种统计单位。逐样本预测风险回答哪些测试分子应优先复核；原等权 selection-risk 只保留为描述性负验证；严格跨端点元风险则在选择单元层面提供有限迁移证据。标签条件保形、低相似度、MoleculeACE 和 bRo5 进一步说明，风险下降或总体覆盖接近标称值均不是部署充分条件。")
    add_p(doc, "本研究仍受公开离线基准、九个确认终点和轻量候选范围限制。P-gp 是严格元风险 50% 门控中唯一未改善的终点；公开 bRo5 时间划分不是独立前瞻性验证；片段与失败案例只能支持关联性解释。真正的时间外 ADMET 盲测、重型候选共享划分重训、公开仓库发布和第三方冷启动是后续验证，而不是当前完成性结论。")

    add_h(doc, "5 结论")
    add_p(doc, "FZYC-Mol 将候选登记、验证侧选择、策略冻结、外层审计和负结果记录组织为可复核的模型选择治理流程。")
    add_p(doc, "重复嵌套、端点配对推断和正负对照确认候选扩张增加选择损失；严格跨端点风险检验和共享划分多视图重训分别提供有限预警与收益兑现证据。")
    add_p(doc, "最稳健的含义是，候选扩张同时提高可达到上界和选择不确定性，冻结治理能够暴露并部分管理这一权衡，但不能保证普遍性能提升。该结论限于公开离线基准和本文重训的轻量多视图候选。")

    add_h(doc, "声明")
    add_h(doc, "数据与材料可用性", 2)
    add_p(doc, "本文使用 MoleculeNet、TDC、MoleculeACE 和公共 bRo5 数据。处理数据、划分索引、候选登记、逐样本预测、source data 和 SHA-256 清单已在本地复现包整理。公开仓库与 Zenodo 尚未发布：[公开仓库 URL]；[Zenodo DOI]。受许可证限制的数据仅提供下载脚本与校验值。")
    add_h(doc, "代码可用性", 2)
    add_p(doc, "数据清洗、候选池随机化、3×3×5 重复嵌套、AutoGluon、多视图确认、风险与保形分析及绘图脚本已形成工作流。公开 release 标签尚待作者补充：[代码仓库 URL 与 release 标签]。")
    add_h(doc, "伦理批准与参与同意", 2); add_p(doc, "不适用。本研究仅使用公开分子数据和离线计算实验，不涉及人体参与者、动物实验或可识别个人信息。")
    add_h(doc, "发表同意", 2); add_p(doc, "不适用。")
    add_h(doc, "利益冲突", 2); add_p(doc, "作者声明不存在利益冲突。")
    add_h(doc, "经费支持", 2); add_p(doc, "[请作者补充基金名称、编号及资助方角色；若无资助，请明确声明“本研究未获得专项资助”。]")
    add_h(doc, "作者贡献", 2); add_p(doc, "[请作者按照 CRediT 补充概念构思、方法学、软件、验证、形式分析、数据整理、初稿撰写、审阅与编辑和监督等贡献。]")
    add_h(doc, "致谢", 2); add_p(doc, "生成式人工智能使用说明：本研究使用 Codex 辅助代码实现、文稿结构整理、图表生成和数值一致性检查；实验设计、结果解释、引用和最终文字均由作者核验并承担责任。其他致谢请作者补充。")

    add_h(doc, "补充信息")
    add_p(doc, "补表 S1–S10 分别提供清洗与泄漏审计、候选登记、100 次随机候选控制、AutoGluon 三预算、22 个 TDC 门控、17 个 MoleculeACE 任务、治理与候选家族消融、风险与保形逐终点结果、失败案例与片段统计，以及近期研究边界。每幅主图配套 PNG、SVG、PDF、TIFF、source data 和生成脚本。")
    add_h(doc, "参考文献")
    for ref in REFERENCES:
        p = add_p(doc, ref)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)


def apply_global_format(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)
        section.header_distance = Cm(1.0); section.footer_distance = Cm(1.0)
    for p in doc.paragraphs:
        if p.style.name in {"Title", "Heading 1", "Heading 2", "Heading 3", "FigureCaption", "TableCaption", "TableNote", "Equation"}:
            p.paragraph_format.first_line_indent = Cm(0)
        elif p.style.name == "Normal":
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.name = "Arial"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            if p.style.name == "Normal": run.font.size = Pt(10.5)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "FigureCaption", "TableCaption", "TableNote", "Equation"]:
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def write_audit_artifacts() -> None:
    AUDIT.mkdir(parents=True, exist_ok=True)
    terminology = [
        ["验证治理", "候选登记、内层选择、冻结、外层审计与负结果记录", "不得等同于逐样本风险门控"],
        ["候选池扩张", "K=4→8/16/32，且说明组成控制", "不得泛指模型数量增加"],
        ["test oracle", "测试事后上界，仅用于遗憾参照", "不得称最终最佳模型"],
        ["promoted/retained/rejected", "终点级治理状态", "样本级使用 accept/manual review/outside AD"],
        ["selection-risk", "验证侧选择单元风险", "逐样本风险写作“逐样本预测风险”"],
        ["TDC 外部面板", "跨来源公开基准", "不得称前瞻性外部盲测"],
        ["异质候选", "Morgan、MACCS、RDKit2D、拼接视图×3 学习器", "不外推到未同折重训深度模型"],
        ["负结果", "未晋级、规则失效、运行失败或数据缺失", "必须区分类型"],
    ]
    pd.DataFrame(terminology, columns=["canonical_term", "definition", "prohibited_or_boundary"]).to_csv(AUDIT / "terminology_ledger.csv", index=False, encoding="utf-8-sig")
    fusion = [
        ["核心问题", "小论文-4", "keep", "候选扩张与验证治理"],
        ["应用背景与决策卡", "初稿-7", "merge", "并入引言、方法2.9和图9"],
        ["3×3旧嵌套与固定Top-3", "初稿-7", "retire", "由3×3×5和CAHit/MRR替换"],
        ["旧TDC Caco2/HIA/Pgp晋级", "初稿-7", "retire", "采用5 promoted/17 retained"],
        ["旧风险AUROC", "初稿-7", "replace", "采用AURC/E-AURC"],
        ["等权selection-risk", "两稿", "downgrade", "保留负验证"],
        ["严格元风险", "小论文-4", "keep", "确认性扩展"],
        ["共享划分多视图", "小论文-4", "keep", "异质候选确认"],
        ["MoleculeACE/bRo5/低相似度", "两稿", "merge", "化学边界"],
        ["片段与失败案例", "两稿", "supplement", "主文只留代表案例"],
    ]
    pd.DataFrame(fusion, columns=["content", "source", "action", "final_role"]).to_csv(AUDIT / "fusion_matrix.csv", index=False, encoding="utf-8-sig")
    claims = [
        ["C1", "候选扩张增加选择损失", "3×3×5、配对K32-K4", "+0.122；CI 0.072–0.175", "公开离线、9终点"],
        ["C2", "指标能恢复排序信号", "随机负对照+六级正对照", "零偏差<0.004；完全信号遗憾0", "仅校准指标"],
        ["C3", "等权风险不能独立门控", "层内置换+LOEO", "P=0.953；门控CI跨0", "描述性负结果"],
        ["C4", "元风险具有有限迁移价值", "严格嵌套LOEO", "MAE 0.112；AUC 0.648；8/9改善", "终点数有限"],
        ["C5", "多视图收益可被验证选择兑现", "共享划分6,480次拟合", "+0.343；9/9；P=0.0039", "轻量候选"],
        ["C6", "应用性能具有终点异质性", "MoleculeNet+TDC+AutoGluon", "5 promoted/17 retained；AutoML 7/0/2", "不宣称全面性能优势"],
        ["C7", "样本与化学边界仍存在", "风险/保形/MoleculeACE/bRo5", "类别1覆盖不足；cliff RMSE 0.813", "非部署充分条件"],
    ]
    pd.DataFrame(claims, columns=["claim_id", "claim", "evidence", "decisive_value", "boundary"]).to_csv(AUDIT / "claim_evidence_map.csv", index=False, encoding="utf-8-sig")

    master_rows = []
    for mid, endpoint, metric, value, ci, source, status in [
        ("pool_regret_k4","all9","fixed_regret",0.049,"0.015–0.090","results/source_data/repeated_nested_bootstrap.csv","confirmatory"),
        ("pool_regret_k32","all9","fixed_regret",0.171,"0.113–0.237","results/source_data/repeated_nested_bootstrap.csv","confirmatory"),
        ("paired_regret_k32_k4","all9","delta_fixed_regret",0.122,"0.072–0.175","results/reviewer_core_20260624/paired_pool_effects.csv","confirmatory"),
        ("equal_risk_perm","all9","permutation_p",0.953,"","results/reviewer_core_20260624/risk_component_summary.csv","negative_validation"),
        ("meta_risk_mae","all9","LOEO_MAE",0.112,"","results/reviewer_core_20260624/cross_endpoint_meta_risk_summary.json","confirmatory_extension"),
        ("meta_risk_auc","all9","high_regret_AUC",0.648,"","results/reviewer_core_20260624/cross_endpoint_meta_risk_summary.json","confirmatory_extension"),
        ("multiview_regret","all9","validation_best_regret",0.043,"0.021–0.067","results/reviewer_core_20260624/multiview_nested/multiview_values.json","confirmatory"),
        ("multiview_gain","all9","utility_gain",0.343,"0.210–0.483","results/reviewer_core_20260624/multiview_nested/paired_multiview_effects.csv","confirmatory"),
        ("clintox","ClinTox","ROC-AUC",0.9496,"SD 0.0262","results/source_data/moleculenet_main_results.csv","application"),
        ("tdc_promoted","TDC22","count",5,"","results/source_data/tdc_gate_audit.csv","application"),
        ("classification_eaurc","classification","E-AURC",0.044,"median","results/source_data/risk_coverage_metrics.csv","boundary"),
        ("moleculeace_cliff_rmse","MoleculeACE17","RMSE",0.813,"","results/source_data/moleculeace_inclusion.csv","boundary"),
    ]:
        master_rows.append({"metric_id":mid,"endpoint":endpoint,"split":"frozen","seed":"aggregate","candidate_set":"registered","selector":"as_specified","value":value,"CI":ci,"source_file":source,"status":status})
    pd.DataFrame(master_rows).to_csv(PACKAGE / "manuscript_master_values.csv", index=False, encoding="utf-8-sig")
    with (AUDIT / "source_manifest.csv").open("w", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f); w.writerow(["source_document", "role", "priority"])
        w.writerow([str(ROOT / "output" / "小论文-4.docx"), "confirmation, statistics and current values", 1])
        w.writerow([r"C:\Users\Administrator\Desktop\修改\初稿-7.docx", "application narrative and chemical boundaries", 2])
        w.writerow([r"C:\Users\Administrator\Downloads\FZYC-Mol_双文档全量融合说明书_内容实验结果图表版.docx", "governing fusion specification", 0])


SUPP_SOURCES = [
    ("S1", "逐数据集清洗、重复、冲突与泄漏审计", ROOT / "results" / "source_data" / "data_cleaning_flow.csv"),
    ("S2", "确认候选登记与配置", ROOT / "results" / "reviewer_core_20260624" / "multiview_nested" / "candidate_registry.csv"),
    ("S3", "随机候选池控制汇总", ROOT / "results" / "source_data" / "candidate_pool_summary.csv"),
    ("S4", "AutoGluon 三预算性能与成本", ROOT / "results" / "source_data" / "autogluon_budget.csv"),
    ("S5", "22 个 TDC 终点门控与事后类别", ROOT / "results" / "source_data" / "tdc_gate_audit.csv"),
    ("S6", "MoleculeACE 逐任务纳入与结果", ROOT / "results" / "source_data" / "moleculeace_inclusion.csv"),
    ("S7", "治理规则与候选家族消融", ROOT / "results" / "source_data" / "ablation_summary.csv"),
    ("S8", "保形、低相似度与风险覆盖", ROOT / "results" / "source_data" / "conformal_long.csv"),
    ("S9", "失败案例与片段统计", ROOT / "reports" / "supplement_experiment_revision_20260606" / "maintext_table_failure_cases_compact.csv"),
    ("S10", "近期研究边界与文献核验", ROOT / "reports" / "draft10_literature_review_20260621" / "extracted_recent_literature.md"),
]


def build_supplement() -> None:
    doc = Document(); ensure_styles(doc)
    add_p(doc, f"{DOC_LABEL} 补充表 S1–S10", "Title", WD_ALIGN_PARAGRAPH.CENTER)
    out_dir = PACKAGE / "supplement_tables"; out_dir.mkdir(parents=True, exist_ok=True)
    for code, title, path in SUPP_SOURCES:
        add_h(doc, f"补表 {code} | {title}")
        if path.suffix.lower() == ".csv" and path.exists():
            df = pd.read_csv(path)
            df.to_csv(out_dir / f"Table_{code}.csv", index=False, encoding="utf-8-sig")
            view = df.head(120).copy()
            if view.shape[1] > 9:
                view = view.iloc[:, :9]
            headers = [str(c) for c in view.columns]
            rows = [["" if pd.isna(v) else str(v) for v in row] for row in view.itertuples(index=False, name=None)]
            add_table(doc, int(code[1:]) + 100, title, headers, rows, f"完整 CSV 共 {len(df)} 行；Word 预览最多显示前 120 行和前 9 列。")
        elif path.exists():
            text = path.read_text(encoding="utf-8")
            (out_dir / f"Table_{code}.md").write_text(text, encoding="utf-8")
            add_p(doc, text[:12000])
        else:
            add_p(doc, f"源文件缺失：{path}")
    apply_global_format(doc); doc.save(SUPP)


def build() -> Path:
    write_audit_artifacts()
    doc = Document(); ensure_styles(doc)
    add_manuscript(doc); apply_global_format(doc)
    doc.core_properties.title = "小论文-5"
    doc.core_properties.subject = "FZYC-Mol validation governance"
    doc.core_properties.author = "[请作者补充]"
    doc.core_properties.comments = "依据双文档全量融合说明书、nature-writing 与 nature-polishing 规范生成"
    doc.save(OUTPUT)
    build_supplement()
    return OUTPUT


if __name__ == "__main__":
    print(build())
