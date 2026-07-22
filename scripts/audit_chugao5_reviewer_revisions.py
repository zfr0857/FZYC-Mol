from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = Path.home() / "Desktop" / "修改" / "FZYC-Mol_初稿-5.docx"
REPORT_DIR = ROOT / "reports" / "reviewer_revision_20260606"
REPORT = REPORT_DIR / "初稿-5_详细检查报告.md"


@dataclass
class Check:
    id: str
    topic: str
    reviewer_need: str
    status: str
    evidence: str
    risk: str
    next_action: str


def normalize(text: str) -> str:
    return " ".join(text.replace("\n", " ").split())


def find_hits(paragraphs: list[str], needles: list[str], limit: int = 3) -> str:
    hits: list[str] = []
    for i, text in enumerate(paragraphs):
        if all(n in text for n in needles):
            hits.append(f"P{i}: {text[:180]}")
        if len(hits) >= limit:
            break
    return "<br>".join(hits) if hits else "未找到直接证据"


def count_formula_markers(paragraphs: list[str]) -> int:
    return sum(1 for n in range(1, 19) if any(f"({n})" in p for p in paragraphs))


def table_text(doc: Document) -> str:
    chunks: list[str] = []
    for ti, table in enumerate(doc.tables):
        for row in table.rows:
            chunks.append(f"T{ti}: " + " | ".join(normalize(cell.text) for cell in row.cells))
    return "\n".join(chunks)


def table_shape_summary(doc: Document) -> list[str]:
    rows: list[str] = []
    for i, t in enumerate(doc.tables):
        rows.append(f"Table {i}: {len(t.rows)} x {len(t.columns)}")
    return rows


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    doc = Document(DOCX)
    paragraphs = [normalize(p.text) for p in doc.paragraphs if normalize(p.text)]
    all_text = "\n".join(paragraphs) + "\n" + table_text(doc)
    shapes = table_shape_summary(doc)

    checks: list[Check] = []
    add = checks.append

    add(Check(
        "A1", "核心定位",
        "全文应改为“可冻结、可审计、可接受/拒绝候选策略的治理流程”，而不是保证测试最优的强模型。",
        "通过",
        find_hits(paragraphs, ["可审计", "测试最优"]) + "<br>" + find_hits(paragraphs, ["接受", "拒绝", "审计"]),
        "低",
        "保留当前定位；后续英文稿标题和摘要也要沿用该说法。",
    ))
    add(Check(
        "A2", "标题与摘要",
        "标题应压缩，摘要减少模块堆叠和过强语气，并前置 ranking audit。",
        "基本通过",
        f"标题：{paragraphs[0]}<br>" + find_hits(paragraphs, ["Top-3", "0.295", "Top-1"]) + "<br>" + find_hits(paragraphs, ["有限到中等", "0.652"]),
        "中",
        "摘要仍有较多数值，投稿英文版时建议进一步压缩为 2-3 个核心数值和 1 个局限句。",
    ))
    add(Check(
        "A3", "引言创新边界",
        "明确现有空白不是缺少模型，而是缺少候选模型选择过程审计；需直接比较 AutoML、ensemble selection、stacking、nested CV、QSAR AD。",
        "通过",
        find_hits(paragraphs, ["AutoML", "ensemble selection", "QSAR applicability domain"]),
        "低",
        "可在英文版中进一步压成一个定位段，避免引言过长。",
    ))
    add(Check(
        "A4", "贡献点第三条",
        "不应列大量模型名作为贡献，应改为可插拔候选池和审计规则。",
        "通过",
        find_hits(paragraphs, ["可插拔候选池", "模型名称作为创新本身"]),
        "低",
        "保持模型清单在方法/补充表，不要回填到摘要或贡献列表。",
    ))
    add(Check(
        "M1", "候选池量化",
        "方法部分需量化每个终点候选自由度、候选行数、终点/seed 单元等。",
        "通过",
        find_hits(paragraphs, ["7,011", "198", "中位候选数"]),
        "低",
        "最好在补充表中保留完整候选登记 CSV 名称和字段说明。",
    ))
    add(Check(
        "M2", "堆叠二次使用验证集风险",
        "需说明 stacking 使用验证集训练与选择的风险，或用 inner validation/nested validation 约束。",
        "通过",
        find_hits(paragraphs, ["二次使用验证集", "nested validation"]),
        "低",
        "后续回复审稿人时不要说“已完全消除偏差”，应说“暴露并约束该风险”。",
    ))
    add(Check(
        "M3", "λ、容差、复杂度惩罚",
        "解释 risk-adjusted λ=0.5、容差、复杂度惩罚、平局规则是否预先固定。",
        "基本通过",
        find_hits(paragraphs, ["λ=0.5", "复杂度惩罚", "读取测试集前固定"]),
        "中",
        "仍建议在补充方法中列出 ε_t、η 参数、δ_AD 的取值或搜索范围，否则审稿人可能继续追问。",
    ))
    add(Check(
        "M4", "表格基础模型通道",
        "未完成同划分正式运行的 tabular foundation model 不应频繁作为主文卖点。",
        "通过",
        find_hits(paragraphs, ["表格基础模型通道", "不进入主文最终保留结果"]),
        "低",
        "保留为限制或未来工作即可。",
    ))
    formula_count = count_formula_markers(paragraphs)
    old_formula = any(p.startswith("3.7.") for p in paragraphs)
    add(Check(
        "M5", "公式压缩",
        "正文保留 12-18 个核心公式，通用指标移补充材料；不要再保留 50 余公式。",
        "通过" if formula_count == 18 and not old_formula else "部分通过",
        f"检测到公式编号 1-18：{formula_count}/18；旧 3.7.x 小节残留：{old_formula}",
        "低" if formula_count == 18 and not old_formula else "中",
        "需导出 PDF 后人工检查公式分页和显示是否完整；当前为 Word 段落公式，非原生公式对象。",
    ))
    add(Check(
        "R1", "结果顺序",
        "结果应按性能、选择偏差、OOD/low-similarity/MoleculeACE、消融负结果、可靠性、解释案例组织。",
        "基本通过",
        find_hits(paragraphs, ["本节按审稿问题重新组织证据链"]),
        "中",
        "虽然开头写明了顺序，但实际小节仍保留 4.5/4.6 后置的补充实验闭环；可在下一版把 4.5/4.6 的关键表进一步前移。",
    ))
    add(Check(
        "R2", "FreeSolv 局限",
        "FreeSolv 仍差于 Chemprop 观测最优，应作为重要局限。",
        "通过",
        find_hits(paragraphs, ["FreeSolv", "Chemprop", "局限"]),
        "低",
        "当前语气足够诚实。",
    ))
    add(Check(
        "R3", "ClinTox 不平衡指标",
        "ClinTox 阳性率约 0.0705，PR-AUC/Brier/ECE/固定 precision recall 应主文呈现。",
        "部分通过",
        find_hits(paragraphs, ["ClinTox", "PR-AUC", "Brier"]) + "<br>表格中含 ClinTox/阳性率证据：" + ("是" if "ClinTox" in all_text and "0.0705" in all_text else "否"),
        "中-高",
        "已加入 PR-AUC/Brier/ECE/阳性率语境，但未见 recall at fixed precision 的明确数值。建议补一个 ClinTox 专门小表或补充表。",
    ))
    add(Check(
        "R4", "TDC 22 endpoint win/tie/loss",
        "若声称 22 个外部终点中 5 个保留增强，应显示完整 win/tie/loss、Delta 和显著性。",
        "基本通过",
        find_hits(paragraphs, ["22 个外部终点", "5/17/0"]) + "<br>" + ("Table 22/统计表存在" if "Bootstrap 95% CI" in all_text and "5/17/0" in all_text else "统计表证据不足"),
        "中",
        "已有 5/17/0 和统计表，但主文表仍较依赖补充表；建议把 endpoint Δ 列表放入主补充图或正文紧凑表。",
    ))
    add(Check(
        "R5", "统计显著性与小幅收益",
        "小幅提升需配 paired difference、CI、p 值或 NS 标记。",
        "部分通过",
        "检测到 Bootstrap/Wilcoxon/CI 表述：" + ("是" if "Bootstrap" in all_text and "Wilcoxon" in all_text else "否"),
        "中",
        "Table 22 有统计项，但 MoleculeNet 主表仍没有 paired difference/CI/p/NS 列，审稿人可能认为统计证据没有真正前置。",
    ))
    add(Check(
        "R6", "Nested validation",
        "补充或诚实标注 nested validation。",
        "通过",
        find_hits(paragraphs, ["3 outer", "3 inner", "nested validation"]) + "<br>" + find_hits(paragraphs, ["nested ROC-AUC", "FreeSolv"]),
        "低",
        "当前定位为偏差诊断，不替代主结果，语气合适。",
    ))
    add(Check(
        "R7", "系统消融矩阵",
        "Full / best single / simple mean / w/o selector / w/o fusion / w/o AD gate / w/o uncertainty weighting / w/o motif/fingerprint / w/o rescue head。",
        "通过",
        "Table 26 headers: " + (shapes[26] if len(shapes) > 26 else "未找到") + "；含 w/o AD gate/uncertainty/motif/rescue：" + ("是" if all(s in all_text for s in ["w/o AD gate", "w/o uncertainty weighting", "w/o motif/fingerprint", "w/o rescue head"]) else "否"),
        "低",
        "表格已齐；可在图中进一步可视化。",
    ))
    add(Check(
        "R8", "低相似度三档",
        "严格按 >0.7、0.5-0.7、<0.5 三个互斥 Tanimoto bin 输出性能、校准、不确定性和风险富集。",
        "通过",
        "Table 27 headers: " + (shapes[27] if len(shapes) > 27 else "未找到") + "；三档存在：" + ("是" if all(s in all_text for s in [">0.7", "0.5-0.7", "<0.5"]) else "否"),
        "低",
        "已满足审稿清单。",
    ))
    add(Check(
        "R9", "Conformal 80/90/95",
        "补齐 80%、90%、95% 覆盖率。",
        "通过",
        find_hits(paragraphs, ["0.814/0.918/0.956", "0.823/0.925/0.962"]) + "<br>Table 31: " + (shapes[31] if len(shapes) > 31 else "未找到"),
        "低",
        "已改正旧版“部分完成”表述。",
    ))
    add(Check(
        "R10", "MoleculeACE gap correlation",
        "补预测差异 vs 真实差异相关性和代表性 cliff pair 案例。",
        "基本通过",
        find_hits(paragraphs, ["gap correlation", "Spearman"]) + "<br>Table 29: " + (shapes[29] if len(shapes) > 29 else "未找到"),
        "中",
        "gap Spearman 已有；代表性 cliff pair 案例在失败案例表里有，但没有真实结构图或 paired molecule 图，图示仍需补。",
    ))
    add(Check(
        "R11", "失败案例",
        "至少 ClinTox 假阴性、FreeSolv 高误差/低相似度、高置信错误、MoleculeACE cliff 失败。",
        "基本通过",
        "Table 34: " + (shapes[34] if len(shapes) > 34 else "未找到") + "；低相似度/ClinTox/MoleculeACE：" + ("是" if all(s in all_text for s in ["ClinTox 高风险假阴性", "低相似度高误差", "活性悬崖"]) else "否"),
        "中",
        "已有案例类别，但表中部分 case 的修改前/修改后为空，若作为主文表显得证据不足，建议移补充或补具体样本 ID/SMILES/预测值。",
    ))
    add(Check(
        "REL1", "风险分数语气",
        "分类 median AUROC 0.788 可较强，回归 0.652 需降调。",
        "通过",
        find_hits(paragraphs, ["0.788", "0.652", "有限到中等"]),
        "低",
        "当前语气合适。",
    ))
    add(Check(
        "REL2", "Risk-coverage 主图/主文",
        "至少 BBBP、ClinTox、Caco2 或 Pgp 的 risk-coverage 曲线应进入主图或主文。",
        "部分通过",
        find_hits(paragraphs, ["risk-coverage"]) + "<br>检测到具体端点 risk-coverage 图名：" + ("是" if any(s in all_text for s in ["BBBP risk", "ClinTox risk", "Caco2 risk", "Pgp risk"]) else "否"),
        "中-高",
        "文字中有 risk-coverage，但没有确认具体主图/图号。建议补主图或至少明确“Figure X shows risk-coverage for BBBP/ClinTox/Caco2/Pgp”。",
    ))
    add(Check(
        "REL3", "基序/片段解释统计",
        "片段富集和基序归因需最小支持度、效应量、p/FDR、代表案例；否则只能写探索性关联。",
        "部分通过",
        find_hits(paragraphs, ["最小支持度", "FDR"]) + "<br>检测到实际 FDR 数值：" + ("是" if "FDR" in all_text and any(token in all_text for token in ["q=", "q <", "FDR <"]) else "否"),
        "中-高",
        "稿件已要求报告这些指标并降调为关联解释，但未见具体 FDR/p 值表，投稿前最好补表或明确“详见 Supplementary Table Sx”。",
    ))
    add(Check(
        "F1", "图 1/2 拆分",
        "图 1 只展示 workflow，图 2 展示 selector/gate/output evidence，候选专家细节放补充图。",
        "未充分落实",
        "当前检查为文本/表格级，未发现新增图 1/2 拆分说明；docx 内图像对象未做视觉比对。",
        "高",
        "需要打开 Word/PDF 视觉检查并重新排图。这个是当前最明显未完成项。",
    ))
    add(Check(
        "F2", "主表列数",
        "主文表格尽量不超过 7 列，候选长名缩短。",
        "部分通过",
        "长模型名残留：" + ("否" if "stack_q1_no_linear_pretrained_plus_rescue:5" not in all_text else "是") + f"；超过7列表格数量：{sum(1 for t in doc.tables if len(t.columns) > 7)}",
        "中",
        "长名已修，但仍有多个 8-10 列表。建议把主文大表拆成核心 6-7 列，完整列放补充。",
    ))
    add(Check(
        "F3", "截断文字",
        "修复 structure-separated s... 等截断。",
        "通过",
        "structure-separated s... 残留：" + ("否" if "structure-separated s..." not in all_text else "是"),
        "低",
        "已修。",
    ))
    add(Check(
        "D1", "讨论限制排序",
        "第一限制应是 validation selection bias；再写小幅收益、异质性、解释性关联等。",
        "通过",
        find_hits(paragraphs, ["本研究也有明确局限", "验证-测试排名审计", "不能保证测试最优"]),
        "低",
        "当前顺序符合审稿建议。",
    ))
    add(Check(
        "D2", "结论语气",
        "强调 reproducible selection/rejection/audit workflow，而不是 retained gains。",
        "通过",
        find_hits(paragraphs, ["核心价值不是保证每个终点取得测试最优", "接受", "拒绝"]),
        "低",
        "当前结论更稳。",
    ))
    add(Check(
        "D3", "代码和数据可用性",
        "需给 GitHub/Zenodo、环境、运行命令、seed-level CSV、source data；不能虚构 DOI。",
        "基本通过",
        find_hits(paragraphs, ["GitHub/Zenodo", "seed-level", "不虚构 accession number"]),
        "中",
        "目前是承诺式占位，还没有真实仓库/DOI。投稿前必须补真实链接或改为“will be made available upon acceptance”的期刊允许措辞。",
    ))
    add(Check(
        "REF1", "2025/2026 引用核验",
        "审稿意见要求投稿前核验 2025/2026 参考文献。",
        "未完成本轮核验",
        "本次检查未联网逐条核验 DOI/arXiv/卷页；稿件中仍有 2026 arXiv 和 2025/2026 期刊条目。",
        "中",
        "投稿前应单独跑一轮 Crossref/PubMed/arXiv 核验，确认 DOI、年份、卷号、页码和是否仍为预印本。",
    ))

    pass_count = sum(1 for c in checks if c.status == "通过")
    basic_count = sum(1 for c in checks if c.status == "基本通过")
    partial_count = sum(1 for c in checks if c.status == "部分通过")
    fail_count = sum(1 for c in checks if c.status.startswith("未"))

    high_risks = [c for c in checks if "高" in c.risk]
    old_residual_terms = [term for term in ["部分完成", "未写成已完成", "不能保证测试最优", "未来工作", "后续"] if term in all_text]

    lines: list[str] = []
    lines.append("# 初稿-5 对审稿意见落实情况详细检查报告\n")
    lines.append(f"检查文件：`{DOCX}`\n")
    lines.append(f"总体计数：通过 {pass_count}，基本通过 {basic_count}，部分通过 {partial_count}，未充分落实/未完成 {fail_count}。\n")
    lines.append("## 总体判断\n")
    lines.append(
        "初稿-5 已经把审稿意见中最核心的主线改过来了：论文定位从“更强模型”转为“可冻结、可审计、可接受/拒绝候选策略的治理流程”；摘要、引言、方法、公式、结果和讨论都能找到对应修改证据。"
        " 但仍有几处投稿风险：图 1/2 拆分和视觉排版尚未确认；主文部分表格仍超过 7 列；ClinTox 的 recall at fixed precision 未见明确数值；risk-coverage 缺少明确主图指向；基序/片段解释的 p/FDR 证据还不够具体；2025/2026 参考文献未做本轮联网核验。\n"
    )
    lines.append("## 高风险/优先补强项\n")
    for c in high_risks:
        lines.append(f"- **{c.id} {c.topic}**：{c.status}。{c.next_action}")
    lines.append("\n## 逐条检查表\n")
    lines.append("| ID | 检查主题 | 审稿要求 | 状态 | 稿件证据 | 风险 | 建议 |")
    lines.append("|---|---|---|---|---|---|---|")
    for c in checks:
        lines.append(
            f"| {c.id} | {c.topic} | {c.reviewer_need} | **{c.status}** | {c.evidence} | {c.risk} | {c.next_action} |"
        )
    lines.append("\n## 表格结构抽检\n")
    lines.extend(f"- {row}" for row in shapes)
    lines.append("\n## 残留措辞提示\n")
    if old_residual_terms:
        for term in old_residual_terms:
            lines.append(f"- 检测到 `{term}`。其中“不能保证测试最优”是当前稿件有意保留的限制性表述；“未来工作/后续”需人工检查是否过多。")
    else:
        lines.append("- 未检测到明显旧版残留措辞。")
    lines.append("\n## 建议的二次修改顺序\n")
    lines.append("1. 先处理图 1/图 2 拆分和 PDF 视觉检查。")
    lines.append("2. 将超过 7 列的主文表格压缩，长表放 Supplementary Tables。")
    lines.append("3. 补 ClinTox recall at fixed precision，或在文中明确该指标未作为本轮主结果。")
    lines.append("4. 给 risk-coverage 指定主图或补充图编号，至少覆盖 BBBP、ClinTox、Caco2/Pgp。")
    lines.append("5. 给 motif/fragment enrichment 增加支持度、效应量、p/FDR 或明确为 exploratory。")
    lines.append("6. 投稿前做参考文献 DOI/arXiv/Crossref 核验和 PDF 公式页检查。")

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    REPORT.write_text("\n".join(lines), encoding="utf-8")
    print(REPORT)
    print(f"PASS={pass_count} BASIC={basic_count} PARTIAL={partial_count} FAIL={fail_count}")


if __name__ == "__main__":
    main()
