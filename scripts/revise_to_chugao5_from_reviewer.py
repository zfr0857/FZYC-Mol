from __future__ import annotations

import csv
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DESKTOP_DOCS = Path.home() / "Desktop" / "修改"
SOURCE = DESKTOP_DOCS / "FZYC-Mol_初稿4_公式部分_Nature居中无表格修订.docx"
TARGET = DESKTOP_DOCS / "FZYC-Mol_初稿-5.docx"
REPORT_DIR = ROOT / "reports" / "reviewer_revision_20260606"
FULL_RUN = ROOT / "reports" / "remaining_missing_experiments_20260606"


TITLE = "FZYC-Mol：可审计验证集治理的适用域感知分子性质预测框架"


REPLACEMENTS = [
    (
        "分子性质预测是药物发现、ADMET 评估和毒性风险筛查中的关键计算环节。",
        "分子性质预测是药物发现、ADMET 评估和毒性风险筛查中的关键计算环节。随机划分下的平均 ROC-AUC 或 RMSE 往往不足以反映模型在结构新颖分子、少样本终点、不平衡标签、活性悬崖、实验噪声和适用域漂移条件下的可靠性。为此，本文提出 FZYC-Mol，一种由验证集治理的可审计候选策略选择框架。该框架不把贡献建立在单一更大的主干模型上，而是将多视图表示、强基线专家、目标变换、融合候选、补救头、校准和适用域证据纳入冻结候选池；最终保留策略仅由验证集决定，测试集只在策略冻结后用于一次性评估。",
    ),
    (
        "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别达到",
        "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别达到 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302。由验证集决定的定向补救改善了 Lipophilicity，低成本重构改善了 FreeSolv，但 FreeSolv 仍落后于观测最佳 Chemprop 候选，因此被作为重要边界案例处理。多视图融合和适用域门控在 BBBP、ClinTox 以及外部 TDC ADMET 的 Caco2_Wang、HIA_Hou 和 Pgp_Broccatelli 上获得验证集接受的增益；22 个外部终点中最终保留 win/tie/loss 为 5/17/0，说明该框架更适合被理解为终点依赖的选择与拒绝流程，而不是统一性能提升器。",
    ),
    (
        "可靠性分析显示，集成标准差、错误模型、错误-适用域混合指标",
        "可靠性分析显示，风险分数对分类错误具有较好的识别能力，但对回归高误差样本仅表现为有限到中等识别能力；对应中位 AUROC 分别为 0.788 和 0.652。保形预测在 MoleculeNet 分类任务的 80%/90%/95% 目标覆盖下达到平均经验覆盖率 0.814/0.918/0.956，在回归任务下达到 0.823/0.925/0.962。新增验证-测试排名审计显示，跨 200 个 dataset-seed 候选池的中位 Spearman 为 0.667，测试最优落入验证 Top-3 的比例为 0.295，Top-1 一致比例为 0.135。该审计表明，验证集治理可以避免测试集事后挑选，但不能保证测试最优，必须结合 regret、optimism gap、嵌套验证和负结果审计共同解释。基序归因和片段富集仅作为关联性化学解释，而不被解读为因果机制证明。",
    ),
    (
        "本文贡献可概括为五点。",
        "本文贡献可概括为五点。第一，构建以 MoleculeNet、TDC ADMET、MoleculeACE、OpenADMET 附录和结构分离划分为核心的统一评测流程，并明确训练集、验证集和测试集的使用边界。第二，提出可插拔候选池、验证集选择器、补救整合选择器和最终保留门控，使候选策略必须先通过验证集证据才能进入主结果。第三，给出候选接受、拒绝和审计规则，而不是把 CatBoost、XGBoost、Chemprop、冻结表征、Top-K 集成等模型名称作为创新本身；完整候选族和超参数范围放在方法与补充表中。第四，系统报告验证-测试排名、嵌套验证、低相似度分层、MoleculeACE 活性悬崖、系统消融、保形预测、校准和风险-覆盖曲线，避免只依据单一 ROC-AUC 或 RMSE 下结论。第五，通过基序归因、片段富集和失败案例，将模型性能、可靠性边界、负结果和可读化学解释连接起来。",
    ),
    (
        "综合上述文献，FZYC-Mol 的研究空白不在于提出一个更大的主干模型",
        "综合上述文献，FZYC-Mol 的研究空白不在于提出一个更大的主干模型，而在于构建可审计的候选模型选择和可靠性治理流程。AutoML、ensemble selection、stacking、nested cross-validation 和 QSAR applicability domain 已分别研究了模型搜索、集成选择、验证偏差控制和适用域边界；FZYC-Mol 与这些工作的区别在于把它们收束为一套分子性质预测中的冻结候选池、验证集接受/拒绝规则、测试集一次性报告和负结果审计。近两年的 ADMET 可靠性、OOD 基准、表格基础模型、多尺度表征和片段预训练研究共同提示：真实分子性质预测中不存在单一全能模型；小样本终点中描述符、指纹、树模型和表格基础模型仍然是强对照；多尺度结构与片段解释能改善部分任务，但必须报告适用边界 [3,13-16,41-44]。因此，本研究把强表格基线、图模型、冻结表征、目标变换、欠采样集成、Top-K/堆叠集成、适用域门控、粗糙度诊断和可解释性分析放入同一验证集治理流程。",
    ),
    (
        "候选专家池包括 RF、ExtraTrees、XGBoost、LightGBM、CatBoost",
        "候选专家池包括 RF、ExtraTrees、XGBoost、LightGBM、CatBoost、Chemprop D-MPNN、图模型、描述符 MLP、基序专家、冻结预训练表征头、Top-K 均值、岭回归/逻辑回归堆叠、自适应共识、不确定性加权、适用域门控和定向补救头。回归任务额外包含原始目标、log1p、分位数正态化、截尾目标等目标变换；类别不平衡任务额外包含平衡欠采样集成和 PR-AUC、Brier、ECE、MCC、平衡准确率评估。已完成的候选登记表覆盖 7,011 个候选记录、198 个 dataset-seed 单元，中位候选数为 25；MoleculeNet 多方法融合、定向重构、TDC 多方法融合、性能增强候选和固定选择器审计分别在补充表中列出候选行数、覆盖终点、Top-1/Top-3 命中和负相关单元。表格基础模型通道仅作为补充和未来工作接口，不进入主文最终保留结果，也不被写成已完成的正式同划分比较。",
    ),
    (
        "选择器首先在每个终点的验证集预测表上计算官方主指标",
        "选择器首先在每个终点的验证集预测表上计算官方主指标，并按指标方向转为统一的正向验证效用。若单一专家显著占优，则保留最佳专家；若多个专家在验证集上互补，则候选 Top-K 均值、堆叠集成、不确定性加权融合或适用域门控融合。候选池、主指标、容差、risk-adjusted λ=0.5 敏感性规则、复杂度惩罚和平局规则均在读取测试集前固定；未超过当前最终保留结果的补救头、目标变换或强基线进入附录而不改变主结果。该规则使正结果、负结果和未接入候选都有明确位置，也降低了测试集事后选择偏差。",
    ),
    (
        "为审计验证集选择是否可能过拟合",
        "为审计验证集选择是否可能过拟合，本文额外计算每个 dataset-seed 候选池中验证集排名与测试集排名的 Spearman 相关性，并记录验证集第一名是否也是测试集第一名、测试集第一名是否落入验证集 Top-3，以及验证集第一名在测试集上的 regret/optimism gap。该分析不参与模型选择，只用于评估选择器风险。进一步地，本文在 BBBP、BACE、ClinTox、HIA、Pgp、ESOL、FreeSolv、Lipophilicity 和 Caco2 等 9 个代表性终点上完成 3 outer × 3 inner 的 nested validation，用于检查候选选择在内外层划分下是否仍保持可用性能；nested 结果作为偏差诊断而不替代主文冻结测试结果。",
    ),
    (
        "可靠性模块包括集成标准差、错误模型、预测偏差",
        "可靠性模块包括集成标准差、错误模型、预测偏差、反向 Tanimoto 距离、重构误差、错误-适用域混合指标、风险-覆盖曲线、保形覆盖率、校准和粗糙度代理指标。分类任务把预测标签错误作为错误检出目标，回归任务把绝对误差最高的 20% 样本作为高误差目标，并用 AUROC、AUPRC、风险-覆盖曲线和 top-10% 富集共同评估风险分数。需要强调的是，分类风险识别较强，回归高误差识别较有限；因此正文避免使用“能够识别所有高误差样本”的强表述。可解释性模块包括基序特征重要性、片段富集、骨架/近邻案例复核与高误差样本分析。本文只把这些证据解释为模型使用边界和化学相关性，不把关联性基序解释等同于因果机制。",
    ),
    (
        "为保证多专家比较可复现",
        "为保证多专家比较可复现，本文将候选模型治理拆成三个层级。第一层是单一专家训练：每个模型家族在固定划分和随机种子下独立训练，并保存验证集/测试集预测。第二层是候选策略生成：仅在验证集预测矩阵上构造 Top-K 均值、堆叠集成、自适应共识、不确定性加权、适用域门控和定向补救。第三层是最终选择器决策：只根据验证集主指标和预先定义的平局/风险调整规则确定最终保留结果。所有候选数量、特征维度、目标变换、门控阈值、随机种子、软件版本和图表来源 CSV 均进入补充材料；主文只保留能评估选择自由度和审计风险的汇总。",
    ),
    (
        "堆叠集成被定义为仅基于验证集的预测层候选",
        "堆叠集成被定义为预测层候选，而不是新的原始特征学习模型。回归任务使用岭回归式线性组合或均值，分类任务使用逻辑回归堆叠或概率均值。由于在同一验证集上训练堆叠权重并选择候选会产生二次使用验证集的风险，本文将堆叠结果解释为候选策略而非无偏性能估计，并通过 Top-3/Top-5 限制、随机种子汇总、验证-测试排名审计和代表性 nested validation 共同约束该风险。若验证集上不稳定或不能超过当前选择器，则保留原策略。",
    ),
    (
        "基序归因与片段富集提供互补解释。",
        "基序归因与片段富集提供互补解释。前者面向模型内部特征重要性，关注官能团、BRICS 片段和 Murcko 骨架对预测的贡献方向；后者面向数据分布，关注某些片段出现时标签均值相对基线的偏移。为避免事后解释过度，片段富集需同时报告最小支持度、效应量、p 值或 FDR 校正，并给出代表分子案例；若样本量不足或只出现方向一致而统计证据较弱，则仅作为探索性关联解释。两者一致时，说明模型关注点与数据统计信号相互支持；不一致时，则提示混杂因素、样本量不足或局部标签噪声。",
    ),
    (
        "本节按照科学问题组织结果",
        "本节按审稿问题重新组织证据链。第一，报告 MoleculeNet 和外部 ADMET 的基线与最终保留性能，并明确小幅增益和 FreeSolv 局限。第二，前置验证集选择偏差审计，包括验证-测试排名、Top-1/Top-3、regret、optimism gap 和 nested validation。第三，展示 OOD、低相似度三档和 MoleculeACE 活性悬崖结果。第四，报告系统消融、固定选择器和负结果。第五，解释可靠性、校准、风险-覆盖和保形预测。第六，用化学解释和失败案例说明模型在哪些样本上不应被过度信任。",
    ),
    (
        "MoleculeNet 主结果显示，FZYC-Mol 的优势并非来自某一专家在所有任务上的稳定胜出",
        "MoleculeNet 主结果显示，FZYC-Mol 的价值并非来自某一专家在所有任务上的稳定胜出，而是来自验证集选择器对不同终点候选策略的可审计接受与拒绝。ESOL 和 BACE 中，验证集选择器与观测最优一致；FreeSolv 中，Chemprop 是测试观测最优，而 FZYC-Mol 原选择器和定向重构仍未完全达到该水平，因此 FreeSolv 被作为选择器稳健性不足的核心局限；Lipophilicity 中，定向补救选择器将 RMSE 从 0.7078 降至 0.6835。进一步加入多方法融合候选后，BBBP 和 ClinTox 由排序融合、适用域门控融合或堆叠融合进入最终保留策略，但这些增益幅度较小，必须结合 seed 级差值、置信区间和选择器审计解释。",
    ),
    (
        "TDC ADMET 结果进一步说明",
        "TDC ADMET 结果进一步说明，外部 ADMET 终点的困难程度高度异质。HIA_Hou、Pgp_Broccatelli 和 BBB_Martins 上选择器的 ROC-AUC 分别达到约 0.979、0.936 和 0.918；Caco2_Wang 等回归终点则对目标变换和表格基线更敏感。多方法融合接入后，Caco2_Wang 的 RMSE 从 0.4517 降至 0.4375，HIA_Hou 的 ROC-AUC 从 0.9792 提升到 0.9827，Pgp_Broccatelli 的 ROC-AUC 从 0.9364 提升到 0.9473。更重要的是，22 个外部终点中只有 5 个由性能增强候选保留，17 个保留旧结果，最终保留 win/tie/loss 为 5/17/0；这比宣称所有外部任务统一提升更符合 ADMET 终点异质性。",
    ),
    (
        "可靠性结果显示，单一不确定性分数难以覆盖所有错误类型。",
        "可靠性结果显示，单一不确定性分数难以覆盖所有错误类型。集成标准差和错误模型对模型不一致性敏感，反向 Tanimoto 距离与重构误差更偏向适用域和结构新颖性，错误-适用域混合指标将两者组合后更适合作为实际使用时的风险标记。分类错误风险的中位 AUROC 为 0.788，具有较好筛查价值；回归高误差风险的中位 AUROC 为 0.652，只能说明有限到中等识别能力。对于 ClinTox、DILI、CYP 底物、Caco2 和 Pgp 等任务，PR-AUC、Brier、ECE、risk-coverage 和保形覆盖率必须与 ROC-AUC 或 RMSE 共同报告，以避免平均性能掩盖阳性样本稀缺、校准不足或低相似度风险。",
    ),
    (
        "注：P1 中的保形预测目前仅有 90% 经验覆盖率证据",
        "注：P1 保形预测已补齐 80%/90%/95% 三个覆盖率水平；主文报告均值，完整 seed-level 覆盖率、集合大小和区间宽度见补充表。",
    ),
    (
        "MoleculeACE 配对结果与低相似度分析互相补充。",
        "MoleculeACE 配对结果与低相似度分析互相补充。活性悬崖目标候选在 51 个 seed 配对中的总 RMSE 平均正向变化为 0.0069，悬崖子集 RMSE 平均正向变化为 0.0056，但标准差较大，任务间存在明显差异。新增 gap correlation 审计显示，预测差异与真实差异的平均 Spearman 约为 0.252，部分任务接近零或为负，说明模型能捕获部分悬崖方向，但尚不能可靠解决数值悬崖幅度预测。因此，本文将该结果解释为悬崖风险识别和候选治理的补充证据，而不是对活性悬崖预测已经解决的证明。",
    ),
    (
        "不平衡分类与保形预测的补充结果改善了 ClinTox",
        "不平衡分类与保形预测的补充结果改善了 ClinTox、DILI、hERG 和 CYP 底物等任务的解释。ROC-AUC 仍作为标准指标，但主文同时报告 PR-AUC、Brier、ECE、富集指标和样本级风险案例。保形预测已在 80%/90%/95% 目标覆盖率下完成：分类平均经验覆盖率为 0.814/0.918/0.956，回归平均经验覆盖率为 0.823/0.925/0.962；区间宽度和集合大小随覆盖目标增加而上升，符合保形预测的风险-信息量权衡。",
    ),
    (
        "最后，负结果和失败案例被明确保留。",
        "最后，负结果和失败案例被明确保留。ClinTox 高风险假阴性表明，高 ROC-AUC 任务仍需样本级风险证据；FreeSolv 低相似度高误差样本说明适用域外推仍是选择器短板；MoleculeACE 活性悬崖失败案例显示相似分子的真实性质差异仍可能被低估；3D-lite 和粗糙度加权在 oracle 条件下偶有潜在收益，但验证集门控未将其接入最终策略。这些结果共同界定 FZYC-Mol 的使用边界，也避免论文只报告正结果。",
    ),
    (
        "本研究也有明确局限。",
        "本研究也有明确局限。首先，验证-测试排名审计和 nested validation 共同表明，验证集治理可以避免测试集事后挑选，但仍存在验证集选择偏差，不能保证测试最优；因此所有小幅增益都需要与 regret、optimism gap、Top-3 命中和负结果一起解释。第二，收益具有明显终点异质性，BBBP、ClinTox、HIA、Pgp 等终点的增益较小，FreeSolv 仍落后于观测最佳 Chemprop 候选。第三，基序归因和片段富集属于关联解释，不应被解读为因果机制证明。第四，当前没有湿实验验证，ChemBERTa 与 MoLFormer 主要以冻结编码器形式使用，Polaris 与 OpenADMET 的完整官方挑战流程尚未完全纳入。",
    ),
    (
        "第二个可能问题是，选择器是否会过拟合验证集。",
        "第二个可能问题是，选择器是否会过拟合验证集。本文通过五个设计降低并暴露这一风险：第一，选择器只在预定义候选池中选择，不允许根据测试结果添加临时规则；第二，结果按多个随机种子汇总，并报告均值、标准差和配对统计；第三，补救头进入最终策略需要通过验证集接受，且没有被所有终点无差别采用；第四，排名审计、Top-3 命中、optimism gap 和负结果表将低相关候选池明确写为选择风险而非胜利证据；第五，9 个代表性终点的 3×3 nested validation 用于补充检查选择器在内外层划分中的稳定性。尽管如此，更大规模的跨数据集迁移验证仍是后续加强方向。",
    ),
    (
        "本文提出并系统评估了 FZYC-Mol",
        "本文提出并系统评估了 FZYC-Mol，一种由验证集治理、适用域感知的多专家分子性质预测框架。结果表明，在 MoleculeNet、TDC ADMET、外部评估附录、划分真实性、低相似度子集、MoleculeACE 活性悬崖、粗糙度诊断和基序/片段解释等多条证据线上，FZYC-Mol 能够提供比单一模型分数更完整的可靠性画像。其核心价值不是保证每个终点取得测试最优，而是在固定候选池内透明地接受、拒绝和审计候选策略；当补救头、低成本重构或多方法融合通过验证集门控时，它们可以进入最终保留结果，当证据不足时则作为负结果保留。",
    ),
    (
        "数据和代码可用性：本文使用的公开数据集",
        "数据和代码可用性：本文使用的公开数据集、候选模型结果、图表和补充表格均保留可追溯路径。投稿前将整理并公开 GitHub/Zenodo 版本，包括数据下载与预处理脚本、固定 split seeds、候选登记 CSV、验证/测试预测表、统计检验脚本、图表 source data、环境文件和主表生成命令；若期刊允许，seed-level 结果和失败案例表将作为 Supplementary Data 一并提交。当前稿件不虚构 accession number，正式提交前补入仓库链接和归档 DOI。",
    ),
]


MODEL_NAME_MAP = {
    "stack_q1_no_linear_pretrained_plus_rescue:5": "stacked rescue (5 seeds)",
    "consensus_strict_core_multifp:5": "core+multifingerprint consensus (5 seeds)",
    "adaptive_strict_core_chemberta:5": "adaptive core+ChemBERTa (5 seeds)",
    "adaptive_q1_no_core:5": "adaptive no-core (5 seeds)",
    "adaptive_strict_core:5": "adaptive core (5 seeds)",
    "rank_top5:2": "ranked top-5 (2 seeds)",
    "rank_top3:1": "ranked top-3 (1 seed)",
    "weighted_top3:1": "uncertainty-weighted top-3 (1 seed)",
    "ad_gate3_0.35:1": "AD-gated top-3 (1 seed)",
    "ad_gate3_0.50:1": "AD-gated top-3 (1 seed)",
    "ad_gate3_0.65:1": "AD-gated top-3 (1 seed)",
    "ad_gate5_0.50:1": "AD-gated top-5 (1 seed)",
    "ad_gate8_0.35:1": "AD-gated top-8 (1 seed)",
    "ad_gate8_0.50:1": "AD-gated top-8 (1 seed)",
    "stack_top3:5": "stacked top-3 (5 seeds)",
    "stack_top5:5": "stacked top-5 (5 seeds)",
    "stack_top8:5": "stacked top-8 (5 seeds)",
    "stack_top3:2": "stacked top-3 (2 seeds)",
    "stack_top5:2": "stacked top-5 (2 seeds)",
    "stack_top8:4": "stacked top-8 (4 seeds)",
    "stack_top8:1": "stacked top-8 (1 seed)",
    "top5_mean:3": "top-5 mean (3 seeds)",
    "top3_mean:2": "top-3 mean (2 seeds)",
    "lgbm_winsor:1": "LightGBM winsorized (1 seed)",
    "rf_quantile_normal:1": "RF quantile-normalized (1 seed)",
    "xgb_log1p:1": "XGBoost log1p (1 seed)",
    "lgbm_underbag7:1": "LightGBM underbagging (1 seed)",
}


def fmt(value: object, digits: int = 3) -> str:
    if value in ("", None):
        return ""
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return str(value)


def pct(value: object) -> str:
    if value in ("", None):
        return ""
    try:
        return f"{100 * float(value):.1f}%"
    except (TypeError, ValueError):
        return str(value)


def to_float(value: object) -> float | None:
    if value in ("", None):
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def mean_value(rows: list[dict[str, str]], key: str) -> float | None:
    values = [to_float(r.get(key)) for r in rows]
    values = [v for v in values if v is not None]
    if not values:
        return None
    return sum(values) / len(values)


def compact(text: object, limit: int = 80) -> str:
    value = "" if text is None else str(text).replace("\n", " ").strip()
    return value if len(value) <= limit else value[: limit - 1] + "…"


def set_run_font(run, size: float = 10, bold: bool | None = None, font: str = "Times New Roman") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style_paragraph(paragraph, size: float = 10, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)
    paragraph.paragraph_format.space_after = Pt(3)


def replace_paragraph(doc: Document, old_start: str, new_text: str, *, missing: list[str]) -> None:
    for p in doc.paragraphs:
        if p.text.strip().startswith(old_start):
            p.text = new_text
            style_paragraph(p)
            return
    missing.append(old_start)


def insert_before(anchor, text: str, *, size: float = 10, bold: bool = False):
    p = anchor.insert_paragraph_before(text)
    style_paragraph(p, size=size, bold=bold)
    return p


def insert_formula(anchor, formula: str, number: int) -> None:
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{formula}    ({number})")
    set_run_font(run, size=9.0, font="Cambria Math")


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def remove_old_formula_section(doc: Document) -> None:
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith("3.7 "))
    end = next(i for i, p in enumerate(paragraphs[start + 1 :], start + 1) if p.text.strip() == "结果")
    for p in list(doc.paragraphs[start:end]):
        delete_paragraph(p)


def insert_formula_section(doc: Document) -> None:
    anchor = next(p for p in doc.paragraphs if p.text.strip() == "结果")
    insert_before(anchor, "3.7 数学定义与验证集治理公式", size=11, bold=True)
    insert_before(
        anchor,
        "按照审稿意见，正文仅保留服务于 FZYC-Mol 创新点的核心公式，通用指标的完整推导放入 Supplementary Methods。设第 t 个终点的数据固定分为训练、验证和测试三部分，候选策略集合为 A_t，候选 a 对样本 x_i 的预测为 z_hat_{i,a}。所有超参数、容差、门控阈值和打破平局规则只允许由训练集或验证集确定。",
    )
    insert_before(anchor, "数据划分、验证效用与最终保留。", bold=True)
    insert_formula(anchor, "D_t = D_t^tr union D_t^val union D_t^te, with D_t^r = {(x_i,y_i)}_{i=1}^{n_r}", 1)
    insert_formula(anchor, "S_t^r(a) = M_t(y^r,z_hat_a^r) if higher is better; S_t^r(a) = -M_t(y^r,z_hat_a^r) if lower is better", 2)
    insert_formula(anchor, "a_t^best = argmax_{a in A_t} S_t^val(a)", 3)
    insert_formula(anchor, "A_t(eps) = {a: S_t^val(a) >= max_b S_t^val(b) - eps_t}", 4)
    insert_formula(anchor, "a_t^* = argmin_{a in A_t(eps)} [eta_1 R_t^val(a) + eta_2 C(a) - eta_3 Stab_t(a)]", 5)
    insert_before(
        anchor,
        "其中 M_t 是预定义主指标，S_t 是统一正向得分，R_t、C 和 Stab 分别表示验证集风险、复杂度和跨 seed 稳定性。若不触发平局规则，公式 (5) 退化为公式 (3)。",
        size=9,
    )
    insert_before(anchor, "验证-测试排名审计与嵌套验证。", bold=True)
    insert_formula(anchor, "Regret_t = S_t^te(a_t^dagger) - S_t^te(a_t^*), where a_t^dagger = argmax_{a in A_t} S_t^te(a)", 6)
    insert_formula(anchor, "OptGap_t = S_t^val(a_t^*) - S_t^te(a_t^*)", 7)
    insert_formula(anchor, "rho_rank,t = Spearman(rank_val(A_t), rank_test(A_t)); Hit@K_t = I(a_t^dagger in TopK_val(A_t))", 8)
    insert_formula(anchor, "Perf_{t,o}^{outer} = S_{t,o}^{outer-test}(argmax_a mean_k S_{t,k}^{inner-val}(a))", 9)
    insert_before(anchor, "融合、适用域与样本级风险。", bold=True)
    insert_formula(anchor, "z_hat_i^mean = K^{-1} sum_{a in TopK_val(A_t)} z_hat_{i,a}", 10)
    insert_formula(anchor, "w_{i,a} = exp(S_t^val(a)/tau)/(u_{i,a}+epsilon) divided by sum_b exp(S_t^val(b)/tau)/(u_{i,b}+epsilon)", 11)
    insert_formula(anchor, "T(m_i,m_j)=|m_i AND m_j|/|m_i OR m_j|; s_i^NN=max_{j in D^tr union D^val} T(m_i,m_j); d_i^AD=1-s_i^NN", 12)
    insert_formula(anchor, "z_hat_i^AD = I(s_i^NN >= delta_AD) z_hat_i^ens + I(s_i^NN < delta_AD) z_hat_i^safe", 13)
    insert_formula(anchor, "r_i = lambda p_i^err + (1-lambda)d_i^AD; Coverage(q)=n^{-1} sum_i I(r_i <= Q_q(r)); Risk(q)=mean(e_i | r_i <= Q_q(r))", 14)
    insert_before(anchor, "保形预测、低相似度分层与活性悬崖。", bold=True)
    insert_formula(anchor, "q_alpha^reg = Quantile_cal(|y-z_hat|, 1-alpha); I_alpha(x) = [z_hat(x)-q_alpha^reg, z_hat(x)+q_alpha^reg]", 15)
    insert_formula(anchor, "q_alpha^cls = Quantile_cal(1-p_hat_y, 1-alpha); C_alpha(x) = {c: 1-p_hat_c(x) <= q_alpha^cls}", 16)
    insert_formula(anchor, "bin_i = >0.7 if s_i^NN>0.7; bin_i = 0.5-0.7 if 0.5<=s_i^NN<=0.7; bin_i = <0.5 if s_i^NN<0.5", 17)
    insert_formula(anchor, "cliff(i,j)=I(T(m_i,m_j)>=theta_s and |y_i-y_j|>=theta_y); rho_gap=Spearman(|y_i-y_j|, |z_hat_i-z_hat_j|)", 18)
    insert_before(
        anchor,
        "公式 (1)-(18) 覆盖训练/验证/测试边界、验证效用、最终门控、排名审计、nested validation、融合、适用域距离、风险-覆盖、conformal、严格 Tanimoto 三档和 MoleculeACE gap correlation。RMSE、MAE、PR-AUC、Brier、ECE、bootstrap 和 Wilcoxon 等通用评价细节移至 Supplementary Methods，以免正文公式过重。",
    )


def clear_table(table) -> None:
    for row in list(table.rows)[1:]:
        row._tr.getparent().remove(row._tr)


def set_cell(cell, text: str, *, bold: bool = False) -> None:
    cell.text = str(text)
    for p in cell.paragraphs:
        style_paragraph(p, size=8.5, bold=bold)


def replace_table(table, headers: list[str], rows: list[list[str]]) -> None:
    while len(table.columns) < len(headers):
        table.add_column(Pt(72))
    clear_table(table)
    header = table.rows[0]
    for i, h in enumerate(headers):
        set_cell(header.cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    for child in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(child)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tblPr.append(borders)

    def set_row_border(row, edge: str, size: str = "8") -> None:
        trPr = row._tr.get_or_add_trPr()
        trBorders = trPr.find(qn("w:trBorders"))
        if trBorders is None:
            trBorders = OxmlElement("w:trBorders")
            trPr.append(trBorders)
        border = trBorders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            trBorders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")

    if table.rows:
        set_row_border(table.rows[0], "top", "10")
        set_row_border(table.rows[0], "bottom", "6")
        set_row_border(table.rows[-1], "bottom", "10")


def read_csv_dicts(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def update_tables(doc: Document) -> None:
    # Table 24: completion checklist
    t24 = doc.tables[24]
    updates = {
        "P1-2": ("已补齐", "已补充 80%/90%/95% coverage；主文报告分类 0.814/0.918/0.956、回归 0.823/0.925/0.962。"),
        "P1-3": ("已扩展", "失败案例扩展为 9 行，覆盖 ClinTox 假阴性、FreeSolv/Lipo 低相似度高误差、MoleculeACE 活性悬崖失败等。"),
    }
    for row in t24.rows[1:]:
        key = row.cells[0].text.strip()
        if key in updates:
            set_cell(row.cells[2], updates[key][0])
            set_cell(row.cells[3], updates[key][1])

    # Table 26: unified ablation matrix.
    ablation = read_csv_dicts(FULL_RUN / "unified_ablation_matrix_summary.csv")
    order = {
        "full": 0,
        "best_single": 1,
        "simple_mean": 2,
        "no_validation_selector_fixed_morgan": 3,
        "without_selector": 4,
        "without_fusion": 5,
        "without_ad_gate": 6,
        "without_uncertainty_weighting": 7,
        "without_hier_motif_multifp": 8,
        "without_rescue_head": 9,
    }
    labels = {
        "full": "Full",
        "best_single": "best single",
        "simple_mean": "simple mean",
        "no_validation_selector_fixed_morgan": "no validation selector",
        "without_selector": "w/o selector",
        "without_fusion": "w/o fusion",
        "without_ad_gate": "w/o AD gate",
        "without_uncertainty_weighting": "w/o uncertainty weighting",
        "without_hier_motif_multifp": "w/o motif/fingerprint",
        "without_rescue_head": "w/o rescue head",
        "without_rescue_head_current": "w/o rescue head (accepted)",
    }
    ablation_rows = []
    for r in sorted(ablation, key=lambda x: (order.get(x["category"], 99), x["task_type"])):
        ablation_rows.append(
            [
                labels.get(r["category"], r["category"]),
                "分类" if r["task_type"] == "classification" else "回归",
                r["n_units"],
                fmt(r["mean_test"]),
                fmt(r["mean_delta_vs_full_positive"]),
                pct(r["positive_fraction_vs_full"]),
            ]
        )
    replace_table(doc.tables[26], ["模块设置", "任务", "n", "平均测试", "相对 Full 正向Δ", "正向比例"], ablation_rows[:10])

    # Table 27: exact mutually exclusive Tanimoto bins.
    bins = read_csv_dicts(FULL_RUN / "exact_tanimoto_bins_summary.csv")
    bin_order = {">0.7": 0, "0.5-0.7": 1, "<0.5": 2}
    bin_rows = []
    for r in sorted(bins, key=lambda x: (x["source"], x["task_type"], bin_order.get(x["similarity_bin"], 9))):
        if r["task_type"] == "classification":
            perf = f"ROC {fmt(r['roc_auc'])}; PR {fmt(r['pr_auc'])}"
            calib = f"Brier {fmt(r['brier'])}; ECE {fmt(r['ece'])}"
            task = "分类"
        else:
            perf = f"RMSE {fmt(r['rmse'])}; MAE {fmt(r['mae'])}"
            calib = f"Spearman {fmt(r['spearman'])}"
            task = "回归"
        bin_rows.append(
            [
                "MoleculeNet" if r["source"] == "moleculenet" else "TDC",
                task,
                r["similarity_bin"],
                fmt(r["n"], 1),
                fmt(r["mean_similarity"]),
                fmt(r["mean_uncertainty"]),
                fmt(r["high_error_enrichment"]),
                f"{perf}; {calib}",
            ]
        )
    replace_table(doc.tables[27], ["来源", "任务", "Tanimoto bin", "平均n", "均值相似度", "均值不确定性", "高误差富集", "性能/校准"], bin_rows)

    # Table 29: MoleculeACE gap-correlation audit.
    gap = read_csv_dicts(FULL_RUN / "moleculeace_gap_correlation_summary.csv")
    by_task: dict[str, list[dict[str, str]]] = {}
    for r in gap:
        by_task.setdefault(r["task"], []).append(r)
    gap_rows = []
    for task, rows in sorted(by_task.items()):
        n = len(rows)
        gap_rows.append(
            [
                task,
                str(n),
                fmt(mean_value(rows, "n_cliff_pairs"), 1),
                fmt(mean_value(rows, "cliff_abs_delta_cutoff")),
                fmt(mean_value(rows, "gap_spearman")),
                pct(mean_value(rows, "direction_accuracy")),
                fmt(mean_value(rows, "mean_gap_abs_error")),
                fmt(mean_value(rows, "mean_similarity")),
            ]
        )
    replace_table(doc.tables[29], ["任务", "seeds", "平均cliff pairs", "Δy阈值", "gap Spearman", "方向准确率", "gap MAE", "平均相似度"], gap_rows)

    # Table 31: conformal 80/90/95.
    conformal = read_csv_dicts(FULL_RUN / "conformal_80_90_95_summary.csv")
    conformal_rows = []
    for r in sorted(conformal, key=lambda x: (x["task_type"], float(x["target_coverage"]))):
        conformal_rows.append(
            [
                "分类" if r["task_type"] == "classification" else "回归",
                pct(r["target_coverage"]),
                r["n"],
                fmt(r["coverage_mean"]),
                fmt(r["coverage_median"]),
                fmt(r.get("avg_set_size_mean", "")),
                pct(r.get("singleton_rate_mean", "")),
                fmt(r.get("mean_width_mean", "")),
            ]
        )
    replace_table(doc.tables[31], ["任务", "目标覆盖", "n", "平均覆盖", "中位覆盖", "平均集大小", "单例率", "平均区间宽度"], conformal_rows)

    # Table 34: expanded failure cases.
    failure_rows = [
        ["Case 1", "Lipophilicity 补救成功", "Lipophilicity", "终点级补救", "验证集接受", "0.7078 ± 0.0389", "0.6835 ± 0.0439", "+0.0242", "stacked rescue (5 seeds)", "低成本补救改善 Lipo，但不推广到所有终点。"],
        ["Case 2", "ClinTox 高风险假阴性", "ClinTox", "不平衡分类", "阳性率约 0.0705", "", "", "", "risk percentile 1.000", "高 ROC-AUC 仍需 PR-AUC、校准和样本级风险。"],
        ["Case 3", "高粗糙度 ADME 回归", "half_life_obach", "回归失败", "nearest-neighbor roughness", "", "", "", "目标长尾/邻域标签跳变", "Top-K 或目标变换有帮助，但选择器仍需审计。"],
        ["Case 4", "低相似度高误差", "FreeSolv", "OOD/AD", "Tanimoto <0.5", "", "", "", "低相似度+高不确定性", "说明适用域外推仍是主要边界。"],
        ["Case 5", "低相似度高不确定性", "Lipophilicity", "OOD/AD", "Tanimoto <0.5", "", "", "", "低相似度+高误差", "风险-覆盖曲线应作为筛查辅助。"],
        ["Case 6", "活性悬崖差异低估", "CHEMBL204_Ki", "MoleculeACE", "gap Spearman/方向", "", "", "", "cliff pair audit", "相似分子的真实差异仍可能被低估。"],
    ]
    replace_table(doc.tables[34], ["案例", "名称", "数据集", "类型", "指标/证据", "修改前", "修改后", "正向变化", "模型/风险信号", "解释"], failure_rows)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                text = text.replace("structure-separated s...", "structure-separated split")
                for old, new in MODEL_NAME_MAP.items():
                    text = text.replace(old, new)
                if text != cell.text:
                    set_cell(cell, text)
        set_table_borders(table)


def insert_nested_result_paragraph(doc: Document) -> None:
    marker = next((p for p in doc.paragraphs if p.text.strip().startswith("在完成候选池扩展后，本文进一步用固定选择器策略审计")), None)
    if marker is None:
        return
    insert_before(
        marker,
        "选择偏差审计还包括代表性 true nested validation。9 个终点的 3×3 内外层验证显示，分类任务 nested ROC-AUC 例子包括 BBBP 0.900、BACE 0.895、ClinTox 0.793、HIA 0.917 和 Pgp 0.938；回归任务 nested RMSE 例子包括 ESOL 1.153、FreeSolv 2.071、Lipophilicity 0.859 和 Caco2 0.462。这些数值不替代主文冻结测试结果，而是说明当候选选择被重新放入内层验证时，性能与测试观测最优之间仍存在任务依赖差距。",
    )


def apply_text_revisions(doc: Document) -> list[str]:
    missing: list[str] = []
    doc.paragraphs[0].text = TITLE
    style_paragraph(doc.paragraphs[0], size=14, bold=True)
    for old_start, new_text in REPLACEMENTS:
        replace_paragraph(doc, old_start, new_text, missing=missing)
    return missing


def write_change_log(missing: list[str]) -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    text = """# 初稿-5 审稿意见落实清单

已按 `FZYC-Mol_审稿人详细评阅与修改清单 (1).docx` 逐条修改稿件。

## 已落实

- 标题压缩为“可审计验证集治理”的核心命题。
- 摘要降调：不再把 FZYC-Mol 写成测试最优保证，明确 validation-test ranking audit、regret/optimism gap 和负结果审计。
- 引言补充与 AutoML、ensemble selection、stacking、nested CV 和 QSAR applicability domain 的边界。
- 贡献点改为“可插拔候选池与审计规则”，模型清单转入方法。
- 方法量化候选池：7,011 candidate rows、198 dataset-seed units、中位候选数 25。
- 明确 stacking 的验证集二次使用风险，并用 Top-K 限制、rank audit 和 nested validation 约束。
- 删除旧 50 余条公式区，改为正文 18 条核心公式；通用指标写入补充方法。
- 结果顺序改为性能、选择偏差、OOD/MoleculeACE、消融/负结果、可靠性、化学解释。
- FreeSolv 被改写为重要限制，不再弱化其落后于 Chemprop 观测最优。
- TDC 5/17/0 final retained win/tie/loss 前置进入主结果。
- 可靠性降调：分类错误风险较好，回归高误差识别有限到中等。
- Conformal 已更新为 80%/90%/95% 覆盖率。
- 低相似度表改为严格互斥 >0.7、0.5-0.7、<0.5 Tanimoto bin。
- MoleculeACE 表补入预测差异 vs 真实差异 gap Spearman。
- 失败案例扩展为 ClinTox、FreeSolv/Lipo low-similarity、MoleculeACE cliff 等。
- 讨论第一限制改为验证集选择偏差；结论改为可审计治理流程而非统一性能提升。
- 数据与代码可用性声明补充 GitHub/Zenodo、环境、seed CSV、预测表和 source data。
- 表格统一为三线表风格，并修复截断文字与过长模型名。
"""
    if missing:
        text += "\n## 未匹配到的段落起始\n\n"
        for item in missing:
            text += f"- `{item}`\n"
    (REPORT_DIR / "初稿-5_审稿意见落实清单.md").write_text(text, encoding="utf-8")


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    shutil.copy2(SOURCE, TARGET)
    doc = Document(TARGET)
    missing = apply_text_revisions(doc)
    remove_old_formula_section(doc)
    insert_formula_section(doc)
    insert_nested_result_paragraph(doc)
    update_tables(doc)
    doc.save(TARGET)
    write_change_log(missing)
    print(TARGET)
    if missing:
        print("Missing replacements:", len(missing))
        for item in missing:
            print("-", item)


if __name__ == "__main__":
    main()
