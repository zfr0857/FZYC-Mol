from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop" / "修改"
IN_DOCX = DESKTOP / "初稿-16_补充实验整理版.docx"
OUT_DOCX = DESKTOP / "初稿-17_补齐LinPept版.docx"
REPORT_DOCX = DESKTOP / "LinPept数据补齐与bRo5实验更新报告_20260611.docx"
RUN_DIR = ROOT / "reports" / "full_missing_experiment_run_20260611"
LINPEPT_DIR = ROOT / "reports" / "bro5_linpept_20260611"


def mean_sd(series: pd.Series, digits: int = 3) -> str:
    values = pd.to_numeric(series, errors="coerce").dropna()
    if values.empty:
        return ""
    if len(values) == 1:
        return f"{values.iloc[0]:.{digits}f}"
    return f"{values.mean():.{digits}f} ± {values.std(ddof=1):.{digits}f}"


def load_linpept() -> tuple[pd.DataFrame, pd.DataFrame, dict[str, str]]:
    selected = pd.read_csv(LINPEPT_DIR / "validation_selected_results.csv")
    status = pd.read_csv(LINPEPT_DIR / "linpept_data_status.csv")
    selected["split_order"] = selected["split"].map({"random": 0, "scaffold": 1, "perimeter": 2}).fillna(99)
    rows = []
    for (dataset, split), group in selected.sort_values(["dataset", "split_order", "seed"]).groupby(["dataset", "split"], sort=False):
        rows.append(
            {
                "dataset": dataset,
                "split": split,
                "n_seed": len(group),
                "test_ROC_AUC": mean_sd(group["test_roc_auc"]),
                "test_PR_AUC": mean_sd(group["test_pr_auc"]),
                "test_Brier": mean_sd(group["test_brier"]),
                "balanced_accuracy": mean_sd(group["test_balanced_accuracy"]),
                "MCC": mean_sd(group["test_mcc"]),
                "recall_at_precision_0.80": mean_sd(group["test_recall_at_p80"]),
                "recall_at_precision_0.90": mean_sd(group["test_recall_at_p90"]),
                "mean_test_nn_Tanimoto": mean_sd(group["mean_test_nn_tanimoto"]),
                "selected_models": "; ".join(f"{k}:{v}" for k, v in group["model"].value_counts().sort_index().items()),
            }
        )
    summary = pd.DataFrame(rows)
    summary["dataset_order"] = summary["dataset"].map({"linpept_cellpen": 0, "linpept_nonfouling": 1}).fillna(99)
    summary["split_order"] = summary["split"].map({"random": 0, "scaffold": 1, "perimeter": 2}).fillna(99)
    summary = summary.sort_values(["dataset_order", "split_order"]).drop(columns=["dataset_order", "split_order"]).reset_index(drop=True)
    summary.to_csv(RUN_DIR / "linpept_compact_summary_20260611.csv", index=False, encoding="utf-8-sig")
    status.to_csv(RUN_DIR / "linpept_data_status_20260611.csv", index=False, encoding="utf-8-sig")

    def split_line(dataset: str) -> str:
        sub = summary[summary["dataset"].eq(dataset)].set_index("split")
        order = [s for s in ["random", "scaffold", "perimeter"] if s in sub.index]
        parts = [
            f"{split} ROC-AUC {sub.loc[split, 'test_ROC_AUC']}, PR-AUC {sub.loc[split, 'test_PR_AUC']}"
            for split in order
        ]
        n_row = status[status["dataset"].eq(dataset)].iloc[0]
        label = "CellPen" if dataset == "linpept_cellpen" else "NonFouling"
        return f"LinPept {label} n={int(n_row['n'])}, positive rate {float(n_row['positive_rate']):.3f}; " + "; ".join(parts) + "."

    lines = {
        "cellpen": split_line("linpept_cellpen"),
        "nonfouling": split_line("linpept_nonfouling"),
    }
    lines["combined"] = lines["cellpen"] + " " + lines["nonfouling"]
    return summary, status, lines


def load_cycpept_line() -> str:
    path = RUN_DIR / "experiment_completion_matrix_20260611.csv"
    if not path.exists():
        return "CycPept-PAMPA public stress test was completed in the previous run."
    matrix = pd.read_csv(path)
    hit = matrix[matrix["module"].eq("bRo5 CycPept-PAMPA")]
    if hit.empty:
        return "CycPept-PAMPA public stress test was completed in the previous run."
    return str(hit.iloc[0]["key_result"])


def update_completion_matrix(lines: dict[str, str]) -> pd.DataFrame:
    path = RUN_DIR / "experiment_completion_matrix_20260611.csv"
    matrix = pd.read_csv(path)
    linpept_result = lines["combined"]
    for i, row in matrix.iterrows():
        module = str(row["module"])
        if module == "bRo5 CycPept-PAMPA":
            matrix.loc[i, "status"] = "completed"
            matrix.loc[i, "manuscript_use"] = "Report as completed CycPept-PAMPA regression bRo5 stress test; combine only at the narrative level with LinPept classification benchmarks."
        if module == "LinPept NonFouling / LinPept CellPen":
            matrix.loc[i, "status"] = "completed"
            matrix.loc[i, "scope"] = "Benchmark-ADMET-2025 origin_data; standardized smiles/y files; random/scaffold/perimeter splits; 3 seeds"
            matrix.loc[i, "key_result"] = linpept_result
            matrix.loc[i, "evidence"] = "reports/bro5_linpept_20260611/"
            matrix.loc[i, "manuscript_use"] = "Report as completed bRo5 classification stress tests; cite Benchmark-ADMET-2025 and the original LinPept dataset papers."
    out = RUN_DIR / "experiment_completion_matrix_with_linpept_20260611.csv"
    matrix.to_csv(out, index=False, encoding="utf-8-sig")
    return matrix


def set_border(element, edge: str, val: str, size: str = "8", color: str = "000000") -> None:
    tag = "w:tblBorders" if element.tag.endswith("tblPr") else "w:tcBorders"
    borders = element.find(qn(tag))
    if borders is None:
        borders = OxmlElement(tag)
        element.append(borders)
    edge_el = borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        borders.append(edge_el)
    edge_el.set(qn("w:val"), val)
    if val != "nil":
        edge_el.set(qn("w:sz"), size)
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)


def apply_three_line_table(table) -> None:
    tbl_pr = table._tbl.tblPr
    for edge in ("left", "right", "insideH", "insideV"):
        set_border(tbl_pr, edge, "nil")
    set_border(tbl_pr, "top", "single", "12")
    set_border(tbl_pr, "bottom", "single", "12")
    if table.rows:
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for edge in ("left", "right", "top"):
                set_border(tc_pr, edge, "nil")
            set_border(tc_pr, "bottom", "single", "8")
        for row in table.rows[1:]:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                for edge in ("left", "right", "top", "bottom"):
                    set_border(tc_pr, edge, "nil")


def add_docx_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    apply_three_line_table(table)


def set_paragraph_text(paragraph: Paragraph, text: str, size: float = 10.5) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.size = Pt(size)


def iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_known_text(doc: Document, lines: dict[str, str]) -> int:
    combined_sentence = (
        "bRo5 模块已新增 CycPept-PAMPA、LinPept CellPen 和 LinPept NonFouling 三项公开压力测试。"
        "其中 CycPept-PAMPA 作为回归型渗透性任务报告 RMSE/MAE，LinPept 两项作为分类型肽性质任务报告 ROC-AUC、PR-AUC、校准和适用域分层；"
        "三者均以冻结测试集作一次性评估，不再作为缺失数据处理。"
    )
    replacements = {
        "bRo5 模块已新增 CycPept-PAMPA 公开子集压力测试，但 LinPept NonFouling 和 LinPept CellPen 仍为 missing_data，因此该模块不能表述为完整 bRo5 性能评估。": combined_sentence,
        "bRo5 当前仅能提供数据状态审计，不能作为性能结论": "bRo5 当前已提供 CycPept-PAMPA 回归压力测试和 LinPept 分类压力测试，但仍需更多盲测式外部验证",
        "LinPept 两项、官方盲测提交和大规模 full fine-tuning 尚不具备完整可核验证据。": "官方盲测提交和大规模 full fine-tuning 尚不具备完整可核验证据；LinPept 两项已补齐为公开数据压力测试。",
        "bRo5 数据状态审计": "bRo5 公开压力测试",
        "bRo5 数据缺失审计": "bRo5 公开压力测试",
        "bRo5 当前新增 CycPept-PAMPA 公开子集结果；LinPept NonFouling 和 LinPept CellPen 仍保留为 missing_data，因此不作为完整 bRo5 全量结论。": combined_sentence,
        "bRo5 相关条目当前以 missing_data 状态报告，不作为已公开或已完成性能数据存档。": "bRo5 相关条目已补齐 CycPept-PAMPA、LinPept CellPen 和 LinPept NonFouling 的可追溯公开数据实验；完整 seed-level 预测、字段字典、脚本和图表源数据保存在 reports/bro5_cycpept_pampa_20260611/ 与 reports/bro5_linpept_20260611/。",
        "bRo5 数据状态": "bRo5 公开压力测试",
    }
    changed = 0
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text
        new = text
        for old, replacement in replacements.items():
            new = new.replace(old, replacement)
        if new != text:
            set_paragraph_text(paragraph, new)
            changed += 1

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.startswith("bRo5 部分在当前版本中仅作为数据状态审计"):
            set_paragraph_text(
                paragraph,
                "bRo5 部分已由数据状态审计更新为可追溯的公开压力测试。CycPept-PAMPA 继续作为规则五以外渗透性回归任务报告 RMSE、MAE、Spearman、适用域覆盖和风险富集；LinPept CellPen 与 LinPept NonFouling 已从 Benchmark-ADMET-2025 公开仓库下载原始 CSV，核验 smiles 与终点标签字段后统一标准化为 smiles/y 格式。两项 LinPept 任务均按 random、scaffold 和 perimeter split 以及 3 个随机种子运行，候选模型仅由验证集 ROC-AUC 冻结选择，测试集用于最终一次性评估。因此，正文可以报告 bRo5 公开压力测试结果，但应区分 CycPept-PAMPA 的回归指标与 LinPept 的分类指标，避免混用 RMSE 与 ROC-AUC。",
            )
            changed += 1
        elif text.startswith("根据补充实验清单，本文进一步补齐了能够在当前数据和依赖条件下运行的关键模块。"):
            set_paragraph_text(
                paragraph,
                "根据补充实验清单，本文进一步补齐了能够在当前数据和依赖条件下运行的关键模块。bRo5 公开压力测试现包括 CycPept-PAMPA 回归任务和 LinPept CellPen/NonFouling 分类任务。CycPept-PAMPA 覆盖 random、scaffold、perimeter 和 time split；LinPept 两项来自 Benchmark-ADMET-2025 origin_data，标准化后分别包含 CellPen 1,960 条样本和 NonFouling 7,239 条样本。验证集选择后，"
                + lines["cellpen"]
                + " "
                + lines["nonfouling"]
                + " 轻量分子大模型适配器、MoleculeACE 可用任务子集和统一消融矩阵仍按各自预定义边界解释；所有新增结果均保留 seed-level 输出和负结果记录，不作为未运行模块的替代证据。",
            )
            changed += 1
        elif "bRo5 审计已新增 CycPept-PAMPA 公开子集实验" in text and "LinPept NonFouling 和 LinPept CellPen 仍为 missing_data" in text:
            set_paragraph_text(
                paragraph,
                text.replace(
                    "bRo5 审计已新增 CycPept-PAMPA 公开子集实验，覆盖 random、scaffold、perimeter 和 time split；但 LinPept NonFouling 和 LinPept CellPen 仍为 missing_data，因此不能表述为完整 bRo5 全量结论。",
                    "bRo5 审计已新增 CycPept-PAMPA、LinPept CellPen 和 LinPept NonFouling 公开压力测试。CycPept-PAMPA 覆盖 random、scaffold、perimeter 和 time split；LinPept 两项覆盖 random、scaffold 和 perimeter split，并以分类指标报告。",
                ),
            )
            changed += 1
    return changed


def update_tables(doc: Document, lines: dict[str, str]) -> int:
    changed = 0
    for table in doc.tables:
        apply_three_line_table(table)
        for row in table.rows:
            cells = row.cells
            first = cells[0].text if cells else ""
            if first == "bRo5 CycPept-PAMPA" and len(cells) >= 5:
                cells[1].text = "completed"
                cells[4].text = "Report as completed CycPept-PAMPA regression bRo5 stress test; combine with LinPept only at the narrative level."
                changed += 1
            elif first == "LinPept NonFouling / LinPept CellPen" and len(cells) >= 5:
                cells[1].text = "completed"
                cells[2].text = lines["combined"]
                cells[3].text = "reports/bro5_linpept_20260611/"
                cells[4].text = "Report as completed bRo5 classification stress tests; cite Benchmark-ADMET-2025 and original LinPept dataset papers."
                changed += 1
            elif first == "bRo5 化学空间" and len(cells) >= 5:
                cells[1].text = "公开压力测试/已补齐"
                cells[4].text = "数据可用性、样本量、ROC-AUC/PR-AUC 或 RMSE/MAE、适用域覆盖和风险富集；回归与分类指标分开报告"
                changed += 1
    return changed


def insert_linpept_table(doc: Document, summary: pd.DataFrame) -> None:
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("根据补充实验清单，本文进一步补齐了能够在当前数据和依赖条件下运行的关键模块。"):
            anchor = paragraph
            break
    if anchor is None:
        return
    rows = []
    label_map = {
        "linpept_cellpen": "LinPept CellPen",
        "linpept_nonfouling": "LinPept NonFouling",
    }
    for _, row in summary.iterrows():
        rows.append(
            [
                label_map.get(row["dataset"], row["dataset"]),
                row["split"],
                row["n_seed"],
                row["test_ROC_AUC"],
                row["test_PR_AUC"],
                row["test_Brier"],
                row["recall_at_precision_0.80"],
            ]
        )
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["dataset", "split", "n_seed", "ROC-AUC", "PR-AUC", "Brier", "recall@P0.80"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    apply_three_line_table(table)
    anchor._p.addnext(table._tbl)


def write_markdown_report(summary: pd.DataFrame, status: pd.DataFrame, matrix: pd.DataFrame, lines: dict[str, str]) -> None:
    lines_md = [
        "# LinPept data fill and bRo5 update, 2026-06-11",
        "",
        "LinPept CellPen and LinPept NonFouling were downloaded from the public Benchmark-ADMET-2025 repository, standardized to `smiles,y`, and run as bRo5 classification stress tests.",
        "",
        "## Data status",
        "",
        status.to_markdown(index=False),
        "",
        "## Validation-selected test performance",
        "",
        summary.to_markdown(index=False),
        "",
        "## Manuscript-safe summary",
        "",
        lines["combined"],
        "",
        "## Updated completion matrix",
        "",
        matrix.to_markdown(index=False),
        "",
    ]
    (RUN_DIR / "linpept_bro5_update_20260611.md").write_text("\n".join(lines_md), encoding="utf-8")


def write_docx_report(summary: pd.DataFrame, status: pd.DataFrame, lines: dict[str, str]) -> None:
    doc = Document()
    doc.add_heading("LinPept 数据补齐与 bRo5 实验更新报告", level=1)
    doc.add_paragraph("日期：2026-06-11")
    doc.add_paragraph(
        "本轮已从 Benchmark-ADMET-2025 公开仓库下载 LinPept CellPen 和 LinPept NonFouling 原始数据，核验 smiles 与标签字段后标准化为 smiles/y，并完成 random、scaffold 和 perimeter split 下的 3-seed 分类压力测试。"
    )
    doc.add_heading("数据状态", level=2)
    add_docx_table(
        doc,
        ["dataset", "status", "n", "positive_rate", "standardized_path", "source_file"],
        status[["dataset", "status", "n", "positive_rate", "path", "source_file"]].values.tolist(),
    )
    doc.add_heading("验证集冻结选择后的测试结果", level=2)
    add_docx_table(
        doc,
        ["dataset", "split", "n_seed", "ROC-AUC", "PR-AUC", "Brier", "MCC", "recall@P0.80", "selected_models"],
        summary[
            [
                "dataset",
                "split",
                "n_seed",
                "test_ROC_AUC",
                "test_PR_AUC",
                "test_Brier",
                "MCC",
                "recall_at_precision_0.80",
                "selected_models",
            ]
        ].values.tolist(),
    )
    doc.add_heading("可写入论文的边界表述", level=2)
    doc.add_paragraph(lines["combined"])
    doc.add_paragraph(
        "写作时应将 LinPept 两项称为公开数据压力测试，而不是盲测或外部未见实验；应将 CellPen/NonFouling 的分类指标与 CycPept-PAMPA 的回归指标分开报告。"
    )
    doc.save(REPORT_DOCX)
    shutil.copy2(REPORT_DOCX, RUN_DIR / REPORT_DOCX.name)


def update_manuscript(summary: pd.DataFrame, lines: dict[str, str]) -> tuple[int, int]:
    doc = Document(IN_DOCX)
    text_changes = replace_known_text(doc, lines)
    table_changes = update_tables(doc, lines)
    insert_linpept_table(doc, summary)
    doc.save(OUT_DOCX)
    shutil.copy2(OUT_DOCX, RUN_DIR / OUT_DOCX.name)
    return text_changes, table_changes


def main() -> None:
    summary, status, lines = load_linpept()
    matrix = update_completion_matrix(lines)
    write_markdown_report(summary, status, matrix, lines)
    write_docx_report(summary, status, lines)
    text_changes, table_changes = update_manuscript(summary, lines)
    print(f"Wrote {RUN_DIR / 'linpept_compact_summary_20260611.csv'}")
    print(f"Wrote {RUN_DIR / 'experiment_completion_matrix_with_linpept_20260611.csv'}")
    print(f"Wrote {RUN_DIR / 'linpept_bro5_update_20260611.md'}")
    print(f"Wrote {REPORT_DOCX}")
    print(f"Wrote {OUT_DOCX}")
    print(f"Manuscript text changes: {text_changes}; table changes: {table_changes}")


if __name__ == "__main__":
    main()
