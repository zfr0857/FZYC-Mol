from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SUPP = ROOT / "reports" / "supplement_experiment_revision_20260606"


def find_source_docx() -> Path:
    marker = "\u521d\u7a3f4"
    candidates = [p for p in (Path.home() / "Desktop").rglob("FZYC-Mol*.docx") if marker in p.name]
    if not candidates:
        raise FileNotFoundError("Cannot find FZYC-Mol draft 4 under Desktop.")
    candidates.sort(key=lambda p: ("~$" in p.name, len(str(p))))
    return candidates[0]


def output_path(src: Path) -> Path:
    return src.with_name(f"{src.stem}_\u8865\u5145\u5b9e\u9a8c\u4fee\u8ba2.docx")


def fmt(x, digits: int = 3) -> str:
    if x is None:
        return ""
    try:
        if pd.isna(x):
            return ""
        value = float(x)
    except Exception:
        return str(x)
    if abs(value) >= 100:
        return f"{value:.1f}"
    return f"{value:.{digits}f}"


def pct(x) -> str:
    try:
        if pd.isna(x):
            return ""
        return f"{100 * float(x):.1f}%"
    except Exception:
        return str(x)


def set_run_font(run, size: int = 10, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style_paragraph(p, size: int = 10, bold: bool = False) -> None:
    for run in p.runs:
        set_run_font(run, size=size, bold=bold)


def replace_paragraph(doc: Document, old_start: str, new_text: str) -> bool:
    for p in doc.paragraphs:
        if p.text.strip().startswith(old_start):
            p.text = new_text
            style_paragraph(p, 10)
            return True
    return False


def clear_borders(cell) -> None:
    set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = f"w:{edge}"
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def make_three_line(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            clear_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    border = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=border, bottom=border)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=border)


def fill_table(table, headers: list[str], rows: list[list[str]]) -> None:
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.cell(i, j).text = "" if value is None else str(value)
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, size=8, bold=(r == 0))
    make_three_line(table)


def insert_para(anchor, text: str, size: int = 10, bold: bool = False, style: str | None = None):
    p = anchor.insert_paragraph_before(text)
    if style:
        try:
            p.style = style
        except Exception:
            pass
    style_paragraph(p, size=size, bold=bold)
    return p


def insert_table(anchor, doc: Document, caption: str, headers: list[str], rows: list[list[str]], note: str | None = None) -> None:
    insert_para(anchor, caption, size=9, bold=True)
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    fill_table(table, headers, rows)
    anchor._p.addprevious(table._tbl)
    if note:
        insert_para(anchor, note, size=8)


def source_label(source: str) -> str:
    labels = {
        "overall": "\u603b\u4f53",
        "validation_selector": "\u539f\u9a8c\u8bc1\u96c6\u9009\u62e9\u5668",
        "validation_selector_expanded": "\u6269\u5c55\u9009\u62e9\u5668",
        "moleculenet_targeted_rebuild": "MoleculeNet \u5b9a\u5411\u91cd\u6784",
        "moleculenet_multimethod_fusion": "MoleculeNet \u591a\u65b9\u6cd5\u878d\u5408",
        "tdc_multimethod_fusion": "TDC \u591a\u65b9\u6cd5\u878d\u5408",
        "structure_full_selector": "\u7ed3\u6784\u5206\u79bb\u9009\u62e9\u5668",
        "three_d_roughness": "3D-lite/\u7c97\u7cd9\u5ea6",
        "strong_tabpfn_pilot": "\u8868\u683c\u57fa\u7840\u6a21\u578b\u5148\u5bfc",
    }
    return labels.get(source, source)


def comparison_label(name: str) -> str:
    labels = {
        "chemprop_only": "\u4ec5 Chemprop",
        "multifp_only": "\u4ec5\u591a\u6307\u7eb9",
        "pretrained_only": "\u4ec5\u51bb\u7ed3\u9884\u8bad\u7ec3\u8868\u5f81",
        "core_only": "\u4ec5\u6838\u5fc3\u5bb6\u65cf",
        "no_core": "\u53bb\u6838\u5fc3\u5bb6\u65cf",
        "no_multifp": "\u53bb\u591a\u6307\u7eb9",
        "no_chemprop": "\u53bb Chemprop",
        "no_pretrained": "\u53bb\u9884\u8bad\u7ec3\u8868\u5f81",
        "chemprop_baseline::chemprop_dmpnn_ens3": "\u5bf9 Chemprop D-MPNN",
        "strict_core_fast::dmpnn": "\u5bf9 D-MPNN core",
    }
    return labels.get(name, name)


def build_tables() -> dict[str, list[list[str]]]:
    tables: dict[str, list[list[str]]] = {}

    tracker = pd.read_csv(SUPP / "supplement_experiment_coverage_tracker.csv")
    tables["tracker"] = tracker.values.tolist()

    rank = pd.read_csv(SUPP / "maintext_table_validation_bias_extended.csv")
    tables["rank"] = [
        [
            source_label(r.source),
            fmt(r.n_dataset_seed_units, 0),
            fmt(r.median_spearman),
            pct(r.top1_match_rate),
            pct(r.test_top_in_valid_top3_rate),
            fmt(r.negative_rank_units, 0),
            fmt(r.median_optimism_gap_native),
            fmt(r.median_selected_test_regret_native),
        ]
        for r in rank.itertuples(index=False)
    ]

    ab = pd.read_csv(SUPP / "maintext_table_systematic_ablation.csv")
    tables["ablation"] = [
        [
            comparison_label(r.comparison),
            fmt(r.mean_positive_delta),
            f"{int(r.wins)}/{int(r.ties)}/{int(r.losses)}",
            fmt(r.median_bootstrap_p),
            fmt(r.median_wilcoxon_p),
            "\u5b8c\u6574\u9009\u62e9\u5668\u5bf9\u7167" if r.section == "family_ablation" else "\u5f3a\u57fa\u7ebf\u914d\u5bf9\u5bf9\u7167",
        ]
        for r in ab.itertuples(index=False)
    ]

    low = pd.read_csv(SUPP / "maintext_table_low_similarity_bins.csv")
    tables["low_similarity"] = [
        [
            "\u5206\u7c7b" if r.task_type == "classification" else "\u56de\u5f52",
            fmt(r.tanimoto_threshold, 1),
            fmt(r.mean_n, 1),
            fmt(r.mean_similarity),
            fmt(r.roc_auc_mean) if r.task_type == "classification" else "",
            fmt(r.pr_auc_mean) if r.task_type == "classification" else "",
            fmt(r.rmse_mean) if r.task_type == "regression" else "",
            fmt(r.mae_mean) if r.task_type == "regression" else "",
        ]
        for r in low.itertuples(index=False)
    ]

    shift = pd.read_csv(SUPP / "maintext_table_structure_shift_compact.csv").head(8)
    tables["structure_shift"] = [
        [
            r.dataset,
            "\u5206\u7c7b" if r.task_type == "classification" else "\u56de\u5f52",
            r.metric,
            fmt(r.random_value),
            fmt(r.scaffold_value),
            fmt(r.structure_value),
            fmt(r.scaffold_to_structure_drop),
        ]
        for r in shift.itertuples(index=False)
    ]

    ace_examples = pd.read_csv(SUPP / "maintext_table_moleculeace_cliff_examples.csv").head(6)
    tables["moleculeace"] = [
        [
            r.task,
            fmt(r.n_seed, 0),
            fmt(r.baseline_rmse_mean),
            fmt(r.cliff_rmse_mean),
            fmt(r.delta_rmse_positive_mean),
            fmt(r.baseline_cliff_rmse_mean),
            fmt(r.cliff_cliff_rmse_mean),
            fmt(r.delta_cliff_rmse_positive_mean),
        ]
        for r in ace_examples.itertuples(index=False)
    ]

    imb = pd.read_csv(SUPP / "maintext_table_imbalanced_metrics_compact.csv")
    tables["imbalanced"] = [
        [
            r.dataset,
            pct(r.positive_rate_or_query_positive_rate),
            fmt(r.roc_auc_mean),
            fmt(r.pr_auc_mean),
            fmt(r.brier_mean),
            fmt(r.ece_mean),
            fmt(r.ef1_mean),
            fmt(r.ef5_mean),
        ]
        for r in imb.itertuples(index=False)
    ]

    conf = pd.read_csv(SUPP / "maintext_table_conformal_coverage_compact.csv")
    tables["conformal"] = [
        [
            "\u5206\u7c7b" if r.task_type == "classification" else "\u56de\u5f52",
            f"{int((1 - float(r.alpha)) * 100)}%",
            fmt(r.n, 0),
            fmt(r.coverage_mean),
            fmt(r.coverage_median),
            fmt(r.avg_set_size_mean),
            fmt(r.singleton_rate_mean),
            fmt(r.mean_width_mean),
        ]
        for r in conf.itertuples(index=False)
    ]

    cases = pd.read_csv(SUPP / "maintext_table_failure_cases_compact.csv")
    tables["cases"] = [
        [
            r.case_id,
            r.dataset,
            r.case_type,
            r.primary_metric,
            r.before,
            r.after,
            fmt(r.delta_positive),
            r.interpretation[:130],
        ]
        for r in cases.itertuples(index=False)
    ]

    neg = pd.read_csv(SUPP / "maintext_table_negative_result_audit.csv")
    neg_pick = pd.concat([neg.head(8), neg[neg["module"].isin(["3d_lite", "roughness_weighted"])].head(6)])
    tables["negative"] = [
        [
            r.module,
            r.scope,
            pct(r.positive_rate) if not pd.isna(r.positive_rate) else "",
            fmt(r.mean_delta_vs_current),
            str(r.decision)[:90],
        ]
        for r in neg_pick.itertuples(index=False)
    ]
    return tables


def insert_supplement_section(doc: Document) -> None:
    anchor = next((p for p in doc.paragraphs if p.text.strip() == "4.6 \u53ef\u89e3\u91ca\u6027\u4e0e\u6848\u4f8b\u5206\u6790"), None)
    if anchor is None:
        raise RuntimeError("Cannot find section 4.6 anchor.")
    anchor.text = "4.7 \u53ef\u89e3\u91ca\u6027\u4e0e\u6848\u4f8b\u5206\u6790"
    style_paragraph(anchor, 11, bold=True)

    tables = build_tables()
    insert_para(anchor, "4.6 \u8865\u5145\u5b9e\u9a8c\u95ed\u73af\uff1a\u9a8c\u8bc1\u6cbb\u7406\u3001\u6d88\u878d\u3001\u4f4e\u76f8\u4f3c\u5ea6\u4e0e\u6d3b\u6027\u60ac\u5d16", size=11, bold=True, style="Heading 2")
    insert_para(
        anchor,
        "\u6839\u636e\u8865\u5145\u5b9e\u9a8c\u6e05\u5355\uff0c\u672c\u8f6e\u4fee\u8ba2\u5c06 P0/P1 \u8bc1\u636e\u4ece\u8865\u5145\u6750\u6599\u62c9\u56de\u4e3b\u6587\u903b\u8f91\uff1a\u5148\u5ba1\u8ba1\u9a8c\u8bc1\u96c6\u9009\u62e9\u662f\u5426\u5b58\u5728\u6392\u540d\u504f\u5dee\uff0c\u518d\u7528\u914d\u5bf9\u7edf\u8ba1\u548c\u7cfb\u7edf\u6d88\u878d\u754c\u5b9a\u54ea\u4e9b\u6a21\u5757\u771f\u6b63\u6709\u6548\uff0c\u6700\u540e\u5728\u4f4e\u76f8\u4f3c\u5ea6\u3001\u7ed3\u6784\u5206\u79bb\u3001MoleculeACE \u6d3b\u6027\u60ac\u5d16\u3001\u4e0d\u5e73\u8861\u5206\u7c7b\u548c\u5931\u8d25\u6848\u4f8b\u4e2d\u68c0\u67e5\u9002\u7528\u8fb9\u754c\u3002\u8fd9\u4e9b\u8865\u5145\u5b9e\u9a8c\u4e0d\u6539\u53d8\u6d4b\u8bd5\u96c6\u4e00\u6b21\u6027\u62a5\u544a\u539f\u5219\uff0c\u800c\u662f\u5c06\u201c\u4e3a\u4ec0\u4e48\u63a5\u53d7\u6216\u62d2\u7edd\u67d0\u4e2a\u5019\u9009\u201d\u5199\u6210\u53ef\u8ffd\u6eaf\u7684\u8bc1\u636e\u94fe\u3002",
    )

    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S1. \u8865\u5145\u5b9e\u9a8c\u6e05\u5355\u7684\u4e3b\u6587\u843d\u5b9e\u60c5\u51b5\u3002",
        ["\u4f18\u5148\u7ea7", "\u8865\u5145\u9879", "\u72b6\u6001", "\u4e3b\u6587\u5904\u7406"],
        tables["tracker"],
        "\u6ce8\uff1aP1 \u4e2d\u7684\u4fdd\u5f62\u9884\u6d4b\u76ee\u524d\u4ec5\u6709 90% \u7ecf\u9a8c\u8986\u76d6\u7387\u8bc1\u636e\uff0c80% \u548c 95% \u8986\u76d6\u7387\u672a\u5199\u6210\u5df2\u5b8c\u6210\u7ed3\u679c\u3002",
    )

    insert_para(
        anchor,
        "\u6269\u5c55\u9a8c\u8bc1\u96c6\u9009\u62e9\u5ba1\u8ba1\u663e\u793a\uff0c\u8de8 200 \u4e2a dataset-seed \u5019\u9009\u6c60\u5355\u5143\u7684\u9a8c\u8bc1\u96c6\u4e0e\u6d4b\u8bd5\u96c6\u6392\u540d\u4e2d\u4f4d Spearman \u4e3a 0.667\uff0c\u6d4b\u8bd5\u6700\u4f18\u5019\u9009\u843d\u5165\u9a8c\u8bc1 Top-3 \u7684\u6bd4\u4f8b\u4e3a 29.5%\uff0c\u9a8c\u8bc1\u96c6\u7b2c\u4e00\u540d\u4e0e\u6d4b\u8bd5\u96c6\u7b2c\u4e00\u540d\u5b8c\u5168\u4e00\u81f4\u7684\u6bd4\u4f8b\u4e3a 13.5%\u3002\u56e0\u6b64\uff0c\u9009\u62e9\u5668\u7684\u79d1\u5b66\u542b\u4e49\u4e0d\u662f\u4fdd\u8bc1\u6d4b\u8bd5\u6700\u4f18\uff0c\u800c\u662f\u5728\u6d4b\u8bd5\u96c6\u9501\u5b9a\u524d\u63d0\u4f9b\u53ef\u51bb\u7ed3\u3001\u53ef\u590d\u6838\u7684\u6cbb\u7406\u89c4\u5219\u3002",
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S2. \u9a8c\u8bc1\u96c6\u9009\u62e9\u504f\u5dee\u7684\u589e\u5f3a\u5ba1\u8ba1\u3002",
        ["\u6765\u6e90", "n", "\u4e2d\u4f4dSpearman", "Top-1\u4e00\u81f4", "\u6d4b\u8bd5\u6700\u4f18\u5728\u9a8c\u8bc1Top-3", "\u8d1f\u76f8\u5173\u6570", "optimism gap", "test regret"],
        tables["rank"],
        "\u6ce8\uff1aoptimism gap \u548c test regret \u5747\u6309\u5404\u4efb\u52a1\u539f\u59cb\u4e3b\u6307\u6807\u65b9\u5411\u8ba1\u7b97\uff0c\u53ea\u7528\u4e8e\u9009\u62e9\u98ce\u9669\u89e3\u91ca\uff0c\u4e0d\u7528\u4e8e\u91cd\u65b0\u9009\u6a21\u578b\u3002",
    )

    insert_para(
        anchor,
        "\u7cfb\u7edf\u6d88\u878d\u4e0e\u914d\u5bf9\u7edf\u8ba1\u8fdb\u4e00\u6b65\u9650\u5b9a\u4e86\u6a21\u5757\u8d21\u732e\u3002\u5b8c\u6574\u9009\u62e9\u5668\u76f8\u5bf9\u5355\u4e00 Chemprop\u3001\u591a\u6307\u7eb9\u3001\u51bb\u7ed3\u9884\u8bad\u7ec3\u8868\u5f81\u548c\u6838\u5fc3\u5bb6\u65cf\u5019\u9009\u5747\u4fdd\u6301\u51c0\u6b63\u5411\uff0c\u4f46 no_chemprop \u548c no_pretrained \u7684\u8d1f\u5411\u6216\u4e0d\u7a33\u5b9a\u7ed3\u679c\u8bf4\u660e\uff0c\u5019\u9009\u5bb6\u65cf\u5e76\u975e\u8d8a\u591a\u8d8a\u597d\uff0c\u5fc5\u987b\u7ecf\u8fc7\u9a8c\u8bc1\u95e8\u63a7\u540e\u624d\u80fd\u8fdb\u5165\u6700\u7ec8\u4fdd\u7559\u7b56\u7565\u3002",
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S3. \u7cfb\u7edf\u6d88\u878d\u4e0e\u5f3a\u57fa\u7ebf\u914d\u5bf9\u7edf\u8ba1\u3002",
        ["\u5bf9\u7167", "\u5e73\u5747\u6b63\u5411\u53d8\u5316", "win/tie/loss", "bootstrap p", "Wilcoxon p", "\u89e3\u91ca"],
        tables["ablation"],
        "\u6ce8\uff1awin/tie/loss \u6309 seed \u6216\u7ec8\u70b9\u914d\u5bf9\u5355\u5143\u7edf\u8ba1\uff1b\u672a\u7ed9\u51fa p \u503c\u7684\u884c\u4f5c\u4e3a\u5bb6\u65cf\u6d88\u878d\u63cf\u8ff0\u6027\u7ed3\u679c\u3002",
    )

    insert_para(
        anchor,
        "\u4f4e\u76f8\u4f3c\u5ea6\u548c\u7ed3\u6784\u5206\u79bb\u7ed3\u679c\u63d0\u9192\uff0c\u5e73\u5747\u6027\u80fd\u5e76\u4e0d\u80fd\u7b49\u540c\u4e8e\u65b0\u9aa8\u67b6\u5916\u63a8\u80fd\u529b\u3002\u5728\u66f4\u4f4e Tanimoto \u9608\u503c\u4e0b\uff0c\u56de\u5f52\u4efb\u52a1\u7684 RMSE \u660e\u663e\u5347\u9ad8\uff1b\u7ed3\u6784\u5206\u79bb\u5212\u5206\u4e2d FreeSolv\u3001Pgp\u3001BACE \u7b49\u4efb\u52a1\u51fa\u73b0\u660e\u663e\u5212\u5206\u60e9\u7f5a\u3002\u8fd9\u4e00\u7ed3\u679c\u652f\u6301\u5c06\u9002\u7528\u57df\u3001\u98ce\u9669-\u8986\u76d6\u66f2\u7ebf\u548c\u6700\u8fd1\u90bb\u8bc1\u636e\u4e0e\u6027\u80fd\u5206\u6570\u540c\u65f6\u62a5\u544a\u3002",
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S4. \u4f4e\u76f8\u4f3c\u5ea6\u56f0\u96be\u5b50\u96c6\u7684\u6027\u80fd\u53d8\u5316\u3002",
        ["\u4efb\u52a1", "Tanimoto\u9608\u503c", "\u5e73\u5747n", "\u5e73\u5747\u76f8\u4f3c\u5ea6", "ROC-AUC", "PR-AUC", "RMSE", "MAE"],
        tables["low_similarity"],
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S5. random/scaffold/structure-separated \u5212\u5206\u4e0b\u7684\u7ed3\u6784\u504f\u79fb\u538b\u529b\u3002",
        ["\u6570\u636e\u96c6", "\u4efb\u52a1", "\u6307\u6807", "random", "scaffold", "structure", "scaffold\u2192structure"],
        tables["structure_shift"],
        "\u6ce8\uff1a\u5206\u7c7b\u4efb\u52a1\u4e2d\u6b63\u503c\u8868\u793a AUC \u4e0b\u964d\uff1b\u56de\u5f52\u4efb\u52a1\u4e2d\u6b63\u503c\u8868\u793a\u8bef\u5dee\u589e\u52a0\u3002",
    )

    insert_para(
        anchor,
        "MoleculeACE \u914d\u5bf9\u7ed3\u679c\u4e0e\u4f4e\u76f8\u4f3c\u5ea6\u5206\u6790\u4e92\u76f8\u8865\u5145\u3002\u6d3b\u6027\u60ac\u5d16\u76ee\u6807\u5019\u9009\u5728 51 \u4e2a seed \u914d\u5bf9\u4e2d\u7684\u603b RMSE \u5e73\u5747\u6b63\u5411\u53d8\u5316\u4e3a 0.0069\uff0c\u60ac\u5d16\u5b50\u96c6 RMSE \u5e73\u5747\u6b63\u5411\u53d8\u5316\u4e3a 0.0056\uff0c\u4f46\u6807\u51c6\u5dee\u8f83\u5927\uff0c\u4efb\u52a1\u95f4\u5b58\u5728\u660e\u663e\u5dee\u5f02\u3002\u56e0\u6b64\uff0c\u672c\u6587\u5c06\u8be5\u7ed3\u679c\u89e3\u91ca\u4e3a\u60ac\u5d16\u98ce\u9669\u8bc6\u522b\u548c\u5019\u9009\u6cbb\u7406\u7684\u8865\u5145\u8bc1\u636e\uff0c\u800c\u4e0d\u662f\u5bf9\u6570\u503c\u60ac\u5d16\u9884\u6d4b\u5df2\u7ecf\u89e3\u51b3\u7684\u8bc1\u660e\u3002",
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S6. MoleculeACE \u4ee3\u8868\u6027\u6d3b\u6027\u60ac\u5d16\u4efb\u52a1\u7684\u914d\u5bf9\u7ed3\u679c\u3002",
        ["\u4efb\u52a1", "seed", "\u57fa\u7ebfRMSE", "\u60ac\u5d16\u5019\u9009RMSE", "\u0394RMSE", "\u57fa\u7ebfcliff RMSE", "\u60ac\u5d16cliff RMSE", "\u0394cliff RMSE"],
        tables["moleculeace"],
        "\u6ce8\uff1a\u0394 \u4e3a\u6b63\u65f6\u8868\u793a\u60ac\u5d16\u76ee\u6807\u5019\u9009\u8f83\u57fa\u7ebf\u8bef\u5dee\u66f4\u4f4e\u3002",
    )

    insert_para(
        anchor,
        "\u4e0d\u5e73\u8861\u5206\u7c7b\u4e0e\u4fdd\u5f62\u9884\u6d4b\u7684\u8865\u5145\u7ed3\u679c\u6539\u5584\u4e86 ClinTox\u3001DILI\u3001hERG \u548c CYP \u5e95\u7269\u7b49\u4efb\u52a1\u7684\u89e3\u91ca\u3002ROC-AUC \u4ecd\u4f5c\u4e3a\u6807\u51c6\u6307\u6807\uff0c\u4f46\u4e3b\u6587\u540c\u65f6\u62a5\u544a PR-AUC\u3001Brier\u3001ECE \u548c\u5bcc\u96c6\u6307\u6807\uff0c\u5e76\u5c06 90% \u7ecf\u9a8c\u8986\u76d6\u7387\u4f5c\u4e3a\u4e0d\u786e\u5b9a\u6027\u8fb9\u754c\u8bc1\u636e\u3002\u7531\u4e8e\u5f53\u524d\u7ed3\u679c\u6ca1\u6709 80% \u548c 95% \u8986\u76d6\u7387\u5b8c\u6574\u8868\uff0c\u672c\u6587\u4e0d\u5c06\u591a\u8986\u76d6\u7387\u4fdd\u5f62\u7ed3\u679c\u5199\u6210\u5df2\u5b8c\u6210\u5b9e\u9a8c\u3002",
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S7. \u4e0d\u5e73\u8861\u5206\u7c7b\u7684 PR-AUC\u3001\u6821\u51c6\u548c\u5bcc\u96c6\u6307\u6807\u3002",
        ["\u6570\u636e\u96c6", "\u9633\u6027\u7387", "ROC-AUC", "PR-AUC", "Brier", "ECE", "EF1", "EF5"],
        tables["imbalanced"],
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S8. \u4fdd\u5f62\u9884\u6d4b 90% \u76ee\u6807\u8986\u76d6\u7387\u6458\u8981\u3002",
        ["\u4efb\u52a1", "\u76ee\u6807\u8986\u76d6", "n", "\u5e73\u5747\u8986\u76d6", "\u4e2d\u4f4d\u8986\u76d6", "\u5e73\u5747\u96c6\u5927\u5c0f", "\u5355\u4f8b\u7387", "\u5e73\u5747\u533a\u95f4\u5bbd\u5ea6"],
        tables["conformal"],
        "\u6ce8\uff1a\u5206\u7c7b\u4efb\u52a1\u62a5\u544a\u9884\u6d4b\u96c6\u5927\u5c0f\uff0c\u56de\u5f52\u4efb\u52a1\u62a5\u544a\u9884\u6d4b\u533a\u95f4\u5bbd\u5ea6\u3002",
    )

    insert_para(
        anchor,
        "\u6700\u540e\uff0c\u8d1f\u7ed3\u679c\u548c\u5931\u8d25\u6848\u4f8b\u88ab\u660e\u786e\u4fdd\u7559\u3002ClinTox \u9ad8\u98ce\u9669\u5047\u9634\u6027\u8868\u660e\uff0c\u9ad8 ROC-AUC \u4efb\u52a1\u4ecd\u9700\u6837\u672c\u7ea7\u98ce\u9669\u8bc1\u636e\uff1b3D-lite \u548c\u7c97\u7cd9\u5ea6\u52a0\u6743\u5728 oracle \u6761\u4ef6\u4e0b\u5076\u6709\u6f5c\u5728\u6536\u76ca\uff0c\u4f46\u9a8c\u8bc1\u96c6\u95e8\u63a7\u672a\u5c06\u5176\u63a5\u5165\u6700\u7ec8\u7b56\u7565\u3002\u8fd9\u4e9b\u7ed3\u679c\u5171\u540c\u754c\u5b9a FZYC-Mol \u7684\u4f7f\u7528\u8fb9\u754c\uff0c\u4e5f\u907f\u514d\u8bba\u6587\u53ea\u62a5\u544a\u6b63\u7ed3\u679c\u3002",
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S9. \u5b9a\u5411\u6848\u4f8b\u4e0e\u5931\u8d25\u6a21\u5f0f\u3002",
        ["\u6848\u4f8b", "\u6570\u636e\u96c6", "\u7c7b\u578b", "\u6307\u6807", "\u4fee\u6539\u524d", "\u4fee\u6539\u540e", "\u0394", "\u89e3\u91ca"],
        tables["cases"],
    )
    insert_table(
        anchor,
        doc,
        "\u8865\u5145\u8868 S10. \u8d1f\u7ed3\u679c\u548c\u672a\u63a5\u5165\u5019\u9009\u7684\u5ba1\u8ba1\u3002",
        ["\u6a21\u5757", "\u8303\u56f4", "\u6b63\u5411\u7387", "\u5e73\u5747\u53d8\u5316", "\u51b3\u7b56"],
        tables["negative"],
        "\u6ce8\uff1a\u8d1f\u7ed3\u679c\u4fdd\u7559\u4e3a\u8bca\u65ad\u8bc1\u636e\uff0c\u4e0d\u53c2\u4e0e\u6700\u7ec8\u4fdd\u7559\u7b56\u7565\u7684\u4e8b\u540e\u8c03\u6574\u3002",
    )


def update_core_text(doc: Document) -> None:
    replace_paragraph(
        doc,
        "\u53ef\u9760\u6027\u5206\u6790\u663e\u793a",
        "\u53ef\u9760\u6027\u5206\u6790\u663e\u793a\uff0c\u96c6\u6210\u6807\u51c6\u5dee\u3001\u9519\u8bef\u6a21\u578b\u3001\u9519\u8bef-\u9002\u7528\u57df\u6df7\u5408\u6307\u6807\u3001Tanimoto \u9002\u7528\u57df\u548c\u91cd\u6784\u8bef\u5dee\u80fd\u591f\u8bc6\u522b\u9ad8\u8bef\u5dee\u6837\u672c\uff1b\u98ce\u9669\u5206\u6570\u5bf9\u5206\u7c7b\u9519\u8bef\u7684\u4e2d\u4f4d AUROC \u4e3a 0.788\uff0c\u5bf9\u56de\u5f52\u9ad8\u8bef\u5dee\u6837\u672c\u7684\u4e2d\u4f4d AUROC \u4e3a 0.652\u3002\u4fdd\u5f62\u9884\u6d4b\u5728 MoleculeNet \u5206\u7c7b\u548c\u56de\u5f52\u4efb\u52a1\u4e0a\u7684 90% \u76ee\u6807\u7ecf\u9a8c\u8986\u76d6\u7387\u5206\u522b\u4e3a 0.918 \u548c 0.925\u3002\u65b0\u589e\u6392\u540d\u5ba1\u8ba1\u8fdb\u4e00\u6b65\u663e\u793a\uff0c\u8de8 200 \u4e2a dataset-seed \u5019\u9009\u6c60\u7684\u9a8c\u8bc1-\u6d4b\u8bd5\u6392\u540d\u4e2d\u4f4d Spearman \u4e3a 0.667\uff0c\u6d4b\u8bd5\u6700\u4f18\u843d\u5165\u9a8c\u8bc1 Top-3 \u7684\u6bd4\u4f8b\u4e3a 0.295\uff0c\u63d0\u793a\u9a8c\u8bc1\u96c6\u6cbb\u7406\u662f\u53ef\u51bb\u7ed3\u3001\u53ef\u5ba1\u8ba1\u7684\u9009\u62e9\u6d41\u7a0b\uff0c\u800c\u4e0d\u662f\u6d4b\u8bd5\u6700\u4f18\u4fdd\u8bc1\u3002\u57fa\u5e8f\u5f52\u56e0\u548c\u7247\u6bb5\u5bcc\u96c6\u4e3a BBBP\u3001BACE\u3001ClinTox \u53ca\u82e5\u5e72 ADME \u56de\u5f52\u4efb\u52a1\u63d0\u4f9b\u4e86\u53ef\u8bfb\u7684\u5316\u5b66\u89e3\u91ca\u3002\u672a\u88ab\u9a8c\u8bc1\u96c6\u63a5\u53d7\u7684 3D-lite\u3001\u7c97\u7cd9\u5ea6\u52a0\u6743\u56de\u5f52\u548c\u672a\u6b63\u5f0f\u8ba1\u5165\u4e3b\u7ed3\u679c\u7684\u8868\u683c\u57fa\u7840\u6a21\u578b\u5148\u5bfc\u901a\u9053\u88ab\u4fdd\u7559\u4e3a\u8d1f\u7ed3\u679c\uff0c\u4ee5\u754c\u5b9a\u65b9\u6cd5\u8fb9\u754c\u3002",
    )
    replace_paragraph(
        doc,
        "\u4e3a\u5ba1\u8ba1\u9a8c\u8bc1\u96c6\u9009\u62e9\u662f\u5426\u53ef\u80fd\u8fc7\u62df\u5408",
        "\u4e3a\u5ba1\u8ba1\u9a8c\u8bc1\u96c6\u9009\u62e9\u662f\u5426\u53ef\u80fd\u8fc7\u62df\u5408\uff0c\u672c\u6587\u989d\u5916\u8ba1\u7b97\u6bcf\u4e2a dataset-seed \u5019\u9009\u6c60\u4e2d\u9a8c\u8bc1\u96c6\u6392\u540d\u4e0e\u6d4b\u8bd5\u96c6\u6392\u540d\u7684 Spearman \u76f8\u5173\u6027\uff0c\u5e76\u8bb0\u5f55\u9a8c\u8bc1\u96c6\u7b2c\u4e00\u540d\u662f\u5426\u4e5f\u662f\u6d4b\u8bd5\u96c6\u7b2c\u4e00\u540d\u3001\u6d4b\u8bd5\u96c6\u7b2c\u4e00\u540d\u662f\u5426\u843d\u5165\u9a8c\u8bc1\u96c6 Top-3\uff0c\u4ee5\u53ca\u9a8c\u8bc1\u96c6\u7b2c\u4e00\u540d\u5728\u6d4b\u8bd5\u96c6\u4e0a\u7684 regret/optimism gap\u3002\u8be5\u5206\u6790\u4e0d\u53c2\u4e0e\u6a21\u578b\u9009\u62e9\uff0c\u53ea\u7528\u4e8e\u8bc4\u4f30\u9009\u62e9\u5668\u98ce\u9669\uff1a\u82e5\u9a8c\u8bc1-\u6d4b\u8bd5\u6392\u540d\u76f8\u5173\u6027\u8f83\u4f4e\u3001Top-3 \u547d\u4e2d\u4e0d\u8db3\u6216\u51fa\u73b0\u8d1f\u76f8\u5173\uff0c\u76f8\u5173\u7ed3\u679c\u5728\u6b63\u6587\u4e2d\u6309\u8fb9\u754c\u6216\u8d1f\u7ed3\u679c\u89e3\u91ca\uff0c\u800c\u4e0d\u4f5c\u4e3a\u65e0\u6761\u4ef6\u6027\u80fd\u63d0\u5347\u8bc1\u636e\u3002",
    )
    replace_paragraph(
        doc,
        "\u7b2c\u4e8c\u4e2a\u53ef\u80fd\u95ee\u9898\u662f",
        "\u7b2c\u4e8c\u4e2a\u53ef\u80fd\u95ee\u9898\u662f\uff0c\u9009\u62e9\u5668\u662f\u5426\u4f1a\u8fc7\u62df\u5408\u9a8c\u8bc1\u96c6\u3002\u672c\u6587\u901a\u8fc7\u56db\u4e2a\u8bbe\u8ba1\u964d\u4f4e\u8fd9\u4e00\u98ce\u9669\u3002\u9996\u5148\uff0c\u9009\u62e9\u5668\u53ea\u5728\u9884\u5b9a\u4e49\u5019\u9009\u6c60\u4e2d\u9009\u62e9\uff0c\u4e0d\u5141\u8bb8\u6839\u636e\u6d4b\u8bd5\u7ed3\u679c\u6dfb\u52a0\u4e34\u65f6\u89c4\u5219\u3002\u5176\u6b21\uff0c\u7ed3\u679c\u6309\u591a\u4e2a\u968f\u673a\u79cd\u5b50\u6c47\u603b\uff0c\u5e76\u62a5\u544a\u5747\u503c\u3001\u6807\u51c6\u5dee\u548c\u914d\u5bf9\u7edf\u8ba1\u3002\u7b2c\u4e09\uff0c\u8865\u6551\u5934\u8fdb\u5165\u6700\u7ec8\u7b56\u7565\u9700\u8981\u901a\u8fc7\u9a8c\u8bc1\u96c6\u63a5\u53d7\uff0c\u4e14\u6ca1\u6709\u88ab\u6240\u6709\u7ec8\u70b9\u65e0\u5dee\u522b\u91c7\u7528\u3002\u7b2c\u56db\uff0c\u672c\u8f6e\u589e\u52a0\u6392\u540d\u5ba1\u8ba1\u3001Top-3 \u547d\u4e2d\u3001optimism gap \u548c\u8d1f\u7ed3\u679c\u8868\uff0c\u5c06\u4f4e\u76f8\u5173\u5019\u9009\u6c60\u660e\u786e\u5199\u4e3a\u9009\u62e9\u98ce\u9669\u800c\u975e\u80dc\u5229\u8bc1\u636e\u3002\u5c3d\u7ba1\u5982\u6b64\uff0c\u5b8c\u6574\u5d4c\u5957\u9a8c\u8bc1\u548c\u8de8\u6570\u636e\u96c6\u8fc1\u79fb\u9a8c\u8bc1\u4ecd\u662f\u540e\u7eed\u52a0\u5f3a\u65b9\u5411\u3002",
    )
    replace_paragraph(
        doc,
        "\u8fdb\u4e00\u6b65\u9a8c\u8bc1\u4f18\u5148\u6269\u5c55\u5916\u90e8\u8bc4\u4f30",
        "\u8fdb\u4e00\u6b65\u9a8c\u8bc1\u4ecd\u5e94\u4f18\u5148\u6269\u5c55\u5916\u90e8\u8bc4\u4f30\u9644\u5f55\uff0c\u4f46\u4e0d\u5e94\u628a\u76ee\u6807\u7b80\u5316\u4e3a\u7ee7\u7eed\u5806\u53e0\u66f4\u591a\u6a21\u578b\u3002\u66f4\u6709\u4ef7\u503c\u7684\u8def\u7ebf\u662f\u5728\u66f4\u591a\u516c\u5f00 ADMET \u7ec8\u70b9\u4e0a\u7cfb\u7edf\u8bc4\u4f30 CatBoost\u3001XGBoost\u3001ExtraTrees\u3001LightGBM\u3001RF\u3001Top-K \u96c6\u6210\u3001\u76ee\u6807\u53d8\u6362\u548c\u6b20\u91c7\u6837\u96c6\u6210\uff0c\u540c\u65f6\u5c06\u6700\u7ec8\u4fdd\u7559\u51b3\u7b56\u4e0e\u7c97\u7cd9\u5ea6\u3001\u4f4e\u76f8\u4f3c\u5ea6\u548c\u9a8c\u8bc1-\u6d4b\u8bd5\u6392\u540d\u4e00\u81f4\u6027\u5173\u8054\u3002",
    )
    replace_paragraph(
        doc,
        "\u7b2c\u4e8c\u4f18\u5148\u7ea7\u662f\u7cfb\u7edf\u5316\u4ec5\u57fa\u4e8e\u9a8c\u8bc1\u96c6\u7684 Top-K",
        "\u7b2c\u4e8c\u4f18\u5148\u7ea7\u662f\u628a\u76ee\u524d\u5df2\u6709\u7684 Top-K\u3001\u5806\u53e0\u3001\u98ce\u9669\u8c03\u6574\u548c\u7a33\u5b9a\u6027\u6253\u7834\u5e73\u5c40\u89c4\u5219\u6539\u6210\u66f4\u4e25\u683c\u7684\u9884\u6ce8\u518c\u5019\u9009\u65cf\u3002\u5f53\u524d\u8865\u6551\u6574\u5408\u9009\u62e9\u5668\u5df2\u663e\u793a Lipophilicity \u80fd\u53d7\u76ca\uff0cMoleculeACE \u60ac\u5d16\u76ee\u6807\u5019\u9009\u4e5f\u6709\u5c0f\u5e45\u5e73\u5747\u6536\u76ca\uff0c\u4f46\u8fd9\u4e9b\u7ed3\u679c\u4ecd\u9700\u5728\u5d4c\u5957\u9a8c\u8bc1\u6216\u5916\u90e8\u65f6\u95f4\u5212\u5206\u4e0a\u590d\u6838\u3002",
    )
    replace_paragraph(
        doc,
        "\u7b2c\u4e09\u4f18\u5148\u7ea7\u662f\u89e3\u91ca\u6027\u6848\u4f8b\u7814\u7a76",
        "\u7b2c\u4e09\u4f18\u5148\u7ea7\u662f\u7ee7\u7eed\u52a0\u5f3a\u89e3\u91ca\u6027\u6848\u4f8b\u7814\u7a76\u3002\u672c\u6587\u5df2\u5c06 Lipophilicity \u8865\u6551\u3001ClinTox \u9ad8\u98ce\u9669\u5047\u9634\u6027\u548c\u9ad8\u7c97\u7cd9\u5ea6 ADME \u56de\u5f52\u5199\u5165\u4e3b\u6587\uff1b\u540e\u7eed\u53ef\u6269\u5c55 BACE\u3001clearance_microsome_az \u548c CYP \u5e95\u7269\u7ec8\u70b9\uff0c\u5e76\u5c06\u7ed3\u6784\u7247\u6bb5\u3001\u6700\u8fd1\u90bb\u3001\u6807\u7b7e\u5dee\u5f02\u3001\u4e0d\u786e\u5b9a\u6027\u548c\u57fa\u5e8f\u5f52\u56e0\u8fde\u63a5\u5230\u540c\u4e00\u4e2a\u6837\u672c\u7ea7\u5ba1\u8ba1\u8868\u4e2d\u3002",
    )
    replace_paragraph(
        doc,
        "\u4fee\u8ba2\u540e\u7684\u7ed3\u679c\u5448\u73b0\u91c7\u7528\u4e3b\u7ebf\u878d\u5408\u7ed3\u6784",
        "\u4fee\u8ba2\u540e\u7684\u7ed3\u679c\u5448\u73b0\u91c7\u7528\u4e3b\u7ebf\u878d\u5408\u7ed3\u6784\uff1aMoleculeNet \u6027\u80fd\u3001TDC \u5916\u90e8\u6cdb\u5316\u3001\u5212\u5206\u771f\u5b9e\u6027\u3001\u53ef\u9760\u6027/\u9002\u7528\u57df\u3001\u9009\u62e9\u5668\u6cbb\u7406\u3001\u8865\u5145\u5b9e\u9a8c\u95ed\u73af\u548c\u57fa\u5e8f/\u6848\u4f8b\u89e3\u91ca\u5206\u522b\u627f\u8f7d\u76f8\u5e94\u8bc1\u636e\u3002\u6b63\u5f0f\u5916\u90e8\u9644\u5f55\u3001\u6027\u80fd\u6a21\u5f0f\u3001\u8865\u6551\u6574\u5408\u9009\u62e9\u5668\u3001\u591a\u65b9\u6cd5\u878d\u5408\u30013D-lite/\u7c97\u7cd9\u5ea6\u8d1f\u7ed3\u679c\u3001MoleculeACE \u6d3b\u6027\u60ac\u5d16\u548c\u9009\u62e9\u5668\u7b56\u7565\u5ba1\u8ba1\u5747\u5d4c\u5165\u4e0a\u8ff0\u4e3b\u7ebf\u4e2d\u89e3\u91ca\uff0c\u800c\u4e0d\u518d\u4f5c\u4e3a\u6309\u65f6\u95f4\u8ffd\u52a0\u7684\u72ec\u7acb\u5b9e\u9a8c\u5757\u3002",
    )
    replace_paragraph(
        doc,
        "\u672c\u6587\u63d0\u51fa\u5e76\u7cfb\u7edf\u8bc4\u4f30\u4e86 FZYC-Mol",
        "\u672c\u6587\u63d0\u51fa\u5e76\u7cfb\u7edf\u8bc4\u4f30\u4e86 FZYC-Mol\uff0c\u4e00\u79cd\u7531\u9a8c\u8bc1\u96c6\u6cbb\u7406\u3001\u9002\u7528\u57df\u611f\u77e5\u7684\u591a\u4e13\u5bb6\u5206\u5b50\u6027\u8d28\u9884\u6d4b\u6846\u67b6\u3002\u7ed3\u679c\u8868\u660e\uff0c\u5728 MoleculeNet\u3001TDC ADMET\u3001\u5916\u90e8\u8bc4\u4f30\u9644\u5f55\u3001\u5212\u5206\u771f\u5b9e\u6027\u3001\u4f4e\u76f8\u4f3c\u5ea6\u5b50\u96c6\u3001MoleculeACE \u6d3b\u6027\u60ac\u5d16\u3001\u7c97\u7cd9\u5ea6\u8bca\u65ad\u548c\u57fa\u5e8f/\u7247\u6bb5\u89e3\u91ca\u7b49\u591a\u6761\u8bc1\u636e\u7ebf\u4e0a\uff0cFZYC-Mol \u80fd\u591f\u63d0\u4f9b\u6bd4\u5355\u4e00\u6a21\u578b\u5206\u6570\u66f4\u5b8c\u6574\u7684\u53ef\u9760\u6027\u753b\u50cf\u3002\u5728\u540c\u4e00\u9a8c\u8bc1\u96c6\u5019\u9009\u6c60\u4e2d\uff0c\u8865\u6551\u5934\u3001\u4f4e\u6210\u672c\u5b9a\u5411\u91cd\u6784\u548c\u591a\u65b9\u6cd5\u878d\u5408\u5206\u522b\u5728 Lipophilicity\u3001FreeSolv\u3001BBBP/ClinTox \u4ee5\u53ca\u5916\u90e8 TDC \u5b98\u65b9\u9762\u677f\u7684 Caco2\u3001HIA \u548c Pgp \u4e0a\u5f62\u6210\u6700\u7ec8\u4fdd\u7559\u589e\u76ca\uff1b\u4f46\u6392\u540d\u5ba1\u8ba1\u3001\u8d1f\u7ed3\u679c\u548c\u60ac\u5d16\u5b50\u96c6\u4e5f\u8868\u660e\uff0c\u8be5\u6846\u67b6\u66f4\u9002\u5408\u88ab\u7406\u89e3\u4e3a\u53ef\u5ba1\u8ba1\u7684\u6cbb\u7406\u6d41\u7a0b\uff0c\u800c\u4e0d\u662f\u4efb\u4f55\u7ec8\u70b9\u4e0a\u90fd\u80fd\u81ea\u52a8\u8fbe\u5230\u6d4b\u8bd5\u6700\u4f18\u7684\u5355\u4e00\u6a21\u578b\u3002",
    )


def main() -> None:
    src = find_source_docx()
    out = output_path(src)
    shutil.copy2(src, out)
    doc = Document(out)
    update_core_text(doc)
    insert_supplement_section(doc)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
