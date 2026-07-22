from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
DOCX = OUT / "小论文-14.docx"
AUDIT = OUT / "paper14_final_audit.json"

TITLE = "冻结验证揭示分子性质预测中候选池扩张的收益与选择损失"
REQUIRED = [
    "K=32 相对 K=4",
    "完整池范围归一化选择损失增加 0.122",
    "实际兑现效用增益为 0.343",
    "高遗憾 AUC 为 0.648",
    "平均遗憾降低 0.034",
    "有效候选数约为 1.01",
    "103,025 条测试样本-候选记录",
    "doi:10.1021/ci100050t",
    "doi:10.1023/A:1010933404324",
    "doi:10.1145/2939672.2939785",
    "doi:10.1002/minf.201000061",
    "doi:10.1021/ci400084k",
]
FORBIDDEN = [
    "有效候选数约为 1.11",
    "TabPFN 完成",
    "九终点全量深度模型面板已完成",
    "通用 SOTA 预测器",
]
COLLOQUIAL = [
    "很",
    "单纯",
    "这样做",
    "外部读者",
    "只是",
]


def all_text(doc: Document) -> str:
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        parts.append(p.text.strip())
    return "\n".join(parts)


def main() -> None:
    doc = Document(DOCX)
    text = all_text(doc)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    with zipfile.ZipFile(DOCX) as z:
        bad = z.testzip()
        xml = ET.fromstring(z.read("word/document.xml"))
        blips = xml.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
        tables = xml.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl")

    abstract_idx = paragraphs.index("摘要")
    intro_idx = paragraphs.index("1 引言")
    abstract_paras = [p for p in paragraphs[abstract_idx + 1 : intro_idx] if not p.startswith("关键词")]
    ref_idx = paragraphs.index("参考文献")
    refs = paragraphs[ref_idx + 1 :]

    result = {
        "docx": str(DOCX),
        "exists": DOCX.exists(),
        "size": DOCX.stat().st_size,
        "zip_test": bad or "ok",
        "paragraphs": len(paragraphs),
        "tables": len(tables),
        "figures": len(blips),
        "title_ok": paragraphs[0] == TITLE,
        "abstract_paragraph_count": len(abstract_paras),
        "required_hits": {s: (s in text) for s in REQUIRED},
        "forbidden_hits": {s: (s in text) for s in FORBIDDEN},
        "colloquial_hits": {s: len(re.findall(re.escape(s), text)) for s in COLLOQUIAL if s in text},
        "reference_count": len(refs),
        "figure_caption_1": next((p for p in paragraphs if p.startswith("图 1")), ""),
    }
    result["passed"] = (
        result["exists"]
        and result["zip_test"] == "ok"
        and result["tables"] == 9
        and result["figures"] == 9
        and result["title_ok"]
        and result["abstract_paragraph_count"] == 4
        and all(result["required_hits"].values())
        and not any(result["forbidden_hits"].values())
        and result["reference_count"] == 34
    )

    AUDIT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
