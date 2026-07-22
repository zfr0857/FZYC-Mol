# -*- coding: utf-8 -*-
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports" / "draft5_table_fix_20260616"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def locate_source() -> Path:
    candidates = [p for p in (Path.home() / "Desktop").rglob("初稿-5.docx")]
    if not candidates:
        raise FileNotFoundError("Could not locate 初稿-5.docx.")
    return max(candidates, key=lambda p: p.stat().st_mtime)


SRC_DOCX = locate_source()
DEST_DOCX = SRC_DOCX.parent / "初稿-5_表格修订版.docx"
REPORT_DOCX = SRC_DOCX.parent / "初稿-5_表格与整体QA报告.docx"


def set_border(element, edge: str, val: str, size: str = "8", color: str = "000000") -> None:
    tag = "w:tblBorders" if element.tag.endswith("tblPr") else "w:tcBorders"
    borders = element.find(qn(tag))
    if borders is None:
        borders = OxmlElement(tag)
        element.append(borders)
    edge_el = borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        borders.append(edge_el)
    edge_el.set(qn("w:val"), val)
    edge_el.set(qn("w:sz"), size)
    edge_el.set(qn("w:space"), "0")
    edge_el.set(qn("w:color"), color)


def set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))


def set_cell_margins(table, top=45, start=70, bottom=45, end=70) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for edge, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = cell_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width_and_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_header_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def apply_three_line_table(table) -> None:
    tbl_pr = table._tbl.tblPr
    for edge in ["top", "bottom"]:
        set_border(tbl_pr, edge, "single", "12")
    for edge in ["left", "right", "insideH", "insideV"]:
        set_border(tbl_pr, edge, "nil")

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for edge in ["left", "right", "insideH", "insideV"]:
                set_border(tc_pr, edge, "nil")
    if table.rows:
        for cell in table.rows[0].cells:
            set_border(cell._tc.get_or_add_tcPr(), "bottom", "single", "8")


def set_text_style(cell, size_pt: float, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.bold = bold


def set_row_values(table, row_idx: int, values: list[str]) -> None:
    cells = table.rows[row_idx].cells
    for i, value in enumerate(values):
        cells[i].text = value


def compact_table_content(doc) -> list[str]:
    changes: list[str] = []

    if len(doc.tables) >= 1:
        t = doc.tables[0]
        set_row_values(t, 0, ["数据集", "来源/任务", "n 或阳性率", "主指标", "划分与种子"])
        rows = [
            ["ESOL", "MoleculeNet; 回归", "1117", "RMSE", "Scaffold/random/结构分离/低相似度; 5 seeds"],
            ["FreeSolv", "MoleculeNet; 回归", "642", "RMSE", "Scaffold/random/结构分离/低相似度; 5 seeds"],
            ["Lipophilicity", "MoleculeNet; 回归", "4200", "RMSE", "Scaffold/random/结构分离/低相似度; 5 seeds"],
            ["BBBP", "MoleculeNet; 分类", "1975; 0.76", "ROC-AUC", "Scaffold/random/结构分离/低相似度; 5 seeds"],
            ["BACE", "MoleculeNet; 分类", "1513; 0.4567", "ROC-AUC", "Scaffold/random/结构分离/低相似度; 5 seeds"],
            ["ClinTox", "MoleculeNet; 分类", "1461; 0.0705", "ROC-AUC", "Scaffold/random/结构分离/低相似度; 5 seeds"],
            ["TDC ADMET", "外部 ADMET", "578-13130", "RMSE / ROC-AUC", "PyTDC 官方划分; scaffold 审计"],
            ["MoleculeACE", "活性悬崖", "可用任务子集", "RMSE / R2 / cliff", "cliff-pair 与 roughness 诊断"],
            ["bRo5", "规则五以外压力测试", "CycPept-PAMPA; LinPept", "RMSE / ROC-AUC / PR-AUC", "公开数据外推与 AD 审计; 非盲测"],
        ]
        for i, row in enumerate(rows, start=1):
            set_row_values(t, i, row)
        changes.append("表1：压缩 bRo5 与划分列，避免长句撑高行距。")

    if len(doc.tables) >= 2:
        t = doc.tables[1]
        set_row_values(t, 0, ["数据集", "指标", "选择器", "最终保留", "对照/观测最优", "说明"])
        rows = [
            ["ESOL", "RMSE", "0.5829 ± 0.0352", "0.5829 ± 0.0352", "Multi-fp 0.6352; best 0.5829", "低值优"],
            ["FreeSolv", "RMSE", "1.0678 ± 0.1883", "1.0286 ± 0.1761", "Chemprop 0.9518; best 0.9518", "边界案例"],
            ["Lipophilicity", "RMSE", "0.7078 ± 0.0389", "0.6835 ± 0.0439", "D-MPNN 0.7456; best 0.6835", "低值优"],
            ["BBBP", "ROC-AUC", "0.9165 ± 0.0290", "0.9243 ± 0.0247", "Multi-fp 0.9215; best 0.9243", "高值优"],
            ["BACE", "ROC-AUC", "0.8753 ± 0.0230", "0.8753 ± 0.0230", "D-MPNN 0.8571; best 0.8753", "高值优"],
            ["ClinTox", "ROC-AUC", "0.9489 ± 0.0302", "0.9496 ± 0.0262", "Multi-fp 0.9479; best 0.9496", "高值优"],
        ]
        for i, row in enumerate(rows, start=1):
            set_row_values(t, i, row)
        changes.append("表2：缩短模型对照列，保留关键数值并提升主文可读性。")

    if len(doc.tables) >= 3:
        t = doc.tables[2]
        set_row_values(t, 0, ["证据模块", "主要发现", "解释"])
        set_row_values(t, 1, ["TDC 多方法融合", "Caco2、HIA、Pgp 出现选择性增益", "多视图和 AD 门控对部分外部终点有效"])
        set_row_values(t, 2, ["22 个外部终点", "win/tie/loss = 5/17/0", "多数终点保留原策略，避免普遍提升主张"])
        set_row_values(t, 3, ["划分压力测试", "官方划分与 scaffold 划分差异明显", "划分真实性属于评价协议的一部分"])
        set_row_values(t, 4, ["后续验证", "部分模块需冻结候选池后再测试", "未就绪模块不进入主性能主张"])
        changes.append("表3：压缩外部 ADMET 证据模块表述。")

    if len(doc.tables) >= 5:
        t = doc.tables[4]
        set_row_values(t, 0, ["模块", "状态", "解释"])
        set_row_values(t, 1, ["FreeSolv 低成本重构", "缩小差距但未超过 Chemprop", "保留为物理相互作用相关边界案例"])
        set_row_values(t, 2, ["bRo5 化学空间", "公共压力测试已纳入", "仅支持公开数据外推与 AD 审计；非独立盲测"])
        set_row_values(t, 3, ["轻量适配器与 3D-lite", "受控扩展方向", "需继续用嵌套验证限制过拟合"])
        set_row_values(t, 4, ["粗糙度加权", "诊断优先", "相关性不稳定时作为负结果"])
        set_row_values(t, 5, ["基序归因与片段富集", "关联性解释", "不作为因果机制证据"])
        changes.append("表5：重点压缩 bRo5 行，修复截图所示跨页断裂风险。")

    return changes


def apply_table_layout(doc) -> list[str]:
    widths = {
        0: [0.72, 1.20, 1.40, 1.10, 1.88],
        1: [0.72, 0.70, 1.05, 1.05, 1.75, 0.72],
        2: [1.40, 2.15, 2.55],
        3: [1.85, 1.65, 2.60],
        4: [1.45, 2.00, 2.65],
        5: [1.30, 4.70],
    }
    notes = []
    for idx, table in enumerate(doc.tables):
        set_table_width_and_layout(table)
        set_cell_margins(table)
        apply_three_line_table(table)

        for row_idx, row in enumerate(table.rows):
            set_row_cant_split(row)
            if row_idx == 0:
                set_header_repeat(row)
            for col_idx, cell in enumerate(row.cells):
                if idx in widths and col_idx < len(widths[idx]):
                    set_cell_width(cell, widths[idx][col_idx])
                font_size = 7.5 if idx == 1 else 8.5
                set_text_style(cell, font_size, bold=(row_idx == 0))
        notes.append(f"表{idx + 1}：已设置固定列宽、表头重复、三线表和禁止跨页拆分。")
    return notes


def update_table_notes_and_caption_style(doc) -> list[str]:
    notes: list[str] = []
    replacements = {
        "注：回归任务以 RMSE 为主指标，数值越低越优；分类任务以 ROC-AUC 为主指标，数值越高越优。Scaffold、structure-separated 与 low-similarity 子集用于评估结构外推和低相似度压力。":
            "注：回归任务以 RMSE 为主指标，分类任务以 ROC-AUC 为主指标。bRo5 行中的 LinPept 包括 CellPen 与 NonFouling，相关结果仅作为公开压力测试。",
        "注：未通过验证集门控的候选策略作为负结果或后续验证接口保留，不纳入主性能结论。":
            "注：未通过验证集门控的候选策略作为负结果或后续验证接口保留，不纳入主性能结论；bRo5 公共压力测试不等同于独立盲测。",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            paragraph.text = replacements[text]
            notes.append(f"更新表注：{text[:20]}...")
        if text.startswith("表") or text.startswith("注："):
            paragraph.paragraph_format.keep_with_next = text.startswith("表")
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return notes


def whole_manuscript_cleanup(doc) -> list[str]:
    notes: list[str] = []
    replacements = {
        "图1. FZYC-Mol 模型结构：多源表示、专家预测矩阵、验证集选择器与证据输出。":
            "图1. FZYC-Mol 模型结构：多源表示、专家预测矩阵、验证集选择器与证据输出。",
        "图2. FZYC-Mol 整体工作流：数据划分、多视图表示、候选专家、验证集选择与可靠性输出。":
            "图2. FZYC-Mol 整体工作流：数据划分、多视图表示、候选专家、验证集选择与可靠性输出。",
    }
    for paragraph in doc.paragraphs:
        original = paragraph.text
        updated = original
        updated = updated.replace("—", "，")
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != original:
            paragraph.text = updated
            notes.append(f"全稿标点/术语清理：{original[:20]}...")
    return notes


def collect_qa(doc) -> list[str]:
    text = "\n".join(p.text for p in doc.paragraphs)
    qa = []
    qa.append(f"段落数：{len(doc.paragraphs)}；表格数：{len(doc.tables)}；图片数：{len(doc.inline_shapes)}。")
    qa.append(f"最大表格列数：{max((len(t.columns) for t in doc.tables), default=0)}。")
    qa.append(f"禁止跨页拆分行数：{sum(1 for t in doc.tables for r in t.rows if r._tr.get_or_add_trPr().find(qn('w:cantSplit')) is not None)}。")
    qa.append(f"长单元格检查（>90 字）：{sum(1 for t in doc.tables for r in t.rows for c in r.cells if len(c.text.strip()) > 90)} 个。")
    qa.append(f"内部自审章节残留：{text.count('6 投稿前自我审查')} 处。")
    for term in ["弱 基线模型", "每个 数据集", "backbone", "benchmark 和离线评估", "KPGT representation", "class weight", "validation metric", "selected candidate", "数据状态审计"]:
        qa.append(f"术语/旧表述检查 {term}：{text.count(term)} 处。")
    for heading in ["摘要", "1 引言", "2 材料与方法", "3 结果", "4 讨论", "5 结论", "List of abbreviations", "Declarations", "References"]:
        qa.append(f"章节检查 {heading}：{'存在' if heading in text else '未检出'}。")
    return qa


def build_report(changes: list[str], layout_notes: list[str], caption_notes: list[str], cleanup_notes: list[str], qa: list[str]) -> None:
    doc = Document()
    doc.add_heading("初稿-5 表格与整体 QA 报告", level=1)
    doc.add_paragraph(f"源文档：{SRC_DOCX}")
    doc.add_paragraph(f"输出文档：{DEST_DOCX}")
    doc.add_heading("Nature 风格处理原则", level=2)
    doc.add_paragraph(
        "本轮按 Nature 技能中的版式诊断原则处理表格：主文表应短、密、可读，不在单元格内放长段说明；长解释移入表注或正文；表格使用三线表、紧凑字号、固定列宽，并禁止行跨页拆分。"
    )
    doc.add_heading("表格内容修订", level=2)
    for item in changes:
        doc.add_paragraph(item)
    doc.add_heading("表格版式修订", level=2)
    for item in layout_notes:
        doc.add_paragraph(item)
    doc.add_heading("表注与全文检查", level=2)
    for item in caption_notes + cleanup_notes:
        doc.add_paragraph(item)
    doc.add_heading("QA 结果", level=2)
    for item in qa:
        doc.add_paragraph(item)
    doc.save(REPORT_DOCX)


def main() -> None:
    doc = Document(SRC_DOCX)
    content_changes = compact_table_content(doc)
    layout_notes = apply_table_layout(doc)
    caption_notes = update_table_notes_and_caption_style(doc)
    cleanup_notes = whole_manuscript_cleanup(doc)
    qa = collect_qa(doc)

    doc.save(DEST_DOCX)
    shutil.copy2(DEST_DOCX, OUT_DIR / DEST_DOCX.name)
    build_report(content_changes, layout_notes, caption_notes, cleanup_notes, qa)
    shutil.copy2(REPORT_DOCX, OUT_DIR / REPORT_DOCX.name)

    print(f"Wrote {DEST_DOCX}")
    print(f"Wrote {REPORT_DOCX}")
    for item in qa:
        print(item)


if __name__ == "__main__":
    main()
