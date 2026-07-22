from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from build_paper19_manuscript import (  # noqa: E402
    add_body,
    add_figure,
    add_heading,
    add_table,
    add_table_caption,
    configure_document,
    set_run_font,
)


OUT = ROOT / "output"
REV = OUT / "paper19_jcheminformatics_revision_20260712"
EXP = OUT / "paper19_rejection_driven_experiments_20260712"
FIG = REV / "figures"
PAPER = OUT / "小论文-19.docx"
DESKTOP = Path("C:/Users/Administrator/Desktop/小论文-19.docx")
CHANGELOG = OUT / "小论文-19_逐段修改说明.md"
REORG = OUT / "小论文-19_图表重组与投稿核查.md"
AUDIT = OUT / "小论文-19_Journal_of_Cheminformatics重构审计.json"


PARAGRAPH_LOG: list[dict[str, str]] = []


def body(doc: Document, section: str, text: str, reason: str, issue: str, *, no_indent: bool = False) -> None:
    add_body(doc, text, no_indent=no_indent)
    PARAGRAPH_LOG.append(
        {
            "section": section,
            "original_issue": issue,
            "revised_paragraph": text,
            "reason": reason,
        }
    )


def add_labelled_abstract(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{label}: ")
    set_run_font(r, 10.5, bold=True)
    r = p.add_run(text)
    set_run_font(r, 10.5)
    PARAGRAPH_LOG.append(
        {
            "section": f"Abstract - {label}",
            "original_issue": "The previous abstract mixed background, methods, many secondary numerical results and interpretation in a continuous list.",
            "revised_paragraph": f"{label}: {text}",
            "reason": "Recast the abstract into the requested background-methods-results-conclusions structure and retained only decision-critical results.",
        }
    )


def set_landscape_table_section(doc: Document) -> None:
    section = doc.add_section()
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def add_reference(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.left_indent = Cm(0.63)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run_font(r, 9)


def load_data() -> dict[str, pd.DataFrame]:
    return {
        "datasets": pd.read_csv(REV / "dataset_characteristics.csv"),
        "diversity": pd.read_csv(EXP / "paper19_effective_diversity.csv"),
        "ranking": pd.read_csv(EXP / "paper19_ranking_fidelity_units.csv"),
        "effects": pd.read_csv(EXP / "paper19_k32_vs_k4_endpoint_effects.csv"),
        "raw_loss": pd.read_csv(EXP / "paper19_raw_selection_loss_summary.csv"),
        "compute": pd.read_csv(EXP / "paper19_compute_budget.csv"),
        "composition": pd.read_csv(REV / "candidate_composition_controls.csv"),
        "multiview": pd.read_csv(REV / "multiview_endpoint_raw_gain.csv"),
        "baseline_settings": pd.read_csv(REV / "strong_baseline_settings.csv"),
        "reliability": pd.read_csv(REV / "reliability_summary.csv"),
        "strong_score": pd.read_csv(OUT / "sci1_hardening_20260707" / "six_task_strong_selection_scorecard.csv"),
        "overlap": pd.read_csv(OUT / "sci1_hardening_20260707" / "six_task_error_overlap_pairwise_summary.csv"),
        "clintox": pd.read_csv(OUT / "sci1_mechanism_uq_decision_20260707" / "clintox_minority_negative_result.csv"),
        "failure": pd.read_csv(OUT / "sci1_mechanism_uq_decision_20260707" / "failure_case_category_summary.csv"),
    }


def build_tables(doc: Document, data: dict[str, pd.DataFrame], which: int) -> None:
    if which == 1:
        rows = []
        for _, r in data["datasets"].iterrows():
            if r.task_type == "classification":
                target = f"positive {int(r.positive_n)} ({100*r.positive_rate:.1f}%)"
            else:
                target = f"{r.target_min:.2f} to {r.target_max:.2f}"
            rows.append(
                [
                    r.display_name,
                    r.task_type,
                    f"{int(r.raw_n)}/{int(r.analysis_n)}",
                    target,
                    r.target_unit,
                    r.evidence_role,
                ]
            )
        add_table_caption(doc, "Table 1. Datasets, analysis populations and evidence roles.")
        add_table(doc, ["Endpoint", "Task", "Raw/analysis n", "Class balance or range", "Unit", "Evidence role"], rows)
    elif which == 2:
        compute = data["compute"]
        total_fits = int(compute["total_fits"].sum())
        total_seconds = float(compute["total_fit_seconds"].sum())
        rows = [
            ["Near-duplicate expansion", "4, 8, 16, 32", "Morgan-512; linear, bagging and boosting variants", "validation-best; frozen alternatives", f"{total_fits:,} fits; {total_seconds:,.1f} recorded CPU fit-s"],
            ["Composition controls", "4, 8, 16, 32", "random order, random subset, family-balanced", "same frozen scoring", "100 frozen subset/order seeds per endpoint-K-mode"],
            ["Multiview confirmation", "12", "Morgan, MACCS, RDKit2D, concatenated view x 3 learners", "shared 3 x 3 x 5 splits", "6,480 fits"],
            ["Modern-baseline stress test", "4 completed", "RDKit-RF, GCN, ChemBERTa probe, MoLFormer probe", "shared 3 x 3 x 5 splits on six tasks", "360 outer; 1,080 inner; 220,040 predictions"],
            ["Unavailable full-audit candidates", "2", "Chemprop/D-MPNN; TabPFN", "not entered into confirmatory comparison", "needs new analysis"],
        ]
        add_table_caption(doc, "Table 2. Candidate pools, selection rules and computational exposure.")
        add_table(doc, ["Analysis", "Nominal K", "Candidates", "Selection/audit rule", "Exposure or status"], rows)
    elif which == 3:
        div = data["diversity"].groupby("candidate_count", as_index=False).agg(
            effective_rank=("outer_entropy_effective_rank", "mean"),
            correlation=("outer_median_pairwise_correlation", "mean"),
        )
        rank = data["ranking"].groupby("candidate_count", as_index=False).agg(
            spearman=("spearman_validation_vs_audit", "mean"),
            top1=("top1_hit", "mean"),
            mrr=("mrr", "mean"),
        )
        tab = div.merge(rank, on="candidate_count")
        rows = [[int(r.candidate_count), f"{r.effective_rank:.3f}", f"{r.correlation:.3f}", f"{r.spearman:.3f}", f"{r.top1:.3f}", f"{r.mrr:.3f}"] for _, r in tab.iterrows()]
        add_table_caption(doc, "Table 3. Nominal candidate scale, effective diversity and validation-ranking fidelity.")
        add_table(doc, ["K", "Effective rank", "Median correlation", "Validation-audit Spearman", "Top-1 hit", "MRR"], rows)
    elif which == 4:
        effects = data["effects"].loc[data["effects"].policy == "validation_best"].copy()
        names = {
            "tdc_caco2_wang": "Caco2 Wang",
            "tdc_hia_hou": "HIA Hou",
            "tdc_pgp_broccatelli": "P-gp Broccatelli",
        }
        rows = []
        for _, r in effects.iterrows():
            unit = "ROC-AUC loss" if r.task_type == "classification" else "RMSE loss"
            rows.append([names.get(r.task, r.task.upper() if r.task in {"bbbp", "bace"} else r.task.title()), unit, "15", f"{r.mean_delta_raw_loss_k32_minus_k4:.4f}", f"{r.ci95_low:.4f} to {r.ci95_high:.4f}", f"{100*r.direction_increase_fraction:.0f}%"])
        add_table_caption(doc, "Table 4. Endpoint-specific raw-scale selection-loss change from K=4 to K=32.")
        add_table(doc, ["Endpoint", "Effect scale", "Paired units", "Mean change", "95% bootstrap CI", "Units with increase"], rows)
    elif which == 5:
        rows = [
            ["Candidate-composition controls", "9 endpoints; 100 frozen subset/order seeds", "Normalized loss increased from approximately 0.09 at K=4 to 0.214 at K=32", "Sensitivity analysis; normalized scale"],
            ["Finite-audit simulation", "Equal-truth candidates; correlated Gaussian audit noise", "At correlation 0.9 and n_eff=50, optimism increased from 0.046 to 0.092 SD units", "Mechanical winner optimism, not empirical model bias"],
            ["Multiview confirmation", "9 endpoints; 135 paired outer units", "Normalized realized gain 0.343 (95% CI 0.210-0.483); 9/9 endpoints", "Shared-split representational confirmation"],
            ["Modern baselines", "6 MoleculeNet tasks; 4 completed candidates", "Top-1 0.878; normalized selection loss 0.0048; error-overlap 0.168-0.296", "Stress test, not nine-endpoint confirmation"],
            ["Conditional conformal", "Classification; 90% target", "Minority coverage 0.788 split, 0.887 label-conditional, 0.891 Mondrian", "Cross-fold reliability analysis"],
            ["CQR", "ESOL, FreeSolv, Lipophilicity", "Mean coverage 0.882; raw mean width 7.25", "Widths are not comparable across endpoint units"],
            ["Chemical boundaries", "Tanimoto, scaffold, activity cliffs and bRo5", "Performance and calibration deteriorated in low-support strata", "Exploratory boundary evidence"],
            ["Chemprop/D-MPNN and TabPFN", "Full nine-endpoint shared-split audit", "Not available", "needs new analysis"],
        ]
        add_table_caption(doc, "Table 5. Secondary stress tests and interpretation boundaries.")
        add_table(doc, ["Analysis", "Scope", "Key result", "Interpretation boundary"], rows)


def build_document(data: dict[str, pd.DataFrame]) -> Document:
    doc = Document()
    configure_document(doc)
    doc.sections[0].header.paragraphs[0].clear()
    h = doc.sections[0].header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(h.add_run("Validation-ranking distortion under candidate-pool expansion"), 8)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.first_line_indent = Cm(0)
    set_run_font(
        title.add_run("Validation-ranking distortion and selection loss under candidate-pool expansion in molecular property prediction: a frozen audit study"),
        16,
        bold=True,
    )
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.first_line_indent = Cm(0)
    set_run_font(sub.add_run("Research Article | Journal of Cheminformatics"), 10, italic=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.first_line_indent = Cm(0)
    set_run_font(sub.add_run("Authors and affiliations: to be completed by the authors before submission"), 9)

    add_heading(doc, "Abstract", 1)
    add_labelled_abstract(
        doc,
        "Background",
        "Molecular property studies increasingly compare many representations, learners and tuning variants on limited validation data. Candidate expansion can raise the best observed audit performance while simultaneously increasing the opportunity to select a validation-specific winner. The magnitude of this trade-off, and the extent to which nominal candidate count reflects genuinely distinct predictive behaviour, remain insufficiently quantified.",
    )
    add_labelled_abstract(
        doc,
        "Methods",
        "We conducted a retrospective frozen audit across nine molecular-property endpoints. Candidate pools of K=4, 8, 16 and 32 were evaluated with five seeds, three outer scaffold folds and three inner folds. Primary estimands were endpoint-specific raw-scale selection loss, effective candidate rank and validation-audit ranking fidelity. Frozen composition controls, an equal-truth finite-audit simulation, a shared-split 12-candidate multiview experiment, four completed modern baselines, and reliability and chemical-boundary analyses were used to test mechanism and scope.",
    )
    add_labelled_abstract(
        doc,
        "Results",
        "At K=32, the mean effective rank was 2.00 and the median candidate-utility correlation was 0.888, showing that nominal expansion greatly exceeded behavioural diversity. Relative to K=4, validation-best selection loss increased in eight of nine endpoints; the mean increase was 0.0096 ROC-AUC for classification, whereas regression effects were retained on their endpoint-specific RMSE scales. Shared-split multiview selection produced a normalized realized gain of 0.343 (95% CI 0.210-0.483) over Morgan-only selection across all nine endpoints. An exploratory leave-one-endpoint-out risk score reduced normalized loss by 0.034 (95% CI 0.020-0.047), but the small endpoint count precludes a general meta-selector claim.",
    )
    add_labelled_abstract(
        doc,
        "Conclusions",
        "In these endpoints, adding highly correlated candidates increased selection loss and weakened validation ranking even when all candidates used a frozen audit protocol. Heterogeneous representations could yield realizable gains, but those gains were only interpretable after accounting for selection opportunity and compute exposure. The outer folds estimate an observed audit-set upper bound, not a true oracle or independent confirmation result.",
    )
    body(doc, "Keywords", "Keywords: molecular property prediction; model selection; candidate-pool expansion; nested cross-validation; selection loss; effective diversity; uncertainty quantification; chemical boundaries", "Standardized index terms and removed promotional wording.", "The previous keyword list was broader than the narrowed methodological claim.", no_indent=True)

    add_heading(doc, "1 Introduction", 1)
    intro = [
        ("Molecular property prediction now spans fingerprint models, graph neural networks, chemical-language models and automated ensembles. Broader search spaces can improve the best performance observed on a finite benchmark, but they also increase the number of comparisons made against the same validation information. Consequently, a reported winner may reflect both genuine predictive advantage and the opportunity created by repeated selection [1-4].", "Establishes the problem before introducing the framework.", "The earlier opening moved too quickly from application context to governance terminology."),
        ("This problem is distinct from direct test-set leakage. Even when outer labels remain hidden during training, the validation ranking can be repeatedly consumed by choices among representations, learners, hyperparameters, calibration procedures and rescue rules. Classical work on model-selection bias shows that nested evaluation is required to separate selection from performance estimation, yet nested evaluation alone does not reveal how candidate count, candidate correlation and finite audit size shape the observed loss [3,4].", "Clarifies the estimand and prevents a leakage straw man.", "The earlier text conflated general leakage control with validation overuse."),
        ("Candidate count is also an incomplete description of search complexity. Thirty-two hyperparameter variants may behave as only a few effective candidates when their predictions or utilities are strongly correlated, whereas a smaller multiview pool can be more diverse. Reporting nominal K without effective diversity therefore obscures whether an observed expansion effect is driven by behavioural breadth, redundant tuning freedom or both.", "Introduces effective diversity as a necessary co-estimand.", "The previous manuscript alternated between nominal and effective pool size without a consistent definition."),
        ("Finite audit sets introduce a second complication. Selecting the maximum among noisy audit estimates creates mechanical winner optimism even when all candidates have equal true performance. We use the term observed audit upper bound for this finite-sample maximum. It is not a true oracle, because it is itself noisy, and it is not an independent confirmation estimate, because the same outer audit is used to characterize candidate performance.", "Defines terminology before results and removes oracle overclaiming.", "The earlier manuscript sometimes described the outer best candidate as a test oracle or independent confirmation."),
        ("Recent molecular-learning studies reinforce the need to evaluate strong baselines, distribution shift, uncertainty and chemical discontinuities under common splits [5-14,21-23,29,33-35]. However, these analyses answer different questions. A modern-baseline panel tests whether the audit conclusion is confined to simple models; conformal and uncertainty analyses test predictive reliability; activity cliffs and beyond-rule-of-five data define chemical boundaries. Combining them into one leaderboard would erase their distinct evidence roles.", "Positions secondary experiments without diluting the central claim.", "The previous version treated many benchmark and boundary analyses as parallel main claims."),
        ("Here we test a narrow question: under fixed validation information, how does candidate-pool expansion affect effective diversity, validation-audit ranking fidelity and raw-scale selection loss? We use a retrospective analysis lock, repeated nested scaffold splits, frozen candidate-composition controls, a finite-audit simulation and a shared-split multiview confirmation. Modern baselines and reliability analyses are reported as stress tests and boundary evidence. The resulting framework is an audit of model selection, not a new universal predictor.", "Ends the Introduction with a precise question, design and boundary.", "The earlier aim paragraph mixed framework advocacy, benchmarking and prospective claims."),
    ]
    for text, reason, issue in intro:
        body(doc, "1 Introduction", text, reason, issue)

    add_heading(doc, "2 Methods", 1)
    add_heading(doc, "2.1 Study design and claim hierarchy", 2)
    body(doc, "2.1", "This study was a retrospective frozen audit of completed molecular-property experiments. The analysis lock specified the nine primary endpoints, candidate order, K values, seeds, outer and inner split indices, selection rules, outcomes and exclusion logic before the present Journal of Cheminformatics reconstruction. It was not a prospective preregistration, and we therefore describe it as an analysis lock rather than as a preregistered protocol.", "States the temporal status of the protocol precisely.", "The prior manuscript risked implying prospective preregistration.")
    body(doc, "2.1", "The evidence hierarchy comprised four levels. The primary analysis was the nine-endpoint near-duplicate candidate expansion. Candidate-composition controls and finite-audit simulation tested the mechanism. A shared-split multiview pool provided representational confirmation. Four completed modern baselines, uncertainty analyses and chemical-boundary studies were stress tests or exploratory boundary evidence. No outer audit result was treated as an independent confirmation set.", "Separates confirmatory, mechanistic and exploratory evidence.", "The previous narrative gave secondary panels the same status as the primary estimand.")

    add_heading(doc, "2.2 Datasets and molecular standardization", 2)
    body(doc, "2.2", "The primary panel contained ESOL, FreeSolv, Lipophilicity, BBBP, BACE, ClinTox, Caco2 Wang, HIA Hou and P-gp Broccatelli. Regression endpoints were evaluated by root mean squared error (RMSE); classification endpoints were evaluated by ROC-AUC, with PR-AUC and calibration quantities used in secondary analyses. The public datasets served as repeated audit domains rather than as independent external validation cohorts [1,2].", "Defines tasks, outcomes and evidence role in one place.", "Dataset roles and metrics were previously distributed across several sections.")
    body(doc, "2.2", "SMILES were parsed with RDKit, cleaned and canonicalized under a fixed rule. The largest fragment was retained; charge normalization was applied when chemically valid. Classification duplicates were merged only when labels agreed, and conflicting label groups were excluded. Regression duplicates were averaged and their replicate counts retained. The nine analysis populations ranged from 578 molecules for HIA Hou to 4,200 for Lipophilicity; ClinTox retained 58 positive molecules among 1,376 standardized structures, making minority-class results intrinsically imprecise.", "Connects cleaning rules to endpoint-specific interpretation.", "The earlier manuscript reported aggregate cleaning totals that obscured ClinTox class scarcity.")
    build_tables(doc, data, 1)

    add_heading(doc, "2.3 Candidate pools and computational exposure", 2)
    body(doc, "2.3", "The controlled expansion pool contained 32 Morgan-512 candidates in a fixed registry order. For classification, candidates 1-4 were logistic regressions, 5-12 random-forest or extremely randomized-tree variants, 13-16 histogram or gradient boosting, 17-24 LightGBM, 25-28 XGBoost and 29-32 CatBoost. Regression used ridge or elastic-net candidates in positions 1-4 and the same tree-family structure thereafter. Prefixes of this registry defined K=4, 8, 16 and 32.", "Makes K reproducible and exposes the correlated nature of the pool.", "The prior text described the pool generically and made candidate expansion difficult to reconstruct.")
    body(doc, "2.3", "The split indices, per-candidate preprocessing, training protocol and tuning opportunity were fixed across K. Total compute exposure was not fixed: larger K required more model fits. Across the near-duplicate audit, 17,280 fits and 4,437.95 recorded fit-seconds were logged. These measurements exclude unlogged environment overhead and do not support a claim of equal total compute across pool sizes.", "Corrects the compute-budget interpretation.", "The earlier manuscript stated or implied a fixed compute budget despite exposure increasing with K.")
    body(doc, "2.3", "The modern-baseline panel included RDKit-RF, a two-layer GCN, frozen ChemBERTa embeddings with a linear probe and frozen MoLFormer embeddings with a linear probe on six MoleculeNet tasks. Chemprop/D-MPNN did not complete the full nine-endpoint 3 x 3 x 5 audit, and TabPFN was unavailable in the frozen runtime. Both are therefore labelled needs new analysis and excluded from completed-baseline claims.", "Documents completed and unavailable candidates without silent substitution.", "Previous versions risked presenting partial or unavailable baselines as completed.")
    build_tables(doc, data, 2)

    add_heading(doc, "2.4 Nested selection and outer auditing", 2)
    body(doc, "2.4", "For each endpoint and seed (11, 23, 37, 53 and 71), molecules were assigned to three outer scaffold folds. Within each outer training partition, three inner scaffold folds estimated candidate validation utility. The validation-best rule selected the candidate with the highest mean inner utility, with deterministic registry-order tie breaking. The selected candidate was then refitted on the full outer training partition and evaluated once on the held-out outer fold.", "Provides a complete, sequential selection protocol.", "The earlier method description scattered split, tie and refit rules across sections.")
    body(doc, "2.4", "The outer fold was an audit set for the frozen selection decision. It was not fed back into candidate eligibility, hyperparameters or selection rules. Nevertheless, because the same outer results were used to estimate the best observed candidate and selection loss, they do not constitute an independent confirmation cohort. Five seeds and three outer folds yielded 15 paired audit units per endpoint and K.", "Makes the no-feedback rule and remaining limitation explicit.", "The previous manuscript used confirmation language too broadly.")

    add_heading(doc, "2.5 Outcomes and statistical estimands", 2)
    body(doc, "2.5", "Utility was defined so that larger values were better: ROC-AUC for classification and negative RMSE for regression. For audit unit u and candidate set C_K, the observed audit upper bound was max_j in C_K U_audit(u,j). Selection loss was this maximum minus the audit utility of the candidate selected by inner validation. Thus, classification loss is measured in ROC-AUC units and regression loss in the endpoint's RMSE units.", "Defines all primary estimands with consistent direction.", "The previous paper alternated between regret, loss and utility without always stating units.")
    body(doc, "2.5", "Validation-ranking fidelity was quantified by Spearman correlation between inner-validation and outer-audit candidate utilities, Top-1 recovery, Top-3 recovery and mean reciprocal rank. Effective candidate rank was calculated from the entropy of the normalized eigenvalue spectrum of the candidate-utility correlation matrix. Median pairwise correlation was reported alongside effective rank because the two summaries capture complementary aspects of redundancy.", "Links ranking and diversity metrics to the mechanism.", "The earlier manuscript reported isolated ranking metrics without a unified estimand section.")
    body(doc, "2.5", "Raw-scale effects were primary. Range-normalized loss was used only for cross-endpoint sensitivity analyses because RMSE values for ESOL, FreeSolv, Lipophilicity and Caco2 are not commensurate physical quantities. No pooled raw RMSE effect was interpreted as a common chemical-unit change.", "Prevents invalid pooling across regression endpoints.", "Earlier drafts averaged raw regression losses across heterogeneous units without sufficient qualification.")

    add_heading(doc, "2.6 Candidate-composition controls", 2)
    body(doc, "2.6", "To distinguish prefix composition from nominal K, we analysed three frozen controls: random registry orders, random subsets and family-balanced subsets. Each mode used 100 predetermined subset or order seeds for every endpoint and K. The endpoint, rather than each subset seed, remained the primary inferential cluster; subset seeds characterized sensitivity to composition.", "Explains both the controls and their statistical unit.", "The prior manuscript could be read as treating thousands of subsets as independent datasets.")
    body(doc, "2.6", "We also compared the validation-best rule with a single fixed candidate, one-standard-error stable selection, a risk-adjusted rule and the expected utility of uniform random selection. These alternatives were frozen before outer evaluation and were used to determine whether any expansion effect depended on one selector.", "Separates selector sensitivity from candidate-composition sensitivity.", "Governance-rule and pool-composition ablations were previously intermingled.")

    add_heading(doc, "2.7 Finite-audit winner-optimism simulation", 2)
    body(doc, "2.7", "A Monte Carlo simulation isolated the optimism created by maximizing noisy audit estimates. Candidate true utilities were equal, while audit noise followed a multivariate Gaussian distribution with pairwise correlations of 0, 0.5 or 0.9 and effective audit sample sizes of 25, 50, 100 or 200. Candidate counts were 4, 8, 16 and 32, with 30,000 replicates per configuration.", "Creates a transparent null mechanism for winner optimism.", "The earlier text used simulation results without fully specifying the data-generating process.")
    body(doc, "2.7", "Winner optimism was the difference between the largest observed audit utility and the common true utility, expressed in noise-standard-deviation units. The simulation does not estimate bias in any empirical endpoint; it demonstrates the mechanical consequence of maximizing finite, correlated estimates.", "Prevents overgeneralization from the simulation.", "The prior wording risked treating simulated optimism as an empirical correction factor.")

    add_heading(doc, "2.8 Multiview and modern-baseline stress tests", 2)
    body(doc, "2.8", "The multiview experiment crossed four representations (Morgan-512, MACCS, RDKit2D and their concatenation) with three learners (linear, random forest and LightGBM), yielding 12 candidates. All candidates used the same nine endpoints, seeds and 3 x 3 nested scaffold splits. Realized gain compared validation-best selection from the full pool with validation-best selection from the Morgan-only pool in each paired outer unit.", "Defines the confirmation contrast without relying on post hoc oracle comparisons.", "The previous manuscript mixed attainable and realized multiview gains.")
    body(doc, "2.8", "The six-task modern-baseline panel used the same outer and inner structure. RDKit-RF used Morgan fingerprints and 120 trees. The GCN used two 32-unit graph-convolution layers, global mean pooling and five fixed epochs. ChemBERTa and MoLFormer were frozen encoders followed by logistic or ridge probes. Candidate errors were exported per molecule and pairwise Jaccard overlap was computed on candidate-specific error sets.", "Makes the strong-baseline stress test reproducible and enables error-overlap interpretation.", "The previous description lacked a concise common specification for the four completed candidates.")

    add_heading(doc, "2.9 Reliability and chemical-boundary analyses", 2)
    body(doc, "2.9", "Classification reliability was evaluated with split conformal, label-conditional conformal and Mondrian label-and-similarity conformal prediction at 80%, 90% and 95% targets. Regression used conformalized quantile regression (CQR). Ensemble uncertainty was assessed by uncertainty-error Spearman correlation, risk-coverage behaviour and enrichment of the top 10% highest-error samples [26-28,33,36].", "Organizes complementary uncertainty methods by outcome type.", "The earlier reliability section listed methods without a shared evaluation framework.")
    body(doc, "2.9", "Chemical-boundary analyses stratified predictions by maximum train-set Morgan Tanimoto, scaffold novelty, extreme labels, activity-cliff pairs and beyond-rule-of-five perimeter status. ClinTox minority recall and false-negative cost were retained as negative results. MoleculeACE and bRo5 analyses were interpreted as public chemical-boundary stress tests, not prospective external validation [29,30].", "States the purpose and limits of each boundary analysis.", "The previous version sometimes treated public boundary panels as external confirmation.")

    add_heading(doc, "2.10 Statistical analysis", 2)
    body(doc, "2.10", "Primary K=32 versus K=4 effects were paired within endpoint, seed and outer fold. Means and 95% percentile bootstrap intervals used resampling of the 15 paired outer units within each endpoint. Endpoint-level direction counts and leave-one-endpoint-out sensitivity were used to assess consistency without treating outer folds as independent datasets. Exact sign tests and Holm adjustment were used where multiple endpoint-level hypotheses were explicitly evaluated [31].", "Aligns uncertainty estimates with the experimental unit.", "Earlier drafts risked pseudoreplication by emphasizing 135 outer units as independent domains.")
    body(doc, "2.10", "Secondary cross-endpoint summaries used endpoint means as the analysis unit. Missing candidates were not imputed. All random procedures used recorded seeds, and source CSV files were hashed in the revision audit. Statistical intervals describe repeat and fold variability under the frozen design; they do not quantify uncertainty for deployment in an unseen medicinal-chemistry programme.", "Defines missingness, reproducibility and inferential scope.", "The previous statistical language occasionally implied broader population inference than the design supports.")
    add_figure(doc, FIG / "Figure_1_frozen_audit_workflow.png", 16.0, "Figure 1. Frozen audit workflow and claim hierarchy. (A) Candidate registration preceded nested selection and outer auditing, with no outer-label feedback. (B) Registry prefixes generated K=4, 8, 16 and 32 pools. (C) Each endpoint used five seeds, three outer scaffold folds and three inner folds. (D) Primary estimands were the observed audit upper bound, raw-scale selection loss, ranking fidelity and effective candidate rank. (E) Evidence was ordered from the nine-endpoint primary audit through multiview confirmation to stress tests and chemical boundaries. (F) The observed outer best is neither a true oracle nor an independent confirmation estimate.")

    add_heading(doc, "3 Results", 1)
    add_heading(doc, "3.1 Nominal candidate count exceeded effective diversity", 2)
    body(doc, "3.1", "Nominal pool size substantially overstated behavioural diversity. Mean effective rank increased from 1.58 at K=4 to 2.07 at K=16 and was 2.00 at K=32. Thus, adding 28 registry entries beyond K=4 produced an average increase of only 0.42 effective candidates. The result identifies the primary pool as a controlled near-duplicate search-space experiment rather than a broad comparison of independent model classes.", "Reports the diversity result before performance consequences.", "The earlier results used inconsistent effective-rank values across versions.")
    body(doc, "3.1", "Candidate correlation was non-monotonic at small K but high after the boosting families entered the registry: mean endpoint-level median correlation was 0.618 at K=8, 0.845 at K=16 and 0.888 at K=32. Effective rank and pairwise correlation therefore converged on the same interpretation: most nominal additions supplied correlated selection opportunities rather than proportionate new predictive behaviour.", "Explains the apparent K=8 irregularity and triangulates diversity metrics.", "The previous narrative claimed monotonic correlation despite the observed K=8 decrease.")
    build_tables(doc, data, 3)

    add_heading(doc, "3.2 Candidate expansion increased selection loss in the studied endpoints", 2)
    body(doc, "3.2", "Validation-audit ranking fidelity weakened as K increased. Mean Spearman correlation declined from 0.760 at K=4 to 0.629 at K=32; Top-1 recovery declined from 0.778 to 0.081, and mean reciprocal rank declined from 0.869 to 0.264. The observed audit upper bound could therefore increase while the validation-selected candidate captured a smaller fraction of that opportunity.", "Connects ranking distortion to selection loss.", "The earlier results presented ranking metrics as a separate benchmark rather than as the mechanism of loss.")
    body(doc, "3.2", "For validation-best selection, K=32 minus K=4 raw selection loss was positive in eight of nine endpoints. Classification increases were 0.0053 ROC-AUC for BACE, 0.0035 for BBBP, 0.0203 for ClinTox and 0.0195 for HIA Hou; P-gp Broccatelli was the exception (-0.0003; 95% CI -0.0043 to 0.0035). Regression increases were 0.0940 RMSE for ESOL, 0.0743 for FreeSolv, 0.0077 for Lipophilicity and 0.0328 on the Caco2 dataset-provided log-permeability scale.", "Reports endpoint-specific raw effects and preserves the negative result.", "The prior manuscript overemphasized pooled means and underreported the P-gp inconsistency.")
    body(doc, "3.2", "Averaged only within task type, classification selection loss increased from 0.0042 at K=4 to 0.0138 at K=32. Regression means increased from 0.0178 to 0.0700 RMSE, but this numerical average is descriptive because the four endpoints use different target scales. The endpoint-specific estimates in Table 4 are the primary regression results.", "Allows concise task-type orientation while preserving unit boundaries.", "The earlier draft could be read as interpreting pooled RMSE as one chemical quantity.")
    build_tables(doc, data, 4)
    add_figure(doc, FIG / "Figure_2_candidate_expansion_selection_loss.png", 16.0, "Figure 2. Candidate expansion, effective diversity and selection loss. (A) Effective candidate rank remained near two despite nominal expansion to 32. (B) Pairwise utility correlation was high at K=16 and K=32. (C) Validation-audit Spearman correlation declined with K. (D) K=32 minus K=4 effects are shown in endpoint-specific raw units; blue denotes ROC-AUC loss and orange denotes RMSE loss, which are not pooled. Error bars are 95% bootstrap intervals across paired outer units. (E) Random-order, random-subset and family-balanced controls preserved the increasing normalized-loss trend. (F) Descriptive task-type means increased for classification and regression; the separate axes indicate non-commensurate units.")

    add_heading(doc, "3.3 Candidate-composition controls supported the expansion trend", 2)
    body(doc, "3.3", "The expansion pattern was not confined to the original registry prefix. Mean normalized selection loss for random orders increased from 0.089 at K=4 to 0.214 at K=32; random subsets increased from 0.090 to 0.214, and family-balanced subsets increased from 0.087 to 0.214. Differences among composition modes were smaller than the change across K.", "Demonstrates robustness to candidate composition.", "The earlier text mentioned composition controls without integrating their quantitative result.")
    body(doc, "3.3", "At K=32, validation-best remained better than uniform random selection: mean classification loss was 0.0138 versus 0.0295, and mean regression loss was 0.0700 versus 0.1782. Candidate expansion therefore did not make validation ranking uninformative. Instead, it reduced the fraction of the observed audit opportunity recovered by the frozen selector.", "Adds a calibrated negative control and avoids an all-or-none conclusion.", "The previous narrative could be misread as claiming validation selection became equivalent to chance.")

    add_heading(doc, "3.4 Finite audit sets induced mechanical winner optimism", 2)
    body(doc, "3.4", "Under equal true candidate utility, the maximum observed audit estimate became more optimistic as K increased and as effective audit size decreased. With pairwise correlation 0.9 and n_eff=50, mean optimism increased from 0.046 SD units at K=4 to 0.092 at K=32. At n_eff=25, the corresponding K=32 optimism was 0.132 SD units.", "Quantifies the finite-audit mechanism on a unitless scale.", "The earlier discussion invoked winner optimism without a concise null result.")
    body(doc, "3.4", "This simulation explains why the observed audit best must not be called a true oracle. It does not account for all empirical selection loss, because real candidates differ in true utility and their correlation structure is endpoint-specific. Its role is narrower: even a perfectly fair audit produces an upward-biased maximum when many finite estimates are compared.", "Keeps the simulation as mechanism rather than correction.", "The prior version risked extending the equal-truth model beyond its assumptions.")

    add_heading(doc, "3.5 Multiview candidates produced auditable representational gains", 2)
    body(doc, "3.5", "The 12-candidate multiview pool completed 6,480 fits across nine endpoints and 135 paired outer units. Validation-best selection from the full pool achieved a mean normalized realized gain of 0.343 over Morgan-only validation-best selection (95% CI 0.210-0.483), with positive endpoint means in all nine endpoints. Concatenated multiview candidates added a further normalized gain of 0.035 (95% CI 0.017-0.053) over the separate-view pool.", "Separates realized selection gain from the observed upper bound.", "The earlier manuscript sometimes mixed oracle and validation-selected multiview effects.")
    body(doc, "3.5", "Raw effects were heterogeneous. Mean realized gains were 0.0873 ROC-AUC for ClinTox, 0.0809 for HIA Hou, 0.0249 for BBBP, 0.0227 for P-gp and 0.0039 for BACE. Regression RMSE reductions were 1.1949 for FreeSolv, 0.9008 for ESOL, 0.2401 for Lipophilicity and 0.1058 for Caco2. BACE included only seven positive paired units and its confidence interval crossed zero, preserving an endpoint-specific null despite the positive cross-endpoint mean.", "Reports raw multiview gains and the BACE limitation.", "The previous results emphasized the all-endpoint normalized mean and hid weak BACE realization.")
    body(doc, "3.5", "Among 135 validation-best selections, the concatenated representation was selected 84 times, RDKit2D 44 times, MACCS four times and Morgan three times. The gain was therefore not attributable to one universally superior learner; it arose from representation choice under a common frozen selection rule.", "Links gain to selected representation rather than model branding.", "The earlier narrative risked presenting concatenation as universally dominant.")
    add_figure(doc, FIG / "Figure_3_ranking_winner_optimism_multiview.png", 16.0, "Figure 3. Ranking controls, finite-audit winner optimism and multiview confirmation. (A) Observed chance-adjusted Top-3 recovery exceeded the endpoint-wise permutation null at every K. (B) A signal-recovery control increased chance-adjusted recovery monotonically with injected validation-audit correlation. (C) In the equal-truth simulation, winner optimism increased with K and decreased with effective audit size. (D) Realized multiview gains are reported in endpoint-specific raw units. (E) Concatenated multiview and RDKit2D representations dominated the frozen selections. (F) Normalized attainable, realized and concatenation gains are summarized across endpoints; horizontal ranges show endpoint minima and maxima, not confidence intervals.")

    add_heading(doc, "3.6 Modern baselines showed correlated and partially complementary errors", 2)
    body(doc, "3.6", "RDKit-RF, GCN, ChemBERTa and MoLFormer completed the shared six-task audit. Across 90 endpoint-seed-outer units, the frozen selector recovered the observed audit best in 0.878 of units, with mean range-normalized selection loss 0.0048 and mean validation-audit Spearman 0.840. RDKit-RF was selected in all six endpoint summaries, while the relative performance of the three learned representations varied by endpoint.", "Presents the completed stress test without claiming a universal model ranking.", "The earlier strong-baseline section mixed completed six-task and partial three-task evidence.")
    body(doc, "3.6", "Pairwise error-overlap Jaccard values averaged 0.215 and ranged from 0.168 for GCN versus RDKit-RF to 0.296 for ChemBERTa versus RDKit-RF. The non-zero overlap confirms correlated failures, whereas the range indicates partial complementarity. Candidate count alone therefore overstates independent opportunity in both the lightweight and modern-baseline pools.", "Uses exported per-sample predictions to connect error structure to effective diversity.", "Earlier versions could only describe error overlap as pending or qualitative.")
    body(doc, "3.6", "These results do not complete a modern-baseline confirmation across all nine primary endpoints. Chemprop/D-MPNN and TabPFN remain unavailable under the full frozen design, and the modern panel contains six rather than nine tasks. The panel is therefore a stress test of error correlation and selector behaviour, not a replacement for the primary expansion experiment.", "Preserves the explicit missing-analysis boundary.", "The prior manuscript risked overstating baseline completeness.")

    add_heading(doc, "3.7 Reliability deteriorated near chemical boundaries", 2)
    body(doc, "3.7", "At 90% target coverage, mean minority-class coverage across classification endpoints was 0.788 for split conformal, 0.887 for label-conditional conformal and 0.891 for Mondrian label-and-similarity conformal prediction. Conditional methods therefore reduced the aggregate minority undercoverage, but neither achieved 0.90 exactly. For ClinTox, only 58 positives remained after standardization, and candidate-specific false-negative rates remained substantial; this endpoint is retained as a negative reliability result.", "Reports both improvement and residual undercoverage.", "The earlier reliability narrative emphasized coverage gains without equally visible negative evidence.")
    body(doc, "3.7", "CQR achieved mean 90% coverage of 0.882 across ESOL, FreeSolv and Lipophilicity, with a raw mean interval width of 7.25. The width average spans incompatible endpoint units and is descriptive only. Ensemble uncertainty enriched the top 10% highest-error samples by 1.54-fold on average and had mean uncertainty-error Spearman 0.321, indicating useful but incomplete error ranking.", "Preserves endpoint units and calibrates the practical value of uncertainty.", "The previous version reported a pooled CQR width without a strong unit warning.")
    body(doc, "3.7", "Classification ROC-AUC averaged 0.803 among molecules with maximum train-set Tanimoto below 0.5 and 0.924 above 0.7. Scaffold novelty, extreme labels, activity-cliff pairs and bRo5 perimeter cases also concentrated large errors or misclassifications. These analyses support a chemical-boundary warning rather than a universal uncertainty guarantee.", "Integrates similarity, scaffold and discontinuity evidence into one bounded conclusion.", "The earlier manuscript listed failure categories without a unifying interpretation.")
    build_tables(doc, data, 5)
    add_figure(doc, FIG / "Figure_4_modern_baselines_reliability_boundaries.png", 16.0, "Figure 4. Modern-baseline error structure, predictive reliability and chemical boundaries. (A) Candidate utilities were normalized within each endpoint for visualization; the heat map is not a cross-endpoint performance average. (B) Pairwise Jaccard overlap was calculated from exported per-sample errors. (C) Minority-class conformal coverage at the 90% target improved under label-conditional and Mondrian methods but remained slightly below target. (D) CQR coverage varied across regression endpoints. (E) ROC-AUC and expected calibration error were stratified by maximum train-set Tanimoto. (F) Systematic failures included novel scaffolds, extreme labels, low similarity, bRo5 perimeter cases, activity cliffs and classification false negatives; error scores are category-specific and should not be compared as a common physical unit.")

    add_heading(doc, "4 Discussion", 1)
    add_heading(doc, "4.1 Candidate count is not model capacity", 2)
    body(doc, "4.1", "The central result is not that larger model searches are intrinsically harmful. Rather, nominal expansion from four to 32 highly correlated candidates produced little additional effective diversity while creating many more opportunities for validation-specific ranking. The distinction matters: a pool can be computationally large, statistically redundant and still yield a higher observed audit maximum.", "Interprets K through effective diversity rather than a simplistic larger-is-worse claim.", "The previous discussion occasionally generalized from one registry to all model expansion.")
    body(doc, "4.1", "Accordingly, candidate-pool reports should include at least four quantities: nominal K, effective candidate rank, total computational exposure and the number of distinct representation or learner families. These quantities answer different questions and cannot be replaced by a single model count.", "Turns the result into a concrete reporting recommendation.", "The earlier framework discussion remained abstract and process-heavy.")

    add_heading(doc, "4.2 Validation-ranking distortion is the operative failure mode", 2)
    body(doc, "4.2", "Selection loss increased because validation ranking recovered the audit winner less reliably as K grew, not because validation became uninformative. Validation-best consistently outperformed random selection, and permutation and signal-recovery controls behaved as expected. The practical failure mode is therefore ranking distortion under limited information, rather than complete absence of predictive signal.", "Uses positive and negative controls to sharpen the mechanism.", "The earlier wording could imply that validation selection broadly failed.")
    body(doc, "4.2", "The finite-audit simulation adds a second warning: even the best observed outer result is optimistic relative to equal truth when several finite estimates are maximized. Nested cross-validation protects the selected model estimate from direct reuse of inner validation, but it does not transform the maximum outer estimate into a noiseless oracle. An additional untouched cohort would be required for independent confirmation.", "Clarifies what nested CV can and cannot solve.", "The prior manuscript overstated the status of the outer maximum.")

    add_heading(doc, "4.3 Heterogeneous candidates can produce realizable gains", 2)
    body(doc, "4.3", "The multiview experiment shows why the answer is not to freeze a permanently small pool. A genuinely heterogeneous 12-candidate pool produced positive realized gains under the same nested selection protocol, and the concatenated view was selected frequently. Candidate expansion is defensible when added candidates increase effective diversity and when the realized gain is measured against a frozen smaller-pool selector.", "Balances the risk result with evidence for beneficial expansion.", "The previous discussion framed governance defensively and underplayed realizable representation gains.")
    body(doc, "4.3", "The modern-baseline panel reached a compatible conclusion from another direction. Error overlap was neither zero nor complete, so learned representations contributed partially distinct failures but did not create independent draws from a performance distribution. Strong baselines should therefore be judged under common splits, candidate-level prediction export and explicit compute accounting, rather than by isolated literature values.", "Connects modern baselines to the same correlation mechanism.", "The earlier comparison relied too heavily on leaderboard-style endpoint scores.")

    add_heading(doc, "4.4 Reliability is conditional on chemical support", 2)
    body(doc, "4.4", "Conditional conformal methods improved minority coverage, and ensemble uncertainty enriched high-error samples, but neither result removes the need for chemical-support checks. Low Tanimoto similarity, scaffold novelty, activity cliffs and extreme labels marked different failure mechanisms. A deployment-facing prediction should therefore be accompanied by class-conditional coverage, nearest-neighbour similarity, scaffold status and an explicit review flag.", "Translates reliability results into auditable outputs.", "The prior text treated conformal coverage as a largely self-contained reliability claim.")
    body(doc, "4.4", "ClinTox illustrates the boundary. Its cleaned positive rate was 4.2%, and false-negative behaviour varied substantially across candidates. Improved prediction-set coverage does not guarantee an acceptable thresholded toxicity screen, particularly when false-negative costs dominate. ClinTox is therefore reported as a negative result rather than as evidence of a production-ready toxicology filter.", "Preserves a clinically relevant negative result and avoids threshold overclaiming.", "Earlier performance summaries foregrounded ROC-AUC and could obscure minority-class failure.")

    add_heading(doc, "4.5 Limitations and practical implications", 2)
    body(doc, "4.5", "The study has five principal limitations. First, the primary registry is deliberately near-duplicate and does not represent the full diversity of current molecular models. Second, only nine endpoints support the primary cross-endpoint inference. Third, the outer audit is not independent confirmation. Fourth, compute logs are fit-level CPU records and do not reconstruct GPU utilization or all overhead. Fifth, Chemprop/D-MPNN and TabPFN require new analysis under the complete frozen protocol.", "Consolidates limitations that were previously scattered.", "The earlier manuscript buried several major evidence gaps in separate sections.")
    body(doc, "4.5", "The immediate practical implication is procedural. Before adding a candidate, investigators should register its representation, learner, tuning opportunity, expected compute and required prediction export; after selection, they should report raw-scale loss, effective diversity, error overlap and boundary failures. This protocol does not guarantee improved prediction, but it makes the cost and evidential value of candidate expansion visible.", "Ends the Discussion with an actionable and appropriately limited contribution.", "The previous conclusion drifted toward broad governance claims without specifying minimum outputs.")

    add_heading(doc, "5 Conclusions", 1)
    body(doc, "5 Conclusions", "Across nine frozen molecular-property audits, expanding a highly correlated candidate registry increased endpoint-specific selection loss in eight endpoints and reduced validation-audit ranking fidelity. The nominal increase to 32 candidates corresponded to only about two effective candidates. Candidate count should therefore be interpreted jointly with effective diversity, compute exposure and raw-scale selection loss.", "Provides a direct answer to the primary research question.", "The previous conclusion repeated framework components rather than the central empirical finding.")
    body(doc, "5 Conclusions", "Shared-split multiview candidates demonstrated that heterogeneous expansion can deliver realizable gains, while modern-baseline, uncertainty and chemical-boundary analyses defined where those gains remain fragile. The present outer audit supplies an observed upper bound, not a true oracle or independent confirmation. Full-audit Chemprop/D-MPNN and TabPFN comparisons, an untouched external cohort and third-party reproduction remain future work.", "Closes with the positive result and explicit remaining work.", "The earlier final paragraph implied a more complete modern-model and external-validation evidence base than was available.")

    add_heading(doc, "Supplementary information map", 1)
    body(doc, "Supplementary map", "The supplementary files should be deposited with machine-readable source data and the frozen code archive. Items marked needs new analysis must not be represented as completed in the submission package.", "Creates a transparent submission map and preserves missing-analysis status.", "The earlier supplementary plan was dispersed and did not distinguish completed from unavailable analyses.")
    add_table_caption(doc, "Supplementary Tables S1-S10: titles and required fields.")
    supp_tables = [
        ["Table S1", "Dataset provenance and cleaning audit", "endpoint; source; version; raw n; analysis n; invalid SMILES; consistent duplicates; conflicts"],
        ["Table S2", "Complete 32-candidate registry", "candidate_id; order; family; representation; hyperparameters; eligibility; status"],
        ["Table S3", "Frozen split manifest", "endpoint; seed; outer fold; inner fold; split hash; scaffold counts"],
        ["Table S4", "Fold-level candidate utilities", "endpoint; seed; fold; K; candidate; inner utility; audit utility; runtime"],
        ["Table S5", "Candidate-composition controls", "mode; subset seed; endpoint; K; normalized loss; ranking metrics"],
        ["Table S6", "Finite-audit simulation grid", "truth scenario; n_eff; correlation; K; replicates; optimism; interval"],
        ["Table S7", "Multiview paired effects", "endpoint; seed; fold; comparison; raw gain; normalized gain; selected representation"],
        ["Table S8", "Modern-baseline predictions and overlap", "endpoint; candidate pair; unit; Jaccard overlap; prediction-file hash"],
        ["Table S9", "Reliability and boundary strata", "endpoint; method; target; coverage; width/set size; similarity; scaffold status"],
        ["Table S10", "Unavailable and failed analyses", "candidate or analysis; scope; failure reason; evidence impact; required next action"],
    ]
    add_table(doc, ["Item", "Title", "Suggested fields"], supp_tables)
    add_table_caption(doc, "Supplementary Figures S1-S7: planned content.")
    supp_figs = [
        ["Figure S1", "Endpoint cleaning and class-balance flow"],
        ["Figure S2", "Candidate correlation matrices and eigenvalue spectra by K"],
        ["Figure S3", "Fold-level raw selection-loss distributions"],
        ["Figure S4", "Random-order, random-subset and family-balanced control distributions"],
        ["Figure S5", "Finite-audit optimism across all correlations and effective sample sizes"],
        ["Figure S6", "Multiview endpoint-by-candidate selection map and paired raw effects"],
        ["Figure S7", "Conformal, uncertainty, ClinTox negative and chemical-boundary diagnostics"],
    ]
    add_table(doc, ["Item", "Content"], supp_figs)

    add_heading(doc, "Declarations", 1)
    for heading, text in [
        ("Ethics approval and consent to participate", "Not applicable; all analyses used public molecular datasets."),
        ("Consent for publication", "Not applicable."),
        ("Availability of data and materials", "Machine-readable source tables, figure files and the frozen local code package accompany this revision. A public repository URL and archival DOI need to be added before submission. Third-party cold-start reproduction was outside the scope of this revision."),
        ("Competing interests", "Author confirmation required before submission."),
        ("Funding", "Author confirmation required before submission."),
        ("Authors' contributions", "Author contribution roles must be completed before submission."),
        ("Acknowledgements", "To be completed by the authors, if applicable."),
    ]:
        add_heading(doc, heading, 2)
        add_body(doc, text, no_indent=True)

    add_heading(doc, "References", 1)
    references = [
        "1. Wu Z, Ramsundar B, Feinberg EN, et al. MoleculeNet: a benchmark for molecular machine learning. Chem Sci. 2018;9:513-530. doi:10.1039/C7SC02664A.",
        "2. Huang K, Fu T, Gao W, et al. Therapeutics Data Commons: machine learning datasets and tasks for drug discovery and development. NeurIPS Datasets and Benchmarks Track. 2021. arXiv:2102.09548.",
        "3. Cawley GC, Talbot NLC. On over-fitting in model selection and subsequent selection bias in performance evaluation. J Mach Learn Res. 2010;11:2079-2107.",
        "4. Varma S, Simon R. Bias in error estimation when using cross-validation for model selection. BMC Bioinformatics. 2006;7:91. doi:10.1186/1471-2105-7-91.",
        "5. Zhao D, Zhu Y, Wu Z, et al. Revisiting ADMET prediction reliability under real-world challenges in the foundation model era. J Cheminform. 2026. doi:10.1186/s13321-026-01217-2.",
        "6. Zhang L, Zeng Y, Qi Y, et al. DCPM-ADMET: fusion of dual-component pre-trained model and molecular fingerprints to enhance drug ADMET properties prediction. J Cheminform. 2026. doi:10.1186/s13321-026-01244-z.",
        "7. Jang Y, Lee J, Jeong K, Kim J. Multimodal graph fusion with statistically guided parsimonious descriptor selection for molecular property prediction. J Cheminform. 2026;18:18. doi:10.1186/s13321-025-01140-y.",
        "8. Zhang Y, Liu W, Zhao H, et al. MolGramTreeNet: a multimodal molecular property prediction model via grammar tree-constrained molecular representation. iScience. 2026;29:114928. doi:10.1016/j.isci.2026.114928.",
        "9. Wen X, Liu H, Long W, Wei S, Zhu R. Consistent semantic representation learning for out-of-distribution molecular property prediction. Brief Bioinform. 2025;26:bbaf147. doi:10.1093/bib/bbaf147.",
        "10. Yin T, Gao P, Panapitiya G, Saldanha EG. Out-of-distribution evaluation of active learning pipelines for molecular property prediction. RSC Adv. 2026;16:5281-5295. doi:10.1039/D5RA08055J.",
        "11. Uchibori Y, Kaneko H. Generation of molecules near the applicability domain boundaries of property prediction models. J Chem Inf Model. 2026;66:6866-6879. doi:10.1021/acs.jcim.5c03220.",
        "12. Kim JY, Vlachos DG. Distance-aware molecular property prediction in nonlinear structure-property space. J Chem Inf Model. 2025;65:6744-6756. doi:10.1021/acs.jcim.5c01037.",
        "13. Tang H, Yue T, Li Y. Assessing uncertainty in machine learning for polymer property prediction: a benchmark study. J Chem Inf Model. 2025;65:6585-6598. doi:10.1021/acs.jcim.5c00550.",
        "14. Fralish Z, Reker D. Pairwise learning for molecular property prediction and optimization. Front Drug Discov. 2026;6:1859068. doi:10.3389/fddsv.2026.1859068.",
        "15. Landrum G. RDKit: open-source cheminformatics software. https://www.rdkit.org/.",
        "16. Rogers D, Hahn M. Extended-connectivity fingerprints. J Chem Inf Model. 2010;50:742-754. doi:10.1021/ci100050t.",
        "17. Breiman L. Random forests. Mach Learn. 2001;45:5-32. doi:10.1023/A:1010933404324.",
        "18. Ke G, Meng Q, Finley T, et al. LightGBM: a highly efficient gradient boosting decision tree. Adv Neural Inf Process Syst. 2017;30.",
        "19. Chen T, Guestrin C. XGBoost: a scalable tree boosting system. Proc ACM SIGKDD. 2016:785-794. doi:10.1145/2939672.2939785.",
        "20. Prokhorenkova L, Gusev G, Vorobev A, Dorogush AV, Gulin A. CatBoost: unbiased boosting with categorical features. Adv Neural Inf Process Syst. 2018;31:6638-6648.",
        "21. Yang K, Swanson K, Jin W, et al. Analyzing learned molecular representations for property prediction. J Chem Inf Model. 2019;59:3370-3388. doi:10.1021/acs.jcim.9b00237.",
        "22. Chithrananda S, Grand G, Ramsundar B. ChemBERTa: large-scale self-supervised pretraining for molecular property prediction. arXiv:2010.09885. 2020.",
        "23. Ross J, Belgodere B, Chenthamarakshan V, et al. Large-scale chemical language representations capture molecular structure and properties. Nat Mach Intell. 2022;4:1256-1264. doi:10.1038/s42256-022-00580-7.",
        "24. Erickson N, Mueller J, Shirkov A, et al. AutoGluon-Tabular: robust and accurate AutoML for structured data. arXiv:2003.06505. 2020.",
        "25. Tropsha A. Best practices for QSAR model development, validation, and exploitation. Mol Inform. 2010;29:476-488. doi:10.1002/minf.201000061.",
        "26. Vovk V, Gammerman A, Shafer G. Algorithmic Learning in a Random World. New York: Springer; 2005.",
        "27. Shafer G, Vovk V. A tutorial on conformal prediction. J Mach Learn Res. 2008;9:371-421.",
        "28. Guo C, Pleiss G, Sun Y, Weinberger KQ. On calibration of modern neural networks. Proc ICML. 2017;70:1321-1330.",
        "29. van Tilborg D, Alenicheva A, Grisoni F. Exposing the limitations of molecular machine learning with activity cliffs. J Chem Inf Model. 2022;62:5938-5951. doi:10.1021/acs.jcim.2c01073.",
        "30. Sheridan RP. Time-split cross-validation as a method for estimating prospective prediction performance. J Chem Inf Model. 2013;53:783-790. doi:10.1021/ci400084k.",
        "31. Demsar J. Statistical comparisons of classifiers over multiple data sets. J Mach Learn Res. 2006;7:1-30.",
        "32. Hoyt CT, Zdrazil B, Guha R, et al. Improving reproducibility and reusability in the Journal of Cheminformatics. J Cheminform. 2023;15:62. doi:10.1186/s13321-023-00730-y.",
        "33. Parrondo-Pizarro R, Lanini J, Rodriguez-Perez R. Uncertainty quantification in molecular machine learning for property predictions under data shifts. J Chem Inf Model. 2026;66:923-935. doi:10.1021/acs.jcim.5c02381.",
        "34. Deng J, Yang Z, Wang H, Ojima I, Samaras D, Wang F. A systematic study of key elements underlying molecular property prediction. Nat Commun. 2023;14:6395. doi:10.1038/s41467-023-41948-6.",
        "35. Li Z, Chen X, Wen H, et al. A systematic survey and benchmark of deep learning for molecular property prediction in the foundation model era. arXiv:2604.16586. 2026.",
        "36. Romano Y, Patterson E, Candes EJ. Conformalized quantile regression. Adv Neural Inf Process Syst. 2019;32.",
    ]
    for ref in references:
        add_reference(doc, ref)
    return doc


def write_changelog() -> None:
    lines = [
        "# 小论文-19逐段修改说明",
        "",
        "本表以本轮完整重构前的主要问题为对照，不将未改变的句子伪装为逐字修订。每行给出重写后的完整段落及其学术理由。",
        "",
        "| Section | Original issue | Revised paragraph | Reason |",
        "|---|---|---|---|",
    ]
    for row in PARAGRAPH_LOG:
        vals = [row["section"], row["original_issue"], row["revised_paragraph"], row["reason"]]
        vals = [v.replace("|", "\\|").replace("\n", " ") for v in vals]
        lines.append("| " + " | ".join(vals) + " |")
    CHANGELOG.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_reorganization() -> None:
    text = """# 小论文-19图表重组与投稿核查

## 主图

1. Figure 1: frozen workflow, split architecture, estimands, claim hierarchy and interpretation boundary.
2. Figure 2: nominal/effective diversity, candidate correlation, ranking fidelity, endpoint raw effects and composition controls.
3. Figure 3: permutation and signal controls, finite-audit winner optimism and multiview confirmation.
4. Figure 4: modern-baseline utility/error overlap, conditional conformal, CQR, similarity reliability and failure modes.

每幅主图均由源 CSV 自动生成，并导出 PNG、SVG 和 PDF；不是截图拼接。

## 主表

1. Table 1: datasets, sample sizes and evidence roles.
2. Table 2: candidate pools, selection rules and computational exposure.
3. Table 3: candidate scale, effective diversity and ranking fidelity.
4. Table 4: endpoint-specific raw-scale selection loss.
5. Table 5: secondary stress tests and interpretation boundaries.

全部 Word 表格仅保留表顶线、表头下横线和表底线，无竖线，符合三线表规则。

## Needs new analysis

- Chemprop/D-MPNN full nine-endpoint 3 x 3 x 5 audit.
- TabPFN full frozen audit after runtime resolution.
- Untouched independent external confirmation cohort.
- Prospective time-split deployment audit.
- Public archival DOI and third-party cold-start reproduction.

## Citation and language audit

- Primary methodological and model references retain DOI or stable bibliographic identifiers.
- The 2026 UQ-under-data-shifts citation was checked against the ACS record: J Chem Inf Model. 2026;66:923-935; doi:10.1021/acs.jcim.5c02381.
- Outer-fold language was standardized to observed audit upper bound; true oracle and independent confirmation claims were removed.
- Cross-endpoint raw RMSE pooling is explicitly prohibited in interpretation.
"""
    REORG.write_text(text, encoding="utf-8")


def audit_docx(path: Path) -> dict[str, object]:
    doc = Document(path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    tables = []
    for i, table in enumerate(doc.tables, 1):
        vertical = 0
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                borders = tc_pr.first_child_found_in("w:tcBorders")
                if borders is not None:
                    for edge in ("left", "right", "insideV"):
                        node = borders.find(qn(f"w:{edge}"))
                        if node is not None and node.get(qn("w:val")) not in {None, "nil", "none"}:
                            vertical += 1
        tables.append({"table": i, "rows": len(table.rows), "columns": len(table.columns), "visible_vertical_borders": vertical})
    checks = {
        "title_present": "Validation-ranking distortion and selection loss" in all_text,
        "abstract_four_labels": all(x in all_text for x in ["Background:", "Methods:", "Results:", "Conclusions:"]),
        "methods_10_subsections": all(f"2.{i} " in all_text for i in range(1, 11)),
        "results_7_subsections": all(f"3.{i} " in all_text for i in range(1, 8)),
        "discussion_5_subsections": all(f"4.{i} " in all_text for i in range(1, 6)),
        "figure_captions_1_to_4": all(f"Figure {i}." in all_text for i in range(1, 5)),
        "table_captions_1_to_5": all(f"Table {i}." in all_text for i in range(1, 6)),
        "needs_new_analysis_present": "needs new analysis" in all_text,
        "no_prospective_preregistration_claim": "prospective preregistration" in all_text and "not a prospective preregistration" in all_text,
        "three_line_no_vertical_borders": all(t["visible_vertical_borders"] == 0 for t in tables),
        "replacement_character_absent": "�" not in all_text,
    }
    return {
        "file": str(path),
        "paragraphs": len(doc.paragraphs),
        "tables": tables,
        "inline_shapes": len(doc.inline_shapes),
        "checks": checks,
        "all_checks_pass": all(checks.values()),
    }


def main() -> None:
    data = load_data()
    doc = build_document(data)
    doc.save(PAPER)
    shutil.copy2(PAPER, DESKTOP)
    write_changelog()
    write_reorganization()
    audit = audit_docx(PAPER)
    AUDIT.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
