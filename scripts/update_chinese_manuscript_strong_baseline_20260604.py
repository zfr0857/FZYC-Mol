from __future__ import annotations

import shutil
import textwrap
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
TABLE_DIR = ROOT / "reports" / "manuscript_tables"
FIG_DIR = ROOT / "reports" / "manuscript_figures_polished"
PACKAGE = ROOT / "reports" / "submission_package"

BASE_DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260603.docx"
BASE_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260603.md"
OUT_DOCX = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604.docx"
OUT_MD = DOCS / "manuscript_draft_full_zh_complete_integrated_single_paper_20260604.md"

FIG24 = FIG_DIR / "fig24_strong_baseline_selector_governance.png"
FIG24_SVG = FIG_DIR / "fig24_strong_baseline_selector_governance.svg"
TABLE47 = TABLE_DIR / "table47_strong_baseline_model_coverage.csv"
TABLE48 = TABLE_DIR / "table48_low_performance_targeted_actions.csv"
TABLE49 = TABLE_DIR / "table49_same_split_model_comparison_registry.csv"
TABLE50 = TABLE_DIR / "table50_tdc_performance_mode_custom_retained_best.csv"


def read_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path) if path.exists() else pd.DataFrame()


def fmt(value: object, digits: int = 4) -> str:
    try:
        number = float(value)
    except Exception:
        return "NA"
    if not np.isfinite(number):
        return "NA"
    return f"{number:.{digits}f}"


def draw_strong_baseline_figure() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(13.5, 7.2), dpi=180)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")
    fig.patch.set_facecolor("#f8fafc")

    def box(x: float, y: float, w: float, h: float, title: str, body: str, color: str, edge: str = "#334155") -> None:
        patch = plt.Rectangle((x, y), w, h, facecolor=color, edgecolor=edge, linewidth=1.5, joinstyle="round")
        ax.add_patch(patch)
        ax.text(x + 2.2, y + h - 5.2, title, fontsize=12, fontweight="bold", color="#0f172a", va="top")
        wrapped = body if "\n" in body else "\n".join(textwrap.wrap(body, width=22))
        ax.text(x + 2.2, y + h - 13.2, wrapped, fontsize=7.4, color="#334155", va="top", linespacing=1.08)

    ax.text(4, 96, "Strong Baseline and Selector Governance", fontsize=18, fontweight="bold", color="#0f172a")
    ax.text(
        4,
        91,
        "All candidates remain validation-only. Pilot candidates are retained only when they beat the existing retained-best gate.",
        fontsize=9.8,
        color="#475569",
    )

    box(5, 68, 23, 20, "Tier-1 Tabular", "CatBoost / XGBoost\nLightGBM / ExtraTrees / RF\nsame-split baselines", "#e0f2fe")
    box(38, 68, 23, 20, "TabPFNv2", "code connected\nlicense/token/cache pending\nno numeric claim", "#fef3c7")
    box(71, 68, 23, 20, "Deep Baselines", "Chemprop / D-MPNN\nfrozen ChemBERTa\nfrozen MoLFormer", "#dcfce7")

    for x0, x1 in [(28, 38), (61, 71)]:
        ax.annotate("", xy=(x1, 78), xytext=(x0, 78), arrowprops=dict(arrowstyle="->", lw=1.6, color="#64748b"))

    box(5, 37, 26, 20, "Candidate Pool", "tree experts + embeddings\ntarget transforms + underbag\nTop-K mean + stacking", "#f1f5f9")
    box(37, 37, 26, 20, "Validation Gate", "primary metric on validation\nfixed split and seed audit\ntest used after selection", "#ede9fe")
    box(69, 37, 26, 20, "Retained-Best Rule", "promote only if better\nunder metric direction\notherwise keep previous best", "#fee2e2")

    for x0, x1 in [(31, 37), (63, 69)]:
        ax.annotate("", xy=(x1, 47), xytext=(x0, 47), arrowprops=dict(arrowstyle="->", lw=1.8, color="#475569"))

    box(8, 10, 18, 18, "Promote", "validation accepted\nmain result tables", "#bbf7d0", "#15803d")
    box(31, 10, 18, 18, "Appendix", "pilot checks reported\nno main-table overwrite", "#dbeafe", "#2563eb")
    box(54, 10, 18, 18, "Guard", "skip unavailable TabPFN\navoid login hangs", "#fef9c3", "#ca8a04")
    box(77, 10, 18, 18, "Limit", "FreeSolv and rough ADME\nremain limitations", "#fecaca", "#dc2626")

    for x in [17, 40, 63, 86]:
        ax.annotate("", xy=(x, 28), xytext=(82, 37), arrowprops=dict(arrowstyle="->", lw=1.2, color="#94a3b8", alpha=0.75))

    ax.text(
        4,
        4,
        "Validation-only governance: strong baselines enter retained-best results only after validation acceptance.",
        fontsize=8.8,
        color="#475569",
    )
    fig.tight_layout()
    fig.savefig(FIG24, bbox_inches="tight")
    fig.savefig(FIG24_SVG, bbox_inches="tight")
    plt.close(fig)


def compact_model_table() -> pd.DataFrame:
    frame = read_csv(TABLE47)
    if frame.empty:
        return frame
    selected = frame[
        [
            "priority",
            "model_or_direction",
            "current_status",
            "paper_location",
            "next_action",
        ]
    ].copy()
    selected = selected.rename(
        columns={
            "priority": "优先级",
            "model_or_direction": "模型/方向",
            "current_status": "当前状态",
            "paper_location": "论文位置",
            "next_action": "下一步",
        }
    )
    return selected


def low_module_sentence() -> str:
    frame = read_csv(TABLE48)
    if frame.empty:
        return "低分模块的具体补强动作见 Table S43。"
    parts = []
    for row in frame.itertuples(index=False):
        if str(row.endpoint_group).lower() in {"freesolv", "lipo", "clintox"}:
            parts.append(
                f"{row.endpoint_group}: retained/reference={row.current_retained_or_reference}, "
                f"pilot={row.pilot_strong_baseline}, decision={row.decision}"
            )
    return "；".join(parts) + "。"


def registry_sentence() -> str:
    registry = read_csv(TABLE49)
    if registry.empty:
        return "same-split 对照注册表见 Table S44。"
    pieces = []
    for row in registry.itertuples(index=False):
        pieces.append(f"{row.source} n={row.n_rows}")
    return "当前 same-split/model-registry 覆盖：" + "；".join(pieces) + "。"


def table50_sentence() -> str:
    table = read_csv(TABLE50)
    if table.empty:
        return "TDC guard smoke 尚未生成结果。"
    row = table.iloc[0]
    return (
        f"TDC guard smoke 在 {row['dataset']} 上选择 {row['performance_model_counts']}，"
        f"test primary={fmt(row['performance_primary_mean'])}，"
        f"相对上一版 delta={fmt(row['performance_delta_vs_previous'])}，因此不覆盖正式 table15。"
    )


def set_run_font(run, size: float = 9.0, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str, style: str | None = None, size: float = 10.5):
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    set_run_font(run, size=size)
    return para


def add_heading(doc: Document, text: str, level: int) -> None:
    para = doc.add_heading(level=level)
    run = para.add_run(text)
    set_run_font(run, size=14 if level == 2 else 12, bold=True, color="0f172a")


def add_docx_table(doc: Document, caption: str, frame: pd.DataFrame) -> None:
    cap = add_paragraph(doc, caption, size=10)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if frame.empty:
        add_paragraph(doc, "未找到对应 CSV 表。", size=9)
        return
    table = doc.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, col in enumerate(frame.columns):
        hdr[idx].text = str(col)
    for row in frame.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            text = str(value)
            if len(text) > 95:
                text = text[:92] + "..."
            cells[idx].text = text
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=6.3)


def move_appended_block_before(doc: Document, marker_text: str, before_text: str) -> None:
    body = doc._body._element
    marker = None
    target = None
    for child in list(body):
        if child.tag != qn("w:p"):
            continue
        text = "".join(node.text or "" for node in child.iter() if node.tag == qn("w:t"))
        if marker_text in text:
            marker = child
        if target is None and text.strip().startswith(before_text):
            target = child
    if marker is None or target is None:
        raise RuntimeError("Could not find marker or insertion target in docx.")
    moving = []
    found = False
    for child in list(body):
        if child is marker:
            found = True
            continue
        if found and child.tag != qn("w:sectPr"):
            moving.append(child)
    body.remove(marker)
    for child in moving:
        body.remove(child)
    for child in moving:
        target.addprevious(child)


def replace_docx_text(doc: Document) -> None:
    replacements = {
        "表 19. 主文图表对应关系。": "表 20. 主文图表对应关系。",
        "表 20. 截至 2026-06-03 的全部实验结果与论文图表对应关系。": "表 21. 截至 2026-06-04 的全部实验结果与论文图表对应关系。",
    }
    for para in doc.paragraphs:
        text = para.text
        for old, new in replacements.items():
            if old in text:
                for run in para.runs:
                    run.text = ""
                para.add_run(text.replace(old, new))


def update_docx() -> None:
    shutil.copy2(BASE_DOCX, OUT_DOCX)
    doc = Document(OUT_DOCX)
    replace_docx_text(doc)
    marker = "__STRONG_BASELINE_UPDATE_20260604__"
    doc.add_paragraph(marker)
    add_heading(doc, "3.15 强模型对照与性能补强更新", level=2)
    add_paragraph(
        doc,
        "根据最新实验方向，本稿进一步把“增加模型复杂度”和“冲性能”拆成两条可审计路线。第一条路线是补齐同类型论文会期待的强模型对照，包括 TabPFNv2、CatBoost/XGBoost/LightGBM/ExtraTrees/RF、Chemprop/D-MPNN、冻结 ChemBERTa/MoLFormer 表征以及 Top-K/stacking ensemble。第二条路线是只在 validation-only 规则内改善 selector，包括 target transform、balanced undersampling、risk-adjusted selector 和 stability tie-breaker。这样写可以避免把论文叙事变成无边界堆模型，而是强调每个新增候选都必须先通过验证集治理边界。",
        size=10.5,
    )
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(FIG24), width=Cm(16.8))
    add_paragraph(
        doc,
        "图 12. 强模型对照与 selector 治理更新。新增 TabPFNv2、树模型全集、Chemprop/D-MPNN 和冻结分子语言模型表征后，所有候选仍通过 validation-only gate 与 retained-best rule 决定是否接入主结果；未授权 TabPFN 和未超过 retained-best 的 pilot 结果只进入附录。",
        size=9.2,
    )
    add_docx_table(doc, "表 19. 强模型对照与性能补强的当前状态。", compact_model_table())
    add_paragraph(
        doc,
        "本轮 pilot 的结果不建议覆盖主结果。MoleculeNet strong pilot 在 FreeSolv、Lipo 和 ClinTox 上分别选择 stacking 或 rank-fusion 候选，但均未超过已有 retained-best；因此它们更适合作为附录证据，说明本文确实比较了更强候选，但不会为了追求表面涨分而违反 retained-best gate。"
        + low_module_sentence(),
        size=10.5,
    )
    add_paragraph(
        doc,
        "TabPFNv2 的处理需要特别写清楚：当前本机已经安装 TabPFN 包，但尚未完成 PriorLabs/TabPFN license、token 或本地模型缓存准备。为避免自动实验触发交互式浏览器登录并导致实验中断，脚本现在使用 tabpfn_ready guard；在 tabpfn_ready=false 时自动跳过 TabPFN 数值实验。授权完成后，可以直接使用同一脚本重跑完整 validation-only integration，再决定是否进入主文或强附录。"
        + table50_sentence(),
        size=10.5,
    )
    add_paragraph(
        doc,
        registry_sentence()
        + " 这个 registry 的作用是把 same-split head-to-head 对照和 literature-reported reference 分开，避免把无法完全复现 split 的模型误写成直接可比结果。正文可以强调 CatBoost/XGBoost/LightGBM/ExtraTrees/RF、Chemprop/D-MPNN 与冻结 embedding heads 已经纳入可追溯候选池；TabPFN 是代码就绪但结果未就绪的强 baseline。",
        size=10.5,
    )
    move_appended_block_before(doc, marker, "4. 讨论")
    doc.save(OUT_DOCX)


def markdown_table(frame: pd.DataFrame, max_rows: int = 12) -> str:
    if frame.empty:
        return "_未找到对应 CSV 表。_"
    frame = frame.head(max_rows).copy()
    headers = [str(col) for col in frame.columns]
    rows = [[str(value).replace("\n", " ")[:110] for value in row] for row in frame.itertuples(index=False)]
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    out.extend("| " + " | ".join(row) + " |" for row in rows)
    return "\n".join(out)


def update_markdown() -> None:
    text = BASE_MD.read_text(encoding="utf-8")
    text = text.replace("表 19. 主文图表对应关系。", "表 20. 主文图表对应关系。")
    text = text.replace(
        "表 20. 截至 2026-06-03 的全部实验结果与论文图表对应关系。",
        "表 21. 截至 2026-06-04 的全部实验结果与论文图表对应关系。",
    )
    block = f"""
## 3.15 强模型对照与性能补强更新

根据最新实验方向，本稿进一步把“增加模型复杂度”和“冲性能”拆成两条可审计路线。第一条路线是补齐同类型论文会期待的强模型对照，包括 TabPFNv2、CatBoost/XGBoost/LightGBM/ExtraTrees/RF、Chemprop/D-MPNN、冻结 ChemBERTa/MoLFormer 表征以及 Top-K/stacking ensemble。第二条路线是只在 validation-only 规则内改善 selector，包括 target transform、balanced undersampling、risk-adjusted selector 和 stability tie-breaker。

![图 12. 强模型对照与 selector 治理更新。]({FIG24.as_posix()})

图 12. 强模型对照与 selector 治理更新。新增 TabPFNv2、树模型全集、Chemprop/D-MPNN 和冻结分子语言模型表征后，所有候选仍通过 validation-only gate 与 retained-best rule 决定是否接入主结果；未授权 TabPFN 和未超过 retained-best 的 pilot 结果只进入附录。

表 19. 强模型对照与性能补强的当前状态。

{markdown_table(compact_model_table())}

本轮 pilot 的结果不建议覆盖主结果。MoleculeNet strong pilot 在 FreeSolv、Lipo 和 ClinTox 上分别选择 stacking 或 rank-fusion 候选，但均未超过已有 retained-best；因此它们更适合作为附录证据，说明本文确实比较了更强候选，但不会为了追求表面涨分而违反 retained-best gate。{low_module_sentence()}

TabPFNv2 的处理需要特别写清楚：当前本机已经安装 TabPFN 包，但尚未完成 PriorLabs/TabPFN license、token 或本地模型缓存准备。为避免自动实验触发交互式浏览器登录并导致实验中断，脚本现在使用 tabpfn_ready guard；在 tabpfn_ready=false 时自动跳过 TabPFN 数值实验。授权完成后，可以直接使用同一脚本重跑完整 validation-only integration，再决定是否进入主文或强附录。{table50_sentence()}

{registry_sentence()} 这个 registry 的作用是把 same-split head-to-head 对照和 literature-reported reference 分开，避免把无法完全复现 split 的模型误写成直接可比结果。

"""
    text = text.replace("\n# 4. 讨论\n", "\n" + block + "\n# 4. 讨论\n")
    OUT_MD.write_text(text, encoding="utf-8")


def main() -> None:
    draw_strong_baseline_figure()
    update_docx()
    update_markdown()
    print(f"wrote {FIG24.relative_to(ROOT)}")
    print(f"wrote {FIG24_SVG.relative_to(ROOT)}")
    print(f"wrote {OUT_DOCX.relative_to(ROOT)}")
    print(f"wrote {OUT_MD.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
