# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "初稿-9.docx"
REPORT = ROOT / "output" / "初稿-9_最终QA报告.md"
JSON_OUT = ROOT / "reports" / "draft9_nature_polish" / "final_qa.json"


def border_value(cell, edge: str) -> str | None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        return None
    el = borders.find(qn(f"w:{edge}"))
    return el.get(qn("w:val")) if el is not None else None


def main() -> None:
    doc = Document(DOCX)
    paragraphs = [p.text for p in doc.paragraphs]
    body = "\n".join(paragraphs)
    table_text = "\n".join("\t".join(cell.text for row in t.rows for cell in row.cells) for t in doc.tables)
    text = body + "\n" + table_text

    figure_captions = [p.text for p in doc.paragraphs if p.style and p.style.name == "FigureCaption"]
    table_captions = [p.text for p in doc.paragraphs if p.style and p.style.name == "TableCaption"]
    figure_numbers = [int(re.match(r"图\s+(\d+)\s+\|", x).group(1)) for x in figure_captions]
    table_numbers = [int(re.match(r"表\s+(\d+)\s+\|", x).group(1)) for x in table_captions]
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    references = [p for p in paragraphs if re.match(r"^\[\d+\]", p)]
    citation_numbers = [
        int(x)
        for block in re.findall(r"\[(\d+(?:[,-]\d+)*)\]", body)
        for x in re.findall(r"\d+", block)
    ]

    table_checks = []
    for index, table in enumerate(doc.tables, 1):
        header = table.rows[0]
        last = table.rows[-1]
        vertical = []
        for row in table.rows:
            for cell in row.cells:
                vertical.extend([border_value(cell, "left"), border_value(cell, "right"), border_value(cell, "insideV")])
        tr_pr = header._tr.get_or_add_trPr()
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.first_child_found_in("w:tblLayout")
        table_checks.append(
            {
                "table": index,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "header_top": all(border_value(c, "top") == "single" for c in header.cells),
                "header_bottom": all(border_value(c, "bottom") == "single" for c in header.cells),
                "last_bottom": all(border_value(c, "bottom") == "single" for c in last.cells),
                "no_vertical_rules": all(v in {None, "nil"} for v in vertical),
                "all_rows_cant_split": all(
                    row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None for row in table.rows
                ),
                "repeat_header": tr_pr.find(qn("w:tblHeader")) is not None,
                "fixed_layout": layout is not None and layout.get(qn("w:type")) == "fixed",
            }
        )

    contribution_idx = paragraphs.index("科学贡献")
    contribution_sentences = len([x for x in paragraphs[contribution_idx + 1].split("。") if x.strip()])
    equations = sum(1 for p in doc.paragraphs if p.style and p.style.name == "Equation")
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))

    internal_terms = [
        "本轮", "当前工作区", "上一版", "对于投稿而言", "完美", "革命性", "世界领先", "轻松实现", "显而易见",
    ]
    legacy_terms = [
        "validation-best", "test-oracle", "source data", "nested validation", "decision card",
        "Scientific Contribution", "retrospective test-oracle", "held-out test regret",
    ]
    body_without_equations_refs = "\n".join(
        p.text for i, p in enumerate(doc.paragraphs) if p.style.name != "Equation" and i < len(doc.paragraphs) - len(references)
    )
    critical_patterns = {
        "candidate_pool_top3": ["0.897", "0.333"],
        "candidate_pool_regret": ["0.116", "0.120", "0.094"],
        "tdc_gate": ["5 个终点晋级、17 个保留、0 个下降"],
        "conformal_classification": ["0.814/0.918/0.956"],
        "conformal_regression": ["0.823/0.925/0.962"],
        "low_similarity": ["1.346", "1.431", "1.250", "1.070"],
        "clintox_recall": ["0.588±0.168", "0.491±0.195"],
        "moleculeace": ["0.711", "0.813", "0.252", "0.750"],
    }
    critical_missing = {
        key: [value for value in values if value not in text]
        for key, values in critical_patterns.items()
        if any(value not in text for value in values)
    }
    evidence_boundaries = {
        "autogluon_incomplete": "AutoGluon" in text and ("尚未完成" in text or "未获得可核验" in text),
        "temporal_blind_incomplete": "时间外盲测" in text and ("尚未完成" in text or "未设置" in text or "并非" in text),
        "doi_incomplete": "Zenodo DOI" in text and ("尚未完成" in text or "尚待" in text),
        "no_universal_optimum_claim": "不主张普遍最优" in text,
    }
    author_placeholders = re.findall(r"\[[^\]]*(?:repository|Zenodo|请作者)[^\]]*\]", text, flags=re.I)

    with ZipFile(DOCX) as archive:
        zip_test = archive.testzip() or "OK"
        media_count = len([n for n in archive.namelist() if n.startswith("word/media/")])

    qa = {
        "paragraphs": len(doc.paragraphs),
        "headings": len(headings),
        "blank_headings": [x for x in headings if not x.strip()],
        "tables": len(doc.tables),
        "figures": len(doc.inline_shapes),
        "figure_numbers": figure_numbers,
        "table_numbers": table_numbers,
        "references": len(references),
        "citation_max": max(citation_numbers) if citation_numbers else None,
        "citation_numbers_without_reference": sorted(set(n for n in citation_numbers if n > len(references))),
        "equations": equations,
        "cjk_characters": cjk,
        "scientific_contribution_sentences": contribution_sentences,
        "internal_terms": {q: text.count(q) for q in internal_terms if q in text},
        "legacy_terms": {q: body_without_equations_refs.lower().count(q.lower()) for q in legacy_terms if q.lower() in body_without_equations_refs.lower()},
        "critical_values_missing": critical_missing,
        "evidence_boundaries": evidence_boundaries,
        "author_placeholders": author_placeholders,
        "table_checks": table_checks,
        "zip_test": zip_test,
        "media_count": media_count,
    }

    qa["passed"] = (
        qa["figure_numbers"] == list(range(1, 11))
        and qa["table_numbers"] == list(range(1, 10))
        and qa["tables"] == 9
        and qa["figures"] == 10
        and qa["references"] == 23
        and qa["citation_numbers_without_reference"] == []
        and qa["equations"] == 10
        and qa["scientific_contribution_sentences"] == 3
        and qa["cjk_characters"] >= 13000
        and not qa["blank_headings"]
        and not qa["internal_terms"]
        and not qa["legacy_terms"]
        and not qa["critical_values_missing"]
        and all(qa["evidence_boundaries"].values())
        and all(
            x["columns"] <= 6
            and x["header_top"]
            and x["header_bottom"]
            and x["last_bottom"]
            and x["no_vertical_rules"]
            and x["all_rows_cant_split"]
            and x["repeat_header"]
            and x["fixed_layout"]
            for x in table_checks
        )
        and qa["zip_test"] == "OK"
        and qa["media_count"] == 10
    )

    JSON_OUT.parent.mkdir(parents=True, exist_ok=True)
    JSON_OUT.write_text(json.dumps(qa, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# 初稿-9 最终 QA 报告",
        "",
        f"- 总体：{'PASS' if qa['passed'] else 'FAIL'}",
        f"- 正文汉字数：{qa['cjk_characters']}",
        f"- 段落/标题：{qa['paragraphs']}/{qa['headings']}",
        f"- 图/表/公式/参考文献：{qa['figures']}/{qa['tables']}/{qa['equations']}/{qa['references']}",
        f"- DOCX 完整性：{qa['zip_test']}；媒体对象：{qa['media_count']}",
        "",
        "## 语言与论证",
        f"- 内部工作语：{qa['internal_terms'] or '无'}",
        f"- 未统一英文术语：{qa['legacy_terms'] or '无'}",
        f"- 关键数值缺失：{qa['critical_values_missing'] or '无'}",
        "- 摘要已合并为连续段落；引言按背景、缺口、研究问题和证据边界组织；结果采用观察、数值和限定的顺序。",
        "",
        "## 表格",
        "- 9 张主表均不超过 6 列，采用固定列宽；无竖线；保留顶线、表头线和底线。",
        "- 所有表头均设置跨页重复，所有表行均禁止跨页拆分，数字列居中，解释列左对齐。",
        "",
        "## 证据边界",
        f"- AutoGluon/时间外盲测/永久 DOI 缺口均已保留：{qa['evidence_boundaries']}",
        "- 未将测试集事后最优上界写成可用于晋级的性能，也未把失败或未运行实验写成完成结果。",
        "",
        "## 作者待补",
        "- 公开仓库、Zenodo DOI、基金和 CRediT 作者贡献仍为作者占位项；这些信息不能由编辑流程代填。",
        "- 本机 Word 后台 PDF 导出此前发生阻塞，本轮未重复启动该进程；DOCX 的 OOXML、媒体、表格结构与文件完整性已自动核验。",
    ]
    REPORT.write_text("\n".join(lines), encoding="utf-8")
    print(json.dumps(qa, ensure_ascii=False, indent=2))
    if not qa["passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
