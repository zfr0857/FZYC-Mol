# -*- coding: utf-8 -*-
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports" / "draft2_revision_20260616"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def locate_source() -> Path:
    desktop = Path.home() / "Desktop"
    exact = [p for p in desktop.rglob("*.docx") if p.name == "初稿-2.docx"]
    if exact:
        return max(exact, key=lambda p: p.stat().st_mtime)

    candidates = [p for p in desktop.rglob("*.docx") if p.stat().st_size > 20_000_000]
    if not candidates:
        raise FileNotFoundError("Could not locate 初稿-2.docx.")
    return max(candidates, key=lambda p: p.stat().st_mtime)


SRC_DOCX = locate_source()
DEST_DOCX = SRC_DOCX.parent / "初稿-4.docx"
REPORT_DOCX = SRC_DOCX.parent / "初稿-4_润色与投稿前自审报告.docx"


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def set_style_safe(paragraph: Paragraph, style_name: str | None) -> None:
    if not style_name:
        return
    try:
        paragraph.style = style_name
    except Exception:
        pass


def insert_before(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    set_style_safe(new_para, style)
    set_paragraph_text(new_para, text)
    return new_para


def insert_table_before(paragraph: Paragraph, headers: list[str], rows: list[list[str]]) -> None:
    marker = OxmlElement("w:p")
    paragraph._p.addprevious(marker)
    table = paragraph._parent.add_table(rows=1, cols=len(headers), width=Inches(6.3))
    marker.addprevious(table._tbl)
    marker.getparent().remove(marker)

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    apply_three_line_table(table)


def set_border(element, edge: str, val: str, size: str = "8", color: str = "000000") -> None:
    tag = "w:tblBorders" if element.tag.endswith("tblPr") else "w:tcBorders"
    borders = element.find(qn(tag))
    if borders is None:
        borders = OxmlElement(tag)
        element.append(borders)
    edge_el = borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        borders.append(edge_el)
    edge_el.set(qn("w:val"), val)
    edge_el.set(qn("w:sz"), size)
    edge_el.set(qn("w:space"), "0")
    edge_el.set(qn("w:color"), color)


def apply_three_line_table(table) -> None:
    tbl_pr = table._tbl.tblPr
    for edge in ["top", "bottom"]:
        set_border(tbl_pr, edge, "single", "12")
    for edge in ["left", "right", "insideH", "insideV"]:
        set_border(tbl_pr, edge, "nil")

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for edge in ["left", "right", "insideH", "insideV"]:
                set_border(tc_pr, edge, "nil")

    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        set_border(tc_pr, "bottom", "single", "8")


def replace_by_prefix(doc, prefix: str, new_text: str, changes: list[str], label: str) -> bool:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            set_paragraph_text(paragraph, new_text)
            changes.append(label)
            return True
    changes.append(f"未匹配：{label}")
    return False


def polish_core_sections(doc) -> list[str]:
    changes: list[str] = []

    replacements = [
        (
            "分子性质预测是药物发现",
            "分子性质预测是药物发现、ADMET 评估和毒性风险筛查中的关键计算环节。然而，单纯报告随机划分下的平均 ROC-AUC 或 RMSE，往往难以反映模型在新骨架外推、低相似度分子、不平衡毒性标签、规则五以外化学空间和活性悬崖等真实应用场景中的可靠性。本文提出 FZYC-Mol，一种由验证集治理驱动的适用域感知分子性质预测框架。该框架将分子图、指纹、二维描述符、片段/骨架信息和冻结预训练表征纳入统一候选池，并以验证集证据决定专家模型、融合策略、补救头和适用域门控是否进入最终测试。",
            "摘要：重写研究背景、问题缺口与方法边界",
        ),
        (
            "在 MoleculeNet 主面板中",
            "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别取得 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302。验证集接受的定向补救策略改善了 Lipophilicity 预测性能；低成本重构缩小了 FreeSolv 误差，但仍未超过 Chemprop，因此被保留为物理相互作用相关任务的边界案例。在 22 个外部 TDC ADMET 终点中，最终保留结果为 win/tie/loss = 5/17/0。风险评分、保形预测、低相似度分层、MoleculeACE 活性悬崖和负结果审计进一步表明，FZYC-Mol 的主要价值在于形成可追踪的模型选择与可靠性证据链，而不是宣称所有终点均获得统一提升。",
            "摘要：压缩结果叙述并降低过度泛化风险",
        ),
        (
            "分子性质预测已成为药物发现早期筛选",
            "分子性质预测已成为药物发现早期筛选、ADMET 评估、毒性风险预警和候选化合物优先级排序中的基础计算任务。近年来，模型形式从传统 QSAR、扩展连通性指纹、分子描述符和树模型，扩展到图神经网络、消息传递网络、分子语言模型和跨任务预训练模型。尽管这些方法在公开基准上推动了性能提升，但真实药物发现更关注模型能否在结构外推、低相似度分子、稀有毒性标签和化学空间转移中保持可解释、可校准且可审计的预测行为。",
            "引言：增强领域背景与真实应用问题",
        ),
        (
            "现有分子机器学习研究多以固定数据集",
            "现有分子机器学习研究多以固定数据集上的平均指标进行比较，这种评价方式存在三方面局限。第一，随机划分可能使训练集与测试集共享高度相似分子，从而低估结构外推难度。第二，不平衡 ADMET 终点上的高 ROC-AUC 并不必然对应可靠的阳性召回、概率校准或筛选富集。第三，在活性悬崖、高粗糙度终点或规则五以外化学空间中，结构相似分子的性质可能发生非连续变化，平均性能指标难以直接支持样本级决策。因此，分子性质预测研究需要同时报告性能、选择过程、适用域边界和失败模式。",
            "引言：明确三层问题链",
        ),
        (
            "针对上述问题，本文提出 FZYC-Mol",
            "针对上述问题，本文提出 FZYC-Mol，一种由验证集治理驱动的适用域感知分子性质预测框架。与强调单一大型主干模型的路线不同，FZYC-Mol 将多类分子表示、强基线模型、图模型、冻结分子表征、目标变换、集成策略、适用域门控、校准模块和风险证据头统一登记为候选策略。候选策略能否进入最终结果完全由验证集证据决定，测试集仅在策略冻结后用于一次性报告，从而降低候选池扩展过程中产生隐性测试集泄漏和事后选择偏差的风险。",
            "引言：收束为本文方法学主张",
        ),
        (
            "本文的主要贡献包括以下五个方面",
            "本文的主要贡献包括五个方面。第一，建立覆盖 MoleculeNet、TDC ADMET、bRo5 公共压力测试、MoleculeACE 活性悬崖任务以及 Random/Scaffold/Perimeter 划分的统一评测流程。第二，将 CatBoost、XGBoost、LightGBM、ExtraTrees、RF、Chemprop、冻结 ChemBERTa/MolT5/KPGT 表征和多视图融合候选纳入同一验证集选择框架。第三，构建适用域、风险评分、校准和保形预测模块，用于报告模型何时可信、何时应谨慎使用。第四，通过验证-测试排名审计、消融实验、负结果和失败案例分析，显式暴露选择器的不确定性和方法边界。第五，将化学片段、骨架和最近邻案例作为关联性解释证据，避免将统计富集直接表述为因果机制。",
            "引言：重写贡献列表并避免过强表述",
        ),
        (
            "本文使用 MoleculeNet、Therapeutics Data Commons ADMET",
            "本文使用 MoleculeNet、Therapeutics Data Commons ADMET、MoleculeACE 以及 bRo5 相关公开数据作为评估基础。MoleculeNet 主面板包括 ESOL、FreeSolv、Lipophilicity、BBBP、BACE 和 ClinTox；TDC ADMET 用于检验框架在外部 ADMET 终点上的可迁移性；MoleculeACE 用于评估活性悬崖和局部结构-性质非连续性；bRo5 相关数据作为规则五以外化学空间的公共压力测试。需要强调的是，bRo5 结果用于检验框架在可获得公开数据上的外推边界，不被表述为独立盲测验证。",
            "方法：澄清数据来源、用途与 bRo5 证据边界",
        ),
        (
            "为模拟不同真实应用难度",
            "为模拟不同真实应用难度，本文采用多层次数据划分策略。MoleculeNet 任务以 scaffold split 为主要划分方式，并补充 random split、structure-separated split 和 low-similarity hard subset。TDC ADMET 任务使用官方 PyTDC 划分，并补充全终点 scaffold 审计；MoleculeACE 采用任务预设划分评估活性悬崖样本；bRo5 公共压力测试用于观察规则五以外分子空间中的误差、校准和适用域变化。所有划分、随机种子和候选登记均在测试集评估前固定。",
            "方法：强化划分流程与冻结原则",
        ),
        (
            "FZYC-Mol 的验证集治理流程包括六个步骤",
            "FZYC-Mol 的验证集治理流程包括六个步骤。首先，预先定义每个终点的数据划分、主指标、指标方向、随机种子和候选专家集合。其次，仅使用训练集拟合单专家模型，包括指纹/描述符树模型、图模型、Chemprop、冻结表征头、目标变换模型和不平衡分类专家。第三，在验证集上生成专家预测矩阵，并基于同一矩阵构造 Top-K 均值、堆叠集成、自适应共识、不确定性加权融合和适用域门控。第四，根据验证集主指标、稳定性、风险调整和复杂度惩罚确定保留策略。第五，冻结候选权重、阈值、平局规则和补救头。第六，在测试集上进行一次性评估，并将未通过门控的候选作为负结果记录。",
            "方法：重写验证集治理步骤，使流程更可复核",
        ),
        (
            "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL",
            "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上表现稳定，分别取得 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302。ESOL 和 BACE 中，验证集选择结果与测试观测最优结果基本一致，提示验证集治理能够在部分终点上有效识别稳健候选。ClinTox 虽然 ROC-AUC 较高，但由于阳性样本稀少，结果仍需结合 PR-AUC、召回、校准和风险覆盖分析解释，不能仅凭 ROC-AUC 推断毒性筛查性能。",
            "结果：重写 MoleculeNet 主结果并补充 ClinTox 解释边界",
        ),
        (
            "在外部 TDC ADMET 任务上",
            "在外部 TDC ADMET 任务上，FZYC-Mol 显示出选择性而非普遍性的增益。多视图融合与适用域门控在 Caco2_Wang、HIA_Hou 和 Pgp_Broccatelli 等任务上带来改进；在多数终点中，验证集未支持更复杂候选替代原策略。整体上，22 个外部终点的最终保留结果为 win/tie/loss = 5/17/0，其中 5 个终点由性能增强候选进入最终保留，17 个终点保留原结果。该结果支持将 FZYC-Mol 定位为受验证集约束的审计框架，而不是对所有 ADMET 终点均保证提升的通用模型。",
            "结果：外部 ADMET 结论降调并突出选择性增益",
        ),
        (
            "在 OOD、Perimeter 和低相似度压力测试中",
            "在 OOD、Perimeter 和低相似度压力测试中，模型性能通常随测试分子与训练集最近邻相似度下降而降低。适用域距离、风险分位和最近邻标签差异能够帮助识别部分高风险样本，但不能完全消除结构外推误差。MoleculeACE 活性悬崖分析显示，相似分子之间的真实标签差异仍可能被平均指标低估，尤其是在高粗糙度终点中。cliff-pair RMSE、gap Spearman 和代表性失败案例共同说明，FZYC-Mol 的可靠性输出应被视为复核依据，而非自动纠错机制。",
            "结果：低相似度与活性悬崖结果改为证据链表达",
        ),
        (
            "本文仍存在一定局限",
            "本文仍存在一定局限。首先，FZYC-Mol 目前仍以公开 benchmark 和离线评估为主，尚未经过真实药物发现项目中的前瞻性验证。其次，bRo5 化学空间仅作为 CycPept-PAMPA、LinPept CellPen 和 LinPept NonFouling 等公开数据上的压力测试，不能替代独立盲测或实验验证。再次，FreeSolv 等物理相互作用相关任务提示，二维描述符和低成本重构仍难以充分捕捉构象、溶剂化和长程相互作用。最后，基序和片段富集只能作为关联性解释，仍需结合统计校正、化学合理性和实验知识进行复核。",
            "讨论：重写局限性，避免把未完成验证写成完成结论",
        ),
        (
            "本文提出并评估了 FZYC-Mol",
            "本文提出并评估了 FZYC-Mol，一种由验证集治理驱动、适用域感知的多专家分子性质预测框架。该框架将多视图表示、强基线专家、适用域判断、风险评分、校准、保形预测、活性悬崖分析和负结果审计整合到统一候选池中，并严格遵守测试集冻结原则。实验结果表明，FZYC-Mol 能够在部分 MoleculeNet 和外部 TDC ADMET 终点上提供选择性增益，同时以失败案例和负结果界定模型边界。总体而言，FZYC-Mol 的核心贡献是为分子性质预测提供一套可复核的模型选择与可靠性报告流程。",
            "结论：收束核心贡献与证据边界",
        ),
        (
            "未来工作将继续扩展强基线",
            "未来工作将进一步扩展强基线同划分比较、bRo5 化学空间、MoleculeACE 活性悬崖、外部留出集和盲测风格验证，并在代表性低分任务上开展受控适配器、轻量微调或 3D-lite 物理代理实验。同时，FZYC-Mol 还应发展为可复现软件包或在线审计工具，使模型预测、适用域证据、风险分位、最近邻解释和拒用理由能够以决策卡形式服务于药物化学和 ADMET 筛选。",
            "结论：改写未来工作，使用正式中文并保留可执行方向",
        ),
        (
            "本文使用的公开数据集可通过",
            "本文使用的公开数据集可通过 MoleculeNet、Therapeutics Data Commons、MoleculeACE、Benchmark-ADMET-2025 及相应原始平台获得。与本文结果对应的 split seeds、候选登记表、验证/测试预测、统计检验脚本、图表 source data、环境文件和表格生成脚本应在投稿或接收版本中同步存档于 GitHub、Zenodo 或机构数据仓库。任何尚未公开的数据、未完成的盲测或仅作为方案设计的实验，不应写入完成性结果。",
            "数据可用性：补充学术诚信边界",
        ),
    ]

    for prefix, new_text, label in replacements:
        replace_by_prefix(doc, prefix, new_text, changes, label)

    # Light language cleanup for a more formal Chinese manuscript style.
    cleanup = {
        "baseline": "基线模型",
        "holdout": "留出集",
        "dataset-seed": "数据集-随机种子",
    }
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("["):
            continue
        updated = paragraph.text
        for old, new in cleanup.items():
            updated = updated.replace(old, new)
        if updated != paragraph.text:
            set_paragraph_text(paragraph, updated)

    return changes


def revise_tables(doc) -> list[str]:
    changes: list[str] = []
    for table in doc.tables:
        apply_three_line_table(table)

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if not cells:
                continue
            first = cells[0]
            if first == "bRo5" and len(row.cells) >= 5:
                row.cells[1].text = "规则五以外公共压力测试"
                row.cells[2].text = "CycPept-PAMPA、LinPept CellPen、LinPept NonFouling"
                row.cells[3].text = "RMSE、ROC-AUC、PR-AUC、适用域审计"
                row.cells[4].text = "用于观察规则五以外化学空间中的误差、校准与适用域边界；不表述为独立盲测验证"
                changes.append("表格：更新 bRo5 数据集用途与证据边界")
            elif "bRo5 化学空间" in first:
                if len(row.cells) > 1:
                    row.cells[1].text = "CycPept-PAMPA 与 LinPept 相关公共压力测试已纳入；结论限定为可获得公开数据上的外推与适用域审计，不等同于独立实验盲测。"
                changes.append("表格：修订 bRo5 边界案例描述")

    changes.append("表格：全文表格已统一为三线表边框样式")
    return changes


def add_self_review(doc) -> None:
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "List of abbreviations":
            anchor = paragraph
            break
    if anchor is None:
        return

    insert_before(anchor, "6 投稿前自我审查", "Heading 1")
    insert_before(
        anchor,
        "依据 Research-Paper-Writing-Skills 中的摘要、引言、方法、实验、结论、段落流畅性和论文审查清单，本节从投稿前审稿人视角检查论文的逻辑链、主张-证据一致性、方法边界和潜在拒稿风险。该部分用于内部自审；正式投稿时可根据期刊要求移至补充材料或作者修订说明。",
    )
    insert_before(anchor, "6.1 主张-证据一致性", "Heading 2")
    insert_table_before(
        anchor,
        ["主张", "对应证据", "一致性判断", "修订处理"],
        [
            [
                "FZYC-Mol 在部分终点提供性能增益",
                "MoleculeNet 主结果、Lipophilicity 补救、TDC win/tie/loss = 5/17/0",
                "基本一致",
                "统一表述为选择性增益，避免写成所有任务全面优于基线",
            ],
            [
                "验证集治理可降低测试集事后选择风险",
                "候选登记、验证选择、测试冻结、验证-测试排名审计",
                "证据支持但存在边界",
                "补充说明验证集排序并不稳定，不能保证验证第一即测试最优",
            ],
            [
                "风险分数可支持可靠性判断",
                "错误识别 AUROC、risk-coverage、低相似度与失败案例",
                "部分支持",
                "将其定位为复核提示，不表述为自动拒用或自动纠错机制",
            ],
            [
                "bRo5 结果支持规则五以外泛化",
                "CycPept-PAMPA、LinPept CellPen、LinPept NonFouling 公共压力测试",
                "需降调",
                "改为公共压力测试和适用域边界观察，不写成独立盲测验证",
            ],
            [
                "基序/片段解释具有化学意义",
                "基序归因、片段富集、最近邻案例和统计校正要求",
                "关联性支持",
                "明确其为关联解释，不等同于因果机制",
            ],
        ],
    )
    insert_before(anchor, "6.2 三位审稿人视角自审", "Heading 2")
    insert_before(
        anchor,
        "审稿人 1（方法学严谨性）：稿件的优点是将候选登记、验证集选择和测试集冻结写成可审计流程，并报告负结果和边界案例。主要风险是候选池较大时仍可能产生验证集过拟合，因此正文已保留验证-测试排名审计、Top-3 一致性和 nested validation 边界表述。",
    )
    insert_before(
        anchor,
        "审稿人 2（新颖性与意义）：稿件不应将创新点放在单一模型结构上，而应强调验证集治理、适用域证据、可靠性输出和失败模式记录形成的整体审计框架。当前修订已将贡献定位为可复核流程，避免把选择性增益包装成普遍性能优势。",
    )
    insert_before(
        anchor,
        "审稿人 3（写作与可读性）：摘要和引言原稿信息密度较高，容易削弱主线。当前修订将论文主线调整为“真实场景挑战—验证集治理框架—多层证据审计—边界与失败模式”，并在方法、结果和讨论中保持同一逻辑顺序。",
    )
    insert_before(anchor, "6.3 主要拒稿风险与投稿前处理", "Heading 2")
    insert_table_before(
        anchor,
        ["风险点", "可能审稿意见", "当前处理", "投稿前建议"],
        [
            [
                "贡献定位",
                "像工程集成而非机制创新",
                "强调验证集治理和可靠性审计",
                "在 cover letter 中突出可复核候选选择流程",
            ],
            [
                "候选池选择偏差",
                "验证集可能被反复使用而过拟合",
                "加入排名审计和测试冻结原则",
                "保留完整候选登记表和种子级预测文件",
            ],
            [
                "实验边界",
                "bRo5 和 MoleculeACE 不能替代盲测",
                "已降调为公共压力测试和活性悬崖审计",
                "若目标期刊要求更强证据，补充独立外部留出集",
            ],
            [
                "语言与结构",
                "信息过密、结论跳跃",
                "重写摘要、引言、方法、结果和结论主段",
                "投稿前再按期刊字数压缩图表和补充材料",
            ],
            [
                "学术诚信",
                "未完成实验被写成完成结果",
                "明确所有公共数据、压力测试和未完成验证的边界",
                "逐项核对 source data、脚本、seed 和图表是否可追溯",
            ],
        ],
    )


def build_report(changes: list[str], table_changes: list[str], qa_notes: list[str]) -> None:
    doc = Document()
    doc.add_heading("初稿-4 润色与投稿前自审报告", level=1)
    doc.add_paragraph(f"源文档：{SRC_DOCX}")
    doc.add_paragraph(f"输出文档：{DEST_DOCX}")
    doc.add_heading("采用的写作规范", level=2)
    doc.add_paragraph(
        "本次修订调用 D:\\skill\\Research-Paper-Writing-Skills 中的 abstract、introduction、method、experiments、conclusion、paper-review 与 paragraph-flow 指南，并结合 Nature 风格写作原则，对摘要、引言、方法、实验结果、讨论和结论进行中文润色、逻辑整理与投稿前审稿人视角自审。"
    )
    doc.add_heading("主要修改", level=2)
    for item in changes + table_changes:
        doc.add_paragraph(item)
    doc.add_heading("QA 检查", level=2)
    for item in qa_notes:
        doc.add_paragraph(item)
    doc.add_heading("自审结论", level=2)
    doc.add_paragraph(
        "稿件主线已调整为“真实应用挑战—验证集治理—多层证据审计—边界与失败模式”。核心贡献应定位为可复核的模型选择与可靠性报告流程，而不是单一模型在所有终点上的普遍性能优势。当前版本已对 bRo5、MoleculeACE、风险分数和基序解释等容易被质疑的部分进行降调和边界化处理。投稿前仍需逐项核对 source data、seed-level 预测、图表源文件和参考文献元数据。"
    )
    doc.save(REPORT_DOCX)


def collect_qa(doc) -> list[str]:
    notes: list[str] = []
    text = "\n".join(p.text for p in doc.paragraphs)
    required = [
        "6 投稿前自我审查",
        "6.1 主张-证据一致性",
        "6.2 三位审稿人视角自审",
        "6.3 主要拒稿风险与投稿前处理",
    ]
    for heading in required:
        notes.append(f"{heading}：{'已插入' if heading in text else '未检出'}")

    max_cols = max((len(t.columns) for t in doc.tables), default=0)
    notes.append(f"表格列数检查：全文最大列数 {max_cols}。")
    notes.append(f"图像保留检查：inline shapes 数量 {len(doc.inline_shapes)}。")
    if "数据状态审计" in text and "bRo5" in text:
        notes.append("提示：正文仍含“数据状态审计”字样，请人工复核 bRo5 表述。")
    else:
        notes.append("bRo5 表述检查：未发现旧版“数据状态审计”主张。")
    if "未完成的盲测或仅作为方案设计的实验，不应写入完成性结果" in text:
        notes.append("学术诚信边界：已写入数据可用性与自审段落。")
    return notes


def main() -> None:
    doc = Document(SRC_DOCX)
    changes = polish_core_sections(doc)
    table_changes = revise_tables(doc)
    add_self_review(doc)
    qa_notes = collect_qa(doc)

    doc.save(DEST_DOCX)
    shutil.copy2(DEST_DOCX, OUT_DIR / DEST_DOCX.name)

    build_report(changes, table_changes, qa_notes)
    shutil.copy2(REPORT_DOCX, OUT_DIR / REPORT_DOCX.name)

    print(f"Wrote {DEST_DOCX}")
    print(f"Wrote {REPORT_DOCX}")
    print(f"Paragraph changes: {len(changes)}")
    print(f"Table changes: {len(table_changes)}")
    for note in qa_notes:
        print(note)


if __name__ == "__main__":
    main()
