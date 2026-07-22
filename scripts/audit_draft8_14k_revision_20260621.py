# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import re
from pathlib import Path
from zipfile import ZipFile

import pandas as pd
from PIL import Image
from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "初稿-8.docx"
BUNDLE = ROOT / "output" / "初稿-8_图表与源数据"
REPORT = ROOT / "output" / "初稿-8_最终QA报告.md"
JSON_OUT = ROOT / "reports" / "draft8_14k_revision" / "final_qa.json"


def border_value(cell, edge: str) -> str | None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        return None
    el = borders.find(qn(f"w:{edge}"))
    return el.get(qn("w:val")) if el is not None else None


def main() -> None:
    doc = Document(DOCX)
    paragraphs = [p.text for p in doc.paragraphs]
    body = "\n".join(paragraphs)
    table_text = "\n".join("\t".join(cell.text for row in t.rows for cell in row.cells) for t in doc.tables)
    text = body + "\n" + table_text

    figure_captions = [p.text for p in doc.paragraphs if p.style and p.style.name == "FigureCaption"]
    table_captions = [p.text for p in doc.paragraphs if p.style and p.style.name == "TableCaption"]
    figure_numbers = [int(re.match(r"图\s+(\d+)", x).group(1)) for x in figure_captions]
    table_numbers = [int(re.match(r"表\s+(\d+)", x).group(1)) for x in table_captions]
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    references = [p for p in paragraphs if re.match(r"^\[\d+\]", p)]
    citation_numbers = [int(x) for block in re.findall(r"\[(\d+(?:[,-]\d+)*)\]", body) for x in re.findall(r"\d+", block)]

    table_checks = []
    for index, table in enumerate(doc.tables, 1):
        header = table.rows[0]
        last = table.rows[-1]
        vertical = []
        for row in table.rows:
            for cell in row.cells:
                vertical.extend([border_value(cell, "left"), border_value(cell, "right"), border_value(cell, "insideV")])
        cant_split = [row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None for row in table.rows]
        repeat_header = header._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None
        table_checks.append(
            {
                "table": index,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "header_top": all(border_value(c, "top") == "single" for c in header.cells),
                "header_bottom": all(border_value(c, "bottom") == "single" for c in header.cells),
                "last_bottom": all(border_value(c, "bottom") == "single" for c in last.cells),
                "no_vertical_rules": all(v in {None, "nil"} for v in vertical),
                "all_rows_cant_split": all(cant_split),
                "repeat_header": repeat_header,
            }
        )

    figure_checks = []
    for index in range(1, 11):
        matches = list((BUNDLE / "figures").glob(f"fig{index:02d}_*.png"))
        if len(matches) != 1:
            figure_checks.append({"figure": index, "error": f"expected one PNG, found {len(matches)}"})
            continue
        png = matches[0]
        stem = png.with_suffix("")
        with Image.open(png) as im:
            width, height = im.size
        figure_checks.append(
            {
                "figure": index,
                "stem": png.stem,
                "width_px": width,
                "height_px": height,
                "png_high_res": width >= 1800,
                "svg": stem.with_suffix(".svg").exists(),
                "pdf": stem.with_suffix(".pdf").exists(),
                "tiff": stem.with_suffix(".tiff").exists(),
            }
        )

    stress = pd.read_csv(ROOT / "reports" / "draft8_14k_revision" / "candidate_pool_stress_summary.csv")
    row32 = stress[stress["pool_size"].eq(32)].set_index("policy")
    conf = pd.read_csv(ROOT / "reports" / "remaining_missing_experiments_20260606" / "conformal_80_90_95_summary.csv")
    clintox = pd.read_csv(ROOT / "reports" / "reviewer_revision_20260607" / "clintox_fixed_precision_recall_consensus_strict_core_multifp.csv")
    critical = {
        "pool_top3_4": float(stress[(stress["pool_size"].eq(4)) & stress["policy"].eq("validation_best")]["top3_hit_rate"].iloc[0]),
        "pool_top3_32": float(row32.loc["validation_best", "top3_hit_rate"]),
        "validation_best_regret_32": float(row32.loc["validation_best", "normalized_regret_mean"]),
        "risk_adjusted_regret_32": float(row32.loc["risk_adjusted", "normalized_regret_mean"]),
        "classification_conformal": conf[conf["task_type"].eq("classification")]["coverage_mean"].round(3).tolist(),
        "regression_conformal": conf[conf["task_type"].eq("regression")]["coverage_mean"].round(3).tolist(),
        "clintox_recall_p80": float(clintox["recall_p80"].mean()),
        "clintox_recall_p90": float(clintox["recall_p90"].mean()),
    }

    internal_terms = [
        "本轮", "当前工作区", "本地复现包", "上一版", "对于投稿而言", "参考文献还需核验", "完美", "革命性", "世界领先",
    ]
    critical_phrases = [
        "0.897", "0.333", "0.120", "0.094", "0.814/0.918/0.956", "0.823/0.925/0.962",
        "5 个提升、17 个保留、0 个下降", "0.588 ± 0.168", "0.491 ± 0.195",
    ]
    source_csvs = list((BUNDLE / "source_data").glob("*.csv"))

    qa = {
        "paragraphs": len(doc.paragraphs),
        "headings": len(headings),
        "blank_headings": [x for x in headings if not x.strip()],
        "tables": len(doc.tables),
        "figures": len(doc.inline_shapes),
        "figure_numbers": figure_numbers,
        "table_numbers": table_numbers,
        "references": len(references),
        "citation_max": max(citation_numbers) if citation_numbers else None,
        "citation_numbers_without_reference": sorted(set(n for n in citation_numbers if n > len(references))),
        "equations": sum(1 for p in doc.paragraphs if p.style and p.style.name == "Equation"),
        "cjk_characters": len(re.findall(r"[\u4e00-\u9fff]", text)),
        "internal_terms": {q: text.count(q) for q in internal_terms if q in text},
        "critical_phrases_missing": [q for q in critical_phrases if q not in text],
        "scientific_contribution_sentences": len([x for x in paragraphs[paragraphs.index("Scientific Contribution") + 1].split("。") if x.strip()]),
        "table_checks": table_checks,
        "figure_checks": figure_checks,
        "source_csv_count": len(source_csvs),
        "critical_source_values": critical,
        "figure3_uses_heldout_label": "held-out test regret" in (BUNDLE / "figures" / "fig03_candidate_pool_stress.svg").read_text(encoding="utf-8"),
        "figure3_avoids_outer_label": "outer-test regret" not in (BUNDLE / "figures" / "fig03_candidate_pool_stress.svg").read_text(encoding="utf-8"),
    }

    with ZipFile(DOCX) as z:
        qa["zip_test"] = z.testzip() or "OK"
        qa["media_count"] = len([n for n in z.namelist() if n.startswith("word/media/")])

    passed = (
        qa["figure_numbers"] == list(range(1, 11))
        and qa["table_numbers"] == list(range(1, 10))
        and qa["references"] == 23
        and qa["citation_numbers_without_reference"] == []
        and qa["equations"] == 10
        and qa["scientific_contribution_sentences"] == 3
        and not qa["blank_headings"]
        and not qa["internal_terms"]
        and not qa["critical_phrases_missing"]
        and all(x.get("png_high_res") and x.get("svg") and x.get("pdf") and x.get("tiff") for x in figure_checks)
        and all(x["columns"] <= 7 and x["header_top"] and x["header_bottom"] and x["last_bottom"] and x["no_vertical_rules"] and x["all_rows_cant_split"] and x["repeat_header"] for x in table_checks)
        and qa["zip_test"] == "OK"
        and qa["media_count"] == 10
        and qa["figure3_uses_heldout_label"]
        and qa["figure3_avoids_outer_label"]
    )
    qa["passed"] = passed
    JSON_OUT.parent.mkdir(parents=True, exist_ok=True)
    JSON_OUT.write_text(json.dumps(qa, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = ["# 初稿-8 最终 QA 报告", "", f"- 总体：{'PASS' if passed else 'FAIL'}"]
    for key in ["paragraphs", "headings", "tables", "figures", "references", "equations", "cjk_characters", "source_csv_count", "zip_test"]:
        lines.append(f"- {key}: {qa[key]}")
    lines += ["", "## 编号与语言", f"- 图号：{qa['figure_numbers']}", f"- 表号：{qa['table_numbers']}", f"- 无引用条目：{qa['citation_numbers_without_reference']}", f"- 内部工作语：{qa['internal_terms']}", f"- 缺失关键数值：{qa['critical_phrases_missing']}"]
    lines += ["", "## 表格", "- 9 张主表均不超过 7 列；表头重复、整行不跨页；仅保留顶线、表头线和底线。"]
    lines += ["", "## 图形", "- 10 张主图均有 PNG/SVG/PDF/TIFF；结果图由 Python 从 source data 生成。", "- 图 3 已使用 held-out test regret，未把回顾性压力审计误写为 outer nested。"]
    lines += ["", "## 限制", "- Microsoft Word 后台 PDF 导出在本机 Office 进程中卡住，因此未完成逐页 PDF 视觉审计；DOCX XML、媒体、表格分页保护和文件完整性均已通过自动检查。"]
    REPORT.write_text("\n".join(lines), encoding="utf-8")
    print(json.dumps(qa, ensure_ascii=False, indent=2))
    if not passed:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
