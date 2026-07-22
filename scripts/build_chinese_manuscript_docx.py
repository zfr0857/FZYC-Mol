from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
PACKAGE = ROOT / "reports" / "submission_package"
MAIN_FIGS = PACKAGE / "main_figures"
SUPP_FIGS = PACKAGE / "supplementary_figures"
MAIN_TABLES = PACKAGE / "main_tables"
SUPP_TABLES = PACKAGE / "supplementary_tables"
OUT = DOCS / "manuscript_draft_full_zh_integrated_20260531.docx"


def set_normal_style(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Microsoft YaHei"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Microsoft YaHei"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.name = "Microsoft YaHei"
    styles["Heading 3"].font.size = Pt(11.5)


def set_margins(section) -> None:
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)


def add_para(doc: Document, text: str = "", style: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(31, 41, 55)
    p.paragraph_format.space_after = Pt(8)


def add_figure(doc: Document, path: Path, caption: str, width_in: float = 6.4) -> None:
    if not path.exists():
        add_para(doc, f"[缺失图片：{path}]", bold=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    add_caption(doc, caption)


def clean_cell(value: object, max_len: int = 90) -> str:
    if pd.isna(value):
        return ""
    text = str(value).replace("\n", " ").replace("\r", " ").strip()
    if len(text) > max_len:
        return text[: max_len - 1] + "…"
    return text


def style_table(table, font_size: float = 7.0) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(font_size)
                    if row_idx == 0:
                        run.bold = True


def add_dataframe_table(
    doc: Document,
    frame: pd.DataFrame,
    font_size: float = 7.0,
    max_text: int = 90,
) -> None:
    table = doc.add_table(rows=1, cols=len(frame.columns))
    hdr = table.rows[0].cells
    for i, col in enumerate(frame.columns):
        hdr[i].text = clean_cell(col, max_len=60)
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(frame.columns):
            cells[i].text = clean_cell(row[col], max_len=max_text)
    style_table(table, font_size=font_size)


def add_csv_table(
    doc: Document,
    label: str,
    caption: str,
    path: Path,
    key_cols: list[str] | None = None,
    chunk_size: int = 7,
    font_size: float = 6.8,
    max_text: int = 80,
) -> None:
    add_para(doc, f"{label}. {caption}", bold=True)
    if not path.exists():
        add_para(doc, f"[缺失表格文件：{path}]", bold=True)
        return
    frame = pd.read_csv(path).fillna("")
    if len(frame.columns) <= 10:
        add_dataframe_table(doc, frame, font_size=font_size, max_text=max_text)
    else:
        key_cols = [col for col in (key_cols or list(frame.columns[:3])) if col in frame.columns]
        other_cols = [col for col in frame.columns if col not in key_cols]
        for idx in range(0, len(other_cols), chunk_size):
            chunk = other_cols[idx : idx + chunk_size]
            sub = frame[key_cols + chunk]
            add_para(doc, f"{label} 分栏 {idx // chunk_size + 1}：{', '.join(chunk)}", bold=True)
            add_dataframe_table(doc, sub, font_size=font_size, max_text=max_text)
    add_para(doc, f"完整 CSV 文件：{path}", style=None)


def add_landscape_section(doc: Document):
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    set_margins(section)
    return section


def build_doc() -> Document:
    doc = Document()
    set_normal_style(doc)
    set_margins(doc.sections[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FZYC-Mol：面向结构分布偏移的验证集选择与适用域感知多专家分子性质预测框架")
    run.bold = True
    run.font.size = Pt(18)
    add_caption(doc, "完整中文整合初稿，2026-05-31")

    add_para(
        doc,
        "说明：本文为计算分子 AI / 化学信息学 / 药物发现机器学习论文初稿，不包含湿实验。所有结果来自公开分子性质数据集、官方划分协议、结构外推评估、活性悬崖评估、校准/不确定性/适用域分析、可解释性分析和外部 ADMET appendix benchmark。",
    )

    doc.add_heading("摘要", level=1)
    add_para(
        doc,
        "分子性质预测是药物发现、ADMET 评估、毒性筛选和先导化合物优化中的核心计算任务。现有研究仍存在三个可靠性缺口：随机或单一 scaffold 划分可能高估结构新颖分子的泛化能力；不同 endpoint 对表示类型和模型家族的偏好明显不同；许多 benchmark 主要报告点预测性能，而对不确定性、校准、适用域、活性悬崖、早期富集和化学可解释性的系统证据不足。",
    )
    add_para(
        doc,
        "本文提出 FZYC-Mol，一个严格基于验证集选择且具备适用域感知能力的多专家分子性质预测框架。FZYC-Mol 整合图神经网络、D-MPNN/Chemprop、随机森林、XGBoost、LightGBM、ExtraTrees、多类型分子指纹、RDKit 描述符模型、BRICS/Murcko/官能团 motif 专家、ChemBERTa 与 MoLFormer 冻结 embedding、validation stacking、adaptive consensus、不确定性模型和适用域加权机制。所有最终预测策略均只使用 validation split 进行选择，test split 仅用于最终报告。",
    )
    add_para(
        doc,
        "在 MoleculeNet 上，FZYC-Mol 表现出端点特异性的竞争性能：ESOL 达到 RMSE 0.5829 +/- 0.0352，BACE 达到 ROC-AUC 0.8753 +/- 0.0230，ClinTox 达到 ROC-AUC 0.9489 +/- 0.0302。进一步地，基于 validation-only 原则的 rescue-integrated selector 在 Lipo 上将 RMSE 从 0.7078 +/- 0.0389 降低到 0.6835 +/- 0.0439。TDC ADMET 结果显示，validation selector 在 HIA_Hou、Pgp_Broccatelli 和 BBB_Martins 上分别达到 ROC-AUC 0.9792、0.9357 和 0.9182。",
    )
    add_para(
        doc,
        "结构分布偏移实验显示，随机划分性能可能显著高估实际泛化能力。重构陌生度提供了与传统 Tanimoto 相似性互补的适用域信号，hybrid reconstruction-plus-AD score 与绝对误差的平均 Spearman 相关达到 0.2225。Motif attribution 进一步提供化学可解释信号。总体而言，FZYC-Mol 将分子性质预测从单一 benchmark 分数比较推进到结构分布偏移下的验证集控制、可靠性评估和化学解释闭环。",
    )

    doc.add_heading("关键词", level=1)
    add_para(doc, "分子性质预测；ADMET；适用域；不确定性估计；validation stacking；结构分布偏移；活性悬崖；分子指纹；Chemprop；分子预训练模型")

    doc.add_heading("图文摘要", level=1)
    add_figure(
        doc,
        MAIN_FIGS / "Figure_1_FZYC_Mol_framework.png",
        "图 1. FZYC-Mol 框架总览。分子输入被转化为图、指纹、描述符、motif 和冻结预训练 embedding 等多视图表示，并进入异质专家池；最终预测策略只由验证集选择。",
    )

    doc.add_heading("1. 引言", level=1)
    add_para(
        doc,
        "分子性质预测是现代计算药物发现的基础任务。给定分子结构，模型可以在实验测量之前估计溶解度、脂溶性、膜渗透性、血脑屏障穿透、酶抑制、毒性或其他 ADMET 相关性质。这类模型能够降低实验成本、提高虚拟筛选效率、支持 hit-to-lead 优化，并帮助研究者尽早识别高风险化合物。",
    )
    add_para(
        doc,
        "分子机器学习已经从手工描述符和传统机器学习快速发展到图神经网络、消息传递神经网络和分子预训练模型。然而，分子性质预测的核心挑战不仅是平均性能，更是可靠性。在真实药物发现中，测试化合物常常来自新的 scaffold、新的化学系列或低相似度区域。随机划分可能让训练集和测试集共享高度相似结构，从而高估泛化能力。",
    )
    add_para(
        doc,
        "因此，本文提出 FZYC-Mol：一个验证集选择、适用域感知、多专家融合的分子性质预测框架。本文的核心不是声称某个单一架构普遍优于所有模型，而是构建一个可复现的分子 AI 可靠性评估闭环。",
    )
    add_numbered(
        doc,
        [
            "提出严格 validation-only 的分子专家选择协议，避免测试集策略选择泄漏。",
            "构建整合图模型、Chemprop、指纹、描述符、motif 专家和冻结预训练编码器的多专家框架。",
            "在 MoleculeNet、TDC ADMET、官方 PyTDC split、MoleculeACE、低相似度 hard subset 和 structure-separated split 上进行广泛评估。",
            "引入不确定性、校准、适用域、conformal prediction、validation-trained error model 和 reconstruction-based unfamiliarity 等可靠性模块。",
            "通过 BRICS fragment、Murcko scaffold 和官能团 motif attribution 加强化学解释。",
            "对低分模块进行 targeted rescue，并通过完整 validation-only selector integration 验证是否应纳入主结果。",
        ],
    )

    doc.add_heading("2. 相关工作", level=1)
    add_para(doc, "MoleculeNet 是分子机器学习中最常用的公开基准之一。本文使用 ESOL、FreeSolv、Lipo、BBBP、BACE 和 ClinTox 作为核心 endpoint。Therapeutics Data Commons 提供了面向药物研发任务的数据和 benchmark，本文使用八个 TDC ADMET endpoint，并进一步扩展到 22 个 PyTDC ADMET/Tox endpoint 的 full-panel appendix benchmark。")
    add_para(doc, "图神经网络、D-MPNN/Chemprop、分子指纹、RDKit 描述符和树模型分别捕获不同结构信号。分子预训练编码器如 ChemBERTa 和 MoLFormer 能提供互补表示，但近期 benchmark 显示，大模型或预训练模型并不自动在所有 drug-discovery endpoint 上胜出。因此本文将它们作为 validation-selected expert pool 中的候选专家。")
    add_para(doc, "适用域、校准、活性悬崖、早期富集和 motif 解释是分子预测从 benchmark 走向实际药物发现的重要条件。FZYC-Mol 将这些可靠性和解释性模块纳入统一实验闭环。")

    doc.add_heading("3. 方法", level=1)
    add_para(doc, "给定分子数据集 D = {(x_i, y_i)}，其中 x_i 为 SMILES 或分子图，y_i 为连续性质或二分类标签。FZYC-Mol 考虑多个专家 f_1, f_2, ..., f_M，并学习组合策略 g(f_1(x), ..., f_M(x))。g 可以是最佳验证专家、固定 consensus、validation stacking、adaptive consensus 或适用域加权 consensus。关键约束是：g 只由 validation split 选择，test split 不参与策略选择。")
    add_csv_table(
        doc,
        "表 1",
        "数据集与评估协议。",
        MAIN_TABLES / "Table_1_Dataset_protocol.csv",
        key_cols=["dataset", "source", "task_type"],
        chunk_size=5,
        font_size=6.4,
    )
    add_para(doc, "FZYC-Mol 包含传统 Morgan 指纹模型、多指纹树模型、GIN/D-MPNN/FZYC graph 模型、Chemprop D-MPNN 与 Chemprop-RDKit、冻结 ChemBERTa/MoLFormer embedding heads、descriptor MLP、BRICS/Murcko/functional-group motif experts，以及不确定性和适用域加权模块。")
    add_para(doc, "在定位到 frozen pretrained linear heads 表现较弱后，本文新增 pretrained rescue heads：保持 ChemBERTa/MoLFormer embedding 冻结，将其与 RDKit descriptors 拼接，并训练 ExtraTrees/LightGBM。随后将 rescue-aware candidates 接回完整 validation-only selector pool，只有验证集支持时才采用。")

    doc.add_heading("4. 实验设置", level=1)
    add_para(doc, "所有实验均为公开数据集上的计算实验，不包含湿实验。RDKit 用于分子解析、描述符、指纹、scaffold 和官能团识别。传统模型基于 scikit-learn、XGBoost 和 LightGBM；神经模型使用 PyTorch 相关实现；Chemprop baseline 使用官方 Chemprop workflow；TDC 数据和官方 split 通过 PyTDC 获取。")

    doc.add_heading("5. 结果", level=1)
    doc.add_heading("5.1 MoleculeNet 主性能", level=2)
    add_csv_table(
        doc,
        "表 2",
        "MoleculeNet 主结果。该表同时保留原 FZYC-Mol validation selector 和 conservative targeted rescue selector；后者仅在 Lipo 上通过 validation-only integration 采用 rescue-aware candidate。",
        MAIN_TABLES / "Table_2_MoleculeNet_main_results.csv",
        key_cols=["dataset", "task_type", "primary_metric"],
        chunk_size=5,
        font_size=6.8,
    )
    add_figure(doc, MAIN_FIGS / "Figure_2_MoleculeNet_model_family_ranks.png", "图 2. MoleculeNet endpoint 内模型家族 rank heatmap。")
    add_figure(doc, MAIN_FIGS / "Figure_3_MoleculeNet_main_performance.png", "图 3. MoleculeNet 主性能散点图。")
    add_para(doc, "Table 2 和 Figure 2/3 共同说明，分子 endpoint 之间存在明显模型偏好差异。ESOL、BACE 和 ClinTox 由 validation selector 表现最佳；FreeSolv 中 Chemprop-RDKit 仍然最强；BBBP 中 multi-fingerprint ExtraTrees 最强；Lipo 则由 rescue-integrated targeted selector 将 RMSE 降至 0.6835 +/- 0.0439。")

    doc.add_heading("5.2 TDC ADMET 迁移与官方 split", level=2)
    add_csv_table(
        doc,
        "表 3",
        "TDC ADMET 迁移与官方 PyTDC random/scaffold split 结果。",
        MAIN_TABLES / "Table_3_TDC_ADMET_official_splits.csv",
        key_cols=["dataset", "task_type", "primary_metric"],
        chunk_size=5,
        font_size=6.2,
    )
    add_figure(doc, MAIN_FIGS / "Figure_5_Official_TDC_ADMET_scaffold_delta.png", "图 5. 官方 PyTDC random-to-scaffold 性能变化。")
    add_para(doc, "TDC ADMET 结果显示，validation selector 在 HIA_Hou、Pgp_Broccatelli 和 BBB_Martins 上分别达到 ROC-AUC 0.9792、0.9357 和 0.9182。官方 split baseline 进一步显示 random split 到 scaffold split 的性能变化依赖 endpoint 和模型家族。")

    doc.add_heading("5.3 Split-realism 与结构分布偏移", level=2)
    add_csv_table(
        doc,
        "表 4",
        "随机划分、scaffold 划分和 structure-separated 划分下的 split-realism 结果。",
        MAIN_TABLES / "Table_4_Split_realism.csv",
        key_cols=["source", "dataset", "task_type", "metric"],
        chunk_size=5,
        font_size=6.4,
    )
    add_figure(doc, MAIN_FIGS / "Figure_4_Split_realism_structure_shift.png", "图 4. 结构分布偏移下的 split-realism 曲线。")
    add_para(doc, "Split-realism 结果显示，ESOL、BACE、TDC HIA_Hou、Bioavailability_Ma 和 Pgp_Broccatelli 等 endpoint 在结构更严格的划分下表现明显下降。总体结论是：分子预测论文不能只依赖 random split 报告性能。")

    doc.add_heading("5.4 不确定性、校准、适用域和 reconstruction unfamiliarity", level=2)
    add_csv_table(
        doc,
        "表 6",
        "可靠性和适用域分析。",
        MAIN_TABLES / "Table_6_Reliability_AD.csv",
        key_cols=["family", "score"],
        chunk_size=5,
        font_size=6.8,
    )
    add_figure(doc, MAIN_FIGS / "Figure_6_Uncertainty_AD_reliability.png", "图 6. 高风险分子预测的可靠性信号。")
    add_para(doc, "重构陌生度与绝对误差的平均 Spearman 相关为 0.2032，top-10% high-error enrichment 为 2.0767；hybrid reconstruction-plus-AD score 的平均 Spearman 相关达到 0.2225。该结果说明，基于训练分布特征流形的 reconstruction error 可以作为传统相似性适用域指标的补充。")

    doc.add_heading("5.5 活性悬崖、早期富集与 motif 解释", level=2)
    add_figure(doc, MAIN_FIGS / "Figure_7_Motif_fragment_interpretation.png", "图 7. Motif attribution 与 fragment-level 化学解释。")
    add_figure(doc, SUPP_FIGS / "Figure_S7_Enrichment_activity_cliffs.png", "图 S7. 早期富集和活性悬崖诊断。")
    add_para(doc, "Motif experts 在 BBBP、BACE 和 ClinTox 上分别达到 ROC-AUC 0.9133、0.8281 和 0.8592。BBBP 中极性片段、羟基、羧酸盐、内酰胺和芳香羟基等 motif 与血脑屏障穿透表现出合理关联；BACE 中卤素、芳基甲基、醚和疏水芳香片段较突出；ClinTox 中季铵、苯胺、芳香氮、羰基和稠环结构值得关注。")

    doc.add_heading("5.6 Ablation、显著性与 selector 行为", level=2)
    add_csv_table(
        doc,
        "表 5",
        "Ablation 与统计证据。",
        MAIN_TABLES / "Table_5_Ablation_significance.csv",
        key_cols=["section", "comparison"],
        chunk_size=5,
        font_size=6.3,
    )
    add_figure(doc, SUPP_FIGS / "Figure_S3_Ablation_significance.png", "图 S3. Ablation 与显著性分析。")
    add_figure(doc, SUPP_FIGS / "Figure_S4_Validation_selector_map.png", "图 S4. Validation selector 选择图谱。")
    add_para(doc, "Ablation 结果显示，单一 expert family 通常弱于完整 validation selector。与此同时，移除某些专家在个别 endpoint 上可能略有改善，说明候选池扩张本身并不必然带来性能提升。")

    doc.add_heading("5.7 Conformal prediction、risk coverage 与校准", level=2)
    add_figure(doc, SUPP_FIGS / "Figure_S1_Risk_coverage_curves.png", "图 S1. Risk-coverage curves。")
    add_figure(doc, SUPP_FIGS / "Figure_S2_Calibration_curves.png", "图 S2. 分类校准曲线。")
    add_figure(doc, SUPP_FIGS / "Figure_S5_Conformal_diagnostics.png", "图 S5. Conformal prediction 诊断。")
    add_csv_table(
        doc,
        "表 S6",
        "Reliability summary。",
        SUPP_TABLES / "Table_S6_Reliability_summary.csv",
        key_cols=["analysis", "dataset", "task_type"] if "analysis" in pd.read_csv(SUPP_TABLES / "Table_S6_Reliability_summary.csv", nrows=0).columns else None,
        chunk_size=5,
        font_size=5.8,
    )

    doc.add_heading("5.8 效率与实用性", level=2)
    add_csv_table(doc, "表 7", "计算效率。", MAIN_TABLES / "Table_7_Efficiency.csv", key_cols=["report", "model_family"], chunk_size=5, font_size=6.5)
    add_figure(doc, SUPP_FIGS / "Figure_S6_Efficiency_tradeoff.png", "图 S6. 性能与计算效率权衡。")
    add_para(doc, "效率结果支持 practical molecular AI 的叙事。框架并不要求每个 endpoint 都由大型神经网络解决，强传统模型与可靠性诊断结合即可形成较低成本的实验闭环。")

    doc.add_heading("5.9 外部 ADMET appendix 与 roughness 分析", level=2)
    add_csv_table(doc, "表 S8", "OpenADMET-ExpansionRx fast external appendix benchmark。", SUPP_TABLES / "Table_S8_OpenADMET_ExpansionRx_fast_external.csv", key_cols=["endpoint", "task_type"], chunk_size=5, font_size=6.0)
    add_csv_table(doc, "表 S9", "PyTDC full-panel ADMET/Tox fast appendix benchmark。", SUPP_TABLES / "Table_S9_TDC_full_panel_fast_appendix.csv", key_cols=["dataset", "family", "task_type"], chunk_size=5, font_size=5.6)
    add_csv_table(doc, "表 S10", "TDC performance-mode retained-best appendix。", SUPP_TABLES / "Table_S10_TDC_performance_mode_retained_best.csv", key_cols=["dataset", "family", "task_type", "official_metric"], chunk_size=5, font_size=5.4)
    add_csv_table(doc, "表 S11", "TDC roughness and literature-alignment appendix。", SUPP_TABLES / "Table_S11_TDC_roughness_literature_alignment.csv", key_cols=["dataset", "family", "task_type", "official_metric"], chunk_size=5, font_size=5.2)
    add_para(doc, "TDC performance-mode retained-best appendix 相比 full-panel fast appendix 改善了五个 endpoint：half_life_obach、vdss_lombardo、ppbr_az、clearance_microsome_az 和 clearance_hepatocyte_az。Roughness 诊断显示，高 local roughness 的 ADME regression endpoint 也是性能模式补强更容易奏效的 endpoint。")

    doc.add_heading("5.10 低分模型模块补强与 rescue-integrated selector", level=2)
    add_csv_table(doc, "表 S12", "MoleculeNet meta-pool selector diagnostic。", SUPP_TABLES / "Table_S12_MoleculeNet_meta_pool_selector.csv", key_cols=["dataset", "task_type"], chunk_size=5, font_size=5.8)
    add_csv_table(doc, "表 S13", "Pretrained rescue heads。", SUPP_TABLES / "Table_S13_Pretrained_rescue_heads.csv", key_cols=["dataset", "task_type", "primary_metric"], chunk_size=5, font_size=5.7)
    add_csv_table(doc, "表 S14", "MoleculeNet rescue-integrated selector。", SUPP_TABLES / "Table_S14_MoleculeNet_rescue_integrated_selector.csv", key_cols=["dataset", "task_type", "primary_metric"], chunk_size=5, font_size=5.7)
    add_para(doc, "低分模块分析显示，frozen pretrained linear heads 是当前最弱的 standalone family。Rescue heads 在 5/6 个 MoleculeNet endpoint 上超过原 best frozen-linear pretrained head。然而，真正决定主文是否采用的是 full selector integration：只有 Lipo 被 validation-only selector 接纳 rescue-aware candidate，RMSE 从 0.7078 降至 0.6835；其他 endpoint 保留原 selector。")

    doc.add_heading("6. 讨论", level=1)
    add_para(doc, "本文结果显示，不同 endpoint 的最优模型家族并不一致。Chemprop-RDKit 在 FreeSolv 上最强，multi-fingerprint ExtraTrees 在 BBBP 上最强，validation selector 在 ESOL、BACE 和 ClinTox 上最强，targeted rescue selector 在 Lipo 上最强。")
    add_para(doc, "Graph Transformer probes 和 frozen pretrained encoders 并未稳定超过强传统 baseline。Rescue-head 实验说明，弱的 pretrained module 可以通过轻量 nonlinear tabular heads 被修复，但 full validation-only integration 只在 Lipo 上接受该补强。因此，模型复杂度本身不是胜利条件，验证集控制和端点级选择才是关键。")
    add_para(doc, "Random split 往往高估真实泛化能力。Scaffold 和 structure-separated split 更接近 scaffold hopping、新化学系列和未探索化学空间中的实际应用场景。本文的 split-realism 结果显示，结构分布偏移会显著改变模型结论。")
    add_para(doc, "Motif attribution 和 fragment enrichment 将模型行为连接到可识别的化学结构单元。虽然这些解释不能直接证明因果机制，但可以帮助审稿人和药物化学用户理解模型为何在某些 endpoint 上给出特定预测。")

    doc.add_heading("7. 局限性", level=1)
    add_numbered(
        doc,
        [
            "本研究不包含湿实验验证，所有结论均来自公开数据集上的计算评估。",
            "Exact Polaris benchmark 未在本轮纳入，尽管官方 PyTDC ADMET split 和外部 ADMET appendix 已补强。",
            "TDC structure-separated stress test 的 seed 数少于主 MoleculeNet 实验。",
            "ChemBERTa 和 MoLFormer 主要以 frozen encoder 形式使用，full fine-tuning 可能改善部分 endpoint，但也会增加计算成本和过拟合风险。",
            "ClinTox 等小样本不平衡数据集存在 seed-level instability，ROC-AUC 应与 PR-AUC、Brier、ECE 和 conformal diagnostics 一起解释。",
            "Motif attribution 和 fragment enrichment 是关联解释，不能被解读为因果机制证明。",
        ],
    )

    doc.add_heading("8. 结论", level=1)
    add_para(doc, "本文提出 FZYC-Mol，一个验证集选择、适用域感知、多专家分子性质预测框架。FZYC-Mol 不依赖单一模型家族，而是整合图模型、Chemprop、分子指纹、描述符模型、motif experts、冻结预训练编码器、validation stacking、adaptive consensus、不确定性建模和适用域加权。所有最终策略选择均只使用 validation split，test split 保留用于最终评估。")
    add_para(doc, "跨 MoleculeNet、TDC ADMET、官方 PyTDC split、MoleculeACE、低相似度 hard subset 和 structure-separated split 的结果表明，FZYC-Mol 能够提供比单一 leaderboard score 更完整的可靠性证据。主表显示，框架在 ESOL、BACE、ClinTox 和 Lipo 等 endpoint 上通过不同 validation-accepted 策略取得强结果；外部 ADMET appendix 和 roughness 诊断进一步说明，强 tabular experts、target transform 和 validation-only ensemble 对困难 ADME regression endpoint 具有实际价值；rescue-integrated selector 则展示了如何在不破坏 validation-only 原则的前提下补强低分模块。")

    doc.add_heading("9. 图表对应关系总览", level=1)
    mapping = pd.DataFrame(
        [
            ["图文摘要", "图 1", "Figure_1_FZYC_Mol_framework.png", "框架总览"],
            ["方法", "表 1", "Table_1_Dataset_protocol.csv", "数据集与协议"],
            ["结果 5.1", "表 2", "Table_2_MoleculeNet_main_results.csv", "MoleculeNet 主结果"],
            ["结果 5.1", "图 2", "Figure_2_MoleculeNet_model_family_ranks.png", "MoleculeNet rank heatmap"],
            ["结果 5.1", "图 3", "Figure_3_MoleculeNet_main_performance.png", "MoleculeNet 性能散点"],
            ["结果 5.2", "表 3 / 图 5", "Table_3_TDC_ADMET_official_splits.csv / Figure_5_Official_TDC_ADMET_scaffold_delta.png", "TDC ADMET 官方 split"],
            ["结果 5.3", "表 4 / 图 4", "Table_4_Split_realism.csv / Figure_4_Split_realism_structure_shift.png", "结构分布偏移"],
            ["结果 5.4", "表 6 / 图 6", "Table_6_Reliability_AD.csv / Figure_6_Uncertainty_AD_reliability.png", "可靠性与适用域"],
            ["结果 5.5", "图 7 / 图 S7", "Figure_7_Motif_fragment_interpretation.png / Figure_S7_Enrichment_activity_cliffs.png", "motif 与 enrichment"],
            ["结果 5.9", "表 S8-S11", "supplementary_tables", "外部 ADMET、TDC full-panel、performance-mode、roughness"],
            ["结果 5.10", "表 S12-S14", "supplementary_tables", "meta-pool、pretrained rescue、rescue-integrated selector"],
        ],
        columns=["章节", "图/表", "文件", "对应内容"],
    )
    add_dataframe_table(doc, mapping, font_size=7.2, max_text=120)

    doc.add_heading("10. 参考文献草稿", level=1)
    add_numbered(
        doc,
        [
            "Wu Z, Ramsundar B, Feinberg EN, et al. MoleculeNet: a benchmark for molecular machine learning. Chemical Science, 2018.",
            "Huang K, Fu T, Gao W, et al. Therapeutics Data Commons: machine learning datasets and tasks for drug discovery and development. NeurIPS Datasets and Benchmarks, 2021.",
            "Yang K, Swanson K, Jin W, et al. Analyzing learned molecular representations for property prediction. Journal of Chemical Information and Modeling, 2019.",
            "Rogers D, Hahn M. Extended-connectivity fingerprints. Journal of Chemical Information and Modeling, 2010.",
            "Landrum G. RDKit: open-source cheminformatics software.",
            "Chithrananda S, Grand G, Ramsundar B. ChemBERTa: large-scale self-supervised pretraining for molecular property prediction.",
            "Ross J, Belgodere B, Chenthamarakshan V, et al. Large-scale chemical language representations capture molecular structure and properties. MoLFormer.",
            "van Tilborg D, Alenicheva A, Grisoni F. Exposing the limitations of molecular machine learning with activity cliffs. MoleculeACE.",
            "Therapeutics Data Commons and PyTDC benchmark documentation.",
            "Recent ADMET reliability, OOD, tabular foundation model, and molecular foundation model benchmark papers discussed in docs/recent_literature_competitive_improvement_20260531.md.",
        ],
    )
    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
