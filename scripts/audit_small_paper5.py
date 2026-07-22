from __future__ import annotations

import hashlib
import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "小论文-5.docx"
TRACKED = ROOT / "output" / "小论文-5_修订痕迹.docx"
SUPP = ROOT / "output" / "小论文-5_补充表.docx"
PACKAGE = ROOT / "output" / "小论文-5_图表包"
AUDIT = ROOT / "results" / "audits" / "small_paper_5_audit.json"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def xml_errors(path: Path) -> list[str]:
    errors = []
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad: errors.append(f"bad_member:{bad}")
        for name in archive.namelist():
            if name.endswith(".xml"):
                try: ElementTree.fromstring(archive.read(name))
                except Exception as exc: errors.append(f"{name}:{type(exc).__name__}")
    return errors


def audit() -> Path:
    doc = Document(DOCX)
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n" + "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    expected = [
        "配对固定分母遗憾增加 0.122",
        "层内置换 P=0.953",
        "MAE 0.112",
        "6,480 次拟合",
        "相对 Morgan-only 的配对效用增益为 0.343",
        "5 promoted 和 17 retained",
        "类别 1 覆盖仅为 0.673/0.807/0.832",
        "公开 release、Zenodo DOI 和独立第三方冷启动尚未完成",
    ]
    prohibited = [
        "win/tie/loss = 5/17/0",
        "分类错误风险的中位 AUROC 为 0.788",
        "回归高误差样本的中位 AUROC 为 0.652",
        "3×3的27个外层单元",
        "Caco2、HIA、Pgp是TDC主要晋级终点",
        "风险调整或稳定性规则普遍降低遗憾",
    ]
    captions = [p.text.strip() for p in doc.paragraphs if re.match(r"^图\s*\d+\s*\|", p.text.strip())]
    table_captions = [p.text.strip() for p in doc.paragraphs if re.match(r"^表\s*\d+\s*\|", p.text.strip())]
    figures = {}
    format_sets = {}
    for ext in ["png","svg","pdf","tiff"]:
        format_sets[ext] = {p.stem for p in (PACKAGE/"figures").glob(f"fig*.{ext}")}
    for path in sorted((PACKAGE/"figures").glob("fig*.png")):
        svg = PACKAGE/"figures"/f"{path.stem}.svg"
        with Image.open(path) as im:
            figures[path.stem] = {"width_px":im.width,"height_px":im.height,"bytes":path.stat().st_size,"svg_editable_text":"<text" in svg.read_text(encoding="utf-8")}
    shapes = [{"width_cm":s.width/360000,"height_cm":s.height/360000} for s in doc.inline_shapes]
    with zipfile.ZipFile(TRACKED) as z:
        tracked_xml = z.read("word/document.xml"); settings = z.read("word/settings.xml")
    tracked = {"insertions":len(re.findall(br"<w:ins(?:\s|>)",tracked_xml)),"deletions":len(re.findall(br"<w:del(?:\s|>)",tracked_xml)),"enabled":b"trackRevisions" in settings}
    report = {
        "docx":{"path":str(DOCX),"sha256":sha256(DOCX),"bytes":DOCX.stat().st_size},
        "tracked_docx":{"path":str(TRACKED),"sha256":sha256(TRACKED),"bytes":TRACKED.stat().st_size},
        "supplement":{"path":str(SUPP),"exists":SUPP.exists(),"bytes":SUPP.stat().st_size if SUPP.exists() else 0},
        "paragraphs":len(doc.paragraphs),"tables":len(doc.tables),"inline_shapes":len(doc.inline_shapes),
        "equations":sum(p.style and p.style.name=="Equation" for p in doc.paragraphs),
        "caption_numbers":[int(re.search(r"\d+",x).group()) for x in captions],
        "table_caption_numbers":[int(re.search(r"\d+",x).group()) for x in table_captions],
        "expected_text":{x:x in text for x in expected},"prohibited_text":{x:x in text for x in prohibited},
        "clean_xml_errors":xml_errors(DOCX),"tracked_xml_errors":xml_errors(TRACKED),"supp_xml_errors":xml_errors(SUPP),
        "tracked":tracked,"figure_formats":{k:len(v) for k,v in format_sets.items()},"figure_stems_match":len({frozenset(v) for v in format_sets.values()})==1,
        "figure_checks":figures,"inline_shape_cm":shapes,
        "source_data_csv":len(list((PACKAGE/"source_data").glob("*.csv"))),
        "supplement_tables":len(list((PACKAGE/"supplement_tables").glob("Table_S*"))),
        "fusion_audit_files":len(list((ROOT/"output"/"小论文-5_融合审计").glob("*"))),
    }
    report["passed"] = all([
        report["tables"]==8, report["inline_shapes"]==9, report["equations"]==7,
        report["caption_numbers"]==list(range(1,10)), report["table_caption_numbers"]==list(range(1,9)),
        all(report["expected_text"].values()), not any(report["prohibited_text"].values()),
        not report["clean_xml_errors"], not report["tracked_xml_errors"], not report["supp_xml_errors"],
        tracked["insertions"]>0, tracked["deletions"]>0, tracked["enabled"],
        all(v==9 for v in report["figure_formats"].values()), report["figure_stems_match"],
        all(v["svg_editable_text"] and v["width_px"]>=2000 and v["bytes"]>=50000 for v in figures.values()),
        all(v["width_cm"]<=16.35 and v["height_cm"]<=13.5 for v in shapes),
        report["source_data_csv"]>=45, report["supplement_tables"]>=10, report["fusion_audit_files"]>=4,
    ])
    AUDIT.parent.mkdir(parents=True,exist_ok=True)
    AUDIT.write_text(json.dumps(report,ensure_ascii=False,indent=2),encoding="utf-8")
    if not report["passed"]: raise SystemExit(json.dumps(report,ensure_ascii=False,indent=2))
    return AUDIT


if __name__=="__main__": print(audit())
