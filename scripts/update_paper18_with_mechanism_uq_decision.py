from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path("D:/fzyc")
OUT = ROOT / "output"
EXP = OUT / "sci1_mechanism_uq_decision_20260707"
SOURCE = max([p for p in OUT.glob("*.docx") if p.name.endswith("-17.docx")], key=lambda p: p.stat().st_mtime)
TARGET = OUT / "小论文-18.docx"
REPORT = OUT / "小论文-18_新增实验执行报告.md"
AUDIT = OUT / "paper18_mechanism_uq_decision_audit.json"


def find_paragraph(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def insert_before(anchor, text: str = "", style: str | None = None):
    paragraph = anchor.insert_paragraph_before(text)
    if style:
        paragraph.style = style
    return paragraph


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if edge is None:
            continue
        tag = "w:" + edge_name
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge.items():
            element.set(qn("w:" + key), str(value))


def make_three_line(table):
    visible = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    none = {"val": "nil"}
    rows = list(table.rows)
    for row_idx, row in enumerate(rows):
        for cell in row.cells:
            set_cell_border(cell, left=none, right=none, top=none, bottom=none)
            if row_idx == 0:
                set_cell_border(cell, top=visible, bottom=visible)
            if row_idx == len(rows) - 1:
                set_cell_border(cell, bottom=visible)


def add_table_before(doc: Document, anchor, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = value
    make_three_line(table)
    anchor._p.addprevious(table._tbl)
    return table


def values() -> dict[str, object]:
    mech = pd.read_csv(EXP / "mechanism_controlled_simulation_summary.csv")
    conf = pd.read_csv(EXP / "conformal_crossfold_summary.csv")
    cqr = pd.read_csv(EXP / "cqr_regression_summary.csv")
    cal = pd.read_csv(EXP / "calibration_ood_scaffold_summary.csv")
    clin = pd.read_csv(EXP / "clintox_minority_negative_result.csv")
    dec = pd.read_csv(EXP / "decision_enrichment_summary.csv")
    tox = pd.read_csv(EXP / "toxicity_false_negative_cost.csv")
    fails = pd.read_csv(EXP / "failure_case_category_summary.csv")

    def mech_value(regime: str, frac: float, k: int, column: str) -> float:
        row = mech[
            mech["correlation_regime"].eq(regime)
            & mech["validation_information_fraction"].eq(frac)
            & mech["candidate_count"].eq(k)
        ].iloc[0]
        return float(row[column])

    rdkit_conf = conf[(conf["candidate"].eq("rdkit_rf")) & (conf["alpha"].eq(0.10))]
    conf_group = rdkit_conf.groupby(["task_type", "method"]).agg(
        mean_coverage=("mean_coverage", "mean"),
        mean_class1=("mean_class_1_coverage", "mean"),
        mean_set_size=("mean_set_size", "mean"),
        mean_width=("mean_interval_width", "mean"),
    )
    cal_group = cal.groupby("tanimoto_bin").agg(
        mean_roc_auc=("mean_roc_auc", "mean"),
        mean_ece=("mean_ece", "mean"),
    )
    fzyc_dec = dec[dec["candidate"].eq("fzyc_selected")].groupby("budget_fraction").agg(
        enrichment=("mean_enrichment", "mean"),
        regret=("mean_regret_vs_oracle", "mean"),
    )
    mol_dec = dec[dec["candidate"].eq("molformer_linear_probe")].groupby("budget_fraction").agg(
        enrichment=("mean_enrichment", "mean")
    )
    tox_group = tox.groupby("candidate").agg(cost100=("cost_per_100_molecules", "mean"))
    clin_group = clin.groupby("candidate").agg(
        recall=("minority_recall", "first"),
        fnr=("minority_false_negative_rate", "first"),
        class1_cov=("mean_class_1_coverage", "mean"),
    )

    return {
        "high_25_k4": mech_value("high_correlated_lightweight", 0.25, 4, "fixed_k64_normalized_selection_loss"),
        "high_25_k64": mech_value("high_correlated_lightweight", 0.25, 64, "fixed_k64_normalized_selection_loss"),
        "mid_25_k4": mech_value("medium_correlated_multiview", 0.25, 4, "fixed_k64_normalized_selection_loss"),
        "mid_25_k64": mech_value("medium_correlated_multiview", 0.25, 64, "fixed_k64_normalized_selection_loss"),
        "low_25_k4": mech_value("low_correlated_deep_foundation", 0.25, 4, "fixed_k64_normalized_selection_loss"),
        "low_25_k64": mech_value("low_correlated_deep_foundation", 0.25, 64, "fixed_k64_normalized_selection_loss"),
        "high_100_k64": mech_value("high_correlated_lightweight", 1.0, 64, "fixed_k64_normalized_selection_loss"),
        "mid_100_k64": mech_value("medium_correlated_multiview", 1.0, 64, "fixed_k64_normalized_selection_loss"),
        "low_100_k64": mech_value("low_correlated_deep_foundation", 1.0, 64, "fixed_k64_normalized_selection_loss"),
        "rdkit_label_cov": float(conf_group.loc[("classification", "label_conditional_conformal"), "mean_coverage"]),
        "rdkit_mondrian_cov": float(conf_group.loc[("classification", "mondrian_label_similarity_conformal"), "mean_coverage"]),
        "rdkit_split_class1": float(conf_group.loc[("classification", "split_conformal"), "mean_class1"]),
        "rdkit_label_class1": float(conf_group.loc[("classification", "label_conditional_conformal"), "mean_class1"]),
        "rdkit_reg_cov": float(conf_group.loc[("regression", "split_conformal_residual"), "mean_coverage"]),
        "rdkit_reg_mondrian_cov": float(conf_group.loc[("regression", "mondrian_similarity_residual"), "mean_coverage"]),
        "cqr_90_cov": float(cqr[cqr["alpha"].eq(0.10)]["mean_coverage"].mean()),
        "cqr_90_width": float(cqr[cqr["alpha"].eq(0.10)]["mean_interval_width"].mean()),
        "low_sim_auc": float(cal_group.loc["<0.5", "mean_roc_auc"]),
        "high_sim_auc": float(cal_group.loc[">0.7", "mean_roc_auc"]),
        "low_sim_ece": float(cal_group.loc["<0.5", "mean_ece"]),
        "high_sim_ece": float(cal_group.loc[">0.7", "mean_ece"]),
        "clintox_rate": float(pd.read_csv(EXP / "clintox_minority_negative_result.csv")["minority_positive_rate"].iloc[0]),
        "rdkit_clintox_recall": float(clin_group.loc["rdkit_rf", "recall"]),
        "rdkit_clintox_fnr": float(clin_group.loc["rdkit_rf", "fnr"]),
        "rdkit_clintox_cov": float(clin_group.loc["rdkit_rf", "class1_cov"]),
        "gnn_clintox_recall": float(clin_group.loc["gnn_gcn", "recall"]),
        "fzyc_ef1": float(fzyc_dec.loc[0.01, "enrichment"]),
        "fzyc_ef5": float(fzyc_dec.loc[0.05, "enrichment"]),
        "fzyc_ef10": float(fzyc_dec.loc[0.10, "enrichment"]),
        "mol_ef1": float(mol_dec.loc[0.01, "enrichment"]),
        "mol_ef5": float(mol_dec.loc[0.05, "enrichment"]),
        "rdkit_tox_cost": float(tox_group.loc["rdkit_rf", "cost100"]),
        "mol_tox_cost": float(tox_group.loc["molformer_linear_probe", "cost100"]),
        "gnn_tox_cost": float(tox_group.loc["gnn_gcn", "cost100"]),
        "failure_categories": int(fails["category"].nunique()),
        "failure_pool_cases": int(fails["n_cases"].sum()),
    }


def revise_docx(v: dict[str, object]) -> None:
    shutil.copy2(SOURCE, TARGET)
    doc = Document(str(TARGET))
    anchor = find_paragraph(doc, "4 讨论")

    insert_before(anchor, "3.11 新增机制、可靠性与决策价值实验", "Heading 2")
    insert_before(
        anchor,
        (
            "为进一步检验候选池扩张机制、可靠性边界、真实筛选价值和化学边界系统化四个关键问题，本研究新增一组"
            "补充机制与边界实验，并将其与既有同划分强基线预测逐样本连接。机制实验采用以真实结果为锚的 controlled "
            "simulation，在验证信息量 25%、50%、75% 和 100%、候选相关性三档以及 K=4、8、16、32、64 的完整"
            "网格上分解 selection loss。该实验不作为模型排行榜，而用于检验固定验证样本量下的排序噪声如何随 K "
            "和候选相关性共同变化。"
        ),
    )
    insert_before(
        anchor,
        (
            f"在 25% 验证信息下，K=4 到 K=64 使 fixed-scale selection loss 在高相关轻量池中由 "
            f"{v['high_25_k4']:.3f} 增至 {v['high_25_k64']:.3f}，在中等相关多视图池中由 "
            f"{v['mid_25_k4']:.3f} 增至 {v['mid_25_k64']:.3f}，在低相关深度/基础模型池中由 "
            f"{v['low_25_k4']:.3f} 增至 {v['low_25_k64']:.3f}。当验证信息升至 100% 时，K=64 的对应损失"
            f"降为 {v['high_100_k64']:.3f}、{v['mid_100_k64']:.3f} 和 {v['low_100_k64']:.3f}。这说明 selection "
            "loss 不是偶然数值，而是由候选规模、候选相关性和验证样本量共同决定的统计后果。"
        ),
    )
    p = insert_before(anchor)
    p.add_run().add_picture(str(EXP / "fig_mechanism_selection_loss.png"), width=Inches(6.5))
    insert_before(
        anchor,
        "图 10  验证集大小、候选相关性与候选规模共同决定 selection loss。曲线来自真实结果锚定的 controlled mechanism experiment；纵轴采用 K=64 固定尺度，避免随 K 改变分母。",
    )

    insert_before(
        anchor,
        (
            "可靠性实验从总体覆盖扩展到标签条件、Mondrian 相似度分层、CQR、ensemble uncertainty 和 scaffold/OOD "
            "校准。RDKit-RF 在 90% 目标覆盖下的分类 split conformal 总覆盖接近目标，但类别 1 覆盖仅为 "
            f"{v['rdkit_split_class1']:.3f}；label-conditional 和 Mondrian label-similarity conformal 将类别 1 "
            f"覆盖提高到 {v['rdkit_label_class1']:.3f} 左右，同时总体覆盖分别为 {v['rdkit_label_cov']:.3f} 和 "
            f"{v['rdkit_mondrian_cov']:.3f}。回归 split conformal 和 Mondrian residual 覆盖为 "
            f"{v['rdkit_reg_cov']:.3f} 和 {v['rdkit_reg_mondrian_cov']:.3f}；CQR 的 90% 平均覆盖为 "
            f"{v['cqr_90_cov']:.3f}，但平均区间宽度为 {v['cqr_90_width']:.2f}，提示其在当前特征和样本量下并未优于"
            "残差式保形。"
        ),
    )
    insert_before(
        anchor,
        (
            f"scaffold/OOD 校准进一步显示，最近邻 Tanimoto <0.5 的分类子集平均 ROC-AUC 为 {v['low_sim_auc']:.3f}，"
            f"低于 >0.7 子集的 {v['high_sim_auc']:.3f}；对应 ECE 为 {v['low_sim_ece']:.3f} 与 "
            f"{v['high_sim_ece']:.3f}。ClinTox 少数类阳性率仅为 {v['clintox_rate']:.3f}。RDKit-RF 的少数类召回为 "
            f"{v['rdkit_clintox_recall']:.3f}，假阴性率为 {v['rdkit_clintox_fnr']:.3f}，即使保形类别 1 覆盖可达 "
            f"{v['rdkit_clintox_cov']:.3f}，阈值式毒性筛选仍构成负结果；GNN-GCN 召回为 {v['gnn_clintox_recall']:.3f}，"
            "但假阳性成本显著升高。"
        ),
    )
    p = insert_before(anchor)
    p.add_run().add_picture(str(EXP / "fig_conformal_coverage.png"), width=Inches(6.2))
    insert_before(
        anchor,
        "图 11  90% 目标覆盖下的保形加厚结果。标签条件和 Mondrian 版本主要改善少数类覆盖；CQR 未在当前回归设置中形成一致优势。",
    )

    insert_before(
        anchor,
        (
            "真实决策价值实验将模型分数转化为固定预算筛选收益。FZYC-selected 在该六任务强基线面板中等同于 RDKit-RF，"
            f"其 top-1%、top-5% 和 top-10% 平均富集分别为 {v['fzyc_ef1']:.2f}、{v['fzyc_ef5']:.2f} 和 "
            f"{v['fzyc_ef10']:.2f}；MoLFormer 对应为 {v['mol_ef1']:.2f} 和 {v['mol_ef5']:.2f}。然而毒性假阴性"
            f"成本给出不同结论：RDKit-RF/FZYC 的 ClinTox 阈值成本为每 100 个分子 {v['rdkit_tox_cost']:.1f}，"
            f"MoLFormer 为 {v['mol_tox_cost']:.1f}，GNN-GCN 为 {v['gnn_tox_cost']:.1f}。因此，同一模型在富集筛选"
            "和毒性排除中可能具有相反的风险排序。"
        ),
    )
    p = insert_before(anchor)
    p.add_run().add_picture(str(EXP / "fig_decision_enrichment.png"), width=Inches(6.2))
    insert_before(
        anchor,
        "图 12  固定实验预算下的筛选富集。FZYC-selected 在本面板中选择 RDKit-RF，因此图中以单条线报告；毒性成本另见新增源数据表。",
    )

    insert_before(
        anchor,
        (
            f"化学边界失败案例被统一为 {v['failure_categories']} 类，包括 activity-cliff pair、最近邻 Tanimoto 分层、"
            "scaffold novelty、bRo5 外缘、极端标签和少数类假阴性。新增失败案例池记录 3-5 个代表结构、SMILES、真实值、"
            "预测值、最近邻相似度、骨架状态和错误解释；其目的不是再次计算平均性能，而是为读者提供可复核的失败机制。"
        ),
    )
    insert_before(anchor, "表 10 | 新增机制、可靠性、决策价值和失败案例实验摘要")
    add_table_before(
        doc,
        anchor,
        ["新增实验", "完成范围", "主要结论", "证据边界"],
        [
            [
                "验证集大小×候选相关性×K",
                "4×3×5 controlled grid；真实结果锚定",
                f"K=64、25% 验证信息时 fixed-scale selection loss 为 {v['high_25_k64']:.3f}/{v['mid_25_k64']:.3f}/{v['low_25_k64']:.3f}",
                "机制实验，不作为模型排行榜",
            ],
            [
                "标签条件、Mondrian 与 CQR",
                "6 任务、4 候选、5 种子、3 外层",
                f"分类少数类覆盖由 split {v['rdkit_split_class1']:.3f} 提升至约 {v['rdkit_label_class1']:.3f}",
                "CQR 未优于残差式保形",
            ],
            [
                "OOD/scaffold 校准与 ClinTox 负结果",
                "最近邻 Tanimoto 与 scaffold novelty 分层",
                f"低相似子集 ROC-AUC {v['low_sim_auc']:.3f}，高相似子集 {v['high_sim_auc']:.3f}",
                "ClinTox 阈值召回仍不足",
            ],
            [
                "固定预算筛选价值",
                "top-1/5/10%、毒性成本、队列模拟",
                f"FZYC top-1% 富集 {v['fzyc_ef1']:.2f}；毒性成本 {v['rdkit_tox_cost']:.1f}/100 molecules",
                "富集与毒性排除目标不完全一致",
            ],
            [
                "系统失败案例",
                f"{v['failure_categories']} 类失败；代表结构与错误原因",
                "活性悬崖、bRo5 外缘、低相似和极端标签被分开归档",
                "用于边界解释，不用于平均性能外推",
            ],
        ],
    )
    insert_before(
        anchor,
        (
            "所有新增实验的逐行结果、审计 JSON、PNG/SVG 图和代表结构表保存在 "
            f"{EXP.as_posix()}。这些结果将论文的证据链从平均性能比较扩展到模型选择机制、校准边界和实验排队价值，"
            "但不改变本文的限定性结论：FZYC-Mol 是冻结治理协议，而不是新的通用最优预测模型。"
        ),
    )

    doc.save(str(TARGET))


def audit_docx() -> dict[str, object]:
    with ZipFile(TARGET) as zf:
        bad = zf.testzip()
    doc = Document(str(TARGET))
    text = "\n".join(p.text for p in doc.paragraphs)
    checks = {
        "zip_ok": bad is None,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "figures": len(doc.inline_shapes),
        "mechanism_written": "验证集大小、候选相关性与候选规模共同决定 selection loss" in text,
        "conformal_written": "label-conditional" in text and "Mondrian" in text and "CQR" in text,
        "decision_written": "固定实验预算下的筛选富集" in text,
        "failure_written": "activity-cliff pair" in text and "bRo5 外缘" in text,
    }
    checks["passed"] = all(v for k, v in checks.items() if isinstance(v, bool))
    return checks


def write_report(v: dict[str, object], checks: dict[str, object]) -> None:
    lines = [
        "# 小论文-18 新增实验执行报告",
        "",
        f"- 生成时间：{datetime.now().isoformat(timespec='seconds')}",
        f"- 输入文档：`{SOURCE}`",
        f"- 输出文档：`{TARGET}`",
        f"- 新实验目录：`{EXP}`",
        "",
        "## 已补实验",
        "",
        "| 模块 | 完成情况 | 写入位置 |",
        "|---|---|---|",
        f"| 机制实验 | 4×3×5 grid；K=4–64；三档候选相关性 | 3.11、图10、表10 |",
        f"| UQ/保形 | label-conditional、Mondrian、CQR、ensemble uncertainty、OOD/scaffold calibration | 3.11、图11、表10 |",
        f"| 决策价值 | top-1/5/10 enrichment、fixed-budget utility、toxicity false-negative cost、queue simulation | 3.11、图12、表10 |",
        f"| 化学边界失败 | {v['failure_categories']} 类失败机制；代表结构与原因 | 3.11、表10、source data |",
        "",
        "## 关键数值",
        "",
        f"- 25% 验证信息、K=64：高/中/低相关池 fixed-scale selection loss = {v['high_25_k64']:.3f}/{v['mid_25_k64']:.3f}/{v['low_25_k64']:.3f}。",
        f"- RDKit-RF 90% 分类保形：split 类别1覆盖 {v['rdkit_split_class1']:.3f}，label-conditional 类别1覆盖 {v['rdkit_label_class1']:.3f}。",
        f"- Tanimoto <0.5 分类子集 ROC-AUC {v['low_sim_auc']:.3f}，>0.7 子集 {v['high_sim_auc']:.3f}。",
        f"- FZYC/RDKit top-1/5/10% 富集 {v['fzyc_ef1']:.2f}/{v['fzyc_ef5']:.2f}/{v['fzyc_ef10']:.2f}。",
        f"- ClinTox RDKit-RF 少数类召回 {v['rdkit_clintox_recall']:.3f}，FNR {v['rdkit_clintox_fnr']:.3f}，作为明确负结果保留。",
        "",
        "## 审计",
        "",
        "```json",
        json.dumps(checks, ensure_ascii=False, indent=2),
        "```",
    ]
    REPORT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    v = values()
    revise_docx(v)
    checks = audit_docx()
    AUDIT.write_text(json.dumps({"values": v, "checks": checks}, ensure_ascii=False, indent=2), encoding="utf-8")
    write_report(v, checks)
    print(json.dumps({"target": str(TARGET), "report": str(REPORT), "passed": checks["passed"]}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
