# -*- coding: utf-8 -*-
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports" / "draft5_nature_revision_20260616"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def locate_source() -> Path:
    candidates = [p for p in (Path.home() / "Desktop").rglob("初稿-4.docx")]
    if not candidates:
        raise FileNotFoundError("Could not locate 初稿-4.docx.")
    return max(candidates, key=lambda p: p.stat().st_mtime)


SRC_DOCX = locate_source()
DEST_DOCX = SRC_DOCX.parent / "初稿-5.docx"
REPORT_DOCX = SRC_DOCX.parent / "初稿-5_Nature技能修订说明.docx"


def set_run_font(run) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run)


def replace_by_prefix(doc, prefix: str, new_text: str, changes: list[str], label: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            set_paragraph_text(paragraph, new_text)
            changes.append(label)
            return
    changes.append(f"未匹配：{label}")


def set_nth_nonempty_after_heading(doc, heading: str, nth: int, new_text: str) -> bool:
    """Replace the nth non-empty paragraph after a heading, stopping at the next numbered heading."""
    found = False
    count = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not found:
            if text == heading:
                found = True
            continue
        if text and text[0].isdigit() and " " in text and count > 0:
            break
        if text:
            count += 1
            if count == nth:
                set_paragraph_text(paragraph, new_text)
                return True
    return False


def set_border(element, edge: str, val: str, size: str = "8", color: str = "000000") -> None:
    tag = "w:tblBorders" if element.tag.endswith("tblPr") else "w:tcBorders"
    borders = element.find(qn(tag))
    if borders is None:
        borders = OxmlElement(tag)
        element.append(borders)
    edge_el = borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        borders.append(edge_el)
    edge_el.set(qn("w:val"), val)
    edge_el.set(qn("w:sz"), size)
    edge_el.set(qn("w:space"), "0")
    edge_el.set(qn("w:color"), color)


def apply_three_line_table(table) -> None:
    tbl_pr = table._tbl.tblPr
    for edge in ["top", "bottom"]:
        set_border(tbl_pr, edge, "single", "12")
    for edge in ["left", "right", "insideH", "insideV"]:
        set_border(tbl_pr, edge, "nil")
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for edge in ["left", "right", "insideV"]:
                set_border(tc_pr, edge, "nil")
    if table.rows:
        for cell in table.rows[0].cells:
            set_border(cell._tc.get_or_add_tcPr(), "bottom", "single", "8")


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_internal_review_section(doc) -> list[str]:
    removed: list[str] = []
    in_review = False
    paragraphs_to_remove = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "6 投稿前自我审查":
            in_review = True
        if in_review and text == "List of abbreviations":
            break
        if in_review:
            paragraphs_to_remove.append(paragraph)

    for paragraph in paragraphs_to_remove:
        if paragraph.text.strip():
            removed.append(paragraph.text.strip())
        remove_element(paragraph._element)

    for table in list(doc.tables):
        if not table.rows:
            continue
        first_cell = table.rows[0].cells[0].text.strip()
        if first_cell in {"主张", "风险点"}:
            remove_element(table._element)
            removed.append(f"移出内部自审表：{first_cell}")

    return removed


def polish_nature_style(doc) -> list[str]:
    changes: list[str] = []

    replacements = [
        (
            "FZYC-Mol：",
            "FZYC-Mol：面向可靠性审计的验证集治理分子性质预测框架",
            "题名：压缩为模型名、核心方法和应用对象",
        ),
        (
            "分子性质预测是药物发现、ADMET 评估",
            "分子性质预测已成为药物发现和 ADMET 评估的基础工具，但公开基准上的高平均分并不总能转化为可靠的化学决策。随机划分、低相似度分子、不平衡毒性标签、规则五以外化学空间和活性悬崖，都会削弱单一 ROC-AUC 或 RMSE 对真实应用风险的解释力。本研究提出 FZYC-Mol，一种由验证集治理驱动的适用域感知框架。该框架将分子图、指纹、二维描述符、片段/骨架信息和冻结预训练表征登记为候选证据，并仅允许通过验证集门控的专家模型、融合策略、补救头和适用域模块进入最终测试。",
            "摘要：按 context/problem-gap-approach 重写",
        ),
        (
            "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上表现稳定",
            "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别取得 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302；在 22 个外部 TDC ADMET 终点中，最终保留结果为 win/tie/loss = 5/17/0。风险评分、保形预测、验证-测试排名审计、低相似度分层和 MoleculeACE 活性悬崖分析显示，该框架的主要价值不是在所有终点上追求统一增益，而是使模型选择、适用域边界和失败模式能够被追踪和复核。",
            "摘要：集中关键结果并以意义句收束",
        ),
        (
            "分子性质预测已成为药物发现早期筛选",
            "分子性质预测正在从药物发现早期筛选工具转变为候选化合物优先级排序、ADMET 预警和毒性风险筛查中的决策组件。传统 QSAR、分子指纹和树模型仍具有可解释性与计算效率优势，图神经网络、消息传递模型和分子预训练模型则扩展了结构表征能力。然而，模型规模和公开基准分数并不能直接回答实际应用中的关键问题：预测何时可信，何时应拒用，何时需要实验或人工复核。",
            "引言：第一段从领域意义转入可靠性问题",
        ),
        (
            "现有分子机器学习研究多以固定数据集",
            "这一缺口来自评价方式与应用场景之间的错位。随机划分可能保留训练集和测试集之间的近邻关系，从而低估骨架外推难度；不平衡毒性终点上的高 ROC-AUC 可能掩盖阳性召回、概率校准或筛选富集不足；活性悬崖、高粗糙度终点和 bRo5 化学空间中的局部性质跳变，也会使平均性能指标难以支持样本级判断。因此，可靠的分子性质预测不仅需要比较模型得分，还需要审计模型选择过程、适用域边界和失败模式。",
            "引言：明确 gap，并减少教科书式铺垫",
        ),
        (
            "针对上述问题，本文提出 FZYC-Mol",
            "FZYC-Mol 针对这一问题将多专家候选池、验证集选择、适用域门控、校准和风险证据输出整合到同一流程中。与依赖单一主干模型的设计不同，FZYC-Mol 先登记候选策略，再由验证集决定接受、拒绝或保留状态。测试集仅在策略、权重、阈值和平局规则冻结后用于一次性报告，以降低候选池扩展带来的隐性测试集泄漏和事后选择偏差。",
            "引言：突出方法与偏倚控制，而非堆叠模块",
        ),
        (
            "本文的主要贡献包括五个方面",
            "本研究的贡献集中在三个层面。第一，建立覆盖 MoleculeNet、TDC ADMET、bRo5 公共压力测试、MoleculeACE 活性悬崖和多种划分策略的评测流程。第二，将强基线、图模型、冻结分子表征、多视图融合、定向补救头和适用域模块纳入同一验证集治理框架。第三，通过风险评分、保形预测、验证-测试排名审计、消融实验和负结果记录，给出可复核的可靠性证据链，并将基序、片段和最近邻案例限定为关联性解释证据。",
            "引言：将贡献从五项列表压缩为三层证据结构",
        ),
        (
            "本文使用 MoleculeNet、Therapeutics Data Commons ADMET",
            "本研究使用 MoleculeNet、Therapeutics Data Commons ADMET、MoleculeACE 以及 bRo5 相关公开数据作为评估基础。MoleculeNet 主面板包括 ESOL、FreeSolv、Lipophilicity、BBBP、BACE 和 ClinTox；TDC ADMET 用于检验外部 ADMET 终点上的可迁移性；MoleculeACE 用于评估活性悬崖和局部结构-性质非连续性；bRo5 相关数据用于规则五以外化学空间的公共压力测试。bRo5 结果仅用于观察可获得公开数据上的外推边界，不被表述为独立盲测验证。",
            "方法：数据来源与边界表述更可复核",
        ),
        (
            "为模拟不同真实应用难度",
            "为模拟不同应用难度，本研究采用多层次划分策略。MoleculeNet 任务以 scaffold split 为主要划分方式，并补充 random split、structure-separated split 和 low-similarity hard subset。TDC ADMET 任务使用官方 PyTDC 划分，并补充全终点 scaffold 审计；MoleculeACE 采用任务预设划分评估活性悬崖样本；bRo5 公共压力测试用于观察规则五以外分子空间中的误差、校准和适用域变化。所有划分、随机种子和候选登记均在测试集评估前固定。",
            "方法：删除冗余表达，保留可重复性信息",
        ),
        (
            "候选专家池包括 RF",
            "候选专家池包括 RF、ExtraTrees、XGBoost、LightGBM、CatBoost、Chemprop D-MPNN、图模型、描述符 MLP、基序专家、冻结预训练表征头、Top-K 均值、岭回归或逻辑回归堆叠、自适应共识、不确定性加权融合、适用域门控和定向补救头。回归任务额外比较原始目标、log1p、分位数正态化和截尾目标等目标变换；类别不平衡任务额外比较类别权重、欠采样集成和阈值移动，并报告 PR-AUC、Brier、ECE、MCC 和 balanced accuracy 等辅助指标。",
            "方法：术语中文化并强化任务对应关系",
        ),
        (
            "FZYC-Mol 的验证集治理流程包括六个步骤",
            "FZYC-Mol 的验证集治理流程包括六个预先定义的步骤。首先，固定每个终点的数据划分、主指标、指标方向、随机种子和候选专家集合。其次，仅使用训练集拟合单专家模型。第三，在验证集上生成专家预测矩阵，并在同一矩阵上构造 Top-K 均值、堆叠集成、自适应共识、不确定性加权融合和适用域门控。第四，根据验证集主指标、稳定性、风险调整和复杂度惩罚确定保留策略。第五，冻结候选权重、阈值、平局规则和补救头。第六，在测试集上进行一次性评估，并将未通过门控的候选作为负结果记录。",
            "方法：验证集治理写成可执行流程",
        ),
        (
            "跨随机种子的结果以均值",
            "跨随机种子的结果以均值 ± 标准差表示。候选之间的差异通过随机种子级配对差值、bootstrap 置信区间和 Wilcoxon signed-rank test 评估。验证集选择器风险通过验证-测试 Spearman 排名相关、Top-1 一致率、Top-3 命中率、regret 和 optimism gap 审计。对于 BBBP、BACE、ClinTox、HIA、Pgp、ESOL、FreeSolv、Lipophilicity 和 Caco2 等代表性终点，本研究补充 3 outer × 3 inner nested validation，用于评估候选选择对验证集波动的敏感性。",
            "方法：统计分析术语统一",
        ),
        (
            "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL",
            "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别取得 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302。ESOL 和 BACE 中，验证集选择器与测试观测最优结果基本一致。Lipophilicity 中，验证集接受的定向补救策略将 RMSE 从 0.7078 ± 0.0389 降至 0.6835 ± 0.0439。",
            "结果：仅报告观察到的主结果，解释移至讨论",
        ),
        (
            "FreeSolv 呈现出不同模式",
            "FreeSolv 呈现出不同模式。低成本重构候选将 RMSE 降至 1.0286 ± 0.1761，缩小了与当前选择器之间的差距，但仍未超过观测最佳 Chemprop 候选 0.9518 ± 0.1314。因此，FreeSolv 被保留为物理相互作用相关任务的边界案例。",
            "结果：保留负结果，删去过度解释句",
        ),
        (
            "分类任务中，BBBP 和 ClinTox",
            "分类任务中，BBBP 和 ClinTox 的多方法融合候选进入最终保留策略。BBBP 的最终 ROC-AUC 为 0.9243 ± 0.0247，ClinTox 的最终 ROC-AUC 为 0.9496 ± 0.0262。由于 ClinTox 阳性样本稀少，正文同时报告 PR-AUC、校准指标和固定精度条件下的召回，以降低单一 ROC-AUC 指标造成的解释偏差。",
            "结果：补充 ClinTox 结果解释的必要边界",
        ),
        (
            "为避免结论依赖弱 基线模型",
            "为避免结论依赖弱基线模型，本研究将 CatBoost、XGBoost、LightGBM、ExtraTrees、RF、Chemprop、TabPFNv2-RDKit、AutoGluon-RDKit、XGBoost-RDKit/Mordred/MorganCount 以及可获得的 KPGT 表征纳入同一候选登记。每个数据集-随机种子单元均输出验证集指标、测试集指标、候选排名、最终保留策略和 regret。结果表明，FZYC-Mol 的目标不是在每个终点上击败所有强基线，而是在强基线进入候选池后透明地接受、拒绝或保留候选策略。",
            "结果：修复术语空格与中英混排",
        ),
        (
            "在外部 TDC ADMET 任务上",
            "在外部 TDC ADMET 任务上，FZYC-Mol 显示出选择性增益。多视图融合与适用域门控在 Caco2_Wang、HIA_Hou 和 Pgp_Broccatelli 等任务上改善了验证集接受的最终策略；在多数终点中，验证集未支持更复杂候选替代原策略。22 个外部终点的最终保留结果为 win/tie/loss = 5/17/0，其中 5 个终点由性能增强候选进入最终保留，17 个终点保留原结果。",
            "结果：删去讨论式定位句，保留数据事实",
        ),
        (
            "可靠性分析显示",
            "可靠性分析显示，风险分数识别分类错误的能力高于识别回归高误差样本的能力。分类任务中，风险分数识别错误样本的中位 AUROC 为 0.788；回归任务中，识别高误差样本的中位 AUROC 为 0.652。该结果将风险分数限定为辅助可靠性证据，而不是可独立识别所有高误差样本的机制。",
            "结果：压缩可靠性结果并匹配证据强度",
        ),
        (
            "保形预测结果显示",
            "保形预测结果显示，在 MoleculeNet 分类任务 80%、90% 和 95% 目标覆盖率下，平均经验覆盖率分别为 0.814、0.918 和 0.956；在回归任务中，对应平均经验覆盖率分别为 0.823、0.925 和 0.962。类别不平衡终点同时报告 PR-AUC、Brier、ECE、MCC、balanced accuracy 和固定精度条件下的 recall。",
            "结果：去除重复意义解释",
        ),
        (
            "验证-测试排名审计显示",
            "验证-测试排名审计显示，跨 200 个数据集-随机种子候选池的中位 Spearman 相关为 0.667，测试最佳候选落入验证 Top-3 的比例为 0.295，Top-1 一致比例为 0.135。该结果表明，验证集排序与测试集排序存在相关性，但不足以保证验证集第一名就是测试集最佳候选。因此，本研究将 regret、optimism gap、nested validation 和负结果审计作为性能分数之外的必要补充。",
            "结果：修复空格并给出边界化解释",
        ),
        (
            "在 OOD、Perimeter 和低相似度压力测试中",
            "在 OOD、Perimeter 和低相似度压力测试中，模型性能通常随测试分子与训练集最近邻相似度下降而降低。适用域距离、风险分位和最近邻标签差异识别出部分高风险样本，但未完全消除结构外推误差。MoleculeACE 活性悬崖分析显示，相似分子之间的真实标签差异仍可能被平均指标低估；cliff-pair RMSE、gap Spearman 和代表性失败案例共同支持这一观察。",
            "结果：将可靠性结论写成观察而非自动纠错主张",
        ),
        (
            "消融实验和负结果审计表明",
            "消融实验和负结果审计显示，并非所有复杂模块都稳定带来收益。补救头加入完整候选池后，只有 Lipophilicity 被验证集接受，其他终点保持原保留策略。部分 3D-lite 物理代理、轻量适配器、目标变换或粗糙度加权候选未通过验证集门控，因此作为负结果保留。样本级案例显示，Lipophilicity 体现了终点级补救的验证集接受过程；ClinTox 提示高 ROC-AUC 仍需结合高风险假阴性分析；FreeSolv 和高粗糙度 ADME 回归任务则暴露了二维表示模型在物理相互作用和局部目标跳变中的边界。",
            "结果：负结果段落压实为三类证据",
        ),
        (
            "本研究表明，分子性质预测模型的价值",
            "本研究将分子性质预测的评价重点从单一排行榜分数转向模型选择过程、结构外推能力、适用域边界、可靠性输出和失败模式。FZYC-Mol 的主要贡献不在于扩大模型规模，而在于把多专家候选、验证集选择、外部 ADMET 评估、可靠性分析和化学解释纳入同一证据链。",
            "讨论：开头改为本文对领域的增量",
        ),
        (
            "与单一 backbone 模型不同",
            "与单一主干模型不同，FZYC-Mol 将不同模型家族视为候选专家，并通过验证集证据决定其接受、拒绝或保留状态。这一设计使正结果和负结果都能进入同一审计框架。Lipophilicity 的定向补救通过验证集门控进入最终结果，FreeSolv 的低成本重构虽缩小误差但未超过 Chemprop，多数外部 ADMET 终点在验证证据不足时保留原策略。这种结果模式说明，FZYC-Mol 更适合作为可追踪的模型治理流程，而不是被解读为所有终点上的性能增强器。",
            "讨论：减少术语混杂并强化边界",
        ),
        (
            "可靠性分析进一步强调了模型审计的重要性",
            "可靠性分析进一步限定了该框架的使用方式。风险分数对分类错误的识别能力较强，但对回归高误差样本的识别能力有限；保形预测可提供接近目标覆盖率的预测集合或区间，但不能替代任务特异性校准；验证-测试排名审计显示，验证集选择仍受候选池规模、终点噪声和小样本波动影响。因此，FZYC-Mol 的输出应作为模型复核和决策支持证据，而不是自动化筛选结论。",
            "讨论：用边界句替代重复结果",
        ),
        (
            "本文仍存在一定局限",
            "本研究仍存在局限。首先，FZYC-Mol 以公开基准数据集和离线评估为主，尚未经过真实药物发现项目中的前瞻性验证。其次，bRo5 化学空间仅作为 CycPept-PAMPA、LinPept CellPen 和 LinPept NonFouling 等公开数据上的压力测试，不能替代独立盲测或实验验证。再次，FreeSolv 等物理相互作用相关任务提示，二维描述符和低成本重构仍难以充分捕捉构象、溶剂化和长程相互作用。最后，基序和片段富集只能作为关联性解释，仍需结合统计校正、化学合理性和实验知识进行复核。",
            "讨论：局限性语言更正式",
        ),
        (
            "尽管存在上述局限",
            "在这些边界内，FZYC-Mol 为 ADMET 预测、虚拟筛选、分子优化和毒性风险筛查提供了一种可复核的报告模板。该模板要求模型同时报告性能、适用域证据、风险分位、最近邻解释和拒用理由，从而使计算预测更接近药物化学中的可审计决策流程。",
            "讨论：最后一段以适用范围收束",
        ),
        (
            "本文提出并评估了 FZYC-Mol",
            "本研究提出 FZYC-Mol，一种由验证集治理驱动、适用域感知的多专家分子性质预测框架。该框架将多视图表示、强基线专家、适用域判断、风险评分、校准、保形预测、活性悬崖分析和负结果审计整合到统一候选池中，并遵守测试集冻结原则。结果表明，FZYC-Mol 能够在部分 MoleculeNet 和外部 TDC ADMET 终点上提供选择性增益，同时用失败案例和负结果界定模型边界。其核心价值在于提供一套可复核的模型选择与可靠性报告流程。",
            "结论：三段式收束，避免泛化过强",
        ),
        (
            "未来工作将进一步扩展",
            "后续工作应在独立外部留出集、盲测风格验证、bRo5 化学空间和 MoleculeACE 活性悬崖任务上继续评估该框架，并在代表性低分任务上开展受控适配器、轻量微调或 3D-lite 物理代理实验。进一步的软件化实现还应保留候选登记、随机种子级预测、适用域证据和拒用理由，使 FZYC-Mol 能够作为可复现的分子性质预测审计工具使用。",
            "结论：未来工作改为具体、边界清楚的任务",
        ),
        (
            "本文使用的公开数据集可通过",
            "本研究使用的公开数据集可通过 MoleculeNet、Therapeutics Data Commons、MoleculeACE、Benchmark-ADMET-2025 及相应原始平台获得。与结果对应的 split seeds、候选登记表、验证/测试预测、统计检验脚本、图表 source data、环境文件和表格生成脚本应在投稿或接收版本中同步存档于 GitHub、Zenodo 或机构数据仓库。尚未公开的数据、未完成的盲测或仅作为方案设计的实验，不应写入完成性结果。",
            "数据可用性：Nature 风格下保留数据追溯边界",
        ),
    ]

    for prefix, new_text, label in replacements:
        replace_by_prefix(doc, prefix, new_text, changes, label)

    # Some sections begin with similar wording. Apply final, heading-scoped
    # replacements so the abstract and Results opening cannot overwrite each other.
    if set_nth_nonempty_after_heading(
        doc,
        "摘要",
        1,
        "分子性质预测已成为药物发现和 ADMET 评估的基础工具，但公开基准上的高平均分并不总能转化为可靠的化学决策。随机划分、低相似度分子、不平衡毒性标签、规则五以外化学空间和活性悬崖，都会削弱单一 ROC-AUC 或 RMSE 对真实应用风险的解释力。本研究提出 FZYC-Mol，一种由验证集治理驱动的适用域感知框架。该框架将分子图、指纹、二维描述符、片段/骨架信息和冻结预训练表征登记为候选证据，并仅允许通过验证集门控的专家模型、融合策略、补救头和适用域模块进入最终测试。",
    ):
        changes.append("摘要：按标题定位重写第一段")
    if set_nth_nonempty_after_heading(
        doc,
        "摘要",
        2,
        "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别取得 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302；在 22 个外部 TDC ADMET 终点中，最终保留结果为 win/tie/loss = 5/17/0。风险评分、保形预测、验证-测试排名审计、低相似度分层和 MoleculeACE 活性悬崖分析显示，该框架的主要价值不是在所有终点上追求统一增益，而是使模型选择、适用域边界和失败模式能够被追踪和复核。",
    ):
        changes.append("摘要：按标题定位重写结果与意义段")
    if set_nth_nonempty_after_heading(
        doc,
        "3.1 MoleculeNet 主结果",
        1,
        "在 MoleculeNet 主面板中，FZYC-Mol 在 ESOL、BACE 和 ClinTox 上分别取得 RMSE 0.5829 ± 0.0352、ROC-AUC 0.8753 ± 0.0230 和 ROC-AUC 0.9489 ± 0.0302。ESOL 和 BACE 中，验证集选择器与测试观测最优结果基本一致。Lipophilicity 中，验证集接受的定向补救策略将 RMSE 从 0.7078 ± 0.0389 降至 0.6835 ± 0.0439。",
    ):
        changes.append("结果：按标题定位重写 MoleculeNet 主结果开头")

    cleanup_pairs = {
        "弱 基线模型": "弱基线模型",
        "每个 数据集-随机种子 单元": "每个数据集-随机种子单元",
        "backbone": "主干模型",
        "公开 benchmark": "公开基准数据集",
        "benchmark 和离线评估": "基准数据集和离线评估",
        "KPGT representation": "KPGT 表征",
        "class weight": "类别权重",
        "validation metric": "验证集指标",
        "test metric": "测试集指标",
        "selected candidate": "最终保留策略",
        "seed-level paired difference": "随机种子级配对差值",
    }
    for paragraph in doc.paragraphs:
        original = paragraph.text
        updated = original
        for old, new in cleanup_pairs.items():
            updated = updated.replace(old, new)
        if updated != original:
            set_paragraph_text(paragraph, updated)

    for table in doc.tables:
        apply_three_line_table(table)

    return changes


def build_report(changes: list[str], removed_review: list[str], qa_notes: list[str]) -> None:
    doc = Document()
    doc.add_heading("初稿-5 Nature 技能修订说明", level=1)
    doc.add_paragraph(f"源文档：{SRC_DOCX}")
    doc.add_paragraph(f"输出文档：{DEST_DOCX}")
    doc.add_heading("调用的 Nature 技能片段", level=2)
    doc.add_paragraph(
        "本轮依据 nature-polishing 技能中的 reader-workflow、ethics、terminology-ledger、failure-modes、research paper、abstract、introduction、methods、results、discussion、conclusion、zh-to-en 和 Nature journal 片段修订。修订原则为：先修正论文类型与章节功能，再处理段落逻辑、主张-证据边界和句子层面表达。"
    )
    doc.add_heading("术语账本", level=2)
    table = doc.add_table(rows=1, cols=4)
    for i, header in enumerate(["标准术语", "首次使用含义", "常见变体", "本轮处理"]):
        table.rows[0].cells[i].text = header
    rows = [
        ["FZYC-Mol", "验证集治理驱动的适用域感知分子性质预测框架", "模型/框架/系统", "统一为框架或 FZYC-Mol"],
        ["验证集治理", "候选策略由验证集接受、拒绝或保留", "验证选择、选择器治理", "统一为验证集治理或验证集选择"],
        ["适用域", "训练/验证分布与测试分子的相似性及风险边界", "AD gate、适用域门控", "按模块语境分别使用"],
        ["公共压力测试", "基于可获得公开数据的边界评估", "盲测、外推验证", "避免写成独立盲测"],
        ["风险评分", "错误或高误差样本的辅助识别信号", "风险分数", "正文统一为风险分数或风险评分，避免机制化表述"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    apply_three_line_table(table)

    doc.add_heading("主要修改", level=2)
    for item in changes:
        doc.add_paragraph(item)
    doc.add_heading("从主文移出的内部自审内容", level=2)
    doc.add_paragraph(
        "Nature 风格主稿不宜包含内部审稿清单。本轮已将“投稿前自我审查”章节及其两张内部表格从主文结构中移出；如需保留，可作为作者内部修订材料或补充说明另存。"
    )
    for item in removed_review[:8]:
        doc.add_paragraph(item)
    doc.add_heading("QA 检查", level=2)
    for note in qa_notes:
        doc.add_paragraph(note)
    doc.save(REPORT_DOCX)


def collect_qa(doc) -> list[str]:
    text = "\n".join(p.text for p in doc.paragraphs)
    notes = []
    notes.append(f"段落数：{len(doc.paragraphs)}；表格数：{len(doc.tables)}；图片数：{len(doc.inline_shapes)}。")
    notes.append(f"最大表格列数：{max((len(t.columns) for t in doc.tables), default=0)}。")
    notes.append(f"内部自审章节：{'仍存在' if '6 投稿前自我审查' in text else '已从主文移出'}。")
    for term in ["弱 基线模型", "每个 数据集", "backbone", "benchmark 和离线评估", "KPGT representation", "class weight", "validation metric", "selected candidate"]:
        notes.append(f"术语检查 {term}：{text.count(term)} 处。")
    notes.append(f"List of abbreviations：{'存在' if 'List of abbreviations' in text else '未检出'}。")
    return notes


def main() -> None:
    doc = Document(SRC_DOCX)
    removed_review = remove_internal_review_section(doc)
    changes = polish_nature_style(doc)
    qa_notes = collect_qa(doc)

    doc.save(DEST_DOCX)
    shutil.copy2(DEST_DOCX, OUT_DIR / DEST_DOCX.name)
    build_report(changes, removed_review, qa_notes)
    shutil.copy2(REPORT_DOCX, OUT_DIR / REPORT_DOCX.name)

    print(f"Wrote {DEST_DOCX}")
    print(f"Wrote {REPORT_DOCX}")
    print(f"Nature-style paragraph changes: {len(changes)}")
    print(f"Removed internal review items: {len(removed_review)}")
    for note in qa_notes:
        print(note)


if __name__ == "__main__":
    main()
