from __future__ import annotations

import hashlib
import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "小论文-2.docx"
TRACKED = ROOT / "output" / "小论文-2_修订痕迹.docx"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def audit() -> Path:
    values = json.loads((ROOT / "results/manuscript_values.json").read_text(encoding="utf-8"))
    document = Document(DOCX)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    expected = [
        "机会校正命中率由 0.723 降至 0.142",
        "固定分母归一化遗憾由 0.090 增至 0.214",
        "3 外层×3 内层×5 重复",
        "相对 32 候选验证最优策略均为 7/0/2",
        "5 个晋级、17 个保留",
        "类别 1 仅为 0.673/0.807/0.832",
        "验证最优、冻结 one-SE、one-SE 低方差和 one-SE 低成本",
        "确认性结论限于公开离线基准中的轻量 Morgan-512 候选池",
    ]
    forbidden = [
        "真实误差排序上界",
        "TDC 的 5/17/0",
        "27 个外层测试单元",
        "同一历史结果被重复导入且验证与测试指标完全一致，只保留运行时间较短",
    ]
    zip_errors: list[str] = []
    with zipfile.ZipFile(DOCX) as archive:
        bad = archive.testzip()
        if bad:
            zip_errors.append(f"bad_member:{bad}")
        for name in archive.namelist():
            if name.endswith(".xml"):
                try:
                    ElementTree.fromstring(archive.read(name))
                except Exception as exc:  # pragma: no cover - audit path
                    zip_errors.append(f"{name}:{type(exc).__name__}")
    tracked_counts = {"insertions": 0, "deletions": 0}
    if TRACKED.exists():
        with zipfile.ZipFile(TRACKED) as archive:
            xml = archive.read("word/document.xml")
            tracked_counts = {
                "insertions": len(re.findall(br"<w:ins(?:\s|>)", xml)),
                "deletions": len(re.findall(br"<w:del(?:\s|>)", xml)),
            }
    report = {
        "docx": str(DOCX),
        "sha256": sha256(DOCX),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "equation_paragraphs": sum(paragraph.style.name == "Equation" for paragraph in document.paragraphs),
        "expected_text": {item: item in text for item in expected},
        "forbidden_text": {item: item in text for item in forbidden},
        "zip_xml_errors": zip_errors,
        "tracked_changes": tracked_counts,
        "manuscript_value_difference_count": json.loads((ROOT / "results/audits/manuscript_value_verification.json").read_text(encoding="utf-8"))["difference_count"],
        "source_values_hash": sha256(ROOT / "results/manuscript_values.json"),
        "source_reference": {
            "repeated_k32_fixed_regret": values["repeated_nested"]["32"]["fixed_normalized_regret"]["mean"],
            "autogluon_1800_validation_best_wins": values["autogluon_budget"]["1800"]["comparisons"]["validation_best"]["wins"],
        },
    }
    passed = (
        report["tables"] == 9
        and report["inline_shapes"] == 10
        and report["equation_paragraphs"] == 17
        and all(report["expected_text"].values())
        and not any(report["forbidden_text"].values())
        and not zip_errors
        and report["manuscript_value_difference_count"] == 0
    )
    report["passed"] = passed
    output = ROOT / "results/audits/small_paper_2_audit.json"
    output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    if not passed:
        raise SystemExit(json.dumps(report, ensure_ascii=False, indent=2))
    return output


if __name__ == "__main__":
    print(audit())
