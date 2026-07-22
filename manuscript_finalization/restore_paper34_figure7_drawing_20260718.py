from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


OLD = Path(r"D:\fzyc\output\paper33_final_minor_revision_20260718")
NEW = Path(r"D:\fzyc\output\paper34_submission_ready_20260718")


def chinese_doc(root: Path) -> Path:
    return next(path for path in root.glob("*.docx") if not path.name.startswith("Candidate"))


def restore(source: Path, target: Path, caption_prefix: str) -> None:
    old = Document(source)
    new = Document(target)
    drawings = [paragraph for paragraph in old.paragraphs if paragraph._p.find(".//" + qn("w:drawing")) is not None]
    if len(drawings) != 7:
        raise RuntimeError(f"Expected seven source drawings, found {len(drawings)}")
    candidates = [paragraph for paragraph in new.paragraphs if paragraph.text.strip().startswith(caption_prefix)]
    caption = next(
        (paragraph for paragraph in reversed(candidates) if paragraph._p.find(".//" + qn("w:drawing")) is None),
        candidates[-1],
    )
    for paragraph in candidates:
        if paragraph._p is not caption._p and paragraph._p.find(".//" + qn("w:drawing")) is not None:
            paragraph._p.getparent().remove(paragraph._p)
    current_drawings = [paragraph for paragraph in new.paragraphs if paragraph._p.find(".//" + qn("w:drawing")) is not None]
    for paragraph in current_drawings[7:]:
        paragraph._p.getparent().remove(paragraph._p)
    previous = caption._p.getprevious()
    if previous is not None and previous.find(".//" + qn("w:drawing")) is not None:
        drawing_node = previous
    else:
        drawing_node = deepcopy(drawings[-1]._p)
        caption._p.addprevious(drawing_node)
    for child in list(drawing_node):
        if child.tag != qn("w:pPr") and child.find(".//" + qn("w:drawing")) is None:
            drawing_node.remove(child)
    drawing_paragraph = next(paragraph for paragraph in new.paragraphs if paragraph._p is drawing_node)
    drawing_paragraph.style = "Normal"
    drawing_paragraph.paragraph_format.keep_with_next = True
    for paragraph in list(new.paragraphs):
        if paragraph.style.name == "Caption" and not paragraph.text.strip():
            next_node = paragraph._p.getnext()
            if next_node is not None and next_node.find(".//" + qn("w:drawing")) is not None:
                paragraph._p.getparent().remove(paragraph._p)
    new.save(target)


def main() -> None:
    restore(
        OLD / "Candidate_pool_expansion_Journal_of_Cheminformatics_FINAL_CLEAN.docx",
        NEW / "Candidate_pool_expansion_Journal_of_Cheminformatics_FINAL_CLEAN.docx",
        "Figure 7.",
    )
    restore(chinese_doc(OLD), chinese_doc(NEW), "图7")


if __name__ == "__main__":
    main()
